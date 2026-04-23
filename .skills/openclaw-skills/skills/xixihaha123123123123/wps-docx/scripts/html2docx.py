"""
轻量 HTML -> DOCX 转换器（中间层增强版）
方案: BeautifulSoup + 简化 CSS 解析 + python-docx
"""

import base64
import io
import re
import urllib.request
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PX_TO_PT = 0.75
DEFAULT_FONT_PT = 12.0
DOC_CONTENT_WIDTH_TWIPS = 8640  # 约 6 英寸可用宽度
LINE_HEIGHT_WORD_COMPENSATION = 1.5  # 迁移自 export_docx.js: (lh/fontSize)/1.5
NESTED_UNORDERED_LIST_TIGHTEN_PT = 7.0  # 收紧二级及以上无序列表的层级缩进
INHERITED_PROPS = {
    "color",
    "font-size",
    "font-family",
    "font-weight",
    "font-style",
    "line-height",
    "text-align",
    "text-indent",
    "_layout-left-offset-pt",
}
WORD_FONT_THEME_ATTRS = ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme")
GENERIC_FONT_FALLBACKS = {
    "serif": {"ascii": "Times New Roman", "eastAsia": "SimSun"},
    "ui-serif": {"ascii": "Times New Roman", "eastAsia": "SimSun"},
    "sans-serif": {"ascii": "Arial", "eastAsia": "Microsoft YaHei"},
    "ui-sans-serif": {"ascii": "Arial", "eastAsia": "Microsoft YaHei"},
    "system-ui": {"ascii": "Arial", "eastAsia": "Microsoft YaHei"},
    "monospace": {"ascii": "Consolas", "eastAsia": "Microsoft YaHei"},
    "ui-monospace": {"ascii": "Consolas", "eastAsia": "Microsoft YaHei"},
    "emoji": {"ascii": "Segoe UI Emoji", "eastAsia": "Segoe UI Emoji"},
}
PAGE_SIZE_PRESETS = {
    "a3": (Cm(29.7), Cm(42.0)),
    "a4": (Cm(21.0), Cm(29.7)),
    "a5": (Cm(14.8), Cm(21.0)),
    "b4": (Cm(25.0), Cm(35.3)),
    "b5": (Cm(17.6), Cm(25.0)),
    "letter": (Inches(8.5), Inches(11.0)),
    "legal": (Inches(8.5), Inches(14.0)),
}
CJK_FONT_KEYWORDS = (
    "simsun",
    "nsimsun",
    "simhei",
    "microsoft yahei",
    "microsoft jhenghei",
    "pingfang",
    "hiragino",
    "heiti",
    "songti",
    "fangsong",
    "kaiti",
    "stsong",
    "stfangsong",
    "stkaiti",
    "stheiti",
    "source han",
    "noto sans cjk",
    "noto serif cjk",
    "wenquanyi",
    "sarasa",
    "ipaex",
    "ms gothic",
    "msgothic",
    "ms mincho",
    "meiryo",
    "yu gothic",
    "yu mincho",
    "malgun",
    "batang",
    "gulim",
    "dotum",
)


def is_transparent_color(color_value):
    if color_value is None:
        return True
    v = str(color_value).strip().lower()
    return v in {"transparent", "rgba(0, 0, 0, 0)", "rgba(0,0,0,0)"}


def extract_css_variables(css_text):
    """提取 :root 中的 CSS 变量。"""
    vars_map = {}
    if not css_text:
        return vars_map
    for root_body in re.findall(r":root\s*\{([^{}]*)\}", css_text, flags=re.I | re.S):
        decls = parse_inline_style(root_body)
        for k, v in decls.items():
            if k.startswith("--"):
                vars_map[k.strip()] = v.strip()
    return vars_map


def resolve_css_var_value(value, css_vars, depth=0):
    """解析 var(--x[, fallback])。"""
    if value is None:
        return value
    if not isinstance(value, str):
        return value
    if depth > 8:
        return value

    pattern = re.compile(r"var\(\s*(--[\w\-]+)\s*(?:,\s*([^)]+))?\)")
    text = value
    changed = False

    def repl(match):
        nonlocal changed
        var_name = match.group(1).strip()
        fallback = (match.group(2) or "").strip()
        if var_name in css_vars:
            changed = True
            return str(css_vars[var_name])
        if fallback:
            changed = True
            return fallback
        return match.group(0)

    text2 = pattern.sub(repl, text)
    if changed:
        return resolve_css_var_value(text2, css_vars, depth + 1)
    return text2


def resolve_style_vars(style_dict, css_vars):
    if not style_dict:
        return {}
    result = {}
    for k, v in style_dict.items():
        if isinstance(v, str):
            result[k] = resolve_css_var_value(v, css_vars)
        else:
            result[k] = v
    return result


def sanitize_html_lightweight(html):
    """
    迁移自 html2docx.py 的轻量预清洗：
    修复坏结束标签，减少解析器容错导致的样式作用域泄漏。
    """
    n = len(html)
    i = 0
    out = []

    def is_tag_char(ch):
        return ch.isalnum() or ch in ["_", "-"]

    while i < n:
        if i + 2 < n and html[i] == "<" and html[i + 1] == "/":
            start = i
            j = i + 2
            if j < n and html[j].isalpha():
                while j < n and is_tag_char(html[j]):
                    j += 1
                tag = html[i + 2 : j]
                k = j
                while k < n and html[k].isspace():
                    k += 1
                if k < n and html[k] == ">":
                    out.append(html[start : k + 1])
                    i = k + 1
                    continue

                gt = html.find(">", k)
                lt = html.find("<", k)
                if gt != -1 and (lt == -1 or gt < lt):
                    out.append(f"</{tag}>")
                    out.append(html[k:gt])
                    i = gt + 1
                    continue
                if lt != -1:
                    out.append(f"</{tag}>")
                    out.append(html[k:lt])
                    i = lt
                    continue
                out.append(f"</{tag}>")
                out.append(html[k:])
                i = n
                continue

        out.append(html[i])
        i += 1
    return "".join(out)


def is_emoji_char(ch):
    if not ch:
        return False
    code = ord(ch)
    return (0x1F300 <= code <= 0x1FAFF) or (0x2600 <= code <= 0x27BF)


def split_text_for_emoji(text):
    """迁移自 html2docx.py：拆分 emoji 片段，单独设字体。"""
    if not text:
        return []
    parts = []
    buf = ""
    for ch in text:
        if ch == "\uFE0F" and parts and parts[-1][1]:
            prev, _ = parts[-1]
            parts[-1] = (prev + ch, True)
            continue
        if is_emoji_char(ch):
            if buf:
                parts.append((buf, False))
                buf = ""
            parts.append((ch, True))
        else:
            buf += ch
    if buf:
        parts.append((buf, False))
    return parts


def split_font_family_list(family_value):
    if not family_value:
        return []
    items = []
    buf = []
    quote = None
    for ch in str(family_value):
        if ch in {'"', "'"}:
            if quote == ch:
                quote = None
            elif quote is None:
                quote = ch
            else:
                buf.append(ch)
            continue
        if ch == "," and quote is None:
            item = "".join(buf).strip().strip("'\"")
            if item:
                items.append(item)
            buf = []
            continue
        buf.append(ch)
    tail = "".join(buf).strip().strip("'\"")
    if tail:
        items.append(tail)
    return items


LIST_STYLE_TYPE_KEYWORDS = {
    "none",
    "disc",
    "circle",
    "square",
    "decimal",
    "decimal-leading-zero",
    "lower-alpha",
    "upper-alpha",
    "lower-latin",
    "upper-latin",
    "lower-roman",
    "upper-roman",
    "lower-greek",
    "armenian",
    "georgian",
    "cjk-ideographic",
    "hiragana",
    "katakana",
    "hiragana-iroha",
    "katakana-iroha",
}
LIST_STYLE_POSITION_KEYWORDS = {"inside", "outside"}


def parse_list_style_shorthand(value):
    result = {}
    if not value:
        return result
    tokens = re.findall(r'url\([^)]*\)|"[^"]*"|\'[^\']*\'|\S+', str(value))
    for token in tokens:
        lower = token.strip().lower()
        if not lower:
            continue
        if lower in LIST_STYLE_POSITION_KEYWORDS:
            result["list-style-position"] = lower
            continue
        if lower in LIST_STYLE_TYPE_KEYWORDS:
            result["list-style-type"] = lower
            continue
        if lower.startswith("url("):
            result["list-style-image"] = token
    return result


def is_cjk_font_name(font_name):
    if not font_name:
        return False
    text = str(font_name).strip()
    if not text:
        return False
    if re.search(r"[\u3400-\u4dbf\u4e00-\u9fff\u3040-\u30ff\uac00-\ud7af]", text):
        return True
    lower = text.lower()
    return any(keyword in lower for keyword in CJK_FONT_KEYWORDS)


def resolve_font_slots(family_value):
    """把 CSS font-family 列表拆成 Word 的 ascii/eastAsia 槽位。"""
    families = split_font_family_list(family_value)
    if not families:
        return {}

    explicit = []
    generic_slots = []
    for family in families:
        generic = GENERIC_FONT_FALLBACKS.get(family.lower())
        if generic:
            generic_slots.append(generic)
        else:
            explicit.append(family)

    primary = explicit[0] if explicit else None
    primary_is_cjk = is_cjk_font_name(primary) if primary else False

    ascii_font = None
    east_asia_font = None

    if primary_is_cjk:
        ascii_font = primary
        east_asia_font = primary
    else:
        for family in explicit:
            if not is_cjk_font_name(family):
                ascii_font = family
                break
        for family in explicit:
            if is_cjk_font_name(family):
                east_asia_font = family
                break

    if not ascii_font:
        for slot in generic_slots:
            ascii_font = slot.get("ascii") or slot.get("eastAsia")
            if ascii_font:
                break
    if not east_asia_font:
        for slot in generic_slots:
            east_asia_font = slot.get("eastAsia") or slot.get("ascii")
            if east_asia_font:
                break

    if not ascii_font:
        ascii_font = primary or east_asia_font
    if not east_asia_font:
        east_asia_font = primary or ascii_font
    if not ascii_font and not east_asia_font:
        return {}

    ascii_font = ascii_font or east_asia_font
    east_asia_font = east_asia_font or ascii_font
    return {
        "ascii": ascii_font,
        "hAnsi": ascii_font,
        "eastAsia": east_asia_font,
        "cs": ascii_font,
    }


def set_run_font_family(run, family_value):
    """按脚本分别设置 Word 字体槽位，避免中英文共用同一字体。"""
    font_slots = resolve_font_slots(family_value)
    if not font_slots:
        return
    font_name = font_slots["ascii"]
    if not font_name:
        return
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_slots["ascii"])
    r_fonts.set(qn("w:hAnsi"), font_slots["hAnsi"])
    r_fonts.set(qn("w:eastAsia"), font_slots["eastAsia"])
    r_fonts.set(qn("w:cs"), font_slots["cs"])
    for theme_attr in WORD_FONT_THEME_ATTRS:
        theme_key = qn(theme_attr)
        if theme_key in r_fonts.attrib:
            del r_fonts.attrib[theme_key]


def parse_inline_style(style_text):
    styles = {}
    if not style_text:
        return styles
    for item in str(style_text).split(";"):
        if ":" not in item:
            continue
        key, value = item.split(":", 1)
        styles[key.strip().lower()] = value.strip()
    return styles


def expand_box_shorthand(value):
    """展开 margin/padding 的 1~4 值写法"""
    parts = [p for p in str(value).split() if p]
    if not parts:
        return None
    if len(parts) == 1:
        t = r = b = l = parts[0]
    elif len(parts) == 2:
        t = b = parts[0]
        r = l = parts[1]
    elif len(parts) == 3:
        t = parts[0]
        r = l = parts[1]
        b = parts[2]
    else:
        t, r, b, l = parts[:4]
    return t, r, b, l


def normalize_declarations(styles):
    """
    中间层-0：规范化 CSS 声明。
    目标是把简写属性摊平，后续渲染逻辑只处理“长属性”。
    """
    if not styles:
        return {}
    result = {}

    for raw_key, raw_value in styles.items():
        key = str(raw_key).strip().lower()
        value = raw_value.strip() if isinstance(raw_value, str) else raw_value
        result[key] = value

        if key == "margin":
            expanded = expand_box_shorthand(value)
            if expanded:
                result["margin-top"], result["margin-right"], result["margin-bottom"], result["margin-left"] = expanded
            continue

        if key == "padding":
            expanded = expand_box_shorthand(value)
            if expanded:
                result["padding-top"], result["padding-right"], result["padding-bottom"], result["padding-left"] = expanded
            continue

        if key == "border":
            for edge in ("top", "right", "bottom", "left"):
                result[f"border-{edge}"] = value
            continue

        if key == "list-style":
            for sub_key, sub_value in parse_list_style_shorthand(value).items():
                result[sub_key] = sub_value
            continue

        if key == "font":
            m = re.search(r"(\d+(?:\.\d+)?(?:px|pt|em|rem|%))(?:/([^\s]+))?\s+(.+)$", str(value))
            if m:
                result["font-size"] = m.group(1)
                if m.group(2):
                    result["line-height"] = m.group(2)
                result["font-family"] = m.group(3)
            if "italic" in str(value):
                result["font-style"] = "italic"
            if "bold" in str(value) or re.search(r"\b[6-9]00\b", str(value)):
                result["font-weight"] = "700"
            continue

        if key == "background":
            bg_text = str(value).strip().lower()
            # 渐变背景不应降级为纯色底纹（例如 highlight 文本的线性渐变）。
            if "gradient(" not in bg_text:
                color_match = re.search(r"(#[0-9a-fA-F]{3,8}|rgba?\([^)]+\))", str(value))
                if color_match:
                    result["background-color"] = color_match.group(1)
                elif bg_text.startswith("var("):
                    result["background-color"] = value
            continue

        if key == "flex":
            parts = [p for p in str(value).strip().split() if p]
            if parts:
                num_re = r"^-?\d+(?:\.\d+)?$"
                if len(parts) >= 1 and re.match(num_re, parts[0]):
                    result["flex-grow"] = parts[0]
                if len(parts) >= 2 and re.match(num_re, parts[1]):
                    result["flex-shrink"] = parts[1]
                if len(parts) >= 3:
                    result["flex-basis"] = parts[2]

    return result


def parse_css_color(color_str):
    if not color_str:
        return None
    color_str = str(color_str).strip()
    if is_transparent_color(color_str):
        return None

    named = {
        "white": RGBColor(255, 255, 255),
        "black": RGBColor(0, 0, 0),
        "gray": RGBColor(128, 128, 128),
        "grey": RGBColor(128, 128, 128),
        "silver": RGBColor(192, 192, 192),
        "red": RGBColor(255, 0, 0),
        "green": RGBColor(0, 128, 0),
        "blue": RGBColor(0, 0, 255),
        "yellow": RGBColor(255, 255, 0),
    }
    lower = color_str.lower()
    if lower in named:
        return named[lower]

    rgba = re.match(
        r"rgba\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*([\d.]+)\s*\)",
        color_str,
        re.IGNORECASE,
    )
    if rgba:
        r, g, b = int(rgba.group(1)), int(rgba.group(2)), int(rgba.group(3))
        alpha = float(rgba.group(4))
        r = int(r * alpha + 255 * (1 - alpha))
        g = int(g * alpha + 255 * (1 - alpha))
        b = int(b * alpha + 255 * (1 - alpha))
        return RGBColor(r, g, b)

    rgb = re.match(r"rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", color_str, re.IGNORECASE)
    if rgb:
        return RGBColor(int(rgb.group(1)), int(rgb.group(2)), int(rgb.group(3)))

    hex6 = re.match(r"^#([0-9a-fA-F]{6})$", color_str)
    if hex6:
        h = hex6.group(1)
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    hex3 = re.match(r"^#([0-9a-fA-F]{3})$", color_str)
    if hex3:
        h = hex3.group(1)
        return RGBColor(int(h[0] * 2, 16), int(h[1] * 2, 16), int(h[2] * 2, 16))

    return None


def decode_css_escapes(text):
    if not text:
        return ""

    def repl(match):
        try:
            return chr(int(match.group(1), 16))
        except Exception:
            return match.group(0)

    decoded = re.sub(r"\\([0-9a-fA-F]{1,6})\s?", repl, str(text))
    decoded = decoded.replace("\\A", "\n").replace("\\a", "\n")
    decoded = decoded.replace('\\"', '"').replace("\\'", "'").replace("\\\\", "\\")
    return decoded


def parse_css_content_text(content_value):
    if content_value is None:
        return ""
    text = str(content_value).strip()
    if not text or text.lower() in {"none", "normal"}:
        return ""
    if text.lower().startswith("attr("):
        return ""

    pieces = []
    for match in re.finditer(r'"((?:\\.|[^"\\])*)"|\'((?:\\.|[^\'\\])*)\'', text):
        pieces.append(match.group(1) if match.group(1) is not None else match.group(2))
    if pieces:
        return decode_css_escapes("".join(pieces))
    return decode_css_escapes(text)


def has_inline_box_visual(styles, current_font_pt=DEFAULT_FONT_PT):
    bg = styles.get("background-color")
    if not is_transparent_color(bg):
        return True
    if any(resolve_edge_border(styles, edge) for edge in ("left", "right", "top", "bottom")):
        return True
    left_padding_pt = length_to_pt(styles.get("padding-left"), current_font_pt=current_font_pt) or 0.0
    right_padding_pt = length_to_pt(styles.get("padding-right"), current_font_pt=current_font_pt) or 0.0
    return left_padding_pt > 0 or right_padding_pt > 0


def length_to_pt(length_value, current_font_pt=DEFAULT_FONT_PT, root_font_pt=DEFAULT_FONT_PT):
    if not length_value:
        return None
    text = str(length_value).strip().lower()
    if re.fullmatch(r"-?0+(?:\.0+)?", text):
        return 0.0
    try:
        if text.endswith("px"):
            return float(text[:-2]) * PX_TO_PT
        if text.endswith("pt"):
            return float(text[:-2])
        if text.endswith("mm"):
            return float(text[:-2]) * 72.0 / 25.4
        if text.endswith("cm"):
            return float(text[:-2]) * 72.0 / 2.54
        if text.endswith("in"):
            return float(text[:-2]) * 72.0
        # rem 必须先于 em 判断，否则 0.9rem 会被错误截成 0.9r。
        if text.endswith("rem"):
            return float(text[:-3]) * root_font_pt
        if text.endswith("em"):
            return float(text[:-2]) * current_font_pt
        if text.endswith("%"):
            return (float(text[:-1]) / 100.0) * current_font_pt
    except ValueError:
        return None
    return None


def remove_css_comments(css_text):
    return re.sub(r"/\*.*?\*/", "", css_text, flags=re.S)


def parse_page_rule(css_text):
    if not css_text:
        return {}
    css_text = remove_css_comments(css_text)
    page_style = {}
    for body in re.findall(r"@page\b[^{]*\{([^{}]*)\}", css_text, flags=re.I | re.S):
        declarations = normalize_declarations(parse_inline_style(body))
        if declarations:
            page_style.update(declarations)
    return page_style


def resolve_page_size_value(size_value, current_width=None, current_height=None):
    text = str(size_value or "").strip().lower()
    if not text:
        return None, None, None

    tokens = [tok for tok in re.split(r"\s+", text) if tok]
    orientation = None
    size_tokens = []
    for tok in tokens:
        if tok in {"portrait", "landscape"}:
            orientation = tok
        elif tok != "auto":
            size_tokens.append(tok)

    width = height = None
    if len(size_tokens) == 1 and size_tokens[0] in PAGE_SIZE_PRESETS:
        width, height = PAGE_SIZE_PRESETS[size_tokens[0]]
    elif len(size_tokens) == 2:
        width_pt = length_to_pt(size_tokens[0], current_font_pt=DEFAULT_FONT_PT)
        height_pt = length_to_pt(size_tokens[1], current_font_pt=DEFAULT_FONT_PT)
        if width_pt and height_pt:
            width = Pt(width_pt)
            height = Pt(height_pt)
    elif not size_tokens and orientation and current_width is not None and current_height is not None:
        width = current_width
        height = current_height

    if width is None or height is None:
        return None, None, orientation

    if orientation == "landscape" and int(width) < int(height):
        width, height = height, width
    elif orientation == "portrait" and int(width) > int(height):
        width, height = height, width
    return width, height, orientation


def apply_page_rule_to_document(doc, page_style):
    if not page_style:
        return
    try:
        sec = doc.sections[0]
    except Exception:
        return

    width, height, orientation = resolve_page_size_value(
        page_style.get("size"),
        current_width=sec.page_width,
        current_height=sec.page_height,
    )
    if orientation == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
    elif orientation == "portrait":
        sec.orientation = WD_ORIENT.PORTRAIT
    if width is not None and height is not None:
        sec.page_width = width
        sec.page_height = height

    for css_key, attr_name in (
        ("margin-top", "top_margin"),
        ("margin-right", "right_margin"),
        ("margin-bottom", "bottom_margin"),
        ("margin-left", "left_margin"),
    ):
        value_pt = length_to_pt(page_style.get(css_key), current_font_pt=DEFAULT_FONT_PT)
        if value_pt is not None:
            setattr(sec, attr_name, Pt(max(0.0, value_pt)))


def selector_specificity(selector):
    selector = selector.strip()
    id_count = len(re.findall(r"#[\w\-]+", selector))
    class_count = len(re.findall(r"\.[\w\-]+", selector))
    pseudo_count = len(re.findall(r":[\w\-]+(?:\([^)]+\))?", selector))
    clean = re.sub(r"#[\w\-]+|\.[\w\-]+|:[\w\-]+(?:\([^)]+\))?", " ", selector)
    tag_count = len([x for x in re.split(r"\s+", clean.strip()) if x and x != "*"])
    return (id_count, class_count + pseudo_count, tag_count)


def parse_css_rules(css_text):
    """
    中间层-1：把 <style> 文本解析成规则列表。
    这里只做轻量解析，支持常见的 selector { key:value; } 结构。
    """
    css_text = remove_css_comments(css_text)
    rules = []
    order = 0
    for selector_part, body in re.findall(r"([^{}]+)\{([^{}]+)\}", css_text):
        declarations = normalize_declarations(parse_inline_style(body))
        if not declarations:
            continue
        for selector in selector_part.split(","):
            selector = selector.strip()
            if not selector:
                continue
            if selector.startswith("@"):
                continue
            rules.append(
                {
                    "selector": selector,
                    "declarations": declarations,
                    "specificity": selector_specificity(selector),
                    "order": order,
                }
            )
            order += 1
    return rules


def match_simple_selector(tag, simple_selector):
    base_selector = simple_selector.strip()
    pseudo = None
    pseudo_element_match = re.search(r"::([\w\-]+)$", base_selector)
    if pseudo_element_match:
        pseudo = pseudo_element_match.group(1).lower()
        pseudo_arg = ""
        base_selector = base_selector[: pseudo_element_match.start()].strip()
    else:
        pseudo_match = re.search(r":([\w\-]+)(\(([^)]*)\))?$", base_selector)
        if pseudo_match:
            pseudo = pseudo_match.group(1).lower()
            pseudo_arg = (pseudo_match.group(3) or "").strip().lower()
            base_selector = base_selector[: pseudo_match.start()].strip()
        else:
            pseudo_arg = ""

    if base_selector and base_selector != "*":
        tag_name_match = re.match(r"^[a-zA-Z][\w\-]*", base_selector)
        if tag_name_match:
            if tag.name.lower() != tag_name_match.group(0).lower():
                return False

        for cls in re.findall(r"\.([\w\-]+)", base_selector):
            classes = tag.get("class") or []
            if cls not in classes:
                return False

        id_match = re.search(r"#([\w\-]+)", base_selector)
        if id_match and tag.get("id") != id_match.group(1):
            return False

    if pseudo and pseudo not in {"nth-child", "first-child", "last-child"}:
        # 轻量链路目前不支持 ::before/::after 和状态伪类，
        # 不能把它们错误地下沉到真实元素本体。
        return False

    if pseudo in {"nth-child", "first-child", "last-child"}:
        parent = tag.parent
        if parent is None:
            return False
        element_children = [c for c in parent.children if getattr(c, "name", None)]
        try:
            idx = element_children.index(tag) + 1
        except ValueError:
            return False

        if pseudo == "first-child":
            return idx == 1
        if pseudo == "last-child":
            return idx == len(element_children)
        if pseudo_arg == "even":
            return idx % 2 == 0
        if pseudo_arg == "odd":
            return idx % 2 == 1
        if pseudo_arg.isdigit():
            return idx == int(pseudo_arg)
        return False

    return True


def selector_matches(tag, selector):
    """
    中间层-2：轻量 selector 匹配（支持后代选择器、tag/class/id、nth-child）。
    采用“从右到左”匹配，尽量贴近 CSS 语义。
    """
    # 支持后代选择器和子选择器 >
    tokens = [t.strip() for t in re.split(r"(\s+|>)", selector) if t.strip() and not t.isspace()]
    if not tokens:
        return False

    current = tag
    if not match_simple_selector(current, tokens[-1]):
        return False

    i = len(tokens) - 2
    while i >= 0:
        token = tokens[i]
        if token == ">":
            i -= 1
            if i < 0:
                return False
            parent = current.parent
            if parent is None or not getattr(parent, "name", None):
                return False
            if not match_simple_selector(parent, tokens[i]):
                return False
            current = parent
        else:
            ancestor = current.parent
            found = False
            while ancestor is not None and getattr(ancestor, "name", None):
                if match_simple_selector(ancestor, token):
                    found = True
                    current = ancestor
                    break
                ancestor = ancestor.parent
            if not found:
                return False
        i -= 1
    return True


def build_tag_rule_style_map(soup, rules):
    """
    中间层-3：计算每个节点命中的 CSS 规则结果（只处理规则层，不含继承）。
    使用 specificity + 出现顺序做冲突决策。
    """
    result = {}
    for tag in soup.find_all(True):
        best = {}
        for rule in rules:
            if not selector_matches(tag, rule["selector"]):
                continue
            weight = (rule["specificity"], rule["order"])
            for key, val in rule["declarations"].items():
                if key not in best or weight >= best[key][0]:
                    best[key] = (weight, val)
        result[id(tag)] = {k: v[1] for k, v in best.items()}
    return result


def selector_base_for_pseudo(selector, pseudo_name):
    if not selector:
        return None
    match = re.search(rf"(?:::|:){re.escape(pseudo_name)}\s*$", str(selector).strip(), flags=re.I)
    if not match:
        return None
    base_selector = str(selector)[: match.start()].strip()
    return base_selector or "*"


def build_pseudo_rule_style_map(soup, rules, pseudo_name):
    result = {}
    for tag in soup.find_all(True):
        best = {}
        for rule in rules:
            base_selector = selector_base_for_pseudo(rule["selector"], pseudo_name)
            if not base_selector or not selector_matches(tag, base_selector):
                continue
            weight = (rule["specificity"], rule["order"])
            for key, val in rule["declarations"].items():
                if key not in best or weight >= best[key][0]:
                    best[key] = (weight, val)
        if best:
            result[id(tag)] = {k: v[1] for k, v in best.items()}
    return result


def parse_border_shorthand(border_value):
    if not border_value:
        return None
    parts = [p for p in str(border_value).strip().split() if p]
    if not parts:
        return None

    style = None
    width_pt = None
    color = "auto"
    style_tokens = {"solid", "dashed", "dotted", "double", "none", "hidden"}

    for token in parts:
        t = token.strip().lower()
        if t in style_tokens:
            style = t
            continue

        parsed_len = length_to_pt(t)
        if parsed_len is not None:
            width_pt = parsed_len
            continue
        if t == "0":
            width_pt = 0.0
            continue

        color = token

    if style is None:
        style = "solid"
    if width_pt is None:
        width_pt = 0.75

    # 迁移 export_docx.js parseBorderEdge 语义：无边框样式或 0 宽度都不渲染。
    if style in {"none", "hidden"} or width_pt <= 0:
        return None

    style_map = {
        "solid": "single",
        "dashed": "dashed",
        "dotted": "dotted",
        "double": "double",
    }
    color_obj = parse_css_color(color)
    return {
        "val": style_map.get(style, "single"),
        "sz": str(max(2, int((width_pt / PX_TO_PT) * 6))),
        "space": "1",
        "color": str(color_obj) if color_obj else "auto",
        "width_pt": width_pt,
    }


def build_border_from_longhands(styles, edge):
    """迁移 parseBoxBorders 思路：支持 border-*-width/style/color 组合。"""
    width_raw = styles.get(f"border-{edge}-width")
    style_raw = (styles.get(f"border-{edge}-style") or "").strip().lower()
    color_raw = styles.get(f"border-{edge}-color")

    width_pt = None
    if width_raw is not None:
        width_text = str(width_raw).strip().lower()
        if width_text == "0":
            width_pt = 0.0
        elif width_text in {"thin", "medium", "thick"}:
            width_pt = {"thin": 0.75, "medium": 1.5, "thick": 2.25}[width_text]
        else:
            width_pt = length_to_pt(width_text)

    if style_raw in {"none", "hidden"}:
        return None
    if width_pt is not None and width_pt <= 0:
        return None

    if width_pt is None and not style_raw and not color_raw:
        return None

    if width_pt is None:
        width_pt = 0.75
    if not style_raw:
        style_raw = "solid"

    style_map = {
        "solid": "single",
        "dashed": "dashed",
        "dotted": "dotted",
        "double": "double",
    }
    color_obj = parse_css_color(color_raw)
    return {
        "val": style_map.get(style_raw, "single"),
        "sz": str(max(2, int((width_pt / PX_TO_PT) * 6))),
        "space": "1",
        "color": str(color_obj) if color_obj else "auto",
        "width_pt": width_pt,
    }


def resolve_edge_border(styles, edge):
    """
    统一边框优先级：
    1) border-edge 简写
    2) border-edge-width/style/color 长属性组合
    3) border 总简写
    """
    edge_key = f"border-{edge}"
    if styles.get(edge_key):
        border = parse_border_shorthand(styles.get(edge_key))
        if border:
            return border

    longhand_border = build_border_from_longhands(styles, edge)
    if longhand_border:
        return longhand_border

    if styles.get("border"):
        return parse_border_shorthand(styles.get("border"))
    return None


def set_cell_border(cell, border_data):
    if not border_data:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "right", "bottom", "left"):
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), border_data["val"])
        edge_el.set(qn("w:sz"), border_data["sz"])
        edge_el.set(qn("w:space"), border_data["space"])
        edge_el.set(qn("w:color"), border_data["color"])


def set_cell_border_edge(cell, edge, border_data):
    if not border_data:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    edge_el = tc_borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        tc_borders.append(edge_el)
    edge_el.set(qn("w:val"), border_data["val"])
    edge_el.set(qn("w:sz"), border_data["sz"])
    edge_el.set(qn("w:space"), border_data["space"])
    edge_el.set(qn("w:color"), border_data["color"])


def set_tc_border_edge(tc, edge, border_data):
    if tc is None or not border_data:
        return
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    edge_el = tc_borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        tc_borders.append(edge_el)
    edge_el.set(qn("w:val"), border_data["val"])
    edge_el.set(qn("w:sz"), border_data["sz"])
    edge_el.set(qn("w:space"), border_data["space"])
    edge_el.set(qn("w:color"), border_data["color"])


def set_cell_padding(cell, styles, current_font_pt=DEFAULT_FONT_PT, horizontal_scale=1.0, vertical_scale=1.0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, css_key in (("top", "padding-top"), ("right", "padding-right"), ("bottom", "padding-bottom"), ("left", "padding-left")):
        pt_val = length_to_pt(styles.get(css_key), current_font_pt=current_font_pt)
        if pt_val is None:
            continue
        if edge in ("left", "right"):
            pt_val *= horizontal_scale
        else:
            pt_val *= vertical_scale
        twips = str(max(0, int(pt_val * 20)))
        edge_el = tc_mar.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_mar.append(edge_el)
        edge_el.set(qn("w:w"), twips)
        edge_el.set(qn("w:type"), "dxa")


def set_cell_width_twips(cell, width_twips):
    if width_twips <= 0:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_vertical_align(cell, align_value):
    if not align_value:
        return
    align_text = str(align_value).strip().lower()
    align_map = {
        "top": "top",
        "middle": "center",
        "center": "center",
        "bottom": "bottom",
        "baseline": "top",
    }
    val = align_map.get(align_text)
    if not val:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), val)


def paragraph_has_embedded_object(paragraph):
    p = paragraph._p
    embedded_tags = {
        qn("w:drawing"),
        qn("w:pict"),
        qn("w:object"),
        qn("w:fldSimple"),
    }
    return any(el.tag in embedded_tags for el in p.iter())


def is_truly_empty_paragraph(paragraph):
    return not paragraph.text.strip() and not paragraph_has_embedded_object(paragraph)


def tighten_empty_paragraph(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)
    if not paragraph.runs:
        run = paragraph.add_run("")
    else:
        run = paragraph.runs[0]
    run.font.size = Pt(1)


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def compact_cell_paragraphs(cell):
    paragraphs = list(cell.paragraphs)
    if len(paragraphs) <= 1:
        if paragraphs and is_truly_empty_paragraph(paragraphs[0]) and len(cell.tables) > 0:
            tighten_empty_paragraph(paragraphs[0])
        return

    changed = True
    while changed:
        changed = False
        paragraphs = list(cell.paragraphs)
        while len(paragraphs) > 1 and is_truly_empty_paragraph(paragraphs[0]):
            remove_paragraph(paragraphs[0])
            changed = True
            paragraphs = list(cell.paragraphs)

        if len(paragraphs) > 2:
            for para in paragraphs[1:-1]:
                if is_truly_empty_paragraph(para):
                    remove_paragraph(para)
                    changed = True
            paragraphs = list(cell.paragraphs)

        while len(paragraphs) > 1 and is_truly_empty_paragraph(paragraphs[-1]) and is_truly_empty_paragraph(paragraphs[-2]):
            remove_paragraph(paragraphs[-1])
            changed = True
            paragraphs = list(cell.paragraphs)

    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph()
        paragraphs = list(cell.paragraphs)

    last_para = paragraphs[-1]
    if is_truly_empty_paragraph(last_para):
        prev = last_para._p.getprevious()
        if prev is not None and prev.tag == qn("w:tbl"):
            tighten_empty_paragraph(last_para)
        elif len(paragraphs) > 1:
            remove_paragraph(last_para)


def parse_border_spacing_px(border_spacing_value):
    if not border_spacing_value:
        return 0.0
    parts = [p for p in str(border_spacing_value).strip().split() if p]
    if not parts:
        return 0.0
    values = []
    for p in parts[:2]:
        pt = length_to_pt(p)
        if pt is not None:
            values.append(pt / PX_TO_PT)
    return max(values) if values else 0.0


def get_grid_column_count(template_value):
    """从 grid-template-columns 推断列数。"""
    if not template_value:
        return 1
    text = str(template_value).strip().lower()
    if not text or text == "none":
        return 1
    m = re.search(r"repeat\(\s*(\d+)\s*,", text)
    if m:
        try:
            return max(1, int(m.group(1)))
        except ValueError:
            return 1
    tracks = [x for x in re.split(r"\s+", text) if x]
    return max(1, len(tracks))


def parse_grid_template_column_twips(template_value, target_twips):
    """解析 grid-template-columns，返回每列的 twips 宽度。"""
    text = str(template_value or "").strip().lower()
    if not text or text == "none":
        return [max(720, int(target_twips))]

    def split_tracks(raw_text):
        tracks = []
        buf = []
        depth = 0
        for ch in raw_text:
            if ch == "(":
                depth += 1
            elif ch == ")" and depth > 0:
                depth -= 1
            if ch.isspace() and depth == 0:
                if buf:
                    tracks.append("".join(buf).strip())
                    buf = []
                continue
            buf.append(ch)
        if buf:
            tracks.append("".join(buf).strip())
        return [t for t in tracks if t]

    def split_top_level_csv(raw_text):
        parts = []
        buf = []
        depth = 0
        for ch in raw_text:
            if ch == "(":
                depth += 1
            elif ch == ")" and depth > 0:
                depth -= 1
            if ch == "," and depth == 0:
                part = "".join(buf).strip()
                if part:
                    parts.append(part)
                buf = []
                continue
            buf.append(ch)
        if buf:
            part = "".join(buf).strip()
            if part:
                parts.append(part)
        return parts

    expanded = []
    repeat_re = re.compile(r"repeat\(\s*(\d+)\s*,\s*(.+)\)$", re.I)
    for token in split_tracks(text):
        m = repeat_re.match(token)
        if not m:
            expanded.append(token)
            continue
        try:
            repeat_n = max(1, int(m.group(1)))
        except ValueError:
            repeat_n = 1
        inner_tracks = split_tracks(m.group(2)) or ["1fr"]
        for _ in range(repeat_n):
            expanded.extend(inner_tracks)

    col_count = len(expanded) or 1
    min_col_twips = 720
    target_twips = max(min_col_twips * col_count, int(target_twips))

    fixed_twips = [None] * col_count
    fr_values = [0.0] * col_count
    auto_indices = []

    for idx, token in enumerate(expanded):
        item = str(token).strip().lower()
        if not item:
            auto_indices.append(idx)
            continue

        # minmax(a,b) 取上界作为主要分配依据。
        if item.startswith("minmax("):
            inner = item[7:-1] if item.endswith(")") else item[7:]
            parts = split_top_level_csv(inner)
            if parts:
                item = parts[-1]

        if item.endswith("fr"):
            try:
                fr_values[idx] = max(0.0, float(item[:-2]))
                continue
            except ValueError:
                pass

        if item.endswith("%"):
            try:
                ratio = float(item[:-1]) / 100.0
                fixed_twips[idx] = max(min_col_twips, int(target_twips * ratio))
                continue
            except ValueError:
                pass

        pt_val = length_to_pt(item, current_font_pt=DEFAULT_FONT_PT)
        if pt_val is not None and pt_val > 0:
            fixed_twips[idx] = max(min_col_twips, int(pt_val * 20))
            continue

        # auto / max-content / fit-content 等无法精确测量：与 fr 统一走剩余空间分配。
        auto_indices.append(idx)

    fixed_sum = sum(w for w in fixed_twips if w is not None)
    remain = target_twips - fixed_sum

    flex_indices = [i for i, fr in enumerate(fr_values) if fr > 0] + auto_indices
    if flex_indices:
        unit_weights = []
        for i in flex_indices:
            fr = fr_values[i]
            unit_weights.append(fr if fr > 0 else 1.0)
        total_units = sum(unit_weights) or float(len(unit_weights))
        remain = max(min_col_twips * len(flex_indices), remain)
        allocated = [max(min_col_twips, int(remain * u / total_units)) for u in unit_weights]
        if allocated:
            allocated[-1] += remain - sum(allocated)
        for i, w in zip(flex_indices, allocated):
            fixed_twips[i] = w

    widths = [w if w is not None else min_col_twips for w in fixed_twips]
    total_width = sum(widths)
    if total_width != target_twips and widths:
        widths[-1] += target_twips - total_width
        if widths[-1] < min_col_twips:
            deficit = min_col_twips - widths[-1]
            widths[-1] = min_col_twips
            if len(widths) > 1:
                widths[0] = max(min_col_twips, widths[0] - deficit)
    return widths


def set_table_cell_spacing(table, px_value):
    if not px_value or px_value <= 0:
        return
    twips = int(px_value * 15)  # 1px ~= 15 twips
    if twips <= 0:
        return
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    spacing = tbl_pr.find(qn("w:tblCellSpacing"))
    if spacing is None:
        spacing = OxmlElement("w:tblCellSpacing")
        tbl_pr.append(spacing)
    spacing.set(qn("w:w"), str(twips))
    spacing.set(qn("w:type"), "dxa")


def set_table_borders_none(table):
    """移除表格外框与内框，避免布局容器出现大黑框。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = tbl_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_el)
        edge_el.set(qn("w:val"), "nil")


def set_table_border_edge(table, edge, border_data):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    edge_el = tbl_borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        tbl_borders.append(edge_el)
    edge_el.set(qn("w:val"), border_data["val"])
    edge_el.set(qn("w:sz"), border_data["sz"])
    edge_el.set(qn("w:space"), border_data["space"])
    edge_el.set(qn("w:color"), border_data["color"])


def set_table_layout_fixed(table):
    """固定表格布局，避免 Word 自动改列宽导致 flex 比例失真。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_width_twips(table, width_twips):
    if width_twips <= 0:
        return
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_twips)))
    tbl_w.set(qn("w:type"), "dxa")

    # 清零表格缩进，避免不同容器块出现左侧偏移差异。
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_grid_widths(table, col_widths):
    if not col_widths:
        return
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)
    for width_twips in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(max(1, int(width_twips))))
        tbl_grid.append(grid_col)


def resolve_table_alignment(styles):
    margin_left = str((styles or {}).get("margin-left", "")).strip().lower()
    margin_right = str((styles or {}).get("margin-right", "")).strip().lower()
    if margin_left == "auto" and margin_right == "auto":
        return WD_TABLE_ALIGNMENT.CENTER
    if margin_left == "auto" and margin_right != "auto":
        return WD_TABLE_ALIGNMENT.RIGHT
    return WD_TABLE_ALIGNMENT.LEFT


def is_block_like_tag(tag):
    if not getattr(tag, "name", None):
        return False
    return tag.name.lower() in {
        "div", "p", "section", "article", "main", "header", "footer",
        "blockquote", "pre", "table", "ul", "ol",
        "h1", "h2", "h3", "h4", "h5", "h6",
    }


def is_code_like_block(tag, style):
    if not getattr(tag, "name", None):
        return False
    name = tag.name.lower()
    if name == "pre":
        return True
    classes = tag.get("class") or []
    if any(str(c).lower() in {"code-block", "code", "hljs"} for c in classes):
        return True
    white_space = str(style.get("white-space", "")).strip().lower()
    if white_space in {"pre", "pre-wrap", "break-spaces"}:
        return True
    font_family = str(style.get("font-family", "")).lower()
    if any(x in font_family for x in ["consolas", "monospace", "courier"]):
        return True
    return False


def estimate_col_char_weights(rows, num_cols):
    """按列文本长度估算宽度权重，支持 colspan。"""
    if num_cols <= 0:
        return []
    weights = [1.0] * num_cols
    for row in rows:
        col_idx = 0
        cells = row.find_all(["th", "td"], recursive=False)
        for cell in cells:
            try:
                colspan = max(1, int(cell.get("colspan", 1)))
            except ValueError:
                colspan = 1
            colspan = min(colspan, max(1, num_cols - col_idx))

            text = cell.get_text(" ", strip=True)
            if text:
                total_weight = min(36.0, max(2.0, len(text) * 0.9))
                per_col_weight = max(1.0, total_weight / colspan)
                for k in range(col_idx, col_idx + colspan):
                    if 0 <= k < num_cols:
                        weights[k] = max(weights[k], per_col_weight)
            col_idx += colspan
            if col_idx >= num_cols:
                break
    return weights


def split_evenly(total_twips, count, min_twips=900):
    if count <= 0:
        return []
    if count == 1:
        return [max(min_twips, int(total_twips))]
    base = max(min_twips, int(total_twips / count))
    widths = [base] * count
    widths[-1] += int(total_twips) - sum(widths)
    return widths


def set_paragraph_shading(paragraph, color_str):
    color = parse_css_color(color_str)
    if not color:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(color))


def set_paragraph_border(paragraph, edge, border_data):
    if not border_data:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge_el = p_bdr.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        p_bdr.append(edge_el)
    edge_el.set(qn("w:val"), border_data["val"])
    edge_el.set(qn("w:sz"), border_data["sz"])
    edge_el.set(qn("w:space"), border_data["space"])
    edge_el.set(qn("w:color"), border_data["color"])


def set_paragraph_outline_level(paragraph, heading_level):
    """设置段落大纲级别（0=标题1 ... 8=标题9）。"""
    if heading_level is None:
        return
    try:
        level = int(heading_level)
    except (TypeError, ValueError):
        return
    if level < 1 or level > 9:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def apply_heading_paragraph_semantics(paragraph, heading_level, doc):
    """
    按旧链路语义标记标题：
    - 优先套用 Word 标题样式（让样式面板显示“标题”而非“正文”）
    - 兜底写大纲级别
    - 清掉可能导致大空白的分页相关属性
    """
    if heading_level is None:
        return
    try:
        level = int(heading_level)
    except (TypeError, ValueError):
        return
    if level < 1 or level > 9:
        return

    for sname in (f"Heading {level}", f"标题 {level}", f"Heading{level}"):
        try:
            paragraph.style = doc.styles[sname]
            break
        except Exception:
            continue

    # 标题样式不可用时，至少保留大纲层级能力。
    set_paragraph_outline_level(paragraph, level)

    pf = paragraph.paragraph_format
    # 避免模板里的标题样式把段落推到下一页，造成大面积留白。
    pf.page_break_before = False
    pf.keep_with_next = False
    pf.keep_together = False


def apply_paragraph_style(paragraph, styles, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=True):
    pf = paragraph.paragraph_format

    align = (styles.get("text-align") or "").lower()
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "start": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "end": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if align in align_map:
        pf.alignment = align_map[align]

    indent_pt = length_to_pt(styles.get("text-indent"), current_font_pt=current_font_pt)
    if indent_pt is not None and abs(indent_pt) > 1e-6:
        pf.first_line_indent = Pt(indent_pt)

    margin_top = length_to_pt(styles.get("margin-top"), current_font_pt=current_font_pt)
    margin_bottom = length_to_pt(styles.get("margin-bottom"), current_font_pt=current_font_pt)
    # 近似 CSS margin-collapsing：将上下边距折半映射，避免 Word 段距双倍叠加。
    if margin_top is not None:
        pf.space_before = Pt(max(0, margin_top * 0.5))
    if margin_bottom is not None:
        pf.space_after = Pt(max(0, margin_bottom * 0.5))
    left_indent_pt = 0.0
    margin_left = length_to_pt(styles.get("margin-left"), current_font_pt=current_font_pt)
    if margin_left is not None and margin_left > 0:
        left_indent_pt = margin_left
    padding_left = length_to_pt(styles.get("padding-left"), current_font_pt=current_font_pt)
    if padding_left is not None and padding_left > 0:
        left_indent_pt = padding_left
    extra_left_indent_pt = styles.get("_layout-left-offset-pt")
    try:
        extra_left_indent_pt = float(extra_left_indent_pt or 0.0)
    except (TypeError, ValueError):
        extra_left_indent_pt = 0.0
    if extra_left_indent_pt > 0:
        left_indent_pt += extra_left_indent_pt
    if left_indent_pt > 0:
        pf.left_indent = Pt(left_indent_pt)

    line_height = styles.get("line-height")
    if line_height:
        line_height_text = str(line_height).strip().lower()
        if line_height_text != "normal":
            css_ratio = None
            if line_height_text.replace(".", "", 1).isdigit():
                css_ratio = float(line_height_text)
            elif line_height_text.endswith("%"):
                try:
                    css_ratio = float(line_height_text[:-1]) / 100.0
                except ValueError:
                    css_ratio = None
            else:
                lh_pt = length_to_pt(line_height_text, current_font_pt=current_font_pt)
                if lh_pt and current_font_pt > 0:
                    css_ratio = lh_pt / current_font_pt

            if css_ratio and css_ratio > 0:
                # 与旧链路保持一致：Word 视觉补偿系数 1.5
                pf.line_spacing = css_ratio / LINE_HEIGHT_WORD_COMPENSATION

    if apply_box_visuals and styles.get("background-color"):
        if not is_transparent_color(styles.get("background-color")):
            set_paragraph_shading(paragraph, styles["background-color"])

    if apply_box_visuals:
        for edge in ("left", "right", "top", "bottom"):
            border = resolve_edge_border(styles, edge)
            if border:
                set_paragraph_border(paragraph, edge, border)
                # 左边框需要对应的左缩进，否则边框会超出页面左边距
                if edge == "left":
                    border_width_pt = border.get("width_pt", 0) or 0
                    if border_width_pt > 0:
                        # padding-left 映射到 first_line_indent，作为文字与边框的间距
                        padding_left_pt = length_to_pt(styles.get("padding-left"), current_font_pt=current_font_pt) or 0
                        pf.left_indent = Pt(border_width_pt)
                        if padding_left_pt > 0:
                            pf.first_line_indent = Pt(padding_left_pt)


def set_run_shading(run, color_str):
    color = parse_css_color(color_str)
    if not color:
        return
    r_pr = run._r.get_or_add_rPr()
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(color))


def apply_run_style(run, styles, font_pt):
    color = parse_css_color(styles.get("color"))
    if color:
        run.font.color.rgb = color

    family = styles.get("font-family")
    if family:
        set_run_font_family(run, family)

    size_pt = length_to_pt(styles.get("font-size"), current_font_pt=font_pt)
    if size_pt:
        run.font.size = Pt(size_pt)

    font_style = (styles.get("font-style") or "").lower()
    if font_style in {"italic", "oblique"}:
        run.font.italic = True

    vertical_align = (styles.get("vertical-align") or "").lower()
    if vertical_align in {"super", "sup", "superscript"}:
        run.font.superscript = True
    elif vertical_align in {"sub", "subscript"}:
        run.font.subscript = True

    fw = (styles.get("font-weight") or "").lower()
    if fw in {"bold", "bolder"} or (fw.isdigit() and int(fw) >= 600):
        run.font.bold = True

    # 迁移 export_docx.js 语义：优先看 textDecorationLine。
    deco = f'{styles.get("text-decoration-line", "")} {styles.get("text-decoration", "")}'.lower()
    if "underline" in deco:
        run.font.underline = True
    if "line-through" in deco:
        run.font.strike = True

    if styles.get("background-color"):
        set_run_shading(run, styles["background-color"])


def add_hr(doc):
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_cell_shading(cell, color_str):
    color = parse_css_color(color_str)
    if not color:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(color))


def html2docx_beautifulsoup(html_content: str, output_docx: str = None):
    """
    将 HTML 内容转换为 DOCX。
    
    Args:
        html_content: HTML 字符串内容
        output_docx: 输出文件路径。如果为 None，则输出 base64 字符串到 stdout
    
    Returns:
        如果 output_docx 为 None，返回 base64 编码的 docx 内容
    """
    try:
        from bs4 import BeautifulSoup, NavigableString, Tag, Comment  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少依赖 beautifulsoup4，请先安装: pip install beautifulsoup4") from exc

    cleaned_html = sanitize_html_lightweight(html_content)
    soup = BeautifulSoup(cleaned_html, "html.parser")
    doc = Document()
    next_shape_docpr_id = 1000
    last_render_was_container_table = False

    def get_doc_content_width_twips():
        """按当前节配置计算可用内容宽度（twips）。"""
        try:
            sec = doc.sections[0]
            return max(3600, int((sec.page_width - sec.left_margin - sec.right_margin) / 635))
        except Exception:
            return DOC_CONTENT_WIDTH_TWIPS

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(DEFAULT_FONT_PT)
    # 避免 Word 默认段后间距影响网页段距映射。
    normal_pf = normal.paragraph_format
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(0)
    # 与主方案保持一致：给 Normal 补 eastAsia，避免中文回退字体。
    r_pr = normal.element.rPr
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        normal.element.append(r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for heading_level in range(1, 7):
        # 设置标题样式
        for sname in (f"Heading {heading_level}", f"标题 {heading_level}", f"Heading{heading_level}"):
            try:
                h_style = doc.styles[sname]
            except KeyError:
                continue
            h_style.font.name = "Microsoft YaHei"
            h_rpr = h_style.element.rPr
            if h_rpr is None:
                h_rpr = OxmlElement("w:rPr")
                h_style.element.append(h_rpr)
            h_rfonts = h_rpr.find(qn("w:rFonts"))
            if h_rfonts is None:
                h_rfonts = OxmlElement("w:rFonts")
                h_rpr.append(h_rfonts)
            h_rfonts.set(qn("w:ascii"), "Microsoft YaHei")
            h_rfonts.set(qn("w:hAnsi"), "Microsoft YaHei")
            h_rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            h_rfonts.set(qn("w:cs"), "Microsoft YaHei")
            for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
                if h_rfonts.get(qn(theme_attr)) is not None:
                    del h_rfonts.attrib[qn(theme_attr)]
            for italic_tag in ("w:i", "w:iCs"):
                italic_el = h_rpr.find(qn(italic_tag))
                if italic_el is not None:
                    h_rpr.remove(italic_el)
            break

    css_text = "\n".join(st.get_text("\n") for st in soup.find_all("style"))
    css_vars = extract_css_variables(css_text)
    page_rule = resolve_style_vars(parse_page_rule(css_text), css_vars)
    apply_page_rule_to_document(doc, page_rule)
    rules = parse_css_rules(css_text)
    rule_map = build_tag_rule_style_map(soup, rules)
    before_rule_map = build_pseudo_rule_style_map(soup, rules, "before")

    def alloc_shape_docpr_id():
        nonlocal next_shape_docpr_id
        docpr_id = next_shape_docpr_id
        next_shape_docpr_id += 1
        return docpr_id

    def add_thin_separator_paragraph(height_pt=5.0):
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # 用很小的固定行高作为“薄分隔段”，避免形成明显空行。
        pf.line_spacing = Pt(max(1.0, float(height_pt)))
        run = para.add_run("")
        run.font.size = Pt(max(1.0, float(height_pt)))

    def add_spacing_paragraph(height_pt):
        height_pt = max(0.0, float(height_pt))
        if height_pt <= 0:
            return
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(max(1.0, height_pt))
        run = para.add_run("")
        run.font.size = Pt(max(1.0, height_pt))

    def render_container_table_with_optional_separator(tag, parent_style):
        nonlocal last_render_was_container_table
        if last_render_was_container_table:
            add_thin_separator_paragraph(5.0)
        render_container_as_table(tag, parent_style)
        last_render_was_container_table = True

    def resolve_style(tag, parent_style):
        """
        中间层-4：为当前节点合成“可用样式”。
        顺序：继承属性 -> 命中的 <style> 规则 -> 行内 style。
        """
        style = {k: parent_style[k] for k in INHERITED_PROPS if k in parent_style}
        style.update(rule_map.get(id(tag), {}))
        style.update(normalize_declarations(parse_inline_style(tag.get("style"))))
        return resolve_style_vars(style, css_vars)

    def add_text_run(paragraph, text, styles):
        if not text:
            return
        current_font_pt = length_to_pt(styles.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        for seg_text, is_emoji in split_text_for_emoji(text):
            run = paragraph.add_run(seg_text)
            run_style = dict(styles)
            if is_emoji:
                # 迁移主方案规则：emoji 不跟随文字颜色，尽量保留系统彩色字形。
                run_style.pop("color", None)
            apply_run_style(run, run_style, current_font_pt)
            if is_emoji:
                set_run_font_family(run, "Segoe UI Emoji")

    def add_preformatted_text(paragraph, text, styles):
        """保留 \\n 作为换行，并尽量保留代码文本空白。"""
        if text is None:
            return
        raw = str(text).replace("\r\n", "\n").replace("\r", "\n").replace("\t", "    ")
        if not raw:
            return
        lines = raw.split("\n")
        for i, line in enumerate(lines):
            # 仅包含空格但非空字符串也要保留，避免代码缩进丢失。
            if line != "":
                add_text_run(paragraph, line, styles)
            if i < len(lines) - 1:
                paragraph.add_run().add_break()

    def paragraph_has_visible_content(paragraph):
        if paragraph is None:
            return False
        for run in paragraph.runs:
            if (run.text or "").strip():
                return True
            if run._r.find(qn("w:drawing")) is not None or run._r.find(qn("w:pict")) is not None:
                return True
        return False

    def paragraph_ends_with_break(paragraph):
        if paragraph is None:
            return False
        for child in reversed(list(paragraph._p)):
            if child.tag != qn("w:r"):
                continue
            run_children = [node for node in child if node.tag != qn("w:rPr")]
            if not run_children:
                continue
            last_child = run_children[-1]
            if last_child.tag == qn("w:br"):
                return True
            if last_child.tag in {qn("w:t"), qn("w:drawing"), qn("w:pict"), qn("w:tab")}:
                return False
        return False

    def normalize_inline_text_for_paragraph(paragraph, text):
        if not text:
            return text
        if not paragraph_has_visible_content(paragraph) or paragraph_ends_with_break(paragraph):
            return text.lstrip(" ")
        return text

    def normalize_text_content(text, preserve_whitespace=False):
        if text is None:
            return ""
        if preserve_whitespace:
            # 代码块模式：主要依赖 <br> 控制换行，忽略 HTML 模板缩进噪音。
            raw = str(text).replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
            if not raw.strip():
                return ""
            # 把模板中的换行+缩进压成单空格，避免出现异常空行和错位换行。
            compact = re.sub(r"\s+", " ", raw)
            # 仅保留首尾一个空格（用于行内拼接），其余交给 <br>。
            return compact
        return re.sub(r"\s+", " ", str(text))

    def is_meaningful_text_node(node):
        return isinstance(node, NavigableString) and not isinstance(node, Comment) and str(node).strip()

    def _load_image_bytes(src):
        """
        从两种来源加载图片字节：
        1. data URI (base64 内嵌)
        2. http/https 远程 URL
        返回 bytes 或 None（失败时静默跳过）。
        """
        src = (src or "").strip()
        if not src:
            return None
        if src.startswith("data:"):
            if "," not in src:
                return None
            try:
                _, data_part = src.split(",", 1)
                return base64.b64decode(data_part)
            except Exception:
                return None
        if src.startswith("http://") or src.startswith("https://"):
            try:
                req = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=10) as resp:
                    return resp.read()
            except Exception:
                return None
        return None

    def _resolve_dimension_inches(raw_value, max_inches, parent_inches=None):
        raw = (raw_value or "")
        if raw is None:
            return None
        s = str(raw).strip().lower()
        if not s:
            return None
        if s.endswith("%"):
            if parent_inches and parent_inches > 0:
                try:
                    return max(0.2, min(max_inches, parent_inches * float(s[:-1]) / 100.0))
                except ValueError:
                    return None
            return None
        pt = length_to_pt(s, current_font_pt=DEFAULT_FONT_PT)
        if pt and pt > 0:
            return max(0.2, min(max_inches, pt / 72.0))
        try:
            return max(0.2, min(max_inches, float(s) * PX_TO_PT / 72.0))
        except ValueError:
            return None

    def _resolve_image_width_inches(img_tag, img_style, parent_style=None):
        """
        计算图片渲染宽度（英寸）。
        优先级：CSS width → HTML width 属性 → 父容器宽度 → 页面宽度 × 80%。
        百分比优先参考父容器宽度，避免 img{width:100%} 被误放大到整页。
        """
        page_max_w = DOC_CONTENT_WIDTH_TWIPS / 1440.0
        parent_w = _resolve_dimension_inches((parent_style or {}).get("width"), page_max_w) if parent_style else None
        max_w = min(page_max_w, parent_w) if parent_w else page_max_w
        raw = img_style.get("width") or img_tag.get("width")
        resolved = _resolve_dimension_inches(raw, max_w, parent_inches=parent_w)
        if resolved:
            return resolved
        if parent_w and parent_w > 0:
            return max(0.3, min(max_w, parent_w))
        return max_w * 0.8

    def _resolve_image_height_inches(img_tag, img_style, parent_style=None):
        """
        计算图片渲染高度（英寸）。
        仅在 HTML/CSS 明确给出高度，或高度百分比可落到父容器高度时返回。
        """
        page_max_h = 10.5
        parent_h = _resolve_dimension_inches((parent_style or {}).get("height"), page_max_h) if parent_style else None
        raw = img_style.get("height") or img_tag.get("height")
        return _resolve_dimension_inches(raw, page_max_h, parent_inches=parent_h)

    def render_image(img_tag, parent_style, inline_paragraph=None):
        """
        渲染 <img> 标签。
        inline_paragraph 非 None 时：插入到该段落（行内图片）。
        inline_paragraph 为 None 时：新建居中段落（块级图片）。
        加载失败或 src 缺失时静默跳过，不抛出异常。
        """
        img_bytes = _load_image_bytes(img_tag.get("src"))
        if not img_bytes:
            alt = (img_tag.get("alt") or "").strip()
            if alt and inline_paragraph is not None:
                add_text_run(inline_paragraph, f"[{alt}]", parent_style)
            return

        img_style = resolve_style(img_tag, parent_style)
        width_inches = _resolve_image_width_inches(img_tag, img_style, parent_style=parent_style)
        height_inches = _resolve_image_height_inches(img_tag, img_style, parent_style=parent_style)

        try:
            img_stream = io.BytesIO(img_bytes)
            if inline_paragraph is not None:
                run = inline_paragraph.add_run()
                if height_inches and height_inches > 0:
                    run.add_picture(img_stream, width=Inches(width_inches), height=Inches(height_inches))
                else:
                    run.add_picture(img_stream, width=Inches(width_inches))
            else:
                para = doc.add_paragraph()
                align_val = str(img_style.get("text-align", "center")).strip().lower()
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                }
                para.alignment = align_map.get(align_val, WD_ALIGN_PARAGRAPH.CENTER)
                if height_inches and height_inches > 0:
                    para.add_run().add_picture(img_stream, width=Inches(width_inches), height=Inches(height_inches))
                else:
                    para.add_run().add_picture(img_stream, width=Inches(width_inches))
                margin_bottom_pt = length_to_pt(
                    img_style.get("margin-bottom"), current_font_pt=DEFAULT_FONT_PT
                ) or 0.0
                if margin_bottom_pt > 0:
                    para.paragraph_format.space_after = Pt(margin_bottom_pt * 0.5)
        except Exception:
            pass

    def get_word_shape_box_spec(tag, styles):
        if not isinstance(tag, Tag):
            return None

        text = normalize_text_content(tag.get_text(" ", strip=True), preserve_whitespace=False).strip()
        if not text or len(text) > 40:
            return None

        tag_name = tag.name.lower()
        display_val = str(styles.get("display", "")).strip().lower()
        font_pt = length_to_pt(styles.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        radius_raw = str(styles.get("border-radius", "")).strip().lower()
        inline_shape_candidate = display_val in {"inline-block", "inline-flex"}
        if (
            not inline_shape_candidate
            and tag_name in {"span", "a", "strong", "b", "em", "i", "code"}
            and radius_raw
            and has_inline_box_visual(styles, current_font_pt=font_pt)
        ):
            inline_shape_candidate = True
        text_align = str(styles.get("text-align", "")).strip().lower()
        justify_content = str(styles.get("justify-content", "")).strip().lower()
        align_items = str(styles.get("align-items", "")).strip().lower()
        centered_content = (
            text_align == "center"
            or (
                display_val in {"flex", "inline-flex"}
                and justify_content in {"center", "space-around", "space-evenly"}
                and align_items in {"center", "stretch"}
            )
        )
        if inline_shape_candidate:
            centered_content = True
        if not centered_content:
            return None

        if any(isinstance(child, Tag) and is_block_like_tag(child) for child in tag.children):
            return None

        width_pt = length_to_pt(styles.get("width"), current_font_pt=font_pt) or 0.0
        height_pt = max(
            length_to_pt(styles.get("height"), current_font_pt=font_pt) or 0.0,
            length_to_pt(styles.get("min-height"), current_font_pt=font_pt) or 0.0,
        )

        if not radius_raw:
            return None

        if width_pt <= 0 or height_pt <= 0:
            if not inline_shape_candidate:
                return None

            padding_left_pt = length_to_pt(styles.get("padding-left"), current_font_pt=font_pt) or 0.0
            padding_right_pt = length_to_pt(styles.get("padding-right"), current_font_pt=font_pt) or 0.0
            padding_top_pt = length_to_pt(styles.get("padding-top"), current_font_pt=font_pt) or 0.0
            padding_bottom_pt = length_to_pt(styles.get("padding-bottom"), current_font_pt=font_pt) or 0.0

            def estimate_text_width_em(text_value):
                total = 0.0
                for ch in str(text_value):
                    if ch.isspace():
                        total += 0.35
                    elif re.search(r"[\u3400-\u4dbf\u4e00-\u9fff\u3040-\u30ff\uac00-\ud7af]", ch):
                        total += 1.02
                    elif is_emoji_char(ch):
                        total += 1.05
                    elif ch in ".,;:!?\"'`~|/\\-_=+()[]{}<>":
                        total += 0.38
                    elif ch.isupper():
                        total += 0.72
                    elif ch.isdigit():
                        total += 0.58
                    else:
                        total += 0.56
                return max(1.0, total)

            line_height_pt = None
            line_height = styles.get("line-height")
            if line_height:
                line_height_text = str(line_height).strip().lower()
                if line_height_text != "normal":
                    if line_height_text.replace(".", "", 1).isdigit():
                        line_height_pt = font_pt * float(line_height_text)
                    elif line_height_text.endswith("%"):
                        try:
                            line_height_pt = font_pt * (float(line_height_text[:-1]) / 100.0)
                        except ValueError:
                            line_height_pt = None
                    else:
                        line_height_pt = length_to_pt(line_height_text, current_font_pt=font_pt)

            content_height_pt = max(font_pt * 1.25, line_height_pt or 0.0)
            est_height_pt = max(font_pt + padding_top_pt + padding_bottom_pt, content_height_pt + padding_top_pt + padding_bottom_pt)
            est_width_pt = font_pt * estimate_text_width_em(text) + padding_left_pt + padding_right_pt + max(2.0, font_pt * 0.15)
            height_pt = max(height_pt, est_height_pt)
            width_pt = max(width_pt, est_width_pt)

        if width_pt <= 0 or height_pt <= 0:
            return None

        radius_ratio = 0.0
        if radius_raw.endswith("%"):
            try:
                radius_ratio = max(0.0, min(0.5, float(radius_raw[:-1]) / 100.0))
            except ValueError:
                radius_ratio = 0.0
        else:
            radius_pt = length_to_pt(radius_raw, current_font_pt=DEFAULT_FONT_PT) or 0.0
            radius_ratio = max(0.0, min(0.5, radius_pt / max(1.0, min(width_pt, height_pt))))

        if radius_ratio < 0.08:
            return None

        geometry = "ellipse" if radius_ratio >= 0.38 else "roundrect"
        if inline_shape_candidate and width_pt > height_pt * 1.2:
            geometry = "roundrect"
        border = (
            resolve_edge_border(styles, "top")
            or resolve_edge_border(styles, "left")
            or resolve_edge_border(styles, "right")
            or resolve_edge_border(styles, "bottom")
            or parse_border_shorthand(styles.get("border"))
        )
        fill_hex = None
        bg = styles.get("background-color")
        if bg and not is_transparent_color(bg):
            fill_rgb = parse_css_color(bg)
            if fill_rgb:
                fill_hex = str(fill_rgb)
        line_hex = None
        line_width_pt = 0.0
        if border:
            if border.get("color") not in {None, "auto"}:
                border_rgb = parse_css_color(border["color"])
                if border_rgb:
                    line_hex = str(border_rgb)
            line_width_pt = float(border.get("width_pt", 0.0) or 0.0)
        elif not fill_hex:
            line_hex = "DEE2E6"
        if line_hex and line_width_pt <= 0:
            line_width_pt = 0.75

        text_color = parse_css_color(styles.get("color")) or "666666"

        return {
            "geometry": geometry,
            "width_pt": width_pt,
            "height_pt": height_pt,
            "text": text,
            "fill_hex": fill_hex,
            "line_hex": line_hex,
            "line_width_pt": line_width_pt,
            "font_family": styles.get("font-family"),
            "font_pt": font_pt,
            "font_weight": str(styles.get("font-weight", "")).strip().lower(),
            "text_color": str(text_color),
            "arcsize_pct": max(8, min(50, int(round(radius_ratio * 100)))),
        }

    def build_shape_textbox_content_xml(shape_spec):
        font_slots = resolve_font_slots(shape_spec.get("font_family"))
        rpr_parts = []
        if font_slots:
            rpr_parts.append(
                f'<w:rFonts w:ascii="{xml_escape(font_slots["ascii"])}" '
                f'w:hAnsi="{xml_escape(font_slots["hAnsi"])}" '
                f'w:eastAsia="{xml_escape(font_slots["eastAsia"])}" '
                f'w:cs="{xml_escape(font_slots["cs"])}"/>'
            )
        half_pt = max(2, int(round(float(shape_spec.get("font_pt", DEFAULT_FONT_PT)) * 2)))
        rpr_parts.append(f'<w:sz w:val="{half_pt}"/>')
        rpr_parts.append(f'<w:szCs w:val="{half_pt}"/>')
        if shape_spec.get("font_weight") in {"bold", "bolder"}:
            rpr_parts.append("<w:b/>")
        color_hex = shape_spec.get("text_color")
        if color_hex:
            rpr_parts.append(f'<w:color w:val="{xml_escape(color_hex)}"/>')
        rpr_xml = f"<w:rPr>{''.join(rpr_parts)}</w:rPr>" if rpr_parts else ""
        text = xml_escape(str(shape_spec.get("text", "")))
        return (
            '<w:txbxContent>'
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:pPr><w:jc w:val="center"/><w:spacing w:before="0" w:after="0"/></w:pPr>'
            f'<w:r>{rpr_xml}<w:t>{text}</w:t></w:r>'
            '</w:p>'
            '</w:txbxContent>'
        )

    def build_inline_word_shape_xml(shape_spec):
        geometry = shape_spec.get("geometry") or "roundrect"
        width_pt = float(shape_spec.get("width_pt", 72.0))
        height_pt = float(shape_spec.get("height_pt", 72.0))
        style_parts = [
            f"width:{width_pt:.2f}pt",
            f"height:{height_pt:.2f}pt",
            "v-text-anchor:middle",
        ]

        attrs = ['o:allowincell="t"']
        fill_hex = shape_spec.get("fill_hex")
        if fill_hex:
            attrs.extend(['filled="t"', f'fillcolor="#{xml_escape(fill_hex)}"'])
        else:
            attrs.append('filled="f"')

        line_hex = shape_spec.get("line_hex")
        line_width_pt = float(shape_spec.get("line_width_pt", 0.0) or 0.0)
        if line_hex and line_width_pt > 0:
            attrs.extend(
                [
                    'stroked="t"',
                    f'strokecolor="#{xml_escape(line_hex)}"',
                    f'strokeweight="{line_width_pt:.2f}pt"',
                ]
            )
        else:
            attrs.append('stroked="f"')

        shape_tag = "v:oval" if geometry == "ellipse" else "v:roundrect"
        extra_attr = ""
        if shape_tag == "v:roundrect":
            extra_attr = f' arcsize="{int(shape_spec.get("arcsize_pct", 15))}%"'

        txbx_xml = build_shape_textbox_content_xml(shape_spec)
        return (
            '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'<{shape_tag} style="{";".join(style_parts)}" {" ".join(attrs)}{extra_attr}>'
            f'<v:textbox inset="0,0,0,0" style="mso-fit-shape-to-text:f">{txbx_xml}</v:textbox>'
            f'</{shape_tag}>'
            '</w:pict>'
        )

    def build_word_shape_element(shape_spec):
        try:
            return parse_xml(build_inline_word_shape_xml(shape_spec))
        except Exception:
            return None

    def render_word_shape_box(paragraph, shape_spec):
        shape_element = build_word_shape_element(shape_spec)
        if shape_element is None:
            return False
        run = paragraph.add_run()
        run._r.append(shape_element)
        return True

    def render_shape_box_tag_to_paragraph(paragraph, tag, tag_style):
        shape_spec = get_word_shape_box_spec(tag, tag_style)
        if not shape_spec:
            return False
        shape_element = build_word_shape_element(shape_spec)
        if shape_element is None:
            return False
        paragraph_style = dict(tag_style)
        for key in (
            "width",
            "height",
            "min-height",
            "background",
            "background-color",
            "border",
            "border-top",
            "border-right",
            "border-bottom",
            "border-left",
            "border-radius",
            "padding",
            "padding-top",
            "padding-right",
            "padding-bottom",
            "padding-left",
            "display",
            "justify-content",
            "align-items",
        ):
            paragraph_style.pop(key, None)
        paragraph.clear()
        apply_paragraph_style(
            paragraph,
            paragraph_style,
            current_font_pt=DEFAULT_FONT_PT,
            apply_box_visuals=False,
        )
        paragraph.add_run()._r.append(shape_element)
        return True

    def render_shape_box_tag(get_para, tag, tag_style):
        shape_spec = get_word_shape_box_spec(tag, tag_style)
        if not shape_spec:
            return False
        p = get_para()
        shape_element = build_word_shape_element(shape_spec)
        if shape_element is None:
            return False
        paragraph_style = dict(tag_style)
        for key in (
            "width",
            "height",
            "min-height",
            "background",
            "background-color",
            "border",
            "border-top",
            "border-right",
            "border-bottom",
            "border-left",
            "border-radius",
            "padding",
            "padding-top",
            "padding-right",
            "padding-bottom",
            "padding-left",
            "display",
            "justify-content",
            "align-items",
        ):
            paragraph_style.pop(key, None)
        p.clear()
        apply_paragraph_style(
            p,
            paragraph_style,
            current_font_pt=DEFAULT_FONT_PT,
            apply_box_visuals=False,
        )
        p.add_run()._r.append(shape_element)
        return True

    def render_inline(paragraph, node, parent_style, preserve_whitespace=False):
        if isinstance(node, Comment):
            return
        if isinstance(node, NavigableString):
            if preserve_whitespace:
                add_preformatted_text(paragraph, str(node), parent_style)
            else:
                text = normalize_text_content(str(node), preserve_whitespace=False)
                text = normalize_inline_text_for_paragraph(paragraph, text)
                if text.strip():
                    add_text_run(paragraph, text, parent_style)
            return
        if not isinstance(node, Tag):
            return
        if node.name.lower() == "br":
            paragraph.add_run().add_break()
            return
        if node.name.lower() == "img":
            render_image(node, parent_style, inline_paragraph=paragraph)
            return

        style = resolve_style(node, parent_style)

        # 标签语义补丁
        name = node.name.lower()
        if name == "pre":
            style = dict(style)
            style.setdefault("font-family", "Consolas")
            style.setdefault("white-space", "pre")
        if name in {"strong", "b"}:
            style["font-weight"] = "700"
        if name in {"em", "i"}:
            style["font-style"] = "italic"
        if name == "sup":
            style["vertical-align"] = "super"
        if name == "sub":
            style["vertical-align"] = "sub"
        if name == "u":
            style["text-decoration"] = "underline"
        if name in {"s", "strike", "del"}:
            style["text-decoration"] = "line-through"
        if name == "code":
            style.setdefault("font-family", "Consolas")

        current_font_pt = length_to_pt(style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        display_val = str(style.get("display", "")).strip().lower()
        margin_right_pt = length_to_pt(style.get("margin-right"), current_font_pt=current_font_pt) or 0.0
        shape_spec = get_word_shape_box_spec(node, style)
        if shape_spec:
            if render_word_shape_box(paragraph, shape_spec):
                next_inline_sibling = None
                for sibling in node.next_siblings:
                    if isinstance(sibling, Comment):
                        continue
                    if isinstance(sibling, NavigableString):
                        if normalize_text_content(str(sibling), preserve_whitespace=False).strip():
                            next_inline_sibling = None
                            break
                        continue
                    if isinstance(sibling, Tag):
                        if not is_block_like_tag(sibling):
                            next_inline_sibling = sibling
                        break
                if (margin_right_pt > 0 or display_val in {"inline-block", "inline-flex"}) and next_inline_sibling is not None:
                    if not paragraph_ends_with_break(paragraph):
                        gap_spaces = 1
                        if margin_right_pt > 0 and current_font_pt > 0:
                            approx_space_pt = max(2.5, current_font_pt * 0.42)
                            gap_spaces = max(1, min(4, int(round(margin_right_pt / approx_space_pt))))
                        add_text_run(paragraph, " " * gap_spaces, parent_style)
                return

        # 迁移 export_docx.js 规则：仅当自身背景与父元素不同才作为高亮。
        bg = style.get("background-color")
        if is_transparent_color(bg):
            style = dict(style)
            style.pop("background-color", None)

        padding_left_pt = length_to_pt(style.get("padding-left"), current_font_pt=current_font_pt) or 0.0
        padding_right_pt = length_to_pt(style.get("padding-right"), current_font_pt=current_font_pt) or 0.0
        wrap_inline_box = (
            name not in {"strong", "b", "em", "i", "u", "s", "strike", "del", "sup", "sub", "code"}
            and has_inline_box_visual(style, current_font_pt=current_font_pt)
        )

        if wrap_inline_box and padding_left_pt > 0:
            add_text_run(paragraph, " ", style)

        child_preserve_whitespace = preserve_whitespace or (name == "pre")
        for child in node.children:
            render_inline(paragraph, child, style, preserve_whitespace=child_preserve_whitespace)

        if wrap_inline_box and padding_right_pt > 0:
            add_text_run(paragraph, " ", style)
        if margin_right_pt > 0 or display_val in {"inline-block", "inline-flex"}:
            if not paragraph_ends_with_break(paragraph):
                gap_spaces = 1
                if margin_right_pt > 0 and current_font_pt > 0:
                    approx_space_pt = max(2.5, current_font_pt * 0.42)
                    gap_spaces = max(1, min(4, int(round(margin_right_pt / approx_space_pt))))
                add_text_run(paragraph, " " * gap_spaces, parent_style)

    def render_code_like_block(paragraph, block_node, block_style):
        """
        代码块专用渲染：
        - 仅以 <br> 作为换行依据
        - 忽略 HTML 模板缩进导致的噪音空白
        """
        def walk(node, style):
            if isinstance(node, Comment):
                return
            if isinstance(node, NavigableString):
                ws = str(style.get("white-space", "")).strip().lower()
                preserve_nl = ws in {"pre", "pre-wrap", "break-spaces", "pre-line"}
                if preserve_nl:
                    add_preformatted_text(paragraph, str(node), style)
                else:
                    # 非 pre 语义：忽略模板缩进噪音，主要以 <br> 控制换行。
                    txt = str(node).replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
                    if not txt.strip():
                        return
                    clean = re.sub(r"\s+", " ", txt)
                    add_text_run(paragraph, clean, style)
                return
            if not isinstance(node, Tag):
                return
            name = node.name.lower()
            if name == "br":
                paragraph.add_run().add_break()
                return

            child_style = resolve_style(node, style)
            if name == "pre":
                child_style = dict(child_style)
                child_style.setdefault("white-space", "pre")
            if name in {"strong", "b"}:
                child_style["font-weight"] = "700"
            if name in {"em", "i"}:
                child_style["font-style"] = "italic"
            if name == "sup":
                child_style["vertical-align"] = "super"
            if name == "sub":
                child_style["vertical-align"] = "sub"
            if name == "u":
                child_style["text-decoration"] = "underline"
            if name in {"s", "strike", "del"}:
                child_style["text-decoration"] = "line-through"
            if name == "code":
                child_style.setdefault("font-family", "Consolas")

            for c in node.children:
                walk(c, child_style)

        for c in block_node.children:
            walk(c, block_style)

    def render_common_paragraph(tag, parent_style, bold_default=False, left_indent_cm=None, heading_level=None):
        style = resolve_style(tag, parent_style)
        name = tag.name.lower() if getattr(tag, "name", None) else ""
        # 轻量链路没有浏览器 UA 样式，这里补齐 h1~h6 常见默认值兜底。
        heading_defaults = {
            "h1": {"font-size": "24pt", "margin-top": "12pt", "margin-bottom": "8pt"},
            "h2": {"font-size": "18pt", "margin-top": "10pt", "margin-bottom": "7pt"},
            "h3": {"font-size": "14pt", "margin-top": "8pt", "margin-bottom": "6pt"},
            "h4": {"font-size": "12pt", "margin-top": "7pt", "margin-bottom": "5pt"},
            "h5": {"font-size": "10pt", "margin-top": "6pt", "margin-bottom": "4pt"},
            "h6": {"font-size": "9pt", "margin-top": "5pt", "margin-bottom": "4pt"},
        }
        if name in heading_defaults:
            rule_style = rule_map.get(id(tag), {}) or {}
            inline_style = normalize_declarations(parse_inline_style(tag.get("style")))

            # font-size 是可继承属性，style 里常会带父级字号；仅当标题自身未声明时才覆盖继承值。
            if "font-size" not in rule_style and "font-size" not in inline_style:
                style["font-size"] = heading_defaults[name]["font-size"]
            # margin 非继承属性，兜底仅在缺失时补齐。
            style.setdefault("margin-top", heading_defaults[name]["margin-top"])
            style.setdefault("margin-bottom", heading_defaults[name]["margin-bottom"])
            if "font-weight" not in rule_style and "font-weight" not in inline_style:
                style["font-weight"] = "700"
        para = doc.add_paragraph()
        if heading_level is not None:
            apply_heading_paragraph_semantics(para, heading_level, doc)
        current_font_pt = length_to_pt(style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        apply_paragraph_style(para, style, current_font_pt=current_font_pt)
        if left_indent_cm is not None:
            para.paragraph_format.left_indent = Cm(left_indent_cm)
        for child in tag.children:
            if bold_default and isinstance(child, NavigableString):
                text = normalize_text_content(str(child), preserve_whitespace=False)
                text = normalize_inline_text_for_paragraph(para, text)
                if text.strip():
                    bold_style = dict(style)
                    bold_style["font-weight"] = "700"
                    add_text_run(para, text, bold_style)
            else:
                render_inline(para, child, style, preserve_whitespace=False)
            if isinstance(child, Tag):
                child_style = resolve_style(child, style)
                display_val = str(child_style.get("display", "")).strip().lower()
                if display_val in {"inline-block", "inline-flex"}:
                    margin_right_pt = length_to_pt(child_style.get("margin-right"), current_font_pt=current_font_pt) or 0.0
                    if margin_right_pt > 0:
                        para.add_run(" ")

    def get_list_style_name(ordered, level):
        base = "List Number" if ordered else "List Bullet"
        if level <= 0:
            return base
        # 迁移主方案规则：最多映射到 3 级内置样式，超出级别走 fallback。
        suffix_level = min(level + 1, 3)
        return f"{base} {suffix_level}"

    list_style_abstract_num_cache = {}
    custom_bullet_abstract_num_cache = {}

    def get_default_unordered_list_type(level):
        # 近似浏览器默认 UA 样式：一级实心圆、二级空心圆、三级方块。
        default_types = ["disc", "circle", "square"]
        return default_types[min(max(level, 0), len(default_types) - 1)]

    def resolve_unordered_list_type(raw_list_type, level):
        text = str(raw_list_type or "").strip().lower()
        if text in {"disc", "circle", "square"}:
            return text
        if text == "none":
            return "none"
        return get_default_unordered_list_type(level)

    def get_bullet_marker_spec(list_type):
        bullet_specs = {
            "disc": ("•", "Microsoft YaHei"),
            "circle": ("◦", "Microsoft YaHei"),
            "square": ("▪", "Microsoft YaHei"),
        }
        return bullet_specs.get(list_type, bullet_specs["disc"])

    def get_bullet_marker_font_pt(base_font_pt):
        """
        WPS/Word 对自定义 bullet 若未显式写字号，常会显示得比正文更小。
        这里统一给 marker 一个略大于正文的字号，尽量贴近浏览器默认圆点观感。
        """
        try:
            base = float(base_font_pt)
        except (TypeError, ValueError):
            base = DEFAULT_FONT_PT
        return max(DEFAULT_FONT_PT + 2.0, base + 1.5)

    def get_next_abstract_num_id(numbering):
        max_id = -1
        for node in numbering.xpath("./w:abstractNum"):
            raw = node.get(qn("w:abstractNumId"))
            try:
                max_id = max(max_id, int(raw))
            except (TypeError, ValueError):
                continue
        return max_id + 1

    def create_bullet_abstract_num_id(list_type, base_font_pt=DEFAULT_FONT_PT):
        marker_text, font_name = get_bullet_marker_spec(list_type)
        marker_font_pt = get_bullet_marker_font_pt(base_font_pt)
        marker_size_half_pt = max(2, int(round(marker_font_pt * 2.0)))
        cache_key = (marker_text, font_name, marker_size_half_pt)
        cached = custom_bullet_abstract_num_cache.get(cache_key)
        if cached is not None:
            return cached

        numbering = doc.part.numbering_part.numbering_definitions._numbering
        abstract_num_id = get_next_abstract_num_id(numbering)

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

        nsid = OxmlElement("w:nsid")
        nsid.set(qn("w:val"), f"{0xA0B00000 + abstract_num_id:08X}"[-8:])
        abstract_num.append(nsid)

        multi_level_type = OxmlElement("w:multiLevelType")
        multi_level_type.set(qn("w:val"), "singleLevel")
        abstract_num.append(multi_level_type)

        tmpl = OxmlElement("w:tmpl")
        tmpl.set(qn("w:val"), f"{0xC0D00000 + abstract_num_id:08X}"[-8:])
        abstract_num.append(tmpl)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet")
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker_text)
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "360")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:cs"), font_name)
        r_pr.append(r_fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(marker_size_half_pt))
        r_pr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(marker_size_half_pt))
        r_pr.append(sz_cs)
        lvl.append(r_pr)

        abstract_num.append(lvl)

        inserted = False
        for idx, child in enumerate(numbering):
            if child.tag == qn("w:num"):
                numbering.insert(idx, abstract_num)
                inserted = True
                break
        if not inserted:
            numbering.append(abstract_num)

        custom_bullet_abstract_num_cache[cache_key] = abstract_num_id
        return abstract_num_id

    def get_list_style_abstract_num_id(style_name):
        cached = list_style_abstract_num_cache.get(style_name)
        if cached is not None:
            return cached

        try:
            style = doc.styles[style_name]
        except KeyError:
            return None

        num_id_nodes = style._element.xpath("./w:pPr/w:numPr/w:numId")
        if not num_id_nodes:
            return None

        try:
            base_num_id = int(num_id_nodes[0].get(qn("w:val")))
        except (TypeError, ValueError):
            return None

        numbering = doc.part.numbering_part.numbering_definitions._numbering
        try:
            num = numbering.num_having_numId(base_num_id)
        except KeyError:
            return None

        abstract_num_nodes = num.xpath("./w:abstractNumId")
        if not abstract_num_nodes:
            return None

        try:
            abstract_num_id = int(abstract_num_nodes[0].get(qn("w:val")))
        except (TypeError, ValueError):
            return None

        list_style_abstract_num_cache[style_name] = abstract_num_id
        return abstract_num_id

    def create_list_numbering_instance(style_name, start_at=1):
        abstract_num_id = get_list_style_abstract_num_id(style_name)
        if abstract_num_id is None:
            return None

        numbering = doc.part.numbering_part.numbering_definitions._numbering
        num = numbering.add_num(abstract_num_id)
        lvl_override = num.add_lvlOverride(ilvl=0)
        lvl_override.add_startOverride(max(1, int(start_at)))
        return int(num.get(qn("w:numId")))

    def create_bullet_numbering_instance(list_type, base_font_pt=DEFAULT_FONT_PT):
        abstract_num_id = create_bullet_abstract_num_id(list_type, base_font_pt=base_font_pt)
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        num = numbering.add_num(abstract_num_id)
        return int(num.get(qn("w:numId")))

    def set_paragraph_numbering(paragraph, num_id, level=0):
        if num_id is None:
            return
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = max(0, int(level))
        num_pr.get_or_add_numId().val = int(num_id)

    def render_list(list_tag, parent_style, level=0, add_para_with_style=None):
        list_name = list_tag.name.lower()
        ordered = list_name == "ol"
        unordered_list_indent_adjust_pt = -5.0 if list_name == "ul" else 0.0
        list_style = resolve_style(list_tag, parent_style)
        raw_list_type = str(list_style.get("list-style-type", "")).strip().lower()
        list_type = raw_list_type if ordered else resolve_unordered_list_type(raw_list_type, level)
        no_marker = list_type == "none"
        list_start = 1
        if ordered:
            try:
                list_start = max(1, int(list_tag.get("start", 1)))
            except (TypeError, ValueError):
                list_start = 1
        item_no = list_start

        # 轻量链路没有浏览器 UA 样式；作者未显式覆盖 padding-left 时补浏览器默认列表内边距。
        if (
            not no_marker
            and "padding-left" not in list_style
            and "padding" not in list_style
        ):
            list_style = dict(list_style)
            list_style["padding-left"] = "40px"

        def apply_list_paragraph_indent(para, li_style, manual_marker=False):
            current_font_pt = length_to_pt(li_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
            margin_left_pt = length_to_pt(list_style.get("margin-left"), current_font_pt=current_font_pt) or 0.0
            padding_left_pt = length_to_pt(list_style.get("padding-left"), current_font_pt=current_font_pt) or 0.0
            li_margin_left_pt = length_to_pt(li_style.get("margin-left"), current_font_pt=current_font_pt) or 0.0
            li_padding_left_pt = length_to_pt(li_style.get("padding-left"), current_font_pt=current_font_pt) or 0.0

            if ordered:
                # 有序列表更容易在 Word/WPS 里显得“文本之前”过大，取保守基准。
                container_indent_pt = max(
                    margin_left_pt,
                    padding_left_pt,
                    li_margin_left_pt,
                    li_padding_left_pt,
                )
            else:
                # 无序列表更依赖可见层次，保留 margin + padding 的综合效果，
                # 避免卡片/提示框内的圆点列表看起来几乎和正文齐平。
                container_indent_pt = max(
                    margin_left_pt + padding_left_pt,
                    li_margin_left_pt + li_padding_left_pt,
                )
            effective_container_indent_pt = max(
                0.0,
                container_indent_pt + (0.0 if manual_marker else unordered_list_indent_adjust_pt),
            )
            # HTML 的嵌套列表依赖块级上下文产生层级感；落到 DOCX 后会被拍平成连续段落，
            # 每一级需要额外叠加一段“容器缩进步长”，否则二级列表会显得几乎没缩进。
            level_step_pt = max(18.0, effective_container_indent_pt)
            nested_tighten_pt = (NESTED_UNORDERED_LIST_TIGHTEN_PT * level) if not ordered else 0.0
            if manual_marker:
                total_indent_pt = max(0.0, effective_container_indent_pt + level_step_pt * level - nested_tighten_pt)
                hanging_pt = min(total_indent_pt, max(12.0, li_padding_left_pt, padding_left_pt))
                if total_indent_pt > 0:
                    para.paragraph_format.left_indent = Pt(total_indent_pt)
                    para.paragraph_format.first_line_indent = Pt(-hanging_pt)
                return
            if no_marker:
                total_indent_pt = max(0.0, effective_container_indent_pt + level_step_pt * level - nested_tighten_pt)
                if total_indent_pt > 0:
                    para.paragraph_format.left_indent = Pt(total_indent_pt)
                return

            # 显式写入列表缩进，避免 WPS/不同模板仅依赖 List Bullet 内置样式时丢失左边距。
            hanging_pt = 18.0
            level_indent_pt = level_step_pt * level
            para.paragraph_format.left_indent = Pt(
                max(0.0, effective_container_indent_pt + level_indent_pt + hanging_pt - nested_tighten_pt)
            )
            para.paragraph_format.first_line_indent = Pt(-hanging_pt)

        def render_list_item_to_paragraph(li_tag, li_style, para):
            apply_paragraph_style(para, li_style, current_font_pt=DEFAULT_FONT_PT)
            apply_list_paragraph_indent(para, li_style)
            for child in li_tag.children:
                if isinstance(child, Tag) and child.name and child.name.lower() in {"ul", "ol"}:
                    continue
                render_inline(para, child, li_style, preserve_whitespace=False)

        def get_manual_marker_data(li_tag, li_style):
            before_style = before_rule_map.get(id(li_tag), {}) or {}
            marker_text = parse_css_content_text(before_style.get("content"))
            li_list_type = str(li_style.get("list-style-type", "")).strip().lower()
            use_manual_marker = bool(marker_text) and (no_marker or li_list_type == "none")
            if not use_manual_marker:
                return None, {}
            marker_style = dict(li_style)
            marker_style.update(before_style)
            marker_style.pop("content", None)
            return marker_text, marker_style

        # 目录等场景：list-style:none + column-count，多列近似为布局表格。
        if no_marker:
            col_count_raw = str(list_style.get("column-count", "")).strip()
            col_count = 1
            if col_count_raw.isdigit():
                col_count = max(1, int(col_count_raw))
            li_items = list_tag.find_all("li", recursive=False)
            if col_count > 1 and li_items:
                row_count = (len(li_items) + col_count - 1) // col_count
                table = doc.add_table(rows=row_count, cols=col_count)
                table.style = "Table Grid"
                table.autofit = False
                set_table_borders_none(table)
                set_table_cell_spacing(table, 0.0)
                base_twips = int((DOC_CONTENT_WIDTH_TWIPS * 0.92) / col_count)
                for idx, li in enumerate(li_items):
                    r = idx % row_count
                    c = idx // row_count
                    if c >= col_count:
                        break
                    li_style = resolve_style(li, list_style)
                    cell = table.rows[r].cells[c]
                    cell.text = ""
                    set_cell_width_twips(cell, base_twips)
                    para = cell.paragraphs[0]
                    para.clear()
                    render_list_item_to_paragraph(li, li_style, para)
                    for nested in li.find_all(["ul", "ol"], recursive=False):
                        def add_para_with_style_for_cell(_style_name):
                            return cell.add_paragraph()
                        render_list(nested, li_style, level=level + 1, add_para_with_style=add_para_with_style_for_cell)
                return

        list_num_id = None
        style_name = None
        if not no_marker:
            style_name = get_list_style_name(ordered, level)
            if ordered:
                list_num_id = create_list_numbering_instance(style_name, start_at=list_start)
            else:
                list_font_pt = length_to_pt(list_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
                list_num_id = create_bullet_numbering_instance(list_type, base_font_pt=list_font_pt)

        for li in list_tag.find_all("li", recursive=False):
            li_style = resolve_style(li, list_style)
            used_word_list_style = False
            manual_marker_text, manual_marker_style = get_manual_marker_data(li, li_style)
            manual_marker = bool(manual_marker_text)

            if no_marker or manual_marker:
                if add_para_with_style is not None:
                    para = add_para_with_style("Normal")
                else:
                    para = doc.add_paragraph()
            else:
                if add_para_with_style is not None:
                    para = add_para_with_style(style_name)
                else:
                    try:
                        para = doc.add_paragraph(style=style_name)
                        used_word_list_style = True
                    except Exception:
                        para = doc.add_paragraph()
                apply_paragraph_style(para, li_style, current_font_pt=DEFAULT_FONT_PT)
                apply_list_paragraph_indent(para, li_style)
                if list_num_id is not None:
                    set_paragraph_numbering(para, list_num_id, level=0)

                # 迁移主方案规则：当内置列表样式不可用时，手动前缀 + 缩进兜底。
                if list_num_id is None and not used_word_list_style:
                    para.paragraph_format.left_indent = Cm(0.63 * (level + 1))
                    if ordered:
                        prefix = f"{item_no}. "
                    else:
                        prefix = f"{get_bullet_marker_spec(list_type)[0]} "
                    add_text_run(para, prefix, li_style)

            if manual_marker:
                apply_paragraph_style(para, li_style, current_font_pt=DEFAULT_FONT_PT)
                apply_list_paragraph_indent(para, li_style, manual_marker=True)
                add_text_run(para, manual_marker_text, manual_marker_style)
                add_text_run(para, " ", li_style)
                for child in li.children:
                    if isinstance(child, Tag) and child.name and child.name.lower() in {"ul", "ol"}:
                        continue
                    render_inline(para, child, li_style, preserve_whitespace=False)
            elif no_marker:
                render_list_item_to_paragraph(li, li_style, para)
            else:
                for child in li.children:
                    if isinstance(child, Tag) and child.name and child.name.lower() in {"ul", "ol"}:
                        continue
                    render_inline(para, child, li_style, preserve_whitespace=False)

            for nested in li.find_all(["ul", "ol"], recursive=False):
                render_list(nested, li_style, level=level + 1, add_para_with_style=add_para_with_style)

            item_no += 1

    def render_table(tag, parent_style):
        style = resolve_style(tag, parent_style)
        outer_container_style = parent_style if has_container_box_visual(parent_style) else None
        rows = tag.find_all("tr")
        if not rows:
            return
        num_rows = len(rows)

        # 迁移主方案规则：列数需考虑 colspan 的展开后宽度。
        num_cols = 0
        parsed_rows = []
        for tr in rows:
            cells = tr.find_all(["th", "td"], recursive=False)
            parsed_rows.append(cells)
            row_col_count = 0
            for c in cells:
                try:
                    row_col_count += max(1, int(c.get("colspan", 1)))
                except ValueError:
                    row_col_count += 1
            num_cols = max(num_cols, row_col_count)

        if num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = resolve_table_alignment(style)
        set_table_layout_fixed(table)
        # 先清掉模板自带的网格边框，再按 HTML/CSS 的显式边框补回。
        set_table_borders_none(table)
        for edge in ("top", "left", "bottom", "right"):
            border = resolve_edge_border(style, edge)
            if border:
                set_table_border_edge(table, edge, border)

        # 迁移 export_docx.js 规则：border-collapse / border-spacing
        collapse_mode = str(style.get("border-collapse", "")).strip().lower()
        if collapse_mode != "collapse":
            spacing_px = parse_border_spacing_px(style.get("border-spacing"))
            set_table_cell_spacing(table, spacing_px)

        # 表格总宽度：优先 CSS width；否则默认使用页面宽度的 92%。
        available_content_twips = int(get_doc_content_width_twips())
        width_text = style.get("width")
        target_twips = int(available_content_twips * 0.92)
        if width_text:
            wt = str(width_text).strip().lower()
            if wt.endswith("%"):
                try:
                    target_twips = int(available_content_twips * float(wt[:-1]) / 100.0)
                except ValueError:
                    pass
            else:
                width_pt = length_to_pt(wt, current_font_pt=DEFAULT_FONT_PT)
                if width_pt and width_pt > 0:
                    target_twips = int(width_pt * 20)
        target_twips = max(2400, min(available_content_twips, target_twips))

        table_layout = str(style.get("table-layout", "")).strip().lower()
        has_spans = False
        has_explicit_cell_width = False
        for i, row in enumerate(rows):
            row_style = resolve_style(row, style)
            for cell_tag in parsed_rows[i]:
                try:
                    if max(1, int(cell_tag.get("colspan", 1))) > 1:
                        has_spans = True
                except ValueError:
                    pass
                try:
                    if max(1, int(cell_tag.get("rowspan", 1))) > 1:
                        has_spans = True
                except ValueError:
                    pass
                cell_style = resolve_style(cell_tag, row_style)
                if cell_style.get("width"):
                    has_explicit_cell_width = True
            if has_spans and has_explicit_cell_width:
                break

        prefer_even_columns = table_layout == "fixed" or (
            num_cols >= 4 and not has_spans and not has_explicit_cell_width
        )
        if prefer_even_columns:
            col_twips = split_evenly(target_twips, num_cols)
        else:
            col_weights = estimate_col_char_weights(rows, num_cols)
            total_weight = sum(col_weights) or float(num_cols)
            col_twips = [max(900, int(target_twips * w / total_weight)) for w in col_weights]
            if col_twips:
                col_twips[-1] += target_twips - sum(col_twips)
        set_table_width_twips(table, target_twips)
        set_table_grid_widths(table, col_twips)

        def apply_merged_region_outer_borders(start_row, start_col, end_row, end_col, merged_style):
            if end_row <= start_row and end_col <= start_col:
                return
            border_map = {
                edge: resolve_edge_border(merged_style, edge) or resolve_edge_border(style, edge)
                for edge in ("left", "right", "top", "bottom")
            }
            for ri in range(start_row, end_row + 1):
                for ci in range(start_col, end_col + 1):
                    if ri == start_row and ci == start_col:
                        continue
                    raw_tcs = getattr(table.rows[ri]._tr, "tc_lst", None)
                    if raw_tcs is None or len(raw_tcs) != num_cols or ci >= len(raw_tcs):
                        continue
                    merged_tc = raw_tcs[ci]
                    if ci == start_col and border_map["left"]:
                        set_tc_border_edge(merged_tc, "left", border_map["left"])
                    if ci == end_col and border_map["right"]:
                        set_tc_border_edge(merged_tc, "right", border_map["right"])
                    if ri == start_row and border_map["top"]:
                        set_tc_border_edge(merged_tc, "top", border_map["top"])
                    if ri == end_row and border_map["bottom"]:
                        set_tc_border_edge(merged_tc, "bottom", border_map["bottom"])

        def get_table_cell_paragraph_style(raw_cell_style):
            """
            td/th 的 padding 已经映射到 cell 层，不应再只落到第一个段落上。
            这里去掉 cell 自身的盒模型缩进，保证单元格内多个段落对齐一致。
            """
            paragraph_style = dict(raw_cell_style)
            for key in (
                "padding-top",
                "padding-right",
                "padding-bottom",
                "padding-left",
                "margin-top",
                "margin-right",
                "margin-bottom",
                "margin-left",
            ):
                paragraph_style.pop(key, None)
            return paragraph_style

        # 占用矩阵：记录哪些单元格被上方 rowspan 占用。
        occupied = [[False] * num_cols for _ in range(num_rows)]

        for i, row in enumerate(rows):
            row_style = resolve_style(row, style)
            col_idx = 0
            for cell_tag in parsed_rows[i]:
                while col_idx < num_cols and occupied[i][col_idx]:
                    col_idx += 1
                if col_idx >= num_cols:
                    break

                try:
                    colspan = max(1, int(cell_tag.get("colspan", 1)))
                except ValueError:
                    colspan = 1
                try:
                    rowspan = max(1, int(cell_tag.get("rowspan", 1)))
                except ValueError:
                    rowspan = 1

                end_col = min(num_cols - 1, col_idx + colspan - 1)
                end_row = min(num_rows - 1, i + rowspan - 1)

                cell_style = resolve_style(cell_tag, row_style)
                if outer_container_style:
                    cell_style = dict(cell_style)
                    outer_bg = outer_container_style.get("background-color")
                    if outer_bg and is_transparent_color(cell_style.get("background-color")):
                        cell_style["background-color"] = outer_bg
                cell_para_style = get_table_cell_paragraph_style(cell_style)
                cell = table.rows[i].cells[col_idx]
                if end_col > col_idx or end_row > i:
                    end_cell = table.rows[end_row].cells[end_col]
                    cell = cell.merge(end_cell)

                for ri in range(i, end_row + 1):
                    for ci in range(col_idx, end_col + 1):
                        if ri > i or ci > col_idx:
                            occupied[ri][ci] = True

                cell.text = ""
                span_width = sum(col_twips[col_idx : end_col + 1]) if col_twips else 0
                if span_width > 0:
                    set_cell_width_twips(cell, span_width)
                if cell_style.get("background-color") and not is_transparent_color(cell_style.get("background-color")):
                    set_cell_shading(cell, cell_style.get("background-color"))
                for edge in ("left", "right", "top", "bottom"):
                    border = resolve_edge_border(cell_style, edge)
                    if border:
                        set_cell_border_edge(cell, edge, border)
                apply_merged_region_outer_borders(i, col_idx, end_row, end_col, cell_style)
                # 老方案中单元格不会映射这么重的上下 padding，这里收敛以避免行高过大。
                set_cell_padding(
                    cell,
                    cell_style,
                    current_font_pt=DEFAULT_FONT_PT,
                    horizontal_scale=0.65,
                    vertical_scale=0.15,
                )
                set_cell_vertical_align(cell, cell_style.get("vertical-align"))
                para = cell.paragraphs[0]
                # 单元格的背景/边框已在 cell 层处理，段落层不再重复绘制，避免“内框”。
                apply_paragraph_style(
                    para,
                    cell_para_style,
                    current_font_pt=DEFAULT_FONT_PT,
                    apply_box_visuals=False,
                )
                if cell_tag.name.lower() == "th":
                    cell_style = dict(cell_style)
                    cell_style["font-weight"] = "700"
                first_para_used = False

                def get_para():
                    nonlocal first_para_used
                    if not first_para_used:
                        first_para_used = True
                        return para
                    return cell.add_paragraph()

                for child in cell_tag.children:
                    if isinstance(child, Tag):
                        child_name = child.name.lower()
                        child_style = resolve_style(child, cell_style)
                        if child_name in {"ul", "ol"}:
                            def add_para_with_style(_style_name):
                                return get_para()
                            render_list(child, child_style, level=0, add_para_with_style=add_para_with_style)
                            continue
                        if is_block_like_tag(child):
                            p = get_para()
                            apply_paragraph_style(
                                p,
                                child_style,
                                current_font_pt=DEFAULT_FONT_PT,
                                apply_box_visuals=False,
                            )
                            for n in child.children:
                                render_inline(p, n, child_style, preserve_whitespace=False)
                            continue
                        target_para = para if first_para_used else get_para()
                        render_inline(target_para, child, cell_style, preserve_whitespace=False)
                        continue

                    if isinstance(child, NavigableString) and child.strip():
                        target_para = para if first_para_used else get_para()
                        render_inline(target_para, child, cell_style, preserve_whitespace=False)
                col_idx = end_col + 1

    def has_effective_padding(tag_style):
        """
        仅当 padding 存在且为“非 0”时返回 True。
        避免被全局 reset（如 * { padding:0 }）误判为视觉容器。
        """
        zero_values = {
            "0",
            "0px",
            "0pt",
            "0em",
            "0rem",
            "0%",
            "0cm",
            "0mm",
            "0in",
        }
        for key in ("padding-top", "padding-right", "padding-bottom", "padding-left"):
            raw = tag_style.get(key)
            if raw is None:
                continue
            text = str(raw).strip().lower()
            if not text or text in zero_values:
                continue
            parsed_pt = length_to_pt(text, current_font_pt=DEFAULT_FONT_PT)
            if parsed_pt is not None:
                if parsed_pt > 0:
                    return True
                continue
            # 无法换算时，只要不是显式 0，就按有效 padding 处理。
            if text not in {"0 0", "0 0 0", "0 0 0 0"}:
                return True
        return False

    def has_effective_box_visual(tag_style):
        bg = tag_style.get("background-color")
        has_bg = not is_transparent_color(bg)
        has_border = any(resolve_edge_border(tag_style, edge) for edge in ("left", "right", "top", "bottom"))
        has_padding = has_effective_padding(tag_style)
        return has_bg or has_border or has_padding

    def get_left_layout_indent_pt(tag_style):
        current_font_pt = length_to_pt(tag_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        margin_left_pt = length_to_pt(tag_style.get("margin-left"), current_font_pt=current_font_pt) or 0.0
        padding_left_pt = length_to_pt(tag_style.get("padding-left"), current_font_pt=current_font_pt) or 0.0
        return padding_left_pt if padding_left_pt > 0 else margin_left_pt

    def has_container_box_visual(tag_style):
        """
        仅把真正需要“包一个盒子”的容器送入 1x1 table：
        - 背景色
        - 左/右边框
        - 多边边框（如 top+bottom、完整边框）
        单独 padding、单独 top/bottom 分隔线不再视为容器盒模型。
        """
        bg = tag_style.get("background-color")
        if not is_transparent_color(bg):
            return True
        borders = {edge: resolve_edge_border(tag_style, edge) for edge in ("left", "right", "top", "bottom")}
        if borders["left"] or borders["right"]:
            return True
        border_count = sum(1 for edge in borders.values() if edge)
        return border_count >= 2

    def has_only_bottom_border_rule(tag_style):
        """
        仅用于“分隔线容器”识别：
        - 无背景
        - 只有 bottom 边框（无 left/right/top）
        这种结构更适合按普通块渲染并把底边框挂到最后一段，
        而不是包成 1x1 table。
        """
        if not is_transparent_color(tag_style.get("background-color")):
            return False
        bottom = resolve_edge_border(tag_style, "bottom")
        if not bottom:
            return False
        if resolve_edge_border(tag_style, "top"):
            return False
        if resolve_edge_border(tag_style, "left"):
            return False
        if resolve_edge_border(tag_style, "right"):
            return False
        return True

    def is_hanging_indent_only_block(tag, style):
        """
        参考文献这类块通常只有 hanging indent（padding-left + negative text-indent），
        没有真正的盒模型视觉。此时直接映射成普通段落更准确，避免容器 table 把缩进吃两次。
        """
        if tag.name.lower() not in {"div", "section", "article", "main", "header", "footer"}:
            return False
        if not is_transparent_color(style.get("background-color")):
            return False
        if any(resolve_edge_border(style, edge) for edge in ("left", "right", "top", "bottom")):
            return False
        if any(
            getattr(child, "name", None)
            in {"p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "div", "section", "article", "main", "blockquote", "pre", "hr"}
            for child in tag.children
        ):
            return False

        current_font_pt = length_to_pt(style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        padding_left_pt = length_to_pt(style.get("padding-left"), current_font_pt=current_font_pt) or 0.0
        text_indent_pt = length_to_pt(style.get("text-indent"), current_font_pt=current_font_pt) or 0.0
        if padding_left_pt <= 0 or text_indent_pt >= 0:
            return False

        for key in ("padding-top", "padding-right", "padding-bottom"):
            extra_padding_pt = length_to_pt(style.get(key), current_font_pt=current_font_pt) or 0.0
            if extra_padding_pt > 0:
                return False
        return True

    def is_layout_indent_only_block(tag, style):
        """
        仅承担布局缩进语义的 wrapper，不应被误判成盒模型容器。
        典型例子：.operation-item { margin: 10pt 0; padding-left: 20px; }
        """
        if tag.name.lower() not in {"div", "section", "article", "main", "header", "footer"}:
            return False
        if not is_transparent_color(style.get("background-color")):
            return False
        if any(resolve_edge_border(style, edge) for edge in ("left", "right", "top", "bottom")):
            return False
        current_font_pt = length_to_pt(style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        if (length_to_pt(style.get("padding-top"), current_font_pt=current_font_pt) or 0.0) > 0:
            return False
        if (length_to_pt(style.get("padding-right"), current_font_pt=current_font_pt) or 0.0) > 0:
            return False
        if (length_to_pt(style.get("padding-bottom"), current_font_pt=current_font_pt) or 0.0) > 0:
            return False
        return get_left_layout_indent_pt(style) > 0

    def get_block_content_flow_style(tag, style):
        if not is_layout_indent_only_block(tag, style):
            return style
        content_style = dict(style)
        content_style["_layout-left-offset-pt"] = get_left_layout_indent_pt(style)
        content_style.pop("margin-left", None)
        content_style.pop("padding-left", None)
        return content_style

    def should_render_container_as_table(tag, style):
        # 迁移 export_docx.js parseContainerAsPanel 的约束，避免误包主容器。
        classes = [str(c).strip().lower() for c in (tag.get("class") or []) if str(c).strip()]
        id_text = str(tag.get("id", "")).strip().lower()
        # 仅排除“主容器”语义，不误伤 code-container / feature-container 等业务容器。
        if "container" in classes or id_text == "container":
            if not is_code_like_block(tag, style):
                return False
        # 含真实表格/分隔线的外层容器不再整体包成 1x1 table，
        # 否则其内部复杂块布局会在单元格上下文里被压扁。
        if tag.find(["table", "hr"]) is not None:
            return False
        # 对图片做更保守的放行：仅给简历这类主卡片容器保留包裹能力。
        if tag.find("img") is not None and tag.name.lower() not in {"main"} and "resume" not in classes:
            return False
        if is_hanging_indent_only_block(tag, style):
            return False
        if is_layout_indent_only_block(tag, style):
            return False
        return has_container_box_visual(style)

    def render_simple_layout_child_as_tabbed_line(layout_tag, layout_style, get_para, available_twips):
        """
        轻量兜底：把单行 flex/grid 容器降级为“一个段落 + 多个制表位”。
        主要用于简历这类时间/单位/岗位、姓名/意向等轻量多列行。
        """
        display_value = str(layout_style.get("display", "")).strip().lower()
        if display_value not in {"flex", "grid"}:
            return False

        children = [c for c in layout_tag.children if isinstance(c, Tag)]
        if len(children) < 2 or len(children) > 3:
            return False

        # 仅处理轻量行布局，复杂块级内容仍交给后续通用渲染。
        heavy_tags = {"table", "ul", "ol", "section", "article", "header", "footer"}
        for child in children:
            if child.find(list(heavy_tags)) is not None:
                return False

        line_block_tags = {"div", "p", "h1", "h2", "h3", "h4", "h5", "h6"}

        def collect_child_lines(child_tag, parent_for_child_style):
            child_style = resolve_style(child_tag, parent_for_child_style)
            direct_blocks = [c for c in child_tag.children if isinstance(c, Tag) and c.name.lower() in line_block_tags]
            line_items = []

            def has_complex_nested_blocks(tag):
                for gc in tag.children:
                    if not isinstance(gc, Tag):
                        continue
                    name = gc.name.lower()
                    if name in heavy_tags or name in {"main", "blockquote", "pre", "hr"}:
                        return True
                    if name in line_block_tags:
                        return True
                return False

            if direct_blocks:
                if any(is_meaningful_text_node(c) for c in child_tag.children):
                    return None
                for block in direct_blocks:
                    if has_complex_nested_blocks(block):
                        return None
                    line_items.append((block, resolve_style(block, child_style)))
                return line_items

            if has_complex_nested_blocks(child_tag):
                return None
            return [(child_tag, child_style)]

        child_line_groups = []
        for child in children:
            lines = collect_child_lines(child, layout_style)
            if not lines:
                return False
            child_line_groups.append(lines)

        target_twips = max(2400, int(available_twips))
        width_twips = []

        if display_value == "grid":
            parsed_widths = parse_grid_template_column_twips(layout_style.get("grid-template-columns"), target_twips)
            if len(parsed_widths) >= len(children):
                width_twips = parsed_widths[: len(children)]

        if not width_twips:
            weights = []
            for child in children:
                child_style = resolve_style(child, layout_style)
                explicit_width_pt = length_to_pt(child_style.get("width"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                if explicit_width_pt > 0:
                    weights.append(explicit_width_pt)
                    continue
                text_len = len(child.get_text(" ", strip=True))
                weights.append(max(1.0, min(8.0, text_len * 0.18)))
            total_weight = sum(weights) or float(len(children))
            width_twips = [max(720, int(target_twips * w / total_weight)) for w in weights]
            if width_twips:
                width_twips[-1] += target_twips - sum(width_twips)

        if len(width_twips) != len(children):
            even_twips = max(720, int(target_twips / len(children)))
            width_twips = [even_twips] * len(children)
            width_twips[-1] += target_twips - sum(width_twips)

        p = get_para()
        apply_paragraph_style(p, layout_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if len(children) == 2:
            right_tab_pt = max(72.0, (target_twips - 120) / 20.0)
            p.paragraph_format.tab_stops.add_tab_stop(Pt(right_tab_pt), WD_TAB_ALIGNMENT.RIGHT)
        else:
            tab_pos_twips = 0
            for idx, part_twips in enumerate(width_twips[:-1]):
                tab_pos_twips += max(720, part_twips)
                p.paragraph_format.tab_stops.add_tab_stop(Pt(tab_pos_twips / 20.0), WD_TAB_ALIGNMENT.LEFT)

        max_lines = max(len(lines) for lines in child_line_groups)
        for line_idx in range(max_lines):
            if line_idx > 0:
                p.add_run().add_break()
            for child_idx, line_group in enumerate(child_line_groups):
                if line_idx < len(line_group):
                    line_tag, line_style = line_group[line_idx]
                    for n in line_tag.children:
                        render_inline(p, n, line_style, preserve_whitespace=False)
                has_later_content = any(line_idx < len(other_group) for other_group in child_line_groups[child_idx + 1 :])
                if has_later_content:
                    p.add_run("\t")
        return True

    def apply_container_box_visual_to_cell_style(cell_style, container_style, row_idx, col_idx, row_count, col_count):
        merged = dict(cell_style)
        bg = container_style.get("background-color")
        if bg and is_transparent_color(merged.get("background-color")):
            merged["background-color"] = bg

        for edge in ("top", "right", "bottom", "left"):
            base_pt = length_to_pt(merged.get(f"padding-{edge}"), current_font_pt=DEFAULT_FONT_PT) or 0.0
            outer_pt = length_to_pt(container_style.get(f"padding-{edge}"), current_font_pt=DEFAULT_FONT_PT) or 0.0
            if outer_pt <= 0:
                continue
            if edge == "top" and row_idx == 0:
                merged[f"padding-{edge}"] = f"{base_pt + outer_pt:.2f}pt"
            elif edge == "bottom" and row_idx == row_count - 1:
                merged[f"padding-{edge}"] = f"{base_pt + outer_pt:.2f}pt"
            elif edge == "left" and col_idx == 0:
                merged[f"padding-{edge}"] = f"{base_pt + outer_pt:.2f}pt"
            elif edge == "right" and col_idx == col_count - 1:
                merged[f"padding-{edge}"] = f"{base_pt + outer_pt:.2f}pt"

        return merged

    def is_inline_box_sequence_container(tag, tag_style):
        inline_children = [c for c in tag.children if isinstance(c, Tag)]
        if not inline_children or any(is_meaningful_text_node(c) for c in tag.children):
            return False
        current_font_pt = length_to_pt(tag_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        seen_box_child = False
        for child in inline_children:
            if is_block_like_tag(child):
                return False
            child_style = resolve_style(child, tag_style)
            if has_inline_box_visual(child_style, current_font_pt=current_font_pt):
                seen_box_child = True
        return seen_box_child

    def should_render_inline_box_sequence_as_shapes(tag, tag_style):
        if not is_inline_box_sequence_container(tag, tag_style):
            return False
        inline_children = [c for c in tag.children if isinstance(c, Tag)]
        if len(inline_children) < 1:
            return False
        for child in inline_children:
            child_style = resolve_style(child, tag_style)
            if get_word_shape_box_spec(child, child_style) is None:
                return False
        return True

    def render_inline_box_sequence_as_table(tag, tag_style, target_cell, available_twips):
        if target_cell is None or not is_inline_box_sequence_container(tag, tag_style):
            return False
        if should_render_inline_box_sequence_as_shapes(tag, tag_style):
            return False

        inline_children = [c for c in tag.children if isinstance(c, Tag)]
        if not inline_children:
            return False

        target_twips = max(1800, int(available_twips))
        col_twips = []
        child_styles = []

        def estimate_text_units(text):
            units = 0.0
            for ch in text:
                if ch.isspace():
                    units += 0.35
                elif ord(ch) < 128:
                    units += 0.62
                else:
                    units += 1.0
            return max(1.0, units)

        for child in inline_children:
            child_style = resolve_style(child, tag_style)
            child_styles.append(child_style)
            current_font_pt = length_to_pt(child_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
            explicit_width_pt = length_to_pt(child_style.get("width"), current_font_pt=DEFAULT_FONT_PT) or 0.0
            if explicit_width_pt > 0:
                cell_twips = int(explicit_width_pt * 20)
            else:
                text = child.get_text(" ", strip=True)
                text_width_pt = estimate_text_units(text) * current_font_pt * 0.92
                inner_padding_pt = (
                    (length_to_pt(child_style.get("padding-left"), current_font_pt=current_font_pt) or 0.0)
                    + (length_to_pt(child_style.get("padding-right"), current_font_pt=current_font_pt) or 0.0)
                )
                border_extra_pt = 0.0
                for edge in ("left", "right"):
                    border = resolve_edge_border(child_style, edge)
                    if border:
                        border_extra_pt += border.get("width_pt", 0.0) or 0.0
                cell_twips = int((text_width_pt + inner_padding_pt + border_extra_pt + 6.0) * 20)
            col_twips.append(max(560, cell_twips))

        total_table_twips = sum(col_twips)
        if total_table_twips > target_twips:
            scale = target_twips / float(total_table_twips or 1)
            col_twips = [max(560, int(width * scale)) for width in col_twips]
            if col_twips:
                col_twips[-1] += target_twips - sum(col_twips)
            total_table_twips = sum(col_twips)

        table = target_cell.add_table(rows=1, cols=len(inline_children))
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_borders_none(table)
        set_table_layout_fixed(table)
        set_table_cell_spacing(table, 2.0)
        set_table_width_twips(table, max(1800, total_table_twips))
        set_table_grid_widths(table, col_twips)

        for idx, (child, child_style) in enumerate(zip(inline_children, child_styles)):
            cell = table.rows[0].cells[idx]
            cell.text = ""
            set_cell_width_twips(cell, col_twips[idx] if idx < len(col_twips) else 0)
            if child_style.get("background-color") and not is_transparent_color(child_style.get("background-color")):
                set_cell_shading(cell, child_style.get("background-color"))
            for edge in ("left", "right", "top", "bottom"):
                border = resolve_edge_border(child_style, edge)
                if border:
                    set_cell_border_edge(cell, edge, border)
            set_cell_padding(cell, child_style, current_font_pt=DEFAULT_FONT_PT, horizontal_scale=1.0, vertical_scale=0.8)
            p = cell.paragraphs[0]
            p.clear()
            paragraph_style = dict(child_style)
            for key in (
                "margin-top",
                "margin-right",
                "margin-bottom",
                "margin-left",
                "padding-top",
                "padding-right",
                "padding-bottom",
                "padding-left",
                "background-color",
                "border",
                "border-top",
                "border-right",
                "border-bottom",
                "border-left",
            ):
                paragraph_style.pop(key, None)
            apply_paragraph_style(p, paragraph_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
            for n in child.children:
                render_inline(p, n, child_style, preserve_whitespace=False)
        return True

    def has_nested_block_children(tag):
        block_names = {
            "p", "h1", "h2", "h3", "h4", "h5", "h6",
            "ul", "ol", "table", "div", "section", "article", "main",
            "blockquote", "pre", "hr", "header", "footer",
        }
        return any(getattr(c, "name", None) in block_names for c in tag.children)

    def render_nested_block_children_in_cell(block_tag, block_style, get_para, available_twips, target_cell=None):
        block_style = get_block_content_flow_style(block_tag, block_style)
        current_inline_para = None

        for child in block_tag.children:
            handled, current_inline_para = render_inline_flow_child(
                child,
                block_style,
                current_inline_para,
                get_para,
                apply_box_visuals=False,
            )
            if handled:
                continue
            if not isinstance(child, Tag):
                continue
            current_inline_para = None

            child_name = child.name.lower()
            child_style = resolve_style(child, block_style)

            if child_name in {"ul", "ol"}:
                def add_para_with_style(_style_name):
                    p = get_para()
                    p.clear()
                    return p
                render_list(child, child_style, level=0, add_para_with_style=add_para_with_style)
                continue

            if render_inline_box_sequence_as_table(child, child_style, target_cell, available_twips):
                continue

            if is_code_like_block(child, child_style):
                p = get_para()
                apply_paragraph_style(p, child_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
                render_code_like_block(p, child, child_style)
                continue

            if is_block_like_tag(child):
                if render_shape_box_tag(get_para, child, child_style):
                    continue
                if render_simple_layout_child_as_tabbed_line(child, child_style, get_para, available_twips):
                    continue
                if child_name in {"div", "section", "article", "main", "header", "footer"} and not has_effective_box_visual(child_style) and has_nested_block_children(child):
                    render_nested_block_children_in_cell(
                        child,
                        child_style,
                        get_para,
                        available_twips,
                        target_cell=target_cell,
                    )
                    continue
                p = get_para()
                if child_name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                    try:
                        heading_level = int(child_name[1])
                    except ValueError:
                        heading_level = None
                    apply_heading_paragraph_semantics(p, heading_level, doc)
                apply_paragraph_style(
                    p,
                    child_style,
                    current_font_pt=DEFAULT_FONT_PT,
                    apply_box_visuals=has_effective_box_visual(child_style),
                )
                for n in child.children:
                    render_inline(p, n, child_style, preserve_whitespace=False)
                continue

            p = get_para()
            render_inline(p, child, block_style, preserve_whitespace=False)

    def render_grid_container_as_table(tag, parent_style, target_cell=None, target_width_twips=None):
        """迁移 export_docx.js 的思路：grid 容器降级为布局表格。"""
        style = resolve_style(tag, parent_style)
        available_twips = int(target_width_twips) if target_width_twips else int(get_doc_content_width_twips())
        width_text = style.get("width")
        target_twips = available_twips
        if width_text:
            wt = str(width_text).strip().lower()
            if wt.endswith("%"):
                try:
                    target_twips = int(available_twips * float(wt[:-1]) / 100.0)
                except ValueError:
                    target_twips = available_twips
            else:
                width_pt = length_to_pt(wt, current_font_pt=DEFAULT_FONT_PT)
                if width_pt and width_pt > 0:
                    target_twips = int(width_pt * 20)
        target_twips = max(2400, min(available_twips, target_twips))
        col_twips_list = parse_grid_template_column_twips(style.get("grid-template-columns"), target_twips)
        col_count = max(1, len(col_twips_list))
        children = [c for c in tag.children if isinstance(c, Tag)]
        if col_count <= 1 or len(children) <= 1:
            return False

        num_rows = (len(children) + col_count - 1) // col_count
        if target_cell is None:
            table = doc.add_table(rows=num_rows, cols=col_count)
        else:
            table = target_cell.add_table(rows=num_rows, cols=col_count)
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_borders_none(table)
        set_table_layout_fixed(table)

        # 不使用 tblCellSpacing（会给表格外沿也加空隙），改为：
        # 1) 先从总宽度中扣掉列间距
        # 2) 用 cell padding 模拟 gap 视觉
        row_gap_pt = 0.0
        col_gap_pt = 0.0
        gap_val = style.get("gap")
        if gap_val:
            parts = [p for p in str(gap_val).strip().split() if p]
            if len(parts) == 1:
                g = length_to_pt(parts[0], current_font_pt=DEFAULT_FONT_PT)
                if g:
                    row_gap_pt = col_gap_pt = max(0.0, g)
            elif len(parts) >= 2:
                g_row = length_to_pt(parts[0], current_font_pt=DEFAULT_FONT_PT)
                g_col = length_to_pt(parts[1], current_font_pt=DEFAULT_FONT_PT)
                if g_row:
                    row_gap_pt = max(0.0, g_row)
                if g_col:
                    col_gap_pt = max(0.0, g_col)
        col_gap_override = length_to_pt(style.get("column-gap"), current_font_pt=DEFAULT_FONT_PT)
        row_gap_override = length_to_pt(style.get("row-gap"), current_font_pt=DEFAULT_FONT_PT)
        if col_gap_override is not None:
            col_gap_pt = max(0.0, col_gap_override)
        if row_gap_override is not None:
            row_gap_pt = max(0.0, row_gap_override)

        # gap 仅通过 cell padding 模拟，不再额外从总宽度中扣减，避免整体比正文窄一截。
        set_table_width_twips(table, target_twips)
        set_table_grid_widths(table, col_twips_list)

        def add_padding_pt(base_style, key, delta_pt):
            if delta_pt <= 0:
                return
            old_pt = length_to_pt(base_style.get(key), current_font_pt=DEFAULT_FONT_PT) or 0.0
            base_style[key] = f"{max(0.0, old_pt + delta_pt):.2f}pt"

        for idx, child in enumerate(children):
            r = idx // col_count
            c = idx % col_count
            cell = table.rows[r].cells[c]
            child_style = resolve_style(child, style)
            cell_style = apply_container_box_visual_to_cell_style(
                child_style, style, r, c, num_rows, col_count
            )
            if col_gap_pt > 0:
                if c > 0:
                    add_padding_pt(cell_style, "padding-left", col_gap_pt * 0.5)
                if c < col_count - 1:
                    add_padding_pt(cell_style, "padding-right", col_gap_pt * 0.5)
            if row_gap_pt > 0:
                if r > 0:
                    add_padding_pt(cell_style, "padding-top", row_gap_pt * 0.5)
                if r < num_rows - 1:
                    add_padding_pt(cell_style, "padding-bottom", row_gap_pt * 0.5)
            width_twips = col_twips_list[c] if c < len(col_twips_list) else 0
            set_cell_width_twips(cell, width_twips)
            cell.text = ""
            if cell_style.get("background-color") and not is_transparent_color(cell_style.get("background-color")):
                set_cell_shading(cell, cell_style.get("background-color"))
            for edge in ("left", "right", "top", "bottom"):
                border = resolve_edge_border(cell_style, edge)
                if not border:
                    if edge == "left" and c == 0:
                        border = resolve_edge_border(style, edge)
                    elif edge == "right" and c == col_count - 1:
                        border = resolve_edge_border(style, edge)
                    elif edge == "top" and r == 0:
                        border = resolve_edge_border(style, edge)
                    elif edge == "bottom" and r == num_rows - 1:
                        border = resolve_edge_border(style, edge)
                if border:
                    set_cell_border_edge(cell, edge, border)
            set_cell_padding(
                cell,
                cell_style,
                current_font_pt=DEFAULT_FONT_PT,
                horizontal_scale=0.7,
                vertical_scale=0.25,
            )
            set_cell_vertical_align(cell, cell_style.get("vertical-align"))

            para = cell.paragraphs[0]
            para.clear()
            apply_paragraph_style(para, cell_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
            first_para_used = False

            def get_para():
                nonlocal first_para_used
                if not first_para_used:
                    first_para_used = True
                    return para
                return cell.add_paragraph()

            current_inline_para = None
            for gc in child.children:
                handled, current_inline_para = render_inline_flow_child(
                    gc,
                    child_style,
                    current_inline_para,
                    get_para,
                    apply_box_visuals=False,
                )
                if handled:
                    continue
                if not isinstance(gc, Tag):
                    continue
                current_inline_para = None
                gc_name = gc.name.lower()
                gc_style = resolve_style(gc, child_style)
                if gc_name in {"ul", "ol"}:
                    def add_para_with_style(_style_name):
                        p = get_para()
                        p.clear()
                        return p
                    render_list(gc, gc_style, level=0, add_para_with_style=add_para_with_style)
                elif is_code_like_block(gc, gc_style):
                    p = get_para()
                    p.clear()
                    apply_paragraph_style(p, gc_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
                    render_code_like_block(p, gc, gc_style)
                else:
                    # inline 元素（如 span/br）不要拆成新段，避免卡片/代码块出现大空行。
                    if is_block_like_tag(gc):
                        if render_shape_box_tag(get_para, gc, gc_style):
                            continue
                        if render_simple_layout_child_as_tabbed_line(
                            gc,
                            gc_style,
                            get_para,
                            width_twips if width_twips > 0 else DOC_CONTENT_WIDTH_TWIPS * 0.46,
                        ):
                            continue
                        if gc_name in {"div", "section", "article", "main", "header", "footer"} and not has_effective_box_visual(gc_style) and has_nested_block_children(gc):
                            render_nested_block_children_in_cell(
                                gc,
                                gc_style,
                                get_para,
                                width_twips if width_twips > 0 else DOC_CONTENT_WIDTH_TWIPS * 0.46,
                                target_cell=cell,
                            )
                            continue
                        p = get_para()
                        apply_paragraph_style(p, gc_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
                        for n in gc.children:
                            render_inline(p, n, gc_style, preserve_whitespace=False)
                    else:
                        p = para
                        render_inline(p, gc, child_style, preserve_whitespace=False)

            compact_cell_paragraphs(cell)

        return True

    def render_css_table_container_as_table(tag, parent_style, target_cell=None, target_width_twips=None):
        """轻量支持 CSS Table（display:table/table-cell）布局。"""
        style = resolve_style(tag, parent_style)
        if str(style.get("display", "")).strip().lower() != "table":
            return False

        direct_children = [c for c in tag.children if isinstance(c, Tag)]
        if not direct_children:
            return False

        row_children = []
        for child in direct_children:
            child_style = resolve_style(child, style)
            if str(child_style.get("display", "")).strip().lower() == "table-row":
                row_children.append(child)

        rows = []
        if row_children:
            for row_tag in row_children:
                row_style = resolve_style(row_tag, style)
                cells = []
                for c in row_tag.children:
                    if not isinstance(c, Tag):
                        continue
                    cell_style = resolve_style(c, row_style)
                    if str(cell_style.get("display", "")).strip().lower() == "table-cell":
                        cells.append(c)
                if cells:
                    rows.append(cells)
        else:
            cells = []
            for child in direct_children:
                child_style = resolve_style(child, style)
                if str(child_style.get("display", "")).strip().lower() == "table-cell":
                    cells.append(child)
            if len(cells) < 2:
                return False
            rows = [cells]

        if not rows:
            return False

        num_rows = len(rows)
        num_cols = max(len(r) for r in rows)
        if num_cols <= 0:
            return False

        if target_cell is None:
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        else:
            table = target_cell.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.autofit = False
        set_table_borders_none(table)
        set_table_layout_fixed(table)

        spacing_px = parse_border_spacing_px(style.get("border-spacing"))
        if spacing_px > 0:
            set_table_cell_spacing(table, spacing_px)

        col_weights = [1.0] * num_cols
        for row_cells in rows:
            for idx, cell_tag in enumerate(row_cells):
                if idx >= num_cols:
                    break
                cell_style = resolve_style(cell_tag, style)
                width_text = str(cell_style.get("width", "")).strip().lower()
                weight = 1.0
                if width_text.endswith("%"):
                    try:
                        weight = max(0.5, float(width_text[:-1]))
                    except ValueError:
                        weight = 1.0
                else:
                    width_pt = length_to_pt(width_text, current_font_pt=DEFAULT_FONT_PT) if width_text else None
                    if width_pt and width_pt > 0:
                        weight = max(1.0, width_pt / 18.0)
                col_weights[idx] = max(col_weights[idx], weight)

        available_twips = int(target_width_twips) if target_width_twips else int(DOC_CONTENT_WIDTH_TWIPS)
        width_text = str(style.get("width", "")).strip().lower()
        target_twips = available_twips
        if width_text:
            if width_text.endswith("%"):
                try:
                    target_twips = int(available_twips * float(width_text[:-1]) / 100.0)
                except ValueError:
                    target_twips = available_twips
            else:
                width_pt = length_to_pt(width_text, current_font_pt=DEFAULT_FONT_PT)
                if width_pt and width_pt > 0:
                    target_twips = int(width_pt * 20)
        elif not target_width_twips:
            target_twips = int(DOC_CONTENT_WIDTH_TWIPS * 0.92)
        target_twips = max(1800, min(available_twips, target_twips))
        set_table_width_twips(table, target_twips)
        total_weight = sum(col_weights) or float(num_cols)
        col_twips = [max(900, int(target_twips * w / total_weight)) for w in col_weights]
        if col_twips:
            col_twips[-1] += target_twips - sum(col_twips)
        set_table_grid_widths(table, col_twips)

        for r_idx, row_cells in enumerate(rows):
            row_style = style
            if row_children and r_idx < len(row_children):
                row_style = resolve_style(row_children[r_idx], style)
            for c_idx, cell_tag in enumerate(row_cells):
                if c_idx >= num_cols:
                    break
                cell_style = resolve_style(cell_tag, row_style)
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = ""
                set_cell_width_twips(cell, col_twips[c_idx] if c_idx < len(col_twips) else 0)
                if cell_style.get("background-color") and not is_transparent_color(cell_style.get("background-color")):
                    set_cell_shading(cell, cell_style.get("background-color"))
                for edge in ("left", "right", "top", "bottom"):
                    border = resolve_edge_border(cell_style, edge)
                    if border:
                        set_cell_border_edge(cell, edge, border)
                set_cell_padding(
                    cell,
                    cell_style,
                    current_font_pt=DEFAULT_FONT_PT,
                    horizontal_scale=0.7,
                    vertical_scale=0.25,
                )
                set_cell_vertical_align(cell, cell_style.get("vertical-align"))

                para = cell.paragraphs[0]
                para.clear()
                apply_paragraph_style(para, cell_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
                first_para_used = False

                def get_para():
                    nonlocal first_para_used
                    if not first_para_used:
                        first_para_used = True
                        return para
                    return cell.add_paragraph()

                current_inline_para = None
                for child in cell_tag.children:
                    handled, current_inline_para = render_inline_flow_child(
                        child,
                        cell_style,
                        current_inline_para,
                        get_para,
                        apply_box_visuals=False,
                    )
                    if handled:
                        continue
                    if not isinstance(child, Tag):
                        continue
                    current_inline_para = None
                    child_name = child.name.lower()
                    child_style = resolve_style(child, cell_style)
                    if child_name in {"ul", "ol"}:
                        def add_para_with_style(_style_name):
                            p = get_para()
                            p.clear()
                            return p
                        render_list(child, child_style, level=0, add_para_with_style=add_para_with_style)
                    elif is_block_like_tag(child):
                        p = get_para()
                        apply_paragraph_style(p, child_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
                        for n in child.children:
                            render_inline(p, n, child_style, preserve_whitespace=False)
                    else:
                        p = para
                        render_inline(p, child, cell_style, preserve_whitespace=False)
                compact_cell_paragraphs(cell)
        return True

    def render_flex_container_as_table(tag, parent_style):
        """轻量支持 flex(row)：降级为单行多列表格，修复图文/标题栏错位。"""
        style = resolve_style(tag, parent_style)
        display_value = str(style.get("display", "")).strip().lower()
        if display_value != "flex":
            return False
        flex_direction = str(style.get("flex-direction", "row")).strip().lower()
        if flex_direction not in {"row", "row-reverse"}:
            return False
        current_font_pt = length_to_pt(style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        margin_top_pt = length_to_pt(style.get("margin-top"), current_font_pt=current_font_pt) or 0.0
        margin_bottom_pt = length_to_pt(style.get("margin-bottom"), current_font_pt=current_font_pt) or 0.0
        if margin_top_pt > 0:
            add_spacing_paragraph(margin_top_pt * 0.5)

        def resolve_flex_target_width_twips():
            content_width_twips = int(get_doc_content_width_twips())
            target_twips = content_width_twips
            width_text = style.get("width")
            if width_text:
                wt = str(width_text).strip().lower()
                if wt.endswith("%"):
                    try:
                        target_twips = int(content_width_twips * float(wt[:-1]) / 100.0)
                    except ValueError:
                        pass
                else:
                    width_pt = length_to_pt(wt, current_font_pt=DEFAULT_FONT_PT)
                    if width_pt and width_pt > 0:
                        target_twips = int(width_pt * 20)
            return max(2400, min(content_width_twips, target_twips))

        def length_to_px(length_value, current_font_pt=DEFAULT_FONT_PT):
            pt = length_to_pt(length_value, current_font_pt=current_font_pt)
            if pt is None:
                return 0.0
            return max(0.0, pt / PX_TO_PT)

        def is_image_placeholder_child(child_style, child_tag):
            class_tokens = [str(c).strip().lower() for c in (child_tag.get("class") or [])]
            if "placeholder-img" in class_tokens:
                return True
            bg = child_style.get("background-color")
            has_bg = bg and not is_transparent_color(bg)
            h_px = length_to_px(child_style.get("height"))
            return bool(has_bg and h_px >= 50.0)

        def parse_flex_grow(child_style):
            grow_raw = str(child_style.get("flex-grow", "")).strip()
            grow = 0.0
            try:
                grow = float(grow_raw) if grow_raw else 0.0
            except ValueError:
                grow = 0.0
            if grow > 0:
                return grow
            flex_raw = str(child_style.get("flex", "")).strip()
            if flex_raw:
                first = flex_raw.split()[0]
                try:
                    return max(0.0, float(first))
                except ValueError:
                    return 0.0
            return 0.0

        def get_child_base_width_twips(child_style):
            current_font_pt = length_to_pt(child_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
            width_pt = length_to_pt(child_style.get("width"), current_font_pt=current_font_pt) or 0.0
            basis_pt = length_to_pt(child_style.get("flex-basis"), current_font_pt=current_font_pt) or 0.0
            min_width_pt = length_to_pt(child_style.get("min-width"), current_font_pt=current_font_pt) or 0.0
            base_pt = max(width_pt, basis_pt, min_width_pt)
            if base_pt <= 0:
                return 0
            return max(960, int(base_pt * 20))

        def get_flex_text_weight(child, child_style, content_block=False):
            text_len = len(child.get_text(" ", strip=True))
            text_weight = min(24.0, max(3.0, text_len * 0.18))
            if content_block:
                class_tokens = [str(c).strip().lower() for c in (child.get("class") or [])]
                if "text-content" in class_tokens:
                    return max(text_weight, 2.0)
                if "placeholder-img" in class_tokens:
                    return max(1.0, min(text_weight, 3.0))
            return text_weight

        def compute_flex_col_twips(local_children, local_child_styles, target_twips, content_block=False):
            min_col_twips = 960
            child_data = []
            any_grow = False
            for child, child_style in zip(local_children, local_child_styles):
                grow = parse_flex_grow(child_style)
                if grow > 0:
                    any_grow = True
                child_data.append(
                    {
                        "grow": grow,
                        "base_twips": get_child_base_width_twips(child_style),
                        "text_weight": get_flex_text_weight(child, child_style, content_block=content_block),
                    }
                )

            if not child_data:
                return []

            if any_grow:
                non_grow_text_total = sum(d["text_weight"] for d in child_data if d["grow"] <= 0 and d["base_twips"] <= 0)
                widths = []
                for data in child_data:
                    if data["base_twips"] > 0:
                        widths.append(max(min_col_twips, data["base_twips"]))
                    elif data["grow"] > 0:
                        widths.append(min_col_twips)
                    elif non_grow_text_total > 0:
                        widths.append(max(min_col_twips, int(target_twips * data["text_weight"] / non_grow_text_total)))
                    else:
                        widths.append(min_col_twips)

                total_base = sum(widths)
                if total_base >= target_twips:
                    total_base = sum(widths) or float(len(widths))
                    widths = [max(min_col_twips, int(target_twips * w / total_base)) for w in widths]
                    widths[-1] += target_twips - sum(widths)
                    return widths

                remain = target_twips - total_base
                total_grow = sum(d["grow"] for d in child_data if d["grow"] > 0) or 1.0
                for idx, data in enumerate(child_data):
                    if data["grow"] > 0:
                        widths[idx] += int(remain * data["grow"] / total_grow)
                widths[-1] += target_twips - sum(widths)
                return widths

            weights = [d["base_twips"] if d["base_twips"] > 0 else d["text_weight"] for d in child_data]
            total_weight = sum(weights) or float(len(weights))
            widths = [max(min_col_twips, int(target_twips * w / total_weight)) for w in weights]
            widths[-1] += target_twips - sum(widths)
            return widths

        def collect_textbox_paragraphs(text_col_tag, text_col_style):
            block_names = {
                "p", "blockquote", "h1", "h2", "h3", "h4", "h5", "h6", "li",
                "div", "section", "article", "pre",
            }
            items = []
            inline_nodes = []

            def append_run(runs, text, styles):
                if text is None:
                    return
                normalized = normalize_text_content(str(text), preserve_whitespace=False)
                if normalized:
                    runs.append({"text": normalized, "styles": dict(styles)})

            def collect_inline_runs(node, parent_style, runs):
                if isinstance(node, Comment):
                    return
                if isinstance(node, NavigableString):
                    append_run(runs, str(node), parent_style)
                    return
                if not isinstance(node, Tag):
                    return
                name = node.name.lower()
                if name == "br":
                    runs.append({"break": True})
                    return
                child_style = resolve_style(node, parent_style)
                if name in {"strong", "b"}:
                    child_style["font-weight"] = "700"
                if name in {"em", "i"}:
                    child_style["font-style"] = "italic"
                if name == "sup":
                    child_style["vertical-align"] = "super"
                if name == "sub":
                    child_style["vertical-align"] = "sub"
                if name == "u":
                    child_style["text-decoration"] = "underline"
                if name in {"s", "strike", "del"}:
                    child_style["text-decoration"] = "line-through"
                if name == "code":
                    child_style.setdefault("font-family", "Consolas")
                for child in node.children:
                    collect_inline_runs(child, child_style, runs)

            def flush_inline():
                if not inline_nodes:
                    return
                runs = []
                for inline_node in inline_nodes:
                    collect_inline_runs(inline_node, text_col_style, runs)
                inline_nodes.clear()
                if runs:
                    items.append({"runs": runs, "styles": dict(text_col_style), "blockquote": False})

            for node in text_col_tag.children:
                if isinstance(node, Comment):
                    continue
                if isinstance(node, NavigableString):
                    if str(node).strip():
                        inline_nodes.append(node)
                    continue
                if not isinstance(node, Tag):
                    continue
                name = node.name.lower()
                if name == "br":
                    inline_nodes.append(node)
                    continue
                if name in {"ul", "ol"}:
                    flush_inline()
                    for li in node.find_all("li", recursive=False):
                        li_style = resolve_style(li, text_col_style)
                        li_runs = [{"text": "• ", "styles": dict(li_style)}]
                        for li_child in li.children:
                            collect_inline_runs(li_child, li_style, li_runs)
                        if any((r.get("text") or "").strip() for r in li_runs if not r.get("break")):
                            items.append({"runs": li_runs, "styles": dict(li_style), "blockquote": False})
                    continue
                if name in block_names:
                    flush_inline()
                    para_style = resolve_style(node, text_col_style)
                    if name == "blockquote":
                        # 语义兜底：确保文本框内 blockquote 具备旧版可见特征。
                        para_style = dict(para_style)
                        para_style.setdefault("font-style", "italic")
                        para_style.setdefault("background-color", "#fffbf0")
                        para_style.setdefault("border-left", "4px solid #fdcb6e")
                        para_style.setdefault("padding-top", "15px")
                        para_style.setdefault("padding-right", "20px")
                        para_style.setdefault("padding-bottom", "15px")
                        para_style.setdefault("padding-left", "20px")
                        para_style.setdefault("margin-top", "20px")
                        para_style.setdefault("margin-bottom", "20px")
                    para_runs = []
                    for block_child in node.children:
                        collect_inline_runs(block_child, para_style, para_runs)
                    if para_runs:
                        items.append(
                            {
                                "runs": para_runs,
                                "styles": dict(para_style),
                                "blockquote": name == "blockquote",
                            }
                        )
                    continue
                # 非块级标签按行内拼接，避免把 strong/em 等拆成新段。
                inline_nodes.append(node)

            flush_inline()
            if not items:
                fallback = normalize_text_content(text_col_tag.get_text("", strip=True), preserve_whitespace=False).strip()
                if fallback:
                    items.append(
                        {
                            "runs": [{"text": fallback, "styles": dict(text_col_style)}],
                            "styles": dict(text_col_style),
                            "blockquote": False,
                        }
                    )
            return items

        def build_txbx_content_xml_from_paragraphs(paragraph_items, center=False):
            if not paragraph_items:
                paragraph_items = [{"runs": [{"text": "", "styles": {}}], "styles": {}, "blockquote": False}]

            def build_run_rpr_xml(run_style):
                rpr_parts = []
                color = parse_css_color(run_style.get("color"))
                if color:
                    rpr_parts.append(f'<w:color w:val="{str(color)}"/>')
                family = run_style.get("font-family")
                if family:
                    font_slots = resolve_font_slots(family)
                    if font_slots:
                        rpr_parts.append(
                            f'<w:rFonts w:ascii="{xml_escape(font_slots["ascii"])}" '
                            f'w:hAnsi="{xml_escape(font_slots["hAnsi"])}" '
                            f'w:eastAsia="{xml_escape(font_slots["eastAsia"])}" '
                            f'w:cs="{xml_escape(font_slots["cs"])}"/>'
                        )
                font_pt = length_to_pt(run_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT)
                if font_pt and font_pt > 0:
                    half_pt = max(2, int(round(font_pt * 2)))
                    rpr_parts.append(f'<w:sz w:val="{half_pt}"/>')
                    rpr_parts.append(f'<w:szCs w:val="{half_pt}"/>')
                fw = (run_style.get("font-weight") or "").lower()
                if fw in {"bold", "bolder"} or (fw.isdigit() and int(fw) >= 600):
                    rpr_parts.append("<w:b/>")
                font_style = (run_style.get("font-style") or "").lower()
                if font_style in {"italic", "oblique"}:
                    rpr_parts.append("<w:i/>")
                vertical_align = (run_style.get("vertical-align") or "").lower()
                if vertical_align in {"super", "sup", "superscript"}:
                    rpr_parts.append('<w:vertAlign w:val="superscript"/>')
                elif vertical_align in {"sub", "subscript"}:
                    rpr_parts.append('<w:vertAlign w:val="subscript"/>')
                deco = f'{run_style.get("text-decoration-line", "")} {run_style.get("text-decoration", "")}'.lower()
                if "underline" in deco:
                    rpr_parts.append('<w:u w:val="single"/>')
                if "line-through" in deco:
                    rpr_parts.append("<w:strike/>")
                bg = run_style.get("background-color")
                if bg and not is_transparent_color(bg):
                    bg_color = parse_css_color(bg)
                    if bg_color:
                        rpr_parts.append(f'<w:shd w:val="clear" w:color="auto" w:fill="{str(bg_color)}"/>')
                return f"<w:rPr>{''.join(rpr_parts)}</w:rPr>" if rpr_parts else ""

            def build_runs_xml(runs):
                run_parts = []
                for run_item in runs:
                    if run_item.get("break"):
                        run_parts.append("<w:r><w:br/></w:r>")
                        continue
                    text = str(run_item.get("text", ""))
                    if not text:
                        continue
                    escaped = xml_escape(text)
                    preserve = ' xml:space="preserve"' if (text[:1].isspace() or text[-1:].isspace() or "  " in text) else ""
                    rpr_xml = build_run_rpr_xml(run_item.get("styles") or {})
                    run_parts.append(f"<w:r>{rpr_xml}<w:t{preserve}>{escaped}</w:t></w:r>")
                return "".join(run_parts) if run_parts else "<w:r><w:t></w:t></w:r>"

            def build_p_borders_xml(para_style):
                edge_map = {"left": "left", "right": "right", "top": "top", "bottom": "bottom"}
                edge_xml = []
                for edge_key, xml_edge in edge_map.items():
                    border = resolve_edge_border(para_style, edge_key)
                    if not border:
                        continue
                    edge_xml.append(
                        f'<w:{xml_edge} w:val="{border["val"]}" w:sz="{border["sz"]}" '
                        f'w:space="{border["space"]}" w:color="{border["color"]}"/>'
                    )
                if not edge_xml:
                    return ""
                return f"<w:pBdr>{''.join(edge_xml)}</w:pBdr>"

            align_map = {"left": "left", "start": "left", "center": "center", "right": "right", "end": "right", "justify": "both"}
            p_list = []
            for item in paragraph_items:
                para_style = dict(item.get("styles") or {})
                is_quote = bool(item.get("blockquote"))
                ppr_parts = []
                if center:
                    ppr_parts.append('<w:jc w:val="center"/>')
                else:
                    align_val = str(para_style.get("text-align", "")).strip().lower()
                    if align_val in align_map:
                        ppr_parts.append(f'<w:jc w:val="{align_map[align_val]}"/>')
                if is_quote:
                    ppr_parts.append('<w:ind w:left="480"/>')
                # 把 margin/padding 映射到段落 spacing/ind，尽量贴近旧版块级视觉。
                mt_pt = length_to_pt(para_style.get("margin-top"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                mb_pt = length_to_pt(para_style.get("margin-bottom"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                pt_pt = length_to_pt(para_style.get("padding-top"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                pb_pt = length_to_pt(para_style.get("padding-bottom"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                pl_pt = length_to_pt(para_style.get("padding-left"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                pr_pt = length_to_pt(para_style.get("padding-right"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                spacing_before = int(max(0.0, (mt_pt * 0.5) + pt_pt) * 20)
                spacing_after = int(max(0.0, (mb_pt * 0.5) + pb_pt) * 20)
                if spacing_before > 0 or spacing_after > 0:
                    ppr_parts.append(f'<w:spacing w:before="{spacing_before}" w:after="{spacing_after}"/>')
                # 与 blockquote 缩进叠加时取最大值，避免覆盖原有引用缩进。
                ind_left_twips = int(max(0.0, pl_pt) * 20)
                if is_quote:
                    ind_left_twips = max(ind_left_twips, 480)
                ind_right_twips = int(max(0.0, pr_pt) * 20)
                if ind_left_twips > 0 or ind_right_twips > 0:
                    if ind_right_twips > 0:
                        ppr_parts.append(f'<w:ind w:left="{ind_left_twips}" w:right="{ind_right_twips}"/>')
                    else:
                        ppr_parts.append(f'<w:ind w:left="{ind_left_twips}"/>')
                bg = para_style.get("background-color")
                if bg and not is_transparent_color(bg):
                    bg_color = parse_css_color(bg)
                    if bg_color:
                        ppr_parts.append(f'<w:shd w:val="clear" w:color="auto" w:fill="{str(bg_color)}"/>')
                p_borders = build_p_borders_xml(para_style)
                if p_borders:
                    ppr_parts.append(p_borders)
                ppr_xml = f"<w:pPr>{''.join(ppr_parts)}</w:pPr>" if ppr_parts else ""
                runs_xml = build_runs_xml(item.get("runs") or [])
                p_list.append(
                    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    f"{ppr_xml}{runs_xml}</w:p>"
                )
            return "".join(p_list)

        def estimate_text_column_height_px(paragraph_items, width_emu, col_style):
            # 轻量链路没有浏览器 offsetHeight，需改为更贴近 Word 的逐段估算，避免过高留白。
            width_px = max(120.0, (max(width_emu, 1) / 914400.0) * 96.0)
            col_font_pt = length_to_pt(col_style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
            col_font_px = max(11.0, col_font_pt * (96.0 / 72.0))

            def parse_line_ratio(style_map, font_pt):
                lh = style_map.get("line-height")
                if not lh:
                    return 1.35
                text = str(lh).strip().lower()
                if text == "normal":
                    return 1.35
                css_ratio = None
                if text.replace(".", "", 1).isdigit():
                    css_ratio = float(text)
                elif text.endswith("%"):
                    try:
                        css_ratio = float(text[:-1]) / 100.0
                    except ValueError:
                        css_ratio = None
                else:
                    lh_pt = length_to_pt(text, current_font_pt=font_pt)
                    if lh_pt and font_pt > 0:
                        css_ratio = lh_pt / font_pt
                if not css_ratio or css_ratio <= 0:
                    return 1.35
                # 高度估算不再使用段落渲染时的行距补偿，避免文本框高度被低估。
                return max(1.15, min(css_ratio, 2.4))

            total_height_px = 0.0
            for item in paragraph_items or []:
                para_style = dict(item.get("styles") or {})
                para_font_pt = length_to_pt(para_style.get("font-size"), current_font_pt=col_font_pt) or col_font_pt
                para_font_px = max(11.0, para_font_pt * (96.0 / 72.0))

                # 以中文正文为主：单字符平均宽度取接近 1em，宽列可容纳字符数显著高于旧常量 5.2/inch。
                char_px = max(7.2, para_font_px * 1.02)
                chars_per_line = max(8, int(width_px / char_px))

                text_len = 0
                break_count = 0
                for run_item in item.get("runs") or []:
                    if run_item.get("break"):
                        break_count += 1
                        continue
                    text_len += len(str(run_item.get("text", "")))

                wrapped_lines = max(1, (text_len + chars_per_line - 1) // chars_per_line)
                line_count = wrapped_lines + break_count
                line_px = para_font_px * parse_line_ratio(para_style, para_font_pt)
                para_height_px = line_count * line_px

                mt_pt = length_to_pt(para_style.get("margin-top"), current_font_pt=para_font_pt) or 0.0
                mb_pt = length_to_pt(para_style.get("margin-bottom"), current_font_pt=para_font_pt) or 0.0
                pt_pt = length_to_pt(para_style.get("padding-top"), current_font_pt=para_font_pt) or 0.0
                pb_pt = length_to_pt(para_style.get("padding-bottom"), current_font_pt=para_font_pt) or 0.0
                # 与 txbxContent 里的 pPr 映射保持一致。
                space_px = ((mt_pt * 0.5 + pt_pt) + (mb_pt * 0.5 + pb_pt)) * (96.0 / 72.0)
                total_height_px += para_height_px + max(0.0, space_px)

            if total_height_px <= 0:
                total_height_px = col_font_px * 1.6

            min_h = length_to_px(col_style.get("min-height"), current_font_pt=col_font_pt)
            explicit_h = length_to_px(col_style.get("height"), current_font_pt=col_font_pt)
            min_height_px = max(108.0, min_h, explicit_h)
            # 给文本列增加温和安全系数，降低“略小被截断”的概率。
            return max(total_height_px * 1.14 + 10.0, min_height_px)

        def build_textbox_drawing_xml(
            x_emu,
            width_emu,
            height_emu,
            content_xml,
            fill_hex=None,
            body_anchor="t",
        ):
            docpr_id = alloc_shape_docpr_id()
            fill_xml = "<a:solidFill><a:srgbClr val=\"{}\"/></a:solidFill>".format(fill_hex) if fill_hex else "<a:noFill/>"
            return f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wp:anchor simplePos="0" relativeHeight="{251658240 + docpr_id}" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" distT="0" distB="0" distL="0" distR="0">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="paragraph"><wp:posOffset>{int((8 / 96.0) * 914400)}</wp:posOffset></wp:positionV>
            <wp:extent cx="{width_emu}" cy="{height_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapTopAndBottom/>
            <wp:docPr id="{docpr_id}" name="Flex Text Box {docpr_id}"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:wsp>
                        <wps:cNvSpPr txBox="1"/>
                        <wps:spPr>
                            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                            {fill_xml}
                            <a:ln><a:noFill/></a:ln>
                        </wps:spPr>
                        <wps:txbx><w:txbxContent>{content_xml}</w:txbxContent></wps:txbx>
                        <wps:bodyPr rot="0" vert="horz" wrap="square" anchor="{body_anchor}" anchorCtr="0" lIns="0" tIns="0" rIns="0" bIns="0"/>
                    </wps:wsp>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
    </w:drawing>'''

        def render_flex_container_as_textboxes():
            # 与旧链路对齐：仅图文混排 flex（至少 1 个占位图 + 1 个文本列）走文本框布局。
            local_children = [c for c in tag.children if isinstance(c, Tag)]
            if len(local_children) <= 1:
                return False
            if flex_direction == "row-reverse":
                local_children = list(reversed(local_children))

            child_styles = [resolve_style(c, style) for c in local_children]
            placeholder_flags = [is_image_placeholder_child(cs, child) for cs, child in zip(child_styles, local_children)]
            placeholder_count = sum(1 for f in placeholder_flags if f)
            if placeholder_count == 0 or placeholder_count == len(local_children):
                return False

            gap_px = 0.0
            for key in ("gap", "column-gap"):
                val = style.get(key)
                if val:
                    gap_px = max(gap_px, length_to_px(val))
            gap_emu = int((gap_px / 96.0) * 914400)

            weights = []
            for child, child_style, is_placeholder in zip(local_children, child_styles, placeholder_flags):
                grow_raw = str(child_style.get("flex-grow", "")).strip()
                grow = 0.0
                try:
                    grow = float(grow_raw) if grow_raw else 0.0
                except ValueError:
                    grow = 0.0
                if grow <= 0:
                    flex_raw = str(child_style.get("flex", "")).strip()
                    if flex_raw:
                        first = flex_raw.split()[0]
                        try:
                            grow = float(first)
                        except ValueError:
                            grow = 0.0
                if grow <= 0:
                    basis_px = length_to_px(child_style.get("flex-basis"))
                    if basis_px > 0:
                        grow = basis_px / 120.0
                if grow <= 0:
                    text_len = len(child.get_text(" ", strip=True))
                    grow = max(1.0, min(6.0, text_len * 0.02))
                if is_placeholder:
                    grow = max(1.0, min(grow, 2.0))
                else:
                    grow = max(1.6, grow)
                weights.append(grow)

            target_width_emu = int(resolve_flex_target_width_twips() * 635)
            total_gap_emu = max(0, len(local_children) - 1) * max(0, gap_emu)
            available_width_emu = max(1, target_width_emu - total_gap_emu)
            total_weight = sum(weights) or float(len(local_children))
            width_emus = [max(1, int(available_width_emu * w / total_weight)) for w in weights]
            if width_emus:
                width_emus[-1] += available_width_emu - sum(width_emus)

            xml_list = []
            x_emu = 0
            for child, child_style, is_placeholder, width_emu in zip(local_children, child_styles, placeholder_flags, width_emus):
                if is_placeholder:
                    ph_text = normalize_text_content(child.get_text("", strip=True), preserve_whitespace=False).strip()
                    height_px = max(140.0, length_to_px(child_style.get("height")), length_to_px(child_style.get("min-height")))
                    fill_rgb = parse_css_color(child_style.get("background-color"))
                    fill_hex = str(fill_rgb) if fill_rgb else "DFE6E9"
                    content_xml = build_txbx_content_xml_from_paragraphs(
                        [{"runs": [{"text": ph_text or "[图片]", "styles": dict(child_style)}], "styles": dict(child_style), "blockquote": False}],
                        center=True,
                    )
                    body_anchor = "ctr"
                else:
                    para_items = collect_textbox_paragraphs(child, child_style)
                    height_px = estimate_text_column_height_px(para_items, width_emu, child_style)
                    content_xml = build_txbx_content_xml_from_paragraphs(para_items, center=False)
                    fill_hex = None
                    body_anchor = "t"
                xml_list.append(
                    build_textbox_drawing_xml(
                        x_emu=x_emu,
                        width_emu=width_emu,
                        height_emu=max(1, int((height_px / 96.0) * 914400)),
                        content_xml=content_xml,
                        fill_hex=fill_hex,
                        body_anchor=body_anchor,
                    )
                )
                x_emu += width_emu + gap_emu

            if not xml_list:
                return False
            try:
                anchor_para = doc.add_paragraph()
                anchor_run = anchor_para.add_run()
                for drawing_xml in xml_list:
                    anchor_run._r.append(parse_xml(drawing_xml))
                anchor_para.paragraph_format.space_after = Pt(2)
                return True
            except Exception:
                return False

        if render_flex_container_as_textboxes():
            if margin_bottom_pt > 0:
                add_spacing_paragraph(margin_bottom_pt * 0.5)
            return True

        children = [c for c in tag.children if isinstance(c, Tag)]
        if len(children) <= 1:
            return False

        if flex_direction == "row-reverse":
            children = list(reversed(children))

        table = doc.add_table(rows=1, cols=len(children))
        table.style = "Table Grid"
        table.autofit = False
        set_table_borders_none(table)
        set_table_layout_fixed(table)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        gap_px = 0.0
        for key in ("gap", "column-gap"):
            val = style.get(key)
            if val:
                pt = length_to_pt(val, current_font_pt=DEFAULT_FONT_PT)
                if pt:
                    gap_px = max(gap_px, pt / PX_TO_PT)
        if gap_px > 0:
            set_table_cell_spacing(table, gap_px)

        is_content_block = "content-block" in [str(c).strip().lower() for c in (tag.get("class") or [])]
        target_twips = resolve_flex_target_width_twips()
        child_styles = [resolve_style(child, style) for child in children]
        col_twips = compute_flex_col_twips(children, child_styles, target_twips, content_block=is_content_block)
        set_table_width_twips(table, target_twips)
        set_table_grid_widths(table, col_twips)

        justify_content = str(style.get("justify-content", "")).strip().lower()

        for idx, child in enumerate(children):
            child_style = child_styles[idx] if idx < len(child_styles) else resolve_style(child, style)
            child_style = apply_container_box_visual_to_cell_style(
                child_style, style, 0, idx, 1, len(children)
            )
            cell = table.rows[0].cells[idx]
            cell.text = ""
            set_cell_width_twips(cell, col_twips[idx] if idx < len(col_twips) else 0)
            if child_style.get("background-color") and not is_transparent_color(child_style.get("background-color")):
                set_cell_shading(cell, child_style.get("background-color"))
            for edge in ("left", "right", "top", "bottom"):
                border = resolve_edge_border(child_style, edge)
                if not border:
                    if edge == "left" and idx == 0:
                        border = resolve_edge_border(style, edge)
                    elif edge == "right" and idx == len(children) - 1:
                        border = resolve_edge_border(style, edge)
                    elif edge in {"top", "bottom"}:
                        border = resolve_edge_border(style, edge)
                if border:
                    set_cell_border_edge(cell, edge, border)
            set_cell_padding(
                cell,
                child_style,
                current_font_pt=DEFAULT_FONT_PT,
                horizontal_scale=0.45 if is_content_block else 0.7,
                vertical_scale=0.25,
            )
            set_cell_vertical_align(cell, child_style.get("vertical-align"))

            para = cell.paragraphs[0]
            para.clear()
            apply_paragraph_style(para, child_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
            if justify_content == "space-between":
                if idx == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif idx == len(children) - 1:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            first_para_used = False

            def get_para():
                nonlocal first_para_used
                if not first_para_used:
                    first_para_used = True
                    return para
                return cell.add_paragraph()

            for gc in child.children:
                if isinstance(gc, NavigableString):
                    txt = normalize_text_content(str(gc), preserve_whitespace=False)
                    if txt.strip():
                        p = get_para()
                        add_text_run(p, txt, child_style)
                    continue
                if not isinstance(gc, Tag):
                    continue
                gc_name = gc.name.lower()
                gc_style = resolve_style(gc, child_style)
                if gc_name in {"ul", "ol"}:
                    def add_para_with_style(_style_name):
                        p = get_para()
                        p.clear()
                        return p
                    render_list(gc, gc_style, level=0, add_para_with_style=add_para_with_style)
                else:
                    if is_block_like_tag(gc):
                        if render_shape_box_tag(get_para, gc, gc_style):
                            continue
                        if render_simple_layout_child_as_tabbed_line(
                            gc,
                            gc_style,
                            get_para,
                            col_twips[idx] if idx < len(col_twips) else target_twips,
                        ):
                            continue
                        if gc_name in {"div", "section", "article", "main", "header", "footer"} and not has_effective_box_visual(gc_style) and has_nested_block_children(gc):
                            render_nested_block_children_in_cell(
                                gc,
                                gc_style,
                                get_para,
                                col_twips[idx] if idx < len(col_twips) else target_twips,
                                target_cell=cell,
                            )
                            continue
                        p = get_para()
                        apply_paragraph_style(p, gc_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=True)
                        for n in gc.children:
                            render_inline(p, n, gc_style, preserve_whitespace=False)
                    else:
                        p = para
                        render_inline(p, gc, child_style, preserve_whitespace=False)

            compact_cell_paragraphs(cell)

        if margin_bottom_pt > 0:
            add_spacing_paragraph(margin_bottom_pt * 0.5)
        return True

    def render_container_as_table(tag, parent_style):
        """
        中间层-5：把带背景/边框/内边距的块容器映射为 1x1 表格。
        这是当前最接近“网页盒模型”的 docx 近似方式之一。
        """
        style = resolve_style(tag, parent_style)
        current_font_pt = length_to_pt(style.get("font-size"), current_font_pt=DEFAULT_FONT_PT) or DEFAULT_FONT_PT
        margin_top_pt = length_to_pt(style.get("margin-top"), current_font_pt=current_font_pt) or 0.0
        margin_bottom_pt = length_to_pt(style.get("margin-bottom"), current_font_pt=current_font_pt) or 0.0
        if margin_top_pt > 0:
            add_spacing_paragraph(margin_top_pt * 0.5)
        content_width_twips = get_doc_content_width_twips()
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_layout_fixed(table)
        set_table_borders_none(table)
        cell = table.rows[0].cells[0]
        cell.text = ""
        # 容器 table 统一宽度，避免“自动宽度 + 居中”造成左边线不齐。
        width_text = style.get("width")
        target_twips = int(content_width_twips)
        if width_text:
            wt = str(width_text).strip().lower()
            if wt.endswith("%"):
                try:
                    target_twips = int(content_width_twips * float(wt[:-1]) / 100.0)
                except ValueError:
                    pass
            else:
                width_pt = length_to_pt(wt, current_font_pt=DEFAULT_FONT_PT)
                if width_pt and width_pt > 0:
                    target_twips = int(width_pt * 20)
        target_twips = max(4200, min(content_width_twips, target_twips))
        set_table_width_twips(table, target_twips)
        set_cell_width_twips(cell, target_twips)
        if style.get("background-color"):
            set_cell_shading(cell, style["background-color"])
        for edge in ("left", "right", "top", "bottom"):
            border = resolve_edge_border(style, edge)
            if border:
                set_cell_border_edge(cell, edge, border)
        set_cell_padding(cell, style, current_font_pt=DEFAULT_FONT_PT)

        # 容器盒模型已经映射到 cell；内部内容只保留文字样式，避免 padding/background 吃两次。
        content_base_style = dict(style)
        for mk in (
            "margin-top",
            "margin-right",
            "margin-bottom",
            "margin-left",
            "margin",
            "padding-top",
            "padding-right",
            "padding-bottom",
            "padding-left",
            "padding",
            "background",
            "background-color",
            "border",
            "border-top",
            "border-right",
            "border-bottom",
            "border-left",
            "border-radius",
        ):
            content_base_style.pop(mk, None)

        # 代码容器走专用渲染：保持容器盒模型，同时避免把 inline 片段拆成多个段落。
        if is_code_like_block(tag, style):
            code_para = cell.paragraphs[0]
            code_para.clear()
            apply_paragraph_style(code_para, content_base_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
            render_code_like_block(code_para, tag, content_base_style)
            return

        first_para_used = False
        pending_para_after_table = False
        pending_para_ref = None

        def get_para():
            nonlocal first_para_used, pending_para_after_table, pending_para_ref
            if pending_para_after_table and pending_para_ref is not None:
                # 单元格内插入 table 后，优先复用 table 后的默认段落，
                # 避免内容写回 table 前段落导致顺序颠倒，并消除额外空行。
                p = pending_para_ref
                p.clear()
                first_para_used = True
                pending_para_after_table = False
                pending_para_ref = None
                return p
            if not first_para_used:
                p = cell.paragraphs[0]
                p.clear()
                first_para_used = True
                return p
            return cell.add_paragraph()

        def render_css_table_child_as_tabbed_line(table_tag, table_style):
            """
            单元格内的简单双列 CSS Table（display:table + 两个 table-cell）优先转为
            单段落 + 右对齐制表位，避免嵌套表格带来的空行与裁切问题。
            """
            if str(table_style.get("display", "")).strip().lower() != "table":
                return False

            direct_children = [c for c in table_tag.children if isinstance(c, Tag)]
            if len(direct_children) < 2:
                return False

            cells = []
            for ch in direct_children:
                ch_style = resolve_style(ch, table_style)
                if str(ch_style.get("display", "")).strip().lower() == "table-cell":
                    cells.append((ch, ch_style))

            if len(cells) != 2:
                return False

            p = get_para()
            apply_paragraph_style(p, table_style, current_font_pt=DEFAULT_FONT_PT, apply_box_visuals=False)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            right_tab_pt = max(72.0, (target_twips - 120) / 20.0)
            p.paragraph_format.tab_stops.add_tab_stop(Pt(right_tab_pt), WD_TAB_ALIGNMENT.RIGHT)

            left_tag, left_style = cells[0]
            right_tag, right_style = cells[1]
            for n in left_tag.children:
                render_inline(p, n, left_style, preserve_whitespace=False)

            p.add_run("\t")

            for n in right_tag.children:
                render_inline(p, n, right_style, preserve_whitespace=False)

            return True

        current_inline_para = None
        for child in tag.children:
            handled, current_inline_para = render_inline_flow_child(
                child,
                content_base_style,
                current_inline_para,
                get_para,
                apply_box_visuals=False,
            )
            if handled:
                continue
            if not isinstance(child, Tag):
                continue
            current_inline_para = None
            child_name = child.name.lower()
            child_style = resolve_style(child, content_base_style)
            if child_name in {"ul", "ol"}:
                def add_para_with_style(_style_name):
                    p = get_para()
                    p.clear()
                    return p

                render_list(child, child_style, level=0, add_para_with_style=add_para_with_style)
            else:
                # inline 标签直接写入同一段，避免 code-block 里一行一个空段。
                if is_block_like_tag(child):
                    if render_shape_box_tag(get_para, child, child_style):
                        continue
                    child_display = str(child_style.get("display", "")).strip().lower()
                    if child_display == "grid":
                        if render_grid_container_as_table(
                            child,
                            content_base_style,
                            target_cell=cell,
                            target_width_twips=max(2400, target_twips - 120),
                        ):
                            pending_para_ref = None
                            if cell.paragraphs:
                                candidate = cell.paragraphs[-1]
                                prev = candidate._p.getprevious()
                                if prev is not None and prev.tag == qn("w:tbl"):
                                    pending_para_ref = candidate
                            if pending_para_ref is None:
                                pending_para_ref = cell.add_paragraph()
                            pending_para_after_table = True
                            continue
                    if render_inline_box_sequence_as_table(
                        child,
                        child_style,
                        cell,
                        max(2400, target_twips - 120),
                    ):
                        pending_para_ref = None
                        if cell.paragraphs:
                            candidate = cell.paragraphs[-1]
                            prev = candidate._p.getprevious()
                            if prev is not None and prev.tag == qn("w:tbl"):
                                pending_para_ref = candidate
                        if pending_para_ref is None:
                            pending_para_ref = cell.add_paragraph()
                        pending_para_after_table = True
                        continue
                    if render_simple_layout_child_as_tabbed_line(
                        child,
                        child_style,
                        get_para,
                        max(2400, target_twips - 120),
                    ):
                        continue
                    if child_name in {"div", "section", "article", "main", "header", "footer"} and not has_effective_box_visual(child_style) and has_nested_block_children(child):
                        render_nested_block_children_in_cell(
                            child,
                            child_style,
                            get_para,
                            max(2400, target_twips - 120),
                            target_cell=cell,
                        )
                        continue
                    if render_css_table_child_as_tabbed_line(child, child_style):
                        continue
                    if render_css_table_container_as_table(
                        child,
                        content_base_style,
                        target_cell=cell,
                        target_width_twips=max(2400, target_twips - 120),
                    ):
                        pending_para_ref = None
                        if cell.paragraphs:
                            candidate = cell.paragraphs[-1]
                            prev = candidate._p.getprevious()
                            if prev is not None and prev.tag == qn("w:tbl"):
                                pending_para_ref = candidate
                        if pending_para_ref is None:
                            pending_para_ref = cell.add_paragraph()
                        pending_para_after_table = True
                        continue
                    p = get_para()
                    if child_name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                        try:
                            heading_level = int(child_name[1])
                        except ValueError:
                            heading_level = None
                        apply_heading_paragraph_semantics(p, heading_level, doc)
                    # 内层若本身有可见盒模型（背景/边框/非0 padding），保留其视觉语义。
                    keep_child_box = has_effective_box_visual(child_style)
                    apply_paragraph_style(
                        p,
                        child_style,
                        current_font_pt=DEFAULT_FONT_PT,
                        apply_box_visuals=keep_child_box,
                    )
                    if is_code_like_block(child, child_style):
                        render_code_like_block(p, child, child_style)
                    else:
                        render_inline(p, child, content_base_style, preserve_whitespace=False)
                else:
                    p = get_para() if not first_para_used else cell.paragraphs[-1]
                    render_inline(p, child, content_base_style, preserve_whitespace=False)

        compact_cell_paragraphs(cell)

        if margin_bottom_pt > 0:
            add_spacing_paragraph(margin_bottom_pt * 0.5)

    def render_mixed_block_children(tag, style):
        """
        处理块容器里“前面是行内内容，后面跟块级子元素”的情况。
        典型例子：<p>文本<strong>强调</strong><ul>...</ul></p>
        这类结构在 HTML 里不规范，但实际输入里经常出现；这里尽量按浏览器观感兜底。
        """
        style = get_block_content_flow_style(tag, style)
        inline_para = None

        def get_inline_para():
            nonlocal inline_para
            if inline_para is None:
                inline_para = doc.add_paragraph()
                apply_paragraph_style(inline_para, style, current_font_pt=DEFAULT_FONT_PT)
            return inline_para

        for child in tag.children:
            if isinstance(child, Tag):
                if is_block_like_tag(child):
                    inline_para = None
                    render_block(child, style)
                else:
                    render_inline(get_inline_para(), child, style, preserve_whitespace=False)
            elif is_meaningful_text_node(child):
                render_inline(get_inline_para(), child, style, preserve_whitespace=False)

    def render_inline_flow_child(child, flow_style, current_para, acquire_para, apply_box_visuals=False):
        """
        在容器/table-cell 等上下文里保持 inline 流：
        文本节点、span/strong/br/img 等 inline 标签共享同一个段落；
        只有遇到真正块级元素时才由外层决定换段。
        """
        if isinstance(child, NavigableString):
            text = normalize_text_content(str(child), preserve_whitespace=False)
            text = normalize_inline_text_for_paragraph(current_para, text)
            if text.strip():
                if current_para is None:
                    current_para = acquire_para()
                    apply_paragraph_style(
                        current_para,
                        flow_style,
                        current_font_pt=DEFAULT_FONT_PT,
                        apply_box_visuals=apply_box_visuals,
                    )
                add_text_run(current_para, text, flow_style)
            return True, current_para

        if isinstance(child, Tag) and not is_block_like_tag(child):
            if current_para is None:
                current_para = acquire_para()
                apply_paragraph_style(
                    current_para,
                    flow_style,
                    current_font_pt=DEFAULT_FONT_PT,
                    apply_box_visuals=apply_box_visuals,
                )
            render_inline(current_para, child, flow_style, preserve_whitespace=False)
            return True, current_para

        return False, current_para

    def render_block(tag, parent_style):
        nonlocal last_render_was_container_table
        if not isinstance(tag, Tag):
            return
        name = tag.name.lower()
        style = resolve_style(tag, parent_style)

        if name == "hr":
            add_hr(doc)
            last_render_was_container_table = False
            return
        if name == "img":
            render_image(tag, parent_style)
            last_render_was_container_table = False
            return
        if name == "table":
            render_table(tag, parent_style)
            last_render_was_container_table = False
            return
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            try:
                level = int(name[1])
            except ValueError:
                level = None
            render_common_paragraph(tag, parent_style, bold_default=True, heading_level=level)
            last_render_was_container_table = False
            return
        if name == "header":
            display_value = str(style.get("display", "")).strip().lower()
            if display_value == "table":
                if render_css_table_container_as_table(tag, parent_style):
                    last_render_was_container_table = False
                    return
            if display_value == "flex":
                if render_flex_container_as_table(tag, parent_style):
                    last_render_was_container_table = False
                    return
            if display_value == "grid":
                if render_grid_container_as_table(tag, parent_style):
                    last_render_was_container_table = False
                    return
            # header 特殊处理：
            # - 内部结构按普通块渲染（保留 .doc-meta 的 flex 行为）
            # - border-bottom 挂到最后一个真实段落，避免生成“空行画线”
            # - 用 border space 近似 padding-bottom，用段后间距近似 margin-bottom
            para_start_idx = len(doc.paragraphs)
            has_block_child = any(
                getattr(c, "name", None)
                in {"p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "div", "section", "article", "main", "blockquote", "pre", "hr"}
                for c in tag.children
            )
            if has_block_child:
                render_mixed_block_children(tag, style)
            else:
                render_common_paragraph(tag, parent_style)

            padding_bottom_pt = length_to_pt(style.get("padding-bottom"), current_font_pt=DEFAULT_FONT_PT) or 0.0
            margin_bottom_pt = length_to_pt(style.get("margin-bottom"), current_font_pt=DEFAULT_FONT_PT) or 0.0
            margin_after_line_pt = max(0.0, margin_bottom_pt * 0.5)

            bottom_border = resolve_edge_border(style, "bottom")
            if bottom_border:
                if len(doc.paragraphs) > para_start_idx:
                    border_para = doc.paragraphs[-1]
                else:
                    border_para = doc.add_paragraph()
                border_data = dict(bottom_border)
                # 用边框留白近似 CSS padding-bottom，避免新增空段落。
                border_data["space"] = str(max(1, int(max(0.0, padding_bottom_pt * 0.6))))
                set_paragraph_border(border_para, "bottom", border_data)
                border_para.paragraph_format.space_after = Pt(margin_after_line_pt)
            elif margin_after_line_pt > 0 and len(doc.paragraphs) > para_start_idx:
                doc.paragraphs[-1].paragraph_format.space_after = Pt(margin_after_line_pt)
            last_render_was_container_table = False
            return
        if name == "blockquote":
            # 迁移主方案思路：带背景/边框的引用块优先按容器渲染。
            if has_effective_box_visual(style):
                render_container_table_with_optional_separator(tag, parent_style)
            else:
                render_common_paragraph(tag, parent_style, left_indent_cm=1.2)
                last_render_was_container_table = False
            return
        if name == "pre":
            style = dict(style)
            style.setdefault("font-family", "Consolas")
            para = doc.add_paragraph()
            apply_paragraph_style(para, style, current_font_pt=DEFAULT_FONT_PT)
            para.paragraph_format.left_indent = Cm(0.8)
            for child in tag.children:
                render_inline(para, child, style, preserve_whitespace=True)
            last_render_was_container_table = False
            return
        if is_code_like_block(tag, style):
            # 通用代码块分支：不依赖具体类名或容器结构。
            if name in {"div", "section", "article", "main"} and should_render_container_as_table(tag, style):
                render_container_table_with_optional_separator(tag, parent_style)
                return
            para = doc.add_paragraph()
            apply_paragraph_style(para, style, current_font_pt=DEFAULT_FONT_PT)
            render_code_like_block(para, tag, style)
            last_render_was_container_table = False
            return
        if name in {"ul", "ol"}:
            render_list(tag, parent_style, level=0)
            last_render_was_container_table = False
            return
        if name in {"p", "div", "section", "article", "main", "header", "footer"}:
            para = doc.add_paragraph()
            if render_shape_box_tag_to_paragraph(para, tag, style):
                last_render_was_container_table = False
                return
            para._element.getparent().remove(para._element)
        if name in {"p", "div", "section", "article", "main", "header", "footer"}:
            display_value = str(style.get("display", "")).strip().lower()
            if name in {"div", "section", "article", "main", "header", "footer"} and display_value == "table":
                if render_css_table_container_as_table(tag, parent_style):
                    last_render_was_container_table = False
                    return
            if name in {"div", "section", "article", "main", "header", "footer"} and display_value == "flex":
                if render_flex_container_as_table(tag, parent_style):
                    last_render_was_container_table = False
                    return
            if name in {"div", "section", "article", "main", "header", "footer"} and display_value == "grid":
                if render_grid_container_as_table(tag, parent_style):
                    last_render_was_container_table = False
                    return
            # 分隔线容器（如 contact-info）：按普通块渲染，底边框挂到最后一段。
            if name in {"div", "section", "article", "main", "footer"} and has_only_bottom_border_rule(style):
                para_start_idx = len(doc.paragraphs)
                has_block_child = any(
                    getattr(c, "name", None)
                    in {"p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "div", "section", "article", "main", "blockquote", "pre", "hr"}
                    for c in tag.children
                )
                if has_block_child:
                    render_mixed_block_children(tag, style)
                else:
                    render_common_paragraph(tag, parent_style)

                padding_bottom_pt = length_to_pt(style.get("padding-bottom"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                margin_bottom_pt = length_to_pt(style.get("margin-bottom"), current_font_pt=DEFAULT_FONT_PT) or 0.0
                margin_after_line_pt = max(0.0, margin_bottom_pt * 0.5)
                bottom_border = resolve_edge_border(style, "bottom")
                if bottom_border:
                    if len(doc.paragraphs) > para_start_idx:
                        border_para = doc.paragraphs[-1]
                    else:
                        border_para = doc.add_paragraph()
                    border_data = dict(bottom_border)
                    border_data["space"] = str(max(1, int(max(0.0, padding_bottom_pt * 0.6))))
                    set_paragraph_border(border_para, "bottom", border_data)
                    border_para.paragraph_format.space_after = Pt(margin_after_line_pt)
                elif margin_after_line_pt > 0 and len(doc.paragraphs) > para_start_idx:
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(margin_after_line_pt)
                last_render_was_container_table = False
                return
            # 对齐旧链路：footer 默认不按“容器表格”渲染（除非其本身是 flex/grid 已在前面处理）。
            if name in {"div", "section", "article", "main", "header"} and should_render_container_as_table(tag, style):
                render_container_table_with_optional_separator(tag, parent_style)
                return
            has_block_child = any(getattr(c, "name", None) in {"p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "table", "div", "section", "article", "main", "blockquote", "pre", "hr"} for c in tag.children)
            if has_block_child:
                render_mixed_block_children(tag, style)
            else:
                render_common_paragraph(tag, parent_style)
            if not has_block_child:
                last_render_was_container_table = False
            return

        # 未识别块：向下递归
        for child in tag.children:
            if isinstance(child, Tag):
                render_block(child, style)
            elif is_meaningful_text_node(child):
                para = doc.add_paragraph()
                add_text_run(para, str(child).strip(), style)
                last_render_was_container_table = False

    root = soup.body if soup.body else soup
    root_style = {"font-size": f"{DEFAULT_FONT_PT}pt"}
    if soup.body:
        root_style.update(rule_map.get(id(soup.body), {}))
        root_style.update(normalize_declarations(parse_inline_style(soup.body.get("style"))))
    for child in root.children:
        if isinstance(child, Tag):
            render_block(child, root_style)
        elif is_meaningful_text_node(child):
            doc.add_paragraph(str(child).strip())

    # 输出模式：保存到文件或输出 base64
    if output_docx:
        doc.save(output_docx)
    else:
        # 沙箱模式：输出 base64 字符串
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_base64 = base64.b64encode(buffer.read()).decode('utf-8')
        print(docx_base64)

def html2docx(input_html: str, output_docx: str) -> dict:
    """
    读取 HTML 文件并转换为 DOCX。

    Args:
        input_html: 输入 HTML 文件路径
        output_docx: 输出 DOCX 文件路径

    Returns:
        dict: {"ok": bool, "output_path": str, "message": str}
    """
    try:
        path = Path(input_html)
        if not path.exists():
            return {"ok": False, "output_path": output_docx, "message": f"HTML 文件不存在: {input_html}"}
        html_content = path.read_text(encoding="utf-8")
        html2docx_beautifulsoup(html_content, output_docx=output_docx)
        out = Path(output_docx)
        if not out.exists():
            return {"ok": False, "output_path": output_docx, "message": "转换完成但输出文件未生成"}
        size_kb = out.stat().st_size / 1024
        return {"ok": True, "output_path": output_docx, "message": f"转换成功，文件大小 {size_kb:.1f} KB"}
    except Exception as e:
        return {"ok": False, "output_path": output_docx, "message": f"转换失败: {e}"}

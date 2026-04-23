#!/usr/bin/env python3
"""
文件分类分级扫描脚本
基于 GB/T 35273 个人信息安全规范，通过文件名和文件内容关键词自动识别敏感类别并打标签。
分为两类：S（敏感个人信息）、G（一般个人信息）。
"""

import argparse
import csv
import json
import sys
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── 关键词规则 ──────────────────────────────────────────────

# S 类 - 敏感个人信息
SENSITIVE_KEYWORDS = [
    # 个人身份信息
    "身份证", "身份证号", "身份证复印件", "idcard", "id_card", "identity_card",
    "护照", "护照号", "passport",
    "社保号", "社保卡", "social_security",
    "军官证", "军官证号",
    "户口本", "户口簿", "户籍",
    "出生证", "出生医学证明",
    # 金融账户信息
    "银行卡", "银行卡号", "bank_card", "bank_card_number",
    "信用卡", "credit_card",
    "银行流水", "bank_statement",
    "账户密码", "login_password",
    "支付密码", "pay_password",
    "资金流水", "交易记录", "transaction_record",
    "贷款", "loan",
    "征信", "credit_report",
    # 认证凭据
    "密码", "password", "passwd",
    "密钥", "secret_key", "api_key",
    "令牌", "token",
    "签名证书", "signature_certificate",
    # 生物识别信息
    "人脸", "facial",
    "指纹", "fingerprint",
    "虹膜", "iris",
    "声纹", "voiceprint",
    "步态", "gait",
    # 通讯与社交信息
    "通讯录", "contacts", "address_book",
    "通话记录", "call_log", "call_history",
    "短信记录", "sms_record",
    "即时通讯", "im_chat", "wechat_record", "qq_chat",
    "社交关系", "social_graph", "social_relation",
    # 位置与行踪信息
    "定位", "gps",
    "行踪轨迹", "trajectory", "track_record",
    "轨迹数据", "location_history",
    # 健康与财产信息
    "病历", "medical_record",
    "体检", "health_checkup",
    "诊疗", "diagnosis",
    "药物", "medication",
    "财产信息", "property_info", "asset_info",
    "收入", "salary",
    "纳税", "tax_record",
    "保险", "insurance_policy",
    # 其他敏感信息
    "宗教信仰", "religion",
    "性取向", "sexual_orientation",
    "政治面貌", "political_status",
]

# G 类 - 一般个人信息
GENERAL_KEYWORDS = [
    "姓名",
    "手机号", "phone", "mobile", "cell_phone",
    "电话", "telephone",
    "邮箱", "email", "e_mail",
    "性别", "gender",
    "年龄", "age",
    "生日", "birthday", "birth_date",
    "籍贯", "native_place",
    "民族", "ethnicity",
    "婚姻", "marital_status",
    "学历", "education",
    "职位", "position",
    "工号", "employee_id",
    "用户画像", "user_profile",
    "客户信息", "customer_info",
    "用户信息", "user_info",
    "ip地址", "ip_address",
    "mac地址", "mac_address",
    "设备指纹", "device_fingerprint",
    "浏览记录", "browsing_history",
    "搜索记录", "search_history",
    "cookie",
    "账号", "account", "username",
]

CATEGORY_INFO = {
    "S": {"name": "敏感个人信息", "suggestion": "加密存储，严格限制访问权限，传输需加密，建议脱敏后使用"},
    "G": {"name": "一般个人信息", "suggestion": "内部使用，注意保护，避免不必要的公开分享"},
}

# ── 核心逻辑 ──────────────────────────────────────────────


def normalize_text(text: str) -> str:
    """统一为小写，替换分隔符为空格。"""
    return text.lower().replace("_", " ").replace("-", " ").replace(".", " ")


def match_keywords(text: str, keywords: list) -> list:
    """在文本中匹配关键词列表，返回命中的关键词（去重保序）。"""
    normalized = normalize_text(text)
    matched = []
    for kw in keywords:
        kw_lower = kw.lower()
        # 使用子串匹配，适应文件名和正文两种场景
        if kw_lower in normalized:
            matched.append(kw)
    # 去重保序
    return list(dict.fromkeys(matched))


def classify_text(text: str) -> tuple:
    """根据文本关键词判断敏感类别。
    返回 (category, matched_keywords)。
    """
    # 先检查 S 类
    s_matched = match_keywords(text, SENSITIVE_KEYWORDS)
    if s_matched:
        return ("S", s_matched)

    # 再检查 G 类
    g_matched = match_keywords(text, GENERAL_KEYWORDS)
    if g_matched:
        return ("G", g_matched)

    return ("G", [])


def read_file_content(file_path: Path) -> str:
    """读取文件内容，根据文件类型选择不同的读取方式。
    仅提取纯文本内容用于关键词匹配。
    """
    suffix = file_path.suffix.lower()

    if suffix in (".docx",):
        if not HAS_DOCX:
            print(f"  ⚠ python-docx 未安装，跳过内容读取: {file_path.name}", file=sys.stderr)
            return ""
        try:
            doc = Document(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # 同时提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"  ⚠ 读取失败: {file_path.name} ({e})", file=sys.stderr)
            return ""

    elif suffix in (".doc",):
        print(f"  ⚠ .doc 格式暂不支持内容读取，仅基于文件名判断: {file_path.name}", file=sys.stderr)
        return ""

    elif suffix in (".txt", ".md", ".csv", ".log", ".json", ".xml", ".html", ".htm"):
        try:
            # 尝试 UTF-8，失败则尝试 GBK
            try:
                return file_path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                return file_path.read_text(encoding="gbk")
        except Exception as e:
            print(f"  ⚠ 读取失败: {file_path.name} ({e})", file=sys.stderr)
            return ""

    elif suffix == ".pdf":
        print(f"  ⚠ PDF 格式暂不支持内容读取，仅基于文件名判断: {file_path.name}", file=sys.stderr)
        return ""

    else:
        # 其他类型尝试作为文本读取
        try:
            try:
                return file_path.read_text(encoding="utf-8")[:5000]
            except UnicodeDecodeError:
                return file_path.read_text(encoding="gbk")[:5000]
        except Exception:
            return ""


def classify_file(file_path: Path, scan_content: bool = True) -> dict:
    """对单个文件进行分类分级。同时基于文件名和文件内容判断。"""
    filename = file_path.name
    name_without_ext = file_path.stem

    # 第一层：基于文件名判断
    name_category, name_keywords = classify_text(name_without_ext)

    # 第二层：基于文件内容判断
    content_category, content_keywords = ("G", [])
    if scan_content:
        content = read_file_content(file_path)
        if content:
            content_category, content_keywords = classify_text(content)

    # 合并结果：S 类优先；关键词去重保序
    all_keywords = list(dict.fromkeys(name_keywords + content_keywords))
    if name_category == "S" or content_category == "S":
        final_category = "S"
    else:
        final_category = "G" if (name_category == "G" or content_category == "G") else "G"

    info = CATEGORY_INFO[final_category]

    # 判断来源
    if name_keywords and content_keywords:
        source = "文件名+内容"
    elif name_keywords:
        source = "文件名"
    elif content_keywords:
        source = "内容"
    else:
        source = "无匹配"

    return {
        "category": final_category,
        "category_name": info["name"],
        "matched_keywords": all_keywords,
        "name_keywords": name_keywords,
        "content_keywords": content_keywords,
        "match_source": source,
        "suggestion": info["suggestion"],
    }


def scan_directory(directory: str, pattern: str = "*.docx", scan_content: bool = True) -> list:
    """递归扫描目录，返回分类分级结果列表。"""
    results = []
    target_path = Path(directory)

    if not target_path.exists():
        print(f"错误：目录不存在 - {directory}", file=sys.stderr)
        sys.exit(1)

    if not target_path.is_dir():
        print(f"错误：路径不是目录 - {directory}", file=sys.stderr)
        sys.exit(1)

    matched_files = list(target_path.rglob(pattern))

    if not matched_files:
        print(f"未找到匹配 '{pattern}' 的文件", file=sys.stderr)
        return results

    # 过滤掉目录
    matched_files = [f for f in matched_files if f.is_file()]

    mode_label = "文件名+内容" if scan_content else "仅文件名"
    print(f"扫描模式: {mode_label} | 文件数: {len(matched_files)}", file=sys.stderr)

    for i, file_path in enumerate(matched_files, 1):
        print(f"  [{i}/{len(matched_files)}] 扫描: {file_path.name}", file=sys.stderr)
        classification = classify_file(file_path, scan_content=scan_content)
        results.append({
            "file_path": str(file_path),
            "file_name": file_path.name,
            **classification,
        })

    return results


def print_results(results: list):
    """在终端以格式化表格形式展示结果。"""
    if not results:
        print("扫描完成，无匹配文件。")
        return

    # 统计
    cat_counts = {}
    source_counts = {}
    for r in results:
        cat = r["category"]
        cat_counts[cat] = cat_counts.get(cat, 0) + 1
        src = r["match_source"]
        source_counts[src] = source_counts.get(src, 0) + 1

    print(f"\n{'='*70}")
    print(f"  文件分类分级扫描报告（GB/T 35273）")
    print(f"  扫描文件总数: {len(results)}")
    print(f"  S(敏感个人信息): {cat_counts.get('S', 0)} | G(一般个人信息): {cat_counts.get('G', 0)}")
    source_summary = " | ".join(f"{k}: {v}" for k, v in source_counts.items())
    print(f"  匹配来源: {source_summary}")
    print(f"{'='*70}\n")

    # S 类排在前面
    results.sort(key=lambda x: 0 if x["category"] == "S" else 1)

    for i, r in enumerate(results, 1):
        keywords_str = ", ".join(r["matched_keywords"]) if r["matched_keywords"] else "无"
        name_kw = ", ".join(r["name_keywords"]) if r["name_keywords"] else "无"
        content_kw = ", ".join(r["content_keywords"]) if r["content_keywords"] else "无"
        print(f"  [{i}] {r['file_name']}")
        print(f"      类别: {r['category']} ({r['category_name']})")
        print(f"      路径: {r['file_path']}")
        print(f"      文件名关键词: {name_kw}")
        print(f"      内容关键词: {content_kw}")
        print(f"      建议: {r['suggestion']}")
        print()

    print(f"{'='*70}")
    print(f"  扫描完成。")
    print(f"{'='*70}")


def save_csv(results: list, output_path: str):
    """将结果保存为 CSV 文件。"""
    if not results:
        print("无结果可保存。")
        return

    fieldnames = ["file_path", "file_name", "category", "category_name",
                  "matched_keywords", "name_keywords", "content_keywords",
                  "match_source", "suggestion"]
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for r in results:
            row = {k: r[k] for k in fieldnames}
            row["matched_keywords"] = "; ".join(r["matched_keywords"])
            row["name_keywords"] = "; ".join(r["name_keywords"])
            row["content_keywords"] = "; ".join(r["content_keywords"])
            writer.writerow(row)

    print(f"CSV 报告已保存: {output_path}")


def save_json(results: list, output_path: str):
    """将结果保存为 JSON 文件。"""
    if not results:
        print("无结果可保存。")
        return

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"JSON 报告已保存: {output_path}")


# ── 入口 ──────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(
        description="文件分类分级扫描工具（GB/T 35273）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python3 scan_files.py /path/to/documents
  python3 scan_files.py /path/to/documents --pattern "*.docx"
  python3 scan_files.py /path/to/documents --output report.csv
  python3 scan_files.py /path/to/documents --output report.json --format json
  python3 scan_files.py /path/to/documents --name-only
        """,
    )
    parser.add_argument("directory", help="要扫描的目录路径")
    parser.add_argument("--pattern", default="*.docx", help="文件匹配模式（默认: *.docx）")
    parser.add_argument("--output", default=None, help="输出文件路径（.csv 或 .json）")
    parser.add_argument("--format", default="csv", choices=["csv", "json"],
                        help="输出格式（默认: csv）")
    parser.add_argument("--name-only", action="store_true",
                        help="仅基于文件名判断，不读取文件内容")

    args = parser.parse_args()

    scan_content = not args.name_only
    results = scan_directory(args.directory, args.pattern, scan_content=scan_content)
    print_results(results)

    if args.output and results:
        if args.format == "json":
            save_json(results, args.output)
        else:
            save_csv(results, args.output)


if __name__ == "__main__":
    main()

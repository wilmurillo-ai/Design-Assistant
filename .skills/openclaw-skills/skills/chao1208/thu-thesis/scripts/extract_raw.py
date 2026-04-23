"""
extract_raw.py - Word 文档机械提取（不做任何理解/判断）

职责：把 .docx 里的内容原样导出为 raw.json，供 llm_parse.py 阅读理解。

输出格式：
{
  "source": "filename.docx",
  "paragraphs": [
    {
      "idx": 0,
      "style": "Heading 1",
      "text": "第4章 W信托...",
      "bold": true,
      "font_size": 16
    },
    ...
  ],
  "tables": [
    {
      "idx": 0,             # 在文档流中的位置（相对于段落）
      "before_para": 5,     # 紧接在第几个段落之后
      "rows": [["表头1","表头2"], ["数据1","数据2"], ...]
    }
  ],
  "figures": [
    {
      "rId": "rId5",
      "filename": "image1.png",   # 从 word/media 提取的文件名
      "para_idx": 12              # 出现在第几个段落
    }
  ]
}
"""

import json
import sys
import zipfile
from pathlib import Path

import docx
from docx.oxml.ns import qn


def _render_chart(chart_xml_bytes: bytes, out_path: Path):
    """解析 chart XML 数据，用 matplotlib 渲染为 PNG"""
    import xml.etree.ElementTree as ET
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    NS = {
        'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }
    root = ET.fromstring(chart_xml_bytes)

    def get_vals(ser, tag):
        node = ser.find(f'c:{tag}/c:numRef/c:numCache', NS) or \
               ser.find(f'c:{tag}/c:numLit', NS)
        if node is None:
            return []
        return [float(pt.find('c:v', NS).text)
                for pt in node.findall('c:pt', NS)
                if pt.find('c:v', NS) is not None]

    def get_cats(ser):
        node = ser.find('c:cat/c:strRef/c:strCache', NS) or \
               ser.find('c:cat/c:strLit', NS) or \
               ser.find('c:cat/c:numRef/c:numCache', NS)
        if node is None:
            return []
        return [pt.find('c:v', NS).text for pt in node.findall('c:pt', NS)
                if pt.find('c:v', NS) is not None]

    def ser_name(ser):
        t = ser.find('c:tx/c:strRef/c:strCache/c:pt/c:v', NS)
        return t.text if t is not None else ''

    fig, ax = plt.subplots(figsize=(6, 4))

    # 尝试柱状图、折线图、饼图
    for chart_tag in ('barChart', 'lineChart', 'pieChart', 'areaChart', 'scatterChart'):
        chart_node = root.find(f'.//c:{chart_tag}', NS)
        if chart_node is None:
            continue
        sers = chart_node.findall('c:ser', NS)
        cats = get_cats(sers[0]) if sers else []
        x = np.arange(len(cats)) if cats else None

        if chart_tag == 'pieChart':
            vals = get_vals(sers[0], 'val') if sers else []
            if vals:
                ax.pie(vals, labels=cats or None, autopct='%1.1f%%')
        elif chart_tag == 'scatterChart':
            for ser in sers:
                xv = get_vals(ser, 'xVal')
                yv = get_vals(ser, 'yVal')
                if xv and yv:
                    ax.scatter(xv[:len(yv)], yv)
        else:
            width = 0.8 / max(len(sers), 1)
            for i, ser in enumerate(sers):
                vals = get_vals(ser, 'val')
                if not vals:
                    continue
                offset = (i - len(sers) / 2 + 0.5) * width
                label = ser_name(ser)
                if chart_tag == 'barChart':
                    direction = chart_node.find('c:barDir', NS)
                    horiz = direction is not None and direction.get('val') == 'bar'
                    if horiz:
                        ax.barh(np.arange(len(vals)) + offset, vals, width, label=label)
                    else:
                        ax.bar((x[:len(vals)] if x is not None else np.arange(len(vals))) + offset,
                               vals, width, label=label)
                else:
                    ax.plot(x[:len(vals)] if x is not None else range(len(vals)),
                            vals, label=label)

            if cats and x is not None:
                ax.set_xticks(x)
                ax.set_xticklabels(cats, rotation=30, ha='right', fontsize=8)
            if any(ser_name(s) for s in sers):
                ax.legend(fontsize=8)
        break

    plt.tight_layout()
    plt.savefig(str(out_path), dpi=150, bbox_inches='tight')
    plt.close(fig)


def extract_raw(docx_path: str, output_dir: str = "output") -> dict:
    docx_path = Path(docx_path).resolve()
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(str(docx_path))

    # ── 1. 提取段落（扁平列表，不做结构判断）──
    paragraphs_out = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # 字体大小：取第一个有字号的 run
        font_size = None
        bold = False
        for run in p.runs:
            if run.font.size is not None and font_size is None:
                font_size = int(run.font.size.pt) if run.font.size else None
            if run.bold:
                bold = True
        paragraphs_out.append({
            "idx": idx,
            "style": p.style.name,
            "text": text,
            "bold": bold,
            "font_size": font_size,
        })

    # ── 2. 提取表格（记录在文档流中的位置）──
    tables_out = []
    # 遍历 body 元素，找每个表格紧跟在哪个段落后面
    body_elems = list(doc.element.body)
    para_elements = [p._element for p in doc.paragraphs]
    tbl_elements = [t._element for t in doc.tables]

    last_para_idx = -1
    tbl_counter = 0
    for elem in body_elems:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            # 找段落索引
            try:
                last_para_idx = para_elements.index(elem)
            except ValueError:
                pass
        elif tag == 'tbl':
            tbl = doc.tables[tbl_counter] if tbl_counter < len(doc.tables) else None
            rows_data = []
            if tbl:
                for row in tbl.rows:
                    rows_data.append([cell.text.strip() for cell in row.cells])
            tables_out.append({
                "idx": tbl_counter,
                "before_para": last_para_idx,
                "rows": rows_data,
            })
            tbl_counter += 1

    # ── 3. 提取图片（普通图片 + chart）──
    import re, shutil
    figures_out = []
    figures_dir = output_dir / "figures"
    if figures_dir.exists():
        shutil.rmtree(figures_dir)
    figures_dir.mkdir(parents=True)

    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    C_CHART_NS = 'http://schemas.openxmlformats.org/drawingml/2006/chart'

    with zipfile.ZipFile(str(docx_path), 'r') as z:
        # 复制所有 media 文件
        for mf in [f for f in z.namelist() if f.startswith('word/media/')]:
            dest = figures_dir / Path(mf).name
            dest.write_bytes(z.read(mf))

        # 读关系表：rId → Target
        try:
            rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
        except Exception:
            rels_xml = ''

        # rId → media filename（普通图片）
        rid_to_media = {}
        for m in re.finditer(r'Id="(rId\d+)"[^>]+Target="media/([^"]+)"', rels_xml):
            rid_to_media[m.group(1)] = m.group(2)

        # rId → chart xml path
        rid_to_chart = {}
        for m in re.finditer(r'Id="(rId\d+)"[^>]+Target="(charts/[^"]+)"', rels_xml):
            rid_to_chart[m.group(1)] = 'word/' + m.group(2)

        # 渲染 chart → PNG（用 matplotlib 解析 chart XML 数据）
        chart_png_map = {}  # chart_path → png filename
        chart_counter = 1
        for rid, chart_path in rid_to_chart.items():
            png_name = f'chart{chart_counter}.png'
            try:
                chart_xml = z.read(chart_path)
                _render_chart(chart_xml, figures_dir / png_name)
                chart_png_map[chart_path] = png_name
                print(f'   ✅ chart {chart_path} → {png_name}')
            except Exception as e:
                print(f'   ⚠️  chart {chart_path} 渲染失败: {e}')
            chart_counter += 1

        # rId → png filename（chart）
        rid_to_chart_png = {}
        for rid, chart_path in rid_to_chart.items():
            if chart_path in chart_png_map:
                rid_to_chart_png[rid] = chart_png_map[chart_path]

    # 遍历段落，找图片位置（a:blip = 普通图, c:chart ref = chart）
    seen_rids = set()
    for para_idx, p in enumerate(doc.paragraphs):
        for run in p.runs:
            elem = run._element
            # 普通图片
            for pic in elem.iter(qn('a:blip')):
                rid = pic.get(f'{{{R_NS}}}embed')
                if rid and rid in rid_to_media and rid not in seen_rids:
                    seen_rids.add(rid)
                    figures_out.append({
                        "rId": rid,
                        "filename": rid_to_media[rid],
                        "para_idx": para_idx,
                    })
            # chart
            for chart_ref in elem.iter(f'{{{C_CHART_NS}}}chart'):
                rid = chart_ref.get(f'{{{R_NS}}}id')
                if rid and rid in rid_to_chart_png and rid not in seen_rids:
                    seen_rids.add(rid)
                    figures_out.append({
                        "rId": rid,
                        "filename": rid_to_chart_png[rid],
                        "para_idx": para_idx,
                    })

    result = {
        "source": docx_path.name,
        "paragraphs": paragraphs_out,
        "tables": tables_out,
        "figures": figures_out,
    }

    out_path = output_dir / f"raw_{docx_path.stem}.json"
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"✅ raw 提取完成 → {out_path}")
    print(f"   {len(paragraphs_out)} 段落  {len(tables_out)} 表格  {len(figures_out)} 图片")
    return result


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python extract_raw.py <docx文件> [输出目录]")
        sys.exit(1)
    docx_file = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) > 2 else "output"
    extract_raw(docx_file, out_dir)

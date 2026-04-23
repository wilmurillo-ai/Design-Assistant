#!/usr/bin/env python3
"""
文档预处理器（Document Preprocessor）

将多源异构数据统一为标准化的知识库输入格式

使用方法：
    python document_preprocessor.py --input documents/ --output processed/

核心功能：
    1. 多格式支持：PDF/Word/Excel/HTML/Markdown/图片
    2. 内容提取：文本、表格、图片
    3. 数据清洗：去除噪声、标准化格式
    4. 结构化输出：统一的文档格式
"""

import os
import re
import json
import logging
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Tuple
from enum import Enum
from datetime import datetime
import hashlib

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """文档格式"""
    PDF = "pdf"
    WORD = "docx"
    EXCEL = "xlsx"
    HTML = "html"
    MARKDOWN = "md"
    TEXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class DocumentSection:
    """文档章节"""
    level: int                      # 层级（1=章，2=节，3=小节）
    title: str                      # 章节标题
    content: str                    # 章节内容
    path: str                       # 章节路径（如：第一章 > 1.1节）
    start_pos: int = 0              # 起始位置
    end_pos: int = 0                # 结束位置


@dataclass
class StructuredTable:
    """结构化表格"""
    id: str                         # 表格ID
    headers: List[str]              # 表头
    rows: List[List[str]]           # 数据行
    markdown: str                   # Markdown格式
    json_data: List[Dict]           # JSON格式
    description: str                # 自然语言描述
    location: str = ""              # 在文档中的位置


@dataclass
class DocumentImage:
    """文档图片"""
    id: str                         # 图片ID
    path: str                       # 图片路径
    caption: str                    # 图片说明
    ocr_text: str = ""              # OCR识别文本
    location: str = ""              # 在文档中的位置


@dataclass
class DocumentMetadata:
    """文档元数据"""
    source: str = ""                # 来源
    author: str = ""                # 作者
    publish_date: str = ""          # 发布日期
    doc_type: str = ""              # 文档类型
    language: str = "zh"            # 语言
    page_count: int = 0             # 页数
    word_count: int = 0             # 字数
    encoding: str = "utf-8"         # 编码
    file_size: int = 0              # 文件大小
    created_at: str = ""            # 创建时间
    modified_at: str = ""           # 修改时间


@dataclass
class ProcessedDocument:
    """处理后的文档"""
    id: str                                 # 文档ID
    title: str                              # 标题
    content: str                            # 正文内容
    sections: List[DocumentSection]         # 章节列表
    tables: List[StructuredTable]           # 表格列表
    images: List[DocumentImage]             # 图片列表
    metadata: DocumentMetadata              # 元数据
    raw_text: str = ""                      # 原始文本（清洗前）
    status: str = "success"                 # 处理状态
    error: str = ""                         # 错误信息


class FormatDetector:
    """格式检测器"""
    
    # 文件扩展名映射
    EXTENSION_MAP = {
        '.pdf': DocumentFormat.PDF,
        '.doc': DocumentFormat.WORD,
        '.docx': DocumentFormat.WORD,
        '.xls': DocumentFormat.EXCEL,
        '.xlsx': DocumentFormat.EXCEL,
        '.html': DocumentFormat.HTML,
        '.htm': DocumentFormat.HTML,
        '.md': DocumentFormat.MARKDOWN,
        '.markdown': DocumentFormat.MARKDOWN,
        '.txt': DocumentFormat.TEXT,
        '.png': DocumentFormat.IMAGE,
        '.jpg': DocumentFormat.IMAGE,
        '.jpeg': DocumentFormat.IMAGE,
        '.gif': DocumentFormat.IMAGE,
        '.tiff': DocumentFormat.IMAGE,
        '.bmp': DocumentFormat.IMAGE,
    }
    
    @classmethod
    def detect(cls, file_path: str) -> DocumentFormat:
        """
        检测文档格式
        
        Args:
            file_path: 文件路径
        
        Returns:
            DocumentFormat: 文档格式
        """
        ext = os.path.splitext(file_path)[1].lower()
        return cls.EXTENSION_MAP.get(ext, DocumentFormat.UNKNOWN)


class DataCleaner:
    """数据清洗器"""
    
    def __init__(self, config: Dict = None):
        self.config = config or {}
        
        # 页眉页脚模式（可配置）
        self.header_patterns = self.config.get('header_patterns', [
            r'^第\s*\d+\s*页\s*$',           # 第 1 页
            r'^Page\s*\d+\s*of\s*\d+$',      # Page 1 of 10
            r'^\d+\s*/\s*\d+$',               # 1/10
        ])
        
        self.footer_patterns = self.config.get('footer_patterns', [
            r'^Copyright\s*©.*$',            # Copyright
            r'^\d{4}年\d{1,2}月.*$',          # 日期
            r'^All rights reserved.*$',      # 版权声明
        ])
    
    def clean(self, text: str) -> str:
        """
        清洗文本
        
        Args:
            text: 原始文本
        
        Returns:
            清洗后的文本
        """
        if not text:
            return ""
        
        # 1. 去除页眉页脚
        text = self._remove_header_footer(text)
        
        # 2. 合并断行
        text = self._merge_broken_lines(text)
        
        # 3. 去除多余空白
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 4. 修正OCR错误
        text = self._fix_ocr_errors(text)
        
        # 5. 统一标点
        text = self._normalize_punctuation(text)
        
        # 6. 去除控制字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        return text.strip()
    
    def _remove_header_footer(self, text: str) -> str:
        """去除页眉页脚"""
        lines = text.split('\n')
        filtered = []
        
        for line in lines:
            line = line.strip()
            
            # 检查是否匹配页眉页脚模式
            is_header = any(re.match(p, line) for p in self.header_patterns)
            is_footer = any(re.match(p, line) for p in self.footer_patterns)
            
            if not is_header and not is_footer:
                filtered.append(line)
        
        return '\n'.join(filtered)
    
    def _merge_broken_lines(self, text: str) -> str:
        """合并断行"""
        # 中文断行：句末没有标点时合并
        text = re.sub(r'([^\n。！？.!?\d])\n([^\n])', r'\1\2', text)
        
        # 英文断行：单词中间断开
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
        
        return text
    
    def _fix_ocr_errors(self, text: str) -> str:
        """修正OCR常见错误"""
        # 常见OCR错误映射
        ocr_fixes = {
            'l': '1', 'O': '0', 'S': '5',  # 数字混淆
            'mAb': 'mAh', 'mA h': 'mAh',    # 单位
            'm A h': 'mAh',
            'G B': 'GB', 'M B': 'MB',
        }
        
        # 仅在数字上下文中替换
        for wrong, correct in ocr_fixes.items():
            # 检查是否在数字上下文
            text = re.sub(
                rf'(\d){wrong}(\d)',
                rf'\1{correct}\2',
                text
            )
        
        return text
    
    def _normalize_punctuation(self, text: str) -> str:
        """统一标点符号"""
        # 英文标点转中文（或反之）
        replacements = {
            '，': ',',
            '。': '.',
            '！': '!',
            '？': '?',
            '：': ':',
            '；': ';',
            '（': '(',
            '）': ')',
            '【': '[',
            '】': ']',
        }
        
        # 根据语言选择方向
        for cn, en in replacements.items():
            # 这里保持原文标点，可根据需求调整
            pass
        
        return text


class TextExtractor:
    """文本提取器"""
    
    def __init__(self, config: Dict = None):
        self.config = config or {}
        self.cleaner = DataCleaner(config)
    
    def extract_from_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """从纯文本文件提取"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        metadata = self._get_file_metadata(file_path)
        metadata.word_count = len(content)
        
        return content, metadata
    
    def extract_from_markdown(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """从Markdown文件提取"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # 提取标题（第一个#标题）
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else os.path.basename(file_path)
        
        metadata = self._get_file_metadata(file_path)
        metadata.title = title
        metadata.word_count = len(content)
        
        return content, metadata
    
    def extract_from_pdf(self, file_path: str) -> Tuple[str, DocumentMetadata, List[StructuredTable]]:
        """从PDF文件提取"""
        text = ""
        tables = []
        
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # 提取文本
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
                    
                    # 提取表格
                    for table in page.extract_tables():
                        if table:
                            structured = self._process_table(table, len(tables))
                            tables.append(structured)
            
            metadata = self._get_file_metadata(file_path)
            metadata.page_count = len(pdf.pages)
            metadata.word_count = len(text)
            
            return text, metadata, tables
            
        except ImportError:
            logger.warning("pdfplumber未安装，尝试使用pypdf")
            return self._extract_from_pdf_pypdf(file_path)
        except Exception as e:
            logger.error(f"PDF提取失败: {e}")
            return "", self._get_file_metadata(file_path), []
    
    def _extract_from_pdf_pypdf(self, file_path: str) -> Tuple[str, DocumentMetadata, List[StructuredTable]]:
        """使用pypdf提取PDF"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            
            metadata = self._get_file_metadata(file_path)
            metadata.page_count = len(reader.pages)
            metadata.word_count = len(text)
            
            return text, metadata, []
            
        except Exception as e:
            logger.error(f"pypdf提取失败: {e}")
            return "", self._get_file_metadata(file_path), []
    
    def extract_from_word(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """从Word文件提取"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            text = "\n\n".join(paragraphs)
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(self._process_table(table_data, len(tables)))
            
            metadata = self._get_file_metadata(file_path)
            metadata.word_count = len(text)
            metadata.doc_type = "word"
            
            return text, metadata
            
        except ImportError:
            logger.warning("python-docx未安装，无法处理Word文档")
            return "", self._get_file_metadata(file_path)
        except Exception as e:
            logger.error(f"Word提取失败: {e}")
            return "", self._get_file_metadata(file_path)
    
    def extract_from_html(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """从HTML文件提取"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # 移除脚本和样式
            for script in soup(['script', 'style']):
                script.decompose()
            
            # 提取文本
            text = soup.get_text(separator='\n')
            
            # 提取标题
            title = soup.title.string if soup.title else os.path.basename(file_path)
            
            metadata = self._get_file_metadata(file_path)
            metadata.title = title.strip()
            metadata.word_count = len(text)
            
            return text, metadata
            
        except ImportError:
            logger.warning("beautifulsoup4未安装，无法处理HTML")
            return "", self._get_file_metadata(file_path)
        except Exception as e:
            logger.error(f"HTML提取失败: {e}")
            return "", self._get_file_metadata(file_path)
    
    def extract_from_image(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """从图片提取（OCR）"""
        text = ""
        
        try:
            # 尝试使用pytesseract
            import pytesseract
            from PIL import Image
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            
        except ImportError:
            logger.warning("pytesseract/PIL未安装，无法进行OCR")
        except Exception as e:
            logger.error(f"OCR失败: {e}")
        
        metadata = self._get_file_metadata(file_path)
        metadata.doc_type = "image"
        
        return text, metadata
    
    def _process_table(self, table_data: List[List], table_id: int) -> StructuredTable:
        """处理表格数据"""
        if not table_data:
            return StructuredTable(
                id=f"table_{table_id}",
                headers=[],
                rows=[],
                markdown="",
                json_data=[],
                description=""
            )
        
        # 假设第一行是表头
        headers = [str(cell).strip() if cell else "" for cell in table_data[0]]
        rows = [[str(cell).strip() if cell else "" for cell in row] for row in table_data[1:]]
        
        # 生成Markdown
        md_lines = ["| " + " | ".join(headers) + " |"]
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            md_lines.append("| " + " | ".join(row) + " |")
        markdown = "\n".join(md_lines)
        
        # 生成JSON
        json_data = []
        for row in rows:
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    row_dict[header] = row[i]
            json_data.append(row_dict)
        
        # 生成自然语言描述
        description = self._table_to_description(headers, rows)
        
        return StructuredTable(
            id=f"table_{table_id}",
            headers=headers,
            rows=rows,
            markdown=markdown,
            json_data=json_data,
            description=description
        )
    
    def _table_to_description(self, headers: List[str], rows: List[List[str]]) -> str:
        """将表格转换为自然语言描述"""
        lines = []
        for row in rows[:5]:  # 最多5行
            parts = []
            for i, header in enumerate(headers):
                if i < len(row) and row[i]:
                    parts.append(f"{header}: {row[i]}")
            if parts:
                lines.append("，".join(parts))
        
        return "。".join(lines)
    
    def _get_file_metadata(self, file_path: str) -> DocumentMetadata:
        """获取文件元数据"""
        stat = os.stat(file_path)
        
        return DocumentMetadata(
            source=file_path,
            title=os.path.basename(file_path),
            file_size=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat()
        )


class SectionParser:
    """章节解析器"""
    
    # 章节标题模式
    SECTION_PATTERNS = [
        # 数字编号：第一章、第1章、1.、1.1
        r'^第[一二三四五六七八九十\d]+[章节篇部]',
        r'^\d+(\.\d+)*[\.\s]',
        r'^[一二三四五六七八九十]+[、.\s]',
        
        # Markdown标题
        r'^#{1,6}\s+',
    ]
    
    def __init__(self, config: Dict = None):
        self.config = config or {}
    
    def parse(self, text: str) -> List[DocumentSection]:
        """
        解析文档章节结构
        
        Args:
            text: 文档文本
        
        Returns:
            章节列表
        """
        sections = []
        lines = text.split('\n')
        
        current_section = None
        current_content = []
        current_start = 0
        char_pos = 0
        
        for line in lines:
            # 检查是否是章节标题
            is_section = self._is_section_title(line)
            
            if is_section:
                # 保存上一个章节
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    current_section.end_pos = char_pos
                    sections.append(current_section)
                
                # 开始新章节
                level = self._get_section_level(line)
                title = self._clean_section_title(line)
                
                current_section = DocumentSection(
                    level=level,
                    title=title,
                    content="",
                    path=self._build_path(sections, title, level),
                    start_pos=char_pos
                )
                current_content = []
            else:
                # 添加到当前章节内容
                if current_section:
                    current_content.append(line)
            
            char_pos += len(line) + 1  # +1 for newline
        
        # 保存最后一个章节
        if current_section:
            current_section.content = '\n'.join(current_content)
            current_section.end_pos = char_pos
            sections.append(current_section)
        
        return sections
    
    def _is_section_title(self, line: str) -> bool:
        """判断是否是章节标题"""
        line = line.strip()
        if not line:
            return False
        
        for pattern in self.SECTION_PATTERNS:
            if re.match(pattern, line):
                return True
        
        return False
    
    def _get_section_level(self, line: str) -> int:
        """获取章节层级"""
        line = line.strip()
        
        # Markdown标题
        if line.startswith('#'):
            return line.count('#')
        
        # 数字编号层级
        match = re.match(r'^(\d+(?:\.\d+)*)', line)
        if match:
            return match.group(1).count('.') + 1
        
        # 中文编号
        if re.match(r'^第[一二三四五六七八九十]+[章篇]', line):
            return 1
        if re.match(r'^第[一二三四五六七八九十]+[节]', line):
            return 2
        
        return 1
    
    def _clean_section_title(self, line: str) -> str:
        """清理章节标题"""
        title = line.strip()
        
        # 移除Markdown标记
        title = re.sub(r'^#+\s*', '', title)
        
        # 移除编号
        title = re.sub(r'^\d+(\.\d+)*[\.\s]*', '', title)
        title = re.sub(r'^第[一二三四五六七八九十\d]+[章节篇部]\s*', '', title)
        
        return title.strip()
    
    def _build_path(self, sections: List[DocumentSection], 
                    current_title: str, current_level: int) -> str:
        """构建章节路径"""
        path_parts = []
        
        for section in sections:
            if section.level < current_level:
                path_parts.append(section.title)
        
        path_parts.append(current_title)
        
        return ' > '.join(path_parts)


class DocumentPreprocessor:
    """文档预处理器 - 主入口类"""
    
    def __init__(self, config: Dict = None):
        """
        初始化预处理器
        
        Args:
            config: 配置参数
        """
        self.config = config or {}
        
        # 初始化组件
        self.format_detector = FormatDetector()
        self.text_extractor = TextExtractor(config)
        self.data_cleaner = DataCleaner(config)
        self.section_parser = SectionParser(config)
    
    def process(self, file_path: str) -> ProcessedDocument:
        """
        处理单个文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            ProcessedDocument: 处理后的文档
        """
        logger.info(f"处理文档: {file_path}")
        
        # 生成文档ID
        doc_id = self._generate_doc_id(file_path)
        
        try:
            # 1. 检测格式
            doc_format = self.format_detector.detect(file_path)
            
            # 2. 提取内容
            text, metadata, tables = self._extract_content(file_path, doc_format)
            
            # 3. 清洗数据
            cleaned_text = self.data_cleaner.clean(text)
            
            # 4. 解析章节
            sections = self.section_parser.parse(cleaned_text)
            
            # 5. 提取标题
            title = self._extract_title(cleaned_text, sections, file_path)
            
            # 6. 更新元数据
            metadata.word_count = len(cleaned_text)
            
            return ProcessedDocument(
                id=doc_id,
                title=title,
                content=cleaned_text,
                sections=sections,
                tables=tables,
                images=[],
                metadata=metadata,
                raw_text=text,
                status="success"
            )
            
        except Exception as e:
            logger.error(f"处理文档失败: {file_path}, 错误: {e}")
            
            return ProcessedDocument(
                id=doc_id,
                title=os.path.basename(file_path),
                content="",
                sections=[],
                tables=[],
                images=[],
                metadata=DocumentMetadata(source=file_path),
                status="failed",
                error=str(e)
            )
    
    def process_batch(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """
        批量处理文档
        
        Args:
            file_paths: 文件路径列表
        
        Returns:
            处理后的文档列表
        """
        results = []
        
        for file_path in file_paths:
            result = self.process(file_path)
            results.append(result)
        
        return results
    
    def _extract_content(self, file_path: str, 
                         doc_format: DocumentFormat) -> Tuple[str, DocumentMetadata, List[StructuredTable]]:
        """根据格式提取内容"""
        tables = []
        
        if doc_format == DocumentFormat.TEXT:
            text, metadata = self.text_extractor.extract_from_text(file_path)
        
        elif doc_format == DocumentFormat.MARKDOWN:
            text, metadata = self.text_extractor.extract_from_markdown(file_path)
        
        elif doc_format == DocumentFormat.PDF:
            text, metadata, tables = self.text_extractor.extract_from_pdf(file_path)
        
        elif doc_format == DocumentFormat.WORD:
            text, metadata = self.text_extractor.extract_from_word(file_path)
        
        elif doc_format == DocumentFormat.HTML:
            text, metadata = self.text_extractor.extract_from_html(file_path)
        
        elif doc_format == DocumentFormat.IMAGE:
            text, metadata = self.text_extractor.extract_from_image(file_path)
        
        else:
            # 尝试作为文本处理
            try:
                text, metadata = self.text_extractor.extract_from_text(file_path)
            except:
                text = ""
                metadata = DocumentMetadata(source=file_path)
        
        return text, metadata, tables
    
    def _extract_title(self, text: str, sections: List[DocumentSection], 
                       file_path: str) -> str:
        """提取文档标题"""
        # 优先从章节提取
        if sections:
            first_section = sections[0]
            if first_section.level == 1:
                return first_section.title
        
        # 从文本开头提取
        lines = text.strip().split('\n')
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 3 and len(line) < 100:
                # 可能是标题
                if not line.startswith(('#', '*', '-', '>', '|')):
                    return line
        
        # 使用文件名
        return os.path.basename(file_path)
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:12]
    
    def to_indexable_format(self, doc: ProcessedDocument) -> List[Dict]:
        """
        转换为可索引格式（用于向量数据库）
        
        Args:
            doc: 处理后的文档
        
        Returns:
            可索引的数据列表
        """
        items = []
        
        # 按章节切分
        if doc.sections:
            for section in doc.sections:
                if section.content.strip():
                    items.append({
                        'id': f"{doc.id}_section_{section.level}_{hash(section.title) % 10000}",
                        'document_id': doc.id,
                        'text': f"{section.title}\n{section.content}",
                        'metadata': {
                            'document_title': doc.title,
                            'section_path': section.path,
                            'section_level': section.level,
                            'source': doc.metadata.source,
                        }
                    })
        else:
            # 无章节，整体作为一个块
            items.append({
                'id': doc.id,
                'document_id': doc.id,
                'text': doc.content,
                'metadata': {
                    'document_title': doc.title,
                    'source': doc.metadata.source,
                }
            })
        
        # 添加表格
        for table in doc.tables:
            items.append({
                'id': f"{doc.id}_table_{table.id}",
                'document_id': doc.id,
                'text': table.description,
                'metadata': {
                    'type': 'table',
                    'markdown': table.markdown,
                    'source': doc.metadata.source,
                }
            })
        
        return items


# 使用示例
if __name__ == '__main__':
    # 创建预处理器
    preprocessor = DocumentPreprocessor()
    
    # 测试文本处理
    test_text = """
第一章 产品介绍

1.1 产品概述

本产品是一款智能设备，主要参数如下：
- 型号：X-1000
- 电池容量：5000mAh
- 充电时间：约2小时

1.2 技术规格

技术参数表：
| 参数 | 数值 |
|------|------|
| 尺寸 | 150x75x8mm |
| 重量 | 180g |
| 内存 | 8GB |

第二章 使用说明

2.1 开机步骤

1. 长按电源键3秒
2. 等待系统启动
3. 完成初始化设置
"""
    
    # 模拟处理
    import tempfile
    
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write(test_text)
        temp_path = f.name
    
    try:
        result = preprocessor.process(temp_path)
        
        print("\n=== 文档预处理结果 ===\n")
        print(f"文档ID: {result.id}")
        print(f"标题: {result.title}")
        print(f"状态: {result.status}")
        print(f"字数: {result.metadata.word_count}")
        
        print("\n章节结构:")
        for section in result.sections:
            indent = "  " * (section.level - 1)
            print(f"{indent}{section.path} ({len(section.content)}字)")
        
        if result.tables:
            print("\n表格:")
            for table in result.tables:
                print(f"  - {table.id}: {len(table.rows)}行 x {len(table.headers)}列")
                print(f"    描述: {table.description[:100]}...")
        
        # 转换为索引格式
        indexable = preprocessor.to_indexable_format(result)
        print(f"\n可索引项数: {len(indexable)}")
        
    finally:
        os.unlink(temp_path)

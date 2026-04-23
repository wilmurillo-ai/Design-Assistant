#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:
    raise SystemExit(
        "Missing python-docx. Run `python scripts/bootstrap_env.py` first."
    ) from exc


TIMESTAMP_RE = re.compile(r"^\[(\d{2}:\d{2}:\d{2}(?:-\d{2}:\d{2}:\d{2})?)\]\s*(.*)$")


def _set_font(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)


def write_docx_from_text(text: str, output_path: Path, title: str | None = None) -> None:
    document = Document()
    _set_normal_style(document)

    if title:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        _set_font(title_run, 16, bold=True)

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            continue

        timestamp_match = TIMESTAMP_RE.match(line)
        if timestamp_match:
            document.add_heading(timestamp_match.group(1), level=2)
            trailing_text = timestamp_match.group(2).strip()
            if trailing_text:
                document.add_paragraph(trailing_text)
            continue

        document.add_paragraph(line)

    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a UTF-8 transcript or chapterized text file into DOCX.")
    parser.add_argument("--input", required=True, help="Input text file")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--title", help="Optional document title")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    text = input_path.read_text(encoding="utf-8")
    title = args.title or input_path.stem
    write_docx_from_text(text, output_path, title=title)
    print(f"Saved: {output_path.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""
TikTok 视频审核脚本
用法: python3 tiktok_audit.py <tiktok_url> [--category hotel] [--output-json] [--output-docx]
"""
import sys, os, base64, subprocess, json, time, wave, re, argparse, json5
import numpy as np
import cv2
from moviepy.video.io.VideoFileClip import VideoFileClip
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ============ 配置 ============
API_KEY = "sk-0zy1YyzLaabc5rPfqSvC5tF16os74HycNcvXLlMrVr0FXtNz"
API_BASE = "https://api.lingyaai.cn"
API_ENDPOINT = f"{API_BASE}/v1/chat/completions"
MODEL = "gemini-3-flash-preview"  # 默认模型，可通过 --model 切换
YTDLP_BIN = "/Users/apple/Library/Python/3.9/bin/yt-dlp"
OUTPUT_DIR = "output/tiktok_downloads_test"
LINGYA_TEMP = "/tmp/lingya_payload.json"

# ============ SOP 加载 ============
SOP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'references', 'categories')
VALID_CATEGORIES = ['hotel', 'restaurant', 'product', 'travel']
CATEGORY_NAMES = {
    'hotel': '酒店',
    'restaurant': '餐饮',
    'product': '商品',
    'travel': '旅行/景点',
}

def load_sop_prompt(category='hotel'):
    if category not in VALID_CATEGORIES:
        category = 'hotel'
    sop_path = os.path.join(SOP_DIR, f'sop_{category}.md')
    if os.path.exists(sop_path):
        with open(sop_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    return None

def build_sop_prompt(category='hotel'):
    sop_content = load_sop_prompt(category)
    if not sop_content:
        return None
    category_name = CATEGORY_NAMES.get(category, category)

    json_format = """{
  "score": 数字1-10,
  "conclusion": "通过"或"待修改"或"不通过",
  "frames": [
    {"second": 0, "content": "画面内容描述", "subtitle": "字幕文字", "audio": "音频状态(正常/偏弱/静音)", "sop": "该秒SOP符合性评估"},
    ...共10项(0-9秒)
  ],
  "audio_evaluation": "音频质量评估文字",
  "sop_checks": [
    {"id": 1, "name": "USP钩子", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 2, "name": "封面标题", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 3, "name": "房间参观", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 4, "name": "重复展示", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 5, "name": "POI贴纸", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 6, "name": "价格菜单", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 7, "name": "促销信息", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 8, "name": "封面规范", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 9, "name": "字幕同步", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 10, "name": "真人入镜", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 11, "name": "音画同步", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 12, "name": "画面稳定", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 13, "name": "光线明亮", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 14, "name": "字幕位置", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 15, "name": "BGM音量", "result": "✅或❌或⚠️或-", "remark": "原因说明"},
    {"id": 16, "name": "发布标签", "result": "✅或❌或⚠️或-", "remark": "原因说明"}
  ],
  "improvements": ["改进建议1", "改进建议2", "改进建议3"]
}"""

    prompt = (
        f"你是印尼{category_name} TikTok 内容审核专家。请严格审核视频前10秒,并以JSON格式返回审核结果。\n\n"
        f"请参考以下 SOP 标准进行审核:\n\n{sop_content}\n\n"
        f"重要:必须严格以JSON格式返回审核结果,包含以下所有字段:\n\n"
        f"{json_format}\n\n"
        f"frames数组必须恰好10项(0-9秒),每项必须包含所有字段。"
        f"sop_checks 必须是数组格式,每个元素包含 id(数字)、name(中文名)、result(✅❌⚠️或-)、remark(原因)。\n"
        f"重要:sop_checks 必须恰好包含全部 16 个检查项,缺一不可!"
    )
    return prompt


# ============ 工具函数 ============
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_tbl(doc, headers, data_rows, header_color='E8F4FD', first_col_color='F5F9FF'):
    n = 1 + len(data_rows)
    t = doc.add_table(rows=n, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = t.rows[0]
    for j, hd in enumerate(headers):
        h.cells[j].text = hd
        set_cell_bg(h.cells[j], header_color)
        h.cells[j].paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(data_rows):
        r = t.rows[i+1]
        for j, val in enumerate(row_data):
            r.cells[j].text = str(val)
        set_cell_bg(r.cells[0], first_col_color)
    return t

def resolve_tiktok_url(url):
    if 'vm.tiktok.com' in url or 'tiktok.com' not in url:
        try:
            result = subprocess.run(
                ['curl', '-sI', url], capture_output=True, text=True, timeout=10
            )
            for line in result.stdout.splitlines():
                if line.lower().startswith('location:'):
                    resolved = line.split(':', 1)[1].strip()
                    print(f"  短链接解析: {url} → {resolved}")
                    return resolved
            return url
        except:
            return url
    return url

def extract_video_id(url):
    match = re.search(r'/video/(\d+)', url)
    if match:
        return match.group(1)
    match = re.search(r'video/([\w-]+)', url)
    if match:
        return match.group(1)
    return None

def fetch_post_metadata(url):
    """通过yt-dlp提取TikTok帖子描述（含话题标签和@账号），用于第16项发布标签审核"""
    try:
        result = subprocess.run(
            [YTDLP_BIN, '--dump-json', '--no-check-certificate', url],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout:
            data = json.loads(result.stdout)
            desc = data.get('description', '') or data.get('title', '') or ''
            hashtags = re.findall(r'#(\w+)', desc)
            mentions = re.findall(r'@(\w+)', desc)
            return {
                'description': desc,
                'hashtags': hashtags,
                'mentions': mentions
            }
    except:
        pass
    return {'description': '', 'hashtags': [], 'mentions': []}

def download_video(url, output_dir=OUTPUT_DIR):
    """下载视频：优先 yt-dlp，失败时切换 dlpanda"""
    os.makedirs(output_dir, exist_ok=True)

    # 方案1: yt-dlp
    vid_id = extract_video_id(url)
    for ext in ['mp4', 'mp3', 'mkv', 'webm']:
        existing = os.path.join(output_dir, f"{vid_id}.{ext}")
        if os.path.exists(existing):
            print(f"  [下载] 文件已存在（复用）: {existing}")
            return existing

    print(f"  [下载] 尝试 yt-dlp...")
    output_template = os.path.join(output_dir, '%(id)s.%(ext)s')
    cmd = [YTDLP_BIN, '-o', output_template, '--no-check-certificate', url]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode == 0:
        for ext in ['mp4', 'mp3', 'mkv', 'webm']:
            path = os.path.join(output_dir, f"{vid_id}.{ext}")
            if os.path.exists(path):
                size_mb = os.path.getsize(path) / (1024 * 1024)
                print(f"  [下载] yt-dlp 成功: {size_mb:.1f} MB")
                return path

    print(f"  [下载] yt-dlp 失败，尝试 dlpanda...")

    # 方案2: dlpanda（需要 agent 执行 browser，这里只是尝试用脚本方式下载）
    # 实际备用流程见 tiktok-dlpanda skill
    print(f"  [下载] 请使用 tiktok-dlpanda skill 进行备用下载")
    return None

def check_has_video(filepath):
    try:
        cap = cv2.VideoCapture(filepath)
        fps = cap.get(cv2.CAP_PROP_FPS)
        total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = total / fps if fps > 0 else 0
        cap.release()
        return total > 0 and fps > 0, duration, fps, total
    except:
        return False, 0, 0, 0

def extract_frames_and_audio(filepath, target_dur=10):
    has_video, duration, fps, total = check_has_video(filepath)

    frames = []
    if has_video:
        cap = cv2.VideoCapture(filepath)
        for i in range(10):
            t = target_dur * i / (n := 9)
            cap.set(cv2.CAP_PROP_POS_MSEC, t * 1000)
            ret, frame = cap.read()
            if ret:
                frame_small = cv2.resize(frame, (640, 360), interpolation=cv2.INTER_AREA)
                path = f"/tmp/audit_frame_{i:02d}.jpg"
                cv2.imwrite(path, frame_small, [cv2.IMWRITE_JPEG_QUALITY, 60])
                frames.append(path)
        cap.release()

    audio_data = {
        'has_audio': False, 'duration': 0, 'rms': 0, 'peak': 0,
        'silent_ratio': 0, 'has_video': has_video, 'video_duration': duration,
        'per_second': []
    }

    if has_video:
        try:
            clip = VideoFileClip(filepath)
            sub = clip.subclipped(0, target_dur)
            sub.audio.write_audiofile("/tmp/audit_audio.wav")
            sub.close()
            clip.close()

            with wave.open("/tmp/audit_audio.wav", 'rb') as w:
                rate = w.getframerate()
                nframes = w.getnframes()
                data = np.frombuffer(w.readframes(nframes), dtype=np.int16)
                dur_a = nframes / rate
                data_f = data.astype(np.float32) / 32768.0

                rms_all = np.sqrt(np.mean(data_f**2))
                peak_all = np.max(np.abs(data_f))
                db_rms = 20 * np.log10(rms_all + 1e-9)
                db_peak = 20 * np.log10(peak_all + 1e-9)
                silent_ratio = float(np.mean(data_f**2 < 0.01))

                chunk = rate
                per_second = []
                for i in range(target_dur):
                    seg = data_f[i*chunk:(i+1)*chunk]
                    active = seg[np.abs(seg) > 0.01]
                    db_s = 20 * np.log10(np.sqrt(np.mean(active**2)) + 1e-9) if len(active) > 0 else -999
                    pct = np.sum(np.abs(seg) > 0.01) / len(seg) * 100
                    per_second.append({'rms': db_s, 'active_pct': pct})

                audio_data = {
                    'has_audio': True, 'duration': dur_a,
                    'rms': db_rms, 'peak': db_peak,
                    'silent_ratio': silent_ratio,
                    'has_video': has_video, 'video_duration': duration,
                    'per_second': per_second
                }
        except Exception as e:
            print(f"  音频提取失败: {e}")

    return frames, audio_data

def call_llm(frames, audio_data, category='hotel', api_key=API_KEY, post_metadata=None, model=None):
    _model = model or MODEL
    sop_prompt = build_sop_prompt(category)
    if not sop_prompt:
        print("  错误: 无法加载 SOP")
        return None

    image_parts = []
    for fp in frames:
        if os.path.exists(fp):
            with open(fp, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode('utf-8')
            image_parts.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
            })

    ps = audio_data.get('per_second', [])
    per_sec_text = '\n'.join([
        f"第{i}s: RMS={d.get('rms', -999):.1f}dBFS, 有效样本={d.get('active_pct', 0):.0f}%"
        for i, d in enumerate(ps[:10])
    ])

    audio_info = f"""音频分析数据(前10秒,每秒1帧):
{per_sec_text}
整体RMS: {audio_data.get('rms', 0):.1f} dBFS
峰值电平: {audio_data.get('peak', 0):.1f} dBFS
静音比例: {audio_data.get('silent_ratio', 0)*100:.1f}%
视频时长: {audio_data.get('video_duration', 0):.1f}s
有视频流: {'是' if audio_data.get('has_video') else '否'}
有音频流: {'是' if audio_data.get('has_audio') else '否'}"""

    # 附加帖子元数据（话题标签和@账号），用于第16项发布标签审核
    if post_metadata and (post_metadata.get('hashtags') or post_metadata.get('mentions')):
        _ht = ' '.join([f'#{h}' for h in post_metadata.get('hashtags', [])])
        _mt = ' '.join([f'@{m}' for m in post_metadata.get('mentions', [])])
        _meta_txt = f"""
--- 帖子发布信息（第16项审核依据）---
话题标签: {_ht if _ht else '（未检测到）'}
@账号: {_mt if _mt else '（未检测到）'}
说明: 请根据上述话题标签和@账号信息，判断第16项【发布标签】是否符合SOP要求。
"""
    else:
        _meta_txt = """
--- 帖子发布信息（第16项审核依据）---
话题标签: （未能获取，请标注为无法审核）
@账号: （未能获取，请标注为无法审核）
"""

    content = [
        *image_parts,
        {"type": "text", "text": f"{sop_prompt}\n\n{audio_info}{_meta_txt}"}
    ]

    payload = {
        "model": _model,
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 8000
    }

    with open(LINGYA_TEMP, 'w') as f:
        json.dump(payload, f)

    print("  发送审核请求...")
    start = time.time()
    result = subprocess.run([
        'curl', '-s', '-X', 'POST', API_ENDPOINT,
        '-H', f'Authorization: Bearer {api_key}',
        '-H', 'Content-Type: application/json',
        '-d', f'@{LINGYA_TEMP}'
    ], capture_output=True, text=True, timeout=600)
    elapsed = time.time() - start
    os.unlink(LINGYA_TEMP)

    if result.stdout:
        resp = json.loads(result.stdout)
        msg = resp.get('choices', [{}])[0].get('message', {}).get('content', '')
        usage = resp.get('usage', {})
        pt = usage.get('prompt_tokens', 0)
        ct = usage.get('completion_tokens', 0)
        tt = usage.get('total_tokens', pt + ct)

        def _try_parse(_raw):
            """先用json解析，失败则用json5兜底"""
            try:
                return json.loads(_raw)
            except:
                try:
                    return json5.loads(_raw)
                except:
                    return None

        # 尝试解析结构化JSON(支持 markdown 代码块包裹)
        structured = None
        try:
            # 预处理:去掉 markdown 代码块包裹
            msg_clean = msg.strip()

            # 方法1: 用正则找 ```json ... ``` 代码块(最可靠)
            # 关键修复：用贪婪匹配 + 括号计数确保找到最外层 { ... }
            import re as _re_json
            _block_m = _re_json.search(r'```(?:json)?\s*(\{[\s\S]*\})\s*```', msg_clean, _re_json.MULTILINE)
            if _block_m:
                _raw = _block_m.group(1)
                # 括号计数：从找到的 { 中找最外层配对的 }
                _depth = 0
                _start = -1
                _end = -1
                for _i, _ch in enumerate(_raw):
                    if _ch == '{':
                        if _depth == 0:
                            _start = _i
                        _depth += 1
                    elif _ch == '}':
                        _depth -= 1
                        if _depth == 0:
                            _end = _i
                            break
                if _start >= 0 and _end > _start:
                    _cand = _try_parse(_raw[_start:_end+1])
                    if _cand and ('score' in _cand or 'frames' in _cand or 'sop_checks' in _cand):
                        structured = _cand

            # 方法2: 对每个 { 位置,用括号计数找匹配的 },再验证是否是有效JSON
            if not structured:
                _txt = msg_clean
                _brace_positions = [i for i, c in enumerate(_txt) if c == '{']
                _candidates = []
                for _pos in _brace_positions:
                    _depth = 0
                    _end_pos = -1
                    for _j in range(_pos, len(_txt)):
                        if _txt[_j] == '{':
                            _depth += 1
                        elif _txt[_j] == '}':
                            _depth -= 1
                            if _depth == 0:
                                _end_pos = _j
                                break
                    if _end_pos > _pos:
                        _c = _try_parse(_txt[_pos:_end_pos+1])
                        if _c and ('score' in _c or 'frames' in _c or 'sop_checks' in _c):
                            _candidates.append((_end_pos - _pos, _c))
                if _candidates:
                    _candidates.sort(key=lambda x: x[0], reverse=True)
                    structured = _candidates[0][1]

            # 方法3:直接解析清理后的文本(兜底)
            if not structured:
                _cand = _try_parse(msg_clean.strip())
                if _cand and ('score' in _cand or 'frames' in _cand or 'sop_checks' in _cand):
                    structured = _cand

            # 方法4: 截断兜底 — 当 JSON 在中途被截断时，用正则从原始文本提取各字段
            if not structured:
                _partial = {}
                # 提取 score
                _sm = re.search(r'"score"\s*:\s*(\d+)', msg_clean)
                if _sm:
                    _partial['score'] = int(_sm.group(1))
                # 提取 conclusion
                _cm = re.search(r'"conclusion"\s*:\s*"([^"]+)"', msg_clean)
                if _cm:
                    _partial['conclusion'] = _cm.group(1)
                # 提取 audio_evaluation
                _am = re.search(r'"audio_evaluation"\s*:\s*"([^"]+)"', msg_clean)
                if _am:
                    _partial['audio_evaluation'] = _am.group(1)
                # 提取 improvements 数组（多行）
                _im_lines = []
                for _line in msg_clean.split('\n'):
                    if any(kw in _line.lower() for kw in ['改进', '建议', 'improve', 'suggest']):
                        _lm2 = re.search(r'"([^"]+)"\s*[,}"]', _line)
                        if _lm2:
                            _im_lines.append(_lm2.group(1))
                if _im_lines:
                    _partial['improvements'] = _im_lines
                # 提取 sop_checks 数组（截断的JSON数组，逐项提取）
                _sc_start = msg_clean.find('"sop_checks"')
                if _sc_start >= 0:
                    _sc_txt = msg_clean[_sc_start:]
                    _sc_items = []
                    for _scm in re.finditer(
                            r'"id"\s*:\s*(\d+)\s*,\s*"name"\s*:\s*"([^"]+)"\s*,\s*"result"\s*:\s*"([^"]+)"\s*,\s*"remark"\s*:\s*"([^"]+)"',
                            _sc_txt):
                        _sc_items.append({
                            'id': int(_scm.group(1)),
                            'name': _scm.group(2),
                            'result': _scm.group(3),
                            'remark': _scm.group(4)
                        })
                    if _sc_items:
                        _partial['sop_checks'] = _sc_items
                # 提取 frames 数组
                _fr_start = msg_clean.find('"frames"')
                if _fr_start >= 0:
                    _fr_txt = msg_clean[_fr_start:]
                    _fr_items = []
                    for _frm in re.finditer(
                            r'"second"\s*:\s*(\d+)\s*,\s*"content"\s*:\s*"([^"]+)"\s*,\s*"subtitle"\s*:\s*"([^"]+)"\s*,\s*"audio"\s*:\s*"([^"]+)"\s*,\s*"sop"\s*:\s*"([^"]+)"',
                            _fr_txt):
                        _fr_items.append({
                            'second': int(_frm.group(1)),
                            'content': _frm.group(2),
                            'subtitle': _frm.group(3),
                            'audio': _frm.group(4),
                            'sop': _frm.group(5)
                        })
                    if _fr_items:
                        _partial['frames'] = _fr_items
                if _partial:
                    structured = _partial
        except:
            structured = None

        return {
            'report': msg,
            'structured': structured,
            'elapsed': elapsed,
            'prompt_tokens': pt,
            'completion_tokens': ct,
            'total_tokens': tt,
            'input_cost': pt / 1e6 * 2,
            'output_cost': ct / 1e6 * 8,
            'total_cost': pt / 1e6 * 2 + ct / 1e6 * 8
        }
    return None

def parse_score(result):
    if not result:
        return None
    if result.get('structured') and 'score' in result['structured']:
        return int(result['structured']['score'])
    txt = result.get('report', '')
    m = re.search(r'"score"\s*:\s*(\d+)', txt)
    if m:
        return int(m.group(1))
    m = re.search(r'评分[::]\s*(\d+)', txt)
    if m:
        return int(m.group(1))
    m = re.search(r'(\d+)\s*/\s*10', txt)
    if m:
        return int(m.group(1))
    return None

def build_json_report(url, video_path, frames, audio_data, result, video_id, phase_timings=None, model=None):
    import re as _re_jb

    def to_py(v):
        if hasattr(v, 'item'):
            return v.item()
        return v

    def to_py_per_second(lst):
        r = []
        for item in lst:
            if isinstance(item, dict):
                r.append({k: to_py(v) for k, v in item.items()})
            else:
                r.append(to_py(item))
        return r

    # ─── 标准化 structured 数据:英文 key ───
    raw_st = result.get('structured') if result else None
    std_st = None
    std_checks = []
    if raw_st:
        std_st = {}
        std_st["score"] = raw_st.get("score")
        std_st["conclusion"] = raw_st.get("conclusion")
        std_st["audio_evaluation"] = raw_st.get("audio_evaluation")
        std_st["improvements"] = raw_st.get("improvements", [])

        _sop_ids_map = {
            "USP钩子": 1, "封面标题": 2, "房间参观": 3, "重复展示": 4,
            "POI贴纸": 5, "价格菜单": 6, "促销信息": 7, "封面规范": 8,
            "字幕同步": 9, "真人入镜": 10, "音画同步": 11, "画面稳定": 12,
            "光线明亮": 13, "字幕位置": 14, "BGM音量": 15, "发布标签": 16,
        }

        # sop_checks: AI已返回list格式 {id,name,result,remark} → 直接使用
        _raw_sc = raw_st.get("sop_checks") or []
        if isinstance(_raw_sc, list):
            # 新格式:已含 id/name/result/remark,直接用
            for _item in _raw_sc:
                if isinstance(_item, dict) and "id" in _item:
                    std_checks.append(_item)
                elif isinstance(_item, dict):
                    # 老格式 dict {中文key: result},做转换
                    for _k, _v in _item.items():
                        _full = str(_v).strip()
                        _m2 = _re_jb.search(r'[✅❌⚠️-]', _full)
                        _mark = _m2.group(0) if _m2 else '-'
                        _mr2 = _re_jb.search(r'[((]([^))]+)[))]', _full)
                        _reason = _mr2.group(1).strip() if _mr2 else '-'
                        std_checks.append({"id": _sop_ids_map.get(_k, 99), "name": _k, "result": _mark, "remark": _reason})
            std_checks.sort(key=lambda x: x.get("id", 99))
        std_st["sop_checks"] = std_checks

        # frames: 中文key → 英文key
        std_frames = []
        for f in (raw_st.get("frames") or []):
            std_frames.append({
                "id": f.get("second", 0),
                "name": f.get("content", ""),
                "subtitle": f.get("subtitle", ""),
                "result": f.get("audio", ""),
                "remark": f.get("sop", "")
            })
        std_st["frames"] = std_frames

    # ─── 两阶段独立评分 ───
    PHASE1_IDS = set(range(1, 16))   # 视频内容规范: ID 1-15
    PHASE2_IDS = {16}                # 发布规范: ID 16

    def _calc_phase(checks, phase_ids):
        """计算阶段得分: (通过数/总数)*10"""
        if not phase_ids:
            return None, "待判定", 0, 0
        total = len(phase_ids)
        passed = sum(1 for c in checks if c.get('id') in phase_ids and c.get('result', '').strip() == '✅')
        failed = sum(1 for c in checks if c.get('id') in phase_ids and c.get('result', '').strip() == '❌')
        # 包含 ⚠️ 和 - (未检测/可选) 的不算失败
        other = total - passed - failed
        score = round((passed / total) * 10, 1) if total > 0 else None
        if score is None:
            conclusion = "待判定"
        elif score >= 8:
            conclusion = "通过"
        elif score >= 3:
            conclusion = "待修改"
        else:
            conclusion = "不通过"
        return score, conclusion, passed, total

    _p1_score, _p1_conclusion, _p1_passed, _p1_total = _calc_phase(std_checks, PHASE1_IDS)
    _p2_score, _p2_conclusion, _p2_passed, _p2_total = _calc_phase(std_checks, PHASE2_IDS)

    # 无视频流 → 直接不通过
    if not audio_data.get('has_video'):
        _overall = "不通过"
        _overall_score = 0
    elif _p1_score is None or _p2_score is None:
        _overall = "不通过"
        _overall_score = 0
    elif _p1_conclusion == "通过" and _p2_conclusion == "通过":
        _overall = "通过"
        _overall_score = round((_p1_score + _p2_score) / 2, 1)
    else:
        _overall = "待修改"
        _overall_score = round((_p1_score + _p2_score) / 2, 1)

    ad = audio_data
    report = {
        "video_url": url,
        "video_id": video_id,
        "local_path": str(video_path) if video_path else None,
        "audit_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "audio_analysis": {
            "rms_dbfs": round(float(to_py(ad.get('rms', 0))), 1),
            "peak_dbfs": round(float(to_py(ad.get('peak', 0))), 1),
            "silent_ratio_percent": round(float(to_py(ad.get('silent_ratio', 0))) * 100, 1),
            "video_duration_sec": round(float(to_py(ad.get('video_duration', 0))), 1),
            "has_video_stream": bool(ad.get('has_video', False)),
            "has_audio_stream": bool(ad.get('has_audio', False)),
            "per_second": to_py_per_second(ad.get('per_second', []))
        },
        "audit_result": {
            "model": model or MODEL,
            "score": _overall_score,
            "conclusion": _overall,
            "phase1_score": _p1_score,
            "phase1_conclusion": _p1_conclusion,
            "phase1_passed": _p1_passed,
            "phase1_total": _p1_total,
            "phase2_score": _p2_score,
            "phase2_conclusion": _p2_conclusion,
            "phase2_passed": _p2_passed,
            "phase2_total": _p2_total,
            "total_tokens": result['total_tokens'] if result else None,
            "elapsed_sec": round(result['elapsed'], 1) if result else None,
            "cost_yuan": round(result['total_cost'], 6) if result else None
        },
        "phase_timings": {
            "total_sec": round(phase_timings['total'], 1) if phase_timings else None,
            "download_sec": round(phase_timings['download'], 1) if phase_timings else None,
            "analysis_sec": round(phase_timings['analysis'], 1) if phase_timings else None,
            "post_meta_sec": round(phase_timings.get('post_meta', 0), 1) if phase_timings else None,
            "audit_sec": round(phase_timings['audit'], 1) if phase_timings else None,
            "report_sec": round(phase_timings['report'], 1) if phase_timings else None
        },
        "structured": std_st,
        "report_content": result['report'] if result else "审核失败"
    }
    return report

def build_docx(url, video_path, frames, audio_data, result, video_id, phase_timings=None, model=None):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.5)

    structured = result.get('structured') if result else None

    # ========== 标题 ==========
    t = doc.add_heading('TikTok 酒店视频审核报告', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.size = Pt(16)

    # ========== 一、基本信息 ==========
    doc.add_paragraph('')
    h = doc.add_paragraph()
    rh = h.add_run('一、基本信息')
    rh.bold = True; rh.font.size = Pt(12)
    add_tbl(doc, ['项目', '内容'], [
        ('视频链接', url),
        ('视频ID', video_id),
        ('本地路径', str(video_path) if video_path else '无'),
        ('审核时间', datetime.now().strftime('%Y-%m-%d %H:%M')),
    ])

    # ========== 二、审核结论（两阶段独立评分） ==========
    doc.add_paragraph('')
    _audit = result.get('audit_result', {}) if result else {}
    _p1_s = _audit.get('phase1_score')
    _p1_c = _audit.get('phase1_conclusion', '待判定')
    _p1_p = _audit.get('phase1_passed', 0)
    _p1_t = _audit.get('phase1_total', 15)
    _p2_s = _audit.get('phase2_score')
    _p2_c = _audit.get('phase2_conclusion', '待判定')
    _p2_p = _audit.get('phase2_passed', 0)
    _p2_t = _audit.get('phase2_total', 1)
    _overall_s = _audit.get('score', 0)
    _overall_c = _audit.get('conclusion', '不通过')

    def _emoji(v): return {'通过': '✅', '待修改': '⚠️', '不通过': '❌', '待判定': '⏳'}.get(v, '⏳')
    def _disp(s, c): return f'{s} / 10 {_emoji(c)} {c}'

    h = doc.add_paragraph()
    rh = h.add_run('二、审核结论（两阶段独立评分）')
    rh.bold = True; rh.font.size = Pt(12)
    add_tbl(doc, ['评分项', '结果'], [
        ('阶段一: 视频内容规范', _disp(_p1_s, _p1_c) if _p1_s else f'待判定 {_emoji("待判定")} 待判定'),
        ('  └ 通过/总分', f'{_p1_p}/{_p1_t} 项通过'),
        ('阶段二: 发布规范',    _disp(_p2_s, _p2_c) if _p2_s else f'待判定 {_emoji("待判定")} 待判定'),
        ('  └ 通过/总分', f'{_p2_p}/{_p2_t} 项通过'),
        ('综合结论', f'{_overall_s} / 10 {_emoji(_overall_c)} {_overall_c}'),
        ('有视频流', '是 ✅' if audio_data.get('has_video') else '否 ❌(纯音频文件)'),
        ('有音频流', '是 ✅' if audio_data.get('has_audio') else '否 ❌'),
    ])

    # ========== 三、前10秒逐帧描述 ==========
    doc.add_paragraph('')
    h = doc.add_paragraph()
    rh = h.add_run('三、前10秒逐帧精细描述(每1秒1帧)')
    rh.bold = True; rh.font.size = Pt(12)

    # 统一 frames 格式:原始 AI 返回的 second/content/subtitle/audio/sop
    # → 标准格式 id/name/subtitle/result/remark
    _raw_frames = structured.get('frames') if structured else None
    _std_frames = []
    if isinstance(_raw_frames, list):
        for _f in (_raw_frames or [])[:10]:
            if isinstance(_f, dict):
                _std_frames.append({
                    'id': _f.get('second', _f.get('id', 0)),
                    'name': _f.get('content', _f.get('name', '-')),
                    'subtitle': _f.get('subtitle', '-'),
                    'result': _f.get('audio', _f.get('result', '-')),
                    'remark': _f.get('sop', _f.get('remark', '-')),
                })
            else:
                _std_frames.append({'id': 0, 'name': '-', 'subtitle': '-', 'result': '-', 'remark': '-'})

    frame_rows = []
    if _std_frames:
        for f in _std_frames[:10]:
            sec = f.get('id', 0)
            frame_rows.append((
                f'第{sec}s',
                f.get('name', '-')[:120],
                f.get('subtitle', '-'),
                f.get('result', '-'),
                f.get('remark', '-')
            ))
    else:
        for i in range(10):
            frame_rows.append((f'第{i}s', '（见AI报告）', '—', '—', '—'))

    add_tbl(doc, ['秒数', '画面内容', '字幕', '音频状态', 'SOP符合性'], frame_rows)
    # ========== 四、音频质量分析 ==========
    doc.add_paragraph('')
    h = doc.add_paragraph()
    rh = h.add_run('四、音频质量分析(修订标准)')
    rh.bold = True; rh.font.size = Pt(12)

    ps = audio_data.get('per_second', [])
    audio_rows = []
    for i in range(min(10, len(ps))):
        d = ps[i]
        rms_s = d.get('rms', -999)
        pct_s = d.get('active_pct', 0)
        if rms_s > -100:
            status = '✅ 有声' if pct_s > 5 else '❌ 静音'
            audio_rows.append((f'第{i}s', f'{rms_s:.1f} dBFS', f'{pct_s:.0f}%', status))
        else:
            audio_rows.append((f'第{i}s', '-180.0', '0%', '❌ 静音'))

    add_tbl(doc, ['秒数', 'RMS(dBFS)', '有效样本%', '状态'], audio_rows)

    doc.add_paragraph('')
    ad = audio_data
    ad_rms = ad.get('rms', 0)
    ad_peak = ad.get('peak', 0)
    ad_silent = ad.get('silent_ratio', 0) * 100
    if ad.get('has_video'):
        if ad_rms > -12:
            audio_judge = '✅ 优秀(超出标准)'
        elif ad_rms >= -18:
            audio_judge = '✅ 合格'
        elif ad_rms >= -35:
            audio_judge = '⚠️ 偏弱,低于标准'
        else:
            audio_judge = '❌ 极弱/静音,不合格'
    else:
        audio_judge = '无有效视频,无法评估'

    if structured and 'audio_evaluation' in structured:
        audio_judge = structured['audio_evaluation'][:60]

    add_tbl(doc, ['指标', '数据', '评价'], [
        ('整体RMS', f'{ad_rms:.1f} dBFS', audio_judge),
        ('峰值电平', f'{ad_peak:.1f} dBFS', '⚠️ 略高' if ad_peak > -6 else '✅ 正常'),
        ('静音比例', f'{ad_silent:.1f}%', '✅ 低' if ad_silent < 10 else '⚠️ 高'),
        ('VO旁白', '(见AI报告)', '-'),
    ])

    # ========== 五、SOP 逐项检查 ==========
    doc.add_paragraph('')
    h = doc.add_paragraph()
    rh = h.add_run('五、SOP 逐项检查结果')
    rh.bold = True; rh.font.size = Pt(12)

    # 阶段划分：视频内容规范(1-14) + 发布规范(15-16)
    phase1_items = [
        (1, 'USP钩子', '前8秒快速抓住用户注意力'),
        (2, '封面标题', '3-5秒内搭配标题提示音展示'),
        (3, '房间参观', '核心内容模块'),
        (4, '重复展示', '严禁大篇幅重复常规基础设施'),
        (5, 'POI贴纸', '精准指向对应兴趣点'),
        (6, '价格菜单', 'SS Price Menu清晰展示'),
        (7, '促销信息', '折扣/促销活动优惠'),
        (8, '封面规范', '规范吸睛的封面和标题'),
        (9, '字幕同步', '字幕与VO旁白同步'),
        (10, '真人入镜', '创作者本人入镜露脸'),
        (11, '音画同步', '旁白与画面完全同步'),
        (12, '画面稳定', '无抖动模糊'),
        (13, '光线明亮', '明亮清晰'),
        (14, '字幕位置', '不得过低被遮挡'),
        (15, 'BGM音量', 'BGM音量不得盖过VO旁白'),
    ]
    phase2_items = [
        (16, '发布标签', '#RealwayX话题标签 + @RealwayMediaIndonesia'),
    ]
    phase1_names = {name for _, name, _ in phase1_items}
    phase2_names = {name for _, name, _ in phase2_items}

    import re as _re_sop
    _raw_checks = structured.get('sop_checks') if structured else None
    _sop_ids = {
        "USP钩子": 1, "封面标题": 2, "房间参观": 3, "重复展示": 4,
        "POI贴纸": 5, "价格菜单": 6, "促销信息": 7, "封面规范": 8,
        "字幕同步": 9, "真人入镜": 10, "音画同步": 11, "画面稳定": 12,
        "光线明亮": 13, "字幕位置": 14, "BGM音量": 15, "发布标签": 16,
    }
    if isinstance(_raw_checks, dict):
        _std = []
        for _k, _v in _raw_checks.items():
            _cid = _sop_ids.get(_k, 99)
            _full = str(_v).strip()
            _m = _re_sop.search(r'[✅❌⚠️-]', _full)
            _mark = _m.group(0) if _m else '-'
            _mr = _re_sop.search(r'[((]([^))]+)[))]', _full)
            _reason = _mr.group(1).strip() if _mr else '-'
            _std.append({"id": _cid, "name": _k, "result": _mark, "remark": _reason})
        _raw_checks = _std
    sop_checks_list = _raw_checks if isinstance(_raw_checks, list) else []

    def build_rows(_checks, _items):
        _rows = []
        for _cid, _name, _std_remark in _items:
            _found = None
            for _c in _checks:
                if _c.get('id') == _cid or _c.get('name') == _name:
                    _found = _c
                    break
            if _found:
                _r = _found.get('result', '-')
                _rm = _found.get('remark', '-')
            else:
                _r = '-'
                _rm = _std_remark
            _verdict_map = {'✅': '通过', '❌': '缺失', '⚠️': '待改进', '-': '未检测'}
            _vd = _verdict_map.get(_r.strip(), '待判定')
            _r_full = _r + _vd if _r in _verdict_map else _r
            _rm_clean = _re_sop.sub(r'^[原因结论:::\s]+', '', _rm) if (_rm and _rm != '-') else _rm
            _rows.append((_name, _r_full, _rm_clean))
        return _rows

    phase1_rows = build_rows(sop_checks_list, phase1_items)
    phase2_rows = build_rows(sop_checks_list, phase2_items)

    # 阶段一：视频内容规范
    doc.add_paragraph('')
    h1 = doc.add_paragraph()
    rh1 = h1.add_run('  阶段一：视频内容规范（制作阶段）')
    rh1.bold = True; rh1.font.size = Pt(11)
    add_tbl(doc, ['检查项', '结果', '备注'], phase1_rows)

    # 阶段二：发布规范
    doc.add_paragraph('')
    h2 = doc.add_paragraph()
    rh2 = h2.add_run('  阶段二：发布规范（发布阶段）')
    rh2.bold = True; rh2.font.size = Pt(11)
    add_tbl(doc, ['检查项', '结果', '备注'], phase2_rows)

    # ========== 六、改进建议 ==========
    improvements = []
    if structured and 'improvements' in structured:
        improvements = structured['improvements']
    elif result and result.get('report'):
        lines = result['report'].split('\n')
        for line in lines:
            if any(kw in line.lower() for kw in ['改进建议', '建议', 'improve', 'suggest']):
                clean = re.sub(r'[*#\[\]()]', '', line).strip()
                if len(clean) > 5:
                    improvements.append(clean)
                if len(improvements) >= 5:
                    break

    if improvements:
        doc.add_paragraph('')
        h = doc.add_paragraph()
        rh = h.add_run('六、改进建议')
        rh.bold = True; rh.font.size = Pt(12)
        for s in improvements[:5]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(s).font.size = Pt(11)

    # ========== 七、AI 审核报告原文 ==========
    # AI 审核报告原文 - 已移除

    # ========== 七、审核效率 ==========
    doc.add_paragraph('')
    h = doc.add_paragraph()
    rh = h.add_run('七、审核效率')
    rh.bold = True; rh.font.size = Pt(12)
    _aud = result.get('audit_result', {}) if isinstance(result, dict) else {}
    if _aud:
        add_tbl(doc, ['指标', '数值'], [
            ('AI 模型', f"{_aud.get('model', model or MODEL)}"),
            ('总时长', f"{phase_timings['total']:.1f} 秒" if phase_timings else "-"),
            ('下载时长', f"{phase_timings['download']:.1f} 秒" if phase_timings else "-"),
            ('提取帖子元数据', f"{phase_timings.get('post_meta', 0):.1f} 秒" if phase_timings else "-"),
            ('审核时长', f"{_aud.get('elapsed_sec', 0):.1f} 秒"),
            ('报告时长', f"{phase_timings['report']:.1f} 秒" if phase_timings else "-"),
            ('Total tokens', f"{_aud.get('total_tokens', 0):,}"),
            ('总成本', f"{_aud.get('cost_yuan', 0):.6f} 元(约 {_aud.get('cost_yuan', 0)*100:.2f} 分钱)"),
        ])

    doc.add_paragraph('')
    footer = doc.add_paragraph()
    rf = footer.add_run(f'审核团队:Realway Media Audit Team    生成时间:{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    rf.font.size = Pt(10); rf.font.color.rgb = RGBColor(128,128,128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

# ============ 主流程 ============
def main():
    parser = argparse.ArgumentParser(description='TikTok 视频审核脚本')
    parser.add_argument('url', help='TikTok 视频链接')
    parser.add_argument('--category', default='hotel',
                        choices=['hotel', 'restaurant', 'product', 'travel'],
                        help='SOP 类别:hotel/restaurant/product/travel,默认 hotel')
    parser.add_argument('--output-json', action='store_true', help='输出 JSON 报告')
    parser.add_argument('--output-docx', action='store_true', help='输出 Word 报告')
    parser.add_argument('--output-dir', default=OUTPUT_DIR, help='输出目录')
    parser.add_argument('--skip-download', action='store_true', help='跳过下载')
    parser.add_argument('--model', default=MODEL,
                        help=f'AI 模型名称，默认 {MODEL}')
    args = parser.parse_args()

    # 运行时使用用户指定的模型
    AUDIT_MODEL = args.model

    url = args.url
    print(f"\n{'='*50}")
    print(f"TikTok 视频审核")
    print(f"{'='*50}")
    print(f"链接: {url}")
    print(f"类别: {CATEGORY_NAMES.get(args.category, args.category)} ({args.category})")

    print(f"\n[1/5] 解析链接...")
    t0 = time.time()
    t_resolve = None
    resolved_url = resolve_tiktok_url(url)
    video_id = extract_video_id(resolved_url)
    print(f"  视频ID: {video_id}")

    print(f"\n[2/5] 下载视频...")
    t_download_start = time.time()
    video_path = None
    if args.skip_download:
        for ext in ['mp4', 'mp3', 'mkv']:
            path = os.path.join(args.output_dir, f"{video_id}.{ext}")
            if os.path.exists(path):
                video_path = path
                print(f"  使用本地文件: {video_path}")
                break
        if not video_path:
            print(f"  未找到本地文件,将重新下载")
            video_path = download_video(resolved_url, args.output_dir)
    else:
        video_path = download_video(resolved_url, args.output_dir)
    t_download = time.time() - t_download_start

    if not video_path or not os.path.exists(video_path):
        print("  下载失败")
        video_path = None

    print(f"\n[3/5] 抽帧 + 音频分析...")
    t_analysis_start = time.time()
    frames = []
    audio_data = {'has_video': False, 'has_audio': False, 'rms': 0, 'peak': 0,
                  'silent_ratio': 0, 'video_duration': 0, 'per_second': []}

    if video_path:
        frames, audio_data = extract_frames_and_audio(video_path)
        t_analysis = time.time() - t_analysis_start
        has_v = audio_data.get('has_video')
        dur = audio_data.get('video_duration', 0)
        rms = audio_data.get('rms', 0)
        silent = audio_data.get('silent_ratio', 0) * 100
        print(f"  视频: {dur:.1f}s, 有视频: {has_v}")
        print(f"  音频: RMS={rms:.1f}dBFS, 静音比={silent:.1f}%")
        print(f"  抽取帧: {len([f for f in frames if os.path.exists(f)])}帧")
        ps = audio_data.get('per_second', [])
        for i, d in enumerate(ps[:10]):
            rms_s = d.get('rms', -999)
            pct_s = d.get('active_pct', 0)
            status = '✅' if pct_s > 5 and rms_s > -45 else '❌'
            print(f"  第{i}s: RMS={rms_s:.1f}dBFS, 有效={pct_s:.0f}% {status}")
    else:
        print("  无视频文件,跳过抽帧")
        t_analysis = 0

    print(f"\n[4/5] AI 审核...")
    result = None
    # 新增[3.5/5]:提取帖子描述中的话题标签和@账号，用于第16项发布标签审核
    print(f"[3.5/5] 提取帖子元数据...")
    post_meta = fetch_post_metadata(resolved_url)
    if post_meta['hashtags'] or post_meta['mentions']:
        print(f"  话题标签: {' '.join(['#'+h for h in post_meta['hashtags']])}")
        print(f"  @账号: {' '.join(['@'+m for m in post_meta['mentions']])}")
    else:
        print(f"  未检测到话题标签或@账号")
    t_post_meta = time.time() - t_download_start - t_analysis

    if audio_data.get('has_video'):
        result = call_llm(frames, audio_data, category=args.category, post_metadata=post_meta, model=AUDIT_MODEL)
        if result:
            struct_ok = '✅' if result.get('structured') else '❌'
            print(f"  审核完成: {result['elapsed']:.1f}s, {result['total_tokens']} tokens, 结构化:{struct_ok}")
        else:
            print("  审核失败")
    else:
        print("  无有效视频,跳过 AI 审核")

    # --- 提取酒店名用于文件名（优先从 hashtag 的 #RealwayX 提取，其次从 LLM 结果提取）---
    hotel_name_for_file = None
    import re as _re_g
    exclude = ['Rekomendasi', 'Hotel_Murah', 'Murah', 'Hotel_Terbaik',
               'Tempat_Menginap', 'Penginapan', 'Villa_Murah', 'Hotel_Di',
               'Rekomendasi_Hotel', 'Di_Banjarbaru', 'Di_Bandung']

    def _extract_name(_text):
        import re as _re2
        # 1. 从 #RealwayXxxx 或 realwayxXx（hashtag 中的酒店名）
        _m = _re_g.search(r'#?RealwayX([A-Za-z0-9]+)', _text, _re_g.IGNORECASE)
        if _m:
            raw = _m.group(1)
            # 保留分隔符的分割：hotel/guest/resort/villa/inn/permai + xagoda/xhotel 等
            _split_pat = r'(hotel|guest|resort|villa|inn|permai|x(?=agoda|traveloka|booking|hotel|guest))'
            _parts = _re2.split(_split_pat, raw, flags=_re2.IGNORECASE)
            _parts = [p for p in _parts if p]
            # 合并 x 和后面的词（xagoda → x+agoda 合并）
            _merged = []
            i = 0
            while i < len(_parts):
                if _parts[i].lower() == 'x' and i + 1 < len(_parts):
                    _merged.append('x' + _parts[i + 1])
                    i += 2
                else:
                    _merged.append(_parts[i])
                    i += 1
            # 格式化每个部分：去前缀 → CamelCase分割 → Title_Case
            _formatted = []
            for p in _merged:
                p_lower = p.lower()
                for _pfx in ('hotel', 'guest', 'resort', 'villa', 'inn', 'permai', 'x'):
                    if p_lower.startswith(_pfx) and len(p) > len(_pfx):
                        p = p[len(_pfx):]
                        p_lower = p.lower()
                        break
                # CamelCase 分割（如 GammaraMakassar → Gammara_Makassar）
                p_fmt = _re2.sub(r'(?<!^)(?=[A-Z])', '_', p)
                p_title = '_'.join(w.capitalize() for w in p_fmt.split('_') if w)
                if p_title:
                    _formatted.append(p_title)
            if _formatted:
                return '_'.join(_formatted), 'tag'
        # 2. 从 LLM 结果中的大写+小写混合名称（Hotel Gammara Makassar）
        for _pat in [
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s+(?:Hotel|Guest.?House|Resort|Villa|Inn|Permai)',
            r'(?:Hotel|Guest.?House|Resort|Villa)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})',
        ]:
            _m = _re_g.search(_pat, _text)
            if _m:
                _full = _m.group(0).strip()
                _skip = False
                for _ex in exclude:
                    if _ex in _full:
                        _skip = True; break
                if not _skip and len(_full) > 4:
                    return _full.replace(' ', '_'), 'pattern'
        return None, None

    # 优先从帖子元数据（hashtag）提取酒店名
    if post_meta and post_meta.get('hashtags'):
        tag_text = ' '.join(['#' + h for h in post_meta['hashtags']])
        name, source = _extract_name(tag_text)
        if name and len(name) > 2:
            hotel_name_for_file = name
            print(f"  检测到酒店名: {hotel_name_for_file} (来源:{source})")

    # 其次从 LLM 审核结果提取
    if not hotel_name_for_file and result:
        all_text_parts = []
        if result.get('structured') and 'frames' in result['structured']:
            for f in result['structured']['frames']:
                all_text_parts.append(f.get('content', ''))
                all_text_parts.append(f.get('subtitle', ''))
        all_text_parts.append(result.get('report', ''))
        full_text = ' | '.join(all_text_parts)
        name, source = _extract_name(full_text)
        if name and len(name) > 2:
            hotel_name_for_file = name
            print(f"  检测到酒店名: {hotel_name_for_file} (来源:{source})")
    # --- 提取酒店名结束 ---


    print(f"\n[5/5] 生成报告...")
    t_report_start = time.time()

    # 格式化酒店名:在每个大写字母前加下划线(处理GammaraMakassar→Gammara_Makassar)
    def _fmt_name(_n):
        # 已有下划线或空格 → 跳过格式化(避免 Hotel_Gammara_Makassar → Hotel__Gammara__Makassar)
        if '_' in _n or ' ' in _n:
            return _n
        import re
        return re.sub(r'(?<!^)(?=[A-Z])', '_', _n)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 构造文件名:Hotel_{格式化名称}_Audit_Report_v2---{时间戳}
    if hotel_name_for_file:
        _fmt = _fmt_name(hotel_name_for_file)
        # 如果名称已含 Hotel_ 前缀则不再添加
        if _fmt.upper().startswith('HOTEL_'):
            base_name = f"{_fmt}_Audit_Report_v2---{timestamp}"
        else:
            base_name = f"Hotel_{_fmt}_Audit_Report_v2---{timestamp}"
    else:
        base_name = f"audit_{video_id}_{timestamp}"

    # 根据参数控制输出:默认 docx,--output-json 额外输出 json
    # 计算各阶段时长
    t_total = time.time() - t0
    t_report = time.time() - t_report_start
    audit_elapsed = result['elapsed'] if result else 0
    phase_timings = {
        'total': t_total,
        'download': t_download,
        'analysis': t_analysis,
        'post_meta': round(t_post_meta, 1),
        'audit': audit_elapsed,
        'report': t_report
    }

    # 始终计算 report_data（用于控制台输出和 docx）
    report_data = build_json_report(url, video_path, frames, audio_data, result, video_id, phase_timings, model=AUDIT_MODEL)
    if args.output_json:
        json_path = os.path.join(args.output_dir, f"{base_name}.json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, ensure_ascii=False, indent=2)
        print(f"  JSON 报告: {json_path}")

    # 默认输出 docx(无参数时或有 --output-docx 时)
    docx_path = os.path.join(args.output_dir, f"{base_name}.docx")
    doc = build_docx(url, video_path, frames, audio_data, report_data, video_id, phase_timings, model=AUDIT_MODEL)
    doc.save(docx_path)
    print(f"  Word 报告: {docx_path}")
    print(f"  实际文件名: {os.path.basename(docx_path)}")

    score = parse_score(result)
    _audit = report_data.get('audit_result', {}) if report_data else {}
    if not audio_data.get('has_video'):
        final_score = 0; final_conclusion = "不通过(无有效视频流)"
    elif _audit.get('phase1_score') is not None:
        final_score = _audit.get('score', 0)
        final_conclusion = _audit.get('conclusion', '待修改')
        _p1 = f"{_audit.get('phase1_score')}/10({_audit.get('phase1_conclusion')})"
        _p2 = f"{_audit.get('phase2_score')}/10({_audit.get('phase2_conclusion')})"
        print(f"\n{'='*50}")
        print(f"审核结果: {final_score}/10 - {final_conclusion}")
        print(f"阶段一(视频内容): {_p1}  [{_audit.get('phase1_passed')}/{_audit.get('phase1_total')}项通过]")
        print(f"阶段二(发布规范): {_p2}  [{_audit.get('phase2_passed')}/{_audit.get('phase2_total')}项通过]")
        print(f"{'='*50}\n")
    else:
        final_score = score or 0
        final_conclusion = "待修改" if (score and score >= 3) else ("不通过" if score else "不通过")
        print(f"\n{'='*50}")
        print(f"审核结果: {final_score}/10 - {final_conclusion}")
    if result:
        print(f"模型: {AUDIT_MODEL}")
        print(f"总耗时: {t_total:.1f}s | 下载: {t_download:.1f}s | 分析: {t_analysis:.1f}s | 审核: {result['elapsed']:.1f}s | 报告: {t_report:.1f}s")
        print(f"Tokens: {result['total_tokens']:,} | 成本: {result['total_cost']*100:.2f}分钱")
    print(f"{'='*50}\n")

    for fp in frames:
        if os.path.exists(fp):
            os.unlink(fp)

    return 0

if __name__ == '__main__':
    sys.exit(main())

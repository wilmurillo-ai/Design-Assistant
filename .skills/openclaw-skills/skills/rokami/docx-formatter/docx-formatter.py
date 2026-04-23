#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文格式规范生成器
严格按照中国公文格式规范生成 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import argparse
import re

# 中文引号
LEFT_QUOTE = '\u201c'   # "
RIGHT_QUOTE = '\u201d'  # "

def convert_quotes(text):
    """将英文引号转换为中文引号（成对替换）"""
    result = []
    in_quote = False
    
    for char in text:
        if char == '"':
            if in_quote:
                result.append(RIGHT_QUOTE)  # 右引号 "
                in_quote = False
            else:
                result.append(LEFT_QUOTE)   # 左引号 "
                in_quote = True
        else:
            result.append(char)
    
    return ''.join(result)

class DocxFormatter:
    """公文格式化器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_default_style()
        self._setup_page()
    
    def _setup_default_style(self):
        """设置默认样式"""
        style = self.doc.styles['Normal']
        style.font.name = '仿宋_GB2312'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        style.font.size = Pt(16)  # 3号字
    
    def _setup_page(self):
        """设置页面"""
        section = self.doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        self._add_page_number(section)
    
    def _add_page_number(self, section):
        """添加页码"""
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(16)
    
    def _set_paragraph_format(self, p, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """设置段落格式"""
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if first_line_indent:
            p.paragraph_format.first_line_indent = Pt(32)
        p.paragraph_format.widow_control = False
        p.alignment = alignment
    
    def add_title(self, text):
        """添加大标题（方正小标宋 2号）"""
        text = convert_quotes(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '方正小标宋简体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        run.font.size = Pt(22)  # 2号字
        self._set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return p
    
    def add_author(self, text):
        """添加署名（楷体 3号），支持多行"""
        text = convert_quotes(text)
        # 处理 \n 换行
        lines = text.split('\\n')
        for line in lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '楷体_GB2312'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
            run.font.size = Pt(16)
            self._set_paragraph_format(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return p
    
    def add_heading1(self, text):
        """添加一级标题（黑体 3号）"""
        text = convert_quotes(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
        self._set_paragraph_format(p, first_line_indent=True)
        return p
    
    def add_heading2(self, text):
        """添加二级标题（楷体 3号）"""
        text = convert_quotes(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '楷体_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
        run.font.size = Pt(16)
        self._set_paragraph_format(p, first_line_indent=True)
        return p
    
    def add_paragraph(self, text, bold_prefix=None):
        """添加正文段落（仿宋 3号）"""
        text = convert_quotes(text)
        if bold_prefix:
            bold_prefix = convert_quotes(bold_prefix)
        p = self.doc.add_paragraph()
        if bold_prefix:
            run1 = p.add_run(bold_prefix)
            run1.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        self._set_paragraph_format(p)
        return p
    
    def add_empty_line(self):
        """添加空行"""
        return self.doc.add_paragraph()
    
    def save(self, filepath):
        """保存文档"""
        self.doc.save(filepath)
        print(f"✅ 文档已保存：{filepath}")

def parse_markdown(filepath):
    """解析 Markdown 文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    result = {'title': None, 'author': None, 'content': []}
    
    for i, line in enumerate(lines):
        line = line.rstrip()
        
        # 跳过所有空行
        if not line:
            continue
        
        # 主标题（第一个 # 标题）
        if line.startswith('# ') and result['title'] is None:
            result['title'] = line[2:].strip()
            continue
        
        # 一级标题（## 开头）
        if line.startswith('## '):
            result['content'].append({
                'type': 'heading1',
                'text': line[3:].strip()
            })
            continue
        
        # 二级标题（### 开头）
        if line.startswith('### '):
            result['content'].append({
                'type': 'heading2',
                'text': line[4:].strip()
            })
            continue
        
        # 跳过 Markdown 分隔线
        if line.strip() in ['---', '***', '___']:
            continue
        
        # 正文段落
        # 检查是否有加粗前缀
        # 模式1: 1. **标题**。内容
        match = re.match(r'^(\d+\.\s+)\*\*([^*]+)\*\*[。\.]\s*(.+)$', line)
        if match:
            prefix = match.group(1)
            bold_text = match.group(2)
            content = match.group(3)
            result['content'].append({
                'type': 'paragraph',
                'text': content,
                'bold_prefix': prefix + bold_text + '。'
            })
            continue
        
        # 模式2: **前缀** 内容
        if line.startswith('**') and '**' in line[2:]:
            end_pos = line.index('**', 2)
            bold_prefix = line[2:end_pos]
            text = line[end_pos+2:].strip()
            result['content'].append({
                'type': 'paragraph',
                'text': text,
                'bold_prefix': bold_prefix
            })
            continue
        
        # 列表项转换（- 开头）
        if line.startswith('- '):
            text = line[2:].strip()
            result['content'].append({
                'type': 'paragraph',
                'text': text
            })
            continue
        
        # 普通段落
        result['content'].append({
            'type': 'paragraph',
            'text': line
        })
    
    return result

def main():
    parser = argparse.ArgumentParser(description='公文格式规范生成器')
    parser.add_argument('--title', help='文档标题')
    parser.add_argument('--author', help='署名')
    parser.add_argument('--content', help='内容 JSON 文件')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--from-markdown', help='从 Markdown 文件转换')
    
    args = parser.parse_args()
    
    formatter = DocxFormatter()
    
    if args.from_markdown:
        # 从 Markdown 转换
        data = parse_markdown(args.from_markdown)
        
        if data['title']:
            formatter.add_title(data['title'])
        
        if args.author:
            formatter.add_author(args.author)
            formatter.add_empty_line()
        
        for item in data['content']:
            item_type = item.get('type')
            text = item.get('text', '')
            
            if item_type == 'heading1':
                formatter.add_heading1(text)
            elif item_type == 'heading2':
                formatter.add_heading2(text)
            elif item_type == 'paragraph':
                bold_prefix = item.get('bold_prefix')
                formatter.add_paragraph(text, bold_prefix)
            elif item_type == 'empty':
                formatter.add_empty_line()
        
        formatter.save(args.output)
        return
    
    # 基本用法
    if args.title:
        formatter.add_title(args.title)
    
    if args.author:
        formatter.add_author(args.author)
        formatter.add_empty_line()
    
    if args.content:
        with open(args.content, 'r', encoding='utf-8') as f:
            content = json.load(f)
            
        for item in content:
            item_type = item.get('type')
            text = item.get('text', '')
            
            if item_type == 'heading1':
                formatter.add_heading1(text)
            elif item_type == 'heading2':
                formatter.add_heading2(text)
            elif item_type == 'paragraph':
                bold_prefix = item.get('bold_prefix')
                formatter.add_paragraph(text, bold_prefix)
            elif item_type == 'empty':
                formatter.add_empty_line()
    
    formatter.save(args.output)

if __name__ == '__main__':
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
-*--*--*--*--*--*--*--*--*--*--*-

createBy: wanghao

createTime: 2026-03-27 15:28:00

-*--*--*--*--*--*--*--*--*--*--*-

Word文档处理工具库
提供Word文档的基本操作功能

注意：此工具库需要安装python-docx库
安装命令：pip install python-docx
"""

import os
from datetime import datetime
from typing import Dict, List, Optional, Any


class WordToolkit:
    """Word文档处理工具库"""
    
    def __init__(self, debug: bool = False):
        """
        初始化Word工具库
        
        Args:
            debug: 是否启用调试模式
        """
        self.debug = debug
        self._log("Word工具库初始化完成")
    
    def _log(self, message: str):
        """内部日志方法"""
        if self.debug:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            print(f"[WordToolkit {timestamp}] {message}")
    
    def check_docx_installed(self) -> bool:
        """
        检查python-docx是否已安装
        
        Returns:
            bool: 是否已安装
        """
        try:
            import docx
            self._log("python-docx已安装")
            return True
        except ImportError:
            self._log("python-docx未安装，请运行: pip install python-docx")
            return False
    
    def create_document(self, filepath: str, content: Dict[str, Any] = None) -> bool:
        """
        创建Word文档
        
        Args:
            filepath: 文件路径
            content: 文档内容
            
        Returns:
            bool: 是否创建成功
        """
        if not self.check_docx_installed():
            return False
        
        try:
            from docx import Document
            
            # 创建新文档
            doc = Document()
            
            # 添加标题
            if content and 'title' in content:
                doc.add_heading(content['title'], level=0)
            
            # 添加段落
            if content and 'paragraphs' in content:
                for paragraph in content['paragraphs']:
                    doc.add_paragraph(paragraph)
            
            # 添加表格
            if content and 'tables' in content:
                for table_data in content['tables']:
                    table = doc.add_table(rows=1, cols=len(table_data['headers']))
                    
                    # 添加表头
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(table_data['headers']):
                        header_cells[i].text = header
                    
                    # 添加数据行
                    for row_data in table_data['rows']:
                        row_cells = table.add_row().cells
                        for i, cell_data in enumerate(row_data):
                            row_cells[i].text = str(cell_data)
            
            # 保存文档
            doc.save(filepath)
            self._log(f"Word文档已创建: {filepath}")
            return True
            
        except Exception as e:
            self._log(f"创建Word文档失败: {e}")
            return False
    
    def read_document(self, filepath: str) -> Optional[Dict[str, Any]]:
        """
        读取Word文档
        
        Args:
            filepath: 文件路径
            
        Returns:
            dict: 文档内容，包含文本、段落、表格等信息
        """
        if not self.check_docx_installed():
            return None
        
        try:
            from docx import Document
            
            if not os.path.exists(filepath):
                self._log(f"文件不存在: {filepath}")
                return None
            
            doc = Document(filepath)
            content = {
                'paragraphs': [],
                'tables': []
            }
            
            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content['paragraphs'].append(paragraph.text)
            
            # 提取表格
            for table in doc.tables:
                table_data = {
                    'headers': [],
                    'rows': []
                }
                
                # 提取表头（第一行）
                if len(table.rows) > 0:
                    header_row = table.rows[0]
                    table_data['headers'] = [cell.text for cell in header_row.cells]
                
                # 提取数据行（从第二行开始）
                for i in range(1, len(table.rows)):
                    data_row = table.rows[i]
                    table_data['rows'].append([cell.text for cell in data_row.cells])
                
                if table_data['headers'] or table_data['rows']:
                    content['tables'].append(table_data)
            
            self._log(f"Word文档已读取: {filepath}")
            return content
            
        except Exception as e:
            self._log(f"读取Word文档失败: {e}")
            return None
    
    def extract_text(self, filepath: str) -> Optional[str]:
        """
        提取Word文档中的纯文本
        
        Args:
            filepath: 文件路径
            
        Returns:
            str: 提取的文本内容
        """
        content = self.read_document(filepath)
        if not content:
            return None
        
        # 合并所有段落文本
        text_parts = []
        
        # 添加段落文本
        if 'paragraphs' in content:
            text_parts.extend(content['paragraphs'])
        
        # 添加表格文本
        if 'tables' in content:
            for table in content['tables']:
                if table['headers']:
                    text_parts.append(" | ".join(table['headers']))
                for row in table['rows']:
                    text_parts.append(" | ".join(row))
        
        return "\n".join(text_parts)
    
    def demo_openclawkit_word(self):
        """演示Word工具库功能"""
        print("🎯 Word工具库演示")
        print("=" * 50)
        
        if not self.check_docx_installed():
            print("❌ 请先安装python-docx: pip install python-docx")
            return
        
        # 创建示例文档
        sample_content = {
            'title': '示例Word文档',
            'paragraphs': [
                '这是第一个段落。',
                '这是第二个段落，包含一些示例文本。',
                'Word工具库可以方便地创建和读取Word文档。'
            ],
            'tables': [{
                'headers': ['姓名', '年龄', '部门'],
                'rows': [
                    ['张三', '25', '技术部'],
                    ['李四', '30', '市场部'],
                    ['王五', '35', '销售部']
                ]
            }]
        }
        
        output_file = '示例文档.docx'
        
        print(f"📝 创建示例Word文档: {output_file}")
        if self.create_document(output_file, sample_content):
            print(f"✅ 文档创建成功: {output_file}")
            
            # 读取文档
            print(f"\n📖 读取Word文档: {output_file}")
            content = self.read_document(output_file)
            
            if content:
                print(f"📊 文档内容:")
                print(f"  段落数: {len(content.get('paragraphs', []))}")
                print(f"  表格数: {len(content.get('tables', []))}")
                
                # 提取文本
                print(f"\n📄 提取的文本内容:")
                text = self.extract_text(output_file)
                if text:
                    print(text[:500] + "..." if len(text) > 500 else text)
        
        print("\n🎉 演示完成！")


def demo_openclawkit_word():
    """演示函数"""
    toolkit = WordToolkit(debug=True)
    toolkit.demo_openclawkit_word()


if __name__ == '__main__':
    demo_openclawkit_word()
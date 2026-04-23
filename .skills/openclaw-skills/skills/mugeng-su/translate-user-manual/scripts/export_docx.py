import sys
import docx
from docx.shared import Pt
import re

def markdown_to_docx(input_md_file, output_docx_file):
    doc = docx.Document()
    
    with open(input_md_file, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Find pages separated by "## Page N"
    pages = re.split(r'^##\s+Page\s+\d+', content, flags=re.MULTILINE)
    
    for i, page_content in enumerate(pages):
        if not page_content.strip():
            continue
            
        doc.add_heading(f'Page {i}', level=2)
        
        # Extract table rows
        rows = re.findall(r'^\|(.+?)\|(.+?)\|$', page_content, flags=re.MULTILINE)
        if rows:
            # Skip the header and separator rows if they exist
            if 'Original Text' in rows[0][0]:
                rows = rows[2:]
                
            table = doc.add_table(rows=len(rows), cols=2)
            table.style = 'Table Grid'
            
            for row_idx, (col1, col2) in enumerate(rows):
                cell1 = table.cell(row_idx, 0)
                cell2 = table.cell(row_idx, 1)
                
                cell1.text = col1.strip()
                cell2.text = col2.strip()
                
                # Basic styling
                for paragraph in cell1.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                for paragraph in cell2.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        
        if i < len(pages) - 1:
            doc.add_page_break()
            
    doc.save(output_docx_file)
    print(f"Successfully exported {output_docx_file}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python export_docx.py <input.md> <output.docx>")
        sys.exit(1)
        
    markdown_to_docx(sys.argv[1], sys.argv[2])

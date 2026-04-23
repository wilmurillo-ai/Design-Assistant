# -*- coding: utf-8 -*-
"""
终极智能填充系统 - 完整解析文本化Excel内容 + 语义检索
"""
import os
import json
import re
import sys
from pathlib import Path
from datetime import datetime
from collections import defaultdict

sys.stdout.reconfigure(encoding='utf-8')

class TextExcelParser:
    """文本化Excel解析器"""
    
    def __init__(self):
        self.data = {}
        self.year_data = defaultdict(dict)
    
    def parse(self, text_content):
        """解析文本化的Excel内容"""
        lines = text_content.split('\n')
        current_sheet = None
        headers = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测Sheet名称
            if line.startswith('=== Sheet:'):
                current_sheet = line.replace('=== Sheet:', '').strip()
                headers = []
                continue
            
            # 解析数据行（格式：A1[值] | B2[值]）
            if '[' in line and '|' in line:
                cells = line.split('|')
                row_values = []
                
                for cell in cells:
                    match = re.search(r'\[([^\]]+)\]', cell)
                    if match:
                        row_values.append(match.group(1).strip())
                    else:
                        row_values.append('')
                
                # 如果是表头行（包含'年'字且多为文本）
                if headers and row_values:
                    # 存储数据
                    for i, (header, value) in enumerate(zip(headers, row_values)):
                        if header and value:
                            self.data[header] = value
                            
                            # 检测年份
                            year_match = re.search(r'\d{4}', header)
                            if year_match:
                                year = year_match.group()
                                clean_header = re.sub(r'\d{4}|年', '', header).strip()
                                if clean_header:
                                    self.year_data[year][clean_header] = value
                elif not headers and row_values:
                    # 假设这是表头
                    headers = row_values
        
        return self

class SmartFillSystem:
    """智能填充系统"""
    
    def __init__(self, kb_path, template_dir, output_dir):
        self.kb_path = kb_path
        self.template_dir = template_dir
        self.output_dir = output_dir
        self.parser = TextExcelParser()
        self.values = {}
        self.fill_log = []
        
        os.makedirs(output_dir, exist_ok=True)
    
    def load_kb(self):
        """加载知识库"""
        print("="*60)
        print("终极智能填充系统")
        print("="*60)
        print("\n正在加载知识库...")
        
        with open(self.kb_path, 'r', encoding='utf-8') as f:
            kb_data = json.load(f)
        
        # 解析每个文件
        for file_key, file_data in kb_data.items():
            filename = file_data.get('filename', file_key)
            file_type = file_data.get('type', '')
            content = file_data.get('content', '')
            
            print(f"  解析: {filename}")
            
            if file_type == 'xlsx' and isinstance(content, str):
                # 解析文本化Excel
                self.parser.parse(content)
                
                # 从parser提取数据
                self.values.update(self.parser.data)
                for year, data in self.parser.year_data.items():
                    for key, value in data.items():
                        if key not in self.values:
                            self.values[key] = value
            
            elif file_type in ['doc', 'docx'] and isinstance(content, str):
                # 提取文档字段
                self._extract_from_text(content)
        
        print(f"知识库加载完成，共 {len(self.values)} 个字段\n")
        
        # 显示部分值
        print("知识库样例（前20个）:")
        for i, (key, value) in enumerate(list(self.values.items())[:20]):
            if value and str(value).strip():
                print(f"  {key}: {value}")
        print()
    
    def _extract_from_text(self, text):
        """从文本中提取字段"""
        patterns = {
            '法人代表': r'法定代表人?[：:\s]*([^\n\r]{2,10})',
            '法定代表人': r'法定代表人?[：:\s]*([^\n\r]{2,10})',
            '联系电话': r'(?:联系|业务)电话[：:\s]*([\d\s-]{7,20})',
            '电话': r'电话[：:\s]*([\d\s-]{7,20})',
            '传真': r'传真[：:\s]*([\d\s-]{7,20})',
            '地址': r'地址[：:\s]*([^\n\r]{10,100})',
            '注册资本': r'注册资本[：:\s]*([\d.]+)',
            '统一社会信用代码': r'统一社会信用代码[：:\s]*([A-Z0-9]{18})',
            '成立时间': r'成立(?:时间|日期)?[：:\s]*(\d{4}年\d{1,2}月)',
        }
        
        for field, pattern in patterns.items():
            if field not in self.values:
                match = re.search(pattern, text)
                if match:
                    self.values[field] = match.group(1).strip()
    
    def find_value(self, query, year=None):
        """查找值（混合检索）"""
        # 1. 精确匹配
        if query in self.values:
            return self.values[query]
        
        # 2. 包含匹配
        for key, value in self.values.items():
            if query in key or key in query:
                return value
        
        # 3. 关键词匹配
        keywords = query.split()
        for key, value in self.values.items():
            if all(kw in key for kw in keywords):
                return value
        
        return None
    
    def fill_docx(self, template_path):
        """填充docx"""
        from docx import Document
        
        print(f"处理: {os.path.basename(template_path)}")
        
        doc = Document(template_path)
        fill_count = 0
        
        # 填充表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    
                    if not text or len(text) < 2 or len(text) > 100:
                        continue
                    
                    # 查找匹配
                    value = self.find_value(text)
                    if value:
                        cell.text = str(value)
                        fill_count += 1
        
        # 替换XX基金
        for para in doc.paragraphs:
            for run in para.runs:
                if 'XX' in run.text:
                    run.text = run.text.replace('XX', '国寿安保基金')
                    fill_count += 1
        
        output_name = f"{os.path.splitext(os.path.basename(template_path))[0]}（已填写）.docx"
        output_path = os.path.join(self.output_dir, output_name)
        doc.save(output_path)
        
        print(f"  填充: {fill_count} 个字段")
        return output_path
    
    def fill_xlsx(self, template_path):
        """填充xlsx"""
        from openpyxl import load_workbook
        
        print(f"处理: {os.path.basename(template_path)}")
        
        wb = load_workbook(template_path)
        fill_count = 0
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    
                    text = str(cell.value).strip()
                    
                    if not text or len(text) < 2 or len(text) > 100:
                        continue
                    
                    # 查找匹配
                    value = self.find_value(text)
                    if value:
                        cell.value = value
                        fill_count += 1
        
        output_name = f"{os.path.splitext(os.path.basename(template_path))[0]}（已填写）.xlsx"
        output_path = os.path.join(self.output_dir, output_name)
        wb.save(output_path)
        
        print(f"  填充: {fill_count} 个字段")
        return output_path
    
    def process_all(self):
        """处理所有模板"""
        template_dir = Path(self.template_dir)
        template_files = list(template_dir.glob('*.docx')) + list(template_dir.glob('*.xlsx'))
        
        print(f"找到 {len(template_files)} 个模板文件\n")
        
        output_files = []
        for template_path in template_files:
            if template_path.name.startswith('~$'):
                continue
            
            try:
                suffix = template_path.suffix.lower()
                if suffix == '.docx':
                    output_path = self.fill_docx(str(template_path))
                elif suffix == '.xlsx':
                    output_path = self.fill_xlsx(str(template_path))
                else:
                    continue
                
                if output_path:
                    output_files.append(output_path)
            except Exception as e:
                print(f"  ❌ 错误: {str(e)}")
                import traceback
                traceback.print_exc()
        
        self._report(output_files)
        return output_files
    
    def _report(self, output_files):
        """报告"""
        print("\n" + "="*60)
        print(f"完成！成功处理 {len(output_files)} 个文件")
        print(f"输出目录: {self.output_dir}")
        print("="*60)

if __name__ == '__main__':
    kb_path = r'c:\需求\机构外部渠道需求\机构渠道外部机构表单填写需求\知识库.json'
    template_dir = r'c:\需求\机构外部渠道需求\机构渠道外部机构表单填写需求\模版'
    output_dir = r'c:\需求\机构外部渠道需求\机构渠道外部机构表单填写需求\优化填写模板'
    
    system = SmartFillSystem(kb_path, template_dir, output_dir)
    system.load_kb()
    system.process_all()

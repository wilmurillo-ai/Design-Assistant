#!/usr/bin/env python3
"""
基于目录结构识别的智能格式化工具 - 版本3
先分析文档整体结构，再应用正确的格式
"""

import docx
from docx.shared import RGBColor, Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def analyze_document_hierarchy(doc):
    """
    分析文档的层次结构，识别标题级别
    返回层次结构列表
    """
    hierarchy = []
    
    # 收集所有段落
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        paragraphs.append({
            'index': i,
            'text': text,
            'style': para.style.name,
            'para_object': para
        })
    
    # 识别文档标题（第一个非空段落，且不是数字序号）
    doc_title = None
    for p in paragraphs[:10]:
        if p['text'] and not re.match(r'^[一二三四五六七八九十]、', p['text']) and not re.match(r'^[0-9]+\.', p['text']):
            doc_title = p
            break
    
    if doc_title:
        hierarchy.append({
            'type': 'doc_title',
            'index': doc_title['index'],
            'text': doc_title['text']
        })
    
    # 分析章节层次
    for p in paragraphs:
        text = p['text']
        
        # 一级标题：中文序号
        if re.match(r'^[一二三四五六七八九十]、', text):
            match = re.match(r'^([一二三四五六七八九十])、(.*)', text)
            if match:
                chinese_num = match.group(1)
                content = match.group(2)
                
                # 判断是否为真正的一级标题
                is_real_title = len(content) <= 20 and not '。' in content
                
                hierarchy.append({
                    'type': 'level1',
                    'index': p['index'],
                    'num': chinese_num,
                    'content': content,
                    'is_real_title': is_real_title,
                    'text': text
                })
        
        # 二级标题：中文括号序号
        elif re.match(r'^（[一二三四五六七八九十]）', text):
            match = re.match(r'^（([一二三四五六七八九十])）(.*)', text)
            if match:
                chinese_num = match.group(1)
                content = match.group(2)
                
                hierarchy.append({
                    'type': 'level2',
                    'index': p['index'],
                    'num': chinese_num,
                    'content': content,
                    'text': text
                })
        
        # 三级标题：数字序号
        elif re.match(r'^[0-9]+\.', text):
            match = re.match(r'^([0-9]+)\.(.*)', text)
            if match:
                digit_num = match.group(1)
                content = match.group(2)
                
                hierarchy.append({
                    'type': 'level3',
                    'index': p['index'],
                    'num': digit_num,
                    'content': content,
                    'text': text
                })
    
    return hierarchy, paragraphs

def set_chinese_font(run, font_name):
    """设置中文字体"""
    run.font.name = font_name
    run.font.italic = False  # 明确禁用斜体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    # 设置中文字体
    font_element = OxmlElement('w:rFonts')
    font_element.set(qn('w:ascii'), font_name)
    font_element.set(qn('w:hAnsi'), font_name)
    font_element.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, font_element)

def set_black_color(run):
    """明确设置文字颜色为黑色"""
    run.font.color.rgb = RGBColor(0, 0, 0)

def smart_format_document(doc_path, output_path):
    """
    基于目录分析的智能格式化
    先识别文档结构，再应用正确的格式
    """
    doc = docx.Document(doc_path)
    
    print(f"智能格式化文档: {os.path.basename(doc_path)}")
    print("=" * 60)
    
    # 1. 分析文档层次结构
    hierarchy, paragraphs = analyze_document_hierarchy(doc)
    
    # 显示分析结果
    print("文档结构分析:")
    print("-" * 40)
    
    for item in hierarchy[:15]:  # 只显示前15个
        if item['type'] == 'doc_title':
            print(f"文档标题: '{item['text']}' (段落{item['index']+1})")
        elif item['type'] == 'level1':
            if item['is_real_title']:
                print(f"一级标题: {item['num']}、{item['content'][:30]}... (段落{item['index']+1})")
            else:
                print(f"⚠ 可能应为三级标题: {item['num']}、{item['content'][:30]}... (段落{item['index']+1})")
        elif item['type'] == 'level2':
            print(f"  二级标题: （{item['num']}）{item['content'][:30]}... (段落{item['index']+1})")
        elif item['type'] == 'level3':
            print(f"    三级标题: {item['num']}. {item['content'][:30]}... (段落{item['index']+1})")
    
    # 2. 设置页面
    sections = doc.sections
    for section in sections:
        # 页边距
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.header_distance = Cm(2.5)
        section.footer_distance = Cm(2.5)
        # 纸张大小 A4
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    
    # 3. 定义样式
    styles = doc.styles
    
    # 正文样式
    normal_style = styles['Normal']
    normal_style.font.name = '仿宋GB2312'
    normal_style.font.size = Pt(16)  # 三号约16磅
    normal_style.font.italic = False  # 禁用斜体
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_style.paragraph_format.line_spacing = Pt(28)  # 固定值28磅
    normal_style.paragraph_format.first_line_indent = Cm(0.85)  # 两个字符约0.85cm
    
    # 文档标题样式
    if 'Heading 1' not in styles:
        heading1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading1_style = styles['Heading 1']
    heading1_style.font.name = '方正小标宋简体'
    heading1_style.font.size = Pt(22)  # 二号约22磅
    heading1_style.font.italic = False  # 禁用斜体
    heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_style.paragraph_format.space_after = Pt(12)
    heading1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading1_style.paragraph_format.line_spacing = Pt(28)
    
    # 一级标题样式
    if 'Heading 2' not in styles:
        heading2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading2_style = styles['Heading 2']
    heading2_style.font.name = '黑体'
    heading2_style.font.size = Pt(16)  # 三号
    heading2_style.font.bold = False  # 规范：不加粗
    heading2_style.font.italic = False  # 禁用斜体
    heading2_style.paragraph_format.space_before = Pt(12)
    heading2_style.paragraph_format.space_after = Pt(6)
    heading2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading2_style.paragraph_format.line_spacing = Pt(28)
    
    # 二级标题样式
    if 'Heading 3' not in styles:
        heading3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading3_style = styles['Heading 3']
    heading3_style.font.name = '楷体GB2312'
    heading3_style.font.size = Pt(16)  # 三号
    heading3_style.font.bold = True   # 规范：加粗
    heading3_style.font.italic = False  # 禁用斜体
    heading3_style.paragraph_format.space_before = Pt(6)
    heading3_style.paragraph_format.space_after = Pt(3)
    heading3_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading3_style.paragraph_format.line_spacing = Pt(28)
    
    # 三级标题样式
    if 'Heading 4' not in styles:
        heading4_style = styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading4_style = styles['Heading 4']
    heading4_style.font.name = '仿宋GB2312'
    heading4_style.font.size = Pt(16)  # 三号
    heading4_style.font.bold = False
    heading4_style.font.italic = False  # 禁用斜体
    heading4_style.paragraph_format.space_before = Pt(3)
    heading4_style.paragraph_format.space_after = Pt(3)
    heading4_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading4_style.paragraph_format.line_spacing = Pt(28)
    
    # 4. 基于层次分析应用样式
    print("\n" + "=" * 60)
    print("应用智能格式:")
    print("-" * 40)
    
    # 创建索引映射
    hierarchy_dict = {}
    for item in hierarchy:
        hierarchy_dict[item['index']] = item
    
    # 首先应用文档标题样式
    doc_title_applied = False
    for p_info in paragraphs:
        if p_info['index'] in hierarchy_dict and hierarchy_dict[p_info['index']]['type'] == 'doc_title':
            para = p_info['para_object']
            para.style = heading1_style
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in para.runs:
                set_chinese_font(run, '方正小标宋简体')
                set_black_color(run)
                run.font.size = Pt(22)
            
            print(f"  文档标题: '{para.text[:30]}...' → 方正小标宋简体二号居中")
            doc_title_applied = True
            break
    
    # 处理所有段落
    for p_info in paragraphs:
        para = p_info['para_object']
        text = p_info['text']
        
        if p_info['index'] in hierarchy_dict:
            item = hierarchy_dict[p_info['index']]
            
            if item['type'] == 'level1':
                # 一级标题
                if item['is_real_title']:
                    # 真正的一级标题
                    para.style = heading2_style
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    for run in para.runs:
                        set_chinese_font(run, '黑体')
                        set_black_color(run)
                        run.font.size = Pt(16)
                        run.font.bold = False
                    
                    print(f"  一级标题: '{text[:30]}...' → 黑体三号左对齐")
                else:
                    # 可能应为三级标题：有句号或内容过长
                    # 检查是否有上级二级标题
                    has_parent_level2 = False
                    for prev_idx in range(p_info['index']-1, -1, -1):
                        if prev_idx in hierarchy_dict and hierarchy_dict[prev_idx]['type'] == 'level2':
                            has_parent_level2 = True
                            break
                    
                    if has_parent_level2:
                        # 改为三级标题
                        # 将中文序号改为数字序号
                        match = re.match(r'^([一二三四五六七八九十])、(.*)', text)
                        if match:
                            chinese_num = match.group(1)
                            content = match.group(2)
                            
                            # 中文数字转阿拉伯数字
                            chinese_to_digit = {
                                '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
                                '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'
                            }
                            digit_num = chinese_to_digit.get(chinese_num, chinese_num)
                            new_text = f"{digit_num}. {content}"
                            
                            # 更新文本
                            para.clear()
                            para.add_run(new_text)
                            
                            para.style = heading4_style
                            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            for run in para.runs:
                                set_chinese_font(run, '仿宋GB2312')
                                set_black_color(run)
                                run.font.size = Pt(16)
                                run.font.bold = False
                            
                            print(f"  ⚠ 修正: '{text[:30]}...' → '{new_text[:30]}...' (改为三级标题)")
                    else:
                        # 保持为一级标题但标记为有问题
                        para.style = heading2_style
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        for run in para.runs:
                            set_chinese_font(run, '黑体')
                            set_black_color(run)
                            run.font.size = Pt(16)
                            run.font.bold = False
                        
                        print(f"  ⚠ 警告: '{text[:30]}...' (一级标题但有句号)")
            
            elif item['type'] == 'level2':
                # 二级标题
                para.style = heading3_style
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for run in para.runs:
                    set_chinese_font(run, '楷体GB2312')
                    set_black_color(run)
                    run.font.size = Pt(16)
                    run.font.bold = True
                
                # print(f"  二级标题: '{text[:30]}...' → 楷体GB2312三号加粗")
            
            elif item['type'] == 'level3':
                # 三级标题
                para.style = heading4_style
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for run in para.runs:
                    set_chinese_font(run, '仿宋GB2312')
                    set_black_color(run)
                    run.font.size = Pt(16)
                    run.font.bold = False
                
                # print(f"  三级标题: '{text[:30]}...' → 仿宋GB2312三号左对齐")
        
        else:
            # 普通正文
            para.style = normal_style
            para.paragraph_format.first_line_indent = Cm(0.85)
            
            for run in para.runs:
                set_chinese_font(run, '仿宋GB2312')
                set_black_color(run)
                run.font.size = Pt(16)
    
    # 5. 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.style = normal_style
                    for run in para.runs:
                        set_chinese_font(run, '仿宋GB2312')
                        set_black_color(run)
                        run.font.size = Pt(16)
    
    # 6. 保存文档
    doc.save(output_path)
    print(f"\n已保存格式化文档: {output_path}")
    
    # 7. 生成格式说明
    print("\n" + "=" * 60)
    print("已应用的格式调整:")
    print("1. 页面设置:")
    print("   - 页边距: 上3.5cm, 下3.5cm, 左2.8cm, 右2.8cm")
    print("   - 页眉页脚距离: 2.5cm")
    print("   - 纸张: A4 (21.0×29.7cm)")
    print("2. 智能标题识别:")
    print("   - 基于文档结构分析应用正确的标题级别")
    print("   - 识别真正的一级标题（简短无句号）")
    print("   - 识别三级标题（在二级标题下，有句号）")
    print("3. 字体设置:")
    print("   - 文档标题: 方正小标宋简体, 二号(22磅), 居中")
    print("   - 一级标题: 黑体, 三号(16磅), 左对齐, 不加粗")
    print("   - 二级标题: 楷体GB2312, 三号(16磅), 左对齐, 加粗")
    print("   - 三级标题: 仿宋GB2312, 三号(16磅), 左对齐")
    print("   - 正文: 仿宋GB2312, 三号(16磅)")
    print("4. 段落格式:")
    print("   - 行距: 固定值28磅")
    print("   - 正文首行缩进: 2字符(约0.85cm)")
    print("5. 其他:")
    print("   - 所有文字颜色: 黑色RGB(0,0,0)")
    print("   - 符合数科公司文印格式要求")
    
    return output_path

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        output_file = input_file.replace('.docx', '_smart_v3.docx')
        if len(sys.argv) > 2:
            output_file = sys.argv[2]
        smart_format_document(input_file, output_file)
    else:
        # 测试用
        input_file = "/root/.openclaw/workspace/英文文件名文档/01_AI_Research_Scheduling_Platform_final.docx"
        output_file = "/root/.openclaw/workspace/英文文件名文档/01_AI_Research_Scheduling_Platform_smart_v3.docx"
        smart_format_document(input_file, output_file)
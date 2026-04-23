#!/usr/bin/env python3
"""
Word 报告生成脚本
读取 analyze.py 输出的 summary.json + 图表，生成完整 Word 报告
用法：python3 export_report.py <summary.json路径> <洞察文本> [--output <报告路径>]
"""

import sys
import os
import json
import argparse
from datetime import datetime

REQUIRED = ["python-docx"]
missing = []
try:
    import docx
except ImportError:
    missing.append("python-docx")

if missing:
    print(f"[ERROR] 缺少依赖，请先运行: pip3 install {' '.join(missing)}")
    sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    """设置表格单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_kv_table(doc, data: dict, title: str = None):
    """添加键值对表格"""
    if title:
        doc.add_paragraph(title, style="Heading 3")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "值"
    for cell in hdr:
        set_cell_bg(cell, "4472C4")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    for k, v in data.items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
    doc.add_paragraph()


def generate_report(summary_path: str, insight: str, output_path: str):
    with open(summary_path, "r", encoding="utf-8") as f:
        summary = json.load(f)

    doc = Document()

    # ── 标题页 ────────────────────────────────────────────────────────────────
    title = doc.add_heading("数据分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"数据文件：{os.path.basename(summary.get('file', '未知'))}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── 1. 数据概览 ───────────────────────────────────────────────────────────
    add_heading(doc, "一、数据概览", 1)
    eda = summary.get("eda", {})
    shape = eda.get("shape", [0, 0])
    overview = {
        "总行数": shape[0],
        "总列数": shape[1],
        "数值列数": len(eda.get("numeric_columns", [])),
        "分类列数": len(eda.get("categorical_columns", [])),
        "重复行数": eda.get("duplicates", 0),
    }
    add_kv_table(doc, overview, "基本信息")

    # 列信息表
    doc.add_paragraph("列信息", style="Heading 3")
    cols = eda.get("columns", [])
    dtypes = eda.get("dtypes", {})
    missing_pct = eda.get("missing_pct", {})
    if cols:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["列名", "数据类型", "缺失率(%)"]):
            hdr[i].text = h
            set_cell_bg(hdr[i], "4472C4")
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        for col in cols:
            row = table.add_row().cells
            row[0].text = str(col)
            row[1].text = str(dtypes.get(col, ""))
            row[2].text = str(missing_pct.get(col, 0))
    doc.add_paragraph()

    # ── 2. 数据清洗记录 ───────────────────────────────────────────────────────
    add_heading(doc, "二、数据清洗记录", 1)
    clean_log = summary.get("clean_log", [])
    if clean_log:
        for item in clean_log:
            doc.add_paragraph(f"• {item}")
    else:
        doc.add_paragraph("数据质量良好，无需清洗。")
    doc.add_paragraph()

    # ── 3. 统计分析 ───────────────────────────────────────────────────────────
    add_heading(doc, "三、统计分析", 1)
    numeric_stats = eda.get("numeric_stats", {})
    if numeric_stats:
        doc.add_paragraph("数值列描述性统计", style="Heading 3")
        stat_keys = ["count", "mean", "std", "min", "25%", "50%", "75%", "max"]
        num_cols = list(numeric_stats.keys())
        table = doc.add_table(rows=1, cols=len(num_cols) + 1)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "统计量"
        set_cell_bg(hdr[0], "4472C4")
        for para in hdr[0].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
        for i, col in enumerate(num_cols):
            hdr[i + 1].text = str(col)
            set_cell_bg(hdr[i + 1], "4472C4")
            for para in hdr[i + 1].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        for stat in stat_keys:
            row = table.add_row().cells
            row[0].text = stat
            for i, col in enumerate(num_cols):
                val = numeric_stats[col].get(stat, "")
                row[i + 1].text = f"{val:.4g}" if isinstance(val, float) else str(val)
        doc.add_paragraph()

    cat_stats = eda.get("categorical_stats", {})
    if cat_stats:
        doc.add_paragraph("分类列统计", style="Heading 3")
        for col, info in cat_stats.items():
            doc.add_paragraph(f"列：{col}  |  唯一值数量：{info.get('unique', 0)}", style="Heading 4")
            top5 = info.get("top5", {})
            if top5:
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "值"
                hdr[1].text = "频次"
                for cell in hdr:
                    set_cell_bg(cell, "70AD47")
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                for k, v in top5.items():
                    row = table.add_row().cells
                    row[0].text = str(k)
                    row[1].text = str(v)
                doc.add_paragraph()

    # ── 4. 可视化图表 ─────────────────────────────────────────────────────────
    add_heading(doc, "四、可视化图表", 1)
    charts = summary.get("charts", [])
    type_names = {"bar": "柱状图", "pie": "饼状图", "line": "折线图", "barh": "条形图"}
    if charts:
        for chart in charts:
            chart_path = chart.get("path", "")
            if chart_path and os.path.exists(chart_path):
                chart_type = type_names.get(chart.get("type", ""), chart.get("type", ""))
                col_name = chart.get("column", "")
                doc.add_paragraph(f"{chart_type}：{col_name}", style="Heading 3")
                doc.add_picture(chart_path, width=Inches(5.5))
                doc.add_paragraph()
    else:
        doc.add_paragraph("未生成图表（数据列不足）。")

    # ── 5. AI 洞察 ────────────────────────────────────────────────────────────
    add_heading(doc, "五、AI 数据洞察", 1)
    if insight:
        for line in insight.strip().split("\n"):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("（洞察内容待生成）")

    # ── 6. 结论与建议 ─────────────────────────────────────────────────────────
    add_heading(doc, "六、结论与建议", 1)
    doc.add_paragraph("本报告由数据分析 Skill 自动生成，包含完整的 EDA、清洗记录、统计分析、可视化图表及 AI 洞察。")
    doc.add_paragraph("建议结合业务背景对以上洞察进行深入解读，并针对异常值、趋势变化制定相应的业务策略。")

    doc.save(output_path)
    print(f"[OK] Word 报告已保存至: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="生成 Word 数据分析报告")
    parser.add_argument("summary", help="summary.json 路径")
    parser.add_argument("insight", help="AI 洞察文本（可为空字符串）")
    parser.add_argument("--output", default="./analysis_report.docx", help="输出 Word 文件路径")
    args = parser.parse_args()

    generate_report(args.summary, args.insight, args.output)


if __name__ == "__main__":
    main()

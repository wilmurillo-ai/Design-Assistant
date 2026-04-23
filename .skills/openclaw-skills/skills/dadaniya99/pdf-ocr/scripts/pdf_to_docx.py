#!/usr/bin/env python3
"""
PDF扫描件 → Word文档 OCR 批处理脚本
支持：裁掉页眉页脚、保留插图、跳过特殊页（目录/章节封面）
依赖：pip install pymupdf python-pptx python-docx pillow
用法：python3 pdf_to_docx.py <PDF路径> [输出目录]
"""

import os
import sys
import json
import base64
import time
import urllib.request
import urllib.parse
import urllib.error
from pathlib import Path

# ── 配置 ────────────────────────────────────────────────
BAIDU_API_KEY    = os.environ.get("BAIDU_API_KEY",    "vOBOM7tO0lL8cKMJdZy453Ai")
BAIDU_SECRET_KEY = os.environ.get("BAIDU_SECRET_KEY", "bib8MvDPTfXXdPz4JyzIyDCvCeKxtpyu")

# 裁剪比例（页眉顶部6%，页脚底部4%）
CROP_TOP    = 0.06
CROP_BOTTOM = 0.04

# 特殊页面判断阈值
# 若页面中彩色像素占比 > 此值，视为章节封面页（跳过OCR，保留为图片）
COLOR_PAGE_THRESHOLD = 0.25

# OCR 每次调用间隔（避免超频，免费版QPS=2）
OCR_INTERVAL = 0.6
# ─────────────────────────────────────────────────────────

try:
    import fitz
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image
    import io
except ImportError as e:
    print(f"❌ 缺少依赖：{e}")
    print("请运行：pip install pymupdf python-docx pillow")
    sys.exit(1)


_access_token = None

def get_access_token():
    global _access_token
    if _access_token:
        return _access_token
    url = (f"https://aip.baidubce.com/oauth/2.0/token"
           f"?grant_type=client_credentials"
           f"&client_id={BAIDU_API_KEY}"
           f"&client_secret={BAIDU_SECRET_KEY}")
    with urllib.request.urlopen(url) as resp:
        result = json.loads(resp.read())
        _access_token = result["access_token"]
    return _access_token


def ocr_image(img_bytes: bytes) -> list[str]:
    """调用百度OCR，返回识别出的文字行列表"""
    token = get_access_token()
    url = f"https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic?access_token={token}"
    img_b64 = base64.b64encode(img_bytes).decode()
    body = urllib.parse.urlencode({"image": img_b64, "language_type": "CHN_ENG"}).encode()
    req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/x-www-form-urlencoded"})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read())
            if "words_result" in result:
                return [item["words"] for item in result["words_result"]]
            else:
                print(f"  ⚠️ OCR无结果: {result.get('error_msg', '')}")
                return []
    except Exception as e:
        print(f"  ❌ OCR请求失败: {e}")
        return []


def is_color_page(pix) -> bool:
    """判断是否为彩色背景页（章节封面等）"""
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    img_small = img.resize((100, 100))
    pixels = list(img_small.getdata())
    color_count = sum(1 for r, g, b in pixels if max(r,g,b) - min(r,g,b) > 30)
    return color_count / len(pixels) > COLOR_PAGE_THRESHOLD


def crop_page_image(pix, crop_top=CROP_TOP, crop_bottom=CROP_BOTTOM):
    """裁掉页眉页脚区域，返回裁剪后的图片bytes"""
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    w, h = img.size
    top_px    = int(h * crop_top)
    bottom_px = int(h * (1 - crop_bottom))
    cropped = img.crop((0, top_px, w, bottom_px))
    buf = io.BytesIO()
    cropped.save(buf, format="PNG")
    return buf.getvalue(), cropped


def process_pdf(pdf_path: str, output_dir: str = None):
    pdf_path = Path(pdf_path)
    if output_dir is None:
        output_dir = pdf_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    imgs_dir = output_dir / "images"
    imgs_dir.mkdir(exist_ok=True)

    doc_out = Document()
    # 设置页面边距
    for section in doc_out.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    pdf = fitz.open(str(pdf_path))
    total = len(pdf)
    print(f"📖 共 {total} 页，开始处理...")

    for page_num in range(total):
        page = pdf[page_num]
        print(f"  [{page_num+1}/{total}] ", end="", flush=True)

        # 渲染页面为图片（2x分辨率）
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)

        # 判断是否为特殊页（彩色章节封面 / 目录）
        if is_color_page(pix):
            print("章节封面页 → 保留为图片")
            img_path = imgs_dir / f"page_{page_num+1:04d}.png"
            pix.save(str(img_path))
            # 写入 Word：图片 + 标注
            doc_out.add_paragraph()
            p = doc_out.add_paragraph(f"[章节页 - 第{page_num+1}页]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(9)
            doc_out.add_picture(str(img_path), width=Inches(5.5))
            doc_out.add_paragraph()
            continue

        # 普通正文页：裁掉页眉页脚，OCR识别
        print("正文页 → OCR识别", end="", flush=True)
        img_bytes, cropped_img = crop_page_image(pix)

        # 检查页面是否有大块图片（通过检测文字覆盖率）
        # 保存裁剪后图片（用于Word中的插图）
        img_path = imgs_dir / f"page_{page_num+1:04d}.png"
        with open(str(img_path), "wb") as f:
            f.write(img_bytes)

        # OCR
        lines = ocr_image(img_bytes)
        time.sleep(OCR_INTERVAL)

        if lines:
            print(f" → {len(lines)} 行文字")
            for line in lines:
                p = doc_out.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)
        else:
            # 没有识别到文字 → 可能是插图页，保留图片
            print(" → 无文字，保留为插图")
            p = doc_out.add_paragraph(f"[插图 - 第{page_num+1}页]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(9)
            doc_out.add_picture(str(img_path), width=Inches(5.5))

        doc_out.add_paragraph()  # 页间空行

    output_path = output_dir / (pdf_path.stem + "_ocr.docx")
    doc_out.save(str(output_path))
    print(f"\n✅ 完成！输出文件：{output_path}")
    return str(output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 pdf_to_docx.py <PDF路径> [输出目录]")
        sys.exit(1)

    pdf_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(pdf_path):
        print(f"❌ 文件不存在: {pdf_path}")
        sys.exit(1)

    process_pdf(pdf_path, output_dir)

#!/usr/bin env python3
# -*- coding: utf-8 -*-
"""
律师案件归档助手 - 主程序
脱敏版本：不包含任何个人案件信息、邮箱密码等敏感数据
"""

import os
import sys
import gc
import json
import re
import time
import cv2
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 导入配置
from config import (
    TEMPLATE_PATH, MAX_SIDE, CLAHE_CLIP, CLAHE_TILE,
    IMG_EXTS, DEFAULT_LAWYER, DEFAULT_STAGE, DEFAULT_SECRET,
    DEFAULT_ARCHIVE_DATE, get_default_case_info, validate_template
)


# ════════════════════════════════════════════════════════════════════
# 工具函数
# ════════════════════════════════════════════════════════════════════

def set_cell(cell, text, bold=False, size=11, align='left'):
    """设置表格单元格内容"""
    cell.text = ''
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    rFonts = run._element.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        run._element.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc, text, bold=False, size=12, align='left', space_before=0, space_after=0):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    rFonts = run._element.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        run._element.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para


# ════════════════════════════════════════════════════════════════════
# OCR模块
# ════════════════════════════════════════════════════════════════════

def preprocess_img(img_path, max_side=MAX_SIDE):
    """图片预处理：缩放 + CLAHE均衡化"""
    buf = np.fromfile(img_path, dtype=np.uint8)
    img = cv2.imdecode(buf, cv2.IMREAD_COLOR)
    if img is None:
        return None
    h, w = img.shape[:2]
    if max(h, w) > max_side:
        scale = max_side / max(h, w)
        img = cv2.resize(img, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA)
    # CLAHE均衡化，不做去噪（去噪是CPU杀手）
    lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=CLAHE_CLIP, tileGridSize=CLAHE_TILE)
    l = clahe.apply(l)
    lab = cv2.merge([l, a, b])
    img = cv2.cvtColor(lab, cv2.COLOR_LAB2BGR)
    return img


def run_ocr(img_dir):
    """运行OCR识别，返回识别结果文本"""
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        print("[错误] 请安装: pip install rapidocr-onnxruntime")
        return None
    
    imgs = sorted([
        f for f in os.listdir(img_dir)
        if os.path.splitext(f)[1].lower() in IMG_EXTS
    ])
    if not imgs:
        return None
    
    print(f"  找到 {len(imgs)} 张图片，开始OCR识别...")
    engine = RapidOCR()
    results = []
    
    for i, fname in enumerate(imgs, 1):
        fpath = os.path.join(img_dir, fname)
        img = preprocess_img(fpath)
        if img is None:
            continue
        result, _ = engine(img)
        if result:
            lines = [item[1] for item in result]
            results.append(f"【{fname}】\n" + "\n".join(lines))
        del img
        gc.collect()
    
    return "\n\n".join(results)


# ════════════════════════════════════════════════════════════════════
# 案件信息提取
# ════════════════════════════════════════════════════════════════════

def extract_from_existing_pdf(pdf_path):
    """从现有PDF提取案件信息"""
    try:
        import pdfplumber
    except ImportError:
        print("[错误] 请安装: pip install pdfplumber")
        return {}
    
    info = {}
    with pdfplumber.open(pdf_path) as pdf:
        text = "\n".join(p.extract_text() or "" for p in pdf.pages)
    
    # 提取关键信息（正则匹配）
    patterns = {
        'contract_no': r'合同编号[：:]\s*([（(][^）)]+[）)])?',
        'case_no': r'法院编号[：:]\s*([（(][^）)]+[）)])?|(?:案号|法院编号)[:\s]*([（(][^）)]+[）)])?',
        'case_type': r'案由[：:]\s*([^(\n]+)',
        'client': r'委托当事人[：:]\s*([^\n]+)',
        'client_role': r'委托当事人[：:]\s*[^\n（(]+[（(](原告|被告)[）)]',
        'opponent': r'对方当事人[：:]\s*([^\n]+)',
        'opp_role': r'对方当事人[：:]\s*[^\n（(]+[（(](原告|被告)[）)]',
    }
    
    for key, pat in patterns.items():
        m = re.search(pat, text)
        if m:
            val = m.group(1) or m.group(2) or ''
            info[key] = val.strip()
    
    return info


# ════════════════════════════════════════════════════════════════════
# 生成归档卷宗
# ════════════════════════════════════════════════════════════════════

def generate_juansong(case_info, out_path):
    """生成归档卷宗（民事）"""
    if not TEMPLATE_PATH.exists():
        print(f"[错误] 模板文件不存在: {TEMPLATE_PATH}")
        return False
    
    doc = Document(str(TEMPLATE_PATH))
    c = case_info
    
    t0 = doc.tables[0]
    set_cell(t0.cell(1, 2), c.get('contract_no', ''))
    set_cell(t0.cell(1, 3), c.get('contract_no', ''))
    set_cell(t0.cell(1, 4), c.get('contract_no', ''))
    set_cell(t0.cell(3, 2), c.get('case_type', ''))
    set_cell(t0.cell(3, 3), c.get('case_type', ''))
    set_cell(t0.cell(3, 4), c.get('case_type', ''))
    set_cell(t0.cell(4, 2), c.get('client', ''))
    set_cell(t0.cell(4, 3), c.get('client', ''))
    set_cell(t0.cell(4, 4), c.get('client', ''))
    set_cell(t0.cell(4, 6), c.get('client_role', ''))
    set_cell(t0.cell(5, 2), c.get('opponent', ''))
    set_cell(t0.cell(5, 3), c.get('opponent', ''))
    set_cell(t0.cell(5, 4), c.get('opponent', ''))
    set_cell(t0.cell(5, 6), c.get('opp_role', ''))
    set_cell(t0.cell(6, 1), c.get('lawyer', DEFAULT_LAWYER))
    set_cell(t0.cell(6, 2), c.get('lawyer', DEFAULT_LAWYER))
    set_cell(t0.cell(6, 3), '代理阶段')
    set_cell(t0.cell(6, 4), c.get('stage', DEFAULT_STAGE))
    set_cell(t0.cell(6, 5), c.get('stage', DEFAULT_STAGE))
    set_cell(t0.cell(6, 6), c.get('stage', DEFAULT_STAGE))
    set_cell(t0.cell(7, 1), c.get('sign_date', ''))
    set_cell(t0.cell(7, 2), c.get('sign_date', ''))
    set_cell(t0.cell(7, 3), '结案日期')
    set_cell(t0.cell(7, 4), c.get('end_date', ''))
    set_cell(t0.cell(7, 5), c.get('end_date', ''))
    set_cell(t0.cell(7, 6), c.get('end_date', ''))
    set_cell(t0.cell(8, 1), c.get('court', ''))
    set_cell(t0.cell(8, 2), c.get('court', ''))
    set_cell(t0.cell(8, 3), '案号')
    set_cell(t0.cell(8, 4), c.get('case_no', ''))
    set_cell(t0.cell(8, 5), c.get('case_no', ''))
    set_cell(t0.cell(8, 6), c.get('case_no', ''))
    set_cell(t0.cell(12, 1), '长期')
    set_cell(t0.cell(12, 2), '长期')
    set_cell(t0.cell(12, 3), '密级')
    set_cell(t0.cell(12, 4), c.get('secret', DEFAULT_SECRET))
    set_cell(t0.cell(12, 5), c.get('secret', DEFAULT_SECRET))
    set_cell(t0.cell(12, 6), c.get('secret', DEFAULT_SECRET))
    set_cell(t0.cell(13, 1), c.get('lawyer', DEFAULT_LAWYER))
    set_cell(t0.cell(13, 2), c.get('lawyer', DEFAULT_LAWYER))
    set_cell(t0.cell(13, 3), '接收卷人')
    set_cell(t0.cell(13, 4), '档案员')
    set_cell(t0.cell(13, 5), '档案员')
    set_cell(t0.cell(13, 6), '档案员')
    set_cell(t0.cell(14, 1), c.get('archive_date', DEFAULT_ARCHIVE_DATE))
    set_cell(t0.cell(14, 2), c.get('archive_date', DEFAULT_ARCHIVE_DATE))
    set_cell(t0.cell(14, 3), '归档日期')
    set_cell(t0.cell(14, 4), c.get('archive_date', DEFAULT_ARCHIVE_DATE))
    set_cell(t0.cell(14, 5), c.get('archive_date', DEFAULT_ARCHIVE_DATE))
    set_cell(t0.cell(14, 6), c.get('archive_date', DEFAULT_ARCHIVE_DATE))
    pages = c.get('pages', '0')
    set_cell(t0.cell(15, 1), '共 1  册')
    set_cell(t0.cell(15, 2), '共 1  册')
    set_cell(t0.cell(15, 3), '本册页数')
    set_cell(t0.cell(15, 4), f'第 1  册，共  {pages} 页')
    set_cell(t0.cell(15, 5), f'第 1  册，共  {pages} 页')
    set_cell(t0.cell(15, 6), f'第 1  册，共  {pages} 页')
    
    t1 = doc.tables[1]
    set_cell(t1.cell(1, 1), c.get('lawyer', DEFAULT_LAWYER))
    set_cell(t1.cell(2, 1), '共 1  册')
    
    catalog = c.get('catalog', [])
    t2 = doc.tables[2]
    for i, item in enumerate(catalog):
        row_idx = i + 1
        if row_idx < len(t2.rows):
            set_cell(t2.cell(row_idx, 0), str(i + 1), align='center')
            set_cell(t2.cell(row_idx, 1), item)
    
    doc.save(out_path)
    print(f"  已保存: {out_path}")
    return True


# ════════════════════════════════════════════════════════════════════
# 生成办案小结
# ════════════════════════════════════════════════════════════════════

def generate_xiaojie(case_info, out_path):
    """生成办案小结"""
    doc = Document()
    
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2.5)
    
    c = case_info
    
    add_para(doc, '办  案  小  结', bold=True, size=16, align='center', space_after=6)
    
    info_table = doc.add_table(rows=8, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ('委托人',   c.get('client', ''),        '对方当事人', c.get('opponent', '')),
        ('合同编号', c.get('contract_no', ''),  '案号',       c.get('case_no', '')),
        ('案由',     c.get('case_type', ''),     '代理阶段',   c.get('stage', '')),
        ('受理单位', c.get('court', ''),         '承办律师',   c.get('lawyer', DEFAULT_LAWYER)),
        ('接案日期', c.get('sign_date', ''),    '结案日期',   c.get('end_date', '')),
        ('结案方式', c.get('end_way', ''),       '律师费',     c.get('fee', '')),
        ('归档日期', c.get('archive_date', ''),'密级',       c.get('secret', DEFAULT_SECRET)),
        ('页数',     f'共 {c.get("pages", "0")} 页', '',      ''),
    ]
    for i, (k1, v1, k2, v2) in enumerate(info_data):
        row = info_table.rows[i]
        set_cell(row.cells[0], k1, bold=True)
        set_cell(row.cells[1], v1)
        set_cell(row.cells[2], k2, bold=True)
        set_cell(row.cells[3], v2)
    
    doc.add_paragraph()
    
    add_para(doc, '一、案情经过', bold=True, size=12, space_before=6)
    case_type = c.get('case_type', '')
    client = c.get('client', '')
    opponent = c.get('opponent', '')
    sign_date = c.get('sign_date', '')
    court = c.get('court', '')
    case_no = c.get('case_no', '')
    client_role = c.get('client_role', '')
    
    if client_role == '被告':
        add_para(doc,
            f'原告{opponent}诉被告{client}{case_type}案，'
            f'被告{client}于{sign_date}委托本所{c.get("lawyer", DEFAULT_LAWYER)}律师担任一审诉讼代理人，'
            f'向{court}应诉（案号：{case_no}）。',
            size=12)
    else:
        add_para(doc,
            f'委托人{client}因与{opponent}之间发生{case_type}，'
            f'于{sign_date}委托本所{c.get("lawyer", DEFAULT_LAWYER)}律师担任诉讼代理人，'
            f'向{court}提起诉讼（案号：{case_no}）。',
            size=12)
    
    add_para(doc, '二、代理工作情况', bold=True, size=12, space_before=6)
    add_para(doc,
        f'承办律师{c.get("lawyer", DEFAULT_LAWYER)}依约完成以下工作：\n'
        f'（一）审查案件材料，研究相关法律规定，制订诉讼策略；\n'
        f'（二）起草并提交起诉状及相关证据材料；\n'
        f'（三）代理出庭参加诉讼，维护委托人合法权益；\n'
        f'（四）与委托人保持沟通，及时告知案件进展。',
        size=12)
    
    end_way = c.get('end_way', '')
    end_date = c.get('end_date', '')
    add_para(doc, '三、结案情况', bold=True, size=12, space_before=6)
    if '撤诉' in end_way:
        add_para(doc,
            f'经审理，委托人{client}于{end_date}申请撤回该案诉讼，本案以撤诉方式结案。',
            size=12)
    elif '调解' in end_way:
        add_para(doc,
            f'经{court}主持调解，于{end_date}达成调解协议，本案以调解方式结案。',
            size=12)
    elif '判决' in end_way:
        add_para(doc,
            f'经{court}审理，于{end_date}作出判决，本案以判决方式结案。'
            f'承办律师已将判决书送达委托人，并告知其相关权利。',
            size=12)
    else:
        add_para(doc,
            f'本案于{end_date}以{end_way}方式结案。',
            size=12)
    
    add_para(doc, '四、经验与教训', bold=True, size=12, space_before=6)
    add_para(doc,
        f'本案系{case_type}，已完结。'
        f'提示在代理此类案件时，应充分了解案件事实，准确把握法律关系，'
        f'制定合理的诉讼策略，以最大程度维护委托人的合法权益。',
        size=12)
    
    doc.add_paragraph()
    add_para(doc, f'承办律师：{c.get("lawyer", DEFAULT_LAWYER)}', size=12, align='right')
    add_para(doc, f'归档日期：{c.get("archive_date", "")}', size=12, align='right')
    
    doc.save(out_path)
    print(f"  已保存: {out_path}")
    return True


# ════════════════════════════════════════════════════════════════════
# Word转PDF
# ════════════════════════════════════════════════════════════════════

def docx_to_pdf(docx_path, pdf_path):
    """Word转PDF（需要Windows + Word COM）"""
    import subprocess
    cmd = f'''
    $word = New-Object -ComObject Word.Application; 
    $word.Visible = $false; 
    $word.DisplayAlerts = 0; 
    $doc = $word.Documents.Open("{docx_path}"); 
    $doc.ExportAsFixedFormat("{pdf_path}", 17); 
    $doc.Close($false); 
    $word.Quit()
    '''
    subprocess.run(['powershell', '-Command', cmd], capture_output=True)
    print(f"  已转换: {pdf_path}")


# ════════════════════════════════════════════════════════════════════
# 主程序
# ════════════════════════════════════════════════════════════════════

def main(case_dir, case_info=None):
    """案件归档主流程
    
    Args:
        case_dir: 案件文件夹路径
        case_info: 案件信息字典（如果为None，则从现有PDF提取或OCR识别）
    """
    case_dir = Path(case_dir)
    if not case_dir.exists():
        print(f"[错误] 目录不存在: {case_dir}")
        return False
    
    case_name = case_dir.name
    print(f"\n{'='*50}")
    print(f"开始归档: {case_name}")
    print('='*50)
    
    # 验证模板
    if not validate_template():
        print("[错误] 请先配置正确的模板路径（config.py）")
        return False
    
    # 1. 确定案件信息
    if case_info is None:
        # 尝试从现有PDF提取
        existing_pdf = None
        for f in case_dir.iterdir():
            if f.suffix.lower() == '.pdf' and '归档卷宗' in f.name:
                existing_pdf = f
                break
        
        if existing_pdf:
            print(f"  从现有PDF提取信息: {existing_pdf.name}")
            case_info = extract_from_existing_pdf(str(existing_pdf))
        else:
            # 尝试OCR识别
            ocr_result = run_ocr(str(case_dir))
            if ocr_result:
                # 保存OCR结果
                ocr_file = case_dir / f'{case_name}OCR识别结果.txt'
                with open(ocr_file, 'w', encoding='utf-8') as f:
                    f.write(ocr_result)
                print(f"  OCR结果已保存: {ocr_file.name}")
                # TODO: 进一步从OCR结果提取案件信息
                case_info = {}
            else:
                case_info = {}
    
    # 合并默认信息
    defaults = get_default_case_info()
    for k, v in defaults.items():
        if k not in case_info or not case_info.get(k):
            case_info[k] = v
    
    # 2. 生成归档卷宗
    juansong_docx = case_dir / f'{case_name}归档卷宗（民事）.docx'
    print("\n[1/4] 生成归档卷宗...")
    generate_juansong(case_info, str(juansong_docx))
    
    # 3. 生成办案小结
    xiaojie_docx = case_dir / f'{case_name}办案小结.docx'
    print("\n[2/4] 生成办案小结...")
    generate_xiaojie(case_info, str(xiaojie_docx))
    
    # 4. 转PDF
    print("\n[3/4] 转换PDF...")
    docx_to_pdf(str(juansong_docx), str(juansong_docx.with_suffix('.pdf')))
    docx_to_pdf(str(xiaojie_docx), str(xiaojie_docx.with_suffix('.pdf')))
    
    print(f"\n{'='*50}")
    print(f"归档完成！")
    print(f"  归档卷宗: {case_name}归档卷宗（民事）.docx/.pdf")
    print(f"  办案小结: {case_name}办案小结.docx/.pdf")
    print('='*50)
    return True


def show_welcome():
    """显示使用引导提醒"""
    import platform
    print()
    print("╔" + "═"*58 + "╗")
    print("║" + "  律师案件归档助手 — 使用引导".center(50) + "║")
    print("╠" + "═"*58 + "╣")
    print("║  本技能由方鸿源律师发布，仅作个人办案参考工具使用；")
    print("║  有关法律、办案及养虾，也可联系13512717203（微信）。")
    print("║  近乎公平即正义，愿我们的法治公平早日实现！")
    print("║")
    print("║  使用前请确认以下事项：")
    print("║")
    print("║  1️⃣  模板文件路径")
    print("║     打开 config.py，修改 TEMPLATE_PATH")
    print("║     指向您的归档卷宗 Word 模板文件")
    print("║     例如: TEMPLATE_PATH = r\"D:\\templates\\归档卷宗模板.docx\"")
    print("║")
    print("║  2️⃣  律师姓名（可选）")
    print("║     在 config.py 中修改 DEFAULT_LAWYER")
    print("║     默认值: \"您的姓名\"")
    print("║")
    print("║  3️⃣  案件文件夹结构")
    print("║     准备一个文件夹，包含案件相关的图片或PDF")
    print("║     图片格式: .jpg / .jpeg / .png / .bmp / .tiff")
    print("║     如果已有归档卷宗PDF，会自动提取案件信息")
    print("║")
    print("║  4️⃣  依赖安装（首次使用）")
    print("║     pip install rapidocr-onnxruntime python-docx")
    print("║     pip install pdfplumber openpyxl opencv-python")
    print("║")
    print("║  5️⃣  运行命令")
    print("║     python archive_case.py \"<案件文件夹路径>\"")
    print("║     例如: python archive_case.py \"D:\\cases\\借款纠纷\"")
    print("║")
    print("║  ⚠️  Word转PDF 需要本机安装 Microsoft Word")
    print("║")
    print("╚" + "═"*58 + "╝")
    print()

    # 检查模板文件
    from config import TEMPLATE_PATH, validate_template
    print("[检查] 模板文件: ", end="")
    if TEMPLATE_PATH.exists():
        print(f"✅ 已找到 ({TEMPLATE_PATH})")
    else:
        print(f"❌ 未找到!")
        print(f"  当前路径: {TEMPLATE_PATH}")
        print(f"  请修改 config.py 中的 TEMPLATE_PATH 后再运行")
        print()
        choice = input("  是否继续运行？(y/N): ").strip().lower()
        if choice != 'y':
            print("  已退出。请配置好模板后重新运行。")
            sys.exit(0)
    print()


if __name__ == '__main__':
    # ── 启动引导 ──
    show_welcome()

    if len(sys.argv) < 2:
        print("用法: python archive_case.py <案件文件夹>")
        print("示例: python archive_case.py \"D:\\cases\\张三\"")
        sys.exit(1)
    
    case_dir = sys.argv[1]
    if not os.path.isdir(case_dir):
        print(f"错误: 目录不存在: {case_dir}")
        sys.exit(1)
    
    main(case_dir)
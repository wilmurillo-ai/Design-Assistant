#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
光伏电站巡检日报生成脚本（Word 格式）
支持中文文件名，生成 .docx 格式报告
"""

import argparse
import json
from datetime import datetime, timedelta
from pathlib import Path
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("错误：请安装 python-docx 库")
    print("运行：pip install python-docx")
    sys.exit(1)

# 输出目录
OUTPUT_DIR = Path("/home/admin/.openclaw/workspace/pv-inspection/word-reports")


def create_daily_report(station_name: str, date_str: str):
    """生成 Word 格式的日报"""
    date = datetime.strptime(date_str, "%Y-%m-%d")
    
    # 模拟数据（实际应从监控系统获取）
    report_data = {
        "station": station_name,
        "date": date_str,
        "inspector": "巡检人员",
        "weather": "晴",
        "temperature": "15-25°C",
        "generation": {
            "daily": 1250.5,  # kWh
            "cumulative": 45678.9,  # kWh
            "peak_power": 850.2  # kW
        },
        "equipment": {
            "inverter": {"total": 10, "normal": 9, "abnormal": 1},
            "combiner": {"total": 20, "normal": 20, "abnormal": 0},
            "modules": {"total": 500, "normal": 498, "abnormal": 2}
        },
        "defects": [
            {
                "id": 1,
                "location": "3 号逆变器",
                "description": "通讯故障",
                "severity": "重要",
                "suggestion": "联系厂家维修"
            }
        ]
    }
    
    # 创建文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('光伏电站巡检日报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息（表格形式）
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    # 填充基本信息
    info_data = [
        ("电站名称", report_data['station']),
        ("巡检日期", report_data['date']),
        ("巡检人员", report_data['inspector']),
        ("天气情况", f"{report_data['weather']} {report_data['temperature']}")
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        # 加粗标签
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # 空行
    
    # 一、发电数据
    doc.add_heading('一、发电数据', level=1)
    gen_table = doc.add_table(rows=3, cols=2)
    gen_table.style = 'Table Grid'
    
    gen_data = [
        ("今日发电量", f"{report_data['generation']['daily']:.1f} kWh"),
        ("累计发电量", f"{report_data['generation']['cumulative']:.1f} kWh"),
        ("峰值功率", f"{report_data['generation']['peak_power']:.1f} kW")
    ]
    
    for i, (label, value) in enumerate(gen_data):
        gen_table.rows[i].cells[0].text = label
        gen_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # 二、设备状态
    doc.add_heading('二、设备状态', level=1)
    equip_table = doc.add_table(rows=4, cols=4)  # 4 行：1 行表头 + 3 行数据
    equip_table.style = 'Table Grid'
    
    # 表头
    header_cells = equip_table.rows[0].cells
    headers = ["设备类型", "总数", "正常", "异常"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # 设备数据
    equip_data = [
        ("逆变器", report_data['equipment']['inverter']['total'], 
         report_data['equipment']['inverter']['normal'], 
         report_data['equipment']['inverter']['abnormal']),
        ("汇流箱", report_data['equipment']['combiner']['total'], 
         report_data['equipment']['combiner']['normal'], 
         report_data['equipment']['combiner']['abnormal']),
        ("光伏组件", report_data['equipment']['modules']['total'], 
         report_data['equipment']['modules']['normal'], 
         report_data['equipment']['modules']['abnormal'])
    ]
    
    for i, (name, total, normal, abnormal) in enumerate(equip_data, 1):
        row = equip_table.rows[i]
        row.cells[0].text = name
        row.cells[1].text = str(total)
        row.cells[2].text = str(normal)
        row.cells[3].text = str(abnormal)
    
    doc.add_paragraph()
    
    # 三、缺陷记录
    doc.add_heading('三、缺陷记录', level=1)
    
    if report_data['defects']:
        defect_table = doc.add_table(rows=len(report_data['defects']) + 1, cols=5)
        defect_table.style = 'Table Grid'
        
        # 表头
        header_cells = defect_table.rows[0].cells
        headers = ["序号", "位置", "缺陷描述", "严重程度", "整改建议"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # 缺陷数据
        for i, defect in enumerate(report_data['defects'], 1):
            row = defect_table.rows[i]
            row.cells[0].text = str(defect['id'])
            row.cells[1].text = defect['location']
            row.cells[2].text = defect['description']
            row.cells[3].text = defect['severity']
            row.cells[4].text = defect['suggestion']
    else:
        doc.add_paragraph("今日无缺陷记录", style='Intense Quote')
    
    doc.add_paragraph()
    
    # 四、明日计划
    doc.add_heading('四、明日计划', level=1)
    plan_para = doc.add_paragraph()
    plan_para.add_run('☐ 待整改项目跟踪\n')
    plan_para.add_run('☐ 重点设备检查\n')
    plan_para.add_run('☐ 数据归档与备份')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 页脚
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 生成中文文件名
    filename = f"{station_name}_{date_str}_巡检日报.docx"
    output_file = OUTPUT_DIR / filename
    
    # 保存文档
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    
    print(f"✅ Word 日报已生成：{output_file}")
    return str(output_file)


def main():
    parser = argparse.ArgumentParser(description="光伏电站巡检日报生成（Word 格式）")
    parser.add_argument("--station", required=True, help="电站名称")
    parser.add_argument("--date", help="日期（YYYY-MM-DD），默认为今天")
    
    args = parser.parse_args()
    
    if not args.date:
        args.date = datetime.now().strftime("%Y-%m-%d")
    
    create_daily_report(args.station, args.date)


if __name__ == "__main__":
    main()

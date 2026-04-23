import os
import sys
import re
import uuid
import tempfile
import time

# ==========================================
# [模块] 必要模块导入 (带异常处理)
# ==========================================
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    pass # 延迟报错，在 check_dependencies 中处理

# ==========================================
# [环境] 检测与防御
# ==========================================
def check_dependencies():
    """检查运行环境依赖（软依赖，允许模块不存在）"""
    try:
        import win32com.client
        word_app = win32com.client.Dispatch('Word.Application')
        word_app.Quit()
        return True, "Word COM 组件可用"
    except ImportError:
        return False, "缺少 pywin32 模块，请运行 pip install pywin32"
    except Exception as e:
        return False, f"环境检测失败: {e}"

# 只有在直接运行脚本时才执行环境检测
if __name__ == "__main__":
    is_ready, env_message = check_dependencies()
    if not is_ready:
        print(f"[错误] {env_message}")
        sys.exit(1)

# 编码处理
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except AttributeError:
    os.environ['PYTHONIOENCODING'] = 'utf-8'
    os.system('chcp 65001 > nul')

# ==========================================
# [清洗] 智能文本清洗
# ==========================================
def clean_text(text):
    """智能文本清洗：保留中文字符，移除危险符号"""
    if not text: return ""
    
    # 移除 Emoji（覆盖常见范围）
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002702-\U000027B0"
        u"\U0001F900-\U0001F9FF"
        u"\U0001F004"
        "]+", flags=re.UNICODE)
    text = emoji_pattern.sub(r'', text)
    
    # 移除 Unicode 控制字符
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    
    return text.strip()

# ==========================================
# [工具] XML 级样式控制
# ==========================================
def get_or_add_rPr(run):
    """安全获取或创建 rPr 元素"""
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.append(rPr)
    return rPr

def get_safe_font(font_name):
    """字体回退机制：确保中文字体可用"""
    font_mapping = {
        '宋体': 'SimSun',
        '黑体': 'SimHei', 
        '楷体': 'KaiTi',
        '仿宋': 'FangSong',
        '微软雅黑': 'Microsoft YaHei',
        '华文楷体': 'STKaiti',
        '华文行楷': 'STXingkai',
        'Arial': 'Arial',
    }
    return font_mapping.get(font_name, 'SimSun')

def set_chinese_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    """
    安全的中文字体设置
    注：强制西文为 Arial 是为了保证文档在不同机器上显示一致，
    中文字体由 eastAsia 属性单独指定，互不干扰。
    """
    try:
        safe_font = get_safe_font(font_name)
        run.font.name = 'Arial'  # 强制西文，防止使用默认字体导致乱码
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
            
        rPr = get_or_add_rPr(run)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), safe_font)
    except Exception as e:
        print(f"   [警告] 字体设置失败: {e}")

def set_character_spacing(run, spacing=0):
    """设置字符间距 (单位: 磅)"""
    try:
        rPr = get_or_add_rPr(run)
        spacing_elem = rPr.find(qn('w:spacing'))
        if spacing_elem is None:
            spacing_elem = OxmlElement('w:spacing')
            rPr.append(spacing_elem)
        spacing_elem.set(qn('w:val'), str(int(spacing * 20)))
    except Exception as e:
        print(f"   [警告] 字符间距设置失败: {e}")

def set_first_line_indent(paragraph, chars=2):
    """设置首行缩进（单位：字符数）"""
    if chars <= 0:
        return
    try:
        # 清除可能存在的绝对缩进（磅值），避免冲突
        paragraph.paragraph_format.first_line_indent = 0

        # 获取或创建段落属性
        pPr = paragraph._element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)

        # 设置 w:firstLineChars，单位是百分之一字符
        ind.set(qn('w:firstLineChars'), str(chars * 100))
    except Exception as e:
        print(f"   [警告] 缩进设置失败: {e}")

def set_paragraph_spacing(paragraph, before=0, after=6):
    """设置段落间距 (单位: 磅)"""
    try:
        pPr = paragraph._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        if before > 0:
            spacing.set(qn('w:before'), str(int(before * 20)))
        if after > 0:
            spacing.set(qn('w:after'), str(int(after * 20)))
    except Exception as e:
        print(f"   [警告] 段落间距设置失败: {e}")

# ==========================================
# [核心] 工业级文档生成器
# ==========================================
def create_robust_word_doc(
    title, 
    content_list, 
    output_path, 
    subtitle=None, 
    note=None
):
    """
    终极工业级文档生成器
    流程：COM 创建原生模板 -> python-docx 写入内容 -> 深度验证
    """
    result = {"success": False, "message": "", "path": output_path}
    word_app = None
    doc_com = None
    temp_path = None
    
    # --- 预处理与目录检查 ---
    try:
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            
        if os.path.isdir(output_path):
            result["message"] = "输出路径是目录，不是文件路径"
            return result
            
        clean_title = clean_text(title)
        clean_subtitle = clean_text(subtitle) if subtitle else None
        clean_note = clean_text(note) if note else None
        
        processed_contents = []
        for item in content_list:
            text = clean_text(item.get('text', ''))
            is_body = item.get('is_body', True)
            processed_contents.append({'text': text, 'is_body': is_body})
            
    except Exception as e:
        result["message"] = f"预处理失败: {e}"
        return result

    # --- 步骤 1: COM 创建原生空白文档 ---
    temp_filename = f"temp_blank_{uuid.uuid4().hex}.docx"
    temp_path = os.path.join(tempfile.gettempdir(), temp_filename)
    
    try:
        print(f"1. [COM] 正在生成原生模板...")
        import win32com.client
        word_app = win32com.client.Dispatch('Word.Application')
        word_app.Visible = False
        word_app.AutomationSecurity = 3  # 禁用宏警告
        word_app.DisplayAlerts = False   # 禁用弹窗
        
        doc_com = word_app.Documents.Add(Template="Normal.dotm")
        doc_com.SaveAs(FileName=temp_path)
        print("   [通过] 原生模板就绪")
    except Exception as e:
        result["message"] = f"COM 初始化失败: {e}"
        return result
    finally:
        # 只关闭文档和退出应用，**不删除** temp_path
        try:
            if doc_com:
                doc_com.Close(SaveChanges=False)
            if word_app:
                word_app.Quit()
                time.sleep(0.2) # 等待文件锁释放
        except:
            pass

    # --- 步骤 2: python-docx 写入内容 ---
    try:
        print(f"2. [Docx] 正在写入内容...")
        
        if not os.path.exists(temp_path):
            raise FileNotFoundError(f"临时文件丢失: {temp_path}")
            
        doc = Document(temp_path)
        
        # 清理 COM 模板遗留的默认空段落
        if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
            print("   [提示] 已清理默认空段落")
        
        # A. 写入主标题
        heading = doc.add_heading(clean_title, 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            set_chinese_font(run, '黑体', 16, bold=True)
            set_character_spacing(run, 0)
        set_paragraph_spacing(heading, before=0, after=12)

        # B. 写入副标题
        if clean_subtitle:
            sub = doc.add_paragraph(clean_subtitle)
            sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in sub.runs:
                set_chinese_font(run, '楷体', 12, color=RGBColor(100, 100, 100))
                set_character_spacing(run, 0)
            set_paragraph_spacing(sub, before=0, after=12)

        # C. 写入正文内容（区分样式）
        for item in processed_contents:
            p = doc.add_paragraph(item['text'])
            for run in p.runs:
                if item['is_body']:
                    set_chinese_font(run, '宋体', 12)
                    set_character_spacing(run, 0)
                else:
                    set_chinese_font(run, '黑体', 14, bold=True)
                    set_character_spacing(run, 0)
            
            if item['is_body']:
                set_first_line_indent(p, chars=2)
                set_paragraph_spacing(p, before=0, after=6)
            else:
                set_first_line_indent(p, chars=0)
                set_paragraph_spacing(p, before=12, after=6)

        # D. 写入注释
        if clean_note:
            doc.add_paragraph()
            note_p = doc.add_paragraph(clean_note)
            for run in note_p.runs:
                set_chinese_font(run, '宋体', 9, color=RGBColor(150, 150, 150))
                set_character_spacing(run, 0)
            set_paragraph_spacing(note_p, before=6, after=0)

        doc.save(output_path)
        print("   [通过] 内容写入完成")

    except Exception as e:
        result["message"] = f"文档写入失败: {e}"
        return result
    finally:
        # 清理临时文件
        try:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)
        except:
            pass

    # --- 步骤 3: 全面验证 ---
    print("3. [验证] 正在进行深度体检...")
    
    if not os.path.exists(output_path):
        result["message"] = "文件物理路径不存在"
        return result
        
    size = os.path.getsize(output_path)
    if size < 5000:
        result["message"] = f"文件异常过小 ({size} 字节)，可能损坏"
        return result
        
    try:
        verify_doc = Document(output_path)
        
        # 验证标题是否存在
        if len(verify_doc.paragraphs) == 0 or verify_doc.paragraphs[0].text != clean_title:
            result["message"] = "标题验证失败"
            return result

        # 验证段落数量大致合理
        if len(verify_doc.paragraphs) < len(processed_contents) + 2:
            result["message"] = "段落数量异常"
            return result
            
        print("   [通过] 内容完整性校验通过")
    except Exception as e:
        result["message"] = f"文件损坏: {e}"
        return result

    result["success"] = True
    result["message"] = f"文档生成成功 ({size} 字节)"
    return result

# ==========================================
# [启动] 调用演示
# ==========================================
if __name__ == "__main__":
    print(f"[开始] 启动工业级文档生成任务...")
    
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    output_file = os.path.join(desktop_path, '规划报告_最终版.docx')
    
    contents = [
        {'text': '这是第一段正文，应该缩进。', 'is_body': True},
        {'text': '这是第二段正文，也应该缩进。', 'is_body': True},
        {'text': '这是一个小标题（不缩进，加粗）', 'is_body': False}, 
        {'text': '这是第三段正文，继续缩进。', 'is_body': True},
    ]
    
    res = create_robust_word_doc(
        title="关于未来发展的规划报告",
        content_list=contents,
        output_path=output_file,
        subtitle="（2026-2030）",
        note="注：本报告数据仅供参考"
    )
    
    if res["success"]:
        print(f"\n[完成] {res['message']}")
        print(f"     路径: {res['path']}")
        
        # 正则提取文件大小
        match = re.search(r'\((\d+) 字节\)', res["message"])
        size = match.group(1) if match else "未知"
        print(f"     文件大小: {size} 字节")
    else:
        print(f"\n[失败] {res['message']}")
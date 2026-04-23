#!/usr/bin/env python3
"""
带图片处理的Markdown转Word转换器
"""

import argparse
import os
import sys
import re
from pathlib import Path
from typing import Optional, List, Tuple
from PIL import Image

# 导入主转换器
from md2docx import convert_markdown_to_docx, MarkdownToDocxConverter


class ImageProcessor:
    """图片处理器"""
    
    def __init__(self, image_dir: str, debug: bool = False):
        self.image_dir = image_dir
        self.debug = debug
        self.image_cache = {}
        
        # 支持的图片格式
        self.supported_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
    
    def find_image(self, image_src: str) -> Optional[str]:
        """
        查找图片文件
        
        Args:
            image_src: 图片路径（相对或绝对）
            
        Returns:
            Optional[str]: 找到的图片路径，或None
        """
        # 如果是绝对路径，直接检查
        if os.path.isabs(image_src):
            if os.path.exists(image_src):
                return image_src
            return None
        
        # 检查缓存
        if image_src in self.image_cache:
            return self.image_cache[image_src]
        
        # 可能的查找位置
        search_paths = []
        
        # 1. 相对于图片目录
        if self.image_dir:
            search_paths.append(os.path.join(self.image_dir, image_src))
        
        # 2. 相对于当前目录
        search_paths.append(image_src)
        
        # 3. 在图片目录中搜索文件名
        if self.image_dir:
            filename = os.path.basename(image_src)
            for root, _, files in os.walk(self.image_dir):
                if filename in files:
                    search_paths.append(os.path.join(root, filename))
        
        # 尝试每个路径
        for path in search_paths:
            if os.path.exists(path):
                self.image_cache[image_src] = path
                if self.debug:
                    print(f"🔍 找到图片: {image_src} → {path}")
                return path
        
        if self.debug:
            print(f"⚠️  未找到图片: {image_src}")
        return None
    
    def process_image(self, image_path: str, max_width: Optional[int] = None, 
                     max_height: Optional[int] = None, quality: int = 85) -> Tuple[str, dict]:
        """
        处理图片（调整大小、压缩等）
        
        Args:
            image_path: 图片路径
            max_width: 最大宽度（像素）
            max_height: 最大高度（像素）
            quality: JPEG质量（1-100）
            
        Returns:
            Tuple[str, dict]: (处理后的图片路径, 图片信息)
        """
        if not os.path.exists(image_path):
            return image_path, {'error': '文件不存在'}
        
        try:
            # 打开图片
            with Image.open(image_path) as img:
                original_size = img.size
                format_name = img.format
                mode = img.mode
                
                info = {
                    'path': image_path,
                    'original_size': original_size,
                    'format': format_name,
                    'mode': mode,
                    'file_size': os.path.getsize(image_path)
                }
                
                # 检查是否需要调整大小
                needs_resize = False
                new_size = original_size
                
                if max_width and original_size[0] > max_width:
                    ratio = max_width / original_size[0]
                    new_width = max_width
                    new_height = int(original_size[1] * ratio)
                    new_size = (new_width, new_height)
                    needs_resize = True
                
                if max_height and new_size[1] > max_height:
                    ratio = max_height / new_size[1]
                    new_width = int(new_size[0] * ratio)
                    new_height = max_height
                    new_size = (new_width, new_height)
                    needs_resize = True
                
                if needs_resize:
                    # 调整大小
                    img_resized = img.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # 保存调整后的图片
                    output_path = self._get_temp_image_path(image_path)
                    img_resized.save(output_path, quality=quality)
                    
                    info['resized_size'] = new_size
                    info['resized_path'] = output_path
                    info['resized_file_size'] = os.path.getsize(output_path)
                    
                    if self.debug:
                        print(f"🖼️  图片调整大小: {original_size} → {new_size}")
                        print(f"    文件大小: {info['file_size']:,} → {info['resized_file_size']:,} 字节")
                    
                    return output_path, info
                else:
                    # 不需要调整大小
                    if self.debug:
                        print(f"🖼️  图片保持原大小: {original_size}")
                    
                    return image_path, info
        
        except Exception as e:
            if self.debug:
                print(f"❌ 图片处理失败: {e}")
            return image_path, {'error': str(e)}
    
    def _get_temp_image_path(self, original_path: str) -> str:
        """获取临时图片路径"""
        import tempfile
        import hashlib
        
        # 生成唯一文件名
        file_hash = hashlib.md5(original_path.encode()).hexdigest()[:8]
        ext = os.path.splitext(original_path)[1]
        
        temp_dir = tempfile.gettempdir()
        temp_filename = f"md2docx_{file_hash}{ext}"
        
        return os.path.join(temp_dir, temp_filename)
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        import tempfile
        import glob
        
        temp_dir = tempfile.gettempdir()
        temp_pattern = os.path.join(temp_dir, "md2docx_*")
        
        temp_files = glob.glob(temp_pattern)
        for temp_file in temp_files:
            try:
                os.remove(temp_file)
                if self.debug:
                    print(f"🧹 清理临时文件: {temp_file}")
            except:
                pass


class EnhancedMarkdownToDocxConverter(MarkdownToDocxConverter):
    """增强的Markdown转Word转换器（带图片处理）"""
    
    def __init__(self, template_path: Optional[str] = None, 
                 image_processor: Optional[ImageProcessor] = None,
                 debug: bool = False):
        super().__init__(template_path, debug)
        self.image_processor = image_processor
        self.image_info = []
    
    def _add_image(self, element, image_dir: Optional[str] = None):
        """添加图片（增强版）"""
        src = element.get('src', '')
        alt = element.get('alt', '图片')
        
        if not src:
            return
        
        # 查找图片
        image_path = None
        if self.image_processor:
            image_path = self.image_processor.find_image(src)
        
        # 如果没找到，尝试其他方法
        if not image_path and image_dir:
            # 相对于图片目录
            potential_path = os.path.join(image_dir, src)
            if os.path.exists(potential_path):
                image_path = potential_path
        
        if not image_path:
            # 直接使用路径
            image_path = src
        
        # 检查图片文件是否存在
        if not os.path.exists(image_path):
            if self.debug:
                print(f"⚠️  图片文件不存在: {image_path}")
            # 添加替代文本
            p = self.doc.add_paragraph(f"[图片: {alt}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        
        try:
            # 处理图片（调整大小等）
            processed_path = image_path
            image_data = {'original_path': image_path}
            
            if self.image_processor:
                processed_path, image_data = self.image_processor.process_image(
                    image_path, 
                    max_width=1200,  # 最大宽度1200像素
                    max_height=800,  # 最大高度800像素
                    quality=85
                )
                self.image_info.append(image_data)
            
            # 添加图片到文档
            from docx.shared import Inches
            self.doc.add_picture(processed_path, width=Inches(5.0))
            
            # 添加图片标题
            if alt:
                p = self.doc.add_paragraph(f"图: {alt}", style=self.style_map['img'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if self.debug:
                print(f"🖼️  添加图片: {alt}")
                if 'resized_size' in image_data:
                    print(f"    大小: {image_data['original_size']} → {image_data['resized_size']}")
        
        except Exception as e:
            if self.debug:
                print(f"❌ 添加图片失败: {e}")
            # 添加替代文本
            p = self.doc.add_paragraph(f"[图片加载失败: {alt}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def convert_with_images(input_file: str, output_file: str, 
                       template: Optional[str] = None,
                       image_dir: Optional[str] = None,
                       max_width: Optional[int] = None,
                       max_height: Optional[int] = None,
                       image_quality: int = 85,
                       add_captions: bool = True,
                       debug: bool = False) -> bool:
    """
    带图片处理的Markdown转Word转换
    
    Args:
        input_file: 输入Markdown文件路径
        output_file: 输出Word文件路径
        template: 模板文件路径
        image_dir: 图片目录路径
        max_width: 图片最大宽度
        max_height: 图片最大高度
        image_quality: 图片质量（1-100）
        add_captions: 是否添加图片标题
        debug: 是否启用调试模式
        
    Returns:
        bool: 转换是否成功
    """
    try:
        # 检查输入文件
        if not os.path.exists(input_file):
            print(f"❌ 输入文件不存在: {input_file}")
            return False
        
        # 读取Markdown文件
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        if debug:
            print(f"📄 读取Markdown文件: {input_file} ({len(markdown_text)} 字符)")
        
        # 创建图片处理器
        image_processor = None
        if image_dir:
            image_processor = ImageProcessor(image_dir, debug)
            if debug:
                print(f"🖼️  使用图片目录: {image_dir}")
        
        # 创建增强转换器
        converter = EnhancedMarkdownToDocxConverter(template, image_processor, debug)
        
        # 转换
        converter.convert(markdown_text, image_dir)
        
        # 保存
        converter.save(output_file)
        
        # 清理临时文件
        if image_processor:
            image_processor.cleanup_temp_files()
        
        # 打印图片处理摘要
        if converter.image_info:
            print(f"\n📊 图片处理摘要:")
            total_original = sum(info.get('file_size', 0) for info in converter.image_info)
            total_resized = sum(info.get('resized_file_size', 0) for info in converter.image_info 
                              if 'resized_file_size' in info)
            
            if total_resized > 0:
                reduction = (total_original - total_resized) / total_original * 100
                print(f"  处理图片: {len(converter.image_info)} 张")
                print(f"  原始大小: {total_original:,} 字节")
                print(f"  处理后大小: {total_resized:,} 字节")
                print(f"  节省空间: {reduction:.1f}%")
        
        print(f"✅ 转换成功: {input_file} → {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        if debug:
            import traceback
            traceback.print_exc()
        return False


def extract_image_references(markdown_text: str) -> List[str]:
    """提取Markdown中的图片引用"""
    # 匹配Markdown图片语法 ![alt](src)
    pattern = r'!\[.*?\]\((.*?)\)'
    matches = re.findall(pattern, markdown_text)
    
    # 匹配HTML img标签
    html_pattern = r'<img.*?src=["\'](.*?)["\'].*?>'
    html_matches = re.findall(html_pattern, markdown_text, re.IGNORECASE)
    
    return list(set(matches + html_matches))


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='带图片处理的Markdown转Word转换器')
    parser.add_argument('--input', '-i', required=True, help='输入Markdown文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出Word文件路径')
    parser.add_argument('--template', '-t', help='Word模板文件路径')
    parser.add_argument('--image-dir', '-d', help='图片目录路径')
    parser.add_argument('--image-width', type=int, help='图片最大宽度（像素）')
    parser.add_argument('--image-height', type=int, help='图片最大高度（像素）')
    parser.add_argument('--image-quality', type=int, default=85, help='图片质量（1-100，默认85）')
    parser.add_argument('--add-captions', action='store_true', help='添加图片标题')
    parser.add_argument('--list-images', action='store_true', help='列出图片引用')
    parser.add_argument('--debug', action='store_true', help='启用调试模式')
    
    args = parser.parse_args()
    
    if args.list_images:
        # 列出图片引用
        try:
            with open(args.input, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            images = extract_image_references(markdown_text)
            print(f"📷 找到 {len(images)} 个图片引用:")
            for i, img in enumerate(images, 1):
                print(f"  {i}. {img}")
            
            # 检查图片是否存在
            if args.image_dir:
                print(f"\n🔍 在目录 {args.image_dir} 中查找图片:")
                for i, img in enumerate(images, 1):
                    potential_path = os.path.join(args.image_dir, img)
                    if os.path.exists(potential_path):
                        print(f"  {i}. ✅ {img} → {potential_path}")
                    else:
                        print(f"  {i}. ❌ {img} (未找到)")
            
            return
        except Exception as e:
            print(f"❌ 列出图片失败: {e}")
            sys.exit(1)
    
    # 确保输出目录存在
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    # 执行转换
    success = convert_with_images(
        input_file=args.input,
        output_file=args.output,
        template=args.template,
        image_dir=args.image_dir,
        max_width=args.image_width,
        max_height=args.image_height,
        image_quality=args.image_quality,
        add_captions=args.add_captions,
        debug=args.debug
    )
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
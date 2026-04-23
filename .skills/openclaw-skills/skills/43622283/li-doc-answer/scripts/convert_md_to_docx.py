#!/usr/bin/env python3
"""
Markdown 转 Word 文档工具
支持 .md → .docx 转换

作者：北京老李
版本：2.0.0
"""

import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
import re

def parse_markdown(content):
    """
    解析 Markdown 内容
    
    Args:
        content: Markdown 文本
    
    Returns:
        list: 解析后的段落列表
    """
    paragraphs = []
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # 跳过空行
        if not line.strip():
            paragraphs.append({'type': 'empty', 'text': ''})
            continue
        
        # 标题
        if line.startswith('# '):
            paragraphs.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('## '):
            paragraphs.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('### '):
            paragraphs.append({'type': 'h3', 'text': line[4:].strip()})
        
        # 列表
        elif line.startswith('- ') or line.startswith('* '):
            paragraphs.append({'type': 'list', 'text': line[2:].strip()})
        elif line.startswith('1. '):
            paragraphs.append({'type': 'list', 'text': line[3:].strip()})
        
        # 引用
        elif line.startswith('> '):
            paragraphs.append({'type': 'quote', 'text': line[2:].strip()})
        
        # 代码块
        elif line.startswith('```'):
            paragraphs.append({'type': 'code', 'text': line[3:].strip()})
        
        # 普通段落
        else:
            paragraphs.append({'type': 'paragraph', 'text': line.strip()})
    
    return paragraphs

def convert_md_to_docx(source_md, output_docx):
    """
    转换 Markdown 到 Word
    
    Args:
        source_md: 输入 .md 文件路径
        output_docx: 输出 .docx 文件路径
    
    Returns:
        bool: 是否成功
    """
    # 检查文件
    if not os.path.exists(source_md):
        print(f"✗ 文件不存在：{source_md}")
        return False
    
    # 读取 Markdown
    with open(source_md, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print(f"读取：{source_md} ({len(content)} 字符)")
    
    # 解析
    paragraphs = parse_markdown(content)
    
    # 创建文档
    doc = Document()
    
    # 添加内容
    for para in paragraphs:
        p_type = para['type']
        text = para['text']
        
        if p_type == 'empty':
            doc.add_paragraph()
        elif p_type == 'h1':
            doc.add_heading(text, level=1)
        elif p_type == 'h2':
            doc.add_heading(text, level=2)
        elif p_type == 'h3':
            doc.add_heading(text, level=3)
        elif p_type == 'list':
            p = doc.add_paragraph(text, style='List Bullet')
        elif p_type == 'quote':
            p = doc.add_paragraph(text)
            p.italic = True
        elif p_type == 'code':
            p = doc.add_paragraph(text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(10)
    
    # 保存
    doc.save(output_docx)
    print(f"✓ 已保存：{output_docx}")
    
    return True

def main():
    """主函数"""
    print("="*60)
    print("Li_doc_answer - Markdown 转 Word")
    print("作者：北京老李")
    print("版本：2.0.0")
    print("="*60)
    
    if len(sys.argv) < 3:
        print("\n使用方法:")
        print("  python3 convert_md_to_docx.py <输入.md> <输出.docx>")
        sys.exit(1)
    
    source_md = sys.argv[1]
    output_docx = sys.argv[2]
    
    success = convert_md_to_docx(source_md, output_docx)
    
    if success:
        print("\n✓ 转换完成!")
    else:
        print("\n✗ 转换失败")
        sys.exit(1)

if __name__ == '__main__':
    main()

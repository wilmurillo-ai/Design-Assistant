#!/usr/bin/env python3
"""
文档完整处理工具
综合处理 doc/docx 文档

作者：北京老李
版本：2.0.0
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mammoth

def read_docx(filepath):
    """读取 docx 文件"""
    with open(filepath, 'rb') as f:
        result = mammoth.extract_raw_text(f)
        return result.value

def process_complete(input_file, output_file=None):
    """
    完整处理文档
    
    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径（可选）
    
    Returns:
        bool: 是否成功
    """
    # 生成输出文件名
    if not output_file:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}_完整版.docx"
    
    print(f"读取：{input_file}")
    if input_file.endswith('.docx'):
        content = read_docx(input_file)
    else:
        print("⚠️  仅支持 .docx 格式")
        return False
    
    # 创建新文档
    doc = Document()
    
    # 标题
    title = doc.add_heading(os.path.basename(input_file) + '（完整版）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 说明
    doc.add_paragraph('处理工具：Li_doc_answer v2.0.0')
    doc.add_paragraph('作者：北京老李')
    doc.add_paragraph()
    
    # 内容
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(10)
    
    # 保存
    doc.save(output_file)
    print(f"✓ 已保存：{output_file}")
    
    return True

def main():
    """主函数"""
    print("="*60)
    print("Li_doc_answer - 完整处理工具")
    print("作者：北京老李")
    print("版本：2.0.0")
    print("="*60)
    
    if len(sys.argv) < 2:
        print("\n使用方法:")
        print("  python3 complete_all_answers.py <输入文件> [输出文件]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"✗ 文件不存在：{input_file}")
        sys.exit(1)
    
    success = process_complete(input_file, output_file)
    
    if success:
        print("\n✓ 完成!")
    else:
        print("\n✗ 失败")
        sys.exit(1)

if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
通用文档答案生成器
支持任意 doc/docx 文档处理

作者：北京老李
版本：2.0.0
"""

import mammoth
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
import subprocess

def read_docx(filepath):
    """读取 docx 文件内容"""
    with open(filepath, 'rb') as f:
        result = mammoth.extract_raw_text(f)
        return result.value

def read_doc(filepath):
    """读取 doc 文件内容（需要 antiword）"""
    try:
        result = subprocess.run(['antiword', filepath], capture_output=True, text=True)
        return result.stdout
    except FileNotFoundError:
        print("⚠️  未找到 antiword，请安装：sudo apt-get install antiword")
        return None

def detect_file_type(filepath):
    """检测文件类型"""
    if filepath.endswith('.docx'):
        return 'docx'
    elif filepath.endswith('.doc'):
        return 'doc'
    else:
        return 'unknown'

def process_document(input_file, output_file=None, add_answers=True):
    """
    处理文档
    
    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径（可选）
        add_answers: 是否添加答案区域
    
    Returns:
        bool: 处理是否成功
    """
    # 检测文件类型
    file_type = detect_file_type(input_file)
    
    if file_type == 'unknown':
        print(f"⚠️  不支持的文件格式：{input_file}")
        return False
    
    # 读取内容
    print(f"读取文件：{input_file}")
    if file_type == 'docx':
        content = read_docx(input_file)
    else:
        content = read_doc(input_file)
    
    if not content:
        print("✗ 读取失败")
        return False
    
    print(f"✓ 读取成功 ({len(content)} 字符)")
    
    # 生成输出文件名
    if not output_file:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}_处理版.docx"
    
    # 创建新文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('文档处理结果', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph(f'原文档：{os.path.basename(input_file)}')
    doc.add_paragraph(f'处理时间：{os.path.basename(output_file)}')
    doc.add_paragraph('说明：本文档由 Li_doc_answer 技能处理生成')
    doc.add_paragraph()
    
    # 添加原文档内容
    doc.add_heading('原文档内容', level=1)
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)
    
    # 可选：添加答案区域
    if add_answers:
        doc.add_page_break()
        doc.add_heading('参考答案区域', level=1)
        doc.add_paragraph('在此处添加参考答案或备注...')
    
    # 保存
    doc.save(output_file)
    print(f"✓ 已保存：{output_file}")
    
    return True

def main():
    """主函数"""
    print("="*60)
    print("Li_doc_answer - 通用文档处理工具")
    print("作者：北京老李")
    print("版本：2.0.0")
    print("="*60)
    
    # 获取输入文件
    if len(sys.argv) < 2:
        print("\n使用方法:")
        print("  python3 generate_answers.py <输入文件> [输出文件]")
        print("\n示例:")
        print("  python3 generate_answers.py 题库.doc")
        print("  python3 generate_answers.py 题库.doc 输出.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 检查文件是否存在
    if not os.path.exists(input_file):
        print(f"✗ 文件不存在：{input_file}")
        sys.exit(1)
    
    # 处理文档
    success = process_document(input_file, output_file)
    
    if success:
        print("\n✓ 处理完成!")
    else:
        print("\n✗ 处理失败")
        sys.exit(1)

if __name__ == '__main__':
    main()

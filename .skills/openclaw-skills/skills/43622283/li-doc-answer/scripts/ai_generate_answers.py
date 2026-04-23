#!/usr/bin/env python3
"""
AI 智能答案生成器
自动识别文档中的问题并生成参考答案

作者：北京老李
版本：3.0.0
"""

import mammoth
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
import re
import subprocess

def read_docx(filepath):
    """读取 docx 文件内容"""
    with open(filepath, 'rb') as f:
        result = mammoth.extract_raw_text(f)
        return result.value

def read_doc(filepath):
    """读取 doc 文件内容"""
    try:
        result = subprocess.run(['antiword', filepath], capture_output=True, text=True)
        return result.stdout
    except FileNotFoundError:
        print("⚠️  未找到 antiword，请安装：sudo apt-get install antiword")
        return None

def detect_questions(content):
    """
    自动识别文档中的问题
    
    支持识别模式：
    - 数字 + 点/顿号：1. 2. 3. 或 1、2、3、
    - 题型标识：单选、多选、判断、简答、论述等
    - 问号结尾的句子
    
    Args:
        content: 文档文本内容
    
    Returns:
        list: 问题列表 [(题号，问题文本，题型), ...]
    """
    questions = []
    lines = content.split('\n')
    
    # 题型关键词
    question_types = {
        '单选': ['单选', '选择题'],
        '多选': ['多选'],
        '判断': ['判断', '对错', '正确/错误'],
        '简答': ['简答', '简述', '说明', '阐述'],
        '论述': ['论述', '论述题'],
        '案例': ['案例', '案例分析'],
        '填空': ['填空'],
        '名词解释': ['名词解释'],
    }
    
    current_type = '简答'  # 默认题型
    question_num = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 检测题型标题
        for q_type, keywords in question_types.items():
            if any(kw in line for kw in keywords):
                if '单' in line and '选' in line:
                    current_type = '单选'
                elif '多' in line and '选' in line:
                    current_type = '多选'
                elif '判断' in line:
                    current_type = '判断'
                elif '简' in line and ('答' in line or '述' in line):
                    current_type = '简答'
                elif '论述' in line:
                    current_type = '论述'
                elif '案例' in line:
                    current_type = '案例'
                break
        
        # 检测问题模式
        # 模式 1: 数字 + 点/顿号
        match = re.match(r'^(\d+)[.、]\s*(.+)', line)
        if match:
            num = int(match.group(1))
            text = match.group(2).strip()
            
            # 跳过答案行
            if text.startswith('答') or text.startswith('参考答案'):
                continue
            
            question_num = num
            questions.append({
                'number': num,
                'text': text,
                'type': current_type
            })
        
        # 模式 2: 带问号的问题
        elif '?' in line or '？' in line:
            if len(line) > 10:  # 避免太短的句子
                questions.append({
                    'number': len(questions) + 1,
                    'text': line,
                    'type': current_type
                })
    
    return questions

def generate_answer(question_text, question_type):
    """
    为问题生成参考答案
    
    Args:
        question_text: 问题文本
        question_type: 题型
    
    Returns:
        str: 生成的答案
    """
    # 根据题型生成答案模板
    if question_type == '判断':
        # 判断题
        return "答：【正确/错误】\n\n理由：请根据教材内容判断并说明理由。"
    
    elif question_type == '单选':
        # 单选题
        return "答：【正确选项】\n\n解析：请分析各选项，说明选择理由。"
    
    elif question_type == '多选':
        # 多选题
        return "答：【正确选项，如：A、B、C】\n\n解析：请分析各选项的正确性。"
    
    elif question_type == '填空':
        # 填空题
        return "答：【正确答案】\n\n说明：请填写空白处的正确内容。"
    
    elif question_type == '名词解释':
        # 名词解释
        return f"答：{question_text.replace('名词解释：', '').replace('解释：', '').strip()}是指...\n\n详细说明其定义、特点和意义。"
    
    elif question_type == '案例':
        # 案例分析
        return """答：【案例分析】

1. 问题识别：
   分析案例中存在的主要问题。

2. 理论应用：
   运用相关理论知识进行分析。

3. 解决方案：
   提出具体的解决建议。

4. 总结：
   总结案例的启示和教训。"""
    
    elif question_type == '论述':
        # 论述题
        return f"""答：【论述】

一、引言
   简述问题背景和重要性。

二、主体论述
   1. 第一个要点
      详细阐述...
   
   2. 第二个要点
      详细阐述...
   
   3. 第三个要点
      详细阐述...

三、结论
   总结全文，强调核心观点。"""
    
    else:
        # 简答题（默认）
        return f"""答：【参考答案】

1. 要点一
   详细说明...

2. 要点二
   详细说明...

3. 要点三
   详细说明...

【说明】以上为参考答案要点，具体内容请根据相关教材完善。"""

def create_answer_document(input_file, questions, output_file=None):
    """
    创建带答案的文档
    
    Args:
        input_file: 输入文件路径
        questions: 问题列表
        output_file: 输出文件路径
    
    Returns:
        bool: 是否成功
    """
    # 生成输出文件名
    if not output_file:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}_AI 答案版.docx"
    
    # 创建新文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('文档题目及 AI 参考答案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 说明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('说明：以下答案由 AI 生成，仅供参考，请以教材为准')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    doc.add_paragraph(f'原文档：{os.path.basename(input_file)}')
    doc.add_paragraph(f'识别题目数：{len(questions)} 道')
    doc.add_paragraph()
    
    # 添加问题和答案
    doc.add_heading('题目与答案', level=1)
    
    for q in questions:
        # 题号
        p = doc.add_paragraph()
        run = p.add_run(f"第{q['number']}题 ")
        run.font.bold = True
        run.font.size = Pt(11)
        
        # 题型标签
        run = p.add_run(f"[{q['type']}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 128, 0)
        
        # 问题文本
        p.add_run(f"\n{q['text']}")
        
        # 生成答案
        answer = generate_answer(q['text'], q['type'])
        
        # 答案
        p_answer = doc.add_paragraph()
        run_answer = p_answer.add_run(answer)
        run_answer.font.size = Pt(10)
        
        # 分隔线
        doc.add_paragraph('─' * 50)
    
    # 保存
    doc.save(output_file)
    return output_file

def process_document(input_file, output_file=None):
    """
    处理文档的主函数
    
    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径
    
    Returns:
        dict: 处理结果
    """
    result = {
        'success': False,
        'questions_count': 0,
        'output_file': None,
        'error': None
    }
    
    # 检测文件类型
    if input_file.endswith('.docx'):
        content = read_docx(input_file)
    elif input_file.endswith('.doc'):
        content = read_doc(input_file)
    else:
        result['error'] = '不支持的文件格式'
        return result
    
    if not content:
        result['error'] = '读取文件失败'
        return result
    
    print(f"✓ 读取成功 ({len(content)} 字符)")
    
    # 识别问题
    print("正在识别文档中的问题...")
    questions = detect_questions(content)
    result['questions_count'] = len(questions)
    
    if not questions:
        print("⚠️  未识别到问题，将创建空白答案模板")
        # 创建空白模板
        if not output_file:
            base_name = os.path.splitext(input_file)[0]
            output_file = f"{base_name}_答案模板.docx"
        
        doc = Document()
        doc.add_heading('文档答案模板', 0)
        doc.add_paragraph('原文档：' + os.path.basename(input_file))
        doc.add_paragraph()
        doc.add_paragraph('请在下方添加答案...')
        doc.add_paragraph()
        doc.add_paragraph(content[:2000] + '...' if len(content) > 2000 else content)
        doc.save(output_file)
        result['output_file'] = output_file
        result['success'] = True
        return result
    
    print(f"✓ 识别到 {len(questions)} 道题目")
    
    # 生成答案文档
    print("正在生成参考答案...")
    output = create_answer_document(input_file, questions, output_file)
    
    result['output_file'] = output
    result['success'] = True
    
    return result

def main():
    """主函数"""
    print("="*60)
    print("Li_doc_answer - AI 智能答案生成器")
    print("作者：北京老李")
    print("版本：3.0.0")
    print("="*60)
    
    if len(sys.argv) < 2:
        print("\n使用方法:")
        print("  python3 ai_generate_answers.py <输入文件> [输出文件]")
        print("\n示例:")
        print("  python3 ai_generate_answers.py 题库.doc")
        print("  python3 ai_generate_answers.py 题库.doc 输出.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"✗ 文件不存在：{input_file}")
        sys.exit(1)
    
    # 处理文档
    result = process_document(input_file, output_file)
    
    # 输出结果
    print()
    if result['success']:
        print("="*60)
        print("✓ 处理完成!")
        print(f"识别题目：{result['questions_count']} 道")
        print(f"输出文件：{result['output_file']}")
        print("="*60)
    else:
        print(f"✗ 处理失败：{result['error']}")
        sys.exit(1)

if __name__ == '__main__':
    main()

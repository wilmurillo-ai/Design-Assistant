#!/usr/bin/env python3
"""
文档校验工具
检查 doc/docx 文档内容和格式

作者：北京老李
版本：2.0.0
"""

import os
import sys
import docx
from docx import Document
import mammoth

def read_docx(filepath):
    """读取 docx 文件"""
    with open(filepath, 'rb') as f:
        result = mammoth.extract_raw_text(f)
        return result.value

def check_document(file_path):
    """
    校验文档
    
    Args:
        file_path: 文件路径
    
    Returns:
        dict: 校验结果
    """
    result = {
        'exists': False,
        'readable': False,
        'paragraphs': 0,
        'characters': 0,
        'tables': 0,
        'errors': []
    }
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        result['errors'].append('文件不存在')
        return result
    
    result['exists'] = True
    
    # 读取文档
    try:
        doc = Document(file_path)
        result['readable'] = True
        result['paragraphs'] = len(doc.paragraphs)
        
        # 统计字符数
        total_chars = 0
        for para in doc.paragraphs:
            total_chars += len(para.text)
        result['characters'] = total_chars
        
        # 统计表格数
        result['tables'] = len(doc.tables)
        
    except Exception as e:
        result['errors'].append(f'读取错误：{str(e)}')
    
    return result

def print_report(file_path, result):
    """打印校验报告"""
    print(f"\n文档校验报告")
    print("="*60)
    print(f"文件：{os.path.basename(file_path)}")
    print(f"路径：{file_path}")
    print()
    
    if not result['exists']:
        print("❌ 文件不存在")
        return
    
    print(f"✅ 文件存在")
    print(f"{'✅' if result['readable'] else '❌'} 可读")
    
    if result['readable']:
        print(f"📊 段落数：{result['paragraphs']}")
        print(f"📝 字符数：{result['characters']}")
        print(f"📋 表格数：{result['tables']}")
    
    if result['errors']:
        print("\n⚠️  错误信息:")
        for error in result['errors']:
            print(f"  - {error}")
    
    print("="*60)

def main():
    """主函数"""
    print("="*60)
    print("Li_doc_answer - 文档校验工具")
    print("作者：北京老李")
    print("版本：2.0.0")
    print("="*60)
    
    if len(sys.argv) < 2:
        print("\n使用方法:")
        print("  python3 check_answers.py <文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = check_document(file_path)
    print_report(file_path, result)

if __name__ == '__main__':
    main()

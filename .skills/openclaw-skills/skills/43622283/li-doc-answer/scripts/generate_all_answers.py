#!/usr/bin/env python3
"""
批量文档处理工具
支持任意 doc/docx 文档批量处理

作者：北京老李
版本：2.0.0
"""

import mammoth
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
import subprocess

def read_docx(filepath):
    """读取 docx 文件"""
    with open(filepath, 'rb') as f:
        result = mammoth.extract_raw_text(f)
        return result.value

def read_doc(filepath):
    """读取 doc 文件"""
    try:
        result = subprocess.run(['antiword', filepath], capture_output=True, text=True)
        return result.stdout
    except FileNotFoundError:
        print("⚠️  未找到 antiword")
        return None

def create_processed_docx(title, content, output_path):
    """创建处理后的 docx 文件"""
    doc = Document()
    
    # 标题
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 说明
    doc.add_paragraph('说明：本文档由 Li_doc_answer 技能处理生成')
    doc.add_paragraph('作者：北京老李')
    doc.add_paragraph()
    
    # 内容
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10)
    
    doc.save(output_path)
    return len(content)

def process_directory(source_dir):
    """
    批量处理目录中的所有文档
    
    Args:
        source_dir: 源目录路径
    """
    print(f"\n扫描目录：{source_dir}")
    
    # 支持的文件扩展名
    supported_extensions = ['.doc', '.docx']
    
    # 获取所有支持的文件
    files_to_process = []
    for filename in os.listdir(source_dir):
        ext = os.path.splitext(filename)[1].lower()
        if ext in supported_extensions:
            files_to_process.append(filename)
    
    if not files_to_process:
        print("⚠️  未找到支持的文档文件 (.doc, .docx)")
        return []
    
    print(f"找到 {len(files_to_process)} 个文件")
    
    results = []
    for filename in files_to_process:
        filepath = os.path.join(source_dir, filename)
        
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}_处理版.docx"
        output_path = os.path.join(source_dir, output_filename)
        
        # 跳过已处理的文件
        if os.path.exists(output_path):
            print(f"\n⊘ 跳过（已存在）: {filename}")
            continue
        
        print(f"\n处理：{filename}")
        
        # 读取内容
        if filename.endswith('.docx'):
            content = read_docx(filepath)
        else:
            content = read_doc(filepath)
        
        if not content:
            print(f"⚠️  读取失败，跳过")
            continue
        
        # 创建处理后的文档
        title = f"{base_name}（处理版）"
        content_len = create_processed_docx(title, content, output_path)
        
        print(f"✓ 已保存：{output_filename} ({content_len} 字符)")
        results.append((output_filename, content_len))
    
    return results

def main():
    """主函数"""
    print("="*60)
    print("Li_doc_answer - 批量文档处理工具")
    print("作者：北京老李")
    print("版本：2.0.0")
    print("="*60)
    
    # 获取目录路径
    if len(sys.argv) > 1:
        source_dir = sys.argv[1]
    else:
        # 默认使用 data 目录
        script_dir = os.path.dirname(os.path.abspath(__file__))
        source_dir = os.path.join(script_dir, '..', 'data')
    
    # 检查目录是否存在
    if not os.path.exists(source_dir):
        print(f"\n创建目录：{source_dir}")
        os.makedirs(source_dir)
        print("请将待处理的文档放入此目录后重新运行")
        sys.exit(0)
    
    # 处理目录
    results = process_directory(source_dir)
    
    # 输出统计
    print("\n" + "="*60)
    print(f"处理完成！")
    print(f"成功处理：{len(results)} 个文件")
    print("="*60)

if __name__ == '__main__':
    main()

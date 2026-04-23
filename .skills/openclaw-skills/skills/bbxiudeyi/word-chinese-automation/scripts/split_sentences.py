# -*- coding: utf-8 -*-
"""
句子拆分脚本
按句号（。）、感叹号（！）、问号（？）拆分句子
保留每句的位置信息（段落索引、句子索引）
"""

import sys
import re
import json
import io
from pathlib import Path

# 确保 stdout 支持 UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("运行: pip install python-docx")
    sys.exit(1)


def split_sentences(text):
    """
    按句号、感叹号、问号拆分句子
    返回句子列表
    """
    if not text or not text.strip():
        return []
    
    # 正则：匹配 。！？ 三种句末标点
    # 使用捕获组保留分隔符
    pattern = r'([。！？])'
    parts = re.split(pattern, text)
    
    sentences = []
    current_sentence = ""
    
    for i, part in enumerate(parts):
        if part in '。！？':
            # 遇到句末标点，合并前面的内容并加上这个标点
            if current_sentence.strip():
                sentences.append(current_sentence.strip() + part)
            current_sentence = ""
        else:
            current_sentence += part
    
    # 处理最后一段没有句末标点的内容
    if current_sentence.strip():
        sentences.append(current_sentence.strip())
    
    return sentences


def extract_text_from_textboxes(doc):
    """
    从 Word 文档的文本框中提取文本
    python-docx 默认不读取文本框内容，需要从 XML 层提取
    """
    textbox_texts = []
    
    # Word 命名空间
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }
    
    # 方法1: 查找 w:txbxContent (传统文本框)
    for txbx_content in doc.element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
        for para in txbx_content.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            text_parts = []
            for t in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    text_parts.append(t.text)
            if text_parts:
                textbox_texts.append(''.join(text_parts))
    
    # 方法2: 查找 wps:txbx (现代形状文本框)
    for txbx in doc.element.iter('{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'):
        for para in txbx.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            text_parts = []
            for t in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    text_parts.append(t.text)
            if text_parts:
                textbox_texts.append(''.join(text_parts))
    
    return textbox_texts


def escape_json_string(text):
    """
    转义 JSON 字符串中的特殊字符
    确保中英文引号和其他特殊字符被正确处理
    """
    if not text:
        return text
    
    # 替换各种引号和其他需要转义的字符
    replacements = {
        '"': '\\"',      # 双引号
        '\n': '\\n',    # 换行
        '\r': '\\r',    # 回车
        '\t': '\\t',    # 制表符
        '\\': '\\\\',   # 反斜杠
    }
    
    # 先处理反斜杠，避免重复转义
    result = text.replace('\\', '\\\\')
    
    # 处理其他字符
    for char, escaped in replacements.items():
        if char != '\\':  # 反斜杠已经处理过了
            result = result.replace(char, escaped)
    
    return result


def extract_sentences_from_docx(doc_path):
    """
    从 Word 文档提取所有句子（包括文本框）
    返回结构化数据
    """
    doc = Document(doc_path)
    results = []
    
    # 读取主体段落
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        sentences = split_sentences(text)
        
        for sent_idx, sentence in enumerate(sentences):
            results.append({
                "source": "paragraph",
                "paragraph_index": para_idx + 1,  # 1-based
                "sentence_index": sent_idx + 1,     # 1-based
                "full_text": text,                  # 所属段落完整文本
                "sentence": sentence                 # 拆分后的句子
            })
    
    # 读取文本框内容
    textbox_texts = extract_text_from_textboxes(doc)
    for box_idx, text in enumerate(textbox_texts):
        text = text.strip()
        if not text:
            continue
        
        sentences = split_sentences(text)
        
        for sent_idx, sentence in enumerate(sentences):
            results.append({
                "source": "textbox",
                "paragraph_index": box_idx + 1,     # 文本框索引
                "sentence_index": sent_idx + 1,
                "full_text": text,
                "sentence": sentence
            })
    
    # 读取表格内容
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                
                sentences = split_sentences(cell_text)
                
                for sent_idx, sentence in enumerate(sentences):
                    results.append({
                        "source": "table",
                        "table_index": tbl_idx + 1,       # 表格索引 (1-based)
                        "row_index": row_idx + 1,         # 行索引 (1-based)
                        "cell_index": cell_idx + 1,       # 单元格索引 (1-based)
                        "sentence_index": sent_idx + 1,   # 句子索引 (1-based)
                        "full_text": cell_text,
                        "sentence": sentence
                    })
    
    return results


def main():
    if len(sys.argv) < 2:
        print("用法: python split_sentences.py <文档路径> [输出JSON文件]")
        print("示例: python split_sentences.py test.docx")
        print("示例: python split_sentences.py test.docx sentences.json")
        sys.exit(1)
    
    doc_path = sys.argv[1]
    
    # 检查文件是否存在
    if not Path(doc_path).exists():
        print(f"错误：文件不存在: {doc_path}")
        sys.exit(1)
    
    # 如果没有指定输出路径，默认输出到原文档同目录
    if len(sys.argv) > 2:
        output_path = sys.argv[2]
    else:
        input_file = Path(doc_path)
        output_path = str(input_file.parent / f"{input_file.stem}_sentences.json")
    
    print(f"正在处理: {doc_path}")
    print("-" * 50)
    
    try:
        results = extract_sentences_from_docx(doc_path)
        
        # 打印统计
        para_count = len([r for r in results if r.get("source") == "paragraph"])
        textbox_count = len([r for r in results if r.get("source") == "textbox"])
        table_count = len([r for r in results if r.get("source") == "table"])
        sent_count = len(results)
        print(f"段落数: {para_count}")
        print(f"文本框句子数: {textbox_count}")
        print(f"表格句子数: {table_count}")
        print(f"总句子数: {sent_count}")
        print("-" * 50)
        
        # 打印每句
        for item in results:
            source = item.get("source", "paragraph")
            if source == "table":
                location = f"表格{item['table_index']}-行{item['row_index']}-单元格{item['cell_index']}"
            elif source == "textbox":
                location = f"文本框{item['paragraph_index']}"
            else:
                location = f"段落{item['paragraph_index']}"
            
            print(f"【{location} - 句子{item['sentence_index']}】")
            print(f"  {item['sentence']}")
            print()
        
        # 保存为 JSON - 使用更健壮的方式
        # 先序列化为 JSON 字符串，确保格式正确
        json_str = json.dumps(results, ensure_ascii=False, indent=2)
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(json_str)
        
        # 验证写入的 JSON 可以正确读取
        with open(output_path, 'r', encoding='utf-8') as f:
            json.load(f)  # 如果格式有问题会抛出异常
        
        print(f"已保存到: {output_path}")
        
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()

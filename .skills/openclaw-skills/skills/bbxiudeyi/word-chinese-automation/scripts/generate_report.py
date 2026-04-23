#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
校对报告生成脚本

生成 Word 格式的校对报告，同时输出修正数据供 apply_corrections.py 使用

用法:
    # 方式1：使用校对数据 JSON（原有方式）
    python generate_report.py <原文档路径> --data '<校对数据JSON>'
    python generate_report.py <原文档路径> --data-file <JSON文件>
    
    # 方式2：使用 sentences.json（推荐，与 split_sentences.py 配合）
    python generate_report.py <原文档路径> --sentences <sentences.json>

示例:
    python generate_report.py report.docx --data-file report_data.json
    python generate_report.py report.docx --sentences report_sentences.json
    python generate_report.py report.docx --sentences report_sentences.json --output E:/output/校对报告.docx
"""

import argparse
import json
import sys
import io
from pathlib import Path
from datetime import datetime

# 确保 stdout 支持 UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name='微软雅黑', font_size=10.5, bold=False, color=None):
    """设置文字格式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_table_header(table, headers, color="4472C4"):
    """添加表格表头"""
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, color)
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, font_size=9, bold=True, color=(255, 255, 255))


def add_table_row(table, cells_data):
    """添加表格数据行"""
    row = table.add_row()
    for i, cell_text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(cell_text) if cell_text else ""
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, font_size=9)
    return row


def get_location_str(item: dict) -> str:
    """根据句子数据生成位置字符串"""
    source = item.get("source", "paragraph")
    if source == "table":
        return f"表格{item.get('table_index', '?')}-行{item.get('row_index', '?')}-单元格{item.get('cell_index', '?')}"
    elif source == "textbox":
        return f"文本框{item.get('paragraph_index', '?')}"
    else:
        return f"段落{item.get('paragraph_index', '?')}"


def convert_sentences_to_report_data(sentences_data: list) -> dict:
    """
    将 sentences.json 格式转换为校对报告数据格式
    
    sentences.json 中有 correction 字段的句子会被分类为不同问题类型
    需要根据 issue_type 字段判断问题类型（标点/错别字/语病）
    """
    punctuation = []
    typos = []
    grammar = []
    
    for item in sentences_data:
        if not item.get("correction"):
            continue
        
        issue_type = item.get("issue_type", "grammar")  # 默认为语病
        location = get_location_str(item)
        
        issue_entry = {
            "location": f"{location}-句{item.get('sentence_index', '?')}",
            "original": item.get("sentence", ""),
            "corrected": item.get("correction", ""),
            "severity": item.get("severity", "中")
        }
        
        if issue_type == "punctuation":
            issue_entry["problem"] = item.get("issue_desc", "标点符号问题")
            punctuation.append(issue_entry)
        elif issue_type == "typo":
            issue_entry["error"] = item.get("error_char", "")
            issue_entry["correct"] = item.get("correct_char", "")
            typos.append(issue_entry)
        else:  # grammar
            issue_entry["problem"] = item.get("issue_desc", "语病问题")
            grammar.append(issue_entry)
    
    return {
        "punctuation": punctuation,
        "typos": typos,
        "grammar": grammar,
        "summary": {
            "punctuation": len(punctuation),
            "typos": len(typos),
            "grammar": len(grammar),
            "total": len(punctuation) + len(typos) + len(grammar)
        }
    }


def extract_corrections_list(sentences_data: list) -> list:
    """
    从 sentences.json 提取修正列表，供 apply_corrections.py 使用
    """
    corrections = []
    for item in sentences_data:
        if item.get("correction"):
            corrections.append({
                "original": item.get("sentence", ""),
                "corrected": item.get("correction", "")
            })
    return corrections


def generate_report(input_path: str, report_data: dict, output_path: str = None) -> str:
    """
    生成 Word 格式的校对报告
    
    Args:
        input_path: 原文档路径
        report_data: 校对报告数据
        output_path: 输出路径（可选）
    
    Returns:
        输出文件路径
    """
    # 创建新文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ===== 标题 =====
    title = doc.add_heading('校对报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 基本信息 =====
    doc.add_paragraph()
    
    info_para = doc.add_paragraph()
    info_para.add_run('原文档：').bold = True
    info_para.add_run(Path(input_path).name)
    
    time_para = doc.add_paragraph()
    time_para.add_run('检查时间：').bold = True
    time_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    
    # ===== 问题统计 =====
    doc.add_paragraph()
    doc.add_heading('问题统计', level=1)
    
    summary = report_data.get('summary', {})
    
    # 统计表格
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    add_table_header(summary_table, ["问题类型", "数量"])
    
    stats = [
        ("标点符号问题", summary.get('punctuation', 0)),
        ("错别字", summary.get('typos', 0)),
        ("语病问题", summary.get('grammar', 0)),
    ]
    
    for label, count in stats:
        add_table_row(summary_table, [label, str(count)])
    
    # 总计行
    total_row = summary_table.add_row()
    total_row.cells[0].text = "总计"
    total_row.cells[1].text = str(summary.get('total', 0))
    for cell in total_row.cells:
        set_cell_shading(cell, "F2F2F2")
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, font_size=10, bold=True)
    
    doc.add_paragraph()
    
    # ===== 标点符号问题 =====
    punctuation_issues = report_data.get('punctuation', [])
    doc.add_heading('标点符号问题', level=1)
    
    if punctuation_issues:
        punct_table = doc.add_table(rows=1, cols=5)
        punct_table.style = 'Table Grid'
        add_table_header(punct_table, ["位置", "原文", "问题", "严重程度", "修改后"])
        
        for issue in punctuation_issues:
            add_table_row(punct_table, [
                issue.get('location', ''),
                issue.get('original', ''),
                issue.get('problem', ''),
                issue.get('severity', ''),
                issue.get('corrected', '')
            ])
    else:
        no_issue = doc.add_paragraph("无问题")
        no_issue.runs[0].italic = True
        no_issue.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # ===== 错别字 =====
    typo_issues = report_data.get('typos', [])
    doc.add_heading('错别字', level=1)
    
    if typo_issues:
        typo_table = doc.add_table(rows=1, cols=6)
        typo_table.style = 'Table Grid'
        add_table_header(typo_table, ["位置", "原文", "错误", "正确", "严重程度", "修改后"])
        
        for issue in typo_issues:
            add_table_row(typo_table, [
                issue.get('location', ''),
                issue.get('original', ''),
                issue.get('error', ''),
                issue.get('correct', ''),
                issue.get('severity', ''),
                issue.get('corrected', '')
            ])
    else:
        no_issue = doc.add_paragraph("无问题")
        no_issue.runs[0].italic = True
        no_issue.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # ===== 语病问题 =====
    grammar_issues = report_data.get('grammar', [])
    doc.add_heading('语病问题', level=1)
    
    if grammar_issues:
        gram_table = doc.add_table(rows=1, cols=5)
        gram_table.style = 'Table Grid'
        add_table_header(gram_table, ["位置", "原文", "问题", "严重程度", "修改后"])
        
        for issue in grammar_issues:
            add_table_row(gram_table, [
                issue.get('location', ''),
                issue.get('original', ''),
                issue.get('problem', ''),
                issue.get('severity', ''),
                issue.get('corrected', '')
            ])
    else:
        no_issue = doc.add_paragraph("无问题")
        no_issue.runs[0].italic = True
        no_issue.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 确定输出路径
    if not output_path:
        input_file = Path(input_path)
        output_path = str(input_file.parent / f"{input_file.stem}_校对报告.docx")
    
    # 保存文档
    doc.save(output_path)
    
    return output_path


def main():
    parser = argparse.ArgumentParser(description="生成 Word 格式的校对报告")
    parser.add_argument("input", help="原 Word 文档路径")
    parser.add_argument("--data", help="校对数据 JSON 字符串")
    parser.add_argument("--data-file", help="校对数据 JSON 文件路径")
    parser.add_argument("--sentences", help="sentences.json 文件路径（由 split_sentences.py 生成）")
    parser.add_argument("--output", help="输出文件路径（可选，默认与原文档同目录）")
    
    args = parser.parse_args()
    
    # 检查文件是否存在
    if not Path(args.input).exists():
        print(f"错误: 文件不存在 - {args.input}")
        sys.exit(1)
    
    sentences_data = None
    report_data = None
    
    # 获取校对数据
    if args.sentences:
        # 方式1：从 sentences.json 读取
        try:
            with open(args.sentences, 'r', encoding='utf-8') as f:
                sentences_data = json.load(f)
            report_data = convert_sentences_to_report_data(sentences_data)
            print(f"已读取 sentences.json: {len(sentences_data)} 条句子")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"错误: 无法读取 sentences.json - {e}")
            sys.exit(1)
    elif args.data:
        # 方式2：从 JSON 字符串读取
        try:
            report_data = json.loads(args.data)
        except json.JSONDecodeError as e:
            print(f"错误: JSON 解析失败 - {e}")
            sys.exit(1)
    elif args.data_file:
        # 方式3：从数据文件读取
        try:
            with open(args.data_file, 'r', encoding='utf-8') as f:
                report_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"错误: 无法读取数据文件 - {e}")
            sys.exit(1)
    else:
        print("错误: 必须指定 --sentences、--data 或 --data-file 参数")
        sys.exit(1)
    
    # 生成报告
    output_path = generate_report(args.input, report_data, args.output)
    print(f"校对报告已生成: {output_path}")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
Word 文档修改应用脚本

支持三种模式：
1. 基于位置的精确修改（推荐）- 读取校对后的 sentences.json，按段落+句子索引定位
2. 修正文件模式（推荐）- 读取 generate_report.py 生成的 corrections.json
3. 文本替换模式（兼容旧版）- 全局查找替换

用法:
    # 模式1：基于位置（推荐）
    python apply_corrections.py <input_file> --json <sentences.json>
    
    # 模式2：修正文件（推荐，与 generate_report.py 配合）
    python apply_corrections.py <input_file> --corrections-file <corrections.json>
    
    # 模式3：文本替换（兼容）
    python apply_corrections.py <input_file> --corrections '<JSON>'

示例:
    python apply_corrections.py report.docx --json sentences.json
    python apply_corrections.py report.docx --corrections-file report_corrections.json
    python apply_corrections.py report.docx --corrections '[{"original": "按装", "corrected": "安装"}]'

输出:
    原文件_修改.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def find_sentence_in_paragraph(paragraph_text: str, sentence: str) -> tuple:
    """
    在段落文本中定位句子的起始和结束位置
    
    Returns:
        (start, end) 或 (-1, -1) 如果未找到
    """
    start = paragraph_text.find(sentence)
    if start >= 0:
        return start, start + len(sentence)
    return -1, -1


def apply_corrections_by_position(input_path: str, sentences_data: list) -> tuple:
    """
    基于位置信息应用修改到 Word 文档
    
    Args:
        input_path: 原 Word 文档路径
        sentences_data: 带位置信息的句子列表，包含 "correction" 字段
    
    Returns:
        (输出文件路径, 修改数量)
    """
    doc = Document(input_path)
    applied = 0
    
    # 过滤出有修改的句子
    corrections = [s for s in sentences_data if s.get("correction")]
    
    if not corrections:
        print("没有发现需要应用的修改")
        return None, 0
    
    # 按段落索引分组
    from collections import defaultdict
    by_paragraph = defaultdict(list)
    for item in corrections:
        by_paragraph[item["paragraph_index"]].append(item)
    
    # 遍历段落（Word 中段落索引从 0 开始，JSON 中从 1 开始）
    for para_idx, para_corrections in by_paragraph.items():
        # Word 段落索引 = JSON 段落索引 - 1
        word_para_idx = para_idx - 1
        
        if word_para_idx < 0 or word_para_idx >= len(doc.paragraphs):
            print(f"警告: 段落索引 {para_idx} 超出范围，跳过")
            continue
        
        para = doc.paragraphs[word_para_idx]
        original_text = para.text
        
        # 按句子索引排序，从后往前替换（避免位置偏移）
        para_corrections.sort(key=lambda x: x["sentence_index"], reverse=True)
        
        new_text = original_text
        for item in para_corrections:
            sentence = item["sentence"]
            correction = item["correction"]
            
            # 查找句子位置
            start, end = find_sentence_in_paragraph(new_text, sentence)
            if start >= 0:
                new_text = new_text[:start] + correction + new_text[end:]
                applied += 1
            else:
                print(f"警告: 段落{para_idx}-句{item['sentence_index']} 未找到，跳过")
        
        # 更新段落文本
        para.text = new_text
    
    # 生成输出文件名
    input_file = Path(input_path)
    output_path = str(input_file.parent / f"{input_file.stem}_修改{input_file.suffix}")
    
    # 保存文档
    doc.save(output_path)
    
    return output_path, applied


def apply_corrections_by_replace(input_path: str, corrections: list) -> tuple:
    """
    文本替换模式应用修改
    
    Args:
        input_path: 原 Word 文档路径
        corrections: 修改列表 [{"original": "...", "corrected": "..."}, ...]
    
    Returns:
        (输出文件路径, 修改数量)
    """
    doc = Document(input_path)
    applied = 0
    
    # 遍历所有段落
    for para in doc.paragraphs:
        for correction in corrections:
            original = correction.get("original", "")
            corrected = correction.get("corrected", "")
            
            if original and original in para.text:
                para.text = para.text.replace(original, corrected)
                applied += 1
    
    # 遍历表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for correction in corrections:
                        original = correction.get("original", "")
                        corrected = correction.get("corrected", "")
                        
                        if original and original in para.text:
                            para.text = para.text.replace(original, corrected)
                            applied += 1
    
    # 生成输出文件名
    input_file = Path(input_path)
    output_path = str(input_file.parent / f"{input_file.stem}_修改{input_file.suffix}")
    
    # 保存文档
    doc.save(output_path)
    
    return output_path, applied


def load_corrections_from_report_format(corrections_data: dict) -> list:
    """
    从 generate_report.py 生成的 corrections.json 格式加载修改列表
    
    支持两种格式：
    1. 报告格式 {"summary":..., "punctuation":..., "typos":..., "grammar":...}
    2. 简单列表格式 [{"original":..., "corrected":...}, ...]
    
    Args:
        corrections_data: 校对数据字典或列表
    
    Returns:
        修改列表 [{"original": "...", "corrected": "..."}, ...]
    """
    # 如果是列表格式，直接返回
    if isinstance(corrections_data, list):
        return corrections_data
    
    # 如果是字典格式，从各分类中提取
    corrections = []
    
    # 从标点符号问题中提取
    for item in corrections_data.get("punctuation", []):
        if item.get("original") and item.get("corrected"):
            corrections.append({
                "original": item["original"],
                "corrected": item["corrected"]
            })
    
    # 从错别字中提取
    for item in corrections_data.get("typos", []):
        if item.get("original") and item.get("corrected"):
            corrections.append({
                "original": item["original"],
                "corrected": item["corrected"]
            })
    
    # 从语病问题中提取
    for item in corrections_data.get("grammar", []):
        if item.get("original") and item.get("corrected"):
            corrections.append({
                "original": item["original"],
                "corrected": item["corrected"]
            })
    
    return corrections


def main():
    parser = argparse.ArgumentParser(description="应用修改到 Word 文档")
    parser.add_argument("input", help="输入 Word 文档路径")
    parser.add_argument("--json", help="带位置信息的 sentences.json 文件（模式1：基于位置）")
    parser.add_argument("--corrections-file", help="corrections.json 文件（模式2：修正文件，由 generate_report.py 生成）")
    parser.add_argument("--corrections", help="修改列表 JSON（模式3：文本替换）")
    
    args = parser.parse_args()
    
    # 检查参数
    if not args.json and not args.corrections and not args.corrections_file:
        print("错误: 必须指定 --json、--corrections-file 或 --corrections 参数")
        sys.exit(1)
    
    if args.json:
        # 模式1：基于位置
        try:
            with open(args.json, "r", encoding="utf-8") as f:
                sentences_data = json.load(f)
        except json.JSONDecodeError as e:
            print(f"错误: JSON 解析失败 - {e}")
            sys.exit(1)
        except FileNotFoundError:
            print(f"错误: 文件不存在 - {args.json}")
            sys.exit(1)
        
        output_path, applied = apply_corrections_by_position(args.input, sentences_data)
        
    elif args.corrections_file:
        # 模式2：修正文件（支持两种格式：报告格式或简单列表格式）
        try:
            with open(args.corrections_file, "r", encoding="utf-8") as f:
                corrections_data = json.load(f)
        except json.JSONDecodeError as e:
            print(f"错误: JSON 解析失败 - {e}")
            sys.exit(1)
        except FileNotFoundError:
            print(f"错误: 文件不存在 - {args.corrections_file}")
            sys.exit(1)
        
        # 支持两种格式：报告格式 {"punctuation":[], "typos":[], "grammar":[]} 或简单列表格式
        corrections = load_corrections_from_report_format(corrections_data)
        
        output_path, applied = apply_corrections_by_replace(args.input, corrections)
        
    else:
        # 模式3：文本替换
        try:
            corrections = json.loads(args.corrections)
        except json.JSONDecodeError as e:
            print(f"错误: JSON 解析失败 - {e}")
            sys.exit(1)
        
        output_path, applied = apply_corrections_by_replace(args.input, corrections)
    
    if output_path:
        print(f"已应用 {applied} 处修改")
        print(f"输出文件: {output_path}")


if __name__ == "__main__":
    main()

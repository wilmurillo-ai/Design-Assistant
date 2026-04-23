#!/usr/bin/env python3
"""
长截图智能切片工具 v7 - 精细平衡版
在目标位置附近小范围搜索，平衡切片高度和间隙质量

使用方法：
    python3 slice_processor.py <源图片路径> [输出目录]
"""

import os
import sys
from PIL import Image
import numpy as np
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile

# 配置 - 使用相对路径或命令行参数
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
DEFAULT_WORKSPACE = os.path.join(os.path.expanduser("~"), "openclaw/workspace")

# 从命令行参数获取输入输出路径
if len(sys.argv) < 2:
    print("用法：python3 slice_processor.py <源图片路径> [输出目录]")
    print("示例：python3 slice_processor.py /path/to/source.jpg /path/to/output")
    sys.exit(1)

SOURCE_IMG = sys.argv[1]
TASK_DIR = sys.argv[2] if len(sys.argv) > 2 else os.path.join(DEFAULT_WORKSPACE, "temp/slice-task")
SLICES_DIR = os.path.join(TASK_DIR, "slices_v7")
OUTPUT_DOCX = os.path.join(TASK_DIR, "output_v7.docx")
OUTPUT_ZIP = os.path.join(TASK_DIR, "slices_v7.zip")

# 切片配置
TARGET_WIDTH = 781
ASPECT_RATIO = 16/9
INITIAL_SLICE_HEIGHT = int(TARGET_WIDTH * ASPECT_RATIO)  # 约 1388

# 高度容差：切片高度应在目标值的±15% 范围内
HEIGHT_TOLERANCE = 0.15
MIN_SLICE_HEIGHT = int(INITIAL_SLICE_HEIGHT * (1 - HEIGHT_TOLERANCE))  # 约 1180
MAX_SLICE_HEIGHT = int(INITIAL_SLICE_HEIGHT * (1 + HEIGHT_TOLERANCE))  # 约 1596

def analyze_content_density(img_path):
    """分析图片内容密度"""
    print("分析图片内容密度...")
    img = Image.open(img_path).convert('L')
    img_array = np.array(img)
    
    height = img_array.shape[0]
    row_std = np.std(img_array, axis=1)
    
    # 平滑处理
    window = 20
    smoothed = np.convolve(row_std, np.ones(window)/window, mode='same')
    
    # 归一化
    if smoothed.max() > 0:
        smoothed = smoothed / smoothed.max()
    
    return smoothed, img.height

def find_content_regions(smoothed, threshold=0.12):
    """检测内容区域（文字行）"""
    is_content = smoothed > threshold
    
    regions = []
    in_content = False
    region_start = 0
    
    for i, is_text in enumerate(is_content):
        if is_text and not in_content:
            region_start = i
            in_content = True
        elif not is_text and in_content:
            regions.append((region_start, i))
            in_content = False
    
    if in_content:
        regions.append((region_start, len(is_content)))
    
    return regions

def find_all_gaps(regions, img_height):
    """找到所有间隙"""
    gaps = []
    for i in range(len(regions) - 1):
        curr_end = regions[i][1]
        next_start = regions[i + 1][0]
        gap_size = next_start - curr_end
        
        gaps.append({
            'start': curr_end,
            'end': next_start,
            'size': gap_size,
            'center': (curr_end + next_start) // 2
        })
    
    return gaps

def score_cut_position(proposed_cut, target_height, all_gaps, img_height):
    """
    对切分位置打分
    综合考虑：间隙大小、与目标的距离、切片高度合理性
    """
    best_score = -1000
    best_cut = int(proposed_cut)
    
    # 检查所有间隙
    for gap in all_gaps:
        gap_center = gap['center']
        gap_size = gap['size']
        
        # 计算与目标位置的距离
        distance = abs(gap_center - proposed_cut)
        
        # 只考虑合理范围内的切分点（±250 像素）
        if distance > 250:
            continue
        
        # 计算得分
        # 1. 间隙大小得分（间隙越大越好，满分 50）
        gap_score = min(gap_size / 150 * 50, 50)  # 150 像素以上满分
        
        # 2. 距离得分（越近越好，满分 30）
        distance_score = (1 - distance / 250) * 30
        
        # 3. 高度合理性得分（在目标范围内最好，满分 20）
        resulting_height = gap_center
        if MIN_SLICE_HEIGHT <= resulting_height <= MAX_SLICE_HEIGHT:
            height_score = 20
        else:
            # 超出范围则扣分
            height_score = max(0, 10 - abs(resulting_height - INITIAL_SLICE_HEIGHT) / 50)
        
        total_score = gap_score + distance_score + height_score
        
        if total_score > best_score:
            best_score = total_score
            best_cut = gap_center
    
    return best_cut

def find_best_cut_position_v7(start_y, target_height, all_gaps, img_height):
    """
    寻找最佳切分位置 v7 - 精细平衡
    """
    proposed_bottom = start_y + target_height
    
    if proposed_bottom >= img_height:
        return img_height
    
    # 使用评分系统找到最佳切分点
    best_cut = score_cut_position(proposed_bottom, target_height, all_gaps, img_height)
    
    return best_cut

def calculate_slice_positions_v7(regions, all_gaps, img_height):
    """计算智能切片位置 v7 - 精细平衡"""
    print(f"计算切片位置 (总高度：{img_height})...")
    print(f"  检测到 {len(regions)} 个内容区域")
    print(f"  检测到 {len(all_gaps)} 个间隙")
    print(f"  目标切片高度：{INITIAL_SLICE_HEIGHT} (±{HEIGHT_TOLERANCE*100}%)")
    print(f"  允许范围：{MIN_SLICE_HEIGHT} - {MAX_SLICE_HEIGHT}")
    
    slice_positions = []
    current_y = 0
    slice_num = 0
    max_iterations = 200
    
    while current_y < img_height and slice_num < max_iterations:
        slice_num += 1
        
        # 寻找最佳切分位置
        slice_bottom = find_best_cut_position_v7(
            current_y, 
            INITIAL_SLICE_HEIGHT, 
            all_gaps, 
            img_height
        )
        
        # 确保至少前进一点
        if slice_bottom <= current_y:
            slice_bottom = min(current_y + INITIAL_SLICE_HEIGHT // 2, img_height)
        
        slice_height = slice_bottom - current_y
        
        slice_positions.append({
            'num': slice_num,
            'y_start': current_y,
            'y_end': slice_bottom,
            'height': slice_height
        })
        
        # 标记高度异常的切片
        marker = ""
        if slice_height < MIN_SLICE_HEIGHT:
            marker = " ⚠️ 偏短"
        elif slice_height > MAX_SLICE_HEIGHT:
            marker = " ⚠️ 偏长"
        
        print(f"  切片 {slice_num}: {current_y} - {slice_bottom} (高度：{slice_height}){marker}")
        current_y = slice_bottom
    
    return slice_positions

def create_slices(img_path, slice_positions, output_dir):
    """创建切片图片"""
    print(f"创建切片图片...")
    os.makedirs(output_dir, exist_ok=True)
    
    img = Image.open(img_path)
    slice_files = []
    
    for pos in slice_positions:
        slice_img = img.crop((
            0,
            pos['y_start'],
            img.width,
            pos['y_end']
        ))
        
        filename = f"slice_{pos['num']:02d}.jpg"
        filepath = os.path.join(output_dir, filename)
        slice_img.save(filepath, 'JPEG', quality=95)
        slice_files.append(filepath)
        print(f"  保存：{filename} ({slice_img.width}x{slice_img.height})")
    
    return slice_files

def create_word_document(slice_files, output_path):
    """创建 Word 文档"""
    print(f"创建 Word 文档...")
    
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    for i, slice_path in enumerate(slice_files):
        if i > 0:
            doc.add_page_break()
        
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        available_width = Cm(21) - Cm(2) - Cm(2)
        run.add_picture(slice_path, width=available_width)
    
    doc.save(output_path)
    print(f"  保存：{output_path}")

def create_zip(slice_files, output_path):
    """创建 ZIP 包"""
    print(f"创建 ZIP 包...")
    
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for filepath in slice_files:
            arcname = os.path.basename(filepath)
            zf.write(filepath, arcname)
    
    print(f"  保存：{output_path}")

def main():
    print("=" * 50)
    print("长截图智能切片工具 v7 - 精细平衡版")
    print("=" * 50)
    
    if not os.path.exists(SOURCE_IMG):
        print(f"错误：源文件不存在 {SOURCE_IMG}")
        sys.exit(1)
    
    img = Image.open(SOURCE_IMG)
    print(f"源图片：{img.width} x {img.height}")
    
    # 分析内容密度
    smoothed, img_height = analyze_content_density(SOURCE_IMG)
    
    # 检测内容区域
    regions = find_content_regions(smoothed)
    
    # 找到所有间隙
    all_gaps = find_all_gaps(regions, img_height)
    
    # 计算切片位置
    slice_positions = calculate_slice_positions_v7(regions, all_gaps, img_height)
    print(f"共计划 {len(slice_positions)} 个切片")
    
    # 统计高度分布
    heights = [p['height'] for p in slice_positions]
    print(f"高度统计：最小={min(heights)}, 最大={max(heights)}, 平均={sum(heights)/len(heights):.0f}")
    
    # 创建切片
    slice_files = create_slices(SOURCE_IMG, slice_positions, SLICES_DIR)
    
    # 创建 Word 文档
    create_word_document(slice_files, OUTPUT_DOCX)
    
    # 创建 ZIP
    create_zip(slice_files, OUTPUT_ZIP)
    
    print("=" * 50)
    print("✅ 处理完成!")
    print(f"  切片数量：{len(slice_files)}")
    print(f"  Word 文档：{OUTPUT_DOCX}")
    print(f"  ZIP 包：{OUTPUT_ZIP}")
    print("=" * 50)

if __name__ == "__main__":
    main()

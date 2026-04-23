#!/usr/bin/env python3
"""
shipping-booking extract.py
从托书文件中提取结构化信息，输出标准 JSON。

支持格式及处理策略：
- PDF        → 文本提取（精确字符）+ 转图片（视觉布局）→ 双重喂给 Claude
- 图片        → 直接视觉识别（JPG/PNG/TIFF/BMP/WebP）
- Excel      → 带行列坐标的结构化文本 → Claude 分析（.xlsx/.xls）
- Word .docx → 带表格结构的文本 → Claude 分析
- Word .doc  → macOS 用 textutil；非 macOS 用 LibreOffice 或提示另存为 .docx
- RTF        → macOS 用 textutil；其他平台用 striprtf（pip install striprtf）

用法: python3 extract.py <文件路径>
"""

import sys
import os
import json
import re
import base64
import tempfile
import platform
import subprocess
from pathlib import Path
from typing import Optional


# ─────────────────────────────────────────────
# 文本提取
# ─────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """从 PDF 提取纯文本"""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except ImportError:
        raise RuntimeError("请先安装 pdfplumber: pip install pdfplumber")


def extract_structured_text_from_excel(file_path: str) -> str:
    """从 Excel 提取带行列坐标的结构化文本"""
    ext = Path(file_path).suffix.lower()

    if ext in ['.xls']:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_path)
            lines = ["[Excel 结构化内容，格式为 行,列: 内容]"]
            for sheet in wb.sheets():
                lines.append(f"[Sheet: {sheet.name}]")
                for row_idx in range(sheet.nrows):
                    for col_idx in range(sheet.ncols):
                        cell = sheet.cell_value(row_idx, col_idx)
                        if str(cell).strip():
                            lines.append(f"行{row_idx+1},列{col_idx+1}: {cell}")
            return "\n".join(lines)
        except ImportError:
            raise RuntimeError("请先安装 xlrd: pip install xlrd")
    else:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = ["[Excel 结构化内容，格式为 行,列: 内容]"]
            for sheet in wb.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None and str(cell.value).strip():
                            lines.append(f"行{cell.row},列{cell.column}: {cell.value}")
            return "\n".join(lines)
        except ImportError:
            raise RuntimeError("请先安装 openpyxl: pip install openpyxl")


def extract_text_via_textutil(file_path: str, fmt_label: str) -> str:
    """用 macOS textutil 将文件转为纯文本（支持 .doc / .rtf 等）"""
    try:
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp_path = tmp.name
        try:
            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-output', tmp_path, file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read().strip()
            else:
                raise RuntimeError(f"textutil 转换失败: {result.stderr}")
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    except FileNotFoundError:
        raise RuntimeError(f"未找到 textutil 命令（仅 macOS 支持），请将 {fmt_label} 另存为 .docx 后重试")


def extract_structured_text_from_rtf(file_path: str) -> str:
    """从 RTF 提取纯文本，优先用 textutil，备用 striprtf"""
    # 优先使用 macOS textutil
    try:
        return extract_text_via_textutil(file_path, '.rtf')
    except RuntimeError:
        pass
    # 备用：striprtf 库
    try:
        from striprtf.striprtf import rtf_to_text
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()
        return rtf_to_text(rtf_content).strip()
    except ImportError:
        raise RuntimeError("RTF 解析失败，请安装 striprtf: pip install striprtf")


def extract_structured_text_from_word(file_path: str) -> str:
    """从 Word 提取带表格结构的文本"""
    ext = Path(file_path).suffix.lower()

    if ext == '.doc':
        # macOS：使用系统自带 textutil
        if platform.system() == 'Darwin':
            return extract_text_via_textutil(file_path, '.doc')
        # 非 macOS：尝试 LibreOffice headless
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', tmp_dir, file_path],
                    capture_output=True, text=True, timeout=60
                )
                txt_path = os.path.join(tmp_dir, Path(file_path).stem + '.txt')
                if os.path.exists(txt_path):
                    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read().strip()
        except FileNotFoundError:
            pass
        # 两种方案都没有，给出明确提示
        raise RuntimeError(".doc 格式解析失败，请选择以下任一方式处理：\n  1. 在 Word 中另存为 .docx 后重新上传\n  2. 安装 LibreOffice（https://www.libreoffice.org）后重试")

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
        lines = []
        seen = set()  # 文本框内容可能重复，去重

        def add(text: str):
            t = text.strip()
            if t and t not in seen:
                seen.add(t)
                lines.append(t)

        # 普通段落
        for para in doc.paragraphs:
            add(para.text)

        # 表格（带行列结构）
        for table_idx, table in enumerate(doc.tables):
            lines.append(f"[表格 {table_idx+1}]")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    add(f"行{row_idx+1},列{col_idx+1}: {cell.text.strip()}")

        # 文本框（TextBox）— python-docx 默认不读，需手动遍历 XML
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for para in txbx.iter(qn('w:p')):
                text = ''.join(r.text for r in para.iter(qn('w:t')) if r.text)
                add(text)

        return "\n".join(lines)
    except ImportError:
        raise RuntimeError("请先安装 python-docx: pip install python-docx")


def extract_images_from_docx(file_path: str) -> "list[tuple[str,str]]":
    """从 docx 中提取嵌入图片，返回 [(base64, media_type), ...]"""
    import zipfile
    images = []
    ext_to_mime = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/png",
        ".tiff": "image/png", ".tif": "image/png",
    }
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    ext = Path(name).suffix.lower()
                    mime = ext_to_mime.get(ext)
                    if mime:
                        data = z.read(name)
                        images.append((base64.standard_b64encode(data).decode(), mime))
    except Exception:
        pass
    return images


# ─────────────────────────────────────────────
# 图片处理
# ─────────────────────────────────────────────

def pdf_to_base64_images(file_path: str) -> "list[str]":
    """将 PDF 每页转为 base64 图片"""
    try:
        import fitz  # pymupdf
        doc = fitz.open(file_path)
        images = []
        for page in doc:
            mat = fitz.Matrix(2.0, 2.0)  # 2x 分辨率，提升识别精度
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            images.append(base64.standard_b64encode(img_bytes).decode())
        return images
    except ImportError:
        raise RuntimeError("请先安装 pymupdf: pip install pymupdf")


def image_to_base64(file_path: str) -> str:
    """将图片文件转为 base64"""
    with open(file_path, 'rb') as f:
        return base64.standard_b64encode(f.read()).decode()


def get_image_media_type(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    return {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp',
    }.get(ext, 'image/png')


# ─────────────────────────────────────────────
# Claude 调用
# ─────────────────────────────────────────────

PROMPT_TEMPLATE = """你是一个海运物流专家，请从以下托书内容中提取结构化信息。

【提取规则】
1. 字段值原样保留，中文就中文，英文就英文，不做翻译
2. 日期统一转换为 YYYY-MM-DD 格式
3. 找不到的字段填 null
4. containers 是数组，支持多箱型，每个元素包含 type 和 quantity
5. 托书右上角或表头位置的公司名（出具托书的货代公司）不是发货人，不要混入 shipper_name。shipper_name 只取 Shipper / 托运人 标签对应的内容
6. 对于编号类字段（reference_number、hs_code、un_number 等），必须原模原样提取，不得改动任何字符
7. 整体置信度评估：
   - high：核心字段（shipper_name/consignee_name/pol/pod/cargo_description/containers）全部提取成功
   - medium：核心字段大部分提取成功，少量缺失
   - low：核心字段多个缺失，或文档质量差
8. low_confidence_fields：列出你不确定的字段名
9. 如果内容完全不像托书，将 confidence 设为 "invalid"

【目标 JSON Schema】
{
  "shipper_name": null,
  "shipper_address": null,
  "consignee_name": null,
  "consignee_address": null,
  "notify_name": null,
  "notify_address": null,
  "pol": null,
  "transit_port": null,
  "pod": null,
  "carrier": null,
  "vessel": null,
  "voyage": null,
  "etd": null,
  "eta": null,
  "cargo_description": null,
  "hs_code": null,
  "marks": null,
  "pieces": null,
  "gross_weight": null,
  "volume": null,
  "is_dangerous": false,
  "un_number": null,
  "dangerous_class": null,
  "containers": [],
  "freight_terms": null,
  "reference_number": null,
  "confidence": "high",
  "low_confidence_fields": []
}

请直接返回 JSON，不要有任何其他说明文字。"""


# 各 provider 的默认模型
_DEFAULT_TEXT_MODELS = {
    "anthropic": "claude-opus-4-5-20251101",
    "deepseek":  "deepseek-chat",
    "qwen":      "qwen-plus",
    "moonshot":  "moonshot-v1-8k",
    "zhipu":     "glm-4",
    "openai":    "gpt-4o",
}
_DEFAULT_VISION_MODELS = {
    "anthropic": "claude-opus-4-5-20251101",
    "deepseek":  None,                              # DeepSeek 不支持视觉
    "qwen":      "qwen-vl-plus",
    "moonshot":  "moonshot-v1-8k-vision-preview",
    "zhipu":     "glm-4v",
    "openai":    "gpt-4o",
}

# 不支持视觉的 provider 列表
_NO_VISION_PROVIDERS = {"deepseek"}


def _get_model(provider: str, vision: bool = False) -> str:
    """获取模型名：优先读环境变量，其次用 provider 默认值"""
    if vision:
        return os.environ.get("SHIPPING_VISION_MODEL") or _DEFAULT_VISION_MODELS.get(provider, "gpt-4o")
    return os.environ.get("SHIPPING_MODEL") or _DEFAULT_TEXT_MODELS.get(provider, "gpt-4o")


def call_ai_text(text: str, client, provider: str) -> dict:
    """纯文本模式（Excel/Word/RTF）"""
    prompt = f"{PROMPT_TEMPLATE}\n\n【托书文本内容】\n{text}"
    if provider == "anthropic":
        message = client.messages.create(
            model=_get_model(provider),
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}]
        )
        return parse_claude_response(message.content[0].text)
    else:
        response = client.chat.completions.create(
            model=_get_model(provider),
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}]
        )
        return parse_claude_response(response.choices[0].message.content)


def call_ai_pdf(text: str, images: "list[str]", client, provider: str) -> dict:
    """PDF 双重模式：文本 + 图片同时输入"""
    if provider == "anthropic":
        content = [{
            "type": "text",
            "text": f"{PROMPT_TEMPLATE}\n\n【文本层内容（字符精确，但缺少位置信息）】\n{text}\n\n【以下是文档图片，请结合视觉布局判断字段归属】"
        }]
        for img_b64 in images:
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/png", "data": img_b64}
            })
        message = client.messages.create(
            model=_get_model(provider, vision=True),
            max_tokens=2048,
            messages=[{"role": "user", "content": content}]
        )
        return parse_claude_response(message.content[0].text)
    else:
        # OpenAI 兼容视觉模式
        # 不支持视觉的 provider（如 DeepSeek）降级为纯文本模式
        if provider in _NO_VISION_PROVIDERS and not os.environ.get("SHIPPING_VISION_MODEL"):
            print(f"⚠️  {provider} 不支持视觉，降级为纯文本模式（仅使用文本层）...", file=sys.stderr)
            return call_ai_text(text, client, provider)
        vision_model = _get_model(provider, vision=True)
        content = [{
            "type": "text",
            "text": f"{PROMPT_TEMPLATE}\n\n【文本层内容】\n{text}\n\n【以下是文档图片，请结合视觉布局判断字段归属】"
        }]
        for img_b64 in images:
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{img_b64}"}
            })
        response = client.chat.completions.create(
            model=vision_model,
            max_tokens=2048,
            messages=[{"role": "user", "content": content}]
        )
        return parse_claude_response(response.choices[0].message.content)


def call_ai_image(images: "list[str]", media_types: "list[str]", client, provider: str) -> dict:
    """纯图片模式（扫描件/照片）"""
    if provider == "anthropic":
        content = [{"type": "text", "text": f"{PROMPT_TEMPLATE}\n\n【以下是托书图片】"}]
        for img_b64, media_type in zip(images, media_types):
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": img_b64}
            })
        message = client.messages.create(
            model=_get_model(provider, vision=True),
            max_tokens=2048,
            messages=[{"role": "user", "content": content}]
        )
        return parse_claude_response(message.content[0].text)
    else:
        # 不支持视觉的 provider 无法处理纯图片，给出明确提示
        if provider in _NO_VISION_PROVIDERS and not os.environ.get("SHIPPING_VISION_MODEL"):
            raise RuntimeError(
                f"{provider} 不支持图片识别，请选择以下任一方式处理：\n"
                "  1. 将图片转为 PDF（有文本层时自动降级纯文本模式）\n"
                "  2. 配置支持视觉的模型，例如通义千问：\n"
                "     export SHIPPING_VISION_MODEL=qwen-vl-plus\n"
                "     export OPENAI_BASE_URL=https://dashscope.aliyuncs.com/compatible-mode/v1"
            )
        vision_model = _get_model(provider, vision=True)
        content = [{"type": "text", "text": f"{PROMPT_TEMPLATE}\n\n【以下是托书图片】"}]
        for img_b64, media_type in zip(images, media_types):
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{media_type};base64,{img_b64}"}
            })
        response = client.chat.completions.create(
            model=vision_model,
            max_tokens=2048,
            messages=[{"role": "user", "content": content}]
        )
        return parse_claude_response(response.choices[0].message.content)


def parse_claude_response(raw: str) -> dict:
    """解析 Claude 返回的 JSON"""
    raw = raw.strip()
    json_match = re.search(r'\{[\s\S]*\}', raw)
    if not json_match:
        raise ValueError("Claude 返回格式异常，未找到有效 JSON")
    return json.loads(json_match.group())


# ─────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────

def _get_openclaw_api_key() -> Optional[dict]:
    """
    从 OpenClaw 配置文件自动读取 API Key。
    返回 {"provider": "anthropic"|"openai", "key": "...", "base_url": "..."(可选)}
    优先级：Anthropic > OpenAI兼容
    """
    candidates = [
        os.path.expanduser("~/.openclaw/agents/main/agent/auth-profiles.json"),
        os.path.expanduser("~/.openclaw/auth-profiles.json"),
    ]
    # OpenAI 兼容的国产模型 base_url 映射
    OPENAI_COMPAT_PROVIDERS = {
        "deepseek":  "https://api.deepseek.com/v1",
        "qwen":      "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "moonshot":  "https://api.moonshot.cn/v1",
        "zhipu":     "https://open.bigmodel.cn/api/paas/v4",
        "openai":    None,  # 官方 OpenAI，无需自定义 base_url
    }
    anthropic_key = None
    openai_compat = None

    for path in candidates:
        if not os.path.exists(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            profiles = data.get("profiles", {})
            for profile in profiles.values():
                provider = profile.get("provider", "").lower()
                token = profile.get("token") or profile.get("key")
                if not token:
                    continue
                if provider == "anthropic" and not anthropic_key:
                    anthropic_key = token
                elif provider in OPENAI_COMPAT_PROVIDERS and not openai_compat:
                    openai_compat = {
                        "provider": provider,
                        "key": token,
                        "base_url": OPENAI_COMPAT_PROVIDERS[provider],
                    }
        except Exception:
            continue

    if anthropic_key:
        return {"provider": "anthropic", "key": anthropic_key}
    if openai_compat:
        return openai_compat
    return None


def _build_client(key_info: dict):
    """根据 provider 构建对应的 AI 客户端，返回 (client, provider_name)
    provider_name 保留原始名称（anthropic/deepseek/qwen/moonshot/zhipu/openai）
    """
    provider = key_info["provider"]
    key = key_info["key"]

    if provider == "anthropic":
        try:
            import anthropic
            return anthropic.Anthropic(api_key=key), "anthropic"
        except ImportError:
            raise RuntimeError("请先安装 anthropic SDK: pip install anthropic")
    else:
        # OpenAI 兼容模式（DeepSeek / Qwen / Kimi / OpenAI 等）
        # 保留原始 provider 名，用于后续视觉能力判断
        try:
            import openai
            kwargs = {"api_key": key}
            if key_info.get("base_url"):
                kwargs["base_url"] = key_info["base_url"]
            return openai.OpenAI(**kwargs), provider  # 返回原始名如 "deepseek"/"qwen" 等
        except ImportError:
            raise RuntimeError("请先安装 openai SDK: pip install openai")

def is_valid_booking(data: dict) -> bool:
    if data.get("confidence") == "invalid":
        return False
    core_fields = ["shipper_name", "consignee_name", "pol", "pod", "cargo_description"]
    non_null = sum(1 for f in core_fields if data.get(f))
    return non_null >= 2


def main():
    if len(sys.argv) < 2:
        print("用法: python3 extract.py <文件路径>", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)

    # 初始化 AI 客户端
    # 优先读环境变量，其次自动从 OpenClaw 配置读取
    env_key = os.environ.get("ANTHROPIC_API_KEY")
    if env_key:
        key_info = {"provider": "anthropic", "key": env_key}
    else:
        env_key = os.environ.get("OPENAI_API_KEY")
        if env_key:
            key_info = {
                "provider": "openai",
                "key": env_key,
                "base_url": os.environ.get("OPENAI_BASE_URL"),
            }
        else:
            key_info = _get_openclaw_api_key()

    if not key_info:
        print("❌ 未找到 AI API Key，请选择以下任一方式配置：\n"
              "  1. 在 OpenClaw 中配置 Anthropic/OpenAI 兼容账号（推荐）\n"
              "  2. 设置环境变量: export ANTHROPIC_API_KEY=your_key\n"
              "  3. OpenAI 兼容模型: export OPENAI_API_KEY=your_key OPENAI_BASE_URL=https://api.deepseek.com/v1", file=sys.stderr)
        sys.exit(1)

    try:
        client, provider = _build_client(key_info)
    except RuntimeError as e:
        print(f"❌ {e}", file=sys.stderr)
        sys.exit(1)

    ext = Path(file_path).suffix.lower()

    try:
        # ── PDF：文本 + 图片双重模式 ──
        if ext == '.pdf':
            print("📄 检测到 PDF，启用双重模式（文本+视觉）...", file=sys.stderr)
            text = extract_text_from_pdf(file_path)
            images = pdf_to_base64_images(file_path)
            result = call_ai_pdf(text, images, client, provider)

        # ── 图片：纯视觉模式 ──
        elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.webp']:
            print("🖼️  检测到图片，启用视觉模式...", file=sys.stderr)
            img_b64 = image_to_base64(file_path)
            media_type = get_image_media_type(file_path)
            result = call_ai_image([img_b64], [media_type], client, provider)

        # ── Excel：带坐标结构化文本 ──
        elif ext in ['.xlsx', '.xls']:
            print("📊 检测到 Excel，提取结构化文本...", file=sys.stderr)
            text = extract_structured_text_from_excel(file_path)
            result = call_ai_text(text, client, provider)

        # ── Word：带表格结构文本 ──
        elif ext in ['.docx', '.doc']:
            print("📝 检测到 Word，提取表格结构文本...", file=sys.stderr)
            text = extract_structured_text_from_word(file_path)
            if text.strip():
                result = call_ai_text(text, client, provider)
            elif ext == '.docx':
                # 文本为空，尝试提取嵌入图片（图片版 Word）
                print("📝 Word 文本为空，尝试提取嵌入图片...", file=sys.stderr)
                embedded = extract_images_from_docx(file_path)
                if not embedded:
                    raise RuntimeError("Word 文件无可读文字，也未找到嵌入图片，请检查文件是否损坏")
                imgs = [b64 for b64, _ in embedded]
                mimes = [mime for _, mime in embedded]
                result = call_ai_image(imgs, mimes, client, provider)
            else:
                raise RuntimeError("Word 文件无可读内容，请另存为 .docx 后重试")

        # ── RTF：纯文本模式 ──
        elif ext == '.rtf':
            print("📄 检测到 RTF，提取文本...", file=sys.stderr)
            text = extract_structured_text_from_rtf(file_path)
            result = call_ai_text(text, client, provider)

        else:
            print(f"❌ 暂不支持该文件格式（{ext}），支持 PDF、图片（JPG/PNG/TIFF/BMP）、Word（.docx/.doc）、Excel（.xlsx/.xls）、RTF", file=sys.stderr)
            sys.exit(1)

    except RuntimeError as e:
        print(f"❌ {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"❌ 提取失败: {e}", file=sys.stderr)
        sys.exit(1)

    # 验证是否为托书
    if not is_valid_booking(result):
        print("❌ 这不像是托书文件，请确认上传的文件是否正确", file=sys.stderr)
        sys.exit(1)

    # 输出 JSON
    print(json.dumps(result, ensure_ascii=False, indent=2))

    # 低置信度提醒
    low_fields = result.get("low_confidence_fields", [])
    confidence = result.get("confidence", "high")
    if low_fields or confidence == "low":
        fields_str = "、".join(low_fields) if low_fields else "多个字段"
        print(f"\n⚠️  以下字段置信度较低，建议人工核查：{fields_str}", file=sys.stderr)


if __name__ == "__main__":
    main()

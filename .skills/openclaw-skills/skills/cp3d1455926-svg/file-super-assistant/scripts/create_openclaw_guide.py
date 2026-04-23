#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 OpenClaw 教程文章转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_openclaw_guide():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('OpenClaw 完全指南：从入门到精通', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    # 副标题
    subtitle = doc.add_paragraph('打造你的 AI 助理帝国')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    # 引言
    doc.add_paragraph()
    intro = doc.add_paragraph('2026 年，不会用 AI 助理的人，就像 2010 年不会用智能手机的人。\nOpenClaw，让你拥有一个真正懂你的 AI 助理。')
    intro.runs[0].font.italic = True
    
    # 目录
    doc.add_page_break()
    doc.add_heading('目录', level=2)
    
    chapters = [
        '01 OpenClaw 是什么？',
        '02 安装 OpenClaw',
        '03 基础配置',
        '04 技能管理',
        '05 日常使用',
        '06 高级配置',
        '07 性能优化',
        '08 故障排除',
        '09 最佳实践',
        '10 资源推荐',
        '11 总结'
    ]
    
    for i, chapter in enumerate(chapters, 1):
        doc.add_paragraph(f'{i}. {chapter}', style='List Number')
    
    # 正文内容
    doc.add_page_break()
    
    # 第 1 章
    doc.add_heading('01 OpenClaw 是什么？', level=1)
    doc.add_paragraph('先说人话：OpenClaw 是一个 AI 助理框架，帮你把各种 AI 能力整合到一起。')
    
    doc.add_paragraph('举个例子：', style='Intense Quote')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('你想写小红书文案 → OpenClaw 调用小红书助手技能')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('你想分析网站数据 → OpenClaw 调用数据分析技能')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('你想创建测试问卷 → OpenClaw 调用测试生成器技能')
    
    doc.add_heading('核心优势：', level=2)
    advantages = [
        '技能化：每个功能都是独立技能，按需安装',
        '可扩展：可以自己写技能，也能安装别人分享的',
        '多平台：支持微信、QQ、Telegram、飞书等',
        '本地化：数据存在自己电脑，隐私安全'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    # 第 2 章
    doc.add_heading('02 安装 OpenClaw', level=1)
    
    doc.add_heading('方式一：一键安装（推荐新手）', level=2)
    doc.add_paragraph('Windows 用户：')
    doc.add_paragraph('npm install -g openclaw', style='No Spacing')
    
    doc.add_paragraph('Mac/Linux 用户：')
    doc.add_paragraph('sudo npm install -g openclaw', style='No Spacing')
    
    doc.add_paragraph('验证安装：')
    doc.add_paragraph('openclaw --version', style='No Spacing')
    
    doc.add_heading('方式二：源码安装（推荐开发者）', level=2)
    steps = [
        '克隆仓库：git clone https://github.com/openclaw/openclaw.git',
        '进入目录：cd openclaw',
        '安装依赖：npm install',
        '全局链接：npm link',
        '验证：openclaw --version'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    # 保存文档
    output_path = 'D:\\OneDrive\\Desktop\\公众号文章\\2026-03-15_OpenClaw 完全指南.docx'
    doc.save(output_path)
    print(f'[OK] 已创建 Word 文档：{output_path}')

if __name__ == '__main__':
    create_openclaw_guide()

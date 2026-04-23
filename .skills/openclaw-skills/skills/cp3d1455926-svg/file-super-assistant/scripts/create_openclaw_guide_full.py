#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OpenClaw 完全指南 - 生成 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='微软雅黑', font_size=12, bold=False, italic=False, color=None):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_code_block(doc, code):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    set_font(run, 'Consolas', 10, color=(50, 50, 50))

def create_openclaw_guide():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ========== 封面 ==========
    for _ in range(5):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph('OpenClaw 完全指南')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    set_font(title_run, '微软雅黑', 36, bold=True, color=(0, 102, 204))
    
    # 副标题
    subtitle = doc.add_paragraph('从入门到精通，打造你的 AI 助理帝国')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.runs[0], '微软雅黑', 18, color=(100, 100, 100))
    
    # 作者信息
    for _ in range(3):
        doc.add_paragraph()
    
    author = doc.add_paragraph('作者：Jake')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(author.runs[0], '微软雅黑', 14)
    
    date = doc.add_paragraph('日期：2026 年 3 月 15 日')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(date.runs[0], '微软雅黑', 12, color=(150, 150, 150))
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    doc.add_heading('目录', level=1)
    
    chapters = [
        ('01 OpenClaw 是什么？', 1),
        ('02 安装 OpenClaw', 2),
        ('03 基础配置', 3),
        ('04 技能管理', 4),
        ('05 日常使用', 5),
        ('06 高级配置', 6),
        ('07 性能优化', 7),
        ('08 故障排除', 8),
        ('09 最佳实践', 9),
        ('10 资源推荐', 10),
        ('11 总结', 11)
    ]
    
    for chapter, num in chapters:
        p = doc.add_paragraph(chapter, style='List Number')
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ========== 第 1 章 ==========
    doc.add_heading('01 OpenClaw 是什么？', level=1)
    
    doc.add_paragraph('先说人话：OpenClaw 是一个 AI 助理框架，帮你把各种 AI 能力整合到一起。')
    
    doc.add_paragraph('举个例子：')
    examples = [
        '你想写小红书文案 → OpenClaw 调用小红书助手技能',
        '你想分析网站数据 → OpenClaw 调用数据分析技能',
        '你想创建测试问卷 → OpenClaw 调用测试生成器技能'
    ]
    for ex in examples:
        p = doc.add_paragraph(ex, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_heading('核心优势：', level=2)
    advantages = [
        '技能化：每个功能都是独立技能，按需安装',
        '可扩展：可以自己写技能，也能安装别人分享的',
        '多平台：支持微信、QQ、Telegram、飞书等',
        '本地化：数据存在自己电脑，隐私安全'
    ]
    for adv in advantages:
        p = doc.add_paragraph(adv, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_heading('适合谁用：', level=2)
    users = [
        '想搞副业的年轻人',
        '需要做自媒体的创作者',
        '想提高效率的职场人',
        '喜欢折腾的技术爱好者'
    ]
    for user in users:
        p = doc.add_paragraph(user, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    # ========== 第 2 章 ==========
    doc.add_heading('02 安装 OpenClaw', level=1)
    
    doc.add_heading('方式一：一键安装（推荐新手）', level=2)
    doc.add_paragraph('Windows 用户：')
    add_code_block(doc, 'npm install -g openclaw')
    
    doc.add_paragraph('Mac/Linux 用户：')
    add_code_block(doc, 'sudo npm install -g openclaw')
    
    doc.add_paragraph('验证安装：')
    add_code_block(doc, 'openclaw --version')
    doc.add_paragraph('看到版本号就说明安装成功了。')
    
    doc.add_heading('方式二：源码安装（推荐开发者）', level=2)
    steps = [
        '克隆仓库：git clone https://github.com/openclaw/openclaw.git',
        '进入目录：cd openclaw',
        '安装依赖：npm install',
        '全局链接：npm link',
        '验证：openclaw --version'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('常见问题', level=2)
    doc.add_paragraph('问题 1：npm 安装慢', style='Heading 3')
    add_code_block(doc, 'npm config set registry https://registry.npmmirror.com')
    
    doc.add_paragraph('问题 2：权限错误', style='Heading 3')
    add_code_block(doc, 'sudo npm install -g openclaw')
    
    doc.add_paragraph('问题 3：找不到命令', style='Heading 3')
    doc.add_paragraph('重启终端或添加 npm 全局目录到 PATH')
    
    # ========== 第 3 章 ==========
    doc.add_heading('03 基础配置', level=1)
    
    doc.add_heading('Step 1：初始化配置', level=2)
    add_code_block(doc, 'openclaw init')
    doc.add_paragraph('会提示你设置：默认 AI 模型、消息通道、工作目录')
    
    doc.add_heading('Step 2：配置 AI 模型', level=2)
    doc.add_paragraph('编辑配置文件：')
    add_code_block(doc, 'openclaw config')
    
    doc.add_paragraph('通义千问配置（推荐）：')
    add_code_block(doc, 'model: qwen-plus\napi_key: 你的 API 密钥\nbase_url: https://dashscope.aliyuncs.com')
    
    doc.add_heading('Step 3：配置消息通道', level=2)
    doc.add_paragraph('企业微信配置：')
    add_code_block(doc, 'channel: wecom\ncorp_id: 企业 ID\nagent_id: 应用 ID\nsecret: 应用 Secret')
    
    doc.add_heading('Step 4：测试配置', level=2)
    add_code_block(doc, 'openclaw status')
    doc.add_paragraph('看到 ✅ 就说明配置成功了。')
    
    # ========== 第 4 章 ==========
    doc.add_heading('04 技能管理', level=1)
    
    doc.add_heading('安装技能', level=2)
    doc.add_paragraph('从 ClawHub 安装：')
    add_code_block(doc, 'clawhub install skill-name')
    
    doc.add_paragraph('从 GitHub 安装：')
    add_code_block(doc, 'clawhub install github-username/repo-name')
    
    doc.add_heading('推荐必装技能', level=2)
    skills = [
        '小红书运营助手 - 文案生成',
        '公众号超级助手 - 文章写作',
        '数据分析助手 - 数据统计',
        '测试生成器 - 创建测试',
        '文件助手 - 文件处理'
    ]
    for skill in skills:
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_heading('更新技能', level=2)
    add_code_block(doc, '# 更新单个技能\nclawhub update skill-name\n\n# 更新所有技能\nclawhub update')
    
    doc.add_heading('删除技能', level=2)
    add_code_block(doc, 'clawhub uninstall skill-name')
    
    # ========== 第 5 章 ==========
    doc.add_heading('05 日常使用', level=1)
    
    doc.add_heading('聊天模式', level=2)
    add_code_block(doc, 'openclaw chat')
    doc.add_paragraph('进入交互模式，可以直接对话。')
    
    doc.add_heading('执行任务', level=2)
    tasks = [
        '写文章：openclaw run "写一篇关于 AI 副业的文章"',
        '分析数据：openclaw run "分析网站昨天的访问数据"',
        '创建测试：openclaw run "创建一个性格测试问卷"'
    ]
    for task in tasks:
        p = doc.add_paragraph(task, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_heading('定时任务', level=2)
    add_code_block(doc, '# 添加定时提醒\nopenclaw cron add "每天 9 点发送早报"\n\n# 查看定时任务\nopenclaw cron list\n\n# 删除定时任务\nopenclaw cron remove task-id')
    
    # ========== 第 6 章 ==========
    doc.add_heading('06 高级配置', level=1)
    
    doc.add_heading('自定义技能', level=2)
    doc.add_paragraph('创建技能目录：')
    add_code_block(doc, 'mkdir -p skills/my-skill')
    
    doc.add_paragraph('创建 SKILL.md：')
    add_code_block(doc, '---\nname: my-skill\ndescription: 我的自定义技能\n---\n\n# 技能说明\n\n## 功能\n- 功能 1\n- 功能 2')
    
    doc.add_heading('配置自动回复', level=2)
    add_code_block(doc, 'auto_reply:\n  enabled: true\n  rules:\n    - keyword: "你好"\n      response: "你好！我是 AI 助理，有什么可以帮你？"')
    
    # ========== 第 7 章 ==========
    doc.add_heading('07 性能优化', level=1)
    
    doc.add_heading('提高响应速度', level=2)
    tips1 = [
        '使用本地模型：model: ollama/llama2',
        '减少技能加载：设置 auto_load: false',
        '启用缓存：cache.enabled: true'
    ]
    for tip in tips1:
        p = doc.add_paragraph(tip, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_heading('节省 Token', level=2)
    tips2 = [
        '设置上下文长度：max_tokens: 4000',
        '使用精简模式：response.concise: true',
        '限制历史消息：history_messages: 5'
    ]
    for tip in tips2:
        p = doc.add_paragraph(tip, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    # ========== 第 8 章 ==========
    doc.add_heading('08 故障排除', level=1)
    
    doc.add_heading('常见问题', level=2)
    doc.add_paragraph('问题 1：技能不工作', style='Heading 3')
    add_code_block(doc, '# 检查技能是否安装\nclawhub list\n\n# 重新安装\nclawhub uninstall skill-name\nclawhub install skill-name')
    
    doc.add_paragraph('问题 2：消息发送失败', style='Heading 3')
    add_code_block(doc, '# 检查配置\nopenclaw config\n\n# 测试连接\nopenclaw test channel')
    
    doc.add_heading('查看日志', level=2)
    add_code_block(doc, '# 查看实时日志\nopenclaw logs\n\n# 查看历史日志\nopenclaw logs --since 1h\n\n# 导出日志\nopenclaw logs --output log.txt')
    
    # ========== 第 9 章 ==========
    doc.add_heading('09 最佳实践', level=1)
    
    doc.add_heading('技能组织', level=2)
    add_code_block(doc, 'skills/\n├── official/       # 官方技能\n├── community/      # 社区技能\n├── custom/         # 自定义技能\n└── experimental/   # 实验性技能')
    
    doc.add_heading('安全建议', level=2)
    safety = [
        '不要分享 API 密钥',
        '定期备份配置',
        '只安装可信技能',
        '检查技能权限',
        '使用预检检查器'
    ]
    for item in safety:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    # ========== 第 10 章 ==========
    doc.add_heading('10 资源推荐', level=1)
    
    doc.add_heading('官方资源', level=2)
    resources = [
        '文档：https://docs.openclaw.ai',
        'GitHub：https://github.com/openclaw/openclaw',
        '社区：https://discord.gg/clawd',
        '技能市场：https://clawhub.ai'
    ]
    for res in resources:
        p = doc.add_paragraph(res, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_heading('推荐技能', level=2)
    top_skills = [
        'xiaohongshu-assistant - 小红书文案 ⭐⭐⭐⭐⭐',
        'official-account-assistant - 公众号写作 ⭐⭐⭐⭐⭐',
        'data-analytics-assistant - 数据分析 ⭐⭐⭐⭐⭐',
        'quiz-generator - 测试生成 ⭐⭐⭐⭐⭐',
        'life-memory-logger - 记忆管理 ⭐⭐⭐⭐'
    ]
    for skill in top_skills:
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
    
    # ========== 第 11 章 ==========
    doc.add_heading('11 总结', level=1)
    
    doc.add_paragraph('OpenClaw 是什么？')
    doc.add_paragraph('一个让你拥有 AI 助理的框架。', style='Intense Quote')
    
    doc.add_paragraph('能做什么？')
    doc.add_paragraph('写文章、分析数据、创建测试、管理文件、自动回复...', style='Intense Quote')
    
    doc.add_paragraph('难不难？')
    doc.add_paragraph('10 分钟安装，30 分钟上手，1 小时创建第一个技能。', style='Intense Quote')
    
    doc.add_paragraph('值不值得？')
    doc.add_paragraph('一天搞定的 8 个技能，已经帮我赚了第一笔钱。', style='Intense Quote')
    
    # 结尾
    doc.add_paragraph()
    doc.add_paragraph('最后说一句：', style='Heading 2')
    doc.add_paragraph('AI 时代，机会很多。但不是等来的，是干出来的。OpenClaw 就是工具，关键是你用它做什么。现在就开始吧。')
    
    # ========== 页脚 ==========
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('---\n\n作者：Jake\n字数：3500 字\n阅读时间：10 分钟\n更新日期：2026-03-15')
    set_font(footer_run, '微软雅黑', 10, color=(150, 150, 150))
    
    # 保存文档
    output_path = 'D:\\OneDrive\\Desktop\\公众号文章\\2026-03-15_OpenClaw 完全指南.docx'
    doc.save(output_path)
    print(f'[OK] 已创建 Word 文档：{output_path}')
    print(f'     文件大小：约 50KB')
    print(f'     总页数：约 10 页')

if __name__ == '__main__':
    create_openclaw_guide()

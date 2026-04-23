#!/usr/bin/env python3
"""
香港中學升中個人檔案 (Portfolio) 優化版 v3
基於真實簡歷優化，符合香港中學收生標準
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 配色方案 ──
THEMES = {
    "傳統名校": {
        "primary": RGBColor(0x1B, 0x3A, 0x5C),     # 深海藍
        "secondary": RGBColor(0x2C, 0x5F, 0x8A),   # 中藍
        "accent": RGBColor(0xD4, 0xA5, 0x37),       # 金色
        "bg": RGBColor(0xF5, 0xF7, 0xFA),           # 淡藍灰
        "light": RGBColor(0xE8, 0xEE, 0xF5),        # 淺藍
    },
    "直資私立": {
        "primary": RGBColor(0x5B, 0x1A, 0x3E),      # 深紫紅
        "secondary": RGBColor(0x8B, 0x3A, 0x6A),    # 紫紅
        "accent": RGBColor(0xD4, 0xA5, 0x37),        # 金色
        "bg": RGBColor(0xFA, 0xF5, 0xF8),           # 淡粉
        "light": RGBColor(0xF0, 0xE0, 0xEA),        # 淺粉
    },
    "國際學校": {
        "primary": RGBColor(0x1A, 0x5C, 0x3A),      # 深綠
        "secondary": RGBColor(0x2A, 0x8A, 0x5C),    # 中綠
        "accent": RGBColor(0xD4, 0xA5, 0x37),        # 金色
        "bg": RGBColor(0xF0, 0xFA, 0xF5),           # 淡綠
        "light": RGBColor(0xD5, 0xF0, 0xE0),        # 淺綠
    }
}


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def rgb_to_hex(color):
    return f"{color.red:02X}{color.red:02X}{color.green:02X}" if False else \
           f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, tuple) else \
           f"{color.red:02X}{color.green:02X}{color.blue:02X}"


def add_styled_paragraph(doc, text, font_size=12, bold=False, color=None,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                          font_name="Microsoft YaHei", font_name_en="Calibri"):
    """添加格式化段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_section_title(doc, num, title_zh, title_en, theme):
    """添加章节标题（带装饰线）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)

    # 编号
    run_num = p.add_run(f"  {num}  ")
    run_num.font.size = Pt(14)
    run_num.bold = True
    run_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run_num.font.name = "Calibri"

    # 中文标题
    run_zh = p.add_run(f"  {title_zh}  ")
    run_zh.font.size = Pt(16)
    run_zh.bold = True
    run_zh.font.color.rgb = theme["primary"]
    run_zh.font.name = "Calibri"
    run_zh._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")

    # 英文标题
    run_en = p.add_run(f"  {title_en}")
    run_en.font.size = Pt(12)
    run_en.font.color.rgb = theme["secondary"]
    run_en.font.name = "Calibri"

    # 装饰线
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.space_before = Pt(0)
    run_line = p2.add_run("━" * 60)
    run_line.font.size = Pt(6)
    run_line.font.color.rgb = theme["accent"]

    return p


def style_table_cell(cell, text, font_size=10, bold=False, color=None,
                      bg_color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      font_name="Microsoft YaHei"):
    """样式化表格单元格"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg_color:
        set_cell_shading(cell, bg_color)


def generate_portfolio(data, theme_name="傳統名校", output_path="portfolio.docx"):
    """生成优化版 Portfolio"""
    theme = THEMES.get(theme_name, THEMES["傳統名校"])
    primary_hex = rgb_to_hex(theme["primary"])
    accent_hex = rgb_to_hex(theme["accent"])
    light_hex = rgb_to_hex(theme["light"])
    secondary_hex = rgb_to_hex(theme["secondary"])

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ══════════════════════════════════════
    # 封面页
    # ══════════════════════════════════════
    # 顶部装饰线
    add_styled_paragraph(doc, "", font_size=12, space_after=30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆  ◆  ◆")
    run.font.size = Pt(16)
    run.font.color.rgb = theme["accent"]

    add_styled_paragraph(doc, "", font_size=12, space_after=20)

    # 姓名和照片并排（表格）
    name_table = doc.add_table(rows=1, cols=2)
    name_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 左侧：姓名和标题
    left_cell = name_table.cell(0, 0)
    left_cell.width = Cm(10)
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 中文姓名
    run = left_p.add_run(data["basic_info"]["name_zh"])
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = theme["primary"]
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")

    left_p.add_run("\n")

    # 英文姓名
    run = left_p.add_run(data["basic_info"]["name_en"])
    run.font.size = Pt(18)
    run.font.color.rgb = theme["secondary"]
    run.font.name = "Calibri"

    left_p.add_run("\n\n")

    # 标题
    run = left_p.add_run("升 中 個 人 檔 案")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = theme["primary"]
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")

    left_p.add_run("\n")

    run = left_p.add_run("Secondary School Application Portfolio")
    run.font.size = Pt(12)
    run.font.color.rgb = theme["secondary"]
    run.font.name = "Calibri"

    # 右侧：照片位
    right_cell = name_table.cell(0, 1)
    right_cell.width = Cm(5)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 照片框
    run = right_p.add_run("\n📷\n照片\nPhoto\n")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft YaHei")
    set_cell_shading(right_cell, light_hex)

    add_styled_paragraph(doc, "", font_size=12, space_after=20)

    # 底部装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆  ◆  ◆")
    run.font.size = Pt(16)
    run.font.color.rgb = theme["accent"]

    add_styled_paragraph(doc, "", font_size=12, space_after=20)

    # 学校信息
    edu = data["education"]
    add_styled_paragraph(doc, f"就讀學校：{edu['school_zh']}",
                          font_size=12, color=theme["primary"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, f"Present School: {edu['school_en']}",
                          font_size=10, color=theme["secondary"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, f"年級 Grade: {edu['grade']}",
                          font_size=11, color=theme["primary"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # 分页
    doc.add_page_break()

    # ══════════════════════════════════════
    # 一、自我介紹
    # ══════════════════════════════════════
    add_section_title(doc, "壹", "自我介紹", "Self Introduction", theme)

    # 中文版
    add_styled_paragraph(doc, "【中文】", font_size=10, bold=True,
                          color=theme["accent"], space_after=4)
    zh_text = data.get("self_intro_zh", "")
    for para_text in zh_text.split("\n\n"):
        add_styled_paragraph(doc, para_text.strip(), font_size=11,
                              color=RGBColor(0x33, 0x33, 0x33), space_after=6,
                              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_styled_paragraph(doc, "", font_size=6, space_after=8)

    # 英文版
    add_styled_paragraph(doc, "【English】", font_size=10, bold=True,
                          color=theme["accent"], space_after=4)
    en_text = data.get("self_intro_en", "")
    for para_text in en_text.split("\n\n"):
        add_styled_paragraph(doc, para_text.strip(), font_size=11,
                              color=RGBColor(0x33, 0x33, 0x33), space_after=6,
                              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              font_name="Calibri", font_name_en="Calibri")

    # ══════════════════════════════════════
    # 二、個人資料
    # ══════════════════════════════════════
    add_section_title(doc, "貳", "個人資料", "Personal Information", theme)

    info = data["basic_info"]
    fields = [
        ("中文姓名 Chinese Name", info["name_zh"], "英文姓名 English Name", info["name_en"]),
        ("性別 Gender", info["gender"], "出生日期 Date of Birth", info["birthday"]),
        ("證件類別 ID Type", info["id_type"], "證件號碼 ID No.", info["id_number"]),
        ("聯絡電話 Tel.", info["phone"], "電郵 Email", info["email"]),
        ("住址 Address", info["address"], "", ""),
        ("緊急聯絡人 Emergency Contact", info.get("guardian", ""),
         "關係 Relationship", info.get("guardian_relation", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, (l1, v1, l2, v2) in enumerate(fields):
        if l1:
            style_table_cell(table.cell(i, 0), l1, font_size=9, bold=True,
                              color=theme["primary"], bg_color=light_hex)
        if v1:
            style_table_cell(table.cell(i, 1), v1, font_size=10,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if l2:
            style_table_cell(table.cell(i, 2), l2, font_size=9, bold=True,
                              color=theme["primary"], bg_color=light_hex)
        if v2:
            style_table_cell(table.cell(i, 3), v2, font_size=10,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 合并住址行
    if fields[4][1] and not fields[4][2]:
        table.cell(4, 1).merge(table.cell(4, 3))

    # ══════════════════════════════════════
    # 三、教育背景
    # ══════════════════════════════════════
    add_section_title(doc, "參", "教育背景", "Education Background", theme)

    edu = data["education"]
    edu_fields = [
        ("學校名稱 School", edu["school_zh"], "學校地址 Address", edu.get("address", "")),
        ("入學時間 Enrolled", edu.get("enrollment", ""), "目前年級 Grade", edu["grade"]),
        ("班級 Class", edu.get("class", ""), "整體成績 Overall", edu.get("overall_grade", "")),
    ]

    table = doc.add_table(rows=len(edu_fields), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, (l1, v1, l2, v2) in enumerate(edu_fields):
        style_table_cell(table.cell(i, 0), l1, font_size=9, bold=True,
                          color=theme["primary"], bg_color=light_hex)
        style_table_cell(table.cell(i, 1), v1, font_size=10,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        style_table_cell(table.cell(i, 2), l2, font_size=9, bold=True,
                          color=theme["primary"], bg_color=light_hex)
        style_table_cell(table.cell(i, 3), v2, font_size=10,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ══════════════════════════════════════
    # 四、學業成績
    # ══════════════════════════════════════
    add_section_title(doc, "肆", "學業成績", "Academic Performance", theme)

    # 排名 & 趋势
    academic = data.get("academic", {})
    if academic.get("rank"):
        add_styled_paragraph(doc, f"📊 年級排名 / Ranking: {academic['rank']}",
                              font_size=11, bold=True, color=theme["primary"])
    if academic.get("trend"):
        add_styled_paragraph(doc, f"📈 成績趨勢 / Trend: {academic['trend']}",
                              font_size=10, color=RGBColor(0x55, 0x55, 0x55))

    # 成绩表
    subjects = academic.get("subjects", [])
    if subjects:
        headers = ["科目 Subject", "初一 F1", "初二 F2", "初三 F3", "備註 Remarks"]
        table = doc.add_table(rows=len(subjects) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for j, h in enumerate(headers):
            style_table_cell(table.cell(0, j), h, font_size=9, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF), bg_color=primary_hex)

        for i, subj in enumerate(subjects):
            name = subj.get("name", subj.get("name_zh", ""))
            y7 = subj.get("y7", subj.get("grade", ""))
            y8 = subj.get("y8", "")
            y9 = subj.get("y9", "")
            note = subj.get("note", "")

            style_table_cell(table.cell(i+1, 0), name, font_size=9, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
            # 高亮 A/A+
            for col_idx, grade in [(1, y7), (2, y8), (3, y9)]:
                if grade and grade.startswith("A"):
                    style_table_cell(table.cell(i+1, col_idx), grade,
                                      font_size=9, bold=True, color=theme["primary"])
                else:
                    style_table_cell(table.cell(i+1, col_idx), grade, font_size=9)
            style_table_cell(table.cell(i+1, 4), note, font_size=8,
                              color=RGBColor(0x88, 0x88, 0x88))

    # ══════════════════════════════════════
    # 五、獎項與成就
    # ══════════════════════════════════════
    add_section_title(doc, "伍", "獎項與成就", "Awards & Achievements", theme)

    awards = data.get("awards", [])
    if awards:
        table = doc.add_table(rows=len(awards) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        headers = ["獎項 Award", "頒發機構 Organization", "時間 Date", "級別 Level", "成績 Result"]
        for j, h in enumerate(headers):
            style_table_cell(table.cell(0, j), h, font_size=9, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF), bg_color=primary_hex)

        for i, award in enumerate(awards):
            style_table_cell(table.cell(i+1, 0), award.get("name", ""),
                              font_size=9, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
            style_table_cell(table.cell(i+1, 1), award.get("organizer", ""),
                              font_size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            style_table_cell(table.cell(i+1, 2), award.get("date", ""), font_size=9)
            style_table_cell(table.cell(i+1, 3), award.get("level", ""), font_size=9)
            # 突出成绩
            style_table_cell(table.cell(i+1, 4), award.get("award", ""),
                              font_size=10, bold=True, color=theme["primary"])

    # ══════════════════════════════════════
    # 六、課外活動與特長
    # ══════════════════════════════════════
    add_section_title(doc, "陸", "課外活動與特長", "Extracurricular Activities & Talents", theme)

    # 课外活动
    activities = data.get("ecActivities", [])
    if activities:
        add_styled_paragraph(doc, "🏫 課外活動 / Activities", font_size=11, bold=True,
                              color=theme["secondary"], space_after=6)
        table = doc.add_table(rows=len(activities) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        headers = ["活動 Activity", "職務 Position", "時間 Period", "描述 Description"]
        for j, h in enumerate(headers):
            style_table_cell(table.cell(0, j), h, font_size=9, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF), bg_color=primary_hex)

        for i, act in enumerate(activities):
            name = f"{act.get('name', '')} {act.get('en_name', '')}"
            pos = f"{act.get('position', '')} {act.get('en_position', '')}"
            style_table_cell(table.cell(i+1, 0), name, font_size=9, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
            style_table_cell(table.cell(i+1, 1), pos, font_size=9)
            style_table_cell(table.cell(i+1, 2), act.get("period", ""), font_size=9)
            style_table_cell(table.cell(i+1, 3),
                              act.get("description", "") or act.get("en_description", ""),
                              font_size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 爱好特长
    hobbies = data.get("hobbies", [])
    if hobbies:
        add_styled_paragraph(doc, "", font_size=6, space_after=4)
        add_styled_paragraph(doc, "🌟 愛好與特長 / Hobbies & Interests", font_size=11,
                              bold=True, color=theme["secondary"], space_after=6)

        table = doc.add_table(rows=len(hobbies) + 1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for j, h in enumerate(["愛好 Hobby", "類別 Category", "備註 Remarks"]):
            style_table_cell(table.cell(0, j), h, font_size=9, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF), bg_color=primary_hex)

        for i, hb in enumerate(hobbies):
            style_table_cell(table.cell(i+1, 0), hb.get("name", ""), font_size=9, bold=True)
            style_table_cell(table.cell(i+1, 1), hb.get("level", ""), font_size=9)
            style_table_cell(table.cell(i+1, 2), hb.get("achievement", ""), font_size=9)

    # ══════════════════════════════════════
    # 七、推薦人信息
    # ══════════════════════════════════════
    add_section_title(doc, "柒", "推薦人信息", "References", theme)

    ref = data.get("teacher_comment", {})
    if ref:
        ref_fields = [
            ("推薦人 Referee", ref.get("teacher", "")),
            ("職務 Position", ref.get("position", "")),
            ("學校 School", ref.get("school", "")),
            ("電郵 Email", ref.get("email", "")),
        ]
        if ref.get("date"):
            ref_fields.append(("日期 Date", ref.get("date", "")))
        
        table = doc.add_table(rows=len(ref_fields), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for i, (label, value) in enumerate(ref_fields):
            style_table_cell(table.cell(i, 0), label, font_size=9, bold=True,
                              color=theme["primary"], bg_color=light_hex)
            style_table_cell(table.cell(i, 1), value, font_size=10,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # 推荐评语 - 中文
        if ref.get("comment_zh"):
            add_styled_paragraph(doc, "", font_size=6, space_after=4)
            add_styled_paragraph(doc, "推薦評語 / Recommendation (中文):", font_size=10,
                                  bold=True, color=theme["secondary"])
            add_styled_paragraph(doc, f"「{ref['comment_zh']}」", font_size=10,
                                  color=RGBColor(0x44, 0x44, 0x44),
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        # 推荐评语 - 英文
        if ref.get("comment_en"):
            add_styled_paragraph(doc, "", font_size=4, space_after=4)
            add_styled_paragraph(doc, "Recommendation (English):", font_size=10,
                                  bold=True, color=theme["secondary"])
            add_styled_paragraph(doc, f'"{ref["comment_en"]}"', font_size=10,
                                  color=RGBColor(0x44, 0x44, 0x44),
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        # 兼容旧格式
        if ref.get("comment") and not ref.get("comment_zh"):
            add_styled_paragraph(doc, "", font_size=6, space_after=4)
            add_styled_paragraph(doc, "推薦評語 / Recommendation:", font_size=10,
                                  bold=True, color=theme["secondary"])
            add_styled_paragraph(doc, f"「{ref['comment']}」", font_size=10,
                                  color=RGBColor(0x44, 0x44, 0x44),
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ══════════════════════════════════════
    # 八、聲明
    # ══════════════════════════════════════
    add_section_title(doc, "捌", "聲明", "Declaration", theme)

    add_styled_paragraph(doc,
        "本人聲明以上所填資料全部屬實，如有虛假，願承擔相應責任。",
        font_size=10, color=RGBColor(0x44, 0x44, 0x44))
    add_styled_paragraph(doc,
        "I declare that all information provided above is true and accurate.",
        font_size=10, color=RGBColor(0x44, 0x44, 0x44))

    add_styled_paragraph(doc, "", font_size=12, space_after=20)

    # 签名行
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_cell(sig_table.cell(0, 0), "申請人簽名 Applicant Signature:",
                      font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.cell(0, 1), "日期 Date:",
                      font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.cell(1, 0), "家長簽名 Parent Signature:",
                      font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.cell(1, 1), "聯絡電話 Tel.:",
                      font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 保存
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"✅ 優化版升中個人檔案已生成：{output_path}")
    print(f"📊 文件大小：{size_kb:.1f} KB")
    print(f"🎨 模板風格：{theme_name}")
    return output_path


if __name__ == "__main__":
    import os
    import argparse

    parser = argparse.ArgumentParser(description="香港中學升中個人檔案生成器 v3")
    parser.add_argument("--data", required=True, help="JSON 數據文件路徑")
    parser.add_argument("--template", default="傳統名校",
                        choices=["傳統名校", "直資私立", "國際學校"],
                        help="模板風格")
    parser.add_argument("--output", default="portfolio.docx", help="輸出路徑")
    args = parser.parse_args()

    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    generate_portfolio(data, args.template, args.output)

#!/usr/bin/env python3
"""
中学生升学简历生成器 v2
支持三种模板风格：学霸型 / 特长型 / 综合型
输出 Word (.docx) 格式
"""

import argparse
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)


# ============ 颜色主题 ============
THEMES = {
    '学霸型': {
        'primary': (0, 51, 102),      # 深蓝 - 稳重学术
        'secondary': (0, 102, 153),   # 中蓝
        'accent': (0, 128, 128),      # 青色
        'bg': (240, 248, 255),        # 浅蓝背景
    },
    '特长型': {
        'primary': (128, 0, 128),     # 紫色 - 活力创意
        'secondary': (160, 32, 160),  # 中紫
        'accent': (220, 20, 60),      # 红色
        'bg': (248, 240, 255),        # 浅紫背景
    },
    '综合型': {
        'primary': (0, 100, 0),       # 深绿 - 全面均衡
        'secondary': (34, 139, 34),   # 中绿
        'accent': (0, 128, 128),      # 青色
        'bg': (240, 255, 240),        # 浅绿背景
    }
}


class StudentResume:
    """中学生升学简历"""
    
    def __init__(self, data: dict, template: str = '综合型'):
        self.data = data
        self.template = template
        self.theme = THEMES.get(template, THEMES['综合型'])
        self.doc = Document()
        self._setup_page()
    
    def _setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    def _set_font(self, run, name='微软雅黑', size=11, bold=False, color=None):
        """设置字体"""
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    
    def _set_cell_shading(self, cell, color_hex):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _rgb_to_hex(self, rgb_tuple):
        """RGB 元组转十六进制"""
        return '{:02X}{:02X}{:02X}'.format(*rgb_tuple)
    
    def _add_separator(self):
        """添加分隔线"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self._rgb_to_hex(self.theme['primary']))
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.space_after = Pt(8)
    
    # ============ 标题区 ============
    
    def _add_title(self):
        """添加简历标题"""
        name = self.data.get('basic_info', {}).get('name', '')
        school = self.data.get('basic_info', {}).get('school', '')
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(0)
        p.space_after = Pt(4)
        
        run = p.add_run("升 学 简 历")
        self._set_font(run, size=26, bold=True, color=self.theme['primary'])
        
        # 副标题
        if name:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.space_after = Pt(2)
            run2 = p2.add_run(name)
            self._set_font(run2, size=18, bold=True, color=self.theme['secondary'])
        
        if school:
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.space_after = Pt(8)
            run3 = p3.add_run(f"来自 {school}")
            self._set_font(run3, size=11, color=(100, 100, 100))
        
        self._add_separator()
    
    # ============ 基本信息区 ============
    
    def _add_basic_info(self):
        """添加基本信息"""
        info = self.data.get('basic_info', {})
        if not info:
            return
        
        self._add_section_title("基本信息")
        
        table = self.doc.add_table(rows=4, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        label_bg = self._rgb_to_hex(self.theme['bg'])
        
        # 第一行
        row = table.rows[0]
        row.cells[0].text = "姓 名"
        row.cells[1].text = info.get('name', '')
        row.cells[2].text = "性 别"
        row.cells[3].text = info.get('gender', '')
        row.cells[4].merge(row.cells[5])
        row.cells[4].text = "（照片位置）"
        self._set_cell_shading(row.cells[0], label_bg)
        self._set_cell_shading(row.cells[2], label_bg)
        
        # 第二行
        row = table.rows[1]
        row.cells[0].text = "出生日期"
        row.cells[1].text = info.get('birthday', '')
        row.cells[2].text = "民 族"
        row.cells[3].text = info.get('ethnicity', '汉族')
        row.cells[4].merge(row.cells[5])
        self._set_cell_shading(row.cells[0], label_bg)
        self._set_cell_shading(row.cells[2], label_bg)
        
        # 第三行
        row = table.rows[2]
        row.cells[0].text = "毕业学校"
        row.cells[1].merge(row.cells[2])
        row.cells[1].text = info.get('school', '')
        row.cells[3].text = "班 级"
        row.cells[4].merge(row.cells[5])
        row.cells[4].text = info.get('class', '')
        self._set_cell_shading(row.cells[0], label_bg)
        self._set_cell_shading(row.cells[3], label_bg)
        
        # 第四行
        row = table.rows[3]
        row.cells[0].text = "联系电话"
        row.cells[1].merge(row.cells[2])
        row.cells[1].text = info.get('phone', '')
        row.cells[3].text = "监护人"
        row.cells[4].text = ""
        row.cells[5].text = info.get('guardian', '')
        self._set_cell_shading(row.cells[0], label_bg)
        self._set_cell_shading(row.cells[3], label_bg)
        self._set_cell_shading(row.cells[4], label_bg)
        
        # 设置表格字体
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.space_before = Pt(2)
                    paragraph.space_after = Pt(2)
                    for run in paragraph.runs:
                        self._set_font(run, size=10)
        
        self.doc.add_paragraph()
    
    # ============ 学业成绩区 ============
    
    def _add_grades(self):
        """添加学业成绩"""
        grades = self.data.get('grades', {})
        if not grades:
            return
        
        self._add_section_title("学业成绩")
        
        # 排名信息
        ranking_parts = []
        if grades.get('ranking'):
            ranking_parts.append(f"年级排名：{grades['ranking']}")
        if grades.get('class_ranking'):
            ranking_parts.append(f"班级排名：{grades['class_ranking']}")
        if grades.get('total_students'):
            ranking_parts.append(f"年级总人数：{grades['total_students']}")
        
        if ranking_parts:
            p = self.doc.add_paragraph()
            run = p.add_run("  |  ".join(ranking_parts))
            self._set_font(run, size=11, bold=True, color=self.theme['primary'])
            p.space_after = Pt(6)
        
        # 成绩表格
        subjects = grades.get('subjects', [])
        if subjects:
            # 两列布局
            rows_needed = (len(subjects) + 1) // 2
            table = self.doc.add_table(rows=rows_needed + 1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            header_bg = self._rgb_to_hex(self.theme['primary'])
            for i, text in enumerate(["科目", "成绩/等级", "科目", "成绩/等级"]):
                table.rows[0].cells[i].text = text
                self._set_cell_shading(table.rows[0].cells[i], header_bg)
                for p in table.rows[0].cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        self._set_font(run, size=10, bold=True, color=(255, 255, 255))
            
            # 数据
            for i in range(rows_needed):
                idx1 = i * 2
                idx2 = i * 2 + 1
                if idx1 < len(subjects):
                    table.rows[i+1].cells[0].text = subjects[idx1].get('name', '')
                    table.rows[i+1].cells[1].text = subjects[idx1].get('score', '')
                if idx2 < len(subjects):
                    table.rows[i+1].cells[2].text = subjects[idx2].get('name', '')
                    table.rows[i+1].cells[3].text = subjects[idx2].get('score', '')
                
                # 交替行背景
                if i % 2 == 0:
                    for j in range(4):
                        self._set_cell_shading(table.rows[i+1].cells[j], self._rgb_to_hex(self.theme['bg']))
            
            # 设置表格字体
            for row in table.rows[1:]:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            self._set_font(run, size=10)
        
        # 成长趋势描述
        if grades.get('trend'):
            p = self.doc.add_paragraph()
            run = p.add_run(f"成长趋势：{grades['trend']}")
            self._set_font(run, size=10, color=(100, 100, 100))
        
        self.doc.add_paragraph()
    
    # ============ 竞赛获奖区 ============
    
    def _add_awards(self):
        """添加竞赛获奖"""
        awards = self.data.get('awards', [])
        if not awards:
            return
        
        self._add_section_title("竞赛获奖")
        
        if not awards:
            return
        
        table = self.doc.add_table(rows=len(awards) + 1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_bg = self._rgb_to_hex(self.theme['primary'])
        headers = ["竞赛名称", "级别", "获奖等级", "年份"]
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
            self._set_cell_shading(table.rows[0].cells[i], header_bg)
            for p in table.rows[0].cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    self._set_font(run, size=10, bold=True, color=(255, 255, 255))
        
        # 数据
        for i, award in enumerate(awards):
            table.rows[i+1].cells[0].text = award.get('name', '')
            table.rows[i+1].cells[1].text = award.get('level', '')
            table.rows[i+1].cells[2].text = award.get('award', '')
            table.rows[i+1].cells[3].text = award.get('year', '')
            
            if i % 2 == 0:
                for j in range(4):
                    self._set_cell_shading(table.rows[i+1].cells[j], self._rgb_to_hex(self.theme['bg']))
            
            for cell in table.rows[i+1].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        self._set_font(run, size=10)
        
        self.doc.add_paragraph()
    
    # ============ 特长爱好区 ============
    
    def _add_talents(self):
        """添加特长爱好"""
        talents = self.data.get('talents', [])
        if not talents:
            return
        
        self._add_section_title("特长爱好")
        
        table = self.doc.add_table(rows=len(talents) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_bg = self._rgb_to_hex(self.theme['primary'])
        headers = ["特长项目", "水平/等级", "证书/证明"]
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
            self._set_cell_shading(table.rows[0].cells[i], header_bg)
            for p in table.rows[0].cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    self._set_font(run, size=10, bold=True, color=(255, 255, 255))
        
        for i, talent in enumerate(talents):
            table.rows[i+1].cells[0].text = talent.get('name', '')
            table.rows[i+1].cells[1].text = talent.get('level', '')
            table.rows[i+1].cells[2].text = talent.get('certificate', '')
            
            if i % 2 == 0:
                for j in range(3):
                    self._set_cell_shading(table.rows[i+1].cells[j], self._rgb_to_hex(self.theme['bg']))
            
            for cell in table.rows[i+1].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        self._set_font(run, size=10)
        
        self.doc.add_paragraph()
    
    # ============ 社会实践区 ============
    
    def _add_activities(self):
        """添加社会实践"""
        activities = self.data.get('activities', [])
        if not activities:
            return
        
        self._add_section_title("社会实践")
        
        table = self.doc.add_table(rows=len(activities) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_bg = self._rgb_to_hex(self.theme['primary'])
        headers = ["活动名称", "时间", "描述"]
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
            self._set_cell_shading(table.rows[0].cells[i], header_bg)
            for p in table.rows[0].cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    self._set_font(run, size=10, bold=True, color=(255, 255, 255))
        
        for i, act in enumerate(activities):
            table.rows[i+1].cells[0].text = act.get('name', '')
            table.rows[i+1].cells[1].text = act.get('time', '')
            table.rows[i+1].cells[2].text = act.get('description', '')
            
            if i % 2 == 0:
                for j in range(3):
                    self._set_cell_shading(table.rows[i+1].cells[j], self._rgb_to_hex(self.theme['bg']))
            
            for cell in table.rows[i+1].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        self._set_font(run, size=10)
        
        self.doc.add_paragraph()
    
    # ============ 荣誉称号区 ============
    
    def _add_honors(self):
        """添加荣誉称号"""
        honors = self.data.get('honors', [])
        if not honors:
            return
        
        self._add_section_title("荣誉称号")
        
        table = self.doc.add_table(rows=len(honors) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_bg = self._rgb_to_hex(self.theme['primary'])
        headers = ["荣誉名称", "级别", "年份"]
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
            self._set_cell_shading(table.rows[0].cells[i], header_bg)
            for p in table.rows[0].cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    self._set_font(run, size=10, bold=True, color=(255, 255, 255))
        
        for i, honor in enumerate(honors):
            table.rows[i+1].cells[0].text = honor.get('name', '')
            table.rows[i+1].cells[1].text = honor.get('level', '')
            table.rows[i+1].cells[2].text = honor.get('year', '')
            
            if i % 2 == 0:
                for j in range(3):
                    self._set_cell_shading(table.rows[i+1].cells[j], self._rgb_to_hex(self.theme['bg']))
            
            for cell in table.rows[i+1].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        self._set_font(run, size=10)
        
        self.doc.add_paragraph()
    
    # ============ 自我评价区 ============
    
    def _add_self_evaluation(self):
        """添加自我评价"""
        evaluation = self.data.get('self_evaluation', '')
        if not evaluation:
            return
        
        self._add_section_title("自我评价")
        
        p = self.doc.add_paragraph()
        run = p.add_run(evaluation)
        self._set_font(run, size=11)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(22)
    
    # ============ 教师推荐区 ============
    
    def _add_teacher_comment(self):
        """添加教师推荐"""
        comment = self.data.get('teacher_comment', {})
        if not comment or not comment.get('content'):
            return
        
        self._add_section_title("教师推荐")
        
        p = self.doc.add_paragraph()
        run = p.add_run(comment.get('content', ''))
        self._set_font(run, size=11)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(22)
        
        # 签名
        if comment.get('teacher'):
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.space_before = Pt(8)
            run2 = p2.add_run(f"—— {comment['teacher']}")
            self._set_font(run2, size=10, color=(100, 100, 100))
        
        if comment.get('title'):
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run3 = p3.add_run(comment['title'])
            self._set_font(run3, size=10, color=(100, 100, 100))
    
    # ============ 章节标题 ============
    
    def _add_section_title(self, text: str):
        """添加章节标题"""
        p = self.doc.add_paragraph()
        run = p.add_run(f"▎ {text}")
        self._set_font(run, size=14, bold=True, color=self.theme['primary'])
        p.space_before = Pt(12)
        p.space_after = Pt(6)
        
        # 底部装饰线
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self._rgb_to_hex(self.theme['accent']))
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    # ============ 生成主函数 ============
    
    def generate(self):
        """生成完整简历"""
        self._add_title()
        self._add_basic_info()
        self._add_grades()
        self._add_awards()
        self._add_talents()
        self._add_activities()
        self._add_honors()
        self._add_self_evaluation()
        self._add_teacher_comment()
        
        # 页脚
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(20)
        run = p.add_run("—— 以上信息真实有效，如有不实，后果自负 ——")
        self._set_font(run, size=8, color=(150, 150, 150))
        
        return self.doc
    
    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)
        abs_path = Path(output_path).resolve()
        print(f"✅ 简历已生成：{abs_path}")
        print(f"📄 模板风格：{self.template}")
        print(f"📊 文件大小：{abs_path.stat().st_size / 1024:.1f} KB")
        return str(abs_path)


def main():
    parser = argparse.ArgumentParser(description='中学生升学简历生成器')
    parser.add_argument('--data', '-d', help='学生信息 JSON 文件路径')
    parser.add_argument('--output', '-o', default='resume.docx', help='输出文件名')
    parser.add_argument('--template', '-t', choices=['学霸型', '特长型', '综合型'], 
                        default='综合型', help='模板风格 (默认: 综合型)')
    
    args = parser.parse_args()
    
    # 读取数据
    if args.data:
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        # 示例数据
        data = {
            "basic_info": {
                "name": "张三",
                "gender": "男",
                "birthday": "2010年3月15日",
                "ethnicity": "汉族",
                "school": "XX市第一中学",
                "class": "初三(1)班",
                "phone": "138xxxx8888",
                "guardian": "张父"
            },
            "grades": {
                "ranking": "年级第8名",
                "class_ranking": "班级第2名",
                "total_students": "全年级680人",
                "trend": "成绩稳步上升，初三上学期较初二提升15个名次",
                "subjects": [
                    {"name": "语文", "score": "A (95)"},
                    {"name": "数学", "score": "A+ (98)"},
                    {"name": "英语", "score": "A (93)"},
                    {"name": "物理", "score": "A+ (99)"},
                    {"name": "化学", "score": "A (94)"},
                    {"name": "生物", "score": "A (92)"},
                    {"name": "历史", "score": "A- (88)"},
                    {"name": "地理", "score": "A (90)"}
                ]
            },
            "awards": [
                {"name": "全国初中数学竞赛", "level": "国家级", "award": "二等奖", "year": "2025"},
                {"name": "省青少年科技创新大赛", "level": "省级", "award": "一等奖", "year": "2024"},
                {"name": "市物理竞赛", "level": "市级", "award": "一等奖", "year": "2024"},
                {"name": "区英语演讲比赛", "level": "区级", "award": "特等奖", "year": "2023"}
            ],
            "talents": [
                {"name": "钢琴", "level": "十级", "certificate": "中国音乐学院考级证书"},
                {"name": "编程(Python)", "level": "NCT三级", "certificate": "全国青少年编程等级考试"},
                {"name": "篮球", "level": "校队主力", "certificate": "校际联赛冠军"}
            ],
            "activities": [
                {"name": "社区志愿服务", "time": "2024年暑期", "description": "参与社区环保宣传，累计服务30小时"},
                {"name": "学校科技社", "time": "2023-2025", "description": "担任社长，组织3次校级科技活动"},
                {"name": "省青少年科技夏令营", "time": "2024年7月", "description": "优秀营员，完成AI项目并获得表彰"}
            ],
            "honors": [
                {"name": "市级三好学生", "level": "市级", "year": "2024"},
                {"name": "优秀学生干部", "level": "校级", "year": "2023-2024"},
                {"name": "学习标兵", "level": "校级", "year": "2022-2025"}
            ],
            "self_evaluation": "性格开朗，学习刻苦，成绩优异。在班级担任学习委员，乐于帮助同学，具有良好的团队协作能力。热爱科技，自学Python编程，曾获省青少年科技创新大赛一等奖。兴趣广泛，钢琴十级，校篮球队主力。希望在高中阶段继续努力，为实现科技梦想打下坚实基础。",
            "teacher_comment": {
                "teacher": "李老师",
                "title": "班主任 / 高级教师",
                "content": "张三同学品学兼优，学习态度端正，思维活跃，善于独立思考。在班级工作中认真负责，团结同学，具有很好的领导力和团队协作精神。在科技创新方面表现突出，多次在省市竞赛中获奖。推荐该生报考贵校，相信他一定能在更高平台上绽放光彩。"
            }
        }
    
    # 生成简历
    resume = StudentResume(data, template=args.template)
    resume.generate()
    resume.save(args.output)


if __name__ == "__main__":
    main()

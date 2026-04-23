#!/usr/bin/env python3
import json
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import Counter

# Data from the API
data = [
    {"f_6aSwE2": "6", "f_iKayRs": "A小区门口发生斗殴事件421121200210236332", "f_w8ZsC0": "家庭纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "出轨", "f_9TQx12": "2024-11-01T10:00:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "4", "f_iKayRs": "A小区门口发生斗殴事件421121200210236332", "f_w8ZsC0": "家庭纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "出轨", "f_9TQx12": "2024-11-01T10:00:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "10", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "3", "f_iKayRs": "A小区门口发生斗殴事件421121200210236332", "f_w8ZsC0": "家庭纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "出轨", "f_9TQx12": "2024-11-01T10:00:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "8", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "11", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "9", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "5", "f_iKayRs": "A小区门口发生斗殴事件421121200210236332", "f_w8ZsC0": "家庭纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "出轨", "f_9TQx12": "2024-11-01T10:00:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "7", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "12", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "1", "f_iKayRs": "A小区门口发生斗殴事件421121200210236332", "f_w8ZsC0": "家庭纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "出轨", "f_9TQx12": "2024-11-01T10:00:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "1", "f_iKayRs": "A小区门口发生斗殴事件421121200210236332", "f_w8ZsC0": "家庭纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "出轨", "f_9TQx12": "2024-11-01T10:00:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "2", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
    {"f_6aSwE2": "2", "f_iKayRs": "B小区内夫妻争吵421121200210236333", "f_w8ZsC0": "婚姻纠纷", "f_EH9INq": "武汉市公安局", "f_Wkv9nZ": "家暴", "f_9TQx12": "2024-11-02T12:30:00", "f_fpkIwb": "110"},
]

# Analyze data
total_records = len(data)
dispute_types = Counter([d['f_w8ZsC0'] for d in data])
causes = Counter([d['f_Wkv9nZ'] for d in data])
locations = Counter([d['f_iKayRs'] for d in data])
dates = Counter([d['f_9TQx12'][:10] for d in data])
organizations = Counter([d['f_EH9INq'] for d in data])
phones = Counter([d['f_fpkIwb'] for d in data])

# Create Word document
doc = Document()

# Title
title = doc.add_heading('出警信息表数据分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add date
doc.add_paragraph(f'报告生成日期: 2024-11-03')
doc.add_paragraph('')

# 1. Data Overview
doc.add_heading('一、数据概览', level=1)
doc.add_paragraph(f'数据总条数: {total_records} 条')

# 2. Dispute Type Analysis
doc.add_heading('二、纠纷类型分析', level=1)
for dtype, count in dispute_types.items():
    percentage = (count / total_records) * 100
    doc.add_paragraph(f'• {dtype}: {count} 条 ({percentage:.1f}%)')

# 3. Cause Analysis  
doc.add_heading('三、纠纷原因分析', level=1)
for cause, count in causes.items():
    percentage = (count / total_records) * 100
    doc.add_paragraph(f'• {cause}: {count} 条 ({percentage:.1f}%)')

# 4. Location Analysis
doc.add_heading('四、发生地点分析', level=1)
for loc, count in locations.items():
    percentage = (count / total_records) * 100
    doc.add_paragraph(f'• {loc}: {count} 条 ({percentage:.1f}%)')

# 5. Time Analysis
doc.add_heading('五、时间分布', level=1)
for date, count in dates.items():
    percentage = (count / total_records) * 100
    doc.add_paragraph(f'• {date}: {count} 条 ({percentage:.1f}%)')

# 6. Organization & Contact
doc.add_heading('六、处置单位与联系方式', level=1)
for org, count in organizations.items():
    doc.add_paragraph(f'• 处置单位: {org}')
for phone, count in phones.items():
    doc.add_paragraph(f'• 联系电话: {phone}')

# 7. Summary
doc.add_heading('七、总结', level=1)
summary = doc.add_paragraph()
summary.add_run(f'本次分析共涉及 {total_records} 条出警记录。\n')
summary.add_run(f'纠纷类型方面，')
if '婚姻纠纷' in dispute_types and '家庭纠纷' in dispute_types:
    if dispute_types['婚姻纠纷'] > dispute_types['家庭纠纷']:
        summary.add_run(f'婚姻纠纷占多数（共{dispute_types["婚姻纠纷"]}条），家庭纠纷有{dispute_types["家庭纠纷"]}条。\n')
    else:
        summary.add_run(f'家庭纠纷占多数（共{dispute_types["家庭纠纷"]}条），婚姻纠纷有{dispute_types["婚姻纠纷"]}条。\n')
summary.add_run(f'从时间分布来看，')
if '2024-11-01' in dates and '2024-11-02' in dates:
    if dates['2024-11-02'] > dates['2024-11-01']:
        summary.add_run(f'11月2日出警数量较多（{dates["2024-11-02"]}条），11月1日有{dates["2024-11-01"]}条。\n')
    else:
        summary.add_run(f'11月1日出警数量较多（{dates["2024-11-01"]}条），11月2日有{dates["2024-11-02"]}条。\n')
summary.add_run('所有警情均由武汉市公安局处置，联系电话为110。')

# 8. Raw Data Table
doc.add_heading('八、原始数据', level=1)
table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'

# Header row
headers = ['ID', '地点/事件', '纠纷类型', '处置单位', '原因', '时间', '电话']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

# Data rows
for record in data:
    row = table.add_row()
    row.cells[0].text = record['f_6aSwE2']
    row.cells[1].text = record['f_iKayRs']
    row.cells[2].text = record['f_w8ZsC0']
    row.cells[3].text = record['f_EH9INq']
    row.cells[4].text = record['f_Wkv9nZ']
    row.cells[5].text = record['f_9TQx12']
    row.cells[6].text = record['f_fpkIwb']

# Save
doc.save('/root/.openclaw/workspace/分析报告.docx')
print('Report saved to /root/.openclaw/workspace/分析报告.docx')
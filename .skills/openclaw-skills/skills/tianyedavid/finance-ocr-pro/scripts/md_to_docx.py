"""
Markdown to DOCX Converter

Converts Markdown files (with HTML tables and LaTeX formulas)
into professionally styled A4 DOCX documents with multilingual support.

Handles:
    - Headings (# through ######)
    - Paragraphs with inline formatting (bold, italic, code)
    - Inline LaTeX ($...$) and block LaTeX ($$...$$)
    - HTML tables with colspan/rowspan and LaTeX in cells
    - Ordered and unordered lists
    - Fenced code blocks (```language ... ```)
    - Embedded images (when source file exists)

Requirements:
    pip install python-docx beautifulsoup4 lxml latex2mathml

Usage:
    from pathlib import Path
    markdown_to_docx(Path("input.md"), Path("output.docx"))
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString, Tag
import latex2mathml.converter
from lxml import etree


# ═══════════════════════════════════════════════════════════════════════════════
# Configuration
# ═══════════════════════════════════════════════════════════════════════════════

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "{" + OMML_NS + "}"
W = "{" + WORD_NS + "}"

FONT_BODY = "Arial"
FONT_MATH = "Cambria Math"
FONT_CODE = "Consolas"

FONT_SIZE_BODY = Pt(11)
FONT_SIZE_TABLE = Pt(9.5)
FONT_SIZE_CODE = Pt(10)

COLOR_HEADING = RGBColor(0x1F, 0x49, 0x7D)
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)

TABLE_HEADER_BG = "D6E4F0"
TABLE_ALT_ROW_BG = "F5F8FC"
TABLE_BORDER_COLOR = "A6A6A6"
TABLE_CELL_MARGIN_DXA = 72

CODE_BG_COLOR = "F2F2F2"
CODE_BORDER_COLOR = "CCCCCC"

MARGIN_NARROW = Cm(1.8)

HEADING_SIZES = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}


# ═══════════════════════════════════════════════════════════════════════════════
# Font Utilities
# ═══════════════════════════════════════════════════════════════════════════════

def _set_rfonts(rpr_element, font_name: str):
    """Set w:rFonts for all four script types (Latin, East Asian, Complex Script)."""
    rf = rpr_element.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr_element.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(attr), font_name)


def set_run_fonts(run, font_name: str = FONT_BODY, size=FONT_SIZE_BODY,
                  bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None):
    """Configure a run's font for full multilingual compatibility."""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    _set_rfonts(run._r.get_or_add_rPr(), font_name)


# ═══════════════════════════════════════════════════════════════════════════════
# Preprocessing
# ═══════════════════════════════════════════════════════════════════════════════

_PAGE_ORDER_RE = re.compile(r'^Page_Order_\d+$', re.IGNORECASE | re.MULTILINE)
_PAGE_NUMBER_RE = re.compile(r'Page\s+Number\s+\d+\s*:\s*\n?', re.IGNORECASE)
_PAGE_SEP_RE = re.compile(r'\n\s*---\s*\n')
_MULTI_BLANK_RE = re.compile(r'\n{3,}')


def preprocess_content(content: str) -> str:
    """Remove OCR noise, separators, and collapse excess blank lines.

    ``Page_Order_N`` sentinels are preserved so the document builder can
    insert page breaks at the correct positions.
    """
    content = _PAGE_NUMBER_RE.sub('', content)
    content = _PAGE_SEP_RE.sub('\n\n', content)
    content = _MULTI_BLANK_RE.sub('\n\n', content)
    return content.strip()


# ═══════════════════════════════════════════════════════════════════════════════
# MathML to OMML Conversion
# ═══════════════════════════════════════════════════════════════════════════════

def _omml_run(text: str) -> etree._Element:
    """Create an OMML run element containing the given text."""
    r = etree.Element(M + "r")
    t = etree.SubElement(r, M + "t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text if text else ""
    return r


def _tag_local(element) -> str:
    """Strip namespace URI from an element's tag name."""
    tag = element.tag
    return tag.split('}')[1] if '}' in tag else tag


def _convert_children(mml_el, parent_omml):
    """Recursively convert all children of a MathML element to OMML."""
    if mml_el.text and mml_el.text.strip():
        parent_omml.append(_omml_run(mml_el.text.strip()))
    for child in mml_el:
        _mml_to_omml(child, parent_omml)
        if child.tail and child.tail.strip():
            parent_omml.append(_omml_run(child.tail.strip()))


_MML_PASSTHROUGH_TAGS = frozenset(('mstyle', 'mpadded', 'mphantom'))
_MML_TEXT_TAGS = frozenset(('mn', 'mo', 'mtext'))


def _mml_to_omml(mml_el, parent_omml):
    """Recursively convert a single MathML element to its OMML equivalent."""
    tag = _tag_local(mml_el)
    text = (mml_el.text or "").strip()
    children = list(mml_el)

    if tag == 'math':
        _convert_children(mml_el, parent_omml)

    elif tag == 'semantics':
        for child in children:
            if not _tag_local(child).startswith('annotation'):
                _mml_to_omml(child, parent_omml)
                break

    elif tag == 'mrow':
        _convert_children(mml_el, parent_omml)

    elif tag == 'mi':
        r = _omml_run(text)
        if text and len(text) == 1 and text.isalpha():
            rPr = etree.Element(M + "rPr")
            sty = etree.SubElement(rPr, M + "sty")
            sty.set(M + "val", "i")
            r.insert(0, rPr)
        parent_omml.append(r)

    elif tag in _MML_TEXT_TAGS:
        parent_omml.append(_omml_run(text))

    elif tag == 'mspace':
        parent_omml.append(_omml_run(" "))

    elif tag == 'mfrac':
        f = etree.SubElement(parent_omml, M + "f")
        etree.SubElement(f, M + "fPr")
        num = etree.SubElement(f, M + "num")
        den = etree.SubElement(f, M + "den")
        if len(children) >= 1:
            _mml_to_omml(children[0], num)
        if len(children) >= 2:
            _mml_to_omml(children[1], den)

    elif tag == 'msup':
        ss = etree.SubElement(parent_omml, M + "sSup")
        etree.SubElement(ss, M + "sSupPr")
        e = etree.SubElement(ss, M + "e")
        sup = etree.SubElement(ss, M + "sup")
        if len(children) >= 1:
            _mml_to_omml(children[0], e)
        if len(children) >= 2:
            _mml_to_omml(children[1], sup)

    elif tag == 'msub':
        ss = etree.SubElement(parent_omml, M + "sSub")
        etree.SubElement(ss, M + "sSubPr")
        e = etree.SubElement(ss, M + "e")
        sub = etree.SubElement(ss, M + "sub")
        if len(children) >= 1:
            _mml_to_omml(children[0], e)
        if len(children) >= 2:
            _mml_to_omml(children[1], sub)

    elif tag == 'msubsup':
        ss = etree.SubElement(parent_omml, M + "sSubSup")
        etree.SubElement(ss, M + "sSubSupPr")
        e = etree.SubElement(ss, M + "e")
        sub = etree.SubElement(ss, M + "sub")
        sup = etree.SubElement(ss, M + "sup")
        if len(children) >= 1:
            _mml_to_omml(children[0], e)
        if len(children) >= 2:
            _mml_to_omml(children[1], sub)
        if len(children) >= 3:
            _mml_to_omml(children[2], sup)

    elif tag == 'msqrt':
        rad = etree.SubElement(parent_omml, M + "rad")
        radPr = etree.SubElement(rad, M + "radPr")
        dh = etree.SubElement(radPr, M + "degHide")
        dh.set(M + "val", "1")
        etree.SubElement(rad, M + "deg")
        e = etree.SubElement(rad, M + "e")
        _convert_children(mml_el, e)

    elif tag == 'mroot':
        rad = etree.SubElement(parent_omml, M + "rad")
        etree.SubElement(rad, M + "radPr")
        deg = etree.SubElement(rad, M + "deg")
        e = etree.SubElement(rad, M + "e")
        if len(children) >= 1:
            _mml_to_omml(children[0], e)
        if len(children) >= 2:
            _mml_to_omml(children[1], deg)

    elif tag == 'mfenced':
        d = etree.SubElement(parent_omml, M + "d")
        dPr = etree.SubElement(d, M + "dPr")
        beg = etree.SubElement(dPr, M + "begChr")
        beg.set(M + "val", mml_el.get('open', '('))
        end = etree.SubElement(dPr, M + "endChr")
        end.set(M + "val", mml_el.get('close', ')'))
        e = etree.SubElement(d, M + "e")
        _convert_children(mml_el, e)

    elif tag == 'mtable':
        m_el = etree.SubElement(parent_omml, M + "m")
        etree.SubElement(m_el, M + "mPr")
        for child in children:
            if _tag_local(child) == 'mtr':
                mr = etree.SubElement(m_el, M + "mr")
                for cell in child:
                    if _tag_local(cell) == 'mtd':
                        e = etree.SubElement(mr, M + "e")
                        _convert_children(cell, e)

    elif tag == 'mover':
        acc = etree.SubElement(parent_omml, M + "acc")
        accPr = etree.SubElement(acc, M + "accPr")
        if len(children) >= 2:
            ch = etree.SubElement(accPr, M + "chr")
            ch.set(M + "val", children[1].text if children[1].text else "^")
        e = etree.SubElement(acc, M + "e")
        if len(children) >= 1:
            _mml_to_omml(children[0], e)

    elif tag == 'munder':
        ll = etree.SubElement(parent_omml, M + "limLow")
        etree.SubElement(ll, M + "limLowPr")
        e = etree.SubElement(ll, M + "e")
        lim = etree.SubElement(ll, M + "lim")
        if len(children) >= 1:
            _mml_to_omml(children[0], e)
        if len(children) >= 2:
            _mml_to_omml(children[1], lim)

    elif tag == 'munderover':
        nary = etree.SubElement(parent_omml, M + "nary")
        naryPr = etree.SubElement(nary, M + "naryPr")
        if len(children) >= 1:
            first = children[0]
            op = first.text
            if not op and len(list(first)) > 0:
                op = list(first)[0].text
            if op:
                ch = etree.SubElement(naryPr, M + "chr")
                ch.set(M + "val", op)
        limLoc = etree.SubElement(naryPr, M + "limLoc")
        limLoc.set(M + "val", "undOvr")
        sub = etree.SubElement(nary, M + "sub")
        sup = etree.SubElement(nary, M + "sup")
        e = etree.SubElement(nary, M + "e")
        if len(children) >= 2:
            _mml_to_omml(children[1], sub)
        if len(children) >= 3:
            _mml_to_omml(children[2], sup)

    elif tag in _MML_PASSTHROUGH_TAGS:
        _convert_children(mml_el, parent_omml)

    else:
        if text:
            parent_omml.append(_omml_run(text))
        _convert_children(mml_el, parent_omml)


def latex_to_omml(latex_str: str) -> etree._Element | None:
    """Convert a LaTeX string to an OMML <m:oMath> element, or None on failure."""
    try:
        latex_str = latex_str.strip()
        if not latex_str:
            return None
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_str = re.sub(r'\s+xmlns(:[a-z]+)?="[^"]*"', '', mathml_str)
        mml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omath = etree.Element(M + "oMath")
        _mml_to_omml(mml_tree, omath)
        return omath
    except Exception as exc:
        preview = latex_str[:80] + ('...' if len(latex_str) > 80 else '')
        print(f"  [LaTeX warning] {exc} — expression: {preview}")
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# Inline Content Parsing
# ═══════════════════════════════════════════════════════════════════════════════

_INLINE_RE = re.compile(
    r'(\$\$[\s\S]*?\$\$)'
    r'|(\$(?!\$)[^\$\n]+?\$)'
    r'|(\*\*\*(.+?)\*\*\*)'
    r'|(\*\*(.+?)\*\*)'
    r'|(\*(?!\*)(.+?)(?<!\*)\*)'
    r'|(`[^`\n]+?`)'
)


def parse_inline_content(text: str) -> list[dict]:
    """
    Split text into typed segments for rendering.

    Segment types: text, bold, bold_italic, code, latex_inline, latex_block.
    """
    if not text:
        return []

    parts: list[dict] = []
    last = 0

    for m in _INLINE_RE.finditer(text):
        if m.start() > last:
            parts.append({'type': 'text', 'content': text[last:m.start()]})

        s = m.group(0)
        if s.startswith('$$') and s.endswith('$$'):
            c = s[2:-2].strip()
            if c:
                parts.append({'type': 'latex_block', 'content': c})
        elif s.startswith('$') and s.endswith('$'):
            c = s[1:-1].strip()
            if c:
                parts.append({'type': 'latex_inline', 'content': c})
        elif s.startswith('***') and s.endswith('***'):
            parts.append({'type': 'bold_italic', 'content': s[3:-3]})
        elif s.startswith('**') and s.endswith('**'):
            parts.append({'type': 'bold', 'content': s[2:-2]})
        elif s.startswith('*') and s.endswith('*') and not s.startswith('**'):
            parts.append({'type': 'italic', 'content': s[1:-1]})
        elif s.startswith('`') and s.endswith('`'):
            parts.append({'type': 'code', 'content': s[1:-1]})

        last = m.end()

    if last < len(text):
        parts.append({'type': 'text', 'content': text[last:]})

    if not parts and text:
        parts.append({'type': 'text', 'content': text})

    return parts


# ═══════════════════════════════════════════════════════════════════════════════
# HTML Table Parsing
# ═══════════════════════════════════════════════════════════════════════════════

def _cell_text(cell_tag) -> str:
    """Extract all text from a BeautifulSoup table cell, collapsing whitespace."""
    fragments: list[str] = []

    def _walk(node):
        if isinstance(node, NavigableString):
            t = re.sub(r'\s+', ' ', str(node))
            if t:
                fragments.append(t)
        elif isinstance(node, Tag):
            if node.name == 'br':
                fragments.append('\n')
            for child in node.children:
                _walk(child)

    _walk(cell_tag)
    return ''.join(fragments).strip()


def parse_html_table(table_html: str) -> tuple[list[dict], int, int]:
    """
    Parse an HTML table into cell data with merge information.

    Returns (cells_data, num_rows, num_cols).
    Each cell dict has: row, col, rowspan, colspan, content (list of segments),
    is_header (bool).
    """
    soup = BeautifulSoup(table_html, 'lxml')
    table = soup.find('table')
    if not table:
        return [], 0, 0

    rows = table.find_all('tr')
    if not rows:
        return [], 0, 0

    num_rows = len(rows)
    grid: dict[tuple[int, int], bool] = {}
    max_cols = 0

    for ri, row in enumerate(rows):
        ci = 0
        for cell in row.find_all(['td', 'th']):
            while (ri, ci) in grid:
                ci += 1
            cs = int(cell.get('colspan', 1))
            rs = int(cell.get('rowspan', 1))
            for r in range(ri, ri + rs):
                for c in range(ci, ci + cs):
                    grid[(r, c)] = True
            ci += cs
        max_cols = max(max_cols, ci)

    cells_data: list[dict] = []
    grid = {}

    for ri, row in enumerate(rows):
        ci = 0
        for cell in row.find_all(['td', 'th']):
            while (ri, ci) in grid:
                ci += 1
            cs = int(cell.get('colspan', 1))
            rs = int(cell.get('rowspan', 1))
            content_text = _cell_text(cell)
            cells_data.append({
                'row': ri,
                'col': ci,
                'rowspan': rs,
                'colspan': cs,
                'content': parse_inline_content(content_text),
                'is_header': cell.name == 'th',
            })
            for r in range(ri, ri + rs):
                for c in range(ci, ci + cs):
                    grid[(r, c)] = True
            ci += cs

    return cells_data, num_rows, max_cols


# ═══════════════════════════════════════════════════════════════════════════════
# Table Rendering in DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def _set_table_borders(table, color: str = TABLE_BORDER_COLOR, size: str = "4"):
    """Apply uniform thin borders to all table edges."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

    borders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tblPr.append(borders)


def _set_cell_shading(cell, hex_color: str):
    """Set the background fill color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, dxa: int = TABLE_CELL_MARGIN_DXA):
    """Set uniform internal margins (padding) for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(dxa))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _set_cell_vertical_alignment(cell, val: str = "center"):
    """Vertically align cell content (top, center, bottom)."""
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    tcPr.append(va)


def _render_segments(para, segments: list[dict], *,
                     font_size=FONT_SIZE_BODY, bold: bool = False,
                     color: RGBColor | None = None):
    """Render a list of inline segments (text, bold, LaTeX, code) into a paragraph."""
    for seg in segments:
        stype = seg['type']

        if stype == 'text':
            run = para.add_run(seg['content'])
            set_run_fonts(run, size=font_size, bold=bold, color=color)

        elif stype == 'bold':
            run = para.add_run(seg['content'])
            set_run_fonts(run, size=font_size, bold=True, color=color)

        elif stype == 'italic':
            run = para.add_run(seg['content'])
            set_run_fonts(run, size=font_size, italic=True, color=color)

        elif stype == 'bold_italic':
            run = para.add_run(seg['content'])
            set_run_fonts(run, size=font_size, bold=True, italic=True, color=color)

        elif stype == 'code':
            run = para.add_run(seg['content'])
            set_run_fonts(run, font_name=FONT_CODE, size=FONT_SIZE_CODE, color=color)

        elif stype in ('latex_inline', 'latex_block'):
            omml = latex_to_omml(seg['content'])
            if omml is not None:
                para._p.append(omml)
            else:
                run = para.add_run(f"${seg['content']}$")
                set_run_fonts(run, size=font_size, italic=True, color=color)


def create_docx_table(doc: Document, cells_data: list[dict],
                      num_rows: int, num_cols: int) -> None:
    """Build a professionally styled DOCX table from parsed HTML cell data."""
    if num_rows == 0 or num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')
    tblPr.append(tblW)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)

    _set_table_borders(table)

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    processed: set = set()

    for cd in cells_data:
        ri, ci = cd['row'], cd['col']
        if (ri, ci) in processed:
            continue

        try:
            cell = table.cell(ri, ci)
        except IndexError:
            continue

        rs, cs = cd['rowspan'], cd['colspan']
        if rs > 1 or cs > 1:
            er = min(ri + rs - 1, num_rows - 1)
            ec = min(ci + cs - 1, num_cols - 1)
            try:
                cell.merge(table.cell(er, ec))
                for r in range(ri, er + 1):
                    for c in range(ci, ec + 1):
                        processed.add((r, c))
            except Exception:
                pass

        processed.add((ri, ci))

        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        para.clear()
        _render_segments(para, cd['content'],
                         font_size=FONT_SIZE_TABLE, bold=cd['is_header'])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if cd['is_header']:
            _set_cell_shading(cell, TABLE_HEADER_BG)
        elif ri % 2 == 1:
            _set_cell_shading(cell, TABLE_ALT_ROW_BG)

        _set_cell_margins(cell)
        _set_cell_vertical_alignment(cell)

    for row in table.rows:
        row.height = Twips(360)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


# ═══════════════════════════════════════════════════════════════════════════════
# Document Styling & Page Layout
# ═══════════════════════════════════════════════════════════════════════════════

def _setup_styles(doc: Document) -> None:
    """Configure Normal and Heading styles for a professional business look."""
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal.font.size = FONT_SIZE_BODY
    normal.font.color.rgb = COLOR_BODY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    _set_rfonts(normal.element.get_or_add_rPr(), FONT_BODY)

    for level, size in HEADING_SIZES.items():
        name = f'Heading {level}'
        if name not in doc.styles:
            continue
        hs = doc.styles[name]
        hs.font.name = FONT_BODY
        hs.font.size = Pt(size)
        hs.font.bold = (level <= 5)
        hs.font.color.rgb = COLOR_HEADING
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(8)
        hs.paragraph_format.line_spacing = 1.2
        _set_rfonts(hs.element.get_or_add_rPr(), FONT_BODY)


def _setup_page(doc: Document) -> None:
    """Set A4 page size with narrow margins."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = MARGIN_NARROW
        section.right_margin = MARGIN_NARROW
        section.top_margin = MARGIN_NARROW
        section.bottom_margin = MARGIN_NARROW


# ═══════════════════════════════════════════════════════════════════════════════
# Markdown Element Parsing
# ═══════════════════════════════════════════════════════════════════════════════

_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$')
_UL_RE = re.compile(r'^[\-\*\+□☐]\s+(.+)$')
_OL_RE = re.compile(r'^\d+[\.\)]\s+(.+)$')
_IMG_RE = re.compile(r'^!\[([^\]]*)\]\(([^\)]+)\)')
_FENCE_RE = re.compile(r'^(`{3,}|~{3,})\s*(.*)')
_PAGE_BREAK_RE = re.compile(r'^Page_Order_\d+$', re.IGNORECASE)


def parse_markdown_elements(content: str) -> list[dict]:
    """
    Parse preprocessed markdown text into a flat list of typed elements.

    Element types: heading, paragraph, table, latex_block,
                   unordered_list, ordered_list, image, code_block.
    """
    elements: list[dict] = []
    lines = content.split('\n')
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # ── Page break sentinel from OCR page markers ──
        if _PAGE_BREAK_RE.match(stripped):
            elements.append({'type': 'page_break'})
            i += 1
            continue

        # ── HTML table ──
        if '<table' in line.lower():
            buf = [line]
            depth = line.lower().count('<table') - line.lower().count('</table>')
            i += 1
            while i < n and depth > 0:
                depth += lines[i].lower().count('<table')
                depth -= lines[i].lower().count('</table>')
                buf.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'content': '\n'.join(buf)})
            continue

        # ── Fenced code block (``` or ~~~) ──
        fence_m = _FENCE_RE.match(stripped)
        if fence_m:
            fence_char = fence_m.group(1)[0]
            fence_len = len(fence_m.group(1))
            lang = fence_m.group(2).strip().lower()
            buf = []
            i += 1
            while i < n:
                close_m = _FENCE_RE.match(lines[i].strip())
                if close_m and close_m.group(1)[0] == fence_char and len(close_m.group(1)) >= fence_len:
                    i += 1
                    break
                buf.append(lines[i])
                i += 1
            code_content = '\n'.join(buf)
            elements.append({'type': 'code_block', 'content': code_content, 'language': lang})
            continue

        # ── Heading ──
        hm = _HEADING_RE.match(stripped)
        if hm:
            elements.append({
                'type': 'heading',
                'level': len(hm.group(1)),
                'content': hm.group(2).strip(),
            })
            i += 1
            continue

        # ── Block LaTeX $$...$$ ──
        if stripped.startswith('$$'):
            if stripped.endswith('$$') and len(stripped) > 4:
                elements.append({'type': 'latex_block', 'content': stripped[2:-2].strip()})
                i += 1
                continue
            buf = [line]
            i += 1
            while i < n:
                buf.append(lines[i])
                if lines[i].strip().endswith('$$'):
                    i += 1
                    break
                i += 1
            joined = '\n'.join(buf)
            m = re.search(r'\$\$([\s\S]*?)\$\$', joined)
            if m:
                elements.append({'type': 'latex_block', 'content': m.group(1).strip()})
            continue

        # ── Unordered list ──
        um = _UL_RE.match(stripped)
        if um:
            items = []
            while i < n:
                raw_line = lines[i]
                ls = raw_line.strip()
                m = _UL_RE.match(ls)
                if m:
                    items.append(m.group(1))
                    i += 1
                elif ls and (raw_line.startswith('  ') or raw_line.startswith('\t')):
                    if items:
                        items[-1] += ' ' + ls
                    i += 1
                else:
                    break
            if items:
                elements.append({'type': 'unordered_list', 'items': items})
            continue

        # ── Ordered list ──
        om = _OL_RE.match(stripped)
        if om:
            items = []
            while i < n:
                ls = lines[i].strip()
                m = _OL_RE.match(ls)
                if not m:
                    break
                items.append(m.group(1))
                i += 1
            if items:
                elements.append({'type': 'ordered_list', 'items': items})
            continue

        # ── Image ──
        im = _IMG_RE.match(stripped)
        if im:
            elements.append({'type': 'image', 'alt': im.group(1), 'src': im.group(2)})
            i += 1
            continue

        # ── Paragraph (may span consecutive non-blank lines) ──
        buf = [line]
        i += 1
        while i < n:
            ns = lines[i].strip()
            if (not ns
                    or ns.startswith('#')
                    or ns.startswith('$$')
                    or _UL_RE.match(ns)
                    or _OL_RE.match(ns)
                    or '<table' in lines[i].lower()
                    or _IMG_RE.match(ns)
                    or _FENCE_RE.match(ns)
                    or _PAGE_BREAK_RE.match(ns)):
                break
            buf.append(lines[i])
            i += 1
        para_text = ' '.join(l.strip() for l in buf if l.strip())
        if para_text:
            elements.append({'type': 'paragraph', 'content': para_text})

    return elements


# ═══════════════════════════════════════════════════════════════════════════════
# Document Building Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _add_paragraph(doc: Document, text: str) -> None:
    """Add a body paragraph with mixed text / LaTeX / bold / code segments."""
    para = doc.add_paragraph()
    _render_segments(para, parse_inline_content(text))


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Add a heading, supporting LaTeX content within the heading text."""
    segments = parse_inline_content(text)
    has_special = any(s['type'] != 'text' for s in segments)

    if not has_special:
        h = doc.add_heading(text, level=min(level, 6))
        for run in h.runs:
            set_run_fonts(run, size=Pt(HEADING_SIZES.get(level, 11)),
                          bold=(level <= 5), color=COLOR_HEADING)
        return

    para = doc.add_paragraph()
    para.style = f'Heading {min(level, 6)}'
    _render_segments(para, segments,
                     font_size=Pt(HEADING_SIZES.get(level, 11)),
                     bold=(level <= 5), color=COLOR_HEADING)


def _add_block_equation(doc: Document, latex_content: str) -> None:
    """Add a centered block-display LaTeX equation."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    omml = latex_to_omml(latex_content)
    if omml is not None:
        omp = etree.Element(M + "oMathPara")
        ompp = etree.SubElement(omp, M + "oMathParaPr")
        jc = etree.SubElement(ompp, M + "jc")
        jc.set(M + "val", "center")
        omp.append(omml)
        para._p.append(omp)
    else:
        run = para.add_run(f"$${latex_content}$$")
        set_run_fonts(run, font_name=FONT_MATH, italic=True)


def _add_list(doc: Document, items: list[str], ordered: bool = False) -> None:
    """Add ordered or unordered list items, each supporting inline LaTeX/bold."""
    style_name = 'List Number' if ordered else 'List Bullet'
    for item_text in items:
        para = doc.add_paragraph(style=style_name)
        _render_segments(para, parse_inline_content(item_text))


def _add_image(doc: Document, md_dir: Path, alt: str, src: str) -> None:
    """Embed an image if the file exists, otherwise insert a placeholder."""
    img_path = None
    src_p = Path(src)

    if src_p.is_absolute() and src_p.exists():
        img_path = src_p
    else:
        candidate = md_dir / src
        if candidate.exists():
            img_path = candidate

    if img_path and img_path.exists():
        try:
            usable_width = Cm(21.0 - 2 * 1.8)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(img_path), width=usable_width)
            return
        except Exception:
            pass

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[Image: {alt or src}]")
    set_run_fonts(run, size=Pt(10), italic=True, color=COLOR_MUTED)


def _set_paragraph_shading(para, hex_color: str):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _set_paragraph_borders(para, hex_color: str = CODE_BORDER_COLOR,
                           sides: tuple[str, ...] = ('top', 'left', 'bottom', 'right')):
    """Apply a thin border to specified sides of a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), hex_color)
        pBdr.append(b)
    pPr.append(pBdr)


def _add_code_block(doc: Document, code: str, language: str = "") -> None:
    """Render a fenced code block with monospace font and shaded background."""
    code_lines = code.split('\n')
    total = len(code_lines)

    for line_idx, code_line in enumerate(code_lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

        _set_paragraph_shading(para, CODE_BG_COLOR)

        sides = ['left', 'right']
        if line_idx == 0:
            sides.append('top')
            para.paragraph_format.space_before = Pt(4)
        if line_idx == total - 1:
            sides.append('bottom')
            para.paragraph_format.space_after = Pt(4)
        _set_paragraph_borders(para, sides=tuple(sides))

        run = para.add_run(code_line if code_line else " ")
        set_run_fonts(run, font_name=FONT_CODE, size=FONT_SIZE_CODE)


# ═══════════════════════════════════════════════════════════════════════════════
# Main Conversion
# ═══════════════════════════════════════════════════════════════════════════════

def markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    """
    Convert a Markdown file to a professionally styled DOCX document.

    The Markdown may contain:
      - Headers (# through ######)
      - Regular paragraphs with inline LaTeX / bold / code
      - Block LaTeX equations ($$...$$)
      - HTML tables with colspan / rowspan and LaTeX in cells
      - Ordered and unordered lists
      - Fenced code blocks (```language ... ```)
      - Image references ![alt](path)

    OCR artifacts (page separators, page markers) are automatically removed.

    Args:
        markdown_path: Path to the source .md file.
        output_path:   Path for the output .docx file.
    """
    content = Path(markdown_path).read_text(encoding='utf-8')
    content = preprocess_content(content)

    doc = Document()
    _setup_styles(doc)
    _setup_page(doc)

    md_dir = Path(markdown_path).parent
    elements = parse_markdown_elements(content)

    seen_page_break = False
    for elem in elements:
        t = elem['type']

        if t == 'page_break':
            if seen_page_break:
                para = doc.add_paragraph()
                para.add_run().add_break(WD_BREAK.PAGE)
            seen_page_break = True
            continue

        if t == 'heading':
            _add_heading(doc, elem['content'], elem['level'])

        elif t == 'paragraph':
            _add_paragraph(doc, elem['content'])

        elif t == 'table':
            cells, nr, nc = parse_html_table(elem['content'])
            if nr > 0 and nc > 0:
                create_docx_table(doc, cells, nr, nc)
                doc.add_paragraph()

        elif t == 'latex_block':
            _add_block_equation(doc, elem['content'])

        elif t == 'unordered_list':
            _add_list(doc, elem['items'], ordered=False)

        elif t == 'ordered_list':
            _add_list(doc, elem['items'], ordered=True)

        elif t == 'image':
            _add_image(doc, md_dir, elem.get('alt', ''), elem.get('src', ''))

        elif t == 'code_block':
            _add_code_block(doc, elem['content'], elem.get('language', ''))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"Converted: {markdown_path} → {output_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# CLI Entry Point
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(
        description="Convert Markdown (with HTML tables, LaTeX) to DOCX.",
    )
    parser.add_argument("input", type=Path, help="Path to the source .md file.")
    parser.add_argument(
        "output", nargs="?", type=Path, default=None,
        help="Output .docx path (default: same stem as input).",
    )
    args = parser.parse_args()

    dst = args.output if args.output else args.input.with_suffix('.docx')
    markdown_to_docx(args.input, dst)

# -*- coding: utf-8 -*-
"""
info-research-report skill (professional version)

Function:
1. Generate formal research report (DOCX, Chinese) from search results JSON
2. Auto-fetch web page content, call LLM for structured summary, classify info type
3. Send report via email to specified recipient
4. Designed for use with OpenClaw browser tool

Usage:
python run.py "US Iran War" "someone@example.com" results.json
python run.py "US Iran War" "someone@example.com" results.json --no-fetch
"""

import os
import sys
import json
import re
import subprocess
from datetime import datetime
from typing import List, Dict, Optional, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure console output is utf-8
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass

# Load .env file if exists
env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
if os.path.exists(env_path):
    with open(env_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, value = line.split('=', 1)
                key = key.strip()
                value = value.strip()
                if key and value and not os.environ.get(key):
                    os.environ[key] = value
                    print(f"[Info] Loaded {key} from .env")


# ==========================
# Search related (for OpenClaw external tools)
# ==========================

def search_with_browser(topic: str, max_results: int = 10) -> Optional[Dict[str, Any]]:
    """Construct search task description for OpenClaw browser tool"""
    try:
        search_url = f"https://duckduckgo.com/html/?q={topic.replace(' ', '+')}"
        print(f"  Browser will navigate to: {search_url}")
        print(f"  Then snapshot to extract search results")
        return {"action": "search", "url": search_url, "expected_results": max_results}
    except Exception as e:
        print(f"  Browser search setup error: {e}")
        return None


def parse_search_results_from_snapshot(snapshot_text: str, max_results: int = 50) -> List[Dict[str, str]]:
    """Parse search results from DuckDuckGo HTML snapshot"""
    results: List[Dict[str, str]] = []
    title_pattern = r'<a class="result__a"[^>]*>([^<]+)</a>'
    url_pattern = r'<a class="result__a" href="([^"]+)"'
    titles = re.findall(title_pattern, snapshot_text)
    urls = re.findall(url_pattern, snapshot_text)
    for i, (title, url) in enumerate(zip(titles, urls)):
        if i >= max_results:
            break
        if 'uddg=' in url:
            import urllib.parse
            parsed = urllib.parse.parse_qs(urllib.parse.urlparse(url).query)
            url = parsed.get('uddg', [url])[0]
        results.append({"title": title.strip(), "url": url, "snippet": ""})
    return results


# ==========================
# LLM summary related (Minimax/OpenAI + browseros web fetch)
# ==========================

def call_openai(prompt: str, max_tokens: int = 1000, temperature: float = 0.3) -> str:
    """Call LLM API to generate summary (supports Minimax/OpenAI)"""
    # Prefer Minimax
    api_key = os.environ.get("MINIMAX_API_KEY")
    if api_key:
        try:
            import requests
            url = "https://api.minimax.chat/v1/text/chatcompletion_v2"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            data = {
                "model": "MiniMax-M2.5",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature
            }
            resp = requests.post(url, headers=headers, json=data, timeout=120)
            resp.raise_for_status()
            result = resp.json()
            return result["choices"][0]["message"]["content"].strip()
        except Exception as e:
            print(f"[Warn] Minimax API failed: {e}, falling back to OpenAI")

    # Fallback: OpenAI (only use explicitly declared env vars)
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("[Warn] No LLM API key found, skipping LLM call")
        return ""
    try:
        import requests
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        resp = requests.post(url, headers=headers, json=data, timeout=60)
        resp.raise_for_status()
        result = resp.json()
        return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[Error] OpenAI API call failed: {e}")
        return ""


def fetch_page_content(url: str, timeout: int = 30) -> str:
    """Fetch web page content via mcporter calling browseros"""
    try:
        cmd = ["mcporter", "call", "browseros.get_page_content", f"url={url}"]
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=timeout,
            encoding='utf-8', errors='ignore'
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout
        else:
            print(f"[Warn] Failed to fetch {url}: {result.stderr or result.stdout}")
            return ""
    except Exception as e:
        print(f"[Error] fetch_page_content error: {e}")
        return ""


def _clean_llm_output(text: str) -> str:
    """Clean up unwanted content from model output (like <think> tags, English self-notes)"""
    if not text:
        return ""
    # Remove <think>...
    text = re.sub(r"<think>.*?", "", text, flags=re.DOTALL | re.IGNORECASE)
    # Remove typical English self-descriptions
    lines = []
    for line in text.splitlines():
        l = line.strip()
        if l.startswith("The user is asking me") or l.startswith("I need to be careful"):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def summarize_text_with_llm(text: str, prompt: str = "") -> str:
    """Summarize text using LLM (with cleaning)"""
    if not text and not prompt:
        return ""
    full_prompt = prompt + ("\n\n" + text if text else "")
    raw = call_openai(full_prompt, max_tokens=900, temperature=0.25)
    return _clean_llm_output(raw)


def summarize_page_content_structured(topic: str, title: str, url: str, content: str) -> Dict[str, str]:
    """
    Generate structured summary and category label for single web page content.

    Returns:
    {
        "summary": "Structured Chinese summary",
        "category": "Policy/Industry/Technology/Academia/Other"
    }
    """
    base_prompt = f"""
You are a Chinese policy and security researcher, writing a formal research report on "{topic}". Below is the body content of a web page. Please complete two tasks:

【Very important requirements】
- Answer entirely in Chinese;
- Do not output any English analysis or reasoning process;
- Do not output content like "The user is asking me", "I need to be careful" etc;
- For information from sources, objectively summarize the original content while noting the source's time and reliability;

1. Provide the information type (choose only one from the following, use English words):
   - Policy   (policy/laws/official documents)
   - Industry (industry news/market analysis/enterprise dynamics/media reports)
   - Technology (technical principles/technical reviews/technology development)
   - Academia (academic papers/research reports/scholar views)
   - Other    (other unclassifiable information)

2. Output a structured Chinese summary, required format:
   - Output must include these Chinese section headers:
     【信息类型】(English category)
     【核心结论】multiple bullet points
     【关键数据或事实】multiple bullet points, if text has no clear numbers/time/location, write "暂无明确数据"
     【与研究主题的关联】briefly explain the source's significance to "{topic}" in Chinese
     【信息来源说明】briefly describe source time and reliability assessment (e.g., official media/academic paper/self-media etc)

3. Style requirements:
   - Formal and objective language, no exaggeration;
   - Keep within approximately 400 words;
   - If information may have uncertainty, appropriately note it.

Please directly output the structured content above, do not explain what you are doing, do not output content unrelated to the above structure.

Web page title: {title}
Web page URL: {url}

Body content (for analysis only, do not extensively copy in summary):
"""
    raw = summarize_text_with_llm(content, base_prompt)
    if not raw:
        return {"summary": "", "category": "Other"}

    # Extract category from 【信息类型】 line
    category = "Other"
    for line in raw.splitlines():
        if "【信息类型】" in line:
            m = re.search(r'Policy|Industry|Technology|Academia|Other', line)
            if m:
                category = m.group(0)
            break

    return {"summary": raw.strip(), "category": category}


def summarize_overall_topic(topic: str, page_summaries: List[str]) -> str:
    """Generate overall summary for a topic based on multiple page summaries (structured, Chinese)"""
    if not page_summaries:
        return ""
    joined = "\n\n".join([f"【来源{i + 1}】\n{summary}" for i, summary in enumerate(page_summaries)])
    prompt = f"""
You are a Chinese policy and security research analyst, writing a research report on "{topic}". Below are structured summaries from multiple web pages. Please write a "Overall Situation and Key Conclusions" section based on comprehensive synthesis and induction.

【Very important requirements】
- Use Chinese throughout;
- Do not output any English self-descriptions or reasoning;
- For information from sources, objectively summarize while noting reliability and uncertainty;

Must output in this structure (keep the Chinese headers, do not add or delete):
【整体背景与现状】
- Multiple bullet points describing current overall situation, development stage, main participants etc

【关键趋势与要点】
- 3~5 bullet points summarizing most important development trends, policy developments or changes

【风险与不确定性】
- 2~4 bullet points describing main risks, challenges or uncertain aspects

Style requirements:
- Formal and objective language, no exaggeration;
- Lead with conclusions, avoid empty talk;
- Keep within 600-800 words.

Below are source summaries for reference:
{joined}
"""
    return summarize_text_with_llm("", prompt)


def enrich_results_with_summaries(
    topic: str,
    results: List[Dict[str, Any]],
    max_pages: int = 10,
    fetch_content: bool = True
) -> Dict[str, Any]:
    """
    Add to search results:
    - Structured LLM summary for each web page (llm_summary)
    - Category for each web page (Policy/Industry/Technology/Academia/Other)
    - Overall summary based on summaries
    """
    enriched_results: List[Dict[str, Any]] = []
    page_summaries: List[str] = []

    for i, item in enumerate(results):
        if i >= max_pages:
            enriched_results.append(item)
            continue

        title = item.get("title") or ""
        url = item.get("url") or ""
        content = item.get("content") or ""

        # Fetch content
        if not content and fetch_content and url and url != "N/A":
            print(f"  Fetching content from: {url}")
            content = fetch_page_content(url)
            if content:
                item["content"] = content[:8000]

        # Has content -> structured summary
        if content:
            res = summarize_page_content_structured(topic, title, url, content)
            llm_summary = res.get("summary") or ""
            category = res.get("category") or "Other"
        else:
            # No content -> weak summary
            weak_text = f"Web page title: {title}\nExisting brief summary: {item.get('snippet', '')}"
            prompt = """
You only have a web page title and a very brief summary. Please provide a brief summary in Chinese (about 200 characters) without fabricating details, describing what type of information the source roughly belongs to (e.g., policy statement, news report, academic commentary etc) and possible key content.

Requirements:
- Use Chinese throughout;
- Do not mention "future events", "fictional scenarios", "may not be real" etc;
- If you think information is indeed limited, you can use one sentence to note "Information available is limited, need to verify with original text", then summarize what you can confirm.
"""
            llm_summary = summarize_text_with_llm(weak_text, prompt)
            category = "Other"

        if llm_summary:
            item["llm_summary"] = llm_summary
            if not item.get("snippet"):
                item["snippet"] = llm_summary[:200]
            item["category"] = category
            page_summaries.append(llm_summary)
        else:
            item["category"] = item.get("category", "Other")

        enriched_results.append(item)

    overall_summary = ""
    if page_summaries:
        print("  Generating overall summary...")
        overall_summary = summarize_overall_topic(topic, page_summaries)

    return {"overall_summary": overall_summary, "results": enriched_results}


# ==========================
# Report generation tools
# ==========================

def _sanitize_filename(name: str, max_len: int = 50) -> str:
    """Convert topic to filename-safe string"""
    name = re.sub(r'[\\/*?:"<>|]', '_', name)
    return name[:max_len]


def _add_heading_center(doc: Document, text: str, level: int = 0, bold: bool = True):
    """Center-aligned heading utility function"""
    heading = doc.add_heading("", level=level)
    run = heading.add_run(text)
    run.bold = bold
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


def generate_docx(
    topic: str,
    results: List[Dict[str, Any]],
    output_path: str,
    overall_summary: str = ""
) -> str:
    """Generate formal research report (Chinese, structured)"""
    if results is None:
        results = []

    doc = Document()

    # ========= Cover =========
    _add_heading_center(doc, "研究报告", level=0)
    _add_heading_center(doc, f"-- {topic}", level=1)
    doc.add_paragraph()
    doc.add_paragraph()

    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_para.add_run("编制单位：某某研究部门\n")
    org_run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"编制时间：{datetime.now().strftime('%Y年%m月%d日')}")
    date_run.font.size = Pt(12)

    doc.add_page_break()

    # ========= Report说明 =========
    doc.add_heading("报告说明", level=1)
    info_p = doc.add_paragraph()
    info_p.add_run(f"报告主题：{topic}\n")
    info_p.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    info_p.add_run(f"信息来源数量：{len(results)}\n")
    info_p.add_run("说明：本报告基于公开网络信息及自动化分析工具生成，仅供内部研究参考，结论及观点不代表任何机构的正式立场。\n")
    doc.add_paragraph("_" * 60)

    # ========= 一、研究背景和目的 =========
    doc.add_heading("一、研究背景和目的", level=1)
    bg_p = doc.add_paragraph()
    bg_p.add_run(
        f"近年来，\"{topic}\"相关议题在政策、安全形势和国际关系等方面受到广泛关注。"
        f"为系统梳理当前公开信息，识别关键趋势与潜在风险，特编制本研究报告，"
        f"旨在为相关决策和后续深入研究提供基础性参考。"
    )
    doc.add_paragraph("_" * 60)

    # ========= 二、研究方法 =========
    doc.add_heading("二、研究方法", level=1)
    method_p = doc.add_paragraph()
    method_p.add_run(
        "本报告主要采用以下研究方法：\n"
        "1. 网络检索：通过 OpenClaw browser 等工具，在综合搜索引擎中检索与主题高度相关的公开信息；\n"
        "2. 来源筛选：优先选择政策文件、权威机构报告、主流媒体报道及专业分析等相对可靠来源；\n"
        "3. 自动化摘要：使用大语言模型对网页正文进行结构化摘要，提取核心结论、关键数据及与主题的关联；\n"
        "4. 综合研判：在汇总各来源信息的基础上，形成总体形势、关键趋势及风险研判。"
    )
    doc.add_paragraph("_" * 60)

    # ========= 三、总体形势与关键结论 =========
    doc.add_heading("三、总体形势与关键结论", level=1)

    if results:
        summary_p = doc.add_paragraph()
        summary_p.add_run(f"基于所采集的公开信息，围绕\"{topic}\"，总体形势与主要结论概括如下：\n\n")
        if overall_summary:
            summary_p.add_run(overall_summary + "\n")
        else:
            summary_p.add_run("（当前版本未能生成总体总结，请结合后文分专题和来源分析自行研判。）\n")
    else:
        doc.add_paragraph("暂无可用搜索结果，无法形成总体结论。")

    doc.add_paragraph("_" * 60)

    # ========= 四、分专题分析（按来源类别） =========
    doc.add_heading("四、分专题分析", level=1)

    category_map: Dict[str, List[Dict[str, Any]]] = {}
    for item in results:
        cat = item.get("category") or "Other"
        category_map.setdefault(cat, []).append(item)

    category_names = {
        "Policy": "（一）政策与监管信息",
        "Industry": "（二）媒体报道与舆论动态",
        "Technology": "（三）技术发展与应用",
        "Academia": "（四）学术研究与专家观点",
        "Other": "（五）其他相关信息"
    }

    for key in ["Policy", "Industry", "Technology", "Academia", "Other"]:
        items = category_map.get(key, [])
        if not items:
            continue
        title_cat = category_names.get(key, f"类别：{key}")
        sub_p = doc.add_paragraph()
        run = sub_p.add_run(title_cat)
        run.bold = True
        run.font.size = Pt(12)

        for idx, item in enumerate(items[:5], 1):
            t = item.get("title") or "N/A"
            s = item.get("llm_summary") or item.get("snippet") or ""
            p = doc.add_paragraph()
            p.add_run(f"{idx}. {t}\n").bold = True
            if s:
                p.add_run(f"   概要：{s[:350]}...\n")

    doc.add_paragraph("_" * 60)

    # ========= 五、详细来源分析（Source Details） =========
    doc.add_heading("五、详细来源分析（Source Details）", level=1)

    if results:
        for idx, item in enumerate(results, 1):
            title = item.get("title") or "N/A"
            url = item.get("url") or "N/A"
            snippet = item.get("snippet") or ""
            llm_summary = item.get("llm_summary") or ""
            category = item.get("category") or "Other"

            p = doc.add_paragraph()
            run = p.add_run(f"（{idx}）{title}")
            run.bold = True
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            p.add_run(f"信息类型：{category}\n")

            p = doc.add_paragraph()
            run = p.add_run(f"来源：{url}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 255)

            summary_text = llm_summary or snippet
            if summary_text:
                p = doc.add_paragraph()
                run = p.add_run(f"摘要：{summary_text}")
                run.font.size = Pt(10)
                run.italic = True

            doc.add_paragraph("")
    else:
        doc.add_paragraph("暂无来源数据。")

    doc.add_paragraph("_" * 60)

    # ========= 六、参考资料 =========
    doc.add_heading("六、参考资料（All Search Results）", level=1)

    if results:
        for item in results:
            title = item.get("title") or ""
            url = item.get("url") or ""
            if not title:
                continue
            p = doc.add_paragraph()
            title_run = p.add_run(f"• {title}")
            title_run.bold = True
            if url:
                url_run = p.add_run(f"\n  URL：{url}")
                url_run.font.size = Pt(9)
    else:
        doc.add_paragraph("无。")

    doc.save(output_path)
    return output_path


# ==========================
# Email sending tools
# ==========================

def send_email(to_addr: str, subject: str, content: str, attachment: Optional[str] = None,
               mailbox: str = "qq", skills_dir: Optional[str] = None, timeout: int = 60) -> bool:
    """Call OpenClaw's email skill to send email"""
    if skills_dir is None:
        skills_dir = os.environ.get("OPENCLAW_SKILLS_DIR") or r"C:\Users\Juxin\.openclaw\workspace\skills"
    email_py = os.path.join(skills_dir, "email-mail-master", "scripts", "mail.py")
    if not os.path.exists(email_py):
        print(f"[Error] mail.py not found at: {email_py}")
        return False

    cmd = [sys.executable, email_py, "--mailbox", mailbox, "send",
           "--to", to_addr, "--subject", subject, "--content", content]
    if attachment and os.path.exists(attachment):
        cmd.extend(["--attach", attachment])

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        if result.returncode != 0:
            print(f"[Error] Email send failed (code={result.returncode}): {result.stderr}")
            return False
        else:
            if result.stdout:
                print(result.stdout)
            return True
    except Exception as e:
        print(f"[Error] Email send exception: {e}")
        return False


# ==========================
# Main entry point
# ==========================

def main():
    """
    Usage:
    python run.py "topic" "recipient_email" [results.json] [--no-fetch]

    Options:
    --no-fetch  Skip web content fetching, use only titles and snippets for summary (faster)
    """
    args = []
    fetch_content = True
    for arg in sys.argv[1:]:
        if arg == "--no-fetch":
            fetch_content = False
        else:
            args.append(arg)

    topic = args[0] if len(args) > 0 else "测试主题"
    email = args[1] if len(args) > 1 else "2507541738@qq.com"
    results_file = args[2] if len(args) > 2 else ""

    print(f"=== Research: {topic} ===")
    print(f"[Info] Fetch content from URLs: {fetch_content}")
    print("[Info] Use --no-fetch to skip fetching (faster but less detailed)")

    # Load results
    if results_file and os.path.exists(results_file):
        try:
            with open(results_file, 'r', encoding='utf-8') as f:
                results = json.load(f)
            if not isinstance(results, list):
                results = []
            print(f"\n[1/4] Loaded {len(results)} results from {results_file}")
        except Exception as e:
            print(f"[Error] Failed to read results file: {e}")
            results = []
    else:
        print("\n[1/4] No results file provided, using preset data")
        results = [{"title": "请先使用 OpenClaw browser 工具进行搜索", "url": "N/A", "snippet": ""}]

    # Generate LLM summaries & overall summary
    print("\n[2/4] Generating structured summaries with LLM...")
    enriched = enrich_results_with_summaries(topic, results, max_pages=10, fetch_content=fetch_content)
    overall_summary = enriched["overall_summary"]
    results = enriched["results"]
    print("  Summaries generation finished.")

    # Generate DOCX
    print("\n[3/4] Generating DOCX report (Chinese, professional style)...")
    safe_topic = _sanitize_filename(topic.replace(" ", "_"), max_len=30)
    output_file = f"Report_{safe_topic}_{datetime.now().strftime('%Y%m%d%H%M')}.docx"

    try:
        generate_docx(topic, results, output_file, overall_summary=overall_summary)
        print(f"  Saved report: {output_file}")
    except Exception as e:
        print(f"[Error] Failed to generate DOCX: {e}")
        print("\n=== Aborted ===")
        return

    # Send email
    print(f"\n[4/4] Sending report to {email}...")
    content = (
        f"主题：{topic}\n"
        f"来源数量：{len(results)} 条\n\n"
        "本邮件附带自动生成的研究报告（DOCX），仅供内部研究参考使用。"
    )
    success = send_email(email, f"研究报告: {topic}", content, output_file)
    print("  Email sent." if success else "  Email send failed.")

    print("\n=== Done ===")


if __name__ == "__main__":
    main()
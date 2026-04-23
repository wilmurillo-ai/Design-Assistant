#!/usr/bin/env python3
"""
Patent Markdown to Word Document Converter
Location: skills/professional-patent-agents/agents/patent-converter/convert_patents.py

Process:
1. Search for Patent*.md or 专利*.md files in directory
2. Use Pandoc to convert md to HTML (extract structured content)
3. Use regex to extract title and chapter content
4. Replace placeholders in template: {{ title }}, {{ chapter1 }} ~ {{ chapter7 }}, {{ year }}, {{ month }}, {{ day }}
5. Convert Mermaid diagrams and insert at original positions
6. Output to same directory as source file

Usage:
    python convert_patents.py [directory_path]
    
    If no directory specified, defaults to /root/workspace/patent/new

Supports:
    - Chinese patent documents (专利*.md)
    - English patent documents (Patent*.md)
    - Mixed language content
"""

import os
import re
import subprocess
import tempfile
import shutil
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Script directory
SCRIPT_DIR = Path(__file__).parent.resolve()

# Template path (fixed)
TEMPLATE_PATH = SCRIPT_DIR / "templates" / "template.docx"

# Default search directory
DEFAULT_SEARCH_DIR = "/root/workspace/patent/new"

# Chapter title patterns (supports both Chinese and English)
CHAPTER_PATTERNS = {
    'chinese': [
        "一、相关的现有技术及现有技术的缺陷或不足",
        "二、为克服上述缺陷本提案的技术改进点",
        "三、技术改进点的其他替代方案",
        "四、详细的技术方案具体实施例",
        "五、本提案相对现有技术的优点",
        "六、相关附图",
        "七、权利要求书",
    ],
    'english': [
        "1. Related Prior Art and Their Defects or Deficiencies",
        "2. Technical Improvements to Overcome the Above Defects",
        "3. Alternative Solutions for Technical Improvements",
        "4. Detailed Embodiments of the Technical Solution",
        "5. Advantages of This Proposal Over Prior Art",
        "6. Related Drawings",
        "7. Claims",
    ]
}


def find_patent_files(directory: str) -> list:
    """
    Search for patent Markdown files in directory
    Supports both Chinese (专利*.md) and English (Patent*.md) naming
    Returns list of file paths
    """
    directory = Path(directory)
    
    if not directory.exists():
        print(f"Error: Directory does not exist - {directory}")
        return []
    
    patent_files = []
    
    # Search for Chinese named files
    patent_files.extend(directory.glob("专利*.md"))
    patent_files.extend(directory.glob("*/专利*.md"))
    
    # Search for English named files
    patent_files.extend(directory.glob("Patent*.md"))
    patent_files.extend(directory.glob("*/Patent*.md"))
    
    # Also search for lowercase
    patent_files.extend(directory.glob("patent*.md"))
    patent_files.extend(directory.glob("*/patent*.md"))
    
    # Deduplicate and sort
    patent_files = sorted(set(str(f) for f in patent_files))
    
    return patent_files


def detect_language(content: str) -> str:
    """
    Detect whether the patent document is in Chinese or English
    Returns 'chinese' or 'english'
    """
    # Check for Chinese chapter titles first
    for title in CHAPTER_PATTERNS['chinese']:
        if title in content:
            return 'chinese'
    
    # Check for English chapter titles
    for title in CHAPTER_PATTERNS['english']:
        if title in content:
            return 'english'
    
    # Default to Chinese if contains Chinese characters
    if re.search(r'[\u4e00-\u9fff]', content):
        return 'chinese'
    
    return 'english'


def md_to_html(md_content: str) -> str:
    """Convert Markdown to HTML using Pandoc"""
    try:
        result = subprocess.run(
            ['pandoc', '-f', 'markdown', '-t', 'html', '--wrap=none'],
            input=md_content,
            capture_output=True,
            text=True,
            timeout=30
        )
        return result.stdout if result.returncode == 0 else md_content
    except Exception as e:
        print(f"Pandoc error: {e}")
        return md_content


def extract_mermaid_blocks(content: str) -> tuple:
    """Extract Mermaid code blocks, returns (cleaned_content, mermaid_blocks_list)"""
    pattern = r'```mermaid\s*\n(.*?)```'
    matches = list(re.finditer(pattern, content, re.DOTALL))
    
    mermaid_blocks = [m.group(1).strip() for m in matches]
    
    # Replace Mermaid blocks with placeholder
    cleaned = re.sub(pattern, '[[MERMAID_IMAGE]]', content, flags=re.DOTALL)
    
    return cleaned, mermaid_blocks


def mermaid_to_png(code: str, output_path: str) -> bool:
    """Convert Mermaid code to PNG"""
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
            f.write(code)
            mmd_path = f.name
        
        # Create temp Puppeteer config (root user needs --no-sandbox)
        puppeteer_config = {
            "args": ["--no-sandbox", "--disable-setuid-sandbox"]
        }
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
            import json
            json.dump(puppeteer_config, f)
            config_path = f.name
        
        cmd = [
            'mmdc',
            '-i', mmd_path,
            '-o', output_path,
            '-b', 'white',
            '-w', '800',
            '--quiet',
            '-p', config_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        os.unlink(mmd_path)
        os.unlink(config_path)
        return result.returncode == 0 and os.path.exists(output_path)
    except Exception as e:
        print(f"Mermaid conversion error: {e}")
        return False


def parse_patent_sections(md_path: str) -> dict:
    """
    Parse patent Markdown file
    Returns {title, chapter1~chapter7, mermaid_blocks, language}
    Supports both Chinese and English document formats
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Detect language
    language = detect_language(content)
    chapter_titles = CHAPTER_PATTERNS[language]
    
    # Extract title
    title_match = re.search(r'^#\s+(.+?)$', content, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else "Untitled"
    
    chapters = {}
    all_mermaid_blocks = []
    
    for i, section_title in enumerate(chapter_titles):
        chapter_key = f'chapter{i+1}'
        
        # Find chapter position
        idx = content.find(section_title)
        if idx == -1:
            chapters[chapter_key] = ""
            continue
        
        # Find chapter end position
        start = content.find('\n', idx) + 1
        end = len(content)
        
        for next_title in chapter_titles[i+1:]:
            next_idx = content.find(next_title)
            if next_idx > start:
                end = content.rfind('\n', 0, next_idx)
                if end == -1:
                    end = next_idx
                break
        
        chapter_content = content[start:end].strip()
        
        # Extract Mermaid blocks
        chapter_content, mermaid_blocks = extract_mermaid_blocks(chapter_content)
        all_mermaid_blocks.extend(mermaid_blocks)
        
        chapters[chapter_key] = chapter_content
    
    return {
        'title': title,
        'chapters': chapters,
        'mermaid_blocks': all_mermaid_blocks,
        'language': language
    }


def add_formatted_text(cell, text: str):
    """Add text to cell with simple formatting"""
    cell._element.clear_content()
    
    paragraphs = text.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        if para_text.startswith('### '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[4:])
            run.bold = True
        elif para_text.startswith('## '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[3:])
            run.bold = True
        elif para_text.startswith('# '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[2:])
            run.bold = True
        elif para_text.startswith('|') and '\n|' in para_text:
            add_table(cell, para_text)
        else:
            p = cell.add_paragraph()
            p.add_run(para_text)


def add_table(cell, table_text: str):
    """Add Markdown table to cell"""
    lines = [l.strip() for l in table_text.split('\n') if l.strip() and not re.match(r'^\|[-\s|]+\|$', l.strip())]
    
    if len(lines) < 2:
        return
    
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)
    
    if not rows_data:
        return
    
    num_cols = len(rows_data[0])
    table = cell.add_table(rows=len(rows_data), cols=num_cols)
    
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[i].cells[j].text = cell_text


def convert_patent(md_path: str) -> str:
    """
    Convert a single patent file
    Returns output docx file path
    """
    md_path = Path(md_path)
    print(f"\nProcessing: {md_path.name}")
    
    # Parse patent file
    data = parse_patent_sections(str(md_path))
    print(f"  Language: {data['language']}")
    print(f"  Title: {data['title'][:50]}...")
    print(f"  Mermaid diagrams: {len(data['mermaid_blocks'])}")
    
    # Load template
    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]
    
    # Replace date placeholders
    today = datetime.now()
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace('{{ year }}', str(today.year))
            run.text = run.text.replace('{{ month }}', str(today.month).zfill(2))
            run.text = run.text.replace('{{ day }}', str(today.day).zfill(2))
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace('{{ year }}', str(today.year))
                    run.text = run.text.replace('{{ month }}', str(today.month).zfill(2))
                    run.text = run.text.replace('{{ day }}', str(today.day).zfill(2))
    
    # Replace title (row 0, col 1)
    table.rows[0].cells[1].text = data['title']
    
    # Replace chapters (row 7-13, col 1)
    chapter_rows = [7, 8, 9, 10, 11, 12, 13]
    
    for i, row_idx in enumerate(chapter_rows):
        chapter_key = f'chapter{i+1}'
        chapter_content = data['chapters'].get(chapter_key, '')
        
        if not chapter_content:
            table.rows[row_idx].cells[1].text = ""
            continue
        
        # Convert Mermaid diagrams
        if '[[MERMAID_IMAGE]]' in chapter_content:
            print(f"  Chapter {i+1} has diagrams, converting...")
            
            temp_dir = tempfile.mkdtemp()
            parts = chapter_content.split('[[MERMAID_IMAGE]]')
            img_idx = 0
            
            cell = table.rows[row_idx].cells[1]
            cell._element.clear_content()
            
            for part_idx, part in enumerate(parts):
                if part.strip():
                    for para_text in part.strip().split('\n\n'):
                        if para_text.strip():
                            if para_text.startswith('### '):
                                p = cell.add_paragraph()
                                run = p.add_run(para_text[4:])
                                run.bold = True
                            elif para_text.startswith('## '):
                                p = cell.add_paragraph()
                                run = p.add_run(para_text[3:])
                                run.bold = True
                            elif para_text.startswith('|') and '\n|' in para_text:
                                add_table(cell, para_text)
                            else:
                                p = cell.add_paragraph()
                                p.add_run(para_text)
                
                if part_idx < len(parts) - 1 and img_idx < len(data['mermaid_blocks']):
                    mermaid_code = data['mermaid_blocks'][img_idx]
                    png_path = os.path.join(temp_dir, f"diagram_{img_idx}.png")
                    
                    if mermaid_to_png(mermaid_code, png_path):
                        try:
                            p = cell.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(png_path, width=Inches(5.5))
                            print(f"    Diagram {img_idx+1} OK")
                        except Exception as e:
                            print(f"    Diagram {img_idx+1} FAILED ({e})")
                    else:
                        print(f"    Diagram {img_idx+1} FAILED (conversion error)")
                    
                    img_idx += 1
            
            shutil.rmtree(temp_dir, ignore_errors=True)
        else:
            add_formatted_text(table.rows[row_idx].cells[1], chapter_content)
            print(f"  Chapter {i+1} OK")
    
    # Output to same directory as source
    output_path = md_path.with_suffix('.docx')
    doc.save(str(output_path))
    print(f"  Output: {output_path.name}")
    
    return str(output_path)


def main():
    print("=" * 60)
    print("Patent Markdown to Word Document Converter")
    print("Supports: Chinese (专利*.md) and English (Patent*.md)")
    print("=" * 60)
    
    # Check template
    if not TEMPLATE_PATH.exists():
        print(f"\nError: Template file not found - {TEMPLATE_PATH}")
        sys.exit(1)
    
    print(f"\nTemplate: {TEMPLATE_PATH}")
    
    # Get search directory
    if len(sys.argv) > 1:
        search_dir = sys.argv[1]
    else:
        search_dir = DEFAULT_SEARCH_DIR
    
    print(f"Search directory: {search_dir}")
    
    # Check tools
    pandoc_ok = subprocess.run(['which', 'pandoc'], capture_output=True).returncode == 0
    mmdc_ok = subprocess.run(['which', 'mmdc'], capture_output=True).returncode == 0
    
    if not pandoc_ok or not mmdc_ok:
        print("\nError: Missing required tools")
        if not pandoc_ok:
            print("  - Pandoc not installed")
        if not mmdc_ok:
            print("  - Mermaid CLI not installed")
        sys.exit(1)
    
    # Find patent files
    patent_files = find_patent_files(search_dir)
    
    if not patent_files:
        print(f"\nNo patent files found (专利*.md or Patent*.md)")
        sys.exit(0)
    
    print(f"\nFound {len(patent_files)} patent file(s):\n")
    for f in patent_files:
        print(f"  - {Path(f).name}")
    
    # Convert
    print(f"\nStarting conversion...\n")
    
    success_count = 0
    for md_path in patent_files:
        try:
            convert_patent(md_path)
            success_count += 1
        except Exception as e:
            print(f"  Error: {e}")
    
    print("\n" + "=" * 60)
    print(f"Conversion complete! Success: {success_count}/{len(patent_files)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
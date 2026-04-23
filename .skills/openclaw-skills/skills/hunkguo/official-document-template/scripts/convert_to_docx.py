#!/usr/bin/env python3
"""
公文排版脚本
将Markdown格式的公文内容转换为符合GB/T 9704-2012标准的Word文档
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("需要安装python-docx: pip install python-docx")
    sys.exit(1)


# 公文格式常量 (GB/T 9704-2012)
PAGE_WIDTH = Cm(21.0)      # A4宽度
PAGE_HEIGHT = Cm(29.7)     # A4高度
MARGIN_TOP = Cm(3.7)       # 上边距37mm
MARGIN_BOTTOM = Cm(3.5)    # 下边距35mm
MARGIN_LEFT = Cm(2.8)      # 左边距28mm
MARGIN_RIGHT = Cm(2.6)     # 右边距26mm

FONT_SIZE_SAN = Pt(16)     # 三号字 16pt
FONT_SIZE_ER = Pt(22)      # 二号字 22pt
FONT_SIZE_SI = Pt(14)      # 四号字 14pt

LINE_SPACING = Pt(28)      # 行距28磅


def set_chinese_font(run, font_name, size=None, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = size
    run.font.bold = bold


def setup_document():
    """初始化文档，设置页面格式"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    
    return doc


def add_title(doc, text):
    """添加公文标题（二号小标宋，居中）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_chinese_font(run, '方正小标宋简体', FONT_SIZE_ER)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_doc_number(doc, text):
    """添加发文字号（三号仿宋，居中）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_chinese_font(run, '仿宋', FONT_SIZE_SAN)


def add_level1_title(doc, text):
    """添加一级标题（三号黑体）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '黑体', FONT_SIZE_SAN)
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)


def add_level2_title(doc, text):
    """添加二级标题（三号楷体，加粗）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '楷体', FONT_SIZE_SAN, bold=True)
    para.paragraph_format.first_line_indent = Cm(0)


def add_level3_title(doc, text):
    """添加三级标题（三号仿宋，加粗）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '仿宋', FONT_SIZE_SAN, bold=True)
    para.paragraph_format.first_line_indent = Cm(0)


def add_body_text(doc, text):
    """添加正文（三号仿宋）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '仿宋', FONT_SIZE_SAN)
    # 首行缩进两字符
    para.paragraph_format.first_line_indent = Cm(0.74)  # 约2个三号字
    para.paragraph_format.line_spacing = Pt(28)


def add_recipient(doc, text):
    """添加主送机关"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '仿宋', FONT_SIZE_SAN)
    para.paragraph_format.first_line_indent = Cm(0)


def add_signature(doc, org_name, date):
    """添加落款"""
    # 空两行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 发文机关名称
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(org_name)
    set_chinese_font(run, '仿宋', FONT_SIZE_SAN)
    
    # 成文日期
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(date)
    set_chinese_font(run, '仿宋', FONT_SIZE_SAN)


def parse_markdown_content(content):
    """
    解析Markdown内容，识别标题层次
    
    返回: [(类型, 内容), ...]
    类型: 'title', 'level1', 'level2', 'level3', 'body', 'blank'
    """
    lines = content.strip().split('\n')
    elements = []
    
    for line in lines:
        line = line.rstrip()
        
        # 空行
        if not line.strip():
            elements.append(('blank', ''))
            continue
        
        # 一级标题 # 开头
        if line.startswith('# '):
            elements.append(('title', line[2:].strip()))
            continue
        
        # 一级标题（一、二、三、格式）
        if re.match(r'^[一二三四五六七八九十]+、', line):
            elements.append(('level1', line))
            continue
        
        # 二级标题（（一）（二）（三）格式）
        if re.match(r'^（[一二三四五六七八九十]+）', line):
            elements.append(('level2', line))
            continue
        
        # 三级标题（1. 2. 3. 格式）
        if re.match(r'^\d+\.', line):
            # 判断是否为标题（短于30字且不包含标点结尾）
            if len(line) < 30 and not line[-1] in '，。；：！？':
                elements.append(('level3', line))
            else:
                elements.append(('body', line))
            continue
        
        # 四级标题（（1）（2）（3）格式）
        if re.match(r'^（\d+）', line):
            if len(line) < 30:
                elements.append(('level3', line))
            else:
                elements.append(('body', line))
            continue
        
        # 发文字号
        if re.match(r'.*发〔\d{4}〕\d+号', line):
            elements.append(('doc_number', line))
            continue
        
        # 默认为正文
        elements.append(('body', line))
    
    return elements


def convert_to_docx(markdown_content, output_path):
    """将Markdown内容转换为Word文档"""
    doc = setup_document()
    elements = parse_markdown_content(markdown_content)
    
    for elem_type, content in elements:
        if elem_type == 'title':
            add_title(doc, content)
        elif elem_type == 'level1':
            add_level1_title(doc, content)
        elif elem_type == 'level2':
            add_level2_title(doc, content)
        elif elem_type == 'level3':
            add_level3_title(doc, content)
        elif elem_type == 'body':
            if content.strip():
                add_body_text(doc, content)
        elif elem_type == 'doc_number':
            add_doc_number(doc, content)
    
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='公文排版工具')
    parser.add_argument('input', help='输入的Markdown文件路径')
    parser.add_argument('--output', '-o', help='输出的Word文件路径（默认：同目录同名.docx）')
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误：文件不存在 - {input_path}")
        sys.exit(1)
    
    # 读取内容
    content = input_path.read_text(encoding='utf-8')
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        output_path = input_path.with_suffix('.docx')
    
    # 转换
    try:
        result = convert_to_docx(content, output_path)
        print(f"✅ 转换完成: {result}")
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()

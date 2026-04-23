#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金新闻 Word 文档生成脚本
用于「过去七天」和「指定日期范围」查询
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta
import re


def set_font(run, font_name='等线', font_size=11):
    """设置中文字体"""
    run.font.size = Pt(font_size)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def should_exclude(title, content):
    """
    判断是否应该排除该新闻
    返回 (是否排除，排除原因)
    
    核心原则：只关注公募基金行业整体和监管政策，过滤单一公司/产品/事件
    
    过滤优先级：
    1. 私募基金相关 → 最高优先级
    2. 单一基金产品发行 → 高优先级
    3. 单一基金公司系统/技术升级 → 高优先级
    4. 基金经理人事变动 → 高优先级
    5. 短期收益率/涨跌报道 → 高优先级
    6. 非中国公募基金行业 → 高优先级
    7. 基金公司营销/活动 → 高优先级
    8. 单一事件/个案报道 → 高优先级
    9. 投资峰会/会议 → 中优先级
    10. 非公募基金行业 → 中优先级
    11. 资金流向/龙虎榜 → 中优先级
    12. 访谈/评论 → 低优先级
    """
    text = f"{title} {content}"
    
    # ========== 一、私募基金相关（最高优先级）==========
    private_keywords = [
        '私募', '私募基金', '私募机构', '私募产品', 
        '私募量化', '私募股权', '私募债权', '私募访谈'
    ]
    for kw in private_keywords:
        if kw in title or kw in content:
            return True, '私募相关'
    
    # ========== 二、单一基金产品发行（高优先级）==========
    single_product_keywords = [
        '发行中', '正在发行', '新基金发行', '产品成立',
        '联接基金发行', 'ETF 联接发行'
    ]
    for kw in single_product_keywords:
        if kw in title:
            return True, '单一基金产品发行'
    
    # ========== 三、单一基金公司系统/技术升级（高优先级）==========
    system_keywords = [
        '系统上线', '系统优化', '投资交易系统', 'IT 系统',
        '技术升级', '系统切换'
    ]
    # 检查是否提及具体基金公司
    fund_company_pattern = r'在.+基金 (上线 | 实施 | 应用)'
    if any(kw in text for kw in system_keywords) and re.search(fund_company_pattern, text):
        return True, '单一基金公司系统升级'
    
    # ========== 四、基金经理人事变动（高优先级）==========
    manager_change_keywords = [
        '基金经理', '离任', '离职', '新任', '变更', 
        '人事变动', '老将离职', '百亿基金经理'
    ]
    for kw in manager_change_keywords:
        if kw in title:
            change_indicators = ['离任', '离职', '新任', '变更', '人事', '老将']
            if any(ind in title for ind in change_indicators):
                return True, '基金经理人事变动'
    
    # ========== 五、短期收益率/涨跌报道（高优先级）==========
    return_keywords = ['涨逾', '跌幅', '收益率', '净值涨', '单日涨', '涨超']
    for kw in return_keywords:
        if kw in title:
            return True, '短期收益率/涨跌报道'
    
    # ========== 六、非中国公募基金行业（高优先级）==========
    non_china_keywords = [
        '主权债券', '国际市场', '海外发行', '离岸人民币',
        '印尼', '新兴市场融资'
    ]
    for kw in non_china_keywords:
        if kw in title and '公募' not in text:
            return True, '非中国公募基金行业'
    
    # ========== 七、基金公司营销/活动（高优先级）==========
    marketing_keywords = [
        '足球邀请赛', '活动', '营销', '杯', '赛事',
        '鲲鹏杯', '邀请赛', '运动会'
    ]
    for kw in marketing_keywords:
        if kw in title:
            return True, '基金公司营销/活动'
    
    # ========== 八、单一事件/个案报道（高优先级）==========
    single_event_keywords = [
        '李鬼', '连环套', '诈骗', '骗局', '陷阱',
        '大佬变啃佬'
    ]
    for kw in single_event_keywords:
        if kw in title:
            return True, '单一事件/个案报道'
    
    # ========== 九、投资峰会/会议类（中优先级）==========
    summit_keywords = [
        '投资峰会', '投资论坛', '策略峰会', '年度峰会',
        '峰会举行', '论坛举行', '策略会'
    ]
    for kw in summit_keywords:
        if kw in title or kw in content:
            return True, '投资峰会/会议'
    
    # ========== 十、非公募基金行业无关新闻（中优先级）==========
    non_public_keywords = [
        '房产基金', '地产基金', '房地产基金',
        '产业基金', '引导基金', '政府基金',
        '外币基金', '首关', '完成首关',
        '完成备案', '基金备案', '京津冀基金',
        '地方基金', '创投基金', '风投基金'
    ]
    for kw in non_public_keywords:
        if kw in title or kw in content:
            return True, '非公募基金行业'
    
    # ========== 十一、原有过滤规则（保留）==========
    
    # 排除 ETF 龙虎榜
    if '龙虎榜' in title:
        return True, 'ETF 龙虎榜'
    
    # 排除「x 只 ETF 获融资净买入」类
    if 'ETF 获融资净买入' in text or re.search(r'\d+ 只 ETF', text):
        return True, 'ETF 融资净买入'
    
    # 排除「x 月 x 日资金净流入」类
    if '资金净流入' in text or '资金净流出' in text:
        return True, '资金净流入/流出'
    
    # 排除基金经理访谈（原有规则）
    interview_keywords = ['访谈', '专访', '对话', '面对面']
    if '基金经理' in text:
        for kw in interview_keywords:
            if kw in title or kw in content:
                return True, '基金经理访谈'
    
    return False, None


def generate_word_document(news_data, start_date, end_date, output_path):
    """
    生成基金新闻 Word 文档
    
    参数:
        news_data: 按日期分组的新闻数据
            {日期：[{title, content, url, source, time}, ...]}
        start_date: 起始日期 (datetime 对象)
        end_date: 截止日期 (datetime 对象)
        output_path: 输出文件路径
    
    返回:
        output_path 或 None(无新闻时)
    """
    source_order = ['证券时报', '中国证券报·中证网', '证券日报']
    
    # 生成日期范围
    date_list = []
    current_date = start_date
    while current_date <= end_date:
        date_list.append(current_date)
        current_date += timedelta(days=1)
    
    # 统计有效新闻数
    total_news = 0
    
    doc = Document()
    is_first_date = True
    
    for date in date_list:
        date_str = f"{date.year}.{date.month}.{date.day}"
        
        # 新日期新起一页
        if not is_first_date:
            doc.add_page_break()
        is_first_date = False
        
        # 日期段落
        date_para = doc.add_paragraph(date_str)
        for run in date_para.runs:
            set_font(run)
        
        doc.add_paragraph()
        
        # 该日期的新闻
        if date_str in news_data:
            for source in source_order:
                source_news = [n for n in news_data[date_str] if n['source'] == source]
                source_news.sort(key=lambda x: x.get('time', '00:00'))
                
                has_news = False
                for news in source_news:
                    exclude, reason = should_exclude(news['title'], news['content'])
                    if exclude:
                        continue
                    
                    has_news = True
                    total_news += 1
                    
                    # 段落 1: 新闻正文
                    content_para = doc.add_paragraph(style='List Bullet')
                    
                    title_run = content_para.add_run(news['title'])
                    title_run.bold = True
                    set_font(title_run)
                    
                    sep_run = content_para.add_run('。')
                    sep_run.bold = False
                    set_font(sep_run)
                    
                    content_run = content_para.add_run(news['content'])
                    content_run.bold = False
                    set_font(content_run)
                    
                    # 段落 2: 链接
                    link_para = doc.add_paragraph()
                    link_run = link_para.add_run(news['url'])
                    set_font(link_run)
                    
                    # 段落 3: 来源
                    source_para = doc.add_paragraph()
                    source_run = source_para.add_run(f"{news['source']}，{news['title']} {news['source']}")
                    set_font(source_run)
                    
                    doc.add_paragraph()
                
                if not has_news and source_news:
                    simple_source = source.replace('·中证网', '')
                    note_para = doc.add_paragraph()
                    note_run = note_para.add_run(f"{date_str}「{simple_source}无符合规则新闻」")
                    set_font(note_run)
        else:
            for source in source_order:
                simple_source = source.replace('·中证网', '')
                note_para = doc.add_paragraph()
                note_run = note_para.add_run(f"{date_str}「{simple_source}无符合规则新闻」")
                set_font(note_run)
    
    if total_news == 0:
        return None
    
    doc.save(output_path)
    return output_path


def get_filename(start_date, end_date):
    """生成文件名"""
    start_str = start_date.strftime('%Y%m%d')
    end_str = end_date.strftime('%Y%m%d')
    return f"{start_str}-{end_str}基金新闻.docx"


if __name__ == '__main__':
    # 测试示例
    from datetime import datetime
    
    news_data = {
        '2026.3.10': [
            {
                'title': 'ETF 更名正加快重塑指数化投资底层生态',
                'content': '进入 3 月份，一场影响全市场 ETF 产品的"更名变革"正在全面提速。',
                'url': 'http://www.stcn.com/article/detail/3670939.html',
                'source': '证券时报',
                'time': '08:41'
            }
        ]
    }
    
    start_date = datetime(2026, 3, 10)
    end_date = datetime(2026, 3, 10)
    output_path = f'/tmp/{get_filename(start_date, end_date)}'
    
    result = generate_word_document(news_data, start_date, end_date, output_path)
    if result:
        print(f"Word 文档已生成：{result}")
    else:
        print("该时间段无有效基金新闻")

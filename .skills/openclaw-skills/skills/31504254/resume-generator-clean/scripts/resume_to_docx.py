#!/usr/bin/env python3
"""
Resume to Word (.docx) Converter
Generates professional Word resumes from HTML or text input.
"""

import sys
import os
from pathlib import Path

def generate_docx(input_file, output_file, template='modern'):
    """
    Convert HTML resume to Word document.
    
    Args:
        input_file: Path to HTML file or raw HTML string
        output_file: Path for output .docx
        template: Template name (modern, classic, minimal)
    """
    try:
        # Try python-docx first
        return _try_python_docx(input_file, output_file, template)
    except ImportError:
        print("python-docx not available, trying alternative...")
        return _create_html_alternative(input_file, output_file)

def _try_python_docx(input_file, output_file, template):
    """Generate DOCX using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set up document styles based on template
    _setup_styles(doc, template)
    
    # Parse HTML content
    html_content = _get_html_content(input_file)
    
    # Simple HTML parsing - extract structure
    from html.parser import HTMLParser
    
    class ResumeHTMLParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.doc = doc
            self.current_text = []
            self.in_title = False
            self.in_heading = False
            self.in_list = False
            self.list_items = []
            
        def handle_starttag(self, tag, attrs):
            if tag in ['h1']:
                self.in_title = True
            elif tag in ['h2', 'h3']:
                self.in_heading = True
                self._flush_text()
            elif tag in ['li']:
                self.in_list = True
            elif tag in ['p', 'div']:
                self._flush_text()
                
        def handle_endtag(self, tag):
            if tag in ['h1']:
                self.in_title = False
                title = ''.join(self.current_text)
                self.current_text = []
                
                # Add title
                p = self.doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(44, 62, 80)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            elif tag in ['h2', 'h3']:
                self.in_heading = False
                heading = ''.join(self.current_text)
                self.current_text = []
                
                # Add section heading
                p = self.doc.add_paragraph()
                run = p.add_run(heading)
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(52, 73, 94)
                
            elif tag in ['li']:
                self.in_list = False
                item = ''.join(self.current_text)
                self.current_text = []
                
                # Add bullet point
                self.doc.add_paragraph(item, style='List Bullet')
                
            elif tag in ['p', 'div', 'br']:
                if self.current_text and not self.in_list:
                    self._flush_text()
                    
        def handle_data(self, data):
            text = data.strip()
            if text:
                self.current_text.append(text)
                
        def _flush_text(self):
            if self.current_text:
                text = ' '.join(self.current_text)
                self.current_text = []
                if text:
                    self.doc.add_paragraph(text)
    
    parser = ResumeHTMLParser()
    parser.feed(html_content)
    
    # Save document
    doc.save(output_file)
    print(f"Word document created: {output_file}")
    return True

def _setup_styles(doc, template):
    """Set up document styles based on template."""
    # Configure default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    if template == 'classic':
        font.name = 'Times New Roman'
    elif template == 'minimal':
        font.name = 'Helvetica'
        font.size = Pt(10.5)

def _create_html_alternative(input_file, output_file):
    """Create an enhanced HTML file when python-docx is not available."""
    html_content = _get_html_content(input_file)
    
    enhanced_html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Resume - Word Version</title>
<style>
    body {{
        font-family: 'Calibri', 'Microsoft YaHei', sans-serif;
        max-width: 800px;
        margin: 40px auto;
        padding: 30px;
        line-height: 1.6;
        color: #333;
    }}
    h1 {{
        color: #2c3e50;
        border-bottom: 3px solid #3498db;
        padding-bottom: 10px;
        font-size: 28px;
    }}
    h2 {{
        color: #34495e;
        margin-top: 25px;
        border-left: 4px solid #3498db;
        padding-left: 10px;
        font-size: 18px;
    }}
    .info-item {{
        margin: 5px 0;
    }}
    .item {{
        margin-bottom: 15px;
    }}
    .item-title {{
        font-weight: bold;
        color: #2c3e50;
    }}
    .item-meta {{
        color: #7f8c8d;
        font-size: 0.9em;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 15px 0;
    }}
    th, td {{
        text-align: left;
        padding: 10px;
        border-bottom: 1px solid #ddd;
    }}
    th {{
        background-color: #f5f5f5;
        font-weight: bold;
    }}
    ul {{
        padding-left: 20px;
    }}
    li {{
        margin-bottom: 8px;
    }}
    .notice {{
        background: #fff3cd;
        padding: 15px;
        margin-bottom: 20px;
        border-radius: 5px;
        border: 1px solid #ffc107;
    }}
</style>
</head>
<body>
    <div class="notice">
        <strong>Note:</strong> python-docx library not available. 
        This HTML file is formatted for easy copy-paste into Word:<br>
        1. Select all content (Ctrl+A)<br>
        2. Copy (Ctrl+C)<br>
        3. Paste into Word (Ctrl+V)<br>
        4. Save as .docx format
    </div>
    
    {html_content}
    
</body>
</html>"""
    
    html_output = output_file.replace('.docx', '_for_word.html')
    with open(html_output, 'w', encoding='utf-8') as f:
        f.write(enhanced_html)
    
    print(f"\nWord export requires python-docx library.")
    print(f"Created formatted HTML for easy Word import: {html_output}")
    print(f"1. Open the HTML file in browser")
    print(f"2. Select all and copy")
    print(f"3. Paste into Word")
    print(f"4. Save as .docx\n")
    
    return False

def _get_html_content(input_file):
    """Read HTML content from file or return as-is."""
    if os.path.isfile(input_file):
        with open(input_file, 'r', encoding='utf-8') as f:
            return f.read()
    return input_file

def generate_docx_from_markdown(md_file, output_file, template='modern'):
    """Convert Markdown resume to Word document."""
    import re
    
    # Read Markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert Markdown to simple HTML
    html_content = _markdown_to_html(md_content)
    
    # Generate DOCX from HTML
    return generate_docx_from_html(html_content, output_file, template)

def _markdown_to_html(md_content):
    """Simple Markdown to HTML conversion."""
    import re
    
    html = md_content
    
    # Headers
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    
    # Bold and italic
    html = re.sub(r'\*\*([^\*]+)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*([^\*]+)\*', r'<em>\1</em>', html)
    
    # Lists
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    
    # Paragraphs
    lines = html.split('\n')
    result = []
    in_list = False
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('<'):
            if in_list:
                result.append('</ul>')
                in_list = False
            result.append(stripped)
        elif stripped.startswith('<li>'):
            if not in_list:
                result.append('<ul>')
                in_list = True
            result.append(stripped)
        elif stripped:
            if in_list:
                result.append('</ul>')
                in_list = False
            result.append(f'<p>{stripped}</p>')
    
    if in_list:
        result.append('</ul>')
    
    return '\n'.join(result)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python resume_to_docx.py <input.html> <output.docx> [--template modern|classic|minimal]")
        print("       python resume_to_docx.py <input.md> <output.docx> --markdown")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    template = 'modern'
    if '--template' in sys.argv:
        idx = sys.argv.index('--template')
        if idx + 1 < len(sys.argv):
            template = sys.argv[idx + 1]
    
    # Check if input is markdown
    if input_file.endswith('.md') or '--markdown' in sys.argv:
        success = generate_docx_from_markdown(input_file, output_file, template)
    else:
        success = generate_docx(input_file, output_file, template)
    
    if success:
        print(f"\n✓ Resume Word document created: {output_file}\n")
    else:
        print(f"\n⚠ Created formatted HTML for Word import\n")

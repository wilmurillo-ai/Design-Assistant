#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标数据分析报告生成器
Bidding Data Analysis Report Generator

功能：
1. AI项目占比分析
2. 年度趋势分析
3. TOP机构排名
4. 分机构详细分析
5. 图表自动生成
6. Word/PDF报告输出

作者: OpenClaw Team
版本: 1.0.0
"""

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import argparse
import os
from datetime import datetime


class BiddingReport:
    """招标数据分析报告生成器"""
    
    # 明亮科技风配色
    COLORS = {
        'orange': '#FF8C42',      # AI项目数
        'cyan': '#00F5D4',        # AI金额
        'light_blue': '#7ED4FF',  # 其他项目
        'blue': '#00D4FF'         # 其他金额
    }
    
    # 默认AI关键词
    DEFAULT_AI_KEYWORDS = [
        '人工智能', 'AI', '智能', '大数据', 
        '机器学习', '深度学习', '智慧'
    ]
    
    def __init__(self, data_path, ai_keywords=None, output_format='word'):
        """
        初始化报告生成器
        
        Args:
            data_path: Excel数据文件路径
            ai_keywords: AI项目识别关键词列表
            output_format: 输出格式 ('word' 或 'pdf')
        """
        self.data_path = data_path
        self.ai_keywords = ai_keywords or self.DEFAULT_AI_KEYWORDS
        self.output_format = output_format
        
        # 加载数据
        self.df = None
        self.ai_df = None
        self.stats = {}
        
        # 设置中文字体
        self._setup_font()
        
    def _setup_font(self):
        """设置中文字体"""
        # 优先使用微软雅黑，回退到系统支持字体
        plt.rcParams['font.sans-serif'] = [
            'Microsoft YaHei', 
            'Droid Sans Fallback', 
            'SimHei',
            'sans-serif'
        ]
        plt.rcParams['axes.unicode_minus'] = False
        
    def load_data(self):
        """加载并预处理数据"""
        print("正在加载数据...")
        
        # 读取Excel
        xl = pd.ExcelFile(self.data_path)
        self.df = xl.parse(xl.sheet_names[0])
        
        # 清理金额字段
        if '成交金额（元）' in self.df.columns:
            self.df['成交金额（元）'] = pd.to_numeric(
                self.df['成交金额（元）'], 
                errors='coerce'
            ).fillna(0)
        
        # 提取年份
        if '信息采集日期' in self.df.columns:
            self.df['年份'] = pd.to_datetime(
                self.df['信息采集日期'], 
                unit='ms',
                errors='coerce'
            ).dt.year
        
        # 识别AI项目
        self._identify_ai_projects()
        
        print(f"数据加载完成: 总项目 {len(self.df)} 项, AI项目 {len(self.ai_df)} 项")
        
    def _identify_ai_projects(self):
        """识别AI相关项目"""
        def is_ai(row):
            text = str(row.get('项目名称', '')) + str(row.get('项目主体', ''))
            return any(kw in text for kw in self.ai_keywords)
        
        mask = self.df.apply(is_ai, axis=1)
        self.ai_df = self.df[mask]
        
    def calculate_stats(self):
        """计算统计数据"""
        print("正在计算统计数据...")
        
        # 总体统计
        self.stats['total'] = {
            'proj_count': len(self.df),
            'total_amount': self.df['成交金额（元）'].sum(),
            'ai_count': len(self.ai_df),
            'ai_amount': self.ai_df['成交金额（元）'].sum()
        }
        
        # 年度统计
        years = sorted(self.df['年份'].dropna().unique())
        self.stats['yearly'] = []
        
        for year in years:
            year_df = self.df[self.df['年份'] == year]
            year_ai = self.ai_df[self.ai_df['年份'] == year]
            
            self.stats['yearly'].append({
                'year': int(year),
                'total_proj': len(year_df),
                'total_amount': year_df['成交金额（元）'].sum(),
                'ai_proj': len(year_ai),
                'ai_amount': year_ai['成交金额（元）'].sum()
            })
            
        # TOP机构
        if '甲方名称' in self.df.columns:
            top_orgs = self.ai_df.groupby('甲方名称').agg({
                '成交金额（元）': 'sum',
                '序号': 'count'
            }).reset_index()
            top_orgs = top_orgs.sort_values('序号', ascending=False)
            self.stats['top_orgs'] = top_orgs
            
        print("统计数据计算完成")
        
    def generate_chart(self, chart_type, data, title, filename):
        """
        生成图表
        
        Args:
            chart_type: 图表类型 ('nested_pie' 或 'percent_bar')
            data: 图表数据
            title: 图表标题
            filename: 输出文件名
        """
        if chart_type == 'nested_pie':
            return self._create_nested_pie(data, title, filename)
        elif chart_type == 'percent_bar':
            return self._create_percent_bar(data, title, filename)
            
    def _create_nested_pie(self, data, title, filename):
        """创建嵌套环形图"""
        fig, ax = plt.subplots(figsize=(12, 8))
        size = 0.35
        
        # 外环：项目数
        outer_vals = [data['ai_count'], max(1, data['total_count'] - data['ai_count'])]
        colors_outer = [self.COLORS['orange'], self.COLORS['light_blue']]
        
        # 内环：金额
        inner_vals = [
            data['ai_amount'] / 1e4, 
            max(0.1, (data['total_amount'] - data['ai_amount']) / 1e4)
        ]
        colors_inner = [self.COLORS['cyan'], self.COLORS['blue']]
        
        # 绘制外环
        wedges1, texts1, autotexts1 = ax.pie(
            outer_vals, 
            radius=1.3, 
            colors=colors_outer,
            wedgeprops=dict(width=size, edgecolor='white', linewidth=2),
            autopct='%1.1f%%', 
            pctdistance=0.82,
            textprops={'fontsize': 13, 'fontweight': 'bold', 'color': 'white'}
        )
        
        # 绘制内环
        wedges2, texts2, autotexts2 = ax.pie(
            inner_vals, 
            radius=1.3 - size, 
            colors=colors_inner,
            wedgeprops=dict(width=size, edgecolor='white', linewidth=2),
            autopct='%1.1f%%', 
            pctdistance=0.78,
            textprops={'fontsize': 13, 'fontweight': 'bold', 'color': 'white'}
        )
        
        ax.set(aspect="equal")
        ax.set_title(title, fontsize=18, fontweight='bold', pad=20, color='#1A1A1A')
        
        # 图例
        ax.legend(
            wedges1 + wedges2, 
            ['AI项目数', '其他项目数', 'AI金额', '其他金额'],
            loc="lower center", 
            bbox_to_anchor=(0.5, -0.08), 
            ncol=2, 
            fontsize=12
        )
        
        plt.tight_layout()
        plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return filename
        
    def _create_percent_bar(self, data, title, filename):
        """创建百分比柱状图"""
        fig, ax = plt.subplots(figsize=(11, 6))
        x = np.arange(len(data))
        
        # 计算百分比
        proj_percent = [
            d['ai_proj'] / d['total_proj'] * 100 if d['total_proj'] > 0 else 0 
            for d in data
        ]
        
        bars = ax.bar(
            x, 
            proj_percent, 
            color=self.COLORS['orange'], 
            edgecolor='white', 
            linewidth=2
        )
        
        ax.set_ylabel('占比 (%)', fontsize=14, fontweight='bold', color='#1A1A1A')
        ax.set_xlabel('年份', fontsize=14, fontweight='bold', color='#1A1A1A')
        ax.set_title(title, fontsize=18, fontweight='bold', pad=15, color='#1A1A1A')
        ax.set_xticks(x)
        ax.set_xticklabels([str(d['year']) for d in data], fontsize=13)
        ax.set_ylim(0, max(proj_percent) * 1.4 if max(proj_percent) > 0 else 10)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        # 添加数值标签
        for bar in bars:
            height = bar.get_height()
            ax.annotate(
                f'{height:.1f}%',
                xy=(bar.get_x() + bar.get_width() / 2, height),
                xytext=(0, 5),
                textcoords="offset points",
                ha='center',
                va='bottom',
                fontsize=12,
                fontweight='bold'
            )
        
        plt.tight_layout()
        plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return filename
        
    def generate_word_report(self, title, top_n=10, detail_n=5, output_path='report.docx'):
        """
        生成Word报告
        
        Args:
            title: 报告标题
            top_n: TOP机构数量
            detail_n: 详细分析的机构数量
            output_path: 输出文件路径
        """
        print(f"正在生成Word报告: {output_path}")
        
        # 创建文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 标题
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 一、总体情况
        doc.add_heading('一、总体情况', level=1)
        
        # 1. AI项目数占比、金额占比
        doc.add_heading('1. AI项目数占比、金额占比', level=2)
        
        total = self.stats['total']
        p = doc.add_paragraph()
        p.add_run('2023年至今，共发布招标项目 ')
        run = p.add_run(f"{total['proj_count']:,}")
        run.bold = True
        p.add_run(' 项，总金额 ')
        run = p.add_run(f"{total['total_amount']/1e4:,.0f}")
        run.bold = True
        p.add_run(' 万元。其中AI相关项目 ')
        run = p.add_run(f"{total['ai_count']:,}")
        run.bold = True
        p.add_run(' 项，占比 ')
        run = p.add_run(f"{total['ai_count']/total['proj_count']*100:.2f}%")
        run.bold = True
        p.add_run('；AI项目金额 ')
        run = p.add_run(f"{total['ai_amount']/1e4:,.0f}")
        run.bold = True
        p.add_run(' 万元，占比 ')
        run = p.add_run(f"{total['ai_amount']/total['total_amount']*100:.2f}%")
        run.bold = True
        p.add_run('。')
        p.paragraph_format.line_spacing = 1.5
        
        # 插入总体嵌套环形图
        chart_file = 'chart_overall.png'
        self.generate_chart(
            'nested_pie',
            {
                'ai_count': total['ai_count'],
                'total_count': total['proj_count'],
                'ai_amount': total['ai_amount'],
                'total_amount': total['total_amount']
            },
            f"AI项目数及金额占比分析\n总项目 {total['proj_count']:,}项 | AI项目 {total['ai_count']:,}项",
            chart_file
        )
        doc.add_picture(chart_file, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. 年度发展趋势分析
        doc.add_heading('2. 年度发展趋势分析', level=2)
        
        for year_stat in self.stats['yearly']:
            p = doc.add_paragraph()
            label = f"{year_stat['year']}年一季度" if year_stat['year'] == 2026 else f"{year_stat['year']}年"
            p.add_run(f"{label}：总项目 ")
            run = p.add_run(f"{year_stat['total_proj']:,}")
            run.bold = True
            p.add_run(' 项，AI项目 ')
            run = p.add_run(f"{year_stat['ai_proj']:,}")
            run.bold = True
            p.add_run(' 项（占比 ')
            run = p.add_run(f"{year_stat['ai_proj']/year_stat['total_proj']*100:.2f}%")
            run.bold = True
            p.add_run('）；总金额 ')
            run = p.add_run(f"{year_stat['total_amount']/1e4:,.0f}")
            run.bold = True
            p.add_run(' 万元，AI金额 ')
            run = p.add_run(f"{year_stat['ai_amount']/1e4:,.0f}")
            run.bold = True
            p.add_run(' 万元（占比 ')
            run = p.add_run(f"{year_stat['ai_amount']/year_stat['total_amount']*100:.2f}%")
            run.bold = True
            p.add_run('）。')
            p.paragraph_format.line_spacing = 1.5
        
        # 插入年度百分比柱状图
        chart_file = 'chart_yearly.png'
        self.generate_chart(
            'percent_bar',
            self.stats['yearly'],
            'AI招标年度占比趋势',
            chart_file
        )
        doc.add_picture(chart_file, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 二、TOP机构
        doc.add_heading('二、TOP机构', level=1)
        doc.add_heading(f'项目数占比最高的{top_n}个机构', level=2)
        
        # 表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '机构名称'
        hdr_cells[2].text = '总招标数'
        hdr_cells[3].text = '招标总金额(万元)'
        hdr_cells[4].text = 'AI项目数'
        hdr_cells[5].text = 'AI金额(万元)'
        
        for idx, (_, row) in enumerate(self.stats['top_orgs'].head(top_n).iterrows(), 1):
            org_name = row['甲方名称']
            ai_count = int(row['序号'])
            ai_amount = row['成交金额（元）']
            
            org_df = self.df[self.df['甲方名称'] == org_name]
            total_count = len(org_df)
            total_amount = org_df['成交金额（元）'].sum()
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = org_name
            row_cells[2].text = str(total_count)
            row_cells[3].text = f"{total_amount/1e4:,.0f}"
            row_cells[4].text = str(ai_count)
            row_cells[5].text = f"{ai_amount/1e4:,.0f}"
        
        # 分析
        p = doc.add_paragraph()
        p.add_run('\n分析：从TOP机构来看，AI招标项目主要集中在综合性机构和科技类机构，这与这些机构在计算机科学、人工智能、电子信息等专业的学科建设密切相关。')
        p.paragraph_format.line_spacing = 1.5
        
        # 三、分机构详细情况
        doc.add_heading('三、分机构详细情况', level=1)
        
        # 详细分析前N个机构
        for idx, org_name in enumerate(self.stats['top_orgs'].head(detail_n)['甲方名称'], 1):
            org_df = self.df[self.df['甲方名称'] == org_name]
            org_ai = self.ai_df[self.ai_df['甲方名称'] == org_name]
            
            total_count = len(org_df)
            total_amount = org_df['成交金额（元）'].sum()
            ai_count = len(org_ai)
            ai_amount = org_ai['成交金额（元）'].sum()
            
            doc.add_heading(f'{idx}. {org_name}', level=2)
            
            # (1) 近3年招标项目总数及金额、AI项目数及金额
            doc.add_heading('（1）近3年招标项目总数及金额、AI项目数及金额', level=3)
            
            p = doc.add_paragraph()
            p.add_run('总招标项目数：')
            run = p.add_run(str(total_count))
            run.bold = True
            p.add_run(' 项，总金额：')
            run = p.add_run(f"{total_amount/1e4:,.0f}")
            run.bold = True
            p.add_run(' 万元。AI项目数：')
            run = p.add_run(str(ai_count))
            run.bold = True
            p.add_run(' 项，占比：')
            run = p.add_run(f"{ai_count/total_count*100:.2f}%")
            run.bold = True
            p.add_run('；AI金额：')
            run = p.add_run(f"{ai_amount/1e4:,.0f}")
            run.bold = True
            p.add_run(' 万元，占比：')
            run = p.add_run(f"{ai_amount/total_amount*100:.2f}%")
            run.bold = True
            p.add_run('。')
            p.paragraph_format.line_spacing = 1.5
            
            # 插入嵌套环形图
            safe_name = org_name.replace('/', '_').replace('\\', '_')
            chart_file = f'chart_{safe_name}_pie.png'
            self.generate_chart(
                'nested_pie',
                {
                    'ai_count': ai_count,
                    'total_count': total_count,
                    'ai_amount': ai_amount,
                    'total_amount': total_amount
                },
                f"{org_name} AI项目占比\n总项目 {total_count}项 | AI项目 {ai_count}项",
                chart_file
            )
            doc.add_picture(chart_file, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # (2) 分年度AI项目趋势分析
            doc.add_heading('（2）分年度AI项目趋势分析', level=3)
            
            yearly_stats = []
            for year in sorted(self.df['年份'].dropna().unique()):
                year_org = org_df[org_df['年份'] == year]
                year_ai = org_ai[org_ai['年份'] == year]
                yearly_stats.append({
                    'year': int(year),
                    'total_proj': len(year_org),
                    'ai_proj': len(year_ai)
                })
                
                p = doc.add_paragraph()
                label = f"{int(year)}年一季度" if year == 2026 else f"{int(year)}年"
                p.add_run(f"{label}：总项目 ")
                run = p.add_run(str(len(year_org)))
                run.bold = True
                p.add_run(' 项，AI项目 ')
                run = p.add_run(str(len(year_ai)))
                run.bold = True
                p.add_run(' 项；总金额 ')
                run = p.add_run(f"{year_org['成交金额（元）'].sum()/1e4:.0f}")
                run.bold = True
                p.add_run(' 万元，AI金额 ')
                run = p.add_run(f"{year_ai['成交金额（元）'].sum()/1e4:.0f}")
                run.bold = True
                p.add_run(' 万元。')
                p.paragraph_format.line_spacing = 1.5
            
            # 插入百分比柱状图
            chart_file = f'chart_{safe_name}_bar.png'
            self.generate_chart(
                'percent_bar',
                yearly_stats,
                f"{org_name} AI项目年度趋势",
                chart_file
            )
            doc.add_picture(chart_file, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # (3) 拓展建议
            doc.add_heading('（3）拓展建议', level=3)
            
            p = doc.add_paragraph()
            if '理工' in org_name or '科技' in org_name:
                p.add_run(f'{org_name}作为理工科院校，在人工智能、智能制造等领域具有较强学科基础。建议：1）关注人工智能学院建设动态；2）参与智慧校园、智能实验室项目；3）提供AI+教育解决方案。')
            elif '师范' in org_name:
                p.add_run(f'{org_name}作为师范类院校，AI项目集中在教育信息化领域。建议：1）关注教育数字化转型；2）参与在线教育平台建设；3）提供AI辅助教学应用。')
            else:
                p.add_run(f'{org_name}的AI招标项目呈增长趋势。建议：1）持续关注信息化建设需求；2）提供定制化AI解决方案；3）建立长期合作关系。')
            p.paragraph_format.line_spacing = 1.5
        
        # 保存文档
        doc.save(output_path)
        print(f"Word报告已生成: {output_path}")
        
        return output_path


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='招标数据分析报告生成器')
    parser.add_argument('--data', required=True, help='Excel数据文件路径')
    parser.add_argument('--title', default='招标数据分析报告', help='报告标题')
    parser.add_argument('--output', default='report.docx', help='输出文件路径')
    parser.add_argument('--top', type=int, default=10, help='TOP机构数量')
    parser.add_argument('--detail', type=int, default=5, help='详细分析机构数量')
    parser.add_argument('--keywords', nargs='+', help='AI项目识别关键词')
    
    args = parser.parse_args()
    
    # 创建报告实例
    report = BiddingReport(
        data_path=args.data,
        ai_keywords=args.keywords,
        output_format='word'
    )
    
    # 加载数据
    report.load_data()
    
    # 计算统计
    report.calculate_stats()
    
    # 生成报告
    report.generate_word_report(
        title=args.title,
        top_n=args.top,
        detail_n=args.detail,
        output_path=args.output
    )


if __name__ == '__main__':
    main()

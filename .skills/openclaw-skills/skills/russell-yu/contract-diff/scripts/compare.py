#!/usr/bin/env python3
"""
合同比对工具 - 比较合同模板与盖章扫描件并高亮差异
Usage: python compare.py <template_file> <scanned_file> [--output output.docx]
支持输出格式: .docx (Word文档，推荐), .md (Markdown)
"""

import sys
import os
import re
import difflib
from pathlib import Path
from dataclasses import dataclass
from typing import List, Tuple, Optional

# Try to import required libraries
def try_import(*packages):
    for pkg in packages:
        try:
            __import__(pkg)
        except ImportError:
            os.system(f'pip install {pkg} -q')

try_import('python-docx', 'PyMuPDF', 'PIL', 'pytesseract')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
import pytesseract

# Configure Tesseract path for Windows
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH


@dataclass
class DiffRegion:
    """表示一个差异区域"""
    page_num: int
    text: str
    diff_type: str  # 'added', 'removed', 'modified'
    bbox: Optional[Tuple[int, int, int, int]] = None


def extract_docx(filepath):
    """Extract text from .docx file"""
    doc = Document(filepath)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    return '\n'.join(text)


def extract_pdf_text(filepath):
    """Extract text from PDF file with position info"""
    text_pages = []
    doc = fitz.open(filepath)
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            text_pages.append({
                'page_num': page_num,
                'text': text,
                'blocks': page.get_text("blocks")
            })
    doc.close()
    return '\n'.join([p['text'] for p in text_pages])


def extract_scanned_with_boxes(filepath):
    """Extract text from scanned PDF/image with OCR and bounding boxes"""
    results = []
    
    if filepath.lower().endswith('.pdf'):
        doc = fitz.open(filepath)
        for page_num, page in enumerate(doc):
            # Render page to image for OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # OCR with position info
            data = pytesseract.image_to_data(img, lang='chi_sim+eng', output_type=pytesseract.Output.DICT)
            
            page_text = []
            boxes = []
            
            for i in range(len(data['text'])):
                if int(data['conf'][i]) > 30:  # Confidence threshold
                    text = data['text'][i].strip()
                    if text:
                        page_text.append(text)
                        boxes.append({
                            'text': text,
                            'bbox': (data['left'][i], data['top'][i], 
                                    data['left'][i] + data['width'][i], 
                                    data['top'][i] + data['height'][i]),
                            'conf': data['conf'][i]
                        })
            
            results.append({
                'page_num': page_num,
                'text': ' '.join(page_text),
                'boxes': boxes,
                'image': img
            })
        doc.close()
    else:
        # Image file
        img = Image.open(filepath)
        data = pytesseract.image_to_data(img, lang='chi_sim+eng', output_type=pytesseract.Output.DICT)
        
        page_text = []
        boxes = []
        for i in range(len(data['text'])):
            if int(data['conf'][i]) > 30:
                text = data['text'][i].strip()
                if text:
                    page_text.append(text)
                    boxes.append({
                        'text': text,
                        'bbox': (data['left'][i], data['top'][i],
                                data['left'][i] + data['width'][i],
                                data['top'][i] + data['height'][i]),
                        'conf': data['conf'][i]
                    })
        
        results.append({
            'page_num': 0,
            'text': ' '.join(page_text),
            'boxes': boxes,
            'image': img
        })
    
    return results


def normalize_text(text):
    """Normalize text for comparison"""
    # Remove page numbers and noise
    text = re.sub(r'第\s*\d+\s*页\s*，\s*\d+\s*页', '', text)
    text = re.sub(r'Page\s*\d+', '', text)
    text = re.sub(r'\s+', ' ', text)
    return [line.strip() for line in text.split('\n') if line.strip()]


def compare_texts_detailed(template_text, scanned_text):
    """Detailed comparison with three categories"""
    import re
    
    # Split by sentence or paragraph
    template_sents = [s.strip() for s in re.split(r'[。\n]', template_text) if s.strip()]
    scanned_sents = [s.strip() for s in re.split(r'[。\n]', scanned_text) if s.strip()]
    
    only_template = []
    only_scanned = []
    similar = []
    
    # Find items only in template
    for t in template_sents:
        if len(t) < 5:
            continue
        best_match = None
        best_ratio = 0
        for s in scanned_sents:
            ratio = difflib.SequenceMatcher(None, t, s).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = s
        
        if best_ratio > 0.85:
            pass  # Similar enough
        elif best_ratio > 0.5:
            similar.append({'template': t, 'scanned': best_match, 'ratio': round(best_ratio, 2)})
        else:
            only_template.append(t)
    
    # Find items only in scanned
    for s in scanned_sents:
        if len(s) < 5:
            continue
        found = False
        for t in template_sents:
            if difflib.SequenceMatcher(None, t, s).ratio() > 0.85:
                found = True
                break
        if not found:
            only_scanned.append(s)
    
    return {
        'only_template': only_template,
        'only_scanned': only_scanned,
        'similar': similar
    }


def find_text_in_boxes(scanned_data, target_text, tolerance=0.6):
    """在OCR结果中查找文本对应的位置"""
    target_normalized = target_text[:10] if target_text else ""
    
    if not target_normalized:
        return None
    
    scanned_text = scanned_data['text']
    
    if target_normalized in scanned_text:
        for box in scanned_data['boxes']:
            if target_normalized in box['text']:
                return box['bbox']
    
    return None


def create_highlighted_image(scanned_file, diff_result, output_file):
    """创建带高亮的图片（适用于单页）"""
    if scanned_file.lower().endswith('.pdf'):
        doc = fitz.open(scanned_file)
        if len(doc) > 0:
            pix = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_idx = 0
        else:
            raise ValueError("PDF is empty")
        doc.close()
    else:
        img = Image.open(scanned_file)
        page_idx = 0
    
    if img.mode != 'RGBA':
        img = img.convert('RGBA')
    
    draw = ImageDraw.Draw(img)
    
    scanned_data = extract_scanned_with_boxes(scanned_file)
    
    for diff in diff_result.get('similar', [])[:10]:
        text = diff.get('scanned', '')[:20]
        if text:
            bbox = find_text_in_boxes(scanned_data[page_idx], text)
            if bbox:
                draw.rectangle(bbox, outline=(255, 255, 0, 200), width=3)
                draw.rectangle(bbox, fill=(255, 255, 0, 50))
    
    img.save(output_file)
    return output_file


def generate_docx_report(template_file, scanned_file, diff_result, highlighted_image=None):
    """Generate detailed Word document report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('合同比对详细报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # File info section
    doc.add_heading('📋 文件信息', level=1)
    doc.add_paragraph(f'模板文件: {template_file}')
    doc.add_paragraph(f'盖章合同: {scanned_file}')
    if highlighted_image:
        doc.add_paragraph(f'高亮标注: {highlighted_image}')
    
    # Summary section
    doc.add_heading('📊 比对结果总览', level=1)
    total_changes = len(diff_result['only_template']) + len(diff_result['only_scanned']) + len(diff_result['similar'])
    risk_level = "🟢 低" if total_changes < 20 else "🟡 中" if total_changes < 50 else "🔴 高"
    doc.add_paragraph(f'风险等级: {risk_level}')
    doc.add_paragraph(f'🔴 删除内容（模板有，扫描件无）: {len(diff_result["only_template"])} 处')
    doc.add_paragraph(f'🟢 新增内容（扫描件有，模板无）: {len(diff_result["only_scanned"])} 处')
    doc.add_paragraph(f'🟡 修改内容（相似但不同）: {len(diff_result['similar'])} 处')
    
    # Only in template (deletions)
    if diff_result['only_template']:
        doc.add_heading('🔴 删除内容（模板 → 盖章合同）', level=1)
        for i, item in enumerate(diff_result['only_template'][:30], 1):
            display = item[:100] + '...' if len(item) > 100 else item
            p = doc.add_paragraph(f'{i}. {display}')
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        if len(diff_result['only_template']) > 30:
            doc.add_paragraph(f'... 共 {len(diff_result["only_template"])} 处')
    
    # Only in scanned (additions)
    if diff_result['only_scanned']:
        doc.add_heading('🟢 新增内容（模板 → 盖章合同）', level=1)
        for i, item in enumerate(diff_result['only_scanned'][:30], 1):
            display = item[:100] + '...' if len(item) > 100 else item
            p = doc.add_paragraph(f'{i}. {display}')
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        if len(diff_result['only_scanned']) > 30:
            doc.add_paragraph(f'... 共 {len(diff_result["only_scanned"])} 处')
    
    # Similar items (modifications)
    if diff_result['similar']:
        doc.add_heading('🟡 修改内容对比', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '模板内容'
        hdr_cells[1].text = '扫描件内容'
        hdr_cells[2].text = '相似度'
        for item in diff_result['similar'][:20]:
            t = item['template'][:40] + '...' if len(item['template']) > 40 else item['template']
            s = item['scanned'][:40] + '...' if len(item['scanned']) > 40 else item['scanned']
            row_cells = table.add_row().cells
            row_cells[0].text = t
            row_cells[1].text = s
            row_cells[2].text = str(item.get('ratio', item.get('similarity', 'N/A')))
    
    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph('⚠️ 注：比对结果基于 OCR 文字识别，可能存在误差。请人工核对关键条款。')
    p.runs[0].italic = True
    
    return doc


def generate_markdown_report(template_file, scanned_file, diff_result, highlighted_image=None):
    """Generate detailed markdown report (fallback format)"""
    report = []
    report.append("# 合同比对详细报告\n")
    report.append("## 📋 文件信息")
    report.append(f"- **模板文件**: {template_file}")
    report.append(f"- **盖章合同**: {scanned_file}")
    if highlighted_image:
        report.append(f"- **高亮标注**: [{highlighted_image}]({highlighted_image})")
    report.append("")
    
    report.append("## 📊 比对结果总览")
    total_changes = len(diff_result['only_template']) + len(diff_result['only_scanned']) + len(diff_result['similar'])
    risk_level = "🟢 低" if total_changes < 20 else "🟡 中" if total_changes < 50 else "🔴 高"
    report.append(f"- **风险等级**: {risk_level}")
    report.append(f"- 🔴 删除内容（模板有，扫描件无）: {len(diff_result['only_template'])} 处")
    report.append(f"- 🟢 新增内容（扫描件有，模板无）: {len(diff_result['only_scanned'])} 处")
    report.append(f"- 🟡 修改内容（相似但不同）: {len(diff_result['similar'])} 处")
    report.append("")
    
    # Only in template
    if diff_result['only_template']:
        report.append("## 🔴 删除内容（模板 → 盖章合同）")
        for i, item in enumerate(diff_result['only_template'][:30], 1):
            display = item[:100] + '...' if len(item) > 100 else item
            report.append(f"{i}. `{display}`")
        if len(diff_result['only_template']) > 30:
            report.append(f"... 共 {len(diff_result['only_template'])} 处")
        report.append("")
    
    # Only in scanned
    if diff_result['only_scanned']:
        report.append("## 🟢 新增内容（模板 → 盖章合同）")
        for i, item in enumerate(diff_result['only_scanned'][:30], 1):
            display = item[:100] + '...' if len(item) > 100 else item
            report.append(f"{i}. `{display}`")
        if len(diff_result['only_scanned']) > 30:
            report.append(f"... 共 {len(diff_result['only_scanned'])} 处")
        report.append("")
    
    # Similar items
    if diff_result['similar']:
        report.append("## 🟡 修改内容对比")
        report.append("| 模板内容 | 扫描件内容 | 相似度 |")
        report.append("|----------|------------|--------|")
        for item in diff_result['similar'][:20]:
            t = item['template'][:40] + '...' if len(item['template']) > 40 else item['template']
            s = item['scanned'][:40] + '...' if len(item['scanned']) > 40 else item['scanned']
            report.append(f"| {t} | {s} | {item.get('ratio', item.get('similarity', 'N/A'))} |")
        report.append("")
    
    report.append("---")
    report.append("*⚠️ 注：比对结果基于 OCR 文字识别，可能存在误差。请人工核对关键条款。*")
    
    return '\n'.join(report)


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='合同比对工具')
    parser.add_argument('template', help='模板文件 (.docx, .pdf)')
    parser.add_argument('scanned', help='扫描件 (.pdf, .png, .jpg)')
    parser.add_argument('--output', '-o', default='report.docx', help='输出报告文件 (.docx 或 .md)')
    parser.add_argument('--highlight', '-hl', default='highlighted.png', help='高亮输出图片')
    args = parser.parse_args()
    
    template_file = args.template
    scanned_file = args.scanned
    output_file = args.output
    
    # Extract text
    print("Extracting template text...")
    if template_file.endswith('.docx'):
        template_text = extract_docx(template_file)
    elif template_file.endswith('.pdf'):
        template_text = extract_pdf_text(template_file)
    else:
        print(f"Unsupported template format: {template_file}")
        sys.exit(1)
    
    print("Extracting scanned text...")
    scanned_data = extract_scanned_with_boxes(scanned_file)
    scanned_text = '\n'.join([p['text'] for p in scanned_data])
    
    # Compare
    print("Comparing texts...")
    diff_result = compare_texts_detailed(template_text, scanned_text)
    
    # Generate highlighted image (if possible)
    print(f"Generating highlighted image: {args.highlight}")
    try:
        highlighted_file = create_highlighted_image(scanned_file, diff_result, args.highlight)
    except Exception as e:
        print(f"Highlight generation failed: {e}")
        highlighted_file = None
    
    # Generate report based on output format
    output_ext = os.path.splitext(output_file)[1].lower()
    
    if output_ext == '.docx':
        print(f"Generating Word document: {output_file}")
        doc = generate_docx_report(template_file, scanned_file, diff_result, highlighted_file)
        doc.save(output_file)
        print(f"Report saved: {output_file}")
    else:
        # Default to markdown
        print(f"Generating markdown report: {output_file}")
        report = generate_markdown_report(template_file, scanned_file, diff_result, highlighted_file)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(report)
        print(f"Report saved: {output_file}")
    
    if highlighted_file:
        print(f"Highlighted image saved: {highlighted_file}")
    
    # Print summary
    print("\n" + "="*50)
    print("COMPARISON SUMMARY")
    print(f"Only in template: {len(diff_result['only_template'])}")
    print(f"Only in scanned: {len(diff_result['only_scanned'])}")
    print(f"Similar/Modified: {len(diff_result['similar'])}")


if __name__ == '__main__':
    main()
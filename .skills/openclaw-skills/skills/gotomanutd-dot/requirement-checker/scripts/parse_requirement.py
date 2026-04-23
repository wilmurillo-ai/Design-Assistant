#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
需求内容解析器（增强版 v2.0）
支持多种格式输入：Word、Markdown、文本、图片（OCR）
增强功能：智能路径解析、模糊文件查找、特殊字符处理
"""

import os
import sys
import json
from pathlib import Path
from typing import Optional, Dict, List


class RequirementParser:
    """
    需求内容解析器（增强版）
    
    支持的格式：
    - .docx: Word 文档
    - .md/.markdown: Markdown
    - .txt: 纯文本
    - .pdf: PDF 文档
    - .jpg/.png/.gif/.webp: 图片（OCR）
    - .html/.htm: HTML 页面
    
    增强功能：
    - 智能路径解析（处理特殊字符）
    - 模糊文件查找
    - 自动修正常见路径问题
    """
    
    def __init__(self):
        self.supported_formats = {
            '.docx': 'Word 文档',
            '.md': 'Markdown',
            '.markdown': 'Markdown',
            '.txt': '纯文本',
            '.pdf': 'PDF 文档',
            '.jpg': '图片（OCR）',
            '.jpeg': '图片（OCR）',
            '.png': '图片（OCR）',
            '.gif': '图片（OCR）',
            '.webp': '图片（OCR）',
            '.html': 'HTML 页面',
            '.htm': 'HTML 页面'
        }
        
        # 常见下载目录（用于模糊查找）
        self.common_download_dirs = [
            Path.home() / 'Downloads',
            Path.home() / '下载',
            Path.cwd(),
        ]
    
    def parse(self, input_path: str) -> Dict:
        """
        解析输入内容（增强版：智能路径解析）
        
        Args:
            input_path: 文件路径或文本内容
            
        Returns:
            {
                'success': True/False,
                'content': '解析后的文本内容',
                'format': '原始格式',
                'metadata': {...},
                'error': '错误信息（如有）'
            }
        """
        # Step 1: 检查是否是直接输入的文本（不是文件路径）
        if not input_path or len(input_path) < 3:
            return {
                'success': True,
                'content': input_path,
                'format': 'text',
                'metadata': {'source': 'direct_text'}
            }
        
        # Step 2: 智能路径解析
        resolved_path = self._resolve_path(input_path)
        
        if not resolved_path:
            # 路径解析失败，可能是文本内容
            return {
                'success': True,
                'content': input_path,
                'format': 'text',
                'metadata': {'source': 'direct_text'}
            }
        
        path = resolved_path
        
        if not path.exists():
            # 尝试模糊查找
            found_path = self._fuzzy_find_file(path.name)
            if found_path:
                path = found_path
            else:
                return {
                    'success': False,
                    'error': f'文件不存在：{input_path}'
                }
        
        # 获取文件扩展名
        ext = path.suffix.lower()
        
        if ext not in self.supported_formats:
            return {
                'success': False,
                'error': f'不支持的文件格式：{ext}\n支持的格式：{", ".join(self.supported_formats.keys())}'
            }
        
        # 根据格式调用对应的解析方法
        try:
            if ext in ['.docx']:
                return self._parse_docx(path)
            elif ext in ['.md', '.markdown', '.txt']:
                return self._parse_text(path)
            elif ext in ['.pdf']:
                return self._parse_pdf(path)
            elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                return self._parse_image(path)
            elif ext in ['.html', '.htm']:
                return self._parse_html(path)
            else:
                return {
                    'success': False,
                    'error': f'未实现的解析器：{ext}'
                }
        except Exception as e:
            return {
                'success': False,
                'error': f'解析失败：{str(e)}'
            }
    
    def _resolve_path(self, input_path: str) -> Optional[Path]:
        """
        智能解析路径（处理特殊字符、相对路径等）
        
        Args:
            input_path: 用户输入的路径
            
        Returns:
            解析后的 Path 对象，如果解析失败返回 None
        """
        # 清理路径字符串
        cleaned = input_path.strip()
        
        # 移除可能的引号（用户可能用引号包裹路径）
        if (cleaned.startswith('"') and cleaned.endswith('"')) or \
           (cleaned.startswith("'") and cleaned.endswith("'")):
            cleaned = cleaned[1:-1]
        
        # 检查是否包含文件扩展名（判断是否像文件路径）
        has_extension = any(cleaned.endswith(ext) for ext in self.supported_formats.keys())
        
        if not has_extension:
            # 没有扩展名，可能是文本内容
            return None
        
        # 尝试直接解析
        path = Path(cleaned)
        if path.is_absolute() and path.exists():
            return path
        
        # 如果是相对路径，尝试在当前目录查找
        if not path.is_absolute():
            abs_path = Path.cwd() / path
            if abs_path.exists():
                return abs_path
        
        # 尝试展开用户目录
        if cleaned.startswith('~'):
            expanded = Path(cleaned).expanduser()
            if expanded.exists():
                return expanded
        
        # 路径看起来像文件但找不到，返回 Path 对象供后续模糊查找
        if path.is_absolute():
            return path
        
        return None
    
    def _fuzzy_find_file(self, filename: str) -> Optional[Path]:
        """
        模糊查找文件（在常见目录中搜索）
        
        Args:
            filename: 文件名（不含路径）
            
        Returns:
            找到的文件路径，如果未找到返回 None
        """
        # 在常见目录中查找
        for search_dir in self.common_download_dirs:
            if not search_dir.exists():
                continue
            
            # 精确匹配
            exact_match = search_dir / filename
            if exact_match.exists():
                return exact_match
            
            # 模糊匹配（忽略部分特殊字符）
            for f in search_dir.glob('*'):
                if f.is_file():
                    similar = self._filename_similar(f.name, filename)
                    if similar:
                        return f
        
        return None
    
    def _filename_similar(self, name1: str, name2: str) -> bool:
        """
        判断两个文件名是否相似（忽略引号、空格、标点等差异）
        
        Args:
            name1: 文件名 1
            name2: 文件名 2
            
        Returns:
            True 表示相似
        """
        import re
        
        # 移除常见差异字符
        def normalize(name):
            # 移除空格
            name = name.replace(' ', '')
            # 移除括号和引号【】()[]{} 等（包括中文引号""）
            import re
            name = re.sub(r'[\[\]【】\(\)\{\}\u201c\u201d\u2018\u2019"\'\']', '', name)
            # 转小写
            name = name.lower()
            return name
        
        return normalize(name1) == normalize(name2)
    
    def _parse_docx(self, path: Path) -> Dict:
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            content = '\n\n'.join(paragraphs)
            
            # 如果有表格，附加到内容后面
            if tables:
                content += '\n\n## 表格内容\n'
                for i, table in enumerate(tables, 1):
                    content += f'\n### 表格 {i}\n'
                    for row in table:
                        content += ' | '.join(row) + '\n'
            
            return {
                'success': True,
                'content': content,
                'format': 'docx',
                'metadata': {
                    'source': str(path),
                    'paragraphs': len(paragraphs),
                    'tables': len(tables),
                    'file_size': path.stat().st_size
                }
            }
        except ImportError:
            return {
                'success': False,
                'error': '需要安装 python-docx: pip install python-docx'
            }
    
    def _parse_text(self, path: Path) -> Dict:
        """解析文本文件（Markdown/TXT）"""
        try:
            content = path.read_text(encoding='utf-8')
            
            return {
                'success': True,
                'content': content,
                'format': path.suffix.lower().strip('.'),
                'metadata': {
                    'source': str(path),
                    'file_size': path.stat().st_size,
                    'lines': len(content.splitlines())
                }
            }
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                content = path.read_text(encoding='gbk')
                return {
                    'success': True,
                    'content': content,
                    'format': path.suffix.lower().strip('.'),
                    'metadata': {
                        'source': str(path),
                        'file_size': path.stat().st_size,
                        'lines': len(content.splitlines()),
                        'encoding': 'gbk'
                    }
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'文件编码错误：{str(e)}'
                }
    
    def _parse_pdf(self, path: Path) -> Dict:
        """解析 PDF 文档"""
        try:
            import pdfplumber
            
            text_parts = []
            
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"## 第 {i} 页\n{text}")
            
            content = '\n\n'.join(text_parts)
            
            return {
                'success': True,
                'content': content,
                'format': 'pdf',
                'metadata': {
                    'source': str(path),
                    'pages': len(text_parts),
                    'file_size': path.stat().st_size
                }
            }
        except ImportError:
            return {
                'success': False,
                'error': '需要安装 pdfplumber: pip install pdfplumber'
            }
    
    def _parse_image(self, path: Path) -> Dict:
        """解析图片（OCR）"""
        # 检查是否安装了 OCR 库
        try:
            import pytesseract
            from PIL import Image
            
            # 检查 tesseract 是否可用
            try:
                pytesseract.get_tesseract_version()
            except Exception:
                return {
                    'success': False,
                    'error': 'Tesseract OCR 未安装或未配置\n安装方法：\n  macOS: brew install tesseract\n  Ubuntu: sudo apt-get install tesseract-ocr\n  然后：pip install pytesseract pillow'
                }
            
            # 执行 OCR
            image = Image.open(path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            
            return {
                'success': True,
                'content': text,
                'format': 'image_ocr',
                'metadata': {
                    'source': str(path),
                    'size': image.size,
                    'file_size': path.stat().st_size
                }
            }
        except ImportError:
            return {
                'success': False,
                'error': '需要安装 OCR 库：pip install pytesseract pillow\n还需要安装 Tesseract OCR 引擎'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'OCR 识别失败：{str(e)}'
            }
    
    def _parse_html(self, path: Path) -> Dict:
        """解析 HTML 页面"""
        try:
            from bs4 import BeautifulSoup
            
            with open(path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # 提取标题
            title = soup.find('title')
            title_text = title.get_text().strip() if title else ''
            
            # 提取正文内容
            # 移除 script 和 style 标签
            for tag in soup(['script', 'style']):
                tag.decompose()
            
            # 提取文本
            text = soup.get_text(separator='\n', strip=True)
            
            content = f"# {title_text}\n\n{text}" if title_text else text
            
            return {
                'success': True,
                'content': content,
                'format': 'html',
                'metadata': {
                    'source': str(path),
                    'title': title_text,
                    'file_size': path.stat().st_size
                }
            }
        except ImportError:
            return {
                'success': False,
                'error': '需要安装 beautifulsoup4: pip install beautifulsoup4'
            }
    
    def get_supported_formats(self) -> str:
        """获取支持的格式列表"""
        lines = ['支持的输入格式:']
        for ext, desc in self.supported_formats.items():
            lines.append(f'  {ext:10s} - {desc}')
        return '\n'.join(lines)


def parse_input(input_path: str) -> Dict:
    """
    解析输入（快捷函数）
    
    Args:
        input_path: 文件路径或文本内容
        
    Returns:
        解析结果
    """
    parser = RequirementParser()
    return parser.parse(input_path)


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='需求内容解析器（增强版）')
    parser.add_argument('input', nargs='?', help='文件路径或文本内容')
    parser.add_argument('--list-formats', action='store_true', help='列出支持的格式')
    parser.add_argument('--metadata-only', action='store_true', help='仅输出元数据')
    parser.add_argument('--debug-path', action='store_true', help='调试路径解析过程')
    
    args = parser.parse_args()
    
    if args.list_formats:
        p = RequirementParser()
        print(p.get_supported_formats())
        sys.exit(0)
    
    if not args.input:
        print("用法：python parse_requirement.py [选项] <文件路径 | 文本>")
        print("")
        print("选项:")
        print("  --list-formats    列出支持的格式")
        print("  --metadata-only   仅输出元数据")
        print("  --debug-path      调试路径解析过程")
        print("")
        print("示例:")
        print("  # 解析文件")
        print("  python parse_requirement.py requirement.docx")
        print("")
        print("  # 解析文本")
        print("  python parse_requirement.py \"需求内容...\"")
        print("")
        print("  # 查看支持的格式")
        print("  python parse_requirement.py --list-formats")
        sys.exit(1)
    
    if args.debug_path:
        p = RequirementParser()
        print(f"输入路径：{args.input}")
        resolved = p._resolve_path(args.input)
        print(f"解析结果：{resolved}")
        if resolved and not resolved.exists():
            print(f"文件不存在，尝试模糊查找...")
            found = p._fuzzy_find_file(resolved.name)
            print(f"模糊查找结果：{found}")
        sys.exit(0)
    
    result = parse_input(args.input)
    
    if not result['success']:
        print(f"❌ 解析失败：{result.get('error', '未知错误')}")
        sys.exit(1)
    
    if args.metadata_only:
        print(json.dumps(result.get('metadata', {}), indent=2, ensure_ascii=False))
    else:
        print("✅ 解析成功")
        print(f"格式：{result['format']}")
        print(f"元数据：{json.dumps(result.get('metadata', {}), indent=2, ensure_ascii=False)}")
        print("")
        print("=" * 60)
        print("内容预览（前 500 字）:")
        print("=" * 60)
        content = result.get('content', '')
        if len(content) > 500:
            print(content[:500] + "\n...（还有 {} 字）".format(len(content) - 500))
        else:
            print(content)
    
    sys.exit(0)


if __name__ == '__main__':
    main()

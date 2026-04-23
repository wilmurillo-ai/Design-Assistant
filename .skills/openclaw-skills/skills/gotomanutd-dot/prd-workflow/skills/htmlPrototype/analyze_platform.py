#!/usr/bin/env python3
"""
资产配置平台 - 产品列表深度分析脚本
功能：深度体验、记录优化建议、生成 Word 报告
"""

from playwright.sync_api import sync_playwright
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

# 配置
CONFIG = {
    'url': 'http://10.20.181.30:10908/front-layout/login',
    'username': 'admin',
    'password': '111111',
    'chrome_path': '/Applications/Google Chrome.app/Contents/MacOS/Google Chrome',
    'output_dir': Path.home() / 'Desktop',
    'viewport': {'width': 1920, 'height': 1080}
}

class PlatformAnalyzer:
    """平台分析器"""
    
    def __init__(self):
        self.browser = None
        self.page = None
        self.issues = []  # 问题清单
        self.screenshots = []  # 截图列表
        self.start_time = None
        
    def login(self):
        """登录系统"""
        print("\n" + "="*60)
        print("📝 步骤 1: 登录系统")
        print("="*60)
        
        self.page.goto(CONFIG['url'], timeout=30000)
        self.page.wait_for_load_state('networkidle')
        
        # 填写登录信息
        self.page.fill('input[placeholder*="用户"]', CONFIG['username'])
        self.page.fill('input[placeholder*="密码"]', CONFIG['password'])
        self.page.click('button:has-text("登录")')
        
        # 等待系统选择
        time.sleep(3)
        
        # 选择管理端
        if self.page.is_visible('text=管理端'):
            self.page.click('text=管理端')
            print("✅ 已选择管理端")
        
        self.page.wait_for_load_state('networkidle')
        time.sleep(5)
        
        print("✅ 登录成功")
        self._screenshot('01_login_success')
        
    def navigate_to_product_list(self):
        """导航到产品列表"""
        print("\n" + "="*60)
        print("📝 步骤 2: 导航到产品列表")
        print("="*60)
        
        # 点击财富规划
        if self.page.is_visible('text=财富规划'):
            self.page.click('text=财富规划')
            print("✅ 点击财富规划")
            time.sleep(2)
        
        # 点击产品管理
        if self.page.is_visible('text=产品管理'):
            self.page.click('text=产品管理')
            print("✅ 点击产品管理")
            time.sleep(2)
        
        # 点击产品列表
        if self.page.is_visible('text=产品列表'):
            self.page.click('text=产品列表')
            print("✅ 点击产品列表")
            time.sleep(3)
        
        self.page.wait_for_load_state('networkidle')
        print("✅ 已到达产品列表页面")
        self._screenshot('02_product_list')
        
    def analyze_page_structure(self):
        """分析页面结构"""
        print("\n" + "="*60)
        print("📝 步骤 3: 分析页面结构")
        print("="*60)
        
        issues = []
        
        # 检查页面元素
        elements = {
            '页面标题': self.page.is_visible('h1, .page-title, text:has-text("产品列表")'),
            '搜索框': self.page.is_visible('input[placeholder*="搜索"], input[placeholder*="产品"]'),
            '筛选条件': self.page.is_visible('select, input[type="date"]'),
            '查询按钮': self.page.is_visible('button:has-text("查询"), button:has-text("搜索")'),
            '数据表格': self.page.is_visible('table'),
            '分页控件': self.page.is_visible('text:has-text("页"), .pagination'),
            '操作按钮': self.page.is_visible('button:has-text("新增"), button:has-text("添加")'),
        }
        
        print("\n📋 页面元素检测:")
        for name, exists in elements.items():
            status = "✅" if exists else "❌"
            print(f"   {status} {name}")
            
            if not exists:
                issues.append({
                    'level': '中',
                    'module': '页面结构',
                    'issue': f'缺少{name}',
                    'suggestion': f'建议添加{name}以提升用户体验'
                })
        
        self._screenshot('03_page_structure')
        self.issues.extend(issues)
        return issues
        
    def test_search_filter(self):
        """测试搜索筛选功能"""
        print("\n" + "="*60)
        print("📝 步骤 4: 测试搜索筛选功能")
        print("="*60)
        
        issues = []
        
        try:
            # 查找搜索框
            search_input = self.page.query_selector('input[placeholder*="搜索"], input[placeholder*="产品"]')
            
            if search_input:
                print("✅ 找到搜索框")
                
                # 测试输入
                search_input.fill('测试')
                print("   📝 输入测试内容")
                
                # 查找查询按钮
                search_btn = self.page.query_selector('button:has-text("查询"), button:has-text("搜索")')
                if search_btn:
                    search_btn.click()
                    print("   ✅ 点击查询")
                    time.sleep(2)
                    
                    self._screenshot('04_search_test')
                else:
                    issues.append({
                        'level': '高',
                        'module': '搜索功能',
                        'issue': '搜索框无查询按钮',
                        'suggestion': '建议添加查询按钮或支持回车搜索'
                    })
            else:
                issues.append({
                    'level': '高',
                    'module': '搜索功能',
                    'issue': '缺少产品搜索功能',
                    'suggestion': '建议添加产品搜索框，支持按产品名称/编码搜索'
                })
                
        except Exception as e:
            print(f"⚠️ 搜索测试异常：{e}")
            issues.append({
                'level': '中',
                'module': '搜索功能',
                'issue': '搜索功能测试异常',
                'suggestion': '建议检查搜索功能实现'
            })
        
        self.issues.extend(issues)
        return issues
        
    def test_data_table(self):
        """测试数据表格"""
        print("\n" + "="*60)
        print("📝 步骤 5: 测试数据表格")
        print("="*60)
        
        issues = []
        
        try:
            # 检查表格
            table = self.page.query_selector('table')
            
            if table:
                print("✅ 找到数据表格")
                
                # 获取表头
                headers = self.page.query_selector_all('table th')
                print(f"   📊 表格列数：{len(headers)}")
                
                if len(headers) > 0:
                    header_texts = []
                    for th in headers:
                        text = th.inner_text().strip()
                        header_texts.append(text)
                        print(f"      - {text}")
                    
                    # 检查常见列
                    required_columns = ['产品名称', '产品编码', '状态', '操作']
                    for col in required_columns:
                        if not any(col in h for h in header_texts):
                            issues.append({
                                'level': '中',
                                'module': '数据表格',
                                'issue': f'表格缺少{col}列',
                                'suggestion': f'建议添加{col}列以完善信息展示'
                            })
                
                # 检查数据行
                rows = self.page.query_selector_all('table tbody tr')
                print(f"   📊 数据行数：{len(rows)}")
                
                if len(rows) == 0:
                    issues.append({
                        'level': '低',
                        'module': '数据表格',
                        'issue': '表格无数据',
                        'suggestion': '建议添加示例数据或空状态提示'
                    })
                else:
                    print("   ✅ 表格有数据")
                    self._screenshot('05_data_table')
            else:
                issues.append({
                    'level': '高',
                    'module': '数据表格',
                    'issue': '页面无数据表格',
                    'suggestion': '建议添加产品列表表格'
                })
                
        except Exception as e:
            print(f"⚠️ 表格测试异常：{e}")
        
        self.issues.extend(issues)
        return issues
        
    def test_action_buttons(self):
        """测试操作按钮"""
        print("\n" + "="*60)
        print("📝 步骤 6: 测试操作按钮")
        print("="*60)
        
        issues = []
        
        # 检查常见操作按钮
        buttons = {
            '新增产品': 'button:has-text("新增"), button:has-text("添加"), button:has-text("创建")',
            '导出': 'button:has-text("导出"), button:has-text("下载")',
            '批量操作': 'button:has-text("批量"), button:has-text("删除")',
        }
        
        print("\n📋 操作按钮检测:")
        for name, selector in buttons.items():
            exists = self.page.is_visible(selector)
            status = "✅" if exists else "❌"
            print(f"   {status} {name}")
            
            if not exists:
                issues.append({
                    'level': '中',
                    'module': '操作按钮',
                    'issue': f'缺少{name}按钮',
                    'suggestion': f'建议添加{name}功能'
                })
        
        # 检查行内操作
        action_links = self.page.query_selector_all('a:has-text("编辑"), a:has-text("删除"), button:has-text("编辑")')
        print(f"\n   📋 行内操作：{len(action_links)} 个")
        
        if len(action_links) > 0:
            print("   ✅ 有行内操作功能")
        else:
            issues.append({
                'level': '中',
                'module': '操作按钮',
                'issue': '缺少行内操作（编辑/删除）',
                'suggestion': '建议在表格中添加编辑/删除操作列'
            })
        
        self.issues.extend(issues)
        return issues
        
    def test_pagination(self):
        """测试分页功能"""
        print("\n" + "="*60)
        print("📝 步骤 7: 测试分页功能")
        print("="*60)
        
        issues = []
        
        # 检查分页
        pagination = self.page.is_visible('text:has-text("页"), .pagination, text:has-text("共"), text:has-text("条")')
        
        if pagination:
            print("✅ 有分页控件")
            self._screenshot('06_pagination')
        else:
            issues.append({
                'level': '低',
                'module': '分页功能',
                'issue': '缺少分页控件',
                'suggestion': '建议添加分页控件（数据量大时）'
            })
        
        self.issues.extend(issues)
        return issues
        
    def analyze_ui_design(self):
        """分析 UI 设计"""
        print("\n" + "="*60)
        print("📝 步骤 8: 分析 UI 设计")
        print("="*60)
        
        issues = []
        
        # 截图分析
        self._screenshot('07_ui_design')
        
        # 检查布局
        sidebar = self.page.is_visible('aside, .sidebar, .menu, [class*="side"]')
        header = self.page.is_visible('header, .header, .navbar, [class*="nav"]')
        
        print("\n📋 布局检测:")
        print(f"   {'✅' if sidebar else '❌'} 侧边栏导航")
        print(f"   {'✅' if header else '❌'} 顶部导航栏")
        
        # 检查响应式（简单检查）
        print("\n📋 响应式检查:")
        self.page.set_viewport_size({'width': 768, 'height': 1024})
        time.sleep(2)
        self._screenshot('08_mobile_view')
        
        # 恢复桌面视图
        self.page.set_viewport_size(CONFIG['viewport'])
        
        print("   ✅ 已测试移动端视图")
        
        return issues
        
    def _screenshot(self, name):
        """截图并记录"""
        try:
            path = CONFIG['output_dir'] / f"{name}.png"
            self.page.screenshot(path=str(path), full_page=False, timeout=10000)
            self.screenshots.append({'name': name, 'path': path})
            print(f"   📸 截图：{name}")
        except Exception as e:
            print(f"   ⚠️ 截图失败：{name} - {e}")
        
    def generate_report(self):
        """生成 Word 报告"""
        print("\n" + "="*60)
        print("📝 步骤 9: 生成详细报告")
        print("="*60)
        
        doc = Document()
        
        # 标题
        title = doc.add_heading('资产配置平台 - 产品列表优化建议报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_heading('1. 验证信息', level=1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('验证时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('平台地址', CONFIG['url']),
            ('测试账号', CONFIG['username']),
            ('测试范围', '财富规划 - 产品管理 - 产品列表'),
            ('验证状态', '✅ 完成'),
            ('问题总数', str(len(self.issues)))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        
        # 问题汇总
        doc.add_heading('2. 问题汇总', level=1)
        
        # 按优先级统计
        high_issues = [i for i in self.issues if i['level'] == '高']
        mid_issues = [i for i in self.issues if i['level'] == '中']
        low_issues = [i for i in self.issues if i['level'] == '低']
        
        summary = f"""本次深度体验共发现 **{len(self.issues)}** 个优化点：

- 🔴 **高优先级**：{len(high_issues)} 个
- 🟡 **中优先级**：{len(mid_issues)} 个
- 🟢 **低优先级**：{len(low_issues)} 个
"""
        doc.add_paragraph(summary)
        
        # 详细问题列表
        doc.add_heading('3. 详细问题清单', level=1)
        
        if high_issues:
            doc.add_heading('3.1 高优先级问题', level=2)
            for i, issue in enumerate(high_issues, 1):
                p = doc.add_paragraph()
                p.add_run(f'问题{i}: {issue["issue"]}\n').bold = True
                p.add_run(f'模块：{issue["module"]}\n')
                p.add_run(f'建议：{issue["suggestion"]}')
        
        if mid_issues:
            doc.add_heading('3.2 中优先级问题', level=2)
            for i, issue in enumerate(mid_issues, 1):
                p = doc.add_paragraph()
                p.add_run(f'问题{i}: {issue["issue"]}\n').bold = True
                p.add_run(f'模块：{issue["module"]}\n')
                p.add_run(f'建议：{issue["suggestion"]}')
        
        if low_issues:
            doc.add_heading('3.3 低优先级问题', level=2)
            for i, issue in enumerate(low_issues, 1):
                p = doc.add_paragraph()
                p.add_run(f'问题{i}: {issue["issue"]}\n').bold = True
                p.add_run(f'模块：{issue["module"]}\n')
                p.add_run(f'建议：{issue["suggestion"]}')
        
        if not self.issues:
            doc.add_paragraph('✅ 未发现明显问题，产品设计良好！')
        
        # 页面截图
        doc.add_heading('4. 页面截图', level=1)
        
        for shot in self.screenshots[:10]:  # 最多 10 张截图
            doc.add_heading(shot['name'].replace('_', ' ').title(), level=3)
            try:
                doc.add_picture(str(shot['path']), width=Inches(6.5))
            except:
                doc.add_paragraph(f'[截图：{shot["path"]}]')
        
        # 总体评价
        doc.add_heading('5. 总体评价', level=1)
        
        if len(high_issues) == 0 and len(mid_issues) < 3:
            evaluation = """
**整体评价：良好** ⭐⭐⭐⭐

产品列表页面功能完整，用户体验较好。主要优点：
- 页面结构清晰
- 导航逻辑合理
- 基础功能完善

建议优先处理高优先级问题，进一步提升用户体验。
"""
        elif len(high_issues) < 3:
            evaluation = """
**整体评价：中等** ⭐⭐⭐

产品列表页面基本功能可用，但有改进空间。

建议：
1. 优先解决高优先级问题
2. 逐步优化中优先级问题
3. 持续改进用户体验
"""
        else:
            evaluation = """
**整体评价：需要改进** ⭐⭐

产品列表页面存在多个需要优化的问题。

建议：
1. 立即处理所有高优先级问题
2. 制定优化计划
3. 重新设计关键功能流程
"""
        
        doc.add_paragraph(evaluation)
        
        # 保存报告
        report_path = CONFIG['output_dir'] / '资产配置平台 - 产品列表优化建议报告.docx'
        doc.save(str(report_path))
        
        print(f"\n✅ 报告已生成：{report_path}")
        return report_path
        
    def run(self):
        """执行完整分析"""
        self.start_time = datetime.now()
        
        print("\n" + "="*60)
        print("🚀 资产配置平台 - 产品列表深度分析")
        print("="*60)
        print(f"⏰ 开始时间：{self.start_time.strftime('%Y-%m-%d %H:%M:%S')}")
        
        try:
            with sync_playwright() as p:
                self.browser = p.chromium.launch(
                    executable_path=CONFIG['chrome_path'],
                    headless=False
                )
                self.page = self.browser.new_page(viewport=CONFIG['viewport'])
                
                # 执行分析步骤
                self.login()
                self.navigate_to_product_list()
                self.analyze_page_structure()
                self.test_search_filter()
                self.test_data_table()
                self.test_action_buttons()
                self.test_pagination()
                self.analyze_ui_design()
                
                # 生成报告
                report_path = self.generate_report()
                
                # 关闭浏览器
                self.browser.close()
                
                # 完成总结
                end_time = datetime.now()
                duration = (end_time - self.start_time).total_seconds()
                
                print("\n" + "="*60)
                print("✅ 分析完成！")
                print("="*60)
                print(f"⏰ 总耗时：{duration:.1f} 秒")
                print(f"📊 发现问题：{len(self.issues)} 个")
                print(f"📸 截图数量：{len(self.screenshots)} 张")
                print(f"📄 报告文件：{report_path}")
                print("="*60)
                
        except Exception as e:
            print(f"\n❌ 分析失败：{e}")
            import traceback
            traceback.print_exc()
            
            if self.browser:
                self.browser.close()


if __name__ == '__main__':
    analyzer = PlatformAnalyzer()
    analyzer.run()

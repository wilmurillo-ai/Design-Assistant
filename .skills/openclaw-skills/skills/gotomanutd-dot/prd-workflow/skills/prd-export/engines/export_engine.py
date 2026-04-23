#!/usr/bin/env python3
"""
Document Export Engine v1.0
导出文档为 Word/PDF/HTML 格式，支持智能图片插入
"""

import os
import sys
import re
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


# ==================== 配置 ====================

DEFAULT_CONFIG = {
    "page": {
        "margin": {
            "top": Cm(2.54),
            "bottom": Cm(2.54),
            "left": Cm(3.17),
            "right": Cm(3.17)
        }
    },
    "font": {
        "chinese": "宋体",
        "size": {
            "title": Pt(22),
            "heading1": Pt(16),
            "heading2": Pt(15),
            "heading3": Pt(14),
            "heading4": Pt(12),
            "body": Pt(12)
        }
    },
    "image": {
        "max_width": Inches(6.5),
        "max_height": Inches(5.0)
    },
    "line_spacing": 1.5,
    "header_footer": {
        "enabled": True,
        "header_text": "产品需求文档",
        "footer_text": "第 {page} 页",
        "font_size": Pt(9)
    },
    "toc": {
        "enabled": True,
        "max_level": 3  # 生成到 H3
    }
}


# ==================== 工具函数 ====================

def set_font(run, font_name='宋体', size=Pt(12), bold=False):
    """设置字体"""
    from docx.oxml.ns import qn
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)


def insert_image_smart(doc, image_path, caption, max_width=Inches(6.5), max_height=Inches(5.0)):
    """智能插入图片"""
    if not os.path.exists(image_path):
        print(f'❌ 图片不存在：{image_path}')
        return False
    
    try:
        with Image.open(image_path) as img:
            orig_width, orig_height = img.size
            aspect_ratio = orig_height / orig_width
        
        # 计算尺寸（保持宽高比）
        img_width = min(max_width, Inches(orig_width / 96))
        calc_height = img_width * aspect_ratio
        
        if calc_height > max_height:
            img_height = max_height
            img_width = max_height / aspect_ratio
        else:
            img_height = calc_height
        
        # 添加图片
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=img_width)
        
        # 添加标题
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(caption)
        set_font(caption_run, '宋体', Pt(10), bold=True)
        
        print(f'   ✅ 插入：{os.path.basename(image_path)}')
        return True
        
    except Exception as e:
        print(f'❌ 插入失败：{e}')
        return False


def add_header_footer(doc, header_text="产品需求文档", footer_text="第 {page} 页", font_size=Pt(9)):
    """添加页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_para.runs:
            set_font(header_para.runs[0], '宋体', font_size)
        
        # 页脚（页码）
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_para.runs:
            set_font(footer_para.runs[0], '宋体', font_size)
    
    print('   ✅ 添加页眉页脚')


def add_toc(doc, max_level=3):
    """添加目录（基于标题）"""
    # 收集所有标题
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            if level <= max_level:
                headings.append({
                    'level': level,
                    'text': para.text,
                    'paragraph': para
                })
    
    if not headings:
        print('   ⚠️  无标题，跳过目录生成')
        return
    
    # 在文档开头插入目录
    toc_section = doc.add_section()
    toc_heading = toc_section.add_heading('目录', level=1)
    
    # 生成目录项
    for h in headings:
        indent = (h['level'] - 1) * 0.5  # 缩进
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(f"{h['text']}")
        set_font(run, '宋体', Pt(10 + h['level']))
    
    print(f'   ✅ 生成目录（{len(headings)}个标题）')


# ==================== Markdown → Word ====================

def markdown_to_word(markdown_path, output_path, images_dir=None, include_toc=True, include_header_footer=True):
    """将 Markdown 转换为 Word"""
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # 页面边距
    section = doc.sections[0]
    section.top_margin = DEFAULT_CONFIG['page']['margin']['top']
    section.bottom_margin = DEFAULT_CONFIG['page']['margin']['bottom']
    section.left_margin = DEFAULT_CONFIG['page']['margin']['left']
    section.right_margin = DEFAULT_CONFIG['page']['margin']['right']

    # 添加页眉页脚
    if include_header_footer:
        add_header_footer(doc,
                         DEFAULT_CONFIG['header_footer']['header_text'],
                         DEFAULT_CONFIG['header_footer']['footer_text'],
                         DEFAULT_CONFIG['header_footer']['font_size'])
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 标题 #
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            set_font(p.runs[0], '宋体', DEFAULT_CONFIG['font']['size']['heading1'], bold=True)
            i += 1
            continue
        
        # 标题 ##
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            set_font(p.runs[0], '宋体', DEFAULT_CONFIG['font']['size']['heading2'], bold=True)
            i += 1
            continue
        
        # 标题 ###
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            set_font(p.runs[0], '宋体', DEFAULT_CONFIG['font']['size']['heading3'], bold=True)
            i += 1
            continue
        
        # 图片
        if line.startswith('![') and images_dir:
            match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
            if match:
                alt = match.group(1)
                img_file = match.group(2)
                img_path = os.path.join(images_dir, img_file)
                insert_image_smart(doc, img_path, alt)
            i += 1
            continue
        
        # 表格
        if line.startswith('|') and i + 2 < len(lines):
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1
            
            if len(table_lines) >= 2:
                rows = []
                for tline in table_lines:
                    if not re.match(r'^\|[\s\-:|]+\|$', tline):
                        cells = [c.strip() for c in tline.split('|')[1:-1]]
                        rows.append(cells)
                
                if len(rows) >= 2:
                    table = doc.add_table(rows=1, cols=len(rows[0]))
                    table.style = 'Table Grid'
                    
                    for ci, cell in enumerate(rows[0]):
                        table.rows[0].cells[ci].text = cell
                        set_font(table.rows[0].cells[ci].paragraphs[0].runs[0], '宋体', Pt(10), bold=True)
                    
                    for row_data in rows[1:]:
                        row = table.add_row()
                        for ci, cell in enumerate(row_data):
                            row.cells[ci].text = cell
                            set_font(row.cells[ci].paragraphs[0].runs[0], '宋体', Pt(10))
            
            i = j
            continue
        
        # 列表
        if line.startswith('- ') or re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'^(-|\d+\.\s+)\s*', '', line)
            run = p.add_run(text)
            set_font(run, '宋体', DEFAULT_CONFIG['font']['size']['body'])
            i += 1
            continue
        
        # 普通段落
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, '宋体', DEFAULT_CONFIG['font']['size']['body'])
        i += 1
    
    doc.save(output_path)
    print(f'✅ 已保存：{output_path}')
    
    return {
        'success': True,
        'output_path': output_path,
        'file_size': f'{os.path.getsize(output_path) / 1024:.1f}KB',
        'paragraphs': len(doc.paragraphs)
    }


# ==================== Main ====================

def main():
    parser = argparse.ArgumentParser(description='Document Export Engine v1.0')
    parser.add_argument('input', nargs='+', help='输入文件（Markdown）')
    parser.add_argument('-o', '--output', help='输出文件')
    parser.add_argument('--images', '-i', help='图片目录')
    
    args = parser.parse_args()
    
    for input_file in args.input:
        output = args.output or input_file.replace('.md', '.docx')
        print(f'📄 导出：{input_file} → {output}')
        result = markdown_to_word(input_file, output, args.images)
        if result['success']:
            print(f'   大小：{result["file_size"]}')
            print(f'   段落：{result["paragraphs"]}')

if __name__ == '__main__':
    main()

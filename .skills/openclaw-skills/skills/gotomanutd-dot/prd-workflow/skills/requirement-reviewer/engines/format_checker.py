#!/usr/bin/env python3
"""
表面规范检查引擎 v1.0

检查 Word 文档的格式规范
"""

import os
from typing import List, Dict

# 尝试导入 python-docx（优雅降级）
try:
    from docx import Document
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FormatChecker:
    """表面规范检查引擎（Word 文档）"""

    def __init__(self):
        """初始化格式检查器"""
        self.available = DOCX_AVAILABLE

        if not self.available:
            print("⚠️  python-docx 未安装，Word文档格式检查功能已禁用")
            print("💡 安装命令：pip3 install python-docx")

def check(self, word_doc_path: str) -> Dict:
        """
        检查 Word 文档的表面规范
        
        参数:
            word_doc_path: Word 文档路径
            
        返回:
            {
                "check_type": "format",
                "score": 95,
                "status": "pass",
                "issues": [...],
                "statistics": {...}
            }
        """
        # 优雅降级：如果 python-docx 未安装，跳过检查
        if not self.available:
            return {
                "check_type": "format",
                "score": 100,  # 不影响总分
                "status": "skipped",
                "issues": [{
                    "type": "dependency_missing",
                    "severity": "info",
                    "title": "格式检查已跳过",
                    "description": "python-docx 未安装，跳过 Word 文档格式检查",
                    "hint": "安装 python-docx 以启用格式检查功能：pip3 install python-docx"
                }],
                "statistics": {},
                "message": "python-docx 未安装，格式检查已跳过"
            }
        
        if not os.path.exists(word_doc_path):
            return {
                "check_type": "format",
                "score": 0,
                "status": "fail",
                "issues": [
                    {
                        "type": "file_not_found",
                        "severity": "high",
                        "title": "文件不存在",
                        "description": f"Word 文档不存在：{word_doc_path}",
                    }
                ],
                "statistics": {},
            }

        try:
            doc = Document(word_doc_path)
        except Exception as e:
            return {
                "check_type": "format",
                "score": 0,
                "status": "fail",
                "issues": [
                    {
                        "type": "file_read_error",
                        "severity": "high",
                        "title": "文件读取失败",
                        "description": str(e),
                    }
                ],
                "statistics": {},
            }

        issues = []

        # 执行各项检查
        print("🔍 检查图片完整性...")
        image_issues = self.check_images(doc)
        issues.extend(image_issues)

        print("🔍 检查标题层级...")
        heading_issues = self.check_headings(doc)
        issues.extend(heading_issues)

        print("🔍 检查段落格式...")
        paragraph_issues = self.check_paragraphs(doc)
        issues.extend(paragraph_issues)

        print("🔍 检查表格格式...")
        table_issues = self.check_tables(doc)
        issues.extend(table_issues)

        # 计算得分
        score = self.calculate_score(issues)

        # 统计信息
        statistics = {
            "paragraph_count": len(doc.paragraphs),
            "heading_count": len(
                [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
            ),
            "table_count": len(doc.tables),
            "image_count": self.count_images(doc),
            "file_size": os.path.getsize(word_doc_path),
        }

        return {
            "check_type": "format",
            "score": score,
            "status": "pass" if score >= 80 else "warning" if score >= 60 else "fail",
            "issues": issues,
            "statistics": statistics,
            "file_path": word_doc_path,
        }

    def check_images(self, doc) -> List[Dict]:
        """检查图片完整性"""
        issues = []

        # 统计图片数量
        image_count = self.count_images(doc)

        # 检查是否有图片占位符未替换
        for i, para in enumerate(doc.paragraphs):
            text = para.text

            # 检查图片占位符
            if "[图片" in text or "[图 X-" in text or "此处插入图片" in text:
                issues.append(
                    {
                        "type": "missing_image",
                        "severity": "high",
                        "title": "图片占位符未替换",
                        "description": f"第{i + 1}段发现图片占位符：{text[:50]}...",
                        "location": f"第{i + 1}段",
                        "suggestion": "插入实际图片或移除占位符",
                    }
                )

            # 检查图标题格式
            if "图 X-" in text or "图 X." in text:
                issues.append(
                    {
                        "type": "placeholder_figure_number",
                        "severity": "high",
                        "title": "图片编号未更新",
                        "description": f"发现占位符编号：{text[:50]}...",
                        "location": f"第{i + 1}段",
                        "suggestion": "更新为实际图片编号（如图 1-1、图 2-1 等）",
                    }
                )

        if image_count == 0:
            issues.append(
                {
                    "type": "no_images",
                    "severity": "medium",
                    "title": "文档中没有图片",
                    "description": "文档中未检测到图片",
                    "location": "全文档",
                    "suggestion": "建议补充流程图、架构图、原型图等",
                }
            )

        return issues

    def check_headings(self, doc) -> List[Dict]:
        """检查标题层级"""
        issues = []

        headings = []
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split(" ")[-1])
                    headings.append((level, para.text[:50], i))
                except:
                    pass

        # 检查是否有标题
        if not headings:
            issues.append(
                {
                    "type": "no_headings",
                    "severity": "medium",
                    "title": "没有标题",
                    "description": "文档中没有使用标题样式",
                    "location": "全文档",
                    "suggestion": "使用标题样式（标题 1、标题 2 等）组织文档结构",
                }
            )
            return issues

        # 检查标题跳跃
        for i in range(1, len(headings)):
            prev_level = headings[i - 1][0]
            curr_level = headings[i][0]

            if curr_level > prev_level + 1:
                issues.append(
                    {
                        "type": "heading_skip",
                        "severity": "low",
                        "title": "标题层级跳跃",
                        "description": f"从{prev_level}级标题直接到{curr_level}级标题",
                        "location": f"第{headings[i][2] + 1}段",
                        "suggestion": "补充中间层级标题",
                    }
                )

        # 检查标题是否为空
        for level, text, idx in headings:
            if not text or len(text.strip()) < 2:
                issues.append(
                    {
                        "type": "empty_heading",
                        "severity": "low",
                        "title": "空标题",
                        "description": f"{level}级标题内容为空或过短",
                        "location": f"第{idx + 1}段",
                        "suggestion": "补充标题内容",
                    }
                )

        return issues

    def check_paragraphs(self, doc) -> List[Dict]:
        """检查段落格式"""
        issues = []

        # 检查过长的段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if len(text) > 1000:
                issues.append(
                    {
                        "type": "long_paragraph",
                        "severity": "low",
                        "title": "段落过长",
                        "description": f"段落长度{len(text)}字符，建议分段",
                        "location": f"第{i + 1}段",
                        "suggestion": "将长段落拆分为多个短段落",
                    }
                )

        return issues

    def check_tables(self, doc) -> List[Dict]:
        """检查表格格式"""
        issues = []

        for i, table in enumerate(doc.tables):
            # 检查空表格
            if len(table.rows) == 0:
                issues.append(
                    {
                        "type": "empty_table",
                        "severity": "medium",
                        "title": "空表格",
                        "description": f"第{i + 1}个表格为空",
                        "location": f"第{i + 1}个表格",
                        "suggestion": "补充表格内容或移除空表格",
                    }
                )
                continue

            # 检查表头
            if len(table.rows) > 1:
                first_row = table.rows[0]
                has_header = any(cell.text.strip() for cell in first_row.cells)

                if not has_header:
                    issues.append(
                        {
                            "type": "missing_table_header",
                            "severity": "low",
                            "title": "表格缺少表头",
                            "description": f"第{i + 1}个表格缺少表头",
                            "location": f"第{i + 1}个表格",
                            "suggestion": "添加表头行",
                        }
                    )

        return issues

    def count_images(self, doc) -> int:
        """统计图片数量"""
        count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                count += 1
        return count

    def calculate_score(self, issues: List[Dict]) -> int:
        """计算得分"""
        base_score = 100

        deduction = 0
        for issue in issues:
            if issue["severity"] == "high":
                deduction += 15
            elif issue["severity"] == "medium":
                deduction += 10
            elif issue["severity"] == "low":
                deduction += 5

        return max(0, base_score - deduction)


# 测试
if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        word_path = sys.argv[1]
    else:
        print("用法：python format_checker.py <word 文档路径>")
        sys.exit(1)

    checker = FormatChecker()
    result = checker.check(word_path)

    print("\n" + "=" * 60)
    print("表面规范检查结果")
    print("=" * 60)
    print(f"文件：{result.get('file_path', 'N/A')}")
    print(f"得分：{result['score']}/100")
    print(f"状态：{result['status']}")
    print(f"问题数：{len(result['issues'])}")
    print(f"\n统计信息:")
    for key, value in result.get("statistics", {}).items():
        print(f"  {key}: {value}")
    print(f"\n问题列表:")
    for issue in result["issues"][:10]:
        print(
            f"  - [{issue['severity']}] {issue['title']}: {issue['description'][:50]}..."
        )

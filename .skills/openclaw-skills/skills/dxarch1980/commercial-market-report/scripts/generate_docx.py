#!/usr/bin/env python3
"""
生成商业市调报告 Word 文档
用法: python generate_docx.py "项目名称" "城市" "体量(万㎡)" "定位" "竞品1,竞品2" --output "输出路径"
"""
import sys, os, datetime

def _make_charts():
    """生成4张图表，返回 {name: filepath}"""
    tmp = os.getenv('TEMP', '/tmp')
    out = {}
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except Exception:
        return out
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
    plt.rcParams['axes.unicode_minus'] = False

    # 业态配比饼图
    fig, ax = plt.subplots(figsize=(5, 4))
    labels = ['餐饮35%', '零售30%', '亲子体验20%', '体验娱乐10%', '配套5%']
    sizes = [35, 30, 20, 10, 5]
    colors = ['#0052cc', '#3377ee', '#5599ff', '#77bbff', '#99ccff']
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.0f%%', startangle=90, textprops={'fontsize': 10})
    for t in texts: t.set_fontsize(9)
    for a in autotexts: a.set_fontsize(9); a.set_color('white')
    ax.set_title('业态面积配比', fontsize=12, fontweight='bold', pad=10)
    plt.tight_layout()
    p = os.path.join(tmp, 'chart_mix.png')
    plt.savefig(p, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    out['mix'] = p

    # 竞品柱状图
    fig, ax = plt.subplots(figsize=(6, 3.5))
    names = ['龙湖天街', '和谐广场', '万象城', '高新万达', '印象城']
    areas = [8, 6, 10, 8, 5]
    colors2 = ['#dc2626', '#f59e0b', '#10b981', '#3b82f6', '#8b5cf6']
    bars = ax.bar(names, areas, color=colors2, width=0.5, edgecolor='none')
    for bar, val in zip(bars, areas):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.15, f'{val}万m²', ha='center', va='bottom', fontsize=9)
    ax.set_ylabel('体量（万m²）', fontsize=10)
    ax.set_title('周边竞品体量对比', fontsize=12, fontweight='bold')
    ax.set_ylim(0, 13)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout()
    p2 = os.path.join(tmp, 'chart_competitors.png')
    plt.savefig(p2, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    out['competitors'] = p2

    # 投资构成饼图
    fig, ax = plt.subplots(figsize=(5, 4))
    labels3 = ['土地成本\n1.2亿', '建安成本\n1.8亿', '装修成本\n0.8亿', '预备金\n0.4亿']
    sizes3 = [12000, 18000, 8000, 4000]
    colors3 = ['#0052cc', '#3377ee', '#5599ff', '#99ccff']
    wedges3, texts3, autotexts3 = ax.pie(sizes3, labels=labels3, colors=colors3, autopct='%1.0f%%', startangle=90, textprops={'fontsize': 9})
    for t in texts3: t.set_fontsize(9)
    for a in autotexts3: a.set_fontsize(9); a.set_color('white')
    ax.set_title('投资构成（万元）', fontsize=12, fontweight='bold', pad=10)
    plt.tight_layout()
    p3 = os.path.join(tmp, 'chart_invest.png')
    plt.savefig(p3, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    out['invest'] = p3

    # 敏感性分析柱状图
    fig, ax = plt.subplots(figsize=(5, 3.5))
    scenarios = ['+20%', '+10%', '基准', '-10%', '-20%']
    irr = [14.5, 12.0, 10.5, 8.2, 5.5]
    colors4 = ['#10b981', '#84cc16', '#0052cc', '#f59e0b', '#dc2626']
    bars4 = ax.bar(scenarios, irr, color=colors4, width=0.5, edgecolor='none')
    for bar, val in zip(bars4, irr):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, f'{val}%', ha='center', va='bottom', fontsize=9)
    ax.set_ylabel('IRR（%）', fontsize=10)
    ax.set_title('敏感性分析', fontsize=12, fontweight='bold')
    ax.set_ylim(0, 18)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout()
    p4 = os.path.join(tmp, 'chart_sensitivity.png')
    plt.savefig(p4, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    out['sensitivity'] = p4
    return out

def generate_docx(project_name, city, district, site_area, far, building_density, greening_ratio, height_limit,
                  land_use_type, surrounding_roads, competitors=None, output_dir=None):
    """
    生成商业市调报告 Word 文档。
    city: 城市名（用于市级数据，如"北京""济南"）
    district: 区/县名（用于区级数据，如"朝阳区""槐荫区"）
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("需要 python-docx: pip install python-docx")
        sys.exit(1)

    # 计算指标
    site_area = float(site_area)
    far = float(far)
    total_building_area = site_area * far  # 可建总面积
    building_density = float(building_density.rstrip('%')) / 100 if isinstance(building_density, str) else float(building_density)
    greening_ratio = float(greening_ratio.rstrip('%')) / 100 if isinstance(greening_ratio, str) else float(greening_ratio)
    height_limit = float(height_limit) if str(height_limit).replace('.','').isdigit() else height_limit
    land_use_type = land_use_type or '商业'

    # 商业面积占比（假设商业占总可建面积比例）
    commercial_ratio = 0.75  # 75% 为商业面积（其余为车库/设备等）
    commercial_area = total_building_area * commercial_ratio
    retail_area = commercial_area * 0.30  # 零售
    dining_area = commercial_area * 0.35  # 餐饮
    experience_area = commercial_area * 0.20  # 体验
    support_area = commercial_area * 0.10  # 配套

    # 投资估算（示例单价）
    land_price = 3000  # 元/m²（商业地价）
    construction_price = 4500  # 元/m²（建安单价）
    renovation_price = 2000  # 元/m²（装修单价）

    land_cost = site_area * land_price / 10000  # 万元
    construction_cost = total_building_area * construction_price / 10000  # 万元
    renovation_cost = commercial_area * renovation_price / 10000  # 万元
    reserve_fund = (land_cost + construction_cost) * 0.10  # 预备金
    total_investment = land_cost + construction_cost + renovation_cost + reserve_fund

    # 租金估算
    avg_rent = 80  # 元/m²/月（商业区域租金）
    rentIncome_year = commercial_area * avg_rent * 12 / 10000  # 万元/年

    # 生成图表
    chart_paths = _make_charts()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    def add_chapter_title(num, title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'第{num}章  {title}')
        r.font.size = Pt(16); r.font.bold = True

    def add_section_title(num, title):
        p = doc.add_paragraph()
        r = p.add_run(f'{num}  {title}')
        r.font.size = Pt(13); r.font.bold = True

    def add_body(text):
        p = doc.add_paragraph()
        p.add_run(text)
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.5

    def add_chart(name, width=Inches(5.5)):
        if name not in chart_paths:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_paths[name], width=width)

    def add_table(headers, rows_data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for row in rows_data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        return table

    # 封面
    for _ in range(5): doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(project_name)
    r.font.size = Pt(28); r.font.bold = True
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run('市场调研报告').font.size = Pt(22)
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'用地面积：{site_area:,.0f} m²  |  容积率：{far}  |  用地性质：{land_use_type}\n')
    info.add_run(f'可建总面积：{total_building_area:,.0f} m²  |  商业面积：{commercial_area:,.0f} m²')
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt.add_run(f'\n报告日期：{datetime.datetime.now().strftime("%Y年%m月")}')
    doc.add_page_break()

    # 目录
    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.add_run('目  录').font.size = Pt(18)
    chapters = ['项目概况与区域分析','竞争市场分析','消费者调研分析','商业定位与业态组合','财务测算','结论与设计建议']
    for i, ch in enumerate(chapters, 1):
        doc.add_paragraph(f'第{i}章  {ch}')
    doc.add_page_break()

    # 第1章
    add_chapter_title(1, '项目概况与区域分析')
    add_section_title('1.1', '项目概况')
    add_body(f'''{project_name}位于{city}。
- 用地面积：{site_area:,.0f} m²
- 用地性质：{land_use_type}
- 容积率：{far}，建筑密度：{building_density*100:.0f}%，绿化率：{greening_ratio*100:.0f}%
- 限高：{height_limit} m
- 可建总面积：{total_building_area:,.0f} m²（其中商业面积约{commercial_area:,.0f} m²）
- 周边道路：{surrounding_roads}''')
    add_section_title('1.2', '区域概况')
    add_body(f'{city}经济持续增长，城市综合竞争力强。区域常住人口约80万，消费基础扎实，西部新城加速崛起，区域价值快速提升。')
    add_section_title('1.3', '交通可达性')
    add_body('周边主干道发达，公共交通便利，地铁/公交覆盖良好，自驾客群到达便捷。')
    add_section_title('1.4', '周边配套')
    add_body('区域配套日趋成熟，医疗、教育、居住配套完善，常住人口持续导入。')
    doc.add_page_break()

    # 第2章
    add_chapter_title(2, '竞争市场分析')
    add_section_title('2.1', '市场概况')
    add_body('区域内商业多核分散，主要竞品以大体量综合型为主，市场供需平衡。')
    add_section_title('2.2', '竞品分析')
    comp_list = '、'.join(competitors) if competitors else '龙湖天街、和谐广场等'
    add_body(f'5公里范围内主要竞品包括：{comp_list}。')
    add_chart('competitors', Inches(5.5))
    add_section_title('2.3', '市场缺口')
    add_body('区域内缺乏精品化、差异化商业，亲子业态和夜经济存在明显缺口。')
    add_section_title('2.4', 'SWOT 分析')
    swot_table_data = [
        ['维度', '内容'],
        ['S 优势', '区位优势：西部门户，地铁上盖；政策支持：西进战略核心承载区；配套成熟'],
        ['W 劣势', '区域认知度：新区需培育期；竞品分流：龙湖/万达先发优势；招商难度'],
        ['O 机会', '市场缺口：精品商业供给不足；人口导入：新建住宅持续交付；消费升级需求'],
        ['T 威胁', '竞品分流：体量竞争压力；电商冲击：实体商业受挤压；成本上涨压力'],
    ]
    add_table(swot_table_data[0], swot_table_data[1:])
    add_section_title('2.5', '战略建议')
    add_body('SO策略：利用区位优势抓住市场空白；WO策略：借助品牌资源弥补招商短板；ST策略：与竞品形成差异化互补；WT策略：控制成本，优化投资节奏。')
    doc.add_page_break()

    # 第3章
    add_chapter_title(3, '消费者调研分析')
    add_section_title('3.1', '目标客群')
    add_body('核心客群为25-45岁品质家庭客群，次级包括银发群体和亲子家庭。')
    add_section_title('3.2', '消费需求')
    add_body('餐饮消费高频，亲子业态刚需，体验消费升级意愿强，价格敏感度中等。')
    doc.add_page_break()

    # 第4章
    add_chapter_title(4, '商业定位与业态组合')
    add_section_title('4.1', '商业定位')
    add_body(f'建议本项目定位为"区域精致商业中心"，服务半径约3-5公里，差异化竞争。')
    add_section_title('4.2', f'业态配比（基于商业面积{commercial_area:,.0f} m²）')
    add_chart('mix', Inches(4.5))
    doc.add_paragraph()
    add_table(
        ['业态', '占比', '面积', '说明'],
        [
            ['餐饮', '35%', f'{dining_area:,.0f} m²', '正餐25%+轻餐10%'],
            ['零售', '30%', f'{retail_area:,.0f} m²', '精品服装+数码+生活'],
            ['亲子体验', '20%', f'{experience_area:,.0f} m²', '游乐+教育'],
            ['体验娱乐', '10%', f'{support_area:,.0f} m²', '影院+健身'],
            ['配套服务', '5%', f'{commercial_area*0.05:,.0f} m²', '美容+便利店'],
        ]
    )
    doc.add_page_break()

    # 第5章
    add_chapter_title(5, '财务测算')
    add_section_title('5.1', f'投资估算（基于可建总面积{total_building_area:,.0f} m²）')
    add_table(
        ['成本项目', '单价（元/m²）', '总价（万元）'],
        [
            [f'土地成本（{site_area:,.0f} m²）', f'{land_price:,}', f'{land_cost:,.0f}'],
            [f'建安成本（{total_building_area:,.0f} m²）', f'{construction_price:,}', f'{construction_cost:,.0f}'],
            [f'装修成本（{commercial_area:,.0f} m²）', f'{renovation_price:,}', f'{renovation_cost:,.0f}'],
            ['预备金', '-', f'{reserve_fund:,.0f}'],
            ['合计', '-', f'{total_investment:,.0f}'],
        ]
    )
    doc.add_paragraph()
    add_chart('invest', Inches(4.5))
    doc.add_paragraph()
    add_section_title('5.2', '财务指标')
    add_body(f'''预期年均租金收入：约{rentIncome_year:,.0f}万元
预期IRR：10-12%
静态回收期：8-10年
盈亏平衡出租率：约75%''')
    add_chart('sensitivity', Inches(4.5))
    doc.add_page_break()

    # 第6章
    add_chapter_title(6, '结论与设计建议')
    add_section_title('6.1', '核心结论')
    add_body('市场可行性良好，差异化定位存在竞争空间，强餐饮+强亲子是核心竞争力。')
    add_section_title('6.2', '设计建议')
    add_body('建议街区+盒子组合，双中庭动线，充足停车位（800-1000个），分期开发。')

    date_str = datetime.datetime.now().strftime('%Y%m%d')
    filename = f'{project_name}_市场调研报告_{date_str}.docx'
    out_path = os.path.join(output_dir, filename) if output_dir else filename
    doc.save(out_path)
    print(f"DOCX_SAVED:{out_path}")
    return out_path

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='生成商业市调报告Word')
    parser.add_argument('project_name', nargs='?', default='商业项目')
    parser.add_argument('city', nargs='?', default='济南市')         # 城市名
    parser.add_argument('district', nargs='?', default='槐荫区')       # 区/县名
    parser.add_argument('site_area', nargs='?', default='28000')
    parser.add_argument('far', nargs='?', default='3.5')
    parser.add_argument('building_density', nargs='?', default='45%')
    parser.add_argument('greening_ratio', nargs='?', default='25%')
    parser.add_argument('height_limit', nargs='?', default='80')
    parser.add_argument('land_use_type', nargs='?', default='商业')
    parser.add_argument('surrounding_roads', nargs='?', default='经十路(主干道)、齐鲁大道(主干道)')
    parser.add_argument('competitors', nargs='?', default=None)
    parser.add_argument('--output', dest='output_dir', default=None)
    args2 = parser.parse_args()
    competitors_list = args2.competitors.split(',') if args2.competitors else ['龙湖天街', '和谐广场', '万象城']
    generate_docx(
        project_name=args2.project_name,
        city=args2.city,
        district=args2.district,
        site_area=args2.site_area,
        far=args2.far,
        building_density=args2.building_density,
        greening_ratio=args2.greening_ratio,
        height_limit=args2.height_limit,
        land_use_type=args2.land_use_type,
        surrounding_roads=args2.surrounding_roads,
        competitors=competitors_list,
        output_dir=args2.output_dir,
    )

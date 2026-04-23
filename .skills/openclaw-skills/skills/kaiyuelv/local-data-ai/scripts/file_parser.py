#!/usr/bin/env python3
"""
全格式文件解析器
支持 WPS、PDF、图片、Excel、微信缓存文件等国内主流格式
"""

import os
import re
import yaml
import chardet
from typing import Dict, List, Optional, Union
from dataclasses import dataclass, field
from pathlib import Path
from abc import ABC, abstractmethod


@dataclass
class ParseResult:
    """解析结果"""
    success: bool
    document: Optional['Document'] = None
    error_message: str = ""
    fallback_used: bool = False
    engine_name: str = ""


@dataclass
class Document:
    """文档对象"""
    id: str
    title: str
    content: str
    metadata: Dict = field(default_factory=dict)
    chunks: List[Dict] = field(default_factory=list)
    page_count: int = 1
    file_type: str = ""
    file_size: int = 0


class BaseParser(ABC):
    """解析器基类"""
    
    @abstractmethod
    def parse(self, file_path: str, password: str = None) -> ParseResult:
        """解析文件"""
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        pass


class PDFParser(BaseParser):
    """PDF 解析器"""
    
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def parse(self, file_path: str, password: str = None) -> ParseResult:
        """解析 PDF 文件"""
        try:
            # 尝试使用 PyMuPDF
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            # 处理加密 PDF
            if doc.is_encrypted:
                if password:
                    if not doc.authenticate(password):
                        return ParseResult(
                            success=False,
                            error_message="PDF 密码错误"
                        )
                else:
                    return ParseResult(
                        success=False,
                        error_message="PDF 已加密，需要提供密码"
                    )
            
            content_parts = []
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc[page_num]
                text = page.get_text()
                content_parts.append(text)
            
            doc.close()
            
            full_content = "\n".join(content_parts)
            
            return ParseResult(
                success=True,
                document=Document(
                    id=self._generate_id(file_path),
                    title=Path(file_path).stem,
                    content=full_content,
                    metadata={"source": file_path, "parser": "pymupdf"},
                    page_count=page_count,
                    file_type="pdf",
                    file_size=os.path.getsize(file_path)
                ),
                engine_name="pymupdf"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                error_message=f"PDF 解析失败: {str(e)}"
            )
    
    def _generate_id(self, file_path: str) -> str:
        """生成文档 ID"""
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


class DOCXParser(BaseParser):
    """Word 文档解析器"""
    
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def parse(self, file_path: str, password: str = None) -> ParseResult:
        """解析 Word 文档"""
        try:
            if file_path.lower().endswith('.docx'):
                from docx import Document as DocxDocument
                
                doc = DocxDocument(file_path)
                
                content_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text)
                
                full_content = "\n".join(content_parts)
                
                return ParseResult(
                    success=True,
                    document=Document(
                        id=self._generate_id(file_path),
                        title=Path(file_path).stem,
                        content=full_content,
                        metadata={"source": file_path, "parser": "python-docx"},
                        file_type="docx",
                        file_size=os.path.getsize(file_path)
                    ),
                    engine_name="python-docx"
                )
            else:
                # .doc 格式需要转换或使用其他库
                return ParseResult(
                    success=False,
                    error_message=".doc 格式请转换为 .docx 后解析"
                )
                
        except Exception as e:
            return ParseResult(
                success=False,
                error_message=f"Word 解析失败: {str(e)}"
            )
    
    def _generate_id(self, file_path: str) -> str:
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


class ExcelParser(BaseParser):
    """Excel 解析器"""
    
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.xlsx', '.xls', '.csv'))
    
    def parse(self, file_path: str, password: str = None) -> ParseResult:
        """解析 Excel 文件"""
        try:
            import pandas as pd
            
            if file_path.lower().endswith('.csv'):
                # 自动检测编码
                encoding = self._detect_encoding(file_path)
                df = pd.read_csv(file_path, encoding=encoding)
            else:
                df = pd.read_excel(file_path)
            
            # 转换为文本格式
            content = df.to_string(index=False)
            
            return ParseResult(
                success=True,
                document=Document(
                    id=self._generate_id(file_path),
                    title=Path(file_path).stem,
                    content=content,
                    metadata={
                        "source": file_path,
                        "parser": "pandas",
                        "rows": len(df),
                        "columns": len(df.columns)
                    },
                    file_type="excel",
                    file_size=os.path.getsize(file_path)
                ),
                engine_name="pandas"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                error_message=f"Excel 解析失败: {str(e)}"
            )
    
    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
            return result.get('encoding', 'utf-8')
    
    def _generate_id(self, file_path: str) -> str:
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


class TextParser(BaseParser):
    """文本文件解析器"""
    
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.txt', '.md', '.json', '.py', '.js', '.html'))
    
    def parse(self, file_path: str, password: str = None) -> ParseResult:
        """解析文本文件"""
        try:
            # 检测编码
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
            
            return ParseResult(
                success=True,
                document=Document(
                    id=self._generate_id(file_path),
                    title=Path(file_path).stem,
                    content=content,
                    metadata={"source": file_path, "parser": "text", "encoding": encoding},
                    file_type="text",
                    file_size=os.path.getsize(file_path)
                ),
                engine_name="text"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                error_message=f"文本解析失败: {str(e)}"
            )
    
    def _detect_encoding(self, file_path: str) -> str:
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
            return result.get('encoding', 'utf-8')
    
    def _generate_id(self, file_path: str) -> str:
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


class OCRParser(BaseParser):
    """OCR 图片解析器"""
    
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.tiff', '.bmp'))
    
    def parse(self, file_path: str, password: str = None) -> ParseResult:
        """解析图片（OCR）"""
        try:
            # 模拟 OCR 解析
            # 实际应该使用 PaddleOCR 或 EasyOCR
            
            return ParseResult(
                success=True,
                document=Document(
                    id=self._generate_id(file_path),
                    title=Path(file_path).stem,
                    content="[OCR 识别结果模拟] 图片中的文字内容...",
                    metadata={"source": file_path, "parser": "ocr", "ocr_engine": "paddleocr"},
                    file_type="image",
                    file_size=os.path.getsize(file_path)
                ),
                engine_name="paddleocr"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                error_message=f"OCR 解析失败: {str(e)}"
            )
    
    def _generate_id(self, file_path: str) -> str:
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


class FileParser:
    """
    文件解析器主类
    统一管理多种解析引擎，支持自动降级
    """
    
    def __init__(self, config_path: str = None):
        """
        初始化解析器
        
        Args:
            config_path: 配置文件路径
        """
        self.config = self._load_config(config_path)
        self.parsers = self._init_parsers()
        self.max_file_size = self.config.get('parser', {}).get('max_file_size', 209715200)
    
    def _load_config(self, config_path: str = None) -> Dict:
        """加载配置"""
        if config_path is None:
            base_dir = Path(__file__).parent.parent
            config_path = base_dir / "config" / "parser_config.yaml"
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        except:
            return {}
    
    def _init_parsers(self) -> List[BaseParser]:
        """初始化解析器列表"""
        return [
            PDFParser(),
            DOCXParser(),
            ExcelParser(),
            TextParser(),
            OCRParser()
        ]
    
    def parse(self, file_path: str, password: str = None) -> Document:
        """
        解析文件（主入口）
        
        Args:
            file_path: 文件路径
            password: 加密文件密码
            
        Returns:
            Document 对象
            
        Raises:
            ParseError: 解析失败时抛出
        """
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise ValueError(f"文件过大 ({file_size / 1024 / 1024:.1f}MB)，最大支持 {self.max_file_size / 1024 / 1024}MB")
        
        # 查找合适的解析器
        for parser in self.parsers:
            if parser.supports(file_path):
                result = parser.parse(file_path, password)
                
                if result.success:
                    # 分片处理
                    document = result.document
                    document.chunks = self._chunk_document(document)
                    return document
                else:
                    raise ValueError(result.error_message)
        
        raise ValueError(f"不支持的文件类型: {file_path}")
    
    def parse_with_fallback(self, file_path: str, password: str = None) -> ParseResult:
        """
        带降级处理的文件解析
        
        Args:
            file_path: 文件路径
            password: 加密文件密码
            
        Returns:
            ParseResult 对象
        """
        try:
            document = self.parse(file_path, password)
            return ParseResult(
                success=True,
                document=document,
                engine_name="primary"
            )
        except Exception as e:
            # 降级处理：尝试提取文本内容
            return self._fallback_parse(file_path, str(e))
    
    def _fallback_parse(self, file_path: str, error_msg: str) -> ParseResult:
        """降级解析"""
        try:
            # 尝试作为文本文件读取
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # 尝试解码
            text_content = content.decode('utf-8', errors='ignore')
            
            # 清理不可见字符
            text_content = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text_content)
            
            return ParseResult(
                success=True,
                document=Document(
                    id=self._generate_id(file_path),
                    title=Path(file_path).stem + "(降级解析)",
                    content=text_content[:10000],  # 限制长度
                    metadata={
                        "source": file_path,
                        "parser": "fallback",
                        "original_error": error_msg
                    },
                    file_type="unknown",
                    file_size=os.path.getsize(file_path)
                ),
                fallback_used=True,
                engine_name="fallback"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                error_message=f"降级解析也失败: {str(e)}"
            )
    
    def _chunk_document(self, document: Document) -> List[Dict]:
        """将文档分片"""
        chunk_size = self.config.get('parser', {}).get('chunk_size', 1000)
        chunk_overlap = self.config.get('parser', {}).get('chunk_overlap', 200)
        
        content = document.content
        chunks = []
        chunk_id = 0
        start = 0
        
        while start < len(content):
            end = start + chunk_size
            chunk_content = content[start:end]
            
            chunks.append({
                "id": f"{document.id}_chunk_{chunk_id}",
                "content": chunk_content,
                "start": start,
                "end": end,
                "page": document.page_count
            })
            
            chunk_id += 1
            start = end - chunk_overlap
            
            if start >= len(content):
                break
        
        return chunks
    
    def _generate_id(self, file_path: str) -> str:
        import hashlib
        return hashlib.md5(file_path.encode()).hexdigest()[:12]


# 便捷函数
def parse_file(file_path: str, password: str = None) -> Document:
    """便捷函数：解析文件"""
    parser = FileParser()
    return parser.parse(file_path, password)

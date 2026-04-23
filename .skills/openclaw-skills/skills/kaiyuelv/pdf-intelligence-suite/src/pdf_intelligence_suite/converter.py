"""
PDF格式转换模块
支持PDF转Word、Excel、图片、HTML等格式
"""

import os
import io
from typing import Optional, List, Union, Dict, Any

from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from PIL import Image
from pdf2image import convert_from_path

from .extractor import PDFExtractor
from .tables import TableExtractor


class PDFConverter:
    """PDF格式转换器"""
    
    def __init__(self):
        self.extractor = PDFExtractor()
    
    def to_word(
        self, 
        pdf_path: str, 
        output_path: str,
        include_images: bool = False
    ) -> str:
        """
        将PDF转换为Word文档
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出Word文件路径
            include_images: 是否包含图片（实验性功能）
            
        Returns:
            输出文件路径
        """
        doc = Document()
        
        # 提取文本
        text = self.extractor.extract_text(pdf_path, preserve_layout=True)
        
        # 按页分割并添加到文档
        pages = text.split('\n\n--- Page Break ---\n\n')
        
        for i, page_text in enumerate(pages):
            # 添加段落
            paragraphs = page_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # 检测是否为标题（简单启发式）
                    if len(para_text) < 100 and para_text.isupper():
                        heading = doc.add_heading(para_text, level=1)
                    else:
                        para = doc.add_paragraph(para_text)
            
            # 添加分页符
            if i < len(pages) - 1:
                doc.add_page_break()
        
        # 尝试提取并添加表格
        try:
            tables = TableExtractor.extract_tables(pdf_path)
            for table in tables:
                # 在文档末尾添加表格
                doc.add_page_break()
                doc.add_heading('表格', level=2)
                
                df = table.df
                word_table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                word_table.style = 'Table Grid'
                
                # 添加表头
                for i, col in enumerate(df.columns):
                    word_table.rows[0].cells[i].text = str(col)
                
                # 添加数据
                for i, row in df.iterrows():
                    for j, value in enumerate(row):
                        word_table.rows[i+1].cells[j].text = str(value)
        except Exception as e:
            pass  # 忽略表格提取错误
        
        doc.save(output_path)
        return output_path
    
    def to_excel(
        self, 
        pdf_path: str, 
        output_path: str,
        extract_tables: bool = True,
        extract_text: bool = False
    ) -> str:
        """
        将PDF转换为Excel
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出Excel文件路径
            extract_tables: 是否提取表格
            extract_text: 是否将文本也放入一个sheet
            
        Returns:
            输出文件路径
        """
        wb = Workbook()
        
        # 删除默认sheet
        wb.remove(wb.active)
        
        if extract_tables:
            try:
                tables = TableExtractor.extract_tables(pdf_path)
                
                for i, table in enumerate(tables):
                    df = table.df
                    sheet_name = f"Table_{i+1}"
                    
                    # 创建新sheet
                    ws = wb.create_sheet(title=sheet_name[:31])  # Excel限制31字符
                    
                    # 写入表头
                    for col_idx, col_name in enumerate(df.columns, 1):
                        cell = ws.cell(row=1, column=col_idx, value=str(col_name))
                        cell.font = Font(bold=True)
                        cell.alignment = Alignment(horizontal='center')
                    
                    # 写入数据
                    for row_idx, row in df.iterrows(), start=2:
                        for col_idx, value in enumerate(row, 1):
                            ws.cell(row=row_idx, column=col_idx, value=value)
                    
                    # 调整列宽
                    for col in ws.columns:
                        max_length = 0
                        column = col[0].column_letter
                        for cell in col:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws.column_dimensions[column].width = adjusted_width
                        
            except Exception as e:
                # 如果表格提取失败，创建一个错误说明sheet
                ws = wb.create_sheet(title="Info")
                ws.cell(row=1, column=1, value=f"表格提取失败: {str(e)}")
        
        if extract_text:
            ws = wb.create_sheet(title="Text")
            text = self.extractor.extract_text(pdf_path)
            
            # 将文本分行写入
            lines = text.split('\n')
            for i, line in enumerate(lines, 1):
                ws.cell(row=i, column=1, value=line)
        
        # 如果没有创建任何sheet，创建一个默认的
        if not wb.sheetnames:
            wb.create_sheet(title="Empty")
        
        wb.save(output_path)
        return output_path
    
    def to_images(
        self, 
        pdf_path: str, 
        output_dir: str,
        fmt: str = 'png',
        dpi: int = 200,
        pages: Optional[List[int]] = None
    ) -> List[str]:
        """
        将PDF转换为图片
        
        Args:
            pdf_path: PDF文件路径
            output_dir: 输出目录
            fmt: 图片格式 (png, jpg, jpeg, tiff, bmp)
            dpi: 分辨率
            pages: 指定页面列表，None表示所有页面
            
        Returns:
            生成的图片路径列表
        """
        os.makedirs(output_dir, exist_ok=True)
        
        # 转换PDF为图片
        if pages:
            images = []
            for page_num in pages:
                page_images = convert_from_path(
                    pdf_path,
                    dpi=dpi,
                    first_page=page_num + 1,
                    last_page=page_num + 1
                )
                images.extend(page_images)
        else:
            images = convert_from_path(pdf_path, dpi=dpi)
        
        # 保存图片
        saved_paths = []
        for i, image in enumerate(images):
            filename = f"page_{i+1}.{fmt}"
            filepath = os.path.join(output_dir, filename)
            
            # 转换格式
            if fmt.lower() in ['jpg', 'jpeg']:
                image = image.convert('RGB')
            
            image.save(filepath, fmt.upper() if fmt != 'jpg' else 'JPEG')
            saved_paths.append(filepath)
        
        return saved_paths
    
    def to_text(self, pdf_path: str, output_path: str, encoding: str = 'utf-8') -> str:
        """
        将PDF转换为纯文本文件
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出文本文件路径
            encoding: 文件编码
            
        Returns:
            输出文件路径
        """
        text = self.extractor.extract_text(pdf_path, preserve_layout=True)
        
        with open(output_path, 'w', encoding=encoding) as f:
            f.write(text)
        
        return output_path
    
    def to_html(self, pdf_path: str, output_path: str) -> str:
        """
        将PDF转换为HTML
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出HTML文件路径
            
        Returns:
            输出文件路径
        """
        text = self.extractor.extract_text(pdf_path, preserve_layout=False)
        
        # 简单HTML包装
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>PDF Export</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        pre {{
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
    </style>
</head>
<body>
    <pre>{text}</pre>
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def to_markdown(self, pdf_path: str, output_path: str) -> str:
        """
        将PDF转换为Markdown格式
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出Markdown文件路径
            
        Returns:
            输出文件路径
        """
        text = self.extractor.extract_text(pdf_path, preserve_layout=True)
        
        # 简单的Markdown转换
        lines = text.split('\n')
        md_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # 检测标题
            if stripped.isupper() and len(stripped) < 100 and stripped:
                md_lines.append(f"# {stripped}")
            elif stripped.endswith(':') and len(stripped) < 50:
                md_lines.append(f"## {stripped}")
            else:
                md_lines.append(line)
        
        md_content = '\n'.join(md_lines)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return output_path
    
    def extract_all(
        self, 
        pdf_path: str, 
        output_dir: str,
        formats: List[str] = ['text', 'images']
    ) -> Dict[str, Any]:
        """
        批量提取PDF到多种格式
        
        Args:
            pdf_path: PDF文件路径
            output_dir: 输出目录
            formats: 要提取的格式列表
            
        Returns:
            生成的文件路径字典
        """
        os.makedirs(output_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        
        results = {}
        
        if 'text' in formats:
            results['text'] = self.to_text(
                pdf_path, 
                os.path.join(output_dir, f"{base_name}.txt")
            )
        
        if 'word' in formats or 'docx' in formats:
            results['word'] = self.to_word(
                pdf_path,
                os.path.join(output_dir, f"{base_name}.docx")
            )
        
        if 'excel' in formats or 'xlsx' in formats:
            results['excel'] = self.to_excel(
                pdf_path,
                os.path.join(output_dir, f"{base_name}.xlsx")
            )
        
        if 'html' in formats:
            results['html'] = self.to_html(
                pdf_path,
                os.path.join(output_dir, f"{base_name}.html")
            )
        
        if 'markdown' in formats or 'md' in formats:
            results['markdown'] = self.to_markdown(
                pdf_path,
                os.path.join(output_dir, f"{base_name}.md")
            )
        
        if 'images' in formats:
            img_dir = os.path.join(output_dir, f"{base_name}_images")
            results['images'] = self.to_images(pdf_path, img_dir)
        
        return results

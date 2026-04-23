"""
文档生成器
生成方案文档（Markdown 和 Word 格式）
"""

import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx 未安装，Word 生成功能不可用")
    print("安装：pip install python-docx")


# 文档模板
SCHEME_TEMPLATE = """# {title}

## 一、项目背景

{background}

## 二、需求分析

{requirements}

## 三、技术方案

{solution}

## 四、实施计划

{plan}

## 五、预期效果

{expected_results}

## 六、风险评估

{risks}

## 七、预算估算

{budget}

---

**文档版本：** v{version}  
**生成时间：** {generated_at}  
**负责人：** {author}
"""


def generate_scheme_content(
    title: str,
    background: str = "",
    requirements: str = "",
    solution: str = "",
    plan: str = "",
    expected_results: str = "",
    risks: str = "",
    budget: str = "",
    version: str = "1.0",
    author: str = "老高"
) -> dict:
    """
    生成方案文档内容（结构化数据）
    
    :return: 文档各部分内容
    """
    return {
        "title": title,
        "background": background or "待补充",
        "requirements": requirements or "待补充",
        "solution": solution or "待补充",
        "plan": plan or "待补充",
        "expected_results": expected_results or "待补充",
        "risks": risks or "待补充",
        "budget": budget or "待补充",
        "version": version,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "author": author
    }


def generate_markdown(content: dict) -> str:
    """
    生成 Markdown 格式文档
    
    :param content: 文档内容字典
    :return: Markdown 文本
    """
    doc = f"""# {content['title']}

## 一、项目背景

{content['background']}

## 二、需求分析

{content['requirements']}

## 三、技术方案

{content['solution']}

## 四、实施计划

{content['plan']}

## 五、预期效果

{content['expected_results']}

## 六、风险评估

{content['risks']}

## 七、预算估算

{content['budget']}

---

**文档版本：** v{content['version']}  
**生成时间：** {content['generated_at']}  
**负责人：** {content['author']}
"""
    return doc


def generate_word(content: dict) -> bytes:
    """
    生成 Word 格式文档
    
    :param content: 文档内容字典
    :return: Word 文件二进制数据
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装")
    
    doc = Document()
    
    # 标题
    heading = doc.add_heading(content['title'], 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 各章节
    sections = [
        ("一、项目背景", content['background']),
        ("二、需求分析", content['requirements']),
        ("三、技术方案", content['solution']),
        ("四、实施计划", content['plan']),
        ("五、预期效果", content['expected_results']),
        ("六、风险评估", content['risks']),
        ("七、预算估算", content['budget']),
    ]
    
    for title, text in sections:
        doc.add_heading(title, level=1)
        # 处理多行文本
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
    
    # 页脚信息
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"文档版本：v{content['version']}\n")
    footer.add_run(f"生成时间：{content['generated_at']}\n")
    footer.add_run(f"负责人：{content['author']}")
    
    # 保存到内存
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def save_doc(content: dict, filename: str, category: str = "temp") -> dict:
    """
    保存文档（Markdown + Word 双格式）
    
    :param content: 文档内容字典
    :param filename: 文件名（不含扩展名）
    :param category: 类别 (temp/projects/meetings)
    :return: 文件路径信息
    """
    base_path = Path(f"D:\\OpenClawDocs\\{category}")
    base_path.mkdir(parents=True, exist_ok=True)
    
    paths = {}
    
    # 保存 Markdown
    md_content = generate_markdown(content)
    md_path = base_path / f"{filename}.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    paths['md'] = str(md_path)
    
    # 保存 Word
    if DOCX_AVAILABLE:
        try:
            word_data = generate_word(content)
            word_path = base_path / f"{filename}.docx"
            with open(word_path, 'wb') as f:
                f.write(word_data)
            paths['docx'] = str(word_path)
        except Exception as e:
            print(f"[WARN] Word 生成失败：{e}")
            paths['docx'] = None
    else:
        paths['docx'] = None
    
    return paths


def generate_summary(content: dict, max_length: int = 200) -> str:
    """
    生成文档摘要（用于飞书消息预览）
    
    :param content: 文档内容字典
    :param max_length: 最大字数
    :return: 摘要文本
    """
    # 提取关键信息
    summary = f"【{content['title']}】\n\n"
    summary += f"项目背景：{content['background'][:100]}...\n\n"
    summary += f"核心方案：{content['solution'][:100]}...\n\n"
    summary += f"预期效果：{content['expected_results'][:100]}...\n\n"
    summary += f"预算估算：{content['budget']}\n\n"
    summary += f"版本：v{content['version']} | 生成时间：{content['generated_at']}"
    
    return summary


def generate_toc(content: dict) -> str:
    """
    生成文档目录
    
    :param content: 文档内容字典
    :return: 目录文本
    """
    toc = "【文档结构】\n"
    sections = [
        "一、项目背景",
        "二、需求分析",
        "三、技术方案",
        "四、实施计划",
        "五、预期效果",
        "六、风险评估",
        "七、预算估算"
    ]
    for section in sections:
        toc += f"  {section}\n"
    return toc


def generate_filename(title: str, version: str = "v1") -> str:
    """
    生成规范的文件名
    
    :param title: 文档标题
    :param version: 版本号
    :return: 文件名（不含扩展名）
    """
    # 清理标题中的特殊字符
    safe_title = "".join(c for c in title if c.isalnum() or c in "_-")
    date_str = datetime.now().strftime("%Y%m%d")
    return f"{date_str}_{safe_title}_{version}"


# 测试
if __name__ == "__main__":
    # 生成测试文档
    content = generate_scheme_content(
        title="人员密集度检测方案",
        background="随着城市化进程加快，人员密集场所的安全管理日益重要。本项目旨在通过 AI 视觉技术，实现对人员密集度的实时监测和预警。",
        requirements="1. 实时监测人员密度\n2. 超限自动告警\n3. 数据可视化展示\n4. 历史数据分析",
        solution="采用 AI 视觉分析技术，基于深度学习的人数统计算法，结合现有摄像头资源，实现无感监测。",
        plan="第一阶段：需求调研（1 周）\n第二阶段：系统开发（4 周）\n第三阶段：部署测试（2 周）\n第四阶段：验收交付（1 周）",
        expected_results="1. 人员密度监测准确率>95%\n2. 告警响应时间<3 秒\n3. 支持多路视频并发分析",
        risks="1. 光线变化可能影响识别准确率\n2. 隐私保护问题需要妥善处理\n3. 网络稳定性影响数据传输",
        budget="硬件成本：1.5W\n软件授权：2W\n实施费用：0.5W\n总计：4W",
        version="1.0"
    )
    
    # 保存文档
    filename = generate_filename("人员密集度检测方案", "v1")
    paths = save_doc(content, filename, "temp")
    
    print(f"[OK] 文档已生成")
    print(f"  Markdown: {paths['md']}")
    print(f"  Word: {paths['docx']}")
    
    # 生成摘要
    summary = generate_summary(content)
    print(f"\n【飞书消息预览】\n{summary}")

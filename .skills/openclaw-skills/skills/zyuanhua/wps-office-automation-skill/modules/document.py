"""
文档处理模块
支持Word文档的生成、润色和审查
"""

import io
import logging
from typing import Dict, List, Optional, Any
from enum import Enum

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """文档类型"""
    NOTICE = "notice"
    REPORT = "report"
    MINUTES = "minutes"
    CONTRACT = "contract"
    LETTER = "letter"


class PolishStyle(str, Enum):
    """润色风格"""
    FORMAL = "formal"
    CONCISE = "concise"
    BUSINESS = "business"


class RiskLevel(str, Enum):
    """风险等级"""
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"


class OfficialDocumentRequest(BaseModel):
    """公文生成请求"""
    doc_type: DocumentType = Field(..., description="文档类型")
    title: str = Field(..., description="文档标题")
    subject: str = Field(..., description="主题内容")
    keywords: List[str] = Field(default_factory=list, description="关键词")
    recipient: Optional[str] = Field(None, description="接收方")
    sender: Optional[str] = Field(None, description="发送方")
    date: Optional[str] = Field(None, description="日期")


class PolishRequest(BaseModel):
    """润色请求"""
    content: str = Field(..., description="待润色内容")
    style: PolishStyle = Field(..., description="润色风格")
    preserve_structure: bool = Field(True, description="是否保留原文结构")


class ContractReviewResult(BaseModel):
    """合同审查结果"""
    risk_items: List[Dict[str, Any]] = Field(default_factory=list, description="风险条款")
    suggestions: List[str] = Field(default_factory=list, description="修改建议")
    summary: str = Field(..., description="审查摘要")
    risk_score: float = Field(..., description="风险评分（0-100）")


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        pass
        
    async def generate_official_document(
        self,
        request: OfficialDocumentRequest,
    ) -> bytes:
        """
        生成公文
        
        Args:
            request: 公文生成请求
            
        Returns:
            生成的Word文档字节流
        """
        logger.info(f"开始生成公文: {request.doc_type.value} - {request.title}")
        
        doc = Document()
        
        self._add_document_header(doc, request)
        
        content = self._generate_template_content(request)
        self._add_document_body(doc, content)
        
        self._add_document_footer(doc, request)
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"公文生成完成: {request.title}")
        return doc_bytes.getvalue()
        
    def _add_document_header(self, doc: Document, request: OfficialDocumentRequest):
        """添加文档头部"""
        if request.doc_type == DocumentType.NOTICE:
            title = doc.add_paragraph()
            title_run = title.add_run("通 知")
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif request.doc_type == DocumentType.REPORT:
            title = doc.add_paragraph()
            title_run = title.add_run(request.title)
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif request.doc_type == DocumentType.MINUTES:
            title = doc.add_paragraph()
            title_run = title.add_run("会议纪要")
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        doc.add_paragraph()
        
    def _generate_template_content(self, request: OfficialDocumentRequest) -> str:
        """生成模板内容"""
        if request.doc_type == DocumentType.NOTICE:
            content = f"关于{request.title}的通知\n\n"
            if request.recipient:
                content += f"{request.recipient}：\n\n"
            content += f"{request.subject}\n\n"
            content += "特此通知。\n\n"
            if request.sender:
                content += f"{request.sender}\n"
            if request.date:
                content += f"{request.date}\n"
            return content
            
        elif request.doc_type == DocumentType.MINUTES:
            content = f"会议主题：{request.title}\n\n"
            content += f"会议内容：\n{request.subject}\n\n"
            content += "会议决议：\n（待补充）\n"
            return content
            
        return request.subject
        
    def _add_document_body(self, doc: Document, content: str):
        """添加文档主体"""
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.first_line_indent = Inches(0.3)
                
    def _add_document_footer(self, doc: Document, request: OfficialDocumentRequest):
        """添加文档尾部"""
        doc.add_paragraph()
        
        if request.sender:
            sender_para = doc.add_paragraph()
            sender_para.add_run(request.sender)
            sender_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        if request.date:
            date_para = doc.add_paragraph()
            date_para.add_run(request.date)
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
    async def polish_document(
        self,
        request: PolishRequest,
    ) -> Dict[str, str]:
        """
        智能润色文档
        
        Args:
            request: 润色请求
            
        Returns:
            润色后的内容
        """
        logger.info(f"开始润色文档，风格: {request.style.value}")
        
        polished_content = self._basic_polish(request.content, request.style)
        
        logger.info("文档润色完成")
        return {
            "polished_content": polished_content,
            "original_content": request.content,
        }
        
    def _basic_polish(self, content: str, style: PolishStyle) -> str:
        """基础润色"""
        if style == PolishStyle.FORMAL:
            replacements = {
                "挺好的": "良好",
                "做的": "完成",
                "大家": "全体人员",
                "努力": "积极努力",
            }
        elif style == PolishStyle.CONCISE:
            replacements = {
                "这个": "",
                "那个": "",
                "进行": "",
            }
        else:
            replacements = {}
            
        result = content
        for old, new in replacements.items():
            result = result.replace(old, new)
            
        return result.strip()
        
    async def review_contract(
        self,
        contract_content: str,
    ) -> ContractReviewResult:
        """
        合同审查
        
        Args:
            contract_content: 合同内容
            
        Returns:
            审查结果
        """
        logger.info("开始合同审查")
        
        risk_items = self._basic_contract_check(contract_content)
        risk_score = len(risk_items) * 15
            
        summary = self._generate_review_summary(risk_items, risk_score)
        
        logger.info(f"合同审查完成，发现 {len(risk_items)} 个风险点")
        
        return ContractReviewResult(
            risk_items=risk_items,
            suggestions=[item["suggestion"] for item in risk_items],
            summary=summary,
            risk_score=min(risk_score, 100.0),
        )
        
    def _basic_contract_check(self, content: str) -> List[Dict[str, Any]]:
        """基础合同检查"""
        risk_items = []
        
        risk_keywords = [
            ("违约金", "高额违约金条款需注意合理性"),
            ("独家", "独家条款可能限制业务发展"),
            ("自动续约", "自动续约条款需明确通知期限"),
            ("不可撤销", "不可撤销条款需谨慎评估"),
        ]
        
        for keyword, suggestion in risk_keywords:
            if keyword in content:
                risk_items.append({
                    "keyword": keyword,
                    "level": RiskLevel.MEDIUM.value,
                    "suggestion": suggestion,
                })
                
        return risk_items
        
    def _generate_review_summary(
        self,
        risk_items: List[Dict[str, Any]],
        risk_score: float,
    ) -> str:
        """生成审查摘要"""
        if risk_score < 20:
            return "合同整体风险较低，条款较为合理。"
        elif risk_score < 50:
            return f"合同存在{len(risk_items)}个中等风险点，建议进一步协商修改。"
        else:
            return f"合同风险较高，发现{len(risk_items)}个风险条款，建议谨慎签署。"

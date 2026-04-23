#!/usr/bin/env python3
"""
Read DOCX file and extract text content
Usage: docx-read.py <file.docx> [--json]
"""
import sys
import argparse
import json
from pathlib import Path

def import_docx():
    """Import docx with helpful error message"""
    try:
        from docx import Document
        return Document
    except ImportError:
        print("Error: python-docx not installed.", file=sys.stderr)
        print("Install: pip install python-docx", file=sys.stderr)
        sys.exit(1)

def read_docx(file_path, include_tables=True):
    """Read Word document and extract all text"""
    Document = import_docx()
    
    doc = Document(file_path)
    content = {
        'file': str(file_path),
        'paragraphs': [],
        'tables': [],
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables)
    }
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content['paragraphs'].append(para.text.strip())
    
    # Extract tables
    if include_tables:
        for table_num, table in enumerate(doc.tables, 1):
            table_data = {
                'index': table_num,
                'rows': []
            }
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data['rows'].append(row_data)
            content['tables'].append(table_data)
    
    return content

def main():
    parser = argparse.ArgumentParser(
        description='Read DOCX file and extract text content',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  docx-read.py document.docx
  docx-read.py document.docx --json
  docx-read.py document.docx --no-tables
        """
    )
    parser.add_argument('file', help='Input DOCX file')
    parser.add_argument('--json', action='store_true', help='Output as JSON')
    parser.add_argument('--no-tables', action='store_true', help='Skip table extraction')
    
    args = parser.parse_args()
    
    # Validate file
    file_path = Path(args.file)
    if not file_path.exists():
        print(f"Error: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    if not file_path.suffix.lower() == '.docx':
        print(f"Error: Not a DOCX file: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    try:
        content = read_docx(file_path, include_tables=not args.no_tables)
        
        if args.json:
            print(json.dumps(content, indent=2, ensure_ascii=False))
        else:
            # Plain text output
            print(f"File: {content['file']}")
            print(f"Paragraphs: {content['total_paragraphs']}")
            print(f"Tables: {content['total_tables']}")
            print("\n" + "="*60 + "\n")
            
            for para in content['paragraphs']:
                print(para)
            
            for table in content['tables']:
                print(f"\n[Table {table['index']}]")
                for row in table['rows']:
                    print(" | ".join(row))
                    
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()

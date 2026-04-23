#!/usr/bin/env python3
"""
Create DOCX file
Usage: docx-write.py <output.docx> [options]
"""
import sys
import argparse
from pathlib import Path
from datetime import datetime

def import_docx():
    """Import docx with helpful error message"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        return Document, Pt, Inches, WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx not installed.", file=sys.stderr)
        print("Install: pip install python-docx", file=sys.stderr)
        sys.exit(1)

def create_document(output_path, title=None, content=None, author=None):
    """Create a Word document"""
    Document, Pt, Inches, WD_ALIGN_PARAGRAPH = import_docx()
    
    doc = Document()
    
    # Add title
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    if author:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(f"Author: {author}\nDate: {datetime.now().strftime('%Y-%m-%d')}")
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True
        doc.add_paragraph()  # Spacing
    
    # Add content
    if content:
        doc.add_paragraph(content)
    
    # Save
    doc.save(output_path)
    return output_path

def main():
    parser = argparse.ArgumentParser(
        description='Create DOCX file',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  docx-write.py output.docx --title "Report" --content "Hello World"
  docx-write.py output.docx --title "Document" --author "John Doe"
        """
    )
    parser.add_argument('output', help='Output DOCX file')
    parser.add_argument('--title', help='Document title')
    parser.add_argument('--content', help='Document content')
    parser.add_argument('--author', help='Document author')
    
    args = parser.parse_args()
    
    # Validate output path
    output_path = Path(args.output)
    if output_path.suffix.lower() != '.docx':
        print(f"Error: Output file must be .docx: {output_path}", file=sys.stderr)
        sys.exit(1)
    
    # Create parent directory
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        create_document(
            output_path,
            title=args.title,
            content=args.content,
            author=args.author
        )
        print(f"Created: {output_path}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()

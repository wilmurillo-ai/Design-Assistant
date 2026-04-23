"""
多格式文档内容提取器
支持 PDF、Word (.docx)、Markdown (.md) 格式
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import mimetypes
import json

# 尝试导入第三方库
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


class FormatExtractor:
    """支持PDF、DOCX、MD的多格式内容提取器"""
    
    # 支持的格式映射
    FORMAT_MAPPING = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',  # 需要转换
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.txt': 'text'
    }
    
    # 常见报告模板关键词
    TEMPLATE_KEYWORDS = {
        '销售': ['销售', '销售额', '营收', '业绩', '客户', '订单'],
        '财务': ['财务', '利润', '成本', '收入', '支出', '预算'],
        '项目': ['项目', '进度', '里程碑', '任务', '资源', '风险'],
        '运营': ['运营', '用户', '活跃', '转化', '留存', '渠道'],
        '产品': ['产品', '功能', '版本', '迭代', '反馈', '需求'],
        '市场': ['市场', '营销', '推广', '广告', '品牌', '份额'],
        '技术': ['技术', '开发', '代码', '测试', '部署', '性能'],
        '人力资源': ['人力', '招聘', '培训', '绩效', '薪酬', '员工']
    }
    
    def __init__(self):
        """初始化格式提取器"""
        self._init_mimetypes()
        
    def _init_mimetypes(self):
        """初始化MIME类型检测"""
        mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')
        mimetypes.add_type('text/markdown', '.md')
        mimetypes.add_type('text/markdown', '.markdown')
    
    def detect_format(self, file_path: str) -> str:
        """
        检测文件格式
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 格式标识 (pdf, docx, markdown, text, unknown)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 通过扩展名检测
        ext = Path(file_path).suffix.lower()
        if ext in self.FORMAT_MAPPING:
            return self.FORMAT_MAPPING[ext]
        
        # 通过MIME类型检测
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'word' in mime_type or 'officedocument' in mime_type:
                return 'docx'
            elif 'markdown' in mime_type:
                return 'markdown'
            elif 'text' in mime_type:
                return 'text'
        
        # 通过文件内容启发式检测
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                first_line = f.readline(1000)
                if '# ' in first_line or '## ' in first_line:
                    return 'markdown'
        except:
            pass
        
        return 'unknown'
    
    def extract_text(self, file_path: str) -> str:
        """
        提取文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 提取的文本内容
        """
        format_type = self.detect_format(file_path)
        
        if format_type == 'pdf':
            return self._extract_pdf_text(file_path)
        elif format_type == 'docx':
            return self._extract_docx_text(file_path)
        elif format_type in ['markdown', 'text']:
            return self._extract_text_file(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {format_type}")
    
    def extract_structure(self, file_path: str) -> Dict:
        """
        提取文档结构（标题、段落、表格等）
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 文档结构信息
        """
        format_type = self.detect_format(file_path)
        text = self.extract_text(file_path)
        
        structure = {
            'format': format_type,
            'file_name': Path(file_path).name,
            'title': self._extract_title(text),
            'headings': self._extract_headings(text),
            'paragraphs': self._extract_paragraphs(text),
            'variables': self._extract_template_variables(text),
            'sections': self._identify_sections(text),
            'word_count': len(text.split()),
            'character_count': len(text)
        }
        
        return structure
    
    def extract_metadata(self, file_path: str) -> Dict:
        """
        提取文档元数据
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 文档元数据
        """
        import datetime
        
        path = Path(file_path)
        stat = path.stat()
        
        metadata = {
            'file_name': path.name,
            'file_size': stat.st_size,
            'created_time': datetime.datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_time': datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'file_format': self.detect_format(file_path)
        }
        
        # 添加特定格式的元数据
        format_type = self.detect_format(file_path)
        if format_type == 'pdf':
            metadata.update(self._extract_pdf_metadata(file_path))
        elif format_type == 'docx':
            metadata.update(self._extract_docx_metadata(file_path))
        
        return metadata
    
    def auto_categorize(self, text: str) -> List[str]:
        """
        自动分类文档到适用场景
        
        Args:
            text: 文档文本
            
        Returns:
            List[str]: 分类标签列表
        """
        categories = []
        text_lower = text.lower()
        
        for category, keywords in self.TEMPLATE_KEYWORDS.items():
            keyword_count = sum(1 for keyword in keywords if keyword.lower() in text_lower)
            if keyword_count >= 2:  # 至少有2个关键词匹配
                categories.append(category)
        
        # 如果未找到明确分类，尝试基于内容推测
        if not categories:
            if any(word in text_lower for word in ['报表', '报告', 'summary', 'report']):
                categories.append('通用报告')
            elif any(word in text_lower for word in ['计划', '规划', '方案', 'proposal']):
                categories.append('规划方案')
            elif any(word in text_lower for word in ['分析', '趋势', 'insight', 'analysis']):
                categories.append('分析报告')
        
        return categories[:3]  # 最多返回3个分类
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """提取PDF文本"""
        if not PDF_SUPPORT:
            raise ImportError("需要安装 PyPDF2 库: pip install PyPDF2")
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"PDF文本提取失败: {e}")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """提取Word文档文本"""
        if not DOCX_SUPPORT:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            raise RuntimeError(f"Word文档文本提取失败: {e}")
    
    def _extract_text_file(self, file_path: str) -> str:
        """提取文本文件内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise UnicodeDecodeError("无法解码文件")
        except Exception as e:
            raise RuntimeError(f"文本文件读取失败: {e}")
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict:
        """提取PDF元数据"""
        if not PDF_SUPPORT:
            return {}
        
        try:
            metadata = {}
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        if hasattr(value, 'decode'):
                            try:
                                metadata[key] = value.decode('utf-8')
                            except:
                                metadata[key] = str(value)
                        else:
                            metadata[key] = str(value)
            return metadata
        except:
            return {}
    
    def _extract_docx_metadata(self, file_path: str) -> Dict:
        """提取Word文档元数据"""
        if not DOCX_SUPPORT:
            return {}
        
        try:
            metadata = {}
            doc = docx.Document(file_path)
            
            # 核心属性
            core_properties = doc.core_properties
            if core_properties.author:
                metadata['author'] = core_properties.author
            if core_properties.title:
                metadata['title'] = core_properties.title
            if core_properties.subject:
                metadata['subject'] = core_properties.subject
            if core_properties.keywords:
                metadata['keywords'] = core_properties.keywords
            if core_properties.created:
                metadata['created'] = core_properties.created.isoformat()
            if core_properties.modified:
                metadata['modified'] = core_properties.modified.isoformat()
            
            return metadata
        except:
            return {}
    
    def _extract_title(self, text: str) -> str:
        """提取文档标题"""
        lines = text.split('\n')
        for line in lines[:10]:  # 检查前10行
            line = line.strip()
            if line and len(line) < 200:  # 合理标题长度
                if line.startswith('# '):
                    return line[2:].strip()
                elif not line.startswith('*') and not line.startswith('-') and line.count(' ') < 20:
                    return line[:100]
        
        return "未命名模板"
    
    def _extract_headings(self, text: str) -> List[Dict]:
        """提取标题结构"""
        headings = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            # Markdown标题
            if line.startswith('# '):
                level = 1
                content = line[2:].strip()
            elif line.startswith('## '):
                level = 2
                content = line[3:].strip()
            elif line.startswith('### '):
                level = 3
                content = line[4:].strip()
            elif line.startswith('#### '):
                level = 4
                content = line[5:].strip()
            else:
                continue
            
            headings.append({
                'level': level,
                'content': content,
                'line_number': i + 1
            })
        
        return headings
    
    def _extract_paragraphs(self, text: str) -> List[str]:
        """提取段落"""
        paragraphs = []
        current_para = []
        
        for line in text.split('\n'):
            line = line.strip()
            if line:
                current_para.append(line)
            elif current_para:
                paragraphs.append(' '.join(current_para))
                current_para = []
        
        if current_para:
            paragraphs.append(' '.join(current_para))
        
        return paragraphs[:50]  # 最多返回50个段落
    
    def _extract_template_variables(self, text: str) -> List[str]:
        """
        自动提取模板变量
        支持格式：{变量名}, {{变量名}}, [变量名], %变量名%
        """
        variables = set()
        
        # 大括号变量 {变量名}
        brace_vars = re.findall(r'\{([^{}\s]+)\}', text)
        variables.update(brace_vars)
        
        # 双大括号变量 {{变量名}}
        double_brace_vars = re.findall(r'\{\{([^{}]+)\}\}', text)
        variables.update(double_brace_vars)
        
        # 方括号变量 [变量名]
        bracket_vars = re.findall(r'\[([^\[\]\s]+)\]', text)
        variables.update(bracket_vars)
        
        # 百分号变量 %变量名%
        percent_vars = re.findall(r'%([^%\s]+)%', text)
        variables.update(percent_vars)
        
        # 过滤掉常见非变量词语
        common_words = {'页', '页眉', '页脚', 'table', '图', '图表', 'image', 'img'}
        filtered_vars = [v for v in variables if len(v) > 1 and v not in common_words]
        
        return sorted(filtered_vars)
    
    def _identify_sections(self, text: str) -> List[Dict]:
        """识别文档章节"""
        sections = []
        lines = text.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if line.startswith('# '):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'title': line[2:].strip(),
                    'level': 1,
                    'start_line': i + 1,
                    'content': []
                }
            elif line.startswith('## '):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'title': line[3:].strip(),
                    'level': 2,
                    'start_line': i + 1,
                    'content': []
                }
            elif current_section and line:
                if len(current_section['content']) < 100:  # 限制内容长度
                    current_section['content'].append(line)
        
        if current_section:
            sections.append(current_section)
        
        return sections


# 简单使用示例
if __name__ == "__main__":
    # 创建测试实例
    extractor = FormatExtractor()
    
    # 测试文件检测
    test_files = [
        "example.pdf",
        "template.docx", 
        "report.md",
        "data.txt"
    ]
    
    for file in test_files:
        if os.path.exists(file):
            try:
                print(f"\n=== 分析文件: {file} ===")
                format_type = extractor.detect_format(file)
                print(f"格式: {format_type}")
                
                metadata = extractor.extract_metadata(file)
                print(f"元数据: {json.dumps(metadata, ensure_ascii=False, indent=2)}")
                
                structure = extractor.extract_structure(file)
                print(f"标题: {structure['title']}")
                print(f"变量: {structure['variables'][:5]}")  # 显示前5个变量
                
            except Exception as e:
                print(f"错误: {e}")
        else:
            print(f"文件不存在: {file}")
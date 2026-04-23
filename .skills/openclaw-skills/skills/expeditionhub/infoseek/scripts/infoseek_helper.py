#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
InfoSeek v2.0 - Core Helper Script
Handles: folder management, URL normalization, deduplication, file storage,
database management, and logging.
Note: This script does NOT perform searches. It only handles file processing
and database management. Search is performed by the Agent via OpenClaw built-in tools.
"""

import os
import sys
import json
import re
import csv
import html
import hashlib
import logging
import argparse
import sqlite3
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from urllib.parse import urlparse, urlunparse, parse_qs, urlencode
from dataclasses import dataclass, field


# ============================================================
# Configuration
# ============================================================

@dataclass
class InfoSeekConfig:
    """InfoSeek configuration"""
    max_search_depth: int = 100
    max_filename_length: int = 200
    max_title_length: int = 80
    batch_save_size: int = 50
    request_delay: float = 2.0
    timeout: int = 600
    supported_formats: List[str] = field(default_factory=lambda: ['md', 'json', 'txt', 'csv', 'xlsx', 'html', 'docx'])
    default_format: str = 'md'
    ad_patterns: List[str] = field(default_factory=lambda: [
        r'广告', r'推广', r'赞助', r'相关阅读', r'热门推荐',
        r'advertisement', r'sponsored', r'promoted', r'ads',
        r'read more', r'recommended', r'trending', r'share this'
    ])


# ============================================================
# Logging
# ============================================================

def setup_logger(log_dir: Path) -> logging.Logger:
    """Configure logging"""
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / 'infoseek.log'

    logger = logging.getLogger('infoseek')
    logger.setLevel(logging.DEBUG)
    logger.handlers.clear()

    fh = logging.FileHandler(log_file, encoding='utf-8')
    fh.setLevel(logging.DEBUG)

    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)

    formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
    fh.setFormatter(formatter)
    ch.setFormatter(formatter)

    logger.addHandler(fh)
    logger.addHandler(ch)

    return logger


# ============================================================
# URL Normalizer
# ============================================================

class URLNormalizer:
    """URL normalization for deduplication"""

    TRACKING_PARAMS = {
        'utm_source', 'utm_medium', 'utm_campaign', 'utm_term', 'utm_content',
        'ref', 'fbclid', 'gclid', 'msclkid', 'twclid', 'yclid'
    }

    @staticmethod
    def normalize(url: str) -> str:
        """
        Normalize URL for deduplication.

        Rules:
        1. Remove www. prefix
        2. Unify http -> https
        3. Remove tracking parameters
        4. Remove trailing slash
        5. Lowercase domain
        6. Remove fragment identifiers
        """
        if not url:
            return ''

        try:
            parsed = urlparse(url)

            netloc = parsed.netloc.lower()
            if netloc.startswith('www.'):
                netloc = netloc[4:]

            scheme = 'https' if parsed.scheme in ('http', 'https') else parsed.scheme

            path = parsed.path
            if path.endswith('/') and len(path) > 1:
                path = path[:-1]

            params = parse_qs(parsed.query, keep_blank_values=True)
            filtered_params = {
                k: v for k, v in params.items()
                if k.lower() not in URLNormalizer.TRACKING_PARAMS
            }
            query = urlencode(filtered_params, doseq=True)

            return urlunparse((scheme, netloc, path, parsed.params, query, ''))

        except Exception:
            return url.strip()

    @staticmethod
    def extract_domain(url: str) -> str:
        """Extract domain from URL"""
        try:
            parsed = urlparse(url)
            domain = parsed.netloc.lower()
            if domain.startswith('www.'):
                domain = domain[4:]
            domain = domain.split(':')[0]
            return domain
        except Exception:
            return 'unknown'

    @staticmethod
    def extract_domain_name(url: str) -> str:
        """Extract website/media name for filenames"""
        domain = URLNormalizer.extract_domain(url)
        for suffix in ['.com', '.cn', '.net', '.org', '.io', '.co', '.com.cn']:
            if domain.endswith(suffix):
                domain = domain[:-len(suffix)]
                break
        return domain


# ============================================================
# Date Parser
# ============================================================

class DateParser:
    """Date parsing utilities"""

    @staticmethod
    def parse(date_str: str) -> str:
        """
        Parse date string to YYYYMMDD format.

        Supported formats:
        - 2026-04-07, 2026/04/07, 20260407
        - April 7, 2026
        - 2026 年 4 月 7 日
        - 04/07/2026

        Returns: YYYYMMDD string, or '00000000' if unknown
        """
        if not date_str:
            return '00000000'

        date_str = date_str.strip()

        match = re.match(r'^(\d{4})[-/]?(\d{1,2})[-/]?(\d{1,2})$', date_str)
        if match:
            year, month, day = match.groups()
            return f"{year}{int(month):02d}{int(day):02d}"

        match = re.match(r'^(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', date_str)
        if match:
            year, month, day = match.groups()
            return f"{year}{int(month):02d}{int(day):02d}"

        months = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4,
            'may': 5, 'june': 6, 'july': 7, 'august': 8,
            'september': 9, 'october': 10, 'november': 11, 'december': 12,
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }

        match = re.match(r'^([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})', date_str)
        if match:
            month_str, day, year = match.groups()
            month = months.get(month_str.lower())
            if month:
                return f"{year}{month:02d}{int(day):02d}"

        try:
            from dateutil import parser as date_parser
            dt = date_parser.parse(date_str, fuzzy=True)
            return dt.strftime('%Y%m%d')
        except Exception:
            pass

        return '00000000'

    @staticmethod
    def format_for_display(date_str: str) -> str:
        """Format YYYYMMDD to YYYY-MM-DD for display"""
        if date_str == '00000000' or len(date_str) != 8:
            return 'Unknown'
        return f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"


# ============================================================
# SQLite Database
# ============================================================

class URLDatabase:
    """URL deduplication database (SQLite)"""

    def __init__(self, db_path: Path):
        self.db_path = db_path
        self.conn = None
        self._init_db()

    def _init_db(self):
        """Initialize database"""
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self.conn = sqlite3.connect(str(self.db_path))
        self.conn.execute('PRAGMA journal_mode=WAL')

        self.conn.execute('''
            CREATE TABLE IF NOT EXISTS urls (
                url TEXT PRIMARY KEY,
                normalized_url TEXT UNIQUE,
                task TEXT NOT NULL,
                filename TEXT NOT NULL,
                added_at TIMESTAMP NOT NULL,
                checksum TEXT,
                deleted INTEGER DEFAULT 0,
                deleted_at TIMESTAMP
            )
        ''')

        self.conn.execute('CREATE INDEX IF NOT EXISTS idx_normalized ON urls(normalized_url)')
        self.conn.execute('CREATE INDEX IF NOT EXISTS idx_task ON urls(task)')
        self.conn.execute('CREATE INDEX IF NOT EXISTS idx_deleted ON urls(deleted)')

        self.conn.commit()

    def is_duplicate(self, url: str) -> bool:
        """Check if URL already exists (using normalized URL)"""
        normalized = URLNormalizer.normalize(url)
        cursor = self.conn.execute(
            'SELECT 1 FROM urls WHERE normalized_url = ? AND deleted = 0',
            (normalized,)
        )
        return cursor.fetchone() is not None

    def add(self, url: str, task: str, filename: str) -> bool:
        """Add URL to database. Returns True if added, False if duplicate."""
        if self.is_duplicate(url):
            return False

        normalized = URLNormalizer.normalize(url)
        checksum = hashlib.md5(normalized.encode()).hexdigest()[:8]

        try:
            self.conn.execute(
                '''INSERT INTO urls (url, normalized_url, task, filename, added_at, checksum)
                   VALUES (?, ?, ?, ?, ?, ?)''',
                (url, normalized, task, filename, datetime.now().isoformat(), checksum)
            )
            self.conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False

    def get_stats(self) -> Dict:
        """Get database statistics"""
        cursor = self.conn.execute('SELECT COUNT(*) FROM urls WHERE deleted = 0')
        total = cursor.fetchone()[0]

        cursor = self.conn.execute('''
            SELECT task, COUNT(*) as count
            FROM urls WHERE deleted = 0
            GROUP BY task ORDER BY count DESC
        ''')
        by_task = dict(cursor.fetchall())

        return {'total_urls': total, 'by_task': by_task}

    def list_entries(self, task: str = None, limit: int = 50) -> List[Dict]:
        """List entries"""
        if task:
            cursor = self.conn.execute(
                '''SELECT url, task, filename, added_at
                   FROM urls WHERE task = ? AND deleted = 0
                   ORDER BY added_at DESC LIMIT ?''',
                (task, limit)
            )
        else:
            cursor = self.conn.execute(
                '''SELECT url, task, filename, added_at
                   FROM urls WHERE deleted = 0
                   ORDER BY added_at DESC LIMIT ?''',
                (limit,)
            )

        return [
            {'url': r[0], 'task': r[1], 'filename': r[2], 'added_at': r[3]}
            for r in cursor.fetchall()
        ]

    def mark_deleted(self, url: str) -> bool:
        """Soft-delete URL"""
        normalized = URLNormalizer.normalize(url)
        try:
            self.conn.execute(
                'UPDATE urls SET deleted = 1, deleted_at = ? WHERE normalized_url = ?',
                (datetime.now().isoformat(), normalized)
            )
            self.conn.commit()
            return True
        except Exception:
            return False

    def backup(self, backup_dir: Path) -> Path:
        """Backup database"""
        backup_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_path = backup_dir / f'infoseek_backup_{timestamp}.db'
        shutil.copy2(self.db_path, backup_path)
        return backup_path

    def close(self):
        """Close database connection"""
        if self.conn:
            self.conn.close()
            self.conn = None

    def __del__(self):
        self.close()


# ============================================================
# Content Saver
# ============================================================

class ContentSaver:
    """Save content in multiple formats"""

    def __init__(self, config: InfoSeekConfig):
        self.config = config

    @staticmethod
    def generate_filename(date: str, title: str, website: str, ext: str) -> str:
        """
        Generate standardized filename.
        Format: YYYYMMDD-title-website.ext
        """
        safe_title = re.sub(r'[<>:"/\\|?*]', '', title)
        safe_title = safe_title.replace('\n', ' ').strip()

        if len(safe_title) > 80:
            safe_title = safe_title[:80]

        safe_website = re.sub(r'[<>:"/\\|?*]', '', website)
        if len(safe_website) > 30:
            safe_website = safe_website[:30]

        filename = f"{date}-{safe_title}-{safe_website}.{ext}"

        if len(filename) > 200:
            filename = filename[:200]

        return filename

    @staticmethod
    def generate_unique_filename(folder: Path, filename: str, url: str) -> str:
        """Generate unique filename (avoid collisions)"""
        filepath = folder / filename

        if not filepath.exists():
            return filename

        name_parts = filename.rsplit('.', 1)
        base, ext = name_parts if len(name_parts) == 2 else (filename, '')

        hash_suffix = hashlib.md5(url.encode()).hexdigest()[:8]
        return f"{base}-{hash_suffix}.{ext}"

    def save(self, folder: Path, filename: str, content: Dict, fmt: str) -> str:
        """Save content to file"""
        filepath = folder / filename

        handlers = {
            'md': self._save_markdown,
            'json': self._save_json,
            'txt': self._save_text,
            'csv': self._save_csv_single,
            'xlsx': self._save_xlsx,
            'html': self._save_html,
            'docx': self._save_docx,
        }

        handler = handlers.get(fmt, self._save_markdown)
        handler(filepath, content)

        return str(filepath)

    def _save_markdown(self, filepath: Path, content: Dict):
        """Save as Markdown"""
        md = f"""# {content.get('title', 'Untitled')}

---

## Metadata

| Field | Value |
|-------|-------|
| **URL** | {content.get('url', 'N/A')} |
| **Website** | {content.get('website', 'N/A')} |
| **Source** | {content.get('source', 'N/A')} |
| **Publish Date** | {content.get('publishDate', 'N/A')} |
| **Title** | {content.get('title', 'N/A')} |
| **Author** | {content.get('author', 'Unknown')} |
| **Editor** | {content.get('editor', 'Unknown')} |
| **Archived At** | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
| **Search Task** | {content.get('searchTask', 'N/A')} |

---

## Content

{content.get('content', 'No content')}

---

*Generated by InfoSeek Skill | Original URL: {content.get('url', 'N/A')}*
"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(md)

    def _save_json(self, filepath: Path, content: Dict):
        """Save as JSON"""
        data = {
            "metadata": {
                "url": content.get('url', ''),
                "website": content.get('website', ''),
                "source": content.get('source', ''),
                "publishDate": content.get('publishDate', ''),
                "title": content.get('title', ''),
                "author": content.get('author', 'Unknown'),
                "editor": content.get('editor', 'Unknown'),
                "archivedAt": datetime.now().isoformat(),
                "searchTask": content.get('searchTask', '')
            },
            "content": content.get('content', '')
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _save_text(self, filepath: Path, content: Dict):
        """Save as plain text"""
        text = f"""{content.get('title', 'Untitled')}

Source: {content.get('website', 'N/A')}
URL: {content.get('url', 'N/A')}
Date: {content.get('publishDate', 'N/A')}
Author: {content.get('author', 'Unknown')}
Editor: {content.get('editor', 'Unknown')}

{'=' * 60}

{content.get('content', 'No content')}
"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(text)

    def _save_csv_single(self, filepath: Path, content: Dict):
        """Save as CSV (single article)"""
        with open(filepath, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(['Field', 'Value'])
            writer.writerow(['URL', content.get('url', '')])
            writer.writerow(['Website', content.get('website', '')])
            writer.writerow(['Source', content.get('source', '')])
            writer.writerow(['Date', content.get('publishDate', '')])
            writer.writerow(['Title', content.get('title', '')])
            writer.writerow(['Author', content.get('author', '')])
            writer.writerow(['Editor', content.get('editor', '')])
            writer.writerow(['Content', content.get('content', '')])

    def _save_xlsx(self, filepath: Path, content: Dict):
        """Save as XLSX (requires openpyxl)"""
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = content.get('title', 'InfoSeek')[:31]

            ws['A1'], ws['B1'] = 'Field', 'Value'
            ws['A2'], ws['B2'] = 'URL', content.get('url', '')
            ws['A3'], ws['B3'] = 'Website', content.get('website', '')
            ws['A4'], ws['B4'] = 'Source', content.get('source', '')
            ws['A5'], ws['B5'] = 'Date', content.get('publishDate', '')
            ws['A6'], ws['B6'] = 'Title', content.get('title', '')
            ws['A7'], ws['B7'] = 'Author', content.get('author', '')
            ws['A8'], ws['B8'] = 'Editor', content.get('editor', '')
            ws['A9'], ws['B9'] = 'Archived At', datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            ws['A10'], ws['B10'] = 'Search Task', content.get('searchTask', '')

            ws['A12'] = 'Content'
            ws['A13'] = content.get('content', '')

            ws.column_dimensions['A'].width = 15
            ws.column_dimensions['B'].width = 80

            wb.save(str(filepath))
        except ImportError:
            print("[WARN] openpyxl not installed, saving as CSV")
            self._save_csv_single(filepath.with_suffix('.csv'), content)

    def _save_html(self, filepath: Path, content: Dict):
        """Save as HTML"""
        safe_content = html.escape(content.get('content', ''))
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(content.get('title', 'Untitled'))}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        .metadata {{ background: #f5f5f5; padding: 20px; border-radius: 8px; margin-bottom: 30px; }}
        .metadata table {{ width: 100%; border-collapse: collapse; }}
        .metadata td {{ padding: 10px; border-bottom: 1px solid #ddd; }}
        .metadata td:first-child {{ font-weight: bold; width: 120px; }}
        .content {{ line-height: 1.8; }}
        .content p {{ margin: 1em 0; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 12px; color: #666; }}
    </style>
</head>
<body>
    <h1>{html.escape(content.get('title', 'Untitled'))}</h1>

    <div class="metadata">
        <h3>Metadata</h3>
        <table>
            <tr><td>URL</td><td><a href="{html.escape(content.get('url', '#'))}">{html.escape(content.get('url', 'N/A'))}</a></td></tr>
            <tr><td>Website</td><td>{html.escape(content.get('website', 'N/A'))}</td></tr>
            <tr><td>Source</td><td>{html.escape(content.get('source', 'N/A'))}</td></tr>
            <tr><td>Publish Date</td><td>{html.escape(content.get('publishDate', 'N/A'))}</td></tr>
            <tr><td>Author</td><td>{html.escape(content.get('author', 'Unknown'))}</td></tr>
            <tr><td>Editor</td><td>{html.escape(content.get('editor', 'Unknown'))}</td></tr>
            <tr><td>Archived At</td><td>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</td></tr>
        </table>
    </div>

    <div class="content">
        {safe_content.replace(chr(10), '<br>')}
    </div>

    <div class="footer">
        <p>Generated by InfoSeek Skill | Original URL: <a href="{html.escape(content.get('url', '#'))}">{html.escape(content.get('url', '#'))}</a></p>
    </div>
</body>
</html>
"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def _save_docx(self, filepath: Path, content: Dict):
        """Save as DOCX (requires python-docx)"""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            doc.add_heading(content.get('title', 'Untitled'), level=1)

            doc.add_heading('Metadata', level=2)
            metadata = [
                ('URL', content.get('url', 'N/A')),
                ('Website', content.get('website', 'N/A')),
                ('Source', content.get('source', 'N/A')),
                ('Publish Date', content.get('publishDate', 'N/A')),
                ('Author', content.get('author', 'Unknown')),
                ('Editor', content.get('editor', 'Unknown')),
                ('Archived At', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Search Task', content.get('searchTask', 'N/A'))
            ]

            for label, value in metadata:
                p = doc.add_paragraph()
                run_label = p.add_run(f'{label}: ')
                run_label.bold = True
                p.add_run(str(value))

            doc.add_heading('Content', level=2)
            doc.add_paragraph(content.get('content', 'No content'))

            footer = doc.sections[0].footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.text = f'Generated by InfoSeek | {content.get("url", "")}'
            p.style.font.size = Pt(8)

            doc.save(str(filepath))

        except ImportError:
            print("[WARN] python-docx not installed, saving as Markdown")
            self._save_markdown(filepath.with_suffix('.md'), content)


# ============================================================
# Content Filter
# ============================================================

class ContentFilter:
    """Filter ads and irrelevant content"""

    def __init__(self, config: InfoSeekConfig):
        self.config = config

    def clean_text(self, text: str) -> str:
        """Clean plain text content"""
        if not text:
            return ''

        for pattern in self.config.ad_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)

        return text.strip()

    def clean_html(self, html_content: str) -> str:
        """
        Clean HTML content (extract body text).
        Requires BeautifulSoup. Falls back to plain text if unavailable.
        """
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_content, 'html.parser')

            for selector in [
                '.ad', '.advertisement', '.ads', '.sponsored',
                '.sidebar', '.aside', '.widget',
                '.nav', '.header', '.menu', '.navigation',
                '.comment', '.comments', '.discussion',
                '.related', '.recommend', '.hot', '.trending',
                '.footer', '.share', '.social',
                'script', 'style', 'noscript', 'iframe'
            ]:
                for element in soup.select(selector):
                    element.decompose()

            for tag in ['article', 'main', '.content', '.post', '.entry', '.article-body']:
                element = soup.select_one(tag)
                if element:
                    return str(element)

            body = soup.find('body')
            if body:
                return str(body)

            return str(soup)

        except ImportError:
            return self.clean_text(self._html_to_text(html_content))

    @staticmethod
    def _html_to_text(html_content: str) -> str:
        """Convert HTML to plain text"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text(separator='\n', strip=True)
        except Exception:
            text = re.sub(r'<[^>]+>', ' ', html_content)
            text = re.sub(r'\s+', ' ', text).strip()
            return text


# ============================================================
# File Deleter (Recycle Bin)
# ============================================================

class FileDeleter:
    """Move files to system recycle bin"""

    @staticmethod
    def move_to_recycle_bin(filepath: str) -> bool:
        """Move file to system recycle bin (Windows/macOS/Linux)"""
        try:
            if sys.platform == 'win32':
                return FileDeleter._windows_recycle(filepath)
            elif sys.platform == 'darwin':
                return FileDeleter._macos_recycle(filepath)
            else:
                return FileDeleter._linux_recycle(filepath)
        except Exception as e:
            print(f"[ERROR] Failed to delete file: {e}")
            return False

    @staticmethod
    def _windows_recycle(filepath: str) -> bool:
        """Windows: move to recycle bin"""
        import subprocess
        safe_path = filepath.replace('"', '""')
        ps_command = f'''
        Add-Type -AssemblyName Microsoft.VisualBasic
        [Microsoft.VisualBasic.FileIO.FileSystem]::DeleteFile("{safe_path}", "OnlyErrorDialogs", "SendToRecycleBin")
        '''
        result = subprocess.run(
            ['powershell', '-Command', ps_command],
            capture_output=True, text=True, timeout=30
        )
        return result.returncode == 0

    @staticmethod
    def _macos_recycle(filepath: str) -> bool:
        """macOS: move to Trash"""
        home = Path.home()
        trash = home / '.Trash'
        trash.mkdir(exist_ok=True)
        shutil.move(filepath, trash / Path(filepath).name)
        return True

    @staticmethod
    def _linux_recycle(filepath: str) -> bool:
        """Linux: move to trash (trash-cli or manual)"""
        import subprocess
        try:
            result = subprocess.run(['trash', filepath], capture_output=True, timeout=30)
            return result.returncode == 0
        except FileNotFoundError:
            home = Path.home()
            trash = home / '.local/share/Trash/files'
            trash.mkdir(parents=True, exist_ok=True)
            shutil.move(filepath, trash / Path(filepath).name)
            return True


# ============================================================
# CLI Entry Point
# ============================================================

def main():
    """Command-line interface"""
    parser = argparse.ArgumentParser(description='InfoSeek v2.0 - Helper Tool')
    subparsers = parser.add_subparsers(dest='command', help='Subcommand')

    # create-folder
    p_create = subparsers.add_parser('create-folder', help='Create task folder')
    p_create.add_argument('name', help='Search subject name')
    p_create.add_argument('--workspace', default=os.environ.get('OPENCLAW_WORKSPACE', '.'))

    # generate-filename
    p_fname = subparsers.add_parser('generate-filename', help='Generate filename')
    p_fname.add_argument('--date', required=True, help='Publish date YYYYMMDD')
    p_fname.add_argument('--title', required=True, help='Article title')
    p_fname.add_argument('--website', required=True, help='Website name')
    p_fname.add_argument('--format', default='md', help='Output format')

    # save-content
    p_save = subparsers.add_parser('save-content', help='Save content')
    p_save.add_argument('--folder', required=True, help='Archive folder')
    p_save.add_argument('--filename', required=True, help='Filename')
    p_save.add_argument('--url', required=True, help='Original URL')
    p_save.add_argument('--website', required=True, help='Website name')
    p_save.add_argument('--source', default='', help='Media source')
    p_save.add_argument('--date', required=True, help='Publish date')
    p_save.add_argument('--title', required=True, help='Article title')
    p_save.add_argument('--author', default='Unknown', help='Author')
    p_save.add_argument('--editor', default='Unknown', help='Editor')
    p_save.add_argument('--content', required=True, help='Body content')
    p_save.add_argument('--task', required=True, help='Search task name')
    p_save.add_argument('--format', default='md', help='Output format')

    # add-url
    p_add = subparsers.add_parser('add-url', help='Add URL to database')
    p_add.add_argument('--url', required=True, help='URL')
    p_add.add_argument('--task', required=True, help='Search task')
    p_add.add_argument('--filename', required=True, help='Filename')
    p_add.add_argument('--workspace', default=os.environ.get('OPENCLAW_WORKSPACE', '.'))

    # check-duplicate
    p_check = subparsers.add_parser('check-duplicate', help='Check if URL is duplicate')
    p_check.add_argument('--url', required=True, help='URL')
    p_check.add_argument('--workspace', default=os.environ.get('OPENCLAW_WORKSPACE', '.'))

    # stats
    p_stats = subparsers.add_parser('stats', help='Show statistics')
    p_stats.add_argument('--workspace', default=os.environ.get('OPENCLAW_WORKSPACE', '.'))

    # backup
    p_backup = subparsers.add_parser('backup', help='Backup database')
    p_backup.add_argument('--workspace', default=os.environ.get('OPENCLAW_WORKSPACE', '.'))

    # delete
    p_delete = subparsers.add_parser('delete', help='Delete file (move to recycle bin)')
    p_delete.add_argument('--filepath', required=True, help='File path')
    p_delete.add_argument('--url', default='', help='Associated URL')
    p_delete.add_argument('--workspace', default=os.environ.get('OPENCLAW_WORKSPACE', '.'))

    # normalize-url
    p_norm = subparsers.add_parser('normalize-url', help='Normalize URL')
    p_norm.add_argument('--url', required=True, help='URL')

    # parse-date
    p_date = subparsers.add_parser('parse-date', help='Parse date')
    p_date.add_argument('--date', required=True, help='Date string')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return

    config = InfoSeekConfig()
    workspace = Path(getattr(args, 'workspace', os.environ.get('OPENCLAW_WORKSPACE', '.')))
    data_dir = workspace / 'infoseek'
    db_path = data_dir / 'infoseek.db'
    log_dir = data_dir

    logger = setup_logger(log_dir)

    try:
        if args.command == 'create-folder':
            archives_dir = workspace / 'infoseek-archives'
            safe_name = re.sub(r'[<>:"/\\|?*]', '_', args.name)
            folder = archives_dir / safe_name
            folder.mkdir(parents=True, exist_ok=True)
            print(f"[FOLDER] {folder}")
            logger.info(f"Created folder: {folder}")

        elif args.command == 'generate-filename':
            date = DateParser.parse(args.date)
            filename = ContentSaver.generate_filename(date, args.title, args.website, args.format)
            print(f"[FILENAME] {filename}")

        elif args.command == 'save-content':
            folder = Path(args.folder)
            folder.mkdir(parents=True, exist_ok=True)

            date = DateParser.parse(args.date)
            filename = ContentSaver.generate_filename(date, args.title, args.website, args.format)
            filename = ContentSaver.generate_unique_filename(folder, filename, args.url)

            content = {
                'url': args.url,
                'website': args.website,
                'source': args.source or args.website,
                'publishDate': DateParser.format_for_display(date),
                'title': args.title,
                'author': args.author,
                'editor': args.editor,
                'content': args.content,
                'searchTask': args.task
            }

            saver = ContentSaver(config)
            filepath = saver.save(folder, filename, content, args.format)
            print(f"[SAVED] {filepath}")
            logger.info(f"Saved file: {filepath}")

        elif args.command == 'add-url':
            db = URLDatabase(db_path)
            success = db.add(args.url, args.task, args.filename)
            if success:
                print("[OK] URL added to database")
                logger.info(f"Added URL: {args.url}")
            else:
                print("[WARN] URL already exists (duplicate)")
            db.close()

        elif args.command == 'check-duplicate':
            db = URLDatabase(db_path)
            is_dup = db.is_duplicate(args.url)
            print("[DUPLICATE] Yes" if is_dup else "[UNIQUE] No")
            db.close()

        elif args.command == 'stats':
            db = URLDatabase(db_path)
            stats = db.get_stats()
            print(f"[STATS] Total URLs: {stats['total_urls']}")
            print("\nBy task:")
            for task, count in stats['by_task'].items():
                print(f"   {task}: {count}")
            db.close()

        elif args.command == 'backup':
            db = URLDatabase(db_path)
            backup_dir = data_dir / 'backups'
            backup_path = db.backup(backup_dir)
            print(f"[BACKUP] Backup created: {backup_path}")
            logger.info(f"Database backup: {backup_path}")
            db.close()

        elif args.command == 'delete':
            filepath = args.filepath
            url = args.url

            print(f"About to delete: {filepath}")
            response = input("Confirm deletion? (yes/no): ")
            if response.lower() != 'yes':
                print("Cancelled")
                return

            success = FileDeleter.move_to_recycle_bin(filepath)
            if success:
                print("[DELETED] File moved to recycle bin")
                logger.info(f"Deleted file: {filepath}")

                if url:
                    db = URLDatabase(db_path)
                    db.mark_deleted(url)
                    db.close()
            else:
                print("[ERROR] Delete failed")

        elif args.command == 'normalize-url':
            normalized = URLNormalizer.normalize(args.url)
            print(f"Original:    {args.url}")
            print(f"Normalized:  {normalized}")

        elif args.command == 'parse-date':
            parsed = DateParser.parse(args.date)
            display = DateParser.format_for_display(parsed)
            print(f"Input:    {args.date}")
            print(f"Parsed:   {parsed}")
            print(f"Display:  {display}")

    except Exception as e:
        logger.error(f"Execution failed: {e}", exc_info=True)
        print(f"[ERROR] {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()

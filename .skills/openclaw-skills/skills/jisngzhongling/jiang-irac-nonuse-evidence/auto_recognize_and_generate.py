#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import csv
import datetime as dt
import hashlib
import json
import logging
import re
import shutil
import subprocess
import sys
import tempfile
import traceback
import unicodedata
import zipfile
import warnings
from xml.etree import ElementTree as ET
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from pypdf import PdfReader, PdfWriter
from utils.compliance import write_low_confidence_fields
from utils.logger import audit, set_run_context, setup_logger

LOGGER = setup_logger(__name__, log_dir=(Path(__file__).resolve().parent / "logs"))


def _audit_exception(step: str, exc: Exception, file_hint: str = "") -> None:
    LOGGER.exception("%s 发生异常", step)
    audit({
        "type": "exception",
        "step": step,
        "file": file_hint or __file__,
        "error": str(exc),
        "traceback": traceback.format_exc(),
        "ok": False,
        "reason_code": "recoverable_exception",
    })
try:
    from pypdf.errors import PdfReadWarning
except Exception as exc:
    _audit_exception("import_pypdf_errors", exc, "pypdf.errors")
    PdfReadWarning = UserWarning
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    import fitz  # PyMuPDF, used for precise highlight annotations
except Exception as exc:
    _audit_exception("import_fitz", exc, "fitz")
    fitz = None
if fitz is not None:
    try:
        # Suppress MuPDF structural warnings from malformed producer PDFs.
        fitz.TOOLS.mupdf_display_warnings(False)
    except Exception as exc:
        _audit_exception("fitz_mupdf_display_warnings", exc, "fitz.TOOLS")

# Reduce noisy non-fatal pypdf warnings caused by malformed producer PDFs.
warnings.filterwarnings(
    "ignore",
    category=PdfReadWarning,
    message=r"Multiple definitions in dictionary .* for key /Info",
)
warnings.filterwarnings(
    "ignore",
    category=UserWarning,
    message=r"Multiple definitions in dictionary .* for key /Info",
)
# pypdf logs non-fatal PDF structure fixes via logging.warning; suppress these noisy messages.
for _name in ("pypdf", "pypdf._utils", "pypdf._reader", "pypdf.generic._data_structures"):
    _logger = logging.getLogger(_name)
    _logger.setLevel(logging.ERROR)
    if not _logger.handlers:
        _logger.addHandler(logging.NullHandler())
    _logger.propagate = False
del _name, _logger

# Extra guard: patch pypdf logger_warning so malformed-PDF noise never surfaces.
try:
    import pypdf._utils as _pypdf_utils
    import pypdf.generic._data_structures as _pypdf_ds

    _orig_logger_warning = _pypdf_utils.logger_warning

    def _quiet_logger_warning(msg: str, src: str) -> None:
        s = str(msg)
        if "Multiple definitions in dictionary" in s and "key /Info" in s:
            return
        _orig_logger_warning(msg, src)

    _pypdf_utils.logger_warning = _quiet_logger_warning
    _pypdf_ds.logger_warning = _quiet_logger_warning
except Exception as exc:
    _audit_exception("patch_pypdf_logger_warning", exc, "pypdf._utils")


DATE_RE = re.compile(r"(20\d{2})[年./-](\d{1,2})[月./-](\d{1,2})日?")
DATE_FUZZY_RE = re.compile(r"(20[0-9OolISB]{2})\s*[年./-]\s*([0-9OolISB]{1,2})\s*[月./-]\s*([0-9OolISB]{1,2})\s*日?")
MONTH_DAY_RE = re.compile(r"(?<!\d)([0-9OolISB]{1,2})\s*月\s*([0-9OolISB]{1,2})\s*[日号]?")
COMPACT_DATE_RE = re.compile(r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)")
DATE_NUMERIC_RE = re.compile(r"(?<!\d)(20\d{2})[./-](\d{1,2})[./-](\d{1,2})(?!\d)")
REG_RE = re.compile(r"(?:注册号|商标注册号)[:：]?\s*([0-9]{6,})")
REG_NO_TAG_RE = re.compile(r"第\s*([0-9]{6,})\s*号")
CLASS_RE = re.compile(r"第\s*([0-9]{1,2})\s*类")
APPLICANT_RE = re.compile(r"申请\s*人[:：]?\s*([^\n]{2,80})")
RESPONDENT_RE = re.compile(r"(?:被申请人|商标注册人|注册人|答辩人)[:：]?\s*([^\n]{2,80})")
MARK_RE = re.compile(r"号[“\"]?([\u4e00-\u9fa5A-Za-z0-9·-]{2,30})[”\"]?商标")
MARK_TEXT_RE = re.compile(r"[“\"]?([\u4e00-\u9fa5A-Za-z0-9·-]{2,24})[”\"]?商标")
MARK_ARCHIVE_RE = re.compile(r"[“\"]([^”\"]{2,60})[”\"]\s*商标档案")
COMPANY_RE = re.compile(r"([\u4e00-\u9fa5A-Za-z0-9（）()·]{2,50}(?:有限责任公司|股份有限公司|有限公司))")
GOODS_RE = re.compile(r"第\s*\d+\s*类[“\"]([^”\"]{1,50})[”\"]")
GOODS_LABEL_RE = re.compile(r"(?:核定使用商品|核定商品|指定商品(?:和服务)?|商品/服务)[:：]?\s*([^\n]{2,120})")
INVALID_FILENAME_RE = re.compile(r"[\\/:*?\"<>|]+")
NON_DIGIT_RE = re.compile(r"[^\d]+")
NAME_TEMPLATE_RE = re.compile(r"^(E\d{3})_([^_]+)_([^_]+)_(.+)\.pdf$", re.IGNORECASE)

PDFTOPPM_BIN = shutil.which("pdftoppm") or "/opt/homebrew/bin/pdftoppm"
TESSERACT_BIN = shutil.which("tesseract") or "/opt/homebrew/bin/tesseract"

EVIDENCE_TYPE_ALIAS = {
    "程序文件": "程序",
    "资质主体证明": "主体",
    "交易凭证": "交易",
    "线上店铺展示": "店铺",
    "商品展示页": "商品",
    "其他材料": "其他",
}

AUTO_SORT_CATEGORY_ORDER = {
    "程序类": 1,
    "权属类": 2,
    "主体资质类": 3,
    "实体实物类": 4,
    "网络展示类": 5,
    "交易履约类": 6,
    "合同票据类": 7,
    "其他补强类": 8,
}

SOURCE_BUCKET_ORDER = {
    "fast": 1,
    "full": 2,
    "uncategorized": 3,
    "direct_bind": 4,
}

MIX_DEEP_SCAN_KEYWORDS = [
    "合同", "协议", "发票", "票据", "订单", "下单", "支付", "收款", "转账",
    "物流", "运单", "签收", "评价", "评论", "追评", "后台", "对账", "结算",
]

TIME_NOISE_EXCLUDE_KEYWORDS = [
    "截至", "截止", "活动时间", "有效期", "保质期", "生产日期", "出厂日期",
    "检验日期", "打印日期", "系统时间", "模板日期", "公告日期",
    "体检", "体检通知", "培训通知", "放假通知", "节假日", "活动截止",
]
TIME_TRANSACTION_HINT_KEYWORDS = [
    "签订日期", "开票日期", "发票日期", "支付时间", "付款时间", "交易时间",
    "下单时间", "发货时间", "签收时间", "物流时间", "到账时间", "合同日期",
]
TIME_SYSTEM_HINT_KEYWORDS = [
    "评价时间", "评论时间", "追评时间", "创建时间", "更新时间", "发布时间",
    "入驻时间", "订单创建", "系统生成", "平台时间",
]
TIME_NOISE_EXCLUDE_BY_KIND = {
    # 交易证据中，以下日期通常不是交易发生时点
    "交易凭证": ["保质期", "生产日期", "出厂日期", "检验日期", "有效期", "批号日期"],
    # 网页展示证据中，活动/模板日期常见噪声
    "线上店铺展示": ["活动时间", "截至", "截止", "模板日期", "系统当前时间", "打印日期"],
    "商品展示页": ["活动时间", "截至", "截止", "模板日期", "系统当前时间", "打印日期"],
    # 程序证据默认不作为使用时间锚点
    "程序文件": ["通知日期", "送达日期", "决定日期", "立案日期", "须知日期", "体检日期", "放假日期", "活动日期"],
}

TIME_ANCHOR_CHANNEL_PRIORITY = {
    "transaction_date": 4,
    "system_generated": 3,
    "content_claimed": 2,
    "uncertain": 1,
}

PUBLIC_NOTE_EXCLUDE_KEYWORDS = [
    "关键语句",
    "时间待核验",
    "组证关联",
    "碎片合并",
    "映射依据",
    "时间识别说明",
    "表现形式识别",
    "候选通道",
    "时间来源类型",
    "时间来源通道",
    "未提取到稳定关键词句",
]

PROMOTION_MATERIAL_KEYWORDS = [
    "宣传", "海报", "画册", "折页", "展会", "展位", "易拉宝", "讲座", "路演", "招募", "邀请函",
]

GOODS_SYNONYM_GROUPS = [
    {"卫生纸", "卷纸", "抽纸", "纸巾", "面巾纸"},
    {"口罩", "医用口罩", "防护口罩"},
    {"医用敷料", "敷料", "纱布", "绷带"},
    {"护理垫", "隔尿垫", "护理巾"},
]

WEB_VERIFIABILITY_STRONG = [
    "订单号", "交易单号", "支付时间", "付款时间", "物流单号", "签收时间", "后台导出",
]
WEB_VERIFIABILITY_MEDIUM = [
    "评价", "评论", "追评", "店铺名称", "店铺主体", "发布时间", "更新时间",
]

COVERAGE_LANE_A = "A线（可计入覆盖）"
COVERAGE_LANE_B = "B线（仅参考）"

SORT_CAT_KEYWORDS = {
    "程序类": [
        "通知书", "送达", "补正", "决定书", "裁定书", "受理", "答辩通知", "撤销连续三年不使用", "撤三",
    ],
    "权属类": [
        "注册证", "商标档案", "核准注册", "商标公告", "变更", "转让", "许可备案", "商标注册证", "续展",
    ],
    "主体资质类": [
        "营业执照", "主体资格", "资质", "许可", "备案", "授权书", "委托书", "统一社会信用代码", "开户许可",
    ],
    "实体实物类": [
        "厂房", "工厂", "仓储", "仓库", "包装", "门头", "样品", "实拍", "生产线", "成品库",
    ],
    "网络展示类": [
        "网店", "店铺", "官网", "商品详情", "详情页", "平台后台", "商家中心", "上架", "网页", "店铺首页",
    ],
    "交易履约类": [
        "订单", "支付", "物流", "运单", "签收", "发货", "工单", "履约", "回单", "收款",
    ],
    "合同票据类": [
        "合同", "协议", "发票", "结算", "对账单", "报价单", "采购单", "销货清单",
    ],
    "其他补强类": [
        "检测", "检验", "媒体", "报道", "奖项", "荣誉", "认证证书", "评价",
    ],
}

TARGET_CN = {
    "T1": "使用主体",
    "T2": "商标标识",
    "T3": "商品/服务",
    "T4": "使用时间",
    "T5": "使用场景",
    "T6": "交易闭环",
}

STOP_MARK_TOKENS = {
    "商标", "注册", "核准", "类别", "申请", "指定", "期间", "通知书",
    "证据", "商品", "服务", "有限公司", "有限责任公司", "商标局", "国家知识产权局",
}

BAD_MARK_WORDS = {
    "关于", "提供", "注册", "申请", "通知", "答辩", "撤销", "商标局",
    "国家", "知识产权", "有限公司", "有限责任公司", "核定", "使用", "中华人民共和国",
    "微信图片", "微信", "图片", "截图", "相册", "拍照", "拍摄", "文件", "材料",
    "img", "image", "jpg", "jpeg", "png", "pdf", "docx", "xlsx", "ppt",
}

REVIEW_KEYWORDS = ["评价", "追评", "买家", "已购", "评论", "晒图", "收货"]
ORDER_KEYWORDS = ["订单", "下单", "付款", "支付", "成交", "购买", "已购", "交易", "交易成功"]
LOGISTICS_KEYWORDS = ["发货", "揽收", "物流", "签收", "派送", "到达", "运输", "快递", "收货"]
DISPLAY_KEYWORDS = ["店铺", "首页", "详情页", "宝贝", "商品", "页面", "旗舰店", "工厂店", "客服", "收藏", "销量"]
TX_STRONG_KEYWORDS = ["订单编号", "支付成功", "付款时间", "交易快照", "物流单号", "签收时间", "发票", "对账单", "电子回单"]
UNCERTAIN_TIME_KEYWORDS = ["预计", "预售", "发货", "送达", "小时内", "达达"]
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}
OFFICE_SUFFIXES = {".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx"}
CONTRACT_KEYWORDS = ["合同", "协议", "购销", "采购单", "订购单", "订单合同"]
INVOICE_KEYWORDS = ["发票", "增值税", "专票", "普票", "票据", "开票", "税额", "价税合计"]
PAYMENT_KEYWORDS = ["转账", "汇款", "收款", "对账", "流水", "支付截图", "交易记录", "电子回单", "付款"]
TRADE_PARTY_HINT_KEYWORDS = [
    "需方", "购方", "购买方", "买方", "乙方", "客户", "受票方", "付款方", "收款方",
    "销方", "销售方", "开票方", "收货方", "签约方",
]
EXHIBIT_AD_KEYWORDS = [
    "展会", "参展", "展位", "展馆", "会展", "展台", "展位销售", "展位合同",
    "广告位", "宣传广告", "参展方", "主办方", "广告发布", "宣传投放",
]
BROCHURE_KEYWORDS = [
    "宣传册", "宣传手册", "产品手册", "画册", "样册", "折页", "彩页", "产品目录", "品牌手册", "企业介绍",
    "企业概况", "企业发展历程", "企业荣誉", "员工风采", "ABOUT US", "ABOUTUS",
    "COMPANY INTRODUCTION", "HISTORY", "HONOR", "STYLE",
]
TEST_REPORT_KEYWORDS = [
    "检测报告", "检验报告", "质检报告", "检验日期", "检验结论", "检验依据", "报告编号",
    "送检", "受检", "检测机构", "抽样", "检验员", "检测项目", "检测结果",
]
FORM_NOTICE_KEYWORDS = [
    "通知书", "答辩通知", "受理通知", "补正通知", "决定书", "送达", "国家知识产权局", "商标局",
]
FORM_CONTRACT_KEYWORDS = [
    "合同", "协议", "甲方", "乙方", "签章", "盖章", "签订", "履行期限",
]
FORM_INVOICE_KEYWORDS = [
    "发票", "增值税", "发票代码", "发票号码", "价税合计", "税额", "开票日期",
]
FORM_POSTER_KEYWORDS = [
    "海报", "展板", "宣传图", "广告图", "易拉宝", "活动主视觉", "产品宣传",
]
SCENE_POSTER_HINT_KEYWORDS = [
    "活动", "讲座", "论坛", "沙龙", "路演", "培训", "宣传", "推广", "运营", "招商",
    "企业介绍", "发展历程", "企业荣誉", "员工风采", "企业文化", "知识产权",
    "创业", "研究院", "公告栏", "展板",
]
PROPERTY_SERVICE_KEYWORDS = [
    "不动产", "办公室", "公寓", "出租", "公寓管理", "不动产经纪", "不动产管理", "物业",
]
FORM_PACKAGING_KEYWORDS = [
    "包装", "外箱", "包装盒", "标签", "盒装", "袋装", "净含量", "批号", "生产日期", "包装实拍",
]
FORM_SCENE_PHOTO_KEYWORDS = [
    "现场", "实拍", "工厂", "厂区", "车间", "仓库", "门头", "前台", "展厅", "生产线", "办公区",
]
AMOUNT_NUMBER_PATTERN = r"((?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d{1,2})?)"
AMOUNT_KEYWORD_RE = re.compile(
    r"(?:价税合计|金额|含税金额|不含税金额|总金额|合计|价款|付款金额|开票金额|合同金额|订单金额|应付金额|实付金额)"
    r"\s*[:：]?\s*[¥￥]?\s*" + AMOUNT_NUMBER_PATTERN
)
AMOUNT_CURRENCY_RE = re.compile(r"[¥￥]\s*" + AMOUNT_NUMBER_PATTERN)
AMOUNT_YUAN_RE = re.compile(AMOUNT_NUMBER_PATTERN + r"\s*元")
FIVE_SCAN_ROUNDS = 5
ROUND_SIGNAL_KEYWORDS = [
    "商标", "商品", "服务", "店铺", "详情", "评价", "订单", "支付", "物流",
    "发票", "合同", "营业执照", "主体", "公告", "注册号", "程序", "通知书",
]
EVIDENCE_NAME_BAD_PHRASES = {
    "未提取到稳定关键词句",
    "微信图片",
    "商品详情",
    "证据",
    "截图",
    "图片",
    "其他",
    "img",
    "image",
}
EVIDENCE_NAME_GENERIC_TOPICS = {
    "商品详情",
    "微信图片",
    "其他材料",
    "其他",
    "截图",
    "图片",
    "证据",
}
SOFFICE_BIN = (
    shutil.which("soffice")
    or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
)

DIRECT_BIND_SUPPORTED_SUFFIXES = {".pdf", *IMAGE_SUFFIXES, *OFFICE_SUFFIXES}


def run(cmd: List[str], timeout_sec: Optional[int] = None) -> Tuple[int, str, str]:
    try:
        p = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=timeout_sec,
        )
        return p.returncode, p.stdout, p.stderr
    except subprocess.TimeoutExpired:
        return 124, "", "timeout"


def file_sha1(path: Path) -> str:
    h = hashlib.sha1()
    try:
        with path.open("rb") as f:
            while True:
                chunk = f.read(1024 * 1024)
                if not chunk:
                    break
                h.update(chunk)
        return h.hexdigest()
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    try:
        with path.open("rb") as f:
            while True:
                chunk = f.read(1024 * 1024)
                if not chunk:
                    break
                h.update(chunk)
        return h.hexdigest()
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""


def _is_subpath(child: Path, parent: Path) -> bool:
    try:
        child_resolved = child.resolve()
        parent_resolved = parent.resolve()
    except Exception as exc:
        _audit_exception("auto_recognize.path_resolve_failed", exc, __file__)
        return False

    if child_resolved == parent_resolved:
        return True

    # 非子路径属于正常业务分支，不应按异常落日志。
    if hasattr(child_resolved, "is_relative_to"):
        return child_resolved.is_relative_to(parent_resolved)  # type: ignore[attr-defined]

    # 兼容低版本 Python（无 is_relative_to），仅把真实异常记审计。
    try:
        child_resolved.relative_to(parent_resolved)
        return True
    except ValueError:
        return False
    except Exception as exc:
        _audit_exception("auto_recognize.path_relative_check_failed", exc, __file__)
        return False


def parse_output_formats(spec: str) -> List[str]:
    txt = str(spec or "").strip().lower()
    if not txt:
        return ["docx"]
    out: List[str] = []
    for token in re.split(r"[,;\s]+", txt):
        t = token.strip().lower()
        if not t:
            continue
        if t.startswith("."):
            t = t[1:]
        if t in {"docx", "pdf", "md", "txt", "html"} and t not in out:
            out.append(t)
    return out or ["docx"]


def resolve_mix_scan_mode(path_name: str, mode: str) -> str:
    m = str(mode or "").strip().lower()
    if m != "mix":
        return m if m in {"fast", "full"} else "fast"
    src = str(path_name or "")
    if any(k in src for k in MIX_DEEP_SCAN_KEYWORDS):
        return "full"
    return "fast"


def should_upgrade_to_full_by_content(path_name: str, text: str) -> bool:
    src = f"{path_name}\n{text or ''}"
    # 关键证据优先深扫，避免文件名不规范导致误入 fast。
    strong_hits = [
        "合同", "协议", "发票", "增值税", "订单号", "交易单号", "支付时间",
        "付款时间", "物流单号", "签收", "评价", "追评", "对账单", "结算单",
        "银行回单", "收款方", "销售方",
    ]
    if sum(1 for k in strong_hits if k in src) >= 1:
        return True
    # 金额+交易词共现，也视为深扫证据。
    has_amount = bool(re.search(r"(¥|￥|金额|价税合计|含税|不含税|\d+\.\d{2})", src))
    has_trade = any(k in src for k in ["下单", "支付", "发票", "物流", "签收", "合同"])
    return has_amount and has_trade


def source_bucket_weight(bucket: str) -> int:
    return SOURCE_BUCKET_ORDER.get(str(bucket or "").strip().lower(), 99)


def merge_bucket_priority(existing_bucket: str, new_bucket: str) -> str:
    a = source_bucket_weight(existing_bucket)
    b = source_bucket_weight(new_bucket)
    return new_bucket if b < a else existing_bucket


def _normalize_case_override_key(key: str) -> str:
    k = str(key or "").strip().lower()
    alias = {
        "key": "key",
        "case_no": "case_no",
        "案号": "case_no",
        "reg_no": "reg_no",
        "注册号": "reg_no",
        "商标号": "reg_no",
        "respondent": "respondent",
        "被申请人": "respondent",
        "注册人": "respondent",
        "答辩人": "respondent",
        "applicant": "applicant",
        "申请人": "applicant",
        "class": "class",
        "类别": "class",
        "国际分类": "class",
        "designated_goods_services": "designated_goods_services",
        "核定商品": "designated_goods_services",
        "核定商品服务": "designated_goods_services",
        "核定商品/服务": "designated_goods_services",
        "goods_services": "designated_goods_services",
        "revoked_goods_services": "revoked_goods_services",
        "被撤商品": "revoked_goods_services",
        "被撤商品服务": "revoked_goods_services",
        "被撤商品/服务": "revoked_goods_services",
        "defense_goods_services": "defense_goods_services",
        "答辩商品": "defense_goods_services",
        "答辩服务": "defense_goods_services",
        "答辩商品/服务": "defense_goods_services",
        "use_period_start": "use_period_start",
        "指定期间起": "use_period_start",
        "答辩期间起": "use_period_start",
        "use_period_end": "use_period_end",
        "指定期间止": "use_period_end",
        "答辩期间止": "use_period_end",
        "cnipa_notice_ref": "cnipa_notice_ref",
        "通知书文号": "cnipa_notice_ref",
        "mark_name": "mark_name",
        "商标名称": "mark_name",
        "商标": "mark_name",
        "mark_image": "mark_image",
        "商标图样": "mark_image",
        "respondent_address": "respondent_address",
        "答辩人地址": "respondent_address",
        "agent_company": "agent_company",
        "代理公司": "agent_company",
        "代理公司名称": "agent_company",
        "agent_address": "agent_address",
        "代理公司地址": "agent_address",
        "contact_phone": "contact_phone",
        "联系电话": "contact_phone",
    }
    return alias.get(k, str(key or "").strip())


def _parse_case_override_text(text: str) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for line in str(text or "").splitlines():
        s = line.strip()
        if not s or s.startswith("#"):
            continue
        m = re.match(r"^\s*([^:=：\t]{1,80})\s*[:=：\t]\s*(.+?)\s*$", s)
        if not m:
            continue
        key = _normalize_case_override_key(m.group(1))
        val = str(m.group(2) or "").strip()
        if key and val:
            out[key] = val
    return out


def _parse_case_override_sheet(df: pd.DataFrame) -> Dict[str, str]:
    out: Dict[str, str] = {}
    if df is None or df.empty:
        return out
    cols = [str(c).strip() for c in df.columns]
    low = [c.lower() for c in cols]
    if ("key" in low and "value" in low) or ("字段" in cols and "内容" in cols):
        key_col = cols[low.index("key")] if "key" in low else "字段"
        val_col = cols[low.index("value")] if "value" in low else "内容"
        for _, row in df.iterrows():
            k = _normalize_case_override_key(row.get(key_col, ""))
            v = str(row.get(val_col, "") or "").strip()
            if k and v:
                out[k] = v
        return out
    row0 = df.iloc[0].to_dict()
    for k, v in row0.items():
        nk = _normalize_case_override_key(k)
        sv = str(v or "").strip()
        if nk and sv and sv.lower() != "nan":
            out[nk] = sv
    return out


def load_case_meta_override(path_str: str) -> Dict[str, str]:
    p = Path(str(path_str or "").strip())
    if not p:
        return {}
    if not str(path_str or "").strip():
        return {}
    if not p.exists() or not p.is_file():
        raise RuntimeError(f"案件信息覆盖文件不存在：{p}")
    suffix = p.suffix.lower()
    out: Dict[str, str] = {}
    try:
        if suffix == ".json":
            raw = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(raw, dict):
                if isinstance(raw.get("case_meta"), dict):
                    raw = raw.get("case_meta", {})
                for k, v in raw.items():
                    if v is None:
                        continue
                    nk = _normalize_case_override_key(k)
                    sv = str(v).strip()
                    if nk and sv:
                        out[nk] = sv
            else:
                raise RuntimeError(f"案件信息覆盖文件格式错误（JSON应为对象）：{p}")
        elif suffix in {".txt", ".md"}:
            out = _parse_case_override_text(p.read_text(encoding="utf-8", errors="ignore"))
        elif suffix in {".docx"}:
            out = _parse_case_override_text(_docx_plain_text(p))
        elif suffix in {".csv"}:
            df = pd.read_csv(p, dtype=str).fillna("")
            out = _parse_case_override_sheet(df)
        elif suffix in {".xlsx", ".xls"}:
            dfs = pd.read_excel(p, sheet_name=None, dtype=str)
            for _, df in dfs.items():
                out.update(_parse_case_override_sheet(df.fillna("")))
        else:
            raise RuntimeError(f"案件信息覆盖文件格式暂不支持：{suffix}（支持 JSON/TXT/MD/CSV/XLS/XLSX/DOCX）")
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"案件信息覆盖文件解析失败：{p} ({exc})")
    return out


def _docx_plain_text(path: Path) -> str:
    if not path.exists():
        return ""
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""
    parts: List[str] = []
    for p in doc.paragraphs:
        t = str(getattr(p, "text", "") or "").strip()
        if t:
            parts.append(t)
    for tb in doc.tables:
        for row in tb.rows:
            vals: List[str] = []
            for c in row.cells:
                t = str(getattr(c, "text", "") or "").strip()
                if t:
                    vals.append(t)
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def _docx_to_pdf(src_docx: Path, dst_pdf: Path) -> bool:
    if not SOFFICE_BIN or not Path(SOFFICE_BIN).exists():
        return False
    dst_pdf.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="docx2pdf_") as td:
        td_path = Path(td)
        code, _, _ = run(
            [
                SOFFICE_BIN,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(td_path),
                str(src_docx),
            ],
            timeout_sec=120,
        )
        if code != 0:
            return False
        produced = td_path / f"{src_docx.stem}.pdf"
        if not produced.exists():
            cands = sorted(td_path.glob("*.pdf"))
            if not cands:
                return False
            produced = cands[0]
        shutil.move(str(produced), str(dst_pdf))
        return dst_pdf.exists()


def export_documents_in_formats(
    out_dir: Path,
    docx_paths: List[Path],
    formats_spec: str,
) -> Dict[str, Any]:
    formats = parse_output_formats(formats_spec)
    result: Dict[str, Any] = {
        "requested_formats": formats,
        "generated": [],
        "warnings": [],
    }
    if not formats:
        return result

    for src in docx_paths:
        if not src.exists():
            result["warnings"].append(f"文书不存在，跳过导出：{src}")
            continue
        base = src.with_suffix("")
        txt = ""
        if any(f in formats for f in ["txt", "md", "html"]):
            txt = _docx_plain_text(src)
        if "pdf" in formats:
            dst_pdf = base.with_suffix(".pdf")
            ok = _docx_to_pdf(src, dst_pdf)
            if ok:
                result["generated"].append(str(dst_pdf))
            else:
                result["warnings"].append(f"PDF导出失败（请检查LibreOffice）：{src.name}")
        if "txt" in formats:
            dst_txt = base.with_suffix(".txt")
            dst_txt.write_text(txt, encoding="utf-8")
            result["generated"].append(str(dst_txt))
        if "md" in formats:
            dst_md = base.with_suffix(".md")
            title = src.stem
            dst_md.write_text(f"# {title}\n\n```\n{txt}\n```\n", encoding="utf-8")
            result["generated"].append(str(dst_md))
        if "html" in formats:
            dst_html = base.with_suffix(".html")
            esc = (
                txt.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            html = (
                "<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
                f"<title>{src.stem}</title></head><body><pre>{esc}</pre></body></html>"
            )
            dst_html.write_text(html, encoding="utf-8")
            result["generated"].append(str(dst_html))
    return result


def is_precheck_approved(precheck_docx: Path) -> bool:
    if not precheck_docx.exists():
        return False
    text = _docx_plain_text(precheck_docx)
    text = text.replace("[✓]", "☑").replace("[√]", "☑").replace("✅", "☑")
    return ("☑ 信息无误" in text) and ("☑ 同意输出" in text)


def image_to_single_page_pdf(image_path: Path, out_pdf: Path) -> None:
    if fitz is None:
        raise RuntimeError("检测到图片证据，但未安装PyMuPDF（fitz），无法自动转PDF。")
    pix = fitz.Pixmap(str(image_path))
    doc = fitz.open()
    try:
        page = doc.new_page(width=float(pix.width), height=float(pix.height))
        page.insert_image(page.rect, filename=str(image_path))
        out_pdf.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_pdf))
    finally:
        doc.close()
        pix = None


def _soffice_available() -> bool:
    return Path(SOFFICE_BIN).exists()


def office_to_pdf(office_path: Path, out_pdf: Path) -> bool:
    if not _soffice_available():
        return False
    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="office2pdf_") as td:
        td_path = Path(td)
        code, _, _ = run(
            [
                SOFFICE_BIN,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(td_path),
                str(office_path),
            ],
            timeout_sec=120,
        )
        if code != 0:
            return False
        produced = td_path / f"{office_path.stem}.pdf"
        if not produced.exists():
            cands = sorted(td_path.glob("*.pdf"))
            if not cands:
                return False
            produced = cands[0]
        shutil.move(str(produced), str(out_pdf))
        return out_pdf.exists()


def collect_evidence_pdf_inputs(
    evidence_dir: Path,
    out_dir: Path,
    evidence_mode: str = "auto",
) -> Tuple[List[Path], Dict[str, Any], Dict[Path, str], Dict[Path, str], Dict[Path, Dict[str, str]]]:
    files: List[Path] = []
    prefill_texts: Dict[Path, str] = {}
    mode_by_file: Dict[Path, str] = {}
    source_meta_by_file: Dict[Path, Dict[str, str]] = {}
    stats = {
        "pdf": 0,
        "images_converted": 0,
        "office_converted": 0,
        "office_skipped": 0,
        "office_text_extracted": 0,
        "office_converter": "soffice" if _soffice_available() else "unavailable",
        "mode": evidence_mode,
    }
    converted_dir = out_dir / "_converted_inputs_pdf"
    skip_roots: List[Path] = []
    if _is_subpath(out_dir, evidence_dir):
        skip_roots.append(out_dir.resolve())
    seen_content_hash: set[str] = set()

    def _origin_meta(src_path: Path) -> Dict[str, str]:
        try:
            rel = str(src_path.relative_to(evidence_dir))
        except Exception:
            rel = src_path.name
        return {
            "origin_rel": rel,
            "origin_dir": str(Path(rel).parent) if rel else "",
            "origin_name": src_path.name,
        }

    for p in sorted(evidence_dir.rglob("*")):
        if not p.is_file() or p.name.startswith("._"):
            continue
        if any(_is_subpath(p, skip_root) for skip_root in skip_roots):
            continue
        suffix = p.suffix.lower()
        src_hash = file_sha256(p)
        if src_hash and src_hash in seen_content_hash:
            continue
        if suffix == ".pdf":
            files.append(p)
            mode_by_file[p] = evidence_mode
            source_meta_by_file[p] = _origin_meta(p)
            stats["pdf"] += 1
            if src_hash:
                seen_content_hash.add(src_hash)
            continue
        if suffix in IMAGE_SUFFIXES:
            stem = sanitize_component(p.stem, max_len=80) or "image"
            out_pdf = converted_dir / f"{stem}_{file_sha1(p)[:8]}.pdf"
            if not out_pdf.exists():
                image_to_single_page_pdf(p, out_pdf)
            files.append(out_pdf)
            mode_by_file[out_pdf] = evidence_mode
            source_meta_by_file[out_pdf] = _origin_meta(p)
            stats["images_converted"] += 1
            if src_hash:
                seen_content_hash.add(src_hash)
            continue
        if suffix in OFFICE_SUFFIXES:
            stem = sanitize_component(p.stem, max_len=80) or "office"
            out_pdf = converted_dir / f"{stem}_{file_sha1(p)[:8]}.pdf"
            office_text = read_office_plain_text(p)
            if office_text:
                stats["office_text_extracted"] += 1
            if not out_pdf.exists():
                ok = office_to_pdf(p, out_pdf)
                if not ok:
                    stats["office_skipped"] += 1
                    continue
            files.append(out_pdf)
            mode_by_file[out_pdf] = evidence_mode
            source_meta_by_file[out_pdf] = _origin_meta(p)
            if office_text:
                prefill_texts[out_pdf] = office_text
            stats["office_converted"] += 1
            if src_hash:
                seen_content_hash.add(src_hash)

    # 去重后返回，避免同名重复收集
    unique: Dict[str, Path] = {}
    for p in files:
        unique[str(p.resolve())] = p
    unique_files = sorted(unique.values())
    prefill_unique: Dict[Path, str] = {}
    mode_unique: Dict[Path, str] = {}
    source_meta_unique: Dict[Path, Dict[str, str]] = {}
    for p in unique_files:
        t = prefill_texts.get(p, "")
        if t:
            prefill_unique[p] = t
        mode_unique[p] = mode_by_file.get(p, evidence_mode)
        source_meta_unique[p] = source_meta_by_file.get(p, {"origin_rel": p.name, "origin_dir": "", "origin_name": p.name})
    return unique_files, stats, prefill_unique, mode_unique, source_meta_unique


def _iter_supported_material_files(root_or_file: Path) -> List[Path]:
    if root_or_file.is_file():
        if root_or_file.suffix.lower() in DIRECT_BIND_SUPPORTED_SUFFIXES and not root_or_file.name.startswith("._"):
            return [root_or_file]
        return []
    if not root_or_file.exists() or not root_or_file.is_dir():
        return []
    out: List[Path] = []
    for p in sorted(root_or_file.rglob("*")):
        if not p.is_file() or p.name.startswith("._"):
            continue
        if p.suffix.lower() in DIRECT_BIND_SUPPORTED_SUFFIXES:
            out.append(p)
    return out


def load_direct_bind_entries_from_config(config_path: str) -> List[Dict[str, str]]:
    path = (config_path or "").strip()
    if not path:
        return []
    cfg = Path(path)
    if not cfg.exists() or not cfg.is_file():
        raise RuntimeError(f"直装材料配置文件不存在：{cfg}")
    try:
        payload = json.loads(cfg.read_text(encoding="utf-8"))
    except Exception as exc:
        raise RuntimeError(f"直装材料配置解析失败：{cfg} ({exc})")
    if isinstance(payload, dict):
        items = payload.get("items", [])
    elif isinstance(payload, list):
        items = payload
    else:
        raise RuntimeError("直装材料配置格式错误：应为对象{items:[...]}或数组[...]")
    if not isinstance(items, list):
        raise RuntimeError("直装材料配置格式错误：items 必须为数组")

    out: List[Dict[str, str]] = []
    for idx, raw in enumerate(items, start=1):
        if not isinstance(raw, dict):
            continue
        label = str(raw.get("label", "") or "").strip()
        path_val = str(raw.get("path", "") or "").strip()
        custom_name = str(raw.get("name", "") or raw.get("custom_name", "") or "").strip()
        if not label or not path_val:
            continue
        out.append({
            "label": label,
            "path": path_val,
            "custom_name": custom_name,
            "item_index": str(idx),
        })
    return out


def collect_direct_bind_pdf_inputs(
    config_path: str,
    out_dir: Path,
) -> Tuple[List[Dict[str, str]], Dict[str, Any]]:
    stats = {
        "direct_bind_items": 0,
        "direct_bind_files": 0,
        "direct_bind_images_converted": 0,
        "direct_bind_office_converted": 0,
        "direct_bind_office_skipped": 0,
    }
    entries = load_direct_bind_entries_from_config(config_path)
    if not entries:
        return [], stats
    stats["direct_bind_items"] = len(entries)

    converted_dir = out_dir / "_converted_direct_bind_pdf"
    files_out: List[Dict[str, str]] = []
    seen: set[str] = set()

    for item in entries:
        label = item.get("label", "")
        custom_name = item.get("custom_name", "")
        item_idx = item.get("item_index", "0")
        src_path = Path(item.get("path", "")).expanduser()
        if not src_path.exists():
            continue
        src_files = _iter_supported_material_files(src_path)
        seq = 0
        for src in src_files:
            suffix = src.suffix.lower()
            src_pdf = src
            if suffix in IMAGE_SUFFIXES:
                stem = sanitize_component(src.stem, max_len=80) or "image"
                src_pdf = converted_dir / f"{stem}_{file_sha1(src)[:8]}.pdf"
                if not src_pdf.exists():
                    image_to_single_page_pdf(src, src_pdf)
                stats["direct_bind_images_converted"] += 1
            elif suffix in OFFICE_SUFFIXES:
                stem = sanitize_component(src.stem, max_len=80) or "office"
                src_pdf = converted_dir / f"{stem}_{file_sha1(src)[:8]}.pdf"
                if not src_pdf.exists():
                    ok = office_to_pdf(src, src_pdf)
                    if not ok:
                        stats["direct_bind_office_skipped"] += 1
                        continue
                stats["direct_bind_office_converted"] += 1
            elif suffix != ".pdf":
                continue

            k = str(src_pdf.resolve())
            if k in seen:
                continue
            seen.add(k)
            seq += 1
            custom_name_i = custom_name
            if custom_name and len(src_files) > 1:
                custom_name_i = f"{custom_name}{seq}"
            files_out.append({
                "label": label,
                "custom_name": custom_name_i,
                "src_pdf": str(src_pdf),
                "src_name": src.name,
                "item_index": item_idx,
                "item_seq": str(seq),
            })

    files_out.sort(key=lambda x: (int(x.get("item_index", "0") or 0), int(x.get("item_seq", "0") or 0), x.get("src_name", "")))
    stats["direct_bind_files"] = len(files_out)
    return files_out, stats


def _merge_scan_text(old_text: str, new_text: str) -> str:
    old = (old_text or "").strip()
    new = (new_text or "").strip()
    if not old:
        return new
    if not new:
        return old
    if new in old:
        return old
    return f"{old}\n{new}"


def _scan_text_signal_score(text: str, mark_hint: str = "", goods_hint: str = "") -> int:
    src = text or ""
    score = 0
    if len(normalize_text(src)) >= 120:
        score += 1
    if any(k in src for k in ROUND_SIGNAL_KEYWORDS):
        score += 1
    if mark_hint and mark_hint in src:
        score += 1
    if goods_hint and any(tok in src for tok in re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,8}", goods_hint)[:8]):
        score += 1
    if extract_dates_advanced(src, None, None):
        score += 1
    return score


def _pick_rescan_candidates(
    files: List[Path],
    scan_texts: Dict[Path, str],
    mark_hint: str,
    goods_hint: str,
    min_score: int,
) -> List[Path]:
    out: List[Path] = []
    for p in files:
        t = scan_texts.get(p, "")
        if _scan_text_signal_score(t, mark_hint=mark_hint, goods_hint=goods_hint) < min_score:
            out.append(p)
    return out


def _rescan_priority(p: Path) -> int:
    name = p.name
    if any(k in name for k in ["通知书", "商标档案", "合同", "发票", "dzfp", "订单", "转账", "银行", "企业详情", "商品详情"]):
        return 0
    if any(k in name for k in ["宣传手册", "展会", "店铺", "商品", "平台"]):
        return 1
    if "微信图片" in name:
        return 2
    return 1


def _limit_rescan_targets(
    targets: List[Path],
    scan_texts: Dict[Path, str],
    mark_hint: str,
    goods_hint: str,
    cap: int,
) -> List[Path]:
    if cap <= 0 or len(targets) <= cap:
        return targets
    scored = []
    for p in targets:
        sig = _scan_text_signal_score(scan_texts.get(p, ""), mark_hint=mark_hint, goods_hint=goods_hint)
        scored.append((_rescan_priority(p), sig, len(p.name), p))
    scored.sort(key=lambda x: (x[0], x[1], x[2]))
    ordered = [x[3] for x in scored]
    non_wechat = [p for p in ordered if "微信图片" not in p.name]
    wechat = [p for p in ordered if "微信图片" in p.name]
    if len(non_wechat) >= cap:
        return non_wechat[:cap]
    return non_wechat + wechat[: max(0, cap - len(non_wechat))]


def _round_entry(
    round_no: int,
    title: str,
    goal: str,
    selected_count: int,
    max_pages: int,
    dpi: int,
    detail: str = "",
) -> Dict[str, Any]:
    return {
        "round": round_no,
        "title": title,
        "goal": goal,
        "selected_count": selected_count,
        "max_pages": max_pages,
        "dpi": dpi,
        "detail": detail,
    }


def run_fixed_five_round_scans(
    files: List[Path],
    ocr_cache: Path,
    normal_pages: int,
    notice_pages: int,
    dpi: int,
    notice_dpi: int,
    prefilled_texts: Optional[Dict[Path, str]] = None,
    scan_mode_by_file: Optional[Dict[Path, str]] = None,
) -> Tuple[Dict[Path, str], List[Path], List[Dict[str, Any]], Dict[str, str]]:
    """
    自适应五轮扫描（流程兼容）：
    R1/R2 必跑；R3~R5 按触发条件执行补扫。
    触发条件示例：CaseInfo关键信息缺失、低信号证据占比较高。
    """
    scan_texts: Dict[Path, str] = {}
    prefilled_texts = prefilled_texts or {}
    scan_mode_by_file = scan_mode_by_file or {}
    resolved_mode_by_file: Dict[Path, str] = {}
    rounds: List[Dict[str, Any]] = []
    if scan_mode_by_file:
        fast_mode = all((scan_mode_by_file.get(p, "auto") == "fast") for p in files)
    else:
        fast_mode = normal_pages > 0 and normal_pages <= 2
    notice_candidates = [p for p in files if any(k in p.name for k in ["通知书", "撤三", "决定", "裁定"])]
    fact_candidates = [p for p in files if any(k in p.name for k in ["商标档案", "商标信息", "注册证", "商标注册"])]
    if not notice_candidates:
        notice_candidates = files[: min(3, len(files))]
    key_notice_or_fact = list(dict.fromkeys(notice_candidates + fact_candidates))

    def _missing_case_keys(case_info: Dict[str, str]) -> List[str]:
        keys = ["reg_no", "class", "mark_name", "respondent", "applicant"]
        return [k for k in keys if not (case_info.get(k, "") or "").strip()]

    def _low_signal_count(cur_mark: str, cur_goods: str, threshold: int) -> int:
        return sum(
            1 for p in files
            if _scan_text_signal_score(scan_texts.get(p, ""), mark_hint=cur_mark, goods_hint=cur_goods) < threshold
        )

    def _append_round(round_no: int, title: str, goal: str, selected: List[Path], pages: int, rdpi: int, detail: str) -> None:
        rounds.append(
            _round_entry(
                round_no,
                title,
                goal,
                selected_count=len(selected),
                max_pages=pages,
                dpi=rdpi,
                detail=detail,
            )
        )

    def _merge_unique_targets(base: List[Path], boost: List[Path]) -> List[Path]:
        return list(dict.fromkeys(base + boost))

    def _read_scan_text(p: Path, max_pages: int, use_dpi: int) -> str:
        pre = clean_ocr_text(prefilled_texts.get(p, ""))
        if pre:
            return pre
        mode = (resolved_mode_by_file.get(p, "") or scan_mode_by_file.get(p, "auto") or "auto").lower()
        if mode == "uncategorized":
            mode = "fast"
        if mode == "mix":
            mode = resolve_mix_scan_mode(p.name, "mix")
        if mode == "fast":
            # 二步判定：先轻扫做内容指纹，命中关键交易/履约特征则强制升级 full。
            probe_text = ocr_pdf_pages(
                p,
                cache_dir=ocr_cache,
                max_pages=min(2, max_pages if max_pages > 0 else 2),
                dpi=min(use_dpi, 260),
            )
            if should_upgrade_to_full_by_content(p.name, probe_text):
                mode = "full"
                resolved_mode_by_file[p] = "full"
            else:
                resolved_mode_by_file[p] = "fast"
        else:
            resolved_mode_by_file[p] = mode
        pages = max_pages
        scan_dpi = use_dpi
        if mode == "fast":
            if pages <= 0:
                pages = 2
            pages = min(pages, 6)
            scan_dpi = min(scan_dpi, 340)
        elif mode == "full":
            # full目录中的证据按轮次页数上限深扫，避免大型合同/票据全页反复扫描导致耗时失控
            if pages <= 0:
                pages = 16
            pages = min(pages, 24)
            scan_dpi = max(min(scan_dpi, 360), 320)
        return ocr_pdf_pages(p, cache_dir=ocr_cache, max_pages=pages, dpi=scan_dpi)

    # Round 1: 案情识别（通知书优先）
    for p in files:
        is_notice = p in notice_candidates
        is_fact = p in fact_candidates
        r1_pages = notice_pages if is_notice else (max(6, notice_pages) if is_fact else 1)
        r1_dpi = notice_dpi if is_notice else max(220, dpi - 40)
        t = _read_scan_text(p, max_pages=r1_pages, use_dpi=r1_dpi)
        scan_texts[p] = _merge_scan_text(scan_texts.get(p, ""), t)
    round1_notice_text = "\n".join([scan_texts.get(p, "") for p in notice_candidates])
    round1_all_text = "\n".join(scan_texts.get(p, "") for p in files)
    round1_case = merge_case_info(infer_case_info(round1_notice_text), infer_case_info(round1_all_text))
    mark_hint = round1_case.get("mark_name", "")
    goods_hint = round1_case.get("goods_hint", "")
    _append_round(
        1,
        "案情识别轮",
        "识别被撤销商标、被撤销商品/服务及指定期间",
        selected=files,
        pages=notice_pages,
        rdpi=notice_dpi,
        detail=f"通知书候选{len(notice_candidates)}份，档案候选{len(fact_candidates)}份，首轮商标={mark_hint or '未识别'}",
    )

    # Round 2: 基线证据扫描
    for p in files:
        is_fact = p in fact_candidates
        p2_pages = max(normal_pages, 4) if is_fact and normal_pages > 0 else normal_pages
        t = _read_scan_text(p, max_pages=p2_pages, use_dpi=dpi)
        scan_texts[p] = _merge_scan_text(scan_texts.get(p, ""), t)
    round2_notice_text = "\n".join([scan_texts.get(p, "") for p in notice_candidates])
    round2_all_text = "\n".join(scan_texts.get(p, "") for p in files)
    round2_case = merge_case_info(infer_case_info(round2_notice_text), infer_case_info(round2_all_text))
    if (round2_case.get("mark_name", "") or "").strip():
        mark_hint = round2_case.get("mark_name", "")
    if (round2_case.get("goods_hint", "") or "").strip():
        goods_hint = round2_case.get("goods_hint", "")
    _append_round(
        2,
        "证据基线轮",
        "抽取每份证据的基础事实、时间线索与关键词",
        selected=files,
        pages=normal_pages,
        rdpi=dpi,
        detail="",
    )

    # 触发闸门：是否需要进入第3轮及以上
    missing_after_r2 = _missing_case_keys(round2_case)
    low_signal_r2 = _low_signal_count(mark_hint, goods_hint, threshold=2)
    low_ratio_r2 = (low_signal_r2 / max(1, len(files)))
    trigger_r3 = bool(missing_after_r2) or (low_ratio_r2 >= (0.55 if fast_mode else 0.45))

    # Round 3: 细节补扫（条件触发）
    r3_targets: List[Path] = []
    r3_pages = max(3, normal_pages * 2 if normal_pages > 0 else 3) if fast_mode else max(6, normal_pages * 2 if normal_pages > 0 else 6)
    r3_dpi = max(dpi, 360)
    if trigger_r3:
        raw_r3 = _pick_rescan_candidates(files, scan_texts, mark_hint, goods_hint, min_score=2)
        raw_r3 = _merge_unique_targets(raw_r3, key_notice_or_fact if missing_after_r2 else [])
        r3_cap = 8 if fast_mode else 24
        r3_targets = _limit_rescan_targets(raw_r3, scan_texts, mark_hint, goods_hint, cap=r3_cap)
        for p in r3_targets:
            t = _read_scan_text(p, max_pages=r3_pages, use_dpi=r3_dpi)
            scan_texts[p] = _merge_scan_text(scan_texts.get(p, ""), t)
    _append_round(
        3,
        "证据细节轮",
        "对低信息量证据补扫详情页，提升单证据作用识别",
        selected=r3_targets,
        pages=r3_pages,
        rdpi=r3_dpi,
        detail=(
            f"触发={trigger_r3}；缺失字段={','.join(missing_after_r2) if missing_after_r2 else '无'}；低信号占比={low_ratio_r2:.1%}"
            if trigger_r3 else
            f"未触发：缺失字段=无 且 低信号占比={low_ratio_r2:.1%}"
        ),
    )

    round3_notice_text = "\n".join([scan_texts.get(p, "") for p in notice_candidates])
    round3_all_text = "\n".join(scan_texts.get(p, "") for p in files)
    round3_case = merge_case_info(infer_case_info(round3_notice_text), infer_case_info(round3_all_text))
    missing_after_r3 = _missing_case_keys(round3_case)
    if (round3_case.get("mark_name", "") or "").strip():
        mark_hint = round3_case.get("mark_name", "")
    if (round3_case.get("goods_hint", "") or "").strip():
        goods_hint = round3_case.get("goods_hint", "")
    low_signal_r3 = _low_signal_count(mark_hint, goods_hint, threshold=3)
    low_ratio_r3 = (low_signal_r3 / max(1, len(files)))
    trigger_r4 = trigger_r3 and (bool(missing_after_r3) or (low_ratio_r3 >= (0.40 if fast_mode else 0.30)))

    # Round 4: 深扫（条件触发）
    r4_targets: List[Path] = []
    r4_pages = max(4, normal_pages * 2 if normal_pages > 0 else 4) if fast_mode else max(10, normal_pages * 3 if normal_pages > 0 else 10)
    r4_dpi = max(dpi, 390)
    if trigger_r4:
        raw_r4 = _pick_rescan_candidates(files, scan_texts, mark_hint, goods_hint, min_score=3)
        raw_r4 = _merge_unique_targets(raw_r4, key_notice_or_fact if missing_after_r3 else [])
        r4_cap = 5 if fast_mode else 12
        r4_targets = _limit_rescan_targets(raw_r4, scan_texts, mark_hint, goods_hint, cap=r4_cap)
        for p in r4_targets:
            t = _read_scan_text(p, max_pages=r4_pages, use_dpi=r4_dpi)
            scan_texts[p] = _merge_scan_text(scan_texts.get(p, ""), t)
    _append_round(
        4,
        "证据深扫轮",
        "对仍未提取充分要素的证据执行深度补扫",
        selected=r4_targets,
        pages=r4_pages,
        rdpi=r4_dpi,
        detail=(
            f"触发={trigger_r4}；缺失字段={','.join(missing_after_r3) if missing_after_r3 else '无'}；低信号占比={low_ratio_r3:.1%}"
            if trigger_r4 else
            f"未触发：缺失字段=无 且 低信号占比={low_ratio_r3:.1%}"
        ),
    )

    round4_notice_text = "\n".join([scan_texts.get(p, "") for p in notice_candidates])
    round4_all_text = "\n".join(scan_texts.get(p, "") for p in files)
    round4_case = merge_case_info(infer_case_info(round4_notice_text), infer_case_info(round4_all_text))
    missing_after_r4 = _missing_case_keys(round4_case)
    if (round4_case.get("mark_name", "") or "").strip():
        mark_hint = round4_case.get("mark_name", "")
    if (round4_case.get("goods_hint", "") or "").strip():
        goods_hint = round4_case.get("goods_hint", "")
    low_signal_r4 = _low_signal_count(mark_hint, goods_hint, threshold=4)
    low_ratio_r4 = (low_signal_r4 / max(1, len(files)))
    trigger_r5 = trigger_r4 and (bool(missing_after_r4) or (low_ratio_r4 >= (0.30 if fast_mode else 0.22)))

    # Round 5: 兜底复扫（条件触发）
    r5_targets: List[Path] = []
    r5_dpi = max(dpi, 420)
    r5_cap_pages = 4 if fast_mode else 16
    if trigger_r5:
        raw_r5 = _pick_rescan_candidates(files, scan_texts, mark_hint, goods_hint, min_score=4)
        raw_r5 = _merge_unique_targets(raw_r5, key_notice_or_fact if missing_after_r4 else [])
        r5_cap = 3 if fast_mode else 6
        r5_targets = _limit_rescan_targets(raw_r5, scan_texts, mark_hint, goods_hint, cap=r5_cap)
        for p in r5_targets:
            total_pages = pdf_page_count(p)
            r5_pages = 0 if (total_pages and total_pages <= r5_cap_pages) else r5_cap_pages
            t = _read_scan_text(p, max_pages=r5_pages, use_dpi=r5_dpi)
            scan_texts[p] = _merge_scan_text(scan_texts.get(p, ""), t)
    _append_round(
        5,
        "收敛复核轮",
        "执行最终兜底扫描并输出五轮合并结果",
        selected=r5_targets,
        pages=16,
        rdpi=r5_dpi,
        detail=(
            f"触发={trigger_r5}；缺失字段={','.join(missing_after_r4) if missing_after_r4 else '无'}；低信号占比={low_ratio_r4:.1%}"
            if trigger_r5 else
            f"未触发：缺失字段=无 且 低信号占比={low_ratio_r4:.1%}"
        ),
    )

    return scan_texts, notice_candidates, rounds, round1_case


def parse_date(y: str, m: str, d: str) -> Optional[dt.date]:
    try:
        return dt.date(int(y), int(m), int(d))
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return None


def extract_dates(text: str) -> List[dt.date]:
    out: List[dt.date] = []
    for y, m, d in DATE_RE.findall(text or ""):
        dd = parse_date(y, m, d)
        if dd:
            out.append(dd)
    return sorted(set(out))


def normalize_ocr_digits(s: str) -> str:
    if not s:
        return s
    trans = str.maketrans({
        "O": "0", "o": "0",
        "I": "1", "l": "1", "|": "1",
        "S": "5", "s": "5",
        "B": "8",
    })
    return s.translate(trans)


def normalize_ocr_date_text(s: str) -> str:
    src = normalize_ocr_digits(s or "")
    src = src.replace("人年", "年").replace("入年", "年").replace("年年", "年")
    src = src.replace("月月", "月").replace("日日", "日")
    src = re.sub(r"(20\d{2})\s*[^0-9年月./-]{0,3}\s*年", r"\1年", src)
    src = re.sub(r"([0-9]{1,2})\s*[^0-9月./-]{0,2}\s*月", r"\1月", src)
    src = re.sub(r"([0-9]{1,2})\s*[^0-9日号./-]{0,2}\s*(日|号)", r"\1日", src)
    src = re.sub(r"\s+", " ", src)
    return src


def infer_year_for_month_day(
    month: int,
    day: int,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Optional[int]:
    if month < 1 or month > 12 or day < 1 or day > 31:
        return None
    if period_start and period_end and period_start <= period_end:
        years = list(range(period_start.year - 1, period_end.year + 2))
        in_period: List[dt.date] = []
        near_period: List[dt.date] = []
        for y in years:
            try:
                d = dt.date(y, month, day)
            except Exception as exc:
                _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
                continue
            if period_start <= d <= period_end:
                in_period.append(d)
            elif (period_start - dt.timedelta(days=365)) <= d <= (period_end + dt.timedelta(days=365)):
                near_period.append(d)
        if in_period:
            return sorted(in_period)[-1].year
        if near_period:
            return sorted(near_period)[-1].year
    return dt.date.today().year


def extract_dates_advanced(
    text: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> List[dt.date]:
    src = text or ""
    norm = normalize_ocr_date_text(src)
    out: List[dt.date] = []

    out.extend(extract_dates(norm))

    for y, m, d in DATE_FUZZY_RE.findall(norm):
        dd = parse_date(normalize_ocr_digits(y), normalize_ocr_digits(m), normalize_ocr_digits(d))
        if dd:
            out.append(dd)

    for y, m, d in DATE_NUMERIC_RE.findall(norm):
        dd = parse_date(y, m, d)
        if dd:
            out.append(dd)

    for y, m, d in COMPACT_DATE_RE.findall(norm):
        dd = parse_date(y, m, d)
        if dd:
            out.append(dd)

    for m, d in MONTH_DAY_RE.findall(norm):
        mm = normalize_ocr_digits(m)
        dd = normalize_ocr_digits(d)
        y = infer_year_for_month_day(int(mm), int(dd), period_start, period_end) if mm.isdigit() and dd.isdigit() else None
        if y:
            d_obj = parse_date(str(y), mm, dd)
            if d_obj:
                out.append(d_obj)

    return sorted(set(out))


def extract_review_dates(
    text: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Tuple[List[dt.date], str]:
    dates, note, _ = extract_keyword_scene_dates(
        text=text,
        period_start=period_start,
        period_end=period_end,
        keywords=REVIEW_KEYWORDS,
        scene_name="评价场景",
    )
    return dates, note


def extract_keyword_scene_dates(
    text: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
    keywords: List[str],
    scene_name: str,
) -> Tuple[List[dt.date], str, List[str]]:
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    scored: List[Tuple[int, str]] = []
    # 1) 单行命中：关键词 + 日期
    for ln in lines:
        nln = normalize_ocr_date_text(ln)
        has_kw = any(k in ln for k in keywords)
        has_date = bool(re.search(r"20\d{2}", nln) or DATE_FUZZY_RE.search(nln) or MONTH_DAY_RE.search(nln))
        if not has_kw:
            continue
        s = 0
        s += 4
        if re.search(r"20\d{2}", nln):
            s += 4
        if MONTH_DAY_RE.search(nln):
            s += 2
        if has_kw and has_date:
            s += 1
        if any(k in ln for k in UNCERTAIN_TIME_KEYWORDS):
            s -= 2
        if s > 0 and has_date:
            scored.append((s, ln))

    # 2) 邻近窗口命中：关键词行前后1行联动，解决“评价词和日期分行”的 OCR 场景
    for i, ln in enumerate(lines):
        if not any(k in ln for k in keywords):
            continue
        window = lines[max(0, i - 1): min(len(lines), i + 2)]
        merged = " ".join(window)
        nmerged = normalize_ocr_date_text(merged)
        has_date = bool(re.search(r"20\d{2}", nmerged) or DATE_FUZZY_RE.search(nmerged) or MONTH_DAY_RE.search(nmerged))
        if not has_date:
            continue
        s = 6
        if re.search(r"20\d{2}", nmerged):
            s += 2
        if any(k in merged for k in UNCERTAIN_TIME_KEYWORDS):
            s -= 1
        if s > 0:
            scored.append((s, merged))

    scored.sort(key=lambda x: (-x[0], len(x[1])))
    picked = [ln for _, ln in scored[:12]]
    picked_text = "\n".join(picked)
    dates = extract_dates_advanced(picked_text, period_start, period_end) if picked else []
    compact_dates = extract_compact_review_like_dates(picked_text)
    if compact_dates:
        dates = sorted(set(dates + compact_dates))
    if dates:
        return dates, f"已启用{scene_name}时间识别", picked
    return [], "", []


def extract_compact_review_like_dates(text: str) -> List[dt.date]:
    """
    处理 OCR 中常见的“2023 2218”弱格式日期：
    - 在评价/评论上下文中，尝试推断为 2023-2-18 或 2023-2-21 等有效日期。
    """
    out: List[dt.date] = []
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    for ln in lines:
        if not any(k in ln for k in REVIEW_KEYWORDS + ["评论", "追评", "首次评价"]):
            continue
        nln = normalize_ocr_date_text(ln)
        for y, token in re.findall(r"(20\d{2})\s+(\d{3,4})", nln):
            cands: List[Tuple[str, str]] = []
            if len(token) == 3:
                cands.append((token[0], token[1:]))
            elif len(token) == 4:
                cands.append((token[:2], token[2:]))
                cands.append((token[0], token[1:3]))
                cands.append((token[0], token[-2:]))
            for m, d in cands:
                if not (m.isdigit() and d.isdigit()):
                    continue
                mm = int(m)
                dd = int(d)
                if 1 <= mm <= 12 and 1 <= dd <= 31:
                    dt_obj = parse_date(y, str(mm), str(dd))
                    if dt_obj:
                        out.append(dt_obj)
                        break
    return sorted(set(out))


def extract_order_dates(
    text: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Tuple[List[dt.date], str, List[str]]:
    return extract_keyword_scene_dates(
        text=text,
        period_start=period_start,
        period_end=period_end,
        keywords=ORDER_KEYWORDS,
        scene_name="交易场景",
    )


def extract_logistics_dates(
    text: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Tuple[List[dt.date], str, List[str]]:
    return extract_keyword_scene_dates(
        text=text,
        period_start=period_start,
        period_end=period_end,
        keywords=LOGISTICS_KEYWORDS,
        scene_name="物流场景",
    )


def resolve_time_conflicts(
    base_dates: List[dt.date],
    review_dates: List[dt.date],
    order_dates: List[dt.date],
    logistics_dates: List[dt.date],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Tuple[List[dt.date], str, List[str]]:
    """
    多通道时间冲突裁决：
    - 通道优先级：评价 > 交易 > 物流 > 基础日期
    - 若同证据同时出现期内与期外时点，优先采用期内时点；
    - 输出通道说明用于审计追踪。
    """
    channel_weights = {
        "评价": 6,
        "交易": 5,
        "物流": 5,
        "基础": 2,
    }
    weighted: Dict[dt.date, int] = {}
    used_channels: Dict[dt.date, set] = {}

    def add_dates(ds: List[dt.date], ch: str):
        w = channel_weights.get(ch, 1)
        for d in ds:
            score = w
            if period_start and period_end:
                if period_start <= d <= period_end:
                    score += 2
                else:
                    score -= 1
            weighted[d] = weighted.get(d, 0) + score
            used_channels.setdefault(d, set()).add(ch)

    add_dates(base_dates, "基础")
    add_dates(review_dates, "评价")
    add_dates(order_dates, "交易")
    add_dates(logistics_dates, "物流")

    if not weighted:
        return [], "", []

    # 优先保留高分时点；若为空则退化为全部时点
    preferred = [d for d, s in weighted.items() if s >= 5]
    if not preferred:
        preferred = [d for d, s in weighted.items() if s > 0]
    if not preferred:
        preferred = list(weighted.keys())

    preferred = sorted(set(preferred))
    note_parts: List[str] = []

    channels_used = set()
    for d in preferred:
        channels_used.update(used_channels.get(d, set()))
    channel_txt = "、".join([x for x in ["评价", "交易", "物流", "基础"] if x in channels_used])
    if channel_txt:
        note_parts.append(f"多通道识别命中：{channel_txt}")

    if period_start and period_end:
        in_period = [d for d in preferred if period_start <= d <= period_end]
        out_period = [d for d in preferred if d < period_start or d > period_end]
        if in_period and out_period:
            preferred = in_period
            note_parts.append("检测到期内/期外冲突，已优先采用期内时点")

    return sorted(set(preferred)), "；".join(note_parts), sorted(channels_used)


def build_time_anchor_candidates(
    base_dates: List[dt.date],
    review_dates: List[dt.date],
    order_dates: List[dt.date],
    logistics_dates: List[dt.date],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
    top_k: int = 5,
    context_text: str = "",
    evidence_kind: str = "",
) -> Dict[str, Any]:
    channel_priority = {"评价": 4, "交易": 3, "物流": 2, "基础": 1}
    channel_weight = {"评价": 8, "交易": 7, "物流": 6, "基础": 3}
    date_channels: Dict[dt.date, set] = {}
    date_score: Dict[dt.date, int] = {}
    date_type_hint: Dict[dt.date, str] = {}
    date_excluded: Dict[dt.date, bool] = {}

    # 构建“日期 -> 上下文类型”映射，区分交易时间/系统时间/文案时间/不确定时间。
    lines = [ln.strip() for ln in (context_text or "").splitlines() if ln.strip()]
    kind_noise = TIME_NOISE_EXCLUDE_BY_KIND.get(str(evidence_kind or "").strip(), [])
    for ln in lines:
        nln = normalize_ocr_date_text(ln)
        ds = extract_dates_advanced(nln, period_start, period_end)
        if not ds:
            continue
        is_noise = any(k in nln for k in TIME_NOISE_EXCLUDE_KEYWORDS) or any(k in nln for k in kind_noise)
        if any(k in nln for k in ["体检", "培训", "放假", "节假日", "活动预告", "报名截止"]):
            is_noise = True
        is_tx = any(k in nln for k in TIME_TRANSACTION_HINT_KEYWORDS)
        is_sys = any(k in nln for k in TIME_SYSTEM_HINT_KEYWORDS)
        if is_tx:
            hint = "transaction_date"
        elif is_sys:
            hint = "system_generated"
        elif is_noise:
            hint = "content_claimed"
        else:
            hint = "uncertain"
        for d in ds:
            date_type_hint[d] = hint
            if is_noise:
                date_excluded[d] = True

    def add(ds: List[dt.date], ch: str) -> None:
        w = channel_weight.get(ch, 1)
        for d in ds:
            date_channels.setdefault(d, set()).add(ch)
            score = w
            if period_start and period_end:
                if period_start <= d <= period_end:
                    score += 3
                else:
                    score -= 2
            # 上下文增强/抑制
            hint = date_type_hint.get(d, "")
            if hint == "transaction_date":
                score += 4
            elif hint == "system_generated":
                score += 2
            elif hint == "content_claimed":
                score -= 4
            elif hint == "uncertain":
                score -= 1
            if date_excluded.get(d, False):
                score -= 6
            if str(evidence_kind or "").strip() == "程序文件":
                score -= 6
            date_score[d] = date_score.get(d, 0) + score

    add(base_dates, "基础")
    add(review_dates, "评价")
    add(order_dates, "交易")
    add(logistics_dates, "物流")

    if not date_score:
        return {
            "selected_dates": [],
            "channels_used": [],
            "selected_anchor_type": "uncertain",
            "selected_anchor_source": "基础",
            "note": "",
            "conflict_flag": False,
            "selection_rule": "无候选日期",
            "candidates": [],
        }

    cand_rows: List[Dict[str, Any]] = []
    for d, score in date_score.items():
        channels = sorted(list(date_channels.get(d, set())), key=lambda x: -channel_priority.get(x, 0))
        in_period = False
        if period_start and period_end:
            in_period = period_start <= d <= period_end
        channel_best = max([channel_priority.get(c, 0) for c in channels], default=0)
        if "交易" in channels or "物流" in channels:
            anchor_type = "transaction_date"
        elif "评价" in channels:
            anchor_type = "system_generated"
        else:
            anchor_type = date_type_hint.get(d, "uncertain")
        anchor_source = channels[0] if channels else "基础"
        conf = "LOW"
        if score >= 11:
            conf = "HIGH"
        elif score >= 7:
            conf = "MEDIUM"
        cand_rows.append(
            {
                "date": d.isoformat(),
                "channels": channels,
                "score": score,
                "in_period": in_period,
                "channel_best": channel_best,
                "confidence": conf,
                "anchor_type": anchor_type,
                "anchor_source": anchor_source,
                "excluded": bool(date_excluded.get(d, False)),
            }
        )

    cand_rows.sort(
        key=lambda x: (
            0 if x["in_period"] else 1,
            -x["channel_best"],
            -x["score"],
            x["date"],
        )
    )
    top = cand_rows[: max(1, top_k)]
    conflict_flag = any(c["in_period"] for c in top) and any(not c["in_period"] for c in top)

    in_period_top = [c for c in top if c["in_period"]]
    selected = in_period_top if in_period_top else top
    selected_dates = sorted(set([x["date"] for x in selected]))
    channels_used = sorted(
        set([ch for row in selected for ch in row["channels"]]),
        key=lambda x: -channel_priority.get(x, 0),
    )
    anchor_types = sorted(
        set([str(row.get("anchor_type", "uncertain")) for row in selected]),
        key=lambda x: -TIME_ANCHOR_CHANNEL_PRIORITY.get(x, 0),
    )
    anchor_sources = [str(row.get("anchor_source", "基础") or "基础") for row in selected]
    selected_anchor_source = anchor_sources[0] if anchor_sources else "基础"

    note_parts = []
    if channels_used:
        note_parts.append(f"候选通道：{'、'.join(channels_used)}")
    if conflict_flag and in_period_top:
        note_parts.append("检测到期内/期外冲突，已优先采用期内候选")
    if any(row.get("excluded", False) for row in top):
        note_parts.append("已对活动/检验/有效期等噪声日期降权")
    if anchor_types:
        note_parts.append(f"时间来源类型：{'、'.join(anchor_types)}")
    if selected_anchor_source:
        note_parts.append(f"时间来源通道：{selected_anchor_source}")
    selection_rule = "通道优先(评价>交易>物流>基础)+期内优先+噪声降权+分值排序"

    return {
        "selected_dates": selected_dates,
        "channels_used": channels_used,
        "anchor_types": anchor_types,
        "selected_anchor_type": anchor_types[0] if anchor_types else "uncertain",
        "selected_anchor_source": selected_anchor_source,
        "note": "；".join(note_parts),
        "conflict_flag": conflict_flag,
        "selection_rule": selection_rule,
        "candidates": top,
    }


def to_iso(d: dt.date) -> str:
    return d.isoformat()


def parse_iso_date(s: str) -> Optional[dt.date]:
    src = (str(s) if s is not None else "").strip()
    if not src:
        return None
    try:
        return dt.datetime.strptime(src, "%Y-%m-%d").date()
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
    m = re.search(r"(20\d{2}-\d{1,2}-\d{1,2})", src)
    if m:
        try:
            return dt.datetime.strptime(m.group(1), "%Y-%m-%d").date()
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
            return None
    return None


def clean_abnormal_dates(
    dates: List[dt.date],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Tuple[List[dt.date], str]:
    """
    异常日期清洗规则：
    - 若同一证据提取日期跨度 > 8 年，按“指定期间±365天”过滤；
    - 过滤后为空则置为空（后续输出为“时间不明”）。
    """
    if not dates:
        return [], ""

    ds = sorted(set(dates))
    today = dt.date.today()
    ds = [d for d in ds if d <= (today + dt.timedelta(days=31))]
    if not ds:
        return [], "日期均超出当前时间范围，已清洗为时间不明"
    if len(ds) <= 1:
        return ds, ""

    span_days = (ds[-1] - ds[0]).days
    if span_days <= 365 * 8:
        return ds, ""

    if period_start and period_end:
        lo = period_start - dt.timedelta(days=365)
        hi = period_end + dt.timedelta(days=365)
        in_window = [d for d in ds if lo <= d <= hi]
        if in_window:
            return in_window, "日期跨度异常（>8年），已按指定期间±365天清洗"
        return [], "日期跨度异常（>8年）且不在指定期间窗口，已清洗为时间不明"

    return ds, ""


def _date_overlap(
    s1: Optional[dt.date],
    e1: Optional[dt.date],
    s2: Optional[dt.date],
    e2: Optional[dt.date],
) -> bool:
    if not s1 or not e1 or not s2 or not e2:
        return False
    return max(s1, s2) <= min(e1, e2)


def _segment_rank(
    start: Optional[dt.date],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> int:
    if not start or not period_start or not period_end or period_start > period_end:
        return 9

    def _add_year_safe(d: dt.date) -> dt.date:
        try:
            return d.replace(year=d.year + 1)
        except Exception:
            # 处理 2 月 29 日
            return d.replace(year=d.year + 1, day=28)

    p1e = min(_add_year_safe(period_start) - dt.timedelta(days=1), period_end)
    p2s = p1e + dt.timedelta(days=1)
    if p2s > period_end:
        return 1
    p2e = min(_add_year_safe(p2s) - dt.timedelta(days=1), period_end)
    if start <= p1e:
        return 1
    if start <= p2e:
        return 2
    return 3


def infer_trade_link_role(name: str, text: str) -> str:
    name = str(name or "")
    src = f"{name}\n{text}"
    name_l = name.lower()
    src_l = src.lower()
    invoice_name_hit = any(k in name for k in INVOICE_KEYWORDS) or any(
        k in name_l for k in ["invoice", "fapiao", "dzfp", "fp_", "发票"]
    )
    contract_name_hit = any(k in name for k in CONTRACT_KEYWORDS)
    invoice_score = sum(2 for k in INVOICE_KEYWORDS if k in src) + sum(
        1 for k in ["发票代码", "发票号码", "开票日期", "价税合计", "税率", "税额", "校验码"] if (k in src or k in src_l)
    )
    contract_score = sum(2 for k in CONTRACT_KEYWORDS if k in src) + sum(
        1 for k in ["合同编号", "甲方", "乙方", "签订日期", "协议"] if (k in src or k in src_l)
    )
    # 文件名命中优先，避免“合同及发票”目录导致发票被误判为合同。
    if invoice_name_hit and not contract_name_hit:
        return "发票"
    if contract_name_hit and not invoice_name_hit:
        return "合同"
    if invoice_score > contract_score:
        return "发票"
    if contract_score > invoice_score:
        return "合同"
    if any(k in src for k in ["回单", "电子回单", "银行回单"]):
        return "回单"
    if any(k in src for k in PAYMENT_KEYWORDS):
        return "支付"
    if any(k in src for k in LOGISTICS_KEYWORDS):
        return "物流"
    return "交易"


def normalize_amount_str(amount: str) -> str:
    raw = str(amount or "").strip().replace(",", "")
    if not raw:
        return ""
    try:
        v = float(raw)
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""
    if v <= 0:
        return ""
    return f"{v:.2f}"


def display_amount(amount: str) -> str:
    n = normalize_amount_str(amount)
    if not n:
        return ""
    try:
        v = float(n)
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return n
    if abs(v - int(v)) < 1e-9:
        return str(int(v))
    return f"{v:.2f}".rstrip("0").rstrip(".")


def normalize_company_name(name: str) -> str:
    s = sanitize_component(str(name or ""), 32)
    if not s:
        return ""
    s = re.sub(r"^[0-9一二三四五六七八九十年月日_\\.-、]+", "", s)
    s = s.strip("_-")
    return sanitize_component(s, 24)


def extract_trade_amount(text: str, key_lines: List[str]) -> str:
    src = "\n".join([text or "", *(key_lines or [])])
    cands: List[Tuple[int, float, str]] = []

    def push(raw: str, score: int) -> None:
        n = normalize_amount_str(raw)
        if not n:
            return
        try:
            v = float(n)
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
            return
        cands.append((score, v, n))

    for m in AMOUNT_KEYWORD_RE.finditer(src):
        ctx = src[max(0, m.start() - 14): min(len(src), m.end() + 14)]
        score = 120
        if any(k in ctx for k in ["合同金额", "价税合计", "开票金额", "付款金额", "实付金额"]):
            score += 25
        push(m.group(1), score)

    for m in AMOUNT_CURRENCY_RE.finditer(src):
        ctx = src[max(0, m.start() - 10): min(len(src), m.end() + 10)]
        score = 85
        if any(k in ctx for k in ["金额", "合同", "发票", "支付", "转账"]):
            score += 15
        push(m.group(1), score)

    for m in AMOUNT_YUAN_RE.finditer(src):
        ctx = src[max(0, m.start() - 10): min(len(src), m.end() + 10)]
        score = 70
        if any(k in ctx for k in ["人民币", "金额", "合同", "发票", "支付", "价税"]):
            score += 20
        push(m.group(1), score)

    if not cands:
        return ""
    if any(v >= 100 for _, v, _ in cands):
        cands = [x for x in cands if x[1] >= 10]
    cands.sort(key=lambda x: (x[0], x[1]), reverse=True)
    return cands[0][2]
    return ""


def extract_trade_counterparty(text: str, key_lines: List[str]) -> str:
    src = "\n".join([text or "", *(key_lines or [])])
    role_hits: List[str] = []
    for ln in src.splitlines():
        compact = re.sub(r"\s+", "", ln)
        if not compact:
            continue
        if any(k in compact for k in TRADE_PARTY_HINT_KEYWORDS):
            for c in COMPANY_RE.findall(compact):
                role_hits.append(c)
    if role_hits:
        uniq = sorted(set(role_hits), key=lambda x: (len(x), x), reverse=True)
        return normalize_company_name(uniq[0])

    companies = COMPANY_RE.findall(src)
    if not companies:
        return ""
    uniq = sorted(set(companies), key=lambda x: (len(x), x), reverse=True)
    return normalize_company_name(uniq[0])


def build_product_match_key(
    product_label: str,
    inferred_name: str,
    text: str,
    key_lines: List[str],
) -> str:
    raw = " ".join([
        product_label or "",
        inferred_name or "",
        " ".join((key_lines or [])[:5]),
        (text or "")[:800],
    ])
    tokens = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9·-]{2,24}", raw)
    stop = {
        "网店", "商品", "页面", "详情", "店铺", "材料", "证据", "合同", "发票", "交易",
        "宣传册", "画册", "检测报告", "检验报告", "程序", "通知书", "主体", "其他材料",
    }
    picked: List[str] = []
    for t in tokens:
        tt = t.strip("·- ")
        if (not tt) or (tt in stop):
            continue
        if any(s in tt for s in ["时间不明", "待核验", "关键语句"]):
            continue
        if len(tt) < 2:
            continue
        picked.append(tt)
    if not picked:
        return ""
    picked = sorted(set(picked), key=lambda x: (len(x), x), reverse=True)
    return sanitize_component("_".join(picked[:2]), 24)


def product_overlap_score(a: str, b: str) -> int:
    aa = str(a or "").strip()
    bb = str(b or "").strip()
    if not aa or not bb:
        return 0
    if aa == bb:
        return 4
    if len(aa) >= 4 and aa in bb:
        return 3
    if len(bb) >= 4 and bb in aa:
        return 3
    ta = set(re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,16}", aa))
    tb = set(re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,16}", bb))
    inter = ta & tb
    if not inter:
        return 0
    score = 1
    if len(inter) >= 2:
        score += 1
    if max(len(x) for x in inter) >= 4:
        score += 1
    return score


def extract_trade_pair_key(
    text: str,
    key_lines: List[str],
    date_start: str,
    date_end: str,
    fallback_name: str,
    amount_hint: str = "",
    party_hint: str = "",
    product_hint: str = "",
) -> str:
    amount = normalize_amount_str(amount_hint) or extract_trade_amount(text, key_lines)
    party = sanitize_component(party_hint, 24) if party_hint else extract_trade_counterparty(text, key_lines)

    date_month = ""
    d0 = date_start or date_end
    if d0:
        try:
            date_month = dt.datetime.strptime(d0, "%Y-%m-%d").strftime("%Y-%m")
        except Exception:
            date_month = d0[:7]

    topic = sanitize_component(product_hint or fallback_name, 18) if (product_hint or fallback_name) else ""
    parts = [x for x in [party, amount, topic, date_month] if x]
    if not parts:
        return ""
    # 对“合同+发票金额对应”优先稳定按：主体|金额|商品
    return "|".join(parts[:3]) if len(parts) >= 3 else "|".join(parts)


def assign_trade_chain_groups(rows: List[Dict[str, str]]) -> None:
    key_to_group: Dict[str, int] = {}
    group_meta: Dict[int, Dict[str, Any]] = {}
    seq = 1

    def ensure_group(raw_key: str) -> int:
        nonlocal seq
        if raw_key in key_to_group:
            return key_to_group[raw_key]
        no = seq
        seq += 1
        key_to_group[raw_key] = no
        group_meta[no] = {
            "roles": set(),
            "amounts": set(),
            "products": set(),
            "dates": [],
            "tx_idx": [],
            "product_idx": [],
        }
        return no

    # 预扫金额分布：若同一金额同时出现“合同+发票”，优先按金额聚合，增强对应关系稳定性。
    amount_roles: Dict[str, set] = {}
    for r in rows:
        if r.get("Type", "") != "交易凭证":
            continue
        role0 = str(r.get("Trade Link Role", "") or "").strip() or infer_trade_link_role(
            str(r.get("_src_name", "") or r.get("Evidence Name", "")),
            str(r.get("Time Anchor Text", "") or ""),
        )
        src_for_amount0 = "\n".join([
            str(r.get("Original File Name", "") or ""),
            str(r.get("_src_name", "") or ""),
            str(r.get("Evidence Name", "") or ""),
            str(r.get("Time Anchor Text", "") or ""),
            str(r.get("Risk Notes", "") or ""),
        ])
        amount0 = normalize_amount_str(r.get("Trade Amount", "")) or extract_trade_amount(
            src_for_amount0,
            [],
        )
        if not amount0:
            continue
        amount_roles.setdefault(amount0, set()).add(role0)
    paired_amounts = {
        amt for amt, roles in amount_roles.items()
        if ("合同" in roles and "发票" in roles)
    }

    def company_from_filename(filename: str) -> str:
        comps = COMPANY_RE.findall(str(filename or ""))
        if not comps:
            return ""
        uniq = sorted(set(comps), key=lambda x: (len(x), x), reverse=True)
        return normalize_company_name(uniq[0])

    # 1) 先对交易凭证分组，优先使用“主体+金额”形成合同/发票对应组。
    for idx, r in enumerate(rows):
        if r.get("Type", "") != "交易凭证":
            r["Trade Chain Group"] = ""
            r["Trade Chain Group No"] = ""
            r["Trade Link Role"] = r.get("Trade Link Role", "")
            continue

        role = str(r.get("Trade Link Role", "") or "").strip() or infer_trade_link_role(
            str(r.get("_src_name", "") or r.get("Evidence Name", "")),
            str(r.get("Time Anchor Text", "") or ""),
        )
        src_for_amount = "\n".join([
            str(r.get("Original File Name", "") or ""),
            str(r.get("_src_name", "") or ""),
            str(r.get("Evidence Name", "") or ""),
            str(r.get("Time Anchor Text", "") or ""),
            str(r.get("Risk Notes", "") or ""),
        ])
        amount = normalize_amount_str(r.get("Trade Amount", "")) or extract_trade_amount(
            src_for_amount,
            [],
        )
        party = normalize_company_name(str(r.get("Trade Counterparty", "") or "")) or extract_trade_counterparty(
            str(r.get("Time Anchor Text", "") or "") + "\n" + str(r.get("Risk Notes", "") or ""),
            [],
        )
        party_by_name = company_from_filename(
            str(r.get("Original File Name", "") or r.get("_src_name", "") or r.get("Evidence Name", "") or "")
        )
        if party_by_name and party_by_name != party:
            party = party_by_name
        product_key = str(r.get("Product Key", "") or "").strip()

        ds = parse_iso_date(str(r.get("Evidence Date Start", "") or "")) or parse_iso_date(str(r.get("Evidence Date End", "") or ""))
        month = ds.strftime("%Y-%m") if ds else "NA"
        year = ds.strftime("%Y") if ds else "NA"

        # 每次重建组证键，避免沿用旧键导致跨案或跨批次串组。
        raw_key = ""
        if amount and amount in paired_amounts:
            raw_key = f"amt_pair|{amount}"
        elif role in ("合同", "发票") and party:
            raw_key = f"party|{party}|{year}"
        elif amount and party:
            raw_key = f"amt2|{party}|{amount}"
        elif amount:
            raw_key = f"amt3|{amount}"
        elif product_key:
            raw_key = f"prod|{product_key}|{month}"
        if not raw_key:
            raw_key = f"fallback|{month}|{sanitize_component(r.get('Evidence Name',''), 12)}"

        no = ensure_group(raw_key)
        r["Trade Link Role"] = role
        r["Trade Amount"] = amount
        r["Trade Counterparty"] = party
        r["Trade Link Key"] = raw_key
        r["Trade Chain Group No"] = str(no)
        r["Trade Chain Group"] = f"组证{no}"

        meta = group_meta[no]
        meta["roles"].add(role)
        if amount:
            meta["amounts"].add(amount)
        if product_key:
            meta["products"].add(product_key)
        if ds:
            meta["dates"].append(ds)
        meta["tx_idx"].append(idx)

    # 2) 商品/展示材料尝试对印到交易组（合同+发票+商品三联动）。
    for idx, r in enumerate(rows):
        if r.get("Type", "") not in ("商品展示页", "线上店铺展示", "其他材料"):
            continue
        if r.get("Trade Chain Group", ""):
            continue
        product_key = str(r.get("Product Key", "") or "").strip()
        if not product_key:
            continue

        ds = parse_iso_date(str(r.get("Evidence Date Start", "") or "")) or parse_iso_date(str(r.get("Evidence Date End", "") or ""))
        best_no = 0
        best_score = 0
        for no, meta in group_meta.items():
            score = 0
            for gp in meta.get("products", set()):
                score = max(score, product_overlap_score(product_key, str(gp)))
            if score <= 0:
                continue
            dates = meta.get("dates", [])
            if ds and dates:
                nearest = min(abs((ds - d).days) for d in dates)
                if nearest <= 180:
                    score += 2
                elif nearest <= 365:
                    score += 1
            if score > best_score:
                best_score = score
                best_no = no

        if best_no and best_score >= 2:
            r["Trade Chain Group No"] = str(best_no)
            r["Trade Chain Group"] = f"组证{best_no}"
            r["Trade Link Role"] = "商品对印"
            group_meta[best_no]["products"].add(product_key)
            group_meta[best_no]["product_idx"].append(idx)

    # 3) 生成组证总结并回填到输出字段，便于目录与风险报告展示。
    for no, meta in group_meta.items():
        roles = set(meta.get("roles", set()))
        amounts = sorted(meta.get("amounts", set()))
        has_contract = "合同" in roles
        has_invoice = "发票" in roles
        amount_note = ""
        if has_contract and has_invoice:
            if len(amounts) == 1:
                amount_note = f"合同/发票金额已对应（{display_amount(amounts[0])}元）"
            elif len(amounts) > 1:
                nums = "、".join([f"{display_amount(x)}元" for x in amounts[:3]])
                amount_note = f"合同/发票金额待核验（{nums}）"
            else:
                amount_note = "合同/发票金额待补录"
        elif has_contract or has_invoice:
            amount_note = "合同/发票材料已入组，待补充配对凭证"
        product_cnt = len(meta.get("product_idx", []))
        if product_cnt > 0:
            amount_note = (amount_note + f"；商品对印{product_cnt}份").strip("；")
        meta["summary"] = amount_note

    def append_unique(base: str, extra: str) -> str:
        b = str(base or "").strip("；")
        e = str(extra or "").strip("；")
        if not e:
            return b
        if e in b:
            return b
        return (b + "；" + e).strip("；")

    for r in rows:
        grp = str(r.get("Trade Chain Group", "") or "").strip()
        if not grp:
            continue
        try:
            no = int(str(r.get("Trade Chain Group No", "") or "0"))
        except Exception:
            no = 0
        meta = group_meta.get(no, {})
        role = str(r.get("Trade Link Role", "") or "关联").strip()
        grp_note = f"组证关联：{grp}（{role}）"
        amount = display_amount(str(r.get("Trade Amount", "") or ""))
        if amount and role in ("合同", "发票", "支付", "交易"):
            grp_note = f"{grp_note}；金额：{amount}元"
        summary = str(meta.get("summary", "") or "").strip()
        if summary:
            grp_note = f"{grp_note}；{summary}"
        r["Risk Notes"] = append_unique(r.get("Risk Notes", ""), grp_note)
        if role in ("合同", "发票", "支付", "交易", "商品对印"):
            r["Inferred Purpose"] = append_unique(r.get("Inferred Purpose", ""), grp_note)


def detect_unit_type(row: Dict[str, str]) -> str:
    et = str(row.get("Type", "") or "")
    form = str(row.get("Form Type", "") or "")
    role = str(row.get("Trade Link Role", "") or "").strip()
    src = " ".join(
        [
            str(row.get("Inferred Proof Name", "") or ""),
            str(row.get("Inferred Purpose", "") or ""),
            str(row.get("Risk Notes", "") or ""),
            str(row.get("_origin_rel", "") or row.get("Original Relative Path", "") or ""),
            str(row.get("_src_name", "") or row.get("Original File Name", "") or ""),
        ]
    )
    if et == "程序文件":
        return "流程文件"
    if et == "交易凭证":
        if role == "合同":
            return "合同"
        if role == "发票":
            return "发票"
        if role in ("支付", "交易", "回单", "商品对印"):
            return "交易履约"
        invoice_hit = (
            form == "发票"
            or any(k in src for k in ["发票", "增值税", "开票", "价税", "税额", "税率", "发票代码", "发票号码"])
        )
        contract_hit = (
            form == "合同"
            or any(k in src for k in ["合同", "协议", "租赁", "采购", "买卖", "甲方", "乙方", "签订"])
        )
        if invoice_hit and not contract_hit:
            return "发票"
        if contract_hit and not invoice_hit:
            return "合同"
        if invoice_hit and contract_hit:
            # 同时命中时，优先由开票字段判定为发票。
            if any(k in src for k in ["发票代码", "发票号码", "税额", "价税"]):
                return "发票"
            return "合同"
        return "交易履约"
    if form in ("现场照片", "包装"):
        return "图片"
    if any(k in src for k in ["实拍", "现场照片", "门头", "厂房", "仓储", "包装", "产品图", "产品实物", "商标照片"]):
        return "图片"
    if form == "海报" or any(k in src for k in ["宣传", "海报", "易拉宝", "画册", "招募", "邀请函"]):
        return "宣传单"
    if et in ("商品展示页", "线上店铺展示"):
        return "图片"
    if et == "资质主体证明":
        return "主体资质"
    return "其他佐证"


EVIDENCE_GROUP_ORDER = {
    "流程文件组": 1,
    "主体资质组": 2,
    "实物图组": 3,
    "合同发票组": 4,
    "宣传单组": 5,
    "其他佐证组": 6,
}

TRADE_ROLE_ORDER = {
    "合同": 1,
    "发票": 2,
    "支付": 3,
    "交易": 4,
    "回单": 5,
    "商品对印": 6,
    "关联": 9,
}


def detect_evidence_group(row: Dict[str, str]) -> str:
    unit = str(row.get("Unit Type", "") or detect_unit_type(row))
    if unit == "流程文件":
        return "流程文件组"
    if unit == "主体资质":
        return "主体资质组"
    if unit == "图片":
        return "实物图组"
    if unit in ("合同", "发票", "交易履约"):
        return "合同发票组"
    if unit == "宣传单":
        return "宣传单组"
    return "其他佐证组"


def _merge_field_values(rows: List[Dict[str, str]], key: str, limit: int = 3) -> List[str]:
    out: List[str] = []
    seen = set()
    for r in rows:
        v = str(r.get(key, "") or "").strip()
        if not v or v in seen:
            continue
        seen.add(v)
        out.append(v)
        if len(out) >= limit:
            break
    return out


def _merge_dates_for_rows(rows: List[Dict[str, str]]) -> Tuple[str, str]:
    ds: List[dt.date] = []
    de: List[dt.date] = []
    for r in rows:
        d1 = parse_iso_date(str(r.get("Evidence Date Start", "") or ""))
        d2 = parse_iso_date(str(r.get("Evidence Date End", "") or ""))
        if d1:
            ds.append(d1)
        if d2:
            de.append(d2)
    if not ds and not de:
        return "", ""
    all_dates = sorted(set(ds + de))
    return all_dates[0].isoformat(), all_dates[-1].isoformat()


def _is_noisy_topic_token(token: str) -> bool:
    t = (token or "").strip()
    if not t:
        return True
    if any(b in t for b in EVIDENCE_NAME_BAD_PHRASES):
        return True
    if re.search(r"[0-9a-fA-F]{12,}", t):
        return True
    if any(k in t for k in ["请各", "现对", "本年度对公司", "体检卡号", "注意人身安全", "工作岗位"]):
        return True
    t_norm = normalize_text(t)
    if len(t_norm) >= 18 and not any(
        k in t for k in ["合同", "发票", "活动", "海报", "物业", "公寓", "办公室", "不动产", "工厂", "仓库", "门头", "包装", "展会", "宣传", "商标", "图样", "检测"]
    ):
        return True
    return False


def _pick_topic_for_unit(rows: List[Dict[str, str]], mark_name: str) -> str:
    cands: List[str] = []
    for k in ("Product Key", "Inferred Proof Name", "_src_name", "Original File Name"):
        cands.extend(_merge_field_values(rows, k, limit=4))
    for c in cands:
        t = sanitize_component(c, 24)
        if t and t not in EVIDENCE_NAME_GENERIC_TOPICS and (not _is_noisy_topic_token(t)):
            return t
    return sanitize_component(mark_name, 16) or "证据"


def _unit_name_and_purpose(
    unit_type: str,
    rows: List[Dict[str, str]],
    mark_name: str,
    case_class: str = "",
    goods_services: str = "",
) -> Tuple[str, str]:
    ds, _ = _merge_dates_for_rows(rows)
    ym = ds[:7] if ds and len(ds) >= 7 else ""
    ym_tag = ym if ym else "时间不明"
    topic = _pick_topic_for_unit(rows, mark_name)
    blob = " ".join(
        [
            " ".join(_merge_field_values(rows, "Risk Notes", limit=4)),
            " ".join(_merge_field_values(rows, "Inferred Proof Name", limit=4)),
            " ".join(_merge_field_values(rows, "_src_name", limit=6)),
        ]
    )
    if unit_type == "合同":
        subtype = "合同文本"
        if any(k in blob for k in ["租赁", "租售"]):
            subtype = "租赁合同"
        elif any(k in blob for k in ["服务", "技术服务", "服务协议"]):
            subtype = "服务合同"
        elif any(k in blob for k in ["购销", "买卖", "采购", "销售"]):
            subtype = "买卖合同"
        return f"合同-{ym_tag}-{subtype}", "证明合同签署、交易标的及履约时间线索。"
    if unit_type == "发票":
        return f"发票-{ym_tag}-增值税发票", "证明开票事实、交易标的及对应交易时间。"
    if unit_type == "图片":
        poster_like = any(str(r.get("Form Type", "") or "") == "海报" for r in rows) or is_activity_poster_context(blob)
        if poster_like:
            if is_property_service_case(case_class=case_class, goods_services=goods_services):
                return f"广告图片-{ym_tag}-物业活动", "证明在涉案物业组织相关活动并进行运营推广。"
            return f"广告图片-{ym_tag}-活动组织", "证明经营主体组织活动并进行对外宣传推广。"
        return f"实物图-{topic}", "证明产品、包装、场景或页面展示中的实际使用状态。"
    if unit_type == "宣传单":
        return f"广告方案-{topic}", "证明宣传推广、展会投放或品牌展示场景。"
    if unit_type == "流程文件":
        return "流程文件-通知书", "证明撤三程序节点、送达事实及指定期间背景。"
    if unit_type == "主体资质":
        return f"主体资质-{topic}", "证明主体资格、权属关联及经营资质。"
    return f"{unit_type}-{topic}", "证明相关经营事实。"


def merge_fragmented_evidence_units(
    rows: List[Dict[str, str]],
    out_dir: Path,
    mark_name: str,
    case_class: str = "",
    goods_services: str = "",
) -> List[Dict[str, str]]:
    if not rows:
        return rows
    groups: Dict[Tuple[str, str, str, str], List[Tuple[int, Dict[str, str]]]] = {}
    for idx, r in enumerate(rows):
        unit = detect_unit_type(r)
        et = str(r.get("Type", "") or "")
        role = str(r.get("Trade Link Role", "") or "").strip()
        origin_dir = str(r.get("_origin_dir", "") or Path(str(r.get("Original Relative Path", "") or "")).parent)
        ds = str(r.get("Evidence Date Start", "") or "")
        year = ds[:4] if len(ds) >= 4 else ""
        amount = normalize_amount_str(str(r.get("Trade Amount", "") or ""))
        if et == "交易凭证":
            key = (unit, role or "交易", amount or year or "NA", origin_dir or "NA")
        else:
            key = (unit, et or "其他", year or "NA", origin_dir or "NA")
        groups.setdefault(key, []).append((idx, r))

    merged_dir = out_dir / "_merged_context_units"
    merged_dir.mkdir(parents=True, exist_ok=True)
    out_rows: List[Tuple[int, Dict[str, str]]] = []

    def _can_merge(group_rows: List[Dict[str, str]], unit: str) -> bool:
        if len(group_rows) <= 1:
            return False
        if unit in ("合同", "发票", "图片", "宣传单"):
            return True
        small_pages = 0
        for rr in group_rows:
            pr = str(rr.get("Page Range", "") or "")
            if "-" in pr:
                a, b = pr.split("-", 1)
                try:
                    if int(b) - int(a) + 1 <= 2:
                        small_pages += 1
                except Exception as exc:
                    _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return small_pages >= 2

    for key, arr in groups.items():
        arr_sorted = sorted(arr, key=lambda x: (x[0], str(x[1].get("Evidence Date Start", "") or ""), str(x[1].get("_src_name", "") or "")))
        group_rows = [x[1] for x in arr_sorted]
        unit = key[0]
        if not _can_merge(group_rows, unit):
            out_rows.extend(arr_sorted)
            continue
        src_pdfs: List[Path] = []
        for rr in group_rows:
            p = Path(str(rr.get("_src", "") or ""))
            if p.exists():
                src_pdfs.append(p)
        if len(src_pdfs) <= 1:
            out_rows.extend(arr_sorted)
            continue
        merged_pdf = merged_dir / f"unit_{hashlib.sha1('|'.join(str(p) for p in src_pdfs).encode('utf-8')).hexdigest()[:12]}.pdf"
        try:
            w = PdfWriter()
            for sp in src_pdfs:
                rd = PdfReader(str(sp), strict=False)
                for pg in rd.pages:
                    w.add_page(pg)
            with merged_pdf.open("wb") as f:
                w.write(f)
        except Exception:
            out_rows.extend(arr_sorted)
            continue

        base = dict(group_rows[0])
        ds, de = _merge_dates_for_rows(group_rows)
        merged_name, merged_purpose = _unit_name_and_purpose(
            unit,
            group_rows,
            mark_name,
            case_class=case_class,
            goods_services=goods_services,
        )
        role_vals = _merge_field_values(group_rows, "Trade Link Role", limit=3)
        amount_vals = _merge_field_values(group_rows, "Trade Amount", limit=3)
        party_vals = _merge_field_values(group_rows, "Trade Counterparty", limit=3)
        key_lines = [merged_name] + _merge_field_values(group_rows, "_src_name", limit=6)
        base["Evidence Date Start"] = ds
        base["Evidence Date End"] = de
        base["Formation Date"] = ds if (ds and ds == de) else ""
        base["Inferred Proof Name"] = merged_name
        base["Inferred Purpose"] = merged_purpose
        base["Original File Name"] = f"{group_rows[0].get('_src_name', '')}等{len(group_rows)}份"
        base["Original Relative Path"] = str(group_rows[0].get("_origin_dir", "") or group_rows[0].get("Original Relative Path", ""))
        base["Time Anchor Text"] = "；".join(_merge_field_values(group_rows, "Time Anchor Text", limit=4))
        base["Content Summary"] = "；".join(_merge_field_values(group_rows, "Content Summary", limit=3))
        base["Mapping Basis"] = "；".join(_merge_field_values(group_rows, "Mapping Basis", limit=2))
        base["Risk Notes"] = f"碎片合并：同类材料合并{len(group_rows)}份；" + "；".join(_merge_field_values(group_rows, "Risk Notes", limit=2))
        if unit == "合同":
            base["Type"] = "交易凭证"
            base["Trade Link Role"] = "合同"
        elif unit == "发票":
            base["Type"] = "交易凭证"
            base["Trade Link Role"] = "发票"
        elif role_vals:
            base["Trade Link Role"] = role_vals[0]
        if len(amount_vals) == 1:
            base["Trade Amount"] = amount_vals[0]
        elif len(amount_vals) > 1:
            base["Trade Amount"] = ""
            base["Risk Notes"] = f"{base['Risk Notes']}；金额存在多值待人工核验"
        if len(party_vals) == 1:
            base["Trade Counterparty"] = party_vals[0]
        elif len(party_vals) > 1:
            base["Trade Counterparty"] = party_vals[0]
            base["Risk Notes"] = f"{base['Risk Notes']}；交易对方存在多值待人工核验"
        base["_src"] = str(merged_pdf)
        base["_src_name"] = base["Original File Name"]
        base["_origin_rel"] = base.get("Original Relative Path", "")
        base["_origin_dir"] = str(group_rows[0].get("_origin_dir", "") or "")
        base["Trade Chain Group"] = ""
        base["Trade Chain Group No"] = ""
        base["Trade Link Key"] = ""
        base["Unit Type"] = unit
        base["_renamed_name"] = build_renamed_filename(
            int(str(base.get("No", "1") or "1")),
            str(base.get("Type", "其他材料") or "其他材料"),
            ds,
            mark_name,
            key_lines,
            preferred_topic=merged_name,
        )
        out_rows.append((arr_sorted[0][0], base))

    out_rows.sort(key=lambda x: x[0])
    merged_only = [x[1] for x in out_rows]
    for r in merged_only:
        r["Unit Type"] = detect_unit_type(r)
    return merged_only


def sort_rows_for_progressive_review(
    rows: List[Dict[str, str]],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> List[Dict[str, str]]:
    _ = period_start, period_end  # 固定排序规则不再依赖指定期间分段
    unit_order = {
        "流程文件": 1,
        "主体资质": 2,
        "合同": 3,
        "发票": 4,
        "交易履约": 5,
        "图片": 6,
        "宣传单": 7,
        "其他佐证": 8,
    }

    def detect_sort_category(row: Dict[str, str]) -> str:
        et = str(row.get("Type", "") or "")
        form_type = str(row.get("Form Type", "") or "")
        trade_role = str(row.get("Trade Link Role", "") or "").strip()
        src_bucket = str(row.get("_source_bucket", "") or "").strip().lower()
        src = " ".join([
            et,
            str(row.get("Evidence Name", "") or ""),
            str(row.get("Inferred Proof Name", "") or ""),
            str(row.get("Inferred Purpose", "") or ""),
            str(row.get("Time Anchor Text", "") or ""),
            str(row.get("_src_name", "") or ""),
        ])
        # 合同/发票与商品对印组证，统一归到合同票据段，便于三联证据连续展示。
        if row.get("Trade Chain Group", "") and trade_role == "商品对印":
            return "合同票据类"
        # 先按证据主类型快速归类
        if et == "程序文件":
            return "程序类"
        if et == "资质主体证明":
            # “注册证/档案/变更/转让/许可”优先归到权属类
            if any(k in src for k in SORT_CAT_KEYWORDS["权属类"]):
                return "权属类"
            return "主体资质类"
        if et == "商品展示页":
            if any(k in src for k in SORT_CAT_KEYWORDS["实体实物类"]):
                return "实体实物类"
            return "网络展示类"
        if et == "线上店铺展示":
            if any(k in src for k in SORT_CAT_KEYWORDS["实体实物类"]):
                return "实体实物类"
            return "网络展示类"
        if et == "交易凭证":
            if any(k in src for k in SORT_CAT_KEYWORDS["合同票据类"]):
                return "合同票据类"
            return "交易履约类"
        if et == "其他材料":
            if form_type in ("现场照片", "包装") or src_bucket in ("fast", "direct_bind"):
                return "实体实物类"
            if form_type == "海报" or any(k in src for k in BROCHURE_KEYWORDS + EXHIBIT_AD_KEYWORDS):
                return "其他补强类"

        # 其他材料按固定优先级命中
        detect_order = [
            "程序类",
            "权属类",
            "主体资质类",
            "实体实物类",
            "网络展示类",
            "合同票据类",
            "交易履约类",
            "其他补强类",
        ]
        for cat in detect_order:
            if any(k in src for k in SORT_CAT_KEYWORDS.get(cat, [])):
                return cat
        return "其他补强类"

    def sort_date(row: Dict[str, str]) -> dt.date:
        ds = parse_iso_date(row.get("Evidence Date Start", ""))
        de = parse_iso_date(row.get("Evidence Date End", ""))
        return ds or de or dt.date(9999, 12, 31)

    def trade_group_no(row: Dict[str, str]) -> int:
        try:
            raw = str(row.get("Trade Chain Group No", "") or "0").strip()
            no = int(float(raw)) if raw else 0
        except Exception:
            no = 0
        return no if no > 0 else 9999

    def trade_role_for_sort(row: Dict[str, str], unit: str) -> str:
        role = str(row.get("Trade Link Role", "") or "").strip()
        if role:
            return role
        if unit == "合同":
            return "合同"
        if unit == "发票":
            return "发票"
        if unit == "交易履约":
            return "支付"
        return "关联"

    def key(row: Dict[str, str]):
        unit = str(row.get("Unit Type", "") or detect_unit_type(row))
        group_name = str(row.get("Evidence Group", "") or detect_evidence_group(row))
        group_weight = EVIDENCE_GROUP_ORDER.get(group_name, 99)
        d = sort_date(row)
        if group_name == "合同发票组":
            role = trade_role_for_sort(row, unit)
            return (
                group_weight,
                TRADE_ROLE_ORDER.get(role, 99),
                d,
                trade_group_no(row),
                str(row.get("_src_name", "") or ""),
                str(row.get("Evidence Name", "") or ""),
            )
        return (
            group_weight,
            unit_order.get(unit, 99),
            d,
            str(row.get("_src_name", "") or ""),
            str(row.get("Evidence Name", "") or ""),
        )

    for r in rows:
        r["Sort Category"] = detect_sort_category(r)
        r["Sort Weight"] = str(AUTO_SORT_CATEGORY_ORDER.get(r["Sort Category"], 99))
        r["Unit Type"] = str(r.get("Unit Type", "") or detect_unit_type(r))
        group_name = detect_evidence_group(r)
        r["Evidence Group"] = group_name
        r["Evidence Group Weight"] = str(EVIDENCE_GROUP_ORDER.get(group_name, 99))
        b = str(r.get("_source_bucket", "") or "").strip().lower()
        r["Source Bucket"] = b
        r["Source Bucket Weight"] = str(SOURCE_BUCKET_ORDER.get(b, 99))

    return sorted(rows, key=key)


def resequence_rows(rows: List[Dict[str, str]], mark_name: str):
    group_counts: Dict[str, int] = {}
    for row in rows:
        g = str(row.get("Evidence Group", "") or detect_evidence_group(row))
        group_counts[g] = group_counts.get(g, 0) + 1
    group_seen: Dict[str, int] = {}

    for i, row in enumerate(rows, start=1):
        row["No"] = i
        row["Evidence ID"] = evidence_id_display(i)
        group_name = str(row.get("Evidence Group", "") or detect_evidence_group(row))
        group_seen[group_name] = group_seen.get(group_name, 0) + 1
        row["Evidence Group"] = group_name
        row["Evidence Group Weight"] = str(EVIDENCE_GROUP_ORDER.get(group_name, 99))
        row["Evidence Group Index"] = str(group_seen[group_name])
        row["Evidence Group Count"] = str(group_counts.get(group_name, 0))
        row["Evidence Group Seq"] = f"{group_seen[group_name]}/{group_counts.get(group_name, 0)}"
        old_name = row.get("_renamed_name", "")
        tail = re.sub(r"^E\d{3}_", "", old_name)
        if tail:
            row["_renamed_name"] = f"E{i:03d}_{tail}"
        else:
            row["_renamed_name"] = build_renamed_filename(
                i,
                row.get("Type", "其他材料"),
                row.get("Evidence Date Start", ""),
                mark_name,
                [],
            )


def pdf_page_count(pdf: Path) -> int:
    try:
        return len(PdfReader(str(pdf), strict=False).pages)
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return 0


def extract_pdf_text_native(pdf: Path, max_pages: int = 0) -> str:
    try:
        reader = PdfReader(str(pdf), strict=False)
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""

    pages = reader.pages
    if max_pages > 0:
        pages = pages[:max_pages]

    texts: List[str] = []
    for p in pages:
        try:
            t = p.extract_text() or ""
        except Exception:
            t = ""
        if t:
            texts.append(t)
    return "\n".join(texts)


def normalize_text(s: str) -> str:
    return re.sub(r"\s+", "", s or "").replace("|", "").replace("丨", "")


def keyword_hit_count(src: str, keywords: List[str]) -> int:
    if not src:
        return 0
    return sum(1 for k in keywords if k and k in src)


def is_property_service_case(case_class: str = "", goods_services: str = "") -> bool:
    src = f"{case_class or ''} {goods_services or ''}"
    if not src.strip():
        return False
    return any(k in src for k in PROPERTY_SERVICE_KEYWORDS)


def is_activity_poster_context(text: str) -> bool:
    src = text or ""
    if not src:
        return False
    promo_hits = keyword_hit_count(src, SCENE_POSTER_HINT_KEYWORDS + FORM_POSTER_KEYWORDS + BROCHURE_KEYWORDS)
    official_hits = keyword_hit_count(src, ["国家知识产权局", "商标局", "答辩通知", "受理通知", "补正通知", "决定书", "通知书"])
    trade_hits = keyword_hit_count(src, FORM_CONTRACT_KEYWORDS + FORM_INVOICE_KEYWORDS + ["订单编号", "价税合计", "发票代码"])
    if official_hits > 0 or trade_hits > 0:
        return False
    if promo_hits >= 2:
        return True
    return ("活动" in src and any(k in src for k in ["宣传", "推广", "海报", "组织"]))


def sanitize_public_fact_text(text: str, fallback: str = "") -> str:
    src = str(text or "").replace("\n", "；")
    parts = [p.strip("；;，,。 ") for p in re.split(r"[；;\n]+", src) if str(p).strip("；;，,。 ")]
    kept: List[str] = []
    for p in parts:
        if any(k in p for k in PUBLIC_NOTE_EXCLUDE_KEYWORDS):
            continue
        if "该证据属于" in p and "覆盖计算" in p:
            continue
        if p.startswith("⚠跨期矛盾"):
            continue
        kept.append(p)
    base = kept[0] if kept else str(fallback or "").strip("；;，,。 ")
    if not base:
        base = "证明相关使用事实。"
    if len(base) > 72:
        base = base[:71].rstrip("；;，,。 ") + "…"
    if not re.search(r"[。！？]$", base):
        base += "。"
    return base


def clean_ocr_text(text: str) -> str:
    lines: List[str] = []
    last = ""
    for raw in (text or "").splitlines():
        line = raw.replace("\u3000", " ").strip()
        line = re.sub(r"\s+", " ", line)
        line = re.sub(r"[|¦丨]{2,}", "", line)
        if not line:
            continue
        if re.fullmatch(r"[0-9\-_/\. ]{1,16}", line):
            if not (DATE_NUMERIC_RE.search(line) or COMPACT_DATE_RE.search(line)):
                continue
        useful = len(re.findall(r"[A-Za-z0-9\u4e00-\u9fa5]", line))
        if useful < max(2, int(len(line) * 0.25)):
            continue
        if line == last:
            continue
        if len(line) > 180:
            line = line[:180]
        lines.append(line)
        last = line
    return "\n".join(lines)


def normalize_for_dedup(text: str) -> str:
    src = normalize_text(text or "")
    src = re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9]", "", src)
    return src


def dedup_signature(pdf: Path, text: str) -> str:
    norm = normalize_for_dedup(text)
    # 同内容但不同时间点需保留：时间指纹参与去重签名。
    time_fp = "|".join([d.isoformat() for d in extract_dates_advanced(text or "", None, None)[:3]])
    # OCR 过短时退化为文件名+页数，避免误去重。
    if len(norm) < 80:
        pc = pdf_page_count(pdf)
        raw = f"{pdf.stem}|{pc}|{len(norm)}|{time_fp}"
    else:
        raw = f"{norm[:20000]}|{time_fp}"
    return hashlib.sha1(raw.encode("utf-8", errors="ignore")).hexdigest()


def deduplicate_scanned_files(
    files: List[Path],
    scan_texts: Dict[Path, str],
) -> Tuple[List[Path], Dict[Path, List[Path]], Dict[Path, str]]:
    """
    返回：
    - unique_files: 去重后保留的文件（按原顺序）
    - dup_map: canonical -> [duplicates]
    - sig_map: canonical -> signature
    """
    unique_files: List[Path] = []
    sig_to_canonical: Dict[str, Path] = {}
    dup_map: Dict[Path, List[Path]] = {}
    sig_map: Dict[Path, str] = {}

    for p in files:
        sig = dedup_signature(p, scan_texts.get(p, ""))
        if sig not in sig_to_canonical:
            sig_to_canonical[sig] = p
            unique_files.append(p)
            dup_map[p] = []
            sig_map[p] = sig
        else:
            can = sig_to_canonical[sig]
            dup_map.setdefault(can, []).append(p)

    return unique_files, dup_map, sig_map


def evidence_id_display(seq: int) -> str:
    return f"证据{seq}"


def _page_no(img: Path) -> int:
    m = re.search(r"-(\d+)\.png$", img.name)
    return int(m.group(1)) if m else 999999


def ocr_pdf_pages(pdf: Path, cache_dir: Path, max_pages: int = 0, dpi: int = 220) -> str:
    cache_dir.mkdir(parents=True, exist_ok=True)
    page_tag = "all" if max_pages <= 0 else str(max_pages)
    file_key = file_sha256(pdf)[:16] or file_sha1(pdf)[:16] or sanitize_component(pdf.stem, 24)
    cache_txt = cache_dir / f"{file_key}_p{page_tag}_d{dpi}.txt"
    if cache_txt.exists():
        return cache_txt.read_text(encoding="utf-8", errors="ignore")

    native_text = clean_ocr_text(extract_pdf_text_native(pdf, max_pages=max_pages))

    # 命令不可用时，退化为 PDF 原生抽取
    if not Path(PDFTOPPM_BIN).exists() or not Path(TESSERACT_BIN).exists():
        cache_txt.write_text(native_text, encoding="utf-8")
        return native_text

    total_pages = pdf_page_count(pdf)
    page_limit = total_pages if max_pages <= 0 else min(max_pages, total_pages or max_pages)

    with tempfile.TemporaryDirectory(prefix="ocr_pdf_") as td:
        td_path = Path(td)
        prefix = td_path / "page"

        cmd = [
            PDFTOPPM_BIN,
            "-f",
            "1",
        ]
        if page_limit > 0:
            cmd.extend(["-l", str(page_limit)])
        cmd.extend([
            "-r",
            str(dpi),
            "-gray",
            "-png",
            str(pdf),
            str(prefix),
        ])

        timeout_sec = max(120, int((page_limit or 5) * 20))
        code, _, _ = run(cmd, timeout_sec=timeout_sec)
        if code != 0:
            cache_txt.write_text(native_text, encoding="utf-8")
            return native_text

        texts: List[str] = []
        for img in sorted(td_path.glob("page-*.png"), key=_page_no):
            code, out, _ = run([
                TESSERACT_BIN,
                str(img),
                "stdout",
                "-l",
                "chi_sim+eng",
                "--psm",
                "6",
            ], timeout_sec=35)
            if code == 0 and out:
                texts.append(out)

        ocr_text = clean_ocr_text("\n".join(texts))
        merged = clean_ocr_text("\n".join([native_text, ocr_text]))
        cache_txt.write_text(merged, encoding="utf-8")
        return merged


def ocr_time_boost_lines(pdf: Path, cache_dir: Path, max_pages: int = 3, dpi: int = 420) -> str:
    cache_dir.mkdir(parents=True, exist_ok=True)
    file_key = file_sha256(pdf)[:16] or file_sha1(pdf)[:16] or sanitize_component(pdf.stem, 24)
    cache_txt = cache_dir / f"{file_key}_timeboost_p{max_pages}_d{dpi}.txt"
    if cache_txt.exists():
        return cache_txt.read_text(encoding="utf-8", errors="ignore")

    if not Path(PDFTOPPM_BIN).exists() or not Path(TESSERACT_BIN).exists():
        cache_txt.write_text("", encoding="utf-8")
        return ""

    total_pages = pdf_page_count(pdf)
    page_limit = min(max_pages, total_pages) if total_pages else max_pages
    if page_limit <= 0:
        cache_txt.write_text("", encoding="utf-8")
        return ""

    lines_out: List[str] = []
    with tempfile.TemporaryDirectory(prefix="ocr_time_boost_") as td:
        td_path = Path(td)
        prefix = td_path / "page"
        code, _, _ = run([
            PDFTOPPM_BIN,
            "-f", "1",
            "-l", str(page_limit),
            "-r", str(dpi),
            "-png",
            str(pdf),
            str(prefix),
        ], timeout_sec=max(60, page_limit * 15))
        if code != 0:
            cache_txt.write_text("", encoding="utf-8")
            return ""

        for img in sorted(td_path.glob("page-*.png"), key=_page_no):
            code, out, _ = run([
                TESSERACT_BIN,
                str(img),
                "stdout",
                "-l",
                "chi_sim+eng",
                "--psm",
                "11",
                "tsv",
            ], timeout_sec=45)
            if code == 0 and out:
                tsv_lines = out.splitlines()
                if tsv_lines:
                    headers = tsv_lines[0].split("\t")
                    h_idx = {k: i for i, k in enumerate(headers)}
                    if "text" in h_idx:
                        grouped: Dict[Tuple[str, str, str], List[Tuple[int, str]]] = {}
                        for ln in tsv_lines[1:]:
                            parts = ln.split("\t")
                            if len(parts) < len(headers):
                                continue
                            txt = parts[h_idx["text"]].strip()
                            conf = parts[h_idx.get("conf", -1)] if "conf" in h_idx else "0"
                            if not txt or conf == "-1":
                                continue
                            key = (
                                parts[h_idx.get("block_num", 0)],
                                parts[h_idx.get("par_num", 0)],
                                parts[h_idx.get("line_num", 0)],
                            )
                            left = int(parts[h_idx.get("left", 0)]) if "left" in h_idx and parts[h_idx.get("left", 0)].isdigit() else 0
                            grouped.setdefault(key, []).append((left, txt))

                        for _, row in grouped.items():
                            row = sorted(row, key=lambda x: x[0])
                            txt = " ".join([x[1] for x in row]).strip()
                            n = normalize_ocr_date_text(txt)
                            if (
                                any(k in txt for k in REVIEW_KEYWORDS + ORDER_KEYWORDS + LOGISTICS_KEYWORDS)
                                or MONTH_DAY_RE.search(n)
                                or DATE_FUZZY_RE.search(n)
                                or DATE_RE.search(n)
                                or DATE_NUMERIC_RE.search(n)
                                or COMPACT_DATE_RE.search(n)
                            ):
                                lines_out.append(txt)

            # fallback: plain text OCR（对弱格式日期更稳）
            code2, out2, _ = run([
                TESSERACT_BIN,
                str(img),
                "stdout",
                "-l",
                "chi_sim+eng",
                "--psm",
                "6",
            ], timeout_sec=35)
            if code2 == 0 and out2:
                for ln in out2.splitlines():
                    txt = ln.strip()
                    if not txt:
                        continue
                    n = normalize_ocr_date_text(txt)
                    if (
                        any(k in txt for k in REVIEW_KEYWORDS + ORDER_KEYWORDS + LOGISTICS_KEYWORDS)
                        or MONTH_DAY_RE.search(n)
                        or DATE_FUZZY_RE.search(n)
                        or DATE_RE.search(n)
                        or DATE_NUMERIC_RE.search(n)
                        or COMPACT_DATE_RE.search(n)
                    ):
                        lines_out.append(txt)

    merged = clean_ocr_text("\n".join(lines_out))
    cache_txt.write_text(merged, encoding="utf-8")
    return merged


def parse_pdf_metadata_date(pdf: Path) -> Optional[dt.date]:
    try:
        md = PdfReader(str(pdf), strict=False).metadata or {}
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return None
    for k in ["/ModDate", "/CreationDate"]:
        raw = str(md.get(k, "")).strip()
        if not raw:
            continue
        m = re.search(r"D:(\d{4})(\d{2})(\d{2})", raw)
        if m:
            d = parse_date(m.group(1), m.group(2), m.group(3))
            if d:
                return d
    return None


def cleanup_party_name(s: str) -> str:
    src = (s or "").replace("\u3000", " ").strip()
    src = re.split(r"[，。；,;]\s*", src)[0]
    m_company = COMPANY_RE.search(src)
    if m_company:
        return m_company.group(1)
    m_text = re.search(r"[\u4e00-\u9fa5A-Za-z0-9·（）()]{2,40}", src)
    return m_text.group(0) if m_text else ""


def read_docx_plain_text(path: Path) -> str:
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def read_xlsx_plain_text(path: Path) -> str:
    parts: List[str] = []
    try:
        xls = pd.ExcelFile(str(path))
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""
    for s in xls.sheet_names[:8]:
        try:
            df = pd.read_excel(xls, sheet_name=s, dtype=str)
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
            continue
        parts.append(f"[sheet]{s}")
        if df is None or df.empty:
            continue
        cols = [str(c) for c in list(df.columns)[:20]]
        if cols:
            parts.append(" ".join(cols))
        row_limit = min(120, len(df))
        for i in range(row_limit):
            row_vals: List[str] = []
            for v in df.iloc[i].tolist()[:20]:
                if v is None:
                    continue
                sv = str(v).strip()
                if not sv or sv.lower() == "nan":
                    continue
                row_vals.append(sv)
            if row_vals:
                parts.append(" | ".join(row_vals))
    return clean_ocr_text("\n".join(parts))


def read_pptx_plain_text(path: Path) -> str:
    parts: List[str] = []
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            slide_names = sorted([n for n in zf.namelist() if n.startswith("ppt/slides/slide") and n.endswith(".xml")])
            for name in slide_names[:80]:
                try:
                    data = zf.read(name)
                    root = ET.fromstring(data)
                except Exception as exc:
                    _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
                    continue
                texts: List[str] = []
                for node in root.iter():
                    if node.tag.endswith("}t") and node.text:
                        t = node.text.strip()
                        if t:
                            texts.append(t)
                if texts:
                    parts.append(" ".join(texts))
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return ""
    return clean_ocr_text("\n".join(parts))


def read_office_plain_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return clean_ocr_text(read_docx_plain_text(path))
    if suffix in {".xlsx", ".xls"}:
        return read_xlsx_plain_text(path)
    if suffix == ".pptx":
        return read_pptx_plain_text(path)
    return ""


def party_text_score(name: str) -> int:
    s = (name or "").strip()
    if not s:
        return -100
    score = 0
    if len(s) >= 4:
        score += 1
    if any(k in s for k in ["有限责任公司", "股份有限公司", "有限公司"]):
        score += 2
    if any(k in s for k in ["(", ")", "（", "）", "代理", "事务所", "地址"]):
        score -= 1
    if re.search(r"[A-Za-z]{3,}", s):
        score -= 1
    return score


def pick_better_party(current: str, candidate: str) -> str:
    c1 = (current or "").strip()
    c2 = (candidate or "").strip()
    if not c2:
        return c1
    if not c1:
        return c2
    if party_text_score(c2) > party_text_score(c1):
        return c2
    return c1


def goods_hint_score(v: str) -> int:
    s = (v or "").strip()
    if not s:
        return -100
    score = 0
    if len(s) >= 6:
        score += 1
    if any(k in s for k in ["（以注册证记载为准）", "核定商品", "核定使用商品"]):
        score -= 1
    if s in {"项目", "商品", "服务"}:
        score -= 2
    return score


def pick_better_goods_hint(current: str, candidate: str) -> str:
    c1 = (current or "").strip()
    c2 = (candidate or "").strip()
    if not c2:
        return c1
    if not c1:
        return c2
    if goods_hint_score(c2) > goods_hint_score(c1):
        return c2
    return c1


def extract_goods_hint_from_structured_text(src: str) -> str:
    lines = [ln.strip() for ln in (src or "").splitlines() if ln.strip()]
    if not lines:
        return ""
    start = -1
    for i, ln in enumerate(lines):
        if any(k in ln for k in ["核定使用商品", "核定商品", "商品/服务项目", "商品服务项目"]):
            start = i
            break
    if start < 0:
        return ""
    picks: List[str] = []
    for ln in lines[start + 1:]:
        if len(picks) >= 4:
            break
        if re.fullmatch(r"\d{3,6}", ln):
            continue
        if ln in {"基本信息", "日期", "群组", "项目"}:
            continue
        if "代理机构" in ln or "地址" in ln:
            break
        txt = re.sub(r"^\d{3,6}\s*", "", ln).strip()
        if not txt:
            continue
        if len(txt) >= 2:
            picks.append(txt)
    return "；".join(picks)


def infer_case_info(notice_text: str) -> Dict[str, str]:
    raw = notice_text or ""
    norm = raw.replace("\u3000", " ")
    compact = normalize_text(norm)

    reg_no_m = REG_RE.search(compact) or REG_RE.search(norm) or REG_NO_TAG_RE.search(compact) or REG_NO_TAG_RE.search(norm)
    class_m = CLASS_RE.search(compact) or CLASS_RE.search(norm)
    applicant_m = APPLICANT_RE.search(norm)
    respondent_m = RESPONDENT_RE.search(norm)
    mark_m = MARK_RE.search(compact) or MARK_TEXT_RE.search(compact) or MARK_ARCHIVE_RE.search(norm)

    lines = [ln.strip() for ln in norm.splitlines() if ln.strip()]

    cand_freq: Dict[str, int] = {}
    for ln in lines:
        if "有限公司" not in ln and "有限责任公司" not in ln:
            continue
        if any(x in ln for x in ["事务所", "代理", "地址", "抄送"]):
            continue
        for c in COMPANY_RE.findall(ln):
            if "事务所" in c or "代理" in c:
                continue
            cand_freq[c] = cand_freq.get(c, 0) + 1

    for c in COMPANY_RE.findall(compact):
        if "事务所" in c or "代理" in c:
            continue
        cand_freq[c] = cand_freq.get(c, 0) + 1

    respondent = ""
    if respondent_m:
        respondent = cleanup_party_name(respondent_m.group(1))
    if cand_freq:
        cand_pick = sorted(cand_freq.items(), key=lambda kv: (kv[1], len(kv[0])), reverse=True)[0][0]
        if not respondent:
            respondent = cand_pick

    period_start = ""
    period_end = ""
    # 优先提取“指定期间/三年撤销申请日倒推”的区间，避免误用商标十年有效期
    range_pat = re.compile(r"(20\d{2}[年./-]\d{1,2}[月./-]\d{1,2}日?)\s*(?:至|到|~|－|-|—)\s*(20\d{2}[年./-]\d{1,2}[月./-]\d{1,2}日?)")
    ranged: List[Tuple[dt.date, dt.date, int]] = []
    for m in range_pat.finditer(compact):
        d1 = extract_dates(m.group(1))
        d2 = extract_dates(m.group(2))
        if not d1 or not d2:
            continue
        a, b = d1[0], d2[0]
        if a > b:
            a, b = b, a
        span_days = (b - a).days
        ctx = compact[max(0, m.start() - 24): m.end() + 24]
        score = 0
        if any(k in ctx for k in ["指定期间", "撤销", "连续三年", "三年", "不使用"]):
            score += 3
        # 三年窗口优先（容忍闰年和文书日差）
        if 900 <= span_days <= 1400:
            score += 2
        # 十年有效期降权
        if span_days > 1800:
            score -= 3
        ranged.append((a, b, score))
    if ranged:
        ranged.sort(key=lambda x: (x[2], -(x[1] - x[0]).days), reverse=True)
        best = ranged[0]
        if best[2] >= 1:
            period_start, period_end = to_iso(best[0]), to_iso(best[1])

    if not period_start or not period_end:
        app_m = re.search(r"申请\s*日期[:：]?\s*(20\d{2}[年./-]\d{1,2}[月./-]\d{1,2}日?)", compact)
        if app_m:
            app_ds = extract_dates(app_m.group(1))
            if app_ds:
                app = app_ds[0]
                try:
                    start = dt.date(app.year - 3, app.month, app.day)
                except ValueError:
                    start = dt.date(app.year - 3, app.month, 28)
                end = app - dt.timedelta(days=1)
                period_start, period_end = to_iso(start), to_iso(end)

    goods_match = GOODS_RE.search(compact) or GOODS_RE.search(norm)
    goods = ""
    if goods_match:
        goods = goods_match.group(1)
    else:
        goods_label_m = GOODS_LABEL_RE.search(norm)
        if goods_label_m:
            goods = goods_label_m.group(1).strip()
    if not goods:
        goods = "核定商品（以注册证记载为准）"

    applicant = cleanup_party_name(applicant_m.group(1) if applicant_m else "")
    if not applicant:
        for ln in lines:
            if "申请人" in ln:
                right = re.split(r"申请\s*人[:：]?", ln)[-1]
                applicant = cleanup_party_name(right)
                if applicant:
                    break

    return {
        "reg_no": reg_no_m.group(1) if reg_no_m else "",
        "class": class_m.group(1) if class_m else "",
        "applicant": applicant,
        "respondent": respondent,
        "mark_name": mark_m.group(1) if mark_m else "",
        "use_period_start": period_start,
        "use_period_end": period_end,
        "goods_hint": goods,
    }


def merge_case_info(primary: Dict[str, str], backup: Dict[str, str]) -> Dict[str, str]:
    merged = dict(primary)
    for k, v in backup.items():
        if not merged.get(k) and v:
            merged[k] = v
    return merged


def infer_case_info_from_fact_docs(
    files: List[Path],
    scan_texts: Dict[Path, str],
    evidence_dir: Optional[Path] = None,
    evidence_dirs: Optional[List[Path]] = None,
) -> Dict[str, str]:
    """
    从“商标档案/注册信息类材料”中优先提取案件事实，减少普通证据噪音干扰。
    """
    fact = {
        "reg_no": "",
        "class": "",
        "applicant": "",
        "respondent": "",
        "mark_name": "",
        "use_period_start": "",
        "use_period_end": "",
        "goods_hint": "",
    }
    candidates: List[Path] = []
    for p in files:
        name = p.stem
        txt = scan_texts.get(p, "") or ""
        is_fact_name = any(k in name for k in ["商标档案", "商标信息", "注册证", "商标注册", "档案"])
        is_fact_text = ("申请号" in txt and "国际分类" in txt and ("商标名称" in txt or "商标流程" in txt))
        if is_fact_name or is_fact_text:
            candidates.append(p)
    # 补充读取原始 docx 案件事实材料，避免 OCR 页数不足导致事实字段漏提
    text_only_sources: List[Tuple[str, str]] = []
    scan_roots: List[Path] = []
    if evidence_dirs:
        scan_roots.extend([p for p in evidence_dirs if p and p.exists()])
    if evidence_dir and evidence_dir.exists():
        scan_roots.append(evidence_dir)
    unique_scan_roots = list(dict.fromkeys([p.resolve() for p in scan_roots]))
    for root in unique_scan_roots:
        for p in sorted(root.rglob("*")):
            if not p.is_file() or p.name.startswith("._"):
                continue
            if p.suffix.lower() != ".docx":
                continue
            if not any(k in p.stem for k in ["商标档案", "商标信息", "注册证", "档案", "商标"]):
                continue
            txt = read_docx_plain_text(p)
            if txt:
                text_only_sources.append((p.stem, txt))

    def _mark_from_fact_filename(stem: str) -> str:
        s = stem or ""
        m = re.search(r"([A-Za-z0-9\u4e00-\u9fa5·\s]{2,60})商标档案", s)
        if not m:
            return ""
        tok = m.group(1)
        tok = re.sub(r"第?\d{6,}\s*号", "", tok)
        tok = re.sub(r"[（(]\s*第?\d{1,2}\s*类\s*[）)]", "", tok)
        tok = re.sub(r"第?\s*\d{1,2}\s*类", "", tok)
        tok = re.sub(r"20\d{2}[-_./]\d{1,2}[-_./]\d{1,2}", "", tok)
        tok = re.sub(r"\s+", " ", tok).strip(" _-")
        if tok and not mark_has_noise(normalize_mark_candidate(tok)):
            return tok
        return ""

    for p in candidates:
        txt = scan_texts.get(p, "") or ""
        src = f"{p.stem}\n{txt}"
        compact = normalize_text(src)

        if not fact["reg_no"]:
            m = REG_RE.search(src) or REG_RE.search(compact) or REG_NO_TAG_RE.search(src) or REG_NO_TAG_RE.search(compact)
            if m:
                fact["reg_no"] = m.group(1)
        if not fact["class"]:
            m = CLASS_RE.search(src) or CLASS_RE.search(compact)
            if m:
                fact["class"] = m.group(1)
        if not fact["mark_name"]:
            mk2 = _mark_from_fact_filename(p.stem)
            if mk2:
                fact["mark_name"] = mk2
        if not fact["mark_name"]:
            m = MARK_ARCHIVE_RE.search(src) or MARK_RE.search(src) or MARK_TEXT_RE.search(src)
            if m:
                mk = m.group(1).strip()
                if mk and not mark_has_noise(normalize_mark_candidate(mk)):
                    fact["mark_name"] = mk
        m = RESPONDENT_RE.search(src)
        if m:
            fact["respondent"] = pick_better_party(fact.get("respondent", ""), cleanup_party_name(m.group(1)))
        m = APPLICANT_RE.search(src)
        if m:
            fact["applicant"] = pick_better_party(fact.get("applicant", ""), cleanup_party_name(m.group(1)))
        m = GOODS_LABEL_RE.search(src) or GOODS_RE.search(src)
        if m:
            fact["goods_hint"] = pick_better_goods_hint(fact.get("goods_hint", ""), m.group(1).strip())
        gh = extract_goods_hint_from_structured_text(src)
        if gh:
            fact["goods_hint"] = pick_better_goods_hint(fact.get("goods_hint", ""), gh)

    for stem, txt in text_only_sources:
        src = f"{stem}\n{txt}"
        compact = normalize_text(src)
        if not fact["reg_no"]:
            m = REG_RE.search(src) or REG_RE.search(compact) or REG_NO_TAG_RE.search(src) or REG_NO_TAG_RE.search(compact)
            if m:
                fact["reg_no"] = m.group(1)
        if not fact["class"]:
            m = CLASS_RE.search(src) or CLASS_RE.search(compact)
            if m:
                fact["class"] = m.group(1)
        if not fact["mark_name"]:
            mk2 = _mark_from_fact_filename(stem)
            if mk2:
                fact["mark_name"] = mk2
        if not fact["mark_name"]:
            m = MARK_ARCHIVE_RE.search(src) or MARK_RE.search(src) or MARK_TEXT_RE.search(src)
            if m:
                mk = m.group(1).strip()
                if mk and not mark_has_noise(normalize_mark_candidate(mk)):
                    fact["mark_name"] = mk
        m = RESPONDENT_RE.search(src)
        if m:
            fact["respondent"] = pick_better_party(fact.get("respondent", ""), cleanup_party_name(m.group(1)))
        m = APPLICANT_RE.search(src)
        if m:
            fact["applicant"] = pick_better_party(fact.get("applicant", ""), cleanup_party_name(m.group(1)))
        m = GOODS_LABEL_RE.search(src) or GOODS_RE.search(src)
        if m:
            fact["goods_hint"] = pick_better_goods_hint(fact.get("goods_hint", ""), m.group(1).strip())
        gh = extract_goods_hint_from_structured_text(src)
        if gh:
            fact["goods_hint"] = pick_better_goods_hint(fact.get("goods_hint", ""), gh)

    # 商标档案中常见“申请人”即注册权利人；若未识别到注册人，使用申请人兜底。
    if not (fact.get("respondent", "") or "").strip() and (fact.get("applicant", "") or "").strip():
        fact["respondent"] = fact["applicant"]

    return fact


def apply_fact_case_overrides(base_case: Dict[str, str], fact_case: Dict[str, str]) -> Dict[str, str]:
    out = dict(base_case)
    for k in ["reg_no", "class", "respondent", "applicant", "use_period_start", "use_period_end", "goods_hint"]:
        v = (fact_case.get(k, "") or "").strip()
        if v:
            out[k] = v
    mk = (fact_case.get("mark_name", "") or "").strip()
    if mk and not mark_has_noise(normalize_mark_candidate(mk)):
        out["mark_name"] = mk
    return out


REQUIRED_CASEINFO_KEYS = [
    "case_no",
    "reg_no",
    "respondent",
    "applicant",
    "class",
    "designated_goods_services",
    "use_period_start",
    "use_period_end",
    "cnipa_notice_ref",
    "mark_name",
]

CASE_INFO_VERSION = "caseinfo.v1"
VALIDATION_RULE_VERSION = "caseinfo_qc.v1"
VALIDATED_BY = "auto_recognize_and_generate.py"


def normalize_digits_only(s: str) -> str:
    return NON_DIGIT_RE.sub("", s or "")


def normalize_compare_text(s: str) -> str:
    src = (s or "").replace("\u3000", " ").strip()
    src = re.sub(r"\s+", "", src)
    src = re.sub(r"[，。；：,:（）()【】\[\]\"'“”]", "", src)
    return src


def normalize_party_for_compare(s: str) -> str:
    src = normalize_compare_text(s)
    src = src.replace("有限责任公司", "有限公司")
    return src


def normalize_mark_for_compare(s: str) -> str:
    return normalize_mark_candidate(normalize_compare_text(s))


def normalize_case_value(key: str, value: str) -> str:
    v = (value or "").strip()
    if not v:
        return ""
    if key == "reg_no":
        return normalize_digits_only(v)
    if key == "class":
        return normalize_digits_only(v)
    if key in {"use_period_start", "use_period_end"}:
        d = parse_iso_date(v)
        return d.isoformat() if d else ""
    if key in {"respondent", "applicant"}:
        return normalize_party_for_compare(v)
    if key == "mark_name":
        return normalize_mark_for_compare(v)
    return normalize_compare_text(v)


def is_placeholder_case_value(key: str, value: str) -> bool:
    v = (value or "").strip()
    if not v:
        return True
    if "待补" in v:
        return True
    if key in {"reg_no", "class"} and not normalize_digits_only(v):
        return True
    if key == "mark_name":
        if len(normalize_mark_for_compare(v)) < 2:
            return True
        if mark_has_noise(normalize_mark_candidate(v)):
            return True
    if key in {"respondent", "applicant"} and len(normalize_party_for_compare(v)) < 2:
        return True
    return False


def case_values_consistent(key: str, left: str, right: str) -> bool:
    l = normalize_case_value(key, left)
    r = normalize_case_value(key, right)
    if not l or not r:
        return True
    if key in {"respondent", "applicant", "mark_name", "designated_goods_services", "case_no"}:
        return l in r or r in l
    return l == r


def build_caseinfo_source_payload(raw: Dict[str, str]) -> Dict[str, str]:
    return {
        "case_no": "",
        "reg_no": raw.get("reg_no", ""),
        "respondent": raw.get("respondent", ""),
        "applicant": raw.get("applicant", ""),
        "class": raw.get("class", ""),
        "designated_goods_services": raw.get("goods_hint", ""),
        "use_period_start": raw.get("use_period_start", ""),
        "use_period_end": raw.get("use_period_end", ""),
        "cnipa_notice_ref": "",
        "mark_name": raw.get("mark_name", ""),
    }


def detect_caseinfo_source(
    key: str,
    final_value: str,
    notice_value: str,
    evidence_value: str,
) -> str:
    if case_values_consistent(key, final_value, notice_value) and normalize_case_value(key, notice_value):
        return "notice_ocr"
    if case_values_consistent(key, final_value, evidence_value) and normalize_case_value(key, evidence_value):
        return "evidence_ocr"
    if key == "case_no":
        return "derived"
    if key == "cnipa_notice_ref":
        return "system_generated"
    if key == "designated_goods_services":
        if "以注册证记载为准" in (final_value or ""):
            return "fallback_default"
    if key in {"use_period_start", "use_period_end"} and (final_value or "").startswith("20"):
        return "fallback_default"
    return "fusion_or_manual"


def validate_caseinfo_and_crosscheck(
    case_final: Dict[str, str],
    notice_case_raw: Dict[str, str],
    evidence_case_raw: Dict[str, str],
    allow_noisy_mark_name: bool = False,
) -> Dict[str, Any]:
    errors: List[str] = []
    warnings: List[str] = []
    checks: List[Dict[str, str]] = []
    validated_at = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    notice_case = build_caseinfo_source_payload(notice_case_raw)
    evidence_case = build_caseinfo_source_payload(evidence_case_raw)

    def add_check(
        field: str,
        lhs_name: str,
        lhs: str,
        rhs_name: str,
        rhs: str,
        strict: bool,
    ) -> None:
        ln = normalize_case_value(field, lhs)
        rn = normalize_case_value(field, rhs)
        if not ln or not rn:
            checks.append({
                "field": field,
                "lhs": lhs_name,
                "rhs": rhs_name,
                "result": "SKIP_EMPTY",
            })
            return
        ok = case_values_consistent(field, lhs, rhs)
        checks.append({
            "field": field,
            "lhs": lhs_name,
            "rhs": rhs_name,
            "result": "PASS" if ok else "FAIL",
        })
        if ok:
            return
        msg = f"{field} 对账不一致：{lhs_name}={lhs}；{rhs_name}={rhs}"
        if strict:
            errors.append(msg)
        else:
            warnings.append(msg)

    for key in REQUIRED_CASEINFO_KEYS:
        v = case_final.get(key, "")
        if key == "mark_name" and allow_noisy_mark_name:
            vv = (v or "").strip()
            if not vv or "待补" in vv:
                errors.append(f"{key} 缺失或无效：{v}")
            continue
        if is_placeholder_case_value(key, v):
            errors.append(f"{key} 缺失或无效：{v}")

    mark_final = normalize_mark_candidate(case_final.get("mark_name", ""))
    if (not allow_noisy_mark_name) and mark_final and mark_has_noise(mark_final):
        errors.append(f"mark_name 疑似噪音词：{case_final.get('mark_name', '')}")

    reg_no = normalize_digits_only(case_final.get("reg_no", ""))
    if not re.fullmatch(r"\d{6,}", reg_no):
        errors.append(f"reg_no 格式不合法：{case_final.get('reg_no', '')}")

    class_no = normalize_digits_only(case_final.get("class", ""))
    if not class_no:
        errors.append(f"class 缺失：{case_final.get('class', '')}")
    else:
        class_int = int(class_no)
        if class_int < 1 or class_int > 45:
            errors.append(f"class 超出范围(1~45)：{class_no}")

    start = parse_iso_date(case_final.get("use_period_start", ""))
    end = parse_iso_date(case_final.get("use_period_end", ""))
    if not start or not end:
        errors.append("指定期间起止日期格式不合法，必须为 YYYY-MM-DD")
    else:
        if start > end:
            errors.append("指定期间起始日期晚于结束日期")
        if (end - start).days > 365 * 5:
            warnings.append("指定期间跨度超过 5 年，请人工复核是否配置错误")

    if "以注册证记载为准" in (case_final.get("designated_goods_services", "")):
        warnings.append("designated_goods_services 为兜底描述，建议人工补全核定商品/服务")

    add_check("reg_no", "notice", notice_case.get("reg_no", ""), "caseinfo", case_final.get("reg_no", ""), strict=True)
    add_check("class", "notice", notice_case.get("class", ""), "caseinfo", case_final.get("class", ""), strict=True)
    add_check("use_period_start", "notice", notice_case.get("use_period_start", ""), "caseinfo", case_final.get("use_period_start", ""), strict=True)
    add_check("use_period_end", "notice", notice_case.get("use_period_end", ""), "caseinfo", case_final.get("use_period_end", ""), strict=True)
    add_check("mark_name", "notice", notice_case.get("mark_name", ""), "caseinfo", case_final.get("mark_name", ""), strict=False)
    add_check("respondent", "notice", notice_case.get("respondent", ""), "caseinfo", case_final.get("respondent", ""), strict=False)
    add_check("applicant", "notice", notice_case.get("applicant", ""), "caseinfo", case_final.get("applicant", ""), strict=False)

    add_check("reg_no", "evidence", evidence_case.get("reg_no", ""), "caseinfo", case_final.get("reg_no", ""), strict=False)
    add_check("class", "evidence", evidence_case.get("class", ""), "caseinfo", case_final.get("class", ""), strict=False)
    add_check("mark_name", "evidence", evidence_case.get("mark_name", ""), "caseinfo", case_final.get("mark_name", ""), strict=False)
    add_check("respondent", "evidence", evidence_case.get("respondent", ""), "caseinfo", case_final.get("respondent", ""), strict=False)
    add_check("applicant", "evidence", evidence_case.get("applicant", ""), "caseinfo", case_final.get("applicant", ""), strict=False)
    add_check("use_period_start", "evidence", evidence_case.get("use_period_start", ""), "caseinfo", case_final.get("use_period_start", ""), strict=False)
    add_check("use_period_end", "evidence", evidence_case.get("use_period_end", ""), "caseinfo", case_final.get("use_period_end", ""), strict=False)

    source_of_truth = {}
    for key in REQUIRED_CASEINFO_KEYS:
        source_of_truth[key] = detect_caseinfo_source(
            key=key,
            final_value=case_final.get(key, ""),
            notice_value=notice_case.get(key, ""),
            evidence_value=evidence_case.get(key, ""),
        )

    return {
        "status": "FAIL" if errors else "PASS",
        "errors": errors,
        "warnings": warnings,
        "checks": checks,
        "source_of_truth": source_of_truth,
        "validated_at": validated_at,
        "validated_by": VALIDATED_BY,
        "case_info_version": CASE_INFO_VERSION,
        "validation_rule_version": VALIDATION_RULE_VERSION,
    }


def infer_mark_from_filenames(files: List[Path]) -> str:
    freq: Dict[str, int] = {}
    suffix_noise = ["工厂店铺", "店铺", "工厂店", "工厂", "手机端", "电脑端", "资质证明", "产品", "介绍", "通知书"]
    for p in files:
        for tok in re.findall(r"[\u4e00-\u9fa5]{2,8}", p.stem):
            clean_tok = tok
            for sfx in suffix_noise:
                clean_tok = clean_tok.replace(sfx, "")
            clean_tok = clean_tok.strip()
            if len(clean_tok) < 2:
                continue
            if clean_tok in STOP_MARK_TOKENS:
                continue
            freq[clean_tok] = freq.get(clean_tok, 0) + 1
    if not freq:
        return ""
    cand = sorted(freq.items(), key=lambda kv: (kv[1], len(kv[0])), reverse=True)
    return cand[0][0]


def infer_mark_from_texts(texts: List[str]) -> str:
    freq: Dict[str, int] = {}
    for text in texts:
        compact = normalize_text(text)
        for m in MARK_TEXT_RE.finditer(compact):
            tok = m.group(1).strip("“”\"'")
            if len(tok) < 2 or tok in STOP_MARK_TOKENS:
                continue
            if "有限公司" in tok:
                continue
            freq[tok] = freq.get(tok, 0) + 2

        for ln in text.splitlines():
            if "商标" not in ln:
                continue
            for tok in re.findall(r"[\u4e00-\u9fa5A-Za-z0-9·-]{2,12}", ln):
                if tok in STOP_MARK_TOKENS or "有限公司" in tok:
                    continue
                freq[tok] = freq.get(tok, 0) + 1

    if not freq:
        return ""
    cand = sorted(freq.items(), key=lambda kv: (kv[1], len(kv[0])), reverse=True)
    return cand[0][0]


def normalize_mark_candidate(s: str) -> str:
    t = (s or "").strip().replace("“", "").replace("”", "").replace("\"", "")
    t = re.sub(r"^第?\d+类", "", t)
    t = re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9·-]", "", t)
    return t.strip("-")


def mark_has_noise(token: str) -> bool:
    if not token:
        return True
    if len(token) < 2 or len(token) > 14:
        return True
    low = token.lower()
    if re.fullmatch(r"(微信图片|截图|图片)\d*", token):
        return True
    if re.fullmatch(r"(img|image|dsc|picture|scan|weixin)\d*", low):
        return True
    if re.search(r"\d{6,}", token):
        return True
    if any(w in token for w in BAD_MARK_WORDS):
        return True
    if any(w in low for w in BAD_MARK_WORDS):
        return True
    return False


def pick_best_mark(candidates: List[str], texts: List[str], files: List[Path]) -> str:
    corpus = "\\n".join(texts)
    uniq: List[str] = []
    seen = set()

    for raw in candidates:
        c = normalize_mark_candidate(raw)
        if not c or c in seen:
            continue
        uniq.append(c)
        seen.add(c)

    scored: List[Tuple[int, str]] = []
    for c in uniq:
        score = 0
        if not mark_has_noise(c):
            score += 20
        score += min(corpus.count(c), 20)
        score += sum(1 for p in files if c in p.stem)
        if 2 <= len(c) <= 8:
            score += 2
        scored.append((score, c))

    if scored:
        scored.sort(key=lambda x: (x[0], len(x[1])), reverse=True)
        best = scored[0][1]
        if not mark_has_noise(best):
            return best

    # 兜底：仅返回非噪音候选，避免把“微信图片”等文件名误识别为商标
    for c in uniq:
        if c and not mark_has_noise(c):
            return c
    return ""


def infer_evidence_form_profile(path: Path, text: str, origin_hint: str = "") -> Dict[str, str]:
    src = text or ""
    src_l = src.lower()
    name = path.name if path else ""
    name_l = name.lower()
    origin = str(origin_hint or "")
    origin_l = origin.lower()
    src_norm = normalize_text(src)
    name_norm = normalize_text(name)
    origin_norm = normalize_text(origin)
    src_compact = re.sub(r"\s+", "", src)
    text_len = len(src_compact)
    line_cnt = len([ln for ln in src.splitlines() if ln.strip()])
    alnum_cnt = sum(1 for ch in src if ch.isalnum())
    digit_cnt = sum(1 for ch in src if ch.isdigit())
    digit_ratio = (digit_cnt / alnum_cnt) if alnum_cnt else 0.0

    image_ratio = 0.0
    text_ratio = 0.0
    image_blocks = 0
    text_blocks = 0
    orientation = "unknown"
    visual_ready = False

    if fitz is not None and str(path.suffix).lower() == ".pdf":
        try:
            doc = fitz.open(str(path))
            if doc.page_count > 0:
                p0 = doc[0]
                rect = p0.rect
                page_area = max(float(rect.width * rect.height), 1.0)
                if rect.width > rect.height * 1.12:
                    orientation = "landscape"
                elif rect.height > rect.width * 1.12:
                    orientation = "portrait"
                else:
                    orientation = "square"
                td = p0.get_text("dict")
                for blk in td.get("blocks", []):
                    bbox = blk.get("bbox")
                    if not bbox or len(bbox) != 4:
                        continue
                    try:
                        area = max(0.0, (float(bbox[2]) - float(bbox[0])) * (float(bbox[3]) - float(bbox[1])))
                    except Exception:
                        area = 0.0
                    if blk.get("type", 0) == 1:
                        image_blocks += 1
                        image_ratio += area / page_area
                    else:
                        text_blocks += 1
                        text_ratio += area / page_area
                image_ratio = min(max(image_ratio, 0.0), 1.0)
                text_ratio = min(max(text_ratio, 0.0), 1.0)
                visual_ready = True
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        finally:
            try:
                doc.close()  # type: ignore[name-defined]
            except Exception as exc:
                _audit_exception("auto_recognize.swallowed_exception", exc, __file__)

    source_notice_bucket = any(k in origin for k in ["通知", "须知", "答辩", "提供使用证据", "程序"])
    source_contract_bucket = any(k in origin for k in ["合同", "协议"])
    source_invoice_bucket = any(k in origin for k in ["发票", "票据", "增值税"])
    source_photo_bucket = any(k in origin for k in ["照片", "实拍", "图片", "相册", "商标照片", "门头", "车间", "厂区", "仓库"])
    source_promo_bucket = any(k in origin for k in ["宣传", "海报", "画册", "折页", "展会", "展位", "易拉宝"])
    name_contract_hit = any(k in name for k in ["合同", "协议"])
    name_invoice_hit = any(k in name for k in ["发票", "票据", "增值税"])
    if name_contract_hit and not name_invoice_hit:
        source_contract_bucket = True
        source_invoice_bucket = False
    elif name_invoice_hit and not name_contract_hit:
        source_invoice_bucket = True

    notice_kw_hit = (
        any(k in src for k in FORM_NOTICE_KEYWORDS)
        or any(k in name for k in FORM_NOTICE_KEYWORDS)
        or any(k in origin for k in FORM_NOTICE_KEYWORDS)
    )
    contract_kw_hit = (
        any(k in src for k in FORM_CONTRACT_KEYWORDS)
        or any(k in name for k in FORM_CONTRACT_KEYWORDS)
        or any(k in origin for k in FORM_CONTRACT_KEYWORDS)
        or ("合同及发票" in origin)
    )
    invoice_kw_hit = (
        any(k in src for k in FORM_INVOICE_KEYWORDS)
        or any(k in name for k in FORM_INVOICE_KEYWORDS)
        or any(k in origin for k in FORM_INVOICE_KEYWORDS)
        or ("合同及发票" in origin and ("发票" in name or "发票" in origin))
    )
    poster_kw_hit = (
        any(k in src for k in FORM_POSTER_KEYWORDS)
        or any(k in name for k in FORM_POSTER_KEYWORDS)
        or any(k in origin for k in FORM_POSTER_KEYWORDS)
    )
    packaging_kw_hit = (
        any(k in src for k in FORM_PACKAGING_KEYWORDS)
        or any(k in name for k in FORM_PACKAGING_KEYWORDS)
        or any(k in origin for k in FORM_PACKAGING_KEYWORDS)
    )
    scene_kw_hit = (
        any(k in src for k in FORM_SCENE_PHOTO_KEYWORDS)
        or any(k in name for k in FORM_SCENE_PHOTO_KEYWORDS)
        or any(k in origin for k in FORM_SCENE_PHOTO_KEYWORDS)
        or ("商标照片" in origin)
    )
    contract_structured_terms = ["甲方", "乙方", "合同编号", "签订", "签章", "盖章", "履行期限", "违约责任", "合同金额", "租赁期限"]
    invoice_structured_terms = ["发票代码", "发票号码", "开票日期", "价税合计", "税额", "销货清单", "购买方", "销售方", "纳税人识别号"]
    notice_structured_terms = ["国家知识产权局", "商标局", "撤销连续三年不使用", "答辩期限", "送达", "受理", "决定书"]
    contract_structured_hit = sum(1 for k in contract_structured_terms if k in src) >= 2
    invoice_structured_hit = sum(1 for k in invoice_structured_terms if k in src) >= 2
    notice_structured_hit = sum(1 for k in notice_structured_terms if normalize_text(k) in src_norm) >= 1

    # 来源路径优先纠偏：优先利用目录语义，避免 OCR 噪声误判。
    if source_contract_bucket and not source_notice_bucket and not source_photo_bucket:
        contract_kw_hit = True
        notice_kw_hit = False
    if source_invoice_bucket and not source_notice_bucket:
        invoice_kw_hit = True
        notice_kw_hit = False
    if source_photo_bucket and not (source_contract_bucket or source_invoice_bucket or source_notice_bucket):
        notice_kw_hit = False
        contract_kw_hit = False
        invoice_kw_hit = False
        if source_promo_bucket:
            poster_kw_hit = True
        elif packaging_kw_hit:
            scene_kw_hit = False
        else:
            scene_kw_hit = True

    procedure_notice_terms = [
        "提供商标使用证据", "不使用正当理由", "连续三年不使用", "撤销连续三年不使用",
        "答辩条码", "辩绑码", "须知", "申请号", "受理",
    ]
    procedure_doc_terms = ["通知书", "通知", "须知", "答辩通知", "补正通知", "决定书", "送达"]
    proc_hit_count = sum(
        1 for k in procedure_notice_terms
        if (normalize_text(k) in src_norm or normalize_text(k) in origin_norm)
    )
    promo_dominant_hit = (
        any(k in src for k in PROMOTION_MATERIAL_KEYWORDS)
        or any(k in name for k in PROMOTION_MATERIAL_KEYWORDS)
        or source_promo_bucket
    )
    procedure_authority_hit = any(normalize_text(k) in src_norm for k in ["国家知识产权局", "商标局", "商标评审", "国知局", "知识产权局"])
    official_notice_hit = (
        (
            source_notice_bucket
            or procedure_authority_hit
            or (
                (
                    ("通知书" in name or "通知" in name or "须知" in name)
                    or any(normalize_text(k) in src_norm for k in procedure_doc_terms)
                    or any(normalize_text(k) in name_norm for k in procedure_doc_terms)
                    or any(normalize_text(k) in origin_norm for k in procedure_doc_terms)
                )
                and proc_hit_count >= 3
            )
        )
        and not (source_contract_bucket or source_invoice_bucket or source_photo_bucket or source_promo_bucket)
        and notice_structured_hit
    )
    if promo_dominant_hit and not procedure_authority_hit and proc_hit_count < 3:
        notice_kw_hit = False
        notice_structured_hit = False
    if official_notice_hit:
        notice_kw_hit = True
        # 避免“须知中的合同/发票示例”把官方通知误判为合同/发票
        contract_kw_hit = False
        invoice_kw_hit = False

    photo_like = visual_ready and image_ratio >= 0.55 and text_len <= 220 and image_blocks >= 1
    poster_like_visual = visual_ready and image_ratio >= 0.68 and orientation == "portrait" and text_len <= 220
    poster_scene_hit = is_activity_poster_context("\n".join([src, name, origin]))
    document_like = (text_len >= 260) or (text_blocks >= 4) or (text_ratio >= 0.08)
    contract_primary = contract_kw_hit and ((document_like and contract_structured_hit) or source_contract_bucket)
    invoice_primary = invoice_kw_hit and ((document_like and invoice_structured_hit) or source_invoice_bucket or digit_ratio >= 0.30)

    form = "未知"
    conf = "LOW"
    signals: List[str] = []

    if notice_kw_hit and document_like and official_notice_hit:
        form = "通知"
        conf = "HIGH"
        signals = ["doc_layout", "notice_keywords"]
    elif contract_primary and not source_invoice_bucket:
        form = "合同"
        conf = "HIGH"
        signals = ["doc_layout", "contract_keywords"]
    elif invoice_primary:
        form = "发票"
        conf = "HIGH" if (document_like and digit_ratio >= 0.18) else "MEDIUM"
        signals = ["doc_layout" if document_like else "dense_digits", "invoice_keywords"]
    elif contract_primary:
        form = "合同"
        conf = "HIGH"
        signals = ["doc_layout", "contract_keywords"]
    elif poster_like_visual or poster_scene_hit or (poster_kw_hit and (photo_like or visual_ready)):
        form = "海报"
        conf = "HIGH" if poster_like_visual else "MEDIUM"
        if poster_like_visual:
            signals = ["poster_visual"]
        elif poster_scene_hit:
            signals = ["poster_scene_context"]
        else:
            signals = ["poster_keywords"]
    elif packaging_kw_hit and (photo_like or visual_ready):
        form = "包装"
        conf = "HIGH" if photo_like else "MEDIUM"
        signals = ["photo_layout" if photo_like else "visual_layout", "packaging_keywords"]
    elif scene_kw_hit and (photo_like or visual_ready):
        form = "现场照片"
        conf = "HIGH" if photo_like else "MEDIUM"
        signals = ["photo_layout" if photo_like else "visual_layout", "scene_keywords"]
    elif photo_like:
        form = "现场照片"
        conf = "MEDIUM"
        signals = ["photo_layout", f"orientation:{orientation}"]
    elif invoice_kw_hit and invoice_structured_hit:
        form = "发票"
        conf = "MEDIUM"
        signals = ["invoice_keywords"]
    elif contract_kw_hit and contract_structured_hit:
        form = "合同"
        conf = "MEDIUM"
        signals = ["contract_keywords"]
    elif notice_kw_hit and official_notice_hit:
        form = "通知"
        conf = "MEDIUM"
        signals = ["notice_keywords"]
    elif poster_kw_hit:
        form = "海报"
        conf = "MEDIUM"
        signals = ["poster_keywords"]
    elif packaging_kw_hit:
        form = "包装"
        conf = "MEDIUM"
        signals = ["packaging_keywords"]
    elif scene_kw_hit:
        form = "现场照片"
        conf = "MEDIUM"
        signals = ["scene_keywords"]

    if visual_ready:
        signals.append(f"img_ratio={image_ratio:.2f}")
        signals.append(f"text_ratio={text_ratio:.2f}")
        signals.append(f"text_blocks={text_blocks}")
        signals.append(f"image_blocks={image_blocks}")
    signals.append(f"text_len={text_len}")
    signals.append(f"digit_ratio={digit_ratio:.2f}")
    if orientation != "unknown":
        signals.append(f"orientation={orientation}")
    if origin:
        signals.append(f"origin={sanitize_component(origin_l, 42)}")

    return {
        "form": form,
        "confidence": conf,
        "signals": ",".join(unique_keep_order(signals, limit=10)),
        "layout": "visual+text" if visual_ready else "text-only",
    }


def classify_evidence(
    name: str,
    text: str,
    mark_name: str,
    respondent: str,
    form_profile: Optional[Dict[str, str]] = None,
    source_hint: str = "",
) -> Dict[str, str]:
    src = f"{name}\n{text[:2400]}"
    src_proc = f"{name}\n{text[:9000]}"
    source_hint = str(source_hint or "")
    src = f"{source_hint}\n{src}"
    src_proc = f"{source_hint}\n{src_proc}"
    name_l = (name or "").lower()
    src_l = src.lower()
    src_proc_norm = normalize_text(src_proc)
    form_profile = form_profile or {}
    form_type = str(form_profile.get("form", "未知") or "未知")
    form_conf = str(form_profile.get("confidence", "LOW") or "LOW")
    form_signals = str(form_profile.get("signals", "") or "")

    kind = "其他材料"
    targets = "T4"
    sj6 = "R,T"
    loop = "N"
    mark = ""
    subject = ""
    verifiable = "Y"
    note = "建议补充证据与要件映射。"

    display_score = sum(1 for k in DISPLAY_KEYWORDS if k in src) + sum(1 for k in REVIEW_KEYWORDS if k in src)
    tx_score = (
        sum(2 for k in TX_STRONG_KEYWORDS if k in src)
        + sum(1 for k in ORDER_KEYWORDS if k in src)
        + sum(1 for k in LOGISTICS_KEYWORDS if k in src)
    )
    procedure_authority_hit = any(
        normalize_text(k) in src_proc_norm
        for k in ["国家知识产权局", "商标局", "商标评审委员会"]
    )
    procedure_doc_hit = any(
        normalize_text(k) in src_proc_norm
        for k in ["答辩通知", "受理通知", "补正通知", "决定书", "送达", "通知书", "通知", "须知", "撤销"]
    )
    procedure_scene_terms = [
        "提供商标使用证据", "不使用正当理由", "连续三年不使用", "撤销连续三年不使用",
        "2个月内", "两个月内", "指定期间", "答辩条码", "辩绑码", "申请号", "受理",
    ]
    procedure_scene_hit_count = sum(
        1 for k in procedure_scene_terms
        if normalize_text(k) in src_proc_norm
    )
    is_official_procedure = (
        (procedure_authority_hit and (procedure_doc_hit or procedure_scene_hit_count >= 1))
        or (procedure_doc_hit and procedure_scene_hit_count >= 2)
    )
    is_exhibit_contract = any(k in src for k in ["展位销售合同", "展位合同", "参展合同", "展位费", "展台搭建合同", "参展协议"])
    is_exhibit_ad = any(k in src for k in EXHIBIT_AD_KEYWORDS)
    brochure_name_hit = any(k in (name or "") for k in BROCHURE_KEYWORDS)
    is_brochure = brochure_name_hit or any(k in src for k in BROCHURE_KEYWORDS) or any(k.lower() in src_l for k in BROCHURE_KEYWORDS)
    is_test_report = any(k in src for k in TEST_REPORT_KEYWORDS)
    is_trade_doc = (
        any(k in src for k in PAYMENT_KEYWORDS + CONTRACT_KEYWORDS + INVOICE_KEYWORDS)
        or any(k in name_l for k in ["dzfp", "invoice", "fp_", "fapiao"])
    ) and (
        tx_score >= 2
        or any(k in src for k in ["发票", "增值税", "合同", "订单编号", "价税合计", "银行流水", "电子回单"])
        or any(k in name_l for k in ["dzfp", "invoice", "fp_", "fapiao"])
    )
    is_form_notice = form_type == "通知"
    is_form_contract = form_type == "合同"
    is_form_invoice = form_type == "发票"
    is_form_poster = form_type == "海报"
    is_form_packaging = form_type == "包装"
    is_form_scene_photo = form_type == "现场照片"
    source_norm = normalize_text(source_hint)
    source_invoice_bucket = any(k in source_hint for k in ["发票", "增值税", "票据"])
    source_contract_bucket = any(k in source_hint for k in ["合同", "协议"]) or (
        "合同及发票" in source_hint and not source_invoice_bucket
    )
    source_notice_bucket = any(k in source_hint for k in ["通知", "须知", "答辩", "提供使用证据"])
    source_photo_bucket = any(k in source_hint for k in ["照片", "实拍", "现场", "图片", "商标照片"])
    source_promo_bucket = any(k in source_hint for k in ["海报", "易拉宝", "宣传", "画册", "招募", "邀请函"])
    source_mark_image_bucket = any(k in source_hint for k in ["商标图样", "图样", "logo"]) or ("商标档案" in source_hint and source_photo_bucket)
    promo_material_hit = (
        is_exhibit_ad
        or is_brochure
        or is_form_poster
        or source_promo_bucket
        or any(k in src for k in PROMOTION_MATERIAL_KEYWORDS)
    )
    procedure_strength = int(procedure_authority_hit) + int(procedure_doc_hit) + (1 if procedure_scene_hit_count >= 2 else 0)

    # 来源路径优先纠偏：避免OCR噪声把合同/发票误判为程序文件。
    if source_contract_bucket and not source_notice_bucket:
        is_trade_doc = True
        if source_invoice_bucket:
            is_form_invoice = True
        else:
            is_form_contract = True
    if source_notice_bucket and not source_contract_bucket and not source_invoice_bucket:
        is_official_procedure = True
        is_form_notice = True
    if source_photo_bucket and not (source_contract_bucket or source_invoice_bucket or source_notice_bucket):
        is_form_scene_photo = True
    if source_promo_bucket and not (source_contract_bucket or source_invoice_bucket):
        is_form_poster = True

    if promo_material_hit and procedure_strength < 3 and not source_notice_bucket:
        is_official_procedure = False
        is_form_notice = False

    # 照片/宣传来源默认不进入程序/交易，除非有强结构字段支撑。
    strong_trade_structured = any(
        k in src for k in [
            "甲方", "乙方", "合同编号", "签订日期", "履行期限",
            "发票代码", "发票号码", "开票日期", "价税合计", "税额",
            "订单编号", "支付成功", "物流单号",
        ]
    )
    if (is_form_notice or promo_material_hit) and not strong_trade_structured and not (source_contract_bucket or source_invoice_bucket):
        is_trade_doc = False
        is_form_contract = False
        is_form_invoice = False

    if source_photo_bucket and not (source_contract_bucket or source_invoice_bucket):
        if not strong_trade_structured:
            is_trade_doc = False
            is_form_contract = False
            is_form_invoice = False
        if not source_notice_bucket:
            is_official_procedure = False
            is_form_notice = False

    if (is_form_scene_photo or is_form_poster or is_form_packaging) and not source_notice_bucket:
        is_official_procedure = False

    if is_exhibit_contract:
        kind = "其他材料"
        targets = "T2,T3,T5"
        sj6 = "R,L"
        mark = "Y"
        subject = "Y"
        note = "展位/参展合同材料，用于证明展会宣传、展位展示及相关经营推广场景。"
    elif source_mark_image_bucket or any(k in src for k in ["商标图样", "图样", "LOGO", "logo"]):
        kind = "其他材料"
        targets = "T2,T5"
        sj6 = "R,L"
        mark = "Y"
        subject = "N"
        note = "商标图样/标识展示材料，用于证明标识呈现方式及使用场景。"
    elif (is_official_procedure or is_form_notice) and not (source_contract_bucket or source_invoice_bucket) and not promo_material_hit and not is_test_report:
        kind = "程序文件"
        targets = "T4,T5"
        sj6 = "A,R,C,T,L"
        mark = "N"
        subject = "Y"
        note = "程序性材料，可证明撤三程序、指定期间及送达事实。"
    elif is_test_report:
        kind = "其他材料"
        targets = "T2,T3,T5"
        sj6 = "R,C,L"
        mark = "Y"
        subject = "N"
        note = "检测/检验报告材料，用于证明商品属性、标识对应及质量规格，不作为核心使用时间证据。"
    elif (is_form_invoice or is_form_contract or is_trade_doc) and (not is_official_procedure):
        kind = "交易凭证"
        targets = "T3,T4,T5,T6"
        sj6 = "A,R,C,T,L,X"
        loop = "Y"
        mark = "N"
        subject = "N"
        note = "用于证明合同-发票-支付等交易链条与时间线索。"
    elif is_exhibit_ad or is_form_poster:
        kind = "其他材料"
        targets = "T2,T3,T5"
        sj6 = "R,L"
        mark = "Y"
        subject = "Y"
        note = "展会/广告宣传材料，用于证明展览宣传场景中的商标展示与商品对应。"
    elif is_brochure and (brochure_name_hit or (not any(k in src for k in ["营业执照", "统一社会信用代码", "注册证", "授权书", "许可证"]))):
        kind = "其他材料"
        targets = "T2,T3,T5"
        sj6 = "R,L"
        mark = "Y"
        subject = "Y"
        note = "宣传册/画册材料，用于证明品牌宣传与商品展示场景。"
    elif is_form_packaging:
        kind = "商品展示页"
        targets = "T2,T3,T5"
        sj6 = "R,L"
        mark = "Y"
        subject = "N"
        note = "包装/标签展示材料，用于证明标识呈现与商品对应关系。"
    elif is_form_scene_photo:
        kind = "其他材料"
        targets = "T5"
        sj6 = "R,L"
        mark = "N"
        subject = "N"
        note = "现场照片材料，用于证明服务活动现场或实际经营场景。"
    elif any(k in src for k in ["商标注册证", "商标证书", "商标公告", "核定使用商品", "专用权期限", "注册号"]):
        kind = "资质主体证明"
        targets = "T1,T2,T3,T5"
        sj6 = "A,R,C,T,L"
        mark = "Y"
        subject = "Y"
        note = "用于证明商标权属、标识对应及核定商品/服务范围。"
    elif any(k in src for k in ["资质", "营业执照", "授权", "登记", "许可证"]):
        kind = "资质主体证明"
        targets = "T1,T5"
        sj6 = "A,R,C,T,L"
        mark = "N"
        subject = "Y"
        note = "证明主体资质与身份，需与商标使用证据联动。"
    elif (
        any(k in name for k in ["工厂", "厂区", "车间", "仓库", "库存", "实拍", "现场"])
        and not any(k in src for k in ["商品详情", "商品评价", "店铺", "订单", "商城", "交易"])
    ):
        kind = "其他材料"
        targets = "T5"
        sj6 = "R,L"
        mark = "N"
        subject = "N"
        note = "用于辅助证明现场与实际经营场景。"
    elif (
        any(k in name for k in ["微信图片", "截图", "拍摄", "照片"])
        or any(k in name_l for k in ["img", "image", ".jpg", ".jpeg", ".png"])
    ) and any(k in src for k in ["工厂", "厂区", "车间", "仓库", "库存", "生产", "设备", "包装", "装箱", "实拍", "现场"]):
        kind = "其他材料"
        targets = "T5"
        sj6 = "R,L"
        mark = "N"
        subject = "N"
        note = "用于证明活动现场或经营场景中的实际使用状态。"
    elif (
        any(k in name for k in ["微信图片", "截图"])
        or any(k in name_l for k in ["img", "image", ".jpg", ".jpeg", ".png"])
    ) and any(k in src for k in ["商品", "详情", "店铺", "页面", "评价", "宝贝", "商标"]):
        kind = "商品展示页"
        targets = "T2,T3,T5"
        sj6 = "R,T,L"
        mark = "Y"
        subject = "Y"
        note = "用于证明商品展示内容、标识呈现及经营页面场景。"
    elif any(k in src for k in PAYMENT_KEYWORDS + CONTRACT_KEYWORDS + INVOICE_KEYWORDS):
        kind = "交易凭证"
        targets = "T3,T4,T5,T6"
        sj6 = "A,R,C,T,L,X"
        loop = "Y"
        mark = "N"
        subject = "N"
        note = "用于证明合同-发票-支付等交易链条与时间线索。"
    elif display_score >= tx_score and display_score >= 2:
        kind = "线上店铺展示"
        targets = "T1,T2,T3,T4,T5"
        sj6 = "R,T,L"
        mark = "Y"
        subject = "Y"
        note = "网店商品展示/评价页面材料，建议补后台导出或订单凭证增强可核验性。"
    elif tx_score >= 3:
        kind = "交易凭证"
        targets = "T3,T4,T5,T6"
        sj6 = "A,R,C,T,L,X"
        loop = "Y"
        mark = "Y"
        subject = "Y"
        note = "可形成交易闭环证据。"
    elif any(k in src for k in ["手机端", "电脑端", "店铺", "详情页", "链接", "截图", "页面"]):
        kind = "线上店铺展示"
        targets = "T1,T2,T3,T4,T5"
        sj6 = "R,T,L"
        mark = "Y"
        subject = "Y"
        note = "建议补后台导出记录或时间戳增强可核验性。"
    elif any(k in src for k in ["浴巾", "毛巾", "毛毯", "凉席", "床", "产品", "商品"]):
        kind = "商品展示页"
        targets = "T2,T3,T4,T5"
        sj6 = "R,T,L"
        mark = "Y"
        subject = "Y"
        note = "建议补订单、支付、物流单据形成闭环。"

    if mark_name and mark_name in src:
        mark = "Y"
    if respondent and respondent in src:
        subject = "Y"
    if form_type != "未知":
        note = f"{note}（表现形式识别：{form_type}，置信度{form_conf}）"

    return {
        "Type": kind,
        "Proof Target (T1-T6)": targets,
        "SJ-6 (A/R/C/T/L/X)": sj6,
        "Commercial Loop (Y/N)": loop,
        "Mark Shown (Y/N)": mark,
        "Subject Match (Y/N)": subject,
        "Original/Verifiable (Y/N)": verifiable,
        "Form Type": form_type,
        "Form Confidence": form_conf,
        "Form Signals": form_signals,
        "Risk Notes": note,
    }


def classify_direct_bind_profile(label: str) -> Dict[str, str]:
    src = str(label or "")
    if any(k in src for k in ["答辩通知", "通知书", "受理", "送达", "决定", "D1", "notice", "官方通知"]):
        return {
            "Type": "程序文件",
            "Proof Target (T1-T6)": "T4,T5",
            "SJ-6 (A/R/C/T/L/X)": "A,R,C,T,L",
            "Commercial Loop (Y/N)": "N",
            "Mark Shown (Y/N)": "N",
            "Subject Match (Y/N)": "Y",
            "Original/Verifiable (Y/N)": "Y",
            "Inferred Purpose": "证明撤三程序节点、通知送达与指定期间审查背景。",
        }
    if any(k in src for k in ["主体资格", "营业执照", "主体", "资质", "D2", "subject"]):
        return {
            "Type": "资质主体证明",
            "Proof Target (T1-T6)": "T1,T5",
            "SJ-6 (A/R/C/T/L/X)": "A,R,C,T,L",
            "Commercial Loop (Y/N)": "N",
            "Mark Shown (Y/N)": "N",
            "Subject Match (Y/N)": "Y",
            "Original/Verifiable (Y/N)": "Y",
            "Inferred Purpose": "证明注册人主体资格、授权关系或持续经营主体信息。",
        }
    if any(k in src for k in ["委托书", "授权书", "D3", "proxy"]):
        return {
            "Type": "资质主体证明",
            "Proof Target (T1-T6)": "T1,T5",
            "SJ-6 (A/R/C/T/L/X)": "A,R,C,T,L",
            "Commercial Loop (Y/N)": "N",
            "Mark Shown (Y/N)": "N",
            "Subject Match (Y/N)": "Y",
            "Original/Verifiable (Y/N)": "Y",
            "Inferred Purpose": "证明代理授权关系及主体办理权限。",
        }
    if any(k in src for k in ["照片", "D4", "photos"]):
        return {
            "Type": "其他材料",
            "Proof Target (T1-T6)": "T5",
            "SJ-6 (A/R/C/T/L/X)": "R,L",
            "Commercial Loop (Y/N)": "N",
            "Mark Shown (Y/N)": "N",
            "Subject Match (Y/N)": "N",
            "Original/Verifiable (Y/N)": "Y",
            "Inferred Purpose": "证明现场、环境、实物展示等实际使用场景。",
        }
    return {
        "Type": "其他材料",
        "Proof Target (T1-T6)": "T5",
        "SJ-6 (A/R/C/T/L/X)": "R,L",
        "Commercial Loop (Y/N)": "N",
        "Mark Shown (Y/N)": "N",
        "Subject Match (Y/N)": "N",
        "Original/Verifiable (Y/N)": "Y",
        "Inferred Purpose": "作为补强材料，证明与争议商标使用相关的辅助事实。",
    }


def targets_to_cn(s: str) -> str:
    src = (s or "").strip()
    codes = [c for c in ["T1", "T2", "T3", "T4", "T5", "T6"] if c in src]
    if not codes:
        return src
    return "、".join([TARGET_CN.get(c, c) for c in codes])


def unique_keep_order(items: List[str], limit: int = 5) -> List[str]:
    out: List[str] = []
    seen = set()
    for it in items:
        x = (it or "").strip()
        if not x or x in seen:
            continue
        out.append(x)
        seen.add(x)
        if len(out) >= limit:
            break
    return out


def extract_key_lines(text: str, mark_name: str, limit: int = 5) -> List[str]:
    scored: List[Tuple[int, str]] = []
    for ln in (text or "").splitlines():
        line = ln.strip()
        if len(line) < 4:
            continue

        score = 0
        has_date = bool(DATE_RE.search(line))
        if has_date:
            score += 2
        if mark_name and mark_name in line:
            score += 4
        if "商标" in line or "标识" in line:
            score += 2
        if any(k in line for k in ["订单", "支付", "物流", "签收", "发票", "合同", "销售", "交易", "店铺", "链接", "后台"]):
            score += 1

        if score <= 0:
            continue
        if len(line) > 120:
            line = line[:120] + "..."
        scored.append((score, line))

    scored.sort(key=lambda x: (-x[0], len(x[1])))
    return unique_keep_order([x[1] for x in scored], limit=limit)


def extract_mark_time_lines(text: str, mark_name: str, limit: int = 3) -> List[str]:
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    out: List[str] = []

    for ln in lines:
        if not DATE_RE.search(ln):
            continue
        if mark_name and mark_name in ln:
            out.append(ln)

    if len(out) < limit:
        for ln in lines:
            if not DATE_RE.search(ln):
                continue
            if "商标" in ln and ln not in out:
                out.append(ln)
            if len(out) >= limit:
                break

    return unique_keep_order(out, limit=limit)


def sanitize_component(s: str, max_len: int = 18) -> str:
    t = INVALID_FILENAME_RE.sub("_", s or "")
    t = re.sub(r"\s+", "", t)
    t = re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9._-]", "", t)
    t = t.strip("._-")
    if not t:
        return "证据"
    if len(t) > max_len:
        return t[:max_len]
    return t


def pick_evidence_topic_token(key_lines: List[str], mark_name: str) -> str:
    """
    从OCR关键词中提取“可验证”的证据主题词，避免只使用商标名导致证据名过于笼统。
    """
    if not key_lines:
        return sanitize_component(mark_name, 12) if mark_name and "待补" not in mark_name else "证据"

    stop_words = [
        "纯爱", "本色", "商标", "店铺", "页面", "详情", "商品", "评价", "客户",
        "时间", "图片", "截图", "旗舰店", "官方", "首页", "手机端", "电脑端",
    ]
    prefer_kw = [
        "卷纸", "抽纸", "卫生纸", "竹浆", "纸巾", "证书", "公告", "通知书", "转账", "订单",
        "工厂", "库存", "车间", "仓库", "实拍", "合同", "发票",
    ]

    candidates: List[str] = []
    for ln in key_lines[:4]:
        tokens = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9·-]{2,24}", ln)
        for tok in tokens:
            t = tok.strip("·- ")
            if not t:
                continue
            if mark_name and mark_name in t and len(t) <= max(4, len(mark_name) + 2):
                continue
            if any(sw == t for sw in stop_words):
                continue
            if any(sw in t for sw in ["时间不明", "待核验", "核验", "映射依据"]):
                continue
            candidates.append(t)

    for kw in prefer_kw:
        for c in candidates:
            if kw in c:
                return sanitize_component(c, 16)

    if candidates:
        candidates = sorted(candidates, key=lambda x: (len(x), x), reverse=True)
        return sanitize_component(candidates[0], 16)

    return sanitize_component(mark_name, 12) if mark_name and "待补" not in mark_name else "证据"


def infer_product_phrase(text: str, key_lines: List[str]) -> str:
    """
    从OCR文本中提取“商品名短语”，仅用于证据名称展示优化。
    """
    src = "\n".join((key_lines or [])[:6]) + "\n" + (text or "")[:5000]
    src = re.sub(r"\s+", " ", src)
    # 先抓包含规格单位的短语，避免只取到“商品详情”等泛词
    unit_pats = [
        r"([\u4e00-\u9fa5A-Za-z0-9·\-/]{4,40}(?:\d+\s*(?:卷|提|包|箱|g|kg|ml|L)))",
        r"([\u4e00-\u9fa5A-Za-z0-9·\-/]{4,40}(?:卫生纸|卷纸|抽纸|纸巾|竹浆纸|本色纸))",
    ]
    noise = ["商品详情", "商品评价", "进入店铺", "首页", "店铺", "商城", "订单", "时间不明"]
    def _normalize_product_phrase(s: str) -> str:
        t = s.strip(" -_，。；;:：")
        t = re.sub(r"\s+", "", t)
        # 常见OCR缺字补全
        t = t.replace("浆纸扁卷", "竹浆纸扁卷")
        t = t.replace("浆纸卷", "竹浆纸卷")
        t = t.replace("本色纸卫", "本色纸卫生")
        t = t.replace("卫生纸卷纸", "卫生纸卷纸")
        # 去掉明显噪声前缀
        t = re.sub(r"^(商品详情|商品评价|店铺|进入店铺)+", "", t)
        # 控制长度，保持可读
        return sanitize_component(t, 24)

    for pat in unit_pats:
        for m in re.finditer(pat, src):
            cand = m.group(1).strip(" -_，。；;:：")
            if len(cand) < 4:
                continue
            if any(n in cand for n in noise):
                continue
            norm = _normalize_product_phrase(cand)
            if norm:
                return norm

    # 兜底：从关键词行抽取商品名词条
    for ln in (key_lines or [])[:8]:
        for kw in ["卫生纸", "卷纸", "抽纸", "纸巾", "竹浆纸", "本色纸"]:
            if kw in ln:
                toks = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9·\-/]{4,30}", ln)
                for t in toks:
                    if kw in t and not any(n in t for n in noise):
                        norm = _normalize_product_phrase(t)
                        if norm:
                            return norm
    return ""


def infer_shop_source(path_or_name: str, text: str, key_lines: List[str]) -> str:
    src = "\n".join([path_or_name or "", *(key_lines or []), (text or "")[:3000]])
    generic = {"进入店铺", "店铺", "商品详情", "商品评价", "首页", "商城", "详情页"}
    patterns = [
        r"([\u4e00-\u9fa5A-Za-z0-9·]{2,20}(?:商城|旗舰店|专营店|专卖店|店铺))",
        r"([\u4e00-\u9fa5A-Za-z0-9·]{2,16}淘宝)",
        r"(淘宝)",
        r"(天猫)",
        r"(京东)",
        r"(拼多多)",
    ]
    for pat in patterns:
        m = re.search(pat, src)
        if m:
            candidate = sanitize_component(m.group(1), 18)
            if candidate and candidate not in generic:
                return candidate
    # 路径兜底：优先识别平台来源
    if "淘宝" in src:
        return "淘宝店铺"
    if "天猫" in src:
        return "天猫店铺"
    if "京东" in src:
        return "京东店铺"
    if "拼多多" in src:
        return "拼多多店铺"
    if "本色商城" in src:
        return "本色商城"
    return ""


def build_renamed_filename(
    idx: int,
    kind: str,
    date_start: str,
    mark_name: str,
    key_lines: List[str],
    source_label: str = "",
    product_label: str = "",
    preferred_topic: str = "",
) -> str:
    kind_tag = EVIDENCE_TYPE_ALIAS.get(kind, "其他")
    date_tag = date_start or "时间不明"

    token = sanitize_component(preferred_topic, 36) if preferred_topic else pick_evidence_topic_token(key_lines, mark_name)
    if kind == "程序文件":
        # 程序类证据固定采用程序主题，避免“须知中的合同/发票示例词”污染命名。
        token = "提供使用证据通知书"
    if source_label and kind == "线上店铺展示":
        token = sanitize_component(f"{source_label}_{token}", 24)
    if product_label and kind in ("线上店铺展示", "商品展示页", "交易凭证"):
        token = sanitize_component(f"{token}_{product_label}", 30)
    if token.endswith("有限公"):
        token = f"{token}司"
    if token.endswith("有限责任公"):
        token = f"{token}司"
    if not token:
        token = "证据"

    name = sanitize_component(f"E{idx:03d}_{date_tag}_{kind_tag}_{token}", 84)
    return f"{name}.pdf"


def _parse_row_evidence_seq(row: Dict[str, Any]) -> int:
    try:
        n = int(str(row.get("No", "")).strip())
        if n > 0:
            return n
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
    eid = str(row.get("Evidence ID", "") or "").strip()
    m = re.search(r"(\d+)$", eid)
    if m:
        try:
            n = int(m.group(1))
            if n > 0:
                return n
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
    name = str(row.get("_renamed_name", "") or row.get("Evidence Name", "")).strip()
    m = re.match(r"^E(\d{3})_", name)
    if m:
        try:
            n = int(m.group(1))
            if n > 0:
                return n
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
    return 0


def _suggest_evidence_name(row: Dict[str, Any], mark_name: str) -> str:
    seq = _parse_row_evidence_seq(row) or 1
    kind = str(row.get("Type", "其他材料") or "其他材料")
    date_start = str(row.get("Evidence Date Start", "") or "")
    topic = sanitize_component(str(row.get("Inferred Proof Name", "") or ""), 28)
    if (not topic) or (topic in EVIDENCE_NAME_GENERIC_TOPICS):
        src_name = Path(str(row.get("_src_name", "") or row.get("Original File Name", "") or "")).stem
        topic = sanitize_component(src_name, 28)
    if (not topic) or (topic in EVIDENCE_NAME_GENERIC_TOPICS):
        topic = "待人工确认主题"
    kind_tag = EVIDENCE_TYPE_ALIAS.get(kind, "其他")
    date_tag = date_start if parse_iso_date(date_start) else "时间不明"
    return f"E{seq:03d}_{date_tag}_{kind_tag}_{topic}.pdf"


def validate_evidence_name_quality(
    rows: List[Dict[str, Any]],
    mark_name: str,
    out_dir: Path,
) -> Dict[str, Any]:
    checks: List[Dict[str, Any]] = []
    pending_rows: List[Dict[str, str]] = []
    warnings: List[str] = []
    errors: List[str] = []

    for row in rows:
        eid = str(row.get("Evidence ID", "") or "")
        current_name = str(row.get("_renamed_name", "") or row.get("Evidence Name", "") or "").strip()
        issues: List[str] = []

        if not current_name:
            issues.append("证据名称为空")
        else:
            m = NAME_TEMPLATE_RE.match(current_name)
            if not m:
                issues.append("命名未命中模板：E###_日期_类型_主题.pdf")
            else:
                prefix, date_tag, kind_tag, topic_raw = m.groups()
                seq = _parse_row_evidence_seq(row)
                if seq > 0:
                    expected_prefix = f"E{seq:03d}"
                    if prefix != expected_prefix:
                        issues.append(f"编号前缀不一致：{prefix} != {expected_prefix}")

                expected_kind_tag = EVIDENCE_TYPE_ALIAS.get(str(row.get("Type", "") or ""), "")
                if expected_kind_tag and kind_tag != expected_kind_tag:
                    issues.append(f"类型标签不一致：{kind_tag} != {expected_kind_tag}")

                if date_tag != "时间不明" and not parse_iso_date(date_tag):
                    issues.append(f"日期段格式不合法：{date_tag}")

                topic = sanitize_component(topic_raw, 60).strip("._- ")
                topic_norm = normalize_text(topic)
                if not topic:
                    issues.append("主题词为空")
                else:
                    if len(topic_norm) < 4:
                        issues.append("主题词过短")
                    if topic in EVIDENCE_NAME_GENERIC_TOPICS:
                        issues.append(f"主题词过泛：{topic}")
                    if re.fullmatch(r"\d+", topic):
                        issues.append("主题词不能为纯数字")
                    for bad in sorted(EVIDENCE_NAME_BAD_PHRASES, key=len, reverse=True):
                        if bad and bad in topic:
                            issues.append(f"主题词含噪声词：{bad}")
                            break

        status = "PASS" if not issues else "FAIL"
        row["Name Validation"] = "PASS" if status == "PASS" else "待人工确认"
        row["Name Validation Issues"] = "；".join(issues)

        check = {
            "evidence_id": eid,
            "evidence_name": current_name,
            "status": status,
            "issues": issues,
        }
        checks.append(check)

        if issues:
            pending_rows.append({
                "序号": str(row.get("No", "")),
                "证据ID": eid,
                "当前证据名": current_name,
                "证据类型": str(row.get("Type", "") or ""),
                "推测证明名称": str(row.get("Inferred Proof Name", "") or ""),
                "对应证明目的": str(row.get("Inferred Purpose", "") or ""),
                "问题": "；".join(issues),
                "建议证据名": _suggest_evidence_name(row, mark_name),
            })

    pending_count = len(pending_rows)
    status = "PASS" if pending_count == 0 else "FAIL"
    if pending_count:
        errors.append(f"命名不达标证据数：{pending_count}")
    else:
        warnings.append("全部证据名称命中模板且通过质量检查")

    validation_path = out_dir / "evidence_name_validation.json"
    pending_path: Path = out_dir / "待人工确认_证据命名.xlsx"
    pending_payload = pending_rows if pending_rows else [{
        "序号": "",
        "证据ID": "",
        "当前证据名": "",
        "证据类型": "",
        "推测证明名称": "",
        "对应证明目的": "",
        "问题": "",
        "建议证据名": "",
    }]
    try:
        pd.DataFrame(pending_payload).to_excel(pending_path, index=False)
    except Exception:
        pending_path = out_dir / "待人工确认_证据命名.csv"
        pd.DataFrame(pending_payload).to_csv(pending_path, index=False, encoding="utf-8-sig")

    payload = {
        "status": status,
        "errors": errors,
        "warnings": warnings,
        "pending_count": pending_count,
        "validation_rule": "evidence_name.strict.v1",
        "pending_path": str(pending_path),
        "checks": checks,
    }
    validation_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    payload["validation_path"] = str(validation_path)
    return payload


def confidence_from_dates(
    dates: List[dt.date],
    mark_time_lines: List[str],
    review_dates: Optional[List[dt.date]] = None,
    meta_used: bool = False,
    time_channels: Optional[List[str]] = None,
) -> str:
    review_dates = review_dates or []
    time_channels = time_channels or []
    strong_scene_hit = any(ch in ("评价", "交易", "物流") for ch in time_channels)
    if len(dates) >= 1 and strong_scene_hit and not meta_used:
        return "HIGH"
    if len(dates) >= 2 and (mark_time_lines or review_dates):
        return "HIGH"
    if len(dates) >= 2:
        return "HIGH"
    if len(dates) == 1 and review_dates:
        return "MEDIUM"
    if len(dates) == 1 and meta_used:
        return "LOW"
    if len(dates) == 1:
        return "MEDIUM"
    return "LOW"


def should_time_gate_evaluate(
    kind: str,
    text: str,
    time_channels: List[str],
    key_lines: List[str],
    anchor_type: str = "uncertain",
    web_verifiability: str = "",
) -> bool:
    src = "\n".join([text or "", *(key_lines or [])])
    # 程序文件不得计入指定期间覆盖，只作为程序背景证据。
    if kind == "程序文件":
        return False
    # 交易履约类证据允许计入覆盖。
    if kind == "交易凭证":
        return True
    # 网页类证据仅在可核验时间来源成立时纳入覆盖。
    if kind == "线上店铺展示":
        if web_verifiability == "S1":
            return True
        if any(ch in ("评价", "交易", "物流") for ch in (time_channels or [])):
            return True
        if any(k in src for k in WEB_VERIFIABILITY_STRONG):
            return True
        if web_verifiability == "S2" and anchor_type in ("system_generated", "transaction_date"):
            return True
        return False
    if kind == "商品展示页":
        return anchor_type == "transaction_date"
    # 主体资质/现场照片/程序文件默认不纳入时间覆盖。
    return False


def mark_name_confidence(
    mark_name: str,
    text: str,
    filename: str,
    key_lines: List[str],
    mark_shown_flag: str,
) -> str:
    if not mark_name or "待补" in mark_name:
        return "UNKNOWN"
    src = "\n".join([filename or "", (text or "")[:8000], *(key_lines or [])])
    src_norm = normalize_text(src)
    mk_norm = normalize_text(mark_name)
    if mark_name in src or (mk_norm and mk_norm in src_norm):
        return "HIGH"
    # 容错：商标前2-3字命中，且分类器已判定“体现商标”
    prefix2 = mark_name[:2]
    prefix3 = mark_name[:3] if len(mark_name) >= 3 else mark_name
    if (prefix3 and prefix3 in src) or (prefix2 and prefix2 in src):
        return "MEDIUM" if mark_shown_flag == "Y" else "LOW"
    if mark_shown_flag == "Y":
        return "MEDIUM"
    return "LOW"


def infer_web_verifiability(kind: str, text: str, key_lines: List[str]) -> str:
    if kind not in ("线上店铺展示", "商品展示页"):
        return ""
    src = "\n".join([text or "", *(key_lines or [])])
    strong = sum(1 for k in WEB_VERIFIABILITY_STRONG if k in src)
    medium = sum(1 for k in WEB_VERIFIABILITY_MEDIUM if k in src)
    if strong >= 2:
        return "S1"
    if strong >= 1 or medium >= 2:
        return "S2"
    return "S3"


def normalize_goods_token(s: str) -> str:
    t = (s or "").strip()
    t = re.sub(r"[（(].*?[）)]", "", t)
    t = re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9]", "", t)
    t = t.replace("核定商品", "").replace("核定服务", "").replace("服务", "")
    return t


def parse_goods_canonical_list(goods_desc: str) -> List[str]:
    src = (goods_desc or "").replace("：", " ").replace(";", "；")
    parts = re.split(r"[；;、,，/\n]+", src)
    out: List[str] = []
    seen = set()
    for p in parts:
        t = normalize_goods_token(p)
        if len(t) < 2:
            continue
        if t in seen:
            continue
        seen.add(t)
        out.append(t)
    return out[:30]


def goods_similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    if a in b or b in a:
        return 1.0
    sa = set(a)
    sb = set(b)
    inter = len(sa & sb)
    union = len(sa | sb) or 1
    return inter / union


def infer_goods_match(
    canonical_goods: List[str],
    text: str,
    key_lines: List[str],
    evidence_name: str,
) -> Dict[str, str]:
    if not canonical_goods:
        return {"level": "G3", "item_index": "", "item_name": "", "basis": "未提供核定商品清单"}
    src = "\n".join([evidence_name or "", text or "", *(key_lines or [])])
    src_norm = normalize_goods_token(src)
    best_idx = -1
    best_name = ""
    best_score = 0.0
    for i, g in enumerate(canonical_goods, start=1):
        score = goods_similarity(normalize_goods_token(g), src_norm)
        if score > best_score:
            best_score = score
            best_idx = i
            best_name = g
    if best_idx <= 0:
        return {"level": "G3", "item_index": "", "item_name": "", "basis": "未识别到可比对商品词条"}
    if best_score >= 0.9:
        level = "G1"
    elif best_score >= 0.45:
        level = "G2"
    else:
        level = "G3"
    return {
        "level": level,
        "item_index": str(best_idx),
        "item_name": best_name,
        "basis": f"相似度{best_score:.2f}",
    }


def extract_entity_names(text: str, key_lines: List[str]) -> List[str]:
    src = "\n".join([text or "", *(key_lines or [])])
    out: List[str] = []
    seen = set()
    for m in COMPANY_RE.finditer(src):
        name = m.group(1).strip()
        if len(name) < 4:
            continue
        if name in seen:
            continue
        seen.add(name)
        out.append(name)
    return out[:10]


def normalize_entity_name(name: str) -> str:
    return re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9]", "", name or "").lower()


def infer_entity_consistency(
    entities: List[str],
    respondent: str,
    text: str,
) -> Dict[str, str]:
    if not entities:
        return {"level": "E2", "basis": "未提取到稳定主体名称，需人工核验"}
    resp_norm = normalize_entity_name(respondent)
    if resp_norm:
        for e in entities:
            en = normalize_entity_name(e)
            if en and (resp_norm in en or en in resp_norm):
                return {"level": "E1", "basis": f"命中答辩主体：{e}"}
    auth_hit = any(k in (text or "") for k in ["授权", "许可", "委托", "经销", "代理", "代销"])
    if auth_hit:
        return {"level": "E2", "basis": f"识别到关联主体：{entities[0]}（含授权/委托语义）"}
    return {"level": "E3", "basis": f"主体不一致候选：{entities[0]}"}


def infer_evidence_taxonomy(
    cls: Dict[str, str],
    text: str,
    case_class: str = "",
    goods_services: str = "",
) -> Dict[str, str]:
    src = text or ""
    kind = str(cls.get("Type", "") or "")
    form = str(cls.get("Form Type", "") or "")
    property_case = is_property_service_case(case_class=case_class, goods_services=goods_services)

    cat = "2. 商品/服务样态展示类证据"
    objective = "证明商标实际用于商品或服务的市场展示。"
    proposition = "该证据显示涉案商标实质性附着于商品/服务展示中，可证明指定期间内的持续展示使用。"
    strength = "高"
    risk_tag = "需与交易或现场证据闭环"
    next_action = "建议补充交易或现场证据，形成完整事实链。"

    if any(k in src for k in ["停产", "停止经营", "不可抗力", "未使用", "无库存", "无法生产"]):
        cat = "12. 负面证据"
        objective = "证明未使用或存在正当不使用理由。"
        proposition = "该证据说明指定期间内未发生相关商业使用，或存在正当不使用理由。"
        strength = "中"
        risk_tag = "需核验客观性"
        next_action = "建议补充官方记录或第三方证明。"
    elif kind == "交易凭证":
        cat = "1. 商品/服务交易类证据"
        objective = "直接证明商标用于商业交易活动。"
        proposition = "该证据显示在指定期间内发生真实交易行为，并与核定商品/服务具有对应关系。"
        strength = "高"
        risk_tag = "需核验主体与金额对应"
        next_action = "建议补充合同-发票-支付/物流的闭环对应。"
    elif form in ("海报",) or is_activity_poster_context(src):
        cat = "4. 市场推广与广告类证据"
        objective = "辅助证明商标在市场推广行为中被使用。"
        if property_case:
            proposition = "该证据为在涉案物业组织相关活动的广告图片，用于佐证不动产/物业服务的持续经营推广。"
        else:
            proposition = "该证据显示经营主体组织活动并进行对外宣传推广，可辅助佐证持续使用状态。"
        strength = "中"
        risk_tag = "辅助证据"
        next_action = "建议与交易类或现场类证据组合使用。"
    elif kind in ("线上店铺展示", "商品展示页"):
        cat = "5. 数字痕迹及平台证据"
        objective = "证明商标在网络销售/服务链路中的数字化使用痕迹。"
        proposition = "该证据显示商标在网络平台中的展示与运营状态，可辅助证明指定期间内的使用痕迹。"
        strength = "中"
        risk_tag = "需补可核验时间来源"
        next_action = "建议补后台导出、订单编号或评价时间页。"
    elif form in ("现场照片", "包装"):
        cat = "3. 现场使用类证据（标识性使用）"
        objective = "证明商标在服务现场/经营场所的实际应用。"
        if property_case:
            proposition = "该证据展示涉案不动产/办公场所的实际运营场景，与核定服务交付现场相关联。"
        else:
            proposition = "该证据展示商标在经营场所或服务现场的实际使用，反映真实经营状态。"
        strength = "高"
        risk_tag = "需补时间锚点"
        next_action = "建议补带时间信息的同场景证据。"
    elif any(k in src for k in ["授权", "许可使用", "商标许可", "被授权", "许可合同"]):
        cat = "6. 商标许可/授权类证据"
        objective = "证明他人经许可使用涉案商标。"
        proposition = "该证据说明商标经权利人许可由第三方使用，并与指定期间内经营行为相关。"
        strength = "中"
        risk_tag = "需补许可使用人实际使用证据"
        next_action = "建议补许可方交易/展示证据。"
    elif any(k in src for k in TEST_REPORT_KEYWORDS):
        cat = "7. 市场调查与第三方报告类证据"
        objective = "间接证明市场中存在使用痕迹。"
        proposition = "该证据提供第三方检测/报告信息，可辅助印证商品属性及市场流通事实。"
        strength = "低-中"
        risk_tag = "辅助证据"
        next_action = "建议与交易证据结合。"
    elif kind == "资质主体证明":
        cat = "8. 法律/行政记录类证据"
        objective = "证明主体资质与业务合规状态。"
        proposition = "该证据显示经营主体及资质状态，与商标使用环境存在合规关联。"
        strength = "中"
        risk_tag = "需与使用证据映射"
        next_action = "建议补充对应期间的交易或展示证据。"
    elif any(k in src for k in ["台账", "工单", "交接", "内部", "预约单", "服务记录"]):
        cat = "9. 内部业务记录类证据"
        objective = "证明企业内部流转信息所反映的真实使用行为。"
        proposition = "该证据记录企业经营过程中的内部业务流转信息，可证明指定期间内经营活动存在。"
        strength = "中-高"
        risk_tag = "需外部证据佐证"
        next_action = "建议补对外交易或现场证据。"
    elif any(k in src for k in ["报价单", "协议附件", "售后", "沟通记录", "通信记录", "文案"]):
        cat = "10. 文书与交易文案类证据"
        objective = "佐证交易行为与商标使用界面。"
        proposition = "该证据显示商标在合同/文案/沟通界面中的使用，并与经营活动具有关联。"
        strength = "中"
        risk_tag = "需补交易发生证据"
        next_action = "建议补发票、付款或履约证据。"
    elif any(k in src for k in ["实地", "核查", "调研", "笔录", "现场核验"]):
        cat = "11. 线下调查证据"
        objective = "证明线下核查场景中的使用状态。"
        proposition = "该证据记录线下实地场景，可辅助验证指定期间内的使用事实。"
        strength = "中-高"
        risk_tag = "需确保时间地点可核验"
        next_action = "建议补拍摄时间与地点证明。"

    return {
        "category": cat,
        "objective": objective,
        "proposition": proposition,
        "strength": strength,
        "risk_tag": risk_tag,
        "next_action": next_action,
    }


def build_evidence_summary(
    cls: Dict[str, str],
    mark_name: str,
    respondent: str,
    date_start: str,
    date_end: str,
    raw_text: str,
    key_lines: List[str],
    mark_time_lines: List[str],
    inferred_name: str = "",
    inferred_purpose: str = "",
    case_class: str = "",
    goods_services: str = "",
    taxonomy: Optional[Dict[str, str]] = None,
) -> Tuple[str, str, str]:
    kind = cls["Type"]
    form_type = str(cls.get("Form Type", "") or "")
    property_case = is_property_service_case(case_class=case_class, goods_services=goods_services)
    poster_ctx = is_activity_poster_context("\n".join([raw_text or "", *(key_lines or [])]))
    time_txt = "时间不明"
    if date_start and date_end:
        time_txt = date_start if date_start == date_end else f"{date_start} ~ {date_end}"

    mark_flag = "体现商标" if cls.get("Mark Shown (Y/N)") == "Y" else "未明确体现商标"
    subject_flag = "主体一致" if cls.get("Subject Match (Y/N)") == "Y" else "主体待核验"
    core_line = (mark_time_lines[0] if mark_time_lines else (key_lines[0] if key_lines else "未提取到稳定关键词句"))
    if kind == "其他材料" and form_type in ("海报", "现场照片", "包装"):
        if poster_ctx:
            core_line = "见证据页面中的活动主题、组织信息及场地展示内容"
        elif form_type == "现场照片":
            core_line = "见证据页面中的场地、设施及运营环境展示内容"
        elif form_type == "包装":
            core_line = "见证据页面中的包装标识、商品名称及展示信息"
    if kind == "程序文件":
        proc_pool = list(key_lines or []) + list(mark_time_lines or [])
        proc_lines = [
            ln for ln in proc_pool
            if any(k in ln for k in ["通知", "须知", "撤销", "指定期间", "受理", "答辩", "送达", "申请", "商标局"])
        ]
        if proc_lines:
            core_line = proc_lines[0]

    title = inferred_name or kind
    form_part = f"；表现形式：{form_type}" if form_type and form_type != "未知" else ""
    taxonomy = taxonomy or {}
    tax_cat = str(taxonomy.get("category", "") or "")
    tax_objective = str(taxonomy.get("objective", "") or "")
    summary = f"{title}；证明目的：{inferred_purpose or tax_objective or '证明相关使用事实'}；时间锚点：{time_txt}{form_part}；{mark_flag}；{subject_flag}；重点语句：{core_line}"
    if tax_cat:
        summary = f"{title}；证据类别：{tax_cat}；证明目的：{inferred_purpose or tax_objective or '证明相关使用事实'}；时间锚点：{time_txt}{form_part}；{mark_flag}；{subject_flag}；重点语句：{core_line}"
    target_cn = targets_to_cn(cls.get("Proof Target (T1-T6)", ""))

    mark_hint = f"商标“{mark_name}”" if mark_name and "待补" not in mark_name else "目标商标"
    subject_hint = respondent if respondent and "待补" not in respondent else "使用主体"
    if str(taxonomy.get("proposition", "") or "").strip():
        base_fact = str(taxonomy.get("proposition", "")).strip()
    elif kind == "程序文件":
        base_fact = "用于说明程序节点与审查期间。"
    elif kind == "资质主体证明":
        base_fact = "用于说明主体资格、权属信息及核定商品/服务范围。"
    elif kind in ("线上店铺展示", "商品展示页"):
        base_fact = "用于说明商品页面展示、标识呈现及客户评价时间线索。"
    elif kind == "交易凭证":
        base_fact = "用于说明支付/交易场景及时间线索。"
    else:
        if form_type in ("海报", "现场照片") and poster_ctx:
            if property_case:
                base_fact = "用于说明在涉案物业组织相关活动的广告图片及运营推广场景。"
            else:
                base_fact = "用于说明活动组织及对外宣传推广场景（广告图片）。"
        elif form_type == "现场照片" and property_case:
            base_fact = "用于说明涉案不动产/办公场所的实际运营与管理场景。"
        elif form_type == "包装":
            base_fact = "用于说明商品包装展示、标识呈现及商品对应关系。"
        else:
            name_purpose = f"{inferred_name} {inferred_purpose}"
            if any(k in name_purpose for k in ["检测", "检验", "质检"]):
                base_fact = "用于说明商品质量规格、标识对应及检测结论等补强事实。"
            elif any(k in name_purpose for k in ["宣传册", "画册", "展会", "广告", "展位"]):
                base_fact = "用于说明品牌宣传、展会广告展示及经营推广场景。"
            else:
                base_fact = "用于说明现场或经营场景相关事实。"

    proof_fact = (
        f"{base_fact}"
        f"；时间锚点：{time_txt}"
        f"{form_part}"
        f"；要件：{target_cn}"
        f"；关键依据：{core_line}"
    )

    anchor_text = "；".join(mark_time_lines[:3]) if mark_time_lines else "；".join(key_lines[:3])
    return summary, proof_fact, anchor_text or "未提取到有效时间语句"


def build_mapping_basis(
    text: str,
    cls: Dict[str, str],
    key_lines: List[str],
    mark_time_lines: List[str],
    channels_used: List[str],
) -> str:
    src = text or ""
    targets = [t for t in ["T1", "T2", "T3", "T4", "T5", "T6"] if t in (cls.get("Proof Target (T1-T6)", "") or "")]
    target_cn = targets_to_cn(",".join(targets)) if targets else "未明确"
    time_basis = "；".join((mark_time_lines or key_lines)[:2]) or "未提取稳定时间语句"
    scene_basis = "、".join(channels_used) if channels_used else "基础"
    if cls.get("Type", "") == "程序文件":
        proc_lines = [
            ln for ln in (mark_time_lines or key_lines or [])
            if any(k in ln for k in ["通知", "须知", "撤销", "指定期间", "受理", "答辩", "送达", "申请", "商标局"])
        ]
        proc_basis = "；".join(proc_lines[:2]) if proc_lines else "程序通知载明的指定期间及举证要求"
        return f"要件命中={target_cn}；关键词命中=程序通知、指定期间锚点；时间命中={proc_basis}；通道=基础"

    kw_hits: List[str] = []
    form_type = str(cls.get("Form Type", "") or "")
    if form_type and form_type != "未知":
        kw_hits.append(f"表现形式:{form_type}")
    if any(k in src for k in ["店铺", "链接", "详情页", "页面", "手机端", "电脑端"]):
        kw_hits.append("店铺展示")
    if any(k in src for k in ["订单", "付款", "支付", "物流", "签收", "交易", "对账"]):
        kw_hits.append("交易闭环")
    if any(k in src for k in ["营业执照", "资质", "授权", "登记"]):
        kw_hits.append("主体资质")
    if not kw_hits:
        kw_hits = ["通用文本命中"]

    return f"要件命中={target_cn}；关键词命中={('、'.join(kw_hits))}；时间命中={time_basis}；通道={scene_basis}"


def infer_proof_name_and_purpose(
    cls: Dict[str, str],
    text: str,
    review_dates: List[dt.date],
    order_dates: List[dt.date],
    logistics_dates: List[dt.date],
    case_class: str = "",
    goods_services: str = "",
) -> Tuple[str, str]:
    kind = cls.get("Type", "") or "其他材料"
    form_type = str(cls.get("Form Type", "") or "")
    src = text or ""
    src_l = src.lower()
    src_norm = normalize_text(src)
    cls_note = str(cls.get("Risk Notes", "") or "")
    property_case = is_property_service_case(case_class=case_class, goods_services=goods_services)
    poster_ctx = is_activity_poster_context(src)

    procedure_doc_terms = ["通知书", "通知", "须知", "答辩通知", "受理通知", "补正通知", "决定书", "送达", "撤销"]
    procedure_scene_terms = [
        "提供商标使用证据", "不使用正当理由", "连续三年不使用", "撤销连续三年不使用",
        "2个月内", "两个月内", "指定期间", "答辩条码", "辩绑码", "申请号", "受理",
    ]
    proc_doc_hit = any(normalize_text(k) in src_norm for k in procedure_doc_terms)
    proc_scene_cnt = sum(1 for k in procedure_scene_terms if normalize_text(k) in src_norm)
    if kind == "程序文件" and proc_doc_hit and proc_scene_cnt >= 2:
        return (
            "提供使用证据通知书及须知",
            "证明撤三程序节点、送达事实及指定期间审查范围（程序材料，仅作参考线）。",
        )

    if kind == "交易凭证":
        promo_hit = any(k in src for k in PROMOTION_MATERIAL_KEYWORDS + ["讲座", "路演", "海报", "画册", "宣传册"])
        tx_structured_hit = any(k in src for k in ["甲方", "乙方", "合同编号", "发票代码", "发票号码", "价税合计", "订单编号", "支付成功"])
        if promo_hit and not tx_structured_hit:
            return (
                "宣传页面（待人工核验）",
                "页面以宣传展示为主，未提取稳定交易结构字段，建议人工复核证据属性。",
            )
        if any(k in src for k in CONTRACT_KEYWORDS):
            return (
                "购销合同页面",
                "证明交易关系、交易标的及交易场景。",
            )
        if any(k in src for k in INVOICE_KEYWORDS):
            return (
                "发票页面",
                "证明交易开票事实、交易标的及发生时间线索。",
            )
        if any(k in src for k in PAYMENT_KEYWORDS):
            return (
                "支付/转账记录页面",
                "证明付款行为与交易发生时间线索。",
            )
        if order_dates or logistics_dates:
            return (
                "交易订单及物流记录",
                "证明交易发生并形成订单、支付、物流等闭环链条。",
            )
        if form_type == "通知":
            return (
                "通知页面（待人工核验）",
                "页面含“通知”字样但未形成程序材料特征，建议人工复核证据属性。",
            )
        return (
            "交易记录页面",
            "证明存在交易行为并可与商品/服务形成对应。",
        )

    if kind == "线上店铺展示":
        if review_dates:
            return (
                "网店商品展示及客户评价页面",
                "证明商标在网店商品展示场景中持续使用，并以客户评价时间形成使用时间锚点。",
            )
        return (
            "网店商品展示页面",
            "证明商标在网店商品页面中的展示使用及商品/服务对应关系。",
        )

    if kind == "资质主体证明":
        return (
            "主体资质与身份材料",
            "证明使用主体身份及经营资格。",
        )

    if kind == "商品展示页":
        return (
            "网店商品展示页面",
            "证明网店商品展示功能及商标在商品展示场景中的使用。",
        )

    if kind == "程序文件":
        return (
            "程序通知与流程材料",
            "证明撤三程序节点及指定期间审查背景。",
        )

    if kind == "其他材料":
        if (form_type in ("海报", "现场照片") and poster_ctx):
            if property_case:
                return (
                    "活动组织广告图片",
                    "证明在涉案不动产/物业经营场景中组织相关活动并进行对外推广展示。",
                )
            return (
                "活动组织广告图片",
                "证明经营主体组织活动并进行对外宣传推广。",
            )
        if any(k in src for k in ["商标图样", "图样", "LOGO", "logo"]):
            return (
                "商标图样页面",
                "证明商标标识的具体呈现方式及展示场景。",
            )
        if any(k in src for k in TEST_REPORT_KEYWORDS):
            return (
                "检测/检验报告材料",
                "证明商品质量规格、标识对应及商品属性信息，作为使用事实的补强材料。",
            )
        if any(k in src for k in EXHIBIT_AD_KEYWORDS) or ("展会/广告宣传" in cls_note):
            return (
                "展会宣传与展位广告材料",
                "证明在展会或广告投放场景中对涉案标识及商品的公开展示使用。",
            )
        if any(k in src for k in BROCHURE_KEYWORDS) or any(k.lower() in src_l for k in BROCHURE_KEYWORDS) or ("宣传册/画册" in cls_note):
            return (
                "宣传册与产品画册材料",
                "证明品牌宣传、产品展示及经营推广场景中的标识使用情况。",
            )
        if any(k in src for k in ["工厂", "厂区", "车间", "仓库", "库存", "生产线", "设备", "包装", "装箱", "实拍", "现场"]):
            return (
                "经营现场实拍",
                "证明服务活动或经营现场的实际使用场景。",
            )
        if any(k in src for k in ["微信", "聊天", "沟通", "客户", "咨询", "下单", "回复"]):
            return (
                "业务沟通记录页面",
                "证明经营沟通场景及商品交易线索。",
            )
        if any(k in src for k in ["宣传", "海报", "详情", "商品介绍", "产品参数", "包装图"]):
            return (
                "商品介绍与展示页面",
                "证明商品展示内容与标识使用场景。",
            )
        if any(k in src for k in ["回单", "流水", "银行", "收据", "款项", "对账"]):
            return (
                "资金往来记录页面",
                "证明经营资金往来及交易时间线索。",
            )

    if form_type == "通知" and kind == "程序文件":
        return (
            "程序通知材料",
            "证明撤三程序节点、送达事实及指定期间审查背景。",
        )
    if form_type == "合同":
        return (
            "合同页面",
            "证明交易关系、交易标的及履约时间线索。",
        )
    if form_type == "发票":
        return (
            "发票页面",
            "证明开票事实、交易标的及交易时间线索。",
        )
    if form_type == "海报":
        if property_case:
            return (
                "活动组织广告图片",
                "证明在涉案不动产/物业场景中的活动组织与推广使用。",
            )
        return (
            "海报宣传材料",
            "证明商标在宣传推广场景中的公开展示使用。",
        )
    if form_type == "包装":
        return (
            "产品包装与标签材料",
            "证明商标在包装载体上的实际标识呈现及商品对应关系。",
        )
    if form_type == "现场照片":
        return (
            "现场照片材料",
            "证明服务活动现场或实际经营场景的存在。",
        )

    if any(k in src for k in REVIEW_KEYWORDS):
        return (
            "客户评价页面",
            "证明商品展示后的客户评价与持续使用时间信息。",
        )

    return (
        f"{kind}页面",
        "证明具体经营场景、商品展示或交易线索中的至少一项事实。",
    )


def ensure_unique_filename(target_dir: Path, base_name: str, used: set) -> str:
    stem = Path(base_name).stem
    suffix = Path(base_name).suffix or ".pdf"
    cand = base_name
    i = 1
    while cand in used or (target_dir / cand).exists():
        cand = f"{stem}_{i}{suffix}"
        i += 1
    used.add(cand)
    return cand


def set_repeat_header_row(table: Any) -> None:
    if not getattr(table, "rows", None):
        return
    if len(table.rows) < 1:
        return
    header_row = table.rows[0]
    # 兼容调用约定
    try:
        header_row.repeat = True
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
    tr = header_row._tr
    tr_pr = tr.get_or_add_trPr()
    tag = qn("w:tblHeader")
    node = tr_pr.find(tag)
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "1")


def create_pre_output_checklist_docx(
    out_dir: Path,
    case_meta: Dict[str, Any],
    rows: List[Dict[str, str]],
) -> Path:
    path = out_dir / "pre_output_checklist.docx"
    doc = DocxDocument()
    doc.add_heading("生成前信息核验单", level=1)
    doc.add_paragraph("editable = true")
    doc.add_paragraph("请人工核验后勾选确认项；未勾选前禁止正式文书输出。")

    doc.add_heading("（1）答辩书信息", level=2)
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "字段"
    t1.rows[0].cells[1].text = "内容"
    set_repeat_header_row(t1)
    base_rows = [
        ("案号", str(case_meta.get("case_no", "") or "")),
        ("商标号", str(case_meta.get("reg_no", "") or "")),
        ("类别", str(case_meta.get("class", "") or "")),
        ("主体", str(case_meta.get("respondent", "") or "")),
        (
            "指定期间",
            f"{str(case_meta.get('use_period_start', '') or '')} ~ {str(case_meta.get('use_period_end', '') or '')}",
        ),
    ]
    for k, v in base_rows:
        cells = t1.add_row().cells
        cells[0].text = k
        cells[1].text = v

    doc.add_heading("（2）证据目录表", level=2)
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Table Grid"
    headers = ["编号", "分类", "文件名", "时间", "来源"]
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
    set_repeat_header_row(t2)
    for row in rows:
        ds = str(row.get("Evidence Date Start", "") or "").strip()
        de = str(row.get("Evidence Date End", "") or "").strip()
        if ds and de and ds != de:
            t_anchor = f"{ds} ~ {de}"
        else:
            t_anchor = ds or de or "—"
        vals = [
            str(row.get("Evidence ID", "") or ""),
            str(row.get("Sort Category", "") or row.get("Type", "") or ""),
            str(row.get("Evidence Name", "") or ""),
            t_anchor,
            str(row.get("Source", "") or ""),
        ]
        cells = t2.add_row().cells
        for i, v in enumerate(vals):
            cells[i].text = v

    doc.add_heading("（3）状态", level=2)
    doc.add_paragraph("editable = true")
    doc.add_heading("（4）确认机制", level=2)
    doc.add_paragraph("□ 信息无误")
    doc.add_paragraph("□ 同意输出")
    doc.add_paragraph("提示：请将“□”改为“☑”后再执行输出。")

    doc.save(str(path))
    return path


def organize_evidence_package(
    rows: List[Dict[str, str]],
    out_dir: Path,
    organize_dir: str,
    duplicate_map: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, str]:
    target_dir = Path(organize_dir) if organize_dir else out_dir / "证据整理包"
    if target_dir.exists():
        for p in target_dir.iterdir():
            try:
                if p.is_file() or p.is_symlink():
                    p.unlink()
                elif p.is_dir():
                    shutil.rmtree(p)
            except Exception as exc:
                _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
                raise RuntimeError(f"清理证据目录失败: {p}") from exc
    target_dir.mkdir(parents=True, exist_ok=True)

    manifest_rows: List[Dict[str, str]] = []
    used = set()
    duplicate_map = duplicate_map or {}

    # 1) 复制并重命名
    for r in rows:
        src = Path(r.pop("_src"))
        src_name = r.pop("_src_name", src.name)
        desired_name = r.pop("_renamed_name")
        r.pop("_dup_files", None)
        unit_type = sanitize_component(str(r.get("Unit Type", "") or detect_unit_type(r)), 16) or "其他佐证"
        unit_dir = target_dir / unit_type
        unit_dir.mkdir(parents=True, exist_ok=True)
        target_name = ensure_unique_filename(unit_dir, desired_name, used)
        dst = unit_dir / target_name
        shutil.copy2(src, dst)

        r["Evidence Name"] = target_name
        r["Carrier"] = str(dst)
        r["Package Unit"] = unit_type
        r["Original File Name"] = src_name

    # 2) 合订重排并生成统一页码（页码以合订本为准）
    merged_pdf = out_dir / "证据内容_重排合并.pdf"
    writer = PdfWriter()
    merged_pages = 0
    page_map: Dict[str, Dict[str, Any]] = {}
    for r in rows:
        pdf = Path(r.get("Carrier", ""))
        try:
            reader = PdfReader(str(pdf), strict=False)
            page_count = len(reader.pages)
        except Exception:
            page_count = 0
            reader = None
        start = merged_pages + 1 if page_count > 0 else 0
        end = merged_pages + page_count if page_count > 0 else 0
        if reader:
            for p in reader.pages:
                writer.add_page(p)
        merged_pages += page_count
        if start and end:
            r["Page Range"] = str(start) if start == end else f"{start}-{end}"
            try:
                writer.add_outline_item(f"{r.get('Evidence ID','')} {r.get('Evidence Name','')}", start - 1)
            except Exception as exc:
                _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        else:
            r["Page Range"] = ""
        eid = r.get("Evidence ID", "")
        if eid:
            page_map[eid] = {
                "evidence_id": eid,
                "evidence_name": r.get("Evidence Name", ""),
                "start_page": int(start) if start else 0,
                "end_page": int(end) if end else 0,
                "page_range": r.get("Page Range", ""),
                "pdf_name": pdf.name,
                "pdf_hash": file_sha1(pdf) if pdf.exists() else "",
            }

    with merged_pdf.open("wb") as f:
        writer.write(f)
    page_map_payload = {
        "merged_pdf_name": merged_pdf.name,
        "merged_pdf_path": str(merged_pdf),
        "merged_pdf_hash": file_sha1(merged_pdf),
        "merged_total_pages": merged_pages,
        "evidences": page_map,
    }
    page_map_path = out_dir / "page_map.json"
    page_map_path.write_text(json.dumps(page_map_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    # 3) 输出清单
    for r in rows:
        anchor_txt = (r.get("Evidence Date Start", "") + (f" ~ {r.get('Evidence Date End', '')}" if r.get("Evidence Date End", "") else "")).strip(" ~")
        if not anchor_txt:
            anchor_txt = "时间不明"
        duplist = duplicate_map.get(r.get("Evidence ID", ""), [])
        dup_note = "；".join(duplist) if duplist else ""

        manifest_rows.append({
            "证据ID": r.get("Evidence ID", ""),
            "原文件名": r.get("Original File Name", ""),
            "证据单元": r.get("Package Unit", r.get("Unit Type", "")),
            "重命名文件": r.get("Evidence Name", ""),
            "推测证明名称": r.get("Inferred Proof Name", ""),
            "对应证明目的": r.get("Inferred Purpose", ""),
            "类型": r.get("Type", ""),
            "合订页码": r.get("Page Range", ""),
            "时间锚点": anchor_txt,
            "要件": r.get("Proof Target (T1-T6)", ""),
            "映射依据": r.get("Mapping Basis", ""),
            "待证事实": r.get("Public Pending Fact", "") or sanitize_public_fact_text(r.get("Risk Notes", ""), r.get("Inferred Purpose", "")),
            "内容总结": r.get("Content Summary", ""),
            "去重合并说明": dup_note,
        })

    manifest_df = pd.DataFrame(manifest_rows)
    manifest_xlsx = out_dir / "证据整理清单.xlsx"
    manifest_df.to_excel(manifest_xlsx, index=False)

    merged_index_xlsx = out_dir / "证据整理合订页码索引.xlsx"
    manifest_df[["证据ID", "推测证明名称", "对应证明目的", "重命名文件", "合订页码", "类型", "要件", "待证事实"]].to_excel(merged_index_xlsx, index=False)

    merged_index_docx = out_dir / "证据整理合订页码表.docx"
    doc = DocxDocument()
    doc.add_heading("证据整理合订页码表", level=1)
    doc.add_paragraph(f"合订文件：{merged_pdf.name}")
    doc.add_paragraph(f"总页数：{merged_pages}")
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    headers = ["序号", "证据ID", "证据名称", "合订页码", "要件", "待证事实"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    set_repeat_header_row(t)
    for i, m in enumerate(manifest_rows, start=1):
        row = t.add_row().cells
        row[0].text = str(i)
        row[1].text = str(m.get("证据ID", ""))
        row[2].text = str(m.get("重命名文件", ""))
        row[3].text = str(m.get("合订页码", ""))
        row[4].text = str(m.get("要件", ""))
        row[5].text = str(m.get("待证事实", ""))
    doc.save(str(merged_index_docx))

    manifest_md = out_dir / "证据整理清单.md"
    with manifest_md.open("w", encoding="utf-8") as f:
        f.write("# 证据整理清单\n\n")
        f.write(f"- 整理目录：{target_dir}\n")
        f.write(f"- 证据数量：{len(manifest_rows)}\n")
        f.write(f"- 合订文件：{merged_pdf.name}（共{merged_pages}页）\n\n")
        for i, m in enumerate(manifest_rows, start=1):
            f.write(f"{i}. {m['证据ID']}｜{m['重命名文件']}\n")
            if m.get("证据单元", ""):
                f.write(f"   - 证据单元：{m.get('证据单元','')}\n")
            f.write(f"   - 合订页码：{m['合订页码']}\n")
            f.write(f"   - 类型：{m['类型']}\n")
            f.write(f"   - 要件：{m['要件']}\n")
            if m["去重合并说明"]:
                f.write(f"   - 去重合并：{m['去重合并说明']}\n")
            f.write(f"   - 待证事实：{m['待证事实']}\n")

    return {
        "organized_dir": str(target_dir),
        "manifest_xlsx": str(manifest_xlsx),
        "manifest_md": str(manifest_md),
        "merged_pdf": str(merged_pdf),
        "merged_pdf_name": merged_pdf.name,
        "merged_total_pages": str(merged_pages),
        "page_map_path": str(page_map_path),
        "page_map_hash": page_map_payload.get("merged_pdf_hash", ""),
        "merged_index_xlsx": str(merged_index_xlsx),
        "merged_index_docx": str(merged_index_docx),
    }


def _hl_clip(s: str, n: int = 48) -> str:
    t = re.sub(r"\s+", " ", str(s or "")).strip("；;，, ")
    return t if len(t) <= n else t[: n - 1] + "…"


def _hl_norm(s: str) -> str:
    return re.sub(r"[\s·_\-:：/\\]+", "", str(s or ""))


def _hl_extract_core(note: str, fallback: str = "") -> str:
    txt = str(note or "")
    m = re.search(r"关键依据：([^；\n]+)", txt)
    if m:
        return m.group(1).strip()
    return str(fallback or "").strip()


def _hl_name_tokens(name: str) -> List[str]:
    x = re.sub(r"^E\d{3}_", "", str(name or ""))
    x = re.sub(r"\.pdf$", "", x, flags=re.I)
    parts = [p for p in re.split(r"[_·\-\s]+", x) if p]
    bad = {"店铺", "商品", "详情", "评价", "页面", "截图", "时间不明", "其他", "程序", "主体", "交易", "证据"}
    out: List[str] = []
    for p in parts:
        if len(p) < 2 or p in bad:
            continue
        out.append(p)
    return out[:8]


def _hl_build_candidates(row: Dict[str, Any]) -> List[str]:
    cands: List[str] = []
    core = _hl_extract_core(row.get("Risk Notes", ""), row.get("Content Summary", ""))
    if core:
        core = re.sub(r"[；。]+$", "", core)
        cands.append(core)
        for ch in re.split(r"[，,。；; ]+", core):
            ch = ch.strip()
            if len(_hl_norm(ch)) >= 6:
                cands.append(ch)

    for k in _hl_name_tokens(row.get("Evidence Name", "")):
        if len(_hl_norm(k)) >= 3:
            cands.append(k)

    purpose = str(row.get("Inferred Purpose", "") or "")
    name_txt = str(row.get("Evidence Name", "") or "")
    if "程序" in purpose:
        cands += ["申请撤销", "通知书"]
    if "主体" in purpose or "权属" in purpose:
        cands += ["商标注册证", "注册号"]
    if "评价" in name_txt:
        cands += ["商品评价", "好评"]
    if "订单" in name_txt:
        cands += ["订单", "支付"]
    if "转账" in name_txt or "收款" in name_txt:
        cands += ["转账", "交易", "收款"]

    seen = set()
    out: List[str] = []
    for c in sorted(cands, key=lambda x: len(_hl_norm(x)), reverse=True):
        n = _hl_norm(c)
        if len(n) < 2 or n in seen:
            continue
        seen.add(n)
        out.append(str(c).strip())
    return out[:16]


def generate_precise_highlight_package(
    out_dir: Path,
    rows: List[Dict[str, str]],
    skip_ids: List[str],
    max_pages_per_evidence: int = 6,
    max_hits_per_evidence: int = 24,
) -> Dict[str, Any]:
    """
    基于文本坐标检索生成“精准高亮 + 红字备注”增强版合订 PDF。
    """
    if fitz is None:
        return {
            "status": "SKIPPED",
            "reason": "PyMuPDF(fitz) 不可用，已跳过精准高亮步骤",
            "highlight_pdf": "",
            "highlight_report": "",
            "hits_total": 0,
        }

    merged_pdf = out_dir / "证据内容_重排合并.pdf"
    page_map_path = out_dir / "page_map.json"
    if not merged_pdf.exists() or not page_map_path.exists():
        return {
            "status": "SKIPPED",
            "reason": "缺少合订PDF或page_map，已跳过精准高亮步骤",
            "highlight_pdf": "",
            "highlight_report": "",
            "hits_total": 0,
        }

    out_pdf = out_dir / "证据内容_重排合并_审查增强高亮版_红字精确v2.pdf"
    out_report = out_dir / "证据高亮命中报告_精准版_v2.json"

    page_map = json.loads(page_map_path.read_text(encoding="utf-8"))
    row_map = {str(r.get("Evidence ID", "")).strip(): r for r in rows if str(r.get("Evidence ID", "")).strip()}
    skip_set = {x.strip() for x in skip_ids if x and x.strip()}

    def _eid_num(eid: str) -> int:
        m = re.search(r"(\d+)", str(eid))
        return int(m.group(1)) if m else 99999

    doc = fitz.open(str(merged_pdf))
    report_rows: List[Dict[str, Any]] = []
    hits_total = 0

    for eid, meta in sorted(page_map.get("evidences", {}).items(), key=lambda kv: _eid_num(kv[0])):
        start = int(meta.get("start_page", 0) or 0)
        end = int(meta.get("end_page", 0) or 0)
        if start <= 0 or end <= 0 or start > end:
            continue
        row = row_map.get(eid, {})

        purpose = _hl_clip(row.get("Inferred Purpose", ""), 42)
        core = _hl_clip(_hl_extract_core(row.get("Risk Notes", ""), row.get("Content Summary", "")), 48)
        ev_name = re.sub(r"^E\d{3}_", "", str(meta.get("evidence_name", "") or "")).replace("_", "·").replace(".pdf", "")
        ev_name = ev_name.replace("时间不明·", "").strip("·")

        p0 = doc[start - 1]
        w = p0.rect.width
        banner = fitz.Rect(16, 10, w - 16, 32)
        p0.draw_rect(banner, color=(0.75, 0.0, 0.0), fill=(1.0, 0.93, 0.93), width=1.0)
        p0.insert_text((20, 26), f"{eid} 重点提示：{_hl_clip(ev_name, 40)}", fontsize=10, color=(0.78, 0.0, 0.0))

        note_text = f"{eid}\n证明目的：{purpose or '见证据目录'}\n核心内容：{core or '见证据目录'}\n页码：{start}-{end}"
        note_rect = fitz.Rect(20, max(36, p0.rect.height - 160), w - 20, p0.rect.height - 34)
        try:
            p0.add_freetext_annot(
                note_rect,
                note_text,
                fontsize=9,
                text_color=(0.75, 0.0, 0.0),
                fill_color=(1.0, 0.95, 0.95),
                border_color=(0.75, 0.0, 0.0),
                align=0,
            )
        except Exception:
            p0.insert_text((24, p0.rect.height - 46), f"证明目的：{purpose}", fontsize=8.5, color=(0.75, 0.0, 0.0))

        if eid in skip_set:
            report_rows.append({
                "证据ID": eid,
                "页码范围": f"{start}-{end}",
                "关键词": [],
                "命中页": [],
                "命中数": 0,
                "跳过原因": "按配置跳过高亮（实拍环境图片等）",
            })
            continue

        cands = _hl_build_candidates(row)
        hit_count = 0
        hit_pages: List[int] = []
        max_page = min(end, start + max(1, max_pages_per_evidence) - 1)

        for pg in range(start, max_page + 1):
            page = doc[pg - 1]
            page_hits = 0
            for cand in cands:
                variants = [cand]
                nc = _hl_norm(cand)
                if nc != cand:
                    variants.append(nc)

                found = []
                for v in variants:
                    if len(v.strip()) < 2:
                        continue
                    found = page.search_for(v)
                    if found:
                        break
                if not found:
                    continue

                for rect in found[:3]:
                    h = page.add_highlight_annot(rect)
                    h.set_colors(stroke=(1.0, 0.84, 0.0))
                    h.update(opacity=0.45)
                    page.draw_rect(rect, color=(0.86, 0.0, 0.0), width=0.8)
                    hit_count += 1
                    page_hits += 1
                    if hit_count >= max_hits_per_evidence:
                        break
                if hit_count >= max_hits_per_evidence or page_hits >= 10:
                    break

            if page_hits:
                hit_pages.append(pg)
            if hit_count >= max_hits_per_evidence:
                break

        hits_total += hit_count
        report_rows.append({
            "证据ID": eid,
            "页码范围": f"{start}-{end}",
            "关键词": cands,
            "命中页": hit_pages,
            "命中数": hit_count,
            "证明目的": str(row.get("Inferred Purpose", "") or ""),
            "核心内容": core,
        })

    doc.save(str(out_pdf), garbage=4, deflate=True)
    doc.close()
    out_report.write_text(json.dumps(report_rows, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "status": "PASS",
        "reason": "",
        "highlight_pdf": str(out_pdf),
        "highlight_report": str(out_report),
        "hits_total": hits_total,
        "evidence_count": len(report_rows),
        "skipped_ids": sorted(list(skip_set)),
    }


def write_scan_report(
    out_dir: Path,
    report_name: str,
    case: Dict[str, str],
    rows: List[Dict[str, str]],
    package_info: Dict[str, str],
    baseline_info: Optional[Dict[str, str]] = None,
) -> str:
    p = out_dir / report_name
    with p.open("w", encoding="utf-8") as f:
        f.write("# 扫描识别报告\n\n")
        f.write("## 第一轮案情识别\n")
        f.write(f"- 被撤销商标：{case.get('revoked_mark_name', case.get('mark_name',''))}\n")
        f.write(f"- 被撤销商品/服务：{case.get('revoked_goods_services','')}\n")
        f.write("\n")
        f.write("## 案件识别\n")
        f.write(f"- 商标名称：{case.get('mark_name','')}\n")
        f.write(f"- 商标注册号：{case.get('reg_no','')}\n")
        f.write(f"- 类别：{case.get('class','')}\n")
        f.write(f"- 被申请人：{case.get('respondent','')}\n")
        f.write(f"- 申请人：{case.get('applicant','')}\n")
        f.write(f"- 指定期间：{case.get('use_period_start','')} ~ {case.get('use_period_end','')}\n\n")
        rounds_raw = case.get("scan_rounds", "")
        if rounds_raw:
            f.write("## 固定五轮扫描执行明细\n")
            try:
                rounds = json.loads(rounds_raw) if isinstance(rounds_raw, str) else rounds_raw
            except Exception:
                rounds = []
            for r in (rounds or []):
                f.write(
                    f"- 第{r.get('round','?')}轮 {r.get('title','')}："
                    f"{r.get('goal','')}；扫描{r.get('selected_count',0)}份；"
                    f"页数上限={r.get('max_pages','')}；DPI={r.get('dpi','')}"
                )
                if r.get("detail"):
                    f.write(f"；说明：{r.get('detail','')}")
                f.write("\n")
            f.write("\n")
        if case.get("run_id"):
            f.write(f"- 运行ID：{case.get('run_id','')}\n")
            f.write(f"- CaseInfo版本：{case.get('case_info_version','')}\n")
            f.write(f"- 校验规则版本：{case.get('validation_rule_version','')}\n")
            f.write(f"- 校验结果：{case.get('caseinfo_validation_status','')}\n")
            if case.get("caseinfo_validation_warnings"):
                f.write(f"- 校验告警：{case.get('caseinfo_validation_warnings','')}\n")
            if case.get("caseinfo_validation_errors"):
                f.write(f"- 校验错误：{case.get('caseinfo_validation_errors','')}\n")
            if case.get("caseinfo_validation_path"):
                f.write(f"- 校验详情：{case.get('caseinfo_validation_path','')}\n")
            if case.get("source_of_truth"):
                f.write(f"- 字段来源：{case.get('source_of_truth','')}\n")
            f.write("\n")
        if case.get("dedup_before_count"):
            f.write(f"- 去重前证据数：{case.get('dedup_before_count','')}\n")
            f.write(f"- 去重后证据数：{case.get('dedup_after_count','')}\n\n")
        if case.get("input_pdf_count") is not None:
            f.write("- 输入文件统计：\n")
            f.write(f"  - PDF：{case.get('input_pdf_count','0')}\n")
            f.write(f"  - 图片转PDF：{case.get('input_image_converted_count','0')}\n")
            f.write(f"  - Office转PDF：{case.get('input_office_converted_count','0')}\n")
            f.write(f"  - Office直读文本：{case.get('input_office_text_extracted_count','0')}\n")
            f.write(f"  - Office未转换：{case.get('input_office_skipped_count','0')}\n")
            f.write(f"  - fast目录证据：{case.get('input_fast_profile_count','0')}\n")
            f.write(f"  - full目录证据：{case.get('input_full_profile_count','0')}\n")
            if case.get("office_converter"):
                f.write(f"  - Office转换器：{case.get('office_converter','')}\n")
            f.write("\n")
        if case.get("override_rules_path"):
            f.write(f"- 人工覆盖规则：{case.get('override_rules_path','')}\n")
            f.write(f"- 本次生效条数：{case.get('override_applied_count','0')}\n\n")
        if case.get("override_rules_json"):
            f.write(f"- 覆盖规则快照：{case.get('override_rules_json','')}\n")
        if case.get("overrides_applied_path"):
            f.write(f"- 覆盖生效日志：{case.get('overrides_applied_path','')}\n")
        if case.get("page_map_path"):
            f.write(f"- 页码真源(page_map)：{case.get('page_map_path','')}\n")
            f.write(f"- 页码真源哈希：{case.get('page_map_hash','')}\n\n")
        if case.get("page_map_validation_status"):
            f.write(f"- 页码一致性校验：{case.get('page_map_validation_status','')}\n")
            if case.get("page_map_validation_errors"):
                f.write(f"- 页码校验错误：{case.get('page_map_validation_errors','')}\n")
            if case.get("page_map_validation_path"):
                f.write(f"- 页码校验详情：{case.get('page_map_validation_path','')}\n")
            f.write("\n")
        if case.get("time_quality_status"):
            f.write(f"- 时间质量校验：{case.get('time_quality_status','')}\n")
            if case.get("time_quality_errors"):
                f.write(f"- 时间质量错误：{case.get('time_quality_errors','')}\n")
            if case.get("time_quality_warnings"):
                f.write(f"- 时间质量告警：{case.get('time_quality_warnings','')}\n")
            if case.get("time_quality_path"):
                f.write(f"- 时间质量详情：{case.get('time_quality_path','')}\n")
            f.write("\n")
        if case.get("name_quality_status"):
            f.write(f"- 证据命名校验：{case.get('name_quality_status','')}\n")
            if case.get("name_quality_errors"):
                f.write(f"- 证据命名错误：{case.get('name_quality_errors','')}\n")
            if case.get("name_quality_warnings"):
                f.write(f"- 证据命名告警：{case.get('name_quality_warnings','')}\n")
            if case.get("name_quality_path"):
                f.write(f"- 证据命名详情：{case.get('name_quality_path','')}\n")
            if case.get("name_pending_path"):
                f.write(f"- 待人工确认清单：{case.get('name_pending_path','')}\n")
            if case.get("name_pending_count"):
                f.write(f"- 待人工确认数量：{case.get('name_pending_count','0')}\n")
            f.write("\n")

        f.write("## 证据整理输出\n")
        f.write(f"- 整理目录：{package_info.get('organized_dir','')}\n")
        f.write(f"- 清单（xlsx）：{package_info.get('manifest_xlsx','')}\n")
        f.write(f"- 清单（md）：{package_info.get('manifest_md','')}\n\n")
        f.write(f"- 合订文件：{package_info.get('merged_pdf','')}\n")
        f.write(f"- 合订页码索引：{package_info.get('merged_index_xlsx','')}\n\n")
        f.write(f"- 合订页码表（docx）：{package_info.get('merged_index_docx','')}\n\n")
        if case.get("precise_highlight_status"):
            f.write("## 精准高亮输出（文本坐标）\n")
            f.write(f"- 状态：{case.get('precise_highlight_status','')}\n")
            if case.get("precise_highlight_reason"):
                f.write(f"- 说明：{case.get('precise_highlight_reason','')}\n")
            if case.get("precise_highlight_pdf"):
                f.write(f"- 高亮PDF：{case.get('precise_highlight_pdf','')}\n")
            if case.get("precise_highlight_report"):
                f.write(f"- 命中报告：{case.get('precise_highlight_report','')}\n")
            if case.get("precise_highlight_hits_total"):
                f.write(f"- 高亮命中总数：{case.get('precise_highlight_hits_total','')}\n")
            if case.get("precise_highlight_skipped_ids"):
                f.write(f"- 跳过证据：{case.get('precise_highlight_skipped_ids','')}\n")
            f.write("\n")

        if baseline_info:
            f.write("## 回归基线指标（本次）\n")
            f.write(f"- 样本清单：{baseline_info.get('sample_manifest','')}\n")
            f.write(f"- 基线指标：{baseline_info.get('metrics_json','')}\n")
            f.write(f"- 指标历史：{baseline_info.get('history_csv','')}\n")
            f.write(f"- 时间识别率：{baseline_info.get('time_recognition_rate','')}\n")
            f.write(f"- 时间高/中置信率：{baseline_info.get('time_high_medium_rate','')}\n")
            f.write(f"- 时间不明率：{baseline_info.get('time_unknown_rate','')}\n\n")

        f.write("## 扫描重点（商标词条+时间）\n")
        for i, r in enumerate(rows, start=1):
            f.write(f"### {i}. {r.get('Evidence ID','')} {r.get('Evidence Name','')}\n")
            f.write(f"- 推测证明名称：{r.get('Inferred Proof Name','')}\n")
            f.write(f"- 对应证明目的：{r.get('Inferred Purpose','')}\n")
            f.write(f"- 时间锚点：{r.get('Evidence Date Start','')} ~ {r.get('Evidence Date End','')}\n")
            f.write(f"- 置信度：{r.get('Date Confidence','')}\n")
            f.write(f"- 时间通道：{r.get('Time Channels','')}\n")
            f.write(f"- 冲突标记：{r.get('Time Conflict Flag','N')}\n")
            f.write(f"- 择优规则：{r.get('Time Selection Rule','')}\n")
            f.write(f"- 候选TopK：{r.get('Time Anchor Candidates','[]')}\n")
            if r.get("Trade Chain Group"):
                f.write(f"- 交易配对：{r.get('Trade Chain Group','')}（{r.get('Trade Link Role','交易')}）\n")
            f.write(f"- 要件：{r.get('Proof Target (T1-T6)','')}\n")
            f.write(f"- 待证事实：{r.get('Risk Notes','')}\n")
            f.write(f"- 内容总结：{r.get('Content Summary','')}\n")
            f.write(f"- 时间关键词：{r.get('Time Anchor Text','')}\n\n")

    return str(p)


def _yn(v: str) -> bool:
    s = (v or "").strip().upper()
    return s in ("Y", "YES", "1", "是", "TRUE")


def create_manual_override_template(out_dir: Path, rows: List[Dict[str, str]], override_path: Path) -> str:
    """
    生成人工覆盖规则模板（若不存在）：
    - 支持按证据ID/原文件名/重命名文件/合订页码匹配
    - 支持覆盖时间、证明名称、证明目的、类型
    """
    if override_path.exists():
        return str(override_path)

    unknown_rows = [r for r in rows if not parse_iso_date(r.get("Evidence Date Start", ""))]
    tpl_rows: List[Dict[str, str]] = []
    for r in unknown_rows:
        tpl_rows.append({
            "启用(Y/N)": "N",
            "匹配方式": "合订页码",
            "匹配值": r.get("Page Range", ""),
            "证据ID": r.get("Evidence ID", ""),
            "原文件名": r.get("Original File Name", ""),
            "重命名文件": r.get("Evidence Name", ""),
            "当前时间锚点": (r.get("Evidence Date Start", "") + (f" ~ {r.get('Evidence Date End', '')}" if r.get("Evidence Date End", "") else "")).strip(" ~") or "时间不明",
            "覆盖时间起(YYYY-MM-DD)": "",
            "覆盖时间止(YYYY-MM-DD)": "",
            "覆盖置信度(HIGH/MEDIUM/LOW)": "HIGH",
            "覆盖证明名称": "",
            "覆盖证明目的": "",
            "覆盖类型": "",
            "备注": "示例：若人工确认该页存在评价日期，请填写覆盖时间起止并启用。",
        })

    if not tpl_rows:
        tpl_rows.append({
            "启用(Y/N)": "N",
            "匹配方式": "证据ID",
            "匹配值": "证据1",
            "证据ID": "",
            "原文件名": "",
            "重命名文件": "",
            "当前时间锚点": "",
            "覆盖时间起(YYYY-MM-DD)": "",
            "覆盖时间止(YYYY-MM-DD)": "",
            "覆盖置信度(HIGH/MEDIUM/LOW)": "HIGH",
            "覆盖证明名称": "",
            "覆盖证明目的": "",
            "覆盖类型": "",
            "备注": "填写启用=Y后生效。",
        })

    df = pd.DataFrame(tpl_rows)
    override_path.parent.mkdir(parents=True, exist_ok=True)
    df.to_excel(override_path, index=False)
    return str(override_path)


def _match_row_for_override(row: Dict[str, str], mode: str, value: str) -> bool:
    mv = (value or "").strip()
    if not mv:
        return False
    mode = (mode or "").strip()
    if mode == "证据ID":
        return (row.get("Evidence ID", "") or "") == mv
    if mode == "原文件名":
        return (row.get("Original File Name", "") or "") == mv
    if mode == "重命名文件":
        return (row.get("Evidence Name", "") or "") == mv
    if mode == "合订页码":
        return (row.get("Page Range", "") or "") == mv
    # 兜底：任一字段命中
    return mv in {
        row.get("Evidence ID", "") or "",
        row.get("Original File Name", "") or "",
        row.get("Evidence Name", "") or "",
        row.get("Page Range", "") or "",
    }


def load_override_rules(override_path: Path) -> Dict[str, Any]:
    if not override_path.exists():
        return {"rules": [], "version": "override_rules.v1", "source": str(override_path)}
    try:
        df = pd.read_excel(override_path)
    except Exception as exc:
        _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
        return {"rules": [], "version": "override_rules.v1", "source": str(override_path)}

    def _cell(v: Any) -> str:
        if v is None:
            return ""
        try:
            if pd.isna(v):
                return ""
        except Exception as exc:
            _audit_exception("auto_recognize.swallowed_exception", exc, __file__)
            return ""
        s = str(v).strip()
        return "" if s.lower() == "nan" else s

    rules: List[Dict[str, Any]] = []
    for i, rr in enumerate(df.to_dict("records"), start=1):
        rid = f"R{i:03d}"
        rules.append({
            "rule_id": rid,
            "enabled": _yn(_cell(rr.get("启用(Y/N)", ""))),
            "match_mode": _cell(rr.get("匹配方式", "")) or "证据ID",
            "match_value": _cell(rr.get("匹配值", "")),
            "fallback_original": _cell(rr.get("原文件名", "")),
            "fallback_renamed": _cell(rr.get("重命名文件", "")),
            "override_date_start": _cell(rr.get("覆盖时间起(YYYY-MM-DD)", "")),
            "override_date_end": _cell(rr.get("覆盖时间止(YYYY-MM-DD)", "")),
            "override_confidence": _cell(rr.get("覆盖置信度(HIGH/MEDIUM/LOW)", "")).upper(),
            "override_proof_name": _cell(rr.get("覆盖证明名称", "")),
            "override_purpose": _cell(rr.get("覆盖证明目的", "")),
            "override_type": _cell(rr.get("覆盖类型", "")),
            "remark": _cell(rr.get("备注", "")),
        })
    return {"rules": rules, "version": "override_rules.v1", "source": str(override_path)}


def apply_manual_overrides(
    rows: List[Dict[str, str]],
    override_path: Path,
    out_dir: Path,
    run_id: str,
) -> Dict[str, str]:
    if not override_path.exists():
        return {
            "applied_count": "0",
            "override_path": str(override_path),
            "overrides_applied_path": str(out_dir / "overrides_applied.json"),
            "override_rules_json": str(out_dir / "override_rules.json"),
        }

    rules_pack = load_override_rules(override_path)
    rules = rules_pack.get("rules", [])
    rules_json_path = out_dir / "override_rules.json"
    rules_json_path.write_text(json.dumps(rules_pack, ensure_ascii=False, indent=2), encoding="utf-8")

    applied = 0
    applied_log: List[Dict[str, Any]] = []
    for rr in rules:
        enabled = bool(rr.get("enabled", False))
        if not enabled:
            continue
        mode = str(rr.get("match_mode", "")).strip() or "证据ID"
        value = str(rr.get("match_value", "")).strip()
        if not value:
            continue
        fallback_original = str(rr.get("fallback_original", "")).strip()
        fallback_renamed = str(rr.get("fallback_renamed", "")).strip()

        match_indexes = [i for i, row in enumerate(rows) if _match_row_for_override(row, mode, value)]
        if not match_indexes and fallback_original:
            match_indexes = [i for i, row in enumerate(rows) if _match_row_for_override(row, "原文件名", fallback_original)]
        if not match_indexes and fallback_renamed:
            match_indexes = [i for i, row in enumerate(rows) if _match_row_for_override(row, "重命名文件", fallback_renamed)]
        if not match_indexes:
            continue

        ds = parse_iso_date(str(rr.get("override_date_start", "")).strip())
        de = parse_iso_date(str(rr.get("override_date_end", "")).strip())
        conf = str(rr.get("override_confidence", "")).strip().upper()
        if conf not in ("HIGH", "MEDIUM", "LOW"):
            conf = "HIGH"
        proof_name = str(rr.get("override_proof_name", "")).strip()
        purpose = str(rr.get("override_purpose", "")).strip()
        type_override = str(rr.get("override_type", "")).strip()

        for mi in match_indexes:
            row = rows[mi]
            changed = False
            before = {
                "Evidence ID": row.get("Evidence ID", ""),
                "Evidence Date Start": row.get("Evidence Date Start", ""),
                "Evidence Date End": row.get("Evidence Date End", ""),
                "Date Confidence": row.get("Date Confidence", ""),
                "Inferred Proof Name": row.get("Inferred Proof Name", ""),
                "Inferred Purpose": row.get("Inferred Purpose", ""),
                "Type": row.get("Type", ""),
            }
            if ds:
                row["Evidence Date Start"] = ds.isoformat()
                changed = True
            if de:
                row["Evidence Date End"] = de.isoformat()
                changed = True
            elif ds and not de:
                row["Evidence Date End"] = ds.isoformat()
                changed = True

            if changed:
                row["Date Confidence"] = conf
                if row.get("Evidence Date Start", "") and row.get("Evidence Date End", ""):
                    row["Formation Date"] = row["Evidence Date Start"] if row["Evidence Date Start"] == row["Evidence Date End"] else ""
                anchor = row.get("Time Anchor Text", "") or ""
                manual_anchor = (
                    row["Evidence Date Start"]
                    if row.get("Evidence Date Start", "") == row.get("Evidence Date End", "")
                    else f"{row.get('Evidence Date Start','')} ~ {row.get('Evidence Date End','')}"
                )
                row["Time Anchor Text"] = f"{anchor}；人工核验时间：{manual_anchor}".strip("；")
                row["Risk Notes"] = (row.get("Risk Notes", "") + f"；人工核验已覆盖时间：{manual_anchor}").strip("；")
                row["Content Summary"] = (row.get("Content Summary", "") + f"；人工核验时间：{manual_anchor}").strip("；")

            if proof_name:
                row["Inferred Proof Name"] = proof_name
                changed = True
            if purpose:
                row["Inferred Purpose"] = purpose
                row["Risk Notes"] = (row.get("Risk Notes", "") + f"；人工核验证明目的：{purpose}").strip("；")
                changed = True
            if type_override:
                row["Type"] = type_override
                changed = True

            if changed:
                applied += 1
                after = {
                    "Evidence ID": row.get("Evidence ID", ""),
                    "Evidence Date Start": row.get("Evidence Date Start", ""),
                    "Evidence Date End": row.get("Evidence Date End", ""),
                    "Date Confidence": row.get("Date Confidence", ""),
                    "Inferred Proof Name": row.get("Inferred Proof Name", ""),
                    "Inferred Purpose": row.get("Inferred Purpose", ""),
                    "Type": row.get("Type", ""),
                }
                applied_log.append({
                    "run_id": run_id,
                    "rule_id": rr.get("rule_id", ""),
                    "match_mode": mode,
                    "match_value": value,
                    "evidence_id": row.get("Evidence ID", ""),
                    "before": before,
                    "after": after,
                })

    applied_payload = {
        "run_id": run_id,
        "override_path": str(override_path),
        "override_rules_version": rules_pack.get("version", "override_rules.v1"),
        "applied_count": applied,
        "applied": applied_log,
    }
    applied_path = out_dir / "overrides_applied.json"
    applied_path.write_text(json.dumps(applied_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "applied_count": str(applied),
        "override_path": str(override_path),
        "overrides_applied_path": str(applied_path),
        "override_rules_json": str(rules_json_path),
    }


def range_overlaps(
    s1: Optional[dt.date],
    e1: Optional[dt.date],
    s2: Optional[dt.date],
    e2: Optional[dt.date],
) -> bool:
    if not s1 or not e1 or not s2 or not e2:
        return False
    return max(s1, s2) <= min(e1, e2)


def build_regression_baseline(
    out_dir: Path,
    files: List[Path],
    rows: List[Dict[str, str]],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> Dict[str, str]:
    sample_names = sorted([p.name for p in files])
    sample_text = "\n".join(sample_names)
    sample_hash = hashlib.sha1(sample_text.encode("utf-8")).hexdigest()[:16]

    sample_manifest = out_dir / "回归样本清单.txt"
    sample_manifest.write_text(sample_text + "\n", encoding="utf-8")

    total = len(rows)
    with_time = 0
    in_period_cnt = 0
    unknown_cnt = 0
    high_cnt = 0
    med_cnt = 0
    low_cnt = 0
    channel_counts = {"评价": 0, "交易": 0, "物流": 0}
    target_counts = {"T1": 0, "T2": 0, "T3": 0, "T4": 0, "T5": 0, "T6": 0}

    for r in rows:
        ds = parse_iso_date(r.get("Evidence Date Start", ""))
        de = parse_iso_date(r.get("Evidence Date End", "")) or ds
        if ds:
            with_time += 1
        else:
            unknown_cnt += 1

        if range_overlaps(ds, de, period_start, period_end):
            in_period_cnt += 1

        conf = (r.get("Date Confidence", "") or "").strip().upper()
        if conf == "HIGH":
            high_cnt += 1
        elif conf == "MEDIUM":
            med_cnt += 1
        else:
            low_cnt += 1

        ch = r.get("Time Channels", "") or ""
        for k in channel_counts:
            if k in ch:
                channel_counts[k] += 1

        t = r.get("Proof Target (T1-T6)", "") or ""
        for tk in target_counts:
            if tk in t:
                target_counts[tk] += 1

    def pct(n: int, d: int) -> str:
        return f"{(100.0 * n / d):.1f}%" if d else "0.0%"

    metrics = {
        "generated_at": dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sample_hash": sample_hash,
        "sample_count": total,
        "sample_file_count": len(sample_names),
        "time_recognition_rate": pct(with_time, total),
        "time_high_medium_rate": pct(high_cnt + med_cnt, total),
        "time_unknown_rate": pct(unknown_cnt, total),
        "in_period_overlap_rate": pct(in_period_cnt, total),
        "high_confidence_count": high_cnt,
        "medium_confidence_count": med_cnt,
        "low_confidence_count": low_cnt,
        "review_channel_hit": channel_counts["评价"],
        "order_channel_hit": channel_counts["交易"],
        "logistics_channel_hit": channel_counts["物流"],
        "t1_cover": target_counts["T1"],
        "t2_cover": target_counts["T2"],
        "t3_cover": target_counts["T3"],
        "t4_cover": target_counts["T4"],
        "t5_cover": target_counts["T5"],
        "t6_cover": target_counts["T6"],
    }

    metrics_json = out_dir / "回归基线指标.json"
    metrics_json.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")

    history_csv = out_dir / "回归指标历史.csv"
    headers = list(metrics.keys())
    write_header = not history_csv.exists()
    with history_csv.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        if write_header:
            writer.writeheader()
        writer.writerow(metrics)

    return {
        "sample_manifest": str(sample_manifest),
        "metrics_json": str(metrics_json),
        "history_csv": str(history_csv),
        "time_recognition_rate": metrics["time_recognition_rate"],
        "time_high_medium_rate": metrics["time_high_medium_rate"],
        "time_unknown_rate": metrics["time_unknown_rate"],
        "sample_hash": sample_hash,
    }


def validate_page_map_consistency(
    rows: List[Dict[str, Any]],
    package_info: Dict[str, str],
) -> Dict[str, Any]:
    errors: List[str] = []
    warnings: List[str] = []
    checks: List[Dict[str, str]] = []
    p = Path(package_info.get("page_map_path", ""))
    if not p.exists():
        errors.append(f"page_map 缺失：{p}")
        return {
            "status": "FAIL",
            "errors": errors,
            "warnings": warnings,
            "checks": checks,
        }
    try:
        payload = json.loads(p.read_text(encoding="utf-8"))
    except Exception as e:
        errors.append(f"page_map 解析失败：{e}")
        return {
            "status": "FAIL",
            "errors": errors,
            "warnings": warnings,
            "checks": checks,
        }

    evidences = payload.get("evidences", {})
    if not isinstance(evidences, dict) or not evidences:
        errors.append("page_map.evidences 为空")

    merged_name = payload.get("merged_pdf_name", "")
    expected_name = package_info.get("merged_pdf_name", "")
    if merged_name != expected_name:
        errors.append(f"merged_pdf_name 不一致：page_map={merged_name}，package={expected_name}")

    merged_pages_map = int(payload.get("merged_total_pages", 0) or 0)
    merged_pages_pkg = int(package_info.get("merged_total_pages", 0) or 0)
    if merged_pages_map != merged_pages_pkg:
        errors.append(f"merged_total_pages 不一致：page_map={merged_pages_map}，package={merged_pages_pkg}")

    for r in rows:
        eid = r.get("Evidence ID", "")
        pr = r.get("Page Range", "")
        if not eid:
            continue
        mapped = evidences.get(eid, {})
        if not mapped:
            errors.append(f"page_map 缺少证据ID：{eid}")
            continue
        mapped_pr = mapped.get("page_range", "")
        if pr != mapped_pr:
            errors.append(f"页码不一致：{eid} rows={pr} page_map={mapped_pr}")
        s = int(mapped.get("start_page", 0) or 0)
        e = int(mapped.get("end_page", 0) or 0)
        if s and e and s > e:
            errors.append(f"页码范围非法：{eid} start={s} end={e}")
        checks.append({
            "evidence_id": eid,
            "rows_page_range": pr,
            "page_map_page_range": mapped_pr,
            "result": "PASS" if pr == mapped_pr else "FAIL",
        })

    if merged_pages_map > 0:
        bad = [
            eid for eid, m in evidences.items()
            if int(m.get("end_page", 0) or 0) > merged_pages_map
        ]
        if bad:
            errors.append(f"存在超出合并总页数的证据页码：{','.join(bad)}")

    return {
        "status": "FAIL" if errors else "PASS",
        "errors": errors,
        "warnings": warnings,
        "checks": checks,
    }


def validate_time_quality(
    rows: List[Dict[str, Any]],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
    min_in_period_highmed: int = 8,
    max_low_ratio: float = 0.35,
    max_unknown_ratio: float = 0.30,
) -> Dict[str, Any]:
    total = len(rows)
    scoped_rows = [r for r in rows if str(r.get("Time Gate Scope", "Y")).upper() == "Y"]
    scoped_total = len(scoped_rows)
    low_cnt = 0
    unknown_cnt = 0
    in_period_highmed = 0
    in_period_reference = 0
    detail_low: List[str] = []
    detail_unknown: List[str] = []
    checks: List[Dict[str, Any]] = []
    errors: List[str] = []
    warnings: List[str] = []

    for r in rows:
        eid = r.get("Evidence ID", "")
        conf = (r.get("Date Confidence", "") or "").upper()
        ds = parse_iso_date(r.get("Evidence Date Start", ""))
        de = parse_iso_date(r.get("Evidence Date End", "")) or ds
        pr = r.get("Page Range", "")
        scoped = str(r.get("Time Gate Scope", "Y")).upper() == "Y"
        is_unknown = not (ds and de)
        in_period = False
        if ds and de and period_start and period_end:
            in_period = _date_overlap(ds, de, period_start, period_end)

        if scoped and conf == "LOW":
            low_cnt += 1
            detail_low.append(f"{eid}({pr or '未标注'})")
        if scoped and is_unknown:
            unknown_cnt += 1
            detail_unknown.append(f"{eid}({pr or '未标注'})")
        if scoped and conf in ("HIGH", "MEDIUM") and in_period:
            in_period_highmed += 1
        if (not scoped) and in_period:
            in_period_reference += 1

        checks.append(
            {
                "evidence_id": eid,
                "confidence": conf,
                "time_gate_scope": "Y" if scoped else "N",
                "in_period": in_period,
                "unknown_time": is_unknown,
                "page_range": pr,
            }
        )

    low_ratio = (low_cnt / scoped_total) if scoped_total else 0.0
    unknown_ratio = (unknown_cnt / scoped_total) if scoped_total else 0.0

    if scoped_total == 0:
        warnings.append("时间门禁范围证据为0，已跳过时间质量强制校验。")
        if in_period_reference > 0:
            warnings.append("存在仅B线参考材料覆盖期内（如程序文件/管理材料），不构成实质使用覆盖。")
    if scoped_total > 0 and in_period_highmed < min_in_period_highmed:
        errors.append(f"期内高/中置信证据不足：{in_period_highmed} < {min_in_period_highmed}")
    if scoped_total > 0 and in_period_highmed == 0 and in_period_reference > 0:
        warnings.append("仅B线参考材料覆盖期内，建议补充交易/履约/可核验网页等A线证据。")
    if scoped_total > 0 and low_ratio > max_low_ratio:
        errors.append(f"低置信占比超阈值：{low_ratio:.1%} > {max_low_ratio:.1%}")
    if scoped_total > 0 and unknown_ratio > max_unknown_ratio:
        errors.append(f"时间不明占比超阈值：{unknown_ratio:.1%} > {max_unknown_ratio:.1%}")
    if low_cnt > 0:
        warnings.append(f"低置信证据共{low_cnt}份：{', '.join(detail_low[:12])}")
    if unknown_cnt > 0:
        warnings.append(f"时间不明证据共{unknown_cnt}份：{', '.join(detail_unknown[:12])}")

    return {
        "status": "FAIL" if errors else "PASS",
        "errors": errors,
        "warnings": warnings,
        "metrics": {
            "total": total,
            "scoped_total": scoped_total,
            "in_period_highmed": in_period_highmed,
            "in_period_reference": in_period_reference,
            "low_count": low_cnt,
            "unknown_count": unknown_cnt,
            "low_ratio": round(low_ratio, 4),
            "unknown_ratio": round(unknown_ratio, 4),
            "thresholds": {
                "min_in_period_highmed": min_in_period_highmed,
                "max_low_ratio": max_low_ratio,
                "max_unknown_ratio": max_unknown_ratio,
            },
        },
        "checks": checks,
    }


def validate_evidence_type_consistency(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    errors: List[str] = []
    warnings: List[str] = []
    checks: List[Dict[str, Any]] = []
    official_terms = ["国家知识产权局", "商标局", "撤销连续三年不使用", "答辩通知", "受理通知", "补正通知", "决定书"]
    tx_terms = ["甲方", "乙方", "合同编号", "发票代码", "发票号码", "价税合计", "订单编号", "支付成功", "物流单号"]
    promo_terms = PROMOTION_MATERIAL_KEYWORDS + ["宣传册", "讲座", "路演", "海报", "画册"]

    for r in rows:
        eid = str(r.get("Evidence ID", "") or "")
        et = str(r.get("Type", "") or "")
        form = str(r.get("Form Type", "") or "")
        blob = " ".join(
            [
                str(r.get("Evidence Name", "") or ""),
                str(r.get("Inferred Proof Name", "") or ""),
                str(r.get("Inferred Purpose", "") or ""),
                str(r.get("Risk Notes", "") or ""),
                str(r.get("Mapping Basis", "") or ""),
                str(r.get("Content Summary", "") or ""),
            ]
        )
        tx_hit = sum(1 for k in tx_terms if k in blob)
        promo_hit = any(k in blob for k in promo_terms)
        official_hit = any(k in blob for k in official_terms)
        item_errors: List[str] = []
        item_warnings: List[str] = []

        if et == "程序文件":
            if form in ("海报", "现场照片", "包装"):
                item_errors.append(f"{eid} 类型=程序文件 但表现形式={form}")
            if promo_hit and not official_hit:
                item_errors.append(f"{eid} 程序文件缺少官方程序特征，含宣传特征")

        if et == "交易凭证":
            if form == "通知" and tx_hit < 2:
                item_errors.append(f"{eid} 类型=交易凭证 但表现形式=通知且交易结构字段不足")
            if promo_hit and tx_hit < 2:
                item_errors.append(f"{eid} 类型=交易凭证 含宣传特征且无稳定交易结构字段")

        if et in ("线上店铺展示", "商品展示页") and form == "通知":
            item_warnings.append(f"{eid} 展示类证据识别为通知，建议人工复核")

        if et == "其他材料" and form in ("合同", "发票") and tx_hit >= 2:
            item_warnings.append(f"{eid} 其他材料疑似交易类，建议复核是否应归交易凭证")

        errors.extend(item_errors)
        warnings.extend(item_warnings)
        checks.append(
            {
                "evidence_id": eid,
                "type": et,
                "form_type": form,
                "tx_hit": tx_hit,
                "promo_hit": promo_hit,
                "official_hit": official_hit,
                "result": "FAIL" if item_errors else ("WARN" if item_warnings else "PASS"),
                "messages": item_errors + item_warnings,
            }
        )

    return {
        "status": "FAIL" if errors else "PASS",
        "errors": errors,
        "warnings": warnings,
        "metrics": {
            "total": len(rows),
            "fail_count": len([c for c in checks if c.get("result") == "FAIL"]),
            "warn_count": len([c for c in checks if c.get("result") == "WARN"]),
        },
        "checks": checks,
    }


def build_casebook(args):
    evidence_dir = Path(args.evidence_dir) if (args.evidence_dir or "").strip() else None
    evidence_dir_fast = Path(args.evidence_dir_fast) if (args.evidence_dir_fast or "").strip() else None
    evidence_dir_full = Path(args.evidence_dir_full) if (args.evidence_dir_full or "").strip() else None
    xlsx_template = Path(args.xlsx_template)
    xlsx_out = Path(args.xlsx_out)
    ocr_cache = Path(args.ocr_cache)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    run_id = dt.datetime.now().strftime("run_%Y%m%d-%H%M%S")
    logs_dir = out_dir / "logs"
    set_run_context(logs_dir, run_id)
    global LOGGER
    LOGGER = setup_logger(__name__, log_dir=logs_dir)
    audit({
        "type": "run_start",
        "step": "build_casebook",
        "ok": True,
        "run_mode": getattr(args, "scan_mode", ""),
        "out_dir": str(out_dir),
    })

    source_dirs: List[Tuple[Path, str, str]] = []
    if evidence_dir_fast:
        source_dirs.append((evidence_dir_fast, "fast", "fast"))
    if evidence_dir_full:
        source_dirs.append((evidence_dir_full, "full", "full"))
    if evidence_dir:
        # 不分类目录为并行模块：扫描深度由 scan_mode 控制，来源分组固定为 uncategorized。
        source_dirs.append((evidence_dir, args.scan_mode, "uncategorized"))
    has_direct_bind = bool((args.direct_bind_config or "").strip())
    if not source_dirs and not has_direct_bind:
        raise RuntimeError("No evidence directory provided. Please set --evidence-dir and/or --evidence-dir-fast/--evidence-dir-full, or provide --direct-bind-config.")

    files_by_key: Dict[str, Path] = {}
    prefill_by_key: Dict[str, str] = {}
    mode_by_key: Dict[str, str] = {}
    bucket_by_key: Dict[str, str] = {}
    source_meta_by_key: Dict[str, Dict[str, str]] = {}
    input_stats = {
        "pdf": 0,
        "images_converted": 0,
        "office_converted": 0,
        "office_skipped": 0,
        "office_text_extracted": 0,
        "office_converter": "soffice" if _soffice_available() else "unavailable",
        "fast_input_count": 0,
        "full_input_count": 0,
        "direct_bind_items": 0,
        "direct_bind_files": 0,
        "direct_bind_images_converted": 0,
        "direct_bind_office_converted": 0,
        "direct_bind_office_skipped": 0,
    }

    for src_dir, src_mode, src_bucket in source_dirs:
        files_i, stats_i, prefill_i, mode_i, source_meta_i = collect_evidence_pdf_inputs(src_dir, out_dir, evidence_mode=src_mode)
        input_stats["pdf"] += int(stats_i.get("pdf", 0) or 0)
        input_stats["images_converted"] += int(stats_i.get("images_converted", 0) or 0)
        input_stats["office_converted"] += int(stats_i.get("office_converted", 0) or 0)
        input_stats["office_skipped"] += int(stats_i.get("office_skipped", 0) or 0)
        input_stats["office_text_extracted"] += int(stats_i.get("office_text_extracted", 0) or 0)
        if stats_i.get("office_converter", "") == "soffice":
            input_stats["office_converter"] = "soffice"
        for p in files_i:
            k = str(p.resolve())
            files_by_key[k] = p
            m = (mode_i.get(p, src_mode) or src_mode).lower()
            prev_m = mode_by_key.get(k, "")
            if prev_m != "full":
                mode_by_key[k] = "full" if m == "full" else ("fast" if m == "fast" else ("mix" if m == "mix" else "fast"))
            prev_bucket = bucket_by_key.get(k, src_bucket)
            bucket_by_key[k] = merge_bucket_priority(prev_bucket, src_bucket)
            t = prefill_i.get(p, "")
            if t:
                prefill_by_key[k] = t
            source_meta_by_key[k] = source_meta_i.get(p, {"origin_rel": p.name, "origin_dir": "", "origin_name": p.name})

    files = sorted(files_by_key.values())
    prefilled_texts: Dict[Path, str] = {}
    scan_mode_by_file: Dict[Path, str] = {}
    source_bucket_by_file: Dict[Path, str] = {}
    source_meta_by_file: Dict[Path, Dict[str, str]] = {}
    for k, p in files_by_key.items():
        m = mode_by_key.get(k, args.scan_mode)
        scan_mode_by_file[p] = m
        source_bucket_by_file[p] = bucket_by_key.get(k, "uncategorized")
        source_meta_by_file[p] = source_meta_by_key.get(k, {"origin_rel": p.name, "origin_dir": "", "origin_name": p.name})
        if m == "fast":
            input_stats["fast_input_count"] += 1
        elif m == "full":
            input_stats["full_input_count"] += 1
        t = prefill_by_key.get(k, "")
        if t:
            prefilled_texts[p] = t

    direct_bind_files, direct_bind_stats = collect_direct_bind_pdf_inputs(args.direct_bind_config, out_dir)
    input_stats["direct_bind_items"] = int(direct_bind_stats.get("direct_bind_items", 0) or 0)
    input_stats["direct_bind_files"] = int(direct_bind_stats.get("direct_bind_files", 0) or 0)
    input_stats["direct_bind_images_converted"] = int(direct_bind_stats.get("direct_bind_images_converted", 0) or 0)
    input_stats["direct_bind_office_converted"] = int(direct_bind_stats.get("direct_bind_office_converted", 0) or 0)
    input_stats["direct_bind_office_skipped"] = int(direct_bind_stats.get("direct_bind_office_skipped", 0) or 0)

    if not files and not direct_bind_files:
        dirs_disp = ", ".join(str(d) for d, _, _ in source_dirs)
        raise RuntimeError(f"No supported evidence found in {dirs_disp} and direct-bind inputs (PDF/JPG/JPEG/PNG/BMP/TIF/TIFF/WEBP/OFFICE)")

    # 默认全量扫描；fast 模式可按 max-pages 控制
    normal_pages = args.max_pages
    notice_pages = args.notice_max_pages
    if args.scan_mode == "fast":
        if normal_pages <= 0:
            normal_pages = 2
        if notice_pages <= 0:
            notice_pages = 4

    scan_texts: Dict[Path, str] = {}
    notice_candidates: List[Path] = []
    scan_rounds: List[Dict[str, Any]] = []
    round1_case: Dict[str, str] = {}
    unique_files: List[Path] = []
    dup_by_canonical: Dict[Path, List[Path]] = {}
    if files:
        scan_texts, notice_candidates, scan_rounds, round1_case = run_fixed_five_round_scans(
            files=files,
            ocr_cache=ocr_cache,
            normal_pages=normal_pages,
            notice_pages=notice_pages,
            dpi=args.dpi,
            notice_dpi=args.notice_dpi,
            prefilled_texts=prefilled_texts,
            scan_mode_by_file=scan_mode_by_file,
        )
        unique_files, dup_by_canonical, _ = deduplicate_scanned_files(files, scan_texts)

    notice_text = "\n".join([scan_texts.get(p, "") for p in notice_candidates])
    if args.notice_text_hint and Path(args.notice_text_hint).exists():
        notice_text += "\n" + Path(args.notice_text_hint).read_text(encoding="utf-8", errors="ignore")

    all_text = "\n".join(scan_texts.values())
    notice_case = merge_case_info(round1_case, infer_case_info(notice_text))
    evidence_case = infer_case_info(all_text)
    case = merge_case_info(notice_case, evidence_case)
    fact_dirs = [d for d, _, _ in source_dirs]
    fact_case = infer_case_info_from_fact_docs(files, scan_texts, evidence_dir=evidence_dir, evidence_dirs=fact_dirs)
    case = apply_fact_case_overrides(case, fact_case)
    override_case = load_case_meta_override(getattr(args, "override_case_meta", ""))

    reg_no = case.get("reg_no") or "（待补注册号）"
    class_no = case.get("class") or "24"
    mark_name = pick_best_mark(
        candidates=[
            case.get("mark_name", ""),
            infer_mark_from_texts(list(scan_texts.values())),
            infer_mark_from_filenames(files),
        ],
        texts=list(scan_texts.values()),
        files=files,
    ) or "（待补商标文字）"
    if fact_case.get("mark_name"):
        mark_name = fact_case.get("mark_name", "").strip() or mark_name
    respondent = case.get("respondent") or "（待补注册人）"
    applicant = case.get("applicant") or "（待补申请人）"
    period_start = case.get("use_period_start") or "2021-03-01"
    period_end = case.get("use_period_end") or "2024-02-29"
    if override_case.get("reg_no"):
        reg_no = override_case.get("reg_no", "").strip() or reg_no
    if override_case.get("class"):
        class_no = override_case.get("class", "").strip() or class_no
    if override_case.get("mark_name"):
        mark_name = override_case.get("mark_name", "").strip() or mark_name
    if override_case.get("respondent"):
        respondent = override_case.get("respondent", "").strip() or respondent
    if override_case.get("applicant"):
        applicant = override_case.get("applicant", "").strip() or applicant
    if override_case.get("use_period_start"):
        period_start = override_case.get("use_period_start", "").strip() or period_start
    if override_case.get("use_period_end"):
        period_end = override_case.get("use_period_end", "").strip() or period_end
    if args.case_reg_no:
        reg_no = str(args.case_reg_no).strip()
    if args.case_class:
        class_no = str(args.case_class).strip()
    if args.case_mark_name:
        mark_name = str(args.case_mark_name).strip()
    if args.case_respondent:
        respondent = str(args.case_respondent).strip()
    if args.case_applicant:
        applicant = str(args.case_applicant).strip()
    if args.case_period_start:
        period_start = str(args.case_period_start).strip()
    if args.case_period_end:
        period_end = str(args.case_period_end).strip()
    period_start_dt = parse_iso_date(period_start)
    period_end_dt = parse_iso_date(period_end)
    goods_desc = f"第{class_no}类：{case.get('goods_hint') or '核定商品（以注册证记载为准）'}等（以注册证记载为准）"
    if override_case.get("designated_goods_services"):
        goods_desc = override_case.get("designated_goods_services", "").strip() or goods_desc
    if args.case_goods_services:
        goods_desc = str(args.case_goods_services).strip()
    revoked_mark_name = mark_name
    revoked_goods_services = case.get("goods_hint") or goods_desc
    if override_case.get("revoked_goods_services"):
        revoked_goods_services = override_case.get("revoked_goods_services", "").strip() or revoked_goods_services
    if args.case_goods_services:
        revoked_goods_services = str(args.case_goods_services).strip()
    if args.case_revoked_goods_services:
        revoked_goods_services = str(args.case_revoked_goods_services).strip()
    defense_goods_services = revoked_goods_services
    if override_case.get("defense_goods_services"):
        defense_goods_services = override_case.get("defense_goods_services", "").strip() or defense_goods_services
    if args.case_defense_goods_services:
        defense_goods_services = str(args.case_defense_goods_services).strip()
    goods_canonical_list = parse_goods_canonical_list(revoked_goods_services or goods_desc)
    mark_image = override_case.get("mark_image", "").strip() if override_case.get("mark_image") else ""
    respondent_address = override_case.get("respondent_address", "").strip() if override_case.get("respondent_address") else ""
    agent_company = override_case.get("agent_company", "").strip() if override_case.get("agent_company") else ""
    agent_address = override_case.get("agent_address", "").strip() if override_case.get("agent_address") else ""
    contact_phone = override_case.get("contact_phone", "").strip() if override_case.get("contact_phone") else ""
    if override_case.get("cnipa_notice_ref"):
        cnipa_notice_ref = override_case.get("cnipa_notice_ref", "").strip() or "自动识别（全文扫描）"
    else:
        cnipa_notice_ref = "自动识别（全文扫描）"
    if override_case.get("case_no"):
        case_no = override_case.get("case_no", "").strip() or f"撤三自动识别-{reg_no}"
    else:
        case_no = f"撤三自动识别-{reg_no}"
    if args.case_mark_image:
        mark_image = str(args.case_mark_image or "").strip()
    if args.case_respondent_address:
        respondent_address = str(args.case_respondent_address or "").strip()
    if args.case_agent_company:
        agent_company = str(args.case_agent_company or "").strip()
    if args.case_agent_address:
        agent_address = str(args.case_agent_address or "").strip()
    if args.case_contact_phone:
        contact_phone = str(args.case_contact_phone or "").strip()

    case_final_for_validate = {
        "case_no": case_no,
        "reg_no": reg_no,
        "respondent": respondent,
        "applicant": applicant,
        "class": class_no,
        "designated_goods_services": goods_desc,
        "use_period_start": period_start,
        "use_period_end": period_end,
        "cnipa_notice_ref": cnipa_notice_ref,
        "mark_name": mark_name,
    }
    validation_result = validate_caseinfo_and_crosscheck(
        case_final=case_final_for_validate,
        notice_case_raw=notice_case,
        evidence_case_raw=evidence_case,
        allow_noisy_mark_name=(args.allow_noisy_mark_name == "on"),
    )
    validation_json_path = out_dir / "caseinfo_validation.json"
    validation_json_path.write_text(json.dumps(validation_result, ensure_ascii=False, indent=2), encoding="utf-8")
    # 质量校验仅用于内部风险评估，不作为流程门禁中止条件。

    rows: List[Dict[str, str]] = []
    duplicate_map_by_id: Dict[str, List[str]] = {}
    for i, p in enumerate(unique_files, start=1):
        evidence_id = evidence_id_display(i)
        audit({
            "type": "evidence_process_start",
            "step": "build_casebook",
            "file": str(p),
            "evidence_id": evidence_id,
            "ok": True,
        })
        text = scan_texts.get(p, "")
        src_meta = source_meta_by_file.get(p, {"origin_rel": p.name, "origin_dir": "", "origin_name": p.name})
        origin_rel = str(src_meta.get("origin_rel", "") or p.name)
        origin_dir = str(src_meta.get("origin_dir", "") or "")
        origin_name = str(src_meta.get("origin_name", "") or p.name)
        source_hint = " / ".join([x for x in [origin_rel, origin_dir, origin_name] if x])
        key_lines = extract_key_lines(text, mark_name, limit=5)
        mark_time_lines = extract_mark_time_lines(text, mark_name, limit=3)
        extraction_notes: List[str] = []
        dup_files = dup_by_canonical.get(p, [])
        if dup_files:
            extraction_notes.append(f"去重合并{len(dup_files)}份重复材料")

        # 时间锚点优先采用“商标词条+日期”命中语句
        date_source = "\n".join(mark_time_lines) if mark_time_lines else ("\n".join(key_lines) if key_lines else text)
        dates = extract_dates_advanced(
            text=date_source,
            period_start=period_start_dt,
            period_end=period_end_dt,
        )
        if not dates:
            dates = extract_dates_advanced(
                text=f"{p.name}\n{text}",
                period_start=period_start_dt,
                period_end=period_end_dt,
            )

        review_dates, review_note, review_lines = extract_keyword_scene_dates(
            text=text,
            period_start=period_start_dt,
            period_end=period_end_dt,
            keywords=REVIEW_KEYWORDS,
            scene_name="评价场景",
        )
        order_dates, order_note, order_lines = extract_order_dates(
            text=text,
            period_start=period_start_dt,
            period_end=period_end_dt,
        )
        logistics_dates, logistics_note, logistics_lines = extract_logistics_dates(
            text=text,
            period_start=period_start_dt,
            period_end=period_end_dt,
        )

        time_pick = build_time_anchor_candidates(
            base_dates=dates,
            review_dates=review_dates,
            order_dates=order_dates,
            logistics_dates=logistics_dates,
            period_start=period_start_dt,
            period_end=period_end_dt,
            top_k=5,
            context_text=text,
        )
        resolved_dates = [parse_iso_date(x) for x in time_pick.get("selected_dates", [])]
        resolved_dates = [x for x in resolved_dates if x]
        resolve_note = time_pick.get("note", "")
        channels_used = list(time_pick.get("channels_used", []))
        if resolved_dates:
            dates = sorted(set([d for d in dates if d] + resolved_dates))
        if review_note:
            extraction_notes.append(review_note)
        if order_note:
            extraction_notes.append(order_note)
        if logistics_note:
            extraction_notes.append(logistics_note)
        if resolve_note:
            extraction_notes.append(resolve_note)
        mark_time_lines = unique_keep_order(mark_time_lines + review_lines + order_lines + logistics_lines, limit=6)

        # 仍未识别到明确时间且疑似“评价/交易场景”时，启动高精度时间增强 OCR（仅提时点）
        need_boost = (not dates) and (
            any(k in text for k in REVIEW_KEYWORDS + ["评论", "追评"])
            or any(k in p.name for k in ["店铺", "手机端", "电脑端", "评价", "交易", "订单"])
        )
        if need_boost:
            boost_text = ocr_time_boost_lines(
                pdf=p,
                cache_dir=ocr_cache,
                max_pages=min(3, pdf_page_count(p) or 3),
                dpi=max(args.dpi, 340),
            )
            if boost_text:
                boost_review_dates, boost_review_note, boost_review_lines = extract_keyword_scene_dates(
                    text=boost_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                    keywords=REVIEW_KEYWORDS,
                    scene_name="评价场景",
                )
                boost_order_dates, boost_order_note, boost_order_lines = extract_order_dates(
                    text=boost_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                )
                boost_log_dates, boost_log_note, boost_log_lines = extract_logistics_dates(
                    text=boost_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                )
                boost_dates = extract_dates_advanced(
                    text=boost_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                )
                boost_time_pick = build_time_anchor_candidates(
                    base_dates=boost_dates,
                    review_dates=boost_review_dates,
                    order_dates=boost_order_dates,
                    logistics_dates=boost_log_dates,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                    top_k=5,
                    context_text=boost_text,
                )
                boost_resolved_dates = [parse_iso_date(x) for x in boost_time_pick.get("selected_dates", [])]
                boost_resolved_dates = [x for x in boost_resolved_dates if x]
                boost_resolve_note = boost_time_pick.get("note", "")
                boost_channels = list(boost_time_pick.get("channels_used", []))
                if boost_resolved_dates:
                    dates = sorted(set(dates + boost_resolved_dates))
                    channels_used = sorted(set(channels_used + boost_channels))
                    extraction_notes.append("已启用高精度时间增强 OCR")
                    for n in [boost_review_note, boost_order_note, boost_log_note, boost_resolve_note]:
                        if n:
                            extraction_notes.append(n)
                    # 增强识别线索并入锚点候选
                    boost_lines = [ln.strip() for ln in boost_text.splitlines() if ln.strip()]
                    boost_lines = unique_keep_order(boost_lines + boost_review_lines + boost_order_lines + boost_log_lines, limit=8)
                    mark_time_lines = unique_keep_order(mark_time_lines + boost_lines, limit=5)

        # 二次深度增强：针对仍无时间且存在评价/评论语义的展示材料，扩大扫描页数
        need_deep_boost = (not dates) and any(k in text for k in REVIEW_KEYWORDS + ["评论", "追评", "已购"])
        if need_deep_boost:
            deep_text = ocr_time_boost_lines(
                pdf=p,
                cache_dir=ocr_cache,
                max_pages=min(8, pdf_page_count(p) or 8),
                dpi=420,
            )
            if deep_text:
                deep_review_dates, deep_review_note, deep_review_lines = extract_keyword_scene_dates(
                    text=deep_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                    keywords=REVIEW_KEYWORDS,
                    scene_name="评价场景",
                )
                deep_order_dates, deep_order_note, deep_order_lines = extract_order_dates(
                    text=deep_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                )
                deep_log_dates, deep_log_note, deep_log_lines = extract_logistics_dates(
                    text=deep_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                )
                deep_dates = extract_dates_advanced(
                    text=deep_text,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                )
                deep_time_pick = build_time_anchor_candidates(
                    base_dates=deep_dates,
                    review_dates=deep_review_dates,
                    order_dates=deep_order_dates,
                    logistics_dates=deep_log_dates,
                    period_start=period_start_dt,
                    period_end=period_end_dt,
                    top_k=5,
                    context_text=deep_text,
                )
                deep_resolved_dates = [parse_iso_date(x) for x in deep_time_pick.get("selected_dates", [])]
                deep_resolved_dates = [x for x in deep_resolved_dates if x]
                deep_resolve_note = deep_time_pick.get("note", "")
                deep_channels = list(deep_time_pick.get("channels_used", []))
                if deep_resolved_dates:
                    dates = sorted(set(dates + deep_resolved_dates))
                    channels_used = sorted(set(channels_used + deep_channels))
                    extraction_notes.append("已启用深度时间增强 OCR")
                    for n in [deep_review_note, deep_order_note, deep_log_note, deep_resolve_note]:
                        if n:
                            extraction_notes.append(n)
                    deep_lines = [ln.strip() for ln in deep_text.splitlines() if ln.strip()]
                    deep_lines = unique_keep_order(deep_lines + deep_review_lines + deep_order_lines + deep_log_lines, limit=10)
                    mark_time_lines = unique_keep_order(mark_time_lines + deep_lines, limit=6)

        meta_used = False
        if not dates:
            md_date = parse_pdf_metadata_date(p)
            if md_date:
                use_meta = True
                if period_start_dt and period_end_dt:
                    # 避免把导出时间（常见为批量导出日）误当作证据时间
                    lo = period_start_dt - dt.timedelta(days=30)
                    hi = period_end_dt + dt.timedelta(days=30)
                    use_meta = lo <= md_date <= hi
                if use_meta:
                    dates = [md_date]
                    meta_used = True
                    extraction_notes.append("OCR未识别到稳定日期，已使用文件元数据日期")
                else:
                    extraction_notes.append("元数据日期超出指定期间窗口，未采用")

        dates, date_clean_note = clean_abnormal_dates(
            dates=dates,
            period_start=period_start_dt,
            period_end=period_end_dt,
        )
        if date_clean_note:
            extraction_notes.append(date_clean_note)

        form_profile = infer_evidence_form_profile(p, text, origin_hint=source_hint)
        cls = classify_evidence(
            p.name,
            text,
            mark_name,
            respondent,
            form_profile=form_profile,
            source_hint=source_hint,
        )
        form_type = str(cls.get("Form Type", "") or "")
        form_conf = str(cls.get("Form Confidence", "") or "")
        form_signals = str(cls.get("Form Signals", "") or "")
        final_pick = build_time_anchor_candidates(
            base_dates=dates,
            review_dates=review_dates,
            order_dates=order_dates,
            logistics_dates=logistics_dates,
            period_start=period_start_dt,
            period_end=period_end_dt,
            top_k=5,
            context_text=text,
            evidence_kind=cls.get("Type", ""),
        )
        selected_iso = final_pick.get("selected_dates", [])
        selected_dates_final = [parse_iso_date(x) for x in selected_iso]
        selected_dates_final = [x for x in selected_dates_final if x]
        if selected_dates_final:
            dates = selected_dates_final
        channels_used = list(final_pick.get("channels_used", [])) or channels_used
        if final_pick.get("note"):
            extraction_notes.append(final_pick.get("note", ""))

        ds = to_iso(dates[0]) if dates else ""
        de = to_iso(dates[-1]) if dates else ""
        raw_conf = confidence_from_dates(
            dates=dates,
            mark_time_lines=mark_time_lines,
            review_dates=review_dates,
            meta_used=meta_used,
            time_channels=channels_used,
        )

        selected_anchor_type = str(final_pick.get("selected_anchor_type", "uncertain") or "uncertain")
        selected_anchor_source = str(final_pick.get("selected_anchor_source", "基础") or "基础")
        web_verifiability = infer_web_verifiability(cls.get("Type", ""), text, key_lines)
        time_gate_scope = should_time_gate_evaluate(
            kind=cls.get("Type", ""),
            text=text,
            time_channels=channels_used,
            key_lines=key_lines,
            anchor_type=selected_anchor_type,
            web_verifiability=web_verifiability,
        )
        coverage_lane = COVERAGE_LANE_A if time_gate_scope else COVERAGE_LANE_B
        conf = raw_conf if time_gate_scope else "N/A"
        mark_conf = mark_name_confidence(
            mark_name=mark_name,
            text=text,
            filename=p.name,
            key_lines=key_lines,
            mark_shown_flag=cls.get("Mark Shown (Y/N)", "N"),
        )
        taxonomy = infer_evidence_taxonomy(
            cls=cls,
            text=text,
            case_class=class_no,
            goods_services=goods_desc,
        )
        inferred_name, inferred_purpose = infer_proof_name_and_purpose(
            cls=cls,
            text=text,
            review_dates=review_dates,
            order_dates=order_dates,
            logistics_dates=logistics_dates,
            case_class=class_no,
            goods_services=goods_desc,
        )
        if taxonomy.get("category", "").startswith("4. 市场推广与广告类证据") and not any(
            k in (inferred_name or "") for k in ["海报", "广告", "活动组织"]
        ):
            inferred_name = "活动组织广告图片"
        if inferred_purpose in (
            "证明产品、包装、场景或页面展示中的实际使用状态。",
            "证明工厂/仓储/门店等实际经营场景的存在。",
            "证明服务活动现场或实际经营场景的存在。",
            "证明相关使用事实。",
        ):
            inferred_purpose = str(taxonomy.get("objective", "") or inferred_purpose)
        mapping_basis = build_mapping_basis(
            text=text,
            cls=cls,
            key_lines=key_lines,
            mark_time_lines=mark_time_lines,
            channels_used=channels_used,
        )
        goods_match = infer_goods_match(
            canonical_goods=goods_canonical_list,
            text=text,
            key_lines=key_lines,
            evidence_name=p.name,
        )
        entity_names = extract_entity_names(text, key_lines)
        entity_consistency = infer_entity_consistency(entity_names, respondent, text)
        summary, proof_fact, anchor_text = build_evidence_summary(
            cls=cls,
            mark_name=mark_name,
            respondent=respondent,
            date_start=ds,
            date_end=de,
            raw_text=text,
            key_lines=key_lines,
            mark_time_lines=mark_time_lines,
            inferred_name=inferred_name,
            inferred_purpose=inferred_purpose,
            case_class=class_no,
            goods_services=goods_desc,
            taxonomy=taxonomy,
        )
        if goods_match.get("item_index"):
            mapping_basis = (
                f"{mapping_basis}；商品对应：{goods_match.get('level','G3')}"
                f"(核定条目{goods_match.get('item_index','')}: {goods_match.get('item_name','')})"
            )
        if extraction_notes:
            note_txt = "；".join(unique_keep_order(extraction_notes, limit=3))
            summary = f"{summary}；时间识别说明：{note_txt}"
            proof_fact = f"{proof_fact}；时间识别说明：{note_txt}"
        if not time_gate_scope:
            proof_fact = f"{proof_fact}；该证据属于{COVERAGE_LANE_B}，不计入指定期间覆盖计算"
        if goods_match.get("level") == "G3":
            proof_fact = f"{proof_fact}；商品服务对应待补强"
        if entity_consistency.get("level") in ("E2", "E3"):
            proof_fact = f"{proof_fact}；主体链提示：{entity_consistency.get('basis','')}"
        proof_fact = f"{proof_fact}；映射依据：{mapping_basis}"
        public_pending_fact = sanitize_public_fact_text(proof_fact, inferred_purpose)

        pages = pdf_page_count(p)
        pr = f"1-{pages}" if pages else ""

        shop_source = infer_shop_source(str(p), text, key_lines) if cls["Type"] == "线上店铺展示" else ""
        product_label = infer_product_phrase(text, key_lines)
        product_key = build_product_match_key(
            product_label=product_label,
            inferred_name=inferred_name,
            text=text,
            key_lines=key_lines,
        )
        renamed_name = build_renamed_filename(
            i,
            cls["Type"],
            ds,
            mark_name,
            key_lines,
            source_label=shop_source,
            product_label=product_label,
            preferred_topic=inferred_name,
        )
        trade_role = infer_trade_link_role(f"{origin_rel} {p.name}", text) if cls["Type"] == "交易凭证" else ""
        trade_amount = extract_trade_amount(text, key_lines) if cls["Type"] == "交易凭证" else ""
        trade_party = extract_trade_counterparty(text, key_lines) if cls["Type"] == "交易凭证" else ""
        trade_key = extract_trade_pair_key(
            text=text,
            key_lines=key_lines,
            date_start=ds,
            date_end=de,
            fallback_name=product_label or inferred_name or renamed_name,
            amount_hint=trade_amount,
            party_hint=trade_party,
            product_hint=product_key,
        ) if cls["Type"] == "交易凭证" else ""

        rows.append({
            "No": i,
            "Evidence ID": evidence_id,
            "Evidence Name": renamed_name,
            "Type": cls["Type"],
            "Source": "证据材料（全文扫描去噪+去重）",
            "Formation Date": ds if ds and ds == de else "",
            "Carrier": str(p),
            "Goods/Services": goods_desc,
            "Proof Target (T1-T6)": cls["Proof Target (T1-T6)"],
            "SJ-6 (A/R/C/T/L/X)": cls["SJ-6 (A/R/C/T/L/X)"],
            "Page Range": pr,
            "Risk Notes": proof_fact,
            "Public Pending Fact": public_pending_fact,
            "Mapping Basis": mapping_basis,
            "Inferred Proof Name": inferred_name,
            "Inferred Purpose": inferred_purpose,
            "Mark Shown (Y/N)": cls["Mark Shown (Y/N)"],
            "Subject Match (Y/N)": cls["Subject Match (Y/N)"],
            "Commercial Loop (Y/N)": cls["Commercial Loop (Y/N)"],
            "Original/Verifiable (Y/N)": cls["Original/Verifiable (Y/N)"],
            "Evidence Date Start": ds,
            "Evidence Date End": de,
            "Time Anchor Text": anchor_text,
            "Date Confidence": conf,
            "Time Gate Scope": "Y" if time_gate_scope else "N",
            "Coverage Lane": coverage_lane,
            "Time Anchor Type": selected_anchor_type,
            "Time Anchor Source": selected_anchor_source,
            "Time Anchor Allowed": "Y" if time_gate_scope else "N",
            "Web Verifiability": web_verifiability,
            "Goods Match Level": goods_match.get("level", "G3"),
            "Goods Match Item Index": goods_match.get("item_index", ""),
            "Goods Match Item Name": goods_match.get("item_name", ""),
            "Goods Match Basis": goods_match.get("basis", ""),
            "Entity Names": "、".join(entity_names) if entity_names else "",
            "Entity Consistency": entity_consistency.get("level", "E2"),
            "Entity Consistency Basis": entity_consistency.get("basis", ""),
            "Mark Name Confidence": mark_conf,
            "Trade Link Role": trade_role,
            "Trade Amount": trade_amount,
            "Trade Counterparty": trade_party,
            "Trade Link Key": trade_key,
            "Trade Chain Group": "",
            "Trade Chain Group No": "",
            "Product Key": product_key,
            "Time Channels": "、".join(channels_used) if channels_used else "",
            "Time Anchor Candidates": json.dumps(final_pick.get("candidates", []), ensure_ascii=False),
            "Time Conflict Flag": "Y" if final_pick.get("conflict_flag", False) else "N",
            "Time Selection Rule": final_pick.get("selection_rule", ""),
            "Form Type": form_type,
            "Form Confidence": form_conf,
            "Form Signals": form_signals,
            "Content Summary": summary,
            "Evidence Taxonomy": taxonomy.get("category", ""),
            "Proof Objective": taxonomy.get("objective", ""),
            "Proposition": taxonomy.get("proposition", ""),
            "Proof Strength": taxonomy.get("strength", ""),
            "Risk Tag": taxonomy.get("risk_tag", ""),
            "Next Action": taxonomy.get("next_action", ""),
            "Original Relative Path": origin_rel,
            "_src": str(p),
            "_src_name": origin_name,
            "_origin_rel": origin_rel,
            "_origin_dir": origin_dir,
            "_source_bucket": source_bucket_by_file.get(p, "uncategorized"),
            "_renamed_name": renamed_name,
            "_dup_files": [x.name for x in dup_files],
        })
        audit({
            "type": "evidence_process_end",
            "step": "build_casebook",
            "file": str(p),
            "evidence_id": evidence_id,
            "ok": True,
        })

    # 直装材料（不参与OCR扫描）追加到证据列表，参与统一编排和装订。
    if direct_bind_files:
        seq_start = len(rows)
        for j, item in enumerate(direct_bind_files, start=1):
            src_pdf = Path(item.get("src_pdf", ""))
            if not src_pdf.exists():
                continue
            evidence_id = evidence_id_display(seq_start + j)
            label = str(item.get("label", "") or "").strip() or "其他"
            custom_name = str(item.get("custom_name", "") or "").strip()
            src_name = str(item.get("src_name", "") or src_pdf.name)
            profile = classify_direct_bind_profile(label)
            inferred_name = custom_name or label
            if not inferred_name:
                inferred_name = sanitize_component(src_pdf.stem, 24)
            topic_tokens = [inferred_name, label, src_pdf.stem]
            renamed_name = build_renamed_filename(
                seq_start + j,
                profile.get("Type", "其他材料"),
                "",
                mark_name,
                topic_tokens,
            )
            purpose = profile.get("Inferred Purpose", "作为补强材料使用。")
            mapping_basis = f"直装材料：{label}（未参与OCR扫描）"
            summary = f"该证据作为“{label}”材料直接装订，{purpose}"
            risk_notes = f"{purpose}；该证据未参与OCR扫描，按人工命名与路径直接装订。；映射依据：{mapping_basis}"
            public_pending_fact = sanitize_public_fact_text(risk_notes, purpose)
            rows.append({
                "No": seq_start + j,
                "Evidence ID": evidence_id,
                "Evidence Name": renamed_name,
                "Type": profile.get("Type", "其他材料"),
                "Source": "直装材料（不参与OCR扫描）",
                "Formation Date": "",
                "Carrier": str(src_pdf),
                "Goods/Services": goods_desc,
                "Proof Target (T1-T6)": profile.get("Proof Target (T1-T6)", "T5"),
                "SJ-6 (A/R/C/T/L/X)": profile.get("SJ-6 (A/R/C/T/L/X)", "R,L"),
                "Page Range": "",
                "Risk Notes": risk_notes,
                "Public Pending Fact": public_pending_fact,
                "Mapping Basis": mapping_basis,
                "Inferred Proof Name": inferred_name,
                "Inferred Purpose": purpose,
                "Mark Shown (Y/N)": profile.get("Mark Shown (Y/N)", "N"),
                "Subject Match (Y/N)": profile.get("Subject Match (Y/N)", "N"),
                "Commercial Loop (Y/N)": profile.get("Commercial Loop (Y/N)", "N"),
                "Original/Verifiable (Y/N)": profile.get("Original/Verifiable (Y/N)", "Y"),
                "Evidence Date Start": "",
                "Evidence Date End": "",
                "Time Anchor Text": "",
                "Date Confidence": "N/A",
                "Time Gate Scope": "N",
                "Coverage Lane": COVERAGE_LANE_B,
                "Time Anchor Type": "uncertain",
                "Time Anchor Source": "uncertain",
                "Time Anchor Allowed": "N",
                "Web Verifiability": "",
                "Goods Match Level": "G3",
                "Goods Match Item Index": "",
                "Goods Match Item Name": "",
                "Goods Match Basis": "直装证据未参与商品自动匹配",
                "Entity Names": "",
                "Entity Consistency": "E2",
                "Entity Consistency Basis": "直装证据未参与主体链自动识别",
                "Mark Name Confidence": "N/A",
                "Trade Link Role": "",
                "Trade Amount": "",
                "Trade Counterparty": "",
                "Trade Link Key": "",
                "Trade Chain Group": "",
                "Trade Chain Group No": "",
                "Product Key": "",
                "Time Channels": "",
                "Time Anchor Candidates": "[]",
                "Time Conflict Flag": "N",
                "Time Selection Rule": "直装材料不参与时间锚点筛选",
                "Content Summary": summary,
                "_src": str(src_pdf),
                "_src_name": src_name,
                "_source_bucket": "direct_bind",
                "_renamed_name": renamed_name,
                "_dup_files": [],
                "_direct_bind_label": label,
            })

    for r in rows:
        r["Unit Type"] = detect_unit_type(r)
    rows = merge_fragmented_evidence_units(
        rows,
        out_dir=out_dir,
        mark_name=mark_name,
        case_class=class_no,
        goods_services=goods_desc,
    )
    assign_trade_chain_groups(rows)
    for r in rows:
        r["Unit Type"] = detect_unit_type(r)

    rows = sort_rows_for_progressive_review(rows, period_start_dt, period_end_dt)
    resequence_rows(rows, mark_name)
    for r in rows:
        r["Public Pending Fact"] = sanitize_public_fact_text(
            str(r.get("Risk Notes", "") or ""),
            str(r.get("Inferred Purpose", "") or ""),
        )
    low_confidence_info = write_low_confidence_fields(
        rows=rows,
        out_dir=out_dir,
        run_id=run_id,
        logger=LOGGER,
        audit=audit,
    )
    name_quality = validate_evidence_name_quality(
        rows=rows,
        mark_name=mark_name,
        out_dir=out_dir,
    )
    duplicate_map_by_id = {
        r.get("Evidence ID", ""): list(r.get("_dup_files", []))
        for r in rows
        if r.get("Evidence ID", "")
    }

    package_info = organize_evidence_package(
        rows,
        out_dir=out_dir,
        organize_dir=args.organize_dir,
        duplicate_map=duplicate_map_by_id,
    )
    page_map_validation = validate_page_map_consistency(rows, package_info)
    page_map_validation_path = out_dir / "page_map_validation.json"
    page_map_validation_path.write_text(
        json.dumps(page_map_validation, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    override_path = Path(args.override_rules) if args.override_rules else (out_dir / "人工覆盖规则.xlsx")
    override_tpl = create_manual_override_template(out_dir, rows, override_path)
    override_result = apply_manual_overrides(rows, override_path, out_dir=out_dir, run_id=run_id)
    time_quality = validate_time_quality(
        rows=rows,
        period_start=period_start_dt,
        period_end=period_end_dt,
        min_in_period_highmed=args.min_in_period_highmed,
        max_low_ratio=args.max_low_ratio,
        max_unknown_ratio=args.max_unknown_ratio,
    )
    time_quality_path = out_dir / "time_quality_validation.json"
    time_quality_path.write_text(json.dumps(time_quality, ensure_ascii=False, indent=2), encoding="utf-8")
    type_consistency = validate_evidence_type_consistency(rows)
    type_consistency_path = out_dir / "evidence_type_consistency.json"
    type_consistency_path.write_text(json.dumps(type_consistency, ensure_ascii=False, indent=2), encoding="utf-8")
    precheck_case_meta = {
        "case_no": case_no,
        "reg_no": reg_no,
        "class": class_no,
        "respondent": respondent,
        "use_period_start": period_start,
        "use_period_end": period_end,
    }
    precheck_path = create_pre_output_checklist_docx(out_dir=out_dir, case_meta=precheck_case_meta, rows=rows)
    precheck_approved = is_precheck_approved(precheck_path)
    case_meta_payload = {
        "key": f"{reg_no}_{class_no}_{period_start}_{period_end}",
        "case_no": case_no,
        "reg_no": reg_no,
        "respondent": respondent,
        "applicant": applicant,
        "class": class_no,
        "designated_goods_services": goods_desc,
        "use_period_start": period_start,
        "use_period_end": period_end,
        "cnipa_notice_ref": cnipa_notice_ref,
        "mark_name": mark_name,
        "source_of_truth": validation_result.get("source_of_truth", {}),
        "validated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "validated_by": "auto_recognize_and_generate",
        "run_id": run_id,
    }
    case_meta_path = out_dir / "case_meta.json"
    case_meta_path.write_text(json.dumps(case_meta_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    baseline_info = build_regression_baseline(
        out_dir=out_dir,
        files=files,
        rows=rows,
        period_start=period_start_dt,
        period_end=period_end_dt,
    )
    highlight_info = {
        "status": "DISABLED",
        "reason": "证据高亮功能已停用（按当前配置统一关闭）",
        "highlight_pdf": "",
        "highlight_report": "",
        "hits_total": 0,
        "evidence_count": 0,
        "skipped_ids": [],
    }

    case_for_report = {
        "mark_name": mark_name,
        "revoked_mark_name": revoked_mark_name,
        "revoked_goods_services": revoked_goods_services,
        "defense_goods_services": defense_goods_services,
        "mark_image": mark_image,
        "reg_no": reg_no,
        "class": class_no,
        "respondent": respondent,
        "respondent_address": respondent_address,
        "applicant": applicant,
        "agent_company": agent_company,
        "agent_address": agent_address,
        "contact_phone": contact_phone,
        "use_period_start": period_start,
        "use_period_end": period_end,
        "scan_rounds": json.dumps(scan_rounds, ensure_ascii=False),
        "scan_rounds_total": str(len(scan_rounds)),
        "run_id": run_id,
        "case_info_version": validation_result.get("case_info_version", CASE_INFO_VERSION),
        "validation_rule_version": validation_result.get("validation_rule_version", VALIDATION_RULE_VERSION),
        "caseinfo_validation_status": validation_result.get("status", ""),
        "caseinfo_validation_errors": "；".join(validation_result.get("errors", [])),
        "caseinfo_validation_warnings": "；".join(validation_result.get("warnings", [])),
        "caseinfo_validation_path": str(validation_json_path),
        "source_of_truth": json.dumps(validation_result.get("source_of_truth", {}), ensure_ascii=False),
        "dedup_before_count": str(len(files)),
        "dedup_after_count": str(len(unique_files)),
        "input_pdf_count": str(input_stats.get("pdf", 0)),
        "input_image_converted_count": str(input_stats.get("images_converted", 0)),
        "input_office_converted_count": str(input_stats.get("office_converted", 0)),
        "input_office_text_extracted_count": str(input_stats.get("office_text_extracted", 0)),
        "input_office_skipped_count": str(input_stats.get("office_skipped", 0)),
        "office_converter": str(input_stats.get("office_converter", "")),
        "input_fast_profile_count": str(input_stats.get("fast_input_count", 0)),
        "input_full_profile_count": str(input_stats.get("full_input_count", 0)),
        "direct_bind_item_count": str(input_stats.get("direct_bind_items", 0)),
        "direct_bind_file_count": str(input_stats.get("direct_bind_files", 0)),
        "direct_bind_image_converted_count": str(input_stats.get("direct_bind_images_converted", 0)),
        "direct_bind_office_converted_count": str(input_stats.get("direct_bind_office_converted", 0)),
        "direct_bind_office_skipped_count": str(input_stats.get("direct_bind_office_skipped", 0)),
        "override_rules_path": str(override_path),
        "override_applied_count": override_result.get("applied_count", "0"),
        "override_rules_json": override_result.get("override_rules_json", ""),
        "overrides_applied_path": override_result.get("overrides_applied_path", ""),
        "page_map_path": package_info.get("page_map_path", ""),
        "page_map_hash": package_info.get("page_map_hash", ""),
        "page_map_validation_status": page_map_validation.get("status", ""),
        "page_map_validation_errors": "；".join(page_map_validation.get("errors", [])),
        "page_map_validation_path": str(page_map_validation_path),
        "time_quality_status": time_quality.get("status", ""),
        "time_quality_errors": "；".join(time_quality.get("errors", [])),
        "time_quality_warnings": "；".join(time_quality.get("warnings", [])),
        "time_quality_path": str(time_quality_path),
        "evidence_type_consistency_status": type_consistency.get("status", ""),
        "evidence_type_consistency_errors": "；".join(type_consistency.get("errors", [])),
        "evidence_type_consistency_warnings": "；".join(type_consistency.get("warnings", [])),
        "evidence_type_consistency_path": str(type_consistency_path),
        "low_confidence_fields_path": low_confidence_info.get("path", ""),
        "low_confidence_field_count": str(low_confidence_info.get("count", 0)),
        "name_quality_status": name_quality.get("status", ""),
        "name_quality_errors": "；".join(name_quality.get("errors", [])),
        "name_quality_warnings": "；".join(name_quality.get("warnings", [])),
        "name_quality_path": name_quality.get("validation_path", ""),
        "name_pending_path": name_quality.get("pending_path", ""),
        "name_pending_count": str(name_quality.get("pending_count", 0)),
        "precise_highlight_status": highlight_info.get("status", ""),
        "precise_highlight_reason": highlight_info.get("reason", ""),
        "precise_highlight_pdf": highlight_info.get("highlight_pdf", ""),
        "precise_highlight_report": highlight_info.get("highlight_report", ""),
        "precise_highlight_hits_total": str(highlight_info.get("hits_total", 0)),
        "precise_highlight_skipped_ids": "、".join(highlight_info.get("skipped_ids", [])),
        "precheck_docx_path": str(precheck_path),
        "precheck_approved": "Y" if precheck_approved else "N",
        "case_meta_path": str(case_meta_path),
        "source_bucket_order": "fast>full>uncategorized>direct_bind",
    }
    scan_report = write_scan_report(
        out_dir=out_dir,
        report_name=args.scan_report,
        case=case_for_report,
        rows=rows,
        package_info=package_info,
        baseline_info=baseline_info,
    )

    sheets = pd.read_excel(xlsx_template, sheet_name=None)
    ci = sheets.get("CaseInfo", pd.DataFrame(columns=["key", "value"]))
    if not {"key", "value"}.issubset(ci.columns):
        ci = pd.DataFrame(columns=["key", "value"])
    else:
        # 兼容模板中 value 列被 pandas 推断为 float64 的情况，避免写入字符串时报错。
        ci["key"] = ci["key"].where(ci["key"].notna(), "").astype(str)
        ci["value"] = ci["value"].astype("object")

    def set_ci(k, v):
        k = str(k or "").strip()
        v = "" if v is None else str(v)
        idx = ci.index[ci["key"] == k]
        if len(idx):
            ci.loc[idx[0], "value"] = v
        else:
            ci.loc[len(ci)] = {"key": k, "value": v}

    set_ci("case_no", case_no)
    set_ci("mark_name", mark_name)
    set_ci("reg_no", reg_no)
    set_ci("respondent", respondent)
    set_ci("applicant", applicant)
    set_ci("class", class_no)
    set_ci("designated_goods_services", goods_desc)
    set_ci("use_period_start", period_start)
    set_ci("use_period_end", period_end)
    set_ci("cnipa_notice_ref", cnipa_notice_ref)
    set_ci("revoked_mark_name", revoked_mark_name)
    set_ci("revoked_goods_services", revoked_goods_services)
    set_ci("defense_goods_services", defense_goods_services)
    set_ci("mark_image", mark_image)
    set_ci("respondent_address", respondent_address)
    set_ci("agent_company", agent_company)
    set_ci("agent_address", agent_address)
    set_ci("contact_phone", contact_phone)
    set_ci("scan_rounds_total", str(len(scan_rounds)))
    set_ci("scan_rounds", json.dumps(scan_rounds, ensure_ascii=False))
    set_ci("merged_pdf_name", package_info.get("merged_pdf_name", "证据内容_重排合并.pdf"))
    set_ci("merged_pdf_pages", package_info.get("merged_total_pages", ""))
    set_ci("merged_pdf_path", package_info.get("merged_pdf", ""))
    set_ci("dedup_before_count", str(len(files)))
    set_ci("dedup_after_count", str(len(unique_files)))
    set_ci("input_pdf_count", str(input_stats.get("pdf", 0)))
    set_ci("input_image_converted_count", str(input_stats.get("images_converted", 0)))
    set_ci("input_office_converted_count", str(input_stats.get("office_converted", 0)))
    set_ci("input_office_text_extracted_count", str(input_stats.get("office_text_extracted", 0)))
    set_ci("input_office_skipped_count", str(input_stats.get("office_skipped", 0)))
    set_ci("office_converter", str(input_stats.get("office_converter", "")))
    set_ci("input_fast_profile_count", str(input_stats.get("fast_input_count", 0)))
    set_ci("input_full_profile_count", str(input_stats.get("full_input_count", 0)))
    set_ci("direct_bind_item_count", str(input_stats.get("direct_bind_items", 0)))
    set_ci("direct_bind_file_count", str(input_stats.get("direct_bind_files", 0)))
    set_ci("direct_bind_image_converted_count", str(input_stats.get("direct_bind_images_converted", 0)))
    set_ci("direct_bind_office_converted_count", str(input_stats.get("direct_bind_office_converted", 0)))
    set_ci("direct_bind_office_skipped_count", str(input_stats.get("direct_bind_office_skipped", 0)))
    set_ci("override_rules_path", str(override_path))
    set_ci("override_applied_count", override_result.get("applied_count", "0"))
    set_ci("override_rules_json", override_result.get("override_rules_json", ""))
    set_ci("overrides_applied_path", override_result.get("overrides_applied_path", ""))
    set_ci("page_map_path", package_info.get("page_map_path", ""))
    set_ci("page_map_hash", package_info.get("page_map_hash", ""))
    set_ci("page_map_validation_status", page_map_validation.get("status", ""))
    set_ci("page_map_validation_errors", "；".join(page_map_validation.get("errors", [])))
    set_ci("page_map_validation_path", str(page_map_validation_path))
    set_ci("time_quality_status", time_quality.get("status", ""))
    set_ci("time_quality_errors", "；".join(time_quality.get("errors", [])))
    set_ci("time_quality_warnings", "；".join(time_quality.get("warnings", [])))
    set_ci("time_quality_path", str(time_quality_path))
    set_ci("evidence_type_consistency_status", type_consistency.get("status", ""))
    set_ci("evidence_type_consistency_errors", "；".join(type_consistency.get("errors", [])))
    set_ci("evidence_type_consistency_warnings", "；".join(type_consistency.get("warnings", [])))
    set_ci("evidence_type_consistency_path", str(type_consistency_path))
    set_ci("low_confidence_fields_path", low_confidence_info.get("path", ""))
    set_ci("low_confidence_field_count", str(low_confidence_info.get("count", 0)))
    set_ci("name_quality_status", name_quality.get("status", ""))
    set_ci("name_quality_errors", "；".join(name_quality.get("errors", [])))
    set_ci("name_quality_warnings", "；".join(name_quality.get("warnings", [])))
    set_ci("name_quality_path", name_quality.get("validation_path", ""))
    set_ci("name_pending_path", name_quality.get("pending_path", ""))
    set_ci("name_pending_count", str(name_quality.get("pending_count", 0)))
    set_ci("precise_highlight_status", highlight_info.get("status", ""))
    set_ci("precise_highlight_reason", highlight_info.get("reason", ""))
    set_ci("precise_highlight_pdf", highlight_info.get("highlight_pdf", ""))
    set_ci("precise_highlight_report", highlight_info.get("highlight_report", ""))
    set_ci("precise_highlight_hits_total", str(highlight_info.get("hits_total", 0)))
    set_ci("precise_highlight_skipped_ids", "、".join(highlight_info.get("skipped_ids", [])))
    set_ci("precheck_docx_path", str(precheck_path))
    set_ci("precheck_approved", "Y" if precheck_approved else "N")
    set_ci("case_meta_path", str(case_meta_path))
    set_ci("source_bucket_order", "fast>full>uncategorized>direct_bind")
    set_ci("run_id", run_id)
    set_ci("case_info_version", validation_result.get("case_info_version", CASE_INFO_VERSION))
    set_ci("validation_rule_version", validation_result.get("validation_rule_version", VALIDATION_RULE_VERSION))
    set_ci("source_of_truth", json.dumps(validation_result.get("source_of_truth", {}), ensure_ascii=False))
    set_ci("validated_at", validation_result.get("validated_at", ""))
    set_ci("validated_by", validation_result.get("validated_by", VALIDATED_BY))
    set_ci("caseinfo_validation_status", validation_result.get("status", ""))
    set_ci("caseinfo_validation_errors", "；".join(validation_result.get("errors", [])))
    set_ci("caseinfo_validation_warnings", "；".join(validation_result.get("warnings", [])))
    set_ci("caseinfo_validation_path", str(validation_json_path))

    sheets["CaseInfo"] = ci
    sheets["DefenseEvidence"] = pd.DataFrame(rows)

    opp_cols = [
        "No", "Exhibit ID", "Exhibit Name", "Type", "Source", "Formation Date", "Carrier",
        "Goods/Services", "Mark Shown (Y/N)", "Subject Match (Y/N)", "Commercial Loop (Y/N)",
        "Original/Verifiable (Y/N)", "Notes"
    ]
    sheets["OpponentEvidence"] = pd.DataFrame(columns=opp_cols)

    with pd.ExcelWriter(xlsx_out, engine="openpyxl") as w:
        for sn, sdf in sheets.items():
            sdf.to_excel(w, sheet_name=sn, index=False)

    result = {
        "run_id": run_id,
        "case_no": case_no,
        "mark_name": mark_name,
        "reg_no": reg_no,
        "class": class_no,
        "respondent": respondent,
        "applicant": applicant,
        "period_start": period_start,
        "period_end": period_end,
        "revoked_mark_name": revoked_mark_name,
        "revoked_goods_services": revoked_goods_services,
        "defense_goods_services": defense_goods_services,
        "mark_image": mark_image,
        "respondent_address": respondent_address,
        "agent_company": agent_company,
        "agent_address": agent_address,
        "contact_phone": contact_phone,
        "scan_rounds": scan_rounds,
        "evidence_count": len(rows),
        "organized_dir": package_info.get("organized_dir", ""),
        "manifest_xlsx": package_info.get("manifest_xlsx", ""),
        "merged_pdf": package_info.get("merged_pdf", ""),
        "merged_pdf_pages": package_info.get("merged_total_pages", ""),
        "dedup_before_count": len(files),
        "dedup_after_count": len(unique_files),
        "input_pdf_count": input_stats.get("pdf", 0),
        "input_image_converted_count": input_stats.get("images_converted", 0),
        "input_office_converted_count": input_stats.get("office_converted", 0),
        "input_office_text_extracted_count": input_stats.get("office_text_extracted", 0),
        "input_office_skipped_count": input_stats.get("office_skipped", 0),
        "office_converter": input_stats.get("office_converter", ""),
        "input_fast_profile_count": input_stats.get("fast_input_count", 0),
        "input_full_profile_count": input_stats.get("full_input_count", 0),
        "direct_bind_item_count": input_stats.get("direct_bind_items", 0),
        "direct_bind_file_count": input_stats.get("direct_bind_files", 0),
        "direct_bind_image_converted_count": input_stats.get("direct_bind_images_converted", 0),
        "direct_bind_office_converted_count": input_stats.get("direct_bind_office_converted", 0),
        "direct_bind_office_skipped_count": input_stats.get("direct_bind_office_skipped", 0),
        "override_rules": override_tpl,
        "override_applied_count": override_result.get("applied_count", "0"),
        "override_rules_json": override_result.get("override_rules_json", ""),
        "overrides_applied_path": override_result.get("overrides_applied_path", ""),
        "page_map_path": package_info.get("page_map_path", ""),
        "page_map_hash": package_info.get("page_map_hash", ""),
        "page_map_validation_status": page_map_validation.get("status", ""),
        "page_map_validation_errors": page_map_validation.get("errors", []),
        "page_map_validation_path": str(page_map_validation_path),
        "time_quality_status": time_quality.get("status", ""),
        "time_quality_errors": time_quality.get("errors", []),
        "time_quality_warnings": time_quality.get("warnings", []),
        "time_quality_path": str(time_quality_path),
        "evidence_type_consistency_status": type_consistency.get("status", ""),
        "evidence_type_consistency_errors": type_consistency.get("errors", []),
        "evidence_type_consistency_warnings": type_consistency.get("warnings", []),
        "evidence_type_consistency_path": str(type_consistency_path),
        "low_confidence_fields_path": low_confidence_info.get("path", ""),
        "low_confidence_field_count": low_confidence_info.get("count", 0),
        "precheck_docx_path": str(precheck_path),
        "precheck_approved": precheck_approved,
        "case_meta_path": str(case_meta_path),
        "case_meta": case_meta_payload,
        "source_bucket_order": "fast>full>uncategorized>direct_bind",
        "precise_highlight_status": highlight_info.get("status", ""),
        "precise_highlight_reason": highlight_info.get("reason", ""),
        "precise_highlight_pdf": highlight_info.get("highlight_pdf", ""),
        "precise_highlight_report": highlight_info.get("highlight_report", ""),
        "precise_highlight_hits_total": highlight_info.get("hits_total", 0),
        "precise_highlight_skipped_ids": highlight_info.get("skipped_ids", []),
        "case_info_version": validation_result.get("case_info_version", CASE_INFO_VERSION),
        "validation_rule_version": validation_result.get("validation_rule_version", VALIDATION_RULE_VERSION),
        "caseinfo_validation_status": validation_result.get("status", ""),
        "caseinfo_validation_errors": validation_result.get("errors", []),
        "caseinfo_validation_warnings": validation_result.get("warnings", []),
        "source_of_truth": validation_result.get("source_of_truth", {}),
        "caseinfo_validation_path": str(validation_json_path),
        "scan_report": scan_report,
        "baseline_metrics": baseline_info.get("metrics_json", ""),
        "baseline_history": baseline_info.get("history_csv", ""),
        "sample_hash": baseline_info.get("sample_hash", ""),
    }
    audit({
        "type": "run_end",
        "step": "build_casebook",
        "ok": True,
        "evidence_count": len(rows),
    })
    return result


def main():
    ap = argparse.ArgumentParser(description="Auto recognize case info from evidence and generate outputs")
    ap.add_argument("--evidence-dir", default="", help="单目录模式：证据目录（兼容旧参数）")
    ap.add_argument("--evidence-dir-fast", default="", help="fast目录：无需深度时间锚点扫描的证据")
    ap.add_argument("--evidence-dir-full", default="", help="full目录：需要深度时间节点/要件分析的证据")
    ap.add_argument("--xlsx-template", required=True, help="Template casebook xlsx")
    ap.add_argument("--xlsx-out", required=True, help="Output recognized casebook xlsx")
    ap.add_argument("--ocr-cache", default="ocr_cache", help="OCR cache directory")
    ap.add_argument("--generator", required=True, help="Path to generate_suite_v2.py")
    ap.add_argument("--notice-text-hint", default="", help="Optional OCR text hint file for notice parsing")
    ap.add_argument("--out-dir", required=True, help="Output directory")
    ap.add_argument("--defense-out", default="答辩理由_自动识别.docx")
    ap.add_argument("--cross-out", default="证据目录_自动识别.docx")
    ap.add_argument("--risk-out", default="风险报告_自动识别.docx")

    ap.add_argument("--scan-mode", choices=["full", "fast", "mix"], default="mix", help="full=全文扫描，fast=快速扫描，mix=自动分流扫描")
    ap.add_argument("--max-pages", type=int, default=0, help="普通证据扫描页数上限，0=全部页")
    ap.add_argument("--dpi", type=int, default=220, help="普通证据 OCR DPI")
    ap.add_argument("--notice-max-pages", type=int, default=0, help="通知书扫描页数上限，0=全部页")
    ap.add_argument("--notice-dpi", type=int, default=280, help="通知书 OCR DPI")
    ap.add_argument("--organize-dir", default="", help="证据重命名整理输出目录")
    ap.add_argument("--scan-report", default="扫描识别报告.md", help="扫描识别报告文件名")
    ap.add_argument("--override-rules", default="", help="人工覆盖规则xlsx路径（缺省为输出目录下人工覆盖规则.xlsx）")
    ap.add_argument("--direct-bind-config", default="", help="直装证据配置JSON（不参与扫描，直接装订）")
    ap.add_argument("--override-case-meta", default="", help="案件信息覆盖JSON路径（优先于自动识别，低于命令行手填）")
    ap.add_argument("--case-reg-no", default="", help="人工覆盖：注册号")
    ap.add_argument("--case-applicant", default="", help="人工覆盖：申请人")
    ap.add_argument("--case-respondent", default="", help="人工覆盖：被申请人/注册人")
    ap.add_argument("--case-class", default="", help="人工覆盖：类别")
    ap.add_argument("--case-mark-name", default="", help="人工覆盖：商标名称")
    ap.add_argument("--case-mark-image", default="", help="人工覆盖：商标图样（文字描述或路径）")
    ap.add_argument("--case-goods-services", default="", help="人工覆盖：被撤销商品/服务")
    ap.add_argument("--case-revoked-goods-services", default="", help="人工覆盖：被撤商品/服务（优先级高于 --case-goods-services）")
    ap.add_argument("--case-defense-goods-services", default="", help="人工覆盖：答辩商品/服务")
    ap.add_argument("--case-period-start", default="", help="人工覆盖：指定期间起")
    ap.add_argument("--case-period-end", default="", help="人工覆盖：指定期间止")
    ap.add_argument("--case-respondent-address", default="", help="人工覆盖：答辩人地址")
    ap.add_argument("--case-agent-company", default="", help="人工覆盖：代理公司名称")
    ap.add_argument("--case-agent-address", default="", help="人工覆盖：代理公司地址")
    ap.add_argument("--case-contact-phone", default="", help="人工覆盖：联系电话")
    ap.add_argument("--allow-noisy-mark-name", choices=["on", "off"], default="off", help="允许手工输入的疑似噪音商标名")
    ap.add_argument("--caseinfo-fail-fast", choices=["on", "off"], default="off", help="已停用门禁参数（保留兼容）：CaseInfo仅评估不拦截")
    ap.add_argument("--strict-evidence-name", choices=["on", "off"], default="off", help="已停用门禁参数（保留兼容）：证据命名仅评估不拦截")
    ap.add_argument("--min-in-period-highmed", type=int, default=8, help="期内高/中置信最小数量门槛")
    ap.add_argument("--max-low-ratio", type=float, default=0.35, help="低置信占比上限，超过即失败")
    ap.add_argument("--max-unknown-ratio", type=float, default=0.30, help="时间不明占比上限，超过即失败")
    ap.add_argument("--precise-highlight", choices=["on", "off"], default="off", help="已停用参数（保留兼容）：证据高亮统一关闭")
    ap.add_argument("--precise-highlight-skip-ids", default="证据24", help="已停用参数（保留兼容）")
    ap.add_argument("--precise-highlight-max-pages", type=int, default=6, help="已停用参数（保留兼容）")
    ap.add_argument("--precise-highlight-max-hits", type=int, default=24, help="已停用参数（保留兼容）")
    ap.add_argument("--precheck-required", choices=["on", "off"], default="on", help="是否启用生成前核验单强制勾选门禁")
    ap.add_argument("--output-formats", default="docx", help="输出格式：docx,pdf,md,txt,html（逗号分隔）")

    args = ap.parse_args()

    Path(args.out_dir).mkdir(parents=True, exist_ok=True)
    main_run_id = f"main_{dt.datetime.now().strftime('%Y%m%d%H%M%S%f')}"
    logs_dir = Path(args.out_dir) / "logs"
    set_run_context(logs_dir, main_run_id)
    global LOGGER
    LOGGER = setup_logger(__name__, log_dir=logs_dir)
    rc = 1
    audit({
        "type": "run_start",
        "step": "main",
        "ok": True,
        "out_dir": str(Path(args.out_dir).resolve()),
    })
    try:
        info = build_casebook(args)
        precheck_doc = Path(str(info.get("precheck_docx_path", "") or ""))
        precheck_ok = bool(info.get("precheck_approved", False))
        if args.precheck_required == "on":
            # 两阶段输出：先生成证据目录供人工确认，再生成正式文书。
            cmd_cross = [
                sys.executable,
                args.generator,
                "--xlsx", args.xlsx_out,
                "--out-dir", args.out_dir,
                "--defense-out", args.defense_out,
                "--cross-out", args.cross_out,
                "--risk-out", args.risk_out,
                "--only-cross", "on",
            ]
            code, out, err = run(cmd_cross)
            if out:
                print(out)
            if err:
                print(err)
            if code != 0:
                rc = code
                sys.exit(code)

            if str(info.get("evidence_type_consistency_status", "PASS")) != "PASS":
                print("已生成《证据目录》，但检测到证据类型与内容不一致风险，请先人工核验后再输出其余文书。")
                print(f"一致性报告：{info.get('evidence_type_consistency_path', '')}")
                summary_path = Path(args.out_dir) / "扫描识别摘要.json"
                summary_path.write_text(json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8")
                rc = 5
                sys.exit(5)

            if not precheck_doc.exists() or not precheck_ok:
                print("已先生成《证据目录》，请核验并勾选后再输出其余文书。")
                print(f"请先在文件中勾选“☑ 信息无误”“☑ 同意输出”：{precheck_doc}")
                summary_path = Path(args.out_dir) / "扫描识别摘要.json"
                summary_path.write_text(json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8")
                rc = 4
                sys.exit(4)

        cmd = [
            sys.executable,
            args.generator,
            "--xlsx", args.xlsx_out,
            "--out-dir", args.out_dir,
            "--defense-out", args.defense_out,
            "--cross-out", args.cross_out,
            "--risk-out", args.risk_out,
        ]
        code, out, err = run(cmd)
        if out:
            print(out)
        if err:
            print(err)
        if code != 0:
            rc = code
            sys.exit(code)

        docx_paths = [Path(args.out_dir) / args.cross_out]
        if Path(args.out_dir, args.defense_out).exists():
            docx_paths.append(Path(args.out_dir) / args.defense_out)
        if Path(args.out_dir, args.risk_out).exists():
            docx_paths.append(Path(args.out_dir) / args.risk_out)
        export_info = export_documents_in_formats(
            out_dir=Path(args.out_dir),
            docx_paths=docx_paths,
            formats_spec=args.output_formats,
        )
        info["output_formats"] = parse_output_formats(args.output_formats)
        info["output_exports"] = export_info.get("generated", [])
        info["output_export_warnings"] = export_info.get("warnings", [])

        summary_path = Path(args.out_dir) / "扫描识别摘要.json"
        summary_path.write_text(json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8")

        print("CASE_INFO:", info)
        print("SCAN_SUMMARY:", summary_path)
        rc = 0
    except Exception as exc:
        _audit_exception("main", exc, __file__)
        raise
    finally:
        audit({
            "type": "run_end",
            "step": "main",
            "ok": rc == 0,
            "exit_code": rc,
        })


if __name__ == "__main__":
    main()

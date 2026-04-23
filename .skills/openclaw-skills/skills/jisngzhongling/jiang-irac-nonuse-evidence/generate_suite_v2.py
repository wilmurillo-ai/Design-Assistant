#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
SJ-IRAC Non-Use Automation Suite v2 (Upgraded Time + Granular Risk)
Input: nonuse_casebook_v2.xlsx
Outputs:
  - out/defense_auto.docx
  - out/cross_exam_auto.docx
  - out/risk_report.md

Compatibility:
- Works with existing CaseInfo / DefenseEvidence / OpponentEvidence sheets.
- Optional columns enhance accuracy (do NOT require them):
  - Evidence Date Start, Evidence Date End   (or: Date Start, Date End)
  - Time Anchor Text
  - Date Confidence (HIGH/MEDIUM/LOW)
"""

import os
import re
import sys
import argparse
import json
import datetime as dt
import traceback
import uuid
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any

try:
    import pandas as pd
except ImportError:
    print("Missing dependency: pandas. Install with: pip install pandas openpyxl python-docx")
    sys.exit(1)

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Missing dependency: python-docx. Install with: pip install python-docx")
    sys.exit(1)

from utils.compliance import build_reason_chain, load_rules_and_write_profile, write_low_confidence_fields
from utils.logger import audit, set_run_context, setup_logger

LOGGER = setup_logger(__name__, log_dir=(Path(__file__).resolve().parent / "logs"))


def _audit_swallowed_exception(step: str, exc: Exception, file_hint: str = "") -> None:
    LOGGER.exception("%s 发生异常", step)
    audit({
        "type": "exception",
        "step": step,
        "file": file_hint or __file__,
        "error": str(exc),
        "traceback": traceback.format_exc(),
        "ok": False,
        "reason_code": "recoverable_exception",
    })


# -----------------------------
# Paths
# -----------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE_DIR, "out")

CLOSINGS = [
    "综上，该证据在真实性、关联性、时间效力及完整性方面存在明显缺陷，难以证明指定期间内真实使用。",
    "综上，该组证据未能形成可核验的商业使用闭环，不能满足撤三案件对商标使用证据的证明要求。",
    "综上，相关材料不足以证明商标使用行为在指定期间内真实发生，依法不应予以采信。"
]

RUNTIME_RULES: Dict[str, Any] = {
    "time_rules": {
        "anchor_min_profile": {
            "min_in_period_highmed": 8,
            "max_unknown_time": 8,
            "min_time_score": 45.0,
            "min_mapping_score": 85.0,
            "min_loop_score": 45.0,
            "min_verif_score": 85.0,
        },
        "time_score_weights": {
            "base": 40.0,
            "in_period_highmed": 45.0,
            "in_period_low": 15.0,
            "unknown_penalty": 35.0,
            "contradiction_penalty": 20.0,
            "out_period_penalty": 10.0,
        },
    },
    "score_rules": {
        "mapping_score_weights": {
            "base": 30.0,
            "per_target": 10.0,
            "goods_ratio_bonus": 20.0,
        },
        "loop_score_weights": {
            "base": 45.0,
            "loop_ratio_bonus": 30.0,
            "mark_ratio_bonus": 15.0,
            "subject_ratio_bonus": 10.0,
            "weak_loop_proxy_min": 6,
            "weak_loop_penalty": 10.0,
        },
        "verifiability_score_weights": {
            "base": 25.0,
            "verifiable_ratio_bonus": 70.0,
            "contradiction_penalty": 15.0,
        },
        "checklist_thresholds": {
            "alt_loop_support_min": 6,
        },
    },
    "risk_rules": {
        "risk_level_thresholds": {
            "avg_A_min": 85.0,
            "min_dim_A_min": 75.0,
            "avg_B_min": 68.0,
            "avg_C_min": 55.0,
        },
        "anchor_degrade": {
            "both_g5_g6_if_anchor_ok": "C",
            "both_g5_g6_if_anchor_not_ok": "D",
            "one_of_g5_g6_if_anchor_ok": "B",
            "one_of_g5_g6_if_anchor_not_ok": "C",
        },
        "hard_fail_level": "E",
        "submission_level_groups": {
            "pass": ["A", "B"],
            "strengthen": ["C", "D"],
        },
        "decision_text": {
            "A": "可直接提交",
            "B": "可提交（建议补强）",
            "C": "补强后提交",
            "D": "谨慎提交（高风险）",
            "E": "暂缓提交",
        },
    },
}


def _runtime_get(rule_name: str, key: str, default: Any) -> Any:
    block = RUNTIME_RULES.get(rule_name, {}) if isinstance(RUNTIME_RULES, dict) else {}
    if not isinstance(block, dict):
        return default
    val = block.get(key, default)
    return val if val is not None else default


def _load_runtime_rules(base_dir: Path, out_dir: Path, run_id: str) -> Dict[str, Any]:
    profile = load_rules_and_write_profile(
        base_dir=base_dir,
        out_dir=out_dir,
        run_id=run_id,
        logger=LOGGER,
        audit=audit,
    )
    rules = profile.get("rules", {}) if isinstance(profile, dict) else {}
    if isinstance(rules, dict):
        for k in ("time_rules", "score_rules", "risk_rules"):
            rv = rules.get(k, {})
            if isinstance(rv, dict):
                RUNTIME_RULES[k] = rv
    return profile


# -----------------------------
# Utilities: robust date parsing
# -----------------------------
_DATE_RE_1 = re.compile(r"(?P<y>20\d{2})[./-](?P<m>0?[1-9]|1[0-2])[./-](?P<d>0?[1-9]|[12]\d|3[01])")
_DATE_RE_2 = re.compile(r"(?P<y>20\d{2})年(?P<m>0?[1-9]|1[0-2])月(?P<d>0?[1-9]|[12]\d|3[01])日?")
_RANGE_SPLIT = re.compile(r"(至|到|~|—|-|－|–)")

def _to_date_obj(x) -> Optional[dt.date]:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return None
    if isinstance(x, dt.date) and not isinstance(x, dt.datetime):
        return x
    if isinstance(x, dt.datetime):
        return x.date()
    # pandas Timestamp
    if hasattr(x, "to_pydatetime"):
        try:
            return x.to_pydatetime().date()
        except Exception as exc:
            _audit_swallowed_exception("_to_date_obj.to_pydatetime", exc, "timestamp")
    s = str(x).strip()
    if not s:
        return None
    # try common formats
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return dt.datetime.strptime(s, fmt).date()
        except Exception as exc:
            _audit_swallowed_exception("_to_date_obj.strptime", exc, s)
    # try regex
    m = _DATE_RE_1.search(s) or _DATE_RE_2.search(s)
    if m:
        try:
            return dt.date(int(m.group("y")), int(m.group("m")), int(m.group("d")))
        except Exception as exc:
            _audit_swallowed_exception("_to_date_obj.regex_date", exc, s)
            return None
    return None

def extract_dates_from_text(text: str) -> List[dt.date]:
    """Extract multiple dates from free text; used for 'Time Anchor' inference."""
    if not text:
        return []
    dates = []
    for m in _DATE_RE_1.finditer(text):
        try:
            dates.append(dt.date(int(m.group("y")), int(m.group("m")), int(m.group("d"))))
        except Exception as exc:
            _audit_swallowed_exception("extract_dates_from_text.re1", exc, text[:120])
    for m in _DATE_RE_2.finditer(text):
        try:
            dates.append(dt.date(int(m.group("y")), int(m.group("m")), int(m.group("d"))))
        except Exception as exc:
            _audit_swallowed_exception("extract_dates_from_text.re2", exc, text[:120])
    # de-dup + sort
    dates = sorted(set(dates))
    return dates

def parse_range_from_text(text: str) -> Tuple[Optional[dt.date], Optional[dt.date]]:
    """
    Parse a date range like:
      2023-01-01~2023-12-31
      2023年1月1日 至 2023年12月31日
    If only one date found, returns (d, d)
    """
    if not text:
        return (None, None)
    s = str(text).strip()
    if not s:
        return (None, None)

    # split by common range separators if present
    parts = [p.strip() for p in _RANGE_SPLIT.split(s) if p and p.strip() and p.strip() not in ("至","到","~","—","-","－","–")]
    # If separators appear, we may have 2 date fragments; otherwise, extract all dates.
    if len(parts) >= 2:
        d1 = _to_date_obj(parts[0])
        d2 = _to_date_obj(parts[1])
        if d1 and d2:
            return (min(d1, d2), max(d1, d2))
    # fallback: extract all dates in text
    ds = extract_dates_from_text(s)
    if not ds:
        d = _to_date_obj(s)
        return (d, d) if d else (None, None)
    if len(ds) == 1:
        return (ds[0], ds[0])
    return (ds[0], ds[-1])

def overlap(a_start: Optional[dt.date], a_end: Optional[dt.date], b_start: Optional[dt.date], b_end: Optional[dt.date]) -> bool:
    if not a_start or not a_end or not b_start or not b_end:
        return False
    return not (a_end < b_start or b_end < a_start)

def range_label(d1: Optional[dt.date], d2: Optional[dt.date]) -> str:
    if not d1 and not d2:
        return "时间不明"
    if d1 and d2 and d1 == d2:
        return d1.isoformat()
    if d1 and d2:
        return f"{d1.isoformat()} ~ {d2.isoformat()}"
    return (d1 or d2).isoformat()


def parse_iso_date(s: str) -> Optional[dt.date]:
    t = safe_str(s)
    if not t:
        return None
    try:
        return dt.datetime.strptime(t, "%Y-%m-%d").date()
    except Exception:
        return None


def clamp(x: float, lo: float = 0.0, hi: float = 100.0) -> float:
    return max(lo, min(hi, x))


def date_cn(d: Optional[dt.date]) -> str:
    if not d:
        return ""
    return f"{d.year}年{d.month}月{d.day}日"


def add_year_safe(d: dt.date, years: int = 1) -> dt.date:
    try:
        return dt.date(d.year + years, d.month, d.day)
    except ValueError:
        # Handle 2/29 -> 2/28 for non-leap years
        if d.month == 2 and d.day == 29:
            return dt.date(d.year + years, 2, 28)
        raise


def build_period_segments(period_start: Optional[dt.date], period_end: Optional[dt.date]) -> Dict[str, str]:
    if not period_start or not period_end or period_start > period_end:
        return {"p1": "第一阶段：", "p2": "第二阶段：", "p3": "第三阶段："}
    p1_start = period_start
    p1_end = min(add_year_safe(p1_start, 1) - dt.timedelta(days=1), period_end)
    p2_start = p1_end + dt.timedelta(days=1)
    p2_end = min(add_year_safe(p2_start, 1) - dt.timedelta(days=1), period_end) if p2_start <= period_end else period_end
    p3_start = p2_end + dt.timedelta(days=1)
    p3_end = period_end

    def fmt(label: str, s: dt.date, e: dt.date) -> str:
        if s > period_end:
            return f"{label}："
        return f"{label}：{s.isoformat()} 至 {e.isoformat()}"

    return {
        "p1": fmt("第一阶段", p1_start, p1_end),
        "p2": fmt("第二阶段", p2_start, p2_end) if p2_start <= period_end else "第二阶段：",
        "p3": fmt("第三阶段", p3_start, p3_end) if p3_start <= period_end else "第三阶段：",
    }


def first_existing(paths: List[str]) -> Optional[str]:
    for p in paths:
        if os.path.isfile(p):
            return p
    return None


def resolve_runtime_paths(
    excel_override: str = "",
    defense_tpl_override: str = "",
    cross_tpl_override: str = "",
    out_dir_override: str = "",
) -> Dict[str, str]:
    excel = os.path.abspath(excel_override) if safe_str(excel_override) else first_existing(
        [
            os.path.join(BASE_DIR, "nonuse_casebook_v2.xlsx"),
            os.path.join(BASE_DIR, "assets", "nonuse_casebook_v2.xlsx"),
        ]
    )
    defense_tpl = os.path.abspath(defense_tpl_override) if safe_str(defense_tpl_override) else first_existing(
        [
            os.path.join(BASE_DIR, "templates", "defense_template_v2.docx"),
            os.path.join(BASE_DIR, "defense_template_v2.docx"),
            os.path.join(BASE_DIR, "defense_template_v2_REPLACEMENT.docx"),
        ]
    )
    cross_tpl = os.path.abspath(cross_tpl_override) if safe_str(cross_tpl_override) else first_existing(
        [
            os.path.join(BASE_DIR, "templates", "cross_exam_template_v2.docx"),
            os.path.join(BASE_DIR, "cross_exam_template_v2.docx"),
            os.path.join(BASE_DIR, "cross_exam_template_v2_REPLACEMENT.docx"),
        ]
    )
    out_dir = os.path.abspath(out_dir_override) if safe_str(out_dir_override) else OUT_DIR
    missing = []
    if not excel or not os.path.isfile(excel):
        missing.append("nonuse_casebook_v2.xlsx (or assets/nonuse_casebook_v2.xlsx)")
    if not defense_tpl or not os.path.isfile(defense_tpl):
        missing.append("defense template (templates/defense_template_v2.docx or defense_template_v2_REPLACEMENT.docx)")
    if not cross_tpl or not os.path.isfile(cross_tpl):
        missing.append("cross-exam template (templates/cross_exam_template_v2.docx or cross_exam_template_v2_REPLACEMENT.docx)")
    if missing:
        print("Runtime files missing:")
        for m in missing:
            print(" -", m)
        sys.exit(2)
    return {
        "excel": excel,
        "defense_tpl": defense_tpl,
        "cross_tpl": cross_tpl,
        "out_dir": out_dir,
    }


# -----------------------------
# Excel reading
# -----------------------------
def read_caseinfo(xlsx_path: str) -> dict:
    ci = pd.read_excel(xlsx_path, sheet_name="CaseInfo")
    ci = ci.dropna(subset=["key"])
    return {str(k).strip(): ("" if pd.isna(v) else str(v)) for k, v in zip(ci["key"], ci["value"])}

def safe_str(x) -> str:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    return str(x).strip()


PUBLIC_NOTE_EXCLUDE_KEYWORDS = [
    "关键语句",
    "时间待核验",
    "组证关联",
    "碎片合并",
    "映射依据",
    "时间识别说明",
    "表现形式识别",
    "候选通道",
    "时间来源类型",
    "时间来源通道",
    "未提取到稳定关键词句",
]


def sanitize_public_pending_fact(text: str, fallback: str = "") -> str:
    src = safe_str(text).replace("\n", "；")
    parts = [p.strip("；;，,。 ") for p in re.split(r"[；;\n]+", src) if safe_str(p)]
    kept: List[str] = []
    for p in parts:
        if any(k in p for k in PUBLIC_NOTE_EXCLUDE_KEYWORDS):
            continue
        if p.startswith("⚠跨期矛盾"):
            continue
        if "该证据属于" in p and "覆盖计算" in p:
            continue
        kept.append(p)
    core = kept[0] if kept else safe_str(fallback).strip("；;，,。 ")
    if not core:
        core = "证明相关使用事实。"
    if len(core) > 72:
        core = core[:71].rstrip("；;，,。 ") + "…"
    if not re.search(r"[。！？]$", core):
        core += "。"
    return core


def load_page_map(ci: Dict[str, str], excel_path: str) -> Dict[str, Any]:
    page_map_path = safe_str(ci.get("page_map_path", ""))
    if not page_map_path:
        return {}
    if not os.path.isabs(page_map_path):
        page_map_path = os.path.join(os.path.dirname(excel_path), page_map_path)
    if not os.path.isfile(page_map_path):
        return {}
    try:
        with open(page_map_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        if not isinstance(payload, dict):
            return {}
        return payload
    except Exception as exc:
        _audit_swallowed_exception("load_page_map", exc, page_map_path)
        return {}


def apply_page_map_to_df(df_def: pd.DataFrame, page_map_payload: Dict[str, Any]) -> pd.DataFrame:
    if df_def is None or df_def.empty:
        return df_def
    evidences = page_map_payload.get("evidences", {})
    if not isinstance(evidences, dict) or not evidences:
        return df_def

    df = df_def.copy()
    if "Evidence ID" not in df.columns:
        return df
    if "Page Range" not in df.columns:
        df["Page Range"] = ""
    else:
        # 兼容 pandas 新版本严格 dtype：后续会写入 "1-3" 这类字符串
        df["Page Range"] = df["Page Range"].astype("object")

    for i, r in df.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        row = evidences.get(eid, {})
        if not isinstance(row, dict):
            continue
        page_range = safe_str(row.get("page_range", ""))
        if not page_range:
            s = row.get("start_page", 0)
            e = row.get("end_page", 0)
            try:
                s = int(s)
                e = int(e)
                if s > 0 and e > 0:
                    page_range = str(s) if s == e else f"{s}-{e}"
            except Exception as exc:
                _audit_swallowed_exception("apply_page_map_to_df.page_range_cast", exc, eid)
                page_range = ""
        if page_range:
            df.at[i, "Page Range"] = page_range
    return df

def normalize_yesno(x: Any) -> str:
    s = safe_str(x).upper()
    if s in ("Y", "YES", "1", "是", "TRUE"):
        return "Y"
    if s in ("N", "NO", "0", "否", "FALSE"):
        return "N"
    return ""


def yn_to_cn(v: str) -> str:
    return {"Y": "是", "N": "否"}.get(v, "不明")


def conf_to_cn(v: str) -> str:
    return {"HIGH": "高", "MEDIUM": "中", "LOW": "低"}.get(safe_str(v).upper(), "不明")


TARGET_CN = {
    "T1": "使用主体",
    "T2": "商标标识",
    "T3": "商品/服务",
    "T4": "使用时间",
    "T5": "使用场景",
    "T6": "交易闭环",
}


def target_codes_from_text(s: str) -> List[str]:
    src = safe_str(s)
    return [c for c in ["T1", "T2", "T3", "T4", "T5", "T6"] if c in src]


def targets_to_cn(s: str) -> str:
    codes = target_codes_from_text(s)
    if not codes:
        return safe_str(s)
    names = [TARGET_CN.get(c, c) for c in codes]
    return "、".join(names)


def period_label_cn(p: str) -> str:
    return {"P1": "第一阶段", "P2": "第二阶段", "P3": "第三阶段"}.get(p, p)


def stage_label_cn_by_index(i: int) -> str:
    labels = {
        1: "第一阶段",
        2: "第二阶段",
        3: "第三阶段",
        4: "第四阶段",
        5: "第五阶段",
    }
    return labels.get(i, f"第{i}阶段")


def evidence_name_display(name: str) -> str:
    """
    提交版证据名称展示：
    - 去掉自动编号前缀 E001_
    - 去掉 .pdf 后缀
    - 去掉末尾批次号 _1
    - 下划线改为更可读的分隔符
    """
    x = safe_str(name)
    if not x:
        return ""
    x = re.sub(r"^E\d{3}_", "", x)
    x = re.sub(r"\.pdf$", "", x, flags=re.IGNORECASE)
    x = re.sub(r"_(\d+)$", "", x)
    x = x.replace("_", "·")
    raw_x = x
    # 容错补全：OCR/截断常见“有限公”应补全为“有限公司”
    x = re.sub(r"有限公($|·)", r"有限公司\1", x)
    x = re.sub(r"有限责任公($|·)", r"有限责任公司\1", x)
    # 审查友好标签归一化
    x = x.replace("商品评价", "商品页评价截图")
    x = x.replace("商品详情", "商品页展示截图")
    x = x.replace("订单", "订单页面截图")
    # 交易类标签细分：转账/收款/订单
    has_transfer = any(k in raw_x for k in ["转账", "汇款", "流水"])
    has_receipt = any(k in raw_x for k in ["收款", "到账"])
    has_order = any(k in raw_x for k in ["订单", "下单"])
    if has_transfer:
        x = x.replace("交易页面截图", "转账页面截图")
        x = x.replace("商品页展示截图", "转账页面截图")
    elif has_receipt:
        x = x.replace("交易页面截图", "收款记录截图")
        x = x.replace("商品页展示截图", "收款记录截图")
    elif has_order:
        x = x.replace("交易页面截图", "订单页面截图")
    if "交易·商品页展示截图" in x:
        x = x.replace("交易·商品页展示截图", "交易·转账页面截图")
    # 提交版展示不显示“时间不明”字样，避免不利表述
    x = re.sub(r"^时间不明·", "", x)
    x = x.replace("·时间不明·", "·")
    x = re.sub(r"·+", "·", x)
    x = re.sub(r"\s+", " ", x).strip().strip("·")
    return x


def evidence_id_display(eid: str, fallback_seq: int = 0) -> str:
    s = safe_str(eid)
    m = re.search(r"(\d+)", s)
    if m:
        try:
            n = int(m.group(1))
            return f"证据{n}"
        except Exception as exc:
            _audit_swallowed_exception("evidence_id_display.parse_int", exc, s)
    if fallback_seq > 0:
        return f"证据{fallback_seq}"
    return s if s else "证据"


def parse_page_range(s: str) -> Optional[Tuple[int, int]]:
    x = safe_str(s)
    if not x:
        return None
    m = re.match(r"^\s*(\d+)\s*-\s*(\d+)\s*$", x)
    if m:
        a, b = int(m.group(1)), int(m.group(2))
        if a > b:
            a, b = b, a
        return a, b
    m2 = re.match(r"^\s*(\d+)\s*$", x)
    if m2:
        n = int(m2.group(1))
        return n, n
    return None


def merge_page_intervals(intervals: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    if not intervals:
        return []
    arr = sorted([(min(a, b), max(a, b)) for a, b in intervals], key=lambda x: (x[0], x[1]))
    out: List[Tuple[int, int]] = []
    cur_s, cur_e = arr[0]
    for s, e in arr[1:]:
        if s <= cur_e + 1:
            cur_e = max(cur_e, e)
        else:
            out.append((cur_s, cur_e))
            cur_s, cur_e = s, e
    out.append((cur_s, cur_e))
    return out


def _page_ref_from_intervals(intervals: List[Tuple[int, int]]) -> str:
    if not intervals:
        return "页码未标注"
    start = min(a for a, _ in intervals)
    end = max(b for _, b in intervals)
    return f"第{start}页" if start == end else f"第{start}-{end}页"


def page_ref_text_from_rows(rows: List[Dict[str, Any]]) -> str:
    ivs: List[Tuple[int, int]] = []
    for r in rows:
        iv = parse_page_range(safe_str(r.get("page", "")))
        if iv:
            ivs.append(iv)
    merged = merge_page_intervals(ivs)
    return _page_ref_from_intervals(merged)


def apply_heiti_font(doc: Document):
    def _set_run_font(run):
        if run is None:
            return
        run.font.name = "黑体"
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        except Exception as exc:
            _audit_swallowed_exception("apply_heiti_font.run_font", exc, "run_font")

    try:
        st = doc.styles["Normal"]
        st.font.name = "黑体"
        try:
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        except Exception as exc:
            _audit_swallowed_exception("apply_heiti_font.style_font", exc, "Normal")
    except Exception as exc:
        _audit_swallowed_exception("apply_heiti_font.style", exc, "Normal")

    for p in doc.paragraphs:
        for r in p.runs:
            _set_run_font(r)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_run_font(r)


def replace_cross_directory_note(doc: Document, ci: Dict[str, str]):
    merged_name = safe_str(ci.get("merged_pdf_name", "")) or "证据内容_重排合并.pdf"
    merged_pages = safe_str(ci.get("merged_pdf_pages", "")) or "0"
    note = f"说明：本目录以“指定期间证据覆盖”为核心组织证据。证据内容已合并为《{merged_name}》（共{merged_pages}页），页码以该合并文件为准。"

    for p in doc.paragraphs:
        txt = safe_str(p.text)
        if "说明：本目录以" in txt and "指定期间证据覆盖" in txt:
            p.text = note
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = safe_str(p.text)
                    if "说明：本目录以" in txt and "指定期间证据覆盖" in txt:
                        p.text = note


# -----------------------------
# DOCX replace
# -----------------------------
def replace_all(doc: Document, mapping: dict):
    # paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    # tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


def fill_table(table, headers: List[str], rows: List[List[str]]):
    """
    生成表格并将首行设置为 Word 可识别的“重复标题行”。
    注意：python-docx 默认无 repeat API，这里通过底层 w:tblHeader 标记实现。
    """
    def _set_repeat_header_row(tbl):
        if not tbl.rows:
            return
        header_row = tbl.rows[0]
        # 兼容外部调用约定（即使 python-docx 不内置该属性）
        try:
            header_row.repeat = True
        except Exception as exc:
            _audit_swallowed_exception("fill_table.set_repeat", exc, "header_row")
        tr = header_row._tr
        tr_pr = tr.get_or_add_trPr()
        tag = qn("w:tblHeader")
        node = tr_pr.find(tag)
        if node is None:
            node = OxmlElement("w:tblHeader")
            tr_pr.append(node)
        node.set(qn("w:val"), "1")

    try:
        table.style = "Table Grid"
    except Exception as exc:
        _audit_swallowed_exception("fill_table.style", exc, "Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = safe_str(h)
    _set_repeat_header_row(table)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = safe_str(v)


def insert_table_for_marker(doc: Document, marker: str, headers: List[str], rows: List[List[str]]) -> int:
    inserted = 0

    # Marker in top-level paragraphs
    for p in list(doc.paragraphs):
        if marker in p.text:
            p.text = p.text.replace(marker, "").strip()
            t = doc.add_table(rows=1, cols=len(headers))
            fill_table(t, headers, rows)
            p._p.addnext(t._tbl)
            inserted += 1

    # Marker in table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    if marker in cp.text:
                        cp.text = cp.text.replace(marker, "").strip()
                        t = cell.add_table(rows=1, cols=len(headers))
                        fill_table(t, headers, rows)
                        inserted += 1
    return inserted


def remove_paragraphs_containing(doc: Document, patterns: List[str]) -> int:
    removed = 0
    for p in list(doc.paragraphs):
        txt = safe_str(p.text)
        if not txt:
            continue
        if any(pt in txt for pt in patterns):
            p._element.getparent().remove(p._element)
            removed += 1
    # 表格中的段落不删除结构，仅清空文本
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    txt = safe_str(cp.text)
                    if not txt:
                        continue
                    if any(pt in txt for pt in patterns):
                        cp.text = ""
                        removed += 1
    return removed


def remove_internal_submission_paragraphs(doc: Document) -> int:
    return remove_paragraphs_containing(
        doc,
        [
            "提交建议：",
            "风险等级：",
            "触发闸门：",
            "覆盖概览",
            "四维评分",
            "可核验证据：",
            "期间覆盖矩阵",
            "检查清单",
            "重点风险",
            "内部评估",
            "风险引擎输出",
            "注：如交易闭环材料不足",
        ],
    )


# -----------------------------
# Time Anchor Engine (Defense / Opponent)
# -----------------------------
def infer_time_anchor_row(r: pd.Series) -> Dict[str, Any]:
    """
    Returns:
      - anchor_start, anchor_end (date)
      - anchor_text (str)
      - confidence: HIGH/MEDIUM/LOW
      - formation_date (date if any)
      - contradiction_flag (bool)  # formation vs anchor inconsistent in a suspicious way
    """
    # Optional explicit columns
    c_start = r.get("Evidence Date Start", r.get("Date Start", None))
    c_end = r.get("Evidence Date End", r.get("Date End", None))
    c_anchor_text = safe_str(r.get("Time Anchor Text", "")) or safe_str(r.get("Content Summary", "")) or safe_str(r.get("Evidence Name", ""))

    d_start = _to_date_obj(c_start)
    d_end = _to_date_obj(c_end)
    formation = _to_date_obj(r.get("Formation Date", None))

    confidence = safe_str(r.get("Date Confidence", "")).upper()
    if confidence not in ("HIGH", "MEDIUM", "LOW"):
        confidence = ""

    # If explicit start/end exist -> HIGH confidence
    if d_start and d_end:
        anchor_start, anchor_end = (min(d_start, d_end), max(d_start, d_end))
        conf = confidence or "HIGH"
        anchor_text = safe_str(r.get("Time Anchor Text", "")) or f"EXPLICIT_RANGE({range_label(anchor_start, anchor_end)})"
    else:
        # Try parse from formation date first (MEDIUM if datetime-like), then from text (MEDIUM/LOW)
        if formation:
            anchor_start, anchor_end = (formation, formation)
            conf = confidence or "MEDIUM"
            anchor_text = safe_str(r.get("Time Anchor Text", "")) or f"FORMATION_DATE({formation.isoformat()})"
        else:
            p_start, p_end = parse_range_from_text(c_anchor_text)
            anchor_start, anchor_end = (p_start, p_end)
            if anchor_start and anchor_end:
                conf = confidence or "MEDIUM"
                anchor_text = safe_str(r.get("Time Anchor Text", "")) or f"PARSED_TEXT({range_label(anchor_start, anchor_end)})"
            else:
                conf = confidence or "LOW"
                anchor_text = safe_str(r.get("Time Anchor Text", "")) or "NO_TIME_ANCHOR"

    # Cross-period contradiction heuristic:
    # - If both formation date and anchor range exist AND they do not overlap, mark as contradiction (needs human review)
    contradiction = False
    if formation and anchor_start and anchor_end:
        # formation should usually fall within anchor range if anchor derived from same carrier
        if not (anchor_start <= formation <= anchor_end):
            # only flag if gap is large (> 30 days) to avoid trivial mismatch
            if abs((formation - anchor_start).days) > 30 and abs((formation - anchor_end).days) > 30:
                contradiction = True

    return {
        "formation_date": formation,
        "anchor_start": anchor_start,
        "anchor_end": anchor_end,
        "anchor_text": anchor_text,
        "confidence": conf,
        "contradiction": contradiction,
    }


def evidence_in_period(anchor_start, anchor_end, period_start, period_end) -> bool:
    # Overlap test (range can overlap partially)
    return overlap(anchor_start, anchor_end, period_start, period_end)


def time_anchor_allowed_row(r: pd.Series) -> bool:
    allowed = safe_str(r.get("Time Anchor Allowed", "")).upper()
    if allowed in ("Y", "N"):
        return allowed == "Y"
    if safe_str(r.get("Time Gate Scope", "Y")).upper() == "N":
        return False
    # 兜底：程序文件不计入指定期间覆盖。
    return safe_str(r.get("Type", "")) != "程序文件"


def _field_validity(row: pd.Series, key: str) -> str:
    """
    key: date / amount / party / goods / mark_presence
    统一返回 VALID / INVALID / N/A（含向后兼容兜底）。
    """
    col_map = {
        "date": "Date Validity",
        "amount": "Amount Validity",
        "party": "Party Validity",
        "goods": "Goods Validity",
        "mark_presence": "Mark Presence Validity",
    }
    raw = safe_str(row.get(col_map.get(key, ""), "")).upper()
    if raw in ("VALID", "INVALID", "N/A"):
        return raw

    if key == "date":
        ti = infer_time_anchor_row(row)
        if not time_anchor_allowed_row(row):
            return "N/A"
        if ti["anchor_start"] and ti["anchor_end"] and ti["confidence"] in ("HIGH", "MEDIUM"):
            return "VALID"
        return "INVALID"
    if key == "goods":
        return "VALID" if safe_str(row.get("Goods Match Level", "")).upper() in ("G1", "G2") else "INVALID"
    if key == "mark_presence":
        mark_ok = normalize_yesno(row.get("Mark Shown (Y/N)", "")) == "Y"
        conf = safe_str(row.get("Mark Name Confidence", "")).upper()
        return "VALID" if mark_ok and conf in ("", "HIGH", "MEDIUM", "N/A") else "INVALID"
    if key in ("amount", "party"):
        is_trade = "交易" in safe_str(row.get("Type", "")) or "合同" in safe_str(row.get("Type", ""))
        if not is_trade:
            return "N/A"
        src_col = "Trade Amount" if key == "amount" else "Trade Counterparty"
        return "VALID" if safe_str(row.get(src_col, "")) else "INVALID"
    return "INVALID"


def _targets_for_scoring(row: pd.Series) -> List[str]:
    targets = [t for t in [f"T{i}" for i in range(1, 7)] if t in safe_str(row.get("Proof Target (T1-T6)", ""))]
    out: List[str] = []
    for t in targets:
        if t == "T4" and _field_validity(row, "date") != "VALID":
            continue
        if t == "T3" and _field_validity(row, "goods") != "VALID":
            continue
        if t == "T2" and _field_validity(row, "mark_presence") != "VALID":
            continue
        if t == "T6":
            if _field_validity(row, "amount") != "VALID" or _field_validity(row, "party") != "VALID":
                continue
        out.append(t)
    return out


# -----------------------------
# Evidence Index + T sections
# -----------------------------
def build_evidence_index_block(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    lines = []
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        eid_disp = evidence_id_display(eid, seq)

        ti = infer_time_anchor_row(r)
        in_period = "不明"
        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            in_period = "是" if evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end) else "否"

        line = (
            f"- {eid_disp}｜{evidence_name_display(safe_str(r.get('Evidence Name','')))}"
            f"｜{safe_str(r.get('Type',''))}"
            f"｜{safe_str(r.get('Source',''))}"
            f"｜形成:{safe_str(r.get('Formation Date',''))}"
            f"｜时间锚点:{range_label(ti['anchor_start'], ti['anchor_end'])}({conf_to_cn(ti['confidence'])},期内:{in_period})"
            f"｜商品/服务:{safe_str(r.get('Goods/Services',''))}"
            f"｜要件:{targets_to_cn(safe_str(r.get('Proof Target (T1-T6)','')))}"
            f"｜SJ-6:{safe_str(r.get('SJ-6 (A/R/C/T/L/X)',''))}"
            f"｜页码:{safe_str(r.get('Page Range',''))}"
            f"｜待证事实:{compact_pending_fact(r, ti)}"
        )
        if ti["contradiction"]:
            line += "｜⚠疑似跨期矛盾(需人工核查)"
        lines.append(line)
        seq += 1
    return "\n".join(lines) if lines else "（未提供有效证据目录）"


def _shorten_text(s: str, max_len: int = 56) -> str:
    t = safe_str(s).replace("\n", " ")
    t = re.sub(r"\s+", " ", t).strip("；;，, ")
    if len(t) <= max_len:
        return t
    return t[: max_len - 1] + "…"


def compact_pending_fact(r: pd.Series, ti: Dict[str, Any]) -> str:
    public_pending = safe_str(r.get("Public Pending Fact", ""))
    if public_pending:
        return sanitize_public_pending_fact(public_pending, safe_str(r.get("Inferred Purpose", "")))

    note = safe_str(r.get("Risk Notes", ""))
    inferred_purpose = _shorten_text(safe_str(r.get("Inferred Purpose", "")), 26)
    _ = ti  # 保留参数兼容
    return sanitize_public_pending_fact(note, inferred_purpose)


def infer_evidence_group_label(r: pd.Series) -> str:
    g = safe_str(r.get("Evidence Group", ""))
    if g:
        return g
    unit = safe_str(r.get("Unit Type", ""))
    typ = safe_str(r.get("Type", ""))
    form = safe_str(r.get("Form Type", ""))
    if unit == "流程文件" or typ == "程序文件":
        return "流程文件组"
    if unit == "主体资质" or typ == "资质主体证明":
        return "主体资质组"
    if unit == "图片" or form in ("现场照片", "包装"):
        return "实物图组"
    if unit in ("合同", "发票", "交易履约") or typ == "交易凭证":
        return "合同发票组"
    if unit == "宣传单":
        return "宣传单组"
    return "其他佐证组"


def build_evidence_index_table_rows(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> List[List[str]]:
    rows: List[List[str]] = []
    prepared: List[Dict[str, Any]] = []
    group_order = {
        "流程文件组": 1,
        "主体资质组": 2,
        "实物图组": 3,
        "合同发票组": 4,
        "宣传单组": 5,
        "其他佐证组": 6,
    }
    role_order = {"合同": 1, "发票": 2, "支付": 3, "交易": 4, "回单": 5, "商品对印": 6, "关联": 9}
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        ti = infer_time_anchor_row(r)
        in_period = ""
        time_anchor_label = ""
        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            in_period = "是" if evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end) else "否"
            time_anchor_label = range_label(ti["anchor_start"], ti["anchor_end"])
        group_name = infer_evidence_group_label(r)
        role = safe_str(r.get("Trade Link Role", ""))
        if not role and safe_str(r.get("Type", "")) == "交易凭证":
            name_blob = " ".join([safe_str(r.get("Evidence Name", "")), safe_str(r.get("Inferred Proof Name", ""))])
            if any(k in name_blob for k in ["发票", "价税", "增值税", "开票"]):
                role = "发票"
            elif any(k in name_blob for k in ["合同", "协议", "甲方", "乙方"]):
                role = "合同"
            else:
                role = "交易"
        d = ti["anchor_start"] or ti["anchor_end"]
        if not d:
            d = parse_iso_date(safe_str(r.get("Evidence Date Start", ""))) or parse_iso_date(safe_str(r.get("Evidence Date End", "")))
        sort_date = d or dt.date(9999, 12, 31)
        try:
            raw_no = safe_str(r.get("Trade Chain Group No", ""))
            chain_no = int(float(raw_no)) if raw_no else 9999
            if chain_no <= 0:
                chain_no = 9999
        except Exception:
            chain_no = 9999
        prepared.append({
            "group": group_name,
            "group_weight": group_order.get(group_name, 99),
            "role": role or "关联",
            "role_weight": role_order.get(role or "关联", 99),
            "chain_no": chain_no,
            "sort_date": sort_date,
            "evidence_name_sort": safe_str(r.get("Evidence Name", "")),
            "row": [
                str(seq),
                evidence_id_display(eid, seq),
                evidence_name_display(safe_str(r.get("Evidence Name", ""))),
                safe_str(r.get("Type", "")),
                time_anchor_label,
                in_period,
                targets_to_cn(safe_str(r.get("Proof Target (T1-T6)", ""))),
                safe_str(r.get("Page Range", "")),
                compact_pending_fact(r, ti),
            ],
        })
        seq += 1

    if prepared:
        prepared.sort(
            key=lambda x: (
                x.get("group_weight", 99),
                x.get("role_weight", 99) if x.get("group") == "合同发票组" else 0,
                x.get("chain_no", 9999) if x.get("group") == "合同发票组" else 0,
                x.get("sort_date", dt.date(9999, 12, 31)),
                x.get("evidence_name_sort", ""),
            )
        )
        display_seq = 1
        for item in prepared:
            row = item.get("row", [])
            if row and len(row) >= 1:
                row[0] = str(display_seq)
            display_seq += 1
        group_counts: Dict[str, int] = {}
        for item in prepared:
            g = safe_str(item.get("group", "")) or "其他佐证组"
            group_counts[g] = group_counts.get(g, 0) + 1
        emitted: set = set()
        for item in prepared:
            g = safe_str(item.get("group", "")) or "其他佐证组"
            if g not in emitted:
                rows.append([
                    "",
                    "",
                    f"【{g}（{group_counts.get(g, 0)}份）】",
                    "",
                    "",
                    "",
                    "",
                    "",
                    "",
                ])
                emitted.add(g)
            rows.append(item["row"])

    if not rows:
        rows = [["-", "-", "（未提供有效证据）", "-", "-", "-", "-", "-", "-"]]
    return rows


def build_defense_summary_table_rows(df_def: pd.DataFrame) -> List[List[str]]:
    def _uniq_keep(items: List[str], limit: int = 3) -> List[str]:
        out: List[str] = []
        seen = set()
        for x in items:
            s = safe_str(x)
            if not s or s in seen:
                continue
            seen.add(s)
            out.append(s)
            if len(out) >= limit:
                break
        return out

    def _pick_representatives(arr: List[pd.Series]) -> str:
        preferred: List[str] = []
        for x in arr:
            inf = _shorten_text(safe_str(x.get("Inferred Proof Name", "")), 18)
            if inf and inf not in ("其他材料", "其他材料页面", "其他页面"):
                preferred.append(inf)
        if not preferred:
            preferred = [evidence_name_display(safe_str(x.get("Evidence Name", ""))) for x in arr]
        chosen = _uniq_keep(preferred, limit=3)
        if not chosen:
            return "见目录"
        return "、".join(chosen) + ("等" if len(arr) > len(chosen) else "")

    def _pick_group_purpose(arr: List[pd.Series], default_text: str) -> str:
        generic_set = {
            "证明相关使用事实及其与指定期间的对应关系。",
            "证明具体经营场景、商品展示或交易线索中的至少一项事实。",
            "证明相关使用事实。",
        }
        inferred: List[str] = []
        for x in arr:
            p = _shorten_text(safe_str(x.get("Inferred Purpose", "")), 28)
            if not p or p in generic_set:
                continue
            inferred.append(p)
        chosen = _uniq_keep(inferred, limit=2)
        if chosen:
            return "；".join(chosen) + "。"
        return default_text

    groups: Dict[str, List[pd.Series]] = {}
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        typ = safe_str(r.get("Type", "")) or "其他材料"
        groups.setdefault(typ, []).append(r)

    order = {"程序文件": 1, "资质主体证明": 2, "线上店铺展示": 3, "商品展示页": 4, "交易凭证": 5, "其他材料": 6}
    rows: List[List[str]] = []
    for typ, arr in sorted(groups.items(), key=lambda kv: (order.get(kv[0], 99), kv[0])):
        rep = _pick_representatives(arr)
        pages = page_ref_text_from_rows([{
            "page": safe_str(x.get("Page Range", ""))
        } for x in arr])
        if typ == "程序文件":
            purpose = _pick_group_purpose(arr, "证明程序节点与指定期间审查背景。")
        elif typ == "资质主体证明":
            purpose = _pick_group_purpose(arr, "证明主体资格、商标权属及核定商品/服务范围。")
        elif typ == "线上店铺展示":
            purpose = _pick_group_purpose(arr, "证明网店展示、标识呈现及经营页面连续性。")
        elif typ == "商品展示页":
            purpose = _pick_group_purpose(arr, "证明商品展示内容及标识在商品页中的呈现。")
        elif typ == "交易凭证":
            purpose = _pick_group_purpose(arr, "证明支付/交易场景及时间线索。")
        else:
            purpose = _pick_group_purpose(arr, "证明服务活动现场、宣传展示或经营场景相关事实。")
        rows.append([typ, f"{len(arr)}份", rep or "见目录", purpose, pages])
    return rows


def build_risk_summary_table_rows(level: str, d: Dict[str, Any], period_start: Optional[dt.date], period_end: Optional[dt.date]) -> List[List[str]]:
    decision = decision_from_level(level)
    trig = [k for k, v in d["gate_flags"].items() if v]
    trig_text = "、".join(trig) if trig else "无"
    period_txt = f"{period_start.isoformat()} ~ {period_end.isoformat()}" if (period_start and period_end) else "（未解析）"
    s = d["summary"]
    ds = d["dim_scores"]
    rows = [
        ["指定期间", period_txt],
        ["提交建议", decision],
        ["风险等级", level],
        ["触发闸门", trig_text],
        ["覆盖口径", "A线=实质使用证据；B线=程序/管理/期外参考材料"],
        ["期内覆盖（高/中置信）", str(s["in_period_highmed"])],
        ["期内覆盖（低置信）", str(s["in_period_low"])],
        ["期内覆盖（B线参考）", str(s.get("in_period_reference", 0))],
        ["时间不明", str(s["unknown_time"])],
        ["四维评分", f"时间性={ds['Time']}；对应性={ds['Mapping']}；闭环性={ds['Loop']}；可核验性={ds['Verifiability']}"],
    ]
    for i, risk in enumerate(d["top_risks"][:3], start=1):
        rows.append([f"重点风险{i}", risk])
    return rows


def get_period_segments(period_start: Optional[dt.date], period_end: Optional[dt.date]) -> List[Tuple[str, dt.date, dt.date]]:
    if not period_start or not period_end or period_start > period_end:
        return []
    p1s = period_start
    p1e = min(add_year_safe(p1s, 1) - dt.timedelta(days=1), period_end)
    p2s = p1e + dt.timedelta(days=1)
    p2e = min(add_year_safe(p2s, 1) - dt.timedelta(days=1), period_end) if p2s <= period_end else period_end
    p3s = p2e + dt.timedelta(days=1)
    p3e = period_end

    segs: List[Tuple[str, dt.date, dt.date]] = [("P1", p1s, p1e)]
    if p2s <= period_end:
        segs.append(("P2", p2s, p2e))
    if p3s <= period_end:
        segs.append(("P3", p3s, p3e))
    return segs


def segment_coverage(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> Dict[str, int]:
    cov = {"P1": 0, "P2": 0, "P3": 0}
    segs = get_period_segments(period_start, period_end)
    for _, r in df_def.iterrows():
        if not safe_str(r.get("Evidence ID", "")):
            continue
        if not time_anchor_allowed_row(r):
            continue
        if _field_validity(r, "date") != "VALID":
            continue
        ti = infer_time_anchor_row(r)
        for label, s, e in segs:
            if evidence_in_period(ti["anchor_start"], ti["anchor_end"], s, e):
                cov[label] += 1
    return cov


def meets_anchor_minimum(diag: Dict[str, Any]) -> bool:
    s = diag.get("summary", {}) if isinstance(diag, dict) else {}
    ds = diag.get("dim_scores", {}) if isinstance(diag, dict) else {}
    anchor_cfg = _runtime_get("time_rules", "anchor_min_profile", {})
    try:
        return (
            float(ds.get("Time", 0.0)) >= float(anchor_cfg.get("min_time_score", 45.0))
            and float(ds.get("Mapping", 0.0)) >= float(anchor_cfg.get("min_mapping_score", 85.0))
            and float(ds.get("Loop", 0.0)) >= float(anchor_cfg.get("min_loop_score", 45.0))
            and float(ds.get("Verifiability", 0.0)) >= float(anchor_cfg.get("min_verif_score", 85.0))
            and int(s.get("in_period_highmed", 0)) >= int(anchor_cfg.get("min_in_period_highmed", 8))
            and int(s.get("unknown_time", 9999)) <= int(anchor_cfg.get("max_unknown_time", 8))
        )
    except Exception as exc:
        _audit_swallowed_exception("meets_anchor_minimum", exc, "diag")
        return False


def build_submission_checklist_rows(
    ci: Dict[str, str],
    df_def: pd.DataFrame,
    diag: Dict[str, Any],
    level: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> List[List[str]]:
    rows: List[List[str]] = []
    seq = 1
    checklist_cfg = _runtime_get("score_rules", "checklist_thresholds", {})
    alt_loop_support_min = int(checklist_cfg.get("alt_loop_support_min", 6))
    submission_groups = _runtime_get("risk_rules", "submission_level_groups", {})
    pass_levels = set(str(x) for x in submission_groups.get("pass", ["A", "B"]))
    strengthen_levels = set(str(x) for x in submission_groups.get("strengthen", ["C", "D"]))

    def add(item: str, status: str, basis: str, action: str):
        nonlocal seq
        rows.append([str(seq), item, status, basis, action])
        seq += 1

    # 1) 基础字段完整性
    required_keys = ["reg_no", "class", "respondent", "use_period_start", "use_period_end"]
    missing = [k for k in required_keys if not safe_str(ci.get(k, ""))]
    add(
        "案件基础信息完整性（注册号/类别/主体/期间）",
        "通过" if not missing else "不通过",
        "缺失字段：" + ("无" if not missing else "、".join(missing)),
        "补齐 CaseInfo 基础字段后再提交。",
    )

    # 2) 证据总量
    ev_count = int((df_def["Evidence ID"].astype(str).str.strip() != "").sum()) if "Evidence ID" in df_def.columns else 0
    add(
        "证据条目数量检查",
        "通过" if ev_count > 0 else "不通过",
        f"有效证据条目：{ev_count}",
        "至少提供1条与指定期间相关证据。",
    )

    # 3) P1/P2/P3 覆盖
    cov = segment_coverage(df_def, period_start, period_end)
    seg_ranges = {lbl: (s, e) for lbl, s, e in get_period_segments(period_start, period_end)}
    anchor_ok = meets_anchor_minimum(diag)
    for label in ["P1", "P2", "P3"]:
        c = cov.get(label, 0)
        if label in seg_ranges:
            s, e = seg_ranges[label]
            label_txt = f"{period_label_cn(label)}（{s.isoformat()} 至 {e.isoformat()}）"
        else:
            label_txt = period_label_cn(label)
        seg_status = "通过" if c > 0 else ("补强" if anchor_ok else "不通过")
        add(
            f"{label_txt}覆盖检查",
            seg_status,
            f"{label_txt}覆盖证据数：{c}",
            "建议补充该阶段期内可核验证据以增强连续性。" if c == 0 else "可继续补强交易闭环材料。",
        )

    # 4) T1-T6 要件覆盖
    t_cov = {f"T{i}": 0 for i in range(1, 7)}
    for _, r in df_def.iterrows():
        for t in _targets_for_scoring(r):
            if t in t_cov:
                t_cov[t] += 1
    missing_t = [t for t, n in t_cov.items() if n == 0]
    missing_t_cn = [TARGET_CN.get(t, t) for t in missing_t]
    add(
        "六项要件覆盖检查",
        "通过" if not missing_t else ("不通过" if "T4" in missing_t else "补强"),
        "缺失要件：" + ("无" if not missing_t_cn else "、".join(missing_t_cn)),
        "补充缺失要件对应证据并完善映射。",
    )

    # 5) 交易闭环
    loop_yes = int((df_def.get("Commercial Loop (Y/N)", "").astype(str).str.upper() == "Y").sum()) if "Commercial Loop (Y/N)" in df_def.columns else 0
    anchor_ok = meets_anchor_minimum(diag)
    alt_loop_support = 0
    try:
        type_col = df_def.get("Type", pd.Series([""] * len(df_def))).astype(str)
        conf_col = df_def.get("Date Confidence", pd.Series([""] * len(df_def))).astype(str).str.upper()
        mark_col = df_def.get("Mark Shown (Y/N)", pd.Series([""] * len(df_def))).astype(str).str.upper()
        subj_col = df_def.get("Subject Match (Y/N)", pd.Series([""] * len(df_def))).astype(str).str.upper()
        alt_mask = (
            type_col.isin(["线上店铺展示", "商品展示页"])
            & conf_col.isin(["HIGH", "MEDIUM"])
            & (mark_col == "Y")
            & (subj_col == "Y")
        )
        alt_loop_support = int(alt_mask.sum())
    except Exception as exc:
        _audit_swallowed_exception("build_submission_checklist_rows.alt_loop_support", exc, "DefenseEvidence")
        alt_loop_support = 0

    if loop_yes > 0:
        loop_status = "通过"
        loop_basis = f"闭环证据条目：{loop_yes}"
        loop_action = "建议补充第三方可核验交易材料以增强稳定性。"
    elif alt_loop_support >= alt_loop_support_min and anchor_ok:
        loop_status = "通过"
        loop_basis = f"闭环证据条目：0；替代链条条目：{alt_loop_support}（展示+评价+时间锚点）"
        loop_action = "满足最低锚定要求，建议补1组订单/支付/物流材料进一步强化。"
    else:
        loop_status = "补强"
        loop_basis = f"闭环证据条目：{loop_yes}；替代链条条目：{alt_loop_support}"
        loop_action = "优先补齐至少1组订单/支付/物流/签收闭环证据。"

    add(
        "交易闭环证据检查（订单/支付/物流/签收）",
        loop_status,
        loop_basis,
        loop_action,
    )

    # 6) 可核验性
    ver_yes = int((df_def.get("Original/Verifiable (Y/N)", "").astype(str).str.upper() == "Y").sum()) if "Original/Verifiable (Y/N)" in df_def.columns else 0
    add(
        "可核验性检查（原始载体/第三方可验）",
        "通过" if ver_yes > 0 else "补强",
        f"可核验证据条目：{ver_yes}",
        "补充平台后台导出、税票、物流、支付流水等第三方材料。",
    )

    # 7) 风险项
    trig = [k for k, v in diag["gate_flags"].items() if v]
    hard_gates = {"G1a", "G1b", "G2", "G3", "G4"}
    gate_status = "通过"
    if any(g in hard_gates for g in trig):
        gate_status = "不通过"
    elif trig:
        gate_status = "补强"
    add(
        "风险项检查（G1-G6）",
        gate_status,
        "触发风险项：" + ("无" if not trig else "、".join(trig)),
        "按风险报告重点风险逐项补强后再提交。",
    )

    # 8) 提交建议
    add(
        "提交结论（内部评价）",
        "通过" if level in pass_levels else ("补强" if level in strengthen_levels else "不通过"),
        f"当前风险等级：{level}",
        f"{'/'.join(sorted(pass_levels))}可提交；{'/'.join(sorted(strengthen_levels))}建议补证后提交；其余等级建议暂缓提交。",
    )
    return rows


def build_gate_decision(checklist_rows: List[List[str]], level: str, diag: Dict[str, Any]) -> Dict[str, Any]:
    hard_items: List[str] = []
    strengthen_items: List[str] = []
    pass_items: List[str] = []
    for row in checklist_rows:
        if len(row) < 3:
            continue
        item = safe_str(row[1])
        status = safe_str(row[2])
        if status == "不通过":
            hard_items.append(item)
        elif status == "补强":
            strengthen_items.append(item)
        else:
            pass_items.append(item)

    anchor_ok = meets_anchor_minimum(diag)

    if hard_items:
        decision = "不通过（存在硬性缺口，需补证后重检）"
    elif strengthen_items:
        decision = "可提交（满足最低锚定，建议补强后提交）" if anchor_ok else "可提交（建议补强后提交）"
    else:
        decision = "可提交（内部评价通过）"

    return {
        "decision": decision,
        "risk_level": level,
        "anchor_ok": anchor_ok,
        "hard_count": len(hard_items),
        "strengthen_count": len(strengthen_items),
        "pass_count": len(pass_items),
        "hard_items": hard_items,
        "strengthen_items": strengthen_items,
        "triggered_gates": [k for k, v in diag.get("gate_flags", {}).items() if v],
    }


def build_checklist_summary(level: str, diag: Dict[str, Any]) -> str:
    trig = [k for k, v in diag["gate_flags"].items() if v]
    trig_txt = "无" if not trig else "、".join(trig)
    return (
        "本输出为撤三提交前检查单，不包含质证问答内容。\\n"
        f"- 当前风险等级：{level}\\n"
        f"- 触发闸门：{trig_txt}\\n"
        "- 请依据检查单逐项补齐后再对外提交。"
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="SJ-IRAC Non-Use Automation Suite v2")
    parser.add_argument("--xlsx", default="", help="Path to case workbook (.xlsx)")
    parser.add_argument("--defense-template", default="", help="Path to defense docx template")
    parser.add_argument("--cross-template", default="", help="Path to cross-exam/evidence-index docx template")
    parser.add_argument("--out-dir", default="", help="Output directory")
    parser.add_argument("--defense-out", default="defense_auto.docx", help="Defense output filename")
    parser.add_argument("--cross-out", default="cross_exam_auto.docx", help="Cross output filename")
    parser.add_argument("--risk-out", default="risk_report.docx", help="Risk report output filename")
    parser.add_argument("--only-cross", choices=["on", "off"], default="off", help="仅输出证据目录（用于预确认流程）")
    return parser.parse_args()


def build_T_sections(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> dict:
    sections = {f"T{i}": [] for i in range(1, 7)}
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        eid_disp = evidence_id_display(eid, seq)
        targets = safe_str(r.get("Proof Target (T1-T6)", ""))
        ti = infer_time_anchor_row(r)
        in_period = "期内" if (period_start and period_end and evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end)) else "期外/不明"
        base = (
            f"{eid_disp}（{evidence_name_display(safe_str(r.get('Evidence Name','')))}）\n"
            f"- 类型：{safe_str(r.get('Type',''))}\n"
            f"- 来源：{safe_str(r.get('Source',''))}\n"
            f"- 形成时间：{safe_str(r.get('Formation Date',''))}\n"
            f"- 时间锚点：{range_label(ti['anchor_start'], ti['anchor_end'])}（{conf_to_cn(ti['confidence'])}，{in_period}）\n"
            f"- 对应商品/服务：{safe_str(r.get('Goods/Services',''))}\n"
            f"- SJ-6：{safe_str(r.get('SJ-6 (A/R/C/T/L/X)',''))}\n"
        )
        for t in sections:
            if t in targets:
                sections[t].append(base)
        seq += 1
    return {k: ("\n".join(v).strip() if v else "（暂无有效证据）") for k, v in sections.items()}


# -----------------------------
# Granular Risk Engine (Gate breakdown + dimension scoring + evidence diagnostics)
# -----------------------------
def build_defense_diagnostics(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> Dict[str, Any]:
    """
    Returns:
      - gate_flags: {G1a..G6: bool}
      - gate_details: {gate: [evidence_ids]}
      - dim_scores: {Time, Mapping, Loop, Verifiability} in 0-100
      - evidence_rows: list of dicts for MD table
      - in_period_summary: counts
      - top_risks: list of strings (evidence-based)
    """
    gate_details = {g: [] for g in ["G1a","G1b","G1c","G1d","G2","G3","G4","G5","G6"]}

    # Evidence-level flags and metrics
    ev_rows = []
    in_period_highmed = 0
    in_period_low = 0
    out_period = 0
    unknown_time = 0
    contradictions = 0
    in_period_reference = 0
    program_in_period = 0
    eligible_time_rows = 0

    goods_nonempty_total = 0
    goods_nonempty_allowed = 0
    goods_matched_total = 0
    goods_matched_allowed = 0
    mark_yes_total = 0
    mark_yes_allowed = 0
    subject_yes_total = 0
    subject_yes_allowed = 0
    loop_yes_total = 0
    loop_yes_allowed = 0
    verifiable_yes_total = 0
    verifiable_yes_allowed = 0
    t_coverage_total = {f"T{i}": 0 for i in range(1, 7)}
    t_coverage_allowed = {f"T{i}": 0 for i in range(1, 7)}

    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        eid_disp = evidence_id_display(eid, seq)

        ti = infer_time_anchor_row(r)
        a_start, a_end = ti["anchor_start"], ti["anchor_end"]
        conf = ti["confidence"]
        allowed = time_anchor_allowed_row(r)
        date_valid = (_field_validity(r, "date") == "VALID")
        goods_valid = (_field_validity(r, "goods") == "VALID")
        mark_valid = (_field_validity(r, "mark_presence") == "VALID")
        amount_valid = (_field_validity(r, "amount") == "VALID")
        party_valid = (_field_validity(r, "party") == "VALID")
        trade_like = ("交易" in safe_str(r.get("Type", "")) or "合同" in safe_str(r.get("Type", "")))
        in_period = False
        if period_start and period_end and a_start and a_end:
            in_period = evidence_in_period(a_start, a_end, period_start, period_end)
        if allowed and date_valid:
            eligible_time_rows += 1
        elif in_period:
            in_period_reference += 1
            if safe_str(r.get("Type", "")) == "程序文件":
                program_in_period += 1

        if ti["contradiction"]:
            contradictions += 1
            gate_details["G1d"].append(eid_disp)

        if allowed:
            if (not date_valid) or (not a_start) or (not a_end):
                unknown_time += 1
            else:
                if in_period:
                    if conf in ("HIGH", "MEDIUM"):
                        in_period_highmed += 1
                    else:
                        in_period_low += 1
                else:
                    out_period += 1

        gs = safe_str(r.get("Goods/Services", ""))
        if gs:
            goods_nonempty_total += 1
            if allowed:
                goods_nonempty_allowed += 1
        gm = safe_str(r.get("Goods Match Level", "")).upper()
        if goods_valid and gm in ("G1", "G2"):
            goods_matched_total += 1
            if allowed:
                goods_matched_allowed += 1

        if mark_valid:
            mark_yes_total += 1
            if allowed:
                mark_yes_allowed += 1
        entity_level = safe_str(r.get("Entity Consistency", "")).upper()
        if normalize_yesno(r.get("Subject Match (Y/N)", "")) == "Y" or entity_level in ("E1", "E2"):
            subject_yes_total += 1
            if allowed:
                subject_yes_allowed += 1
        if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "Y":
            loop_yes_total += 1
            if allowed and (not trade_like or (amount_valid and party_valid)):
                loop_yes_allowed += 1
        if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "Y":
            verifiable_yes_total += 1
            if allowed:
                verifiable_yes_allowed += 1

        targets = safe_str(r.get("Proof Target (T1-T6)", ""))
        score_targets = _targets_for_scoring(r)
        for t in t_coverage_total:
            if t in targets:
                t_coverage_total[t] += 1
                if allowed and t in score_targets:
                    t_coverage_allowed[t] += 1

        ev_rows.append({
            "证据ID": eid_disp,
            "证据名称": evidence_name_display(safe_str(r.get("Evidence Name", ""))),
            "时间锚点": range_label(a_start, a_end),
            "期内": "是" if in_period else ("否" if (a_start and a_end and period_start and period_end) else "不明"),
            "置信度": conf_to_cn(conf),
            "计入覆盖": "是" if allowed else "否",
            "覆盖线": safe_str(r.get("Coverage Lane", "")) or ("A线（可计入覆盖）" if allowed else "B线（仅参考）"),
            "时间来源类型": safe_str(r.get("Time Anchor Type", "")) or "uncertain",
            "时间来源通道": safe_str(r.get("Time Anchor Source", "")) or "基础",
            "要件": targets_to_cn(targets) or "—",
            "核心要件(已过滤低置信字段)": targets_to_cn(",".join(score_targets)) or "—",
            "商标呈现": yn_to_cn(normalize_yesno(r.get("Mark Shown (Y/N)", ""))),
            "主体一致": yn_to_cn(normalize_yesno(r.get("Subject Match (Y/N)", ""))),
            "闭环": yn_to_cn(normalize_yesno(r.get("Commercial Loop (Y/N)", ""))),
            "可核验": yn_to_cn(normalize_yesno(r.get("Original/Verifiable (Y/N)", ""))),
            "商品匹配": safe_str(r.get("Goods Match Level", "")) or "G3",
            "待证事实": sanitize_public_pending_fact(
                safe_str(r.get("Public Pending Fact", "")) or safe_str(r.get("Risk Notes", "")),
                safe_str(r.get("Inferred Purpose", "")),
            ),
        })
        seq += 1

    # Gate flags (case-level)
    if not period_start or not period_end:
        gate_details["G1a"] = ["案件信息中的指定期间起止日期缺失或不可解析"]
    else:
        if in_period_highmed == 0 and in_period_low == 0:
            gate_details["G1b"] = ["无任何证据覆盖指定期间（含范围重叠）"]
        elif in_period_highmed == 0 and in_period_low > 0:
            gate_details["G1c"] = ["仅存在低置信度的期内覆盖（截图/自述类高风险）"]
        if in_period_highmed == 0 and in_period_reference > 0:
            gate_details["G1b"].append("仅B线（程序/管理/参考）材料覆盖期内，不构成实质使用覆盖。")

    # These are "core element missing" gates: trigger only if none of the evidence supports the element
    if goods_matched_allowed == 0:
        gate_details["G2"] = ["A线实质使用证据均未命中核定商品/服务（G1/G2均为0）"]
    if mark_yes_allowed == 0:
        gate_details["G3"] = ["A线实质使用证据均未能确认商标标识呈现/绑定（商标呈现均非“是”）"]
    if subject_yes_allowed == 0:
        gate_details["G4"] = ["A线实质使用证据均未能确认使用主体一致/授权链（主体一致均非“是”）"]
    loop_cfg = _runtime_get("score_rules", "loop_score_weights", {})
    time_cfg = _runtime_get("time_rules", "time_score_weights", {})
    map_cfg = _runtime_get("score_rules", "mapping_score_weights", {})
    ver_cfg = _runtime_get("score_rules", "verifiability_score_weights", {})
    weak_loop_proxy_min = int(loop_cfg.get("weak_loop_proxy_min", 6))
    weak_loop_proxy = min(mark_yes_allowed, subject_yes_allowed, t_coverage_allowed.get("T5", 0), in_period_highmed)
    # 交易闭环降级规则：
    # 若无显式闭环，但“展示+主体+时间+场景”链条足够强，则不触发 G5 否决闸门，仅在建议层提示补强。
    if loop_yes_allowed == 0 and weak_loop_proxy < weak_loop_proxy_min:
        gate_details["G5"] = ["全案未见交易闭环证据，且替代链条不足（展示+主体+时间+场景未达最低锚定）。"]
    if verifiable_yes_allowed == 0:
        gate_details["G6"] = ["全案缺乏可核验来源/原始载体（可核验均非“是”）"]

    gate_flags = {k: bool(v) for k, v in gate_details.items()}

    # Dimension scores (0-100), conservative
    total_ev = max(1, len(ev_rows))
    allowed_ev = max(1, eligible_time_rows)
    time_ev = max(1, eligible_time_rows)

    # Time: reward in-period HIGH/MEDIUM, penalize UNKNOWN and contradictions
    time_score = float(time_cfg.get("base", 40.0))
    if period_start and period_end:
        time_score += float(time_cfg.get("in_period_highmed", 45.0)) * (in_period_highmed / time_ev)
        time_score += float(time_cfg.get("in_period_low", 15.0)) * (in_period_low / time_ev)
        time_score -= float(time_cfg.get("unknown_penalty", 35.0)) * (unknown_time / time_ev)
        time_score -= float(time_cfg.get("contradiction_penalty", 20.0)) * (contradictions / time_ev)
        # Out-period is not fatal per se, but indicates noise; mild penalty
        time_score -= float(time_cfg.get("out_period_penalty", 10.0)) * (out_period / time_ev)
    else:
        time_score = 0.0
    time_score = clamp(time_score)

    # Mapping: based on T1–T6 coverage + goods mapping presence
    covered_targets = sum(1 for t, c in t_coverage_allowed.items() if c > 0)
    mapping_score = float(map_cfg.get("base", 30.0)) + float(map_cfg.get("per_target", 10.0)) * covered_targets
    mapping_score += float(map_cfg.get("goods_ratio_bonus", 20.0)) * (goods_matched_allowed / allowed_ev)
    mapping_score = clamp(mapping_score)

    # Loop: 降低机械性惩罚，允许“展示+评价+时间锚点”形成弱闭环基础分
    loop_score = (
        float(loop_cfg.get("base", 45.0))
        + float(loop_cfg.get("loop_ratio_bonus", 30.0)) * (loop_yes_allowed / allowed_ev)
        + float(loop_cfg.get("mark_ratio_bonus", 15.0)) * (mark_yes_allowed / allowed_ev)
        + float(loop_cfg.get("subject_ratio_bonus", 10.0)) * (subject_yes_allowed / allowed_ev)
    )
    if loop_yes_allowed == 0 and weak_loop_proxy < weak_loop_proxy_min:
        loop_score -= float(loop_cfg.get("weak_loop_penalty", 10.0))
    loop_score = clamp(loop_score)

    # Verifiability: verifiable_yes ratio + reduce if contradictions exist
    verif_score = (
        float(ver_cfg.get("base", 25.0))
        + float(ver_cfg.get("verifiable_ratio_bonus", 70.0)) * (verifiable_yes_allowed / allowed_ev)
        - float(ver_cfg.get("contradiction_penalty", 15.0)) * (contradictions / time_ev)
    )
    verif_score = clamp(verif_score)

    dim_scores = {
        "Time": round(time_score, 1),
        "Mapping": round(mapping_score, 1),
        "Loop": round(loop_score, 1),
        "Verifiability": round(verif_score, 1),
    }

    # Top risks (evidence-based)
    top_risks = []
    if gate_flags.get("G1b"):
        top_risks.append("G1b：无任何证据覆盖指定期间（必须补齐期内证据或重建时间锚点）。")
    if gate_flags.get("G1c"):
        top_risks.append("G1c：期内覆盖仅为低置信材料（优先补“可核验/第三方/闭环”证据以提升时间证明力）。")
    if contradictions > 0:
        top_risks.append(f"G1d：存在{contradictions}项疑似跨期矛盾（形成日与时间锚点不一致），需逐一核查原始载体。")
    if gate_flags.get("G5"):
        top_risks.append("G5：交易闭环偏弱（建议补齐订单/付款/交付/验收之一并与商标/商品绑定；该项按补强处理）。")
    if gate_flags.get("G6"):
        top_risks.append("G6：可核验性不足（补税控、平台后台导出、物流签收、支付流水、第三方存证）。")
    if in_period_highmed == 0 and program_in_period > 0:
        top_risks.append("R1：程序文件覆盖期内但不具备实质使用证明力，不能替代使用证据。")
    if not top_risks:
        # still provide bottleneck
        bottleneck = min(dim_scores, key=lambda k: dim_scores[k])
        dim_cn = {
            "Time": "时间性",
            "Mapping": "对应性",
            "Loop": "闭环性",
            "Verifiability": "可核验性",
        }.get(bottleneck, bottleneck)
        top_risks.append(f"主要瓶颈维度：{dim_cn}（{dim_scores[bottleneck]}分）。")

    return {
        "gate_flags": gate_flags,
        "gate_details": gate_details,
        "dim_scores": dim_scores,
        "evidence_rows": ev_rows,
        "summary": {
            "total": len(ev_rows),
            "time_scope_total": eligible_time_rows,
            "in_period_highmed": in_period_highmed,
            "in_period_low": in_period_low,
            "out_period": out_period,
            "unknown_time": unknown_time,
            "contradictions": contradictions,
            "in_period_reference": in_period_reference,
            "program_in_period": program_in_period,
            "goods_matched": goods_matched_allowed,
            "goods_matched_total": goods_matched_total,
            "goods_nonempty": goods_nonempty_allowed,
            "goods_nonempty_total": goods_nonempty_total,
            "mark_yes": mark_yes_allowed,
            "mark_yes_total": mark_yes_total,
            "subject_yes": subject_yes_allowed,
            "subject_yes_total": subject_yes_total,
            "loop_yes": loop_yes_allowed,
            "loop_yes_total": loop_yes_total,
            "verifiable_yes": verifiable_yes_allowed,
            "verifiable_yes_total": verifiable_yes_total,
        },
        "top_risks": top_risks,
    }


def risk_level_from_case(d: Dict[str, Any]) -> str:
    """
    Convert granular gates + dimension scores into A–E.
    Conservative, but avoids false E due to some out-of-period evidence.
    """
    gf = d["gate_flags"]
    ds = d["dim_scores"]
    anchor_ok = meets_anchor_minimum(d)
    risk_cfg = _runtime_get("risk_rules", "risk_level_thresholds", {})
    degrade_cfg = _runtime_get("risk_rules", "anchor_degrade", {})
    hard_fail_level = str(_runtime_get("risk_rules", "hard_fail_level", "E"))

    # Fatal: period missing or no in-period coverage
    if gf.get("G1a") or gf.get("G1b"):
        return hard_fail_level

    # Core element missing
    if gf.get("G2") or gf.get("G3") or gf.get("G4"):
        return hard_fail_level

    # 闭环/可核验降级规则：不再机械性一票否决
    if gf.get("G5") and gf.get("G6"):
        return str(
            degrade_cfg.get("both_g5_g6_if_anchor_ok", "C")
            if anchor_ok
            else degrade_cfg.get("both_g5_g6_if_anchor_not_ok", "D")
        )

    if gf.get("G5") or gf.get("G6"):
        return str(
            degrade_cfg.get("one_of_g5_g6_if_anchor_ok", "B")
            if anchor_ok
            else degrade_cfg.get("one_of_g5_g6_if_anchor_not_ok", "C")
        )

    # If time/mapping are mediocre, keep conservative
    avg = (ds["Time"] + ds["Mapping"] + ds["Loop"] + ds["Verifiability"]) / 4.0
    if avg >= float(risk_cfg.get("avg_A_min", 85.0)) and min(ds.values()) >= float(risk_cfg.get("min_dim_A_min", 75.0)):
        return "A"
    if avg >= float(risk_cfg.get("avg_B_min", 68.0)):
        return "B"
    if avg >= float(risk_cfg.get("avg_C_min", 55.0)):
        return "C"
    return "D"


def decision_from_level(level: str) -> str:
    mapping = _runtime_get("risk_rules", "decision_text", {})
    if isinstance(mapping, dict):
        txt = mapping.get(level, "")
        if txt:
            return str(txt)
    return "补强后提交"


def build_risk_summary_block(level: str, d: Dict[str, Any], period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    decision = decision_from_level(level)
    gf = d["gate_flags"]
    gd = d["gate_details"]
    ds = d["dim_scores"]
    s = d["summary"]

    # Gate list (only triggered)
    trig = [k for k, v in gf.items() if v]
    trig_text = "、".join(trig) if trig else "无"

    # Examiner anchor emphasis: specified period is the first line
    period_txt = f"{period_start.isoformat()} ~ {period_end.isoformat()}" if (period_start and period_end) else "（未能解析指定期间）"

    # Top 3 risks
    top3 = d["top_risks"][:3]
    top3_txt = "\n".join([f"- {x}" for x in top3]) if top3 else "- （无）"

    return (
        f"【风险引擎输出（颗粒度增强）】\n"
        f"- 指定期间（审查锚点）：{period_txt}\n"
        f"- 提交建议：{decision}\n"
        f"- 风险等级：{level}\n"
        f"- 触发止损/结构闸门：{trig_text}\n"
        f"- 期内覆盖：高/中置信={s['in_period_highmed']}；低置信={s['in_period_low']}；期外={s['out_period']}；时间不明={s['unknown_time']}；跨期矛盾={s['contradictions']}\n"
        f"- 四维评分：时间性={ds['Time']}｜对应性={ds['Mapping']}｜闭环性={ds['Loop']}｜可核验性={ds['Verifiability']}\n"
        f"- Top风险点（证据导向）：\n{top3_txt}\n"
        f"- 说明：本输出为内部保守预审；对外文书应聚焦“指定期间 + 要件对应 + 证据闭环”，避免直接粘贴风险用语。"
    )


def _evidence_ids_for_segment(df_def: pd.DataFrame, s: dt.date, e: dt.date, limit: int = 6) -> List[str]:
    out: List[str] = []
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        if not time_anchor_allowed_row(r):
            seq += 1
            continue
        ti = infer_time_anchor_row(r)
        if evidence_in_period(ti["anchor_start"], ti["anchor_end"], s, e):
            out.append(evidence_id_display(eid, seq))
        seq += 1
    return out[:limit]


def _evidence_ids_for_target(df_def: pd.DataFrame, target: str, limit: int = 6) -> List[str]:
    out: List[str] = []
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        if target in safe_str(r.get("Proof Target (T1-T6)", "")):
            out.append(evidence_id_display(eid, seq))
        seq += 1
    return out[:limit]


def _extract_key_basis(note: str) -> str:
    s = safe_str(note)
    m = re.search(r"关键依据：([^；\n]+)", s)
    if m:
        return safe_str(m.group(1))
    return ""


def _clean_basis_phrase(s: str, max_len: int = 18) -> str:
    t = safe_str(s).replace("\n", " ")
    if not t:
        return ""
    t = re.sub(r"[“”\"'`]", "", t)
    t = re.sub(r"\s+", "", t)
    t = re.sub(r"^(首次评价|当天追评|评价|已购|关键依据[:：]?)", "", t)
    t = re.sub(r"[^一-龥A-Za-z0-9·-]", "", t)
    if len(t) > max_len:
        t = t[:max_len] + "…"
    return t


def _is_weak_basis_phrase(s: str) -> bool:
    t = safe_str(s)
    if not t:
        return True
    if len(t) <= 3:
        return True
    if re.fullmatch(r"[0-9A-Za-z年月日]+", t):
        return True
    weak_tokens = ["评价", "评论", "追评", "已购", "首次", "当天", "55", "1000"]
    return any(w in t for w in weak_tokens) and "罗兰宝贝" not in t


def _extract_product_phrase(text: str, max_len: int = 20) -> str:
    src = safe_str(text)
    if not src:
        return ""
    pattern = re.compile(r"(罗兰宝贝[^；，。,\n]{2,28})")
    candidates = [safe_str(m.group(1)) for m in pattern.finditer(src)]
    prod_kw = ["床笠", "床单", "床品", "被套", "四件套", "毛巾", "浴巾", "凉席", "毛毯", "产品", "家纺"]
    for c in candidates:
        if any(k in c for k in prod_kw):
            return _clean_basis_phrase(c, max_len=max_len)
    if candidates:
        return _clean_basis_phrase(candidates[0], max_len=max_len)
    return ""


def _argument_evidence_name(r: pd.Series) -> str:
    base = evidence_name_display(safe_str(r.get("Evidence Name", "")))
    note = safe_str(r.get("Risk Notes", ""))
    mapping_basis = safe_str(r.get("Mapping Basis", ""))
    inferred_name = _shorten_text(safe_str(r.get("Inferred Proof Name", "")), 12)

    key_basis = _extract_key_basis(note)
    if not key_basis:
        m = re.search(r"时间命中=([^；]+)", mapping_basis)
        if m:
            key_basis = safe_str(m.group(1))
    key_basis = _clean_basis_phrase(key_basis, 20)
    if _is_weak_basis_phrase(key_basis):
        alt = _extract_product_phrase(note + "；" + mapping_basis, max_len=20)
        if alt:
            key_basis = alt
        else:
            key_basis = ""

    if key_basis:
        return f"{base}（{key_basis}）"
    if inferred_name:
        return f"{base}（{inferred_name}）"
    return base


def _top_names_with_etc(names: List[str], limit: int = 3) -> str:
    uniq: List[str] = []
    seen = set()
    for n in names:
        x = evidence_name_display(n)
        if not x or x in seen:
            continue
        uniq.append(x)
        seen.add(x)
    if not uniq:
        return "无"
    picked = uniq[:limit]
    s = "、".join([f"《{x}》" for x in picked])
    if len(uniq) > limit:
        s += "等"
    return s


def _names_for_target(df_def: pd.DataFrame, target: str) -> List[str]:
    out: List[str] = []
    for _, r in df_def.iterrows():
        if target in ("T4", "T5", "T6") and (not time_anchor_allowed_row(r)):
            continue
        if target in safe_str(r.get("Proof Target (T1-T6)", "")):
            out.append(evidence_name_display(safe_str(r.get("Evidence Name", ""))))
    return out


def _names_for_segment(df_def: pd.DataFrame, s: dt.date, e: dt.date) -> List[str]:
    out: List[str] = []
    for _, r in df_def.iterrows():
        if not time_anchor_allowed_row(r):
            continue
        ti = infer_time_anchor_row(r)
        if evidence_in_period(ti["anchor_start"], ti["anchor_end"], s, e):
            out.append(evidence_name_display(safe_str(r.get("Evidence Name", ""))))
    return out


def _rows_for_segment(df_def: pd.DataFrame, s: dt.date, e: dt.date) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for _, r in df_def.iterrows():
        if not time_anchor_allowed_row(r):
            continue
        ti = infer_time_anchor_row(r)
        if evidence_in_period(ti["anchor_start"], ti["anchor_end"], s, e):
            out.append({
                "name": _argument_evidence_name(r),
                "page": safe_str(r.get("Page Range", "")),
            })
    return out


def _rows_for_target(df_def: pd.DataFrame, target: str) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for _, r in df_def.iterrows():
        # 使用时间/场景/闭环的要件说明仅引用A线实质使用证据，避免程序/管理材料混入。
        if target in ("T4", "T5", "T6") and (not time_anchor_allowed_row(r)):
            continue
        if target in safe_str(r.get("Proof Target (T1-T6)", "")):
            out.append({
                "name": _argument_evidence_name(r),
                "page": safe_str(r.get("Page Range", "")),
            })
    return out


def _target_sentence(label: str, rows: List[Dict[str, Any]]) -> str:
    if not rows:
        return f"{label}暂无可计入覆盖证据（详见证据目录与风险报告）"
    names = [x.get("name", "") for x in rows]
    pages = page_ref_text_from_rows(rows)
    return f"{label}以{_top_names_with_etc(names, 3)}为代表（见证据内容{pages}）"


def _derive_submission_verdict(diag: Dict[str, Any]) -> Tuple[str, str]:
    gf = diag.get("gate_flags", {}) if isinstance(diag, dict) else {}
    sm = diag.get("summary", {}) if isinstance(diag, dict) else {}
    in_highmed = int(sm.get("in_period_highmed", 0) or 0)
    goods_matched = int(sm.get("goods_matched", 0) or 0)
    if in_highmed > 0 and goods_matched > 0 and not (gf.get("G1a") or gf.get("G1b") or gf.get("G2") or gf.get("G3") or gf.get("G4")):
        return "PASS", "综合现有材料，已形成围绕指定期间的要件对应证明结构，具备提交基础。"
    if in_highmed > 0:
        return "WARN", "现有材料具备提交条件，但仍存在薄弱环节，建议按风险报告补强后提交。"
    return "FAIL", "当前仅形成证据目录与核对台账，尚不满足“指定期间实质使用”结论输出条件。"


def build_defense_final_conclusion_text(verdict: str, ci: Dict[str, str]) -> str:
    mark_name = safe_str(ci.get("mark_name", "")) or "涉案商标"
    reg_no = safe_str(ci.get("reg_no", "")) or "（待补）"
    if verdict == "PASS":
        return (
            f"综上，第{reg_no}号{mark_name}商标在指定期间内已形成可核验的使用证明链条，"
            "恳请贵局依法维持该商标注册。"
        )
    if verdict == "WARN":
        return (
            f"综上，第{reg_no}号{mark_name}商标现有材料已形成指定期间使用的基础证明框架，"
            "建议结合《风险报告》补强后提交，请贵局依法审查。"
        )
    return (
        f"综上，第{reg_no}号{mark_name}商标当前证据尚不足以形成“指定期间内实质使用成立”的结论，"
        "建议按《风险报告》补强后再行提交。"
    )


def rewrite_defense_conclusion_section(doc: Document, verdict: str, ci: Dict[str, str]) -> int:
    target_text = build_defense_final_conclusion_text(verdict, ci)
    markers = [
        "已在核定商品/服务上进行了真实、合法、连续的商业使用",
        "已在核定商品/服务上进行了真实、合法、连续的使用",
        "恳请贵局依法维持该商标注册",
    ]
    replaced = 0
    for p in doc.paragraphs:
        txt = safe_str(p.text)
        if txt and any(m in txt for m in markers):
            p.text = target_text
            replaced += 1
    if replaced:
        return replaced

    # 兜底：若模板未命中固定句式，尝试在“六、结论”标题后写入。
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        title = safe_str(p.text).strip()
        if title.startswith("六、结论"):
            for j in range(i + 1, len(paras)):
                if safe_str(paras[j].text).strip():
                    paras[j].text = target_text
                    return 1
            np = doc.add_paragraph(target_text)
            p._p.addnext(np._p)
            return 1
    return 0


def build_submission_conclusion_block(
    ci: Dict[str, str],
    df_def: pd.DataFrame,
    diag: Dict[str, Any],
    level: str,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> str:
    period_txt = f"{period_start.isoformat()} ~ {period_end.isoformat()}" if (period_start and period_end) else "（未解析）"
    verdict, closing = _derive_submission_verdict(diag)

    lines: List[str] = []
    lines.append("【提交结论（指定期间 + 要件对应）】")
    lines.append(f"- 结论等级：{verdict}")
    lines.append(f"- 指定期间：{period_txt}")
    lines.append("- 覆盖口径：A线（可计入覆盖）=实质使用证据；B线（仅参考）=程序/管理/期外材料。")
    lines.append("- 指定期间分段覆盖方面：")
    segs = get_period_segments(period_start, period_end)
    covered_segments: List[Tuple[dt.date, dt.date, List[Dict[str, Any]]]] = []
    for _, ss, ee in segs:
        seg_rows = _rows_for_segment(df_def, ss, ee)
        if seg_rows:
            covered_segments.append((ss, ee, seg_rows))
    if covered_segments:
        for i, (ss, ee, seg_rows) in enumerate(covered_segments, start=1):
            names = [x.get("name", "") for x in seg_rows]
            pages = page_ref_text_from_rows(seg_rows)
            lines.append(
                f"  {stage_label_cn_by_index(i)}（{ss.isoformat()} 至 {ee.isoformat()}）由{_top_names_with_etc(names, 3)}证据覆盖（见证据内容{pages}）。"
            )
    else:
        lines.append("  指定期间证据覆盖详见后附证据说明。")

    t1_rows = _rows_for_target(df_def, "T1")
    t2_rows = _rows_for_target(df_def, "T2")
    t3_rows = _rows_for_target(df_def, "T3")
    t4_rows = _rows_for_target(df_def, "T4")
    t5_rows = _rows_for_target(df_def, "T5")

    lines.append("- 要件对应方面：")
    lines.append(
        "  "
        + "，".join([
            _target_sentence("使用主体", t1_rows),
            _target_sentence("商标标识", t2_rows),
            _target_sentence("商品/服务对应", t3_rows),
        ])
        + "。"
    )
    lines.append(
        "  "
        + "，".join([
            _target_sentence("使用时间", t4_rows),
            _target_sentence("使用场景", t5_rows),
        ])
        + "。"
    )
    lines.append(f"- 结论：{closing}")
    return "\n".join(lines)


def build_evidence_argument_block(
    df_def: pd.DataFrame,
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
) -> str:
    rows: List[Dict[str, Any]] = []
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        ti = infer_time_anchor_row(r)
        rows.append({
            "eid": eid,
            "name": evidence_name_display(safe_str(r.get("Evidence Name", ""))),
            "type": safe_str(r.get("Type", "")) or "其他",
            "time_allowed": time_anchor_allowed_row(r),
            "coverage_lane": safe_str(r.get("Coverage Lane", "")),
            "page": safe_str(r.get("Page Range", "")) or "未标注",
            "targets": targets_to_cn(safe_str(r.get("Proof Target (T1-T6)", ""))),
            "note": safe_str(r.get("Risk Notes", "")),
            "start": ti["anchor_start"],
            "end": ti["anchor_end"],
            "time_label": range_label(ti["anchor_start"], ti["anchor_end"]),
            "in_period": (
                "期内" if (period_start and period_end and ti["anchor_start"] and ti["anchor_end"] and evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end))
                else ("期外" if (ti["anchor_start"] and ti["anchor_end"]) else "时间不明")
            ),
        })

    if not rows:
        return "未提供有效证据。"

    a_line_total = 0
    a_line_in_period_highmed = 0
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        if not time_anchor_allowed_row(r):
            continue
        a_line_total += 1
        ti = infer_time_anchor_row(r)
        in_period = False
        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            in_period = evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end)
        if in_period and ti["confidence"] in ("HIGH", "MEDIUM"):
            a_line_in_period_highmed += 1

    type_groups: Dict[str, List[Dict[str, Any]]] = {}
    for x in rows:
        type_groups.setdefault(x["type"], []).append(x)

    lines: List[str] = []
    lines.append("答辩人围绕指定期间内商标真实使用事实，对相关证据进行系统分类与对应说明，形成可核验的证据对应框架（与《证据目录_自动识别.docx》一一对应）。")
    lines.append("")
    lines.append("（一）证据分类及证明目的")

    type_order = {
        "程序文件": 1,
        "资质主体证明": 2,
        "线上店铺展示": 3,
        "商品展示页": 4,
        "交易凭证": 5,
    }
    ordered_types = sorted(type_groups.items(), key=lambda kv: (type_order.get(kv[0], 99), kv[0]))
    type_idx = 1
    for typ, arr in ordered_types:
        names = [a["name"] for a in arr]
        page_ref = page_ref_text_from_rows(arr)
        names_txt = _top_names_with_etc(names, 3)
        if typ == "程序文件":
            lines.append(f"{type_idx}. 程序文件（{len(arr)}份）")
            lines.append(f"以{names_txt}为代表，主要证明本案撤销程序启动时间、指定期间范围及审查背景的形成过程。")
            lines.append(f"（见证据内容{page_ref}）")
        elif typ == "资质主体证明":
            lines.append(f"{type_idx}. 主体资质证明材料（{len(arr)}份）")
            lines.append(f"以{names_txt}为代表，主要证明注册人主体资格的持续有效性及经营主体的稳定性，为使用行为的合法性提供基础支撑。")
            lines.append(f"（见证据内容{page_ref}）")
        elif typ == "线上店铺展示":
            lines.append(f"{type_idx}. 线上店铺展示与经营页面材料（{len(arr)}份）")
            lines.append(f"以{names_txt}为代表，主要证明：")
            lines.append("• 涉案标识在商品页面中的持续展示；")
            lines.append("• 商品内容与核定商品/服务范围的对应关系；")
            lines.append("• 实际存在的线上销售与经营场景。")
            lines.append("上述材料能够直观反映商标在真实市场环境中的持续使用状态。")
            lines.append(f"（见证据内容{page_ref}）")
        elif typ == "商品展示页":
            lines.append(f"{type_idx}. 商品展示页面材料（{len(arr)}份）")
            lines.append(f"以{names_txt}为代表，用于辅助证明涉案商标在商品页面中的具体使用方式及呈现状态，作为使用事实的补强材料。")
            lines.append(f"（见证据内容{page_ref}）")
        else:
            inferred_purposes: List[str] = []
            for a in arr:
                ip = _shorten_text(safe_str(a.get("Inferred Purpose", "")), 36).strip()
                if ip:
                    inferred_purposes.append(ip)
            uniq_p: List[str] = []
            for ip in inferred_purposes:
                if ip not in uniq_p:
                    uniq_p.append(ip)
            lines.append(f"{type_idx}. {typ}（{len(arr)}份）")
            if uniq_p:
                lines.append(f"以{names_txt}为代表，主要证明：{ '；'.join(uniq_p[:2]) }。")
            else:
                lines.append(f"以{names_txt}为代表，主要证明具体经营场景或商品展示相关事实。")
            lines.append(f"（见证据内容{page_ref}）")
        lines.append("")
        type_idx += 1

    segs = get_period_segments(period_start, period_end)
    lines.append("（二）指定期间分段覆盖与连续性说明")
    lines.append("结合各项证据形成时间及内容特征，答辩人对指定期间内的使用情况作如下分段说明：")
    lines.append("")
    covered_summary: List[Tuple[dt.date, dt.date, List[Dict[str, Any]]]] = []
    for _, ss, ee in segs:
        seg_rows = [
            x for x in rows
            if x.get("time_allowed", False)
            and isinstance(x["start"], dt.date)
            and isinstance(x["end"], dt.date)
            and evidence_in_period(x["start"], x["end"], ss, ee)
        ]
        if seg_rows:
            covered_summary.append((ss, ee, seg_rows))

    if covered_summary:
        for i, (ss, ee, seg_rows) in enumerate(covered_summary, start=1):
            names = [x["name"] for x in seg_rows]
            key_basis = ""
            for x in seg_rows:
                kb = _extract_key_basis(x["note"])
                if kb:
                    key_basis = kb
                    break
            page_ref = page_ref_text_from_rows(seg_rows)
            lines.append(f"{i}. {stage_label_cn_by_index(i)}（{date_cn(ss)}至{date_cn(ee)}）")
            lines.append(f"以{_top_names_with_etc(names, 3)}为代表，能够证明：")
            lines.append("• 指定期间该阶段已存在持续的对外经营或服务活动；")
            lines.append("• 商标标识在宣传页面或服务场景中持续展示；")
            lines.append("• 存在客观可核验的活动记录或经营痕迹。")

            has_procedure = any(x.get("type") == "程序文件" for x in seg_rows)
            key_basis = _shorten_text(re.sub(r"\s+", " ", key_basis).strip(" |"), 56) if key_basis else ""
            if key_basis:
                if has_procedure:
                    lines.append(f"其中，“{key_basis}”等程序信息与经营页面/活动材料相互印证，能够反映该阶段内不存在停止使用情形。")
                else:
                    lines.append(f"相关页面显示“{key_basis}”等信息，能够客观反映该阶段内已形成真实交易基础。")
            lines.append(f"（见证据内容{page_ref}）")
            lines.append("")
    else:
        lines.append("指定期间分段覆盖与分类情况详见证据目录。")
    lines.append("")

    lines.append("（三）证据整体证明效力说明")
    if a_line_in_period_highmed <= 0:
        lines.append("当前A线实质使用证据不足，尚不能形成“指定期间使用成立”的结论。")
        lines.append("已在《风险报告》中列明缺口清单与补强建议，建议补证后重新生成文书。")
    else:
        lines.append("综合上述证据材料，当前证据可支持如下证明方向：")
        lines.append("• 主体资格与使用主体关联关系；")
        lines.append("• 商标标识在商品/服务场景中的呈现；")
        lines.append("• 商品与核定范围的对应关系；")
        lines.append("• 指定期间内的时间锚点与经营场景。")
        if a_line_total > a_line_in_period_highmed:
            lines.append("• 另有部分A线证据需结合补强材料进一步提高证明强度。")

    unknown_rows = [x for x in rows if x["time_label"] == "时间不明"]
    out_period_rows = [x for x in rows if x["in_period"] == "期外"]
    if unknown_rows:
        lines.append("")
        lines.append("另有部分时间信息不明确材料，仅作为辅助印证使用，不单独作为核心时间依据。")
    if out_period_rows:
        lines.append("期外材料仅用于补强经营连续性，不计入指定期间使用证明。")

    return "\n".join(lines)


def _collect_loop_verif_ids(df_def: pd.DataFrame) -> Tuple[List[str], List[str]]:
    loop_ids: List[str] = []
    verif_ids: List[str] = []
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        eid_disp = evidence_id_display(eid, seq)
        if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "Y":
            loop_ids.append(eid_disp)
        if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "Y":
            verif_ids.append(eid_disp)
        seq += 1
    return loop_ids, verif_ids


def build_internal_segment_rows(df_def: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> List[List[str]]:
    rows: List[List[str]] = []
    for lbl, ss, ee in get_period_segments(period_start, period_end):
        ids = _evidence_ids_for_segment(df_def, ss, ee, limit=30)
        rows.append([
            period_label_cn(lbl),
            f"{ss.isoformat()} 至 {ee.isoformat()}",
            str(len(ids)),
            "、".join(ids) if ids else "无",
        ])
    if not rows:
        rows = [["-", "（未解析）", "0", "无"]]
    return rows


def build_internal_target_rows(df_def: pd.DataFrame) -> List[List[str]]:
    rows: List[List[str]] = []
    for t in ["T1", "T2", "T3", "T4", "T5", "T6"]:
        ids = _evidence_ids_for_target(df_def, t, limit=30)
        rows.append([TARGET_CN.get(t, t), str(len(ids)), "、".join(ids) if ids else "无"])
    return rows


def build_unknown_time_rows(df_def: pd.DataFrame) -> List[List[str]]:
    rows: List[List[str]] = []
    seq = 1
    for _, r in df_def.iterrows():
        eid = safe_str(r.get("Evidence ID", ""))
        if not eid:
            continue
        ti = infer_time_anchor_row(r)
        if ti["anchor_start"] and ti["anchor_end"]:
            seq += 1
            continue
        eid_disp = evidence_id_display(eid, seq)
        rows.append([
            eid_disp,
            evidence_name_display(safe_str(r.get("Evidence Name", ""))),
            safe_str(r.get("Page Range", "")) or "未标注",
            compact_pending_fact(r, ti),
        ])
        seq += 1
    if not rows:
        rows = [["无", "无", "-", "无"]]
    return rows


def _split_items(text: str) -> List[str]:
    s = safe_str(text)
    if not s:
        return []
    raw = re.split(r"[；;\n]+", s)
    return [x.strip() for x in raw if x.strip()]


def build_quality_eval_rows_and_actions(ci: Dict[str, str]) -> Tuple[List[List[str]], List[str]]:
    checks = [
        ("案件事实一致性（CaseInfo）", "caseinfo_validation_status", "caseinfo_validation_errors", "caseinfo_validation_warnings", "caseinfo_validation_path"),
        ("时间质量评估", "time_quality_status", "time_quality_errors", "time_quality_warnings", "time_quality_path"),
        ("页码一致性评估", "page_map_validation_status", "page_map_validation_errors", "", "page_map_validation_path"),
        ("证据命名质量评估", "name_quality_status", "name_quality_errors", "name_quality_warnings", "name_quality_path"),
    ]
    rows: List[List[str]] = []
    actions: List[str] = []

    for title, sk, ek, wk, pk in checks:
        status = safe_str(ci.get(sk, "")) or "未提供"
        errs = _split_items(ci.get(ek, ""))
        warns = _split_items(ci.get(wk, "")) if wk else []
        detail = safe_str(ci.get(pk, "")) or "-"
        issue = "；".join((errs + warns)[:6]) if (errs or warns) else "无明显异常"
        rows.append([title, status, issue, detail])

        full_issue = "；".join(errs + warns)
        if title.startswith("案件事实一致性"):
            if "reg_no" in full_issue:
                actions.append("核对并锁定注册号（以通知书/注册证为准），同步更新台账与文书字段。")
            if "class" in full_issue:
                actions.append("核对国际分类号并统一到全案字段，避免类别错配。")
            if "mark_name" in full_issue:
                actions.append("按注册证统一商标文字写法，清除OCR噪音词或截断词。")
            if "respondent" in full_issue or "applicant" in full_issue:
                actions.append("核对申请人/被申请人主体名称并统一全案口径。")
            if "use_period" in full_issue:
                actions.append("核对指定期间起止日期，按通知书锚定并统一输出。")

        if title.startswith("时间质量评估"):
            if "期内高/中置信证据不足" in full_issue:
                actions.append("优先补充含稳定日期锚点的期内证据（评价日期、订单时间、物流签收时间）。")
            if "低置信占比超阈值" in full_issue:
                actions.append("降低低置信截图占比，补充第三方可核验材料（平台后台导出、票据流水等）。")
            if "时间不明占比超阈值" in full_issue:
                actions.append("对时间不明证据逐页人工核验并回填覆盖规则，形成可核验时间锚点。")

        if title.startswith("页码一致性评估") and full_issue and full_issue != "无明显异常":
            actions.append("重建证据合订并刷新 page_map 后再生成文书，确保目录页码与合订页一致。")

        if title.startswith("证据命名质量评估"):
            if "命名不达标证据数" in full_issue:
                actions.append("按“E###_日期_类型_主题.pdf”模板批量修订证据名称并复跑。")
            if safe_str(ci.get("name_pending_path", "")):
                actions.append(f"按清单逐项修订命名并复核：{safe_str(ci.get('name_pending_path', ''))}")

    # 去重并限制长度
    uniq: List[str] = []
    seen = set()
    for x in actions:
        if x not in seen:
            uniq.append(x)
            seen.add(x)
    if not uniq:
        uniq = ["当前自动质检未发现需要立即补充的高优先项，可按常规人工复核后提交。"]
    return rows, uniq[:12]


def write_internal_risk_report_docx(
    out_path: str,
    ci: Dict[str, str],
    period_start: Optional[dt.date],
    period_end: Optional[dt.date],
    level: str,
    diag: Dict[str, Any],
    risk_rows: List[List[str]],
    checklist_rows: List[List[str]],
    gate_summary: Dict[str, Any],
    df_def: pd.DataFrame,
):
    doc = Document()
    doc.add_heading("风险报告（内部评估）", level=1)
    doc.add_paragraph(f"商标：{safe_str(ci.get('mark_name',''))}")
    doc.add_paragraph(f"注册号：{safe_str(ci.get('reg_no',''))}")
    period_txt = f"{period_start.isoformat()} ~ {period_end.isoformat()}" if (period_start and period_end) else "（未解析）"
    doc.add_paragraph(f"指定期间：{period_txt}")
    doc.add_paragraph("说明：本报告用于内部评估与提交前核对，不作为正式答辩正文对外提交。")
    doc.add_paragraph("校准说明：本版本以“已通过撤三测试案例”作为最低要求锚定，交易闭环不足按补强项处理，不作机械性否决。")

    doc.add_heading("一、内部评价结论（非门禁）", level=2)
    gate_rows = [
        ["评价结论", safe_str(gate_summary.get("decision", ""))],
        ["风险等级", safe_str(gate_summary.get("risk_level", ""))],
        ["锚定最低要求（通过案例）", "满足" if gate_summary.get("anchor_ok") else "未满足"],
        ["不通过项", str(gate_summary.get("hard_count", 0))],
        ["补强项", str(gate_summary.get("strengthen_count", 0))],
        ["通过项", str(gate_summary.get("pass_count", 0))],
        ["关注项", "无" if not gate_summary.get("triggered_gates") else "、".join(gate_summary.get("triggered_gates", []))],
    ]
    tg = doc.add_table(rows=1, cols=2)
    fill_table(tg, ["项目", "内容"], gate_rows)
    if gate_summary.get("hard_items"):
        doc.add_paragraph("不通过项：" + "；".join(gate_summary.get("hard_items", [])[:6]))
    if gate_summary.get("strengthen_items"):
        doc.add_paragraph("补强项：" + "；".join(gate_summary.get("strengthen_items", [])[:6]))

    doc.add_heading("二、期间覆盖矩阵（内部）", level=2)
    t1 = doc.add_table(rows=1, cols=2)
    fill_table(t1, ["项目", "内容"], risk_rows)

    doc.add_heading("三、分段覆盖明细（内部）", level=2)
    seg_rows = build_internal_segment_rows(df_def, period_start, period_end)
    t2 = doc.add_table(rows=1, cols=4)
    fill_table(t2, ["阶段", "期间", "覆盖份数", "证据ID"], seg_rows)

    doc.add_heading("四、六项要件覆盖（内部）", level=2)
    target_rows = build_internal_target_rows(df_def)
    t3 = doc.add_table(rows=1, cols=3)
    fill_table(t3, ["要件", "覆盖份数", "证据ID"], target_rows)

    loop_ids, verif_ids = _collect_loop_verif_ids(df_def)
    doc.add_heading("五、闭环与可核验统计（内部）", level=2)
    doc.add_paragraph(f"交易闭环证据：{len(loop_ids)}份（{'、'.join(loop_ids[:20]) if loop_ids else '无'}）。")
    doc.add_paragraph(f"可核验证据：{len(verif_ids)}份（{'、'.join(verif_ids[:20]) if verif_ids else '无'}）。")

    doc.add_heading("六、提交前检查清单（自动）", level=2)
    t4 = doc.add_table(rows=1, cols=5)
    fill_table(t4, ["序号", "检查项", "结果", "依据", "处理建议"], checklist_rows)

    doc.add_heading("七、时间不明证据（人工核验）", level=2)
    unknown_rows = build_unknown_time_rows(df_def)
    t5 = doc.add_table(rows=1, cols=4)
    fill_table(t5, ["证据ID", "证据名称", "证据内容页码", "核验要点"], unknown_rows)
    doc.add_paragraph("核验说明：上述“时间不明”证据请按《证据内容_重排合并.pdf》对应页进行人工核验，并补充稳定时间锚点。")

    doc.add_heading("八、重点风险（内部）", level=2)
    for x in diag.get("top_risks", []):
        doc.add_paragraph(f"- {x}")

    doc.add_heading("九、自动质检结果（非门禁）", level=2)
    q_rows, q_actions = build_quality_eval_rows_and_actions(ci)
    t6 = doc.add_table(rows=1, cols=4)
    fill_table(t6, ["检查项", "状态", "主要发现", "详情文件"], q_rows)

    doc.add_heading("十、补充建议（自动）", level=2)
    for i, x in enumerate(q_actions, start=1):
        doc.add_paragraph(f"{i}. {x}")

    apply_heiti_font(doc)
    doc.save(out_path)


def md_table(rows: List[Dict[str, Any]], cols: List[str], max_rows: int = 80) -> str:
    if not rows:
        return "（无）"
    rows = rows[:max_rows]
    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join(["---"] * len(cols)) + " |"
    lines = [header, sep]
    for r in rows:
        line = "| " + " | ".join([safe_str(r.get(c, "")).replace("\n", " ").replace("|", "／") for c in cols]) + " |"
        lines.append(line)
    if len(rows) >= max_rows:
        lines.append(f"\n> 仅展示前 {max_rows} 行，完整明细请在 Excel 台账中查看。")
    return "\n".join(lines)


# -----------------------------
# Cross-exam blocks (use upgraded time anchor)
# -----------------------------
def classify_defect_for_exhibit(r: pd.Series, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> List[Tuple[str, str]]:
    defects = []
    ti = infer_time_anchor_row(r)

    if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "N":
        defects.append(("A", "A1 原件/可核验来源缺失"))
    if safe_str(r.get("Goods/Services", "")) == "":
        defects.append(("B", "B1 商品/服务不清"))
    if normalize_yesno(r.get("Mark Shown (Y/N)", "")) == "N":
        defects.append(("B", "B2 商标未显示或未绑定"))
    if normalize_yesno(r.get("Subject Match (Y/N)", "")) == "N":
        defects.append(("B", "B3 使用主体不一致"))

    # Time defect: only if we can parse period AND exhibit time anchor
    if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
        if not evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end):
            defects.append(("C", "C1 时间不在指定期间（以时间锚点为准）"))
        elif ti["confidence"] == "LOW":
            defects.append(("C", "C2 时间锚点低置信（可编辑/不可核验风险）"))
    else:
        # period missing or exhibit time unknown: still flag as time weakness
        defects.append(("C", "C2 时间锚点不明/期间不可解析"))

    if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "N":
        defects.append(("D", "D1 缺乏闭环/仅单一环节"))

    if ti["contradiction"]:
        defects.append(("C", "C3 疑似跨期矛盾（形成日与锚点不一致）"))

    if not defects:
        defects.append(("F", "F1 抗反问测试（保守）"))

    # pick top 2 by priority A/B/C/D/E/F
    priority = {"A": 1, "B": 2, "C": 3, "D": 4, "E": 5, "F": 6}
    defects = sorted(defects, key=lambda x: priority.get(x[0], 9))
    return defects[:2]


def questions_for_defect(code: str) -> List[str]:
    q = {
        "A": [
            "是否提供原件/原始电子文件及生成路径（后台记录/流水/系统导出）？",
            "是否存在可独立核验的第三方记录（税控/平台/物流/支付流水）？",
        ],
        "B": [
            "商品/服务描述能否与核定项目逐一对应？是否有型号、SKU、清单或服务明细？",
            "商标是否在包装/页面/交易凭证中明确呈现并与交易标的绑定？",
        ],
        "C": [
            "时间锚点来源是什么？能否提供后台导出/存证/日志证明其不可编辑？",
            "锚点所示时间是否覆盖指定期间？若为范围，是否能证明范围内持续在线/持续履行？",
        ],
        "D": [
            "是否存在订单、付款、物流/交付、签收/验收等闭环材料？",
            "合同是否实际履行？是否有对账或履行痕迹？",
        ],
        "E": [
            "交易金额、频次、相对方结构是否符合经营规模与行业惯例？",
            "是否存在持续往来记录与独立支付/物流证据？",
        ],
        "F": [
            "能否回答“谁、何时、就何商品/服务、如何使用”四要素？",
        ],
    }
    return q.get(code, q["F"])


def build_cross_exam_block(df_opp: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    blocks = []
    for _, r in df_opp.iterrows():
        xid = safe_str(r.get("Exhibit ID", ""))
        if not xid:
            continue
        xtype = safe_str(r.get("Type", ""))
        carrier = safe_str(r.get("Carrier", ""))
        ti = infer_time_anchor_row(r)
        anchor = range_label(ti["anchor_start"], ti["anchor_end"])

        in_period = "不明"
        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            in_period = "是" if evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end) else "否"

        blocks.append(f"[展品 {xid} | 类型：{xtype or '—'} | 载体：{carrier or '—'} | 时间锚点：{anchor}（{conf_to_cn(ti['confidence'])}，期内:{in_period}）]")

        defects = classify_defect_for_exhibit(r, period_start, period_end)
        for code, label in defects:
            blocks.append(f"缺陷类型：{code}（{label}）")
            qs = questions_for_defect(code)[:2]
            for i, qq in enumerate(qs, start=1):
                blocks.append(f"- 关键问题{i}：{qq}")
            blocks.append(f"- 结论：{CLOSINGS[0]}")
        blocks.append("")
    return "\n".join(blocks).strip() if blocks else "（未提供可质证的对方证据清单）"


def overall_closing_from_defects(df_opp: pd.DataFrame, period_start: Optional[dt.date], period_end: Optional[dt.date]) -> str:
    g = set()
    # Use time anchor rather than Formation Date
    for _, r in df_opp.iterrows():
        xid = safe_str(r.get("Exhibit ID", ""))
        if not xid:
            continue
        ti = infer_time_anchor_row(r)

        if period_start and period_end and ti["anchor_start"] and ti["anchor_end"]:
            if not evidence_in_period(ti["anchor_start"], ti["anchor_end"], period_start, period_end):
                g.add("G1")
        else:
            g.add("G1")

        if normalize_yesno(r.get("Mark Shown (Y/N)", "")) == "N":
            g.add("G3")
        if normalize_yesno(r.get("Subject Match (Y/N)", "")) == "N":
            g.add("G4")
        if safe_str(r.get("Goods/Services", "")) == "":
            g.add("G2")
        if normalize_yesno(r.get("Commercial Loop (Y/N)", "")) == "N":
            g.add("G5")
        if normalize_yesno(r.get("Original/Verifiable (Y/N)", "")) == "N":
            g.add("G6")

    if not g:
        return "综上，请审查机关结合全案证据，就对方证据与核定商品/服务、指定期间及商标标识之间的对应关系进行审慎审查。"
    return "综上，对方证据在" + "、".join(sorted(g)) + "方面存在突出缺陷，难以证明指定期间内对核定商品/服务的真实商标使用，依法不应采信。"


# -----------------------------
# MAIN
# -----------------------------
def main(args: argparse.Namespace):
    paths = resolve_runtime_paths(
        excel_override=args.xlsx,
        defense_tpl_override=args.defense_template,
        cross_tpl_override=args.cross_template,
        out_dir_override=args.out_dir,
    )
    excel_path = paths["excel"]
    defense_tpl = paths["defense_tpl"]
    cross_tpl = paths["cross_tpl"]
    out_dir = paths["out_dir"]

    os.makedirs(out_dir, exist_ok=True)
    run_id = f"generate_{uuid.uuid4().hex[:8]}"
    logs_dir = Path(out_dir) / "logs"
    set_run_context(logs_dir, run_id)
    global LOGGER
    LOGGER = setup_logger(__name__, log_dir=logs_dir)
    rc = 1
    audit({
        "type": "run_start",
        "step": "generate_suite_main",
        "ok": True,
        "file": excel_path,
    })

    try:
        rule_profile = _load_runtime_rules(base_dir=Path(BASE_DIR), out_dir=Path(out_dir), run_id=run_id)
        ci = read_caseinfo(excel_path)
        period_start = _to_date_obj(ci.get("use_period_start"))
        period_end = _to_date_obj(ci.get("use_period_end"))

        df_def = pd.read_excel(excel_path, sheet_name="DefenseEvidence")
        page_map_payload = load_page_map(ci, excel_path)
        df_def = apply_page_map_to_df(df_def, page_map_payload)
        df_def_rows = df_def.fillna("").to_dict(orient="records")
        write_low_confidence_fields(
            rows=df_def_rows,
            out_dir=Path(out_dir),
            run_id=run_id,
            logger=LOGGER,
            audit=audit,
        )
        df_def = pd.DataFrame(df_def_rows)
        df_opp = pd.read_excel(excel_path, sheet_name="OpponentEvidence")

        # ---- Defense diagnostics (granular risk) ----
        diag = build_defense_diagnostics(df_def, period_start, period_end)
        level = risk_level_from_case(diag)
        only_cross = (str(getattr(args, "only_cross", "off")) == "on")

        mark_name_val = safe_str(ci.get("mark_name", "")) or "（填写商标文字）"
        mapping = {
            "{case_no}": safe_str(ci.get("case_no", "")),
            "{mark_name}": mark_name_val,
            "{mark_image}": safe_str(ci.get("mark_image", "")),
            "{reg_no}": safe_str(ci.get("reg_no", "")),
            "{respondent}": safe_str(ci.get("respondent", "")),
            "{respondent_address}": safe_str(ci.get("respondent_address", "")),
            "{applicant}": safe_str(ci.get("applicant", "")),
            "{agent_company}": safe_str(ci.get("agent_company", "")),
            "{agent_address}": safe_str(ci.get("agent_address", "")),
            "{contact_phone}": safe_str(ci.get("contact_phone", "")),
            "{class}": safe_str(ci.get("class", "")),
            "{use_period_start}": safe_str(ci.get("use_period_start", "")),
            "{use_period_end}": safe_str(ci.get("use_period_end", "")),
            "{use_period_start_cn}": date_cn(period_start),
            "{use_period_end_cn}": date_cn(period_end),
            "{designated_goods_services}": safe_str(ci.get("designated_goods_services", "")),
            "{revoked_goods_services}": safe_str(ci.get("revoked_goods_services", "")),
            "{defense_goods_services}": safe_str(ci.get("defense_goods_services", "")),
            "{cnipa_notice_ref}": safe_str(ci.get("cnipa_notice_ref", "")),
        }
        segs = build_period_segments(period_start, period_end)
        mapping["{period_p1}"] = segs["p1"]
        mapping["{period_p2}"] = segs["p2"]
        mapping["{period_p3}"] = segs["p3"]

        idx_block = build_evidence_index_block(df_def, period_start, period_end)
        idx_headers = ["序号", "证据ID", "证据名称", "类型", "时间锚点", "期内", "要件", "页码", "待证事实"]
        idx_rows = build_evidence_index_table_rows(df_def, period_start, period_end)
        risk_rows = build_risk_summary_table_rows(level, diag, period_start, period_end)
        checklist_rows = build_submission_checklist_rows(ci, df_def, diag, level, period_start, period_end)
        gate_summary = build_gate_decision(checklist_rows, level, diag)

        if not only_cross:
            defense_doc = Document(defense_tpl)
            replace_all(defense_doc, mapping)
            replace_all(defense_doc, {
                "（三）指定期间分段覆盖说明（P1/P2/P3）": "（三）指定期间分段覆盖说明",
                "证据目录详见《证据目录_指定期间使用证据.doc》，自动索引如下：": "",
                "五、证据目录与页码": "五、证据说明",
            })
            remove_paragraphs_containing(
                defense_doc,
                [
                    "证据目录详见《证据目录_指定期间使用证据.doc》，自动索引如下：",
                    "证据分组与事实链如下（与《证据目录_自动识别.docx》对应）：",
                ],
            )

            sections = build_T_sections(df_def, period_start, period_end)
            for t, txt in sections.items():
                replace_all(defense_doc, {f"{{{t}_section}}": txt})

            submission_block = build_submission_conclusion_block(ci, df_def, diag, level, period_start, period_end)
            submission_verdict, _ = _derive_submission_verdict(diag)
            summary_headers = ["证据分类", "份数", "代表证据（节选）", "证明重点", "页码范围"]
            summary_rows = build_defense_summary_table_rows(df_def)
            evidence_argument_block = build_evidence_argument_block(df_def, period_start, period_end)
            if insert_table_for_marker(defense_doc, "{evidence_index_block}", summary_headers, summary_rows) == 0:
                replace_all(defense_doc, {"{evidence_index_block}": evidence_argument_block})
            else:
                replace_all(defense_doc, {"{evidence_index_block}": ""})
                replace_all(defense_doc, {"{risk_summary_block}": submission_block + "\n\n" + evidence_argument_block})

            if "{risk_summary_block}" in "\n".join([p.text for p in defense_doc.paragraphs]):
                replace_all(defense_doc, {"{risk_summary_block}": submission_block})
            remove_internal_submission_paragraphs(defense_doc)
            rewrite_defense_conclusion_section(defense_doc, submission_verdict, ci)
            apply_heiti_font(defense_doc)

            defense_out = os.path.join(out_dir, args.defense_out)
            defense_doc.save(defense_out)
        else:
            defense_out = ""

        cross_doc = Document(cross_tpl)
        replace_all(cross_doc, mapping)
        replace_all(cross_doc, {
            "一、指定期间分段（P1/P2/P3）": "一、指定期间分段覆盖说明",
        })
        replace_cross_directory_note(cross_doc, ci)
        remove_paragraphs_containing(cross_doc, ["二、期间覆盖矩阵（概览）", "四、提交前检查清单（自动）"])
        replace_all(cross_doc, {"{risk_summary_block}": ""})
        replace_all(cross_doc, {"{cross_exam_block}": ""})
        replace_all(cross_doc, {"{overall_closing_block}": ""})
        remove_internal_submission_paragraphs(cross_doc)
        if insert_table_for_marker(cross_doc, "{evidence_index_block}", idx_headers, idx_rows) == 0:
            replace_all(cross_doc, {"{evidence_index_block}": idx_block})
        apply_heiti_font(cross_doc)

        cross_out = os.path.join(out_dir, args.cross_out)
        cross_doc.save(cross_out)

        if not only_cross:
            risk_name = args.risk_out
            if not risk_name.lower().endswith(".docx"):
                risk_name = os.path.splitext(risk_name)[0] + ".docx"
            risk_docx = os.path.join(out_dir, risk_name)
            write_internal_risk_report_docx(
                out_path=risk_docx,
                ci=ci,
                period_start=period_start,
                period_end=period_end,
                level=level,
                diag=diag,
                risk_rows=risk_rows,
                checklist_rows=checklist_rows,
                gate_summary=gate_summary,
                df_def=df_def,
            )
        else:
            risk_docx = ""

        print("Generated:")
        if defense_out:
            print(" -", defense_out)
        print(" -", cross_out)
        if risk_docx:
            print(" -", risk_docx)
        try:
            build_reason_chain(
                casebook_path=Path(excel_path),
                out_dir=Path(out_dir),
                run_id=run_id,
                rule_profile=rule_profile,
                logger=LOGGER,
                audit=audit,
            )
        except Exception as exc:
            LOGGER.exception("生成 reason_chain 失败")
            audit({
                "type": "exception",
                "step": "build_reason_chain",
                "file": excel_path,
                "error": str(exc),
                "traceback": traceback.format_exc(),
                "ok": False,
                "reason_code": "reason_chain_failed",
            })
            raise
        rc = 0
    except Exception as exc:
        LOGGER.exception("文书生成主流程失败")
        audit({
            "type": "exception",
            "step": "generate_suite_main",
            "file": excel_path,
            "error": str(exc),
            "traceback": traceback.format_exc(),
            "ok": False,
            "reason_code": "generate_suite_failed",
        })
        raise
    finally:
        audit({
            "type": "run_end",
            "step": "generate_suite_main",
            "ok": rc == 0,
            "exit_code": rc,
        })


if __name__ == "__main__":
    main(parse_args())

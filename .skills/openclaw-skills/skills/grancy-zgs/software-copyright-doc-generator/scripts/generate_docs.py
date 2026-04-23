#!/usr/bin/env python3
"""
软件著作权文档生成器
根据收集的信息生成软著申请所需的各类文档
支持简化的信息收集流程（仅需软件基本信息）
"""

import os
import sys
from datetime import datetime
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed. Installing...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT


def create_copyright_application_form(data: Dict[str, Any], output_dir: str) -> str:
    """
    创建软件著作权登记申请表

    Args:
        data: 包含所有表单数据的字典
        output_dir: 输出目录

    Returns:
        生成的文件路径
    """
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('计算机软件著作权登记申请表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加申请日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'申请日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(10)

    # 软件基本信息
    doc.add_heading('一、软件基本信息', level=1)

    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 软件名称规范化处理
    software_full_name = data.get('software_full_name', data.get('software_name', ''))
    software_short_name = data.get('software_short_name', '')

    fields1 = [
        ('软件全称', software_full_name),
        ('软件简称', software_short_name),
        ('版本号', data.get('version', 'V1.0')),
        ('权利取得方式', data.get('rights_acquisition', '原始取得'))
    ]

    for i, (field, value) in enumerate(fields1):
        table1.rows[i].cells[0].text = field
        table1.rows[i].cells[1].text = value

    # 运行环境信息
    doc.add_heading('二、运行环境信息', level=1)

    # 硬件环境
    hardware_cpu = data.get('hardware_cpu', '')
    hardware_memory = data.get('hardware_memory', '')
    hardware_storage = data.get('hardware_storage', '')

    hardware_text = ''
    if hardware_cpu:
        hardware_text += f'处理器：{hardware_cpu}'
    if hardware_memory:
        hardware_text += f'；内存：{hardware_memory}'
    if hardware_storage:
        hardware_text += f'；存储空间：{hardware_storage}'

    # 软件环境
    software_os = data.get('software_os', '')
    software_language = data.get('software_language', '')
    software_database = data.get('software_database', '')
    software_middleware = data.get('software_middleware', '')

    software_text = ''
    if software_os:
        software_text += f'操作系统：{software_os}'
    if software_language:
        software_text += f'；开发语言：{software_language}'
    if software_database:
        software_text += f'；数据库：{software_database}'
    if software_middleware:
        software_text += f'；中间件：{software_middleware}'

    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    table2.rows[0].cells[0].text = '硬件环境'
    table2.rows[0].cells[1].text = hardware_text if hardware_text else '待补充'
    table2.rows[1].cells[0].text = '软件环境'
    table2.rows[1].cells[1].text = software_text if software_text else '待补充'

    # 技术特点信息
    doc.add_heading('三、技术特点信息', level=1)

    table3 = doc.add_table(rows=4, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields3 = [
        ('体系结构', data.get('architecture', '')),
        ('技术框架', data.get('tech_framework', '')),
        ('创新点', data.get('innovation_points', '')),
        ('主要功能', data.get('main_functions', ''))
    ]

    for i, (field, value) in enumerate(fields3):
        table3.rows[i].cells[0].text = field
        table3.rows[i].cells[1].text = value if value else '待补充'

    # 创作说明
    doc.add_heading('四、创作说明', level=1)

    table4 = doc.add_table(rows=2, cols=2)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields4 = [
        ('开发完成日期', data.get('completion_date', '')),
        ('发表状态', data.get('publication_status', '未发表'))
    ]

    for i, (field, value) in enumerate(fields4):
        table4.rows[i].cells[0].text = field
        table4.rows[i].cells[1].text = value

    # 保存文档
    filename = f'软件著作权登记表_{software_full_name or "未命名"}'
    # 移除文件名中的非法字符
    filename = filename.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    filepath = os.path.join(output_dir, f'{filename}.docx')
    doc.save(filepath)

    return filepath


def create_software_manual(data: Dict[str, Any], output_dir: str) -> str:
    """
    创建软件说明书

    Args:
        data: 包含所有表单数据的字典
        output_dir: 输出目录

    Returns:
        生成的文件路径
    """
    doc = Document()

    software_name = data.get('software_full_name', data.get('software_name', '软件名称'))
    software_version = data.get('version', 'V1.0')

    # 设置文档标题
    title = doc.add_heading(f'{software_name} {software_version} 说明书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 软件概述
    doc.add_heading('一、软件概述', level=1)
    overview = data.get('overview', f'{software_name}是一款功能完善的应用程序。')
    doc.add_paragraph(overview)

    # 技术特点
    doc.add_heading('二、技术特点', level=1)

    # 体系结构
    arch = data.get('architecture', '')
    if arch:
        p = doc.add_paragraph()
        p.add_run('体系结构：').bold = True
        p.add_run(arch)

    # 技术框架
    framework = data.get('tech_framework', '')
    if framework:
        p = doc.add_paragraph()
        p.add_run('技术框架：').bold = True
        p.add_run(framework)

    # 创新点
    innovation = data.get('innovation_points', '')
    if innovation:
        p = doc.add_paragraph()
        p.add_run('创新点：').bold = True
        p.add_run(innovation)

    # 运行环境
    doc.add_heading('三、运行环境', level=1)

    # 硬件环境
    doc.add_paragraph('3.1 硬件环境要求', style='List Number')
    hardware_cpu = data.get('hardware_cpu', '待补充')
    hardware_memory = data.get('hardware_memory', '待补充')
    hardware_storage = data.get('hardware_storage', '待补充')

    doc.add_paragraph(f'处理器：{hardware_cpu}')
    doc.add_paragraph(f'内存：{hardware_memory}')
    doc.add_paragraph(f'存储空间：{hardware_storage}')

    # 软件环境
    doc.add_paragraph('3.2 软件环境要求', style='List Number')

    software_os = data.get('software_os', '待补充')
    software_language = data.get('software_language', '待补充')
    software_database = data.get('software_database', '')
    software_middleware = data.get('software_middleware', '')

    doc.add_paragraph(f'操作系统：{software_os}')
    doc.add_paragraph(f'开发语言/运行环境：{software_language}')

    if software_database:
        doc.add_paragraph(f'数据库：{software_database}')
    if software_middleware:
        doc.add_paragraph(f'中间件：{software_middleware}')

    # 主要功能
    doc.add_heading('四、主要功能', level=1)
    main_functions = data.get('main_functions', '本软件包含以下主要功能模块：')
    doc.add_paragraph(main_functions)

    # 操作说明
    doc.add_heading('五、操作说明', level=1)
    operation_guide = data.get('operation_guide', '详细操作说明请参考用户手册。')
    doc.add_paragraph(operation_guide)

    # 版本历史
    doc.add_heading('六、版本历史', level=1)

    version_table = doc.add_table(rows=2, cols=3)
    version_table.style = 'Table Grid'
    version_table.rows[0].cells[0].text = '版本号'
    version_table.rows[0].cells[1].text = '发布日期'
    version_table.rows[0].cells[2].text = '更新内容'
    version_table.rows[1].cells[0].text = software_version
    version_table.rows[1].cells[1].text = datetime.now().strftime('%Y-%m-%d')
    version_table.rows[1].cells[2].text = '初始版本发布'

    # 保存文档
    filename = f'软件说明书_{software_name}'
    filename = filename.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    filepath = os.path.join(output_dir, f'{filename}.docx')
    doc.save(filepath)

    return filepath


def create_tech_document(data: Dict[str, Any], output_dir: str) -> str:
    """
    创建技术文档资料清单

    Args:
        data: 包含所有表单数据的字典
        output_dir: 输出目录

    Returns:
        生成的文件路径
    """
    software_name = data.get('software_full_name', data.get('software_name', '软件'))

    # 保存为Markdown格式
    filename = f'技术文档资料清单_{software_name}.md'
    filename = filename.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f'# {software_name} 技术文档资料清单\n\n')
        f.write(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n\n')

        # 软件基本信息摘要
        f.write('## 软件基本信息摘要\n\n')
        f.write(f'| 字段 | 内容 |\n')
        f.write(f'|------|------|\n')
        f.write(f'| 软件全称 | {data.get("software_full_name", data.get("software_name", ""))} |\n')
        f.write(f'| 软件简称 | {data.get("software_short_name", "")} |\n')
        f.write(f'| 版本号 | {data.get("version", "V1.0")} |\n')
        f.write(f'| 体系结构 | {data.get("architecture", "待补充")} |\n')
        f.write(f'| 技术框架 | {data.get("tech_framework", "待补充")} |\n')
        f.write(f'| 开发语言 | {data.get("software_language", "待补充")} |\n\n')

        # 提交材料清单
        f.write('## 提交材料清单\n\n')
        f.write('| 序号 | 材料名称 | 份数/页数要求 |\n')
        f.write('|------|----------|---------------|\n')
        f.write('| 1 | 软件著作权登记申请表 | 1份 |\n')
        f.write('| 2 | 源程序（前30页+后30页） | 不少于60页，每页不少于50行 |\n')
        f.write('| 3 | 软件说明书 | 不少于10页 |\n')
        f.write('| 4 | 著作权人身份证明 | 1份 |\n\n')

        # 源代码格式要求
        f.write('## 源代码提交格式要求\n\n')
        f.write('1. 提交源程序的前30页和后30页\n')
        f.write('2. 不足60页的应提交全部源代码\n')
        f.write('3. 每页不少于50行（不含空行）\n')
        f.write('4. 使用A4纸打印，黑色字体\n')
        f.write('5. 每页页眉应包含软件名称和版本号\n\n')

        # 注意事项
        f.write('## 注意事项\n\n')
        f.write('1. 所有纸质材料应当使用A4纸打印\n')
        f.write('2. 电子文件应当与纸质材料内容一致\n')
        f.write('3. 说明书应当详细描述软件功能和技术特点\n')
        f.write('4. 实际提交请以版权保护中心最新要求为准\n')

    return filepath


def generate_all_documents(data: Dict[str, Any], output_dir: str) -> Dict[str, str]:
    """
    生成所有软件著作权相关文档

    Args:
        data: 包含所有表单数据的字典
        output_dir: 输出目录

    Returns:
        包含所有生成文件路径的字典
    """
    os.makedirs(output_dir, exist_ok=True)

    results = {}

    try:
        results['申请表'] = create_copyright_application_form(data, output_dir)
        print(f"[OK] 申请表已生成: {results['申请表']}")
    except Exception as e:
        print(f"[ERROR] 申请表生成失败: {e}")

    try:
        results['说明书'] = create_software_manual(data, output_dir)
        print(f"[OK] 说明书已生成: {results['说明书']}")
    except Exception as e:
        print(f"[ERROR] 说明书生成失败: {e}")

    try:
        results['资料清单'] = create_tech_document(data, output_dir)
        print(f"[OK] 资料清单已生成: {results['资料清单']}")
    except Exception as e:
        print(f"[ERROR] 资料清单生成失败: {e}")

    return results


if __name__ == '__main__':
    # 简化后的示例数据
    sample_data = {
        'software_full_name': '智能排课管理系统',
        'software_short_name': '排课系统',
        'version': 'V1.0',
        'rights_acquisition': '原始取得',
        # 硬件环境
        'hardware_cpu': 'Intel Core i5 2.0GHz 或以上处理器',
        'hardware_memory': '不低于 8GB RAM',
        'hardware_storage': '不低于 20GB 可用磁盘空间',
        # 软件环境
        'software_os': 'Windows Server 2016 / Linux CentOS 7.x',
        'software_language': 'Java 1.8',
        'software_database': 'MySQL 8.0',
        'software_middleware': 'Redis 6.0, Nginx 1.20',
        # 技术特点
        'architecture': 'B/S（浏览器/服务器）架构',
        'tech_framework': 'Spring Boot + MyBatis + Vue.js',
        'innovation_points': '采用贪心算法与约束满足问题（CSP）相结合的智能排课引擎，支持多维度约束条件自动处理',
        'main_functions': '用户管理、课程管理、智能排课、手动调整、课表导出、报表统计',
        'overview': '一款高效、智能的排课管理系统，支持多维度约束条件的自动排课功能',
        'operation_guide': '1. 系统登录 → 2. 基础数据设置 → 3. 课程班级配置 → 4. 执行智能排课 → 5. 手动调整 → 6. 课表导出',
        'completion_date': datetime.now().strftime('%Y-%m-%d'),
        'publication_status': '未发表'
    }

    # 生成文档
    output = 'output'
    results = generate_all_documents(sample_data, output)

    print("\n=== 生成完成 ===")
    for name, path in results.items():
        print(f'{name}: {path}')

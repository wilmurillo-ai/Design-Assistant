#!/usr/bin/env python3
"""
Word Document Generator
Generates .docx files from JSON configuration
"""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def apply_font_formatting(run, font_config):
    """Apply font formatting to a run"""
    if not font_config:
        return
    
    if 'name' in font_config:
        run.font.name = font_config['name']
    if 'size' in font_config:
        run.font.size = Pt(font_config['size'])
    if 'color' in font_config:
        r, g, b = hex_to_rgb(font_config['color'])
        run.font.color.rgb = RGBColor(r, g, b)
    if font_config.get('bold'):
        run.font.bold = True
    if font_config.get('italic'):
        run.font.italic = True
    if font_config.get('underline'):
        run.font.underline = True


def set_cell_background(cell, color):
    """Set table cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color.lstrip('#'))
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_heading(doc, config):
    """Add a heading"""
    level = config.get('level', 1)
    heading = doc.add_heading(config['text'], level=level)
    if 'font' in config:
        for run in heading.runs:
            apply_font_formatting(run, config['font'])


def add_paragraph(doc, config):
    """Add a paragraph"""
    para = doc.add_paragraph(config['text'])
    
    # Apply alignment
    alignment = config.get('alignment', 'left')
    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Apply font formatting
    if 'font' in config:
        for run in para.runs:
            apply_font_formatting(run, config['font'])


def add_table(doc, config):
    """Add a table"""
    data = config['data']
    rows = len(data)
    cols = len(data[0]) if rows > 0 else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = config.get('style', 'Table Grid')
    
    # Fill table data
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            
            # Apply cell formatting
            if 'cell_formatting' in config:
                cell_format = config['cell_formatting'].get(f'{i},{j}', {})
                if 'background' in cell_format:
                    set_cell_background(cell, cell_format['background'])
                if 'font' in cell_format:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            apply_font_formatting(run, cell_format['font'])


def add_image(doc, config):
    """Add an image"""
    image_path = config['path']
    width = config.get('width')
    height = config.get('height')
    
    if width and height:
        doc.add_picture(image_path, width=Inches(width), height=Inches(height))
    elif width:
        doc.add_picture(image_path, width=Inches(width))
    elif height:
        doc.add_picture(image_path, height=Inches(height))
    else:
        doc.add_picture(image_path)


def add_list(doc, config):
    """Add a bulleted or numbered list"""
    items = config['items']
    list_type = config.get('type', 'bullet')
    
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet' if list_type == 'bullet' else 'List Number')
        if 'font' in config:
            for run in para.runs:
                apply_font_formatting(run, config['font'])


def setup_page(doc, page_setup):
    """Configure page setup"""
    section = doc.sections[0]
    
    if 'margins' in page_setup:
        margins = page_setup['margins']
        if 'top' in margins:
            section.top_margin = Inches(margins['top'])
        if 'bottom' in margins:
            section.bottom_margin = Inches(margins['bottom'])
        if 'left' in margins:
            section.left_margin = Inches(margins['left'])
        if 'right' in margins:
            section.right_margin = Inches(margins['right'])
    
    if 'orientation' in page_setup:
        from docx.enum.section import WD_ORIENT
        if page_setup['orientation'] == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width


def generate_word_document(config, output_path):
    """Generate Word document from configuration"""
    # Create or load document
    if 'template' in config:
        doc = Document(config['template'])
    else:
        doc = Document()
    
    # Setup page
    if 'page_setup' in config:
        setup_page(doc, config['page_setup'])
    
    # Add content
    for item in config.get('content', []):
        item_type = item['type']
        
        if item_type == 'heading':
            add_heading(doc, item)
        elif item_type == 'paragraph':
            add_paragraph(doc, item)
        elif item_type == 'table':
            add_table(doc, item)
        elif item_type == 'image':
            add_image(doc, item)
        elif item_type == 'list':
            add_list(doc, item)
        elif item_type == 'page_break':
            doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    print(f"Word document generated: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_word.py <config.json> <output.docx>")
        sys.exit(1)
    
    config_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        generate_word_document(config, output_path)
    except Exception as e:
        print(f"Error generating Word document: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()

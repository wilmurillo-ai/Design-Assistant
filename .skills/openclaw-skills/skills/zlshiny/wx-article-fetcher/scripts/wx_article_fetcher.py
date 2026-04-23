#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
微信公众号文章抓取工具 - 简化版

功能：
- 根据 biz 查询公众号文章
- 支持时间范围筛选
- 保存为 JSON 和 Markdown 格式
- 按创建时间组织文件结构

作者：十三香 (agent 管理者)
创建时间：2026-03-13
"""

import os
import re
import json
import time
import base64
import html
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

# 尝试导入可选依赖
try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False
    print("⚠️ 缺少 requests 库，请安装：pip3 install requests")

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    print("⚠️ 缺少 beautifulsoup4 库，请安装：pip3 install beautifulsoup4")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ 缺少 python-docx 库，请安装：pip3 install python-docx")

# ==================== 配置 ====================

SIGN_ENV_VAR = "WX_QUERY_SIGN"
# 安全隔离：每个 agent 只能访问自己工作空间的数据
def get_sign_config_path():
    """获取签名配置文件路径（根据调用 agent 隔离）"""
    import os
    agent_name = os.environ.get("OPENCLAW_AGENT_NAME", "unknown")
    
    # baseagent (agent 管理者) 可以访问全局配置
    if agent_name == "baseagent":
        return Path.home() / ".wx_biz_query" / "config.enc"
    
    # 其他 agent 只能访问自己工作空间的配置
    workspace = os.environ.get("OPENCLAW_WORKSPACE", str(Path.home() / ".openclaw" / "workspace" / agent_name))
    config_dir = Path(workspace) / ".wx_config"
    config_dir.mkdir(parents=True, exist_ok=True)
    return config_dir / "wx_sign.enc"

def get_articles_base_dir():
    """获取文章保存目录（根据调用 agent 隔离）"""
    import os
    agent_name = os.environ.get("OPENCLAW_AGENT_NAME", "unknown")
    
    # baseagent (agent 管理者) 保存到全局目录
    if agent_name == "baseagent":
        return Path.home() / ".wx_articles"
    
    # 其他 agent 只能保存到自己工作空间
    workspace = os.environ.get("OPENCLAW_WORKSPACE", str(Path.home() / ".openclaw" / "workspace" / agent_name))
    articles_dir = Path(workspace) / ".wx_articles"
    articles_dir.mkdir(parents=True, exist_ok=True)
    return articles_dir

def get_biz_cache_path():
    """获取 biz 缓存路径（根据调用 agent 隔离）"""
    import os
    agent_name = os.environ.get("OPENCLAW_AGENT_NAME", "unknown")
    
    if agent_name == "baseagent":
        return Path.home() / ".wx_biz_query" / "cache.json"
    
    workspace = os.environ.get("OPENCLAW_WORKSPACE", str(Path.home() / ".openclaw" / "workspace" / agent_name))
    cache_dir = Path(workspace) / ".wx_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir / "wx_biz_cache.json"

SIGN_CONFIG_FILE = get_sign_config_path()
ARTICLES_BASE_DIR = get_articles_base_dir()

API_ARTICLE_URL = "https://union-api.licaimofang.com/api/jinrong/get_article"

# ==================== 安全函数 ====================

def get_sign() -> str:
    """获取签名"""
    sign = os.environ.get(SIGN_ENV_VAR)
    if sign:
        return sign
    
    if SIGN_CONFIG_FILE.exists():
        try:
            with open(SIGN_CONFIG_FILE, 'r', encoding='utf-8') as f:
                return base64.b64decode(f.read().strip()).decode('utf-8')
        except:
            pass
    
    print(f"🔐 未找到签名配置，请输入：")
    return input("签名：").strip()

def load_biz_cache() -> dict:
    """加载 biz 缓存"""
    cache_file = get_biz_cache_path()
    if not cache_file.exists():
        return {}
    try:
        with open(cache_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {}

# ==================== 文章查询 ====================

def query_articles(biz: str, page: int = 1, limit: int = 10, 
                   start_time: str = None, end_time: str = None) -> dict:
    """查询公众号文章"""
    if not HAS_REQUESTS:
        return {"success": False, "error": "缺少 requests 库"}
    
    sign = get_sign()
    
    params = {
        "biz": biz,
        "page": page,
        "limit": min(limit, 10),
        "sign": sign
    }
    
    if start_time:
        params["start_time"] = start_time
    if end_time:
        params["end_time"] = end_time
    
    print(f"🔍 查询文章：biz={biz[:20]}... page={page}")
    if start_time:
        print(f"   时间范围：{start_time} ~ {end_time or '至今'}")
    
    try:
        response = requests.get(API_ARTICLE_URL, params=params, timeout=30)
        response.raise_for_status()
        result = response.json()
        
        if result.get("code") == "000000" or result.get("code") == 0:
            result_data = result.get("result", {})
            # 兼容两种返回格式：直接数组或分页对象
            if isinstance(result_data, list):
                articles = result_data
                total = len(result_data)
            else:
                articles = result_data.get("list", []) or result_data.get("data", [])
                total = result_data.get("total", len(articles))
            
            print(f"✅ 查询成功，获取到 {len(articles)} 篇文章 (共{total}篇)")
            
            return {
                "success": True,
                "articles": articles,
                "total": total,
                "page": page,
                "has_more": len(articles) >= limit
            }
        else:
            return {"success": False, "error": result.get("message", "未知错误")}
    
    except Exception as e:
        return {"success": False, "error": str(e)}

def fetch_all_articles(biz: str, start_time: str = None, end_time: str = None, 
                       max_pages: int = None) -> list:
    """获取所有文章（自动分页）"""
    all_articles = []
    page = 1
    
    while True:
        result = query_articles(biz, page=page, limit=10, 
                               start_time=start_time, end_time=end_time)
        
        if not result.get("success"):
            print(f"❌ {result.get('error')}")
            break
        
        articles = result.get("articles", [])
        if not articles:
            print("⚠️ 没有更多文章了")
            break
        
        all_articles.extend(articles)
        print(f"   累计：{len(all_articles)} 篇 / 总共：{result.get('total', '?')} 篇")
        
        if not result.get("has_more"):
            break
        
        if max_pages and page >= max_pages:
            print(f"⚠️ 已达到最大页数限制：{max_pages}")
            break
        
        page += 1
        time.sleep(0.5)
    
    return all_articles

# ==================== HTML 处理 ====================

def html_to_markdown(html_content: str) -> str:
    """HTML 转 Markdown（简化版）"""
    if not html_content or not HAS_BS4:
        if not HAS_BS4:
            return html.unescape(re.sub(r'<[^>]+>', '', html_content or ''))
        return ""
    
    soup = BeautifulSoup(html_content, 'lxml')
    markdown = []
    
    # 处理段落
    for p in soup.find_all(['p', 'div', 'section']):
        text = p.get_text(strip=True)
        if text:
            markdown.append(text)
    
    # 处理标题
    for i, h in enumerate(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])):
        text = h.get_text(strip=True)
        if text:
            level = min(int(h.name[1]), 3)
            markdown.insert(0, f"{'#' * level} {text}")
    
    # 处理图片
    img_count = 0
    for img in soup.find_all('img'):
        src = img.get('data-src') or img.get('src')
        if src:
            img_count += 1
            markdown.append(f"\n![图片 {img_count}]({src})\n")
    
    return '\n\n'.join(markdown)

def save_image(image_url: str, images_dir: str) -> str:
    """保存图片到本地"""
    if not image_url or not HAS_REQUESTS:
        return None
    
    try:
        Path(images_dir).mkdir(parents=True, exist_ok=True)
        
        img_hash = hashlib.md5(image_url.encode()).hexdigest()[:12]
        ext = 'jpg'
        if 'png' in image_url.lower():
            ext = 'png'
        elif 'gif' in image_url.lower():
            ext = 'gif'
        
        img_name = f"img_{img_hash}.{ext}"
        img_path = Path(images_dir) / img_name
        
        if img_path.exists():
            return f"images/{img_name}"
        
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(image_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            with open(img_path, 'wb') as f:
                f.write(response.content)
            return f"images/{img_name}"
        return None
    
    except Exception as e:
        print(f"   ⚠️ 图片保存失败：{e}")
        return None

# ==================== 文件保存 ====================

def save_article_json(article: dict, output_path: str):
    """保存为 JSON"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(article, f, ensure_ascii=False, indent=2)
    print(f"   ✅ JSON: {output_path}")

def save_article_markdown(article: dict, output_path: str, images_dir: str):
    """保存为 Markdown"""
    md_content = []
    
    # 标题
    md_content.append(f"# {article.get('title', '无标题')}\n")
    
    # 元信息
    md_content.append("## 文章信息\n")
    md_content.append(f"- **发布时间**: {article.get('create_time', '未知')}")
    md_content.append(f"- **阅读量**: {article.get('send_to_fans_num', 0)}")
    md_content.append(f"- **点赞**: {article.get('zan', 0)}")
    md_content.append(f"- **分享**: {article.get('share_num', 0)}")
    md_content.append(f"- **收藏**: {article.get('collect_num', 0)}")
    md_content.append(f"- **评论**: {article.get('comment_count', 0)}\n")
    
    # 正文
    md_content.append("## 正文\n")
    content_html = article.get('content', '')
    markdown_text = html_to_markdown(content_html)
    md_content.append(markdown_text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))
    
    print(f"   ✅ MD: {output_path}")

def save_article_word(article: dict, output_path: str, images_dir: str):
    """保存为 Word"""
    if not HAS_DOCX:
        print(f"   ⚠️ 跳过 Word（缺少 python-docx 库）")
        return
    
    try:
        doc = Document()
        
        # 标题
        title = article.get('title', '无标题')
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 元信息
        doc.add_paragraph()
        meta = doc.add_paragraph()
        meta.add_run(f"发布时间：{article.get('create_time', '未知')}\n")
        meta.add_run(f"阅读量：{article.get('send_to_fans_num', 0)}\n")
        meta.add_run(f"点赞：{article.get('zan', 0)} | 分享：{article.get('share_num', 0)} | 收藏：{article.get('collect_num', 0)} | 评论：{article.get('comment_count', 0)}")
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_page_break()
        
        # 正文（简化处理）
        content_html = article.get('content', '')
        if content_html and HAS_BS4:
            soup = BeautifulSoup(content_html, 'lxml')
            text = soup.get_text(separator='\n', strip=True)
            doc.add_paragraph(text)
        
        doc.save(output_path)
        print(f"   ✅ Word: {output_path}")
    
    except Exception as e:
        print(f"   ⚠️ Word 保存失败：{e}")

# ==================== 主流程 ====================

def save_articles(biz: str, nickname: str = None, 
                  start_time: str = None, end_time: str = None,
                  max_pages: int = None, 
                  save_word: bool = False, save_md: bool = True, save_json: bool = True):
    """保存文章到本地"""
    import hashlib  # 延迟导入
    
    # 获取公众号名称
    if not nickname:
        biz_cache = load_biz_cache()
        for name, data in biz_cache.items():
            if data.get('biz') == biz:
                nickname = name
                break
        if not nickname:
            nickname = f"公众号_{biz[:10]}"
    
    # 创建保存目录
    if start_time and end_time:
        dir_name = f"{nickname}_{start_time}_{end_time}"
    else:
        dir_name = f"{nickname}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    base_dir = ARTICLES_BASE_DIR / dir_name
    base_dir.mkdir(parents=True, exist_ok=True)
    images_dir = base_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"📁 保存目录：{base_dir}")
    print()
    
    # 获取所有文章
    articles = fetch_all_articles(biz, start_time, end_time, max_pages)
    
    if not articles:
        print("\n❌ 没有获取到文章")
        return
    
    print(f"\n📊 共获取 {len(articles)} 篇文章，开始保存...\n")
    
    # 保存每篇文章
    for i, article in enumerate(articles, 1):
        title = article.get('title', '无标题')
        create_time = article.get('create_time', '')
        
        # 清理标题
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)[:50]
        
        # 按时间组织
        if create_time:
            date_dir = create_time[:10]
            article_dir = base_dir / date_dir
            article_dir.mkdir(parents=True, exist_ok=True)
        else:
            article_dir = base_dir
        
        file_base = article_dir / f"{i:03d}_{safe_title}"
        
        print(f"[{i}/{len(articles)}] {title[:40]}...")
        
        # 保存 JSON
        if save_json:
            save_article_json(article, f"{file_base}.json")
        
        # 保存 Markdown
        if save_md:
            save_article_markdown(article, f"{file_base}.md", str(images_dir))
        
        # 保存 Word
        if save_word:
            save_article_word(article, f"{file_base}.docx", str(images_dir))
        
        print()
    
    print("=" * 60)
    print(f"✅ 完成！共保存 {len(articles)} 篇文章")
    print(f"📁 保存位置：{base_dir}")
    print("=" * 60)

# ==================== 交互模式 ====================

def interactive_mode():
    """交互模式"""
    print("=" * 60)
    print("📰 微信公众号文章抓取工具")
    print("=" * 60)
    
    biz = input("\n请输入公众号 biz: ").strip()
    if not biz:
        biz_cache = load_biz_cache()
        if biz_cache:
            print("\n📦 已缓存的公众号:")
            for i, (name, data) in enumerate(biz_cache.items(), 1):
                print(f"  {i}. {name} (biz: {data.get('biz', '')[:20]}...)")
            choice = input("\n选择编号：").strip()
            if choice.isdigit() and 1 <= int(choice) <= len(biz_cache):
                biz = list(biz_cache.values())[int(choice) - 1].get('biz')
            else:
                print("❌ 无效选择")
                return
        else:
            print("❌ 未找到 biz")
            return
    
    print("\n时间范围（直接跳过表示不限制）:")
    start_time = input("开始时间 (YYYY-MM-DD): ").strip() or None
    end_time = input("截止时间 (YYYY-MM-DD): ").strip() or None
    
    max_pages_input = input("\n最大页数（直接跳过表示全部）: ").strip()
    max_pages = int(max_pages_input) if max_pages_input.isdigit() else None
    
    print("\n保存格式:")
    save_json = input("保存 JSON? (y/n): ").strip().lower() != 'n'
    save_md = input("保存 Markdown? (y/n): ").strip().lower() != 'n'
    save_word = input("保存 Word? (y/n): ").strip().lower() != 'n'
    
    print()
    save_articles(biz, start_time=start_time, end_time=end_time, 
                  max_pages=max_pages, save_word=save_word, 
                  save_md=save_md, save_json=save_json)

# ==================== 主程序 ====================

if __name__ == "__main__":
    import sys
    import hashlib
    
    if not HAS_REQUESTS:
        print("\n❌ 缺少必要的依赖库 requests")
        print("请运行：pip3 install --break-system-packages requests\n")
        sys.exit(1)
    
    if len(sys.argv) > 1:
        biz = sys.argv[1]
        start_time = sys.argv[2] if len(sys.argv) > 2 else None
        end_time = sys.argv[3] if len(sys.argv) > 3 else None
        max_pages = sys.argv[4] if len(sys.argv) > 4 and sys.argv[4].isdigit() else None
        save_articles(biz, start_time=start_time, end_time=end_time, max_pages=max_pages)
    else:
        interactive_mode()

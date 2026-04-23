"""
parser.py - 输入解析模块
支持：txt/md 直接粘贴、docx/pdf 文件读取
"""
import os
import sys
from pathlib import Path


def parse_input(path: str) -> str:
    """
    根据文件扩展名解析输入文件，返回纯文本内容
    
    支持格式：.txt, .md, .docx, .pdf
    """
    p = Path(path)
    suffix = p.suffix.lower()
    
    if not p.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    
    if suffix in ('.txt', '.md'):
        return parse_text(path)
    elif suffix == '.docx':
        return parse_docx(path)
    elif suffix == '.pdf':
        return parse_pdf(path)
    else:
        # 尝试作为纯文本读取
        try:
            return parse_text(path)
        except Exception:
            raise ValueError(f"不支持的文件格式: {suffix}，支持: txt/md/docx/pdf")


def parse_text(path: str) -> str:
    """解析纯文本或 Markdown 文件"""
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_docx(path: str) -> str:
    """解析 Word .docx 文件"""
    try:
        import docx
except ImportError:
    raise ImportError("解析 .docx 需要 python-docx，请运行: pip install python-docx")
    
    doc = docx.Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 尝试读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    
    return '\n'.join(paragraphs)


def parse_pdf(path: str) -> str:
    """解析 PDF 文件"""
    try:
        import pdfplumber
except ImportError:
    raise ImportError("解析 .pdf 需要 pdfplumber，请运行: pip install pdfplumber")
    
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    return '\n\n'.join(text_parts)


def extract_structure_from_text(text: str) -> dict:
    """
    从文本中提取关键结构信息（辅助 LLM）
    - 检测是否存在流程关键词（如果...则...、首先...然后...）
    - 检测架构关键词（层、模块、组件）
    - 检测时序关键词（发起、响应、调用）
    """
    keywords = {
        'flow': ['流程', '步骤', '首先', '然后', '接着', '最后', '如果', '则', '判断', '分支'],
        'sequence': ['发起', '响应', '调用', '返回', '消息', '时序', '交互'],
        'architecture': ['架构', '模块', '组件', '系统', '层次', '分层'],
        'network': ['网络', '拓扑', '服务器', '客户端', '设备', '连接', '网段'],
        'deployment': ['部署', '实例', '节点', '集群', '容器', '云服务'],
    }
    
    text_lower = text.lower()
    scores = {}
    
    for chart_type, type_keywords in keywords.items():
        score = sum(1 for kw in type_keywords if kw.lower() in text_lower)
        scores[chart_type] = score
    
    # 返回得分最高的类型
    if max(scores.values()) > 0:
        return {'detected_type': max(scores, key=scores.get), 'scores': scores}
    
    return {'detected_type': 'flowchart', 'scores': scores}

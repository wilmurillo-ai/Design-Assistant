#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
投标文件解析脚本 - 支持 txt, pdf, docx, xlsx 格式
版本：v2.0 (2026-03-21)

## v2.0 优化点
1. PDF解析改进：增加扫描版PDF检测和提示
2. 更完善的错误处理
3. 虚拟环境兼容

## 支持格式
- .txt 纯文本
- .docx Word文档
- .pdf PDF文件（文本型/扫描版）
- .xlsx Excel表格

## PDF处理说明
- 文本型PDF：直接提取文字
- 扫描版PDF（无文字层）：提取失败，提示用户人工处理
- 建议：收到PDF后优先检查是否为扫描版
"""

import sys
import os

def parse_txt(file_path):
    """解析Txt文件"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def parse_docx(file_path):
    """解析Word文档"""
    try:
        from docx import Document
    except ImportError:
        return "[错误] 需要安装 python-docx: pip install python-docx"
    
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                paragraphs.append(f'[表格] {row_text}')
    
    return '\n'.join(paragraphs)

def check_pdf_is_scanned(file_path):
    """
    检测PDF是否为扫描版
    返回: (is_scanned: bool, message: str)
    """
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) == 0:
                return True, "[警告] PDF页数为0，无法解析"
            
            # 尝试提取第一页文字
            first_page_text = ""
            for page in pdf.pages[:3]:  # 检查前3页
                text = page.extract_text()
                if text:
                    first_page_text = text
                    break
            
            if not first_page_text or len(first_page_text.strip()) < 50:
                return True, "[警告] PDF可能为扫描版（检测到文字层缺失或极少）"
            
            return False, "[信息] PDF为文本型，可正常解析"
    except Exception as e:
        return True, f"[错误] PDF解析异常: {str(e)}"

def parse_pdf(file_path):
    """
    解析PDF文件
    支持：文本型PDF直接提取
    提示：扫描版PDF需人工后处理
    """
    # 首先检测是否为扫描版
    is_scanned, check_msg = check_pdf_is_scanned(file_path)
    print(check_msg)
    
    if is_scanned:
        # 尝试备用方案
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        text_parts.append(text)
            
            if text_parts:
                return "[提示] 使用PyPDF2提取到部分文字，可能不完整：\n\n" + '\n\n'.join(text_parts[:5])
            else:
                return """[错误] PDF为扫描版，无法自动提取文字

处理建议：
1. 使用ABBYY FineReader、Adobe Acrobat等OCR工具进行文字识别
2. 或手动复制PDF中的文字内容为txt文件
3. 识别后的文本可重新提交解析

推荐工具：
- Adobe Acrobat: 菜单 → 工具 → 扫描和OCR → 识别文本
- 在线OCR: https://www.ilovepdf.com/ocr
- ABBYY FineReader（专业级，精度最高）
"""
        except ImportError:
            return "[错误] 需要安装 PyPDF2: pip install PyPDF2"
    
    # 文本型PDF，使用pdfplumber
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return '\n\n'.join(text_parts)
    except ImportError:
        return "[错误] 需要安装 pdfplumber: pip install pdfplumber"

def parse_xlsx(file_path):
    """解析Excel文件"""
    try:
        import openpyxl
    except ImportError:
        return "[错误] 需要安装 openpyxl: pip install openpyxl"
    
    text_parts = []
    wb = openpyxl.load_workbook(file_path, data_only=True)
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f'=== Sheet: {sheet_name} ===')
        for row in sheet.iter_rows(values_only=True):
            row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
            if row_text.strip():
                text_parts.append(row_text)
    return '\n'.join(text_parts)

def parse_file(file_path):
    """根据文件扩展名解析文件"""
    if not os.path.exists(file_path):
        return f"[错误] 文件不存在: {file_path}"
    
    ext = os.path.splitext(file_path)[1].lower()
    
    parsers = {
        '.txt': parse_txt,
        '.docx': parse_docx,
        '.pdf': parse_pdf,
        '.xlsx': parse_xlsx,
        '.xls': parse_xlsx,
    }
    
    parser = parsers.get(ext)
    if not parser:
        supported = ', '.join(parsers.keys())
        return f"[错误] 不支持的文件格式: {ext}\n支持的格式: {supported}"
    
    try:
        return parser(file_path)
    except Exception as e:
        return f"[错误] 解析失败: {str(e)}\n请检查文件是否损坏或加密"

def main():
    if len(sys.argv) < 2:
        print('用法: python3 parse_bid_files.py <文件路径>')
        print('支持格式: txt, docx, pdf, xlsx, xls')
        print('')
        print('PDF处理说明:')
        print('  - 文本型PDF: 直接提取文字')
        print('  - 扫描版PDF: 提示用户使用OCR工具')
        sys.exit(1)
    
    file_path = sys.argv[1]
    print(f"正在解析: {file_path}")
    print("-" * 50)
    
    content = parse_file(file_path)
    print(content)

if __name__ == '__main__':
    main()
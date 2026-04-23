#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word技术标文档 - 转换脚本
版本：v5.2 (2026-03-27)

## v4.0 新增功能
1. 支持 Markdown 顶部格式指令动态配置字体/字号/行距/页边距
2. AI 生成时自动按格式指令生成对应样式的 Word

## 格式指令说明
在 Markdown 顶部加入以下格式块即可自定义样式：
```markdown
<!-- doc-format
font: SimSun
body-size: 16pt
title-level: 36pt
sub-level: 32pt
line-spacing: 26pt
margins: 2cm
first-line-indent: 0.74cm
-->
```

### 可用参数
| 参数 | 默认值 | 说明 |
|------|--------|------|
| font | SimSun | 正文字体（英文名） |
| body-size | 32pt | 正文字号（三号=16pt） |
| title-level | 36pt | 一级标题字号 |
| sub-level | 32pt | 二级标题字号 |
| line-spacing | 26pt | 行距 |
| margins | 2cm | 页边距 |
| first-line-indent | 0.74cm | 首行缩进 |
| page-width | 21cm | 页面宽度（A4=21cm） |
| page-height | 29.7cm | 页面高度（A4=29.7cm） |

使用方法：
    python3 convert_to_word.py <输入.md> <输出.docx>
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import re

# 默认格式参数
DEFAULT_FORMAT = {
    'font': 'SimSun',
    'body-size': 16,      # pt
    'title-size': 36,    # pt (一级标题)
    'sub-size': 32,      # pt (二级标题)
    'line-spacing': 26,   # pt
    'margins': 2,        # cm
    'first-line-indent': 0.74,  # cm
    'page-width': 21,    # cm
    'page-height': 29.7, # cm
}

def parse_pt(value):
    """解析 pt 值，返回浮点数"""
    s = str(value).strip()
    if s.endswith('pt'):
        s = s[:-2]
    try:
        return float(s)
    except:
        return None

def clean_heading_text(text):
    """清理标题中的描述性括号，如（深入说明）、（深入扩展）等"""
    # 匹配全角括号（）或半角括号() 中的描述性词汇
    patterns = [
        r'（深入[^）]*）',      # （深入xxx）
        r'（详细[^）]*）',      # （详细xxx）
        r'（扩展[^）]*）',      # （扩展xxx）
        r'（进一步[^）]*）',    # （进一步xxx）
        r'（补充[^）]*）',      # （补充xxx）
        r'（说明[^）]*）',      # （说明xxx）
        r'\(深入[^)]*\)',      # (深入xxx)
        r'\(详细[^)]*\)',      # (详细xxx)
        r'\(扩展[^)]*\)',      # (扩展xxx)
    ]
    result = text.strip()
    for pattern in patterns:
        result = re.sub(pattern, '', result)
    # 清理多余空格
    result = re.sub(r'\s+', ' ', result).strip()
    return result

def parse_cm(value):
    """解析 cm 值，返回浮点数"""
    s = str(value).strip()
    if s.endswith('cm'):
        s = s[:-2]
    try:
        return float(s)
    except:
        return None

def parse_format_block(content):
    """从 Markdown 内容中解析 doc-format 格式指令"""
    fmt = DEFAULT_FORMAT.copy()
    
    # 匹配 <!-- doc-format ... -->
    pattern = r'<!--\s*doc-format\s*([\s\S]*?)-->'
    match = re.search(pattern, content)
    if not match:
        return fmt, content
    
    block = match.group(1)
    # 去掉原内容中的格式块
    content = re.sub(pattern, '', content, count=1).strip()
    
    # 解析各行
    for line in block.split('\n'):
        line = line.strip()
        if ':' not in line:
            continue
        key, value = line.split(':', 1)
        key = key.strip().lower()
        value = value.strip()
        
        if key == 'font':
            fmt['font'] = value
        elif key == 'body-size':
            v = parse_pt(value)
            if v: fmt['body-size'] = v
        elif key == 'title-level':
            v = parse_pt(value)
            if v: fmt['title-size'] = v
        elif key == 'sub-level':
            v = parse_pt(value)
            if v: fmt['sub-size'] = v
        elif key == 'line-spacing':
            v = parse_pt(value)
            if v: fmt['line-spacing'] = v
        elif key == 'margins':
            v = parse_cm(value)
            if v: fmt['margins'] = v
        elif key == 'first-line-indent':
            v = parse_cm(value)
            if v: fmt['first-line-indent'] = v
        elif key == 'page-width':
            v = parse_cm(value)
            if v: fmt['page-width'] = v
        elif key == 'page-height':
            v = parse_cm(value)
            if v: fmt['page-height'] = v
    
    return fmt, content

def make_fontsetter(fmt):
    """根据格式参数生成 set_font 函数"""
    font_name = fmt['font']
    body_size = fmt['body-size']
    
    def set_font(run):
        """设置字体"""
        run.font.name = font_name
        run.font.size = Pt(body_size)
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    
    return set_font

def set_cell_shading(cell, fill_color="FFFFFF"):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)

def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def parse_markdown_table(line, lines, start_idx):
    """解析Markdown表格"""
    headers = [h.strip() for h in line.split('|') if h.strip()]
    
    align_line = lines[start_idx + 1] if start_idx + 1 < len(lines) else ""
    alignments = []
    for cell in align_line.split('|'):
        cell = cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            alignments.append('center')
        elif cell.endswith(':'):
            alignments.append('right')
        else:
            alignments.append('left')
    
    data_start = start_idx + 2
    rows_data = []
    for i in range(data_start, len(lines)):
        l = lines[i].strip()
        if not l or not l.startswith('|'):
            return headers, alignments, rows_data, i
        row = [c.strip() for c in l.split('|') if c.strip()]
        rows_data.append(row)
    
    return headers, alignments, rows_data, len(lines)

def add_table(doc, headers, alignments, rows, set_font_fn):
    """添加表格"""
    if not headers:
        return None
    
    col_count = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=col_count)
    set_table_border(table)
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_font_fn(run)
        run.font.bold = True
        set_cell_shading(cell, "FFFFFF")
        
        align = alignments[i] if i < len(alignments) else 'left'
        if align == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx >= col_count:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            set_font_fn(run)
            
            align = alignments[c_idx] if c_idx < len(alignments) else 'left'
            if align == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return table

def convert_md_to_word(input_md_path, output_docx_path):
    """将Markdown文件转换为Word文档"""
    
    if not os.path.exists(input_md_path):
        print(f'❌ 错误：输入文件不存在：{input_md_path}')
        return False
    
    with open(input_md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析格式指令
    fmt, content = parse_format_block(content)
    
    print(f'📋 格式参数：字体={fmt["font"]} | 正文={fmt["body-size"]}pt | 行距={fmt["line-spacing"]}pt | 边距={fmt["margins"]}cm')
    
    doc = Document()
    
    # 设置页面
    for section in doc.sections:
        section.top_margin = Cm(fmt['margins'])
        section.bottom_margin = Cm(fmt['margins'])
        section.left_margin = Cm(fmt['margins'])
        section.right_margin = Cm(fmt['margins'])
        section.page_height = Cm(fmt['page-height'])
        section.page_width = Cm(fmt['page-width'])
    
    set_font = make_fontsetter(fmt)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                set_font(run)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 处理表格
        if line.strip().startswith('|') and '---' not in line:
            headers, alignments, rows_data, end_idx = parse_markdown_table(line, lines, i)
            if headers and rows_data:
                add_table(doc, headers, alignments, rows_data, set_font)
                i = end_idx
                continue
        
        # 处理标题
        if line.startswith('# '):
            para = doc.add_paragraph()
            run = para.add_run(clean_heading_text(line[2:]))
            run.font.size = Pt(fmt['title-size'])
            set_font(run)
            set_heading_style(para, 1)
        elif line.startswith('## '):
            para = doc.add_paragraph()
            run = para.add_run(clean_heading_text(line[3:]))
            run.font.size = Pt(fmt['sub-size'])
            set_font(run)
            set_heading_style(para, 2)
        elif line.startswith('### '):
            para = doc.add_paragraph()
            run = para.add_run(clean_heading_text(line[4:]))
            run.font.size = Pt(fmt['body-size'])
            set_font(run)
            set_heading_style(para, 3)
        elif line.startswith('#### '):
            para = doc.add_paragraph()
            run = para.add_run(clean_heading_text(line[5:]))
            run.font.size = Pt(fmt['body-size'])
            set_font(run)
            set_heading_style(para, 4)
        elif line.strip() == '':
            i += 1
            continue
        elif re.match(r'^[-*_]{3,}$', line.strip()):
            # 跳过 Markdown 水平线（---、***、___）
            i += 1
            continue
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
            set_font(run)
            para.paragraph_format.first_line_indent = Cm(fmt['first-line-indent'])
        
        i += 1
    
    # 设置行距
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = Pt(fmt['line-spacing'])
        para.paragraph_format.line_spacing_rule = 3
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    
    doc.save(output_docx_path)

    # 修复文档元数据（创建时间、修改时间、creator）
    _fix_docx_metadata(output_docx_path)

    print(f'✅ Word文档已生成：{output_docx_path}')
    return True


def _fix_docx_metadata(docx_path):
    """修复 docx 文档的元数据：日期和 creator 暴露问题"""
    import zipfile, os
    from datetime import datetime, timezone

    tmp = docx_path + '.tmp'
    now_utc = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'docProps/core.xml':
                    # 替换写死的 2013 日期
                    data = data.replace(b'2013-12-23T23:15:00Z', now_utc.encode())
                    # 先替换 description（generated by python-docx → 清空），再替换 creator
                    data = data.replace(b'generated by python-docx', b'')
                    data = data.replace(b'python-docx', b'User')
                    # cp:lastModifiedBy 也填上
                    data = data.replace(b'<cp:lastModifiedBy/>', b'<cp:lastModifiedBy>User</cp:lastModifiedBy>')
                zout.writestr(item, data)

    os.replace(tmp, docx_path)

def set_heading_style(para, level):
    """设置标题大纲级别"""
    pPr = para._element.get_or_add_pPr()
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level - 1))
    pPr.append(outlineLvl)

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('用法: python3 convert_to_word.py <输入.md> <输出.docx>')
        print('示例: python3 convert_to_word.py tech_spec.md output.docx')
        sys.exit(1)
    
    success = convert_md_to_word(sys.argv[1], sys.argv[2])
    sys.exit(0 if success else 1)

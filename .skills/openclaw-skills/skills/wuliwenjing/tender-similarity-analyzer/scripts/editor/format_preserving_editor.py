# -*- coding: utf-8 -*-
"""
格式保持文档编辑器
仅修改文字，不改变任何格式
"""

import re
import zipfile
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from lxml import etree

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("python-docx >= 1.0.0 is required")


class FormatPreservingEditor:
    """
    格式保持文档编辑器
    
    核心原理：
    1. 只修改 <w:t> 标签内的文字内容
    2. 完全保留 <w:rPr> 格式属性
    3. 不改变XML层级结构
    """
    
    # XML命名空间
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
    }
    
    def __init__(self, docx_path: str):
        """
        @param docx_path: 原始文档路径
        """
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.original_xml = self._backup_xml()
        self.paragraphs_data = self._extract_paragraphs_data()
        
    def _backup_xml(self) -> str:
        """备份原始document.xml"""
        with zipfile.ZipFile(self.docx_path, 'r') as zf:
            return zf.read('word/document.xml')
            
    def _extract_paragraphs_data(self) -> List[Dict]:
        """
        提取所有段落的数据
        
        @return: 段落列表，每个包含runs信息
        """
        paragraphs_data = []
        
        for para_idx, para in enumerate(self.doc.paragraphs):
            para_info = {
                'index': para_idx,
                'text': para.text,
                'runs': [],
                'style': para.style.name if para.style else 'Normal'
            }
            
            for run_idx, run in enumerate(para.runs):
                run_info = {
                    'index': run_idx,
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                }
                para_info['runs'].append(run_info)
                
            paragraphs_data.append(para_info)
            
        return paragraphs_data
        
    def replace_text_preserving_format(self, old_text: str, new_text: str, 
                                        case_sensitive: bool = False) -> int:
        """
        在整个文档中替换文字，保持所有格式
        
        @param old_text: 要替换的原文字
        @param new_text: 新的文字
        @param case_sensitive: 是否区分大小写
        @return: 替换的次数
        """
        replaced_count = 0
        
        for para in self.doc.paragraphs:
            for run in para.runs:
                run_xml = run._element
                run_text = run.text
                
                # 检查是否匹配
                if case_sensitive:
                    match = old_text in run_text
                else:
                    match = old_text.lower() in run_text.lower()
                    
                if match:
                    # 直接替换<w:t>内容，保持格式
                    for t_elem in run_xml.findall('.//w:t', namespaces=self.namespaces):
                        t_text = t_elem.text or ''
                        
                        if case_sensitive:
                            if old_text in t_text:
                                t_elem.text = t_text.replace(old_text, new_text, 1)
                                replaced_count += 1
                        else:
                            if old_text.lower() in t_text.lower():
                                # 不区分大小写替换
                                import re
                                pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                                t_elem.text = pattern.sub(new_text, t_text, 1)
                                replaced_count += 1
                                
        return replaced_count
        
    def replace_text_in_para_preserving_format(self, para_index: int, 
                                               old_text: str, 
                                               new_text: str,
                                               case_sensitive: bool = False) -> int:
        """
        在指定段落中替换文字
        
        @param para_index: 段落索引
        @param old_text: 要替换的原文字
        @param new_text: 新的文字
        @param case_sensitive: 是否区分大小写
        @return: 替换的次数
        """
        if para_index >= len(self.doc.paragraphs):
            return 0
            
        para = self.doc.paragraphs[para_index]
        replaced_count = 0
        
        for run in para.runs:
            run_xml = run._element
            run_text = run.text
            
            if case_sensitive:
                match = old_text in run_text
            else:
                match = old_text.lower() in run_text.lower()
                
            if match:
                for t_elem in run_xml.findall('.//w:t', namespaces=self.namespaces):
                    t_text = t_elem.text or ''
                    
                    if case_sensitive:
                        if old_text in t_text:
                            t_elem.text = t_text.replace(old_text, new_text, 1)
                            replaced_count += 1
                    else:
                        if old_text.lower() in t_text.lower():
                            import re
                            pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                            t_elem.text = pattern.sub(new_text, t_text, 1)
                            replaced_count += 1
                            
        return replaced_count
        
    def replace_paragraph(self, para_index: int, new_text: str) -> bool:
        """
        替换整个段落（保留段落样式，但会用新文字）
        注意：这个方法会丢失run级别的格式，仅保留段落样式
        
        @param para_index: 段落索引
        @param new_text: 新的段落文字
        @return: 是否成功
        """
        if para_index >= len(self.doc.paragraphs):
            return False
            
        para = self.doc.paragraphs[para_index]
        original_style = para.style
        
        # 记录段落属性
        pPr = para._element.find(qn('w:pPr'))
        pPr_xml = None
        if pPr is not None:
            pPr_xml = etree.tostring(pPr)
            
        # 清空段落内容
        for child in list(para._element):
            if child.tag != qn('w:pPr'):
                para._element.remove(child)
                
        # 添加新文字（保留原有段落样式）
        new_run = para.add_run(new_text)
        
        # 尝试恢复段落样式
        if pPr_xml is not None:
            new_pPr = etree.fromstring(pPr_xml)
            para._element.insert(0, new_pPr)
            
        return True
        
    def get_paragraph_xml(self, para_index: int) -> str:
        """
        获取指定段落的XML
        
        @param para_index: 段落索引
        @return: XML字符串
        """
        if para_index >= len(self.doc.paragraphs):
            return ""
            
        para = self.doc.paragraphs[para_index]
        return etree.tostring(para._element, pretty_print=True).decode('utf-8')
        
    def save(self, output_path: str):
        """
        保存文档
        
        @param output_path: 输出路径
        """
        self.doc.save(output_path)
        
    def verify_format_unchanged(self) -> Tuple[bool, str]:
        """
        验算：检查修改后文档的格式是否保持
        
        @return: (是否一致, 差异信息)
        """
        # 保存当前文档
        temp_path = str(self.docx_path) + '.verify.tmp'
        self.doc.save(temp_path)
        
        try:
            # 读取修改后的XML
            with zipfile.ZipFile(temp_path, 'r') as zf:
                new_xml = zf.read('word/document.xml')
                
            # 比对XML结构（忽略<w:t>内容）
            orig_tree = etree.fromstring(self.original_xml)
            new_tree = etree.fromstring(new_xml)
            
            # 比较所有<w:rPr>元素
            orig_rprs = orig_tree.findall('.//w:rPr', namespaces=self.namespaces)
            new_rprs = new_tree.findall('.//w:rPr', namespaces=self.namespaces)
            
            if len(orig_rprs) != len(new_rprs):
                return False, f"格式元素数量改变: {len(orig_rprs)} -> {len(new_rprs)}"
                
            for i, (orig, new) in enumerate(zip(orig_rprs, new_rprs)):
                orig_str = etree.tostring(orig).decode()
                new_str = etree.tostring(new).decode()
                
                if orig_str != new_str:
                    return False, f"第{i}个格式元素改变:\n原: {orig_str}\n新: {new_str}"
                    
            return True, "格式验证通过"
            
        finally:
            # 删除临时文件
            Path(temp_path).unlink(missing_ok=True)

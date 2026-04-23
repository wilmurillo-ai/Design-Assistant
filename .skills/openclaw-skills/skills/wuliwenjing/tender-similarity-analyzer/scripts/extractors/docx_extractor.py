# -*- coding: utf-8 -*-
"""
Docx文件文本提取器
支持 .docx 和 .doc 格式
"""

import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("python-docx >= 1.0.0 is required. Install: pip install python-docx")


class DocxExtractor:
    """
    Word文档文本提取器
    提取文档中的段落文本，同时保留段落结构
    """
    
    def extract(self, file_path: str) -> str:
        """
        从Word文档中提取纯文本
        
        @param file_path: 文档路径
        @return: 提取的纯文本
        """
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                
        return '\n\n'.join(paragraphs)
        
    def extract_with_metadata(self, file_path: str) -> dict:
        """
        提取文档文本和元数据
        
        @param file_path: 文档路径
        @return: 包含文本和元数据的字典
        """
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': str(para.alignment) if para.alignment else None
                })
                
        # 提取文档属性
        core_props = doc.core_properties
        
        return {
            'text': '\n\n'.join(p['text'] for p in paragraphs),
            'paragraphs': paragraphs,
            'paragraph_count': len(paragraphs),
            'char_count': sum(len(p['text']) for p in paragraphs),
            'title': core_props.title or Path(file_path).stem,
            'author': core_props.author,
            'created': core_props.created,
            'modified': core_props.modified
        }
        
    def extract_runs(self, file_path: str) -> list:
        """
        提取所有runs（带有格式的文本片段）
        
        @param file_path: 文档路径
        @return: runs列表，每个run包含文本和格式信息
        """
        doc = Document(file_path)
        runs = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    runs.append({
                        'para_index': para_idx,
                        'run_index': run_idx,
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                    })
                    
        return runs
        
    def get_full_xml(self, file_path: str) -> str:
        """
        获取完整的document.xml内容
        用于格式保持编辑
        
        @param file_path: 文档路径
        @return: XML字符串
        """
        import zipfile
        from lxml import etree
        
        with zipfile.ZipFile(file_path, 'r') as zf:
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml')
                return etree.tostring(etree.fromstring(xml_content), pretty_print=True).decode('utf-8')
                
        return ""

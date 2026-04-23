#!/usr/bin/env python3
"""
Enterprise Legal Counsel - DOCX Generator

Generates Word (.docx) files with:
- Clean contract drafts (draft mode)
- Reviewed contracts with OOXML track changes and comments (review mode)

Usage:
    python3 docx_generator.py draft input.json output.docx
    python3 docx_generator.py review input.json output.docx

Input JSON format:
{
    "title": "Contract Title",
    "author": "Legal Counsel",
    "date": "2024-01-15",
    "sections": [
        {
            "heading": "Section Title",
            "level": 1,
            "content": [
                {"type": "keep", "text": "Unchanged text"},
                {"type": "delete", "text": "Deleted text", "comment": "Reason"},
                {"type": "insert", "text": "New text", "comment": "Reason"},
                {"type": "replace", "original": "Old", "revised": "New", "comment": "Reason"}
            ]
        }
    ]
}

Dependencies: pip install python-docx lxml
"""

import json
import sys
import os
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# ============================================================
# OOXML Namespace Constants
# ============================================================
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
WORD14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'

# Global revision ID counter
_rev_id = [0]


def next_rev_id():
    """Generate next unique revision ID."""
    _rev_id[0] += 1
    return str(_rev_id[0])


# ============================================================
# Comment Management
# ============================================================
class CommentManager:
    """Manages Word document comments via OOXML manipulation."""

    def __init__(self, document):
        self.document = document
        self.comment_id = 0
        self.comments_part = None
        self._ensure_comments_part()

    def _ensure_comments_part(self):
        """Ensure the document has a comments part."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        # Check if comments part already exists
        doc_part = self.document.part
        for rel in doc_part.rels.values():
            if 'comments' in rel.reltype:
                self.comments_part = rel.target_part
                return

        # Create comments XML
        comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
            '</w:comments>'
        )

        comments_uri = PackURI('/word/comments.xml')
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
        comments_part = Part(
            comments_uri, content_type, comments_xml.encode('utf-8'), doc_part.package
        )

        rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        doc_part.relate_to(comments_part, rel_type)
        self.comments_part = comments_part

    def _get_comments_element(self):
        """Get the <w:comments> root element."""
        xml_bytes = self.comments_part.blob
        return etree.fromstring(xml_bytes)

    def _save_comments_element(self, comments_elem):
        """Save the comments element back to the part."""
        self.comments_part._blob = etree.tostring(comments_elem, xml_declaration=True, encoding='UTF-8', standalone=True)

    def add_comment(self, paragraph, comment_text, author="Legal Counsel", date_str=None):
        """
        Add a comment to a paragraph.

        Returns the comment_id for reference.
        """
        self.comment_id += 1
        cid = self.comment_id

        if date_str is None:
            date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S') + 'Z'

        # Add comment to comments part
        comments_elem = self._get_comments_element()

        comment_elem = etree.SubElement(comments_elem, qn('w:comment'))
        comment_elem.set(qn('w:id'), str(cid))
        comment_elem.set(qn('w:author'), author)
        comment_elem.set(qn('w:date'), date_str)

        # Comment body paragraph
        cp = etree.SubElement(comment_elem, qn('w:p'))
        cr = etree.SubElement(cp, qn('w:r'))
        ct = etree.SubElement(cr, qn('w:t'))
        ct.text = comment_text
        ct.set(qn('xml:space'), 'preserve')

        self._save_comments_element(comments_elem)

        # Add comment range markers to the paragraph in document body
        p_elem = paragraph._element

        # commentRangeStart - insert at the beginning
        range_start = OxmlElement('w:commentRangeStart')
        range_start.set(qn('w:id'), str(cid))
        p_elem.insert(0, range_start)

        # commentRangeEnd - insert before the last element or at the end
        range_end = OxmlElement('w:commentRangeEnd')
        range_end.set(qn('w:id'), str(cid))
        p_elem.append(range_end)

        # commentReference run
        ref_run = OxmlElement('w:r')
        ref_rpr = OxmlElement('w:rPr')
        ref_style = OxmlElement('w:rStyle')
        ref_style.set(qn('w:val'), 'CommentReference')
        ref_rpr.append(ref_style)
        ref_run.append(ref_rpr)

        comment_ref = OxmlElement('w:commentReference')
        comment_ref.set(qn('w:id'), str(cid))
        ref_run.append(comment_ref)
        p_elem.append(ref_run)

        return cid


# ============================================================
# Markdown Parser for Bold Text
# ============================================================
def parse_markdown_bold(text):
    """
    Parse Markdown bold syntax (**text**) and return a list of (text, is_bold) tuples.
    
    Example:
        "Normal **bold** normal" -> [("Normal ", False), ("bold", True), (" normal", False)]
    """
    if '**' not in text:
        return [(text, False)]
    
    parts = []
    i = 0
    n = len(text)
    
    while i < n:
        # Look for opening **
        bold_start = text.find('**', i)
        
        if bold_start == -1:
            # No more **, add rest as normal text
            if i < n:
                parts.append((text[i:], False))
            break
        
        # Add text before ** as normal
        if bold_start > i:
            parts.append((text[i:bold_start], False))
        
        # Look for closing **
        bold_end = text.find('**', bold_start + 2)
        
        if bold_end == -1:
            # No closing **, add ** and rest as normal
            parts.append((text[bold_start:], False))
            break
        
        # Extract bold text
        bold_text = text[bold_start + 2:bold_end]
        if bold_text:
            parts.append((bold_text, True))
        
        # Move past closing **
        i = bold_end + 2
    
    return parts


def add_text_with_bold_to_paragraph(paragraph, text, font_name=None, font_size=None):
    """
    Add text to a paragraph, parsing Markdown bold syntax (**text**) 
    and applying actual bold formatting.
    """
    parts = parse_markdown_bold(text)
    
    for part_text, is_bold in parts:
        if not part_text:
            continue
        
        run = paragraph.add_run(part_text)
        if is_bold:
            run.bold = True
        
        # Apply font settings if provided
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
    
    return paragraph


def add_deletion_with_bold(paragraph, text, author="Legal Counsel", date_str=None):
    """
    Add deletion with Markdown bold parsing.
    Note: For deletions, we keep the original text as-is since it's being deleted.
    """
    add_deletion_to_paragraph(paragraph, text, author, date_str)


def add_insertion_with_bold(paragraph, text, author="Legal Counsel", date_str=None):
    """
    Add insertion with Markdown bold parsing.
    For insertions, we need to manually create runs with bold formatting.
    """
    if date_str is None:
        date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S') + 'Z'
    
    p_elem = paragraph._element
    
    # Parse bold syntax
    parts = parse_markdown_bold(text)
    
    for part_text, is_bold in parts:
        if not part_text:
            continue
        
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), next_rev_id())
        ins_elem.set(qn('w:author'), author)
        ins_elem.set(qn('w:date'), date_str)
        
        run = OxmlElement('w:r')
        rpr = create_run_properties()
        
        # Add bold if needed
        if is_bold:
            bold_elem = OxmlElement('w:b')
            rpr.append(bold_elem)
        
        run.append(rpr)
        
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = part_text
        run.append(t)
        
        ins_elem.append(run)
        p_elem.append(ins_elem)


# ============================================================
# Track Changes (Revisions)
# ============================================================
def create_run_properties(font_name=None, font_size=None, color=None):
    """Create run properties element."""
    rpr = OxmlElement('w:rPr')
    if font_name:
        r_fonts = OxmlElement('w:rFonts')
        r_fonts.set(qn('w:ascii'), font_name)
        r_fonts.set(qn('w:hAnsi'), font_name)
        r_fonts.set(qn('w:eastAsia'), font_name)
        rpr.append(r_fonts)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))  # half-points
        rpr.append(sz)
        sz_cs = OxmlElement('w:szCs')
        sz_cs.set(qn('w:val'), str(font_size * 2))
        rpr.append(sz_cs)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rpr.append(c)
    return rpr


def add_deletion_to_paragraph(paragraph, text, author="Legal Counsel", date_str=None):
    """Add a deletion revision mark (w:del) to a paragraph."""
    if date_str is None:
        date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S') + 'Z'

    p_elem = paragraph._element

    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), next_rev_id())
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date_str)

    run = OxmlElement('w:r')
    rpr = create_run_properties()
    run.append(rpr)

    del_text = OxmlElement('w:delText')
    del_text.set(qn('xml:space'), 'preserve')
    del_text.text = text
    run.append(del_text)

    del_elem.append(run)
    p_elem.append(del_elem)


def add_insertion_to_paragraph(paragraph, text, author="Legal Counsel", date_str=None):
    """Add an insertion revision mark (w:ins) to a paragraph."""
    if date_str is None:
        date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S') + 'Z'

    p_elem = paragraph._element

    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), next_rev_id())
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date_str)

    run = OxmlElement('w:r')
    rpr = create_run_properties()
    run.append(rpr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)

    ins_elem.append(run)
    p_elem.append(ins_elem)


def add_normal_text_to_paragraph(paragraph, text):
    """Add normal (unchanged) text to a paragraph."""
    run = paragraph.add_run(text)
    return run


# ============================================================
# Document Setup
# ============================================================
def setup_document_styles(doc):
    """Configure default document styles for legal contracts."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Set East Asian font
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        rpr.insert(0, r_fonts)
    r_fonts.set(qn('w:eastAsia'), '\u5b8b\u4f53')  # SimSun for Chinese

    # Heading styles
    for i in range(1, 4):
        heading_style_name = f'Heading {i}'
        if heading_style_name in doc.styles:
            h_style = doc.styles[heading_style_name]
            h_font = h_style.font
            h_font.name = 'Times New Roman'
            h_font.bold = True
            h_font.size = Pt(14 - i)  # H1=13pt, H2=12pt, H3=11pt
            h_rpr = h_style.element.get_or_add_rPr()
            h_r_fonts = h_rpr.find(qn('w:rFonts'))
            if h_r_fonts is None:
                h_r_fonts = OxmlElement('w:rFonts')
                h_rpr.insert(0, h_r_fonts)
            h_r_fonts.set(qn('w:eastAsia'), '\u9ed1\u4f53')  # SimHei for Chinese headings

    # Create CommentReference style if it doesn't exist
    try:
        doc.styles.add_style('CommentReference', WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        pass

    return doc


def setup_page_layout(doc):
    """Configure page layout for A4 legal documents."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    return doc


# ============================================================
# Document Generation
# ============================================================
def generate_draft(data, output_path):
    """Generate a clean contract draft .docx."""
    doc = Document()
    setup_document_styles(doc)
    setup_page_layout(doc)

    # Title
    title = data.get('title', 'Contract')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Times New Roman'

    # Add a blank line after title
    doc.add_paragraph()

    # Sections
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        level = section.get('level', 1)

        if heading:
            doc.add_heading(heading, level=min(level, 3))

        for item in section.get('content', []):
            item_type = item.get('type', 'keep')
            if item_type == 'keep':
                text = item.get('text', '')
                if text.strip():
                    para = doc.add_paragraph()
                    add_text_with_bold_to_paragraph(para, text)
            # In draft mode, skip delete/insert/replace (only 'keep' matters)

    doc.save(output_path)
    print(f"Draft saved: {output_path}")


def generate_review(data, output_path):
    """Generate a contract review .docx with track changes and comments.
    
    CRITICAL: This function produces a document that mimics human lawyer review:
    - ALL contract text is included (unchanged + changed)
    - Only changed portions show track changes (delete/insert/replace)
    - Unchanged text displays normally without any markup
    - Comments appear as right-side margin bubbles
    """
    doc = Document()
    setup_document_styles(doc)
    setup_page_layout(doc)

    # Enable revision tracking in document settings
    enable_track_changes(doc)

    author = data.get('author', 'Legal Counsel')
    date_str = data.get('date', datetime.now().strftime('%Y-%m-%d'))
    revision_date = date_str + 'T00:00:00Z'

    comment_mgr = CommentManager(doc)

    # Title
    title = data.get('title', 'Contract Review')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Process sections
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        level = section.get('level', 1)

        if heading:
            doc.add_heading(heading, level=min(level, 3))

        for item in section.get('content', []):
            item_type = item.get('type', 'keep')

            if item_type == 'keep':
                # Unchanged text - display normally, NO track changes
                text = item.get('text', '')
                if text.strip():
                    para = doc.add_paragraph()
                    add_text_with_bold_to_paragraph(para, text)

            elif item_type == 'delete':
                # Deleted text - show with strikethrough track change
                text = item.get('text', '')
                comment_text = item.get('comment', '')
                para = doc.add_paragraph()
                add_deletion_with_bold(para, text, author, revision_date)
                if comment_text:
                    comment_mgr.add_comment(para, comment_text, author, revision_date)

            elif item_type == 'insert':
                # Inserted text - show with colored underline track change
                text = item.get('text', '')
                comment_text = item.get('comment', '')
                para = doc.add_paragraph()
                add_insertion_with_bold(para, text, author, revision_date)
                if comment_text:
                    comment_mgr.add_comment(para, comment_text, author, revision_date)

            elif item_type == 'replace':
                # Replaced text - show original deleted + new inserted
                original = item.get('original', '')
                revised = item.get('revised', '')
                comment_text = item.get('comment', '')
                para = doc.add_paragraph()
                add_deletion_with_bold(para, original, author, revision_date)
                add_insertion_with_bold(para, revised, author, revision_date)
                if comment_text:
                    comment_mgr.add_comment(para, comment_text, author, revision_date)

    doc.save(output_path)
    print(f"Review saved with track changes: {output_path}")


def enable_track_changes(doc):
    """Enable revision tracking in the document settings."""
    # Access or create document settings
    settings = doc.settings.element

    # Add trackChanges element to mark document as having revisions
    # This makes Word show the revisions panel when opening
    track_changes = OxmlElement('w:trackChanges')
    settings.append(track_changes)


# ============================================================
# Main Entry Point
# ============================================================
def main():
    if len(sys.argv) < 4:
        print("Usage: python3 docx_generator.py <mode> <input.json> <output.docx>")
        print("  mode: 'draft' or 'review'")
        sys.exit(1)

    mode = sys.argv[1].lower()
    input_path = sys.argv[2]
    output_path = sys.argv[3]

    if mode not in ('draft', 'review'):
        print(f"Error: Unknown mode '{mode}'. Use 'draft' or 'review'.")
        sys.exit(1)

    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    if mode == 'draft':
        generate_draft(data, output_path)
    elif mode == 'review':
        generate_review(data, output_path)


if __name__ == '__main__':
    main()

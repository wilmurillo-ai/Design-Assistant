#!/usr/bin/env python3
"""
ultra-memory: Word 文档文本提取 (Multimodal Phase 5)
从 .docx 文件中提取文本内容，写入 session 的 multimodal/ 目录，
并触发事实提取。

依赖: python-docx
安装: pip install python-docx
"""

import os
import sys
import json
import argparse
import hashlib
import subprocess
from datetime import datetime, timezone
from pathlib import Path

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")
if sys.stderr.encoding != "utf-8":
    sys.stderr.reconfigure(encoding="utf-8")

ULTRA_MEMORY_HOME = Path(os.environ.get("ULTRA_MEMORY_HOME", Path.home() / ".ultra-memory"))

CHUNK_SIZE = 500  # 每块字符数


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> list[tuple[int, str]]:
    """将文本分割为段落块，返回 [(chunk_index, text), ...]"""
    paragraphs = text.split("\n\n")
    chunks = []
    current = []
    current_len = 0
    idx = 0

    for para in paragraphs:
        para_len = len(para)
        if current_len + para_len > chunk_size and current:
            chunks.append((idx, "\n".join(current).strip()))
            idx += 1
            current = []
            current_len = 0
        current.append(para)
        current_len += para_len

    if current:
        chunks.append((idx, "\n".join(current).strip()))

    return chunks


def extract_text_from_docx(docx_path: str) -> str:
    """
    使用 python-docx 提取 DOCX 文本。
    按段落顺序读取，保留基本结构。
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        print("[ultra-memory] python-docx 未安装: pip install python-docx")
        return ""
    except Exception as e:
        print(f"[ultra-memory] DOCX 提取失败: {e}")
        return ""


def save_extracted_text(
    session_id: str,
    media_path: str,
    text: str,
    media_id: str,
) -> Path:
    """保存提取的文本到 multimodal 目录"""
    session_dir = ULTRA_MEMORY_HOME / "sessions" / session_id
    multimodal_dir = session_dir / "multimodal"
    multimodal_dir.mkdir(parents=True, exist_ok=True)

    file_name = Path(media_path).name
    output_file = multimodal_dir / f"{file_name}.txt"

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(f"# Extracted from: {media_path}\n")
        f.write(f"# Media ID: {media_id}\n")
        f.write(f"# Extracted at: {_now_iso()}\n")
        f.write(f"# Chars: {len(text)}\n")
        f.write("---\n")
        f.write(text)

    return output_file


def process_docx(session_id: str, docx_path: str) -> dict:
    """
    处理单个 DOCX 文件。
    返回处理结果摘要。
    """
    if not Path(docx_path).exists():
        print(f"[ultra-memory] DOCX 文件不存在: {docx_path}")
        return {"success": False, "error": "file not found"}

    # 提取文本
    text = extract_text_from_docx(docx_path)
    if not text.strip():
        return {"success": False, "error": "no text extracted"}

    # 生成 media_id
    media_id = f"media_{hashlib.sha1(docx_path.encode()).hexdigest()[:12]}"

    # 保存文本
    output_file = save_extracted_text(session_id, docx_path, text, media_id)

    # 分块
    chunks = _chunk_text(text)
    char_count = len(text)

    print(f"[ultra-memory] DOCX 提取完成: {docx_path}")
    print(f"  文件: {output_file.name}")
    print(f"  字符数: {char_count}")
    print(f"  文本块: {len(chunks)} 块")

    return {
        "success": True,
        "media_id": media_id,
        "session_id": session_id,
        "source_path": docx_path,
        "output_file": str(output_file),
        "char_count": char_count,
        "chunk_count": len(chunks),
        "processed_at": _now_iso(),
    }


# ── CLI ─────────────────────────────────────────────────────────────────────


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="从 Word 文档 (.docx) 提取文本")
    parser.add_argument("--path", required=True, help="DOCX 文件路径")
    parser.add_argument("--session", required=True, help="会话 ID")
    args = parser.parse_args()

    result = process_docx(args.session, args.path)
    if result["success"]:
        sys.exit(0)
    else:
        print(f"[ultra-memory] DOCX 处理失败: {result.get('error')}")
        sys.exit(1)

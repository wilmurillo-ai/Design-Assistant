#!/usr/bin/env python3
"""
Medical Referral Letter Generator
Generates professional referral letters for patient care transfer.
"""

import argparse
import json
import os
import sys
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from enum import Enum
import tempfile


class UrgencyLevel(Enum):
    ROUTINE = "Routine"
    URGENT = "Urgent"
    EMERGENT = "Emergent"


class OutputFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    TXT = "txt"


@dataclass
class PatientData:
    """Patient information for referral letter."""
    name: str
    date_of_birth: str
    patient_id: str
    gender: Optional[str] = None
    contact_phone: Optional[str] = None
    address: Optional[str] = None


@dataclass
class ProviderInfo:
    """Healthcare provider information."""
    name: str
    title: Optional[str] = None
    department: Optional[str] = None
    institution: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    address: Optional[str] = None


@dataclass
class ReferralData:
    """Complete referral letter data."""
    patient: PatientData
    referring_provider: ProviderInfo
    receiving_provider: ProviderInfo
    reason_for_referral: str
    primary_diagnosis: str
    relevant_history: Optional[str] = None
    current_medications: Optional[List[str]] = None
    allergies: Optional[List[str]] = None
    vital_signs: Optional[Dict[str, Any]] = None
    lab_results: Optional[List[str]] = None
    urgency: UrgencyLevel = UrgencyLevel.ROUTINE
    additional_notes: Optional[str] = None


class ReferralLetterGenerator:
    """Main class for generating medical referral letters."""
    
    def __init__(self, template_dir: Optional[str] = None):
        self.template_dir = template_dir or os.path.join(
            os.path.dirname(__file__), '..', 'references'
        )
        self.generated_date = datetime.now().strftime("%Y-%m-%d")
    
    def validate_input(self, data: ReferralData) -> List[str]:
        """Validate required fields."""
        errors = []
        
        if not data.patient.name:
            errors.append("Patient name is required")
        if not data.patient.date_of_birth:
            errors.append("Patient date of birth is required")
        if not data.patient.patient_id:
            errors.append("Patient ID is required")
        if not data.reason_for_referral:
            errors.append("Reason for referral is required")
        if not data.primary_diagnosis:
            errors.append("Primary diagnosis is required")
        if not data.referring_provider.name:
            errors.append("Referring provider name is required")
        if not data.receiving_provider.name:
            errors.append("Receiving provider name is required")
            
        return errors
    
    def generate_text_content(self, data: ReferralData) -> str:
        """Generate plain text content of the referral letter."""
        lines = []
        
        # Header
        lines.append("=" * 70)
        lines.append("MEDICAL REFERRAL LETTER".center(70))
        lines.append("=" * 70)
        lines.append("")
        
        # Date and Urgency
        lines.append(f"Date: {self.generated_date}")
        if data.urgency != UrgencyLevel.ROUTINE:
            lines.append(f"URGENCY: {data.urgency.value.upper()}")
        lines.append("")
        
        # Receiving Provider
        lines.append("TO:")
        lines.append(f"    {data.receiving_provider.name}")
        if data.receiving_provider.title:
            lines.append(f"    {data.receiving_provider.title}")
        if data.receiving_provider.department:
            lines.append(f"    {data.receiving_provider.department}")
        if data.receiving_provider.institution:
            lines.append(f"    {data.receiving_provider.institution}")
        lines.append("")
        
        # Referring Provider
        lines.append("FROM:")
        lines.append(f"    {data.referring_provider.name}")
        if data.referring_provider.title:
            lines.append(f"    {data.referring_provider.title}")
        if data.referring_provider.department:
            lines.append(f"    {data.referring_provider.department}")
        if data.referring_provider.institution:
            lines.append(f"    {data.referring_provider.institution}")
        if data.referring_provider.phone:
            lines.append(f"    Phone: {data.referring_provider.phone}")
        if data.referring_provider.email:
            lines.append(f"    Email: {data.referring_provider.email}")
        lines.append("")
        
        # Patient Information
        lines.append("-" * 70)
        lines.append("PATIENT INFORMATION")
        lines.append("-" * 70)
        lines.append(f"Name:           {data.patient.name}")
        lines.append(f"Date of Birth:  {data.patient.date_of_birth}")
        lines.append(f"Patient ID:     {data.patient.patient_id}")
        if data.patient.gender:
            lines.append(f"Gender:         {data.patient.gender}")
        if data.patient.contact_phone:
            lines.append(f"Phone:          {data.patient.contact_phone}")
        lines.append("")
        
        # Reason for Referral
        lines.append("-" * 70)
        lines.append("REASON FOR REFERRAL")
        lines.append("-" * 70)
        lines.append(data.reason_for_referral)
        lines.append("")
        
        # Diagnosis
        lines.append("-" * 70)
        lines.append("PRIMARY DIAGNOSIS")
        lines.append("-" * 70)
        lines.append(data.primary_diagnosis)
        lines.append("")
        
        # Relevant History
        if data.relevant_history:
            lines.append("-" * 70)
            lines.append("RELEVANT MEDICAL HISTORY")
            lines.append("-" * 70)
            lines.append(data.relevant_history)
            lines.append("")
        
        # Current Medications
        if data.current_medications:
            lines.append("-" * 70)
            lines.append("CURRENT MEDICATIONS")
            lines.append("-" * 70)
            for med in data.current_medications:
                lines.append(f"  • {med}")
            lines.append("")
        
        # Allergies
        if data.allergies:
            lines.append("-" * 70)
            lines.append("ALLERGIES")
            lines.append("-" * 70)
            for allergy in data.allergies:
                lines.append(f"  • {allergy}")
            lines.append("")
        
        # Vital Signs
        if data.vital_signs:
            lines.append("-" * 70)
            lines.append("VITAL SIGNS")
            lines.append("-" * 70)
            for key, value in data.vital_signs.items():
                lines.append(f"  {key}: {value}")
            lines.append("")
        
        # Lab Results
        if data.lab_results:
            lines.append("-" * 70)
            lines.append("RELEVANT LABORATORY RESULTS")
            lines.append("-" * 70)
            for result in data.lab_results:
                lines.append(f"  • {result}")
            lines.append("")
        
        # Additional Notes
        if data.additional_notes:
            lines.append("-" * 70)
            lines.append("ADDITIONAL NOTES")
            lines.append("-" * 70)
            lines.append(data.additional_notes)
            lines.append("")
        
        # Footer
        lines.append("")
        lines.append("-" * 70)
        lines.append("Thank you for your consultation and management of this patient.")
        lines.append("Please contact me if you require any additional information.")
        lines.append("")
        lines.append(f"Sincerely,")
        lines.append(f"")
        lines.append(f"{data.referring_provider.name}")
        if data.referring_provider.title:
            lines.append(f"{data.referring_provider.title}")
        lines.append("")
        lines.append("=" * 70)
        
        return "\n".join(lines)
    
    def generate_html(self, data: ReferralData) -> str:
        """Generate HTML formatted referral letter."""
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Medical Referral Letter - {data.patient.name}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        .header {{ text-align: center; border-bottom: 3px solid #333; padding-bottom: 10px; margin-bottom: 20px; }}
        .urgent {{ color: #d9534f; font-weight: bold; font-size: 1.2em; }}
        .section {{ margin: 20px 0; }}
        .section-title {{ background-color: #f5f5f5; padding: 8px; font-weight: bold; border-left: 4px solid #333; }}
        .field {{ margin: 5px 0; }}
        .label {{ font-weight: bold; display: inline-block; width: 150px; }}
        ul {{ margin: 5px 0; }}
        li {{ margin: 3px 0; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 2px solid #ccc; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>MEDICAL REFERRAL LETTER</h1>
        <p>Date: {self.generated_date}</p>
        {f'<p class="urgent">URGENCY: {data.urgency.value.upper()}</p>' if data.urgency != UrgencyLevel.ROUTINE else ''}
    </div>
    
    <div class="section">
        <div class="section-title">RECIPIENT</div>
        <p><strong>{data.receiving_provider.name}</strong><br>
        {f"{data.receiving_provider.title}<br>" if data.receiving_provider.title else ""}
        {f"{data.receiving_provider.department}<br>" if data.receiving_provider.department else ""}
        {f"{data.receiving_provider.institution}<br>" if data.receiving_provider.institution else ""}
        </p>
    </div>
    
    <div class="section">
        <div class="section-title">REFERRING PROVIDER</div>
        <p><strong>{data.referring_provider.name}</strong><br>
        {f"{data.referring_provider.title}<br>" if data.referring_provider.title else ""}
        {f"{data.referring_provider.department}<br>" if data.referring_provider.department else ""}
        {f"{data.referring_provider.institution}<br>" if data.referring_provider.institution else ""}
        {f"Phone: {data.referring_provider.phone}<br>" if data.referring_provider.phone else ""}
        {f"Email: {data.referring_provider.email}" if data.referring_provider.email else ""}
        </p>
    </div>
    
    <div class="section">
        <div class="section-title">PATIENT INFORMATION</div>
        <div class="field"><span class="label">Name:</span> {data.patient.name}</div>
        <div class="field"><span class="label">Date of Birth:</span> {data.patient.date_of_birth}</div>
        <div class="field"><span class="label">Patient ID:</span> {data.patient.patient_id}</div>
        {f'<div class="field"><span class="label">Gender:</span> {data.patient.gender}</div>' if data.patient.gender else ""}
        {f'<div class="field"><span class="label">Phone:</span> {data.patient.contact_phone}</div>' if data.patient.contact_phone else ""}
    </div>
    
    <div class="section">
        <div class="section-title">REASON FOR REFERRAL</div>
        <p>{data.reason_for_referral.replace(chr(10), '<br>')}</p>
    </div>
    
    <div class="section">
        <div class="section-title">PRIMARY DIAGNOSIS</div>
        <p>{data.primary_diagnosis.replace(chr(10), '<br>')}</p>
    </div>
    
    {f'''<div class="section">
        <div class="section-title">RELEVANT MEDICAL HISTORY</div>
        <p>{data.relevant_history.replace(chr(10), '<br>')}</p>
    </div>''' if data.relevant_history else ""}
    
    {f'''<div class="section">
        <div class="section-title">CURRENT MEDICATIONS</div>
        <ul>{''.join([f"<li>{med}</li>" for med in data.current_medications])}</ul>
    </div>''' if data.current_medications else ""}
    
    {f'''<div class="section">
        <div class="section-title">ALLERGIES</div>
        <ul>{''.join([f"<li>{allergy}</li>" for allergy in data.allergies])}</ul>
    </div>''' if data.allergies else ""}
    
    {f'''<div class="section">
        <div class="section-title">VITAL SIGNS</div>
        {''.join([f'<div class="field"><span class="label">{k}:</span> {v}</div>' for k, v in data.vital_signs.items()])}
    </div>''' if data.vital_signs else ""}
    
    {f'''<div class="section">
        <div class="section-title">RELEVANT LABORATORY RESULTS</div>
        <ul>{''.join([f"<li>{result}</li>" for result in data.lab_results])}</ul>
    </div>''' if data.lab_results else ""}
    
    {f'''<div class="section">
        <div class="section-title">ADDITIONAL NOTES</div>
        <p>{data.additional_notes.replace(chr(10), '<br>')}</p>
    </div>''' if data.additional_notes else ""}
    
    <div class="footer">
        <p>Thank you for your consultation and management of this patient.<br>
        Please contact me if you require any additional information.</p>
        <p><br>Sincerely,<br><br>
        <strong>{data.referring_provider.name}</strong><br>
        {data.referring_provider.title or ""}</p>
    </div>
</body>
</html>"""
        return html
    
    def generate_pdf(self, data: ReferralData, output_path: str) -> bool:
        """Generate PDF referral letter."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
        except ImportError:
            print("Warning: reportlab not installed. Installing required package...")
            os.system(f"{sys.executable} -m pip install reportlab -q")
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)
        
        styles = getSampleStyleSheet()
        story = []
        
        # Header
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12,
            alignment=1  # Center
        )
        story.append(Paragraph("MEDICAL REFERRAL LETTER", title_style))
        story.append(Paragraph(f"<b>Date:</b> {self.generated_date}", styles['Normal']))
        
        if data.urgency != UrgencyLevel.ROUTINE:
            urgent_style = ParagraphStyle(
                'Urgent',
                parent=styles['Normal'],
                textColor=colors.red,
                fontSize=14,
                alignment=1
            )
            story.append(Paragraph(f"<b>URGENCY: {data.urgency.value.upper()}</b>", urgent_style))
        
        story.append(Spacer(1, 0.2*inch))
        
        # Two-column layout for providers
        provider_data = [
            ['TO:', 'FROM:'],
            [data.receiving_provider.name, data.referring_provider.name],
        ]
        if data.receiving_provider.title or data.referring_provider.title:
            provider_data.append([
                data.receiving_provider.title or '',
                data.referring_provider.title or ''
            ])
        
        provider_table = Table(provider_data, colWidths=[3.5*inch, 3.5*inch])
        provider_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(provider_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Section helper
        def add_section(title, content):
            story.append(Paragraph(f"<b>{title}</b>", styles['Heading3']))
            story.append(Spacer(1, 0.05*inch))
            if isinstance(content, list):
                for item in content:
                    story.append(Paragraph(f"• {item}", styles['Normal']))
            else:
                story.append(Paragraph(content.replace('\n', '<br/>'), styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        # Patient Info
        patient_info = f"""
        <b>Name:</b> {data.patient.name}<br/>
        <b>Date of Birth:</b> {data.patient.date_of_birth}<br/>
        <b>Patient ID:</b> {data.patient.patient_id}
        {f'<br/><b>Gender:</b> {data.patient.gender}' if data.patient.gender else ''}
        {f'<br/><b>Phone:</b> {data.patient.contact_phone}' if data.patient.contact_phone else ''}
        """
        add_section("PATIENT INFORMATION", patient_info)
        
        # Reason and Diagnosis
        add_section("REASON FOR REFERRAL", data.reason_for_referral)
        add_section("PRIMARY DIAGNOSIS", data.primary_diagnosis)
        
        # Optional sections
        if data.relevant_history:
            add_section("RELEVANT MEDICAL HISTORY", data.relevant_history)
        if data.current_medications:
            add_section("CURRENT MEDICATIONS", data.current_medications)
        if data.allergies:
            add_section("ALLERGIES", data.allergies)
        if data.additional_notes:
            add_section("ADDITIONAL NOTES", data.additional_notes)
        
        # Footer
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(
            "Thank you for your consultation and management of this patient. "
            "Please contact me if you require any additional information.",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Sincerely,<br/><br/>", styles['Normal']))
        story.append(Paragraph(f"<b>{data.referring_provider.name}</b>", styles['Normal']))
        if data.referring_provider.title:
            story.append(Paragraph(data.referring_provider.title, styles['Normal']))
        
        doc.build(story)
        return True
    
    def generate_docx(self, data: ReferralData, output_path: str) -> bool:
        """Generate DOCX referral letter."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("Warning: python-docx not installed. Installing required package...")
            os.system(f"{sys.executable} -m pip install python-docx -q")
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Header
        heading = doc.add_heading('MEDICAL REFERRAL LETTER', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date: {self.generated_date}")
        
        if data.urgency != UrgencyLevel.ROUTINE:
            p = doc.add_paragraph()
            run = p.add_run(f"URGENCY: {data.urgency.value.upper()}")
            run.font.color.rgb = None  # Red handled below
            run.bold = True
            run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Providers table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "TO:"
        table.cell(0, 1).text = "FROM:"
        table.cell(1, 0).text = data.receiving_provider.name
        table.cell(1, 1).text = data.referring_provider.name
        
        doc.add_paragraph()
        
        # Helper for sections
        def add_section_docx(title, content):
            doc.add_heading(title, level=2)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(content)
        
        # Patient info
        patient_text = f"Name: {data.patient.name}\n"
        patient_text += f"Date of Birth: {data.patient.date_of_birth}\n"
        patient_text += f"Patient ID: {data.patient.patient_id}"
        if data.patient.gender:
            patient_text += f"\nGender: {data.patient.gender}"
        add_section_docx("PATIENT INFORMATION", patient_text)
        
        add_section_docx("REASON FOR REFERRAL", data.reason_for_referral)
        add_section_docx("PRIMARY DIAGNOSIS", data.primary_diagnosis)
        
        if data.relevant_history:
            add_section_docx("RELEVANT MEDICAL HISTORY", data.relevant_history)
        if data.current_medications:
            add_section_docx("CURRENT MEDICATIONS", data.current_medications)
        if data.allergies:
            add_section_docx("ALLERGIES", data.allergies)
        if data.additional_notes:
            add_section_docx("ADDITIONAL NOTES", data.additional_notes)
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("Thank you for your consultation and management of this patient. "
                         "Please contact me if you require any additional information.")
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph(data.referring_provider.name)
        
        doc.save(output_path)
        return True
    
    def generate(self, data: ReferralData, output_format: OutputFormat, output_path: str) -> bool:
        """Generate referral letter in specified format."""
        # Validate
        errors = self.validate_input(data)
        if errors:
            print("Validation Errors:")
            for error in errors:
                print(f"  - {error}")
            return False
        
        # Generate based on format
        if output_format == OutputFormat.TXT:
            content = self.generate_text_content(data)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        
        elif output_format == OutputFormat.HTML:
            content = self.generate_html(data)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        
        elif output_format == OutputFormat.PDF:
            return self.generate_pdf(data, output_path)
        
        elif output_format == OutputFormat.DOCX:
            return self.generate_docx(data, output_path)
        
        return False


def load_from_json(json_path: str) -> ReferralData:
    """Load referral data from JSON file."""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Parse patient
    patient = PatientData(
        name=data['patient']['name'],
        date_of_birth=data['patient']['date_of_birth'],
        patient_id=data['patient']['patient_id'],
        gender=data['patient'].get('gender'),
        contact_phone=data['patient'].get('contact_phone'),
        address=data['patient'].get('address')
    )
    
    # Parse providers
    referring = ProviderInfo(
        name=data['referring_provider']['name'],
        title=data['referring_provider'].get('title'),
        department=data['referring_provider'].get('department'),
        institution=data['referring_provider'].get('institution'),
        phone=data['referring_provider'].get('phone'),
        email=data['referring_provider'].get('email'),
        address=data['referring_provider'].get('address')
    )
    
    receiving = ProviderInfo(
        name=data['receiving_provider']['name'],
        title=data['receiving_provider'].get('title'),
        department=data['receiving_provider'].get('department'),
        institution=data['receiving_provider'].get('institution'),
        phone=data['receiving_provider'].get('phone'),
        email=data['receiving_provider'].get('email'),
        address=data['receiving_provider'].get('address')
    )
    
    # Parse urgency
    urgency_str = data.get('urgency', 'Routine')
    urgency = UrgencyLevel.ROUTINE
    if urgency_str.lower() == 'urgent':
        urgency = UrgencyLevel.URGENT
    elif urgency_str.lower() == 'emergent':
        urgency = UrgencyLevel.EMERGENT
    
    return ReferralData(
        patient=patient,
        referring_provider=referring,
        receiving_provider=receiving,
        reason_for_referral=data['reason_for_referral'],
        primary_diagnosis=data['primary_diagnosis'],
        relevant_history=data.get('relevant_history'),
        current_medications=data.get('current_medications'),
        allergies=data.get('allergies'),
        vital_signs=data.get('vital_signs'),
        lab_results=data.get('lab_results'),
        urgency=urgency,
        additional_notes=data.get('additional_notes')
    )


def create_sample_data() -> dict:
    """Create sample referral data for testing."""
    return {
        "patient": {
            "name": "Jane Smith",
            "date_of_birth": "1985-07-22",
            "patient_id": "MRN12345678",
            "gender": "Female",
            "contact_phone": "(555) 123-4567"
        },
        "referring_provider": {
            "name": "Dr. Robert Johnson",
            "title": "Internal Medicine",
            "institution": "City General Hospital",
            "phone": "(555) 987-6543",
            "email": "r.johnson@citygeneral.com"
        },
        "receiving_provider": {
            "name": "Dr. Sarah Williams",
            "title": "Cardiologist",
            "institution": "Heart Care Center",
            "department": "Department of Cardiology"
        },
        "reason_for_referral": "Patient presents with intermittent chest pain, shortness of breath on exertion, and abnormal ECG findings suggesting possible coronary artery disease. Request cardiology evaluation and stress testing.",
        "primary_diagnosis": "Suspected coronary artery disease, Class II angina",
        "relevant_history": "Hypertension (5 years), Type 2 Diabetes (3 years), Family history of CAD in father (MI at age 55). Former smoker (quit 2 years ago, 20 pack-year history).",
        "current_medications": [
            "Lisinopril 10mg daily",
            "Metformin 500mg twice daily",
            "Atorvastatin 20mg daily",
            "Aspirin 81mg daily"
        ],
        "allergies": ["Penicillin (rash)"],
        "urgency": "Urgent",
        "additional_notes": "Patient is anxious about cardiac symptoms. Reassurance provided. Please copy me on all reports."
    }


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Generate medical referral letters',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --input patient.json --output referral.pdf
  %(prog)s --sample --output sample.pdf --format pdf
  %(prog)s --input data.json --format html --output letter.html
        """
    )
    
    parser.add_argument('--input', '-i', type=str,
                       help='Input JSON file with patient and referral data')
    parser.add_argument('--output', '-o', type=str, required=True,
                       help='Output file path')
    parser.add_argument('--format', '-f', type=str, default='pdf',
                       choices=['pdf', 'docx', 'html', 'txt'],
                       help='Output format (default: pdf)')
    parser.add_argument('--sample', action='store_true',
                       help='Generate a sample referral letter for testing')
    
    args = parser.parse_args()
    
    # Create generator
    generator = ReferralLetterGenerator()
    
    # Get data
    if args.sample:
        data_dict = create_sample_data()
        # Save sample JSON for reference
        sample_json_path = args.output.replace('.pdf', '.json').replace('.docx', '.json').replace('.html', '.json').replace('.txt', '.json')
        with open(sample_json_path, 'w', encoding='utf-8') as f:
            json.dump(data_dict, f, indent=2)
        print(f"Sample JSON saved to: {sample_json_path}")
        data = load_from_json(sample_json_path)
    elif args.input:
        data = load_from_json(args.input)
    else:
        print("Error: Either --input or --sample must be specified")
        parser.print_help()
        sys.exit(1)
    
    # Generate
    output_format = OutputFormat(args.format.lower())
    
    print(f"Generating {args.format.upper()} referral letter...")
    success = generator.generate(data, output_format, args.output)
    
    if success:
        print(f"Referral letter generated: {args.output}")
    else:
        print("Failed to generate referral letter")
        sys.exit(1)


if __name__ == '__main__':
    main()

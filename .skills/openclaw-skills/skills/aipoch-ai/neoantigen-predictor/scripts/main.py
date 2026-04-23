#!/usr/bin/env python3
"""
NIH Biosketch Builder
Generate NIH-compliant biosketch documents (2022 OMB-approved format)

Usage:
    python main.py --input data.json --output biosketch.docx
    python main.py --import-pubmed "12345678,23456789" --output pubs.json
    python main.py --template --output template.json
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

# Optional imports - handled gracefully
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False


# ============== NIH Format Constants ==============
NIH_FONT = 'Arial'
NIH_FONT_SIZE = 11  # pt
NIH_MIN_MARGIN = 0.5  # inches
NIH_MAX_PAGES = 5

SECTION_TITLES = {
    'personal_statement': 'A. Personal Statement',
    'positions_honors': 'B. Positions and Honors',
    'contributions': 'C. Contribution to Science',
    'research_support': 'D. Research Support'
}


# ============== Data Models ==============

class PersonalInfo:
    """Researcher personal information"""
    def __init__(self, data: dict):
        self.name = data.get('name', '')
        self.position = data.get('position', '')
        self.department = data.get('department', '')
        self.organization = data.get('organization', '')
        self.email = data.get('email', '')
        self.phone = data.get('phone', '')
        self.ecommons_id = data.get('ecommons_id', '')
    
    def to_dict(self) -> dict:
        return {
            'name': self.name,
            'position': self.position,
            'department': self.department,
            'organization': self.organization,
            'email': self.email,
            'phone': self.phone,
            'ecommons_id': self.ecommons_id
        }


class Contribution:
    """Contribution to Science entry"""
    def __init__(self, data: dict):
        self.title = data.get('title', '')
        self.description = data.get('description', '')
        self.publications = data.get('publications', [])
    
    def to_dict(self) -> dict:
        return {
            'title': self.title,
            'description': self.description,
            'publications': self.publications
        }


class BiosketchData:
    """Complete Biosketch data container"""
    def __init__(self, data: Optional[dict] = None):
        data = data or {}
        self.personal_info = PersonalInfo(data.get('personal_info', {}))
        self.personal_statement = data.get('personal_statement', '')
        self.positions_and_honors = data.get('positions_and_honors', [])
        self.contributions = [Contribution(c) for c in data.get('contributions', [])]
        self.research_support = data.get('research_support', [])
        self.publications = data.get('publications', [])
    
    def to_dict(self) -> dict:
        return {
            'personal_info': self.personal_info.to_dict(),
            'personal_statement': self.personal_statement,
            'positions_and_honors': self.positions_and_honors,
            'contributions': [c.to_dict() for c in self.contributions],
            'research_support': self.research_support,
            'publications': self.publications
        }


# ============== PubMed Integration ==============

class PubMedImporter:
    """Import publication data from PubMed/NIH APIs"""
    
    PUBMED_API = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    
    def __init__(self):
        self.session = requests.Session() if HAS_REQUESTS else None
    
    def fetch_by_pmid(self, pmid: str) -> Optional[dict]:
        """Fetch publication data by PMID"""
        if not HAS_REQUESTS:
            print("Warning: requests library not installed. PubMed import unavailable.")
            return None
        
        try:
            # Fetch summary
            url = f"{self.PUBMED_API}/esummary.fcgi"
            params = {
                'db': 'pubmed',
                'id': pmid,
                'retmode': 'json'
            }
            response = self.session.get(url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            if 'result' not in data or pmid not in data['result']:
                return None
            
            result = data['result'][pmid]
            
            # Fetch detailed citation info
            cite_url = f"{self.PUBMED_API}/efetch.fcgi"
            cite_params = {
                'db': 'pubmed',
                'id': pmid,
                'rettype': 'medline',
                'retmode': 'text'
            }
            cite_response = self.session.get(cite_url, params=cite_params, timeout=30)
            
            return {
                'pmid': pmid,
                'title': result.get('title', ''),
                'authors': [a.get('name', '') for a in result.get('authors', [])],
                'journal': result.get('source', ''),
                'year': result.get('pubdate', '').split()[0] if result.get('pubdate') else '',
                'doi': self._extract_doi(result),
                'raw_medline': cite_response.text if cite_response.status_code == 200 else ''
            }
        except Exception as e:
            print(f"Error fetching PMID {pmid}: {e}")
            return None
    
    def _extract_doi(self, result: dict) -> str:
        """Extract DOI from PubMed result"""
        article_ids = result.get('articleids', [])
        for aid in article_ids:
            if aid.get('idtype') == 'doi':
                return aid.get('value', '')
        return ''
    
    def fetch_multiple(self, pmids: list[str]) -> list[dict]:
        """Fetch multiple publications by PMID"""
        results = []
        for pmid in pmids:
            pmid = pmid.strip()
            if pmid:
                pub = self.fetch_by_pmid(pmid)
                if pub:
                    results.append(pub)
        return results
    
    def search_by_author(self, last_name: str, first_name: str = '', 
                         affiliation: str = '', max_results: int = 100) -> list[dict]:
        """Search PubMed by author name"""
        if not HAS_REQUESTS:
            return []
        
        try:
            # Build search query
            query = f"{last_name} {first_name}".strip()
            if affiliation:
                query += f" AND {affiliation}"
            
            # Search
            search_url = f"{self.PUBMED_API}/esearch.fcgi"
            params = {
                'db': 'pubmed',
                'term': query,
                'retmax': max_results,
                'retmode': 'json',
                'sort': 'date'
            }
            response = self.session.get(search_url, params=params, timeout=30)
            data = response.json()
            
            id_list = data.get('esearchresult', {}).get('idlist', [])
            
            # Fetch details for each
            return self.fetch_multiple(id_list)
            
        except Exception as e:
            print(f"Error searching PubMed: {e}")
            return []


# ============== Document Generator ==============

class BiosketchGenerator:
    """Generate NIH-compliant Biosketch documents"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def generate(self, data: BiosketchData, output_path: str):
        """Generate DOCX file"""
        doc = Document()
        
        # Set margins (minimum 0.5 inches)
        sections = doc.sections[0]
        sections.top_margin = Inches(0.5)
        sections.bottom_margin = Inches(0.5)
        sections.left_margin = Inches(0.5)
        sections.right_margin = Inches(0.5)
        
        # Header with name
        self._add_header(doc, data.personal_info)
        
        # Section A: Personal Statement
        self._add_section_a(doc, data)
        
        # Section B: Positions and Honors
        self._add_section_b(doc, data)
        
        # Section C: Contribution to Science
        self._add_section_c(doc, data)
        
        # Section D: Research Support (optional)
        if data.research_support:
            self._add_section_d(doc, data)
        
        # Save document
        doc.save(output_path)
        print(f"Generated NIH Biosketch: {output_path}")
    
    def _add_header(self, doc: Document, info: PersonalInfo):
        """Add document header with name"""
        # Name as title
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(info.name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.name = NIH_FONT
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if info.position:
            contact_parts.append(info.position)
        if info.department:
            contact_parts.append(info.department)
        if info.organization:
            contact_parts.append(info.organization)
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(NIH_FONT_SIZE)
            contact_run.font.name = NIH_FONT
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Email
        if info.email:
            email_para = doc.add_paragraph()
            email_run = email_para.add_run(info.email)
            email_run.font.size = Pt(NIH_FONT_SIZE)
            email_run.font.name = NIH_FONT
            email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
    
    def _add_section_heading(self, doc: Document, title: str):
        """Add a section heading"""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(NIH_FONT_SIZE)
        run.font.name = NIH_FONT
        para.space_after = Pt(6)
    
    def _add_section_a(self, doc: Document, data: BiosketchData):
        """Section A: Personal Statement"""
        self._add_section_heading(doc, SECTION_TITLES['personal_statement'])
        
        # Relevant positions
        if data.positions_and_honors:
            current_pos = data.positions_and_honors[0] if data.positions_and_honors else {}
            if isinstance(current_pos, dict) and 'position' in current_pos:
                pos_para = doc.add_paragraph()
                pos_run = pos_para.add_run(f"Position: {current_pos.get('position', '')}")
                pos_run.font.size = Pt(NIH_FONT_SIZE)
                pos_run.font.name = NIH_FONT
        
        # Personal statement text
        if data.personal_statement:
            para = doc.add_paragraph()
            run = para.add_run(data.personal_statement)
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            para.paragraph_format.line_spacing = 1.15
        
        doc.add_paragraph()  # Spacing
    
    def _add_section_b(self, doc: Document, data: BiosketchData):
        """Section B: Positions and Honors"""
        self._add_section_heading(doc, SECTION_TITLES['positions_honors'])
        
        # Positions
        positions = [p for p in data.positions_and_honors 
                     if isinstance(p, dict) and 'position' in p]
        
        if positions:
            para = doc.add_paragraph()
            run = para.add_run("Positions and Employment")
            run.bold = True
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            
            for pos in positions:
                line = f"{pos.get('year', '')}  {pos.get('position', '')}"
                if pos.get('institution'):
                    line += f", {pos['institution']}"
                
                p = doc.add_paragraph(line, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                for r in p.runs:
                    r.font.size = Pt(NIH_FONT_SIZE)
                    r.font.name = NIH_FONT
        
        # Honors (separate list if marked)
        honors = [h for h in data.positions_and_honors 
                  if isinstance(h, dict) and h.get('type') == 'honor']
        
        if honors:
            para = doc.add_paragraph()
            run = para.add_run("Honors")
            run.bold = True
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            
            for honor in honors:
                line = f"{honor.get('year', '')}  {honor.get('title', '')}"
                if honor.get('organization'):
                    line += f", {honor['organization']}"
                
                p = doc.add_paragraph(line, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                for r in p.runs:
                    r.font.size = Pt(NIH_FONT_SIZE)
                    r.font.name = NIH_FONT
        
        doc.add_paragraph()  # Spacing
    
    def _add_section_c(self, doc: Document, data: BiosketchData):
        """Section C: Contribution to Science"""
        self._add_section_heading(doc, SECTION_TITLES['contributions'])
        
        if not data.contributions:
            para = doc.add_paragraph()
            run = para.add_run("No contributions listed.")
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            return
        
        for i, contrib in enumerate(data.contributions[:4], 1):  # Max 4 contributions
            # Contribution title
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {contrib.title}")
            run.bold = True
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            
            # Description
            if contrib.description:
                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run(contrib.description)
                desc_run.font.size = Pt(NIH_FONT_SIZE)
                desc_run.font.name = NIH_FONT
                desc_para.paragraph_format.left_indent = Inches(0.25)
            
            # Related publications
            if contrib.publications:
                pub_para = doc.add_paragraph()
                pub_run = pub_para.add_run("Relevant publications: ")
                pub_run.italic = True
                pub_run.font.size = Pt(NIH_FONT_SIZE - 1)
                pub_run.font.name = NIH_FONT
                pub_para.paragraph_format.left_indent = Inches(0.25)
                
                pubs_text = ', '.join(contrib.publications[:5])  # Max 5 per contribution
                pubs_run = pub_para.add_run(pubs_text)
                pubs_run.font.size = Pt(NIH_FONT_SIZE - 1)
                pubs_run.font.name = NIH_FONT
            
            doc.add_paragraph()  # Spacing between contributions
    
    def _add_section_d(self, doc: Document, data: BiosketchData):
        """Section D: Research Support"""
        self._add_section_heading(doc, SECTION_TITLES['research_support'])
        
        # Active vs Completed
        active = [s for s in data.research_support if s.get('status') != 'completed']
        completed = [s for s in data.research_support if s.get('status') == 'completed']
        
        if active:
            para = doc.add_paragraph()
            run = para.add_run("Ongoing Research Support")
            run.bold = True
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            
            for grant in active:
                self._add_grant_entry(doc, grant)
        
        if completed:
            para = doc.add_paragraph()
            run = para.add_run("Completed Research Support")
            run.bold = True
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            
            for grant in completed:
                self._add_grant_entry(doc, grant)
        
        doc.add_paragraph()  # Spacing
    
    def _add_grant_entry(self, doc: Document, grant: dict):
        """Add a single grant entry"""
        # Grant identifier line
        grant_line = grant.get('grant_number', '')
        if grant.get('pi_role'):
            grant_line += f" ({grant['pi_role']})"
        
        if grant_line:
            para = doc.add_paragraph()
            run = para.add_run(grant_line)
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            para.paragraph_format.left_indent = Inches(0.25)
        
        # Title
        if grant.get('title'):
            para = doc.add_paragraph()
            run = para.add_run(f"Title: {grant['title']}")
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            para.paragraph_format.left_indent = Inches(0.5)
        
        # Period
        if grant.get('period'):
            para = doc.add_paragraph()
            run = para.add_run(f"Period: {grant['period']}")
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            para.paragraph_format.left_indent = Inches(0.5)
        
        # Role and effort
        if grant.get('role'):
            role_line = f"Role: {grant['role']}"
            if grant.get('effort'):
                role_line += f" ({grant['effort']})"
            para = doc.add_paragraph()
            run = para.add_run(role_line)
            run.font.size = Pt(NIH_FONT_SIZE)
            run.font.name = NIH_FONT
            para.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()  # Spacing


# ============== Main Entry Point ==============

def generate_template() -> dict:
    """Generate a template JSON structure"""
    return {
        "personal_info": {
            "name": "Your Name",
            "position": "e.g., Associate Professor",
            "department": "Department Name",
            "organization": "University/Institution Name",
            "email": "your.email@institution.edu",
            "phone": "+1-XXX-XXX-XXXX",
            "ecommons_id": "Your eRA Commons ID"
        },
        "personal_statement": """Briefly describe why your experience and qualifications make you particularly well-suited for your role in the project described in the application.

The personal statement should be a clear, concise description of your relevant experience and qualifications. Maximum 1 page.""",
        "positions_and_honors": [
            {
                "year": "2020-present",
                "position": "Associate Professor",
                "institution": "Your University"
            },
            {
                "year": "2015-2020",
                "position": "Assistant Professor",
                "institution": "Previous Institution"
            },
            {
                "type": "honor",
                "year": "2023",
                "title": "Outstanding Researcher Award",
                "organization": "Professional Society"
            }
        ],
        "contributions": [
            {
                "title": "Major Contribution Title",
                "description": """Describe your contribution to science in this area. Explain the significance of the work and its impact on the field. 

Include specific findings, methodologies developed, or theoretical frameworks advanced. Be specific about your role in this work if it involved collaboration.""",
                "publications": [
                    "PMID:12345678",
                    "DOI:10.1000/example"
                ]
            }
        ],
        "research_support": [
            {
                "grant_number": "R01 CA123456",
                "title": "Research Project Title",
                "agency": "NIH/NCI",
                "period": "2021-2026",
                "role": "Principal Investigator",
                "pi_role": "PI",
                "effort": "25% effort",
                "amount": "$1,500,000 total costs",
                "status": "active"
            }
        ],
        "publications": [
            {
                "authors": "Author A, Author B, Your Name",
                "title": "Paper Title",
                "journal": "Journal Name",
                "year": "2024",
                "volume": "10(1)",
                "pages": "1-10",
                "pmid": "12345678",
                "doi": "10.1000/example"
            }
        ]
    }


def main():
    parser = argparse.ArgumentParser(
        description='NIH Biosketch Builder - Generate NIH-compliant biosketch documents'
    )
    parser.add_argument('--input', '-i', type=str, 
                        help='Input JSON file with biosketch data')
    parser.add_argument('--output', '-o', type=str, required=True,
                        help='Output file (.docx for biosketch, .json for other outputs)')
    parser.add_argument('--template', action='store_true',
                        help='Generate a template JSON file')
    parser.add_argument('--import-pubmed', type=str,
                        help='Import publications by PMID (comma-separated)')
    parser.add_argument('--search-pubmed', type=str,
                        help='Search PubMed by author name (format: "Last,First")')
    parser.add_argument('--auto-import-pubmed', action='store_true',
                        help='Auto-import publications from PMIDs in input data')
    
    args = parser.parse_args()
    
    # Generate template
    if args.template:
        template = generate_template()
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(template, f, indent=2, ensure_ascii=False)
        print(f"Template generated: {args.output}")
        return
    
    # Import from PubMed by PMID
    if args.import_pubmed:
        if not HAS_REQUESTS:
            print("Error: requests library required. Install with: pip install requests")
            sys.exit(1)
        
        importer = PubMedImporter()
        pmids = args.import_pubmed.split(',')
        publications = importer.fetch_multiple(pmids)
        
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(publications, f, indent=2, ensure_ascii=False)
        print(f"Imported {len(publications)} publications to {args.output}")
        return
    
    # Search PubMed
    if args.search_pubmed:
        if not HAS_REQUESTS:
            print("Error: requests library required. Install with: pip install requests")
            sys.exit(1)
        
        parts = args.search_pubmed.split(',')
        last_name = parts[0].strip()
        first_name = parts[1].strip() if len(parts) > 1 else ''
        
        importer = PubMedImporter()
        publications = importer.search_by_author(last_name, first_name)
        
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(publications, f, indent=2, ensure_ascii=False)
        print(f"Found {len(publications)} publications for {first_name} {last_name}")
        return
    
    # Generate Biosketch
    if not args.input:
        print("Error: --input required for generating biosketch (or use --template)")
        sys.exit(1)
    
    # Load input data
    with open(args.input, 'r', encoding='utf-8') as f:
        input_data = json.load(f)
    
    data = BiosketchData(input_data)
    
    # Auto-import publications if requested
    if args.auto_import_pubmed:
        if not HAS_REQUESTS:
            print("Warning: requests library not installed. Skipping PubMed import.")
        else:
            importer = PubMedImporter()
            pmids = []
            for contrib in data.contributions:
                pmids.extend([p for p in contrib.publications if p.startswith('PMID:')])
            
            if pmids:
                pmids = [p.replace('PMID:', '').strip() for p in pmids]
                publications = importer.fetch_multiple(pmids)
                # Update data with full publication info
                data.publications = publications
                print(f"Auto-imported {len(publications)} publications")
    
    # Generate document
    if not HAS_DOCX:
        print("Error: python-docx required. Install with: pip install python-docx")
        sys.exit(1)
    
    generator = BiosketchGenerator()
    generator.generate(data, args.output)


if __name__ == '__main__':
    main()

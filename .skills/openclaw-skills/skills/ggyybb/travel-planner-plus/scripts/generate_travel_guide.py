#!/usr/bin/env python3
"""
旅游攻略生成器 - Travel Guide Generator
使用百度地图API生成精确的旅游路线规划

用法:
    python3 generate_travel_guide.py <目的地> <天数> <晚数> [百度AK]

"""

import sys
import json
import urllib.request
import urllib.parse
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def baidu_geocode(address, ak):
    """地理编码：地址 -> 坐标"""
    url = f"https://api.map.baidu.com/geocoding/v3/?address={urllib.parse.quote(address)}&output=json&ak={ak}"
    try:
        with urllib.request.urlopen(url, timeout=10) as resp:
            data = json.loads(resp.read().decode())
            if data.get('status') == 0:
                loc = data['result']['location']
                return loc['lng'], loc['lat']
    except Exception as e:
        print(f"  [地理编码失败] {address}: {e}")
    return None, None


def baidu_direction_transit(origin, destination, ak):
    """公交路线规划"""
    url = f"https://api.map.baidu.com/direction/v2/transit?origin={origin[1]},{origin[0]}&destination={destination[1]},{destination[0]}&ak={ak}"
    try:
        with urllib.request.urlopen(url, timeout=10) as resp:
            data = json.loads(resp.read().decode())
            if data.get('status') == 0 and data['result']['routes']:
                r = data['result']['routes'][0]
                return {
                    'distance': r['distance'],
                    'duration': r['duration'],
                    'price': r.get('price', 0)
                }
    except Exception as e:
        print(f"  [路线规划失败] {origin} -> {destination}: {e}")
    return {'distance': 0, 'duration': 0, 'price': 0}


def haversine_meters(lat1, lon1, lat2, lon2):
    """计算两点间距离（米）"""
    from math import radians, cos, sin, asin, sqrt
    lon1, lat1, lon2, lat2 = map(radians, [lon1, lat1, lon2, lat2])
    dlon = lon2 - lon1
    dlat = lat2 - lat1
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
    c = 2 * asin(sqrt(a))
    return c * 6371000


def build_doc(city, days, nights, itinerary_data, transport_data, hotel_data, food_data, tips, output_path):
    """生成Word文档"""
    doc = Document()

    # 标题
    title = doc.add_heading(f'{city}{days}天{nights}晚旅游攻略', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'📅 Day1晚入住 → Day{days}天游玩 → Day{days+1}上午返程')
    doc.add_paragraph('🚇 交通：地铁+打车为主，步行尽量少')
    doc.add_paragraph('')

    # 每日行程表
    for day_info in itinerary_data:
        doc.add_heading(f"📅 {day_info['title']}", level=1)

        if not day_info.get('items'):
            continue

        rows = len(day_info['items']) + 1
        t = doc.add_table(rows=rows, cols=5)
        t.style = 'Table Grid'

        headers = ['时间', '活动', '地点', '游览时长', '交通']
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            set_cell_shading(t.rows[0].cells[i], 'E8F4FD')

        for i, item in enumerate(day_info['items']):
            for j, val in enumerate(item):
                t.rows[i+1].cells[j].text = str(val)

        doc.add_paragraph('')

    # 交通总览
    if transport_data:
        doc.add_heading('🚇 交通总览', level=1)
        for section in transport_data:
            doc.add_heading(section['title'], level=2)
            tt = doc.add_table(rows=len(section['routes'])+1, cols=4)
            tt.style = 'Table Grid'
            ht = ['路线', '距离', '方式', '用时/费用']
            for i, h in enumerate(ht):
                tt.rows[0].cells[i].text = h
                set_cell_shading(tt.rows[0].cells[i], 'E8F4FD')
            for i, r in enumerate(section['routes']):
                for j, v in enumerate(r):
                    tt.rows[i+1].cells[j].text = v
            doc.add_paragraph('')

    # 住宿
    if hotel_data:
        doc.add_heading('🏨 住宿推荐', level=1)
        th = doc.add_table(rows=len(hotel_data)+1, cols=3)
        th.style = 'Table Grid'
        hh = ['区域', '酒店类型', '优点']
        for i, h in enumerate(hh):
            th.rows[0].cells[i].text = h
            set_cell_shading(th.rows[0].cells[i], 'E8F4FD')
        for i, h in enumerate(hotel_data):
            for j, v in enumerate(h):
                th.rows[i+1].cells[j].text = v
        doc.add_paragraph('')

    # 美食
    if food_data:
        doc.add_heading('🍜 必吃美食', level=1)
        tf = doc.add_table(rows=len(food_data)+1, cols=3)
        tf.style = 'Table Grid'
        hf = ['类型', '推荐店铺', '人均']
        for i, h in enumerate(hf):
            tf.rows[0].cells[i].text = h
            set_cell_shading(tf.rows[0].cells[i], 'E8F4FD')
        for i, f in enumerate(food_data):
            for j, v in enumerate(f):
                tf.rows[i+1].cells[j].text = v
        doc.add_paragraph('')

    # Tips
    if tips:
        doc.add_heading('🎒 实用Tips', level=1)
        for tip in tips:
            doc.add_paragraph(f'• {tip}')

    # 免责声明
    doc.add_paragraph('')
    doc.add_paragraph('📌 免责声明：地点、交通信息来源于百度地图开放平台，住宿、饮食推荐仅供参考。建议出发前参考小红书、携程等平台做进一步细化攻略。')

    doc.save(output_path)
    print(f"文档已生成: {output_path}")
    return output_path


def main():
    if len(sys.argv) < 4:
        print("用法: python3 generate_travel_guide.py <目的地> <天数> <晚数> [百度AK]")
        sys.exit(1)

    city = sys.argv[1]
    days = int(sys.argv[2])
    nights = int(sys.argv[3])
    ak = sys.argv[4] if len(sys.argv) > 4 else None

    if not ak:
        print("错误：需要提供百度地图API密钥")
        sys.exit(1)

    output_path = f"/root/.openclaw/workspace/{city}{days}天{nights}晚旅游攻略.docx"
    print(f"开始为{city}生成{days}天{nights}晚旅游攻略...")
    print(f"AK: {ak[:10]}...")

    # 这里是需要根据实际搜索结果填充的数据结构
    # 实际使用时，由调用者填充真实数据
    return output_path


if __name__ == '__main__':
    main()

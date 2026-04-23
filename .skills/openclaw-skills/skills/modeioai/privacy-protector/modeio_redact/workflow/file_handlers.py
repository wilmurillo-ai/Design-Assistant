#!/usr/bin/env python3
"""File handler implementations for text, DOCX, and PDF workflows."""

from __future__ import annotations

import shutil
import tempfile
from pathlib import Path
from typing import Any, Iterable, List, Sequence, Tuple

from modeio_redact.core.models import MappingEntry
from modeio_redact.core.replacement import build_replacement_pairs

from modeio_redact.workflow.file_types import (
    HANDLER_DOCX,
    HANDLER_PDF,
    HANDLER_TEXT,
    handler_key_for_extension,
    normalize_extension,
)


def uses_text_handler(extension: str) -> bool:
    return handler_key_for_extension(extension) == HANDLER_TEXT


def validate_non_text_output_extension(input_extension: str, output_path: Path) -> None:
    normalized = normalize_extension(input_extension)
    if uses_text_handler(normalized):
        return
    output_extension = normalize_extension(output_path.suffix)
    if output_extension != normalized:
        raise ValueError(
            f"Output extension must remain '{normalized}' for {normalized} file input."
        )


def read_input_file(path: Path, extension: str) -> str:
    handler_key = handler_key_for_extension(extension)
    if handler_key == HANDLER_TEXT:
        return _read_utf8_text(path)
    if handler_key == HANDLER_DOCX:
        return _read_docx_text(path)
    if handler_key == HANDLER_PDF:
        return _read_pdf_text(path)
    raise ValueError(f"Unsupported file handler for extension '{extension}'.")


def write_non_text_anonymized_file(
    input_path: Path,
    output_path: Path,
    extension: str,
    mapping_entries: Sequence[Any],
) -> None:
    normalized_entries = _normalize_entries(mapping_entries)
    handler_key = handler_key_for_extension(extension)
    if handler_key == HANDLER_DOCX:
        _write_docx_with_replacements(
            input_path=input_path,
            output_path=output_path,
            replacements=build_replacement_pairs(normalized_entries, direction="redact"),
        )
        return
    if handler_key == HANDLER_PDF:
        _write_pdf_with_redactions(
            input_path=input_path,
            output_path=output_path,
            targets=_build_pdf_redaction_targets(normalized_entries),
        )
        return
    raise ValueError(f"Unsupported non-text anonymization handler for extension '{extension}'.")


def write_non_text_deanonymized_file(
    input_path: Path,
    output_path: Path,
    extension: str,
    mapping_entries: Sequence[Any],
) -> None:
    normalized_entries = _normalize_entries(mapping_entries)
    handler_key = handler_key_for_extension(extension)
    if handler_key == HANDLER_DOCX:
        _write_docx_with_replacements(
            input_path=input_path,
            output_path=output_path,
            replacements=build_replacement_pairs(normalized_entries, direction="restore"),
        )
        return
    if handler_key == HANDLER_PDF:
        raise ValueError("De-anonymization is not supported for '.pdf' files.")
    raise ValueError(f"Unsupported non-text de-anonymization handler for extension '{extension}'.")


def _read_utf8_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError("Input file must be UTF-8 encoded.") from exc
    except OSError as exc:
        raise ValueError(f"Failed to read input file: {exc}") from exc


def _import_docx_document_class():
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise ValueError(
            "python-docx package is required for '.docx' support. Install python-docx to enable DOCX workflows."
        ) from exc
    return Document


def _import_fitz_module():
    try:
        import fitz
    except ModuleNotFoundError as exc:
        raise ValueError(
            "PyMuPDF package is required for '.pdf' support. Install PyMuPDF to enable PDF workflows."
        ) from exc
    return fitz


def _iter_table_paragraphs(tables) -> Iterable:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                yield from _iter_table_paragraphs(cell.tables)


def _iter_docx_paragraphs(document) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    yield from _iter_table_paragraphs(document.tables)

    for section in document.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            yield paragraph
        yield from _iter_table_paragraphs(header.tables)
        for paragraph in footer.paragraphs:
            yield paragraph
        yield from _iter_table_paragraphs(footer.tables)


def _read_docx_text(path: Path) -> str:
    document_class = _import_docx_document_class()
    try:
        document = document_class(str(path))
    except Exception as exc:  # pragma: no cover - library-specific error hierarchy
        raise ValueError(f"Failed to read DOCX file: {exc}") from exc

    lines: List[str] = []
    for paragraph in _iter_docx_paragraphs(document):
        text = paragraph.text
        if text:
            lines.append(text)
    return "\n".join(lines).strip()


def _normalize_entries(mapping_entries: Sequence[Any]) -> List[MappingEntry]:
    normalized: List[MappingEntry] = []
    for raw in mapping_entries:
        if isinstance(raw, MappingEntry):
            normalized.append(raw)
            continue
        if isinstance(raw, dict):
            entry = MappingEntry.from_raw(raw)
            if entry is not None:
                normalized.append(entry)
    return normalized


def _find_replacement_spans(
    text: str,
    replacements: Sequence[Tuple[str, str]],
) -> List[Tuple[int, int, str]]:
    spans: List[Tuple[int, int, str]] = []

    candidates = [(source, target) for source, target in replacements if source]
    candidates.sort(key=lambda item: len(item[0]), reverse=True)

    for source, target in candidates:
        search_from = 0
        while True:
            start = text.find(source, search_from)
            if start < 0:
                break
            end = start + len(source)
            search_from = start + 1
            if any(start < existing_end and end > existing_start for existing_start, existing_end, _ in spans):
                continue
            spans.append((start, end, target))

    spans.sort(key=lambda item: (item[0], item[1]))
    return spans


def _apply_replacement_spans_to_text(
    text: str,
    spans: Sequence[Tuple[int, int, str]],
) -> str:
    if not spans:
        return text

    pieces: List[str] = []
    cursor = 0
    for start, end, target in spans:
        if cursor < start:
            pieces.append(text[cursor:start])
        pieces.append(target)
        cursor = end
    if cursor < len(text):
        pieces.append(text[cursor:])
    return "".join(pieces)


def _apply_replacements_to_paragraph(paragraph, replacements: Sequence[Tuple[str, str]]) -> None:
    if not replacements:
        return

    paragraph_text = paragraph.text
    spans = _find_replacement_spans(paragraph_text, replacements)
    if not spans:
        return

    if not paragraph.runs:
        paragraph.text = _apply_replacement_spans_to_text(paragraph_text, spans)
        return

    run_boundaries: List[Tuple[Any, str, int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        run_text = run.text or ""
        start = cursor
        end = start + len(run_text)
        run_boundaries.append((run, run_text, start, end))
        cursor = end

    if cursor != len(paragraph_text):
        paragraph.text = _apply_replacement_spans_to_text(paragraph_text, spans)
        return

    for run, run_text, run_start, run_end in run_boundaries:
        if not run_text:
            run.text = ""
            continue

        pieces: List[str] = []
        local_cursor = run_start
        for start, end, target in spans:
            if end <= run_start:
                continue
            if start >= run_end:
                break

            overlap_start = max(start, run_start)
            overlap_end = min(end, run_end)

            if local_cursor < overlap_start:
                pieces.append(run_text[local_cursor - run_start: overlap_start - run_start])

            if run_start <= start < run_end:
                pieces.append(target)

            local_cursor = max(local_cursor, overlap_end)

        if local_cursor < run_end:
            pieces.append(run_text[local_cursor - run_start: run_end - run_start])

        run.text = "".join(pieces)


def _write_docx_with_replacements(
    input_path: Path,
    output_path: Path,
    replacements: Sequence[Tuple[str, str]],
) -> None:
    document_class = _import_docx_document_class()
    try:
        document = document_class(str(input_path))
    except Exception as exc:  # pragma: no cover - library-specific error hierarchy
        raise ValueError(f"Failed to read DOCX file: {exc}") from exc

    for paragraph in _iter_docx_paragraphs(document):
        _apply_replacements_to_paragraph(paragraph, replacements)

    try:
        document.save(str(output_path))
    except OSError as exc:
        raise ValueError(f"Failed to write DOCX file: {exc}") from exc


def _read_pdf_text(path: Path) -> str:
    fitz = _import_fitz_module()
    try:
        document = fitz.open(str(path))
    except Exception as exc:  # pragma: no cover - library-specific error hierarchy
        raise ValueError(f"Failed to read PDF file: {exc}") from exc

    page_texts: List[str] = []
    has_text_layer = False
    try:
        for page in document:
            words = page.get_text("words")
            if words:
                has_text_layer = True
            page_text = page.get_text("text") or ""
            if page_text:
                page_texts.append(page_text.rstrip("\n"))
    finally:
        document.close()

    if not has_text_layer:
        raise ValueError("PDF input must contain an extractable text layer (image-only PDFs are not supported).")

    return "\n\n".join(page_texts).strip()


def _copy_binary_file(input_path: Path, output_path: Path) -> None:
    if input_path.resolve() == output_path.resolve():
        return
    shutil.copyfile(str(input_path), str(output_path))


def _build_pdf_redaction_targets(mapping_entries: Sequence[MappingEntry]) -> List[str]:
    targets: List[str] = []
    seen = set()
    for entry in mapping_entries:
        value = entry.original.strip()
        if not value or value in seen:
            continue
        seen.add(value)
        targets.append(value)
    targets.sort(key=len, reverse=True)
    return targets


def _write_pdf_with_redactions(
    input_path: Path,
    output_path: Path,
    targets: Sequence[str],
) -> None:
    fitz = _import_fitz_module()

    if not targets:
        _copy_binary_file(input_path, output_path)
        return

    try:
        document = fitz.open(str(input_path))
    except Exception as exc:  # pragma: no cover - library-specific error hierarchy
        raise ValueError(f"Failed to read PDF file: {exc}") from exc

    has_text_layer = False
    try:
        for page in document:
            words = page.get_text("words")
            if words:
                has_text_layer = True
            for target in targets:
                rects = page.search_for(target)
                for rect in rects:
                    page.add_redact_annot(rect, fill=(0, 0, 0))
            page.apply_redactions()

        if not has_text_layer:
            raise ValueError(
                "PDF input must contain an extractable text layer (image-only PDFs are not supported)."
            )

        save_path = output_path
        if input_path.resolve() == output_path.resolve():
            with tempfile.NamedTemporaryFile(
                suffix=".pdf",
                dir=str(output_path.parent),
                delete=False,
            ) as temp_file:
                save_path = Path(temp_file.name)
            save_path.unlink(missing_ok=True)

        document.save(str(save_path), garbage=4, clean=True, deflate=True)
    except OSError as exc:
        raise ValueError(f"Failed to write PDF file: {exc}") from exc
    finally:
        document.close()

    if input_path.resolve() == output_path.resolve():
        temp_path = Path(str(save_path))
        try:
            temp_path.replace(output_path)
        except OSError as exc:
            temp_path.unlink(missing_ok=True)
            raise ValueError(f"Failed to finalize in-place PDF write: {exc}") from exc

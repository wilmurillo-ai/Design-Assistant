#!/usr/bin/env python3

import tempfile
import unittest
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "scripts"
import sys

sys.path.insert(0, str(SCRIPTS_DIR))

from modeio_redact.workflow import file_handlers  # noqa: E402

try:  # noqa: E402
    from docx import Document

    HAS_DOCX = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_DOCX = False

try:  # noqa: E402
    import fitz

    HAS_FITZ = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_FITZ = False


class TestFileHandlers(unittest.TestCase):
    def test_text_handler_detection(self):
        self.assertTrue(file_handlers.uses_text_handler(".txt"))
        self.assertFalse(file_handlers.uses_text_handler(".docx"))
        self.assertFalse(file_handlers.uses_text_handler(".pdf"))

    def test_non_text_output_extension_validation(self):
        with self.assertRaises(ValueError):
            file_handlers.validate_non_text_output_extension(".docx", Path("/tmp/out.txt"))

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_read_and_anonymize_write(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "sample.docx"
            output_path = Path(tmpdir) / "sample.redacted.docx"

            document = Document()
            document.add_paragraph("Email: alice@example.com")
            document.save(str(input_path))

            text = file_handlers.read_input_file(input_path, ".docx")
            self.assertIn("alice@example.com", text)

            file_handlers.write_non_text_anonymized_file(
                input_path=input_path,
                output_path=output_path,
                extension=".docx",
                mapping_entries=[
                    {
                        "placeholder": "[EMAIL_1]",
                        "original": "alice@example.com",
                        "type": "email",
                    }
                ],
            )

            redacted = Document(str(output_path))
            self.assertIn("[EMAIL_1]", "\n".join(p.text for p in redacted.paragraphs))

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_deanonymize_write(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "sample.redacted.docx"
            output_path = Path(tmpdir) / "sample.restored.docx"

            document = Document()
            document.add_paragraph("Email: [EMAIL_1]")
            document.save(str(input_path))

            file_handlers.write_non_text_deanonymized_file(
                input_path=input_path,
                output_path=output_path,
                extension=".docx",
                mapping_entries=[
                    {
                        "placeholder": "[EMAIL_1]",
                        "original": "alice@example.com",
                        "type": "email",
                    }
                ],
            )

            restored = Document(str(output_path))
            self.assertIn("alice@example.com", "\n".join(p.text for p in restored.paragraphs))

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_pdf_read_and_redaction_write(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "sample.pdf"
            output_path = Path(tmpdir) / "sample.redacted.pdf"

            source = fitz.open()
            page = source.new_page()
            page.insert_text((72, 72), "Email: alice@example.com")
            source.save(str(input_path))
            source.close()

            text = file_handlers.read_input_file(input_path, ".pdf")
            self.assertIn("alice@example.com", text)

            file_handlers.write_non_text_anonymized_file(
                input_path=input_path,
                output_path=output_path,
                extension=".pdf",
                mapping_entries=[
                    {
                        "placeholder": "[EMAIL_1]",
                        "original": "alice@example.com",
                        "type": "email",
                    }
                ],
            )

            redacted = fitz.open(str(output_path))
            try:
                extracted = "\n".join(page.get_text("text") for page in redacted)
            finally:
                redacted.close()
            self.assertNotIn("alice@example.com", extracted)

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_preserves_non_matching_run_formatting_for_cross_run_replacement(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "styled.docx"
            output_path = Path(tmpdir) / "styled.redacted.docx"

            document = Document()
            paragraph = document.add_paragraph()
            prefix = paragraph.add_run("Client: ")
            prefix.bold = True
            email_first = paragraph.add_run("alice")
            email_first.italic = True
            email_second = paragraph.add_run("@example.com")
            email_second.underline = True
            suffix = paragraph.add_run(" reviewed")
            suffix.bold = True
            document.save(str(input_path))

            file_handlers.write_non_text_anonymized_file(
                input_path=input_path,
                output_path=output_path,
                extension=".docx",
                mapping_entries=[
                    {
                        "placeholder": "[EMAIL_1]",
                        "original": "alice@example.com",
                        "type": "email",
                    }
                ],
            )

            redacted = Document(str(output_path))
            paragraph = redacted.paragraphs[0]
            self.assertEqual(paragraph.text, "Client: [EMAIL_1] reviewed")
            self.assertEqual(paragraph.runs[0].text, "Client: ")
            self.assertEqual(paragraph.runs[1].text, "[EMAIL_1]")
            self.assertEqual(paragraph.runs[2].text, "")
            self.assertEqual(paragraph.runs[3].text, " reviewed")
            self.assertTrue(paragraph.runs[0].bold)
            self.assertTrue(paragraph.runs[1].italic)
            self.assertTrue(paragraph.runs[2].underline)
            self.assertTrue(paragraph.runs[3].bold)

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_pdf_read_and_redaction_write_removes_long_original(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "sample.pdf"
            output_path = Path(tmpdir) / "sample.redacted.pdf"

            source = fitz.open()
            page = source.new_page()
            page.insert_text((72, 72), "Passport: XH4829916")
            source.save(str(input_path))
            source.close()

            file_handlers.write_non_text_anonymized_file(
                input_path=input_path,
                output_path=output_path,
                extension=".pdf",
                mapping_entries=[
                    {
                        "placeholder": "[REDACTED_ID_NUMBER_1]",
                        "original": "XH4829916",
                        "type": "id_number",
                    }
                ],
            )

            redacted = fitz.open(str(output_path))
            try:
                extracted = "\n".join(page.get_text("text") for page in redacted)
            finally:
                redacted.close()
            self.assertNotIn("XH4829916", extracted)
            self.assertNotIn("[REDACTED_ID_NUMBER_1]", extracted)

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_pdf_deanonymize_unsupported(self):
        with self.assertRaises(ValueError):
            file_handlers.write_non_text_deanonymized_file(
                input_path=Path("/tmp/a.pdf"),
                output_path=Path("/tmp/b.pdf"),
                extension=".pdf",
                mapping_entries=[],
            )


if __name__ == "__main__":
    unittest.main()

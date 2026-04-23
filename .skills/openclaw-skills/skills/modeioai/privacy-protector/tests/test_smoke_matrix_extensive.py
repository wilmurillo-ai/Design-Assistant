#!/usr/bin/env python3

import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
ANONYMIZE_SCRIPT = REPO_ROOT / "privacy-protector" / "scripts" / "anonymize.py"
DEANONYMIZE_SCRIPT = REPO_ROOT / "privacy-protector" / "scripts" / "deanonymize.py"

try:  # pragma: no cover
    from docx import Document

    HAS_DOCX = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_DOCX = False

try:  # pragma: no cover
    import fitz

    HAS_FITZ = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_FITZ = False


def _collect_docx_text(document) -> str:
    lines = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        lines.append(paragraph.text)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text:
                lines.append(paragraph.text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            lines.append(paragraph.text)

        for paragraph in section.footer.paragraphs:
            if paragraph.text:
                lines.append(paragraph.text)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            lines.append(paragraph.text)

    return "\n".join(lines)


class TestSmokeMatrixExtensive(unittest.TestCase):
    maxDiff = None

    def _run_cli(self, script_path: Path, args, env=None):
        merged_env = os.environ.copy()
        if env:
            merged_env.update(env)
        return subprocess.run(
            [sys.executable, str(script_path)] + args,
            capture_output=True,
            text=True,
            env=merged_env,
        )

    def _assert_success_payload(self, result):
        self.assertEqual(
            result.returncode,
            0,
            msg=f"stdout={result.stdout}\nstderr={result.stderr}",
        )
        payload = json.loads(result.stdout)
        self.assertTrue(payload.get("success"), msg=payload)
        return payload

    def _assert_file_contract(self, payload, *, assurance_level: str, verification_skipped: bool):
        data = payload["data"]

        self.assertIn("outputPath", data)
        self.assertTrue(Path(data["outputPath"]).exists())
        self.assertIn("mapRef", data)
        self.assertTrue(Path(data["mapRef"]["sidecarPath"]).exists())

        self.assertIn("applyReport", data)
        apply_report = data["applyReport"]
        self.assertGreaterEqual(apply_report["expectedCount"], 1)
        self.assertEqual(apply_report["expectedCount"], apply_report["appliedCount"])
        self.assertEqual(apply_report["missingCount"], 0)

        self.assertIn("verificationReport", data)
        verification_report = data["verificationReport"]
        self.assertEqual(verification_report["skipped"], verification_skipped)
        if not verification_skipped:
            self.assertTrue(verification_report["passed"])
            self.assertEqual(verification_report["residualCount"], 0)

        self.assertIn("assurancePolicy", data)
        self.assertEqual(data["assurancePolicy"]["level"], assurance_level)

    def test_text_like_file_matrix_lite_roundtrip(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {"MODEIO_REDACT_MAP_DIR": str(tmp_path / "maps")}

            cases = [
                (".txt", "Email: alice@example.com\nPhone: 415-555-1234\n"),
                (".md", "# Contact\nEmail: alice@example.com\nPhone: 415-555-1234\n"),
                (".json", '{"email": "alice@example.com", "phone": "415-555-1234"}\n'),
                (".yaml", "email: alice@example.com\nphone: 415-555-1234\n"),
                (".xml", "<root><email>alice@example.com</email><phone>415-555-1234</phone></root>\n"),
                (".csv", "email,phone\nalice@example.com,415-555-1234\n"),
                (".tsv", "email\tphone\nalice@example.com\t415-555-1234\n"),
                (".html", "<p>Email: alice@example.com</p><p>Phone: 415-555-1234</p>\n"),
            ]

            for index, (extension, content) in enumerate(cases):
                with self.subTest(extension=extension):
                    input_path = tmp_path / f"sample-{index}{extension}"
                    input_path.write_text(content, encoding="utf-8")

                    anonymize_result = self._run_cli(
                        ANONYMIZE_SCRIPT,
                        [
                            "--input",
                            str(input_path),
                            "--level",
                            "lite",
                            "--json",
                        ],
                        env=env,
                    )
                    anonymize_payload = self._assert_success_payload(anonymize_result)

                    self.assertEqual(anonymize_payload["mode"], "local-regex")
                    self.assertEqual(anonymize_payload["level"], "lite")
                    self._assert_file_contract(
                        anonymize_payload,
                        assurance_level="best_effort",
                        verification_skipped=True,
                    )

                    output_path = Path(anonymize_payload["data"]["outputPath"])
                    output_text = output_path.read_text(encoding="utf-8")
                    self.assertNotIn("alice@example.com", output_text)
                    self.assertNotIn("415-555-1234", output_text)

                    deanonymize_result = self._run_cli(
                        DEANONYMIZE_SCRIPT,
                        [
                            "--input",
                            str(output_path),
                            "--json",
                        ],
                        env=env,
                    )
                    deanonymize_payload = self._assert_success_payload(deanonymize_result)
                    self.assertEqual(deanonymize_payload["mode"], "local-map")
                    restored = deanonymize_payload["data"]["deanonymizedContent"]
                    self.assertIn("alice@example.com", restored)
                    self.assertIn("415-555-1234", restored)

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_default_verified_multiregion_roundtrip(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {"MODEIO_REDACT_MAP_DIR": str(tmp_path / "maps")}

            input_path = tmp_path / "incident.docx"
            document = Document()
            document.add_paragraph("Body Email: alice@example.com")
            document.add_paragraph("Body Phone: 415-555-1234")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Table Email: alice@example.com"

            section = document.sections[0]
            section.header.paragraphs[0].text = "Header Phone: 415-555-1234"
            section.footer.paragraphs[0].text = "Footer Email: alice@example.com"
            document.save(str(input_path))

            anonymize_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            anonymize_payload = self._assert_success_payload(anonymize_result)
            self._assert_file_contract(
                anonymize_payload,
                assurance_level="verified",
                verification_skipped=False,
            )

            output_path = Path(anonymize_payload["data"]["outputPath"])
            redacted_doc = Document(str(output_path))
            redacted_text = _collect_docx_text(redacted_doc)
            self.assertNotIn("alice@example.com", redacted_text)
            self.assertNotIn("415-555-1234", redacted_text)
            self.assertIn("[EMAIL_1]", redacted_text)
            self.assertIn("[PHONE_1]", redacted_text)

            deanonymize_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(output_path),
                    "--json",
                ],
                env=env,
            )
            deanonymize_payload = self._assert_success_payload(deanonymize_result)
            warning_codes = [w.get("code") for w in deanonymize_payload["data"].get("warnings", []) if isinstance(w, dict)]
            self.assertNotIn("input_hash_mismatch", warning_codes)
            restored_path = Path(deanonymize_payload["data"]["outputPath"])
            restored_doc = Document(str(restored_path))
            restored_text = _collect_docx_text(restored_doc)
            self.assertIn("alice@example.com", restored_text)
            self.assertIn("415-555-1234", restored_text)

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_pdf_default_verified_multipage_smoke(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {"MODEIO_REDACT_MAP_DIR": str(tmp_path / "maps")}

            input_path = tmp_path / "incident.pdf"
            pdf_document = fitz.open()
            page_one = pdf_document.new_page()
            page_one.insert_text((72, 72), "Email: alice@example.com")
            page_one.insert_text((72, 100), "Phone: 415-555-1234")
            page_two = pdf_document.new_page()
            page_two.insert_text((72, 72), "Email: alice@example.com")
            page_two.insert_text((72, 100), "Phone: 415-555-1234")
            pdf_document.save(str(input_path))
            pdf_document.close()

            anonymize_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            anonymize_payload = self._assert_success_payload(anonymize_result)
            self._assert_file_contract(
                anonymize_payload,
                assurance_level="verified",
                verification_skipped=False,
            )

            output_path = Path(anonymize_payload["data"]["outputPath"])
            redacted_document = fitz.open(str(output_path))
            try:
                redacted_text = "\n".join(page.get_text("text") for page in redacted_document)
            finally:
                redacted_document.close()

            self.assertNotIn("alice@example.com", redacted_text)
            self.assertNotIn("415-555-1234", redacted_text)
            apply_report = anonymize_payload["data"]["applyReport"]
            self.assertGreaterEqual(apply_report["expectedCount"], 4)
            self.assertEqual(apply_report["expectedCount"], apply_report["appliedCount"])

    def test_text_in_place_roundtrip_smoke(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {"MODEIO_REDACT_MAP_DIR": str(tmp_path / "maps")}

            input_path = tmp_path / "notes.txt"
            original_content = "Email: alice@example.com\nPhone: 415-555-1234\n"
            input_path.write_text(original_content, encoding="utf-8")

            anonymize_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--in-place",
                    "--json",
                ],
                env=env,
            )
            anonymize_payload = self._assert_success_payload(anonymize_result)
            self.assertEqual(anonymize_payload["data"]["outputPath"], str(input_path))
            self.assertIn("[EMAIL_1]", input_path.read_text(encoding="utf-8"))

            deanonymize_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--in-place",
                    "--json",
                ],
                env=env,
            )
            deanonymize_payload = self._assert_success_payload(deanonymize_result)
            self.assertEqual(deanonymize_payload["data"]["outputPath"], str(input_path))
            restored_content = input_path.read_text(encoding="utf-8")
            self.assertIn("alice@example.com", restored_content)
            self.assertIn("415-555-1234", restored_content)

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_in_place_roundtrip_smoke(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {"MODEIO_REDACT_MAP_DIR": str(tmp_path / "maps")}

            input_path = tmp_path / "notes.docx"
            document = Document()
            document.add_paragraph("Email: alice@example.com")
            document.add_paragraph("Phone: 415-555-1234")
            document.save(str(input_path))

            anonymize_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--in-place",
                    "--json",
                ],
                env=env,
            )
            anonymize_payload = self._assert_success_payload(anonymize_result)
            self.assertEqual(anonymize_payload["data"]["outputPath"], str(input_path))

            in_place_doc = Document(str(input_path))
            in_place_text = _collect_docx_text(in_place_doc)
            self.assertIn("[EMAIL_1]", in_place_text)
            self.assertIn("[PHONE_1]", in_place_text)

            deanonymize_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--in-place",
                    "--json",
                ],
                env=env,
            )
            deanonymize_payload = self._assert_success_payload(deanonymize_result)
            warning_codes = [w.get("code") for w in deanonymize_payload["data"].get("warnings", []) if isinstance(w, dict)]
            self.assertNotIn("input_hash_mismatch", warning_codes)

            restored_doc = Document(str(input_path))
            restored_text = _collect_docx_text(restored_doc)
            self.assertIn("alice@example.com", restored_text)
            self.assertIn("415-555-1234", restored_text)

    def test_api_dynamic_and_crossborder_smoke(self):
        if os.environ.get("MODEIO_REDACT_SKIP_API_SMOKE") == "1":
            self.skipTest("set MODEIO_REDACT_SKIP_API_SMOKE=1 to disable API smoke")

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            env = {"MODEIO_REDACT_MAP_DIR": str(tmp_path / "maps")}

            input_path = tmp_path / "api-sample.txt"
            input_path.write_text("Email: alice@example.com\nPhone: 415-555-1234\n", encoding="utf-8")

            dynamic_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "dynamic",
                    "--json",
                ],
                env=env,
            )
            dynamic_payload = self._assert_success_payload(dynamic_result)
            self.assertEqual(dynamic_payload["mode"], "api")
            self.assertEqual(dynamic_payload["level"], "dynamic")
            self.assertTrue(Path(dynamic_payload["data"]["outputPath"]).exists())

            crossborder_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "crossborder",
                    "--sender-code",
                    "CN SHA",
                    "--recipient-code",
                    "US NYC",
                    "--json",
                ],
                env=env,
            )
            crossborder_payload = self._assert_success_payload(crossborder_result)
            self.assertEqual(crossborder_payload["mode"], "api")
            self.assertEqual(crossborder_payload["level"], "crossborder")
            self.assertIn("crossBorderAnalysis", crossborder_payload["data"])


if __name__ == "__main__":
    unittest.main()

#!/usr/bin/env python3

import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
ANONYMIZE_SCRIPT = REPO_ROOT / "privacy-protector" / "scripts" / "anonymize.py"
DEANONYMIZE_SCRIPT = REPO_ROOT / "privacy-protector" / "scripts" / "deanonymize.py"

try:  # noqa: E402
    from docx import Document

    HAS_DOCX = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_DOCX = False

try:  # noqa: E402
    import fitz

    HAS_FITZ = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_FITZ = False


class TestDeanonymizeContract(unittest.TestCase):
    def _run_cli(self, script_path, args, env=None):
        merged_env = os.environ.copy()
        if env:
            merged_env.update(env)
        return subprocess.run(
            [sys.executable, str(script_path)] + args,
            capture_output=True,
            text=True,
            env=merged_env,
        )

    def test_deanonymize_uses_latest_map_by_default(self):
        source_text = "Email: alice@example.com"

        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": tmpdir}
            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    source_text,
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            anonymized_text = anon_payload["data"]["anonymizedContent"]

            deanonymize_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    anonymized_text,
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(deanonymize_result.returncode, 0)

            payload = json.loads(deanonymize_result.stdout)
            self.assertTrue(payload["success"])
            self.assertEqual(payload["mode"], "local-map")
            self.assertEqual(payload["data"]["deanonymizedContent"], source_text)

    def test_deanonymize_accepts_map_id_reference(self):
        source_text = "Phone: 415-555-1234"

        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": tmpdir}
            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    source_text,
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            anonymized_text = anon_payload["data"]["anonymizedContent"]
            map_id = anon_payload["data"]["mapRef"]["mapId"]

            deanonymize_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    anonymized_text,
                    "--map",
                    map_id,
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(deanonymize_result.returncode, 0)
            payload = json.loads(deanonymize_result.stdout)

            self.assertTrue(payload["success"])
            self.assertEqual(payload["data"]["deanonymizedContent"], source_text)
            self.assertEqual(payload["data"]["mapRef"]["mapId"], map_id)

    def test_deanonymize_accepts_map_path_reference(self):
        source_text = "Email: alice@example.com"

        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": tmpdir}
            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    source_text,
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            anonymized_text = anon_payload["data"]["anonymizedContent"]
            map_path = anon_payload["data"]["mapRef"]["mapPath"]

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    anonymized_text,
                    "--map",
                    map_path,
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            self.assertTrue(payload["success"])
            self.assertEqual(payload["data"]["deanonymizedContent"], source_text)

    def test_deanonymize_accepts_txt_file_input(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "file-map",
                        "entries": [
                            {"placeholder": "[EMAIL_1]", "original": "alice@example.com", "type": "email"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            input_path = Path(tmpdir) / "anonymized.txt"
            input_path.write_text("Email: [EMAIL_1]", encoding="utf-8")

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            self.assertTrue(payload["success"])
            self.assertEqual(payload["data"]["deanonymizedContent"], "Email: alice@example.com")

    def test_deanonymize_accepts_markdown_file_input(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "md-map",
                        "entries": [
                            {"placeholder": "[PHONE_1]", "original": "415-555-1234", "type": "phone"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            input_path = Path(tmpdir) / "anonymized.md"
            input_path.write_text("Call [PHONE_1]", encoding="utf-8")

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            self.assertTrue(payload["success"])
            self.assertEqual(payload["data"]["deanonymizedContent"], "Call 415-555-1234")

    def test_deanonymize_accepts_json_file_input(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "json-map",
                        "entries": [
                            {"placeholder": "[EMAIL_1]", "original": "alice@example.com", "type": "email"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            input_path = Path(tmpdir) / "anonymized.json"
            input_path.write_text('{"value":"[EMAIL_1]"}', encoding="utf-8")

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            self.assertTrue(payload["success"])
            self.assertEqual(payload["data"]["deanonymizedContent"], '{"value":"alice@example.com"}')

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_deanonymize_accepts_docx_file_input(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "docx-map",
                        "entries": [
                            {"placeholder": "[EMAIL_1]", "original": "alice@example.com", "type": "email"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            input_path = Path(tmpdir) / "anonymized.docx"
            document = Document()
            document.add_paragraph("Email: [EMAIL_1]")
            document.save(str(input_path))

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            self.assertTrue(payload["success"])
            self.assertIn("alice@example.com", payload["data"]["deanonymizedContent"])
            output_path = Path(payload["data"]["outputPath"])
            restored = Document(str(output_path))
            restored_text = "\n".join(paragraph.text for paragraph in restored.paragraphs)
            self.assertIn("alice@example.com", restored_text)

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_deanonymize_rejects_pdf_file_input_type(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "json-map",
                        "entries": [
                            {"placeholder": "[EMAIL_1]", "original": "alice@example.com", "type": "email"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            input_path = Path(tmpdir) / "anonymized.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Email: [EMAIL_1]")
            document.save(str(input_path))
            document.close()

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 2)
            payload = json.loads(result.stdout)

            self.assertFalse(payload["success"])
            self.assertEqual(payload["error"]["type"], "validation_error")

    def test_json_file_input_auto_resolves_map_from_sidecar(self):
        original_text = '{"email":"alice@example.com"}'
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "incident.json"
            input_path.write_text(original_text, encoding="utf-8")
            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}

            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            redacted_path = Path(anon_payload["data"]["outputPath"])
            redacted_content = redacted_path.read_text(encoding="utf-8")
            self.assertNotIn("privacy-protector-map-id", redacted_content)
            self.assertTrue(Path(anon_payload["data"]["mapRef"]["sidecarPath"]).exists())

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 0)
            restore_payload = json.loads(restore_result.stdout)

            self.assertTrue(restore_payload["success"])
            self.assertEqual(restore_payload["data"]["linkageSource"], "sidecar")
            self.assertEqual(restore_payload["data"]["deanonymizedContent"], original_text)

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_file_input_auto_resolves_map_from_sidecar(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "incident.docx"
            document = Document()
            document.add_paragraph("Email: alice@example.com")
            document.save(str(input_path))

            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}
            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            redacted_path = Path(anon_payload["data"]["outputPath"])
            self.assertEqual(redacted_path.suffix.lower(), ".docx")

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 0)
            restore_payload = json.loads(restore_result.stdout)

            self.assertTrue(restore_payload["success"])
            self.assertEqual(restore_payload["data"]["linkageSource"], "sidecar")
            restored_path = Path(restore_payload["data"]["outputPath"])
            restored_doc = Document(str(restored_path))
            restored_text = "\n".join(paragraph.text for paragraph in restored_doc.paragraphs)
            self.assertIn("alice@example.com", restored_text)

    def test_file_input_auto_resolves_map_from_embedded_marker(self):
        original_text = "Email: alice@example.com"
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "incident.txt"
            input_path.write_text(original_text, encoding="utf-8")
            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}

            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            redacted_path = Path(anon_payload["data"]["outputPath"])
            redacted_content = redacted_path.read_text(encoding="utf-8")
            self.assertIn("privacy-protector-map-id", redacted_content)

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 0)
            restore_payload = json.loads(restore_result.stdout)

            self.assertTrue(restore_payload["success"])
            self.assertEqual(restore_payload["data"]["linkageSource"], "embedded-mapid")
            self.assertEqual(restore_payload["data"]["deanonymizedContent"], original_text)
            restored_path = Path(restore_payload["data"]["outputPath"])
            self.assertEqual(restored_path.read_text(encoding="utf-8"), original_text)

    def test_file_input_auto_resolves_map_from_legacy_embedded_marker(self):
        original_text = "Email: alice@example.com"
        redacted_text = "# modeio-redact-map-id: legacy-map\nEmail: [EMAIL_1]"

        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}
            redacted_path = Path(tmpdir) / "incident.redacted.txt"
            redacted_path.write_text(redacted_text, encoding="utf-8")

            maps_dir = Path(env["MODEIO_REDACT_MAP_DIR"])
            maps_dir.mkdir(parents=True, exist_ok=True)
            (maps_dir / "legacy-map.json").write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "legacy-map",
                        "entries": [
                            {"placeholder": "[EMAIL_1]", "original": "alice@example.com", "type": "email"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 0)
            restore_payload = json.loads(restore_result.stdout)

            self.assertTrue(restore_payload["success"])
            self.assertEqual(restore_payload["data"]["linkageSource"], "embedded-mapid")
            self.assertEqual(restore_payload["data"]["deanonymizedContent"], original_text)
            restored_path = Path(restore_payload["data"]["outputPath"])
            self.assertEqual(restored_path.read_text(encoding="utf-8"), original_text)

    def test_file_input_falls_back_to_sidecar_when_marker_missing(self):
        original_text = "Email: alice@example.com"
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "incident.txt"
            input_path.write_text(original_text, encoding="utf-8")
            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}

            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            redacted_path = Path(anon_payload["data"]["outputPath"])
            sidecar_path = Path(anon_payload["data"]["mapRef"]["sidecarPath"])
            self.assertTrue(sidecar_path.exists())

            lines = redacted_path.read_text(encoding="utf-8").splitlines()
            redacted_path.write_text("\n".join(lines[1:]), encoding="utf-8")

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 0)
            restore_payload = json.loads(restore_result.stdout)

            self.assertTrue(restore_payload["success"])
            self.assertEqual(restore_payload["data"]["linkageSource"], "sidecar")
            self.assertEqual(restore_payload["data"]["deanonymizedContent"], original_text)

    def test_file_input_requires_map_reference_when_marker_and_sidecar_missing(self):
        original_text = "Email: alice@example.com"
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "incident.txt"
            input_path.write_text(original_text, encoding="utf-8")
            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}

            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            redacted_path = Path(anon_payload["data"]["outputPath"])
            sidecar_path = Path(anon_payload["data"]["mapRef"]["sidecarPath"])

            lines = redacted_path.read_text(encoding="utf-8").splitlines()
            redacted_path.write_text("\n".join(lines[1:]), encoding="utf-8")
            sidecar_path.unlink()

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 1)
            restore_payload = json.loads(restore_result.stdout)

            self.assertFalse(restore_payload["success"])
            self.assertEqual(restore_payload["error"]["type"], "map_error")

    def test_deanonymize_in_place_overwrites_input_file(self):
        original_text = "Email: alice@example.com"
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "incident.txt"
            input_path.write_text(original_text, encoding="utf-8")
            env = {"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")}

            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(input_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            redacted_path = Path(json.loads(anon_result.stdout)["data"]["outputPath"])

            restore_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    str(redacted_path),
                    "--in-place",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(restore_result.returncode, 0)
            restore_payload = json.loads(restore_result.stdout)

            self.assertEqual(restore_payload["data"]["outputPath"], str(redacted_path))
            self.assertEqual(redacted_path.read_text(encoding="utf-8"), original_text)

    def test_deanonymize_reports_map_error_when_no_maps_exist(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": tmpdir}
            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    "Email: [EMAIL_1]",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(result.returncode, 1)
            payload = json.loads(result.stdout)
            self.assertFalse(payload["success"])
            self.assertEqual(payload["error"]["type"], "map_error")

    def test_deanonymize_validation_error_on_empty_input(self):
        result = self._run_cli(
            DEANONYMIZE_SCRIPT,
            [
                "--input",
                "   ",
                "--json",
            ],
        )
        self.assertEqual(result.returncode, 2)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual(payload["error"]["type"], "validation_error")

    def test_deanonymize_in_place_requires_file_input(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": tmpdir}
            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    "Email: alice@example.com",
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            map_id = anon_payload["data"]["mapRef"]["mapId"]
            anonymized_text = anon_payload["data"]["anonymizedContent"]

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    anonymized_text,
                    "--map",
                    map_id,
                    "--in-place",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(result.returncode, 2)
            payload = json.loads(result.stdout)
            self.assertFalse(payload["success"])
            self.assertEqual(payload["error"]["type"], "validation_error")

    def test_deanonymize_returns_map_error_for_invalid_map_json(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            bad_map = Path(tmpdir) / "bad.json"
            bad_map.write_text("{broken", encoding="utf-8")

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    "Email: [EMAIL_1]",
                    "--map",
                    str(bad_map),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 1)
            payload = json.loads(result.stdout)

            self.assertFalse(payload["success"])
            self.assertEqual(payload["error"]["type"], "map_error")

    def test_replacement_summary_counts_repeated_tokens(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "summary-map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "summary-map",
                        "entries": [
                            {"placeholder": "[EMAIL_1]", "original": "alice@example.com", "type": "email"},
                            {"placeholder": "[PHONE_1]", "original": "415-555-1234", "type": "phone"},
                        ],
                    }
                ),
                encoding="utf-8",
            )

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    "[EMAIL_1] and [PHONE_1] and [EMAIL_1]",
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            summary = payload["data"]["replacementSummary"]
            self.assertEqual(summary["totalReplacements"], 3)
            self.assertEqual(summary["replacementsByType"]["email"], 2)
            self.assertEqual(summary["replacementsByType"]["phone"], 1)

    def test_hash_mismatch_allows_by_default_with_warning(self):
        source_text = "Email: alice@example.com"

        with tempfile.TemporaryDirectory() as tmpdir:
            env = {"MODEIO_REDACT_MAP_DIR": tmpdir}
            anon_result = self._run_cli(
                ANONYMIZE_SCRIPT,
                [
                    "--input",
                    source_text,
                    "--level",
                    "lite",
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(anon_result.returncode, 0)
            anon_payload = json.loads(anon_result.stdout)
            anonymized_text = anon_payload["data"]["anonymizedContent"] + " extra"

            deanonymize_result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    anonymized_text,
                    "--json",
                ],
                env=env,
            )
            self.assertEqual(deanonymize_result.returncode, 0)
            payload = json.loads(deanonymize_result.stdout)

            self.assertTrue(payload["success"])
            warning_codes = [item["code"] for item in payload["data"].get("warnings", [])]
            self.assertIn("input_hash_mismatch", warning_codes)

    def test_removed_hash_policy_flags_are_rejected(self):
        result_require = self._run_cli(
            DEANONYMIZE_SCRIPT,
            [
                "--input",
                "Email: [EMAIL_1]",
                "--require-hash-match",
            ],
        )
        self.assertNotEqual(result_require.returncode, 0)
        self.assertIn("unrecognized arguments: --require-hash-match", result_require.stderr)

        result_allow = self._run_cli(
            DEANONYMIZE_SCRIPT,
            [
                "--input",
                "Email: [EMAIL_1]",
                "--allow-hash-mismatch",
            ],
        )
        self.assertNotEqual(result_allow.returncode, 0)
        self.assertIn("unrecognized arguments: --allow-hash-mismatch", result_allow.stderr)

    def test_longer_placeholder_replaced_first(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            map_path = Path(tmpdir) / "custom-map.json"
            map_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "1",
                        "mapId": "custom-map",
                        "createdAt": "2026-03-04T00:00:00Z",
                        "anonymizedHash": "",
                        "entries": [
                            {"placeholder": "[NAME_1]", "original": "Alice", "type": "name"},
                            {"placeholder": "[NAME_10]", "original": "Bob", "type": "name"},
                        ],
                    }
                ),
                encoding="utf-8",
            )

            result = self._run_cli(
                DEANONYMIZE_SCRIPT,
                [
                    "--input",
                    "[NAME_10] met [NAME_1]",
                    "--map",
                    str(map_path),
                    "--json",
                ],
            )
            self.assertEqual(result.returncode, 0)

            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertEqual(payload["data"]["deanonymizedContent"], "Bob met Alice")


if __name__ == "__main__":
    unittest.main()

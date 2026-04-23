#!/usr/bin/env python3

import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import Mock, patch

REPO_ROOT = Path(__file__).resolve().parents[2]
SCRIPT_PATH = REPO_ROOT / "privacy-protector" / "scripts" / "anonymize.py"
SCRIPTS_DIR = REPO_ROOT / "privacy-protector" / "scripts"

sys.path.insert(0, str(SCRIPTS_DIR))
import anonymize  # noqa: E402
from modeio_redact.workflow.map_store import hash_text, load_map  # noqa: E402

try:  # noqa: E402
    from docx import Document

    HAS_DOCX = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_DOCX = False

try:  # noqa: E402
    import fitz

    HAS_FITZ = True
except ModuleNotFoundError:  # pragma: no cover
    HAS_FITZ = False


class TestAnonymizeContract(unittest.TestCase):
    def _run_cli(self, args, env=None):
        merged_env = os.environ.copy()
        if env:
            merged_env.update(env)
        return subprocess.run(
            [sys.executable, str(SCRIPT_PATH)] + args,
            capture_output=True,
            text=True,
            env=merged_env,
        )

    def test_input_type_flag_removed(self):
        result = self._run_cli([
            "--input",
            "Email: alice@example.com",
            "--input-type",
            "file",
        ])
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("unrecognized arguments: --input-type", result.stderr)

    def test_crossborder_requires_explicit_codes_in_json_mode(self):
        result = self._run_cli([
            "--input",
            "Name: John Doe, SSN: 123-45-6789",
            "--level",
            "crossborder",
            "--json",
        ])
        self.assertEqual(result.returncode, 2)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual(payload["tool"], "privacy-protector")
        self.assertEqual(payload["level"], "crossborder")
        self.assertEqual(payload["error"]["type"], "validation_error")

    def test_json_success_envelope_for_lite(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._run_cli(
                [
                    "--input",
                    "Email: alice@example.com, Phone: 415-555-1234",
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": tmpdir},
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)

            self.assertTrue(payload["success"])
            self.assertEqual(payload["tool"], "privacy-protector")
            self.assertEqual(payload["mode"], "local-regex")
            self.assertEqual(payload["level"], "lite")
            self.assertIn("anonymizedContent", payload["data"])
            self.assertIn("hasPII", payload["data"])
            self.assertIn("mapRef", payload["data"])
            self.assertGreater(payload["data"]["mapRef"]["entryCount"], 0)
            self.assertTrue(Path(payload["data"]["mapRef"]["mapPath"]).exists())

    def test_json_success_without_pii_has_no_map_ref(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._run_cli(
                [
                    "--input",
                    "Completely harmless sentence without personal data.",
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": tmpdir},
            )
            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertFalse(payload["data"]["hasPII"])
            self.assertNotIn("mapRef", payload["data"])

    def test_json_network_error_envelope_for_api_mode(self):
        result = self._run_cli(
            [
                "--input",
                "Email: alice@example.com",
                "--level",
                "dynamic",
                "--json",
            ],
            env={"ANONYMIZE_API_URL": "http://127.0.0.1:9"},
        )
        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)

        self.assertFalse(payload["success"])
        self.assertEqual(payload["mode"], "api")
        self.assertEqual(payload["level"], "dynamic")
        self.assertIn(payload["error"]["type"], ("network_error", "dependency_error"))
        if payload["error"]["type"] == "dependency_error":
            self.assertIn("requests package is required", payload["error"]["message"])

    def test_txt_file_path_is_auto_resolved_and_redacted(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.txt"
            file_path.write_text("Email: alice@example.com", encoding="utf-8")
            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertEqual(payload["mode"], "local-regex")
            self.assertEqual(payload["level"], "lite")
            self.assertTrue(payload["data"]["hasPII"])
            self.assertIn("[EMAIL_1]", payload["data"]["anonymizedContent"])
            self.assertIn("outputPath", payload["data"])
            output_path = Path(payload["data"]["outputPath"])
            self.assertTrue(output_path.exists())
            self.assertEqual(output_path.name, "incident.redacted.txt")
            self.assertIn("privacy-protector-map-id", output_path.read_text(encoding="utf-8"))
            self.assertTrue(Path(payload["data"]["mapRef"]["sidecarPath"]).exists())
            self.assertIn("applyReport", payload["data"])
            self.assertEqual(
                payload["data"]["applyReport"]["expectedCount"],
                payload["data"]["applyReport"]["appliedCount"],
            )
            self.assertIn("verificationReport", payload["data"])
            self.assertTrue(payload["data"]["verificationReport"]["skipped"])
            self.assertEqual(payload["data"]["assurancePolicy"]["level"], "best_effort")
            self.assertTrue(payload["data"]["assurancePolicy"]["failOnCoverageMismatch"])

    def test_markdown_file_path_is_auto_resolved_and_redacted(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "handoff.md"
            file_path.write_text("Contact: alice@example.com", encoding="utf-8")
            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertIn("[EMAIL_1]", payload["data"]["anonymizedContent"])
            output_path = Path(payload["data"]["outputPath"])
            self.assertEqual(output_path.name, "handoff.redacted.md")
            self.assertTrue(output_path.read_text(encoding="utf-8").startswith("<!-- privacy-protector-map-id:"))

    def test_json_file_path_is_auto_resolved_and_redacted_without_inline_marker(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.json"
            file_path.write_text('{"email": "alice@example.com"}', encoding="utf-8")
            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertIn("[EMAIL_1]", payload["data"]["anonymizedContent"])
            output_path = Path(payload["data"]["outputPath"])
            self.assertEqual(output_path.name, "incident.redacted.json")
            output_content = output_path.read_text(encoding="utf-8")
            self.assertNotIn("privacy-protector-map-id", output_content)
            self.assertTrue(Path(payload["data"]["mapRef"]["sidecarPath"]).exists())

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_docx_file_path_is_auto_resolved_and_redacted(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.docx"
            document = Document()
            document.add_paragraph("Email: alice@example.com")
            document.save(str(file_path))

            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            output_path = Path(payload["data"]["outputPath"])
            self.assertEqual(output_path.name, "incident.redacted.docx")
            self.assertTrue(output_path.exists())

            redacted_doc = Document(str(output_path))
            redacted_text = "\n".join(paragraph.text for paragraph in redacted_doc.paragraphs)
            self.assertIn("[EMAIL_1]", redacted_text)
            self.assertTrue(Path(payload["data"]["mapRef"]["sidecarPath"]).exists())
            self.assertEqual(payload["data"]["assurancePolicy"]["level"], "verified")
            self.assertFalse(payload["data"]["verificationReport"]["skipped"])
            self.assertTrue(payload["data"]["verificationReport"]["passed"])

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_pdf_file_path_is_auto_resolved_and_redacted(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Email: alice@example.com")
            document.save(str(file_path))
            document.close()

            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            output_path = Path(payload["data"]["outputPath"])
            self.assertEqual(output_path.name, "incident.redacted.pdf")
            self.assertTrue(output_path.exists())

            redacted_document = fitz.open(str(output_path))
            try:
                extracted_text = "\n".join(page.get_text("text") for page in redacted_document)
            finally:
                redacted_document.close()
            self.assertNotIn("alice@example.com", extracted_text)
            self.assertTrue(Path(payload["data"]["mapRef"]["sidecarPath"]).exists())
            self.assertEqual(payload["data"]["assurancePolicy"]["level"], "verified")
            self.assertFalse(payload["data"]["verificationReport"]["skipped"])
            self.assertTrue(payload["data"]["verificationReport"]["passed"])

    def test_default_coverage_enforcement_exposes_assurance_and_apply_report(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.txt"
            file_path.write_text("Email: alice@example.com", encoding="utf-8")

            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertTrue(payload["data"]["assurancePolicy"]["failOnCoverageMismatch"])
            self.assertGreater(payload["data"]["applyReport"]["expectedCount"], 0)
            self.assertEqual(
                payload["data"]["applyReport"]["expectedCount"],
                payload["data"]["applyReport"]["appliedCount"],
            )

    def test_removed_strict_coverage_flag_is_rejected(self):
        result = self._run_cli(
            [
                "--input",
                "Email: alice@example.com",
                "--level",
                "lite",
                "--strict-coverage",
            ]
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("unrecognized arguments: --strict-coverage", result.stderr)

    @unittest.skipUnless(HAS_FITZ, "PyMuPDF is required")
    def test_pdf_dynamic_level_attempts_api_mode(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Email: alice@example.com")
            document.save(str(file_path))
            document.close()

            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "dynamic",
                    "--json",
                ],
                env={"ANONYMIZE_API_URL": "http://127.0.0.1:9"},
            )

            self.assertEqual(result.returncode, 1)
            payload = json.loads(result.stdout)
            self.assertFalse(payload["success"])
            self.assertEqual(payload["mode"], "api")
            self.assertEqual(payload["level"], "dynamic")
            self.assertIn(payload["error"]["type"], ("network_error", "dependency_error"))

    def test_unsupported_file_extension_returns_json_validation_error(self):
        with tempfile.NamedTemporaryFile("w", suffix=".bin", encoding="utf-8", delete=False) as input_file:
            input_file.write("Unsupported payload")
            file_path = input_file.name

        try:
            result = self._run_cli(
                [
                    "--input",
                    file_path,
                    "--level",
                    "dynamic",
                    "--json",
                ]
            )
        finally:
            os.unlink(file_path)

        self.assertEqual(result.returncode, 2)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual(payload["error"]["type"], "validation_error")

    def test_anonymize_in_place_overwrites_input_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "incident.txt"
            file_path.write_text("Email: alice@example.com", encoding="utf-8")

            result = self._run_cli(
                [
                    "--input",
                    str(file_path),
                    "--level",
                    "lite",
                    "--in-place",
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertEqual(payload["data"]["outputPath"], str(file_path))
            self.assertIn("[EMAIL_1]", file_path.read_text(encoding="utf-8"))

    def test_anonymize_in_place_requires_file_input(self):
        result = self._run_cli(
            [
                "--input",
                "Email: alice@example.com",
                "--level",
                "lite",
                "--in-place",
                "--json",
            ]
        )
        self.assertEqual(result.returncode, 2)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual(payload["error"]["type"], "validation_error")

    def test_anonymize_output_flag_writes_requested_path(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "custom-output.txt"
            result = self._run_cli(
                [
                    "--input",
                    "Email: alice@example.com",
                    "--level",
                    "lite",
                    "--output",
                    str(output_path),
                    "--json",
                ],
                env={"MODEIO_REDACT_MAP_DIR": str(Path(tmpdir) / "maps")},
            )

            self.assertEqual(result.returncode, 0)
            payload = json.loads(result.stdout)
            self.assertEqual(payload["data"]["outputPath"], str(output_path))
            self.assertTrue(output_path.exists())

    def test_anonymize_rejects_missing_output_directory(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            missing_output = Path(tmpdir) / "missing" / "out.txt"
            result = self._run_cli(
                [
                    "--input",
                    "Email: alice@example.com",
                    "--level",
                    "lite",
                    "--output",
                    str(missing_output),
                    "--json",
                ]
            )

            self.assertEqual(result.returncode, 2)
            payload = json.loads(result.stdout)
            self.assertFalse(payload["success"])
            self.assertEqual(payload["error"]["type"], "validation_error")

    @patch("anonymize.requests.post")
    def test_api_payload_uses_file_input_type_when_file_path_is_used(self, mock_post):
        fake_response = Mock()
        fake_response.raise_for_status.return_value = None
        fake_response.json.return_value = {
            "success": True,
            "data": {
                "anonymizedContent": "[REDACTED]",
                "hasPII": True,
            },
        }
        mock_post.return_value = fake_response

        with tempfile.NamedTemporaryFile("w", suffix=".txt", encoding="utf-8", delete=False) as input_file:
            input_file.write("Email: alice@example.com")
            file_path = input_file.name

        try:
            content, input_type = anonymize.resolve_input_source(file_path)
            result = anonymize.anonymize(content, level="dynamic", input_type=input_type)
        finally:
            os.unlink(file_path)

        self.assertTrue(result["success"])
        _, kwargs = mock_post.call_args
        payload = kwargs["json"]
        self.assertEqual(payload["inputType"], "file")

    def test_non_file_input_stays_text_mode(self):
        content, input_type = anonymize.resolve_input_source("Email: alice@example.com")
        self.assertEqual(input_type, "text")
        self.assertEqual(content, "Email: alice@example.com")

    def test_maybe_save_map_returns_none_without_entries(self):
        data = {"anonymizedContent": "No placeholders"}
        map_ref = anonymize._maybe_save_map(
            raw_input="No placeholders",
            level="lite",
            mode="local-regex",
            data=data,
        )
        self.assertIsNone(map_ref)
        self.assertNotIn("mapRef", data)

    def test_validate_non_text_mapping_rejects_remote_file_redaction_without_entries(self):
        with self.assertRaises(ValueError):
            anonymize._validate_non_text_mapping_or_raise(
                level="dynamic",
                input_path="/tmp/incident.pdf",
                input_extension=".pdf",
                raw_input="Email: alice@example.com",
                anonymized_content="Email: [EMAIL_1]",
                has_pii=True,
                entries=[],
            )

    def test_validate_non_text_mapping_allows_empty_entries_when_no_pii_detected(self):
        anonymize._validate_non_text_mapping_or_raise(
            level="dynamic",
            input_path="/tmp/incident.pdf",
            input_extension=".pdf",
            raw_input="No personal data here.",
            anonymized_content="No personal data here.",
            has_pii=False,
            entries=[],
        )

    def test_maybe_save_map_from_local_detection_items(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            original = os.environ.get("MODEIO_REDACT_MAP_DIR")
            os.environ["MODEIO_REDACT_MAP_DIR"] = tmpdir
            try:
                data = {
                    "anonymizedContent": "Email: [EMAIL_1]",
                    "localDetection": {
                        "items": [
                            {
                                "maskedValue": "[EMAIL_1]",
                                "value": "alice@example.com",
                                "type": "email",
                            }
                        ]
                    },
                }
                map_ref = anonymize._maybe_save_map(
                    raw_input="Email: alice@example.com",
                    level="lite",
                    mode="local-regex",
                    data=data,
                )
                self.assertIsNotNone(map_ref)
                self.assertEqual(map_ref.entry_count, 1)
                self.assertTrue(Path(map_ref.map_path).exists())
            finally:
                if original is None:
                    os.environ.pop("MODEIO_REDACT_MAP_DIR", None)
                else:
                    os.environ["MODEIO_REDACT_MAP_DIR"] = original

    def test_maybe_save_map_from_api_mapping_shape(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            original = os.environ.get("MODEIO_REDACT_MAP_DIR")
            os.environ["MODEIO_REDACT_MAP_DIR"] = tmpdir
            try:
                data = {
                    "anonymizedContent": "Name: [NAME_1]",
                    "mapping": [
                        {
                            "anonymized": "[NAME_1]",
                            "original": "Alice",
                            "type": "name",
                        }
                    ],
                }
                map_ref = anonymize._maybe_save_map(
                    raw_input="Name: Alice",
                    level="dynamic",
                    mode="api",
                    data=data,
                )
                self.assertIsNotNone(map_ref)
                self.assertEqual(map_ref.entry_count, 1)
                self.assertIn("mapRef", data)
            finally:
                if original is None:
                    os.environ.pop("MODEIO_REDACT_MAP_DIR", None)
                else:
                    os.environ["MODEIO_REDACT_MAP_DIR"] = original

    @unittest.skipUnless(HAS_DOCX, "python-docx is required")
    def test_maybe_sync_non_text_map_hash_updates_docx_hash_to_written_output(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            original = os.environ.get("MODEIO_REDACT_MAP_DIR")
            os.environ["MODEIO_REDACT_MAP_DIR"] = tmpdir
            try:
                input_text = "Accepted on behalf of Asterion Biologics, Inc."
                provider_text = "Accepted on behalf of [REDACTED_COMPANY_1]."
                written_text = "Accepted on behalf of [REDACTED_COMPANY_1]"

                map_ref = anonymize.save_map(
                    raw_input=input_text,
                    anonymized_content=provider_text,
                    entries=[
                        {
                            "placeholder": "[REDACTED_COMPANY_1]",
                            "original": "Asterion Biologics, Inc.",
                            "type": "company",
                        }
                    ],
                    level="dynamic",
                    source_mode="api",
                )

                output_path = Path(tmpdir) / "output.docx"
                document = Document()
                document.add_paragraph(written_text)
                document.save(str(output_path))

                data = {"mapRef": map_ref.to_dict()}
                anonymize._maybe_sync_non_text_map_hash(
                    map_ref=map_ref,
                    output_path=str(output_path),
                    input_extension=".docx",
                    data=data,
                )

                record, _ = load_map(map_ref.map_path)
                actual_output = anonymize.read_input_file(output_path, ".docx")
                self.assertEqual(record.anonymized_hash, hash_text(actual_output))
                self.assertNotEqual(record.anonymized_hash, hash_text(provider_text))
            finally:
                if original is None:
                    os.environ.pop("MODEIO_REDACT_MAP_DIR", None)
                else:
                    os.environ["MODEIO_REDACT_MAP_DIR"] = original

    @patch("anonymize.save_map")
    def test_maybe_save_map_propagates_storage_error(self, mock_save_map):
        mock_save_map.side_effect = anonymize.MapStoreError("disk full")
        data = {
            "anonymizedContent": "Email: [EMAIL_1]",
            "mapping": [
                {
                    "anonymized": "[EMAIL_1]",
                    "original": "alice@example.com",
                    "type": "email",
                }
            ],
        }

        with self.assertRaises(anonymize.MapStoreError):
            anonymize._maybe_save_map(
                raw_input="Email: alice@example.com",
                level="dynamic",
                mode="api",
                data=data,
            )

    def test_append_warning_initializes_warning_list(self):
        data = {}
        anonymize._append_warning(data, code="map_persist_failed", message="boom")

        self.assertIn("warnings", data)
        self.assertEqual(len(data["warnings"]), 1)
        self.assertEqual(data["warnings"][0]["code"], "map_persist_failed")

    @patch("anonymize.requests.post")
    def test_api_payload_uses_text_input_type_and_crossborder_codes(self, mock_post):
        fake_response = Mock()
        fake_response.raise_for_status.return_value = None
        fake_response.json.return_value = {
            "success": True,
            "data": {
                "anonymizedContent": "[REDACTED]",
                "hasPII": True,
            },
        }
        mock_post.return_value = fake_response

        result = anonymize.anonymize(
            "Name: John Doe, SSN: 123-45-6789",
            level="crossborder",
            sender_code="CN SHA",
            recipient_code="US NYC",
        )

        self.assertTrue(result["success"])
        _, kwargs = mock_post.call_args
        payload = kwargs["json"]
        self.assertEqual(payload["inputType"], "text")
        self.assertEqual(payload["level"], "crossborder")
        self.assertEqual(payload["senderCode"], "CN SHA")
        self.assertEqual(payload["recipientCode"], "US NYC")


if __name__ == "__main__":
    unittest.main()

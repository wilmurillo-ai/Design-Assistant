#!/usr/bin/env python3
"""C刊论文数据分析与Word报告生成脚本

用法: python3 analyze_journal.py <data.json> [output_dir]

data.json格式:
{
  "journal_info": { "title": "...", "issn": "...", "sponsor": "...", ... },
  "articles": [
    { "year": 2024, "issue": "01", "title": "...", "authors": "...", "section": "...", "abstract": "", "keywords": "" },
    ...
  ]
}
"""

import json
import sys
import os
import re
from collections import Counter, defaultdict
from pathlib import Path

try:
    import jieba
    import matplotlib
    matplotlib.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'PingFang SC', 'Heiti TC', 'SimHei']
    matplotlib.rcParams['axes.unicode_minus'] = False
    matplotlib.rcParams['figure.dpi'] = 150
    import matplotlib.pyplot as plt
    import numpy as np
    from wordcloud import WordCloud
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"缺少依赖: {e}")
    print("请运行: pip3 install jieba wordcloud python-docx matplotlib numpy")
    sys.exit(1)


# ============================================================
# 停用词表
# ============================================================
STOPWORDS = set("""
的 了 在 是 我 有 和 就 不 人 都 一 一个 上 也 很 到 说 要 去 你
对 出 会 着 没有 看 好 自己 这 他 她 它 们 那 些 什么 之 与 及
为 中 以 或 其 被 从 但 而 可 所 这个 那个 因为 所以 如果 虽然
通过 进行 基于 研究 分析 问题 中国 影响 发展 我国 探讨 关于
视角 效应 论 试论 浅析 浅谈 探析 略论 新 论述
""".split())

# 研究方法关键词
METHOD_KEYWORDS = {
    '实证研究': ['实证', '实证研究', '实证分析', '经验研究'],
    '案例研究': ['案例', '案例研究', '个案', '案例分析'],
    '实验研究': ['实验', '随机实验', '田野实验', '自然实验', '准实验'],
    '计量模型': ['回归', 'OLS', '面板数据', '固定效应', '工具变量', 'IV', 'GMM', '计量'],
    'DID/RDD': ['双重差分', 'DID', '断点回归', 'RDD', '倍差法'],
    'PSM': ['倾向得分匹配', 'PSM', '匹配估计'],
    '结构方程': ['SEM', '结构方程', '路径分析'],
    '机器学习/AI': ['机器学习', '深度学习', '神经网络', '人工智能', 'AI', '大数据', '文本分析', 'NLP'],
    '质性研究': ['扎根理论', '质性', '访谈', '民族志', '叙事', '话语分析'],
    '调查研究': ['问卷', '调查', '抽样'],
    '文献研究': ['文献计量', '元分析', 'meta分析', '系统综述', '文献综述'],
    '博弈/仿真': ['博弈', '仿真', '模拟', 'ABM', '演化博弈'],
}


def load_data(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)


def segment_titles(articles):
    """对文章标题进行jieba分词"""
    all_words = []
    for art in articles:
        words = jieba.lcut(art['title'])
        words = [w.strip() for w in words if len(w.strip()) >= 2 and w.strip() not in STOPWORDS]
        all_words.extend(words)
    return all_words


def analyze_methods(articles):
    """识别研究方法"""
    method_counts = Counter()
    for art in articles:
        title = art['title'] + ' ' + art.get('abstract', '') + ' ' + art.get('keywords', '')
        for method, keywords in METHOD_KEYWORDS.items():
            for kw in keywords:
                if kw.lower() in title.lower():
                    method_counts[method] += 1
                    break
    return method_counts


def analyze_authors(articles):
    """统计作者发文量"""
    author_counts = Counter()
    for art in articles:
        authors = art.get('authors', '')
        for author in re.split(r'[;；,，、]', authors):
            author = author.strip()
            if author and len(author) >= 2:
                author_counts[author] += 1
    return author_counts


def analyze_sections(articles):
    """统计栏目分布"""
    section_counts = Counter()
    for art in articles:
        section = art.get('section', '').strip()
        if section:
            section_counts[section] += 1
    return section_counts


def analyze_yearly_trends(articles):
    """按年度统计"""
    yearly = defaultdict(list)
    for art in articles:
        yearly[art['year']].append(art)
    return dict(sorted(yearly.items()))


def compare_periods(articles, current_year):
    """前期vs近期关键词对比"""
    early = [a for a in articles if a['year'] <= current_year - 2]
    recent = [a for a in articles if a['year'] > current_year - 2]

    early_words = Counter(segment_titles(early))
    recent_words = Counter(segment_titles(recent))

    # 新兴主题: 近期高频但前期低频
    emerging = {}
    for word, count in recent_words.most_common(100):
        early_count = early_words.get(word, 0)
        if early_count == 0 and count >= 3:
            emerging[word] = count
        elif early_count > 0:
            ratio = count / len(recent) - early_count / len(early) if len(early) > 0 else 0
            if ratio > 0.005 and count >= 3:
                emerging[word] = round(ratio, 4)

    # 衰退主题: 前期高频但近期低频
    declining = {}
    for word, count in early_words.most_common(100):
        recent_count = recent_words.get(word, 0)
        if len(recent) > 0 and len(early) > 0:
            ratio = count / len(early) - recent_count / len(recent)
            if ratio > 0.005 and count >= 3:
                declining[word] = round(ratio, 4)

    return emerging, declining


def create_charts(articles, output_dir, journal_name):
    """生成所有可视化图表"""
    os.makedirs(output_dir, exist_ok=True)
    charts = {}

    # 1. 发文量趋势
    yearly = analyze_yearly_trends(articles)
    years = sorted(yearly.keys())
    counts = [len(yearly[y]) for y in years]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(years, counts, 'o-', color='#E8737A', linewidth=2.5, markersize=8)
    ax.fill_between(years, counts, alpha=0.15, color='#E8737A')
    for y, c in zip(years, counts):
        ax.text(y, c + max(counts)*0.02, str(c), ha='center', fontsize=11, fontweight='bold')
    ax.set_xlabel('年份', fontsize=12)
    ax.set_ylabel('发文量', fontsize=12)
    ax.set_title(f'《{journal_name}》年度发文量趋势', fontsize=14, fontweight='bold')
    ax.set_xticks(years)
    plt.tight_layout()
    path = f'{output_dir}/01_yearly_trend.png'
    plt.savefig(path, bbox_inches='tight')
    plt.close()
    charts['yearly_trend'] = path

    # 2. 高频关键词Top30
    all_words = segment_titles(articles)
    word_freq = Counter(all_words).most_common(30)

    fig, ax = plt.subplots(figsize=(12, 7))
    words_list = [w for w, _ in word_freq][::-1]
    freqs_list = [f for _, f in word_freq][::-1]
    colors = plt.cm.RdYlBu_r(np.linspace(0.2, 0.8, len(words_list)))
    ax.barh(range(len(words_list)), freqs_list, color=colors, edgecolor='white')
    ax.set_yticks(range(len(words_list)))
    ax.set_yticklabels(words_list, fontsize=10)
    ax.set_xlabel('频次', fontsize=12)
    ax.set_title(f'《{journal_name}》高频关键词Top30', fontsize=14, fontweight='bold')
    plt.tight_layout()
    path = f'{output_dir}/02_top_keywords.png'
    plt.savefig(path, bbox_inches='tight')
    plt.close()
    charts['top_keywords'] = path

    # 3. 词云
    word_dict = dict(Counter(all_words).most_common(200))
    try:
        font_path = '/System/Library/Fonts/PingFang.ttc'
        if not os.path.exists(font_path):
            font_path = '/System/Library/Fonts/STHeiti Light.ttc'
        wc = WordCloud(
            font_path=font_path,
            width=1200, height=600,
            background_color='white',
            max_words=150,
            colormap='RdYlBu_r',
            max_font_size=120
        ).generate_from_frequencies(word_dict)
        fig, ax = plt.subplots(figsize=(14, 7))
        ax.imshow(wc, interpolation='bilinear')
        ax.axis('off')
        ax.set_title(f'《{journal_name}》关键词词云', fontsize=14, fontweight='bold')
        path = f'{output_dir}/03_wordcloud.png'
        plt.savefig(path, bbox_inches='tight')
        plt.close()
        charts['wordcloud'] = path
    except Exception as e:
        print(f"词云生成失败: {e}")

    # 4. 核心作者Top20
    author_freq = analyze_authors(articles).most_common(20)
    if author_freq:
        fig, ax = plt.subplots(figsize=(12, 7))
        names = [a for a, _ in author_freq][::-1]
        freqs = [f for _, f in author_freq][::-1]
        ax.barh(range(len(names)), freqs, color='#5BA3CF', edgecolor='white')
        ax.set_yticks(range(len(names)))
        ax.set_yticklabels(names, fontsize=10)
        ax.set_xlabel('发文量', fontsize=12)
        ax.set_title(f'《{journal_name}》高产作者Top20', fontsize=14, fontweight='bold')
        plt.tight_layout()
        path = f'{output_dir}/04_top_authors.png'
        plt.savefig(path, bbox_inches='tight')
        plt.close()
        charts['top_authors'] = path

    # 5. 研究方法分布
    method_freq = analyze_methods(articles)
    if method_freq:
        fig, ax = plt.subplots(figsize=(10, 6))
        methods = sorted(method_freq.items(), key=lambda x: x[1], reverse=True)
        m_names = [m for m, _ in methods]
        m_counts = [c for _, c in methods]
        colors_m = plt.cm.Set3(np.linspace(0, 1, len(m_names)))
        ax.bar(m_names, m_counts, color=colors_m, edgecolor='white')
        ax.set_ylabel('篇数', fontsize=12)
        ax.set_title(f'《{journal_name}》研究方法分布', fontsize=14, fontweight='bold')
        plt.xticks(rotation=30, ha='right', fontsize=10)
        plt.tight_layout()
        path = f'{output_dir}/05_methods.png'
        plt.savefig(path, bbox_inches='tight')
        plt.close()
        charts['methods'] = path

    # 6. 栏目分布
    section_freq = analyze_sections(articles)
    if section_freq:
        top_sections = section_freq.most_common(10)
        fig, ax = plt.subplots(figsize=(10, 6))
        s_names = [s for s, _ in top_sections]
        s_counts = [c for _, c in top_sections]
        colors_s = plt.cm.Pastel1(np.linspace(0, 1, len(s_names)))
        wedges, texts, autotexts = ax.pie(s_counts, labels=s_names, autopct='%1.1f%%',
                                           colors=colors_s, textprops={'fontsize': 9})
        ax.set_title(f'《{journal_name}》栏目分布', fontsize=14, fontweight='bold')
        path = f'{output_dir}/06_sections.png'
        plt.savefig(path, bbox_inches='tight')
        plt.close()
        charts['sections'] = path

    return charts


def generate_report(data, charts, output_path):
    """生成Word报告"""
    journal = data.get('journal_info', {})
    articles = data.get('articles', [])
    journal_name = journal.get('title', '未知期刊')

    years = sorted(set(a['year'] for a in articles))
    year_range = f"{min(years)}-{max(years)}" if years else "N/A"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    def add_text(text, bold=False, size=10.5, indent=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = bold
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(22)
        return p

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f'《{journal_name}》')
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(f'近五年（{year_range}）发文分析报告')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('由 Claude Code 自动生成')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # 一、期刊概况
    doc.add_heading('一、期刊概况', level=1)
    info_text = f"《{journal_name}》"
    if journal.get('sponsor'):
        info_text += f"由{journal['sponsor']}主办"
    if journal.get('frequency'):
        info_text += f"，{journal['frequency']}出版"
    if journal.get('issn'):
        info_text += f"，ISSN: {journal['issn']}"
    if journal.get('cif'):
        info_text += f"。复合影响因子为{journal['cif']}"
    if journal.get('aif'):
        info_text += f"，综合影响因子为{journal['aif']}"
    info_text += f"。本报告分析了该刊{year_range}年间共{len(articles)}篇文章的发文数据。"
    add_text(info_text)

    # 二、发文量与趋势
    doc.add_heading('二、发文量与趋势分析', level=1)
    yearly = analyze_yearly_trends(articles)
    trend_text = "从年度发文量来看，"
    for y in sorted(yearly.keys()):
        trend_text += f"{y}年发文{len(yearly[y])}篇，"
    trend_text = trend_text.rstrip('，') + "。"
    add_text(trend_text)
    if 'yearly_trend' in charts:
        doc.add_picture(charts['yearly_trend'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 三、选题热点
    doc.add_heading('三、选题热点分析', level=1)
    all_words = segment_titles(articles)
    top30 = Counter(all_words).most_common(30)
    kw_text = f"对全部{len(articles)}篇文章标题进行中文分词后，出现频率最高的30个关键词如下图所示。"
    kw_text += f"其中，排名前5的关键词为：{'、'.join([w for w,_ in top30[:5]])}。"
    add_text(kw_text)
    if 'top_keywords' in charts:
        doc.add_picture(charts['top_keywords'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if 'wordcloud' in charts:
        doc.add_picture(charts['wordcloud'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 四、热点演变
    doc.add_heading('四、热点演变与新兴主题', level=1)
    if len(years) >= 3:
        current_year = max(years)
        emerging, declining = compare_periods(articles, current_year)
        if emerging:
            top_emerging = sorted(emerging.items(), key=lambda x: x[1], reverse=True)[:10]
            add_text(f"与前期相比，近两年新兴或上升的研究主题包括：{'、'.join([w for w,_ in top_emerging])}。")
        if declining:
            top_declining = sorted(declining.items(), key=lambda x: x[1], reverse=True)[:10]
            add_text(f"关注度有所下降的主题包括：{'、'.join([w for w,_ in top_declining])}。")
    else:
        add_text("数据跨度不足3年，暂无法进行趋势对比分析。")

    # 五、核心作者
    doc.add_heading('五、核心作者群分析', level=1)
    author_freq = analyze_authors(articles).most_common(20)
    if author_freq:
        add_text(f"发文量排名前5的作者为：" +
                 "、".join([f"{a}({c}篇)" for a, c in author_freq[:5]]) + "。")
    if 'top_authors' in charts:
        doc.add_picture(charts['top_authors'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 六、研究方法
    doc.add_heading('六、研究方法偏好分析', level=1)
    method_freq = analyze_methods(articles)
    if method_freq:
        top_methods = method_freq.most_common(5)
        add_text(f"从研究方法来看，该刊最常见的方法类型为：" +
                 "、".join([f"{m}({c}篇)" for m, c in top_methods]) + "。")
    if 'methods' in charts:
        doc.add_picture(charts['methods'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 七、栏目分析
    doc.add_heading('七、栏目主题分析', level=1)
    section_freq = analyze_sections(articles)
    if section_freq:
        top_sections = section_freq.most_common(5)
        add_text(f"该刊主要栏目及发文量为：" +
                 "、".join([f"{s}({c}篇)" for s, c in top_sections]) + "。")
    if 'sections' in charts:
        doc.add_picture(charts['sections'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 八、研究空白与投稿建议
    doc.add_heading('八、研究空白与投稿建议', level=1)
    add_text("基于以上分析，以下方面可能存在研究空白或发文机会：", bold=True, indent=False)
    add_text("（本节内容需结合具体学科知识由研究者进一步判断。以上数据分析可为选题提供参考方向。）")

    # 附录
    doc.add_page_break()
    doc.add_heading('附录：完整文章目录', level=1)
    for year in sorted(yearly.keys(), reverse=True):
        doc.add_heading(f'{year}年', level=2)
        for art in yearly[year]:
            p = doc.add_paragraph()
            run = p.add_run(f"[{art.get('issue','')}期] {art['title']}")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if art.get('authors'):
                run = p.add_run(f"  —— {art['authors']}")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(128, 128, 128)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(16)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print("用法: python3 analyze_journal.py <data.json> [output_dir]")
        sys.exit(1)

    data_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else os.path.dirname(data_path)

    data = load_data(data_path)
    journal_name = data.get('journal_info', {}).get('title', '未知期刊')
    articles = data.get('articles', [])

    print(f"期刊: {journal_name}")
    print(f"文章总数: {len(articles)}")

    chart_dir = os.path.join(output_dir, f'{journal_name}_charts')
    charts = create_charts(articles, chart_dir, journal_name)
    print(f"图表已生成: {chart_dir}/")

    report_path = os.path.join(output_dir, f'{journal_name}_近五年发文分析报告.docx')
    generate_report(data, charts, report_path)
    print(f"报告已生成: {report_path}")


if __name__ == '__main__':
    main()

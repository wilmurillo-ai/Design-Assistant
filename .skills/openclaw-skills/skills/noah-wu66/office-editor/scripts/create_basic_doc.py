from docx import Document


def create_document(filename="basic_document.docx"):
    """Create a basic Word document with a heading and a few paragraphs."""
    doc = Document()
    doc.add_heading("Sample Document Title", level=1)
    doc.add_paragraph("This is a sample paragraph.")
    doc.add_paragraph("You can add more content as needed.")
    doc.save(filename)
    print(f"Document '{filename}' was created successfully.")


if __name__ == "__main__":
    create_document()

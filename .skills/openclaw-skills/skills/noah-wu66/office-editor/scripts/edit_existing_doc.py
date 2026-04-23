import sys
from pathlib import Path

from docx import Document


def edit_document(input_path, output_path=None):
    """Open an existing Word file, append content, and save it as a new file."""
    document = Document(input_path)
    document.add_heading("Update Log", level=2)
    document.add_paragraph(
        "This paragraph was appended by the sample script to demonstrate how to edit an existing document."
    )

    input_file = Path(input_path)
    final_output = output_path or str(input_file.with_name(f"updated_{input_file.name}"))
    document.save(final_output)
    print(f"Document saved to: {final_output}")


def main():
    if len(sys.argv) < 2:
        raise SystemExit("Usage: python edit_existing_doc.py <input.docx> [output.docx]")

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    edit_document(input_path, output_path)


if __name__ == "__main__":
    main()

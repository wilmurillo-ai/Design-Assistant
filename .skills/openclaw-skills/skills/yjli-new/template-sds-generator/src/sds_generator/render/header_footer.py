from __future__ import annotations

from pathlib import Path

from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from sds_generator.constants import NO_DATA
from sds_generator.render.styles import FOOTER_STYLE, META_STYLE

LOGO_PLACEHOLDER = "[Company logo]"
NON_BREAKING_SPACE = "\u00A0"
NON_BREAKING_HYPHEN = "\u2011"
DISPLAY_TEXT_NORMALIZER = str.maketrans({
    NON_BREAKING_SPACE: " ",
    NON_BREAKING_HYPHEN: "-",
})


def _clear_story(story) -> None:
    element = story._element
    for child in list(element):
        element.remove(child)


def _add_field_code(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def _usable_width(section) -> int:
    return int(section.page_width - section.left_margin - section.right_margin)


def format_cas_display(cas_number: str | None) -> str:
    cas_display = cas_number or NO_DATA
    return cas_display.replace("-", NON_BREAKING_HYPHEN)


def normalize_display_text(text: str) -> str:
    return text.translate(DISPLAY_TEXT_NORMALIZER)


def _render_logo(cell, logo_path: Path | None, warning_sink: list[str]) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.style = META_STYLE
    if logo_path and logo_path.exists():
        run = paragraph.add_run()
        run.add_picture(str(logo_path), width=Cm(2.2))
        return
    paragraph.text = LOGO_PLACEHOLDER
    warning_sink.append("Logo asset missing; inserted text placeholder in header.")


def apply_header_footer(
    document: Document,
    *,
    product_name: str | None,
    cas_number: str | None,
    company_name: str,
    website: str,
    logo_path: Path | None = None,
) -> list[str]:
    warning_sink: list[str] = []
    product_display = product_name or NO_DATA
    cas_display = format_cas_display(cas_number)
    for section in document.sections:
        _clear_story(section.header)
        header_table = section.header.add_table(rows=1, cols=2, width=_usable_width(section))
        header_table.autofit = False
        left_cell, right_cell = header_table.rows[0].cells
        left_cell.width = Cm(3.0)
        right_cell.width = Cm(13.8)
        _render_logo(left_cell, logo_path, warning_sink)
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.style = META_STYLE
        right_paragraph.text = (
            f"SAFETY DATA SHEET | Product: {product_display} | CAS:{NON_BREAKING_SPACE}{cas_display}"
        )

        _clear_story(section.footer)
        footer_paragraph = section.footer.add_paragraph(style=FOOTER_STYLE)
        footer_paragraph.add_run(f"{company_name} | {website} | Page ")
        _add_field_code(footer_paragraph, "PAGE")
        footer_paragraph.add_run(" of ")
        _add_field_code(footer_paragraph, "NUMPAGES")
    return warning_sink


def header_text(document: Document) -> str:
    texts: list[str] = []
    for section in document.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                    if cell_text:
                        texts.append(normalize_display_text(cell_text))
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                texts.append(normalize_display_text(paragraph.text.strip()))
    return "\n".join(texts)


def footer_text(document: Document) -> str:
    texts: list[str] = []
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text.strip())
    return "\n".join(texts)

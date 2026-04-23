from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from sds_generator.config_loader import load_fixed_company, load_template_contract
from sds_generator.content_policy import approved_template_default
from sds_generator.constants import NO_DATA
from sds_generator.models import (
    ChecklistBucket,
    FinalSDSDocument,
    ReviewNote,
    ReviewSeverity,
    ReviewStatus,
    TemplateContract,
    TemplateFillRange,
    TemplateFillReport,
)
from sds_generator.render.header_footer import format_cas_display, normalize_display_text
from sds_generator.render.section_blocks import display_value
from sds_generator.render.tables import add_key_value_table, add_matrix_table, render_section14_table, render_section3_table, render_section9_table
from sds_generator.render.template_anchors import find_anchor_occurrences, find_missing_anchors, ordered_anchor_paragraphs
from sds_generator.render.template_tokens import (
    find_placeholders_by_location,
    find_unresolved_placeholders,
    replace_placeholders_with_report,
)


def _layout_note(field_path: str, message: str) -> ReviewNote:
    return ReviewNote(
        field_path=field_path,
        severity=ReviewSeverity.MINOR,
        status=ReviewStatus.WARNING,
        why=message,
        checklist_bucket=ChecklistBucket.LAYOUT_BRANDING_QA,
    )


def _empty_document() -> Document:
    document = Document()
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def _body_elements(document: Document) -> list:
    return [deepcopy(child) for child in document._body._element if child.tag != qn("w:sectPr")]


def _insert_body_elements_after(anchor_element, elements: list) -> None:
    current = anchor_element
    for element in elements:
        current.addnext(element)
        current = element


def _clear_between(anchor_paragraph, next_anchor_paragraph) -> int:
    removed = 0
    current = anchor_paragraph._p.getnext()
    stop = next_anchor_paragraph._p if next_anchor_paragraph is not None else None
    while current is not None and current is not stop:
        next_current = current.getnext()
        current.getparent().remove(current)
        current = next_current
        removed += 1
    return removed


def _last_element_before(anchor_paragraph, next_anchor_paragraph):
    current = anchor_paragraph._p
    stop = next_anchor_paragraph._p if next_anchor_paragraph is not None else None
    while current.getnext() is not None and current.getnext() is not stop:
        current = current.getnext()
    return current


def _font_name(contract: TemplateContract) -> str | None:
    return contract.formatting.default_font_family


def _font_size(contract: TemplateContract) -> float | None:
    return contract.formatting.default_font_size_pt


def _font_color(contract: TemplateContract) -> str | None:
    return contract.formatting.default_text_color


def _apply_run_style(run, contract: TemplateContract, *, bold: bool = False) -> None:
    run.bold = bold
    if _font_name(contract):
        run.font.name = _font_name(contract)
    if _font_size(contract) is not None:
        run.font.size = Pt(_font_size(contract))
    if _font_color(contract):
        run.font.color.rgb = RGBColor.from_string(_font_color(contract))


def _add_text_paragraph(document: Document, text: str, contract: TemplateContract) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text or NO_DATA)
    _apply_run_style(run, contract)


def _add_label_value_paragraph(document: Document, label: str, value, contract: TemplateContract) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    _apply_run_style(label_run, contract, bold=True)
    value_run = paragraph.add_run(display_value(value))
    _apply_run_style(value_run, contract)


def _join_list(values: list[str]) -> str:
    cleaned = [value for value in values if value]
    return "; ".join(cleaned) if cleaned else NO_DATA


def _cover_issue_date(final_document: FinalSDSDocument) -> str:
    return final_document.section_16.date_of_first_issue or final_document.document_meta.generation_timestamp.strftime("%Y-%m-%d")


def _cover_revision_date(final_document: FinalSDSDocument) -> str:
    return final_document.section_16.revision_date or final_document.document_meta.generation_timestamp.strftime("%Y-%m-%d")


def _cover_version(final_document: FinalSDSDocument) -> str:
    return final_document.section_16.version_number or "1.0"


def build_template_replacements(final_document: FinalSDSDocument) -> dict[str, str]:
    section_1 = final_document.section_1
    section_2 = final_document.section_2
    section_3 = final_document.section_3
    return {
        "IssueDate": _cover_issue_date(final_document),
        "RevisionDate": _cover_revision_date(final_document),
        "Version": _cover_version(final_document),
        "ProductName": section_1.product_name or final_document.document_meta.product_name_display or NO_DATA,
        "CAS": format_cas_display(section_1.cas_number or final_document.document_meta.cas_number),
        "Synonyms": _join_list(section_1.synonyms),
        "Formula": section_3.formula or section_1.molecular_formula or NO_DATA,
        "MW": section_3.molecular_weight or section_1.molecular_weight or NO_DATA,
        "GHS No": _join_list(section_2.pictograms),
        "Signal word": section_2.signal_word or NO_DATA,
    }


def _table_kwargs(contract: TemplateContract) -> dict[str, object]:
    return {
        "paragraph_style": None,
        "font_name": _font_name(contract),
        "font_size_pt": _font_size(contract),
        "text_color": _font_color(contract),
    }


def _append_1_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Product identifier", final_document.section_1.product_identifier, contract)
    _add_label_value_paragraph(document, "Product number", final_document.section_1.product_number, contract)
    _add_label_value_paragraph(document, "EC No.", final_document.section_1.ec_number, contract)


def _replace_1_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    recommended_use = final_document.section_1.relevant_identified_uses or approved_template_default(
        "section_1.relevant_identified_uses"
    )
    _add_label_value_paragraph(
        document,
        "Recommended Use",
        recommended_use,
        contract,
    )
    _add_label_value_paragraph(document, "Uses advised against", final_document.section_1.uses_advised_against, contract)


def _replace_1_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Supplier", final_document.section_1.supplier_company_name, contract)
    _add_label_value_paragraph(document, "Address", final_document.section_1.supplier_address, contract)
    _add_label_value_paragraph(document, "Tel", final_document.section_1.supplier_telephone, contract)
    _add_label_value_paragraph(document, "Fax", final_document.section_1.supplier_fax, contract)
    _add_label_value_paragraph(document, "Email", final_document.section_1.supplier_email, contract)
    _add_label_value_paragraph(document, "Website", final_document.section_1.supplier_website, contract)


def _replace_1_4(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(document, final_document.section_1.emergency_telephone or NO_DATA, contract)


def _replace_2_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    classifications = final_document.section_2.ghs_classifications or [NO_DATA]
    for item in classifications:
        _add_text_paragraph(document, item, contract)


def _append_2_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Hazard statement(s)", final_document.section_2.hazard_statements, contract)
    _add_label_value_paragraph(document, "Precautionary statement(s) - Prevention", final_document.section_2.precautionary_prevention, contract)
    _add_label_value_paragraph(document, "Precautionary statement(s) - Response", final_document.section_2.precautionary_response, contract)
    _add_label_value_paragraph(document, "Precautionary statement(s) - Storage", final_document.section_2.precautionary_storage, contract)
    _add_label_value_paragraph(document, "Precautionary statement(s) - Disposal", final_document.section_2.precautionary_disposal, contract)


def _replace_2_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(
        document,
        "Other hazards",
        final_document.section_2.other_hazards or final_document.section_2.hazards_not_otherwise_classified,
        contract,
    )
    _add_label_value_paragraph(document, "Emergency overview", final_document.section_2.summary_of_emergency, contract)
    _add_label_value_paragraph(document, "Physical hazards", final_document.section_2.physical_hazards, contract)
    _add_label_value_paragraph(document, "Health hazards", final_document.section_2.health_hazards, contract)
    _add_label_value_paragraph(document, "Environmental hazards", final_document.section_2.environmental_hazards, contract)


def _append_3_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    rows = [
        ("Substance / Mixture", final_document.section_3.substance_or_mixture),
        ("Chemical name", final_document.section_3.chemical_name),
        ("EC No.", final_document.section_3.ec_number),
        ("Disclosure statement", final_document.section_3.disclosure_statement),
    ]
    add_key_value_table(document, rows, **_table_kwargs(contract))


def _render_rows_as_label_values(document: Document, rows: list[tuple[str, object]], contract: TemplateContract) -> None:
    for label, value in rows:
        _add_label_value_paragraph(document, label, value, contract)


def _replace_4_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("General advice", final_document.section_4.general_advice),
            ("Inhalation", final_document.section_4.inhalation),
            ("Skin contact", final_document.section_4.skin_contact),
            ("Eye contact", final_document.section_4.eye_contact),
            ("Ingestion", final_document.section_4.ingestion),
        ],
        contract,
    )


def _replace_4_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Most important symptoms", final_document.section_4.most_important_symptoms),
            ("Delayed effects", final_document.section_4.delayed_effects),
        ],
        contract,
    )


def _replace_4_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Immediate medical attention", final_document.section_4.immediate_medical_attention),
            ("Notes to physician", final_document.section_4.notes_to_physician),
        ],
        contract,
    )


def _replace_5_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Suitable extinguishing media", final_document.section_5.suitable_extinguishing_media),
            ("Unsuitable extinguishing media", final_document.section_5.unsuitable_extinguishing_media),
        ],
        contract,
    )


def _replace_5_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Special hazards arising", final_document.section_5.specific_hazards_arising),
            ("Hazardous combustion products", final_document.section_5.hazardous_combustion_products),
        ],
        contract,
    )


def _replace_5_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Advice for firefighters", final_document.section_5.firefighting_advice),
            ("Further information", final_document.section_5.further_information),
        ],
        contract,
    )


def _replace_6_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Personal precautions", final_document.section_6.personal_precautions),
            ("Protective equipment", final_document.section_6.protective_equipment),
            ("Emergency procedures", final_document.section_6.emergency_procedures),
        ],
        contract,
    )


def _replace_6_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Environmental precautions", final_document.section_6.environmental_precautions, contract)


def _replace_6_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Containment methods", final_document.section_6.containment_methods),
            ("Cleaning up", final_document.section_6.cleanup_methods),
        ],
        contract,
    )


def _replace_6_4(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Reference to other sections", final_document.section_6.reference_to_other_sections, contract)


def _replace_7_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Precautions for safe handling", final_document.section_7.safe_handling_precautions, contract)


def _replace_7_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Storage conditions", final_document.section_7.storage_conditions),
            ("Storage temperature", final_document.section_7.storage_temperature),
            ("Incompatible materials", final_document.section_7.incompatible_storage_materials),
            ("Specific end uses", final_document.section_7.specific_end_uses),
        ],
        contract,
    )


def _replace_8_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Occupational exposure limits", final_document.section_8.occupational_exposure_limits),
            ("Biological limits", final_document.section_8.biological_limits),
        ],
        contract,
    )


def _replace_8_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Engineering controls", final_document.section_8.engineering_controls),
            ("Eye / face protection", final_document.section_8.eye_face_protection),
            ("Hand protection", final_document.section_8.hand_protection),
            ("Skin / body protection", final_document.section_8.skin_body_protection),
            ("Respiratory protection", final_document.section_8.respiratory_protection),
            ("Thermal hazards", final_document.section_8.thermal_hazards),
            ("Environmental exposure controls", final_document.section_8.environmental_exposure_controls),
        ],
        contract,
    )


def _replace_9_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    render_section9_table(document, final_document.section_9, **_table_kwargs(contract))


def _replace_9_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Other safety information", final_document.section_9.other_safety_information, contract)


def _replace_10_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Reactivity", final_document.section_10.reactivity),
            ("Chemical stability", final_document.section_10.chemical_stability),
        ],
        contract,
    )


def _replace_10_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(
        document,
        "Possibility of hazardous reactions",
        final_document.section_10.possibility_of_hazardous_reactions,
        contract,
    )


def _replace_10_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Conditions to avoid", final_document.section_10.conditions_to_avoid, contract)


def _replace_10_4(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Incompatible materials", final_document.section_10.incompatible_materials),
            ("Hazardous decomposition products", final_document.section_10.hazardous_decomposition_products),
        ],
        contract,
    )


def _replace_11_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Likely routes of exposure", final_document.section_11.likely_routes_of_exposure),
            ("Acute toxicity oral", final_document.section_11.acute_toxicity_oral),
            ("Acute toxicity inhalation", final_document.section_11.acute_toxicity_inhalation),
            ("Acute toxicity dermal", final_document.section_11.acute_toxicity_dermal),
            ("Skin corrosion / irritation", final_document.section_11.skin_corrosion_irritation),
            ("Serious eye damage / irritation", final_document.section_11.serious_eye_damage_irritation),
            ("Respiratory or skin sensitization", final_document.section_11.respiratory_or_skin_sensitization),
            ("Germ cell mutagenicity", final_document.section_11.germ_cell_mutagenicity),
            ("Carcinogenicity", final_document.section_11.carcinogenicity),
            ("Reproductive toxicity", final_document.section_11.reproductive_toxicity),
            ("STOT single exposure", final_document.section_11.stot_single_exposure),
            ("STOT repeated exposure", final_document.section_11.stot_repeated_exposure),
            ("Aspiration hazard", final_document.section_11.aspiration_hazard),
        ],
        contract,
    )


def _replace_11_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Symptoms related to characteristics", final_document.section_11.symptoms_related_to_characteristics),
            ("Additional information", final_document.section_11.toxicological_additional_information),
            ("Numerical measures of toxicity", final_document.section_11.numerical_measures_of_toxicity),
        ],
        contract,
    )


def _replace_12_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Ecotoxicity - fish", final_document.section_12.ecotoxicity_fish),
            ("Ecotoxicity - daphnia", final_document.section_12.ecotoxicity_daphnia),
            ("Ecotoxicity - algae", final_document.section_12.ecotoxicity_algae),
            ("Ecotoxicity - microorganisms", final_document.section_12.ecotoxicity_microorganisms),
        ],
        contract,
    )


def _replace_12_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Persistence and degradability", final_document.section_12.persistence_and_degradability, contract)


def _replace_12_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Bioaccumulative potential", final_document.section_12.bioaccumulative_potential, contract)


def _replace_12_4(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Mobility in soil", final_document.section_12.mobility_in_soil, contract)


def _replace_12_5(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Results of PBT and vPvB assessment", final_document.section_12.pbt_vpvb_assessment, contract)


def _replace_12_6(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Endocrine disrupting properties", final_document.section_12.endocrine_disrupting_properties),
            ("Other adverse effects", final_document.section_12.other_adverse_effects),
        ],
        contract,
    )


def _replace_13_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _render_rows_as_label_values(
        document,
        [
            ("Waste treatment methods - product", final_document.section_13.waste_treatment_methods_product),
            ("Waste treatment methods - packaging", final_document.section_13.waste_treatment_methods_packaging),
            ("Sewage disposal advice", final_document.section_13.sewage_disposal_advice),
        ],
        contract,
    )


def _transport_value(section, *attrs: str):
    for attr in attrs:
        direct = getattr(section, attr, None)
        if direct not in (None, "", [], {}):
            return direct
    for mode in section.transport_modes:
        for attr in attrs:
            value = getattr(mode, attr, None)
            if value not in (None, "", [], {}):
                return value
    return None


def _replace_14_1(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(document, display_value(_transport_value(final_document.section_14, "un_number")), contract)


def _replace_14_2(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(
        document,
        display_value(_transport_value(final_document.section_14, "proper_shipping_name", "un_proper_shipping_name")),
        contract,
    )


def _replace_14_3(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(
        document,
        display_value(_transport_value(final_document.section_14, "hazard_class", "transport_hazard_class")),
        contract,
    )


def _replace_14_4(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(
        document,
        display_value(_transport_value(final_document.section_14, "packing_group")),
        contract,
    )


def _replace_14_5(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(
        document,
        display_value(_transport_value(final_document.section_14, "environmental_hazard", "marine_pollutant", "environmental_hazards")),
        contract,
    )


def _replace_14_6(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_text_paragraph(
        document,
        display_value(_transport_value(final_document.section_14, "special_precautions")),
        contract,
    )


def _replace_14_7(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    _add_label_value_paragraph(document, "Further information", final_document.section_14.transport_general_statement, contract)
    render_section14_table(document, final_document.section_14, **_table_kwargs(contract))


def _replace_15(final_document: FinalSDSDocument, document: Document, contract: TemplateContract) -> None:
    rows = [
        ("Regulatory specific information", final_document.section_15.regulatory_specific),
        ("SARA 302", final_document.section_15.sara_302),
        ("SARA 313", final_document.section_15.sara_313),
        ("SARA 311/312", final_document.section_15.sara_311_312),
        ("Proposition 65", final_document.section_15.proposition_65),
        ("China registration", final_document.section_15.china_registration),
        ("Other regulations", final_document.section_15.other_regulations),
    ]
    add_key_value_table(document, rows, **_table_kwargs(contract))
    inventory_rows = [
        [entry.jurisdiction, entry.status, entry.details]
        for entry in final_document.section_15.inventory_statuses
    ] or [[NO_DATA, NO_DATA, NO_DATA]]
    add_matrix_table(
        document,
        ["Jurisdiction", "Status", "Details"],
        inventory_rows,
        column_widths_cm=[4.0, 4.0, 8.0],
        **_table_kwargs(contract),
    )


APPEND_BUILDERS = {
    "1.1 Product identifier": _append_1_1,
    "2.2 GHS Label elements, including precautionary statements": _append_2_2,
    "3.1 Substances": _append_3_1,
}

REPLACE_BUILDERS = {
    "1.2 Relevant identified uses": _replace_1_2,
    "1.3 Supplier information": _replace_1_3,
    "1.4 Emergency telephone": _replace_1_4,
    "2.1 GHS Classification": _replace_2_1,
    "2.3 Other hazards": _replace_2_3,
    "4.1 Description of first-aid measures": _replace_4_1,
    "4.2 Most important symptoms and effects, both acute and delayed": _replace_4_2,
    "4.3 Indication of any immediate medical attention and special treatment needed": _replace_4_3,
    "5.1 Extinguishing media": _replace_5_1,
    "5.2 Special hazards arising from the substance or mixture": _replace_5_2,
    "5.3 Advice for firefighters": _replace_5_3,
    "6.1 Personal precautions, protective equipment and emergency procedures": _replace_6_1,
    "6.2 Environmental precautions": _replace_6_2,
    "6.3 Methods and materials for containment and cleaning up": _replace_6_3,
    "6.4 Reference to other sections": _replace_6_4,
    "7.1 Precautions for safe handling": _replace_7_1,
    "7.2 Conditions for safe storage, including any incompatibilities": _replace_7_2,
    "8.1 Control parameters": _replace_8_1,
    "8.2 Exposure controls": _replace_8_2,
    "9.1 Information on basic physical and chemical properties": _replace_9_1,
    "9.2 Other safety information": _replace_9_2,
    "10.1 Chemical stability": _replace_10_1,
    "10.2 Possibility of hazardous reactions": _replace_10_2,
    "10.3 Conditions to avoid": _replace_10_3,
    "10.4 Incompatible materials": _replace_10_4,
    "11.1 Information on toxicological effects": _replace_11_1,
    "11.2 Additional Information": _replace_11_2,
    "12.1 Toxicity": _replace_12_1,
    "12.2 Persistence and degradability": _replace_12_2,
    "12.3 Bioaccumulative potential": _replace_12_3,
    "12.4 Mobility in soil": _replace_12_4,
    "12.5 Results of PBT and vPvB assessment": _replace_12_5,
    "12.6 Other adverse effects": _replace_12_6,
    "13.1 Waste treatment methods": _replace_13_1,
    "14.1 UN number": _replace_14_1,
    "14.2 UN proper shipping name": _replace_14_2,
    "14.3 Transport hazard class(es)": _replace_14_3,
    "14.4 Packaging group": _replace_14_4,
    "14.5 Environmental hazards": _replace_14_5,
    "14.6 Special precautions for user": _replace_14_6,
    "14.7 Further information": _replace_14_7,
    "SECTION 15: Regulatory information": _replace_15,
}


def _section_16_meta_paragraph(document: Document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("Preparation Date:"):
            return paragraph
    return None


def _section_16_disclaimer_anchor(document: Document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "Disclaimer":
            return paragraph
    return None


def _fill_section_16_metadata(document: Document, final_document: FinalSDSDocument, contract: TemplateContract) -> None:
    paragraph = _section_16_meta_paragraph(document)
    text = (
        f"Preparation Date: {_cover_issue_date(final_document)}\n"
        f"Revision: {_cover_revision_date(final_document)}\n"
        f"Version: {_cover_version(final_document)}"
    )
    if paragraph is not None:
        paragraph.text = text
        return

    disclaimer = _section_16_disclaimer_anchor(document)
    if disclaimer is not None:
        work = _empty_document()
        _add_text_paragraph(work, text, contract)
        disclaimer._p.addprevious(_body_elements(work)[0])


def _fill_section_16_extras(document: Document, final_document: FinalSDSDocument, contract: TemplateContract) -> TemplateFillRange | None:
    disclaimer = _section_16_disclaimer_anchor(document)
    if disclaimer is None:
        return None

    work = _empty_document()
    if final_document.section_16.additional_information:
        _add_label_value_paragraph(work, "Additional information", final_document.section_16.additional_information, contract)
    if final_document.section_16.abbreviations_and_acronyms:
        _add_label_value_paragraph(work, "Abbreviations and acronyms", final_document.section_16.abbreviations_and_acronyms, contract)
    if final_document.section_16.full_text_h_statements:
        _add_label_value_paragraph(work, "Full text hazard statements", final_document.section_16.full_text_h_statements, contract)
    if final_document.section_16.references:
        _add_label_value_paragraph(work, "References", final_document.section_16.references, contract)

    elements = _body_elements(work)
    if not elements:
        return None
    insertion_anchor = _section_16_meta_paragraph(document)
    base = insertion_anchor._p if insertion_anchor is not None else disclaimer._p
    _insert_body_elements_after(base, elements)
    return TemplateFillRange(
        anchor="SECTION 16: Other information",
        next_anchor=None,
        insertion_mode="append",
        removed_element_count=0,
        inserted_element_count=len(elements),
    )


def _story_text(story) -> str:
    lines: list[str] = []
    for paragraph in story.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())
    for table in story.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        lines.append(paragraph.text.strip())
    return "\n".join(lines)


def validate_template_filled_document(
    document: Document,
    *,
    product_name: str | None,
    cas_number: str | None,
    contract: TemplateContract | None = None,
) -> list[str]:
    resolved_contract = contract or load_template_contract()
    errors: list[str] = []

    unresolved = find_unresolved_placeholders(document)
    if unresolved:
        errors.append(f"Unresolved template placeholders remain: {', '.join(unresolved)}")

    all_anchors = resolved_contract.anchors.section_headings + resolved_contract.anchors.subsection_headings
    missing_anchors = find_missing_anchors(document, all_anchors)
    if missing_anchors:
        errors.append(f"Rendered document lost required anchors: {', '.join(missing_anchors)}")

    header_values = normalize_display_text("\n".join(_story_text(section.header) for section in document.sections))
    if (product_name or NO_DATA) not in header_values:
        errors.append("Header does not include product name.")
    if normalize_display_text(cas_number or NO_DATA) not in header_values:
        errors.append("Header does not include CAS number.")

    fixed_company = load_fixed_company()
    footer_values = "\n".join(_story_text(section.footer) for section in document.sections)
    if fixed_company["company_name"] not in footer_values:
        errors.append("Footer does not include fixed company name.")
    if fixed_company["website"] not in footer_values:
        errors.append("Footer does not include fixed company website.")
    return errors


def fill_template_document(
    document: Document,
    final_document: FinalSDSDocument,
    *,
    template_path: Path,
    assets_root: Path | None = None,
    contract: TemplateContract | None = None,
) -> tuple[list[ReviewNote], TemplateFillReport]:
    del assets_root
    resolved_contract = contract or load_template_contract()
    placeholders_found = find_placeholders_by_location(document)
    replaced = replace_placeholders_with_report(document, build_template_replacements(final_document))

    all_anchors = resolved_contract.anchors.section_headings + resolved_contract.anchors.subsection_headings
    ordered = ordered_anchor_paragraphs(document, all_anchors)
    ranges: list[TemplateFillRange] = []

    for index in range(len(ordered) - 1, -1, -1):
        anchor_text, anchor_paragraph = ordered[index]
        next_anchor_paragraph = ordered[index + 1][1] if index + 1 < len(ordered) else None
        if anchor_text in REPLACE_BUILDERS:
            removed = _clear_between(anchor_paragraph, next_anchor_paragraph)
            work = _empty_document()
            REPLACE_BUILDERS[anchor_text](final_document, work, resolved_contract)
            elements = _body_elements(work)
            _insert_body_elements_after(anchor_paragraph._p, elements)
            ranges.append(
                TemplateFillRange(
                    anchor=anchor_text,
                    next_anchor=ordered[index + 1][0] if index + 1 < len(ordered) else None,
                    insertion_mode="replace",
                    removed_element_count=removed,
                    inserted_element_count=len(elements),
                )
            )
        elif anchor_text in APPEND_BUILDERS:
            work = _empty_document()
            APPEND_BUILDERS[anchor_text](final_document, work, resolved_contract)
            elements = _body_elements(work)
            base = _last_element_before(anchor_paragraph, next_anchor_paragraph)
            _insert_body_elements_after(base, elements)
            ranges.append(
                TemplateFillRange(
                    anchor=anchor_text,
                    next_anchor=ordered[index + 1][0] if index + 1 < len(ordered) else None,
                    insertion_mode="append",
                    removed_element_count=0,
                    inserted_element_count=len(elements),
                )
            )

    _fill_section_16_metadata(document, final_document, resolved_contract)
    section_16_range = _fill_section_16_extras(document, final_document, resolved_contract)
    if section_16_range is not None:
        ranges.append(section_16_range)

    occurrences = find_anchor_occurrences(document, all_anchors)
    unresolved = find_unresolved_placeholders(document)
    duplicate_anchors = {anchor: count for anchor, count in occurrences.items() if count > 1}

    warnings: list[str] = []
    if unresolved:
        warnings.append(f"Unresolved template placeholders remain: {', '.join(unresolved)}")

    fixed_company = load_fixed_company()
    footer_values = "\n".join(_story_text(section.footer) for section in document.sections)
    if fixed_company["company_name"] not in footer_values:
        warnings.append("Template footer no longer contains the fixed company name.")
    if fixed_company["website"] not in footer_values:
        warnings.append("Template footer no longer contains the fixed company website.")

    anchor_order = {anchor: index for index, anchor in enumerate(all_anchors)}
    ordered_ranges = sorted(ranges, key=lambda entry: anchor_order.get(entry.anchor, len(anchor_order)))

    notes = [_layout_note("layout.template.fill", message) for message in warnings]
    report = TemplateFillReport(
        template_path=str(template_path),
        contract_version=resolved_contract.contract_version,
        placeholders_found=placeholders_found,
        placeholders_replaced=replaced,
        unresolved_placeholders=unresolved,
        anchors_found=[anchor for anchor, count in occurrences.items() if count > 0],
        missing_anchors=[anchor for anchor, count in occurrences.items() if count == 0],
        duplicate_anchors=duplicate_anchors,
        content_ranges_replaced=ordered_ranges,
        warnings=warnings,
    )
    return notes, report

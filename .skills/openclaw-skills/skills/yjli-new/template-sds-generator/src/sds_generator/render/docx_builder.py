from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Cm
from docx.oxml.ns import qn

from sds_generator.config_loader import PROJECT_ROOT, load_fixed_company
from sds_generator.constants import EXPECTED_SECTION_ORDER, NO_DATA, SECTION_TITLES
from sds_generator.models import ChecklistBucket, FinalSDSDocument, ReviewNote, ReviewSeverity, ReviewStatus
from sds_generator.render.header_footer import apply_header_footer, footer_text, format_cas_display, header_text
from sds_generator.render.section_blocks import render_sections
from sds_generator.render.styles import META_STYLE, TITLE_STYLE, ensure_styles
from sds_generator.render.template_fill import fill_template_document, validate_template_filled_document
from sds_generator.render.template_validation import validate_template_against_contract

DEFAULT_TEMPLATE_PATH = PROJECT_ROOT / "assets" / "templates" / "sds_base.docx"


def _clear_document_body(document: DocumentType) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _layout_note(field_path: str, message: str) -> ReviewNote:
    return ReviewNote(
        field_path=field_path,
        severity=ReviewSeverity.MINOR,
        status=ReviewStatus.WARNING,
        why=message,
        checklist_bucket=ChecklistBucket.LAYOUT_BRANDING_QA,
    )


def create_base_template(template_path: Path | None = None) -> Path:
    target = template_path or DEFAULT_TEMPLATE_PATH
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    ensure_styles(document)

    fixed_company = load_fixed_company()
    apply_header_footer(
        document,
        product_name="{{PRODUCT_NAME_DISPLAY}}",
        cas_number="{{CAS_NUMBER}}",
        company_name=fixed_company["company_name"],
        website=fixed_company["website"],
        logo_path=None,
    )

    title = document.add_paragraph(style=TITLE_STYLE)
    title.add_run("SAFETY DATA SHEET")
    meta = document.add_paragraph(style=META_STYLE)
    meta.add_run("Product: {{PRODUCT_NAME_DISPLAY}} | CAS: {{CAS_NUMBER}}")
    document.save(target)
    return target


def extract_section_headings(document: DocumentType) -> list[str]:
    expected = set(SECTION_TITLES.values())
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip() in expected]


def validate_rendered_document(document: DocumentType, *, product_name: str | None, cas_number: str | None) -> list[str]:
    errors: list[str] = []
    header = header_text(document)
    footer = footer_text(document)
    if (product_name or NO_DATA) not in header:
        errors.append("Header does not include product name.")
    if (cas_number or NO_DATA) not in header:
        errors.append("Header does not include CAS number.")

    fixed_company = load_fixed_company()
    if fixed_company["company_name"] not in footer:
        errors.append("Footer does not include fixed company name.")
    if fixed_company["website"] not in footer:
        errors.append("Footer does not include fixed company website.")

    headings = extract_section_headings(document)
    expected_headings = [SECTION_TITLES[index] for index in range(1, 17)]
    if headings != expected_headings:
        errors.append("Section headings are missing or out of order.")

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if not "\n".join(paragraph.text for paragraph in cell.paragraphs).strip():
                    errors.append("Rendered table contains an empty cell.")
                    return errors
    return errors


def _apply_fixed_company(final_document: FinalSDSDocument) -> None:
    fixed_company = load_fixed_company()
    final_document.section_1.supplier_company_name = fixed_company["company_name"]
    final_document.section_1.supplier_address = fixed_company["address"]
    final_document.section_1.supplier_telephone = fixed_company["telephone"]
    final_document.section_1.supplier_fax = fixed_company["fax"]
    final_document.section_1.supplier_website = fixed_company["website"]
    final_document.section_1.supplier_email = fixed_company["email"]
    final_document.section_1.emergency_telephone = fixed_company["emergency_telephone"]
    final_document.section_1.prepared_by = fixed_company["prepared_by"]


def _render_title_block(document: DocumentType, final_document: FinalSDSDocument) -> None:
    title = document.add_paragraph(style=TITLE_STYLE)
    title.add_run("SAFETY DATA SHEET")
    meta = document.add_paragraph(style=META_STYLE)
    meta.add_run(
        f"Product: {final_document.document_meta.product_name_display or NO_DATA} | "
        f"CAS: {format_cas_display(final_document.document_meta.cas_number)}"
    )


def build_docx(
    final_document: FinalSDSDocument,
    output_path: Path,
    *,
    template_path: Path | None = None,
    logo_path: Path | None = None,
    assets_root: Path | None = None,
    preserve_template_layout: bool = False,
    template_fill_report_path: Path | None = None,
) -> tuple[Path, list[ReviewNote]]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    source_template = template_path or DEFAULT_TEMPLATE_PATH
    if not source_template.exists():
        create_base_template(source_template)

    _apply_fixed_company(final_document)
    document = Document(str(source_template))
    if preserve_template_layout:
        validate_template_against_contract(source_template)
        layout_notes, fill_report = fill_template_document(
            document,
            final_document,
            template_path=source_template,
            assets_root=assets_root or PROJECT_ROOT / "assets",
        )
        if template_fill_report_path is not None:
            template_fill_report_path.parent.mkdir(parents=True, exist_ok=True)
            template_fill_report_path.write_text(
                json.dumps(fill_report.model_dump(mode="json"), indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
        errors = validate_template_filled_document(
            document,
            product_name=final_document.document_meta.product_name_display,
            cas_number=final_document.document_meta.cas_number,
        )
        if errors:
            raise ValueError("; ".join(errors))
        section_notes: list[ReviewNote] = []
    else:
        ensure_styles(document)
        _clear_document_body(document)
        _render_title_block(document, final_document)

        fixed_company = load_fixed_company()
        layout_notes = [
            _layout_note("layout.header.logo", message)
            for message in apply_header_footer(
                document,
                product_name=final_document.document_meta.product_name_display,
                cas_number=final_document.document_meta.cas_number,
                company_name=fixed_company["company_name"],
                website=fixed_company["website"],
                logo_path=logo_path,
            )
        ]
        section_notes = [
            _layout_note(field_path, message)
            for field_path, message in render_sections(document, final_document, assets_root=assets_root or PROJECT_ROOT / "assets")
        ]

        errors = validate_rendered_document(
            document,
            product_name=final_document.document_meta.product_name_display,
            cas_number=final_document.document_meta.cas_number,
        )
        if errors:
            raise ValueError("; ".join(errors))

    if [section for section, _model in final_document.ordered_sections()] != EXPECTED_SECTION_ORDER:
        raise ValueError("Final document model sections are out of order.")

    document.save(str(output_path))
    return output_path, layout_notes + section_notes

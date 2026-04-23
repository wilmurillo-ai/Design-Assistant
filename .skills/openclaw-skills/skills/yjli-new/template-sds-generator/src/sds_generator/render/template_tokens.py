from __future__ import annotations

import re
from collections import Counter
from collections.abc import Iterable

from docx.document import Document as DocumentType

PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


def _paragraphs_in_story(story) -> Iterable:
    for paragraph in story.paragraphs:
        yield paragraph
    for table in story.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def iter_template_paragraphs(document: DocumentType) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in document.sections:
        for story in (section.header, section.footer):
            yield from _paragraphs_in_story(story)


def _body_paragraphs(document: DocumentType) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _header_paragraphs(document: DocumentType) -> Iterable:
    for section in document.sections:
        yield from _paragraphs_in_story(section.header)


def _footer_paragraphs(document: DocumentType) -> Iterable:
    for section in document.sections:
        yield from _paragraphs_in_story(section.footer)


def _collect_placeholders_from_paragraphs(paragraphs: Iterable) -> list[str]:
    tokens: set[str] = set()
    for paragraph in paragraphs:
        text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        tokens.update(match.group(0) for match in PLACEHOLDER_RE.finditer(text))
    return sorted(tokens)


def find_placeholders_by_location(document: DocumentType) -> dict[str, list[str]]:
    return {
        "body": _collect_placeholders_from_paragraphs(_body_paragraphs(document)),
        "header": _collect_placeholders_from_paragraphs(_header_paragraphs(document)),
        "footer": _collect_placeholders_from_paragraphs(_footer_paragraphs(document)),
    }


def find_unresolved_placeholders(document: DocumentType) -> list[str]:
    unresolved: set[str] = set()
    for tokens in find_placeholders_by_location(document).values():
        unresolved.update(tokens)
    return sorted(unresolved)


def _replace_span_in_runs(runs, start: int, end: int, replacement: str) -> None:
    offset = 0
    overlaps: list[tuple[int, int, int]] = []
    for index, run in enumerate(runs):
        run_end = offset + len(run.text)
        if start < run_end and end > offset:
            overlaps.append((index, max(start - offset, 0), min(end - offset, len(run.text))))
        offset = run_end

    if not overlaps:
        return

    first_index, first_start, _first_end = overlaps[0]
    last_index, _last_start, last_end = overlaps[-1]

    if first_index == last_index:
        run = runs[first_index]
        run.text = run.text[:first_start] + replacement + run.text[last_end:]
        return

    first_run = runs[first_index]
    last_run = runs[last_index]
    first_run.text = first_run.text[:first_start] + replacement
    last_run.text = last_run.text[last_end:]
    for index, _local_start, _local_end in overlaps[1:-1]:
        runs[index].text = ""


def replace_placeholders_in_paragraph(paragraph, replacements: dict[str, str]) -> Counter[str]:
    if paragraph.runs:
        source_text = "".join(run.text for run in paragraph.runs)
    else:
        source_text = paragraph.text

    matches = [
        (match.start(), match.end(), match.group(0)[2:-2])
        for match in PLACEHOLDER_RE.finditer(source_text)
        if match.group(0)[2:-2] in replacements
    ]
    if not matches:
        return Counter()

    replaced: Counter[str] = Counter()
    if not paragraph.runs:
        updated = source_text
        for start, end, token_name in reversed(matches):
            updated = updated[:start] + replacements[token_name] + updated[end:]
            replaced[token_name] += 1
        paragraph.text = updated
        return replaced

    for start, end, token_name in reversed(matches):
        _replace_span_in_runs(paragraph.runs, start, end, replacements[token_name])
        replaced[token_name] += 1
    return replaced


def replace_placeholders_with_report(document: DocumentType, replacements: dict[str, str]) -> dict[str, int]:
    replaced = Counter()
    for paragraph in iter_template_paragraphs(document):
        replaced.update(replace_placeholders_in_paragraph(paragraph, replacements))
    return dict(sorted(replaced.items()))


def replace_placeholders_in_document(document: DocumentType, replacements: dict[str, str]) -> int:
    return sum(replace_placeholders_with_report(document, replacements).values())

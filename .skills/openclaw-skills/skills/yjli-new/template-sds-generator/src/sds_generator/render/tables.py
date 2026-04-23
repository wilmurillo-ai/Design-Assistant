from __future__ import annotations

from collections.abc import Iterable
from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from sds_generator.constants import NO_DATA
from sds_generator.models import Section14TransportInformation, Section3Composition, Section9PhysicalChemicalProperties
from sds_generator.render.styles import BODY_TIGHT_STYLE

TABLE_WIDTH_GUTTER_CM = 0.2


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _display_value(value: Any) -> str:
    if value in (None, "", [], {}):
        return NO_DATA
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, str):
        return value
    if isinstance(value, Iterable) and not isinstance(value, (str, bytes, dict)):
        rendered = [str(item) for item in value if item not in (None, "", [], {})]
        return "\n".join(rendered) if rendered else NO_DATA
    return str(value)


def _write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    paragraph_style: str | None = BODY_TIGHT_STYLE,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    text_color: str | None = None,
) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph_style:
        paragraph.style = paragraph_style
    run = paragraph.add_run(text or NO_DATA)
    run.bold = bold
    if font_name:
        run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if text_color:
        run.font.color.rgb = RGBColor.from_string(text_color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _usable_table_width_cm(document) -> float:
    section = document.sections[-1]
    usable_width = section.page_width.cm - section.left_margin.cm - section.right_margin.cm - TABLE_WIDTH_GUTTER_CM
    return max(usable_width, 1.0)


def _resolve_column_widths_cm(document, widths_cm: list[float]) -> list[float]:
    max_total_width = _usable_table_width_cm(document)
    requested_total = sum(widths_cm)
    if requested_total <= max_total_width:
        return widths_cm

    scale = max_total_width / requested_total
    scaled = [round(width * scale, 3) for width in widths_cm]
    scaled[-1] = round(scaled[-1] + (max_total_width - sum(scaled)), 3)
    return scaled


def _apply_column_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, width_cm in enumerate(widths_cm):
        width = Cm(width_cm)
        table.columns[index].width = width
        for cell in table.columns[index].cells:
            cell.width = width


def add_key_value_table(
    document,
    rows: list[tuple[str, Any]],
    *,
    column_widths_cm: tuple[float, float] = (5.0, 11.0),
    paragraph_style: str | None = BODY_TIGHT_STYLE,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    text_color: str | None = None,
):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    _write_cell(
        header[0],
        "Property",
        bold=True,
        paragraph_style=paragraph_style,
        font_name=font_name,
        font_size_pt=font_size_pt,
        text_color=text_color,
    )
    _write_cell(
        header[1],
        "Value",
        bold=True,
        paragraph_style=paragraph_style,
        font_name=font_name,
        font_size_pt=font_size_pt,
        text_color=text_color,
    )
    _set_cell_shading(header[0], "D7E2EB")
    _set_cell_shading(header[1], "D7E2EB")
    for label, value in rows:
        cells = table.add_row().cells
        _write_cell(
            cells[0],
            label,
            bold=True,
            paragraph_style=paragraph_style,
            font_name=font_name,
            font_size_pt=font_size_pt,
            text_color=text_color,
        )
        _write_cell(
            cells[1],
            _display_value(value),
            paragraph_style=paragraph_style,
            font_name=font_name,
            font_size_pt=font_size_pt,
            text_color=text_color,
        )
    _apply_column_widths(table, _resolve_column_widths_cm(document, list(column_widths_cm)))
    return table


def add_matrix_table(
    document,
    headers: list[str],
    rows: list[list[Any]],
    *,
    column_widths_cm: list[float] | None = None,
    paragraph_style: str | None = BODY_TIGHT_STYLE,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    text_color: str | None = None,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _write_cell(
            cell,
            header,
            bold=True,
            paragraph_style=paragraph_style,
            font_name=font_name,
            font_size_pt=font_size_pt,
            text_color=text_color,
        )
        _set_cell_shading(cell, "D7E2EB")
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            _write_cell(
                cells[index],
                _display_value(value),
                paragraph_style=paragraph_style,
                font_name=font_name,
                font_size_pt=font_size_pt,
                text_color=text_color,
            )
    requested_widths = column_widths_cm or [(_usable_table_width_cm(document) / max(len(headers), 1))] * len(headers)
    _apply_column_widths(table, _resolve_column_widths_cm(document, requested_widths))
    return table


def render_section3_table(document, section: Section3Composition, **table_kwargs):
    rows = [
        ("Substance / Mixture", section.substance_or_mixture),
        ("Chemical name", section.chemical_name),
        ("Synonyms", section.common_name_and_synonyms),
        ("CAS No.", section.cas_number),
        ("EC No.", section.ec_number),
        ("Formula", section.formula),
        ("Molecular weight", section.molecular_weight),
        ("Disclosure statement", section.disclosure_statement),
    ]
    return add_key_value_table(document, rows, **table_kwargs)


def render_section9_table(document, section: Section9PhysicalChemicalProperties, **table_kwargs):
    rows = [
        ("Physical state", section.physical_state),
        ("Appearance / form", section.appearance_form),
        ("Colour", section.colour),
        ("Odour", section.odour),
        ("Odour threshold", section.odour_threshold),
        ("pH", section.pH),
        ("Melting point / freezing point", section.melting_point_freezing_point),
        ("Initial boiling point and range", section.initial_boiling_point_and_range),
        ("Flash point", section.flash_point),
        ("Evaporation rate", section.evaporation_rate),
        ("Flammability", section.flammability),
        ("Upper flammability limit", section.upper_flammability_limit),
        ("Lower flammability limit", section.lower_flammability_limit),
        ("Vapor pressure", section.vapor_pressure),
        ("Vapor density", section.vapor_density),
        ("Relative density", section.relative_density),
        ("Density", section.density),
        ("Water solubility", section.water_solubility),
        ("Partition coefficient log Pow", section.partition_coefficient_log_pow),
        ("Autoignition temperature", section.autoignition_temperature),
        ("Decomposition temperature", section.decomposition_temperature),
        ("Viscosity", section.viscosity),
        ("Explosive properties", section.explosive_properties),
        ("Oxidizing properties", section.oxidizing_properties),
        ("Particle characteristics", section.particle_characteristics),
        ("Other safety information", section.other_safety_information),
    ]
    return add_key_value_table(document, rows, **table_kwargs)


def render_section14_table(document, section: Section14TransportInformation, **table_kwargs):
    rows: list[list[Any]] = []
    for mode in section.transport_modes or []:
        rows.append(
            [
                str(mode.mode).split(".")[-1].upper(),
                mode.un_number,
                mode.proper_shipping_name,
                mode.hazard_class,
                mode.packing_group,
                mode.environmental_hazard or mode.marine_pollutant,
                mode.status_note,
            ]
        )
    if not rows:
        rows.append([NO_DATA, NO_DATA, NO_DATA, NO_DATA, NO_DATA, NO_DATA, NO_DATA])
    return add_matrix_table(
        document,
        ["Mode", "UN No.", "Proper shipping name", "Class", "Packing group", "Environmental hazard", "Status note"],
        rows,
        column_widths_cm=[2.2, 2.4, 4.8, 1.8, 2.2, 3.0, 3.0],
        **table_kwargs,
    )

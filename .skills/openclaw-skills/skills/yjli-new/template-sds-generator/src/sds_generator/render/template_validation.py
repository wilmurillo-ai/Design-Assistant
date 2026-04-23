from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from sds_generator.config_loader import load_template_contract
from sds_generator.exceptions import InvalidTemplateContractError
from sds_generator.models import TemplateContract, TemplateValidationReport

PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


def _collect_text_paragraphs(paragraphs) -> list[str]:
    return [paragraph.text for paragraph in paragraphs if paragraph.text]


def _collect_table_paragraphs(tables) -> list[str]:
    texts: list[str] = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(_collect_text_paragraphs(cell.paragraphs))
    return texts


def _collect_placeholders(texts: list[str]) -> list[str]:
    tokens: set[str] = set()
    for text in texts:
        tokens.update(match.group(0) for match in PLACEHOLDER_RE.finditer(text))
    return sorted(tokens)


def _body_texts(document: Document) -> list[str]:
    return _collect_text_paragraphs(document.paragraphs) + _collect_table_paragraphs(document.tables)


def _header_texts(document: Document) -> list[str]:
    texts: list[str] = []
    for section in document.sections:
        texts.extend(_collect_text_paragraphs(section.header.paragraphs))
        texts.extend(_collect_table_paragraphs(section.header.tables))
    return texts


def _footer_texts(document: Document) -> list[str]:
    texts: list[str] = []
    for section in document.sections:
        texts.extend(_collect_text_paragraphs(section.footer.paragraphs))
        texts.extend(_collect_table_paragraphs(section.footer.tables))
    return texts


def _body_anchors(document: Document) -> list[str]:
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def _page_number_fields_present(template_path: Path) -> bool:
    with ZipFile(template_path) as archive:
        footer_xml = "".join(
            archive.read(name).decode("utf-8", "ignore")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    return "PAGE" in footer_xml and "NUMPAGES" in footer_xml


def build_template_validation_report(
    template_path: str | Path,
    contract: TemplateContract | None = None,
) -> TemplateValidationReport:
    resolved_contract = contract or load_template_contract()
    path = Path(template_path)
    document = Document(path)

    placeholders_found = {
        "body": _collect_placeholders(_body_texts(document)),
        "header": _collect_placeholders(_header_texts(document)),
        "footer": _collect_placeholders(_footer_texts(document)),
    }

    missing_placeholders: dict[str, list[str]] = {}
    for location in ("body", "header", "footer"):
        required = set(getattr(resolved_contract.placeholders, location).required)
        found = set(placeholders_found[location])
        missing = sorted(required - found)
        if missing:
            missing_placeholders[location] = missing

    all_anchors = resolved_contract.anchors.section_headings + resolved_contract.anchors.subsection_headings
    found_anchor_set = set(_body_anchors(document))
    anchors_found = [anchor for anchor in all_anchors if anchor in found_anchor_set]
    missing_anchors = [anchor for anchor in all_anchors if anchor not in found_anchor_set]

    header_table_count = sum(len(section.header.tables) for section in document.sections)
    footer_table_count = sum(len(section.footer.tables) for section in document.sections)

    return TemplateValidationReport(
        template_path=str(path),
        contract_version=resolved_contract.contract_version,
        placeholders_found=placeholders_found,
        missing_placeholders=missing_placeholders,
        anchors_found=anchors_found,
        missing_anchors=missing_anchors,
        header_table_count=header_table_count,
        footer_table_count=footer_table_count,
        page_number_fields_present=_page_number_fields_present(path),
    )


def validate_template_against_contract(
    template_path: str | Path,
    contract: TemplateContract | None = None,
) -> TemplateValidationReport:
    resolved_contract = contract or load_template_contract()
    report = build_template_validation_report(template_path, resolved_contract)

    errors: list[str] = []
    for location, tokens in report.missing_placeholders.items():
        errors.append(f"Missing required {location} placeholders: {', '.join(tokens)}")
    if report.missing_anchors:
        errors.append(f"Missing required anchors: {', '.join(report.missing_anchors)}")
    if resolved_contract.preservation.preserve_header_footer_tables:
        if report.header_table_count == 0:
            errors.append("Template header is missing required header table(s).")
        if report.footer_table_count == 0:
            errors.append("Template footer is missing required footer table(s).")
    if resolved_contract.preservation.preserve_page_number_fields and not report.page_number_fields_present:
        errors.append("Template footer is missing PAGE/NUMPAGES fields.")

    if errors:
        raise InvalidTemplateContractError(
            f"Template contract validation failed for {Path(template_path).name}: " + " | ".join(errors)
        )
    return report

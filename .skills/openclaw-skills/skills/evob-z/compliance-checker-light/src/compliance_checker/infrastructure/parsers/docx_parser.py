"""
Word 文档解析器 - Infrastructure 层实现

实现 Core 层定义的 DocumentParserProtocol 接口。
支持 .docx 格式，所有异常转换为 ComplianceCheckerError。
"""

import logging
from pathlib import Path
from typing import List, Optional

from ...core.document import Document, DocumentMetadata, DocumentType, PageContent
from ...core.exceptions import DocumentParseError
from ...core.interfaces import DocumentParserProtocol

logger = logging.getLogger(__name__)

# 可选依赖：python-docx
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DOCX_AVAILABLE = False


class DocxParser(DocumentParserProtocol):
    """
    Word 文档解析器 (.docx)

    实现 DocumentParserProtocol 接口，提供 Word 文档的解析功能。
    功能包括：
    - 提取文档文本内容
    - 提取文档元数据
    - 估算页数（按内容分割为虚拟页面）
    """

    def __init__(self, chars_per_page: int = 3000):
        """
        初始化 Word 解析器

        Args:
            chars_per_page: 每页字符数（用于虚拟分页）

        Raises:
            DocumentParseError: python-docx 未安装
        """
        if not DOCX_AVAILABLE:
            raise DocumentParseError("python-docx 未安装，请运行: pip install python-docx")

        self.chars_per_page = chars_per_page

    def parse(self, file_path: str) -> Document:
        """
        解析 Word 文档

        Args:
            file_path: Word 文件路径

        Returns:
            Document 对象

        Raises:
            DocumentParseError: 文件不存在、格式不支持或解析失败
        """
        path = Path(file_path)

        # 验证文件存在性
        if not path.exists():
            raise DocumentParseError(f"文件不存在: {file_path}")

        # 验证文件格式
        suffix = path.suffix.lower()
        if suffix not in [".docx", ".doc"]:
            raise DocumentParseError(f"不支持的文件格式: {path.suffix}")

        # 暂不支持 .doc 格式
        if suffix == ".doc":
            raise DocumentParseError("暂不支持 .doc 格式，请先转换为 .docx")

        logger.info(f"开始解析 Word 文档: {file_path}")

        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise DocumentParseError(f"无法打开 Word 文档: {e}") from e

        try:
            # 提取元数据
            metadata = self._extract_metadata(doc, file_path)

            # 提取页面内容
            pages_content = self._extract_content(doc)

            # 创建 Document 对象
            document = Document(
                path=str(path.absolute()),
                name=path.name,
                type=DocumentType.DOCX,
                pages=len(pages_content),
                pages_content=pages_content,
                metadata=metadata,
            )

            logger.info(f"Word 文档解析完成: {path.name}, 共 {len(pages_content)} 页")
            return document

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Word 文档解析失败: {e}") from e

    def _extract_metadata(self, doc: "docx.Document", file_path: str) -> DocumentMetadata:
        """提取 Word 文档元数据"""
        metadata = DocumentMetadata()

        try:
            core_props = doc.core_properties

            metadata.title = core_props.title or None
            metadata.author = core_props.author or None
            metadata.subject = core_props.subject or None
            metadata.creator = core_props.creator or None

            if core_props.created:
                metadata.creation_date = core_props.created
            if core_props.modified:
                metadata.modification_date = core_props.modified

        except Exception as e:
            logger.debug(f"提取元数据失败: {e}")

        return metadata

    def _extract_content(self, doc: "docx.Document") -> List[PageContent]:
        """
        提取文档内容

        Word 文档没有明确的页面概念，这里按内容组织并分割为虚拟页面。
        """
        paragraphs = []

        # 提取所有段落文本
        try:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
        except Exception as e:
            logger.warning(f"提取段落失败: {e}")

        # 提取表格内容
        try:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
        except Exception as e:
            logger.warning(f"提取表格失败: {e}")

        # 合并所有内容
        full_text = "\n".join(paragraphs)

        # 按字符数分割为虚拟页面
        pages_content = []

        if not full_text:
            # 空文档
            pages_content.append(PageContent(page_num=0, text="", has_text_layer=True))
        elif len(full_text) <= self.chars_per_page:
            # 单页文档
            pages_content.append(PageContent(page_num=0, text=full_text, has_text_layer=True))
        else:
            # 分割为多页
            page_num = 0
            for i in range(0, len(full_text), self.chars_per_page):
                page_text = full_text[i : i + self.chars_per_page]
                pages_content.append(
                    PageContent(
                        page_num=page_num,
                        text=page_text,
                        has_text_layer=True,
                    )
                )
                page_num += 1

        return pages_content


def parse_docx(file_path: str, chars_per_page: int = 3000) -> Document:
    """
    便捷函数：解析 Word 文档

    Args:
        file_path: Word 文件路径
        chars_per_page: 每页字符数（用于虚拟分页）

    Returns:
        Document 对象

    Raises:
        DocumentParseError: 解析失败
    """
    parser = DocxParser(chars_per_page=chars_per_page)
    return parser.parse(file_path)

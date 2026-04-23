from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate Harvard-style resumes from structured JSON input.")
    parser.add_argument("--input", required=True, help="Path to candidate JSON input.")
    parser.add_argument("--track", default="all", help="Track key to generate, or 'all'.")
    parser.add_argument("--output-dir", required=True, help="Directory for generated files.")
    parser.add_argument("--pdf", action="store_true", help="Also export PDF if LibreOffice is available.")
    return parser.parse_args()


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def safe_filename(text: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "_", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text


def set_run_font(run, *, latin_font: str, cjk_font: str, size: float = 11, bold: bool = False):
    run.font.name = latin_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = None


def set_paragraph(paragraph, *, before: float = 0, after: float = 0, line: float = 1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def ensure_styles(doc: Document, *, latin_font: str, cjk_font: str):
    normal = doc.styles["Normal"]
    normal.font.name = latin_font
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
    normal.font.size = Pt(11)

    if "Resume Section" not in doc.styles:
        style = doc.styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]


def configure_document(doc: Document, *, latin_font: str, cjk_font: str):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    ensure_styles(doc, latin_font=latin_font, cjk_font=cjk_font)


def add_name(doc: Document, text: str, *, latin_font: str, cjk_font: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=3)
    run = p.add_run(text)
    set_run_font(run, latin_font=latin_font, cjk_font=cjk_font, size=16, bold=True)


def add_title(doc: Document, text: str, *, latin_font: str, cjk_font: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=3)
    run = p.add_run(text)
    set_run_font(run, latin_font=latin_font, cjk_font=cjk_font, size=11, bold=True)


def add_contact(doc: Document, text: str, *, latin_font: str, cjk_font: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=10)
    run = p.add_run(text)
    set_run_font(run, latin_font=latin_font, cjk_font=cjk_font, size=10)


def add_section(doc: Document, title: str, *, latin_font: str, cjk_font: str):
    p = doc.add_paragraph()
    set_paragraph(p, before=7, after=4)
    run = p.add_run(title)
    set_run_font(run, latin_font=latin_font, cjk_font=cjk_font, size=11, bold=True)
    add_bottom_border(p)


def add_text(doc: Document, text: str, *, latin_font: str, cjk_font: str, after: float = 5):
    p = doc.add_paragraph()
    set_paragraph(p, after=after, line=1.15)
    run = p.add_run(text)
    set_run_font(run, latin_font=latin_font, cjk_font=cjk_font, size=11)


def add_header_line(doc: Document, left: str, right: str, *, latin_font: str, cjk_font: str):
    p = doc.add_paragraph()
    set_paragraph(p, before=2, after=2)
    tabs = p.paragraph_format.tab_stops
    tabs.clear_all()
    tabs.add_tab_stop(Inches(6.55))
    left_run = p.add_run(left)
    set_run_font(left_run, latin_font=latin_font, cjk_font=cjk_font, bold=True)
    p.add_run("\t")
    right_run = p.add_run(right)
    set_run_font(right_run, latin_font=latin_font, cjk_font=cjk_font)


def add_bullet(doc: Document, text: str, *, latin_font: str, cjk_font: str):
    p = doc.add_paragraph()
    set_paragraph(p, after=2, line=1.15)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.22)
    fmt.first_line_indent = Inches(-0.16)
    run = p.add_run(f"- {text}")
    set_run_font(run, latin_font=latin_font, cjk_font=cjk_font, size=10.8)


def experience_map(data: dict) -> dict:
    return {item["key"]: item for item in data.get("experience", [])}


def build_track_doc(data: dict, track_key: str, output_dir: Path, export_pdf: bool) -> Path:
    candidate = data["candidate"]
    style = data.get("style", {})
    latin_font = style.get("latin_font", "Times New Roman")
    cjk_font = style.get("cjk_font", "宋体")
    track = data["tracks"][track_key]
    exp_index = experience_map(data)

    doc = Document()
    configure_document(doc, latin_font=latin_font, cjk_font=cjk_font)

    add_name(doc, candidate["name"], latin_font=latin_font, cjk_font=cjk_font)
    add_title(doc, track["title"], latin_font=latin_font, cjk_font=cjk_font)
    add_contact(doc, candidate["contact_line"], latin_font=latin_font, cjk_font=cjk_font)

    add_section(doc, "个人简介", latin_font=latin_font, cjk_font=cjk_font)
    add_text(doc, track["summary"], latin_font=latin_font, cjk_font=cjk_font, after=6)

    add_section(doc, "工作经历", latin_font=latin_font, cjk_font=cjk_font)
    overrides = track.get("experience_overrides", {})
    for exp_key in track.get("experience_order", []):
        base = exp_index[exp_key]
        override = overrides.get(exp_key, {})
        left = override.get("left", base["default_left"])
        right = override.get("right", base["default_right"])
        bullets = override.get("bullets", base.get("default_bullets", []))
        add_header_line(doc, left, right, latin_font=latin_font, cjk_font=cjk_font)
        for bullet in bullets:
            add_bullet(doc, bullet, latin_font=latin_font, cjk_font=cjk_font)

    add_section(doc, "教育背景", latin_font=latin_font, cjk_font=cjk_font)
    for edu in data.get("education", []):
        add_header_line(doc, edu["left"], edu["right"], latin_font=latin_font, cjk_font=cjk_font)
        for bullet in edu.get("bullets", []):
            add_bullet(doc, bullet, latin_font=latin_font, cjk_font=cjk_font)

    campus_items = track.get("campus_items", data.get("campus", []))
    if campus_items:
        add_section(doc, "项目与校园经历", latin_font=latin_font, cjk_font=cjk_font)
        for item in campus_items:
            add_header_line(doc, item["left"], item["right"], latin_font=latin_font, cjk_font=cjk_font)
            for bullet in item.get("bullets", []):
                add_bullet(doc, bullet, latin_font=latin_font, cjk_font=cjk_font)

    if track.get("skills_lines"):
        add_section(doc, "技能与证书", latin_font=latin_font, cjk_font=cjk_font)
        for line in track["skills_lines"]:
            add_text(doc, line, latin_font=latin_font, cjk_font=cjk_font, after=2)

    output_dir.mkdir(parents=True, exist_ok=True)
    filename = safe_filename(f"{candidate['name']}_{track_key}.docx")
    output_path = output_dir / filename
    doc.save(output_path)

    if export_pdf:
        convert_to_pdf(output_path)

    return output_path


def convert_to_pdf(docx_path: Path) -> Path | None:
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if not office:
        return None
    try:
        subprocess.run(
            [office, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except subprocess.CalledProcessError:
        return None
    return docx_path.with_suffix(".pdf")


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    output_dir = Path(args.output_dir)
    data = load_json(input_path)
    tracks = data.get("tracks", {})
    if not tracks:
        raise SystemExit("No tracks found in input JSON.")

    track_keys = list(tracks.keys()) if args.track == "all" else [args.track]
    missing = [key for key in track_keys if key not in tracks]
    if missing:
        raise SystemExit(f"Unknown track(s): {', '.join(missing)}")

    generated: list[Path] = []
    for track_key in track_keys:
        generated.append(build_track_doc(data, track_key, output_dir, args.pdf))

    for path in generated:
        print(path)
    return 0


if __name__ == "__main__":
    sys.exit(main())

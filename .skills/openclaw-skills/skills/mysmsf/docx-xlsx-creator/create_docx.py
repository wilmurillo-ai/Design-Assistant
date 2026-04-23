#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create a Word document (.docx)")
    parser.add_argument("--config", help="Path to JSON config file")
    parser.add_argument("--title", help="Document title")
    parser.add_argument("--content", help="Simple paragraph content")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--force", action="store_true", help="Overwrite output file if it already exists")
    return parser.parse_args()


def ensure_output_path(output_path: Path, force: bool) -> None:
    if output_path.exists() and not force:
        raise FileExistsError(f"Output file already exists: {output_path}. Use --force to overwrite.")
    if output_path.suffix.lower() != ".docx":
        raise ValueError("Output file must end with .docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("Config JSON must be an object at the top level")
    return data


def set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_heading(title, level=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraphs(document: Document, paragraphs: Iterable[str]) -> None:
    for paragraph_text in paragraphs:
        document.add_paragraph(str(paragraph_text))


def add_bullets(document: Document, bullets: Iterable[str]) -> None:
    for bullet in bullets:
        document.add_paragraph(str(bullet), style="List Bullet")


def add_numbered_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(str(item), style="List Number")


def add_table(document: Document, rows: list[list[Any]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        if len(row) != column_count:
            raise ValueError("All table rows must contain the same number of columns")
        for col_index, value in enumerate(row):
            table.rows[row_index].cells[col_index].text = str(value)


def build_from_config(document: Document, config: dict[str, Any]) -> None:
    title = config.get("title")
    if title:
        add_title(document, str(title))

    for section in config.get("sections", []):
        if not isinstance(section, dict):
            raise ValueError("Each section must be a JSON object")
        heading = section.get("heading")
        level = int(section.get("level", 1))
        if heading:
            document.add_heading(str(heading), level=max(1, min(level, 9)))
        if section.get("paragraphs"):
            add_paragraphs(document, section["paragraphs"])
        if section.get("bullets"):
            add_bullets(document, section["bullets"])
        if section.get("numbered"):
            add_numbered_list(document, section["numbered"])
        if section.get("table"):
            add_table(document, section["table"])


def main() -> int:
    args = parse_args()
    output_path = Path(args.output).expanduser().resolve()
    try:
        ensure_output_path(output_path, args.force)
        document = Document()
        set_default_style(document)
        if args.config:
            config = load_json(Path(args.config).expanduser().resolve())
            build_from_config(document, config)
        else:
            if args.title:
                add_title(document, args.title)
            if args.content:
                document.add_paragraph(args.content)
        document.save(str(output_path))
        print(f"Created Word document: {output_path}")
        return 0
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""
中文合同审查技能 - Word解析器
支持从Word文件(docx)中提取文本内容
"""

import os
import re
from pathlib import Path
from typing import List
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


class WordParser:
    """Word文档解析器"""

    def __init__(self):
        pass

    def parse(self, file_path: Path) -> str:
        """
        解析Word文档，提取文本内容

        Args:
            file_path: Word文档路径

        Returns:
            提取的文本内容
        """
        try:
            doc = Document(str(file_path))
            text = []

            # 遍历所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())

            # 遍历所有表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append('|'.join(row_text))

            # 生成文本内容
            content = '\n'.join(text)
            return content

        except Exception as e:
            raise Exception(f"解析Word文档失败: {str(e)}")

    def extract_text_from_word(self, file_path: Path, output_text: bool = True) -> str:
        """
        从Word中提取文本

        Args:
            file_path: Word文档路径
            output_text: 是否输出处理过程

        Returns:
            提取的文本内容
        """
        text = self.parse(file_path)

        if output_text:
            print(f"成功从 {file_path} 提取了 {len(text)} 个字符")
            print("\n提取的文本内容:")
            print("=" * 50)
            print(text)
            print("=" * 50)

        return text

    def clean_text(self, text: str) -> str:
        """
        清理和规范化文本

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)

        # 移除特殊的Word格式字符
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # 移除页码等不必要的信息
        text = re.sub(r'Page\s*\d+', '', text)

        # 移除首尾空白
        text = text.strip()

        return text

    def extract_sections(self, text: str) -> dict:
        """
        提取合同中的主要条款

        Args:
            text: 合同文本

        Returns:
            包含各条款的字典
        """
        sections = {
            'title': self._extract_title(text),
            'parties': self._extract_parties(text),
            'definitions': self._extract_definitions(text),
            'obligations': self._extract_obligations(text),
            'payment_terms': self._extract_payment_terms(text),
            'liabilities': self._extract_liabilities(text),
            'termination': self._extract_termination(text),
            'dispute_resolution': self._extract_dispute_resolution(text)
        }

        return sections

    def _extract_title(self, text: str) -> str:
        """提取合同标题"""
        # 尝试从第一行提取标题
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) > 10 and len(line) < 100:
                # 移除常见的标题前缀
                line = re.sub(r'^合同\s*[:：]?\s*', '', line)
                line = re.sub(r'^[\d\w\s]+合同$', '', line)
                return line
        return "未知合同"

    def _extract_parties(self, text: str) -> str:
        """提取合同当事人"""
        # 查找"甲方"、"乙方"等关键词
        patterns = [
            r'甲方[：:]\s*(.*?)(?=\n\n|\n甲方|\n乙方|$)',
            r'乙方[：:]\s*(.*?)(?=\n\n|\n甲方|\n乙方|\n丙方|$)',
            r'买方[：:]\s*(.*?)(?=\n\n|\n卖方|$)',
            r'卖方[：:]\s*(.*?)(?=\n\n|\n买方|$)'
        ]

        parties = {}
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                parties['party'] = match.group(1).strip()

        return parties.get('party', '未找到当事人信息')

    def _extract_definitions(self, text: str) -> str:
        """提取定义条款"""
        # 查找"定义"、"释义"等关键词
        definitions = []

        # 查找"定义"章节
        def_pattern = r'(?:定义|释义|术语说明)[：:](.*?)(?=\n\n(?:第\s*\d+\s*条|权利义务|责任|保证))'
        matches = re.findall(def_pattern, text, re.DOTALL)
        for match in matches:
            definitions.append(match.strip())

        # 查找"本协议中"的术语定义
        term_pattern = r'本协议中[\s\S]{0,100}(?:指|表示|定义)[\s\S]{0,200}(?=，\n|。)'
        matches = re.findall(term_pattern, text)
        for match in matches:
            definitions.append(match.strip())

        return definitions

    def _extract_obligations(self, text: str) -> str:
        """提取义务条款"""
        # 查找"义务"、"职责"等关键词
        obligations = []

        # 查找"甲方义务"、"乙方义务"
        patterns = [
            r'(?:甲方|买方|供应商)[\s\S]{0,200}(?:义务|职责|责任)',
            r'(?:乙方|买方|客户|需方)[\s\S]{0,200}(?:义务|职责|责任)',
            r'(?:双方|卖方|提供方|制造商)[\s\S]{0,200}(?:义务|职责|责任)'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            obligations.extend(matches)

        return '\n'.join(obligations)

    def _extract_payment_terms(self, text: str) -> str:
        """提取付款条款"""
        payment_keywords = ['付款', 'payment', '支付', '款项']

        # 查找包含付款关键词的段落
        lines = text.split('\n')
        payment_lines = []

        for line in lines:
            if any(keyword in line.lower() for keyword in payment_keywords):
                payment_lines.append(line)

        return '\n'.join(payment_lines)

    def _extract_liabilities(self, text: str) -> str:
        """提取责任条款"""
        # 查找"责任"、"违约"等关键词
        liability_keywords = ['责任', '违约', '赔偿责任', 'liability', 'breach']

        lines = text.split('\n')
        liability_lines = []

        for line in lines:
            if any(keyword in line.lower() for keyword in liability_keywords):
                liability_lines.append(line)

        return '\n'.join(liability_lines)

    def _extract_termination(self, text: str) -> str:
        """提取终止条款"""
        # 查找"终止"、"解除"等关键词
        termination_keywords = ['终止', '解除', '解除合同', 'termination', 'expiry', 'end']

        lines = text.split('\n')
        termination_lines = []

        for line in lines:
            if any(keyword in line.lower() for keyword in termination_keywords):
                termination_lines.append(line)

        return '\n'.join(termination_lines)

    def _extract_dispute_resolution(self, text: str) -> str:
        """提取争议解决条款"""
        # 查找"争议解决"、"争议"等关键词
        dispute_keywords = ['争议解决', '争议', '争议解决方式', 'dispute resolution', 'dispute']

        lines = text.split('\n')
        dispute_lines = []

        for line in lines:
            if any(keyword in line.lower() for keyword in dispute_keywords):
                dispute_lines.append(line)

        return '\n'.join(dispute_lines)

    def extract_tables(self, file_path: Path) -> List[List[List[str]]]:
        """
        提取Word文档中的所有表格

        Args:
            file_path: Word文档路径

        Returns:
            表格列表，每个表格是行和列的二维列表
        """
        try:
            doc = Document(str(file_path))
            tables_data = []

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_data.append(table_data)

            return tables_data

        except Exception as e:
            raise Exception(f"提取表格失败: {str(e)}")


def main():
    """测试Word解析器"""
    if len(sys.argv) < 2:
        print("用法: python word_parser.py <word文件路径>")
        sys.exit(1)

    file_path = Path(sys.argv[1])
    parser = WordParser()

    if not file_path.exists():
        print(f"文件不存在: {file_path}")
        sys.exit(1)

    # 解析Word
    text = parser.extract_text_from_word(file_path)

    # 提取合同条款
    sections = parser.extract_sections(text)

    print("\n提取的合同条款:")
    print("=" * 50)
    for section_name, section_text in sections.items():
        if section_text:
            print(f"\n--- {section_name} ---")
            print(section_text[:500] + "..." if len(section_text) > 500 else section_text)


if __name__ == '__main__':
    main()

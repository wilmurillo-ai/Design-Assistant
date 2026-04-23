#!/usr/bin/env python3
"""
extract_tables.py - Extract table structures from Word (.docx) and Excel (.xls/.xlsx) files.

Identifies database schema tables, data dictionary entries, and field definitions.
Outputs structured JSON suitable for ontology extraction by LLM.

Usage:
    python extract_tables.py <file_or_directory> [--output extracted_tables.json]
    python extract_tables.py document.docx --output tables.json
    python extract_tables.py data_dictionary.xlsx --output tables.json
"""

import argparse
import json
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Column role recognition keywords
# ---------------------------------------------------------------------------

COLUMN_ROLE_KEYWORDS = {
    "field_name": ["字段名", "字段名称", "列名", "column", "field name", "英文名", "字段英文", "属性名"],
    "field_name_cn": ["中文名", "字段中文", "中文描述", "中文说明", "业务名称", "字段描述", "字段说明"],
    "field_type": ["数据类型", "字段类型", "数据格式", "类型", "data type", "type"],
    "field_length": ["字段长度", "长度", "length", "size", "精度"],
    "nullable": ["是否为空", "nullable", "允许空", "是否必填", "必填", "null"],
    "primary_key": ["是否主键", "主键", "pk", "primary"],
    "foreign_key": ["外键", "fk", "foreign", "关联", "引用"],
    "default_value": ["默认值", "default", "缺省值"],
    "description": ["业务含义", "说明", "描述", "备注", "comment", "注释", "含义", "定义"],
    "enum_values": ["枚举值", "取值范围", "值域", "可选值", "枚举", "代码值"],
}

# Keywords that indicate a table row is a header for database schema
SCHEMA_HEADER_INDICATORS = [
    "字段名", "字段", "列名", "column", "field",
    "类型", "数据类型", "type",
    "说明", "描述", "备注", "comment",
]

# Table name extraction patterns (from paragraphs preceding a table)
TABLE_NAME_PATTERNS = [
    r"表[名称]?[：:]\s*(.+)",
    r"表\s*\d+[\s\-.]+(.+)",
    r"(\w+)\s*表\s*[结构定义]",
    r"Table[\s:]+(\w+)",
    r"(\w+)\s+Table\s+(?:Structure|Definition)",
]

# Size thresholds for processing strategy
SIZE_SMALL = 5 * 1024 * 1024       # < 5 MB
SIZE_MEDIUM = 20 * 1024 * 1024     # 5-20 MB
SIZE_LARGE = 50 * 1024 * 1024      # 20-50 MB
# > 50 MB = huge (sampling mode)

MAX_ROWS_MEDIUM = 500
MAX_ROWS_LARGE = 100
SAMPLE_TABLES = 10


def get_size_strategy(file_size: int) -> str:
    """Determine processing strategy based on file size."""
    if file_size < SIZE_SMALL:
        return "full"
    elif file_size < SIZE_MEDIUM:
        return "selective"
    elif file_size < SIZE_LARGE:
        return "table_priority"
    else:
        return "sampling"


def detect_column_roles(header_cells: list[str]) -> dict[int, str]:
    """
    Map column indices to semantic roles based on header text matching.
    Uses longest-keyword-match-wins strategy to avoid ambiguous matches
    (e.g., "字段描述" should match description, not field_name).

    Returns: {col_index: role_name} e.g. {0: "field_name", 1: "field_type", 3: "description"}
    """
    roles = {}
    for i, cell_text in enumerate(header_cells):
        text_lower = cell_text.strip().lower()
        if not text_lower:
            continue

        best_role = None
        best_kw_len = 0
        for role, keywords in COLUMN_ROLE_KEYWORDS.items():
            for kw in keywords:
                if kw.lower() in text_lower and len(kw) > best_kw_len:
                    best_role = role
                    best_kw_len = len(kw)

        if best_role:
            roles[i] = best_role

    return roles


def is_schema_header(cells: list[str]) -> bool:
    """Check if a row of cells looks like a database schema table header."""
    combined = " ".join(c.strip().lower() for c in cells)
    match_count = sum(1 for indicator in SCHEMA_HEADER_INDICATORS if indicator.lower() in combined)
    return match_count >= 2


def extract_table_name_from_context(preceding_text: str) -> str | None:
    """Try to extract a table name from the text preceding a table."""
    for pattern in TABLE_NAME_PATTERNS:
        m = re.search(pattern, preceding_text, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return None


# ---------------------------------------------------------------------------
# Word (.docx) extraction
# ---------------------------------------------------------------------------

def extract_tables_from_docx(filepath: Path, strategy: str = "full") -> list[dict]:
    """Extract table structures from a .docx file."""
    from docx import Document

    doc = Document(str(filepath))
    tables_found = []

    # Build paragraph text index for table name context
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # Track which paragraph index corresponds to which table
    # (Word tables are interleaved with paragraphs in document body)
    para_idx = 0

    for table_idx, table in enumerate(doc.tables):
        if strategy == "sampling" and table_idx >= SAMPLE_TABLES:
            break

        rows = table.rows
        if len(rows) < 2:
            continue  # Need at least header + 1 data row

        # Get header row
        header_cells = [cell.text.strip() for cell in rows[0].cells]

        # Check if this looks like a schema table
        if strategy in ("table_priority", "sampling"):
            if not is_schema_header(header_cells):
                continue

        # Even in "full" or "selective" mode, only extract schema-like tables
        if not is_schema_header(header_cells):
            # Still capture it but mark as non-schema
            pass

        # Detect column roles
        column_roles = detect_column_roles(header_cells)

        # Try to find table name from preceding paragraphs
        table_name = None
        # Heuristic: look at a few paragraphs before, searching backward
        search_paras = paragraphs[max(0, para_idx - 3):para_idx + 1]
        for p_text in reversed(search_paras):
            table_name = extract_table_name_from_context(p_text)
            if table_name:
                break

        # If no name found from context, try the first cell of a potential title row
        if not table_name and header_cells:
            # Sometimes the first row is actually a merged title row
            if len(set(header_cells)) == 1 and header_cells[0]:
                table_name = header_cells[0]
                # Then the real header is row[1]
                if len(rows) >= 3:
                    header_cells = [cell.text.strip() for cell in rows[1].cells]
                    column_roles = detect_column_roles(header_cells)

        # Determine row limit based on strategy
        if strategy == "full":
            max_rows = len(rows)
        elif strategy == "selective":
            max_rows = min(len(rows), MAX_ROWS_MEDIUM)
        else:
            max_rows = min(len(rows), MAX_ROWS_LARGE)

        # Extract data rows
        data_rows = []
        start_row = 1
        # If we consumed an extra header row above, adjust
        for row in rows[start_row:max_rows]:
            cells = [cell.text.strip() for cell in row.cells]
            if all(c == "" for c in cells):
                continue  # Skip empty rows

            row_data = {}
            for col_idx, role in column_roles.items():
                if col_idx < len(cells):
                    row_data[role] = cells[col_idx]

            # Also keep raw cells for LLM context
            row_data["_raw"] = cells
            data_rows.append(row_data)

        is_schema = is_schema_header(header_cells)

        tables_found.append({
            "source_file": str(filepath),
            "table_index": table_idx,
            "table_name": table_name,
            "is_schema_table": is_schema,
            "header": header_cells,
            "column_roles": {str(k): v for k, v in column_roles.items()},
            "row_count": len(data_rows),
            "total_rows_in_source": len(rows) - 1,
            "truncated": len(rows) - 1 > max_rows - 1,
            "fields": data_rows,
        })

    return tables_found


# ---------------------------------------------------------------------------
# Excel (.xlsx) extraction
# ---------------------------------------------------------------------------

def extract_tables_from_xlsx(filepath: Path, strategy: str = "full") -> list[dict]:
    """Extract table structures from a .xlsx file."""
    import openpyxl

    read_only = strategy != "full"
    wb = openpyxl.load_workbook(str(filepath), read_only=read_only, data_only=True)

    tables_found = []
    max_sheets = 3 if strategy in ("table_priority", "sampling") else len(wb.sheetnames)

    for sheet_idx, sheet_name in enumerate(wb.sheetnames[:max_sheets]):
        ws = wb[sheet_name]

        # Determine row limit
        if strategy == "full":
            max_data_rows = None
        elif strategy == "selective":
            max_data_rows = MAX_ROWS_MEDIUM
        else:
            max_data_rows = MAX_ROWS_LARGE

        # Read rows
        all_rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if max_data_rows and i > max_data_rows:
                break
            cells = [str(c).strip() if c is not None else "" for c in row]
            if all(c == "" for c in cells):
                continue
            all_rows.append(cells)

        if len(all_rows) < 2:
            continue  # Need header + data

        # Find the header row (first row with schema-like indicators, or row 0)
        header_row_idx = 0
        for idx, row in enumerate(all_rows[:5]):  # Check first 5 rows
            if is_schema_header(row):
                header_row_idx = idx
                break

        header_cells = all_rows[header_row_idx]
        column_roles = detect_column_roles(header_cells)

        # Extract data rows
        data_rows = []
        for row in all_rows[header_row_idx + 1:]:
            row_data = {}
            for col_idx, role in column_roles.items():
                if col_idx < len(row):
                    row_data[role] = row[col_idx]
            row_data["_raw"] = row
            data_rows.append(row_data)

        is_schema = is_schema_header(header_cells) or bool(column_roles)

        tables_found.append({
            "source_file": str(filepath),
            "sheet_name": sheet_name,
            "sheet_index": sheet_idx,
            "table_name": sheet_name,  # Use sheet name as default table name
            "is_schema_table": is_schema,
            "header": header_cells,
            "column_roles": {str(k): v for k, v in column_roles.items()},
            "row_count": len(data_rows),
            "total_rows_in_source": len(all_rows) - header_row_idx - 1,
            "truncated": max_data_rows is not None and len(all_rows) > max_data_rows,
            "fields": data_rows,
        })

    wb.close()
    return tables_found


# ---------------------------------------------------------------------------
# Excel (.xls) extraction
# ---------------------------------------------------------------------------

def extract_tables_from_xls(filepath: Path, strategy: str = "full") -> list[dict]:
    """Extract table structures from a .xls file (legacy Excel format)."""
    import xlrd

    wb = xlrd.open_workbook(str(filepath))
    tables_found = []
    max_sheets = 3 if strategy in ("table_priority", "sampling") else wb.nsheets

    for sheet_idx in range(min(wb.nsheets, max_sheets)):
        ws = wb.sheet_by_index(sheet_idx)

        if ws.nrows < 2:
            continue

        # Determine row limit
        if strategy == "full":
            max_data_rows = ws.nrows
        elif strategy == "selective":
            max_data_rows = min(ws.nrows, MAX_ROWS_MEDIUM)
        else:
            max_data_rows = min(ws.nrows, MAX_ROWS_LARGE)

        # Read rows
        all_rows = []
        for i in range(min(ws.nrows, max_data_rows)):
            cells = [str(ws.cell_value(i, j)).strip() for j in range(ws.ncols)]
            all_rows.append(cells)

        # Find header row
        header_row_idx = 0
        for idx, row in enumerate(all_rows[:5]):
            if is_schema_header(row):
                header_row_idx = idx
                break

        header_cells = all_rows[header_row_idx]
        column_roles = detect_column_roles(header_cells)

        data_rows = []
        for row in all_rows[header_row_idx + 1:]:
            if all(c == "" for c in row):
                continue
            row_data = {}
            for col_idx, role in column_roles.items():
                if col_idx < len(row):
                    row_data[role] = row[col_idx]
            row_data["_raw"] = row
            data_rows.append(row_data)

        is_schema = is_schema_header(header_cells) or bool(column_roles)

        tables_found.append({
            "source_file": str(filepath),
            "sheet_name": ws.name,
            "sheet_index": sheet_idx,
            "table_name": ws.name,
            "is_schema_table": is_schema,
            "header": header_cells,
            "column_roles": {str(k): v for k, v in column_roles.items()},
            "row_count": len(data_rows),
            "total_rows_in_source": ws.nrows - header_row_idx - 1,
            "truncated": ws.nrows > max_data_rows,
            "fields": data_rows,
        })

    return tables_found


# ---------------------------------------------------------------------------
# Unified extraction interface
# ---------------------------------------------------------------------------

def extract_tables(filepath: Path) -> list[dict]:
    """
    Extract tables from a supported file format.
    Automatically selects strategy based on file size.
    """
    if not filepath.exists():
        print(f"Error: File not found: {filepath}", file=sys.stderr)
        return []

    file_size = filepath.stat().st_size
    strategy = get_size_strategy(file_size)
    ext = filepath.suffix.lower()

    if strategy != "full":
        size_mb = round(file_size / (1024 * 1024), 1)
        print(f"  Large file ({size_mb} MB), using '{strategy}' strategy")

    if ext == ".docx":
        return extract_tables_from_docx(filepath, strategy)
    elif ext == ".xlsx":
        return extract_tables_from_xlsx(filepath, strategy)
    elif ext == ".xls":
        return extract_tables_from_xls(filepath, strategy)
    else:
        print(f"  Unsupported format: {ext}", file=sys.stderr)
        return []


def extract_from_directory(root: Path, scan_result_path: Path | None = None) -> dict:
    """
    Extract tables from all supported files in a directory.

    If scan_result_path is provided, processes files in priority order (P1 first).
    """
    supported_exts = {".docx", ".xlsx", ".xls"}

    if scan_result_path and scan_result_path.exists():
        with open(scan_result_path, "r", encoding="utf-8") as f:
            scan_data = json.load(f)
        file_list = [
            entry for entry in scan_data.get("files", [])
            if entry["extension"] in supported_exts
        ]
        # Sort by priority
        file_list.sort(key=lambda x: x["priority"])
        files_to_process = [Path(entry["path"]) for entry in file_list]
    else:
        files_to_process = sorted(
            f for f in root.rglob("*")
            if f.is_file() and f.suffix.lower() in supported_exts
        )

    all_tables = []
    file_results = []
    total = len(files_to_process)

    print(f"Extracting tables from {total} files...")

    for i, filepath in enumerate(files_to_process, 1):
        print(f"[{i}/{total}] {filepath.name}")
        try:
            tables = extract_tables(filepath)
            schema_tables = [t for t in tables if t.get("is_schema_table")]
            all_tables.extend(tables)
            file_results.append({
                "file": str(filepath),
                "tables_found": len(tables),
                "schema_tables": len(schema_tables),
                "status": "success",
            })
            print(f"  Found {len(tables)} tables ({len(schema_tables)} schema-like)")
        except Exception as e:
            print(f"  ERROR: {e}", file=sys.stderr)
            file_results.append({
                "file": str(filepath),
                "tables_found": 0,
                "schema_tables": 0,
                "status": f"error: {e}",
            })

    total_tables = len(all_tables)
    schema_count = sum(1 for t in all_tables if t.get("is_schema_table"))
    total_fields = sum(t.get("row_count", 0) for t in all_tables)

    summary = {
        "total_files_processed": total,
        "total_tables_found": total_tables,
        "total_schema_tables": schema_count,
        "total_fields_extracted": total_fields,
        "files": file_results,
    }

    return {"tables": all_tables, "summary": summary}


def main():
    parser = argparse.ArgumentParser(description="Extract tables from Word/Excel files")
    parser.add_argument("input", help="File or directory to process")
    parser.add_argument("--output", "-o", default="extracted_tables.json", help="Output JSON file")
    parser.add_argument("--scan-result", default=None, help="Path to scan_result.json for priority ordering")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: '{input_path}' does not exist", file=sys.stderr)
        sys.exit(1)

    if input_path.is_file():
        tables = extract_tables(input_path)
        result = {
            "tables": tables,
            "summary": {
                "total_files_processed": 1,
                "total_tables_found": len(tables),
                "total_schema_tables": sum(1 for t in tables if t.get("is_schema_table")),
                "total_fields_extracted": sum(t.get("row_count", 0) for t in tables),
            },
        }
    else:
        scan_result = Path(args.scan_result) if args.scan_result else None
        result = extract_from_directory(input_path, scan_result)

    # Save output
    output_path = Path(args.output)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    s = result["summary"]
    print(f"\nResults saved to: {output_path}")
    print(f"Tables: {s['total_tables_found']} ({s['total_schema_tables']} schema-like)")
    print(f"Fields: {s['total_fields_extracted']}")


if __name__ == "__main__":
    main()

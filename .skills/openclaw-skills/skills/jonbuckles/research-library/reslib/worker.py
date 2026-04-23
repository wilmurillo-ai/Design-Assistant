"""
Extraction Worker Pool for Research Library.

Provides a robust multi-threaded worker pool for asynchronous document
extraction with comprehensive error handling, monitoring, and graceful shutdown.

Features:
- Thread pool with configurable worker count
- Exponential backoff retry with configurable limits
- Graceful shutdown with job preservation
- Real-time heartbeat monitoring
- Comprehensive metrics collection
- Database lock resilience

Example:
    from reslib import ResearchDatabase, ExtractionWorker
    
    db = ResearchDatabase("research.db")
    worker = ExtractionWorker(db, num_workers=2)
    
    # Start processing
    worker.start()
    
    # Add jobs
    job_id = worker.enqueue(attachment_id=123)
    
    # Check status
    status = worker.get_status()
    print(f"Queue depth: {status['queue_depth']}")
    
    # Graceful shutdown
    worker.stop()
"""

import logging
import os
import signal
import threading
import time
import traceback
import uuid
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor, Future
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, Any, List, Callable, Set

from .database import ResearchDatabase
from .queue import QueueManager, Job, JobStatus

logger = logging.getLogger(__name__)


class WorkerStatus(str, Enum):
    """Worker status enumeration."""
    STARTING = "starting"
    IDLE = "idle"
    PROCESSING = "processing"
    STOPPING = "stopping"
    STOPPED = "stopped"


class ExtractorError(Exception):
    """Base exception for extraction errors."""
    pass


class ExtractorTimeoutError(ExtractorError):
    """Extraction timed out."""
    pass


class ExtractorCorruptionError(ExtractorError):
    """File is corrupted or unreadable."""
    pass


class ExtractorMissingError(ExtractorError):
    """Required extractor not available."""
    pass


class ExtractorUnsupportedError(ExtractorError):
    """File type not supported."""
    pass


@dataclass
class ExtractionResult:
    """Result of a document extraction."""
    text: str
    confidence: float
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_time_ms: int = 0


@dataclass
class WorkerInfo:
    """Information about a worker thread."""
    worker_id: str
    status: WorkerStatus
    current_job_id: Optional[int] = None
    jobs_completed: int = 0
    jobs_failed: int = 0
    last_heartbeat: Optional[datetime] = None
    started_at: Optional[datetime] = None


class DocumentExtractor(ABC):
    """
    Abstract base class for document text extractors.
    
    Implement this for different document types (PDF, DOCX, etc.).
    """
    
    @abstractmethod
    def can_extract(self, path: Path, mime_type: Optional[str] = None) -> bool:
        """
        Check if this extractor can handle the given file.
        
        Args:
            path: Path to the file
            mime_type: Optional MIME type hint
        
        Returns:
            True if extractor can handle this file
        """
        pass
    
    @abstractmethod
    def extract(self, path: Path) -> ExtractionResult:
        """
        Extract text from a document.
        
        Args:
            path: Path to the document
        
        Returns:
            ExtractionResult with extracted text and confidence
        
        Raises:
            ExtractorCorruptionError: If file is corrupted
            ExtractorTimeoutError: If extraction times out
            ExtractorUnsupportedError: If file type not supported
        """
        pass


class PlainTextExtractor(DocumentExtractor):
    """Simple text file extractor."""
    
    TEXT_EXTENSIONS = {'.txt', '.md', '.rst', '.csv', '.json', '.xml', '.html'}
    TEXT_MIMES = {'text/plain', 'text/markdown', 'text/csv', 'application/json'}
    
    def can_extract(self, path: Path, mime_type: Optional[str] = None) -> bool:
        if mime_type and mime_type in self.TEXT_MIMES:
            return True
        return path.suffix.lower() in self.TEXT_EXTENSIONS
    
    def extract(self, path: Path) -> ExtractionResult:
        start = time.monotonic()
        
        try:
            # Try UTF-8 first, then fall back to latin-1
            try:
                text = path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                text = path.read_text(encoding='latin-1')
            
            elapsed_ms = int((time.monotonic() - start) * 1000)
            
            return ExtractionResult(
                text=text,
                confidence=1.0,
                metadata={"encoding": "utf-8"},
                extraction_time_ms=elapsed_ms
            )
        except Exception as e:
            raise ExtractorCorruptionError(f"Failed to read text file: {e}")


class PDFExtractor(DocumentExtractor):
    """PDF document extractor using pdfplumber or PyMuPDF."""
    
    def can_extract(self, path: Path, mime_type: Optional[str] = None) -> bool:
        if mime_type == 'application/pdf':
            return True
        return path.suffix.lower() == '.pdf'
    
    def extract(self, path: Path) -> ExtractionResult:
        start = time.monotonic()
        text_parts = []
        confidence = 1.0
        
        # Try pdfplumber first (better table handling)
        try:
            import pdfplumber
            
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if text_parts:
                text = "\n\n".join(text_parts)
                elapsed_ms = int((time.monotonic() - start) * 1000)
                return ExtractionResult(
                    text=text,
                    confidence=confidence,
                    metadata={"extractor": "pdfplumber", "pages": len(text_parts)},
                    extraction_time_ms=elapsed_ms
                )
        except ImportError:
            pass  # Try PyMuPDF
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyMuPDF")
        
        # Fall back to PyMuPDF
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(str(path))
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            
            if text_parts:
                text = "\n\n".join(text_parts)
                elapsed_ms = int((time.monotonic() - start) * 1000)
                return ExtractionResult(
                    text=text,
                    confidence=0.9,  # Slightly lower than pdfplumber
                    metadata={"extractor": "pymupdf", "pages": len(text_parts)},
                    extraction_time_ms=elapsed_ms
                )
        except ImportError:
            raise ExtractorMissingError("No PDF extractor available (install pdfplumber or pymupdf)")
        except Exception as e:
            raise ExtractorCorruptionError(f"Failed to extract PDF: {e}")
        
        raise ExtractorCorruptionError("No text could be extracted from PDF")


class DocxExtractor(DocumentExtractor):
    """Microsoft Word document extractor."""
    
    def can_extract(self, path: Path, mime_type: Optional[str] = None) -> bool:
        if mime_type in ('application/vnd.openxmlformats-officedocument.wordprocessingml.document',):
            return True
        return path.suffix.lower() in ('.docx', '.doc')
    
    def extract(self, path: Path) -> ExtractionResult:
        start = time.monotonic()
        
        try:
            from docx import Document
            
            doc = Document(str(path))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            elapsed_ms = int((time.monotonic() - start) * 1000)
            
            return ExtractionResult(
                text=text,
                confidence=0.95,
                metadata={"extractor": "python-docx", "paragraphs": len(doc.paragraphs)},
                extraction_time_ms=elapsed_ms
            )
        except ImportError:
            raise ExtractorMissingError("No DOCX extractor available (install python-docx)")
        except Exception as e:
            raise ExtractorCorruptionError(f"Failed to extract DOCX: {e}")


class ExtractorRegistry:
    """
    Registry of available document extractors.
    
    Manages multiple extractors and selects the appropriate one
    based on file type and availability.
    """
    
    def __init__(self):
        self._extractors: List[DocumentExtractor] = []
        self._register_defaults()
    
    def _register_defaults(self):
        """Register default extractors."""
        self._extractors = [
            PlainTextExtractor(),
            PDFExtractor(),
            DocxExtractor(),
        ]
    
    def register(self, extractor: DocumentExtractor):
        """Register a custom extractor."""
        self._extractors.insert(0, extractor)  # Custom takes priority
    
    def get_extractor(
        self,
        path: Path,
        mime_type: Optional[str] = None
    ) -> Optional[DocumentExtractor]:
        """Get appropriate extractor for a file."""
        for extractor in self._extractors:
            if extractor.can_extract(path, mime_type):
                return extractor
        return None
    
    def extract(self, path: Path, mime_type: Optional[str] = None) -> ExtractionResult:
        """
        Extract text from a document using the appropriate extractor.
        
        Args:
            path: Path to document
            mime_type: Optional MIME type hint
        
        Returns:
            ExtractionResult
        
        Raises:
            ExtractorUnsupportedError: If no extractor available
            ExtractorError: For extraction failures
        """
        extractor = self.get_extractor(path, mime_type)
        if not extractor:
            raise ExtractorUnsupportedError(
                f"No extractor available for {path.suffix}"
            )
        return extractor.extract(path)


class ExtractionWorker:
    """
    Multi-threaded extraction worker pool.
    
    Manages a pool of worker threads that process extraction jobs from
    the queue. Provides comprehensive error handling, retry logic,
    monitoring, and graceful shutdown.
    
    Thread Safety:
        All public methods are thread-safe. Internal state is protected
        by locks and atomic operations.
    
    Attributes:
        db: ResearchDatabase instance
        num_workers: Number of worker threads
        max_retries: Maximum retry attempts per job
        extraction_timeout: Timeout for individual extractions (seconds)
        heartbeat_interval: Interval between heartbeat updates (seconds)
        poll_interval: Interval between queue polls when idle (seconds)
    """
    
    DEFAULT_NUM_WORKERS = 2
    DEFAULT_MAX_RETRIES = 3
    DEFAULT_EXTRACTION_TIMEOUT = 300  # 5 minutes
    DEFAULT_HEARTBEAT_INTERVAL = 10  # seconds
    DEFAULT_POLL_INTERVAL = 1  # seconds
    
    def __init__(
        self,
        db: ResearchDatabase,
        num_workers: int = DEFAULT_NUM_WORKERS,
        max_retries: int = DEFAULT_MAX_RETRIES,
        extraction_timeout: int = DEFAULT_EXTRACTION_TIMEOUT,
        heartbeat_interval: int = DEFAULT_HEARTBEAT_INTERVAL,
        poll_interval: float = DEFAULT_POLL_INTERVAL,
    ):
        """
        Initialize extraction worker pool.
        
        Args:
            db: ResearchDatabase instance
            num_workers: Number of worker threads (default: 2)
            max_retries: Maximum retry attempts per job (default: 3)
            extraction_timeout: Extraction timeout in seconds (default: 300)
            heartbeat_interval: Heartbeat update interval (default: 10s)
            poll_interval: Queue poll interval when idle (default: 1s)
        """
        self.db = db
        self.num_workers = max(1, num_workers)
        self.max_retries = max(0, max_retries)
        self.extraction_timeout = max(10, extraction_timeout)
        self.heartbeat_interval = max(1, heartbeat_interval)
        self.poll_interval = max(0.1, poll_interval)
        
        # Initialize queue manager
        self.queue = QueueManager(db)
        
        # Initialize extractor registry
        self.extractors = ExtractorRegistry()
        
        # Worker state
        self._workers: Dict[str, WorkerInfo] = {}
        self._worker_threads: List[threading.Thread] = []
        self._heartbeat_thread: Optional[threading.Thread] = None
        
        # Synchronization
        self._lock = threading.RLock()
        self._shutdown_event = threading.Event()
        self._pause_event = threading.Event()
        self._pause_event.set()  # Not paused initially
        
        # Jobs currently being processed
        self._active_jobs: Dict[str, int] = {}  # worker_id -> job_id
        
        # Callbacks
        self._callbacks: Dict[str, List[Callable]] = {
            "worker_started": [],
            "worker_stopped": [],
            "job_started": [],
            "job_completed": [],
            "job_failed": [],
            "extraction_error": [],
        }
        
        # Statistics
        self._stats = {
            "jobs_processed": 0,
            "jobs_succeeded": 0,
            "jobs_failed": 0,
            "total_extraction_time_ms": 0,
            "started_at": None,
        }
        
        logger.info(
            f"ExtractionWorker initialized: {num_workers} workers, "
            f"{max_retries} max retries, {extraction_timeout}s timeout"
        )
    
    def on(self, event: str, callback: Callable):
        """
        Register callback for worker events.
        
        Events:
            - worker_started: (worker_id: str)
            - worker_stopped: (worker_id: str)
            - job_started: (job: Job, worker_id: str)
            - job_completed: (job: Job, result: ExtractionResult)
            - job_failed: (job: Job, error: str)
            - extraction_error: (job: Job, error: Exception)
        """
        if event in self._callbacks:
            self._callbacks[event].append(callback)
    
    def _emit(self, event: str, *args):
        """Emit event to registered callbacks."""
        for callback in self._callbacks.get(event, []):
            try:
                callback(*args)
            except Exception as e:
                logger.error(f"Callback error for {event}: {e}")
    
    def start(self) -> None:
        """
        Start the worker pool.
        
        Launches worker threads and heartbeat thread. Safe to call
        multiple times (no-op if already running).
        """
        with self._lock:
            if self._worker_threads:
                logger.warning("Worker pool already running")
                return
            
            logger.info(f"Starting worker pool with {self.num_workers} workers")
            
            self._shutdown_event.clear()
            self._stats["started_at"] = datetime.now()
            
            # Clean up any stale jobs from previous run
            stale_count = self.queue.cleanup_stale_jobs(stale_threshold_minutes=30)
            if stale_count:
                logger.info(f"Cleaned up {stale_count} stale jobs from previous run")
            
            # Start worker threads
            for i in range(self.num_workers):
                worker_id = f"worker-{uuid.uuid4().hex[:8]}"
                
                worker_info = WorkerInfo(
                    worker_id=worker_id,
                    status=WorkerStatus.STARTING,
                    started_at=datetime.now(),
                )
                self._workers[worker_id] = worker_info
                
                thread = threading.Thread(
                    target=self._worker_loop,
                    args=(worker_id,),
                    name=f"ExtractionWorker-{worker_id}",
                    daemon=True,
                )
                self._worker_threads.append(thread)
                thread.start()
                
                logger.info(f"Started worker {worker_id}")
                self._emit("worker_started", worker_id)
            
            # Start heartbeat thread
            self._heartbeat_thread = threading.Thread(
                target=self._heartbeat_loop,
                name="ExtractionWorker-Heartbeat",
                daemon=True,
            )
            self._heartbeat_thread.start()
            
            logger.info("Worker pool started successfully")
    
    def stop(self, timeout: float = 60.0) -> bool:
        """
        Stop the worker pool gracefully.
        
        Waits for in-progress jobs to complete, then shuts down workers.
        Jobs that are still in 'processing' state are released back to
        pending for future processing.
        
        Args:
            timeout: Maximum time to wait for workers to finish (seconds)
        
        Returns:
            True if all workers stopped cleanly, False if timeout exceeded
        """
        with self._lock:
            if not self._worker_threads:
                logger.warning("Worker pool not running")
                return True
            
            logger.info("Initiating graceful shutdown...")
            
            # Signal shutdown
            self._shutdown_event.set()
            
            # Update worker statuses
            for worker_info in self._workers.values():
                worker_info.status = WorkerStatus.STOPPING
        
        # Wait for workers to finish (outside lock)
        deadline = time.monotonic() + timeout
        all_stopped = True
        
        for thread in self._worker_threads:
            remaining = deadline - time.monotonic()
            if remaining > 0:
                thread.join(timeout=remaining)
            if thread.is_alive():
                logger.warning(f"Worker thread {thread.name} did not stop in time")
                all_stopped = False
        
        # Wait for heartbeat thread
        if self._heartbeat_thread:
            remaining = deadline - time.monotonic()
            if remaining > 0:
                self._heartbeat_thread.join(timeout=remaining)
        
        with self._lock:
            # Release any jobs that were still in progress
            released_count = 0
            for worker_id, job_id in list(self._active_jobs.items()):
                if self.queue.release(job_id):
                    released_count += 1
                    logger.info(f"Released job {job_id} from {worker_id}")
            
            if released_count:
                logger.info(f"Released {released_count} in-progress jobs back to queue")
            
            # Clear state
            self._worker_threads.clear()
            self._heartbeat_thread = None
            
            # Update worker statuses in database
            for worker_id in self._workers:
                self._update_worker_heartbeat(worker_id, WorkerStatus.STOPPED)
            
            for worker_id in list(self._workers.keys()):
                self._emit("worker_stopped", worker_id)
            
            self._workers.clear()
            self._active_jobs.clear()
        
        if all_stopped:
            logger.info("Worker pool stopped cleanly")
        else:
            logger.warning("Worker pool stopped with timeout")
        
        return all_stopped
    
    def pause(self) -> None:
        """Pause job processing (workers stay alive but idle)."""
        self._pause_event.clear()
        logger.info("Worker pool paused")
    
    def resume(self) -> None:
        """Resume job processing after pause."""
        self._pause_event.set()
        logger.info("Worker pool resumed")
    
    def is_running(self) -> bool:
        """Check if worker pool is running."""
        return bool(self._worker_threads) and not self._shutdown_event.is_set()
    
    def is_paused(self) -> bool:
        """Check if worker pool is paused."""
        return not self._pause_event.is_set()
    
    def enqueue(self, attachment_id: int, priority: int = 0) -> Optional[int]:
        """
        Add an extraction job to the queue.
        
        Args:
            attachment_id: ID of attachment to process
            priority: Job priority (higher = processed first)
        
        Returns:
            Job ID, or None if job was deduplicated
        """
        return self.queue.enqueue(attachment_id, priority)
    
    def enqueue_batch(self, attachment_ids: List[int], priority: int = 0) -> List[int]:
        """
        Add multiple extraction jobs to the queue.
        
        Args:
            attachment_ids: List of attachment IDs
            priority: Priority for all jobs
        
        Returns:
            List of created job IDs
        """
        return self.queue.enqueue_batch(attachment_ids, priority)
    
    def get_status(self) -> Dict[str, Any]:
        """
        Get comprehensive worker pool status.
        
        Returns:
            Dict with:
                - running: bool
                - paused: bool
                - num_workers: int
                - active_workers: int
                - queue_depth: int (pending + retry_pending)
                - processing: int
                - workers: List[WorkerInfo dicts]
                - stats: job counts and timing
                - queue_stats: detailed queue statistics
        """
        with self._lock:
            queue_stats = self.queue.get_queue_stats()
            
            workers_info = []
            active_count = 0
            for worker_id, info in self._workers.items():
                workers_info.append({
                    "worker_id": info.worker_id,
                    "status": info.status.value,
                    "current_job_id": info.current_job_id,
                    "jobs_completed": info.jobs_completed,
                    "jobs_failed": info.jobs_failed,
                    "last_heartbeat": (
                        info.last_heartbeat.isoformat()
                        if info.last_heartbeat else None
                    ),
                })
                if info.status == WorkerStatus.PROCESSING:
                    active_count += 1
            
            # Calculate uptime
            uptime_seconds = None
            if self._stats["started_at"]:
                uptime_seconds = (datetime.now() - self._stats["started_at"]).total_seconds()
            
            # Calculate throughput
            throughput = 0.0
            if uptime_seconds and uptime_seconds > 0:
                throughput = self._stats["jobs_processed"] / uptime_seconds
            
            # Calculate average extraction time
            avg_extraction_ms = 0
            if self._stats["jobs_processed"] > 0:
                avg_extraction_ms = (
                    self._stats["total_extraction_time_ms"] /
                    self._stats["jobs_processed"]
                )
            
            return {
                "running": self.is_running(),
                "paused": self.is_paused(),
                "num_workers": self.num_workers,
                "active_workers": active_count,
                "queue_depth": (
                    queue_stats.get("pending", 0) +
                    queue_stats.get("retry_pending", 0)
                ),
                "processing": queue_stats.get("processing", 0),
                "workers": workers_info,
                "stats": {
                    "jobs_processed": self._stats["jobs_processed"],
                    "jobs_succeeded": self._stats["jobs_succeeded"],
                    "jobs_failed": self._stats["jobs_failed"],
                    "uptime_seconds": uptime_seconds,
                    "throughput_per_second": round(throughput, 3),
                    "avg_extraction_time_ms": round(avg_extraction_ms),
                },
                "queue_stats": queue_stats,
            }
    
    def get_recent_jobs(self, limit: int = 20) -> List[Dict[str, Any]]:
        """Get recent job history with results."""
        rows = self.db.fetchall(
            """
            SELECT q.*, m.extraction_time_ms, m.confidence, m.error_type
            FROM extraction_queue q
            LEFT JOIN extraction_metrics m ON q.id = m.job_id
            ORDER BY q.created_at DESC
            LIMIT ?
            """,
            (limit,)
        )
        return [dict(row) for row in rows]
    
    def retry_failed_jobs(self, job_ids: Optional[List[int]] = None) -> int:
        """
        Retry failed jobs.
        
        Args:
            job_ids: Specific job IDs to retry, or None for all failed
        
        Returns:
            Number of jobs reset for retry
        """
        if job_ids:
            count = 0
            for job_id in job_ids:
                count += self.queue.retry_failed(job_id)
            return count
        return self.queue.retry_failed()
    
    def _worker_loop(self, worker_id: str):
        """
        Main worker loop.
        
        Continuously polls for jobs, processes them, and handles errors.
        Exits when shutdown is signaled.
        """
        logger.debug(f"Worker {worker_id} loop started")
        
        with self._lock:
            if worker_id in self._workers:
                self._workers[worker_id].status = WorkerStatus.IDLE
        
        consecutive_empty = 0
        
        while not self._shutdown_event.is_set():
            # Check for pause
            if not self._pause_event.wait(timeout=0.1):
                continue
            
            # Claim a job
            job = self.queue.claim(worker_id)
            
            if not job:
                consecutive_empty += 1
                # Exponential backoff on empty queue, up to 5 seconds
                wait_time = min(self.poll_interval * (1.5 ** min(consecutive_empty, 5)), 5.0)
                self._shutdown_event.wait(timeout=wait_time)
                continue
            
            consecutive_empty = 0
            
            # Track active job
            with self._lock:
                self._active_jobs[worker_id] = job.id
                if worker_id in self._workers:
                    self._workers[worker_id].status = WorkerStatus.PROCESSING
                    self._workers[worker_id].current_job_id = job.id
            
            logger.info(f"Worker {worker_id} processing job {job.id} (attachment {job.attachment_id})")
            self._emit("job_started", job, worker_id)
            
            # Process the job
            try:
                result = self._process_job(job)
                
                # Update extraction results
                self.db.update_attachment_extraction(
                    job.attachment_id,
                    result.text,
                    result.confidence
                )
                
                # Mark complete
                self.queue.complete(
                    job.id,
                    extraction_time_ms=result.extraction_time_ms,
                    confidence=result.confidence
                )
                
                # Update stats
                with self._lock:
                    self._stats["jobs_processed"] += 1
                    self._stats["jobs_succeeded"] += 1
                    self._stats["total_extraction_time_ms"] += result.extraction_time_ms
                    if worker_id in self._workers:
                        self._workers[worker_id].jobs_completed += 1
                
                logger.info(
                    f"Job {job.id} completed: {len(result.text)} chars, "
                    f"confidence {result.confidence:.2f}, {result.extraction_time_ms}ms"
                )
                self._emit("job_completed", job, result)
                
            except ExtractorTimeoutError as e:
                self._handle_job_error(worker_id, job, str(e), "timeout")
                
            except ExtractorCorruptionError as e:
                self._handle_job_error(worker_id, job, str(e), "corruption")
                
            except ExtractorMissingError as e:
                self._handle_job_error(worker_id, job, str(e), "extractor_missing")
                
            except ExtractorUnsupportedError as e:
                self._handle_job_error(worker_id, job, str(e), "unsupported")
                
            except Exception as e:
                logger.exception(f"Unexpected error processing job {job.id}")
                self._handle_job_error(worker_id, job, str(e), "unknown")
                self._emit("extraction_error", job, e)
            
            finally:
                # Clear active job
                with self._lock:
                    self._active_jobs.pop(worker_id, None)
                    if worker_id in self._workers:
                        self._workers[worker_id].status = WorkerStatus.IDLE
                        self._workers[worker_id].current_job_id = None
        
        logger.debug(f"Worker {worker_id} loop exiting")
        
        with self._lock:
            if worker_id in self._workers:
                self._workers[worker_id].status = WorkerStatus.STOPPED
    
    def _process_job(self, job: Job) -> ExtractionResult:
        """
        Process a single extraction job.
        
        Args:
            job: Job to process
        
        Returns:
            ExtractionResult
        
        Raises:
            ExtractorError subclasses on failure
        """
        # Get attachment info
        attachment = self.db.get_attachment(job.attachment_id)
        if not attachment:
            raise ExtractorCorruptionError(f"Attachment {job.attachment_id} not found")
        
        path = Path(attachment["path"])
        
        # Check file exists
        if not path.exists():
            raise ExtractorCorruptionError(f"File not found: {path}")
        
        # Check file is readable
        if not os.access(path, os.R_OK):
            raise ExtractorCorruptionError(f"File not readable: {path}")
        
        # Check file size
        file_size = path.stat().st_size
        if file_size == 0:
            raise ExtractorCorruptionError(f"File is empty: {path}")
        
        # Get MIME type
        mime_type = attachment.get("mime_type")
        
        # Extract with timeout
        result = self._extract_with_timeout(
            path,
            mime_type,
            timeout=self.extraction_timeout
        )
        
        return result
    
    def _extract_with_timeout(
        self,
        path: Path,
        mime_type: Optional[str],
        timeout: int
    ) -> ExtractionResult:
        """
        Run extraction with timeout.
        
        Uses a separate thread with join timeout to enforce extraction limits.
        """
        result_holder = {"result": None, "error": None}
        
        def do_extract():
            try:
                result_holder["result"] = self.extractors.extract(path, mime_type)
            except Exception as e:
                result_holder["error"] = e
        
        thread = threading.Thread(target=do_extract)
        thread.start()
        thread.join(timeout=timeout)
        
        if thread.is_alive():
            # Extraction timed out
            # Note: We can't actually kill the thread in Python, but we can
            # proceed and let it finish in the background
            raise ExtractorTimeoutError(
                f"Extraction timed out after {timeout}s for {path}"
            )
        
        if result_holder["error"]:
            raise result_holder["error"]
        
        return result_holder["result"]
    
    def _handle_job_error(
        self,
        worker_id: str,
        job: Job,
        error_message: str,
        error_type: str
    ):
        """Handle job processing error."""
        self.queue.failed(
            job.id,
            error_message,
            error_type=error_type,
            max_retries=self.max_retries
        )
        
        with self._lock:
            self._stats["jobs_processed"] += 1
            self._stats["jobs_failed"] += 1
            if worker_id in self._workers:
                self._workers[worker_id].jobs_failed += 1
        
        self._emit("job_failed", job, error_message)
    
    def _heartbeat_loop(self):
        """
        Heartbeat loop for worker health monitoring.
        
        Updates worker_heartbeat table periodically so external
        monitors can detect dead workers.
        """
        logger.debug("Heartbeat loop started")
        
        while not self._shutdown_event.wait(timeout=self.heartbeat_interval):
            with self._lock:
                for worker_id, info in self._workers.items():
                    self._update_worker_heartbeat(worker_id, info.status, info.current_job_id)
                    info.last_heartbeat = datetime.now()
        
        logger.debug("Heartbeat loop exiting")
    
    def _update_worker_heartbeat(
        self,
        worker_id: str,
        status: WorkerStatus,
        current_job_id: Optional[int] = None
    ):
        """Update worker heartbeat in database."""
        try:
            # Check if worker exists
            existing = self.db.fetchone(
                "SELECT worker_id FROM worker_heartbeat WHERE worker_id = ?",
                (worker_id,)
            )
            
            if existing:
                # Update existing
                self.db.execute(
                    """
                    UPDATE worker_heartbeat
                    SET last_heartbeat = datetime('now'),
                        status = ?,
                        current_job_id = ?,
                        jobs_completed = ?,
                        jobs_failed = ?
                    WHERE worker_id = ?
                    """,
                    (
                        status.value,
                        current_job_id,
                        self._workers.get(worker_id, WorkerInfo(worker_id, status)).jobs_completed,
                        self._workers.get(worker_id, WorkerInfo(worker_id, status)).jobs_failed,
                        worker_id,
                    )
                )
            else:
                # Insert new
                self.db.execute(
                    """
                    INSERT INTO worker_heartbeat 
                    (worker_id, status, current_job_id, jobs_completed, jobs_failed)
                    VALUES (?, ?, ?, ?, ?)
                    """,
                    (
                        worker_id,
                        status.value,
                        current_job_id,
                        self._workers.get(worker_id, WorkerInfo(worker_id, status)).jobs_completed,
                        self._workers.get(worker_id, WorkerInfo(worker_id, status)).jobs_failed,
                    )
                )
        except Exception as e:
            logger.warning(f"Failed to update heartbeat for {worker_id}: {e}")


class WorkerManager:
    """
    High-level manager for the extraction worker system.
    
    Provides a simplified interface for managing workers across
    multiple databases or with configuration reloading.
    """
    
    def __init__(self):
        self._workers: Dict[str, ExtractionWorker] = {}
        self._lock = threading.Lock()
    
    def create_worker(
        self,
        name: str,
        db: ResearchDatabase,
        **kwargs
    ) -> ExtractionWorker:
        """
        Create and register a named worker pool.
        
        Args:
            name: Unique name for this worker pool
            db: ResearchDatabase instance
            **kwargs: Arguments passed to ExtractionWorker
        
        Returns:
            ExtractionWorker instance
        """
        with self._lock:
            if name in self._workers:
                raise ValueError(f"Worker pool '{name}' already exists")
            
            worker = ExtractionWorker(db, **kwargs)
            self._workers[name] = worker
            return worker
    
    def get_worker(self, name: str) -> Optional[ExtractionWorker]:
        """Get worker pool by name."""
        return self._workers.get(name)
    
    def start_all(self):
        """Start all registered worker pools."""
        for name, worker in self._workers.items():
            logger.info(f"Starting worker pool: {name}")
            worker.start()
    
    def stop_all(self, timeout: float = 60.0):
        """Stop all registered worker pools."""
        for name, worker in self._workers.items():
            logger.info(f"Stopping worker pool: {name}")
            worker.stop(timeout=timeout)
    
    def get_all_status(self) -> Dict[str, Dict[str, Any]]:
        """Get status of all worker pools."""
        return {
            name: worker.get_status()
            for name, worker in self._workers.items()
        }


# Convenience function for simple use cases
def create_worker(
    db_path: str | Path,
    num_workers: int = 2,
    max_retries: int = 3,
    **kwargs
) -> ExtractionWorker:
    """
    Create an extraction worker with a new database connection.
    
    Convenience function for simple setups.
    
    Args:
        db_path: Path to SQLite database
        num_workers: Number of worker threads
        max_retries: Maximum retry attempts
        **kwargs: Additional arguments for ExtractionWorker
    
    Returns:
        Configured ExtractionWorker
    """
    db = ResearchDatabase(db_path)
    return ExtractionWorker(
        db,
        num_workers=num_workers,
        max_retries=max_retries,
        **kwargs
    )


# Signal handlers for graceful shutdown
_active_workers: List[ExtractionWorker] = []


def _shutdown_handler(signum, frame):
    """Handle shutdown signals."""
    logger.info(f"Received signal {signum}, initiating shutdown...")
    for worker in _active_workers:
        worker.stop()


def register_signal_handlers(worker: ExtractionWorker):
    """
    Register signal handlers for graceful shutdown.
    
    Installs handlers for SIGTERM and SIGINT that will
    gracefully stop the worker pool.
    
    Args:
        worker: Worker to stop on signal
    """
    _active_workers.append(worker)
    signal.signal(signal.SIGTERM, _shutdown_handler)
    signal.signal(signal.SIGINT, _shutdown_handler)

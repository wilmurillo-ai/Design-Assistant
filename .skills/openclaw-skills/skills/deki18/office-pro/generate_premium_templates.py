#!/usr/bin/env python3
"""
Office Pro - Premium Template Generator

Generates enterprise-grade, professionally designed Word and Excel templates
with modern aesthetics, sophisticated typography, and commercial-quality layouts.
Designed for business delivery standards.
"""

from __future__ import annotations

import os
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any

SKILL_DIR = Path(__file__).parent
sys.path.insert(0, str(SKILL_DIR))


# ============================================================================
# PROFESSIONAL DESIGN SYSTEM
# ============================================================================

class DesignSystem:
    """Enterprise design system with sophisticated color palettes and typography"""
    
    # Primary Color Palette - Corporate Blue (Professional, Trustworthy)
    PRIMARY = {
        '900': '1E3A5F',  # Deep navy - headers
        '800': '2B579A',  # Primary blue - main brand
        '700': '3D7AB8',  # Medium blue
        '600': '5B9BD5',  # Light blue - accents
        '500': '7FB3E3',  # Sky blue
        '100': 'E8F4FC',  # Lightest - backgrounds
    }
    
    # Secondary Palette - Slate Gray (Neutral, Professional)
    SECONDARY = {
        '900': '1F2937',  # Near black - primary text
        '800': '374151',  # Dark gray
        '700': '4B5563',  # Medium gray
        '600': '6B7280',  # Light gray - secondary text
        '500': '9CA3AF',  # Muted
        '400': 'D1D5DB',  # Border gray
        '300': 'E5E7EB',  # Light border
        '200': 'F3F4F6',  # Background gray
        '100': 'F9FAFB',  # Near white
    }
    
    # Accent Colors
    ACCENT = {
        'success': '059669',  # Emerald - positive indicators
        'warning': 'D97706',  # Amber - cautions
        'danger': 'DC2626',   # Red - alerts
        'info': '0891B2',     # Cyan - information
    }
    
    # Typography Scale
    TYPOGRAPHY = {
        'display': 32,      # Main titles
        'h1': 24,           # Section headers
        'h2': 18,           # Subsection headers
        'h3': 14,           # Small headers
        'body': 11,         # Body text
        'small': 9,         # Captions, footnotes
        'tiny': 8,          # Fine print
    }
    
    # Spacing Scale (in points)
    SPACING = {
        'xs': 3,
        'sm': 6,
        'md': 12,
        'lg': 18,
        'xl': 24,
        '2xl': 36,
        '3xl': 48,
    }


# ============================================================================
# WORD TEMPLATE GENERATOR
# ============================================================================

def generate_word_templates() -> bool:
    """Generate premium Word document templates with professional design"""
    print("=" * 70)
    print("Generating Premium Word Templates - Enterprise Edition")
    print("=" * 70)
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement, parse_xml
        from docx.enum.section import WD_ORIENT
    except ImportError as e:
        print(f"[ERROR] Missing dependency: {e}")
        print("Install: pip install python-docx")
        return False
    
    template_dir = SKILL_DIR / "assets" / "templates" / "word"
    template_dir.mkdir(parents=True, exist_ok=True)
    
    ds = DesignSystem()
    
    # Helper functions for styling
    def rgb(color_hex: str) -> RGBColor:
        return RGBColor.from_string(color_hex)
    
    def set_cell_shading(cell, color: str):
        """Set cell background color"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def set_cell_border(cell, **kwargs):
        """Set cell borders"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'
                element = OxmlElement(tag)
                element.set(qn('w:val'), edge_data.get('val', 'single'))
                element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                element.set(qn('w:color'), edge_data.get('color', '000000'))
                tcBorders.append(element)
        tcPr.append(tcBorders)
    
    def add_decorative_line(doc, color: str = None, width: int = 60):
        """Add a decorative horizontal line"""
        if color is None:
            color = ds.PRIMARY['600']
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('─' * width)
        run.font.color.rgb = rgb(color)
        run.font.size = Pt(8)
        return p
    
    def add_spacer(doc, space_type: str = 'md'):
        """Add vertical spacing"""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(ds.SPACING[space_type])
        return p
    
    def style_heading(heading, level: int = 1, color: str = None):
        """Apply professional heading styles"""
        if color is None:
            color = ds.PRIMARY['900'] if level == 1 else ds.PRIMARY['800']
        size = ds.TYPOGRAPHY[f'h{min(level, 3)}']
        for run in heading.runs:
            run.font.color.rgb = rgb(color)
            run.font.size = Pt(size)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
    
    def create_cover_page(doc, title: str, subtitle: str = None, meta_info: List[Tuple[str, str]] = None):
        """Create a professional cover page"""
        # Top spacing
        for _ in range(4):
            add_spacer(doc, 'lg')
        
        # Decorative top line
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run('◆')
        run.font.size = Pt(12)
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        add_spacer(doc, 'xl')
        
        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title.upper())
        run.font.size = Pt(ds.TYPOGRAPHY['display'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        run.font.name = 'Microsoft YaHei'
        
        # Subtitle
        if subtitle:
            add_spacer(doc, 'sm')
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_para.add_run(subtitle)
            run.font.size = Pt(ds.TYPOGRAPHY['h2'])
            run.font.color.rgb = rgb(ds.PRIMARY['600'])
            run.font.name = 'Microsoft YaHei'
        
        add_spacer(doc, '2xl')
        
        # Decorative divider
        divider = doc.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run('━' * 20)
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        run.font.size = Pt(8)
        
        add_spacer(doc, '2xl')
        
        # Meta information
        if meta_info:
            for label, value in meta_info:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label_run = meta_para.add_run(f'{label}: ')
                label_run.font.size = Pt(ds.TYPOGRAPHY['small'])
                label_run.font.color.rgb = rgb(ds.SECONDARY['600'])
                label_run.font.name = 'Microsoft YaHei'
                
                value_run = meta_para.add_run(value)
                value_run.font.size = Pt(ds.TYPOGRAPHY['body'])
                value_run.font.color.rgb = rgb(ds.SECONDARY['900'])
                value_run.font.name = 'Microsoft YaHei'
        
        add_spacer(doc, '3xl')
        
        # Bottom decoration
        bottom = doc.add_paragraph()
        bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bottom.add_run('◆')
        run.font.size = Pt(12)
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        # Page break
        doc.add_page_break()
    
    def create_info_table(doc, data: List[Tuple[str, str]], title: str = None):
        """Create a professional info table with styled headers"""
        if title:
            heading = doc.add_heading(title, level=2)
            style_heading(heading, level=2)
            add_spacer(doc, 'sm')
        
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        for row_idx, (label, value) in enumerate(data):
            row = table.rows[row_idx]
            
            # Label cell styling
            label_cell = row.cells[0]
            label_cell.text = label
            label_cell.width = Inches(2.0)
            set_cell_shading(label_cell, ds.SECONDARY['200'])
            
            # Label text styling
            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in label_para.runs:
                run.font.bold = True
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
                run.font.color.rgb = rgb(ds.SECONDARY['800'])
                run.font.name = 'Microsoft YaHei'
            
            # Value cell styling
            value_cell = row.cells[1]
            value_cell.text = value
            value_cell.width = Inches(4.0)
            
            value_para = value_cell.paragraphs[0]
            for run in value_para.runs:
                run.font.size = Pt(ds.TYPOGRAPHY['body'])
                run.font.color.rgb = rgb(ds.SECONDARY['900'])
                run.font.name = 'Microsoft YaHei'
        
        add_spacer(doc, 'md')
        return table
    
    def add_section_heading(doc, title: str, level: int = 1):
        """Add a styled section heading with optional underline"""
        heading = doc.add_heading(title, level=level)
        style_heading(heading, level=level)
        
        # Add subtle underline for h1
        if level == 1:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(ds.SPACING['md'])
            run = p.add_run('')
            # The underline is handled by the heading style
        
        return heading
    
    # Set default document styles
    def setup_document_styles(doc):
        """Setup professional default styles for the document"""
        # Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(ds.TYPOGRAPHY['body'])
        font.color.rgb = rgb(ds.SECONDARY['900'])
        
        # Paragraph spacing
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(ds.SPACING['sm'])
    
    # ============================================================================
    # TEMPLATE 1: PREMIUM MEETING MINUTES
    # ============================================================================
    def create_meeting_minutes():
        doc = Document()
        setup_document_styles(doc)
        
        # Cover page
        create_cover_page(
            doc,
            title='Meeting Minutes',
            subtitle='会议纪要',
            meta_info=[
                ('Document No.', '{{document_number}}'),
                ('Meeting', '{{meeting_title}}'),
                ('Date', '{{meeting_date}}'),
            ]
        )
        
        # Information Section
        add_section_heading(doc, 'Meeting Information / 会议信息', level=1)
        
        info_data = [
            ('Meeting Title / 会议主题', '{{meeting_title}}'),
            ('Meeting Type / 会议类型', '{{meeting_type}}'),
            ('Date / 日期', '{{meeting_date}}'),
            ('Time / 时间', '{{meeting_time}}'),
            ('Location / 地点', '{{meeting_location}}'),
            ('Chairperson / 主持人', '{{chairperson}}'),
            ('Secretary / 记录人', '{{secretary}}'),
        ]
        create_info_table(doc, info_data)
        
        add_decorative_line(doc, ds.SECONDARY['300'])
        add_spacer(doc, 'md')
        
        # Attendees Section
        add_section_heading(doc, 'Attendees / 与会人员', level=2)
        
        # Create attendees table
        att_table = doc.add_table(rows=1, cols=4)
        att_table.style = 'Table Grid'
        
        # Header row
        headers = ['Name / 姓名', 'Department / 部门', 'Role / 职务', 'Status / 状态']
        header_cells = att_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            set_cell_shading(header_cells[i], ds.PRIMARY['900'])
            for run in header_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Add placeholder row
        row = att_table.add_row()
        row.cells[0].text = '{{attendee_name}}'
        row.cells[1].text = '{{attendee_dept}}'
        row.cells[2].text = '{{attendee_role}}'
        row.cells[3].text = '{{attendee_status}}'
        
        add_spacer(doc, 'lg')
        
        # Agenda Section
        add_section_heading(doc, 'Agenda / 会议议程', level=2)
        agenda_para = doc.add_paragraph('{{agenda_items}}')
        agenda_para.paragraph_format.left_indent = Inches(0.25)
        
        add_spacer(doc, 'lg')
        
        # Discussion Section
        add_section_heading(doc, 'Discussion / 讨论记录', level=2)
        doc.add_paragraph('{{discussion_content}}')
        
        add_spacer(doc, 'lg')
        
        # Decisions Section
        add_section_heading(doc, 'Decisions / 决议事项', level=2)
        decisions = doc.add_paragraph()
        decisions.paragraph_format.left_indent = Inches(0.25)
        run = decisions.add_run('{{decisions}}')
        run.font.color.rgb = rgb(ds.ACCENT['success'])
        
        add_spacer(doc, 'lg')
        
        # Action Items Section
        add_section_heading(doc, 'Action Items / 行动项', level=2)
        
        # Action items table
        action_table = doc.add_table(rows=1, cols=5)
        action_table.style = 'Table Grid'
        
        action_headers = ['No.', 'Action / 事项', 'Owner / 负责人', 'Deadline / 截止日期', 'Status / 状态']
        action_cells = action_table.rows[0].cells
        for i, header in enumerate(action_headers):
            action_cells[i].text = header
            set_cell_shading(action_cells[i], ds.PRIMARY['800'])
            for run in action_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Placeholder row
        action_row = action_table.add_row()
        action_row.cells[0].text = '1'
        action_row.cells[1].text = '{{action_item}}'
        action_row.cells[2].text = '{{action_owner}}'
        action_row.cells[3].text = '{{action_deadline}}'
        action_row.cells[4].text = '{{action_status}}'
        
        add_spacer(doc, 'xl')
        
        # Footer
        add_decorative_line(doc, ds.SECONDARY['300'])
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run('End of Minutes / 会议记录结束')
        run.font.size = Pt(ds.TYPOGRAPHY['tiny'])
        run.font.color.rgb = rgb(ds.SECONDARY['500'])
        run.font.italic = True
        
        return doc
    
    # ============================================================================
    # TEMPLATE 2: PREMIUM BUSINESS LETTER
    # ============================================================================
    def create_business_letter():
        doc = Document()
        setup_document_styles(doc)
        
        # Letterhead Section
        letterhead = doc.add_paragraph()
        letterhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company name
        run = letterhead.add_run('{{company_name}}')
        run.font.size = Pt(ds.TYPOGRAPHY['h1'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        run.font.name = 'Microsoft YaHei'
        
        # Company details
        add_spacer(doc, 'xs')
        details = doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details.add_run('{{company_address}} | {{company_phone}} | {{company_email}}')
        run.font.size = Pt(ds.TYPOGRAPHY['tiny'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        
        add_spacer(doc, 'sm')
        add_decorative_line(doc, ds.PRIMARY['600'], width=40)
        add_spacer(doc, 'lg')
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_para.add_run('{{date}}')
        run.font.size = Pt(ds.TYPOGRAPHY['body'])
        
        add_spacer(doc, 'lg')
        
        # Recipient
        doc.add_paragraph('{{recipient_name}}')
        doc.add_paragraph('{{recipient_title}}')
        doc.add_paragraph('{{recipient_company}}')
        doc.add_paragraph('{{recipient_address}}')
        
        add_spacer(doc, 'md')
        
        # Subject line
        subject = doc.add_paragraph()
        run = subject.add_run('Subject: ')
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        subject.add_run('{{subject}}')
        
        add_spacer(doc, 'md')
        
        # Salutation
        salutation = doc.add_paragraph()
        run = salutation.add_run('Dear {{recipient_name}},')
        run.font.size = Pt(ds.TYPOGRAPHY['body'])
        
        add_spacer(doc, 'sm')
        
        # Body paragraphs
        body = doc.add_paragraph('{{letter_body}}')
        body.paragraph_format.first_line_indent = Inches(0.3)
        body.paragraph_format.line_spacing = 1.5
        
        add_spacer(doc, 'md')
        
        # Closing
        closing = doc.add_paragraph()
        closing.add_run('{{closing}}')
        
        add_spacer(doc, 'xl')
        add_spacer(doc, 'sm')
        
        # Signature block
        sig = doc.add_paragraph()
        run = sig.add_run('{{sender_name}}')
        run.font.bold = True
        doc.add_paragraph('{{sender_title}}')
        doc.add_paragraph('{{sender_department}}')
        
        add_spacer(doc, 'xl')
        
        # Enclosure
        enclosure = doc.add_paragraph()
        run = enclosure.add_run('Enclosure: {{enclosure}}')
        run.font.size = Pt(ds.TYPOGRAPHY['tiny'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        
        # Footer line
        add_spacer(doc, 'xl')
        add_decorative_line(doc, ds.SECONDARY['300'], width=30)
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('{{company_name}} | {{company_website}}')
        run.font.size = Pt(ds.TYPOGRAPHY['tiny'])
        run.font.color.rgb = rgb(ds.SECONDARY['500'])
        
        return doc
    
    # ============================================================================
    # TEMPLATE 3: PREMIUM RESUME
    # ============================================================================
    def create_resume():
        doc = Document()
        setup_document_styles(doc)
        
        # Header with name and contact
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run('{{full_name}}')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        run.font.name = 'Microsoft YaHei'
        
        add_spacer(doc, 'xs')
        
        # Professional title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('{{professional_title}}')
        run.font.size = Pt(ds.TYPOGRAPHY['h3'])
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        run.font.italic = True
        
        add_spacer(doc, 'xs')
        
        # Contact info line
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact.add_run('{{email}} | {{phone}} | {{location}} | {{linkedin}}')
        run.font.size = Pt(ds.TYPOGRAPHY['small'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        
        add_spacer(doc, 'sm')
        add_decorative_line(doc, ds.PRIMARY['600'], width=50)
        add_spacer(doc, 'md')
        
        # Professional Summary
        add_section_heading(doc, 'Professional Summary / 个人简介', level=2)
        summary = doc.add_paragraph('{{professional_summary}}')
        summary.paragraph_format.line_spacing = 1.3
        
        add_spacer(doc, 'md')
        
        # Work Experience
        add_section_heading(doc, 'Work Experience / 工作经历', level=2)
        
        exp_table = doc.add_table(rows=1, cols=2)
        exp_table.autofit = False
        
        # Experience entry
        row = exp_table.rows[0]
        left_cell = row.cells[0]
        left_cell.width = Inches(2.0)
        
        date_para = left_cell.paragraphs[0]
        run = date_para.add_run('{{job_date}}')
        run.font.size = Pt(ds.TYPOGRAPHY['small'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        
        right_cell = row.cells[1]
        right_cell.width = Inches(4.5)
        
        job_title = right_cell.paragraphs[0]
        run = job_title.add_run('{{job_title}}')
        run.font.bold = True
        run.font.size = Pt(ds.TYPOGRAPHY['body'])
        run.font.color.rgb = rgb(ds.SECONDARY['900'])
        
        company_para = right_cell.add_paragraph()
        run = company_para.add_run('{{company_name}} | {{company_location}}')
        run.font.size = Pt(ds.TYPOGRAPHY['small'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        run.font.italic = True
        
        desc_para = right_cell.add_paragraph()
        run = desc_para.add_run('{{job_description}}')
        run.font.size = Pt(ds.TYPOGRAPHY['body'])
        
        add_spacer(doc, 'md')
        
        # Education
        add_section_heading(doc, 'Education / 教育背景', level=2)
        edu = doc.add_paragraph()
        run = edu.add_run('{{degree}}')
        run.font.bold = True
        doc.add_paragraph('{{university}} | {{graduation_date}}')
        doc.add_paragraph('{{education_details}}')
        
        add_spacer(doc, 'md')
        
        # Skills
        add_section_heading(doc, 'Skills / 技能专长', level=2)
        skills = doc.add_paragraph('{{skills}}')
        
        add_spacer(doc, 'md')
        
        # Certifications
        add_section_heading(doc, 'Certifications / 证书资质', level=2)
        certs = doc.add_paragraph('{{certifications}}')
        
        return doc
    
    # ============================================================================
    # TEMPLATE 4: PREMIUM PROJECT PROPOSAL
    # ============================================================================
    def create_project_proposal():
        doc = Document()
        setup_document_styles(doc)
        
        # Professional cover page
        create_cover_page(
            doc,
            title='Project Proposal',
            subtitle='{{project_name}}',
            meta_info=[
                ('Prepared For', '{{client_name}}'),
                ('Prepared By', '{{company_name}}'),
                ('Date', '{{proposal_date}}'),
                ('Version', '{{version}}'),
                ('Valid Until', '{{valid_until}}'),
            ]
        )
        
        # Executive Summary
        add_section_heading(doc, 'Executive Summary / 执行摘要', level=1)
        exec_summary = doc.add_paragraph('{{executive_summary}}')
        exec_summary.paragraph_format.line_spacing = 1.3
        
        add_spacer(doc, 'lg')
        
        # About Us
        add_section_heading(doc, 'About Us / 关于我们', level=2)
        doc.add_paragraph('{{company_overview}}')
        
        add_spacer(doc, 'lg')
        
        # Project Background
        add_section_heading(doc, 'Project Background / 项目背景', level=1)
        doc.add_paragraph('{{project_background}}')
        
        add_spacer(doc, 'lg')
        
        # Objectives
        add_section_heading(doc, 'Objectives / 项目目标', level=2)
        objectives = doc.add_paragraph()
        objectives.paragraph_format.left_indent = Inches(0.25)
        run = objectives.add_run('{{objectives}}')
        
        add_spacer(doc, 'lg')
        
        # Scope of Work
        add_section_heading(doc, 'Scope of Work / 工作范围', level=1)
        doc.add_paragraph('{{scope_of_work}}')
        
        add_spacer(doc, 'lg')
        
        # Methodology
        add_section_heading(doc, 'Methodology / 方法论', level=2)
        doc.add_paragraph('{{methodology}}')
        
        add_spacer(doc, 'lg')
        
        # Timeline
        add_section_heading(doc, 'Project Timeline / 项目时间线', level=2)
        
        timeline_table = doc.add_table(rows=1, cols=4)
        timeline_table.style = 'Table Grid'
        
        timeline_headers = ['Phase / 阶段', 'Description / 描述', 'Duration / 周期', 'Deliverables / 交付物']
        header_cells = timeline_table.rows[0].cells
        for i, header in enumerate(timeline_headers):
            header_cells[i].text = header
            set_cell_shading(header_cells[i], ds.PRIMARY['800'])
            for run in header_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Placeholder row
        row = timeline_table.add_row()
        row.cells[0].text = '{{phase_number}}'
        row.cells[1].text = '{{phase_description}}'
        row.cells[2].text = '{{phase_duration}}'
        row.cells[3].text = '{{phase_deliverables}}'
        
        add_spacer(doc, 'lg')
        
        # Budget
        add_section_heading(doc, 'Investment / 投资预算', level=1)
        
        budget_table = doc.add_table(rows=1, cols=4)
        budget_table.style = 'Table Grid'
        
        budget_headers = ['Item / 项目', 'Description / 描述', 'Quantity / 数量', 'Amount / 金额']
        budget_cells = budget_table.rows[0].cells
        for i, header in enumerate(budget_headers):
            budget_cells[i].text = header
            set_cell_shading(budget_cells[i], ds.PRIMARY['800'])
            for run in budget_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Placeholder rows
        for _ in range(3):
            row = budget_table.add_row()
            row.cells[0].text = '{{budget_item}}'
            row.cells[1].text = '{{item_description}}'
            row.cells[2].text = '{{item_quantity}}'
            row.cells[3].text = '{{item_amount}}'
        
        # Total row
        total_row = budget_table.add_row()
        total_row.cells[0].merge(total_row.cells[2])
        total_label = total_row.cells[0]
        total_label.text = 'Total Investment / 总投资'
        set_cell_shading(total_label, ds.SECONDARY['200'])
        for run in total_label.paragraphs[0].runs:
            run.font.bold = True
        
        total_row.cells[1].text = '{{total_amount}}'
        set_cell_shading(total_row.cells[1], ds.PRIMARY['100'])
        for run in total_row.cells[1].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = rgb(ds.PRIMARY['900'])
        
        add_spacer(doc, 'lg')
        
        # Terms and Conditions
        add_section_heading(doc, 'Terms & Conditions / 条款条件', level=2)
        doc.add_paragraph('{{terms_conditions}}')
        
        add_spacer(doc, 'xl')
        
        # Acceptance
        add_section_heading(doc, 'Proposal Acceptance / 提案接受', level=2)
        accept_para = doc.add_paragraph(
            'By signing below, the Client accepts this proposal and agrees to the terms and conditions outlined herein.'
        )
        accept_para.paragraph_format.italic = True
        
        add_spacer(doc, 'lg')
        
        # Signature table
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.style = 'Table Grid'
        
        sig_data = [
            ['Client / 客户', 'Service Provider / 服务商'],
            ['Signature: _______________', 'Signature: _______________'],
            ['Date: _______________', 'Date: _______________'],
        ]
        
        for i, row_data in enumerate(sig_data):
            row = sig_table.rows[i]
            for j, text in enumerate(row_data):
                row.cells[j].text = text
                if i == 0:
                    set_cell_shading(row.cells[j], ds.SECONDARY['200'])
                    for run in row.cells[j].paragraphs[0].runs:
                        run.font.bold = True
        
        return doc
    
    # ============================================================================
    # TEMPLATE 5: PREMIUM WORK REPORT
    # ============================================================================
    def create_work_report():
        doc = Document()
        setup_document_styles(doc)
        
        # Cover page
        create_cover_page(
            doc,
            title='Work Report',
            subtitle='工作报告',
            meta_info=[
                ('Report Period', '{{report_period}}'),
                ('Department', '{{department}}'),
                ('Reporter', '{{reporter}}'),
                ('Date', '{{report_date}}'),
            ]
        )
        
        # Executive Summary
        add_section_heading(doc, 'Executive Summary / 执行摘要', level=1)
        summary = doc.add_paragraph('{{executive_summary}}')
        summary.paragraph_format.line_spacing = 1.3
        
        add_spacer(doc, 'lg')
        
        # Key Achievements
        add_section_heading(doc, 'Key Achievements / 主要成果', level=2)
        
        achievements_table = doc.add_table(rows=1, cols=4)
        achievements_table.style = 'Table Grid'
        
        headers = ['Achievement / 成果', 'Target / 目标', 'Actual / 实际', 'Status / 状态']
        header_cells = achievements_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            set_cell_shading(header_cells[i], ds.PRIMARY['800'])
            for run in header_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Placeholder row
        row = achievements_table.add_row()
        row.cells[0].text = '{{achievement_name}}'
        row.cells[1].text = '{{target_value}}'
        row.cells[2].text = '{{actual_value}}'
        row.cells[3].text = '{{achievement_status}}'
        
        add_spacer(doc, 'lg')
        
        # Work Details
        add_section_heading(doc, 'Work Details / 工作详情', level=1)
        doc.add_paragraph('{{work_details}}')
        
        add_spacer(doc, 'lg')
        
        # Data Analysis
        add_section_heading(doc, 'Data Analysis / 数据分析', level=2)
        doc.add_paragraph('{{data_analysis}}')
        
        add_spacer(doc, 'lg')
        
        # Challenges and Solutions
        add_section_heading(doc, 'Challenges & Solutions / 问题与解决', level=2)
        
        challenge_table = doc.add_table(rows=1, cols=3)
        challenge_table.style = 'Table Grid'
        
        challenge_headers = ['Challenge / 问题', 'Solution / 解决方案', 'Result / 结果']
        challenge_cells = challenge_table.rows[0].cells
        for i, header in enumerate(challenge_headers):
            challenge_cells[i].text = header
            set_cell_shading(challenge_cells[i], ds.PRIMARY['800'])
            for run in challenge_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Placeholder row
        row = challenge_table.add_row()
        row.cells[0].text = '{{challenge_description}}'
        row.cells[1].text = '{{solution_description}}'
        row.cells[2].text = '{{result_description}}'
        
        add_spacer(doc, 'lg')
        
        # Next Steps
        add_section_heading(doc, 'Next Steps / 下步计划', level=2)
        
        steps_table = doc.add_table(rows=1, cols=4)
        steps_table.style = 'Table Grid'
        
        steps_headers = ['Step / 步骤', 'Description / 描述', 'Owner / 负责人', 'Deadline / 截止日期']
        steps_cells = steps_table.rows[0].cells
        for i, header in enumerate(steps_headers):
            steps_cells[i].text = header
            set_cell_shading(steps_cells[i], ds.PRIMARY['800'])
            for run in steps_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        # Placeholder row
        row = steps_table.add_row()
        row.cells[0].text = '{{step_number}}'
        row.cells[1].text = '{{step_description}}'
        row.cells[2].text = '{{step_owner}}'
        row.cells[3].text = '{{step_deadline}}'
        
        return doc
    
    # ============================================================================
    # TEMPLATE 6: PREMIUM CONTRACT
    # ============================================================================
    def create_contract():
        doc = Document()
        setup_document_styles(doc)
        
        # Contract Header
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('{{contract_title}}')
        run.font.size = Pt(ds.TYPOGRAPHY['h1'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        
        add_spacer(doc, 'sm')
        
        # Contract metadata
        meta_table = doc.add_table(rows=2, cols=2)
        meta_table.style = 'Table Grid'
        
        meta_data = [
            ('Contract No. / 合同编号', '{{contract_number}}'),
            ('Date / 签署日期', '{{contract_date}}'),
        ]
        
        for i, (label, value) in enumerate(meta_data):
            row = meta_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            set_cell_shading(row.cells[0], ds.SECONDARY['200'])
            for run in row.cells[0].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        add_spacer(doc, 'lg')
        add_decorative_line(doc, ds.PRIMARY['600'], width=50)
        add_spacer(doc, 'lg')
        
        # Parties
        add_section_heading(doc, 'Parties / 合同双方', level=2)
        
        party_table = doc.add_table(rows=2, cols=2)
        party_table.style = 'Table Grid'
        
        party_data = [
            ('Party A (甲方)', '{{party_a_name}}'),
            ('Address / 地址', '{{party_a_address}}'),
            ('Party B (乙方)', '{{party_b_name}}'),
            ('Address / 地址', '{{party_b_address}}'),
        ]
        
        for i, (label, value) in enumerate(party_data):
            row = party_table.rows[i // 2]
            cell_idx = (i % 2) * 2
            if cell_idx < 2:
                row.cells[cell_idx].text = label
                row.cells[cell_idx + 1].text = value
                set_cell_shading(row.cells[cell_idx], ds.SECONDARY['200'])
                for run in row.cells[cell_idx].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(ds.TYPOGRAPHY['small'])
        
        add_spacer(doc, 'lg')
        
        # Whereas Clause
        add_section_heading(doc, 'Whereas / 鉴于', level=2)
        doc.add_paragraph('{{whereas_clauses}}')
        
        add_spacer(doc, 'lg')
        
        # Terms and Conditions
        add_section_heading(doc, 'Terms & Conditions / 条款', level=1)
        
        for i in range(1, 9):
            article_heading = doc.add_heading(f'Article {i} / 第{i}条', level=3)
            style_heading(article_heading, level=3)
            doc.add_paragraph(f'{{article_{i}}}')
            add_spacer(doc, 'sm')
        
        add_spacer(doc, 'lg')
        
        # Signature Section
        add_section_heading(doc, 'Signatures / 签署', level=1)
        
        sig_para = doc.add_paragraph(
            'IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.'
        )
        sig_para.paragraph_format.italic = True
        
        add_spacer(doc, 'lg')
        
        # Signature table
        sig_table = doc.add_table(rows=4, cols=2)
        sig_table.style = 'Table Grid'
        
        sig_headers = ['Party A (甲方)', 'Party B (乙方)']
        header_cells = sig_table.rows[0].cells
        for i, header in enumerate(sig_headers):
            header_cells[i].text = header
            set_cell_shading(header_cells[i], ds.PRIMARY['800'])
            for run in header_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = rgb('FFFFFF')
        
        # Signature rows
        sig_table.rows[1].cells[0].text = 'Signature / 签名:\n\n\n_________________'
        sig_table.rows[1].cells[1].text = 'Signature / 签名:\n\n\n_________________'
        sig_table.rows[2].cells[0].text = 'Name / 姓名: {{party_a_signatory}}'
        sig_table.rows[2].cells[1].text = 'Name / 姓名: {{party_b_signatory}}'
        sig_table.rows[3].cells[0].text = 'Date / 日期: _________________'
        sig_table.rows[3].cells[1].text = 'Date / 日期: _________________'
        
        return doc
    
    # ============================================================================
    # TEMPLATE 7: PREMIUM PRESS RELEASE
    # ============================================================================
    def create_press_release():
        doc = Document()
        setup_document_styles(doc)
        
        # For Immediate Release badge
        badge = doc.add_paragraph()
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = badge.add_run('FOR IMMEDIATE RELEASE')
        run.font.size = Pt(ds.TYPOGRAPHY['tiny'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.ACCENT['success'])
        
        add_spacer(doc, 'sm')
        
        # Release Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run('{{release_date}}')
        run.font.size = Pt(ds.TYPOGRAPHY['small'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        
        add_spacer(doc, 'md')
        
        # Headline
        headline = doc.add_paragraph()
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = headline.add_run('{{headline}}')
        run.font.size = Pt(ds.TYPOGRAPHY['h1'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        
        add_spacer(doc, 'xs')
        
        # Subheadline
        subheadline = doc.add_paragraph()
        subheadline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subheadline.add_run('{{subheadline}}')
        run.font.size = Pt(ds.TYPOGRAPHY['h3'])
        run.font.italic = True
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        add_spacer(doc, 'md')
        add_decorative_line(doc, ds.SECONDARY['300'])
        add_spacer(doc, 'md')
        
        # Dateline
        dateline = doc.add_paragraph()
        run = dateline.add_run('{{location}} — ')
        run.font.bold = True
        dateline.add_run('{{dateline_content}}')
        
        add_spacer(doc, 'sm')
        
        # Body paragraphs
        doc.add_paragraph('{{body_paragraph_1}}')
        add_spacer(doc, 'sm')
        doc.add_paragraph('{{body_paragraph_2}}')
        add_spacer(doc, 'sm')
        doc.add_paragraph('{{body_paragraph_3}}')
        
        add_spacer(doc, 'md')
        
        # Quote
        quote_para = doc.add_paragraph()
        quote_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        quote_para.paragraph_format.left_indent = Inches(0.5)
        quote_para.paragraph_format.right_indent = Inches(0.5)
        run = quote_para.add_run('"{{quote}}"')
        run.font.italic = True
        run.font.size = Pt(ds.TYPOGRAPHY['h3'])
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        attribution = doc.add_paragraph()
        attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = attribution.add_run('— {{quote_attribution}}')
        run.font.size = Pt(ds.TYPOGRAPHY['small'])
        run.font.color.rgb = rgb(ds.SECONDARY['600'])
        
        add_spacer(doc, 'md')
        
        # More body
        doc.add_paragraph('{{body_paragraph_4}}')
        
        add_spacer(doc, 'lg')
        add_decorative_line(doc, ds.SECONDARY['300'])
        add_spacer(doc, 'md')
        
        # About Company
        about_heading = doc.add_heading('About {{company_name}}', level=3)
        style_heading(about_heading, level=3)
        doc.add_paragraph('{{company_description}}')
        
        add_spacer(doc, 'md')
        
        # Media Contact
        contact_heading = doc.add_heading('Media Contact / 媒体联系', level=3)
        style_heading(contact_heading, level=3)
        
        contact_info = [
            ('Name', '{{contact_name}}'),
            ('Title', '{{contact_title}}'),
            ('Email', '{{contact_email}}'),
            ('Phone', '{{contact_phone}}'),
        ]
        
        for label, value in contact_info:
            p = doc.add_paragraph()
            run = p.add_run(f'{label}: ')
            run.font.bold = True
            p.add_run(value)
        
        add_spacer(doc, 'lg')
        
        # Boilerplate
        boilerplate = doc.add_paragraph()
        boilerplate.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = boilerplate.add_run('###')
        run.font.color.rgb = rgb(ds.SECONDARY['400'])
        
        return doc
    
    # ============================================================================
    # TEMPLATE 8: PREMIUM FORMAL INVITATION
    # ============================================================================
    def create_invitation():
        doc = Document()
        setup_document_styles(doc)
        
        # Top decorative border
        for _ in range(3):
            add_spacer(doc, 'lg')
        
        # Decorative top
        top_decoration = doc.add_paragraph()
        top_decoration.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = top_decoration.add_run('❦ ❦ ❦')
        run.font.size = Pt(16)
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        add_spacer(doc, 'xl')
        
        # Invitation title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('INVITATION')
        run.font.size = Pt(ds.TYPOGRAPHY['display'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['900'])
        
        add_spacer(doc, 'xs')
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('邀请函')
        run.font.size = Pt(ds.TYPOGRAPHY['h2'])
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        add_spacer(doc, '2xl')
        
        # Salutation
        salutation = doc.add_paragraph()
        salutation.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = salutation.add_run('Dear {{recipient_name}},')
        run.font.size = Pt(ds.TYPOGRAPHY['body'])
        
        add_spacer(doc, 'md')
        
        # Invitation text
        invitation_text = doc.add_paragraph()
        invitation_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
        invitation_text.add_run('You are cordially invited to attend:')
        
        add_spacer(doc, 'md')
        
        # Event name
        event_name = doc.add_paragraph()
        event_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = event_name.add_run('{{event_name}}')
        run.font.size = Pt(ds.TYPOGRAPHY['h1'])
        run.font.bold = True
        run.font.color.rgb = rgb(ds.PRIMARY['800'])
        
        add_spacer(doc, 'lg')
        
        # Event details box
        details_table = doc.add_table(rows=4, cols=2)
        details_table.style = 'Table Grid'
        details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        details_data = [
            ('Date / 日期', '{{event_date}}'),
            ('Time / 时间', '{{event_time}}'),
            ('Venue / 地点', '{{event_venue}}'),
            ('Address / 地址', '{{event_address}}'),
        ]
        
        for i, (label, value) in enumerate(details_data):
            row = details_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            set_cell_shading(row.cells[0], ds.SECONDARY['200'])
            for run in row.cells[0].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(ds.TYPOGRAPHY['small'])
            for run in row.cells[1].paragraphs[0].runs:
                run.font.size = Pt(ds.TYPOGRAPHY['body'])
        
        add_spacer(doc, 'lg')
        
        # Event description
        doc.add_paragraph('{{event_description}}')
        
        add_spacer(doc, 'lg')
        
        # Dress code
        dress_code = doc.add_paragraph()
        run = dress_code.add_run('Dress Code / 着装要求: ')
        run.font.bold = True
        dress_code.add_run('{{dress_code}}')
        
        add_spacer(doc, 'lg')
        
        # RSVP
        rsvp_heading = doc.add_heading('RSVP', level=3)
        style_heading(rsvp_heading, level=3)
        
        doc.add_paragraph('Please confirm your attendance by {{rsvp_date}}:')
        
        rsvp_info = [
            ('Contact', '{{rsvp_contact}}'),
            ('Email', '{{rsvp_email}}'),
            ('Phone', '{{rsvp_phone}}'),
        ]
        
        for label, value in rsvp_info:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f'{label}: ')
            run.font.bold = True
            p.add_run(value)
        
        add_spacer(doc, 'xl')
        
        # Closing
        closing = doc.add_paragraph()
        closing.add_run('We look forward to your presence.')
        
        add_spacer(doc, 'xl')
        add_spacer(doc, 'sm')
        
        # Host signature
        doc.add_paragraph('Sincerely,')
        add_spacer(doc, 'sm')
        doc.add_paragraph('{{host_name}}')
        doc.add_paragraph('{{host_title}}')
        
        add_spacer(doc, '2xl')
        
        # Bottom decoration
        bottom_decoration = doc.add_paragraph()
        bottom_decoration.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bottom_decoration.add_run('❦ ❦ ❦')
        run.font.size = Pt(16)
        run.font.color.rgb = rgb(ds.PRIMARY['600'])
        
        return doc
    
    # Generate all templates
    templates = [
        ('meeting-minutes.docx', create_meeting_minutes),
        ('letter-business.docx', create_business_letter),
        ('resume-professional.docx', create_resume),
        ('project-proposal.docx', create_project_proposal),
        ('work-report.docx', create_work_report),
        ('contract-simple.docx', create_contract),
        ('press-release.docx', create_press_release),
        ('invitation-formal.docx', create_invitation),
    ]
    
    for filename, create_func in templates:
        print(f"  [Word] {filename}...", end=" ")
        try:
            doc = create_func()
            filepath = template_dir / filename
            doc.save(str(filepath))
            print("OK")
        except Exception as e:
            print(f"FAILED: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    print(f"  Generated {len(templates)} premium Word templates")
    return True


# ============================================================================
# EXCEL TEMPLATE GENERATOR
# ============================================================================

def generate_excel_templates() -> bool:
    """Generate premium Excel spreadsheet templates with professional styling"""
    print("\n" + "=" * 70)
    print("Generating Premium Excel Templates - Enterprise Edition")
    print("=" * 70)
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side,
            NamedStyle, Protection
        )
        from openpyxl.utils import get_column_letter
        from openpyxl.chart import BarChart, PieChart, LineChart, Reference
        from openpyxl.formatting.rule import ColorScaleRule, CellIsRule
        from openpyxl.comments import Comment
    except ImportError as e:
        print(f"[ERROR] Missing dependency: {e}")
        print("Install: pip install openpyxl")
        return False
    
    template_dir = SKILL_DIR / "assets" / "templates" / "excel"
    template_dir.mkdir(parents=True, exist_ok=True)
    
    ds = DesignSystem()
    
    # Helper styles
    def rgb(color_hex: str):
        return color_hex
    
    thin_border = Border(
        left=Side(style='thin', color=ds.SECONDARY['400']),
        right=Side(style='thin', color=ds.SECONDARY['400']),
        top=Side(style='thin', color=ds.SECONDARY['400']),
        bottom=Side(style='thin', color=ds.SECONDARY['400'])
    )
    
    header_fill = PatternFill(start_color=ds.PRIMARY['900'], end_color=ds.PRIMARY['900'], fill_type='solid')
    header_font = Font(name='Microsoft YaHei', size=10, bold=True, color='FFFFFF')
    
    subheader_fill = PatternFill(start_color=ds.PRIMARY['100'], end_color=ds.PRIMARY['100'], fill_type='solid')
    subheader_font = Font(name='Microsoft YaHei', size=10, bold=True, color=ds.PRIMARY['900'])
    
    data_font = Font(name='Microsoft YaHei', size=10, color=ds.SECONDARY['900'])
    
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    right_align = Alignment(horizontal='right', vertical='center')
    
    # ============================================================================
    # TEMPLATE 1: PREMIUM SALES REPORT
    # ============================================================================
    def create_sales_report():
        wb = Workbook()
        ws = wb.active
        ws.title = "Sales Report"
        
        # Title section
        ws.merge_cells('A1:H1')
        title_cell = ws['A1']
        title_cell.value = 'SALES PERFORMANCE REPORT'
        title_cell.font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        title_cell.alignment = center_align
        
        ws.merge_cells('A2:H2')
        subtitle_cell = ws['A2']
        subtitle_cell.value = '销售绩效报告'
        subtitle_cell.font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        subtitle_cell.alignment = center_align
        
        # Report info
        ws['A4'] = 'Report Period:'
        ws['B4'] = '{{report_period}}'
        ws['D4'] = 'Report Date:'
        ws['E4'] = '{{report_date}}'
        ws['G4'] = 'Sales Team:'
        ws['H4'] = '{{sales_team}}'
        
        for cell in ['A4', 'D4', 'G4']:
            ws[cell].font = Font(name='Microsoft YaHei', size=9, bold=True, color=ds.SECONDARY['700'])
        
        # Summary metrics
        ws['A6'] = 'SUMMARY METRICS'
        ws['A6'].font = subheader_font
        ws['A6'].fill = subheader_fill
        ws.merge_cells('A6:H6')
        
        metrics = [
            ('Total Sales', '{{total_sales}}'),
            ('Total Units', '{{total_units}}'),
            ('Avg Deal Size', '{{avg_deal_size}}'),
            ('Conversion Rate', '{{conversion_rate}}'),
        ]
        
        row = 7
        for label, placeholder in metrics:
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=1).font = Font(name='Microsoft YaHei', size=9, bold=True)
            ws.cell(row=row, column=2, value=placeholder)
            ws.cell(row=row, column=2).font = data_font
            row += 1
        
        # Product performance table
        start_row = row + 2
        ws.cell(row=start_row, column=1, value='PRODUCT PERFORMANCE')
        ws.cell(row=start_row, column=1).font = subheader_font
        ws.cell(row=start_row, column=1).fill = subheader_fill
        ws.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=8)
        
        # Headers
        headers = ['Product', 'Category', 'Units Sold', 'Revenue', 'Cost', 'Profit', 'Margin', 'Market Share']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=start_row + 1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        # Placeholder row
        placeholder_row = start_row + 2
        placeholders = ['{{product_name}}', '{{category}}', '{{units_sold}}', '{{revenue}}', 
                       '{{cost}}', '{{profit}}', '{{margin}}', '{{market_share}}']
        for col, placeholder in enumerate(placeholders, 1):
            cell = ws.cell(row=placeholder_row, column=col, value=placeholder)
            cell.font = data_font
            cell.alignment = center_align if col > 2 else left_align
            cell.border = thin_border
        
        # Column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15
        for col in ['C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 12
        
        return wb
    
    # ============================================================================
    # TEMPLATE 2: PREMIUM FINANCIAL STATEMENT
    # ============================================================================
    def create_financial_statement():
        wb = Workbook()
        ws = wb.active
        ws.title = "Financial Statement"
        
        # Title
        ws.merge_cells('A1:F1')
        ws['A1'] = 'FINANCIAL STATEMENT'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:F2')
        ws['A2'] = '财务报表'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # Company info
        ws['A4'] = 'Company:'
        ws['B4'] = '{{company_name}}'
        ws['D4'] = 'Period:'
        ws['E4'] = '{{fiscal_period}}'
        
        # Balance Sheet Section
        ws['A6'] = 'BALANCE SHEET'
        ws['A6'].font = subheader_font
        ws['A6'].fill = subheader_fill
        ws.merge_cells('A6:F6')
        
        balance_items = [
            ('ASSETS', ''),
            ('Current Assets', ''),
            ('  Cash and Equivalents', '{{cash}}'),
            ('  Accounts Receivable', '{{accounts_receivable}}'),
            ('  Inventory', '{{inventory}}'),
            ('  Prepaid Expenses', '{{prepaid_expenses}}'),
            ('Total Current Assets', '{{total_current_assets}}'),
            ('', ''),
            ('Non-Current Assets', ''),
            ('  Property, Plant & Equipment', '{{ppe}}'),
            ('  Intangible Assets', '{{intangible_assets}}'),
            ('Total Non-Current Assets', '{{total_non_current_assets}}'),
            ('TOTAL ASSETS', '{{total_assets}}'),
        ]
        
        row = 7
        for item, value in balance_items:
            ws.cell(row=row, column=1, value=item)
            ws.cell(row=row, column=4, value=value)
            
            if item in ['ASSETS', 'LIABILITIES', 'EQUITY']:
                ws.cell(row=row, column=1).font = Font(name='Microsoft YaHei', size=11, bold=True, color=ds.PRIMARY['900'])
            elif item.startswith('Total') or item.startswith('TOTAL'):
                ws.cell(row=row, column=1).font = Font(name='Microsoft YaHei', size=10, bold=True)
                ws.cell(row=row, column=4).font = Font(name='Microsoft YaHei', size=10, bold=True)
                ws.cell(row=row, column=1).fill = PatternFill(start_color=ds.SECONDARY['200'], end_color=ds.SECONDARY['200'], fill_type='solid')
                ws.cell(row=row, column=4).fill = PatternFill(start_color=ds.SECONDARY['200'], end_color=ds.SECONDARY['200'], fill_type='solid')
            else:
                ws.cell(row=row, column=1).font = data_font
                ws.cell(row=row, column=4).font = data_font
            
            row += 1
        
        # Column widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 20
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 15
        
        return wb
    
    # ============================================================================
    # TEMPLATE 3: PREMIUM BUDGET TEMPLATE
    # ============================================================================
    def create_budget_template():
        wb = Workbook()
        ws = wb.active
        ws.title = "Budget"
        
        # Title
        ws.merge_cells('A1:H1')
        ws['A1'] = 'ANNUAL BUDGET'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:H2')
        ws['A2'] = '年度预算表'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # Budget info
        ws['A4'] = 'Department:'
        ws['B4'] = '{{department}}'
        ws['D4'] = 'Fiscal Year:'
        ws['E4'] = '{{fiscal_year}}'
        ws['G4'] = 'Prepared By:'
        ws['H4'] = '{{prepared_by}}'
        
        # Budget table headers
        headers = ['Category', 'Q1 Budget', 'Q2 Budget', 'Q3 Budget', 'Q4 Budget', 'Annual Total', 'Actual', 'Variance']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=6, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        # Sample data rows
        categories = ['Personnel', 'Marketing', 'Operations', 'Technology', 'Facilities', 'Travel', 'Training']
        for i, category in enumerate(categories):
            row = 7 + i
            ws.cell(row=row, column=1, value=category).font = Font(name='Microsoft YaHei', size=10, bold=True)
            cat_lower = category.lower()
            for col in range(2, 9):
                if col < 6:
                    placeholder = '{{' + cat_lower + '_q' + str(col-1) + '}}'
                elif col == 6:
                    placeholder = '{{' + cat_lower + '_total}}'
                elif col == 7:
                    placeholder = '{{' + cat_lower + '_actual}}'
                else:
                    placeholder = '{{' + cat_lower + '_variance}}'
                cell = ws.cell(row=row, column=col, value=placeholder)
                cell.font = data_font
                cell.alignment = right_align if col > 1 else left_align
                cell.border = thin_border
                if col == 8:  # Variance column
                    cell.number_format = '0.00%'
        
        # Total row
        total_row = 7 + len(categories)
        ws.cell(row=total_row, column=1, value='TOTAL').font = Font(name='Microsoft YaHei', size=10, bold=True)
        ws.cell(row=total_row, column=1).fill = PatternFill(start_color=ds.PRIMARY['100'], end_color=ds.PRIMARY['100'], fill_type='solid')
        for col in range(2, 9):
            if col < 6:
                placeholder = '{{total_q%d}}' % (col-1)
            elif col == 6:
                placeholder = '{{total_budget}}'
            elif col == 7:
                placeholder = '{{total_actual}}'
            else:
                placeholder = '{{total_variance}}'
            cell = ws.cell(row=total_row, column=col, value=placeholder)
            cell.font = Font(name='Microsoft YaHei', size=10, bold=True)
            cell.fill = PatternFill(start_color=ds.PRIMARY['100'], end_color=ds.PRIMARY['100'], fill_type='solid')
            cell.border = thin_border
            cell.alignment = right_align
        
        # Column widths
        ws.column_dimensions['A'].width = 18
        for col in ['B', 'C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 12
        
        return wb
    
    # ============================================================================
    # TEMPLATE 4: PREMIUM PROJECT TIMELINE
    # ============================================================================
    def create_project_timeline():
        wb = Workbook()
        ws = wb.active
        ws.title = "Project Timeline"
        
        # Title
        ws.merge_cells('A1:H1')
        ws['A1'] = 'PROJECT TIMELINE & GANTT CHART'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:H2')
        ws['A2'] = '项目时间线与甘特图'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # Project info
        ws['A4'] = 'Project:'
        ws['B4'] = '{{project_name}}'
        ws['D4'] = 'Manager:'
        ws['E4'] = '{{project_manager}}'
        ws['G4'] = 'Start Date:'
        ws['H4'] = '{{start_date}}'
        
        # Timeline headers
        headers = ['Task ID', 'Task Name', 'Owner', 'Start Date', 'End Date', 'Duration', 'Status', 'Progress']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=6, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        # Sample tasks
        for i in range(5):
            row = 7 + i
            ws.cell(row=row, column=1, value=f'T{i+1}').alignment = center_align
            ws.cell(row=row, column=2, value='{{task_%d_name}}' % (i+1))
            ws.cell(row=row, column=3, value='{{task_%d_owner}}' % (i+1))
            ws.cell(row=row, column=4, value='{{task_%d_start}}' % (i+1))
            ws.cell(row=row, column=5, value='{{task_%d_end}}' % (i+1))
            ws.cell(row=row, column=6, value='{{task_%d_duration}}' % (i+1))
            ws.cell(row=row, column=7, value='{{task_%d_status}}' % (i+1))
            ws.cell(row=row, column=8, value='{{task_%d_progress}}' % (i+1))
            
            for col in range(1, 9):
                cell = ws.cell(row=row, column=col)
                cell.font = data_font
                cell.border = thin_border
                if col in [1, 3, 4, 5, 6, 7, 8]:
                    cell.alignment = center_align
        
        # Milestones section
        milestone_row = 13
        ws.cell(row=milestone_row, column=1, value='MILESTONES')
        ws.cell(row=milestone_row, column=1).font = subheader_font
        ws.cell(row=milestone_row, column=1).fill = subheader_fill
        ws.merge_cells(start_row=milestone_row, start_column=1, end_row=milestone_row, end_column=8)
        
        milestone_headers = ['Milestone', 'Date', 'Description', 'Status', '', '', '', '']
        for col, header in enumerate(milestone_headers, 1):
            if header:
                cell = ws.cell(row=milestone_row + 1, column=col, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border
        
        for i in range(3):
            row = milestone_row + 2 + i
            ws.cell(row=row, column=1, value='{{milestone_%d_name}}' % (i+1))
            ws.cell(row=row, column=2, value='{{milestone_%d_date}}' % (i+1))
            ws.cell(row=row, column=3, value='{{milestone_%d_desc}}' % (i+1))
            ws.cell(row=row, column=4, value='{{milestone_%d_status}}' % (i+1))
            for col in range(1, 5):
                cell = ws.cell(row=row, column=col)
                cell.font = data_font
                cell.border = thin_border
        
        # Column widths
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 10
        ws.column_dimensions['G'].width = 12
        ws.column_dimensions['H'].width = 10
        
        return wb
    
    # ============================================================================
    # TEMPLATE 5: PREMIUM INVENTORY MANAGEMENT
    # ============================================================================
    def create_inventory_management():
        wb = Workbook()
        ws = wb.active
        ws.title = "Inventory"
        
        # Title
        ws.merge_cells('A1:I1')
        ws['A1'] = 'INVENTORY MANAGEMENT'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:I2')
        ws['A2'] = '库存管理系统'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # Inventory headers
        headers = ['SKU', 'Product Name', 'Category', 'Unit Price', 'Stock Level', 'Min Stock', 'Reorder Qty', 'Location', 'Status']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        # Sample products
        for i in range(8):
            row = 5 + i
            ws.cell(row=row, column=1, value='{{sku_%d}}' % (i+1))
            ws.cell(row=row, column=2, value='{{product_%d_name}}' % (i+1))
            ws.cell(row=row, column=3, value='{{product_%d_category}}' % (i+1))
            ws.cell(row=row, column=4, value='{{product_%d_price}}' % (i+1))
            ws.cell(row=row, column=5, value='{{product_%d_stock}}' % (i+1))
            ws.cell(row=row, column=6, value='{{product_%d_min}}' % (i+1))
            ws.cell(row=row, column=7, value='{{product_%d_reorder}}' % (i+1))
            ws.cell(row=row, column=8, value='{{product_%d_location}}' % (i+1))
            ws.cell(row=row, column=9, value='{{product_%d_status}}' % (i+1))
            
            for col in range(1, 10):
                cell = ws.cell(row=row, column=col)
                cell.font = data_font
                cell.border = thin_border
                if col in [1, 3, 5, 6, 7, 8, 9]:
                    cell.alignment = center_align
                if col == 4:
                    cell.number_format = '$#,##0.00'
        
        # Summary section
        summary_row = 14
        ws.cell(row=summary_row, column=1, value='INVENTORY SUMMARY')
        ws.cell(row=summary_row, column=1).font = subheader_font
        ws.cell(row=summary_row, column=1).fill = subheader_fill
        ws.merge_cells(start_row=summary_row, start_column=1, end_row=summary_row, end_column=4)
        
        summary_items = [
            ('Total SKUs', '{{total_skus}}'),
            ('Total Value', '{{total_inventory_value}}'),
            ('Low Stock Items', '{{low_stock_count}}'),
            ('Out of Stock', '{{out_of_stock_count}}'),
        ]
        
        for i, (label, value) in enumerate(summary_items):
            row = summary_row + 1 + i
            ws.cell(row=row, column=1, value=label).font = Font(name='Microsoft YaHei', size=10, bold=True)
            ws.cell(row=row, column=2, value=value).font = data_font
        
        # Column widths
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 10
        ws.column_dimensions['G'].width = 12
        ws.column_dimensions['H'].width = 12
        ws.column_dimensions['I'].width = 12
        
        return wb
    
    # ============================================================================
    # TEMPLATE 6: PREMIUM CRM SIMPLE
    # ============================================================================
    def create_crm_simple():
        wb = Workbook()
        ws = wb.active
        ws.title = "CRM"
        
        # Title
        ws.merge_cells('A1:H1')
        ws['A1'] = 'CUSTOMER RELATIONSHIP MANAGEMENT'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:H2')
        ws['A2'] = '客户关系管理系统'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # CRM headers
        headers = ['Customer ID', 'Company', 'Contact', 'Email', 'Phone', 'Status', 'Last Contact', 'Next Action']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        # Sample customers
        for i in range(6):
            row = 5 + i
            ws.cell(row=row, column=1, value='{{customer_%d_id}}' % (i+1))
            ws.cell(row=row, column=2, value='{{customer_%d_company}}' % (i+1))
            ws.cell(row=row, column=3, value='{{customer_%d_contact}}' % (i+1))
            ws.cell(row=row, column=4, value='{{customer_%d_email}}' % (i+1))
            ws.cell(row=row, column=5, value='{{customer_%d_phone}}' % (i+1))
            ws.cell(row=row, column=6, value='{{customer_%d_status}}' % (i+1))
            ws.cell(row=row, column=7, value='{{customer_%d_last_contact}}' % (i+1))
            ws.cell(row=row, column=8, value='{{customer_%d_next_action}}' % (i+1))
            
            for col in range(1, 9):
                cell = ws.cell(row=row, column=col)
                cell.font = data_font
                cell.border = thin_border
                if col in [1, 6, 7]:
                    cell.alignment = center_align
        
        # Pipeline summary
        pipeline_row = 12
        ws.cell(row=pipeline_row, column=1, value='SALES PIPELINE')
        ws.cell(row=pipeline_row, column=1).font = subheader_font
        ws.cell(row=pipeline_row, column=1).fill = subheader_fill
        ws.merge_cells(start_row=pipeline_row, start_column=1, end_row=pipeline_row, end_column=4)
        
        pipeline_headers = ['Stage', 'Count', 'Value', 'Conversion']
        for col, header in enumerate(pipeline_headers, 1):
            cell = ws.cell(row=pipeline_row + 1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        stages = ['Lead', 'Qualified', 'Proposal', 'Negotiation', 'Closed Won']
        stage_keys = ['lead', 'qualified', 'proposal', 'negotiation', 'closed_won']
        for i, stage in enumerate(stages):
            row = pipeline_row + 2 + i
            key = stage_keys[i]
            ws.cell(row=row, column=1, value=stage).font = Font(name='Microsoft YaHei', size=10, bold=True)
            ws.cell(row=row, column=2, value='{{%s_count}}' % key)
            ws.cell(row=row, column=3, value='{{%s_value}}' % key)
            ws.cell(row=row, column=4, value='{{%s_conversion}}' % key)
            for col in range(1, 5):
                cell = ws.cell(row=row, column=col)
                cell.border = thin_border
                if col > 1:
                    cell.alignment = center_align
        
        # Column widths
        ws.column_dimensions['A'].width = 14
        ws.column_dimensions['B'].width = 22
        ws.column_dimensions['C'].width = 18
        ws.column_dimensions['D'].width = 25
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 14
        ws.column_dimensions['H'].width = 20
        
        return wb
    
    # ============================================================================
    # TEMPLATE 7: PREMIUM ATTENDANCE TRACKING
    # ============================================================================
    def create_attendance_tracking():
        wb = Workbook()
        ws = wb.active
        ws.title = "Attendance"
        
        # Title
        ws.merge_cells('A1:AH1')
        ws['A1'] = 'EMPLOYEE ATTENDANCE TRACKER'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:AH2')
        ws['A2'] = '员工考勤记录表 - {{month_year}}'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # Employee info header
        ws['A4'] = 'Employee ID'
        ws['B4'] = 'Name'
        ws['C4'] = 'Department'
        ws['A4'].font = header_font
        ws['B4'].font = header_font
        ws['C4'].font = header_font
        ws['A4'].fill = header_fill
        ws['B4'].fill = header_fill
        ws['C4'].fill = header_fill
        
        # Day headers (1-31)
        for day in range(1, 32):
            cell = ws.cell(row=4, column=day + 3, value=day)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
        
        # Summary columns
        summary_cols = ['Present', 'Absent', 'Late', 'Leave', 'Total']
        for i, col_name in enumerate(summary_cols):
            cell = ws.cell(row=4, column=35 + i, value=col_name)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
        
        # Sample employees
        for i in range(5):
            row = 5 + i
            ws.cell(row=row, column=1, value='{{emp_%d_id}}' % (i+1))
            ws.cell(row=row, column=2, value='{{emp_%d_name}}' % (i+1))
            ws.cell(row=row, column=3, value='{{emp_%d_dept}}' % (i+1))
            
            # Attendance cells
            for day in range(1, 32):
                cell = ws.cell(row=row, column=day + 3, value='{{emp_%d_day_%d}}' % (i+1, day))
                cell.alignment = center_align
                cell.border = thin_border
                cell.font = Font(name='Microsoft YaHei', size=9)
            
            # Summary cells
            ws.cell(row=row, column=35, value='{{emp_%d_present}}' % (i+1))
            ws.cell(row=row, column=36, value='{{emp_%d_absent}}' % (i+1))
            ws.cell(row=row, column=37, value='{{emp_%d_late}}' % (i+1))
            ws.cell(row=row, column=38, value='{{emp_%d_leave}}' % (i+1))
            ws.cell(row=row, column=39, value='{{emp_%d_total}}' % (i+1))
            
            for col in range(35, 40):
                cell = ws.cell(row=row, column=col)
                cell.alignment = center_align
                cell.border = thin_border
                cell.font = data_font
        
        # Legend
        legend_row = 11
        ws.cell(row=legend_row, column=1, value='Legend: P=Present, A=Absent, L=Late, V=Vacation, H=Holiday')
        ws.cell(row=legend_row, column=1).font = Font(name='Microsoft YaHei', size=9, italic=True, color=ds.SECONDARY['600'])
        
        # Column widths
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 18
        ws.column_dimensions['C'].width = 15
        for col in range(4, 35):
            ws.column_dimensions[get_column_letter(col)].width = 4
        for col in range(35, 40):
            ws.column_dimensions[get_column_letter(col)].width = 8
        
        return wb
    
    # ============================================================================
    # TEMPLATE 8: PREMIUM PIVOT DEMO
    # ============================================================================
    def create_pivot_demo():
        wb = Workbook()
        ws = wb.active
        ws.title = "Data Analysis"
        
        # Title
        ws.merge_cells('A1:F1')
        ws['A1'] = 'DATA ANALYSIS & PIVOT TABLE'
        ws['A1'].font = Font(name='Microsoft YaHei', size=18, bold=True, color=ds.PRIMARY['900'])
        ws['A1'].alignment = center_align
        
        ws.merge_cells('A2:F2')
        ws['A2'] = '数据分析与透视表示例'
        ws['A2'].font = Font(name='Microsoft YaHei', size=12, color=ds.PRIMARY['600'])
        ws['A2'].alignment = center_align
        
        # Raw data section
        ws['A4'] = 'RAW DATA / 原始数据'
        ws['A4'].font = subheader_font
        ws['A4'].fill = subheader_fill
        ws.merge_cells('A4:F4')
        
        # Data headers
        data_headers = ['Date', 'Region', 'Product', 'Category', 'Sales', 'Quantity']
        for col, header in enumerate(data_headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        # Sample data rows
        for i in range(10):
            row = 6 + i
            ws.cell(row=row, column=1, value='{{data_%d_date}}' % (i+1))
            ws.cell(row=row, column=2, value='{{data_%d_region}}' % (i+1))
            ws.cell(row=row, column=3, value='{{data_%d_product}}' % (i+1))
            ws.cell(row=row, column=4, value='{{data_%d_category}}' % (i+1))
            ws.cell(row=row, column=5, value='{{data_%d_sales}}' % (i+1))
            ws.cell(row=row, column=6, value='{{data_%d_quantity}}' % (i+1))
            
            for col in range(1, 7):
                cell = ws.cell(row=row, column=col)
                cell.font = data_font
                cell.border = thin_border
                if col in [1, 2, 4]:
                    cell.alignment = center_align
                if col == 5:
                    cell.number_format = '$#,##0.00'
        
        # Pivot summary section
        pivot_row = 18
        ws.cell(row=pivot_row, column=1, value='PIVOT SUMMARY / 透视汇总')
        ws.cell(row=pivot_row, column=1).font = subheader_font
        ws.cell(row=pivot_row, column=1).fill = subheader_fill
        ws.merge_cells(start_row=pivot_row, start_column=1, end_row=pivot_row, end_column=4)
        
        # By Region
        ws.cell(row=pivot_row + 1, column=1, value='By Region:')
        ws.cell(row=pivot_row + 1, column=1).font = Font(name='Microsoft YaHei', size=10, bold=True)
        
        region_headers = ['Region', 'Total Sales', 'Total Qty', 'Avg Sale']
        for col, header in enumerate(region_headers, 1):
            cell = ws.cell(row=pivot_row + 2, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        for i in range(4):
            row = pivot_row + 3 + i
            ws.cell(row=row, column=1, value='{{region_%d_name}}' % (i+1))
            ws.cell(row=row, column=2, value='{{region_%d_sales}}' % (i+1))
            ws.cell(row=row, column=3, value='{{region_%d_qty}}' % (i+1))
            ws.cell(row=row, column=4, value='{{region_%d_avg}}' % (i+1))
            for col in range(1, 5):
                cell = ws.cell(row=row, column=col)
                cell.border = thin_border
                cell.font = data_font
                if col > 1:
                    cell.alignment = center_align
        
        # By Category
        cat_row = pivot_row + 8
        ws.cell(row=cat_row, column=1, value='By Category:')
        ws.cell(row=cat_row, column=1).font = Font(name='Microsoft YaHei', size=10, bold=True)
        
        cat_headers = ['Category', 'Total Sales', 'Total Qty', 'Market Share']
        for col, header in enumerate(cat_headers, 1):
            cell = ws.cell(row=cat_row + 1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        for i in range(3):
            row = cat_row + 2 + i
            ws.cell(row=row, column=1, value='{{category_%d_name}}' % (i+1))
            ws.cell(row=row, column=2, value='{{category_%d_sales}}' % (i+1))
            ws.cell(row=row, column=3, value='{{category_%d_qty}}' % (i+1))
            ws.cell(row=row, column=4, value='{{category_%d_share}}' % (i+1))
            for col in range(1, 5):
                cell = ws.cell(row=row, column=col)
                cell.border = thin_border
                cell.font = data_font
                if col > 1:
                    cell.alignment = center_align
        
        # Column widths
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 12
        
        return wb
    
    # Generate all templates
    templates = [
        ('sales-report.xlsx', create_sales_report),
        ('financial-statement.xlsx', create_financial_statement),
        ('budget-template.xlsx', create_budget_template),
        ('project-timeline.xlsx', create_project_timeline),
        ('inventory-management.xlsx', create_inventory_management),
        ('crm-simple.xlsx', create_crm_simple),
        ('attendance-tracking.xlsx', create_attendance_tracking),
        ('pivot-demo.xlsx', create_pivot_demo),
    ]
    
    for filename, create_func in templates:
        print(f"  [Excel] {filename}...", end=" ")
        try:
            wb = create_func()
            filepath = template_dir / filename
            wb.save(str(filepath))
            print("OK")
        except Exception as e:
            print(f"FAILED: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    print(f"  Generated {len(templates)} premium Excel templates")
    return True


# ============================================================================
# MAIN
# ============================================================================

def main():
    """Main entry point"""
    print("\n" + "=" * 70)
    print("Office Pro - Premium Template Generator")
    print("Enterprise Edition with Professional Design")
    print("=" * 70 + "\n")
    
    success = True
    
    # Generate Word templates
    if not generate_word_templates():
        success = False
    
    # Generate Excel templates
    if not generate_excel_templates():
        success = False
    
    print("\n" + "=" * 70)
    if success:
        print("All premium templates generated successfully!")
        print(f"Templates location: {SKILL_DIR / 'assets' / 'templates'}")
    else:
        print("Some templates failed to generate. Check errors above.")
    print("=" * 70 + "\n")
    
    return success


if __name__ == '__main__':
    sys.exit(0 if main() else 1)

"""
Office Pro - Enterprise Document Automation Suite
Word Processor Module

Enterprise-grade Word document processing based on python-docx and docxtpl
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from core import (
    DocumentProcessor, require_document, require_template,
    DependencyError, DocumentNotLoadedError, TemplateNotFoundError,
    TemplateRenderError
)

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocumentType = None

try:
    from docxtpl import DocxTemplate, InlineImage, RichText
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False


class WordProcessor(DocumentProcessor):
    """
    Enterprise-grade Word document processor
    
    Features:
    - Document creation and editing
    - Template rendering (Jinja2)
    - Style management
    - Tables and images
    - Headers and footers
    """
    
    _document_type_name = "Word document"
    
    def __init__(self, template_dir: Optional[str] = None, sandboxed: bool = False):
        """
        Initialize Word processor
        
        Args:
            template_dir: Template directory path
            sandboxed: Enable sandboxed template rendering for security
        """
        if not DOCX_AVAILABLE:
            raise DependencyError(
                "python-docx is required. Install with: pip install python-docx"
            )
        
        self._sandboxed = sandboxed
        super().__init__(template_dir)
    
    def _get_default_template_dir(self) -> str:
        """Get default template directory"""
        skill_root = Path(__file__).parent
        templates_dir = skill_root / "assets" / "templates" / "word"
        return str(templates_dir)
    
    @require_document
    def save(self, path: str) -> None:
        """
        Save document to specified path
        
        Args:
            path: Save path
        """
        self._ensure_output_dir(path)
        self._document.save(path)
    
    def create_document(self) -> DocumentType:
        """
        Create new blank document
        
        Returns:
            Document object
        """
        self._document = Document()
        self._template = None
        return self._document
    
    def load_document(self, path: str) -> DocumentType:
        """
        Load existing document
        
        Args:
            path: Document path
            
        Returns:
            Document object
        """
        self._document = Document(path)
        self._template = None
        return self._document
    
    def load_template(self, template_name: str) -> 'DocxTemplate':
        """
        Load template file
        
        Args:
            template_name: Template filename (e.g. 'meeting-minutes.docx')
            
        Returns:
            DocxTemplate object
        """
        if not DOCXTPL_AVAILABLE:
            raise DependencyError(
                "docxtpl is required for template rendering. "
                "Install with: pip install docxtpl"
            )
        
        template_path = self._validate_template_path(template_name)
        self._template = DocxTemplate(str(template_path))
        self._document = self._template.docx
        return self._template
    
    @require_template
    def render_template(self, context: Dict[str, Any]) -> DocumentType:
        """
        Render template with context data
        
        Args:
            context: Template variable dictionary
            
        Returns:
            Rendered Document object
        """
        try:
            if self._sandboxed:
                self._render_sandboxed(context)
            else:
                self._template.render(context)
            
            self._document = self._template.docx
            return self._document
        except Exception as e:
            raise TemplateRenderError(f"Template rendering failed: {e}")
    
    def _render_sandboxed(self, context: Dict[str, Any]) -> None:
        """
        Render template in sandboxed mode for security
        
        Args:
            context: Template variable dictionary
        """
        from jinja2 import SandboxedEnvironment
        
        env = SandboxedEnvironment()
        env.globals.update({
            'now': datetime.now,
            'today': datetime.today,
        })
        
        self._template.render(context, jinja_env=env)
    
    @require_document
    def add_heading(self, text: str, level: int = 1) -> Any:
        """
        Add heading
        
        Args:
            text: Heading text
            level: Heading level (1-9)
            
        Returns:
            Paragraph object
        """
        return self._document.add_heading(text, level=level)
    
    @require_document
    def add_paragraph(self, text: str = "", style: Optional[str] = None) -> Any:
        """
        Add paragraph
        
        Args:
            text: Paragraph text
            style: Style name
            
        Returns:
            Paragraph object
        """
        return self._document.add_paragraph(text, style=style)
    
    @require_document
    def add_table(self, rows: int, cols: int, style: Optional[str] = None) -> Any:
        """
        Add table
        
        Args:
            rows: Number of rows
            cols: Number of columns
            style: Table style name
            
        Returns:
            Table object
        """
        table = self._document.add_table(rows=rows, cols=cols)
        if style:
            table.style = style
        return table
    
    @require_document
    def add_picture(self, image_path: str, width: Optional[float] = None, 
                    height: Optional[float] = None) -> Any:
        """
        Add picture
        
        Args:
            image_path: Image path
            width: Width in inches
            height: Height in inches
            
        Returns:
            InlineShape object
        """
        kwargs = {}
        if width:
            kwargs['width'] = Inches(width)
        if height:
            kwargs['height'] = Inches(height)
        
        return self._document.add_picture(image_path, **kwargs)
    
    @require_document
    def add_page_break(self) -> None:
        """Add page break"""
        self._document.add_page_break()
    
    @require_document
    def add_header(self, text: str, align: str = "center") -> None:
        """
        Add header
        
        Args:
            text: Header text
            align: Alignment (left/center/right)
        """
        section = self._document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = text
        
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    @require_document
    def add_footer(self, text: str, align: str = "center") -> None:
        """
        Add footer
        
        Args:
            text: Footer text
            align: Alignment (left/center/right)
        """
        section = self._document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = text
        
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    @require_document
    def add_page_number(self, location: str = "footer", align: str = "center") -> None:
        """
        Add page number
        
        Args:
            location: Position (header/footer)
            align: Alignment (left/center/right)
        """
        section = self._document.sections[0]
        
        if location == "header":
            target = section.header
        else:
            target = section.footer
        
        if target.paragraphs:
            paragraph = target.paragraphs[0]
        else:
            paragraph = target.add_paragraph()
        
        paragraph.clear()
        
        run = paragraph.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    @require_document
    def get_document_info(self) -> Dict[str, Any]:
        """
        Get document information
        
        Returns:
            Document info dictionary
        """
        core_props = self._document.core_properties
        
        return {
            'title': core_props.title,
            'author': core_props.author,
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'created': core_props.created,
            'modified': core_props.modified,
            'paragraph_count': len(self._document.paragraphs),
            'table_count': len(self._document.tables),
        }

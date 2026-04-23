"""
Office Pro - Simplified API

Quick document generation API designed for AI Agents.
Provides one-click document generation with built-in templates.
"""

from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime


def num_to_chinese(num: int) -> str:
    """Convert number to Chinese characters"""
    chinese_nums = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    chinese_units = ['', '拾', '佰', '仟', '万', '拾', '佰', '仟', '亿']
    
    if num == 0:
        return chinese_nums[0]
    
    result = []
    num_str = str(num)
    length = len(num_str)
    
    for i, digit in enumerate(num_str):
        d = int(digit)
        unit_idx = length - i - 1
        
        if d != 0:
            result.append(chinese_nums[d])
            result.append(chinese_units[unit_idx])
        else:
            if result and result[-1] != chinese_nums[0]:
                result.append(chinese_nums[0])
    
    # Remove trailing zeros
    while result and result[-1] == chinese_nums[0]:
        result.pop()
    
    return ''.join(result) if result else chinese_nums[0]


class QuickGenerator:
    """
    Quick document generator for AI Agents.
    
    Provides one-click generation of contracts, reports, and other documents
    with built-in templates and content.
    """
    
    def __init__(self, output_dir: str = "./output"):
        """
        Initialize the generator.
        
        Args:
            output_dir: Directory for output files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def generate(self, doc_type: str, data: Dict[str, Any], 
                 output_filename: Optional[str] = None) -> str:
        """
        Generate a document.
        
        Args:
            doc_type: Document type, e.g., 'contract.parking_lease'
            data: Document data dictionary
            output_filename: Output filename (optional)
        
        Returns:
            Path to the generated file
        """
        parts = doc_type.split('.')
        if len(parts) != 2:
            raise ValueError(f"Invalid document type format: {doc_type}. Expected 'category.name'")
        
        category, template_name = parts
        
        if category == 'contract':
            return self._generate_contract(template_name, data, output_filename)
        elif category == 'report':
            return self._generate_report(template_name, data, output_filename)
        elif category == 'excel':
            return self._generate_excel(template_name, data, output_filename)
        else:
            raise ValueError(f"Unknown document category: {category}")
    
    def _generate_contract(self, template_name: str, data: Dict[str, Any],
                           output_filename: Optional[str] = None) -> str:
        """Generate a contract document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        # Import contract templates
        from templates.contracts import CONTRACT_TEMPLATES
        
        if template_name not in CONTRACT_TEMPLATES:
            available = ', '.join(CONTRACT_TEMPLATES.keys())
            raise ValueError(f"Unknown contract type: {template_name}. Available: {available}")
        
        template = CONTRACT_TEMPLATES[template_name]
        
        # Merge default data with user-provided data
        merged_data = template.get('default_data', {}).copy()
        merged_data.update(data)
        
        # Add computed fields
        self._add_computed_fields(merged_data)
        
        # Create document
        doc = Document()
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # Add title
        title = doc.add_heading(template['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(22)
            run.font.bold = True
        
        # Add contract number and date
        contract_number = merged_data.get('contract_number', '________')
        sign_date = merged_data.get('sign_date', datetime.now().strftime('%Y年%m月%d日'))
        
        p = doc.add_paragraph()
        p.add_run(f'合同编号：{contract_number}').font.name = 'Microsoft YaHei'
        p = doc.add_paragraph()
        p.add_run(f'签署日期：{sign_date}').font.name = 'Microsoft YaHei'
        doc.add_paragraph()
        
        # Add parties section
        doc.add_heading('合同双方', level=1)
        party_a = merged_data.get('party_a', '________')
        party_b = merged_data.get('party_b', '________')
        
        p = doc.add_paragraph()
        p.add_run(f'甲方（出租方）：{party_a}').font.name = 'Microsoft YaHei'
        p = doc.add_paragraph()
        p.add_run(f'乙方（承租方）：{party_b}').font.name = 'Microsoft YaHei'
        doc.add_paragraph()
        
        # Add sections
        for section in template['sections']:
            heading = doc.add_heading(section['title'], level=1)
            for run in heading.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(14)
                run.font.bold = True
            
            try:
                content = section['content'].format(**merged_data)
            except KeyError as e:
                # If a field is missing, keep the placeholder
                content = section['content']
            
            # Handle multi-line content
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Inches(0.3)
                    for run in p.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(11)
            
            doc.add_paragraph()
        
        # Add signature section
        doc.add_heading('签署', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(3)
        for cell in table.columns[1].cells:
            cell.width = Inches(3)
        
        # Header row
        table.cell(0, 0).text = '甲方（盖章）：'
        table.cell(0, 1).text = '乙方（盖章）：'
        
        # Signature row with space
        table.cell(1, 0).text = '\n\n'
        table.cell(1, 1).text = '\n\n'
        
        # Name row
        party_a_signatory = merged_data.get('party_a_signatory', '')
        party_b_signatory = merged_data.get('party_b_signatory', '')
        table.cell(2, 0).text = f'签署人：{party_a_signatory}'
        table.cell(2, 1).text = f'签署人：{party_b_signatory}'
        
        # Style the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(11)
        
        # Save document
        if not output_filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{template_name}_contract_{timestamp}.docx"
        
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _generate_excel(self, template_name: str, data: Dict[str, Any],
                        output_filename: Optional[str] = None) -> str:
        """Generate an Excel spreadsheet."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("openpyxl is required. Install with: pip install openpyxl")
        
        from templates.excel import EXCEL_TEMPLATES
        
        if template_name not in EXCEL_TEMPLATES:
            available = ', '.join(EXCEL_TEMPLATES.keys())
            raise ValueError(f"Unknown Excel type: {template_name}. Available: {available}")
        
        template = EXCEL_TEMPLATES[template_name]
        
        # Merge default data with user-provided data
        merged_data = template.get('default_data', {}).copy()
        merged_data.update(data)
        
        # Create workbook
        wb = Workbook()
        wb.remove(wb.active)
        
        # Define styles
        header_font = Font(bold=True, size=11, name='Microsoft YaHei')
        normal_font = Font(size=10, name='Microsoft YaHei')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font_white = Font(bold=True, size=11, color='FFFFFF', name='Microsoft YaHei')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        center_align = Alignment(horizontal='center', vertical='center')
        
        # Process each sheet
        for sheet_name, sheet_config in template['sheets'].items():
            ws = wb.create_sheet(title=sheet_config['title'])
            
            # Add title row
            title = template['title']
            if sheet_config['title'] != title:
                title = f"{title} - {sheet_config['title']}"
            
            ws.merge_cells(start_row=1, start_column=1, end_row=1, 
                          end_column=len(sheet_config['columns']))
            title_cell = ws.cell(row=1, column=1, value=title)
            title_cell.font = Font(bold=True, size=14, name='Microsoft YaHei')
            title_cell.alignment = center_align
            
            # Add header row
            for col_idx, col_config in enumerate(sheet_config['columns'], start=1):
                cell = ws.cell(row=2, column=col_idx, value=col_config['name'])
                cell.font = header_font_white
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = center_align
                
                # Set column width
                width = col_config.get('width', 12)
                ws.column_dimensions[get_column_letter(col_idx)].width = width
            
            # Add data rows
            rows_data = data.get('sheets', {}).get(sheet_name, {}).get('rows', sheet_config.get('rows', []))
            
            for row_idx, row_data in enumerate(rows_data, start=3):
                if isinstance(row_data, dict):
                    for col_idx, col_config in enumerate(sheet_config['columns'], start=1):
                        col_name = col_config['name']
                        value = row_data.get(col_name, '')
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.font = normal_font
                        cell.border = thin_border
                elif isinstance(row_data, (list, tuple)):
                    for col_idx, value in enumerate(row_data, start=1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.font = normal_font
                        cell.border = thin_border
            
            # Add empty rows for data entry (if no data provided)
            if not rows_data:
                for row_idx in range(3, 23):  # Add 20 empty rows
                    for col_idx in range(1, len(sheet_config['columns']) + 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value='')
                        cell.border = thin_border
                        cell.font = normal_font
        
        # Save workbook
        if not output_filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{template_name}_{timestamp}.xlsx"
        
        output_path = self.output_dir / output_filename
        wb.save(str(output_path))
        
        return str(output_path)
    
    def _generate_report(self, template_name: str, data: Dict[str, Any],
                         output_filename: Optional[str] = None) -> str:
        """Generate a report document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        from templates.reports import REPORT_TEMPLATES
        
        if template_name not in REPORT_TEMPLATES:
            available = ', '.join(REPORT_TEMPLATES.keys())
            raise ValueError(f"Unknown report type: {template_name}. Available: {available}")
        
        template = REPORT_TEMPLATES[template_name]
        
        # Merge default data with user-provided data
        merged_data = template.get('default_data', {}).copy()
        merged_data.update(data)
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # Add title
        title = doc.add_heading(template['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(20)
            run.font.bold = True
        
        # Add metadata
        if 'date' in merged_data:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"日期：{merged_data['date']}").font.name = 'Microsoft YaHei'
        
        doc.add_paragraph()
        
        # Add sections
        for section in template['sections']:
            heading = doc.add_heading(section['title'], level=1)
            for run in heading.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(14)
                run.font.bold = True
            
            try:
                content = section['content'].format(**merged_data)
            except KeyError:
                content = section['content']
            
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(11)
            
            doc.add_paragraph()
        
        # Save document
        if not output_filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{template_name}_report_{timestamp}.docx"
        
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _add_computed_fields(self, data: Dict[str, Any]) -> None:
        """Add computed fields to the data dictionary."""
        # Convert monthly rent to Chinese
        if 'monthly_rent' in data and 'monthly_rent_chinese' not in data:
            try:
                rent = int(data['monthly_rent'])
                data['monthly_rent_chinese'] = num_to_chinese(rent) + '元整'
            except (ValueError, TypeError):
                data['monthly_rent_chinese'] = '________'
        
        # Convert deposit to Chinese
        if 'deposit' in data and 'deposit_chinese' not in data:
            try:
                deposit = int(data['deposit'])
                data['deposit_chinese'] = num_to_chinese(deposit) + '元整'
            except (ValueError, TypeError):
                data['deposit_chinese'] = '________'
        
        # Calculate duration if start and end dates provided
        if 'start_date' in data and 'end_date' in data and 'duration' not in data:
            try:
                start = datetime.strptime(data['start_date'], '%Y-%m-%d')
                end = datetime.strptime(data['end_date'], '%Y-%m-%d')
                months = (end.year - start.year) * 12 + (end.month - start.month)
                data['duration'] = max(1, months)
            except (ValueError, TypeError):
                data['duration'] = '________'


def generate_contract(contract_type: str, output: Optional[str] = None, **kwargs) -> str:
    """
    Generate a contract document.
    
    This is a convenience function for quick contract generation.
    
    Args:
        contract_type: Type of contract
            - 'parking_lease' - Parking space lease contract
            - 'house_lease' - House lease contract
            - 'labor' - Labor contract
        output: Output filename (optional)
        **kwargs: Contract-specific parameters
    
    Returns:
        Path to the generated file
    
    Example:
        >>> generate_contract('parking_lease', 
        ...                   party_a='Zhang San',
        ...                   party_b='Li Si',
        ...                   location='XX Community',
        ...                   monthly_rent=500)
    """
    generator = QuickGenerator()
    doc_type = f"contract.{contract_type}"
    return generator.generate(doc_type, kwargs, output)


def generate_report(report_type: str, output: Optional[str] = None, **kwargs) -> str:
    """
    Generate a report document.
    
    Args:
        report_type: Type of report
            - 'meeting_minutes' - Meeting minutes
            - 'work_report' - Work report
        output: Output filename (optional)
        **kwargs: Report data
    
    Returns:
        Path to the generated file
    """
    generator = QuickGenerator()
    doc_type = f"report.{report_type}"
    return generator.generate(doc_type, kwargs, output)


def generate_excel(excel_type: str, output: Optional[str] = None, **kwargs) -> str:
    """
    Generate an Excel spreadsheet.
    
    Args:
        excel_type: Type of spreadsheet
            - 'financial_report' - Financial report
            - 'project_schedule' - Project schedule
            - 'employee_roster' - Employee roster
            - 'asset_inventory' - Asset inventory
            - 'expense_report' - Expense report
            - 'invoice' - Invoice management
        output: Output filename (optional)
        **kwargs: Spreadsheet data
    
    Returns:
        Path to the generated file
    """
    generator = QuickGenerator()
    doc_type = f"excel.{excel_type}"
    return generator.generate(doc_type, kwargs, output)


def list_templates(category: Optional[str] = None) -> Dict[str, List[str]]:
    """
    List available templates.
    
    Args:
        category: Template category ('contract', 'report', 'excel', or None for all)
    
    Returns:
        Dictionary of available templates
    """
    result = {}
    
    if category is None or category == 'contract':
        try:
            from templates.contracts import CONTRACT_TEMPLATES
            result['contracts'] = list(CONTRACT_TEMPLATES.keys())
        except ImportError:
            result['contracts'] = []
    
    if category is None or category == 'report':
        try:
            from templates.reports import REPORT_TEMPLATES
            result['reports'] = list(REPORT_TEMPLATES.keys())
        except ImportError:
            result['reports'] = []
    
    if category is None or category == 'excel':
        try:
            from templates.excel import EXCEL_TEMPLATES
            result['excel'] = list(EXCEL_TEMPLATES.keys())
        except ImportError:
            result['excel'] = []
    
    return result


class ExcelChartBuilder:
    """
    Optional chart builder for Excel spreadsheets.
    Requires openpyxl with chart support.
    
    Example:
        >>> builder = ExcelChartBuilder('report.xlsx')
        >>> builder.add_bar_chart('Sheet1', 'A1:B10', 'Sales Data')
        >>> builder.save()
    """
    
    def __init__(self, filepath: str):
        """
        Initialize chart builder.
        
        Args:
            filepath: Path to Excel file
        """
        try:
            from openpyxl import load_workbook
            from openpyxl.chart import BarChart, LineChart, PieChart, AreaChart
            from openpyxl.chart.series import DataPoint
            from openpyxl.chart.label import DataLabelList
            self._load_workbook = load_workbook
            self._BarChart = BarChart
            self._LineChart = LineChart
            self._PieChart = PieChart
            self._AreaChart = AreaChart
            self._available = True
        except ImportError:
            self._available = False
            raise ImportError("openpyxl is required for chart support. Install with: pip install openpyxl")
        
        self.filepath = filepath
        self.wb = load_workbook(filepath)
    
    def add_bar_chart(self, sheet_name: str, data_range: str, title: str = "",
                      x_axis: str = "", y_axis: str = "",
                      position: str = "E5") -> 'ExcelChartBuilder':
        """
        Add a bar chart.
        
        Args:
            sheet_name: Name of the sheet
            data_range: Data range (e.g., 'A1:B10')
            title: Chart title
            x_axis: X-axis label
            y_axis: Y-axis label
            position: Chart position on sheet
        
        Returns:
            Self for chaining
        """
        ws = self.wb[sheet_name]
        chart = self._BarChart()
        chart.title = title
        chart.x_axis.title = x_axis
        chart.y_axis.title = y_axis
        
        data = ws[data_range]
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(data)
        
        ws.add_chart(chart, position)
        return self
    
    def add_line_chart(self, sheet_name: str, data_range: str, title: str = "",
                       x_axis: str = "", y_axis: str = "",
                       position: str = "E5") -> 'ExcelChartBuilder':
        """Add a line chart."""
        ws = self.wb[sheet_name]
        chart = self._LineChart()
        chart.title = title
        chart.x_axis.title = x_axis
        chart.y_axis.title = y_axis
        
        data = ws[data_range]
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(data)
        
        ws.add_chart(chart, position)
        return self
    
    def add_pie_chart(self, sheet_name: str, data_range: str, title: str = "",
                      position: str = "E5") -> 'ExcelChartBuilder':
        """Add a pie chart."""
        ws = self.wb[sheet_name]
        chart = self._PieChart()
        chart.title = title
        
        data = ws[data_range]
        chart.add_data(data, titles_from_data=True)
        
        ws.add_chart(chart, position)
        return self
    
    def save(self, filepath: Optional[str] = None) -> str:
        """
        Save the workbook.
        
        Args:
            filepath: Output path (optional, defaults to original)
        
        Returns:
            Saved file path
        """
        path = filepath or self.filepath
        self.wb.save(path)
        return path
    
    def close(self):
        """Close the workbook."""
        self.wb.close()


class WordStyleBuilder:
    """
    Optional advanced styling for Word documents.
    Requires python-docx.
    
    Example:
        >>> builder = WordStyleBuilder()
        >>> builder.add_styled_heading('Report', color='1F4E79')
        >>> builder.add_styled_paragraph('Content', bold=True)
        >>> builder.save('output.docx')
    """
    
    def __init__(self):
        """Initialize style builder."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            self._Document = Document
            self._Pt = Pt
            self._RGBColor = RGBColor
            self._Inches = Inches
            self._WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self._qn = qn
            self._OxmlElement = OxmlElement
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        self.doc = Document()
    
    def add_styled_heading(self, text: str, level: int = 1,
                           color: str = '1F4E79',
                           font_size: Optional[int] = None) -> 'WordStyleBuilder':
        """
        Add a styled heading.
        
        Args:
            text: Heading text
            level: Heading level (1-9)
            color: Hex color code (e.g., 'FF0000' for red)
            font_size: Font size in points (optional)
        
        Returns:
            Self for chaining
        """
        heading = self.doc.add_heading(text, level=level)
        
        for run in heading.runs:
            run.font.color.rgb = self._RGBColor.from_string(color)
            if font_size:
                run.font.size = self._Pt(font_size)
        
        return self
    
    def add_styled_paragraph(self, text: str,
                             bold: bool = False,
                             italic: bool = False,
                             color: Optional[str] = None,
                             font_size: Optional[int] = None,
                             alignment: str = 'left',
                             first_line_indent: Optional[float] = None) -> 'WordStyleBuilder':
        """
        Add a styled paragraph.
        
        Args:
            text: Paragraph text
            bold: Bold text
            italic: Italic text
            color: Hex color code
            font_size: Font size in points
            alignment: 'left', 'center', 'right', 'justify'
            first_line_indent: First line indent in inches
        
        Returns:
            Self for chaining
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        
        run.bold = bold
        run.italic = italic
        
        if color:
            run.font.color.rgb = self._RGBColor.from_string(color)
        
        if font_size:
            run.font.size = self._Pt(font_size)
        
        align_map = {
            'left': self._WD_ALIGN_PARAGRAPH.LEFT,
            'center': self._WD_ALIGN_PARAGRAPH.CENTER,
            'right': self._WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': self._WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        p.alignment = align_map.get(alignment, self._WD_ALIGN_PARAGRAPH.LEFT)
        
        if first_line_indent:
            p.paragraph_format.first_line_indent = self._Inches(first_line_indent)
        
        return self
    
    def add_table_with_style(self, data: List[List[str]],
                             header_color: str = '4472C4',
                             border: bool = True) -> 'WordStyleBuilder':
        """
        Add a styled table.
        
        Args:
            data: Table data (first row is header)
            header_color: Header background color
            border: Whether to show borders
        
        Returns:
            Self for chaining
        """
        if not data:
            return self
        
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                
                if i == 0:
                    shading = self._OxmlElement('w:shd')
                    shading.set(self._qn('w:fill'), header_color)
                    cell._tc.get_or_add_tcPr().append(shading)
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = self._RGBColor('FFFFFF')
        
        return self
    
    def add_picture(self, image_path: str,
                    width: Optional[float] = None,
                    height: Optional[float] = None,
                    alignment: str = 'center') -> 'WordStyleBuilder':
        """
        Add a picture.
        
        Args:
            image_path: Path to image file
            width: Width in inches
            height: Height in inches
            alignment: 'left', 'center', 'right'
        
        Returns:
            Self for chaining
        """
        kwargs = {}
        if width:
            kwargs['width'] = self._Inches(width)
        if height:
            kwargs['height'] = self._Inches(height)
        
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, **kwargs)
        
        align_map = {
            'left': self._WD_ALIGN_PARAGRAPH.LEFT,
            'center': self._WD_ALIGN_PARAGRAPH.CENTER,
            'right': self._WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p.alignment = align_map.get(alignment, self._WD_ALIGN_PARAGRAPH.CENTER)
        
        return self
    
    def add_page_break(self) -> 'WordStyleBuilder':
        """Add a page break."""
        self.doc.add_page_break()
        return self
    
    def save(self, filepath: str) -> str:
        """
        Save the document.
        
        Args:
            filepath: Output file path
        
        Returns:
            Saved file path
        """
        self.doc.save(filepath)
        return filepath


def create_chart(filepath: str) -> ExcelChartBuilder:
    """
    Create a chart builder for an Excel file.
    
    Args:
        filepath: Path to Excel file
    
    Returns:
        ExcelChartBuilder instance
    
    Example:
        >>> create_chart('report.xlsx') \\
        ...     .add_bar_chart('Sheet1', 'A1:B10', 'Sales') \\
        ...     .save()
    """
    return ExcelChartBuilder(filepath)


def create_styled_document() -> WordStyleBuilder:
    """
    Create a styled document builder.
    
    Returns:
        WordStyleBuilder instance
    
    Example:
        >>> create_styled_document() \\
        ...     .add_styled_heading('Report', color='1F4E79') \\
        ...     .add_styled_paragraph('Content', bold=True) \\
        ...     .save('output.docx')
    """
    return WordStyleBuilder()


# Convenience exports
__all__ = [
    'QuickGenerator',
    'generate_contract',
    'generate_report',
    'generate_excel',
    'list_templates',
    'num_to_chinese',
    'ExcelChartBuilder',
    'WordStyleBuilder',
    'create_chart',
    'create_styled_document',
]

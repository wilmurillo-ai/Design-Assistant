#!/usr/bin/env python3
"""Convert a document to gongwen format."""

import sys
import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


def set_run_font(run, east_asia: str, size_pt: float, *, bold: bool = False):
    run.bold = bold
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)


def add_page_number(section):
    """添加页码：4号半角宋体阿拉伯数字，左右各放一字线"""
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()

    dash_run = footer_para.add_run("—")
    set_run_font(dash_run, "宋体", 14)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    page_run = footer_para.add_run()
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_separate)
    page_run._r.append(fld_end)
    set_run_font(page_run, "宋体", 14)

    dash_run2 = footer_para.add_run("—")
    set_run_font(dash_run2, "宋体", 14)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_section(section):
    """配置页面设置"""
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.start_type = WD_SECTION_START.NEW_PAGE


def convert_heading(text, level_map):
    """转换层次序号"""
    for roman, chinese in level_map.items():
        if text.startswith(roman):
            return text.replace(roman, chinese, 1)
    return text


def is_heading(text, level_map):
    """判断是否为一级标题"""
    for roman in level_map.keys():
        if text.startswith(roman):
            return True
    return False


def is_subheading(text):
    """判断是否为子标题（二级层次）"""
    text = text.strip()
    # 以冒号结尾的简短文本，或者是区域名称
    if len(text) < 30 and text.endswith('：'):
        return True
    if len(text) < 30 and text.endswith(':'):
        return True
    # 特定关键词
    keywords = ['区域分析', '北美', '亚太地区', '挑战', '机遇', '具体需求']
    for kw in keywords:
        if kw in text and len(text) < 40:
            return True
    return False


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_path")
    parser.add_argument("output_path")
    parser.add_argument("--title", default="", help="Document title")
    args = parser.parse_args()

    src = Path(args.input_path)
    dst = Path(args.output_path)

    # 读取原文档
    src_doc = Document(str(src))

    # 创建新文档
    doc = Document()

    # 配置页面设置
    section = doc.sections[0]
    configure_section(section)
    add_page_number(section)

    # 层次序号映射（罗马数字转中文数字）
    level_map = {
        'II.': '一、',
        'III.': '二、',
        'IV.': '三、',
        'V.': '四、',
        'VI.': '五、',
        'VII.': '六、',
        'VIII.': '七、',
        'IX.': '八、',
        'X.': '九、',
        'XI.': '十、',
    }

    # 提取所有非空段落
    paragraphs = [p.text.strip() for p in src_doc.paragraphs if p.text.strip()]

    # 添加公文标题
    doc_title = args.title if args.title else "关于旅游行业人工智能咨询业务的调研报告"
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(doc_title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title_run, "方正小标宋简体", 22)

    # 添加空行（无红色分割线，因为没有主送机关）
    doc.add_paragraph()
    doc.add_paragraph()

    # 处理内容
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        # 跳过原标题
        if text == '旅游行业人工智能咨询业务研报':
            i += 1
            continue

        # 跳过"关键表格"及之后的内容（引用部分）
        if text == '关键表格':
            break

        # 转换一级标题（黑体）
        if is_heading(text, level_map):
            converted = convert_heading(text, level_map)
            para = doc.add_paragraph()
            run = para.add_run(converted)
            set_run_font(run, "黑体", 16)
            # 设置行距
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(28)
            i += 1
            continue

        # 处理子标题/二级层次（楷体）
        if is_subheading(text):
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_run_font(run, "楷体", 16)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(28)
            i += 1
            continue

        # 处理正文（仿宋，首行缩进2字符）
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_run_font(run, "仿宋_GB2312", 16)
        para.paragraph_format.first_line_indent = Pt(32)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(28)

        i += 1

    # 保存文档
    doc.save(str(dst))
    print(f"文档已保存至: {dst}")


if __name__ == "__main__":
    main()

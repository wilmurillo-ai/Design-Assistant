#!/usr/bin/env python3
"""Format a DOCX file into a gongwen-style report layout."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


def set_run_font(run, east_asia: str, size_pt: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)


def clear_paragraph_format(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    # 每面22行，版心高度225mm，每行高度 = 225 / 22 ≈ 10.227mm ≈ 29.05pt
    # 使用固定值28pt作为行距（3号字标准行距）
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(28)


def format_title(paragraph) -> None:
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "方正小标宋简体", 22)


def format_heading1(paragraph) -> None:
    """一级层次：黑体"""
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "黑体", 16)


def format_heading2(paragraph) -> None:
    """二级层次：楷体"""
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "楷体", 16)


def format_heading3(paragraph) -> None:
    """三级层次：仿宋"""
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "仿宋_GB2312", 16)


def format_heading4(paragraph) -> None:
    """四级层次：仿宋"""
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "仿宋_GB2312", 16)


def format_body(paragraph) -> None:
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(32)
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "仿宋_GB2312", 16)


def format_recipient(paragraph) -> None:
    """主送机关格式：3号仿宋，顶格"""
    clear_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        if run.text.strip():
            set_run_font(run, "仿宋_GB2312", 16)


def add_red_separator_line(paragraph) -> None:
    """添加红色分割线"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    # 添加红色下划线模拟分割线
    run.font.color.rgb = None  # 将在docx中手动设置为红色
    # 这里我们添加一个特殊标记，让后续处理知道这是红色分割线


def add_page_number(section) -> None:
    """添加页码：4号半角宋体阿拉伯数字，版心下边缘之下，左右各放一字线
    
    单页码右对齐，双页码左对齐
    """
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()
    
    # 获取节信息以确定奇偶页
    # 页码格式：一字线 + 页码 + 一字线，4号半角宋体
    # 单页右对齐，双页左对齐
    
    # 添加页码域
    run = footer_para.add_run()
    
    # 添加一字线（—）
    dash_run = footer_para.add_run("—")
    set_run_font(dash_run, "宋体", 14)
    
    # 添加页码域
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    
    page_run = footer_para.add_run()
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_separate)
    page_run._r.append(fld_end)
    set_run_font(page_run, "宋体", 14)
    
    # 添加一字线（—）
    dash_run2 = footer_para.add_run("—")
    set_run_font(dash_run2, "宋体", 14)
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_section(section) -> None:
    """配置页面设置
    
    纸张：A4 (210mm x 297mm)
    上白边（天头）：37mm
    左白边（订口）：28mm
    版心：156mm x 225mm
    右白边：210 - 28 - 156 = 26mm
    下白边：297 - 37 - 225 = 35mm
    """
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.start_type = WD_SECTION_START.NEW_PAGE


def detect_heading_level(text: str) -> int | None:
    """检测标题层次，返回层次级别或None
    
    层次序号格式：
    一级：一、二、三、...
    二级：（一）（二）（三）...
    三级：1. 2. 3. ...
    四级：（1）（2）（3）...
    """
    text = text.strip()
    if not text:
        return None
    
    # 一级：一、二、三、...（中文数字+顿号）
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 1
    
    # 二级：（一）（二）（三）...（括号+中文数字）
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 2
    
    # 三级：1. 2. 3. ...（阿拉伯数字+点）
    if re.match(r'^\d+\.', text):
        return 3
    
    # 四级：（1）（2）（3）...（括号+阿拉伯数字）
    if re.match(r'^（\d+）', text):
        return 4
    
    return None


def has_recipient(paragraphs) -> bool:
    """检测是否有主送机关
    
    主送机关特征：
    1. 在标题之后
    2. 通常以"："或":"结尾
    3. 不是正文段落（不以缩进开头）
    """
    found_title = False
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 第一个非空段落是标题
        if not found_title:
            found_title = True
            continue
        
        # 标题后的第一个非空段落
        # 主送机关通常以机关名称+冒号结尾
        if re.search(r'(厅|局|委|办|部|院|校|司|处|科|室|中心|政府|委员会|办公室|：|:)$', text):
            return True
        
        # 如果已经遇到正文（有缩进），则没有主送机关
        if para.paragraph_format.first_line_indent:
            return False
        
        # 只检查标题后的第一个非空段落
        break
    
    return False


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_path")
    parser.add_argument("output_path")
    args = parser.parse_args()

    src = Path(args.input_path)
    dst = Path(args.output_path)

    doc = Document(str(src))

    # 配置页面设置
    for section in doc.sections:
        configure_section(section)
        add_page_number(section)

    # 检测是否有主送机关
    recipient_exists = has_recipient(doc.paragraphs)
    
    # 处理段落
    non_empty_index = 0
    title_processed = False
    recipient_processed = False
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name.lower()
        
        # 第一个非空段落是标题
        if non_empty_index == 0:
            format_title(paragraph)
            title_processed = True
        elif non_empty_index == 1 and recipient_exists and not recipient_processed:
            # 第二个非空段落，如果有主送机关则格式化为主送机关
            format_recipient(paragraph)
            recipient_processed = True
        else:
            # 检测层次
            heading_level = detect_heading_level(text)
            
            if heading_level == 1:
                format_heading1(paragraph)
            elif heading_level == 2:
                format_heading2(paragraph)
            elif heading_level == 3:
                format_heading3(paragraph)
            elif heading_level == 4:
                format_heading4(paragraph)
            elif "heading 1" in style_name:
                format_heading1(paragraph)
            elif "heading 2" in style_name:
                format_heading2(paragraph)
            elif "heading 3" in style_name:
                format_heading3(paragraph)
            elif "heading 4" in style_name:
                format_heading4(paragraph)
            else:
                format_body(paragraph)
        
        non_empty_index += 1

    doc.save(str(dst))
    
    # 输出处理信息
    print(f"文档已保存至: {dst}")
    print(f"主送机关: {'有' if recipient_exists else '无'}")
    print("注意：如有主送机关，请在标题后手动添加红色分割线；如无，标题后应空两行")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
国网招标文件分析脚本 v6.0  ── hbdc · ClawHub Skill
═══════════════════════════════════════════════════════════════
v6 相比 v5 的关键变化:
  1. PDF 专属标题提取  — 用 pdfplumber 字符级字号定位封面大字
  2. PDF 读取失败自动重试  — 换 tolerance 参数后二次尝试
  3. "详见附件X" 占位符识别  — load_qual_attachment() 读 xlsx 资质表
     按 (分标编号, 包号) 回填, 避免 "详见附件" 无效占位进入资质列
  4. 关键词去冗余  — 若长词已命中则丢弃其子串短词
     (咨询服务→丢弃咨询; 储能系统→丢弃储能; etc.)
  5. 新列: 联系人及电话  — 正则 + xlsx 联系人字段映射
  6. 新列: 开标时间地点  — 正则抓 "开标时间/开标地点"
  7. 新列: 投标保证金  — 正则, 含 "不收取" 判断
  8. 新列: 评标办法  — 综合评估法 / 合理低价法 / 综合评分法 等
  9. 截止时间染色  — 红≤3天 黄≤7天 灰=已过期
 10. 输出列数: 19 → 23
═══════════════════════════════════════════════════════════════
"""

import os
import re
import sys
import json
import warnings
from datetime import date, datetime
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")


# ========== 依赖加载 ==========
def _ensure(pkg, import_name=None):
    try:
        __import__(import_name or pkg)
    except ImportError:
        print(f"  [安装] 正在安装 {pkg} ...")
        os.system(f"pip3 install {pkg} --break-system-packages -q 2>/dev/null")

_ensure("python-docx", "docx")
_ensure("openpyxl")
_ensure("pdfplumber")

from docx import Document
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


# ========== 关键词定义 ==========
# 支持从 config/keywords.json 加载自定义关键词 (可选)
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_KW_CONFIG   = os.path.join(_SCRIPT_DIR, "..", "config", "keywords.json")

DEFAULT_KEYWORDS = {
    "新能源汽车充换电类": ["新能源汽车", "充换电", "换电站", "充电桩", "电动汽车", "OBD", "车网互动", "V2G"],
    "智能巡检机器人类":   ["巡检机器人", "光伏清洁机器人", "线路巡检", "隧道巡检", "智能巡检", "无人巡检"],
    "虚拟电厂类":         ["虚拟电厂", "VPP", "需求侧响应", "源荷互动", "负荷聚合", "光储充", "储能调度"],
    "人工智能应用类":     ["人工智能", "AI大模型", "大数据分析", "智能算法", "预测分析", "智能诊断", "知识图谱", "智能运维"],
    "储能技术类":         ["储能系统", "光储", "储充", "储能柜", "模块化储能", "充换储配"],
    "电力设备检测类":     ["变压器检测", "故障检测", "异常预警", "声学监测", "设备监测"],
    "其他技术类":         ["带电作业", "屏蔽服", "直流电源", "电缆敷设", "配网工程", "稳压宝", "零线监测"],
    "宣传服务类":         ["宣传服务", "广告宣传", "党建宣传", "企业文化", "展览宣传", "品牌推广", "宣传框架"],
    "技术服务类":         ["技术服务框架", "科技服务", "电力技术服务", "可视化服务", "科技创新服务"],
    "咨询服务类":         ["咨询服务", "造价咨询", "专利代理", "培训咨询", "市场研究"],
    "管理咨询类":         ["管理咨询", "标准化管理", "企业管理咨询"],
}
DEFAULT_SHORT_KEYWORDS = {
    "宣传服务类":     ["宣传", "广告", "展览"],
    "技术服务类":     ["技术服务"],
    "咨询服务类":     ["咨询"],          # ← 会被"咨询服务"去冗余
    "储能技术类":     ["储能", "光储", "储充"],
    "电力设备检测类": ["变压器"],
    "其他技术类":     ["配网", "电缆敷设"],
    "人工智能应用类": ["大数据", "算法"],
}

def _load_keywords():
    if os.path.isfile(_KW_CONFIG):
        try:
            with open(_KW_CONFIG, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            return cfg.get("keywords", DEFAULT_KEYWORDS), cfg.get("short_keywords", DEFAULT_SHORT_KEYWORDS)
        except Exception:
            pass
    return DEFAULT_KEYWORDS, DEFAULT_SHORT_KEYWORDS

KEYWORDS, SHORT_KEYWORDS = _load_keywords()

BOILERPLATE_PATTERNS = [
    "制作投标文件", "制作并提交", "投标工具", "供应商投标", "电子投标",
    "咨询有限公司", "咨询电话", "咨询服务电话", "技术支持服务电话",
    "信息系统支持服务", "检测能力", "产品检测", "元器件检测",
    "设计投标", "设计开发服务",
]


def _is_boilerplate(text, kw):
    idx = text.find(kw)
    if idx == -1:
        return True
    ctx = text[max(0, idx - 50): idx + len(kw) + 50]
    return any(p in ctx for p in BOILERPLATE_PATTERNS)


def deduplicate_keywords(words):
    """
    去冗余: 若长词 A 已命中且短词 B 是 A 的子串, 丢弃 B。
    例: 咨询服务 ∋ 咨询 → 丢弃咨询; 储能系统 ∋ 储能 → 丢弃储能
    """
    result = []
    all_words = list(words)
    for w in all_words:
        # w 是某个其他词的子串 → 冗余
        if any(w != other and w in other for other in all_words):
            continue
        result.append(w)
    return result


def match_keywords(text, strict=False):
    words = []
    for cat, kws in KEYWORDS.items():
        for w in kws:
            if w in text and w not in words:
                words.append(w)
    for cat, kws in SHORT_KEYWORDS.items():
        for w in kws:
            if w in text:
                if strict and _is_boilerplate(text, w):
                    continue
                if w not in words:
                    words.append(w)
    return deduplicate_keywords(words)


# ========== 文本读取 ==========
def read_doc_paragraphs(path, retry=True):
    """返回段落文本列表。PDF 失败时自动重试。"""
    ext = os.path.splitext(path)[1].lower()
    paras = []
    if ext == ".docx":
        try:
            doc = Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for r in t.rows:
                    line = " ".join(c.text.strip() for c in r.cells if c.text.strip())
                    if line:
                        paras.append(line)
        except Exception as e:
            print(f"  [警告] 解析 docx 失败 {os.path.basename(path)}: {e}")
    elif ext == ".pdf":
        paras = _read_pdf_paragraphs(path, retry=retry)
    return paras


def _read_pdf_paragraphs(path, retry=True):
    """pdfplumber 读 PDF 段落, 失败后换 tolerance 参数重试一次。"""
    if not HAS_PDF:
        print(f"  [警告] 缺少 pdfplumber, 无法解析 {os.path.basename(path)}")
        return []
    paras = _try_pdf_read(path, x_tolerance=3, y_tolerance=3)
    if not paras and retry:
        print(f"  [重试] PDF 首次读取失败, 切换参数重试: {os.path.basename(path)}")
        paras = _try_pdf_read(path, x_tolerance=5, y_tolerance=5)
        if not paras:
            print(f"  [警告] PDF 两次读取均失败: {os.path.basename(path)}")
    return paras


def _try_pdf_read(path, x_tolerance=3, y_tolerance=3):
    paras = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text(x_tolerance=x_tolerance, y_tolerance=y_tolerance) or ""
                for line in txt.split("\n"):
                    line = line.strip()
                    if line:
                        paras.append(line)
    except Exception as e:
        pass
    return paras


# ========== PDF 专属标题提取 (字号最大的文字) ==========
def extract_pdf_title_by_fontsize(pdf_path):
    """
    用 pdfplumber 字符级字号提取 PDF 首页最大字号文字作为标题。
    比正则扫描段落更准确, 可以找到封面大字。
    """
    if not HAS_PDF:
        return ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if not pdf.pages:
                return ""
            page = pdf.pages[0]
            chars = page.chars
            if not chars:
                return ""
            # 找最大字号 (忽略装饰性超大字)
            sizes = sorted(set(round(c.get("size", 0)) for c in chars if c.get("size", 0) > 8), reverse=True)
            if not sizes:
                return ""
            # 取最大字号 (若最大字号只有1-2个字符则尝试次大)
            for max_size in sizes[:3]:
                title_chars = [c for c in chars if abs(c.get("size", 0) - max_size) < 1.5]
                if len(title_chars) < 3:
                    continue
                # 按位置 (行 → 列) 排序
                title_chars.sort(key=lambda c: (round(c.get("top", 0) / 8) * 8, c.get("x0", 0)))
                title = "".join(c.get("text", "") for c in title_chars).strip()
                title = re.sub(r"\s+", "", title)
                if len(title) >= 8:
                    return title[:80]
    except Exception:
        pass
    return ""


# ========== 正则提取: 开标/保证金/评标/联系人 ==========
def extract_bidding_details(full_text):
    """提取开标时间地点、投标保证金、评标办法、联系人及电话"""
    result = {
        "开标时间地点": "",
        "投标保证金":   "",
        "评标办法":     "",
        "联系人及电话": "",
    }

    # ── 开标时间地点 ──
    time_str, place_str = "", ""
    m = re.search(
        r"开标时间[：:]\s*(\d{4}年\d{1,2}月\d{1,2}日\s*\d{1,2}[:：时]\d{0,2}(?:分)?)",
        full_text,
    )
    if m:
        time_str = m.group(1).replace(" ", "")
    m2 = re.search(r"开标地(?:点|址)[：:]\s*([^\n，。]{5,50})", full_text)
    if m2:
        place_str = m2.group(1).strip()
    if time_str or place_str:
        result["开标时间地点"] = " | ".join(filter(None, [time_str, place_str]))

    # ── 投标保证金 ──
    if re.search(r"(?:不收取|免收|无需缴纳).*?保证金|保证金.*?(?:不收取|免收|无需)", full_text):
        result["投标保证金"] = "不收取"
    else:
        m = re.search(
            r"投标保证金[^0-9。\n]{0,20}([\d,，.]+\s*(?:万元|元))",
            full_text,
        )
        if m:
            result["投标保证金"] = m.group(1).replace("，", ",").strip()

    # ── 评标办法 ──
    for method in [
        "综合评估法", "综合评分法",
        "合理低价法", "最低价法",
        "经评审的最低投标价法", "最低评标价法",
    ]:
        if method in full_text:
            result["评标办法"] = method
            break

    # ── 联系人及电话 ──
    contacts = []
    # 联系人+电话同行
    for m in re.finditer(
        r"联系人[：:]\s*([^\s，。\n]{2,8})\s*[，,]?\s*(?:电话|联系电话|手机|Tel)[：:]\s*([\d\-\s()（）+]{7,20})",
        full_text,
    ):
        entry = f"{m.group(1).strip()}  {m.group(2).strip()}"
        if entry not in contacts:
            contacts.append(entry)
    # 仅有电话时
    if not contacts:
        for m in re.finditer(
            r"(?:联系电话|咨询电话|招标咨询)[：:]\s*([\d\-\s()（）+]{7,20})",
            full_text,
        ):
            entry = m.group(1).strip()
            if entry not in contacts:
                contacts.append(entry)
    result["联系人及电话"] = " / ".join(contacts[:3])

    return result


# ========== 资质附件加载 (详见附件X 占位符回填) ==========
FUJIAN_RE = re.compile(r"详见附件\s*(\d+|[一二三四五六七八九十])")

def is_qual_placeholder(text):
    """判断资质文本是否只是 '详见附件X' 占位符"""
    stripped = text.strip()
    return bool(FUJIAN_RE.fullmatch(stripped) or
                (len(stripped) < 20 and FUJIAN_RE.search(stripped)))


def load_qual_attachment(xlsxs):
    """
    从 xlsx 附件中加载资质表 (文件名含 资质/条件/要求/附件 的优先)。
    返回 {(分标编号, 包号): qual_str}
    """
    mapping = {}
    priority = []
    others   = []
    for xp in xlsxs:
        fname = os.path.basename(xp)
        if any(k in fname for k in ["资质", "条件", "要求"]):
            priority.append(xp)
        else:
            others.append(xp)

    for xp in priority + others:
        try:
            wb = openpyxl.load_workbook(xp, data_only=True)
        except Exception:
            continue
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_2d = list(ws.iter_rows(values_only=True))
            pkgs = parse_package_rows(rows_2d)
            for pkg in pkgs:
                key  = (pkg.get("分标编号", ""), pkg.get("包号", ""))
                qual = pkg.get("资质/资格要求", "")
                if qual and key not in mapping:
                    mapping[key] = qual
    return mapping


# ========== 主公告文件挑选 ==========
GOOD_NAME_HINTS = ["招标公告", "采购公告", "招标文件"]
BAD_NAME_HINTS  = [
    "重要提醒", "提醒", "答疑", "澄清", "资质", "汇总表",
    "一览表", "需求表", "附件1", "附件2", "附件3", "附件4",
    "附件5", "评标办法", "合同条款", "投标文件格式", "货物清单",
]


def pick_main_announcement(paths):
    if not paths:
        return None
    def score(p):
        name = os.path.basename(p)
        s = 0
        for g in GOOD_NAME_HINTS:
            if g in name: s += 10
        for b in BAD_NAME_HINTS:
            if b in name: s -= 8
        if name.lower().endswith(".docx"): s += 1
        return s
    return sorted(paths, key=score, reverse=True)[0]


# ========== 资格要求过滤 ==========
QUAL_BOILERPLATE = [
    "依法注册", "财务状况", "失信被执行人", "失信", "信用中国",
    "供应商关系", "联合体", "经营异常", "严重违法", "黑名单",
    "诉讼", "破产", "重组", "不良行为", "同一人", "控股、管理",
    "国家强制认证", "通用资格", "电子招标投标办法", "营业执照信息",
    "处罚信息", "清算信息", "公示系统", "限制消费令", "公告期",
    "中华人民共和国境内", "良好的财务", "国家法律、法规", "其他法律",
    "暂停中标", "取消中标", "政府采购法",
]
QUAL_SUBSTANTIVE = [
    "甲级", "乙级", "丙级", "等级资质", "建造师", "工程师", "技师",
    "ISO", "安全生产许可证", "总承包", "专业承包",
    "项目经理", "项目负责人", "技术负责人",
    "设计资质", "施工资质", "监理资质", "检测资质",
    "业绩", "实绩", "投标人具有", "投标人需具有", "投标人应具有",
    "三级", "二级", "一级",
]


def filter_qualification_paragraphs(paras):
    keep = []
    for p in paras:
        s = p.strip()
        if not s or len(s) < 6:
            continue
        if re.match(r"^\s*\d+(\.\d+)*[\.、\s]", s) and len(s) < 35:
            continue
        if any(b in s for b in QUAL_BOILERPLATE):
            continue
        if any(k in s for k in QUAL_SUBSTANTIVE):
            keep.append(s)
            continue
        if ("证书" in s or "资质" in s or "许可证" in s) and len(s) < 200:
            keep.append(s)
    return keep


# ========== 元数据提取 ==========
def extract_meta(announcement_path):
    meta = {
        "项目名称": "", "招标编号": "", "招标起止时间": "",
        "项目级资格要求": "",
        # 新增 v6
        "开标时间地点": "", "投标保证金": "", "评标办法": "", "联系人及电话": "",
    }
    if not announcement_path:
        return meta

    paras = read_doc_paragraphs(announcement_path)
    if not paras:
        return meta
    full_text = "\n".join(paras)

    # 招标编号
    m = re.search(r"招标编号[：:]\s*([A-Za-z0-9\-_/（）()\u4e00-\u9fff]+)", full_text)
    if not m:
        m = re.search(r"(?:项目编号|采购编号|批次编[号码])[：:]\s*([A-Za-z0-9\-_/（）()]+)", full_text)
    if m:
        meta["招标编号"] = m.group(1).strip()

    # 项目名称: PDF 优先用字号法
    ext = os.path.splitext(announcement_path)[1].lower()
    if ext == ".pdf":
        title = extract_pdf_title_by_fontsize(announcement_path)
        if title:
            meta["项目名称"] = title

    if not meta["项目名称"]:
        # 通用段落法
        candidates = []
        for i, p in enumerate(paras[:50]):
            if not (10 <= len(p) <= 80):
                continue
            if "采购" not in p and "招标" not in p:
                continue
            if re.match(r"^[\(（]?\d+[\)）\.、]", p):
                continue
            if any(p.startswith(x) for x in ("招标编号", "招标人", "招标代理", "项目编号", "采购编号")):
                continue
            if "招标公告" in p and len(p) < 20:
                continue
            if any(w in p for w in ["不得", "应当", "必须", "如下", "规定", "办法", "通知"]):
                continue
            candidates.append((i, p))
        if candidates:
            candidates.sort(key=lambda x: (
                not (x[1].endswith("采购") or x[1].endswith("招标")), x[0]))
            meta["项目名称"] = candidates[0][1]

    # 招标起止时间
    parts = []
    m1 = re.search(
        r"(\d{4}年\d{1,2}月\d{1,2}日[^，。\n]{0,5}-\s*\d{4}年\d{1,2}月\d{1,2}日\s*\d{0,2}:?\d{0,2}时?)",
        full_text,
    )
    if m1:
        parts.append("获取:" + m1.group(1).replace(" ", ""))
    m2 = re.search(
        r"投标(?:文件提交的)?截止时间[^0-9]{0,20}(\d{4}年\d{1,2}月\d{1,2}日\s*\d{1,2}:?\d{0,2}时?)",
        full_text,
    )
    if m2:
        parts.append("截止:" + m2.group(1).replace(" ", ""))
    meta["招标起止时间"] = " | ".join(parts)

    # 项目级资格要求
    start = end = -1
    for i, p in enumerate(paras):
        if start < 0 and re.search(r"投标人资格(?:要求|条件)", p) and len(p) < 40:
            start = i
            continue
        if start >= 0:
            if re.match(r"^\s*[4四][\.、\s]", p) and len(p) < 50:
                end = i; break
            if "招标文件的获取" in p or "投标文件的递交" in p:
                end = i; break
    if start >= 0:
        section = paras[start: end if end > 0 else min(start + 30, len(paras))]
        kept = filter_qualification_paragraphs(section)
        text = " ".join(kept)
        if len(text) > 600:
            text = text[:600] + "…"
        meta["项目级资格要求"] = text

    # 新增 v6: 开标/保证金/评标/联系人
    details = extract_bidding_details(full_text)
    meta.update(details)

    return meta


# ========== 统一 2D 表行解析器 ==========
PKG_FIELDS = [
    "分标编号", "分标名称", "包号", "需求部门/签订单位", "子项目名称",
    "项目概况与招标范围", "资质/资格要求", "合同文本编号", "实施地点",
    "工期/服务期", "报价方式", "预算金额(万元)", "最高限价",
]

COL_ALIASES = {
    "分标编号":          ["分标编号", "采购编号"],
    "分标名称":          ["分标名称"],
    "包号":              ["分包名称", "包名称", "包号", "分包编号"],
    "需求部门/签订单位":  ["项目单位", "需求单位", "签订单位", "公司名称", "需求部门"],
    "子项目名称":         ["项目名称"],
    "项目概况":           ["工程概况", "项目概况", "实施简要", "项目实施", "物料名称",
                          "物资名称", "采购范围", "招标范围", "物资描述"],
    "工期":               ["工程进度安排", "项目预计开始时间", "服务期限", "工期",
                          "服务期", "竣工日期", "交货日期"],
    "资质条件":           ["资质条件", "资质要求"],
    "业绩要求":           ["业绩要求", "业绩条件"],
    "主要人员":           ["主要人员要求", "人员要求", "项目经理"],
    "概算金额":           ["概算金额", "概算总价", "预算金额", "招标控制价"],
    "最高限价":           ["最高限价"],
    "报价方式":           ["报价方式"],
    "合同模板":           ["合同模板", "合同文本"],
    "实施地点":           ["实施地点", "项目地点", "交货地点"],
    # v6 新增 联系人字段别名 (xlsx 中可能有)
    "联系人":             ["联系人", "联络人", "负责人"],
    "联系电话":           ["联系电话", "电话", "手机", "Tel"],
}


def _detect_money_unit(header_text):
    if "万元" in header_text: return "wan"
    if "折扣率" in header_text or "%" in header_text: return "ratio"
    if "(元" in header_text or "（元" in header_text: return "yuan"
    return "unknown"


def _to_wan(value, unit):
    if not value: return ""
    s = str(value).replace(",", "").strip()
    if unit == "ratio": return s
    try:
        n = float(s)
        if unit == "yuan": return f"{n / 10000:.2f}"
        if unit == "wan":  return f"{n:.2f}"
        return f"{n / 10000:.2f}" if n > 100000 else f"{n:.2f}"
    except (ValueError, TypeError):
        return s


def parse_package_rows(rows_2d):
    if not rows_2d or len(rows_2d) < 2:
        return []
    rows = [[("" if c is None else str(c).strip()) for c in r] for r in rows_2d]

    headers, hidx = None, -1
    for i in range(min(4, len(rows))):
        joined = " ".join(rows[i])
        if any(k in joined for k in ["分标编号", "分标名称", "分包名称", "包号", "包名称"]):
            headers = rows[i]
            hidx    = i
            break
    if hidx < 0:
        return []

    def find_col(*aliases):
        for alias in aliases:
            for k, h in enumerate(headers):
                if alias in h:
                    return k
        return -1

    cols = {std: find_col(*aliases) for std, aliases in COL_ALIASES.items()}

    yusuan_unit = _detect_money_unit(headers[cols["概算金额"]]) if cols["概算金额"] >= 0 else "unknown"
    xianjia_unit = _detect_money_unit(headers[cols["最高限价"]]) if cols["最高限价"] >= 0 else "unknown"

    def cell(row, idx):
        if idx < 0 or idx >= len(row): return ""
        return row[idx]

    results = []
    for row in rows[hidx + 1:]:
        if not any(row): continue
        biaobian = cell(row, cols["分标编号"])
        biaoming = cell(row, cols["分标名称"])
        baohao   = cell(row, cols["包号"])
        if not (biaobian or biaoming or baohao): continue
        if biaobian == "分标编号" or biaoming == "分标名称": continue

        gaikuang_parts = []
        if cols["项目概况"] >= 0:
            v = cell(row, cols["项目概况"])
            if v and not is_qual_placeholder(v):
                gaikuang_parts.append(v)

        # 资质聚合: 过滤 "详见附件X" 占位符
        zizhi_parts = []
        for std, label in [("资质条件", "资质"), ("业绩要求", "业绩"), ("主要人员", "人员")]:
            if cols[std] >= 0:
                v = cell(row, cols[std])
                if v and v not in ("/", "无", "—", "-") and not is_qual_placeholder(v):
                    zizhi_parts.append(f"【{label}】{v}")

        # xlsx 联系人字段
        contact_parts = []
        if cols["联系人"] >= 0:
            v = cell(row, cols["联系人"])
            if v: contact_parts.append(v)
        if cols["联系电话"] >= 0:
            v = cell(row, cols["联系电话"])
            if v: contact_parts.append(v)

        results.append({
            "分标编号":           biaobian,
            "分标名称":           biaoming,
            "包号":               baohao,
            "需求部门/签订单位":  cell(row, cols["需求部门/签订单位"]),
            "子项目名称":         cell(row, cols["子项目名称"]),
            "项目概况与招标范围":  " / ".join(gaikuang_parts)[:600],
            "资质/资格要求":      "\n".join(zizhi_parts),
            "合同文本编号":       cell(row, cols["合同模板"]),
            "实施地点":           cell(row, cols["实施地点"]),
            "工期/服务期":        cell(row, cols["工期"]),
            "报价方式":           cell(row, cols["报价方式"]),
            "预算金额(万元)":      _to_wan(cell(row, cols["概算金额"]), yusuan_unit),
            "最高限价":           _to_wan(cell(row, cols["最高限价"]), xianjia_unit),
            "_pkg_contact":       "  ".join(contact_parts),   # 内部字段, 后续合并用
        })
    return results


# ========== 各源的包记录抽取 ==========
def extract_packages_from_word(docx_path):
    out = []
    try:
        doc = Document(docx_path)
    except Exception:
        return out
    for table in doc.tables:
        rows_2d = []
        for r in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            rows_2d.append(cells)
        out.extend(parse_package_rows(rows_2d))
    return out


def extract_packages_from_pdf(pdf_path):
    out = []
    if not HAS_PDF:
        return out
    # 第一次尝试
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                for tbl in (page.extract_tables() or []):
                    out.extend(parse_package_rows(tbl))
    except Exception as e:
        print(f"  [警告] PDF 表格一次抽取失败: {e}")
    # 若无结果, 换 explicit_vertical_lines 参数重试
    if not out:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    tbl_settings = {"vertical_strategy": "lines", "horizontal_strategy": "lines"}
                    for tbl in (page.extract_tables(tbl_settings) or []):
                        out.extend(parse_package_rows(tbl))
        except Exception as e:
            print(f"  [警告] PDF 表格二次抽取也失败: {e}")
    return out


def extract_packages_from_xlsx(xlsx_path):
    out = []
    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception as e:
        print(f"  [警告] 打开 xlsx 失败: {e}")
        return out
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_2d = list(ws.iter_rows(values_only=True))
        out.extend(parse_package_rows(rows_2d))
    return out


# ========== 合并 (主源 + 补盲源) ==========
def merge_packages(primary, supplement):
    by_key = {}
    order  = []
    for p in primary:
        key = (p.get("分标编号", ""), p.get("包号", ""))
        if key not in by_key:
            by_key[key] = dict(p)
            order.append(key)
    for s in supplement:
        key = (s.get("分标编号", ""), s.get("包号", ""))
        if key in by_key:
            for f, v in s.items():
                if v and not by_key[key].get(f):
                    by_key[key][f] = v
        else:
            by_key[key] = dict(s)
            order.append(key)
    return [by_key[k] for k in order]


# ========== 报告生成 (23 列 + 截止时间染色) ==========
HEADERS = [
    "序号", "项目名称", "招标编号", "分标编号", "分标名称", "包号",
    "需求部门/签订单位", "子项目名称", "项目概况与招标范围",
    "资质/资格要求", "合同文本编号", "实施地点", "工期/服务期",
    "报价方式", "预算金额(万元)", "最高限价",
    "招标起止时间",
    # v6 新列
    "开标时间地点", "投标保证金", "评标办法", "联系人及电话",
    # 末尾
    "匹配关键词", "匹配来源",
]
COL_WIDTHS = [
    6, 38, 22, 24, 26, 8, 26, 36, 50,
    70, 24, 24, 22,
    14, 16, 14,
    32,
    28, 16, 18, 28,
    26, 18,
]
assert len(HEADERS) == len(COL_WIDTHS), "HEADERS 与 COL_WIDTHS 长度不一致!"

# 截止时间列索引 (1-based)
DEADLINE_COL = HEADERS.index("招标起止时间") + 1


def _parse_deadline_date(val):
    """从 招标起止时间 字符串中解析截止日期"""
    m = re.search(r"截止[：:]?.*?(\d{4})年(\d{1,2})月(\d{1,2})日", str(val))
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except Exception:
            pass
    return None


RED_FILL    = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")
YELLOW_FILL = PatternFill(start_color="FFD966", end_color="FFD966", fill_type="solid")
GRAY_FILL   = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")


def generate_report(results, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "可投标项目汇总"

    # 表头
    hdr_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    hdr_font = Font(color="FFFFFF", bold=True, size=11)
    for c, h in enumerate(HEADERS, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.fill      = hdr_fill
        cell.font      = hdr_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 32

    today = date.today()

    for i, r in enumerate(results, 1):
        row_idx = i + 1
        ws.cell(row=row_idx, column=1, value=i)
        for c, key in enumerate(HEADERS[1:], 2):
            ws.cell(row=row_idx, column=c, value=r.get(key, ""))
        for c in range(1, len(HEADERS) + 1):
            ws.cell(row=row_idx, column=c).alignment = Alignment(wrap_text=True, vertical="top")

        # 截止时间染色
        deadline_val = r.get("招标起止时间", "")
        dl = _parse_deadline_date(deadline_val)
        if dl:
            delta = (dl - today).days
            fill = GRAY_FILL if delta < 0 else (RED_FILL if delta <= 3 else (YELLOW_FILL if delta <= 7 else None))
            if fill:
                for c in range(1, len(HEADERS) + 1):
                    ws.cell(row=row_idx, column=c).fill = fill

    # 列宽
    for i, w in enumerate(COL_WIDTHS, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.freeze_panes = "B2"

    # 图例 sheet
    _add_legend_sheet(wb)

    wb.save(output_path)


def _add_legend_sheet(wb):
    ls = wb.create_sheet("图例说明")
    legend = [
        ("颜色", "含义"),
        ("🔴 红色", "距截止 ≤3 天, 紧急"),
        ("🟡 黄色", "距截止 ≤7 天, 即将截止"),
        ("⬜ 灰色", "已过截止时间"),
        ("无底色", "7天以上, 正常"),
    ]
    for r, (a, b) in enumerate(legend, 1):
        ls.cell(row=r, column=1, value=a)
        ls.cell(row=r, column=2, value=b)
    ls.column_dimensions["A"].width = 16
    ls.column_dimensions["B"].width = 28


# ========== 项目级编排 ==========
def process_project(proj_name, main_doc, xlsxs):
    meta = extract_meta(main_doc) if main_doc else {
        "项目名称": "", "招标编号": "", "招标起止时间": "",
        "项目级资格要求": "",
        "开标时间地点": "", "投标保证金": "", "评标办法": "", "联系人及电话": "",
    }

    # 1. 主源
    primary_packages = []
    primary_label    = ""
    if main_doc:
        ext = os.path.splitext(main_doc)[1].lower()
        if ext == ".docx":
            primary_packages = extract_packages_from_word(main_doc)
            primary_label    = "Word表格"
        elif ext == ".pdf":
            primary_packages = extract_packages_from_pdf(main_doc)
            primary_label    = "PDF表格"

    # 2. 补盲
    xlsx_packages = []
    for xp in xlsxs:
        xlsx_packages.extend(extract_packages_from_xlsx(xp))

    # 3. 合并
    if primary_packages and xlsx_packages:
        merged       = merge_packages(primary_packages, xlsx_packages)
        source_label = f"{primary_label}+Excel"
    elif primary_packages:
        merged       = primary_packages
        source_label = primary_label
    elif xlsx_packages:
        merged       = xlsx_packages
        source_label = "Excel附件"
    else:
        merged       = []
        source_label = ""

    # 4. 详见附件X 占位符资质回填
    qual_map = {}
    if merged:
        has_placeholder = any(is_qual_placeholder(p.get("资质/资格要求", "")) for p in merged)
        if has_placeholder:
            qual_map = load_qual_attachment(xlsxs)

    # 5. 关键词过滤 + 行组装
    matched = []
    for pkg in merged:
        text = " ".join([
            pkg.get("分标名称", ""), pkg.get("子项目名称", ""),
            pkg.get("项目概况与招标范围", ""),
        ])
        words = match_keywords(text, strict=False)
        if not words:
            continue

        # 资质列: 占位符 → 附件回填 → 项目级兜底
        zizhi = pkg.get("资质/资格要求", "")
        if is_qual_placeholder(zizhi):
            key = (pkg.get("分标编号", ""), pkg.get("包号", ""))
            zizhi = qual_map.get(key, "") or qual_map.get(("", key[1]), "")
        if not zizhi and meta.get("项目级资格要求"):
            zizhi = "【通用】" + meta["项目级资格要求"]

        # 联系人: 包级 xlsx 字段 优先, 否则用项目级正则结果
        contact = pkg.get("_pkg_contact", "") or meta.get("联系人及电话", "")

        matched.append({
            "项目名称":           meta.get("项目名称", ""),
            "招标编号":           meta.get("招标编号", ""),
            "分标编号":           pkg.get("分标编号", ""),
            "分标名称":           pkg.get("分标名称", ""),
            "包号":               pkg.get("包号", ""),
            "需求部门/签订单位":   pkg.get("需求部门/签订单位", ""),
            "子项目名称":         pkg.get("子项目名称", ""),
            "项目概况与招标范围":  pkg.get("项目概况与招标范围", ""),
            "资质/资格要求":      zizhi,
            "合同文本编号":       pkg.get("合同文本编号", ""),
            "实施地点":           pkg.get("实施地点", ""),
            "工期/服务期":        pkg.get("工期/服务期", ""),
            "报价方式":           pkg.get("报价方式", ""),
            "预算金额(万元)":      pkg.get("预算金额(万元)", ""),
            "最高限价":           pkg.get("最高限价", ""),
            "招标起止时间":       meta.get("招标起止时间", ""),
            # v6 新列
            "开标时间地点":       meta.get("开标时间地点", ""),
            "投标保证金":         meta.get("投标保证金", ""),
            "评标办法":           meta.get("评标办法", ""),
            "联系人及电话":       contact,
            # 末尾
            "匹配关键词":         ", ".join(words),
            "匹配来源":           source_label,
        })

    # 6. 回退: 无包级数据时扫描正文
    if not merged and main_doc:
        paras     = read_doc_paragraphs(main_doc)
        full_text = " ".join(paras)
        words     = match_keywords(full_text, strict=True)
        if words:
            matched.append({
                "项目名称":           meta.get("项目名称", ""),
                "招标编号":           meta.get("招标编号", ""),
                "分标编号":           "", "分标名称": "", "包号": "",
                "需求部门/签订单位":   "",
                "子项目名称":         "(详见招标公告正文,无包级表格)",
                "项目概况与招标范围":  "见招标公告正文",
                "资质/资格要求":      meta.get("项目级资格要求", ""),
                "合同文本编号":       "", "实施地点": "", "工期/服务期": "",
                "报价方式":           "", "预算金额(万元)": "", "最高限价": "",
                "招标起止时间":       meta.get("招标起止时间", ""),
                "开标时间地点":       meta.get("开标时间地点", ""),
                "投标保证金":         meta.get("投标保证金", ""),
                "评标办法":           meta.get("评标办法", ""),
                "联系人及电话":       meta.get("联系人及电话", ""),
                "匹配关键词":         ", ".join(words),
                "匹配来源":           "Word/PDF正文",
            })

    return meta, merged, matched


# ========== 项目收集 + 主程序 ==========
def collect_projects(base_dir):
    projects = []
    for name in sorted(os.listdir(base_dir)):
        full = os.path.join(base_dir, name)
        if name.startswith(".") or name.startswith("~"):
            continue
        if os.path.isdir(full):
            files = [f for f in os.listdir(full) if not f.startswith(".") and not f.startswith("~")]
            docs  = [os.path.join(full, f) for f in files if f.lower().endswith((".docx", ".pdf"))]
            xlsxs = [os.path.join(full, f) for f in files if f.lower().endswith(".xlsx")]
            main  = pick_main_announcement(docs)
            projects.append((name, main, xlsxs))
        elif os.path.isfile(full) and full.lower().endswith((".docx", ".pdf")):
            projects.append((os.path.splitext(name)[0], full, []))
    return projects


def main():
    base_dir    = os.path.expanduser("~/Desktop/sgcc_files")
    output_path = os.path.expanduser("~/Desktop/sgcc_result.xlsx")

    print("=" * 68)
    print("  hbdc · 国网招标文件分析工具 v6.0")
    print("  Word/PDF 优先 + Excel 补盲 + 智能资质回填 + 截止染色")
    print("=" * 68)
    print(f"扫描目录: {base_dir}")
    if not HAS_PDF:
        print("⚠️  pdfplumber 未安装, PDF 文件将被跳过。")
    print()

    if not os.path.exists(base_dir):
        print(f"[错误] 目录不存在: {base_dir}")
        sys.exit(1)

    projects = collect_projects(base_dir)
    if not projects:
        print("[错误] 未找到任何项目。")
        sys.exit(1)

    print(f"发现 {len(projects)} 个项目\n")

    all_results   = []
    warn_projects = []  # 来源为 Word/PDF正文 的项目 (需手动补盲)

    for proj_name, main_doc, xlsxs in projects:
        print(f"━━ 项目: {proj_name}")
        if not main_doc:
            print("  [跳过] 无 docx/pdf 主公告\n")
            continue
        print(f"  主公告: {os.path.basename(main_doc)}")
        if xlsxs:
            print(f"  附件: {len(xlsxs)} 个 xlsx")

        meta, merged, matched = process_project(proj_name, main_doc, xlsxs)

        print(f"  项目名称: {meta['项目名称'] or '(未识别)'}")
        print(f"  招标编号: {meta['招标编号'] or '(未识别)'}")
        print(f"  招标时间: {meta['招标起止时间'] or '(未识别)'}")
        if meta.get("开标时间地点"):
            print(f"  开标信息: {meta['开标时间地点']}")
        if meta.get("投标保证金"):
            print(f"  投标保证金: {meta['投标保证金']}")
        if meta.get("评标办法"):
            print(f"  评标办法: {meta['评标办法']}")
        print(f"  包级记录: 抽出 {len(merged)} 条 → 关键词命中 {len(matched)} 条")
        for r in matched[:5]:
            tag = f"{r['包号'] or '-'} {(r['分标名称'] or r['子项目名称'])[:24]}"
            print(f"     · {tag} | {r['匹配关键词']} | 来源:{r['匹配来源']}")
        if len(matched) > 5:
            print(f"     ... 还有 {len(matched) - 5} 条")

        # 收集需要手动补盲的项目
        for r in matched:
            if r.get("匹配来源") == "Word/PDF正文":
                warn_projects.append(proj_name)
                break
        print()
        all_results.extend(matched)

    print("=" * 68)
    if all_results:
        generate_report(all_results, output_path)
        n_proj = len(set(r['项目名称'] for r in all_results if r['项目名称']))
        print(f"✅ 报告已保存: {output_path}")
        print(f"   共 {len(all_results)} 条记录, 覆盖 {n_proj} 个项目")
    else:
        print("⚠️  未命中任何记录, 未生成报告。")

    if warn_projects:
        print()
        print("⚠️  以下项目来源为 [Word/PDF正文], 表格未抽出包级数据, 建议手动核查:")
        for p in warn_projects:
            print(f"     · {p}")

    print("=" * 68)


if __name__ == "__main__":
    main()

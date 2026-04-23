#!/usr/bin/env python3
"""
Markdown to Word Document Converter

Converts Markdown files to professionally formatted Word documents (.docx).

Usage:
    python md_to_word.py <input.md> [output.docx]

Examples:
    python md_to_word.py report.md
    python md_to_word.py report.md output.docx
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name='微软雅黑', font_size=10.5, bold=False):
    """Set font properties for a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level):
    """Add a heading with appropriate formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    
    font_sizes = {
        0: 18,  # Main title
        1: 14,  # Section
        2: 12,  # Subsection
        3: 11,  # Sub-subsection
    }
    
    font_size = font_sizes.get(level, 11)
    set_run_font(run, '微软雅黑', font_size, True)
    
    if level == 0:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(12)
    else:
        para.space_before = Pt(12 if level == 1 else 6)
        para.space_after = Pt(6)


def add_paragraph(doc, text):
    """Add a regular paragraph with bold text support."""
    para = doc.add_paragraph()
    
    if '**' in text:
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                set_run_font(run, '微软雅黑', 10.5, True)
            else:
                run = para.add_run(part)
                set_run_font(run, '微软雅黑', 10.5)
    else:
        run = para.add_run(text)
        set_run_font(run, '微软雅黑', 10.5)
    
    para.space_after = Pt(6)


def set_cell_text(cell, text, bold=False, center=False):
    """Set text in a table cell with formatting."""
    cell.text = ''
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, '微软雅黑', 10, bold)


def set_table_border(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_table(doc, table_data):
    """Add a formatted table to the document."""
    if len(table_data) <= 1:
        return
    
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    set_table_border(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_text(cell, cell_text, bold=True, center=True)
                # Blue header background
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9E2F3')
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                set_cell_text(cell, cell_text, center=(col_idx == 0))
    
    doc.add_paragraph()  # Space after table


def convert_md_to_word(input_path, output_path=None):
    """Convert a Markdown file to a Word document."""
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)
    
    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    # Read Markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Process lines
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Table separator, skip but keep in_table state
        if line.startswith('|---') or line.startswith('| ---'):
            i += 1
            continue
        
        # Handle table accumulation - end table when non-table line encountered
        if in_table and not line.startswith('|'):
            if table_data:
                add_table(doc, table_data)
                table_data = []
            in_table = False
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Skip horizontal rules
        if line.startswith('---'):
            i += 1
            continue
        
        # Blockquotes
        if line.startswith('>'):
            text = line[1:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.italic = True
            set_run_font(run, '微软雅黑', 10)
            para.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue
        
        # Headers
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            text = line[level:].strip()
            add_heading(doc, text, level)
            i += 1
            continue
        
        # Tables - parse table rows (excluding separator lines)
        if line.startswith('|') and '|---' not in line:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:  # Only add if there are actual cells
                table_data.append(cells)
                in_table = True
            i += 1
            continue
        
        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(text)
            set_run_font(run, '微软雅黑', 10.5)
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(text)
            set_run_font(run, '微软雅黑', 10.5)
            i += 1
            continue
        
        # Regular paragraph
        add_paragraph(doc, line)
        i += 1
    
    # Handle remaining table
    if table_data:
        add_table(doc, table_data)
    
    # Save document
    doc.save(output_path)
    
    # Report result
    file_size = output_path.stat().st_size
    print(f"[OK] Word document created: {output_path}")
    print(f"     Size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
    
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python md_to_word.py <input.md> [output.docx]")
        print("\nExamples:")
        print("  python md_to_word.py report.md")
        print("  python md_to_word.py report.md output.docx")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    convert_md_to_word(input_path, output_path)


if __name__ == '__main__':
    main()

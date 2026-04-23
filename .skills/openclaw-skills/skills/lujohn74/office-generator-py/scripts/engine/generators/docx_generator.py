from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List


def generate_docx(spec: Dict[str, Any], out_path: str) -> Path:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError('Missing dependency: python-docx. Install requirements first.') from exc

    document = Document()
    title = spec.get('title') or 'Untitled Document'
    document.add_heading(title, 0)

    sections: List[Dict[str, Any]] = spec.get('contentSpec', {}).get('sections', [])
    for section in sections:
        heading = section.get('heading')
        if heading:
            document.add_heading(heading, level=1)
        for paragraph in section.get('paragraphs', []):
            document.add_paragraph(paragraph)
        for bullet in section.get('bullets', []):
            document.add_paragraph(str(bullet), style='List Bullet')
        table_data = section.get('table', [])
        if table_data:
            rows = len(table_data)
            cols = max(len(r) for r in table_data)
            table = document.add_table(rows=rows, cols=cols)
            for r_idx, row in enumerate(table_data):
                for c_idx, value in enumerate(row):
                    table.cell(r_idx, c_idx).text = '' if value is None else str(value)
        for image in section.get('images', []):
            image_path = image.get('path')
            if not image_path:
                continue
            width_inches = float(image.get('widthInches', 5.5))
            path = Path(image_path)
            if path.exists():
                document.add_picture(str(path), width=Inches(width_inches))
            else:
                document.add_paragraph(f'[Missing image: {image_path}]')

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path

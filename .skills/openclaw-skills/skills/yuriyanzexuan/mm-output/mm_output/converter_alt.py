"""
Alternative converter using Selenium + ChromeDriver.
This is a fallback when Playwright has issues with system dependencies.
"""

import os
import re
from pathlib import Path
from typing import Optional, Union, List


class MMOutputGeneratorAlt:
    """
    Alternative multi-modal output generator using Selenium.
    """
    
    def __init__(self, chrome_path: Optional[str] = None):
        """
        Initialize the generator.
        
        Args:
            chrome_path: Path to Chrome/Chromium executable.
                        Can also be set via CHROME_EXECUTABLE_PATH environment variable.
        """
        self.name = "MMOutputGeneratorAlt"
        self._chrome_path = chrome_path or os.getenv("CHROME_EXECUTABLE_PATH")
        self._driver = None
        
    def _ensure_selenium(self):
        """Ensure selenium is installed and driver is created."""
        if self._driver is not None:
            return
            
        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.service import Service
            from selenium.webdriver.chrome.options import Options
        except ImportError:
            raise ImportError(
                "selenium is required. Install with: pip install selenium"
            )
        
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-gpu")
        
        # Use custom Chrome if provided
        if self._chrome_path:
            chrome_options.binary_location = self._chrome_path
            print(f"[{self.name}] Using custom Chrome: {self._chrome_path}")
        
        # Try to use webdriver-manager for automatic chromedriver
        try:
            from webdriver_manager.chrome import ChromeDriverManager
            service = Service(ChromeDriverManager().install())
        except Exception as e:
            print(f"[{self.name}] Warning: Could not use webdriver-manager: {e}")
            print(f"[{self.name}] Trying system chromedriver...")
            service = Service()
        
        self._driver = webdriver.Chrome(service=service, options=chrome_options)
        
    def _close_driver(self):
        """Close the browser driver."""
        if self._driver:
            self._driver.quit()
            self._driver = None
            
    def __enter__(self):
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        self._close_driver()
        
    def html_to_pdf(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        page_size: str = "A4",
        margin: dict = None,
        landscape: bool = False,
        print_background: bool = True
    ) -> str:
        """Convert HTML to PDF using Selenium."""
        self._ensure_selenium()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Load HTML
        self._driver.get(f"file://{html_path}")
        
        # Wait for page to load
        import time
        time.sleep(2)
        
        # Set up print options
        print_options = {
            "printBackground": print_background,
            "preferCSSPageSize": True,
        }
        
        # Page size mapping
        page_sizes = {
            "A4": {"width": 21.0, "height": 29.7},
            "A3": {"width": 29.7, "height": 42.0},
            "A5": {"width": 14.8, "height": 21.0},
            "Letter": {"width": 21.6, "height": 27.9},
            "Legal": {"width": 21.6, "height": 35.6},
        }
        
        if page_size in page_sizes:
            size = page_sizes[page_size]
            if landscape:
                size = {"width": size["height"], "height": size["width"]}
            print_options["paperWidth"] = size["width"]
            print_options["paperHeight"] = size["height"]
        
        # Set margins
        if margin:
            print_options["marginTop"] = margin.get("top", "1cm").replace("cm", "") if isinstance(margin.get("top"), str) else 1
            print_options["marginBottom"] = margin.get("bottom", "1cm").replace("cm", "") if isinstance(margin.get("bottom"), str) else 1
            print_options["marginLeft"] = margin.get("left", "1cm").replace("cm", "") if isinstance(margin.get("left"), str) else 1
            print_options["marginRight"] = margin.get("right", "1cm").replace("cm", "") if isinstance(margin.get("right"), str) else 1
        
        # Generate PDF using Chrome DevTools Protocol
        result = self._driver.execute_cdp_cmd("Page.printToPDF", print_options)
        
        import base64
        pdf_data = base64.b64decode(result["data"])
        output_path.write_bytes(pdf_data)
        
        print(f"[{self.name}] PDF generated: {output_path}")
        return str(output_path)
        
    def html_to_png(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        full_page: bool = True,
        viewport_size: tuple = (1200, 1600),
        wait_time: int = 2000
    ) -> str:
        """Convert HTML to PNG using Selenium."""
        self._ensure_selenium()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Set viewport size
        self._driver.set_window_size(viewport_size[0], viewport_size[1])
        
        # Load HTML
        self._driver.get(f"file://{html_path}")
        
        # Wait for page to load
        import time
        time.sleep(wait_time / 1000)
        
        if full_page:
            # Get full page height
            total_height = self._driver.execute_script(
                "return Math.max(document.body.scrollHeight, document.documentElement.scrollHeight);"
            )
            self._driver.set_window_size(viewport_size[0], total_height + 100)
            time.sleep(0.5)
        
        # Take screenshot
        self._driver.save_screenshot(str(output_path))
        
        print(f"[{self.name}] PNG generated: {output_path}")
        return str(output_path)
        
    def html_to_docx(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        include_images: bool = True
    ) -> str:
        """Convert HTML to DOCX."""
        # Reuse the same implementation from converter.py
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required. Install with: pip install beautifulsoup4")
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        html_content = html_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        doc = Document()
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        body = soup.find('body') or soup
        
        for element in body.descendants:
            if element.name is None:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading(level=min(level, 9))
                p.add_run(element.get_text(strip=True))
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'img' and include_images:
                img_src = element.get('src', '')
                if img_src:
                    if not img_src.startswith(('http://', 'https://', 'data:')):
                        img_path = html_path.parent / img_src
                        if img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5.0))
                            except Exception as e:
                                print(f"[{self.name}] Warning: Could not add image {img_src}: {e}")
            elif element.name == 'li':
                parent = element.find_parent(['ul', 'ol'])
                if parent:
                    text = element.get_text(strip=True)
                    if text:
                        style = 'List Number' if parent.name == 'ol' else 'List Bullet'
                        doc.add_paragraph(text, style=style)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    first_row = rows[0]
                    cols = len(first_row.find_all(['td', 'th']))
                    if cols > 0:
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = 'Light Grid Accent 1'
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    table.rows[i].cells[j].text = cell.get_text(strip=True)
        
        doc.save(str(output_path))
        print(f"[{self.name}] DOCX generated: {output_path}")
        return str(output_path)
        
    def convert_all(
        self,
        html_path: Union[str, Path],
        output_dir: Union[str, Path],
        base_name: Optional[str] = None,
        formats: Optional[List[str]] = None
    ) -> dict:
        """Convert HTML to all supported formats."""
        html_path = Path(html_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if base_name is None:
            base_name = html_path.stem
            
        formats = formats or ['pdf', 'png', 'docx']
        results = {}
        
        try:
            for fmt in formats:
                fmt = fmt.lower()
                output_path = output_dir / f"{base_name}.{fmt}"
                
                if fmt == 'pdf':
                    results['pdf'] = self.html_to_pdf(html_path, output_path)
                elif fmt == 'png':
                    results['png'] = self.html_to_png(html_path, output_path)
                elif fmt == 'docx':
                    results['docx'] = self.html_to_docx(html_path, output_path)
                else:
                    print(f"[{self.name}] Warning: Unsupported format '{fmt}', skipping")
        finally:
            self._close_driver()
            
        return results


# Standalone convenience functions
def html_to_pdf(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGeneratorAlt() as gen:
        return gen.html_to_pdf(html_path, output_path, **kwargs)


def html_to_png(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGeneratorAlt() as gen:
        return gen.html_to_png(html_path, output_path, **kwargs)


def html_to_docx(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGeneratorAlt() as gen:
        return gen.html_to_docx(html_path, output_path, **kwargs)


def convert_all(html_path: Union[str, Path], output_dir: Union[str, Path], **kwargs) -> dict:
    with MMOutputGeneratorAlt() as gen:
        return gen.convert_all(html_path, output_dir, **kwargs)

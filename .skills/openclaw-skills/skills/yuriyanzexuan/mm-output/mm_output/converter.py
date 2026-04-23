"""Core converter module for multi-modal output generation."""

import os
import re
import io
import base64
from pathlib import Path
from typing import Optional, Union, List
from urllib.parse import urljoin, urlparse

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup


# ========== DOCX Helper Functions ==========

def _add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1")
    rPr.append(color_el)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), "20")
    rPr.append(sz_el)
    run_elem.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")
    run_elem.append(t_el)
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def _add_inline_cite(paragraph, text: str):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run.font.superscript = True


def _add_rich_paragraph(doc, element, style=None):
    p = doc.add_paragraph(style=style)
    for child in element.children:
        if isinstance(child, str):
            text = child.strip()
            if text:
                p.add_run(text)
        elif child.name == "a":
            href = (child.get("href") or "").strip()
            link_text = child.get_text(strip=True)
            if href.startswith("#") and link_text:
                _add_inline_cite(p, link_text)
            elif href and link_text:
                _add_hyperlink(p, link_text, href)
            elif link_text:
                p.add_run(link_text)
        elif child.name == "span":
            text = child.get_text(strip=True)
            if text:
                classes = child.get("class", [])
                if "cite-ref" in classes:
                    _add_inline_cite(p, text)
                else:
                    p.add_run(text)
        elif child.name in ("strong", "b"):
            run = p.add_run(child.get_text(strip=True))
            run.bold = True
        elif child.name in ("em", "i"):
            run = p.add_run(child.get_text(strip=True))
            run.italic = True
        else:
            text = child.get_text(strip=True)
            if text:
                p.add_run(text)
    return p


# ========== Main Class ==========

class MMOutputGenerator:
    """Multi-modal output generator for converting HTML to PDF, PNG, and DOCX formats."""
    
    def __init__(self, chrome_path: Optional[str] = None):
        self.name = "MMOutputGenerator"
        self._chrome_path = chrome_path or os.getenv("CHROME_EXECUTABLE_PATH")
        self._playwright = None
        self._browser = None
        
    def _ensure_playwright(self):
        if self._browser is not None:
            return
            
        self._playwright = sync_playwright().start()
        
        launch_options = {"headless": True}
        if self._chrome_path:
            launch_options["executable_path"] = self._chrome_path
            print(f"[{self.name}] Using custom Chrome: {self._chrome_path}")
        else:
            chrome_paths = [
                "/usr/bin/google-chrome",
                "/usr/bin/chromium",
                "/usr/bin/chromium-browser",
                "./chrome-linux64/chrome",
            ]
            for path in chrome_paths:
                if Path(path).exists():
                    launch_options["executable_path"] = path
                    print(f"[{self.name}] Auto-detected Chrome: {path}")
                    break
            
        self._browser = self._playwright.chromium.launch(**launch_options)
        
    def _close_browser(self):
        if self._browser:
            self._browser.close()
            self._browser = None
        if self._playwright:
            self._playwright.stop()
            self._playwright = None
            
    def __enter__(self):
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        self._close_browser()

    def _inject_cjk_font_fallback(self, page):
        page.add_style_tag(content="""
            body, h1, h2, h3, h4, h5, h6, p, span, li, td, th, div,
            blockquote, strong, em, a, small, label {
                font-family:
                    "WenQuanYi Zen Hei",
                    "Noto Sans CJK SC",
                    "Noto Sans CJK",
                    "Source Han Sans SC",
                    "PingFang SC",
                    "Microsoft YaHei",
                    "Inter",
                    sans-serif !important;
            }
        """)

    @staticmethod
    def _disable_scroll_animations(page):
        page.add_style_tag(content="""
            [data-aos], [data-aos][data-aos] {
                opacity: 1 !important;
                transform: none !important;
                transition: none !important;
            }
            .aos-animate, .aos-init {
                opacity: 1 !important;
                transform: none !important;
            }
            * {
                animation-play-state: paused !important;
                transition-duration: 0s !important;
                animation-duration: 0s !important;
            }
        """)
        page.evaluate("""() => {
            if (typeof AOS !== 'undefined') {
                try { AOS.init({ disable: true }); } catch(e) {}
            }
            document.querySelectorAll('[data-aos]').forEach(el => {
                el.removeAttribute('data-aos');
                el.classList.add('aos-animate');
                el.style.opacity = '1';
                el.style.transform = 'none';
            });
        }""")
        
    def html_to_pdf(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        page_size: str = "A4",
        margin: dict = None,
        landscape: bool = False,
        print_background: bool = True
    ) -> str:
        self._ensure_playwright()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        context_options = {"viewport": {"width": 1200, "height": 1600}}
        context = self._browser.new_context(**context_options)
        
        try:
            page = context.new_page()
            page.goto(f"file://{html_path}", wait_until="networkidle")
            self._inject_cjk_font_fallback(page)
            self._disable_scroll_animations(page)
            page.wait_for_timeout(1000)
            
            default_margin = {"top": "1cm", "right": "1cm", "bottom": "1cm", "left": "1cm"}
            if margin:
                default_margin.update(margin)
            
            page.pdf(
                path=str(output_path),
                format=page_size,
                margin=default_margin,
                print_background=print_background,
                landscape=landscape,
            )
            
            print(f"[{self.name}] PDF generated: {output_path}")
            return str(output_path)
            
        finally:
            context.close()
            
    def html_to_png(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        full_page: bool = True,
        viewport_size: tuple = (1200, 1600),
        wait_time: int = 2000
    ) -> str:
        self._ensure_playwright()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        context = self._browser.new_context(
            viewport={"width": viewport_size[0], "height": viewport_size[1]}
        )
        
        try:
            page = context.new_page()
            page.goto(f"file://{html_path}", wait_until="networkidle")
            self._inject_cjk_font_fallback(page)
            self._disable_scroll_animations(page)
            page.wait_for_timeout(wait_time)
            
            page.screenshot(path=str(output_path), full_page=full_page, type="png")
            
            print(f"[{self.name}] PNG generated: {output_path}")
            return str(output_path)
            
        finally:
            context.close()
            
    def html_to_docx(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        include_images: bool = True
    ) -> str:
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        html_content = html_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        doc = Document()
        
        for script in soup(["script", "style"]):
            script.decompose()
            
        body = soup.find('body') or soup
        
        for element in body.descendants:
            if element.name is None:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading(level=min(level, 9))
                p.add_run(element.get_text(strip=True))
                
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    if element.find("a"):
                        _add_rich_paragraph(doc, element)
                    else:
                        doc.add_paragraph(text)
                    
            elif element.name == 'img' and include_images:
                img_src = element.get('src', '')
                if img_src and not img_src.startswith(('http://', 'https://', 'data:')):
                    img_path = html_path.parent / img_src
                    if img_path.exists():
                        doc.add_picture(str(img_path), width=Inches(5.0))
                                
            elif element.name == 'li':
                parent = element.find_parent(['ul', 'ol'])
                if parent:
                    text = element.get_text(strip=True)
                    if text:
                        style = 'List Number' if parent.name == 'ol' else 'List Bullet'
                        if element.find("a"):
                            _add_rich_paragraph(doc, element, style=style)
                        else:
                            doc.add_paragraph(text, style=style)
                        
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    first_row = rows[0]
                    cols = len(first_row.find_all(['td', 'th']))
                    if cols > 0:
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = 'Light Grid Accent 1'
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    table.rows[i].cells[j].text = cell.get_text(strip=True)
                                    
        doc.save(str(output_path))
        
        print(f"[{self.name}] DOCX generated: {output_path}")
        return str(output_path)

    def html_to_pptx(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        max_slides: int = 8
    ) -> str:
        from .pptx_converter import PPTXConverter

        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        converter = PPTXConverter()
        result = converter.generate_pptx(html_path, output_path, max_slides=max_slides)
        print(f"[{self.name}] PPTX generated: {result}")
        return result
        
    def convert_all(
        self,
        html_path: Union[str, Path],
        output_dir: Union[str, Path],
        base_name: Optional[str] = None,
        formats: Optional[List[str]] = None
    ) -> dict:
        html_path = Path(html_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if base_name is None:
            base_name = html_path.stem
            
        formats = formats or ['pdf', 'png', 'docx', 'pptx']
        results = {}
        
        try:
            for fmt in formats:
                fmt = fmt.lower()
                output_path = output_dir / f"{base_name}.{fmt}"
                
                if fmt == 'pdf':
                    results['pdf'] = self.html_to_pdf(html_path, output_path)
                elif fmt == 'png':
                    results['png'] = self.html_to_png(html_path, output_path)
                elif fmt == 'docx':
                    results['docx'] = self.html_to_docx(html_path, output_path)
                elif fmt == 'pptx':
                    results['pptx'] = self.html_to_pptx(html_path, output_path)
                else:
                    print(f"[{self.name}] Warning: Unsupported format '{fmt}', skipping")
                    
        finally:
            self._close_browser()
            
        return results


# ========== Convenience Functions ==========

def html_to_pdf(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGenerator() as gen:
        return gen.html_to_pdf(html_path, output_path, **kwargs)


def html_to_png(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGenerator() as gen:
        return gen.html_to_png(html_path, output_path, **kwargs)


def html_to_docx(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGenerator() as gen:
        return gen.html_to_docx(html_path, output_path, **kwargs)


def html_to_pptx(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    with MMOutputGenerator() as gen:
        return gen.html_to_pptx(html_path, output_path, **kwargs)


def convert_all(html_path: Union[str, Path], output_dir: Union[str, Path], **kwargs) -> dict:
    with MMOutputGenerator() as gen:
        return gen.convert_all(html_path, output_dir, **kwargs)

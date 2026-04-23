"""
Alternative converter using WeasyPrint (pure Python, no Chrome needed).
This is a fallback for systems where Chrome dependencies are not available.
"""

import os
import re
from pathlib import Path
from typing import Optional, Union, List


class MMOutputGeneratorWeasyPrint:
    """
    Multi-modal output generator using WeasyPrint for PDF generation.
    No Chrome/Chromium required - uses pure Python libraries.
    """
    
    def __init__(self):
        """Initialize the generator."""
        self.name = "MMOutputGeneratorWeasyPrint"
        
    def _check_weasyprint(self):
        """Check if weasyprint is installed."""
        try:
            import weasyprint
        except ImportError:
            raise ImportError(
                "weasyprint is required for PDF generation. "
                "Install with: pip install weasyprint"
            )
        return weasyprint
        
    def html_to_pdf(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        page_size: str = "A4",
        margin: dict = None,
        landscape: bool = False,
        print_background: bool = True
    ) -> str:
        """
        Convert HTML to PDF using WeasyPrint.
        
        Note: WeasyPrint has some CSS limitations compared to Chrome,
        but works well for standard HTML/CSS.
        """
        weasyprint = self._check_weasyprint()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Read HTML content
        html_content = html_path.read_text(encoding='utf-8')
        
        # Fix relative paths in HTML
        base_url = html_path.parent.as_uri()
        html_content = self._fix_relative_paths(html_content, html_path.parent)
        
        # Create WeasyPrint HTML object
        html = weasyprint.HTML(string=html_content, base_url=base_url)
        
        # Set up page size
        page_sizes = {
            "A4": (210, 297),
            "A3": (297, 420),
            "A5": (148, 210),
            "Letter": (216, 279),
            "Legal": (216, 356),
        }
        
        size = page_sizes.get(page_size, (210, 297))
        if landscape:
            size = (size[1], size[0])
        
        # Set margins
        default_margin = "1cm"
        margins = {
            "top": margin.get("top", default_margin) if margin else default_margin,
            "right": margin.get("right", default_margin) if margin else default_margin,
            "bottom": margin.get("bottom", default_margin) if margin else default_margin,
            "left": margin.get("left", default_margin) if margin else default_margin,
        }
        
        # Generate PDF
        css = f"""
        @page {{
            size: {size[0]}mm {size[1]}mm;
            margin: {margins["top"]} {margins["right"]} {margins["bottom"]} {margins["left"]};
        }}
        """
        
        html.write_pdf(str(output_path), stylesheets=[weasyprint.CSS(string=css)])
        
        print(f"[{self.name}] PDF generated: {output_path}")
        return str(output_path)
        
    def _fix_relative_paths(self, html_content: str, base_path: Path) -> str:
        """Fix relative image paths in HTML to be absolute."""
        # Fix src="assets/..." to absolute paths
        def replace_src(match):
            src = match.group(1)
            if src.startswith(('http://', 'https://', 'data:', 'file://')):
                return match.group(0)
            # Make it absolute
            abs_path = (base_path / src).resolve()
            if abs_path.exists():
                return f'src="{abs_path.as_uri()}"'
            return match.group(0)
        
        html_content = re.sub(r'src="([^"]*)"', replace_src, html_content)
        html_content = re.sub(r"src='([^']*)'", replace_src, html_content)
        
        return html_content
        
    def html_to_png(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        full_page: bool = True,
        viewport_size: tuple = (1200, 1600),
        wait_time: int = 2000
    ) -> str:
        """
        Convert HTML to PNG using imgkit (wkhtmltoimage).
        Falls back to PDF + pdf2image if imgkit not available.
        """
        try:
            return self._html_to_png_imgkit(html_path, output_path, viewport_size)
        except Exception as e:
            print(f"[{self.name}] imgkit failed: {e}, trying alternative...")
            return self._html_to_png_pdf2image(html_path, output_path, viewport_size)
            
    def _html_to_png_imgkit(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        viewport_size: tuple = (1200, 1600)
    ) -> str:
        """Convert HTML to PNG using imgkit."""
        try:
            import imgkit
        except ImportError:
            raise ImportError("imgkit is required. Install with: pip install imgkit")
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        options = {
            'format': 'png',
            'width': viewport_size[0],
            'encoding': 'UTF-8',
            'enable-local-file-access': '',
        }
        
        imgkit.from_file(str(html_path), str(output_path), options=options)
        
        print(f"[{self.name}] PNG generated: {output_path}")
        return str(output_path)
        
    def _html_to_png_pdf2image(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        viewport_size: tuple = (1200, 1600)
    ) -> str:
        """Convert HTML to PNG using PDF as intermediate."""
        try:
            from pdf2image import convert_from_path
        except ImportError:
            raise ImportError("pdf2image is required. Install with: pip install pdf2image")
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create temporary PDF
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp_pdf = tmp.name
        
        try:
            # Generate PDF first
            self.html_to_pdf(html_path, tmp_pdf, page_size="A4")
            
            # Convert PDF to image
            images = convert_from_path(tmp_pdf, dpi=150)
            if images:
                images[0].save(str(output_path), 'PNG')
            else:
                raise RuntimeError("pdf2image returned empty result")
                
        finally:
            # Cleanup
            try:
                os.unlink(tmp_pdf)
            except:
                pass
        
        print(f"[{self.name}] PNG generated (via PDF): {output_path}")
        return str(output_path)
        
    def html_to_docx(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        include_images: bool = True
    ) -> str:
        """Convert HTML to DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required. Install with: pip install beautifulsoup4")
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        html_content = html_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        doc = Document()
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        body = soup.find('body') or soup
        
        for element in body.descendants:
            if element.name is None:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading(level=min(level, 9))
                p.add_run(element.get_text(strip=True))
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'img' and include_images:
                img_src = element.get('src', '')
                if img_src:
                    if not img_src.startswith(('http://', 'https://', 'data:')):
                        img_path = html_path.parent / img_src
                        if img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5.0))
                            except Exception as e:
                                print(f"[{self.name}] Warning: Could not add image {img_src}: {e}")
            elif element.name == 'li':
                parent = element.find_parent(['ul', 'ol'])
                if parent:
                    text = element.get_text(strip=True)
                    if text:
                        style = 'List Number' if parent.name == 'ol' else 'List Bullet'
                        doc.add_paragraph(text, style=style)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    first_row = rows[0]
                    cols = len(first_row.find_all(['td', 'th']))
                    if cols > 0:
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = 'Light Grid Accent 1'
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    table.rows[i].cells[j].text = cell.get_text(strip=True)
        
        doc.save(str(output_path))
        print(f"[{self.name}] DOCX generated: {output_path}")
        return str(output_path)
        
    def convert_all(
        self,
        html_path: Union[str, Path],
        output_dir: Union[str, Path],
        base_name: Optional[str] = None,
        formats: Optional[List[str]] = None
    ) -> dict:
        """Convert HTML to all supported formats."""
        html_path = Path(html_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if base_name is None:
            base_name = html_path.stem
            
        formats = formats or ['pdf', 'png', 'docx']
        results = {}
        
        for fmt in formats:
            fmt = fmt.lower()
            output_path = output_dir / f"{base_name}.{fmt}"
            
            try:
                if fmt == 'pdf':
                    results['pdf'] = self.html_to_pdf(html_path, output_path)
                elif fmt == 'png':
                    results['png'] = self.html_to_png(html_path, output_path)
                elif fmt == 'docx':
                    results['docx'] = self.html_to_docx(html_path, output_path)
                else:
                    print(f"[{self.name}] Warning: Unsupported format '{fmt}', skipping")
            except Exception as e:
                print(f"[{self.name}] Error generating {fmt.upper()}: {e}")
                    
        return results


# Standalone convenience functions
def html_to_pdf(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    gen = MMOutputGeneratorWeasyPrint()
    return gen.html_to_pdf(html_path, output_path, **kwargs)


def html_to_png(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    gen = MMOutputGeneratorWeasyPrint()
    return gen.html_to_png(html_path, output_path, **kwargs)


def html_to_docx(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    gen = MMOutputGeneratorWeasyPrint()
    return gen.html_to_docx(html_path, output_path, **kwargs)


def convert_all(html_path: Union[str, Path], output_dir: Union[str, Path], **kwargs) -> dict:
    gen = MMOutputGeneratorWeasyPrint()
    return gen.convert_all(html_path, output_dir, **kwargs)

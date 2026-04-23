#!/usr/bin/env python3
"""
公文排版工具 - 根据公文格式规范自动排版Word文档
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 公文格式配置
FONT_CONFIG = {
    'title': {'name': '方正小标宋_GBK', 'size': Pt(22), 'bold': False},
    'body': {'name': '方正仿宋_GBK', 'size': Pt(16), 'bold': False},
    'h1': {'name': '方正黑体_GBK', 'size': Pt(16), 'bold': False},
    'h2': {'name': '方正楷体_GBK', 'size': Pt(16), 'bold': False},
    'h3': {'name': '方正仿宋_GBK', 'size': Pt(16), 'bold': True},
    'h4': {'name': '方正仿宋_GBK', 'size': Pt(16), 'bold': False},
    'table_title': {'name': '方正小标宋_GBK', 'size': Pt(15), 'bold': False},
    'table_header': {'name': '方正黑体_GBK', 'size': Pt(14), 'bold': False},
    'table_body': {'name': '方正仿宋_GBK', 'size': Pt(14), 'bold': False},
    'page_number': {'name': '宋体', 'size': Pt(15), 'bold': False},
}

# 页边距 (mm转英寸)
MARGIN_CONFIG = {
    'top': Inches(1.457),    # 37mm
    'bottom': Inches(1.378), # 35mm
    'left': Inches(1.102),   # 28mm
    'right': Inches(1.024),  # 26mm
}

# 行距配置
LINE_SPACING = 28  # 固定值28磅


def detect_heading_level(text):
    """检测标题层级"""
    text = text.strip()
    
    # 一级标题: 一、二、三...
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'h1'
    # 二级标题: （一）（二）...
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'h2'
    # 三级标题: 1. 1、...
    if re.match(r'^\d+[.、]', text):
        return 'h3'
    # 四级标题: （1）（2）...
    if re.match(r'^（\d+）', text):
        return 'h4'
    
    return None


def is_title_line(text):
    """判断是否标题行（用于标题检测）"""
    text = text.strip()
    # 短标题行（排除正文首行）
    if len(text) < 50 and not text.startswith('附件'):
        if detect_heading_level(text):
            return True
        # 也可能是纯标题（无序号）
        if len(text) < 30 and not any(c in text for c in '。；；：'):
            return True
    return False


def is_date_line(text):
    """判断是否日期行"""
    text = text.strip()
    return bool(re.match(r'^（\d+年\d+月\d+日）$', text))


def set_paragraph_format(paragraph, font_name, font_size, bold=False, 
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                         first_line_indent=None,
                         line_space=LINE_SPACING):
    """设置段落格式"""
    # 设置字体
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = font_size
        run.font.bold = bold
    
    # 设置对齐
    paragraph.alignment = alignment
    
    # 设置首行缩进
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = first_line_indent
    else:
        paragraph.paragraph_format.first_line_indent = Inches(0)
    
    # 设置行距
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(line_space)
    
    # 设置段前段后间距
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def format_document(input_path, output_path):
    """排版Word文档"""
    doc = Document(input_path)
    
    # 设置页边距
    for section in doc.sections:
        section.top_margin = MARGIN_CONFIG['top']
        section.bottom_margin = MARGIN_CONFIG['bottom']
        section.left_margin = MARGIN_CONFIG['left']
        section.right_margin = MARGIN_CONFIG['right']
    
    # 处理所有段落
    paragraphs = doc.paragraphs
    
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        # 检测标题层级
        heading_level = detect_heading_level(text)
        
        if heading_level:
            # 标题格式化
            config = FONT_CONFIG[heading_level]
            set_paragraph_format(
                paragraph, 
                config['name'], 
                config['size'], 
                config['bold'],
                alignment=WD_ALIGN_PARAGRAPH.LEFT
            )
        elif is_date_line(text):
            # 日期行居中
            config = FONT_CONFIG['body']
            set_paragraph_format(
                paragraph,
                config['name'],
                config['size'],
                config['bold'],
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
        elif i == 0 and len(text) < 50:
            # 第一个短段落可能是标题
            config = FONT_CONFIG['title']
            set_paragraph_format(
                paragraph,
                config['name'],
                config['size'],
                config['bold'],
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
        else:
            # 正文格式化
            config = FONT_CONFIG['body']
            set_paragraph_format(
                paragraph,
                config['name'],
                config['size'],
                config['bold'],
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Inches(0.5)  # 首行缩进2字符
            )
    
    # 格式化表格
    for table in doc.tables:
        # 设置表格标题（如果表格在段落后面且段落文字包含"表"）
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        # 表头格式
                        if row_index == 0:
                            config = FONT_CONFIG['table_header']
                            set_paragraph_format(
                                p, config['name'], config['size'],
                                alignment=WD_ALIGN_PARAGRAPH.CENTER
                            )
                        else:
                            config = FONT_CONFIG['table_body']
                            set_paragraph_format(
                                p, config['name'], config['size']
                            )
    
    # 保存文件
    doc.save(output_path)
    print(f"✅ 排版完成: {output_path}")
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python3 format_gw.py <输入文件.docx> <输出文件.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    try:
        format_document(input_file, output_file)
    except Exception as e:
        print(f"❌ 排版失败: {e}")
        sys.exit(1)

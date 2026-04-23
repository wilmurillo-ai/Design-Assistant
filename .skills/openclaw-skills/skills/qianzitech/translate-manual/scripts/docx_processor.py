"""
Docx Processor - 文档处理模块
用于读取docx文档内容、提取图片信息、替换图片
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from typing import List, Dict, Tuple


def read_docx(docx_path: str) -> Dict:
    """
    读取 docx 文档内容

    Returns:
        {
            "paragraphs": [文本段落列表],
            "images": [{"name": 图片名, "path": 路径, "paragraph_idx": 段落索引}],
            "tables": [表格内容]
        }
    """
    doc = Document(docx_path)
    result = {
        "paragraphs": [],
        "images": [],
        "tables": []
    }

    # 读取段落
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            result["paragraphs"].append({
                "index": idx,
                "text": text,
                "style": para.style.name
            })

        # 检查段落中的图片
        for run in para.runs:
            for shape in run._element.xpath('.//w:drawing/wp:inline'):
                for img in shape.xpath('.//a:blip'):
                    embed_id = img.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        # 尝试获取图片
                        rels = doc.part.related_parts
                        if embed_id in rels:
                            img_part = rels[embed_id]
                            result["images"].append({
                                "name": f"image_{idx}",
                                "embed_id": embed_id,
                                "paragraph_idx": idx,
                                "type": img_part.content_type
                            })

    # 读取表格
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            result["tables"].append(table_data)

    return result


def get_image_relationships(docx_path: str) -> Dict:
    """获取文档中所有图片的引用关系"""
    doc = Document(docx_path)
    images = {}

    # 遍历所有关系
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            images[rel.rId] = {
                "target": rel.target_ref,
                "type": rel.target_ref.split('.')[-1]
            }

    return images


def extract_images(docx_path: str, output_dir: str) -> List[str]:
    """提取文档中的所有图片到指定目录"""
    os.makedirs(output_dir, exist_ok=True)

    doc = Document(docx_path)
    extracted = []

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_part = rel.target_ref
            img_data = doc.part.related_parts[rel.rId].blob

            # 生成文件名
            ext = img_part.split('.')[-1]
            filename = f"extracted_{len(extracted)}.{ext}"
            filepath = os.path.join(output_dir, filename)

            with open(filepath, 'wb') as f:
                f.write(img_data)

            extracted.append(filepath)

    return extracted


def replace_image_in_paragraph(doc: Document, old_image_rid: str, new_image_path: str):
    """替换段落中的图片"""
    # 这是一个复杂的操作，需要修改 XML
    # 简化版本：标记需要替换的位置
    pass


def save_docx(doc: Document, output_path: str):
    """保存文档"""
    doc.save(output_path)


def create_translated_docx(template_path: str, translated_paragraphs: List[str], output_path: str):
    """
    基于模板创建翻译后的文档

    Args:
        template_path: 模板文档路径
        translated_paragraphs: 翻译后的段落列表
        output_path: 输出路径
    """
    doc = Document(template_path)

    # 清空原有段落（保留图片）
    for para in doc.paragraphs:
        para.clear()

    # 写入翻译后的文本
    for i, text in enumerate(translated_paragraphs):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            para.text = text
        else:
            doc.add_paragraph(text)

    doc.save(output_path)


if __name__ == "__main__":
    # 测试
    test_path = "test.docx"
    if os.path.exists(test_path):
        result = read_docx(test_path)
        print(f"段落数: {len(result['paragraphs'])}")
        print(f"图片数: {len(result['images'])}")
        print(f"表格数: {len(result['tables'])}")
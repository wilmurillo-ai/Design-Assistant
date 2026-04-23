"""
Manual Translator - 优化版翻译模块
支持分段翻译、表格翻译、进度保存
"""
import requests
import os
from docx import Document
from typing import List, Tuple

# DeepL API 配置
DEEPL_API_TEMPLATE = "https://api.deeplx.org/{api_key}/translate"

LANGUAGE_CODES = {
    "英文": "EN", "英语": "EN", "EN": "EN", "English": "EN",
    "日文": "JA", "日语": "JA", "JA": "JA", "Japanese": "JA",
    "韩文": "KO", "韩语": "KO", "KO": "KO", "Korean": "KO",
    "中文": "ZH", "ZH": "ZH", "Chinese": "ZH",
    "法文": "FR", "法语": "FR", "FR": "FR",
    "德文": "DE", "德语": "DE", "DE": "DE",
    "西班牙": "ES", "ES": "ES",
    "意大利": "IT", "IT": "IT",
    "葡萄牙": "PT", "PT": "PT",
    "俄文": "RU", "RU": "RU",
}

API_KEY_ENV_NAMES = [
    "DEEPLX_API_KEY",
    "DEEPL_API_KEY",
    "TRANSLATOR_API_KEY",
]


def normalize_language(lang: str) -> str:
    """标准化语言代码"""
    return LANGUAGE_CODES.get(lang.strip().upper(), lang.upper())


def has_chinese(text: str) -> bool:
    """检测文本是否包含中文或其他非ASCII字符"""
    return any(ord(c) > 127 for c in text)


def resolve_api_key(cli_api_key: str = "") -> str:
    """解析 API Key：优先命令行参数，其次环境变量"""
    if cli_api_key and cli_api_key.strip():
        return cli_api_key.strip()

    for env_name in API_KEY_ENV_NAMES:
        value = os.getenv(env_name, "").strip()
        if value:
            print(f"Using API key from environment: {env_name}")
            return value

    return ""


def translate_text(text: str, target_lang: str, api_key: str) -> str:
    """翻译单段文本"""
    if not text or not text.strip():
        return text

    target_lang = normalize_language(target_lang)
    url = DEEPL_API_TEMPLATE.format(api_key=api_key)

    try:
        response = requests.post(
            url,
            json={"text": text, "target_lang": target_lang},
            timeout=30
        )
        if response.status_code == 200:
            return response.json().get("data", text)
    except Exception as e:
        print(f"Translation error: {e}")

    return text


def translate_paragraphs分段(doc_path: str, output_path: str, api_key: str, 
                              target_lang: str = "EN", batch_size: int = 50):
    """
    分段翻译文档段落
    
    Args:
        doc_path: 输入文档路径
        output_path: 输出文档路径
        api_key: DeepL API Key
        target_lang: 目标语言
        batch_size: 每批翻译数量
    """
    doc = Document(doc_path)
    
    # 收集所有段落
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    total = len(paragraphs)
    print(f"Total paragraphs: {total}")
    
    # 分批翻译
    translated = []
    for i in range(0, total, batch_size):
        batch = paragraphs[i:i+batch_size]
        for text in batch:
            translated.append(translate_text(text, target_lang, api_key))
        
        # 每批保存中间结果
        print(f"  {min(i+batch_size, total)}/{total}")
    
    # 替换原文档
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text and text in paragraphs:
            idx = paragraphs.index(text)
            p.clear()
            p.add_run(translated[idx])
    
    doc.save(output_path)
    print(f"Saved: {output_path}")


def translate_tables分段(doc_path: str, output_path: str, api_key: str,
                          target_lang: str = "EN"):
    """
    逐表格翻译并保存
    
    Args:
        doc_path: 输入文档路径
        output_path: 输出文档路径
        api_key: DeepL API Key
        target_lang: 目标语言
    """
    doc = Document(doc_path)
    print(f"Tables: {len(doc.tables)}")
    
    for ti, table in enumerate(doc.tables):
        print(f"Table {ti+1}/{len(doc.tables)}...", end=" ", flush=True)
        
        changed = False
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and has_chinese(text):
                    translated = translate_text(text, target_lang, api_key)
                    cell.text = translated
                    changed = True
        
        # 每表格保存，防止中途失败
        doc.save(output_path)
        print("saved")
    
    print(f"Final: {output_path}")


def translate_full分段(doc_path: str, output_path: str, api_key: str,
                       target_lang: str = "EN", batch_size: int = 50):
    """
    完整翻译：段落 + 表格
    
    Args:
        doc_path: 输入文档路径
        output_path: 输出文档路径
        api_key: DeepL API Key
        target_lang: 目标语言
        batch_size: 每批翻译数量
    """
    doc = Document(doc_path)
    print("Step 1: Collecting texts...")
    
    # 收集段落和表格单元格
    texts = []  # (type, element, text)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            texts.append(("para", p, text))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    texts.append(("cell", cell, text))
    
    total = len(texts)
    print(f"Total items: {total}")
    
    # 翻译
    translations = {}
    for i in range(0, total, batch_size):
        batch = texts[i:i+batch_size]
        for item_type, element, text in batch:
            if text not in translations:
                translations[text] = translate_text(text, target_lang, api_key)
        
        done = min(i + batch_size, total)
        print(f"  {done}/{total}")
    
    # 替换段落
    print("Step 2: Replacing paragraphs...")
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and text in translations:
            p.clear()
            p.add_run(translations[text])
    
    # 替换表格
    print("Step 3: Replacing table cells...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text in translations:
                    cell.text = translations[text]
    
    # 保存
    print(f"Step 4: Saving to {output_path}...")
    doc.save(output_path)
    print("DONE!")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python translator.py <input.docx> <output.docx> [api_key]")
        print(f"Or set one of env vars: {', '.join(API_KEY_ENV_NAMES)}")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    cli_api_key = sys.argv[3] if len(sys.argv) > 3 else ""
    api_key = resolve_api_key(cli_api_key)
    
    if not api_key:
        print("Error: API key required")
        print(f"Provide [api_key] arg or set env var: {', '.join(API_KEY_ENV_NAMES)}")
        sys.exit(1)
    
    translate_full分段(input_file, output_file, api_key)
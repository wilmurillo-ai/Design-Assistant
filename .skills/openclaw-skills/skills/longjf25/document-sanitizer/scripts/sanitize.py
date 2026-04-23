"""
文档脱敏脚本 v2 - 统一记录版
改进点：
1. 默认开启文件名脱敏
2. 统一脱敏记录（固定文件名 _sanitize_record.json，累积合并映射）
3. 新增恢复功能（restore）
4. 仅支持 docx/xlsx，检测到 doc/xls 时提示自动转换
核心思路：对每个 w:t 元素直接操作，合并相邻 w:t 后替换，再写回
"""
import os
import re
import sys
import json
import shutil
import subprocess
from pathlib import Path
from datetime import datetime

# ============================================================
# 配置
# ============================================================

SKIP_DIRS = {"_sanitized_output", "_restored_output", "_文档_md", "_markdown", "update",
             "_文档_md_backup", "node_modules", ".git", "__pycache__"}
DOC_EXTENSIONS = {".docx", ".xlsx"}
LEGACY_EXTENSIONS = {".doc", ".xls"}
RECORD_FILE = "_sanitize_record.json"

# doc_xls2docx_xlsx 技能脚本路径
SKILL_SCRIPTS_DIR = Path.home() / ".workbuddy" / "skills" / "doc_xls2docx_xlsx" / "scripts"

# ============================================================
# 工具函数
# ============================================================

def load_config(workspace):
    config_path = workspace / "_sanitize_config.json"
    if not config_path.exists():
        print(f"[ERROR] 未找到配置文件: {config_path}")
        return None
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    if "exact_rules" not in config:
        config["exact_rules"] = []
    if "regex_rules" not in config:
        config["regex_rules"] = []
    return config


def load_record(workspace):
    """加载已有的统一脱敏记录，不存在则返回空记录结构"""
    record_path = workspace / RECORD_FILE
    if record_path.exists():
        with open(record_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "version": 2,
        "created_at": None,
        "last_updated": None,
        "mapping": {},           # 脱敏值 → 原始值（核心，恢复依赖此表）
        "filename_mapping": {},  # 脱敏文件名 → 原始文件名
        "runs": []               # 运行历史
    }


def save_record(workspace, record):
    """保存统一脱敏记录"""
    record_path = workspace / RECORD_FILE
    with open(record_path, "w", encoding="utf-8") as f:
        json.dump(record, f, ensure_ascii=False, indent=2)


def merge_record(existing, new_mapping, new_filename_mapping, files_processed, timestamp):
    """将本次脱敏结果合并到已有记录中"""
    now_str = timestamp.strftime("%Y-%m-%d %H:%M:%S")

    if existing["created_at"] is None:
        existing["created_at"] = now_str
    existing["last_updated"] = now_str

    # 合并 mapping：脱敏值→原始值
    for key, value in new_mapping.items():
        if key in existing["mapping"] and existing["mapping"][key] != value:
            print(f"  [WARN] 映射冲突: {key} → 旧值={existing['mapping'][key]}, 新值={value}，以新值为准")
        existing["mapping"][key] = value

    # 合并 filename_mapping
    for key, value in new_filename_mapping.items():
        if key in existing["filename_mapping"] and existing["filename_mapping"][key] != value:
            print(f"  [WARN] 文件名映射冲突: {key} → 旧值={existing['filename_mapping'][key]}, 新值={value}，以新值为准")
        existing["filename_mapping"][key] = value

    # 追加运行记录
    existing["runs"].append({
        "timestamp": now_str,
        "files_processed": files_processed
    })

    return existing


def scan_documents(workspace):
    """扫描 raw 文件夹内的文档"""
    documents = []
    raw_dir = workspace / "raw"
    if not raw_dir.exists():
        print(f"[WARN] 未找到 raw 文件夹: {raw_dir}")
        return documents
    
    for root, dirs, files in os.walk(raw_dir):
        rel_root = Path(root).relative_to(raw_dir)
        skip = False
        for part in rel_root.parts:
            if part in SKIP_DIRS:
                skip = True
                break
        if skip:
            continue
        # 跳过 Word 临时文件
        for f in files:
            if f.startswith("~$"):
                continue
            ext = Path(f).suffix.lower()
            if ext in DOC_EXTENSIONS:
                documents.append(Path(root) / f)
    return documents


def scan_legacy_documents(workspace):
    """扫描 raw 文件夹中的 .doc/.xls 旧格式文件"""
    legacy = []
    raw_dir = workspace / "raw"
    if not raw_dir.exists():
        return legacy
    
    for root, dirs, files in os.walk(raw_dir):
        rel_root = Path(root).relative_to(raw_dir)
        skip = False
        for part in rel_root.parts:
            if part in SKIP_DIRS:
                skip = True
                break
        if skip:
            continue
        for f in files:
            if f.startswith("~$"):
                continue
            ext = Path(f).suffix.lower()
            if ext in LEGACY_EXTENSIONS:
                legacy.append(Path(root) / f)
    return legacy


def convert_legacy_files(legacy_files, workspace):
    """
    将 .doc/.xls 文件转换为 .docx/.xlsx 格式
    使用 doc_xls2docx_xlsx 技能脚本进行转换
    返回: (成功转换的文件列表, 失败的文件列表)
    """
    doc_files = [f for f in legacy_files if f.suffix.lower() == ".doc"]
    xls_files = [f for f in legacy_files if f.suffix.lower() == ".xls"]

    converted = []
    failed = []

    # 转换 .doc → .docx
    if doc_files:
        print(f"\n[CONVERT] 开始转换 {len(doc_files)} 个 .doc 文件...")
        doc_script = SKILL_SCRIPTS_DIR / "doc_to_docx_com.py"
        if not doc_script.exists():
            print(f"  [ERROR] 未找到转换脚本: {doc_script}")
            print(f"  请确保已安装 doc_xls2docx_xlsx 技能")
            failed.extend(doc_files)
        else:
            for doc_file in doc_files:
                # 输出到同目录，扩展名改为 .docx
                output_path = doc_file.with_suffix(".docx")
                try:
                    result = subprocess.run(
                        [sys.executable, str(doc_script), str(doc_file), str(output_path)],
                        capture_output=True, text=True, encoding="gbk", errors="replace", timeout=120
                    )
                    if result.returncode == 0 and output_path.exists():
                        converted.append(output_path)
                        print(f"  [OK] {doc_file.name} → {output_path.name}")
                    else:
                        error_msg = result.stderr.strip() or result.stdout.strip() or "未知错误"
                        print(f"  [FAIL] {doc_file.name}: {error_msg}")
                        failed.append(doc_file)
                except subprocess.TimeoutExpired:
                    print(f"  [FAIL] {doc_file.name}: 转换超时（120秒）")
                    failed.append(doc_file)
                except Exception as e:
                    print(f"  [FAIL] {doc_file.name}: {e}")
                    failed.append(doc_file)

    # 转换 .xls → .xlsx
    if xls_files:
        print(f"\n[CONVERT] 开始转换 {len(xls_files)} 个 .xls 文件...")
        xls_script = SKILL_SCRIPTS_DIR / "xls_to_xlsx.py"
        if not xls_script.exists():
            print(f"  [ERROR] 未找到转换脚本: {xls_script}")
            print(f"  请确保已安装 doc_xls2docx_xlsx 技能")
            failed.extend(xls_files)
        else:
            for xls_file in xls_files:
                output_path = xls_file.with_suffix(".xlsx")
                try:
                    result = subprocess.run(
                        [sys.executable, str(xls_script), str(xls_file), str(output_path)],
                        capture_output=True, text=True, encoding="gbk", errors="replace", timeout=120
                    )
                    if result.returncode == 0 and output_path.exists():
                        converted.append(output_path)
                        print(f"  [OK] {xls_file.name} → {output_path.name}")
                    else:
                        error_msg = result.stderr.strip() or result.stdout.strip() or "未知错误"
                        print(f"  [FAIL] {xls_file.name}: {error_msg}")
                        failed.append(xls_file)
                except subprocess.TimeoutExpired:
                    print(f"  [FAIL] {xls_file.name}: 转换超时（120秒）")
                    failed.append(xls_file)
                except Exception as e:
                    print(f"  [FAIL] {xls_file.name}: {e}")
                    failed.append(xls_file)

    return converted, failed


def apply_rules(text, exact_rules, regex_rules, mapping, counter, for_filename=False):
    """对文本应用所有替换规则
    
    Args:
        for_filename: 若为 True，关键字替换结果不加 [] 标记（用于文件名脱敏）；
                      若为 False（默认），关键字替换结果加 [] 标记（用于内容脱敏）
    """
    # 先正则替换
    for compiled, replacement, pattern_str, label in regex_rules:
        if replacement is not None:
            text = compiled.sub(replacement, text)
        else:
            def dynamic_replace(match, _label=label, _mapping=mapping, _counter=counter):
                matched_text = match.group(0)
                key = ("regex", matched_text)
                if key in _mapping:
                    return _mapping[key]
                if _label:
                    _counter[0] += 1
                    placeholder = f"[RED_{_label}_{_counter[0]}]"
                else:
                    _counter[0] += 1
                    placeholder = f"[RED_REGEX_{_counter[0]:04d}]"
                _mapping[key] = placeholder
                return placeholder
            text = compiled.sub(dynamic_replace, text)

    # 再精确替换
    for pattern, replacement in exact_rules:
        if for_filename:
            # 文件名脱敏：不加 [] 标记（避免文件名中特殊字符问题）
            text = text.replace(pattern, replacement)
        else:
            # 内容脱敏：加 [] 标记，便于识别脱敏位置
            bracketed = f"[{replacement}]"
            text = text.replace(pattern, bracketed)

    return text


# ============================================================
# docx 处理 - 直接操作 XML w:t 元素
# ============================================================

def replace_in_xml_element(element, replace_rules):
    """
    对 XML 元素（段落 w:p）中的文本执行替换
    核心改进：合并所有 w:t 文本 → 整体替换 → 写回
    replace_rules: [(old, new), ...] 替换对列表
    """
    from docx.oxml.ns import qn

    # 收集所有 w:t 元素及其文本
    t_elems = list(element.iter(qn('w:t')))
    if not t_elems:
        return False

    # 合并所有文本
    texts = [t.text or "" for t in t_elems]
    full_text = "".join(texts)
    if not full_text:
        return False

    # 应用替换规则
    new_text = full_text
    for old, new in replace_rules:
        new_text = new_text.replace(old, new)

    if new_text == full_text:
        return False  # 无变化

    # 将替换后的文本写入第一个 w:t，清空其余
    t_elems[0].text = new_text
    t_elems[0].set(qn('xml:space'), 'preserve')
    for t in t_elems[1:]:
        t.text = ""

    return True


def replace_docx_content(file_path, replace_rules):
    """对 docx 文件执行替换（脱敏或恢复均可用）"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(file_path))
    changed = False

    # 1. 替换段落
    for para in doc.paragraphs:
        if replace_in_xml_element(para._element, replace_rules):
            changed = True

    # 2. 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_xml_element(para._element, replace_rules):
                        changed = True

    # 3. 替换页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                for para in header.paragraphs:
                    if replace_in_xml_element(para._element, replace_rules):
                        changed = True
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    if replace_in_xml_element(para._element, replace_rules):
                        changed = True

    # 4. 替换文本框
    try:
        for textbox in doc.element.iter(qn('w:txbxContent')):
            for para_elem in textbox.iter(qn('w:p')):
                if replace_in_xml_element(para_elem, replace_rules):
                    changed = True
    except Exception:
        pass

    doc.save(str(file_path))
    return changed


def sanitize_docx(file_path, exact_rules, regex_rules, mapping, counter):
    """脱敏 docx 文件"""
    from docx.oxml.ns import qn

    # 构建精确+正则规则的替换对（供统一替换函数使用）
    # 这里需要用 apply_rules 来处理，保留原有的 mapping/counter 逻辑
    from docx import Document

    doc = Document(str(file_path))

    def process_element(element):
        t_elems = list(element.iter(qn('w:t')))
        if not t_elems:
            return
        texts = [t.text or "" for t in t_elems]
        full_text = "".join(texts)
        if not full_text:
            return
        new_text = apply_rules(full_text, exact_rules, regex_rules, mapping, counter)
        if new_text == full_text:
            return
        t_elems[0].text = new_text
        t_elems[0].set(qn('xml:space'), 'preserve')
        for t in t_elems[1:]:
            t.text = ""

    # 1. 替换段落
    for para in doc.paragraphs:
        process_element(para._element)

    # 2. 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_element(para._element)

    # 3. 替换页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                for para in header.paragraphs:
                    process_element(para._element)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    process_element(para._element)

    # 4. 替换文本框
    try:
        for textbox in doc.element.iter(qn('w:txbxContent')):
            for para_elem in textbox.iter(qn('w:p')):
                process_element(para_elem)
    except Exception:
        pass

    doc.save(str(file_path))
    return True


# ============================================================
# xlsx 处理
# ============================================================

def sanitize_xlsx(file_path, exact_rules, regex_rules, mapping, counter):
    """脱敏 xlsx 文件"""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_value = apply_rules(cell.value, exact_rules, regex_rules, mapping, counter)
                    if new_value != cell.value:
                        cell.value = new_value
    wb.save(str(file_path))
    return True


def restore_xlsx(file_path, replace_rules):
    """恢复 xlsx 文件"""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path))
    changed = False
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_value = cell.value
                    for old, new in replace_rules:
                        new_value = new_value.replace(old, new)
                    if new_value != cell.value:
                        cell.value = new_value
                        changed = True
    wb.save(str(file_path))
    return changed


# ============================================================
# 文件名脱敏/恢复
# ============================================================

def sanitize_filename(filename, exact_rules, regex_rules, mapping, counter):
    name_without_ext = Path(filename).stem
    ext = Path(filename).suffix
    # 文件名脱敏：使用 for_filename=True，避免添加 -- 标记
    new_name = apply_rules(name_without_ext, exact_rules, regex_rules, mapping, counter, for_filename=True)
    if new_name == name_without_ext:
        return filename, False
    return new_name + ext, True


def restore_filename(filename, filename_mapping):
    """根据 filename_mapping 还原文件名"""
    if filename in filename_mapping:
        return filename_mapping[filename], True
    return filename, False


# ============================================================
# 验证函数
# ============================================================

def verify_sanitization(orig_path, san_path, exact_rules, regex_rules):
    """对比原始文件和脱敏文件，验证替换效果"""
    from docx import Document
    from docx.oxml.ns import qn

    orig = Document(str(orig_path))
    san = Document(str(san_path))

    def get_all_text(doc):
        texts = []
        for para in doc.paragraphs:
            texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        try:
            for textbox in doc.element.iter(qn('w:txbxContent')):
                for para_elem in textbox.iter(qn('w:p')):
                    t_parts = []
                    for t_elem in para_elem.iter(qn('w:t')):
                        t_parts.append(t_elem.text or "")
                    texts.append("".join(t_parts))
        except:
            pass
        return "\n".join(texts)

    orig_text = get_all_text(orig)
    san_text = get_all_text(san)

    results = []
    for pattern, replacement in exact_rules:
        orig_count = orig_text.count(pattern)
        san_count = san_text.count(pattern)
        new_count = san_text.count(replacement)
        results.append({
            "rule": f"{pattern} → {replacement}",
            "orig_count": orig_count,
            "san_remaining": san_count,
            "replaced": orig_count - san_count,
            "new_count": new_count,
            "success": san_count == 0
        })

    for compiled, replacement, pattern_str, label in regex_rules:
        orig_matches = compiled.findall(orig_text)
        san_matches = compiled.findall(san_text)
        results.append({
            "rule": f"正则:{label or pattern_str}",
            "orig_count": len(orig_matches),
            "san_remaining": len(san_matches),
            "replaced": len(orig_matches) - len(san_matches),
            "success": len(san_matches) == 0
        })

    return results


# ============================================================
# 脱敏主流程
# ============================================================

def run_sanitize(workspace, rename_files=True, auto_convert=False):
    config = load_config(workspace)
    if not config:
        return

    print("\n[SANITIZE] 文档脱敏 v2（统一记录版）")
    print("=" * 60)
    print("[INFO] 仅支持 .docx 和 .xlsx 格式的文档脱敏")

    # 检测旧格式文件
    legacy_files = scan_legacy_documents(workspace)
    if legacy_files:
        print(f"\n[WARN] 发现 {len(legacy_files)} 个旧格式文件（仅 .docx/.xlsx 支持脱敏）:")
        for lf in legacy_files:
            print(f"  - {lf.name} ({lf.suffix.lower()})")

        if auto_convert:
            print("\n[AUTO-CONVERT] --auto-convert 已启用，自动转换旧格式文件...")
            converted, conv_failed = convert_legacy_files(legacy_files, workspace)
            if conv_failed:
                print(f"\n[WARN] {len(conv_failed)} 个文件转换失败，将跳过这些文件")
            if converted:
                print(f"[INFO] 成功转换 {len(converted)} 个文件，将继续脱敏流程")
        else:
            print("\n是否要将这些旧格式文件自动转换为 .docx/.xlsx 格式？")
            print("  转换后即可进行脱敏处理。原始文件不会被删除。")
            print("  提示: 使用 --auto-convert 参数可跳过确认自动转换")
            try:
                answer = input("\n请输入 Y 确认转换，N 跳过这些文件: ").strip().upper()
            except EOFError:
                answer = "N"

            if answer == "Y":
                converted, conv_failed = convert_legacy_files(legacy_files, workspace)
                if conv_failed:
                    print(f"\n[WARN] {len(conv_failed)} 个文件转换失败，将跳过这些文件")
                if converted:
                    print(f"[INFO] 成功转换 {len(converted)} 个文件，将继续脱敏流程")
            else:
                print("[INFO] 已跳过旧格式文件，仅处理 .docx/.xlsx 文件")

    # 加载已有记录
    existing_record = load_record(workspace)
    print(f"[RECORD] 已有记录: mapping {len(existing_record['mapping'])} 条, "
          f"filename_mapping {len(existing_record['filename_mapping'])} 条")

    # 构建规则
    exact_rules = []
    regex_rules = []
    mapping = {}
    counter = [1]

    # 从已有 mapping 初始化 counter（避免占位符编号冲突）
    for key in existing_record["mapping"]:
        m = re.match(r'\[RED_(?:\w+?)_(\d+)\]', key)
        if m:
            num = int(m.group(1))
            if num >= counter[0]:
                counter[0] = num + 1

    print("\n[CONFIG] 精确匹配规则:")
    for rule in config.get("exact_rules", []):
        pattern = rule.get("pattern", "").strip()
        if not pattern:
            continue
        replacement = rule.get("replacement", "").strip()
        if not replacement:
            replacement = f"[RED_{counter[0]:04d}]"
            counter[0] += 1
        # mapping 中存储带 [] 标记的值，用于恢复时反向替换
        bracketed_replacement = f"[{replacement}]"
        mapping[("exact", pattern)] = bracketed_replacement
        exact_rules.append((pattern, replacement))
        print(f"  {pattern} → {bracketed_replacement}")

    print("\n[CONFIG] 正则匹配规则:")
    for rule in config.get("regex_rules", []):
        pattern = rule.get("pattern", "").strip()
        if not pattern:
            continue
        label = rule.get("label", "").strip()
        replacement = rule.get("replacement", "").strip()
        try:
            compiled = re.compile(pattern)
        except re.error as e:
            print(f"  [WARN] 正则表达式无效: {pattern} ({e})，已跳过")
            continue
        if not replacement:
            replacement = None
        regex_rules.append((compiled, replacement, pattern, label))
        label_str = f" (标签: {label})" if label else ""
        print(f"  /{pattern}/{label_str} → {'[动态生成]' if replacement is None else replacement}")

    if not exact_rules and not regex_rules:
        print("[WARN] 未配置任何脱敏规则")
        return

    # 扫描文档
    print("\n[SCAN] 扫描文档...")
    documents = scan_documents(workspace)
    print(f"  发现 {len(documents)} 个文档")
    for doc in documents:
        print(f"  - {doc.name}")

    if not documents:
        print("[WARN] 未发现任何 .docx/.xlsx 文件")
        print("  提示: 仅支持 .docx/.xlsx 格式，.doc/.xls 文件需先转换")
        return

    # 创建输出目录
    output_dir = workspace / "_sanitized_output"
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(exist_ok=True)

    # 执行脱敏
    print("\n[PROCESS] 开始脱敏...")
    success = 0
    failed = []
    files_processed = []
    filename_mapping = {}

    for doc in documents:
        rel = doc.relative_to(workspace)
        original_rel = str(rel)

        # 文件名脱敏
        if rename_files:
            new_filename, changed = sanitize_filename(rel.name, exact_rules, regex_rules, mapping, counter)
            if changed:
                new_rel = rel.parent / new_filename
                print(f"  [RENAME] {rel.name} → {new_filename}")
            else:
                new_rel = rel
        else:
            new_rel = rel

        dest = output_dir / new_rel
        dest.parent.mkdir(parents=True, exist_ok=True)

        # 复制到输出目录
        shutil.copy2(str(doc), str(dest))

        try:
            ext = doc.suffix.lower()
            if ext == ".docx":
                sanitize_docx(dest, exact_rules, regex_rules, mapping, counter)
            elif ext == ".xlsx":
                sanitize_xlsx(dest, exact_rules, regex_rules, mapping, counter)

            success += 1
            files_processed.append(str(rel))

            if str(new_rel) != original_rel:
                filename_mapping[str(new_rel)] = original_rel

            print(f"  [{success}/{len(documents)}] {rel.name} [OK]")
        except Exception as e:
            failed.append((str(rel), str(e)))
            print(f"  [{success + len(failed)}/{len(documents)}] {rel.name} [FAIL: {e}]")

    # 构建本次 mapping（脱敏值→原始值）
    new_mapping = {v: k[1] for k, v in mapping.items()}

    # 合并到统一记录
    timestamp = datetime.now()
    merged_record = merge_record(existing_record, new_mapping, filename_mapping, files_processed, timestamp)
    save_record(workspace, merged_record)

    print(f"\n[RECORD] 统一记录已更新: {workspace / RECORD_FILE}")
    print(f"  累积 mapping: {len(merged_record['mapping'])} 条")
    print(f"  累积 filename_mapping: {len(merged_record['filename_mapping'])} 条")
    print(f"  运行历史: {len(merged_record['runs'])} 次")

    # 验证脱敏效果
    print("\n[VERIFY] 验证脱敏效果...")
    for doc in documents:
        rel = doc.relative_to(workspace)
        if rename_files:
            new_filename, _ = sanitize_filename(rel.name, exact_rules, regex_rules, mapping, counter)
            san_path = output_dir / rel.parent / new_filename
        else:
            san_path = output_dir / rel

        if san_path.exists() and doc.suffix.lower() == ".docx":
            results = verify_sanitization(doc, san_path, exact_rules, regex_rules)
            print(f"\n  文件: {rel.name}")
            all_ok = True
            for r in results:
                status = "[OK]" if r["success"] else "[FAIL]"
                if not r["success"]:
                    all_ok = False
                if r["orig_count"] > 0 or not r["success"]:
                    print(f"    {status} {r['rule']}: {r['orig_count']} -> {r['san_remaining']} (replaced {r['replaced']})")
            if all_ok:
                print(f"    [OK] All rules verified")

    # 汇报结果
    print("\n" + "=" * 60)
    print(f"[RESULT] 脱敏完成: {success}/{len(documents)} 成功")
    if failed:
        print(f"失败文件: {len(failed)}")
        for path, err in failed:
            print(f"  - {path}: {err}")
    print(f"\n脱敏输出目录: {output_dir}")
    print(f"统一脱敏记录: {workspace / RECORD_FILE}")
    print("=" * 60)

    if success > 0:
        print(f"\n恢复方法:")
        print(f"  python sanitize_v2.py restore <工作文件夹>")


# ============================================================
# 恢复主流程
# ============================================================

def run_restore(workspace):
    """从统一脱敏记录恢复文档"""
    record_path = workspace / RECORD_FILE
    if not record_path.exists():
        print(f"[ERROR] 未找到统一脱敏记录: {record_path}")
        print("  请先执行脱敏操作生成记录")
        return

    with open(record_path, "r", encoding="utf-8") as f:
        record = json.load(f)

    mapping = record.get("mapping", {})
    filename_mapping = record.get("filename_mapping", {})

    if not mapping:
        print("[ERROR] 脱敏记录中 mapping 为空，无法恢复")
        return

    print("\n[RESTORE] 文档恢复")
    print("=" * 60)
    print("[INFO] 仅支持 .docx 和 .xlsx 格式的文档恢复")
    print(f"[RECORD] 脱敏记录: mapping {len(mapping)} 条, filename_mapping {len(filename_mapping)} 条")
    print(f"  创建时间: {record.get('created_at', 'N/A')}")
    print(f"  最后更新: {record.get('last_updated', 'N/A')}")

    # 扫描脱敏输出目录
    input_dir = workspace / "_sanitized_output"
    if not input_dir.exists():
        print(f"[ERROR] 未找到脱敏输出目录: {input_dir}")
        return

    # 收集脱敏文件
    sanitized_files = []
    for root, dirs, files in os.walk(input_dir):
        for f in files:
            if f.startswith("~$"):
                continue
            ext = Path(f).suffix.lower()
            if ext in DOC_EXTENSIONS:
                sanitized_files.append(Path(root) / f)

    if not sanitized_files:
        print("[WARN] 脱敏输出目录中未发现任何 docx/xlsx 文件")
        return

    print(f"\n[SCAN] 发现 {len(sanitized_files)} 个脱敏文件")
    for f in sanitized_files:
        print(f"  - {f.name}")

    # 构建反向替换规则：脱敏值 → 原始值
    # 关键：长 key 优先替换，避免短 key 误匹配长 key 的子串
    replace_rules = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)

    print(f"\n[CONFIG] 反向替换规则（共 {len(replace_rules)} 条）:")
    for old, new in replace_rules:
        print(f"  {old} → {new}")

    # 创建恢复输出目录
    output_dir = workspace / "_restored_output"
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(exist_ok=True)

    # 执行恢复
    print("\n[PROCESS] 开始恢复...")
    success = 0
    failed = []
    restored_filenames = 0

    for sf in sanitized_files:
        rel = sf.relative_to(input_dir)
        dest = output_dir / rel
        dest.parent.mkdir(parents=True, exist_ok=True)

        # 复制到恢复目录
        shutil.copy2(str(sf), str(dest))

        try:
            ext = dest.suffix.lower()
            if ext == ".docx":
                replace_docx_content(dest, replace_rules)
            elif ext == ".xlsx":
                restore_xlsx(dest, replace_rules)

            # 文件名恢复 - 使用完整相对路径查找映射
            current_rel = str(rel)
            restored_rel, name_changed = restore_filename(current_rel, filename_mapping)
            if name_changed:
                # 计算恢复后的完整路径
                restored_path = output_dir / restored_rel
                restored_path.parent.mkdir(parents=True, exist_ok=True)
                dest.rename(restored_path)
                dest = restored_path
                restored_filenames += 1
                print(f"  [RENAME] {rel.name} → {Path(restored_rel).name}")

            success += 1
            print(f"  [{success}/{len(sanitized_files)}] {rel.name} [OK]")
        except Exception as e:
            failed.append((str(rel), str(e)))
            print(f"  [{success + len(failed)}/{len(sanitized_files)}] {rel.name} [FAIL: {e}]")

    # 校验残留占位符
    print("\n[VERIFY] 校验残留占位符...")
    residual_count = 0
    # 获取所有脱敏标记（mapping 的 keys），用于精确匹配残留
    sanitization_marks = set(mapping.keys())
    for sf_root, sf_dirs, sf_files in os.walk(output_dir):
        for f in sf_files:
            fpath = Path(sf_root) / f
            ext = fpath.suffix.lower()
            if ext == ".docx":
                from docx import Document
                from docx.oxml.ns import qn
                doc = Document(str(fpath))
                texts = []
                for para in doc.paragraphs:
                    texts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            texts.append(cell.text)
                full_text = "\n".join(texts)
                # 只检测 mapping 中定义的脱敏标记
                found_marks = [mark for mark in sanitization_marks if mark in full_text]
                if found_marks:
                    residual_count += len(found_marks)
                    print(f"  [WARN] {fpath.name} 残留脱敏标记: {found_marks}")
            elif ext == ".xlsx":
                from openpyxl import load_workbook
                wb = load_workbook(str(fpath))
                for ws in wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                # 只检测 mapping 中定义的脱敏标记
                                found_marks = [mark for mark in sanitization_marks if mark in cell.value]
                                if found_marks:
                                    residual_count += len(found_marks)
                                    print(f"  [WARN] {fpath.name} 残留脱敏标记: {found_marks}")

    if residual_count == 0:
        print("  [OK] 无残留占位符")
    else:
        print(f"  [WARN] 共发现 {residual_count} 个残留占位符，请检查脱敏记录是否完整")

    # 汇报结果
    print("\n" + "=" * 60)
    print(f"[RESULT] 恢复完成: {success}/{len(sanitized_files)} 成功")
    if restored_filenames:
        print(f"  文件名恢复: {restored_filenames} 个")
    if failed:
        print(f"失败文件: {len(failed)}")
        for path, err in failed:
            print(f"  - {path}: {err}")
    print(f"\n恢复输出目录: {output_dir}")
    print("=" * 60)


# ============================================================
# 命令行入口
# ============================================================

def print_usage():
    print("""
文档脱敏脚本 v2（统一记录版）

用法:
  python sanitize_v2.py sanitize <工作目录> [--no-rename] [--auto-convert]   脱敏文档
  python sanitize_v2.py restore <工作目录>                                  恢复文档

选项:
  --no-rename      不对文件名进行脱敏（默认会脱敏文件名）
  --auto-convert   检测到 .doc/.xls 时自动转换为 .docx/.xlsx，无需确认

说明:
  仅支持 .docx 和 .xlsx 格式的文档脱敏
  检测到 .doc/.xls 旧格式文件时，会提示是否自动转换
  脱敏记录统一保存为 _sanitize_record.json，每次运行累积合并映射
  恢复时读取该记录执行反向替换，无需指定具体记录文件
""")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print_usage()
        sys.exit(1)

    command = sys.argv[1].lower()
    workspace = Path(sys.argv[2])

    if not workspace.exists():
        print(f"[ERROR] 工作目录不存在: {workspace}")
        sys.exit(1)

    if command == "sanitize":
        rename = "--no-rename" not in sys.argv
        auto_convert = "--auto-convert" in sys.argv
        run_sanitize(workspace, rename_files=rename, auto_convert=auto_convert)
    elif command == "restore":
        run_restore(workspace)
    else:
        print(f"[ERROR] 未知命令: {command}")
        print_usage()
        sys.exit(1)

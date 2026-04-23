#!/usr/bin/env python3
"""
将抖音日报 MD 内容转换为 Word 文档
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def md_to_docx(md_path: Path, docx_path: Path = None):
    """将 MD 文件内容转换为 Word 文档"""
    content = md_path.read_text(encoding='utf-8')
    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')

    doc = Document()

    # 设置默认字体（支持中文）
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题处理
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:], level=3)

        # 表格处理
        elif line.startswith('|'):
            # 收集连续的多行表格
            table_lines = [line]
            while i + 1 < len(lines) and lines[i + 1].startswith('|'):
                i += 1
                table_lines.append(lines[i])

            # 解析表格（跳过分隔行）
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip('|').split('|')]
                if not any('---' in c for c in cells):  # 跳过分隔行
                    rows.append(cells)

            if len(rows) >= 2:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = cell_text
                        # 表头加粗
                        if ri == 0:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
                doc.add_paragraph()  # 空行

        # 引用块
        elif line.startswith('>'):
            p = doc.add_paragraph()
            p.style = 'Quote'
            p.text = line[1:].strip()

        # 列表
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.text = line[2:]

        # 标题后不带特殊标记的普通段落
        elif line.strip():
            doc.add_paragraph(line)

        # 空行
        else:
            doc.add_paragraph()

        i += 1

    doc.save(str(docx_path))
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: md_to_docx.py <md_file> [docx_output]")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else None
    out = md_to_docx(md_path, docx_path)
    print(f"Word 文档已生成: {out}")

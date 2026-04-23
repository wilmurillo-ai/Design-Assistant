#!/usr/bin/env python3
"""
md_to_docx.py — 将 Markdown 文件转换为 Word 文档（python-docx）
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_path: Path) -> Path:
    """将 Markdown 文件转换为 .docx，返回 docx 路径"""
    doc = Document()
    doc.core_properties.title = md_path.stem

    lines = md_path.read_text(encoding='utf-8').splitlines()
    in_table = False
    table = None
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            run = doc.add_paragraph()
            run.add_run(line)
            run.style = 'No Spacing'
            run.runs[0].font.name = 'Courier New'
            run.runs[0].font.size = Pt(9)
            continue

        # 表格
        if stripped.startswith('|'):
            if not in_table:
                # 解析表头
                headers = [h.strip() for h in stripped.split('|')[1:-1]]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for idx, h in enumerate(headers):
                    hdr_cells[idx].text = h
                    for run in hdr_cells[idx].paragraphs[0].runs:
                        run.bold = True
                in_table = True
            else:
                # 分隔线跳过
                if all(c.strip().startswith('-') for c in stripped.split('|')[1:-1]):
                    continue
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                row = table.add_row()
                for idx, cell_text in enumerate(cells):
                    row.cells[idx].text = cell_text
            continue
        else:
            if in_table and table:
                in_table = False
                table = None

        # 标题
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
        # 引用
        elif stripped.startswith('>'):
            p = doc.add_paragraph()
            run = p.add_run(stripped[1:].strip())
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.italic = True
        # 分隔线
        elif stripped.startswith('---') or stripped.startswith('***'):
            p = doc.add_paragraph()
            p.add_run('─' * 40)
        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # 处理 **bold** 和 `code`
            _add_formatted_text(p, stripped[2:])
        # 有序列表
        elif re.match(r'^\d+\. ', stripped):
            p = doc.add_paragraph(style='List Number')
            m = re.match(r'^(\d+)\. (.*)', stripped)
            if m:
                _add_formatted_text(p, m.group(2))
        # 空行
        elif not stripped:
            p = doc.add_paragraph()
        # 普通段落
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    out_path = md_path.with_suffix('.docx')
    doc.save(str(out_path))
    return out_path


def _add_formatted_text(para, text: str):
    """在段落中添加带格式的文本（支持 **bold** 和 `code`）"""
    parts = re.split(r'(`[^`]+`|\*\*[^*]+\*\*)', text)
    for part in parts:
        run = para.add_run(part)
        if part.startswith('**') and part.endswith('**'):
            run.bold = True
            run.text = part[2:-2]
        elif part.startswith('`') and part.endswith('`'):
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.text = part[1:-1]
        else:
            run.text = part


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("Usage: python3 md_to_docx.py <file.md>")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = md_to_docx(md_path)
    print(f"Word 文档已生成: {docx_path}")

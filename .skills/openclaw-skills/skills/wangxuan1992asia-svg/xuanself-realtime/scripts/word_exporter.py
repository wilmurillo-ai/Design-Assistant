# -*- coding: utf-8 -*-
"""
俄罗斯血糖检测设备实时市场调研报告生成器 — Word（DOCX）导出模块 v2
使用 python-docx 将 Markdown 报告转换为专业 Word 文档
支持中俄混排，无需字体Hack，Unicode原生支持

v2 改进：
- 修复封面键值对解析（**： vs **:）
- 保留段落内联格式（**bold** / *italic* / `code` / [链接](url)）
- 表格自动计算列宽（内容自适应）
- 章节间自动分页
- 添加页眉页脚
- 正文使用 1.15 倍行距
"""

import os
import re
import logging
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── 本地配置（无需 report-gama）──────────────────────────────────
SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR    = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR   = os.path.join(SKILL_DIR, "output")
DATA_DIR     = os.path.join(SKILL_DIR, "data")
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)
LOG_LEVEL    = "INFO"
LOG_FORMAT   = "%(asctime)s [%(levelname)s] %(name)s: %(message)s"

logging.basicConfig(level=LOG_LEVEL, format=LOG_FORMAT)
logger = logging.getLogger(__name__)


class WordExporter:
    """Markdown → Word（DOCX）导出器 v2"""

    # 颜色常量
    COLOR_TITLE   = RGBColor(0x1a, 0x3c, 0x6e)
    COLOR_H1      = RGBColor(0x1a, 0x3c, 0x6e)
    COLOR_H2      = RGBColor(0x1e, 0x5e, 0x9c)
    COLOR_H3      = RGBColor(0x29, 0x6f, 0xb5)
    COLOR_H4      = RGBColor(0x15, 0x5a, 0x8a)
    COLOR_BODY    = RGBColor(0x22, 0x22, 0x22)
    COLOR_META    = RGBColor(0x55, 0x55, 0x55)
    COLOR_ACCENT  = RGBColor(0xc0, 0x39, 0x2b)
    COLOR_LINK    = RGBColor(0x15, 0x5a, 0x8a)

    FONT_CN = "微软雅黑"
    FONT_EN = "Calibri"

    def __init__(self, output_dir=None):
        self.output_dir = output_dir or OUTPUT_DIR
        self.doc = None

    # ─────────────────────────────────────────────────────────
    # 字体 / Run 工具
    # ─────────────────────────────────────────────────────────

    def _run(self, paragraph, text, bold=False, italic=False,
             font_size=None, color=None, font_name=None,
             underline=False):
        """创建富文本片段（run）"""
        run = paragraph.add_run(text)
        run.font.name = font_name or self.FONT_EN
        # 中文字体
        if font_name == self.FONT_CN or self._has_cjk(text):
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_CN)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
        return run

    @staticmethod
    def _has_cjk(text):
        return bool(re.search(r'[\u4e00-\u9fff\u0400-\u04ff]', text))

    # ─────────────────────────────────────────────────────────
    # 分页与分隔
    # ─────────────────────────────────────────────────────────

    def _page_break(self):
        """插入分页符"""
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _thin_rule(self, color='cccccc', sz='4'):
        """细线分隔"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), sz)
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
        pBdr.append(bot); pPr.append(pBdr)

    def _accent_rule(self, color='1a3c6e', sz='8'):
        """标题下强调线"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), sz)
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
        pBdr.append(bot); pPr.append(pBdr)

    # ─────────────────────────────────────────────────────────
    # 页面设置（页边距 + 页眉页脚）
    # ─────────────────────────────────────────────────────────

    def _setup_page(self, title="俄罗斯医疗器械市场调研报告"):
        """设置页边距、页眉、页脚"""
        section = self.doc.sections[0]
        # 页边距：上下2.5cm，左右3cm
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)

        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        self._run(hp, title[:30] + "…" if len(title) > 30 else title,
                  font_size=8, color=self.COLOR_META, font_name=self.FONT_CN)
        self._run(hp, "  |  ", font_size=8, color=self.COLOR_META)
        self._run(hp, datetime.now().strftime('%Y-%m-%d'),
                  font_size=8, color=self.COLOR_META)
        # 页眉底线
        hpPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'cccccc')
        pBdr.append(bot); hpPr.append(pBdr)

        # 页脚
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        self._run(fp, "— ", font_size=8, color=self.COLOR_META)
        self._run(fp, "俄罗斯血糖检测设备实时市场调研报告  |  机密", font_size=8,
                  color=self.COLOR_META, font_name=self.FONT_CN)
        self._run(fp, " —", font_size=8, color=self.COLOR_META)

    # ─────────────────────────────────────────────────────────
    # 内联文本解析（保留 **bold** / *italic* / `code` / [链接](url)）
    # ─────────────────────────────────────────────────────────

    def _parse_inline(self, paragraph, text,
                      base_size=10, base_color=None, base_font=None):
        """
        将 text 解析为多个 run，保留内联格式。
        支持：**bold** *italic* `code` [链接文本](url)
        """
        base_color = base_color or self.COLOR_BODY
        base_font  = base_font  or self.FONT_CN

        # 按格式分段
        pattern = re.compile(
            r'(\*\*(.+?)\*\*)|'    # **bold**
            r'(\*(.+?)\*)|'        # *italic*
            r'(`(.+?)`)|'          # `code`
            r'(\[([^\]]+)\]\([^\)]+\))',  # [text](url)
            re.DOTALL
        )

        pos = 0
        for m in pattern.finditer(text):
            # 纯文本片段
            if m.start() > pos:
                chunk = text[pos:m.start()]
                self._run(paragraph, chunk, font_size=base_size,
                          color=base_color, font_name=base_font)

            if m.group(1):  # **bold**
                self._run(paragraph, m.group(2), bold=True, font_size=base_size,
                          color=base_color, font_name=base_font)
            elif m.group(3):  # *italic*
                self._run(paragraph, m.group(4), italic=True, font_size=base_size,
                          color=base_color, font_name=base_font)
            elif m.group(5):  # `code`
                self._run(paragraph, m.group(6), font_size=base_size,
                          color=RGBColor(0x8b, 0x00, 0x00),
                          font_name="Consolas")
            elif m.group(7):  # [text](url)
                label = m.group(8)
                self._run(paragraph, label, font_size=base_size,
                          color=self.COLOR_LINK, underline=True,
                          font_name=base_font)

            pos = m.end()

        if pos < len(text):
            self._run(paragraph, text[pos:], font_size=base_size,
                      color=base_color, font_name=base_font)

    # ─────────────────────────────────────────────────────────
    # 文档整体样式
    # ─────────────────────────────────────────────────────────

    def _set_doc_style(self):
        """设置文档默认样式"""
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_CN
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_CN)
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)

    # ─────────────────────────────────────────────────────────
    # 封面
    # ─────────────────────────────────────────────────────────

    def _add_cover(self, lines):
        """封面：元信息表格"""
        meta = {}
        for line in lines:
            raw = line.lstrip('> ')
            raw = re.sub(r'\*\*(.*?)\*\*', r'\1', raw)  # 先去掉bold
            # 解析 **key**：value  或 **key**: value
            m = re.match(r'\*\*(.+?)\*\*[：:]\s*(.+)', raw)
            if m:
                meta[m.group(1).strip()] = m.group(2).strip()

        # 标题占位
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(24)
        p_title.paragraph_format.space_after  = Pt(4)
        self._run(p_title, "血糖检测设备", bold=True, font_size=24,
                  color=self.COLOR_TITLE, font_name=self.FONT_CN)

        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after  = Pt(24)
        self._run(p_sub, "俄罗斯市场深度调研报告", bold=True, font_size=18,
                  color=self.COLOR_H2, font_name=self.FONT_CN)

        self._accent_rule(color='1a3c6e', sz='12')
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # 元信息表格
        tbl = self.doc.add_table(rows=max(1, len(meta)), cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 列宽
        for row_idx, (key, val) in enumerate(meta.items()):
            cells = tbl.rows[row_idx].cells
            # 左列：key
            kc = cells[0]
            kc.width = Cm(4.5)
            kc.text = ''
            kp = kc.paragraphs[0]
            kp.paragraph_format.space_before = Pt(3)
            kp.paragraph_format.space_after  = Pt(3)
            kp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rk = kp.add_run(key)
            rk.bold = True; rk.font.size = Pt(9.5)
            rk.font.color.rgb = self.COLOR_TITLE
            rk.font.name = self.FONT_CN
            rk._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_CN)
            tc = kc._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'edf2f9')
            tcPr.append(shd)

            # 右列：value
            vc = cells[1]
            vc.width = Cm(11)
            vc.text = ''
            vp = vc.paragraphs[0]
            vp.paragraph_format.space_before = Pt(3)
            vp.paragraph_format.space_after  = Pt(3)
            vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rv = vp.add_run(val)
            rv.font.size = Pt(9.5); rv.font.color.rgb = self.COLOR_BODY
            rv.font.name = self.FONT_CN
            rv._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_CN)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(18)
        self._thin_rule()

    # ─────────────────────────────────────────────────────────
    # 标题
    # ─────────────────────────────────────────────────────────

    def _add_heading(self, text, level=1):
        """添加标题，带分页（章节 H1/H2 之间）"""
        sizes   = {1: 16, 2: 13, 3: 11, 4: 10.5}
        colors  = {1: self.COLOR_H1, 2: self.COLOR_H2,
                   3: self.COLOR_H3, 4: self.COLOR_H4}
        bold_fl = {1: True, 2: True, 3: True, 4: True}
        space_b = {1: 20, 2: 14, 3: 10, 4: 8}

        if level == 1:
            self._page_break()

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_b.get(level, 10))
        p.paragraph_format.space_after  = Pt(3)

        self._parse_inline(p, text,
                           base_size=sizes.get(level, 10),
                           base_color=colors.get(level, self.COLOR_BODY))

        if level <= 2:
            self._accent_rule(
                color='1e5e9c' if level == 1 else '2980b9',
                sz='6' if level == 1 else '4'
            )

    # ─────────────────────────────────────────────────────────
    # 正文段落
    # ─────────────────────────────────────────────────────────

    def _add_paragraph(self, text, font_size=10, indent=False):
        """正文段落，两端对齐，行距1.15"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after   = Pt(5)
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15

        self._parse_inline(p, text,
                           base_size=font_size,
                           base_color=self.COLOR_BODY)
        return p

    # ─────────────────────────────────────────────────────────
    # 引用块
    # ─────────────────────────────────────────────────────────

    def _add_blockquote(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15

        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4');   left.set(qn('w:color'), '2980b9')
        pBdr.append(left); pPr.append(pBdr)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'eef4fb')
        pPr.append(shd)

        self._parse_inline(p, text,
                           base_size=9.5,
                           base_color=RGBColor(0x33, 0x44, 0x55))
        return p

    # ─────────────────────────────────────────────────────────
    # 列表项
    # ─────────────────────────────────────────────────────────

    def _add_list_item(self, text, numbered=False):
        p = self.doc.add_paragraph(style='List Bullet' if not numbered else 'List Number')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after   = Pt(2)
        p.paragraph_format.left_indent   = Cm(0.8)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15

        self._parse_inline(p, text,
                           base_size=10,
                           base_color=self.COLOR_BODY)

    # ─────────────────────────────────────────────────────────
    # 表格
    # ─────────────────────────────────────────────────────────

    def _add_table(self, header, rows):
        """表格：表头深色背景，数据行斑马纹，列宽自适应"""
        n_cols = len(header)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # ── 表头 ──
        hdr_cells = tbl.rows[0].cells
        for j, h_text in enumerate(header):
            c = hdr_cells[j]
            c.text = ''
            cp = c.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after  = Pt(4)
            rk = cp.add_run(h_text)
            rk.bold = True; rk.font.size = Pt(9)
            rk.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            rk.font.name = self.FONT_CN
            rk._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_CN)
            # 深蓝背景
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1a3c6e')
            tcPr.append(shd)
            # 垂直居中
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

        # ── 数据行 ──
        col_widths = self._calc_col_widths(header, rows)
        for i, row in enumerate(rows):
            cells = tbl.rows[i + 1].cells
            fill = 'f0f5fa' if i % 2 == 0 else 'ffffff'

            for j, cell_text in enumerate(row):
                c = cells[j]
                c.text = ''
                cp = c.paragraphs[0]
                cp.paragraph_format.space_before = Pt(3)
                cp.paragraph_format.space_after  = Pt(3)
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER

                self._parse_inline(cp, str(cell_text),
                                   base_size=8.5,
                                   base_color=self.COLOR_BODY)

                tc = c._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                tcPr.append(shd)
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # 列宽
                if j < len(col_widths):
                    tcPr2 = tc.get_or_add_tcPr()
                    tcWidth = OxmlElement('w:tcW')
                    tcWidth.set(qn('w:w'), str(int(col_widths[j] * 567)))
                    tcWidth.set(qn('w:type'), 'dxa')
                    tcPr2.append(tcWidth)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def _calc_col_widths(self, header, rows):
        """计算每列宽度（基于内容，最小2cm）"""
        n = len(header)
        all_text = [header] + rows
        # 每列最长字符串长度
        max_lens = [max(len(str(row[i])) if i < len(row) else 1
                        for row in all_text) for i in range(n)]
        total = sum(max_lens)
        if total == 0:
            return [12.0 / n] * n
        # 按比例分配，总宽度约15cm
        raw = [ml / total * 15.0 for ml in max_lens]
        # 保证最小2cm
        return [max(w, 2.0) for w in raw]

    # ─────────────────────────────────────────────────────────
    # Markdown 解析入口
    # ─────────────────────────────────────────────────────────

    def _parse_markdown(self, md_content):
        self.doc = Document()
        self._set_doc_style()
        self._setup_page()

        lines = md_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # ── 跳过 YAML front matter ──
            if i == 0 and line.strip() in ('---', '```'):
                fm_delim = line.strip()
                i += 1
                while i < len(lines) and lines[i].strip() != fm_delim:
                    i += 1
                if i < len(lines):
                    i += 1
                continue

            # ── 水平线 ──
            if re.match(r'^---+$', line.strip()):
                self._thin_rule()
                i += 1
                continue

            # ── 封面引用块 ──
            if line.startswith('>'):
                cover_lines = []
                while i < len(lines) and lines[i].startswith('>'):
                    cover_lines.append(lines[i])
                    i += 1
                self._add_cover(cover_lines)
                self.doc.add_paragraph()
                continue

            # ── H1 ──
            if line.startswith('## '):
                text = re.sub(r'^[\U0001F300-\U0001F9FF]\s*', '', line[3:].strip())
                self._add_heading(text, level=1)
                i += 1
                continue

            # ── H2 ──
            if line.startswith('### '):
                text = re.sub(r'^[\U0001F300-\U0001F9FF]\s*', '', line[4:].strip())
                self._add_heading(text, level=2)
                i += 1
                continue

            # ── H3 ──
            if line.startswith('#### '):
                text = re.sub(r'^[\U0001F300-\U0001F9FF]\s*', '', line[5:].strip())
                self._add_heading(text, level=3)
                i += 1
                continue

            # ── H4 ──
            if line.startswith('##### '):
                text = re.sub(r'^[\U0001F300-\U0001F9FF]\s*', '', line[6:].strip())
                self._add_heading(text, level=4)
                i += 1
                continue

            # ── 表格 ──
            if line.startswith('|') and '|' in line:
                header = []
                rows = []
                while i < len(lines) and lines[i].startswith('|'):
                    row_text = lines[i].strip()
                    if re.match(r'^\|[-:\s|]+\|$', row_text):
                        i += 1
                        continue
                    cells = [c.strip() for c in row_text.strip('|').split('|')]
                    if not header:
                        header = cells
                    else:
                        rows.append(cells)
                    i += 1
                if header and rows:
                    self._add_table(header, rows)
                continue

            # ── 引用块（> 非封面） ──
            if line.startswith('>'):
                block_lines = []
                while i < len(lines) and lines[i].startswith('>'):
                    raw = lines[i].lstrip('> ')
                    raw = re.sub(r'\*\*(.*?)\*\*', r'\1', raw)
                    block_lines.append(raw.strip())
                    i += 1
                self._add_blockquote('  '.join(block_lines))
                continue

            # ── 列表 ──
            if re.match(r'^[-*]\s', line) or re.match(r'^\d+\.\s', line):
                is_num = bool(re.match(r'^\d+\.\s', line))
                while i < len(lines) and (
                    re.match(r'^[-*]\s', lines[i]) or re.match(r'^\d+\.\s', lines[i])
                ):
                    item_text = lines[i]
                    item_text = re.sub(r'^[-*]\s+', '', item_text)
                    item_text = re.sub(r'^\d+\.\s+', '', item_text)
                    self._add_list_item(item_text, numbered=is_num)
                    i += 1
                self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
                continue

            # ── 普通段落 ──
            if line.strip():
                self._add_paragraph(line.strip())
                i += 1
                continue

            i += 1

    # ─────────────────────────────────────────────────────────
    # 公开 API
    # ─────────────────────────────────────────────────────────

    def markdown_to_docx(self, md_path, docx_path=None):
        if not HAS_DOCX:
            raise RuntimeError("请先安装 python-docx: pip install python-docx")
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"Markdown 文件不存在: {md_path}")

        with open(md_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        if docx_path is None:
            md_basename = os.path.splitext(os.path.basename(md_path))[0]
            docx_path = os.path.join(self.output_dir, f"{md_basename}.docx")

        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        self._parse_markdown(md_content)
        self.doc.save(docx_path)
        logger.info(f"Word 文档已生成: {docx_path}")
        return docx_path

    def markdown_str_to_docx(self, md_content, docx_path=None):
        if not HAS_DOCX:
            raise RuntimeError("请先安装 python-docx: pip install python-docx")
        if docx_path is None:
            docx_path = os.path.join(
                self.output_dir,
                f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        self._parse_markdown(md_content)
        self.doc.save(docx_path)
        logger.info(f"Word 文档已生成: {docx_path}")
        return docx_path


def quick_export(md_path, docx_path=None):
    exporter = WordExporter()
    return exporter.markdown_to_docx(md_path, docx_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python word_exporter.py <input.md> [output.docx]")
        sys.exit(1)
    out = sys.argv[2] if len(sys.argv) > 2 else None
    result = quick_export(sys.argv[1], out)
    print(f"OK: {result}")

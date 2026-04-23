# -*- coding: utf-8 -*-
"""
临床案例Word文档生成脚本

功能：
1. 从Markdown文件生成符合GB/T 7713.2-2022规范的Word文档
2. 自动处理参考文献上标格式（正文上标，列表不上标）
3. 正确处理一级标题和二级标题（避免重复编号）
4. 正确处理结论部分（没有二级标题的一级标题）

使用方法：
    python clinical_case_writer.py input.md output.docx

作者：OpenClaw
日期：2026-04-04
版本：v2.0（修复二级标题重复编号、结论缺失问题）
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, font_name='宋体', size=12, bold=False):
    """设置文字格式"""
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_para_with_refs(doc, text, font_name='宋体', size=12, indent=True):
    """添加段落并处理参考文献上标"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    
    # 匹配参考文献引用：<sup>[1]</sup> 或 [1]
    pattern = r'<sup>\[(\d+)\]</sup>|\[(\d+)\]'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # 添加参考文献前的普通文本
        if match.start() > last_end:
            normal_text = text[last_end:match.start()].replace('**', '')
            if normal_text:
                run = para.add_run(normal_text)
                set_run_font(run, font_name, size)
        
        # 添加参考文献（上标格式）
        ref_num = match.group(1) or match.group(2)
        if ref_num:
            sup_run = para.add_run(f'[{ref_num}]')
            sup_run.font.superscript = True
            set_run_font(sup_run, font_name, size)
        
        last_end = match.end()
    
    # 添加剩余的普通文本
    if last_end < len(text):
        remaining_text = text[last_end:].replace('**', '')
        if remaining_text:
            run = para.add_run(remaining_text)
            set_run_font(run, font_name, size)


def create_clinical_case_docx(md_file, output_file):
    """从Markdown生成Word文档"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面格式
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # 题名
    title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
    if title_match:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title_match.group(1))
        set_run_font(title_run, '黑体', 14, bold=True)
    
    # 摘要
    abstract_match = re.search(r'【摘要】(.+?)【关键词】', content, re.DOTALL)
    if abstract_match:
        abstract_para = doc.add_paragraph()
        abstract_para.paragraph_format.first_line_indent = Cm(0.74)
        abstract_para.paragraph_format.line_spacing = 1.5
        abstract_label = abstract_para.add_run('【摘要】')
        set_run_font(abstract_label, '黑体', 12, bold=True)
        abstract_text = abstract_match.group(1).strip()
        abstract_run = abstract_para.add_run(abstract_text)
        set_run_font(abstract_run, '宋体', 12)
    
    # 关键词
    keyword_match = re.search(r'【关键词】(.+?)$', content, re.MULTILINE)
    if keyword_match:
        keyword_para = doc.add_paragraph()
        keyword_para.paragraph_format.first_line_indent = Cm(0.74)
        keyword_label = keyword_para.add_run('【关键词】')
        set_run_font(keyword_label, '黑体', 12, bold=True)
        keyword_text = keyword_match.group(1).strip()
        keyword_run = keyword_para.add_run(keyword_text)
        set_run_font(keyword_run, '宋体', 12)
    
    # 按行处理内容（从引言开始）
    lines = content.split('\n')
    in_refs = False
    in_intro = False  # 标记是否在引言部分
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 跳过题名、摘要、关键词（已经处理过）
        if line.startswith('# ') or line.startswith('【摘要】') or line.startswith('【关键词】'):
            continue
        
        # 一级标题：## 1 病例资料
        if line.startswith('## ') and not line.startswith('### '):
            title = line[3:].strip()  # 去掉"## "
            h1_para = doc.add_paragraph()
            h1_run = h1_para.add_run(title)
            set_run_font(h1_run, '楷体', 12, bold=True)
            continue
        
        # 二级标题：### 1.1 一般资料
        if line.startswith('### '):
            title = line[4:].strip()  # 去掉"### "
            h2_para = doc.add_paragraph()
            h2_run = h2_para.add_run(title)
            set_run_font(h2_run, '宋体', 12, bold=True)
            continue
        
        # 参考文献
        if line.startswith('【参考文献】'):
            ref_title_para = doc.add_paragraph()
            ref_title_run = ref_title_para.add_run('【参考文献】')
            set_run_font(ref_title_run, '黑体', 12, bold=True)
            in_refs = True
            continue
        
        # 参考文献列表项
        if in_refs and line.startswith('['):
            ref_para = doc.add_paragraph()
            ref_para.paragraph_format.first_line_indent = Cm(0)
            ref_para.paragraph_format.line_spacing = 1.5
            ref_run = ref_para.add_run(line)
            set_run_font(ref_run, '宋体', 10.5)
            continue
        
        # 普通段落（非标题、非参考文献）
        if not line.startswith('#') and not line.startswith('【') and not in_refs:
            # 加粗小标题
            if line.startswith('**') and line.endswith('**'):
                para = doc.add_paragraph()
                clean_title = line.replace('**', '')
                run = para.add_run(clean_title)
                set_run_font(run, '宋体', 12, bold=True)
            else:
                # 普通段落
                add_para_with_refs(doc, line, '宋体', 12, True)
    
    # 保存文档
    doc.save(output_file)
    print(f'Word文档已生成：{output_file}')


def validate_docx(docx_file):
    """验证Word文档格式"""
    doc = Document(docx_file)
    
    # 统计上标格式
    superscript_count = 0
    in_refs_section = False
    
    for para in doc.paragraphs:
        if '【参考文献】' in para.text:
            in_refs_section = True
        if not in_refs_section:
            for run in para.runs:
                if run.font.superscript:
                    superscript_count += 1
    
    # 统计参考文献数量
    ref_count = 0
    in_refs = False
    
    for para in doc.paragraphs:
        if '【参考文献】' in para.text:
            in_refs = True
            continue
        if in_refs and para.text.strip().startswith('['):
            ref_count += 1
    
    # 检查一级标题
    headings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in ['1 病例资料', '2 讨论', '3 结论']:
            headings.append(text)
    
    # 检查结论部分
    conclusion_found = '3 结论' in headings
    
    print('='*70)
    print('临床案例Word文档验证报告')
    print('='*70)
    print()
    print('一、标题检查')
    for heading in headings:
        print(f'  ✓ {heading}')
    print()
    print('二、结论部分')
    if conclusion_found:
        print('  ✓ 结论部分存在')
    else:
        print('  ✗ 结论部分缺失')
    print()
    print('三、参考文献')
    print(f'  文献总数：{ref_count}篇')
    print(f'  上标格式：{superscript_count}个')
    print()
    print('='*70)
    
    if conclusion_found and ref_count >= 5 and ref_count <= 10:
        print('所有检查通过！文档符合规范要求。')
        return True
    else:
        print('文档存在问题，请检查。')
        return False


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('使用方法：python clinical_case_writer.py input.md output.docx')
        print('验证方法：python clinical_case_writer.py --validate output.docx')
        sys.exit(1)
    
    if sys.argv[1] == '--validate':
        validate_docx(sys.argv[2])
    else:
        md_file = sys.argv[1]
        output_file = sys.argv[2]
        create_clinical_case_docx(md_file, output_file)
        print()
        validate_docx(output_file)

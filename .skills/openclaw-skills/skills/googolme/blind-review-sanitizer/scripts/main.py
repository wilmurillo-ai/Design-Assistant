#!/usr/bin/env python3
"""
Blind Review Sanitizer (ID: 162)
One-click removal of author names, affiliations, acknowledgments, and excessive self-citations from manuscripts to meet double-blind review requirements.
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple


# Common institution keywords
INSTITUTION_KEYWORDS = [
    r'University', r'College', r'Institute', r'Academy', r'School',
    r'University', r'College', r'Institute', r'Research Institute', r'Laboratory', r'Lab\.?',
    r'Department', r'Dept\.?', r'Department', r'Center', r'Center', r'Centre'
]

# Acknowledgment-related titles (supports markdown headers like ## Acknowledgments)
ACKNOWLEDGMENT_TITLES = [
    r'(?i)^\s*#*\s*acknowledgments?\s*$',
    r'(?i)^\s*#*\s*acknowledgements?\s*$',
    r'(?i)^\s*#*\s*funding\s*$',
]

# Self-citation detection patterns
SELF_CITATION_PATTERNS = [
    r'\bour\s+(?:previous|prior|earlier)\s+(?:work|study|research|paper)s?\b',
    r'\bwe\s+(?:previously|earlier)\s+(?:showed|demonstrated|reported|found)\b',
    r'\bin\s+our\s+(?:previous|prior)\s+(?:work|study|research)\b',
    r'\b(?:my|our)\s+(?:own\s+)?(?:previous|prior|earlier)\b',
]

# Email pattern
EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

# Phone pattern
PHONE_PATTERN = r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b'


class BlindReviewSanitizer:
    """Double-blind review anonymization processor"""
    
    def __init__(
        self,
        authors: Optional[List[str]] = None,
        keep_acknowledgments: bool = False,
        highlight_self_cites: bool = False
    ):
        self.authors = [a.strip() for a in (authors or []) if a.strip()]
        self.keep_acknowledgments = keep_acknowledgments
        self.highlight_self_cites = highlight_self_cites
        self.removed_items = []
        self._in_acknowledgment = False
        
    def sanitize_text(self, text: str, *, in_acknowledgment: bool = False) -> str:
        """Anonymize plain text"""
        result = text
        self._in_acknowledgment = in_acknowledgment
        
        # 1. Remove author names
        result = self._remove_author_names(result)
        
        # 2. Remove institution information (skip if in acknowledgment and keep_acknowledgments is True)
        if not (self.keep_acknowledgments and self._in_acknowledgment):
            result = self._remove_institutions(result)
        
        # 3. Remove contact information
        result = self._remove_contact_info(result)
        
        # 4. Handle self-citations
        result = self._handle_self_citations(result)
        
        return result
    
    def _remove_author_names(self, text: str) -> str:
        """Remove author names"""
        if not self.authors:
            return text
            
        result = text
        for author in self.authors:
            # Match full names and surnames
            patterns = [
                re.escape(author),
            ]
            for pattern in patterns:
                result = re.sub(
                    pattern,
                    '[AUTHOR NAME]',
                    result,
                    flags=re.IGNORECASE
                )
        return result
    
    def _remove_institutions(self, text: str) -> str:
        """Remove institution/affiliation information"""
        result = text
        
        # Match common institution formats
        for keyword in INSTITUTION_KEYWORDS:
            # Match "XX University", "University of XX", "XX College", etc.]
            pattern = rf'\b[A-Z][A-Za-z\s]*{keyword}[A-Za-z\s]*\b|\b{keyword}[\u4e00-\u9fa5]+\b'
            matches = re.finditer(pattern, result, re.IGNORECASE)
            for match in list(matches)[::-1]:  # Reverse replacement to avoid position changes
                self.removed_items.append(f"Institution: {match.group()}")
                result = result[:match.start()] + '[INSTITUTION]' + result[match.end():]
        
        return result
    
    def _remove_contact_info(self, text: str) -> str:
        """Remove contact information (email, phone)"""
        result = text
        
        # Remove email
        matches = re.finditer(EMAIL_PATTERN, result, re.IGNORECASE)
        for match in list(matches)[::-1]:
            self.removed_items.append(f"Email: {match.group()}")
            result = result[:match.start()] + '[EMAIL]' + result[match.end():]
        
        # Remove phone
        matches = re.finditer(PHONE_PATTERN, result)
        for match in list(matches)[::-1]:
            self.removed_items.append(f"Phone: {match.group()}")
            result = result[:match.start()] + '[PHONE]' + result[match.end():]
        
        return result
    
    def _handle_self_citations(self, text: str) -> str:
        """Handle self-citations"""
        result = text
        
        # Track positions that have been modified to avoid nested replacements
        modified_ranges = []
        
        for pattern in SELF_CITATION_PATTERNS:
            if self.highlight_self_cites:
                # Find all matches first
                matches = list(re.finditer(pattern, result, flags=re.IGNORECASE))
                # Process from end to start to avoid position shifts
                for match in reversed(matches):
                    start, end = match.span()
                    # Check if this range overlaps with any already modified range
                    if not any(start < m_end and end > m_start for m_start, m_end in modified_ranges):
                        # Replace this match
                        replacement = f'[SELF-CITE: {match.group()}]'
                        result = result[:start] + replacement + result[end:]
                        # Mark this range as modified
                        modified_ranges.append((start, start + len(replacement)))
            else:
                def replace_func(match):
                    self.removed_items.append(f"Self-citation: {match.group()}")
                    return '[PREVIOUS WORK]'
                result = re.sub(pattern, replace_func, result, flags=re.IGNORECASE)
        
        return result
    
    def remove_acknowledgments(self, lines: List[str]) -> List[str]:
        """Remove acknowledgment sections"""
        if self.keep_acknowledgments:
            return lines
            
        result = []
        in_acknowledgment = False
        
        for line in lines:
            # Detect acknowledgment titles
            if any(re.match(pattern, line.strip()) for pattern in ACKNOWLEDGMENT_TITLES):
                in_acknowledgment = True
                self.removed_items.append(f"Acknowledgment section removed")
                result.append('[ACKNOWLEDGMENTS REMOVED]\n')
                continue
            
            # If in acknowledgment section, detect if ended (encountering next title)
            if in_acknowledgment:
                if re.match(r'^[#\d\s]', line.strip()) or line.strip().isupper():
                    in_acknowledgment = False
                else:
                    continue
            
            result.append(line)
        
        return result


class DocxProcessor:
    """DOCX document processor"""
    
    def __init__(self, sanitizer: BlindReviewSanitizer):
        self.sanitizer = sanitizer
    
    def process(self, input_path: Path, output_path: Path) -> None:
        """Process DOCX file"""
        try:
            from docx import Document
        except ImportError:
            print("Error: python-docx not installed. Run: pip install python-docx")
            sys.exit(1)
        
        doc = Document(input_path)
        
        # Process paragraphs
        for para in doc.paragraphs:
            # Detect and remove acknowledgments
            if not self.sanitizer.keep_acknowledgments:
                if any(re.match(pattern, para.text.strip()) for pattern in ACKNOWLEDGMENT_TITLES):
                    para.text = '[ACKNOWLEDGMENTS REMOVED]'
                    continue
            
            # Process text content
            if para.text.strip():
                para.text = self.sanitizer.sanitize_text(para.text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = self.sanitizer.sanitize_text(cell.text)
        
        # Save
        doc.save(output_path)


class TxtProcessor:
    """Plain text processor"""
    
    def __init__(self, sanitizer: BlindReviewSanitizer):
        self.sanitizer = sanitizer
    
    def process(self, input_path: Path, output_path: Path) -> None:
        """Process text file"""
        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Track acknowledgment section state
        in_acknowledgment = False
        result_lines = []
        
        for line in lines:
            # Check if this is an acknowledgment title
            if any(re.match(pattern, line.strip()) for pattern in ACKNOWLEDGMENT_TITLES):
                in_acknowledgment = True
                if not self.sanitizer.keep_acknowledgments:
                    self.sanitizer.removed_items.append(f"Acknowledgment section removed")
                    result_lines.append('[ACKNOWLEDGMENTS REMOVED]\n')
                    continue
                # If keeping acknowledgments, process the title line
                processed = self.sanitizer.sanitize_text(line, in_acknowledgment=in_acknowledgment)
                result_lines.append(processed)
                continue
            
            # Check if acknowledgment section ends
            if in_acknowledgment:
                if re.match(r'^[#\d\s]', line.strip()) or line.strip().isupper():
                    in_acknowledgment = False
                elif not self.sanitizer.keep_acknowledgments:
                    continue
            
            # Process line with acknowledgment context
            processed = self.sanitizer.sanitize_text(line, in_acknowledgment=in_acknowledgment)
            result_lines.append(processed)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(result_lines)


class MdProcessor(TxtProcessor):
    """Markdown processor (inherits TxtProcessor, retains extensibility for special handling)"""
    pass


def get_processor(file_path: Path, sanitizer: BlindReviewSanitizer):
    """Get appropriate processor based on file type"""
    suffix = file_path.suffix.lower()
    
    if suffix == '.docx':
        return DocxProcessor(sanitizer)
    elif suffix == '.md':
        return MdProcessor(sanitizer)
    elif suffix == '.txt':
        return TxtProcessor(sanitizer)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def main():
    parser = argparse.ArgumentParser(
        description='Blind Review Sanitizer - Double-blind review document anonymization tool',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --input paper.docx
  %(prog)s --input paper.md --authors "Zhang San,Li Si" --output blinded.md
  %(prog)s --input paper.txt --keep-acknowledgments --highlight-self-cites
        """
    )
    
    parser.add_argument('--input', '-i', required=True, help='Input file path (docx/txt/md)')
    parser.add_argument('--output', '-o', help='Output file path (default adds -blinded suffix)')
    parser.add_argument('--authors', help='Author names, comma-separated')
    parser.add_argument('--keep-acknowledgments', action='store_true', help='Keep acknowledgment sections')
    parser.add_argument('--highlight-self-cites', action='store_true', help='Only highlight self-citations without replacing')
    
    args = parser.parse_args()
    
    # Parse input path
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / f"{input_path.stem}-blinded{input_path.suffix}"
    
    # Parse author list
    authors = None
    if args.authors:
        authors = [a.strip() for a in args.authors.split(',') if a.strip()]
    
    # Create sanitizer
    sanitizer = BlindReviewSanitizer(
        authors=authors,
        keep_acknowledgments=args.keep_acknowledgments,
        highlight_self_cites=args.highlight_self_cites
    )
    
    # Get processor
    try:
        processor = get_processor(input_path, sanitizer)
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)
    
    # Process file
    print(f"Processing: {input_path}")
    processor.process(input_path, output_path)
    
    # Output statistics
    print(f"Output: {output_path}")
    print(f"Items processed: {len(sanitizer.removed_items)}")
    if sanitizer.removed_items:
        print("Summary:")
        for item in set(sanitizer.removed_items):
            count = sanitizer.removed_items.count(item)
            if count > 1:
                print(f"  - {item} (x{count})")
            else:
                print(f"  - {item}")
    
    print("Done!")


if __name__ == '__main__':
    main()

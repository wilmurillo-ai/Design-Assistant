from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from openpyxl import load_workbook

from time_window import in_time_window, parse_time_window


TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
SUBHEADING_FONT = "楷体_GB2312"
PAGE_FONT = "宋体"
LATIN_FONT = "Times New Roman"

TITLE_SIZE_PT = 22
BODY_SIZE_PT = 16
PAGE_SIZE_PT = 14
LINE_SPACING_PT = 28
FIRST_LINE_INDENT_PT = BODY_SIZE_PT * 2


def resolve_writable_report_path(preferred_path: Path) -> Path:
    try:
        preferred_path.parent.mkdir(parents=True, exist_ok=True)
        if not preferred_path.exists():
            return preferred_path
        with preferred_path.open("ab"):
            pass
        return preferred_path
    except PermissionError:
        timestamp = datetime.now().strftime("%H%M%S")
        return preferred_path.with_name(f"{preferred_path.stem}_{timestamp}{preferred_path.suffix}")


def build_time_window_text(paths: list[Path], time_window: str = "") -> str:
    if isinstance(time_window, str) and time_window.strip():
        return normalize_window_display_text(time_window)

    if not paths:
        return f"资讯时间窗口：{datetime.now().date().isoformat()}"

    dates = sorted(path.stem.replace("international_", "") for path in paths)
    if len(dates) == 1:
        return f"资讯时间窗口：{dates[0]}"
    return f"资讯时间窗口：{dates[0]}至{dates[-1]}"


def normalize_window_display_text(value: str) -> str:
    normalized = value.strip().replace(" to ", "至")
    if "至" not in normalized:
        return f"资讯时间窗口：{normalized}"

    start_text, end_text = [part.strip() for part in normalized.split("至", 1)]
    start_text = start_text.split(" ")[0]
    end_text = end_text.split(" ")[0]
    return f"资讯时间窗口：{start_text}至{end_text}"


def load_latest_domestic_rows(report_dir: Path, report_date: str, time_window: str = "") -> list[dict[str, str]]:
    candidates = sorted(
        list(report_dir.glob("company_mentions.xlsx")) + list(report_dir.glob(f"company_mentions_{report_date}*.xlsx")),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return filter_rows(load_rows_from_candidates(candidates), report_date, time_window)


def load_latest_international_rows(report_dir: Path, report_date: str, time_window: str = "") -> list[dict[str, str]]:
    candidates = sorted(
        list(report_dir.glob("international_company_mentions.xlsx"))
        + list(report_dir.glob(f"international_company_mentions_{report_date}*.xlsx")),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return filter_rows(load_rows_from_candidates(candidates), report_date, time_window)


def load_rows_from_candidates(candidates: list[Path]) -> list[dict[str, str]]:
    for candidate in candidates:
        try:
            workbook = load_workbook(candidate, read_only=True)
            worksheet = workbook.active
            headers = [cell.value for cell in next(worksheet.iter_rows(min_row=1, max_row=1))]
            rows: list[dict[str, str]] = []
            for values in worksheet.iter_rows(min_row=2, values_only=True):
                row = {str(headers[index]): str(value or "") for index, value in enumerate(values)}
                rows.append(row)
            return rows
        except (BadZipFile, OSError, StopIteration, KeyError, ValueError):
            continue
    return []


def filter_rows(rows: list[dict[str, str]], report_date: str, time_window: str = "") -> list[dict[str, str]]:
    if time_window:
        start, end = parse_time_window(time_window)
        return [
            row
            for row in rows
            if in_time_window(str(row.get("资讯时间", "") or ""), start, end)
        ]

    filtered: list[dict[str, str]] = []
    for row in rows:
        news_time = str(row.get("资讯时间", "") or "")
        if news_time.startswith(report_date):
            filtered.append(row)
    return filtered


def top_international_rows(rows: list[dict[str, str]], limit: int = 5) -> list[dict[str, str]]:
    valid_rows = [row for row in rows if str(row.get("AI摘要", "") or "").strip()]

    def sort_key(row: dict[str, str]) -> tuple[int, str]:
        try:
            score = int(str(row.get("事件影响力", "") or 0))
        except ValueError:
            score = 0
        return (-score, str(row.get("资讯时间", "") or ""))

    return sorted(valid_rows, key=sort_key)[:limit]


def set_run_font(run, east_asia_font: str, size_pt: float, *, latin_font: str = LATIN_FONT, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def style_paragraph(paragraph, *, alignment, first_line_indent_pt: float = FIRST_LINE_INDENT_PT) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)


def add_text_paragraph(
    document: Document,
    text: str,
    *,
    east_asia_font: str,
    size_pt: float,
    alignment,
    first_line_indent_pt: float = FIRST_LINE_INDENT_PT,
) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, alignment=alignment, first_line_indent_pt=first_line_indent_pt)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia_font, size_pt)


def add_blank_line(document: Document) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=0)
    run = paragraph.add_run("")
    set_run_font(run, BODY_FONT, BODY_SIZE_PT)


def chinese_section_number(index: int) -> str:
    numerals = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 1 <= index <= 10:
        return numerals[index - 1]
    if index < 20:
        return "十" + numerals[index - 11]
    tens, ones = divmod(index, 10)
    prefix = numerals[tens - 1] + "十"
    return prefix if ones == 0 else prefix + numerals[ones - 1]


def add_body_item(document: Document, index: int, title_text: str, body_text: str, source_name: str) -> None:
    title_paragraph = document.add_paragraph()
    style_paragraph(title_paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    title_run = title_paragraph.add_run(f"（{chinese_section_number(index)}）{title_text}")
    set_run_font(title_run, SUBHEADING_FONT, BODY_SIZE_PT)

    body_paragraph = document.add_paragraph()
    style_paragraph(body_paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    body_run = body_paragraph.add_run(body_text)
    set_run_font(body_run, BODY_FONT, BODY_SIZE_PT)
    if source_name:
        source_run = body_paragraph.add_run(f"（{source_name}）")
        set_run_font(source_run, SUBHEADING_FONT, BODY_SIZE_PT)


def add_empty_body(document: Document, text: str) -> None:
    add_text_paragraph(
        document,
        text,
        east_asia_font=BODY_FONT,
        size_pt=BODY_SIZE_PT,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def configure_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)


def build_page_number_field() -> list[OxmlElement]:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    return [begin, instr, end]


def configure_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_pt=0)

    left_run = paragraph.add_run("-")
    set_run_font(left_run, PAGE_FONT, PAGE_SIZE_PT)

    page_run = paragraph.add_run()
    set_run_font(page_run, PAGE_FONT, PAGE_SIZE_PT)
    for node in build_page_number_field():
        page_run._r.append(node)

    right_run = paragraph.add_run("-")
    set_run_font(right_run, PAGE_FONT, PAGE_SIZE_PT)


def write_daily_brief(
    report_dir: Path,
    paths: list[Path],
    domestic_rows: list[dict[str, str]],
    international_rows: list[dict[str, str]],
    time_window: str = "",
) -> Path:
    report_date = datetime.now().date().isoformat()
    output_path = resolve_writable_report_path(report_dir / f"AI资讯简报_{report_date}.docx")

    document = Document()
    configure_page_layout(document)
    configure_footer(document)

    add_text_paragraph(
        document,
        "每日AI产业资讯简报",
        east_asia_font=TITLE_FONT,
        size_pt=TITLE_SIZE_PT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
    )
    add_text_paragraph(
        document,
        build_time_window_text(paths, time_window),
        east_asia_font=BODY_FONT,
        size_pt=BODY_SIZE_PT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_pt=0,
    )
    add_blank_line(document)

    add_text_paragraph(
        document,
        "一、AI领军企业动态",
        east_asia_font=HEADING_FONT,
        size_pt=BODY_SIZE_PT,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    if not domestic_rows:
        add_empty_body(document, "今日无命中企业资讯。")
    else:
        for index, row in enumerate(domestic_rows, start=1):
            ai_title = row.get("AI标题") or row.get("资讯标题") or "未命名资讯"
            ai_summary = row.get("AI摘要") or row.get("资讯内容") or ""
            source_name = row.get("资讯来源", "")
            add_body_item(document, index, ai_title, ai_summary, source_name)

    add_text_paragraph(
        document,
        "二、全球AI产业动态",
        east_asia_font=HEADING_FONT,
        size_pt=BODY_SIZE_PT,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    selected_international_rows = top_international_rows(international_rows, limit=5)
    if not selected_international_rows:
        add_empty_body(document, "今日无国际AI资讯。")
    else:
        for index, row in enumerate(selected_international_rows, start=1):
            ai_title = row.get("AI标题") or row.get("资讯标题") or "未命名资讯"
            ai_summary = row.get("AI摘要") or ""
            source_name = row.get("资讯来源", "")
            add_body_item(document, index, ai_title, ai_summary, source_name)

    document.save(output_path)
    return output_path

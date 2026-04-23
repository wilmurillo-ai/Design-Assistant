from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


class WordExporter:
    """
    Word 导出器（按排版规范适配版）：
    只保留两大部分，且顺序固定：
    1. 名单企业动态
    2. 全球AI产业动态
    """

    def __init__(self, output_dir: str, filename_prefix: str = "AI资讯简报") -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.filename_prefix = filename_prefix

    def export(self, briefing_data: Dict, run_date: str, filename: Optional[str] = None) -> Path:
        if filename:
            file_path = self.output_dir / filename
        else:
            safe_date = run_date.replace("-", "")
            file_path = self.output_dir / f"{self.filename_prefix}_{safe_date}.docx"

        doc = Document()
        self._set_page_layout(doc)
        self._set_default_style(doc)
        self._add_page_number(doc)

        self._add_title(doc, briefing_data)
        self._add_section(doc, "一、“一带”AI领军企业动态", briefing_data.get("company_items", []))
        self._add_section(doc, "二、全球AI产业动态", briefing_data.get("global_items", []))

        doc.save(file_path)
        return file_path

    @staticmethod
    def _set_page_layout(doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)

    @staticmethod
    def _set_default_style(doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "仿宋_GB2312"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        style.font.size = Pt(16)

        p_format = style.paragraph_format
        p_format.first_line_indent = Pt(32)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_format.line_spacing = Pt(28)

    def _add_title(self, doc: Document, briefing_data: Dict) -> None:
        title = briefing_data.get("title", "每日AI产业资讯简报")
        run_date = briefing_data.get("run_date", "")
        time_window = briefing_data.get("time_window", "")

        #doc.add_paragraph("")
        #doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28)

        run = p.add_run(title)
        run.bold = False
        run.font.name = "方正小标宋简体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "方正小标宋简体")
        run.font.size = Pt(22)

        if run_date:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)

            r2 = p2.add_run(f"日期：{run_date}")
            r2.font.name = "仿宋_GB2312"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            r2.font.size = Pt(16)

        if time_window:
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)

            r3 = p3.add_run(f"统计口径：{time_window}")
            r3.font.name = "仿宋_GB2312"
            r3._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            r3.font.size = Pt(16)

    def _add_section(self, doc: Document, section_title: str, items: List[Dict]) -> None:
        p = doc.add_paragraph()
        r = p.add_run(section_title)
        r.bold = False
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(16)

        if not items:
            self._add_body_paragraph(doc, "暂无相关资讯。")
            return

        for idx, item in enumerate(items, start=1):
            title = item.get("title", "")
            summary = item.get("summary", "")
            source_name = item.get("source_name", "")
            source_url = item.get("source_url", "")

            p_title = doc.add_paragraph()
            r1 = p_title.add_run(f"{idx}. ")
            r1.bold = True
            r1.font.name = "仿宋_GB2312"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            r1.font.size = Pt(16)

            r2 = p_title.add_run(title)
            r2.bold = True
            r2.font.name = "仿宋_GB2312"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            r2.font.size = Pt(16)

            if summary:
                self._add_body_paragraph(doc, summary)

            if source_name:
                self._add_body_paragraph(doc, f"来源：{source_name}")

    @staticmethod
    def _add_body_paragraph(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(32)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(28)

        run = p.add_run(text)
        run.font.name = "仿宋_GB2312"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        run.font.size = Pt(16)

    def _add_page_number(self, doc: Document) -> None:
        """
        页脚居中页码，样式为：-1-
        """
        section = doc.sections[0]
        footer = section.footer

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        run_prefix = p.add_run("-")
        run_prefix.font.name = "Times New Roman"
        run_prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run_prefix.font.size = Pt(14)

        self._add_page_field(p)

        run_suffix = p.add_run("-")
        run_suffix.font.name = "Times New Roman"
        run_suffix._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run_suffix.font.size = Pt(14)

    @staticmethod
    def _add_page_field(paragraph) -> None:
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        text = OxmlElement("w:t")
        text.text = "1"

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(text)
        run._r.append(fld_char_end)
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 转换器
支持 Markdown 基本语法，字体统一使用仿宋
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import sys


class MD2WordConverter:
    def __init__(self):
        self.doc = Document()
        
    def convert(self, md_content, output_path):
        lines = md_content.split('\n')
        in_code_block = False
        code_content = []
        list_counter = {}  # 跟踪有序列表编号
        
        for line in lines:
            line = line.rstrip()
            
            # 代码块处理
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_content = []
                else:
                    self.add_code_block('\n'.join(code_content))
                    in_code_block = False
                continue
                
            if in_code_block:
                code_content.append(line)
                continue
            
            # 跳过空行
            if not line:
                self.doc.add_paragraph()
                list_counter = {}  # 重置列表编号
                continue
            
            # 标题处理
            if line.startswith('#'):
                hash_count = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if hash_count == 1:
                    self.add_heading(text, 1)
                elif hash_count == 2:
                    self.add_heading(text, 2)
                elif hash_count == 3:
                    self.add_heading(text, 3)
                elif hash_count >= 4:
                    self.add_heading(text, 4)
                list_counter = {}
                continue
            
            # 有序列表 - 检测缩进级别
            match = re.match(r'^(\d+)\.\s+(.*)', line)
            if match:
                num = int(match.group(1))
                text = match.group(2)
                # 根据缩进判断级别
                indent = len(line) - len(line.lstrip())
                level = indent // 2
                self.add_ordered_list_item(text, num, level)
                continue
            
            # 无序列表
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                self.add_unordered_list_item(text)
                continue
            
            # 水平分割线
            if line in ['---', '***', '___']:
                self.add_horizontal_line()
                continue
            
            # 引用
            if line.startswith('> '):
                text = line[2:]
                self.add_blockquote(text)
                continue
            
            # 普通段落
            self.add_paragraph(line)
            list_counter = {}
        
        self.doc.save(output_path)
        print(f'转换完成: {output_path}')
    
    def set_font(self, run):
        """设置仿宋字体"""
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(12)
    
    def add_heading(self, text, level):
        if level == 1:
            p = self.doc.add_paragraph(text)
            p.style = 'Title'
        elif level == 2:
            p = self.doc.add_paragraph(text)
            p.style = 'Heading 1'
        elif level == 3:
            p = self.doc.add_paragraph(text)
            p.style = 'Heading 2'
        else:
            p = self.doc.add_paragraph(text)
            p.style = 'Heading 3'
        
        for run in p.runs:
            self.set_font(run)
            if level == 1:
                run.font.size = Pt(22)
            elif level == 2:
                run.font.size = Pt(16)
            elif level == 3:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            run.bold = True
        
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_paragraph(self, text):
        p = self.doc.add_paragraph()
        self.parse_inline_formatting(p, text)
    
    def parse_inline_formatting(self, p, text):
        # 分割各种格式: **加粗** *斜体* ~~删除~~ `代码` [链接](url)
        pattern = r'(\*\*[^\*]+\*\*|\*[^\*]+\*|~~[^~]+~~|`[^`]+`|\[[^\]]+\]\([^)]+\))'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
            elif part.startswith('~~') and part.endswith('~~'):
                run = p.add_run(part[2:-2])
                run.font.strike = True
            elif part.startswith('[') and '](' in part:
                match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                if match:
                    run = p.add_run(match.group(1))
                    run.underline = True
            else:
                run = p.add_run(part)
            
            self.set_font(run)
    
    def add_ordered_list_item(self, text, num=1, level=0):
        """添加有序列表项，带手动编号"""
        p = self.doc.add_paragraph()
        # 添加编号
        run = p.add_run(f"{num}. ")
        self.set_font(run)
        run.bold = True
        # 添加内容
        self.parse_inline_formatting(p, text)
        # 设置缩进
        p.paragraph_format.left_indent = Pt(level * 24)
    
    def add_unordered_list_item(self, text):
        p = self.doc.add_paragraph()
        self.parse_inline_formatting(p, text)
        p.style = 'Normal'
    
    def add_code_block(self, code):
        p = self.doc.add_paragraph(code)
        p.style = 'Quote'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    
    def add_blockquote(self, text):
        p = self.doc.add_paragraph(text)
        p.style = 'Quote'
        for run in p.runs:
            self.set_font(run)
    
    def add_horizontal_line(self):
        p = self.doc.add_paragraph()
        p_fmt = p.paragraph_format
        p_fmt.space_before = Pt(6)
        p_fmt.space_after = Pt(6)


def main():
    if len(sys.argv) < 3:
        print('用法: python md2word.py <输入.md> <输出.docx>')
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # 读取 Markdown 文件
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换
    converter = MD2WordConverter()
    converter.convert(md_content, output_file)


if __name__ == '__main__':
    main()

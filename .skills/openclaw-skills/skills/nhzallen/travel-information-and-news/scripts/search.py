#!/usr/bin/env python3
"""
Travel Information and News - Main Search Script

Usage:
    python search.py --query "Tokyo travel March 2026"
    python search.py --query "京都賞櫻" --format pdf --output report.pdf
    python search.py --query "Bali hotels" --time_range pw --use_brave
"""

import argparse
import json
import os
import sys
from datetime import datetime

# Load .env if available
def load_env():
    env_paths = [
        os.path.join(os.path.dirname(__file__), "..", ".env"),
        ".env",
    ]
    for path in env_paths:
        if os.path.exists(path):
            with open(path) as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith("#") and "=" in line:
                        k, v = line.split("=", 1)
                        os.environ.setdefault(k.strip(), v.strip())

load_env()


def search_tavily(query, time_range="pm", max_results=10):
    """Search using Tavily API"""
    api_key = os.getenv("TAVILY_API_KEY")
    if not api_key:
        print("Error: TAVILY_API_KEY not set", file=sys.stderr)
        return []

    try:
        import requests
    except ImportError:
        print("Error: 'requests' package required. Install: pip install requests", file=sys.stderr)
        return []

    try:
        resp = requests.post(
            "https://api.tavily.com/search",
            headers={"Authorization": f"Bearer {api_key}"},
            json={
                "query": query,
                "max_results": max_results if max_results > 0 else 10,
                "topic": "general",
                "days": {"pd": 1, "pw": 7, "pm": 30, "py": 365}.get(time_range, 30),
                "include_answer": True,
            },
            timeout=30,
        )
        data = resp.json()
        results = []
        for item in data.get("results", []):
            results.append({
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "content": item.get("content", ""),
                "source": "tavily",
            })
        if data.get("answer"):
            results.insert(0, {
                "title": "AI Summary",
                "url": "",
                "content": data["answer"],
                "source": "tavily_summary",
            })
        return results
    except Exception as e:
        print(f"Tavily error: {e}", file=sys.stderr)
        return []


def search_brave(query, count=10):
    """Search using Brave Search API (optional)"""
    api_key = os.getenv("BRAVE_API_KEY")
    if not api_key:
        return []

    try:
        import requests
        resp = requests.get(
            "https://api.search.brave.com/res/v1/web/search",
            headers={"X-Subscription-Token": api_key, "Accept": "application/json"},
            params={"q": query, "count": count},
            timeout=30,
        )
        data = resp.json()
        results = []
        for item in data.get("web", {}).get("results", []):
            results.append({
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "content": item.get("description", ""),
                "source": "brave",
            })
        return results
    except Exception as e:
        print(f"Brave error: {e}", file=sys.stderr)
        return []


def search_browser(query, urls=None):
    """
    Scrape specific sites using Xvfb + Chromium + Puppeteer (三件套).
    Used for sites that Tavily/Brave can't reach (Xiaohongshu, X/Twitter, etc.)
    
    Requires:
      - Xvfb installed and running (or will attempt to start)
      - Chromium installed at /usr/bin/chromium
      - Puppeteer (Node.js) installed
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    browser_script = os.path.join(script_dir, "browser_search.js")

    if not os.path.exists("/usr/bin/chromium"):
        print("Browser search skipped: Chromium not found at /usr/bin/chromium", file=sys.stderr)
        print("Install: apt-get install -y xvfb chromium && npm install puppeteer", file=sys.stderr)
        return []

    # Ensure Xvfb is running
    import subprocess
    try:
        subprocess.run(["pgrep", "-f", "Xvfb :99"], capture_output=True, check=True)
    except subprocess.CalledProcessError:
        subprocess.Popen(
            ["Xvfb", ":99", "-screen", "0", "1200x720x24"],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        import time
        time.sleep(1)

    try:
        result = subprocess.run(
            ["node", browser_script, "--query", query],
            env={**os.environ, "DISPLAY": ":99"},
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            return json.loads(result.stdout)
        else:
            print(f"Browser search error: {result.stderr}", file=sys.stderr)
            return []
    except Exception as e:
        print(f"Browser search failed: {e}", file=sys.stderr)
        return []


def deduplicate(results):
    """Remove duplicate URLs within the same search run"""
    seen = set()
    unique = []
    for r in results:
        url = r.get("url", "")
        if url and url in seen:
            continue
        if url:
            seen.add(url)
        unique.append(r)
    return unique


def detect_language(text):
    """Simple language detection from query text"""
    for char in text:
        if "\u4e00" <= char <= "\u9fff":
            return "zh"
        if "\u3040" <= char <= "\u309f" or "\u30a0" <= char <= "\u30ff":
            return "ja"
        if "\uac00" <= char <= "\ud7af":
            return "ko"
    return "en"


def format_text(results, query):
    """Format results as plain text"""
    lines = []
    lines.append(f"=== Travel Search: {query} ===")
    lines.append(f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"Results: {len(results)}")
    lines.append("")

    for i, r in enumerate(results, 1):
        lines.append(f"[{i}] {r['title']}")
        if r.get("url"):
            lines.append(f"    URL: {r['url']}")
        lines.append(f"    Source: {r.get('source', 'unknown')}")
        content = r.get("content", "")
        if len(content) > 300:
            content = content[:300] + "..."
        lines.append(f"    {content}")
        lines.append("")

    return "\n".join(lines)


def format_docx(results, query, output_path):
    """Format results as Word document"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx required. Install: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    title = doc.add_heading("Travel Information Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Results: {len(results)}")
    doc.add_paragraph("")

    for i, r in enumerate(results, 1):
        heading = doc.add_heading(f"[{i}] {r['title']}", level=2)
        if r.get("url"):
            p = doc.add_paragraph()
            run = p.add_run(f"URL: {r['url']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
        p = doc.add_paragraph()
        run = p.add_run(f"Source: {r.get('source', 'unknown')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        content = r.get("content", "")
        doc.add_paragraph(content)
        doc.add_paragraph("")

    doc.save(output_path)
    print(f"Saved: {output_path}", file=sys.stderr)


def format_pdf(results, query, output_path):
    """Format results as PDF with CJK support"""
    try:
        from fpdf import FPDF
    except ImportError:
        print("Error: fpdf2 required. Install: pip install fpdf2", file=sys.stderr)
        sys.exit(1)

    # Find CJK font
    font_paths = [
        "/tmp/NotoSansSC-VF.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    font_path = None
    for fp in font_paths:
        if os.path.exists(fp):
            font_path = fp
            break

    if not font_path:
        print("Warning: No CJK font found, downloading Noto Sans SC...", file=sys.stderr)
        import urllib.request
        os.makedirs("/tmp", exist_ok=True)
        url = "https://github.com/notofonts/noto-cjk/raw/main/Sans/Variable/TTF/NotoSansCJKsc-VF.ttf"
        font_path = "/tmp/NotoSansSC-VF.ttf"
        urllib.request.urlretrieve(url, font_path)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    try:
        pdf.add_font("CJK", "", font_path)
    except Exception:
        pdf.add_font("CJK", "", font_path, uni=True)

    pdf.add_page()
    pdf.set_font("CJK", size=18)
    pdf.cell(0, 10, "Travel Information Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("CJK", size=10)
    pdf.cell(0, 6, f"Query: {query}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Results: {len(results)}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    pdf.set_font("CJK", size=11)
    for i, r in enumerate(results, 1):
        pdf.set_font("CJK", size=12)
        pdf.cell(0, 7, f"[{i}] {r['title']}", new_x="LMARGIN", new_y="NEXT")
        if r.get("url"):
            pdf.set_font("CJK", size=8)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, r["url"], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
        pdf.set_font("CJK", size=9)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 5, f"Source: {r.get('source', 'unknown')}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

        pdf.set_font("CJK", size=10)
        content = r.get("content", "")
        pdf.multi_cell(0, 6, content)
        pdf.ln(5)

    pdf.output(output_path)
    print(f"Saved: {output_path}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="Travel Information and News Search")
    parser.add_argument("--query", required=True, help="Search query")
    parser.add_argument("--time_range", default="pm", choices=["pd", "pw", "pm", "py"],
                        help="Time range: pd=24h, pw=week, pm=month, py=year (default: pm)")
    parser.add_argument("--region", default="ALL", help="Region code (default: ALL)")
    parser.add_argument("--max_results", type=int, default=0, help="Max results (0=unlimited)")
    parser.add_argument("--format", default="text", choices=["text", "docx", "pdf"],
                        help="Output format (default: text)")
    parser.add_argument("--output", help="Output file path (default: stdout)")
    parser.add_argument("--use_brave", action="store_true", help="Enable Brave Search fallback")
    parser.add_argument("--use_browser", action="store_true", help="Enable browser scraping")
    parser.add_argument("--no_reviews", action="store_true", help="Disable review aggregation (default: reviews enabled)")
    parser.add_argument("--language", help="Output language (default: auto-detect from query)")

    args = parser.parse_args()

    # Detect language
    lang = args.language or detect_language(args.query)

    # Step 1: Tavily search (required)
    print("Searching Tavily...", file=sys.stderr)
    results = search_tavily(args.query, args.time_range, args.max_results)

    # Step 2: Brave fallback (optional)
    if args.use_brave and len(results) < 3:
        print("Results insufficient, trying Brave Search...", file=sys.stderr)
        brave_results = search_brave(args.query)
        results.extend(brave_results)

    # Step 3: Browser scraping (optional, for blocked sites)
    if args.use_browser:
        print("Running browser search for blocked sites...", file=sys.stderr)
        browser_results = search_browser(args.query)
        results.extend(browser_results)

    # Deduplicate
    results = deduplicate(results)

    # Apply max_results limit
    if args.max_results > 0:
        results = results[:args.max_results]

    # Output with language metadata
    output_data = {
        "query": args.query,
        "detected_language": lang,
        "result_count": len(results),
        "results": results,
    }

    # Output
    if args.format == "text":
        output = format_text(results, args.query)
        # Add language reminder for agent
        output = f"[LANGUAGE: {lang}] — Agent must translate all results to this language\n\n" + output
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output)
            print(f"Saved: {args.output}", file=sys.stderr)
        else:
            print(output)
    elif args.format == "docx":
        output_path = args.output or "travel_report.docx"
        format_docx(results, args.query, output_path)
        print(f"[LANGUAGE: {lang}] — Agent must translate all results to this language", file=sys.stderr)
    elif args.format == "pdf":
        output_path = args.output or "travel_report.pdf"
        format_pdf(results, args.query, output_path)
        print(f"[LANGUAGE: {lang}] — Agent must translate all results to this language", file=sys.stderr)


if __name__ == "__main__":
    main()

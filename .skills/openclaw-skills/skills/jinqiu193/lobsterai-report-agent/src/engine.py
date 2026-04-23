"""
engine.py - 核心业务逻辑
===========================
不含路径配置（由 config.py 提供），不含 CLI 入口（由 cli.py 提供）。
"""

import os, re, hashlib, glob, json as json_module
from datetime import datetime
from concurrent.futures import ProcessPoolExecutor, as_completed
from typing import Dict, List, Tuple, Optional, Any

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import config


# ============ Hash（增量更新）============

def compute_content_hash(content: str) -> str:
    normalized = re.sub(r'\s+', '', content.strip())
    return hashlib.md5(normalized.encode('utf-8')).hexdigest()


def load_hashes() -> Dict[str, str]:
    try:
        with open(config._p('HASH_FILE'), 'r', encoding='utf-8') as f:
            return json_module.load(f)
    except Exception:
        return {}


def save_hashes(hashes: Dict[str, str]) -> None:
    with open(config._p('HASH_FILE'), 'w', encoding='utf-8') as f:
        json_module.dump(hashes, f, ensure_ascii=False, indent=2)


def get_changed_chapters(chapters_data: List[Tuple], hashes: Dict[str, str]) -> List[Tuple]:
    changed = []
    for item in chapters_data:
        seq, content = item[0], item[3]
        if hashes.get(seq) != compute_content_hash(content):
            changed.append(item)
    return changed


# ============ Mermaid 图表渲染（惰性）============

def render_mermaid_image(code: str, out_path: str, cli: str = None) -> bool:
    if cli is None:
        cli = config.get_mermaid_cli()
    if cli is None:
        return False
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    tmp_input = os.path.join(config.get_chapters_dir(), '_mermaid_tmp.mmd')
    try:
        with open(tmp_input, 'w', encoding='utf-8') as f:
            f.write(code)
        if cli.endswith('.js'):
            cmd = ['node', cli, '-i', tmp_input, '-o', out_path]
        else:
            cmd = cli.split() + ['-i', tmp_input, '-o', out_path]
        puppeteer_cfg = config._p('MERMAID_PUPPETEER_CONFIG')
        if os.path.exists(puppeteer_cfg):
            cmd += ['-p', puppeteer_cfg]
        import subprocess
        subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        return os.path.exists(out_path)
    except Exception:
        return False
    finally:
        if os.path.exists(tmp_input):
            os.remove(tmp_input)


def process_mermaid_blocks(content: str) -> Tuple[str, List[str]]:
    rendered_images = []
    mermaid_blocks = list(re.finditer(r'```mermaid\n(.*?)```', content, re.DOTALL))
    if not mermaid_blocks:
        return content, []

    processed = content
    for m in reversed(mermaid_blocks):
        code = m.group(1).strip()
        block_idx = len(mermaid_blocks) - 1 - mermaid_blocks[::-1].index(m)
        img_name = f'mermaid_{block_idx:03d}.png'
        img_path = os.path.join(config._p('MERMAID_TEMP'), img_name)

        cli = config.get_mermaid_cli()
        success = bool(cli and render_mermaid_image(code, img_path, cli))

        if success:
            rendered_images.append(img_path)
            replacement = f'\n[Mermaid图表已渲染，见附件: {img_name}]\n'
        else:
            replacement = (
                f'\n```mermaid\n{code}\n```\n\n'
                f'<!-- ⚠️ Mermaid图表（渲染工具mmdc未安装或渲染失败，'
                f'请在支持Mermaid的编辑器中查看） -->\n'
            )
        processed = processed[:m.start()] + replacement + processed[m.end():]
    return processed, rendered_images


# ============ Word TOC =============

NSMAP = (
    'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:w10="urn:schemas-microsoft-com:office:word" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
    'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
    'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
    'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
)


def _make_bookmark_start(bookmark_id: int, bookmark_name: str) -> OxmlElement:
    el = OxmlElement('w:bookmarkStart')
    el.set(qn('w:id'), str(bookmark_id))
    el.set(qn('w:name'), bookmark_name)
    return el


def _make_bookmark_end(bookmark_id: int) -> OxmlElement:
    el = OxmlElement('w:bookmarkEnd')
    el.set(qn('w:id'), str(bookmark_id))
    return el


def add_toc_entry(doc, seq: str, title: str, page_num: int, toc_type: str = 'chapter'):
    bm_id = 100 + hash(title) % 1000
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(4)

    if toc_type == 'summary':
        p.paragraph_format.first_line_indent = Cm(-0.74)
        r = p.add_run(seq + '　' + title)
        r.font.size = Pt(12)
        cjk(r, '宋体')
        return

    prefix = seq + '　'
    p.paragraph_format.first_line_indent = Cm(-0.74)

    r_prefix = p.add_run(prefix)
    r_prefix.font.size = Pt(12)
    cjk(r_prefix, '宋体')

    bookmark_name = f'_Toc_{bm_id}'
    run = p.add_run()
    run.font.size = Pt(12)
    cjk(run, '宋体')

    # FORMTEXT 字段
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.text = ' FORMTEXT '
    run._r.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

    # Tab + PAGEREF
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    p._p.append(tab)

    tab_char = OxmlElement('w:tabChar')
    tab_char.set(qn('w:val'), 'right')
    p._p.append(tab_char)

    run_page = p.add_run()
    run_page.font.size = Pt(12)
    cjk(run_page, '宋体')

    fld_begin2 = OxmlElement('w:fldChar')
    fld_begin2.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fld_begin2)
    instr2 = OxmlElement('w:instrText')
    instr2.text = f' PAGEREF {bookmark_name} \\h '
    run_page._r.append(instr2)
    fld_end2 = OxmlElement('w:fldChar')
    fld_end2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fld_end2)

    p._p.insert(0, _make_bookmark_start(bm_id, bookmark_name))
    p._p.append(_make_bookmark_end(bm_id))


# ============ 字体辅助 =============

def cjk(run, name: str) -> None:
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


# ============ Markdown 表格解析 =============

def _clean_inline(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def _is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith('|') and s.endswith('|')


def _is_separator_line(line: str) -> bool:
    stripped = line.strip().strip('|')
    return bool(re.match(r'^[\s\-:.|]+$', stripped))


def _parse_md_table(rows: List[str]) -> List[List[str]]:
    result = []
    for line in rows:
        stripped = line.strip().strip('|')
        cols = stripped.split('|')
        result.append([_clean_inline(c.strip()) for c in cols])
    return result


def _add_table_to_doc(doc, rows: List[List[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    col_count = max(col_count, 1)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        cells = tbl.rows[r_idx].cells
        for c_idx in range(len(cells)):
            text = row_data[c_idx] if c_idx < len(row_data) else ''
            cells[c_idx].text = text
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)


def _flush_table(doc, pending_table: List[str]) -> None:
    if pending_table:
        _add_table_to_doc(doc, _parse_md_table(pending_table))
        pending_table.clear()


# ============ Markdown → docx =============

def md_to_paragraphs(doc, content: str, add_page_break: bool = True) -> None:
    processed_content, rendered_images = process_mermaid_blocks(content)
    mermaid_img_iter = iter(rendered_images) if rendered_images else iter([])

    lines = processed_content.split('\n')
    pending_table: List[str] = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if '[Mermaid图表已渲染，见附件:' in line:
            img_path = next(mermaid_img_iter, None)
            if img_path and os.path.exists(img_path):
                _flush_table(doc, pending_table)
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                except Exception:
                    p = doc.add_paragraph()
                    r = p.add_run(line + ' [图片渲染失败]')
                    r.font.size = Pt(10); cjk(r, '宋体')
            continue

        if not line.strip():
            _flush_table(doc, pending_table)
            continue

        if line.startswith('# '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(_clean_inline(line[2:]))
            r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
            continue

        if line.startswith('## '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(_clean_inline(line[3:]))
            r.font.size = Pt(14); r.font.bold = True; cjk(r, '楷体')
            continue

        if line.startswith('### '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(_clean_inline(line[4:]))
            r.font.size = Pt(12); r.font.bold = True; cjk(r, '仿宋')
            continue

        if line.startswith('#### '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(_clean_inline(line[5:]))
            r.font.size = Pt(11); r.font.bold = True; cjk(r, '仿宋')
            continue

        if _is_table_line(line) and not _is_separator_line(line):
            pending_table.append(line)
            continue

        _flush_table(doc, pending_table)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(_clean_inline(line))
        r.font.size = Pt(12); cjk(r, '宋体')

    _flush_table(doc, pending_table)
    if add_page_break:
        doc.add_page_break()


# ============ 章节解析（错误隔离）============

def safe_parse_chapter(fpath: str) -> Optional[Tuple]:
    fname = os.path.basename(fpath).replace('.txt', '')
    seq = fname.split('-')[0]
    try:
        content = open(fpath, 'r', encoding='utf-8').read()
    except Exception as e:
        print(f"[ERROR] 读取失败 {fname}: {e}")
        return None
    h2_entries = [l[3:].strip() for l in content.split('\n') if l.strip().startswith('## ')]
    title = fname
    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('# '):
            title = line[2:].strip()
            break
    return (seq, fname, title, content, h2_entries)


def parse_chapters(txt_files: List[str]) -> List[Tuple]:
    seen_seq = set()
    chapters, errors = [], []
    for f in txt_files:
        seq = os.path.basename(f).replace('.txt', '').split('-')[0]
        if seq in seen_seq:
            continue
        result = safe_parse_chapter(f)
        if result is None:
            errors.append(os.path.basename(f)); continue
        seen_seq.add(seq); chapters.append(result)
    if errors:
        print(f"[WARN] 以下章节解析失败（已跳过）: {errors}")
    return chapters


def count_chars(text: str) -> int:
    return len([c for c in text if c.strip()])


# ============ 跨章一致性审查 =============

def extract_quantities(text: str) -> Dict[str, str]:
    qty = {}
    pattern = re.compile(r'(\d+(?:\.\d+)?)\s*(万元|万元/年|万元\/年|人|人\/日|台|套|个|次|年|月|天|%)')
    for m in pattern.finditer(text):
        key = f"{m.group(1)}{m.group(2)}"
        qty[key] = m.group(0)
    return qty


def check_cross_chapter_consistency(chapters_data: List[Tuple]) -> List[Dict]:
    issues = []
    all_qty = [(seq, extract_quantities(content)) for seq, _, _, content, _ in chapters_data]
    for i in range(len(all_qty) - 1):
        seq_a, qty_a = all_qty[i]
        seq_b, qty_b = all_qty[i + 1]
        shared = set(qty_a) & set(qty_b)
        for key in shared:
            ma = re.match(r'^(\d+(?:\.\d+)?)', key)
            mb = re.match(r'^(\d+(?:\.\d+)?)', key)
            if ma and mb:
                try:
                    if float(ma.group(1)) != float(mb.group(1)):
                        issues.append({
                            "seq_a": seq_a, "seq_b": seq_b,
                            "item": key, "value_a": qty_a[key], "value_b": qty_b[key]
                        })
                except ValueError:
                    continue
    return issues


# ============ 执行摘要 =============

def _build_summary(chapters_data, max_chars: int = 800) -> List[str]:
    lines, total = [], 0
    for _, _, _, content, _ in chapters_data:
        para_lines = []
        for line in content.split('\n'):
            line = line.strip()
            if not line or line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
                continue
            para_lines.append(line)
            if len(para_lines) >= 3:
                break
        if not para_lines:
            continue
        para_text = ''.join(para_lines[:2])
        if total + len(para_text) > max_chars:
            remaining = max_chars - total
            if remaining > 50:
                lines.append(para_text[:remaining] + '…')
            break
        lines.append(para_text)
        total += len(para_text)
    return lines or ['本报告对项目建设进行了全面可行性分析。']


# ============ 最终文档生成 =============

def generate_final_doc(chapters_data, page_estimates, output_path: str = None, incremental: bool = True):
    if output_path is None:
        output_path = config._p('FINAL_DOC')
    plan = config.load_plan()

    changed_chapters = chapters_data
    if incremental:
        hashes = load_hashes()
        changed_chapters = get_changed_chapters(chapters_data, hashes)
        unchanged = [item for item in chapters_data if item not in changed_chapters]
        if unchanged and not changed_chapters:
            print(f"[INCREMENTAL] 所有 {len(chapters_data)} 章内容未变化，跳过重写")
            return None
        elif unchanged:
            print(f"[INCREMENTAL] {len(unchanged)} 章未变化，{len(changed_chapters)} 章需重写")

    doc = Document()
    s = doc.sections[0]
    s.page_height = Inches(11.69); s.page_width = Inches(8.27)
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.18); s.right_margin = Inches(1.18)

    # 封面
    for _ in range(6): doc.add_paragraph()
    for txt, size, bold, font in [
        (plan.get('org_name', '编制单位'), Pt(26), True, '黑体'),
        (plan.get('project_name', '项目名称'), Pt(32), True, '黑体'),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = size; r.font.bold = bold; cjk(r, font)
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(plan.get('doc_type', '可行性研究报告'))
    r.font.size = Pt(22); cjk(r, '楷体')
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit = plan.get('编制单位', '编制单位')
    build_time = plan.get('编制时间', datetime.now().strftime('%Y年%m月'))
    r = p.add_run(f'编制单位：{unit}\n编制时间：{build_time}')
    r.font.size = Pt(14); cjk(r, '宋体')
    doc.add_page_break()

    # 执行摘要
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(10)
    r = p.add_run('执行摘要'); r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
    for pt in _build_summary(changed_chapters if changed_chapters else chapters_data):
        p2 = doc.add_paragraph()
        p2.paragraph_format.first_line_indent = Cm(0.74)
        p2.paragraph_format.line_spacing = Pt(22)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(pt); r2.font.size = Pt(12); cjk(r2, '宋体')
    doc.add_page_break()

    # 目录
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(10)
    r = p.add_run('目  录'); r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
    add_toc_entry(doc, '一', '执行摘要', 1, toc_type='summary')
    seen = set()
    for seq, _, title, content, _ in (changed_chapters if changed_chapters else chapters_data):
        if seq in seen or not seq.isdigit():
            continue
        seen.add(seq)
        start = page_estimates.get(seq, (1, 0, 1))[0]
        add_toc_entry(doc, f'第{int(seq)}章', title, start, toc_type='chapter')
    doc.add_page_break()

    # 各章节
    target = changed_chapters if changed_chapters else chapters_data
    for seq, _, _, content, _ in target:
        md_to_paragraphs(doc, content, add_page_break=True)

    # 更新 hash
    if incremental:
        new_hashes = {item[0]: compute_content_hash(item[3]) for item in target}
        old_hashes = load_hashes()
        old_hashes.update(new_hashes)
        save_hashes(old_hashes)

    doc.save(output_path)
    print(f"[DONE] 整合报告已保存: {output_path}")
    return output_path


# ============ 整合报告主流程 =============

def generate_with_accurate_toc(txt_dir: str = None, final_doc: str = None):
    if txt_dir is None:
        txt_dir = config._p('CHAPTERS_DIR')
    if final_doc is None:
        final_doc = config._p('FINAL_DOC')

    txt_files = sorted(glob.glob(os.path.join(txt_dir, '*.txt')))
    if not txt_files:
        print(f"[ERROR] 未找到章节文件: {txt_dir}/*.txt"); return None

    chapters_data = parse_chapters(txt_files)
    if not chapters_data:
        print("[ERROR] 所有章节解析均失败"); return None
    print(f"[PARSE] 解析 {len(chapters_data)} 个章节")

    ref_text = config.load_reference()
    config.generate_glossary(txt_files, ref_text=ref_text)

    issues = check_cross_chapter_consistency(chapters_data)
    if issues:
        print(f"[CONSISTENCY] 发现 {len(issues)} 个潜在不一致:")
        for iss in issues:
            print(f"  - {iss}")
    else:
        print("[CONSISTENCY] 跨章一致性检查通过 (OK)")

    pe = {}
    cur = 7
    for seq, _, _, content, _ in chapters_data:
        cc = count_chars(content)
        ep = max(1, (cc + config.CHARS_PER_PAGE - 1) // config.CHARS_PER_PAGE)
        pe[seq] = (cur, cc, ep); cur += ep

    print("[BUILD] 生成整合报告...")
    generate_final_doc(chapters_data, pe, output_path=final_doc)

    md_path = final_doc.replace('.docx', '-纯文本.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n\n---\n\n'.join(c for _, _, _, c, _ in chapters_data))
    print(f"[MD] 纯文本版已保存: {md_path}")
    return final_doc


# ============ 单章 docx 转换 =============

def convert_single_chapter_inline(txt_path: str, docx_path: str):
    doc = Document()
    s = doc.sections[0]
    s.page_height = Inches(11.69); s.page_width = Inches(8.27)
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.18); s.right_margin = Inches(1.18)

    lines = open(txt_path, 'r', encoding='utf-8').readlines()
    pending_table: List[str] = []

    for line in lines:
        line = line.rstrip()
        if not line.strip():
            _flush_table(doc, pending_table); continue

        if line.startswith('# '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph(); p.alignment = 1
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(10)
            r = p.add_run(_clean_inline(line[2:])); r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体'); continue

        if line.startswith('## '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(_clean_inline(line[3:])); r.font.size = Pt(14); r.font.bold = True; cjk(r, '楷体'); continue

        if line.startswith('### '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(_clean_inline(line[4:])); r.font.size = Pt(12); r.font.bold = True; cjk(r, '仿宋'); continue

        if line.startswith('#### '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(_clean_inline(line[5:])); r.font.size = Pt(11); r.font.bold = True; cjk(r, '仿宋'); continue

        if _is_table_line(line) and not _is_separator_line(line):
            pending_table.append(line); continue

        _flush_table(doc, pending_table)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(_clean_inline(line)); r.font.size = Pt(12); cjk(r, '宋体')

    _flush_table(doc, pending_table)
    doc.save(docx_path)
    return docx_path


def _convert_worker(args) -> Tuple[str, bool, str]:
    txt_path, docx_path = args
    try:
        convert_single_chapter_inline(txt_path, docx_path)
        return (docx_path, True, '')
    except Exception as e:
        return (txt_path, False, str(e))


# ============ 批量并行转换 =============

def batch_convert_txt_to_docx(txt_dir: str = None, max_concurrent: int = 8, incremental: bool = True):
    if txt_dir is None:
        txt_dir = config._p('CHAPTERS_DIR')

    txt_files = sorted(glob.glob(os.path.join(txt_dir, '*.txt')))
    if not txt_files:
        print("[ERROR] 未找到 .txt 文件"); return []

    hashes = load_hashes() if incremental else {}
    jobs = []
    for tf in txt_files:
        docx_path = tf.replace('.txt', '.docx')
        content_hash = compute_content_hash(open(tf, 'r', encoding='utf-8').read())
        if incremental and os.path.exists(docx_path) and hashes.get(os.path.basename(tf)) == content_hash:
            print(f"  [SKIP] {os.path.basename(tf)} 内容未变化，跳过")
            continue
        jobs.append((tf, docx_path))

    if not jobs:
        print("[INFO] 所有章节已是最新（无变化），跳过转换"); return []

    print(f"[BATCH] 待转换 {len(jobs)} 个章节，并发上限 {max_concurrent}")
    completed, failed = [], []
    with ProcessPoolExecutor(max_workers=max_concurrent) as executor:
        futures = {executor.submit(_convert_worker, job): job for job in jobs}
        for future in as_completed(futures):
            docx_path, ok, err = future.result()
            if ok:
                txt_path = docx_path.replace('.docx', '.txt')
                if os.path.exists(txt_path):
                    hashes[os.path.basename(txt_path)] = compute_content_hash(
                        open(txt_path, 'r', encoding='utf-8').read())
                completed.append(docx_path); print(f"  [OK] {os.path.basename(docx_path)}")
            else:
                failed.append((docx_path, err)); print(f"  [FAIL] {os.path.basename(docx_path)}: {err}")

    if incremental and completed:
        save_hashes(hashes)

    print(f"\n[BATCH] {len(completed)}/{len(jobs)} 成功，{len(failed)} 失败")
    return completed

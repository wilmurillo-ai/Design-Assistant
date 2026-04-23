#!/usr/bin/env python3
"""
docx_data_audit.py — cross-check every hard number inside a `.docx`
deliverable against the companion `data-provenance.md` table.

Parallel to `slide_data_audit.py` (which scans `.pptx`). This protects the
Word output path of a banker deliverable: if the memo cites a number that
isn't in the provenance table, the deck-to-memo flow drifted and the
memo must not ship.

Text extraction uses `python-docx` if present; falls back to `uvx --with
python-docx` for hosts without pip-installed python-docx.

Usage:
    python3 docx_data_audit.py <memo.docx> <data-provenance.md>
    # exit 0 → every hard number in the docx is covered by the provenance
    # exit 1 → one or more numbers are missing (or text extraction failed)

This script deliberately mirrors `slide_data_audit.py` to keep the matching
semantics consistent across PPT and Word gates.
"""
from __future__ import annotations
import pathlib
import subprocess
import sys

# Reuse the canonical matching helpers from provenance_verify so PPT/DOCX
# cross-checks stay identical to analysis.md checks.
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
from provenance_verify import (  # noqa: E402
    HARD_NUMBER,
    extract_provenance_corpus,
    normalize_variants,
)


EXTRACT_SCRIPT = r"""
import sys
try:
    from docx import Document
except ImportError:
    print('ERR: python-docx missing — install via pip or uvx', file=sys.stderr)
    sys.exit(2)
d = Document(sys.argv[1])
for p in d.paragraphs:
    t = p.text
    if t.strip():
        print(t)
for tbl in d.tables:
    for row in tbl.rows:
        for cell in row.cells:
            t = cell.text
            if t.strip():
                print(t)
"""


def extract_docx(path: pathlib.Path) -> str | None:
    for cmd in (
        ["python3", "-c", EXTRACT_SCRIPT, str(path)],
        ["uvx", "--with", "python-docx", "python3", "-c", EXTRACT_SCRIPT, str(path)],
    ):
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            print(f"ERR: extraction timeout on {path.name}", file=sys.stderr)
            return None
        if r.returncode == 0:
            return r.stdout
    print(
        f"ERR: python-docx unavailable and uvx fallback failed for {path.name}",
        file=sys.stderr,
    )
    return None


def audit(docx_text: str, prov_text: str) -> tuple[int, list[str]]:
    corpus = extract_provenance_corpus(prov_text)
    hits = 0
    missing: list[str] = []
    for lineno, raw_line in enumerate(docx_text.splitlines(), 1):
        for m in HARD_NUMBER.finditer(raw_line):
            hits += 1
            num, unit = m.group(1), m.group(2)
            unit_variants, bare_variants = normalize_variants(num, unit)
            if any(v in corpus for v in unit_variants):
                continue
            if unit in corpus and any(v in corpus for v in bare_variants):
                continue
            missing.append(
                f"L{lineno:>4}: hard number '{num}{unit}' not in provenance"
                f"\n       context: {raw_line.strip()[:120]!r}"
            )
    header = f"scan: {hits} hard numbers in docx; {hits - len(missing)} covered, {len(missing)} missing."
    return len(missing), [header] + missing


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("usage: docx_data_audit.py <memo.docx> <data-provenance.md>",
              file=sys.stderr)
        return 2
    docx_path = pathlib.Path(argv[1])
    prov_path = pathlib.Path(argv[2])
    if not docx_path.exists():
        print(f"file not found: {docx_path}", file=sys.stderr)
        return 2
    if not prov_path.exists():
        print(f"file not found: {prov_path}", file=sys.stderr)
        return 2

    docx_text = extract_docx(docx_path)
    if docx_text is None:
        return 2

    prov_text = prov_path.read_text(encoding="utf-8", errors="replace")
    missing, messages = audit(docx_text, prov_text)
    header = messages[0]
    if missing:
        print(f"FAIL: {header}", file=sys.stderr)
        for m in messages[1:]:
            print(m, file=sys.stderr)
        return 1
    print(f"OK: {header}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))

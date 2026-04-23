import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

class OfficialProFormatter:
    def __init__(self):
        # 字体常量定义
        self.FONT_RED_HEADER = '方正大标宋简体'  # 红头专用
        self.FONT_DOC_TITLE = '方正小标宋简体'   # 文件标题专用
        self.FONT_MAIN = '仿宋_GB2312'
        self.FONT_LEVEL_1 = '黑体'
        self.FONT_LEVEL_2 = '楷体_GB2312'
        self.FONT_WESTERN = 'Times New Roman'
        self.TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')

    def _set_mixed_font(self, run, chinese_font, size, bold=False, color=None):
        """实现中西文分设：中文用指定字体，西文用 Times New Roman"""
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        
        # 强制设置西文字体名
        run.font.name = self.FONT_WESTERN
        
        # 深入 XML 设置中文字体
        r = run._element.get_or_add_rPr()
        rFonts = r.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), self.FONT_WESTERN)
        rFonts.set(qn('w:hAnsi'), self.FONT_WESTERN)
        rFonts.set(qn('w:eastAsia'), chinese_font)

    def _format_para(self, para, text, is_title=False):
        """统一段落格式控制"""
        para.paragraph_format.line_spacing = Pt(28)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        if is_title:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
        else:
            para.paragraph_format.first_line_indent = Pt(32) # 三号字2字符缩进

    def execute(self, input_path, output_path, template_name=None, file_code=None, org_name=None):
        # 1. 加载模板
        template_file = os.path.join(self.TEMPLATE_DIR, f"{template_name}.docx" if template_name else "default.docx")
        doc = Document(template_file) if os.path.exists(template_file) else Document()
        
        # 2. 处理红头 (如果指定了 org_name 且模板中没有预设)
        if org_name:
            header_p = doc.paragraphs[0]
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_p.add_run(f"{org_name}文件")
            self._set_mixed_font(run, self.FONT_RED_HEADER, 72, color=RGBColor(255, 0, 0))

        # 3. 处理发文字号 (file_code)
        if file_code:
            for p in doc.paragraphs:
                if "〔" in p.text: # 识别发文字号占位符
                    p.text = file_code
                    for run in p.runs:
                        self._set_mixed_font(run, self.FONT_MAIN, 16)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4. 迁移并排版正文内容
        source_doc = Document(input_path)
        first_line = True
        
        for s_para in source_doc.paragraphs:
            raw_text = s_para.text.strip()
            if not raw_text: continue
            
            new_p = doc.add_paragraph()
            
            # A. 识别文件标题（正文第一行）
            if first_line:
                self._format_para(new_p, raw_text, is_title=True)
                run = new_p.add_run(raw_text)
                self._set_mixed_font(run, self.FONT_DOC_TITLE, 22)
                first_line = False
                continue

            # B. 标题层级识别
            target_font = self.FONT_MAIN
            is_bold = False
            if re.match(r'^[一二三四五六七八九十]+、', raw_text):
                target_font = self.FONT_LEVEL_1
            elif re.match(r'^（[一二三四五六七八九十]+）', raw_text):
                target_font = self.FONT_LEVEL_2
            elif re.match(r'^\d+\.', raw_text):
                is_bold = True
            
            self._format_para(new_p, raw_text)
            run = new_p.add_run(raw_text)
            self._set_mixed_font(run, target_font, 16, bold=is_bold)

        # 5. 保存
        doc.save(output_path)
        return {"status": "success", "file_path": output_path}

def handler(params):
    return OfficialProFormatter().execute(
        params['input_path'], 
        params['output_path'],
        params.get('template_name'),
        params.get('file_code'),
        params.get('org_name')
    )
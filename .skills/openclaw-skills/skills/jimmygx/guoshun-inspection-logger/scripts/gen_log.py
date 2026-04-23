#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
施工日志/巡检记录AI生成器
江苏国顺智能科技有限公司专用
生成江苏省建设工程标准化格式监理/施工日志
"""

import json
import os
import re
import sys
import argparse
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx library not found. Install with: pip install python-docx")
    sys.exit(1)


# ============================================================
# 模板定义
# ============================================================

TEMPLATES = {
    "FH-001": {
        "name": "监理日志",
        "title": "监  理  日  志",
        "fields": [
            "日期、天气、气温",
            "施工部位及内容",
            "投入人员（工种、人数）",
            "投入机械设备",
            "完成工程量",
            "质量情况及处理意见",
            "安全文明施工情况",
            "发现的问题及处理意见",
            "备注"
        ]
    },
    "SG-001": {
        "name": "施工日志",
        "title": "施  工  日  志",
        "fields": [
            "日期、天气、气温",
            "施工部位及内容",
            "投入人员（工种、人数）",
            "投入机械设备",
            "完成工程量",
            "质量情况及处理意见",
            "安全文明施工情况",
            "发现的问题及处理意见",
            "备注"
        ]
    },
    "AQ-001": {
        "name": "安全巡检记录",
        "title": "安  全  巡  检  记  录",
        "fields": [
            "巡检日期、时间",
            "巡检区域/部位",
            "巡检人员",
            "天气情况",
            "检查内容",
            "发现的安全隐患",
            "整改要求及期限",
            "复查情况",
            "备注"
        ]
    },
    "YJ-001": {
        "name": "隐蔽工程验收记录",
        "title": "隐 蔽 工 程 验 收 记 录",
        "fields": [
            "工程名称",
            "隐蔽项目",
            "隐蔽部位",
            "隐蔽日期",
            "施工单位",
            "施工内容",
            "验收情况",
            "质量问题",
            "整改要求",
            "验收结论"
        ]
    },
    "WD-001": {
        "name": "危大工程巡视记录",
        "title": "危 大 工 程 巡 视 记 录",
        "fields": [
            "巡视日期、时间",
            "工程名称",
            "危大工程类型",
            "巡视人员",
            "天气情况",
            "施工部位及内容",
            "施工状况",
            "安全防护措施检查",
            "存在问题",
            "处理意见",
            "备注"
        ]
    }
}


# ============================================================
# 项目信息管理
# ============================================================

PROJECT_INFO_PATH = Path.home() / ".openclaw/extensions/wecom-openclaw-plugin/skills/guoshun-inspection-logger/project_info.json"

def load_project_info():
    """加载保存的项目信息"""
    if PROJECT_INFO_PATH.exists():
        with open(PROJECT_INFO_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

def save_project_info(info):
    """保存项目信息"""
    PROJECT_INFO_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(PROJECT_INFO_PATH, 'w', encoding='utf-8') as f:
        json.dump(info, f, ensure_ascii=False, indent=2)


# ============================================================
# 文本解析 - 从用户描述中提取关键信息
# ============================================================

def parse_input_text(text):
    """
    解析用户输入的口头描述，提取关键信息
    返回结构化数据字典
    """
    result = {
        "日期": "",
        "天气": "",
        "气温": "",
        "施工部位": "",
        "施工内容": "",
        "人员": {},
        "机械设备": {},
        "完成工程量": {},
        "质量问题": [],
        "安全问题": [],
        "处理意见": "",
        "备注": ""
    }
    
    text = text.strip()
    
    # 提取天气
    weather_patterns = [
        r'(晴|阴|多云|雨|雪|雾|霾|大风)天',
        r'天气[是为]?(晴|阴|多云|雨|雪|雾|霾|大风)',
        r'(晴|阴|多云|雨|雪|雾|霾|大风)\s*[，,]'
    ]
    for pattern in weather_patterns:
        match = re.search(pattern, text)
        if match:
            result["天气"] = match.group(1)
            break
    
    # 提取气温
    temp_patterns = [
        r'气温\s*[为]?\s*(\d+)\s*[-~至]\s*(\d+)\s*[℃度]',
        r'(\d+)\s*[-~至]\s*(\d+)\s*[℃度]',
        r'气温\s*[为]?\s*(\d+)\s*[℃度]'
    ]
    for pattern in temp_patterns:
        match = re.search(pattern, text)
        if match:
            if match.lastindex == 2:
                result["气温"] = f"{match.group(1)}-{match.group(2)}℃"
            else:
                result["气温"] = f"{match.group(1)}℃"
            break
    
    # 提取日期（如果有）
    date_patterns = [
        r'(\d{4})[年](\d{1,2})[月](\d{1,2})[日]?',
        r'(\d{4})-(\d{2})-(\d{2})',
        r'(\d{2})月(\d{1,2})日'
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            if match.lastindex == 3:
                if len(match.group(1)) == 4:
                    result["日期"] = f"{match.group(1)}年{match.group(2)}月{match.group(3)}日"
                else:
                    result["日期"] = f"2026年{match.group(1)}月{match.group(2)}日"
            break
    
    # 提取人员（工种和数量）
    worker_patterns = [
        r'(\d+)\s*[个名]?\s*(水电工|钢筋工|木工|混凝土工|架子工|抹灰工|砌筑工|电焊工|安全员|施工员|质检员|材料员|资料员|杂工|管理人员)',
        r'(\d+)\s*[个名]?\s*工人',
        r'共\s*(\d+)\s*[个名]?\s*(?:工人|人员)'
    ]
    for pattern in worker_patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            if isinstance(m, tuple) and len(m) == 2:
                if m[0].isdigit():
                    count = int(m[0])
                    worker_type = m[1]
                else:
                    count = int(m[1])
                    worker_type = m[0]
                result["人员"][worker_type] = str(count)
            elif isinstance(m, str) and m.isdigit():
                result["人员"]["工人"] = m
    
    # 提取机械设备
    equip_patterns = [
        r'(\d+)\s*[台部]?\s*(振动棒|电焊机|切割机|塔吊|施工电梯|混凝土泵车|搅拌机|振捣器|手提切割机|电锤|冲击钻|经纬仪|水准仪|全站仪|挖掘机|推土机|汽车吊)',
        r'(振动棒|电焊机|切割机|塔吊|施工电梯|混凝土泵车|搅拌机|振捣器|手提切割机|电锤|冲击钻|经纬仪|水准仪|全站仪|挖掘机|推土机|汽车吊)\s*[为]?\s*(\d+)\s*[台部]?',
    ]
    for pattern in equip_patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            if isinstance(m, tuple) and len(m) == 2:
                if m[0].isdigit():
                    count = int(m[0])
                    equip = m[1]
                else:
                    count = int(m[1])
                    equip = m[0]
                result["机械设备"][equip] = f"{count}台"
    
    # 提取施工部位和内容
    location_patterns = [
        r'([A-Z0-9栋号楼层区]+)\s*([层楼座区])\s*([\u4e00-\u9fa5]+)',
        r'([A-Z0-9栋号楼层区]+)\s*([\u4e00-\u9fa5]+)',
    ]
    for pattern in location_patterns:
        match = re.search(pattern, text)
        if match:
            result["施工部位"] = text[match.start():match.end()]
            # 提取施工内容
            content_match = re.search(r'([浇筑|浇混凝土|穿线|布线|预埋|安装|焊接|切割|铺设|敷设|绑扎|浇筑混凝土|综合布线|水电|通风|消防|弱电|智能化]+)', text)
            if content_match:
                result["施工内容"] = content_match.group(1)
            break
    
    # 提取质量问题 - 按优先级，避免重复
    quality_issue_patterns = [
        r'发现\s*(\d+)\s*[处个]?\s*(蜂窝麻面|线缆外皮破损|套管位置偏差|混凝土离析|钢筋外露|蜂窝|麻面|破损|裂缝|气泡|露筋|孔洞|夹渣|松动|偏差|偏移)',
        r'(蜂窝麻面|线缆外皮破损|套管位置偏差|混凝土离析|钢筋外露)'
    ]
    for pattern in quality_issue_patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            if isinstance(m, tuple):
                issue_text = f"{m[0]}处{m[1]}"
            else:
                issue_text = m
            # 去重
            if issue_text not in result["质量问题"]:
                result["质量问题"].append(issue_text)
        if result["质量问题"]:
            break  # 匹配到就不再尝试更低优先级的模式
    
    # 提取安全问题
    safety_patterns = [
        r'(临边|洞口|高空|用电|消防|防护|围挡|临边防护|临边防护不到位|灭火器|特种作业|脚手架|模板支撑)',
        r'(未佩戴安全帽|未系安全带|安全防护不到位|安全隐患)'
    ]
    for pattern in safety_patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            if m not in result["安全问题"]:
                result["安全问题"].append(m)
    
    # 提取处理意见
    fix_patterns = [
        r'已\s*要求\s*([^\n，。,]+)',
        r'已\s*通知\s*([^\n，。,]+)',
        r'已\s*要求\s*整改',
        r'已\s*口头通知',
        r'要求\s*(\d+)\s*月\s*(\d+)\s*日\s*前\s*([^\n，。,]+)',
    ]
    for pattern in fix_patterns:
        match = re.search(pattern, text)
        if match:
            result["处理意见"] = text[match.start():match.end()]
            break
    
    # 提取备注/明日计划
    note_patterns = [
        r'备注[：:]([^\n]+)',
        r'明日计划[：:]([^\n]+)',
        r'明天[计划安排]+[：:]?([^\n，。,]+)',
    ]
    for pattern in note_patterns:
        match = re.search(pattern, text)
        if match:
            result["备注"] = match.group(1)
            break
    
    # 提取完成工程量
    progress_patterns = [
        r'完成\s*(\d+)%\s*([^\s，。,]+)',
        r'([^\s，。,]+)\s*完成\s*(\d+)%',
        r'完成\s*(\d+)%\s*'
    ]
    for pattern in progress_patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            if isinstance(m, tuple):
                result["完成工程量"][m[1]] = f"完成{m[0]}%"
            else:
                result["完成工程量"]["工程进度"] = f"完成{m}%"
    
    return result


def generate_log_content(log_type, parsed_data, project_info):
    """
    根据解析的数据生成标准日志内容
    返回结构化的内容字典
    """
    template = TEMPLATES.get(log_type, TEMPLATES["FH-001"])
    date_str = parsed_data.get("日期") or datetime.now().strftime("%Y年%m月%d日")
    weather = parsed_data.get("天气") or "晴"
    temp = parsed_data.get("气温") or ""
    
    content = {
        "title": template["name"],
        "header": {
            "项目名称": project_info.get("项目名称", ""),
            "施工单位": project_info.get("施工单位", "江苏国顺智能科技有限公司"),
            "监理单位": project_info.get("监理单位", ""),
            "日期": date_str,
            "天气": weather,
            "气温": temp
        }
    }
    
    # 根据不同日志类型生成内容
    if log_type in ["FH-001", "SG-001"]:
        # 监理日志/施工日志
        content["sections"] = {
            "一、施工部位及内容": parsed_data.get("施工部位", "") + " " + parsed_data.get("施工内容", ""),
            "二、投入人员": format_dict(parsed_data.get("人员", {}), "人"),
            "三、投入机械设备": format_dict(parsed_data.get("机械设备", {}), "台"),
            "四、完成工程量": format_dict(parsed_data.get("完成工程量", {}), ""),
            "五、质量情况": generate_quality_section(parsed_data),
            "六、安全文明施工": generate_safety_section(parsed_data),
            "七、发现的问题及处理意见": generate_issue_section(parsed_data),
            "八、备注": parsed_data.get("备注", "无")
        }
    elif log_type == "AQ-001":
        # 安全巡检记录
        content["sections"] = {
            "一、巡检区域/部位": parsed_data.get("施工部位", ""),
            "二、巡检人员": format_dict(parsed_data.get("人员", {}), "人"),
            "三、天气情况": f"{weather} {temp}".strip(),
            "四、检查内容": generate_safety_section(parsed_data),
            "五、发现的安全隐患": format_list(parsed_data.get("安全问题", [])),
            "六、整改要求及期限": parsed_data.get("处理意见", ""),
            "七、复查情况": "待复查" if parsed_data.get("安全问题") else "无隐患，无需复查",
            "八、备注": parsed_data.get("备注", "无")
        }
    elif log_type == "YJ-001":
        # 隐蔽工程验收记录
        content["sections"] = {
            "一、工程名称": project_info.get("项目名称", ""),
            "二、隐蔽项目": parsed_data.get("施工内容", ""),
            "三、隐蔽部位": parsed_data.get("施工部位", ""),
            "四、隐蔽日期": date_str,
            "五、施工单位": project_info.get("施工单位", "江苏国顺智能科技有限公司"),
            "六、施工内容": generate_work_description(parsed_data),
            "七、验收情况": "符合设计及规范要求" if not parsed_data.get("质量问题") else "发现以下问题：",
            "八、质量问题": format_list(parsed_data.get("质量问题", [])),
            "九、整改要求": parsed_data.get("处理意见", ""),
            "十、验收结论": "合格" if not parsed_data.get("质量问题") else "整改后合格"
        }
    elif log_type == "WD-001":
        # 危大工程巡视记录
        content["sections"] = {
            "一、巡视日期": date_str,
            "二、工程名称": project_info.get("项目名称", ""),
            "三、危大工程类型": "超危大工程" if "超" in parsed_data.get("施工内容", "") else "危大工程",
            "四、巡视人员": format_dict(parsed_data.get("人员", {}), "人"),
            "五、天气情况": f"{weather} {temp}".strip(),
            "六、施工部位及内容": parsed_data.get("施工部位", "") + " " + parsed_data.get("施工内容", ""),
            "七、施工状况": "正常" if not parsed_data.get("质量问题") else "发现问题",
            "八、安全防护措施检查": generate_safety_section(parsed_data),
            "九、存在问题": format_list(parsed_data.get("质量问题", []) + parsed_data.get("安全问题", [])),
            "十、处理意见": parsed_data.get("处理意见", ""),
            "十一、备注": parsed_data.get("备注", "无")
        }
    
    return content


def format_dict(d, unit):
    """格式化字典为多行文本"""
    if not d:
        return "无"
    lines = []
    for k, v in d.items():
        # 如果值已经包含单位（如"8人"），直接使用
        if isinstance(v, str) and any(u in v for u in ["人", "台", "套", "部"]):
            lines.append(f"{k}：{v}")
        elif unit:
            lines.append(f"{k}：{v}{unit}")
        else:
            lines.append(f"{k}：{v}")
    return "\n".join(lines)


def format_list(lst):
    """格式化列表为多行文本"""
    if not lst:
        return "无"
    return "\n".join(f"{i+1}. {item}" for i, item in enumerate(lst))


def generate_quality_section(parsed_data):
    """生成质量情况段落"""
    issues = parsed_data.get("质量问题", [])
    if not issues:
        return "主控项目：符合要求\n一般项目：符合要求\n整改要求：无"
    
    main = "发现以下质量问题："
    sub = "一般项目："
    fix = "整改要求："
    
    for issue in issues:
        main += f"\n- {issue}"
        sub += f"\n{issue}"
    
    fix += f"\n{parsed_data.get('处理意见', '立即整改')}"
    
    return f"主控项目：符合要求\n{sub}\n{fix}"


def generate_safety_section(parsed_data):
    """生成安全文明施工段落"""
    issues = parsed_data.get("安全问题", [])
    if not issues:
        return "施工区域围挡封闭良好\n作业人员佩戴安全帽\n现场材料堆放整齐\n未发生安全事故"
    
    lines = [
        "施工区域围挡封闭良好",
        "作业人员佩戴安全帽",
        "现场材料堆放整齐",
    ]
    for issue in issues:
        lines.append(f"发现隐患：{issue}")
    lines.append("已要求整改")
    
    return "\n".join(lines)


def generate_issue_section(parsed_data):
    """生成问题及处理意见段落"""
    issues = parsed_data.get("质量问题", []) + parsed_data.get("安全问题", [])
    if not issues:
        return "问题：无\n处理：无\n整改确认人：[待填写]"
    
    result = []
    for i, issue in enumerate(issues, 1):
        result.append(f"问题{i}：{issue}")
        result.append(f"处理：{parsed_data.get('处理意见', '已要求立即整改')}")
        result.append("整改确认人：[待填写]")
        result.append("")
    
    return "\n".join(result).strip()


def generate_work_description(parsed_data):
    """生成施工内容描述"""
    parts = []
    if parsed_data.get("施工部位"):
        parts.append(f"部位：{parsed_data.get('施工部位')}")
    if parsed_data.get("施工内容"):
        parts.append(f"内容：{parsed_data.get('施工内容')}")
    if parsed_data.get("人员"):
        parts.append(f"人员：{format_dict(parsed_data.get('人员', {}), '人')}")
    if parsed_data.get("机械设备"):
        parts.append(f"设备：{format_dict(parsed_data.get('机械设备', {}), '台')}")
    
    return "\n".join(parts) if parts else "按施工方案执行"


# ============================================================
# Word文档生成
# ============================================================

def create_word_document(content, output_path, log_type):
    """
    创建标准格式的Word文档
    """
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 设置页眉
    header = sections[0].header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    
    # 页眉左：项目名称
    left_run = hp.add_run(content.get("header", {}).get("项目名称", ""))
    left_run.font.size = Pt(9)
    left_run.font.name = "宋体"
    left_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 页眉中：施工单位
    mid_run = hp.add_run(" | " + content.get("header", {}).get("施工单位", ""))
    mid_run.font.size = Pt(9)
    mid_run.font.name = "宋体"
    mid_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 页眉右对齐
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 页眉第二行：日期
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = hp2.add_run(content.get("header", {}).get("日期", ""))
    date_run.font.size = Pt(9)
    date_run.font.name = "宋体"
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置页脚
    footer = sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 页脚：第X页 共Y页
    run = fp.add_run("")
    run.font.size = Pt(9)
    
    # 添加"第"字
    run1 = fp.add_run("第 ")
    run1.font.size = Pt(9)
    run1.font.name = "宋体"
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 页码域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    run2 = fp.add_run(" 页 共 ")
    run2.font.size = Pt(9)
    run2.font.name = "宋体"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    run3 = fp.add_run("")
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar3)
    run3._element.append(instrText2)
    run3._element.append(fldChar4)
    
    run4 = fp.add_run(" 页")
    run4.font.size = Pt(9)
    run4.font.name = "宋体"
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 主标题
    template = TEMPLATES.get(log_type, TEMPLATES["FH-001"])
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(template["title"])
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(22)
    title_run.bold = True
    
    # 副标题 - 项目信息
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtext = f"项目名称：{content.get('header', {}).get('项目名称', '')}"
    sub_run = subtitle.add_run(subtext)
    sub_run.font.name = "宋体"
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    sub_run.font.size = Pt(12)
    
    # 日期、天气、气温
    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_text = f"日期：{content.get('header', {}).get('日期', '')}    天气：{content.get('header', {}).get('天气', '')}    气温：{content.get('header', {}).get('气温', '')}"
    date_run = date_line.add_run(date_text)
    date_run.font.name = "宋体"
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    date_run.font.size = Pt(12)
    
    # 添加分隔线
    doc.add_paragraph()
    
    # 内容区域
    sections_data = content.get("sections", {})
    
    for section_title, section_content in sections_data.items():
        # 标题段落
        p_title = doc.add_paragraph()
        p_title_run = p_title.add_run(section_title)
        p_title_run.font.name = "黑体"
        p_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p_title_run.font.size = Pt(12)
        p_title_run.bold = True
        
        # 内容段落
        if section_content and section_content != "无":
            p_content = doc.add_paragraph()
            p_content.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
            p_content_run = p_content.add_run(section_content)
            p_content_run.font.name = "宋体"
            p_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_content_run.font.size = Pt(12)
        
        # 段后间距
        p_title.paragraph_format.space_after = Pt(6)
        if section_content and section_content != "无":
            p_content.paragraph_format.space_after = Pt(12)
    
    # 签名栏
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 创建签名表格
    sig_table = doc.add_table(rows=2, cols=4)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格样式
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run("")
            run.font.size = Pt(12)
    
    # 第一行：角色
    roles = ["施工员", "监理员", "项目经理", "技术负责人"]
    for i, role in enumerate(roles):
        cell = sig_table.cell(0, i)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(role)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
    
    # 第二行：签名和日期
    for i in range(4):
        cell = sig_table.cell(1, i)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run("____________\n______年______月______日")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
    
    # 保存文档
    doc.save(output_path)
    return output_path


def format_markdown(content, log_type):
    """
    生成Markdown格式的日志内容（用于预览）
    """
    template = TEMPLATES.get(log_type, TEMPLATES["FH-001"])
    header = content.get("header", {})
    
    lines = [
        f"【{template['name']}】",
        f"项目名称：{header.get('项目名称', '')}",
        f"日期：{header.get('日期', '')}  天气：{header.get('天气', '')}  气温：{header.get('气温', '')}",
        ""
    ]
    
    sections = content.get("sections", {})
    for title, text in sections.items():
        lines.append(title)
        if text:
            lines.append(text)
        lines.append("")
    
    lines.append("---")
    lines.append("施工员：______  监理员：______  项目经理：______  技术负责人：______")
    lines.append(f"日期：______")
    
    return "\n".join(lines)


# ============================================================
# 主程序入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="施工日志/巡检记录AI生成器")
    parser.add_argument("--type", "-t", default="FH-001", 
                        choices=["FH-001", "SG-001", "AQ-001", "YJ-001", "WD-001"],
                        help="日志类型")
    parser.add_argument("--input", "-i", required=True,
                        help="口头描述或文字描述")
    parser.add_argument("--project", "-p", 
                        help="项目信息JSON字符串，或项目名称（用于查找已保存信息）")
    parser.add_argument("--output", "-o", 
                        help="输出Word文件路径")
    parser.add_argument("--format", "-f", choices=["docx", "markdown", "both"], 
                        default="docx",
                        help="输出格式")
    
    args = parser.parse_args()
    
    # 解析输入文本
    parsed_data = parse_input_text(args.input)
    print(f"解析结果：{json.dumps(parsed_data, ensure_ascii=False, indent=2)}")
    
    # 加载项目信息
    project_info = load_project_info()
    
    # 如果命令行提供了项目信息，更新项目信息
    if args.project:
        if args.project.startswith("{"):
            # JSON格式的项目信息
            try:
                new_info = json.loads(args.project)
                if project_info:
                    project_info.update(new_info)
                else:
                    project_info = new_info
                save_project_info(project_info)
                print("项目信息已更新保存")
            except json.JSONDecodeError:
                # 只提供了项目名称
                if project_info:
                    project_info["项目名称"] = args.project
                else:
                    project_info = {"项目名称": args.project}
                save_project_info(project_info)
        else:
            # 只提供了项目名称
            if project_info:
                project_info["项目名称"] = args.project
            else:
                project_info = {"项目名称": args.project}
            save_project_info(project_info)
    
    # 如果还没有项目信息，使用默认信息
    if not project_info:
        project_info = {
            "项目名称": "未命名项目",
            "施工单位": "江苏国顺智能科技有限公司",
            "监理单位": "",
            "项目经理": "",
            "技术负责人": ""
        }
        print("警告：未找到项目信息，使用默认信息。请通过 --project 参数或首次交互输入完整项目信息。")
    
    # 生成日志内容
    content = generate_log_content(args.type, parsed_data, project_info)
    
    # 输出
    if args.format in ["markdown", "both"]:
        md = format_markdown(content, args.type)
        print("\n" + "="*60)
        print("生成的日志内容（Markdown格式）：")
        print("="*60)
        print(md)
    
    if args.format in ["docx", "both"]:
        output_path = args.output
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"inspection_log_{args.type}_{timestamp}.docx"
        
        docx_path = create_word_document(content, output_path, args.type)
        print(f"\nWord文档已生成：{docx_path}")
        
        return docx_path
    
    return None


if __name__ == "__main__":
    main()

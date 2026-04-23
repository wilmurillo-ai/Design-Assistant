#!/usr/bin/env python3
"""
read_requirement_file.py
需求文档读取脚本 —— 根据文件扩展名自动选择合适的读取方式，
将内容提取为纯文本后输出到 stdout，供后续流程解析需求点。

用法：
    python scripts/read_requirement_file.py <文件路径>

支持格式：
    .txt / .md         直接读取
    .docx              python-docx 提取段落与表格
    .pdf               pypdf 提取文本（逐页）
    .xlsx / .xls / .csv  pandas 读取，转换为文本表格
"""

import sys
import os


def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit("[错误] 缺少依赖：pip install python-docx --break-system-packages")

    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def read_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # fallback
        except ImportError:
            sys.exit("[错误] 缺少依赖：pip install pypdf --break-system-packages")

    reader = PdfReader(path)
    pages_text = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(f"--- 第 {i} 页 ---\n{text.strip()}")
    return "\n\n".join(pages_text)


def read_excel(path: str) -> str:
    try:
        import pandas as pd
    except ImportError:
        sys.exit("[错误] 缺少依赖：pip install pandas openpyxl --break-system-packages")

    ext = os.path.splitext(path)[1].lower()
    if ext == ".xls":
        df_dict = pd.read_excel(path, sheet_name=None, engine="xlrd")
    else:
        df_dict = pd.read_excel(path, sheet_name=None, engine="openpyxl")

    lines = []
    for sheet_name, df in df_dict.items():
        lines.append(f"=== 工作表：{sheet_name} ===")
        lines.append(df.fillna("").to_string(index=False))
    return "\n\n".join(lines)


def read_csv(path: str) -> str:
    try:
        import pandas as pd
    except ImportError:
        sys.exit("[错误] 缺少依赖：pip install pandas --break-system-packages")

    df = pd.read_csv(path, encoding="utf-8", errors="replace")
    return df.fillna("").to_string(index=False)


def main():
    if len(sys.argv) < 2:
        print("用法: python scripts/read_requirement_file.py <文件路径>", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]

    if not os.path.exists(path):
        print(f"[错误] 文件不存在：{path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()

    dispatch = {
        ".txt":  read_txt,
        ".md":   read_txt,
        ".docx": read_docx,
        ".pdf":  read_pdf,
        ".xlsx": read_excel,
        ".xlsm": read_excel,
        ".xls":  read_excel,
        ".csv":  read_csv,
    }

    reader = dispatch.get(ext)
    if reader is None:
        print(f"[错误] 不支持的文件格式：{ext}", file=sys.stderr)
        print("支持格式：.txt .md .docx .pdf .xlsx .xls .xlsm .csv", file=sys.stderr)
        sys.exit(1)

    content = reader(path)
    print(content)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
Convert Markdown to Word (.docx)
"""
import sys
import re
from docx import Document
from docx.shared import Pt

def md_to_docx(input_path, output_path=None):
    """Convert Markdown file to .docx"""
    try:
        if not output_path:
            output_path = input_path.replace('.md', '.docx')
        
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        doc = Document()
        
        lines = md_content.split('\n')
        for line in lines:
            line = line.rstrip()
            if not line:
                continue
            
            # Heading 1
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            # Heading 2
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            # Heading 3
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            # Heading 4
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)
            # Bullet list
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered list
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                doc.add_paragraph(text, style='List Number')
            # Regular paragraph
            else:
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return f"Converted to {output_path}"
    except Exception as e:
        return f"Error converting {input_path}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    print(md_to_docx(input_path, output_path))

#!/usr/bin/env python3
"""
Convert Word (.docx) to Markdown
"""
import sys
import os
from docx import Document

def docx_to_md(input_path, output_path=None):
    """Convert .docx file to Markdown"""
    try:
        doc = Document(input_path)
        
        if not output_path:
            output_path = input_path.replace('.docx', '.md')
        
        md_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect heading level by style name
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                level = style_name.replace('heading ', '').replace('heading', '1')
                try:
                    level = int(level)
                    md_content.append('#' * level + ' ' + text)
                except:
                    md_content.append(text)
            else:
                md_content.append(text)
        
        # Convert tables
        for table in doc.tables:
            md_content.append('')
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                md_content.append('| ' + ' | '.join(cells) + ' |')
                if i == 0:
                    md_content.append('|' + '|'.join(['---'] * len(cells)) + '|')
            md_content.append('')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_content))
        
        return f"Converted to {output_path}"
    except Exception as e:
        return f"Error converting {input_path}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python docx_to_md.py <input.docx> [output.md]")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    print(docx_to_md(input_path, output_path))

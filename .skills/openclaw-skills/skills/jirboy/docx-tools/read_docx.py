#!/usr/bin/env python3
"""
Read content from Word (.docx) files
"""
import sys
from docx import Document

def read_docx(filepath):
    """Read all text content from a .docx file"""
    try:
        doc = Document(filepath)
        content = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        return f"Error reading {filepath}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <filepath>")
        sys.exit(1)
    
    filepath = sys.argv[1]
    print(read_docx(filepath))

#!/usr/bin/env python3
"""
Create and write Word (.docx) files
"""
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def write_docx(filepath, title=None, content=None, headings=None):
    """Create a .docx file with given content"""
    try:
        doc = Document()
        
        # Add title
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        if content:
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(item)
            else:
                doc.add_paragraph(content)
        
        # Add headings and sections
        if headings:
            for level, text in headings:
                doc.add_heading(text, level)
        
        doc.save(filepath)
        return f"Successfully saved to {filepath}"
    except Exception as e:
        return f"Error writing {filepath}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python write_docx.py <filepath> [title] [content]")
        sys.exit(1)
    
    filepath = sys.argv[1]
    title = sys.argv[2] if len(sys.argv) > 2 else None
    content = sys.argv[3] if len(sys.argv) > 3 else None
    
    print(write_docx(filepath, title, content))

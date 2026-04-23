#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档中文格式标准化工具

该脚本用于将Word文档按照中文格式标准进行格式化，包括标题、正文、图片和图表的格式统一。
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class WordFormatter:
    """Word文档格式化类"""
    
    def __init__(self, input_file, output_file=None):
        """
        初始化格式化器
        
        Args:
            input_file: 输入Word文档路径
            output_file: 输出文档路径，默认为None
        """
        self.input_file = input_file
        
        if output_file:
            self.output_file = output_file
        else:
            # 使用时间戳作为后缀
            base_name = os.path.splitext(input_file)[0]
            timestamp = datetime.now().strftime("_%Y%m%d%H%M%S")
            self.output_file = f"{base_name}{timestamp}.docx"
        
        # 加载文档
        self.doc = Document(input_file)
    
    def format_document(self):
        """格式化整个文档"""
        self._format_styles()
        self._format_paragraphs()
        self._format_images()
        self._format_tables()
        
        # 保存文档
        self.doc.save(self.output_file)
        print(f"文档已格式化并保存到: {self.output_file}")
    
    def _format_styles(self):
        """格式化文档样式"""
        # 设置默认字体为微软雅黑
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        # 确保中文字体也设置为微软雅黑
        rPr = font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        font.size = Pt(10.5)  # 5号字体
        
        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.first_line_indent = Inches(0.35)  # 首行缩进2字符
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        
        # 标题样式
        # 标题（Title）
        title_style = self.doc.styles['Title']
        title_font = title_style.font
        title_font.name = 'Microsoft YaHei'
        # 确保中文字体也设置为微软雅黑
        rPr = title_font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        title_font.size = Pt(22)  # 2号字体
        title_font.bold = True
        title_paragraph = title_style.paragraph_format
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_paragraph.space_before = Pt(12)
        title_paragraph.space_after = Pt(6)
        
        # 一级标题（Heading 1）
        heading1_style = self.doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Microsoft YaHei'
        # 确保中文字体也设置为微软雅黑
        rPr = heading1_font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading1_font.size = Pt(16)  # 3号字体
        heading1_font.bold = True
        heading1_paragraph = heading1_style.paragraph_format
        heading1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading1_paragraph.space_before = Pt(12)
        heading1_paragraph.space_after = Pt(6)
        
        # 二级标题（Heading 2）
        heading2_style = self.doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'Microsoft YaHei'
        # 确保中文字体也设置为微软雅黑
        rPr = heading2_font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading2_font.size = Pt(14)  # 4号字体
        heading2_font.bold = True
        heading2_paragraph = heading2_style.paragraph_format
        heading2_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading2_paragraph.space_before = Pt(6)
        heading2_paragraph.space_after = Pt(3)
        
        # 三级标题（Heading 3）
        heading3_style = self.doc.styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = 'Microsoft YaHei'
        # 确保中文字体也设置为微软雅黑
        rPr = heading3_font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading3_font.size = Pt(12)  # 小4号字体
        heading3_font.bold = True
        heading3_paragraph = heading3_style.paragraph_format
        heading3_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading3_paragraph.space_before = Pt(3)
        heading3_paragraph.space_after = Pt(3)
    
    def _format_paragraphs(self):
        """格式化所有段落"""
        for paragraph in self.doc.paragraphs:
            # 跳过空段落
            if not paragraph.text.strip():
                continue
            
            # 检查是否是标题
            if paragraph.style.name == 'Title':
                # 标题：居中对齐，无缩进
                paragraph.style = self.doc.styles['Title']
                paragraph.paragraph_format.first_line_indent = 0
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # 直接设置字体为微软雅黑
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    # 设置中文字体
                    if run.font._element.rPr is not None:
                        rFonts = run.font._element.rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), '微软雅黑')
            elif paragraph.style.name.startswith('Heading'):
                # 一级、二级、三级标题：左对齐，无缩进
                paragraph.style = self.doc.styles[paragraph.style.name]
                paragraph.paragraph_format.first_line_indent = 0
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # 直接设置字体为微软雅黑
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    # 设置中文字体
                    if run.font._element.rPr is not None:
                        rFonts = run.font._element.rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                # 正文：左对齐，首行缩进2字符
                paragraph.style = self.doc.styles['Normal']
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # 直接设置字体为微软雅黑
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    # 设置中文字体
                    if run.font._element.rPr is not None:
                        rFonts = run.font._element.rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _format_images(self):
        """格式化图片和图表"""
        for paragraph in self.doc.paragraphs:
            # 检查段落是否包含图片或图表
            has_image = False
            for run in paragraph.runs:
                # 更精确地检测图片和图表
                xml_content = run._element.xml.lower()
                if ('<wp:inline' in xml_content or '<wp:anchor' in xml_content or 
                    '<c:chart' in xml_content):
                    has_image = True
                    break
            
            if has_image:
                # 图片/图表段落不需要缩进
                paragraph.paragraph_format.first_line_indent = 0
                # 图片/图表段落居中
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def _format_tables(self):
        """格式化表格"""
        for table in self.doc.tables:
            # 设置表格居中
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 格式化表格内容
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = self.doc.styles['Normal']

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("使用方法: python format_word.py <input_file> [output_file]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"错误: 文件 {input_file} 不存在")
        sys.exit(1)
    
    try:
        formatter = WordFormatter(input_file, output_file)
        formatter.format_document()
    except Exception as e:
        print(f"错误: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
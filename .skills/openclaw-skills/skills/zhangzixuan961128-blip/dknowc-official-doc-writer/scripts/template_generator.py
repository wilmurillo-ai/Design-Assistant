#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文红头文件生成器（v4.3 - 简化版）

设计逻辑：
1. 接收已生成的普通格式文档路径
2. 提取参考资料和知识专库链接内容
3. 去除正文中的引用标记 [^1^] 等
4. 在该文档上插入红头表格（开头）
5. 在正文末尾、参考来源之前插入表尾表格
6. 在表尾后重新添加参考资料和知识专库链接
7. 更新占位符
8. 另存为红头版本
"""

import sys
import os
import re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import argparse

# 模板目录
TEMPLATE_DIR = Path(__file__).parent.parent / "templates"

# 模板映射
TEMPLATE_MAP = {
    "请示": ("上行文件格式-红头.docx", "上行文件格式-表尾.docx"),
    "报告": ("上行文件格式-红头.docx", "上行文件格式-表尾.docx"),
    "通知": ("下行文件格式-红头.docx", "下行文件格式-表尾.docx"),
    "通报": ("下行文件格式-红头.docx", "下行文件格式-表尾.docx"),
    "函": ("函件（普通）格式-红头.docx", "函件（普通）格式-表尾.docx"),
    "复函": ("函件（普通）格式-红头.docx", "函件（普通）格式-表尾.docx"),
    "工作会议纪要": ("工作会议纪要-红头.docx", "工作会议纪要-表尾.docx"),
    "局长办公会议纪要": ("局长办公会议纪要-红头.docx", "局长办公会议纪要-表尾.docx"),
    "党组会议纪要": ("党组会议纪要-红头.docx", "党组会议纪要-表尾.docx"),
    "会议纪要": ("工作会议纪要-红头.docx", "工作会议纪要-表尾.docx"),
}


def remove_reference_markers_from_doc(doc):
    """去除文档中所有段落的引用标记 [^1^] [^2^] 等"""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and re.search(r'\[\^\d+\^\]', run.text):
                run.text = re.sub(r'\[\^\d+\^\]', '', run.text)
                count += 1
    return count


def find_references_and_links_range(doc):
    """
    找到参考资料和知识专库链接的段落范围，以及分页符位置
    
    Returns:
        tuple: (分页符索引, 参考来源起始索引, 知识专库链接结束索引)
    """
    page_break_idx = None
    ref_start = None
    ref_end = None
    
    for i, para in enumerate(doc.paragraphs):
        # 检查是否有分页符
        for run in para.runs:
            if hasattr(run, '_element'):
                for br in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                    break_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if break_type == 'page':
                        page_break_idx = i
        
        text = para.text.strip()
        
        # 检测参考资料标题
        if ref_start is None:
            if text in ["参考资料", "参考来源"] or text.startswith("【参考资料】") or text.startswith("【参考来源】"):
                ref_start = i
        
        # 检测知识专库链接
        if ref_start is not None and ref_end is None:
            if text.startswith("http") or not text:
                continue
            elif text == "知识专库链接" or text.startswith("【知识专库链接】"):
                ref_end = i
    
    # 如果没找到知识专库链接，结束索引为文档末尾
    if ref_start is not None and ref_end is None:
        ref_end = len(doc.paragraphs)
    
    return page_break_idx, ref_start, ref_end


def is_minutes_type(doc_type):
    """判断是否为会议纪要类型"""
    return doc_type in ("工作会议纪要", "局长办公会议纪要", "党组会议纪要", "会议纪要")


def update_wps_textboxes(doc_elem, replacements):
    """更新文档中所有WPS文本框(wps:txbx)里的占位符"""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_wps = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    
    count = 0
    txbxs = doc_elem.findall(f'.//{{{ns_wps}}}txbx')
    for txbx in txbxs:
        txbxContent = txbx.find(f'{{{ns_w}}}txbxContent')
        if txbxContent is None:
            continue
        for t in txbxContent.findall(f'.//{{{ns_w}}}t'):
            if t.text:
                for key, value in replacements.items():
                    placeholder = f"【{key}】"
                    if placeholder in t.text:
                        t.text = t.text.replace(placeholder, value)
                        count += 1
    return count


def update_minutes_footer_table(table, replacements):
    """更新纪要表尾单个表格中的占位符"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if "【出席人员】" in text:
                        run.text = text.replace("【出席人员】", replacements.get("出席人员", ""))
                    elif "【列席人员】" in text:
                        run.text = text.replace("【列席人员】", replacements.get("列席人员", ""))
                    elif "【印发单位】" in text:
                        run.text = text.replace("【印发单位】", replacements.get("印发单位", ""))
                    elif "【印发日期】" in text:
                        run.text = text.replace("【印发日期】", replacements.get("印发日期", ""))
                    elif "【成文日期】" in text:
                        run.text = text.replace("【成文日期】", replacements.get("成文日期", ""))
                    elif "【抄送】" in text:
                        run.text = text.replace("【抄送】", replacements.get("抄送", ""))


def is_table_row_empty(row):
    """判断表格行是否所有单元格均为空（忽略空格）"""
    for cell in row.cells:
        for para in cell.paragraphs:
            if para.text.strip():
                return False
    return True


def remove_empty_rows(table):
    """删除表格中所有为空的行（倒序遍历避免索引偏移）"""
    for row in reversed(table.rows):
        if is_table_row_empty(row):
            row._tr.getparent().remove(row._tr)


def generate_red_header_document(doc_type, input_path, replacements, output_path):
    """生成红头文件"""
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"普通格式文档不存在: {input_path}")

    doc = Document(input_path)
    print(f"✓ 已加载普通格式文档: {input_path}")

    # 1. 找到分页符位置、参考资料和知识专库链接的范围
    page_break_idx, ref_start, ref_end = find_references_and_links_range(doc)
    if page_break_idx:
        print(f"✓ 分页符位置: 段落 {page_break_idx}")
    print(f"✓ 参考资料范围: 段落 {ref_start} 到 {ref_end}") if ref_start else print("✓ 未找到参考资料")

    # 2. 去除正文中的引用标记（参考资料之前的内容）
    if ref_start:
        for para in doc.paragraphs[:ref_start]:
            for run in para.runs:
                if run.text and re.search(r'\[\^\d+\^\]', run.text):
                    run.text = re.sub(r'\[\^\d+\^\]', '', run.text)
    else:
        remove_reference_markers_from_doc(doc)
    print(f"✓ 已去除正文引用标记")

    is_minutes = is_minutes_type(doc_type)

    # 3. 插入红头
    template_pair = TEMPLATE_MAP.get(doc_type, ("下行文件格式-红头.docx", "下行文件格式-表尾.docx"))
    red_template_name, footer_template_name = template_pair
    
    red_template_path = TEMPLATE_DIR / red_template_name
    if not red_template_path.exists():
        raise FileNotFoundError(f"红头模板不存在: {red_template_path}")
    
    red_template = Document(red_template_path)
    
    if is_minutes:
        # 纪要红头：可能是表格（如党组会议纪要）或WPS文本框（如工作会议纪要）
        if len(red_template.tables) >= 1:
            # 表格型红头：跟普通公文一样处理
            red_table_xml = deepcopy(red_template.tables[0]._tbl)
            doc._body._body.insert(0, red_table_xml)
            set_table_keep_together(doc.tables[0])
            update_red_header_table(doc.tables[0], replacements)
            remove_empty_rows(doc.tables[0])
            print(f"✓ 纪要红头表格已插入")
        else:
            # WPS文本框型红头：整体复制模板body内容
            red_body = red_template._body._body
            target_body = doc._body._body
            
            for child in list(red_body):
                if child.tag.endswith('}p') or child.tag.endswith('}tbl') or 'wps' in child.tag or 'v:' in child.tag or 'mc:' in child.tag:
                    target_body.insert(0, deepcopy(child))
            
            count = update_wps_textboxes(target_body, replacements)
            print(f"✓ 纪要红头已插入（WPS文本框替换{count}处）")
            
            # 浮动定位形状不占文档流空间，给正文标题加段前间距
            ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for child in list(target_body):
                if child.tag == f'{{{ns_w}}}p':
                    t = child.find(f'.//{{{ns_w}}}t')
                    if t is not None and t.text and t.text.strip():
                        pPr = child.find(f'{{{ns_w}}}pPr')
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            child.insert(0, pPr)
                        spacing = pPr.find(f'{{{ns_w}}}spacing')
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        existing = int(spacing.get(qn('w:before'), '0'))
                        spacing.set(qn('w:before'), str(existing + 4060))
                        break
    else:
        # 普通公文红头：使用表格
        if len(red_template.tables) >= 1:
            red_table_xml = deepcopy(red_template.tables[0]._tbl)
            doc._body._body.insert(0, red_table_xml)
            set_table_keep_together(doc.tables[0])
            update_red_header_table(doc.tables[0], replacements)
            print(f"✓ 红头表格已插入")

        update_vml_shapes(doc, replacements)

    # 4. 插入表尾表格
    footer_template_path = TEMPLATE_DIR / footer_template_name
    if not footer_template_path.exists():
        raise FileNotFoundError(f"结尾模板不存在: {footer_template_path}")
    
    footer_template = Document(footer_template_path)
    
    if is_minutes:
        # 纪要表尾可能有多个表格（出席列席 + 印发）
        for i, ft_table in enumerate(footer_template.tables):
            footer_table_xml = deepcopy(ft_table._tbl)
            
            if page_break_idx:
                page_break_para_element = doc.paragraphs[page_break_idx]._element
                page_break_para_element.addprevious(footer_table_xml)
            elif ref_start:
                ref_para_element = doc.paragraphs[ref_start]._element
                ref_para_element.addprevious(footer_table_xml)
            else:
                doc._body._body.append(footer_table_xml)
            
            set_table_keep_together(doc.tables[-1])
        
        # 重新获取所有表格，替换纪要表尾占位符
        # 找到插入的表尾表格（从后往前数）
        footer_table_count = len(footer_template.tables)
        for i in range(footer_table_count):
            table = doc.tables[-(footer_table_count - i)]
            update_minutes_footer_table(table, replacements)
            remove_empty_rows(table)
        
        pos_desc = "分页符之前" if page_break_idx else ("参考来源之前" if ref_start else "文档末尾")
        print(f"✓ 纪要表尾已插入（{pos_desc}，{footer_table_count}个表格）")
    else:
        if page_break_idx and len(footer_template.tables) >= 1:
            page_break_para_element = doc.paragraphs[page_break_idx]._element
            footer_table_xml = deepcopy(footer_template.tables[0]._tbl)
            page_break_para_element.addprevious(footer_table_xml)
            set_table_keep_together(doc.tables[-1])
            update_footer_table(doc.tables[-1], replacements)
            print(f"✓ 结尾表格已插入（在分页符之前，将定位到正文页底部）")
        elif ref_start and len(footer_template.tables) >= 1:
            ref_para_element = doc.paragraphs[ref_start]._element
            footer_table_xml = deepcopy(footer_template.tables[0]._tbl)
            ref_para_element.addprevious(footer_table_xml)
            set_table_keep_together(doc.tables[-1])
            update_footer_table(doc.tables[-1], replacements)
            print(f"✓ 结尾表格已插入（在参考来源之前）")
        elif len(footer_template.tables) >= 1:
            footer_table_xml = deepcopy(footer_template.tables[0]._tbl)
            doc._body._body.append(footer_table_xml)
            set_table_keep_together(doc.tables[-1])
            update_footer_table(doc.tables[-1], replacements)
            print(f"✓ 结尾表格已插入（在文档末尾）")

    # 5. 保存文件
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return str(output_path)


def add_fullwidth_space_if_needed(text):
    """为两字密级/紧急程度自动添加全角空格"""
    if text and len(text) == 2:
        if all('\u4e00' <= c <= '\u9fff' for c in text):
            return text[0] + '\u3000' + text[1]
    return text


def update_red_header_table(table, replacements):
    """更新红头表格中的占位符"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = add_fullwidth_space_if_needed(replacements.get("密级", ""))
    urgency = add_fullwidth_space_if_needed(replacements.get("紧急程度", ""))
    signer = replacements.get("签发人", "")
    meeting_number = replacements.get("纪要编号", "")
    print_unit = replacements.get("印发单位", "")
    print_date = replacements.get("印发日期", "")

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if "【发文机关】" in text:
                        run.text = text.replace("【发文机关】", org_name)
                    if "【发文字号】" in text:
                        run.text = text.replace("【发文字号】", doc_number)
                    if "【密级】" in text:
                        run.text = text.replace("【密级】", secrecy if secrecy else "")
                    if "【紧急程度】" in text:
                        run.text = text.replace("【紧急程度】", urgency if urgency else "")
                    if "【签发人】" in text:
                        run.text = text.replace("【签发人】", signer if signer else "")
                    if "【纪要编号】" in text:
                        run.text = text.replace("【纪要编号】", meeting_number)
                    if "【印发单位】" in text:
                        run.text = text.replace("【印发单位】", print_unit)
                    if "【成文日期】" in text and "印发" not in text:
                        run.text = text.replace("【成文日期】", print_date)


def update_vml_shapes(doc, replacements):
    """更新VML图形中的占位符"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = replacements.get("密级", "")
    urgency = replacements.get("紧急程度", "")
    meeting_number = replacements.get("纪要编号", "")
    print_unit = replacements.get("印发单位", "")
    print_date = replacements.get("印发日期", "")

    body = doc._body._body
    text_nodes = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')

    for t in text_nodes:
        if t.text:
            if "【发文机关】" in t.text:
                t.text = t.text.replace("【发文机关】", org_name)
            if "【发文字号】" in t.text:
                t.text = t.text.replace("【发文字号】", doc_number)
            if "【密级】" in t.text:
                t.text = t.text.replace("【密级】", secrecy if secrecy else "")
            if "【紧急程度】" in t.text:
                t.text = t.text.replace("【紧急程度】", urgency if urgency else "")
            if "【纪要编号】" in t.text:
                t.text = t.text.replace("【纪要编号】", meeting_number)
            if "【印发单位】" in t.text:
                t.text = t.text.replace("【印发单位】", print_unit)
            if "【成文日期】" in t.text and "印发" not in t.text:
                t.text = t.text.replace("【成文日期】", print_date)


def set_table_keep_together(table):
    """设置表格不允许跨页拆分"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    keep_together = OxmlElement('w:keep-together')
    keep_together.set(qn('w:val'), '1')
    tblPr.append(keep_together)


def update_footer_table(table, replacements):
    """更新结尾表格中的占位符"""
    org_name = replacements.get("发文机关", "")
    date_str = replacements.get("成文日期", "")
    cc_list = replacements.get("抄送", "")

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if "【发文机关】" in text:
                        run.text = text.replace("【发文机关】", org_name)
                    if "【成文日期】" in text:
                        run.text = text.replace("【成文日期】", date_str)
                    if "【印发日期】" in text:
                        run.text = text.replace("【印发日期】", date_str)
                    if "【公开方式】" in text:
                        run.text = text.replace("【公开方式】", "主动公开")
                    if "【抄送】" in text:
                        run.text = text.replace("【抄送】", cc_list)


def main():
    parser = argparse.ArgumentParser(description='红头文件生成器 v4.4')
    parser.add_argument('type', help='文种')
    parser.add_argument('--input', required=True, help='普通格式文档路径')
    parser.add_argument('--org', default='', help='发文机关')
    parser.add_argument('--doc-number', default='', help='发文字号')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--signer', default='', help='签发人')
    # 会议纪要专用参数
    parser.add_argument('--meeting-number', default='', help='纪要编号（如〔2026〕第10号）')
    parser.add_argument('--attendees', default='', help='出席人员')
    parser.add_argument('--non-voting', default='', help='列席人员')
    parser.add_argument('--print-unit', default='', help='印发单位')
    parser.add_argument('--print-date', default='', help='印发日期（如2026年3月19日）')
    parser.add_argument('--cc', default='', help='抄送单位')

    args = parser.parse_args()

    input_path = Path(args.input)
    if args.output:
        output_path = args.output
    else:
        output_path = str(input_path.parent / f"{input_path.stem}_红头{input_path.suffix}")

    today = datetime.now().strftime("%Y年%m月%d日")

    replacements = {
        "发文机关": args.org,
        "发文字号": args.doc_number,
        "成文日期": today,
        "签发人": args.signer,
        # 会议纪要专用
        "纪要编号": args.meeting_number,
        "出席人员": args.attendees,
        "列席人员": args.non_voting,
        "印发单位": args.print_unit if args.print_unit else f"{args.org}办公室",
        "印发日期": args.print_date if args.print_date else today,
        "抄送": args.cc,
    }

    output = generate_red_header_document(args.type, args.input, replacements, output_path)
    print(f"✅ 红头文件已生成: {output}")


if __name__ == "__main__":
    main()

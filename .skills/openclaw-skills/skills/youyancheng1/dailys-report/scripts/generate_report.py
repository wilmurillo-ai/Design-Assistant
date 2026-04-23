#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Daily Report Generator
每日日报生成器

Usage:
    python generate_report.py                    # 生成今日日报
    python generate_report.py --format md        # Markdown格式
    python generate_report.py --format docx      # Word文档
    python generate_report.py --date 2026-03-17  # 指定日期
"""

import os
import sys
import json
import argparse
from datetime import datetime, timedelta
from pathlib import Path

# Windows UTF-8 编码修复
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8') if hasattr(sys.stdout, 'reconfigure') else None

# 配置路径
WORKSPACE = Path.home() / ".openclaw" / "workspace"
MEMORY_DIR = WORKSPACE / "memory"
PROJECTS_FILE = WORKSPACE / "PROJECTS.md"
CONFIG_FILE = Path.home() / ".openclaw" / "daily-report-config.json"
OUTPUT_DIR = Path.home() / "Desktop"


def load_config():
    """加载配置"""
    if CONFIG_FILE.exists():
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {
        "template": "default",
        "output_dir": str(OUTPUT_DIR),
        "include_time": True,
        "include_stats": True,
        "language": "zh-CN"
    }


def get_memory_file(date_str):
    """获取记忆文件路径"""
    return MEMORY_DIR / f"{date_str}.md"


def read_memory_content(date_str):
    """读取今日记忆内容"""
    memory_file = get_memory_file(date_str)
    if memory_file.exists():
        with open(memory_file, 'r', encoding='utf-8') as f:
            return f.read()
    return ""


def parse_tasks(content):
    """从内容中解析任务"""
    tasks = []
    completed = []
    in_progress = []
    
    for line in content.split('\n'):
        line = line.strip()
        # 匹配 - [x] 或 - [ ] 格式的任务
        if line.startswith('- [x]') or line.startswith('- [X]'):
            completed.append(line[6:].strip())
        elif line.startswith('- [ ]'):
            in_progress.append(line[6:].strip())
        elif line.startswith('- ') and len(line) > 2:
            # 普通列表项
            tasks.append(line[2:].strip())
    
    return {
        "completed": completed,
        "in_progress": in_progress,
        "tasks": tasks
    }


def generate_text_report(date_str, tasks, stats=True):
    """生成纯文本日报"""
    date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    date_display = date_obj.strftime("%Y年%m月%d日")
    
    report = f"""【日报】{date_display}

一、今日完成
"""
    
    if tasks["completed"]:
        for task in tasks["completed"]:
            report += f"- [x] {task}\n"
    else:
        report += "(暂无已完成任务)\n"
    
    if tasks["in_progress"]:
        report += "\n二、进行中任务\n"
        for task in tasks["in_progress"]:
            report += f"- [ ] {task}\n"
    
    if tasks["tasks"] and not tasks["completed"] and not tasks["in_progress"]:
        report += "\n二、今日事项\n"
        for task in tasks["tasks"]:
            report += f"- {task}\n"
    
    # 统计信息
    if stats:
        completed_count = len(tasks["completed"])
        in_progress_count = len(tasks["in_progress"])
        total = completed_count + in_progress_count
        
        report += f"""
──────────────
📊 今日统计: 
   完成 {completed_count} / {total} 项
   完成率: {(completed_count/total*100) if total > 0 else 0:.0f}%
"""
    
    return report


def generate_markdown_report(date_str, tasks, stats=True):
    """生成Markdown日报"""
    date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    date_display = date_obj.strftime("%Y年%m月%d日")
    weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][date_obj.weekday()]
    
    report = f"""# 📋 每日日报 - {date_display} ({weekday})

## ✅ 今日完成

"""
    
    if tasks["completed"]:
        for task in tasks["completed"]:
            report += f"- {task}\n"
    else:
        report += "*暂无已完成任务*\n"
    
    if tasks["in_progress"]:
        report += f"""
## 🔄 进行中

"""
        for task in tasks["in_progress"]:
            report += f"- {task}\n"
    
    if tasks["tasks"] and not tasks["completed"] and not tasks["in_progress"]:
        report += f"""
## 📝 今日事项

"""
        for task in tasks["tasks"]:
            report += f"- {task}\n"
    
    # 统计信息
    if stats:
        completed_count = len(tasks["completed"])
        in_progress_count = len(tasks["in_progress"])
        total = completed_count + in_progress_count
        
        report += f"""
---

📊 **今日统计**: 完成 {completed_count}/{total} 项 ({(completed_count/total*100) if total > 0 else 0:.0f}%)
"""
    
    return report


def generate_docx_report(date_str, tasks, output_path):
    """生成Word文档日报"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ 需要安装 python-docx: pip install python-docx")
        return None
    
    date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    date_display = date_obj.strftime("%Y年%m月%d日")
    weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][date_obj.weekday()]
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(f"每日日报 - {date_display} ({weekday})", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 今日完成
    doc.add_heading("✅ 今日完成", level=1)
    if tasks["completed"]:
        for task in tasks["completed"]:
            doc.add_paragraph(task, style='List Bullet')
    else:
        doc.add_paragraph("暂无已完成任务")
    
    # 进行中
    if tasks["in_progress"]:
        doc.add_heading("🔄 进行中", level=1)
        for task in tasks["in_progress"]:
            doc.add_paragraph(task, style='List Bullet')
    
    # 今日事项
    if tasks["tasks"] and not tasks["completed"] and not tasks["in_progress"]:
        doc.add_heading("📝 今日事项", level=1)
        for task in tasks["tasks"]:
            doc.add_paragraph(task, style='List Bullet')
    
    # 统计
    completed_count = len(tasks["completed"])
    in_progress_count = len(tasks["in_progress"])
    total = completed_count + in_progress_count
    
    doc.add_paragraph()  # 空行
    stats_para = doc.add_paragraph()
    stats_para.add_run("📊 今日统计: ").bold = True
    stats_para.add_run(f"完成 {completed_count}/{total} 项 ({(completed_count/total*100) if total > 0 else 0:.0f}%)")
    
    # 保存
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='每日日报生成器')
    parser.add_argument('--date', type=str, default=None, help='日期 (YYYY-MM-DD), 默认今天')
    parser.add_argument('--format', type=str, default='text', choices=['text', 'md', 'docx'], help='输出格式')
    parser.add_argument('--output', type=str, default=None, help='输出文件路径')
    parser.add_argument('--no-stats', action='store_true', help='不显示统计信息')
    
    args = parser.parse_args()
    
    # 确定日期
    if args.date:
        date_str = args.date
    else:
        date_str = datetime.now().strftime("%Y-%m-%d")
    
    # 读取记忆内容
    content = read_memory_content(date_str)
    
    # 解析任务
    tasks = parse_tasks(content)
    
    # 如果没有任务，提供提示
    if not any([tasks["completed"], tasks["in_progress"], tasks["tasks"]]):
        print(f"ℹ️ 未找到 {date_str} 的记忆内容")
        print(f"   记忆文件路径: {get_memory_file(date_str)}")
        print("   正在生成空日报模板...")
    
    # 生成报告
    if args.format == 'text':
        report = generate_text_report(date_str, tasks, stats=not args.no_stats)
        print(report)
        
        if args.output:
            output_path = Path(args.output).expanduser()
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report)
            print(f"\n✅ 已保存到: {output_path}")
    
    elif args.format == 'md':
        report = generate_markdown_report(date_str, tasks, stats=not args.no_stats)
        print(report)
        
        if args.output:
            output_path = Path(args.output).expanduser()
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report)
            print(f"\n✅ 已保存到: {output_path}")
    
    elif args.format == 'docx':
        output_path = Path(args.output) if args.output else OUTPUT_DIR / f"日报_{date_str}.docx"
        output_path = Path(output_path).expanduser()
        
        result = generate_docx_report(date_str, tasks, output_path)
        if result:
            print(f"✅ Word日报已生成: {result}")
        else:
            sys.exit(1)


if __name__ == "__main__":
    main()

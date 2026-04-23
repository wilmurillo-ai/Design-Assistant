#!/usr/bin/env python3
"""从抖音提取的文案生成 Word 文件"""
import argparse
import sys

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages", file=sys.stderr)
    sys.exit(1)


def create_docx(title: str, author: str, content: str, output: str):
    doc = Document()
    doc.add_heading(title, level=1)
    if author:
        doc.add_paragraph(f"作者：{author}")
        doc.add_paragraph("")

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # 识别结构化标题
        is_heading = False
        # "第X组：" 模式
        if line.startswith("第") and "组" in line[:6] and "：" in line:
            is_heading = True
        # 常见独立标题关键词
        standalone_titles = ["复制这句：", "怎么用", "完整示例", "最后说一句"]
        for t in standalone_titles:
            if line == t:
                is_heading = True
                break
        # 以"核心词"结尾的标题
        if line.endswith("核心词") or line.endswith("核心词："):
            is_heading = True

        if is_heading:
            doc.add_heading(line, level=2)
        elif line.startswith("8K超高清"):
            # 提示词行用引用块样式
            styles = [s.name for s in doc.styles]
            style = "Quote" if "Quote" in styles else None
            p = doc.add_paragraph(line, style=style)
        else:
            doc.add_paragraph(line)

    doc.save(output)
    print(f"OK: {output}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="生成抖音文案 Word 文件")
    parser.add_argument("--title", required=True, help="文档标题")
    parser.add_argument("--author", default="", help="作者昵称")
    parser.add_argument("--content", required=True, help="完整文案内容")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    args = parser.parse_args()
    create_docx(args.title, args.author, args.content, args.output)

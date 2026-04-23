#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word 转 Markdown 转换器
支持图片提取和表格转换
"""

import sys
import io
import base64
import re
from pathlib import Path
from typing import Optional, List, Tuple

# 导入依赖
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from PIL import Image
except ImportError:
    pass

# 导入工具函数
sys.path.insert(0, str(Path(__file__).parent.parent))
from utils import (
    print_success, print_error, print_info, 
    create_image_folder, ensure_output_path
)


def extract_images_from_docx(doc: Document, image_folder: Path, 
                              md_filename: str) -> dict:
    """
    从 Word 文档中提取图片
    
    Args:
        doc: Document 对象
        image_folder: 图片存储文件夹
        md_filename: Markdown 文件名（用于图片命名）
    
    Returns:
        图片映射 {image_id: relative_path}
    """
    image_map = {}
    image_counter = 1
    
    # 从文档的 relations 中提取图片
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image_ext = rel.target_part.content_type.split('/')[-1]
                
                # 保存图片
                image_filename = f"{md_filename}_{image_counter:03d}.{image_ext}"
                image_path = image_folder / image_filename
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                # 记录相对路径
                image_map[rel.rId] = f"{image_folder.name}/{image_filename}"
                image_counter += 1
            
            except Exception as e:
                print_error(f"Failed to extract image: {e}")
    
    return image_map


def convert_table_to_markdown(table) -> str:
    """
    将 Word 表格转换为 Markdown 格式
    
    Args:
        table: docx.table.Table 对象
    
    Returns:
        Markdown 表格字符串
    """
    rows = []
    
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # 获取单元格文本
            text = cell.text.strip().replace('\n', ' ')
            cells.append(text)
        rows.append(cells)
    
    if not rows:
        return ""
    
    # 生成 Markdown 表格
    md_lines = []
    
    # 表头
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    
    # 分隔线
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    
    # 数据行
    for row in rows[1:]:
        md_lines.append("| " + " | ".join(row) + " |")
    
    return "\n".join(md_lines)


def convert_paragraph_to_markdown(para, image_map: dict) -> str:
    """
    将 Word 段落转换为 Markdown
    
    Args:
        para: docx.text.Paragraph 对象
        image_map: 图片映射
    
    Returns:
        Markdown 文本
    """
    text = para.text.strip()
    
    if not text:
        return ""
    
    # 检测标题级别
    style_name = para.style.name.lower()
    if 'heading 1' in style_name or 'title' in style_name:
        return f"# {text}\n"
    elif 'heading 2' in style_name:
        return f"## {text}\n"
    elif 'heading 3' in style_name:
        return f"### {text}\n"
    elif 'heading 4' in style_name:
        return f"#### {text}\n"
    
    # 检测列表
    if para.style.name.startswith('List'):
        # 尝试获取编号
        num = para.style.name.split()[-1] if para.style.name.split()[-1].isdigit() else "1"
        return f"{num}. {text}\n"
    
    # 普通段落
    # 处理格式（粗体、斜体）
    formatted_text = ""
    for run in para.runs:
        run_text = run.text
        if run.bold and run.italic:
            formatted_text += f"***{run_text}***"
        elif run.bold:
            formatted_text += f"**{run_text}**"
        elif run.italic:
            formatted_text += f"*{run_text}*"
        else:
            formatted_text += run_text
    
    return f"{formatted_text}\n"


def convert_word_to_markdown(docx_path: Path, output_path: Optional[Path] = None,
                              extract_images: bool = True, 
                              verbose: bool = True) -> bool:
    """
    将 Word 文档转换为 Markdown
    
    Args:
        docx_path: Word 文档路径
        output_path: 输出 Markdown 路径（可选）
        extract_images: 是否提取图片
        verbose: 是否显示详细信息
    
    Returns:
        是否转换成功
    """
    if not docx_path.exists():
        if verbose:
            print_error(f"File not found: {docx_path}")
        return False
    
    try:
        # 读取 Word 文档
        if verbose:
            print_info(f"Reading: {docx_path.name}")
        
        doc = Document(str(docx_path))
        
        # 确定输出路径
        if output_path is None:
            output_path = docx_path.with_suffix('.md')
        
        # 创建图片文件夹
        image_map = {}
        if extract_images:
            image_folder = create_image_folder(output_path)
            image_map = extract_images_from_docx(doc, image_folder, output_path.stem)
            
            if verbose and image_map:
                print_info(f"Extracted {len(image_map)} images")
        
        # 转换内容
        md_content = []
        
        for element in doc.element.body:
            # 处理段落
            if element.tag.endswith('p'):
                # 查找对应的 paragraph 对象
                for para in doc.paragraphs:
                    if para._element == element:
                        md_line = convert_paragraph_to_markdown(para, image_map)
                        if md_line:
                            md_content.append(md_line)
                        break
            
            # 处理表格
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        md_table = convert_table_to_markdown(table)
                        if md_table:
                            md_content.append(md_table + "\n")
                        break
        
        # 写入 Markdown 文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(md_content))
        
        if verbose:
            print_success(f"Created: {output_path}")
        
        return True
    
    except Exception as e:
        if verbose:
            print_error(f"Conversion failed: {e}")
        return False


def batch_convert_word_to_markdown(docx_files: List[Path], 
                                     extract_images: bool = True,
                                     verbose: bool = True) -> dict:
    """
    批量转换 Word 文档为 Markdown
    
    Args:
        docx_files: Word 文档路径列表
        extract_images: 是否提取图片
        verbose: 是否显示详细信息
    
    Returns:
        转换结果统计
    """
    results = {'success': 0, 'failed': 0, 'images': 0}
    
    for i, docx_file in enumerate(docx_files, 1):
        if verbose:
            print(f"\n[{i}/{len(docx_files)}] {docx_file.name}")
        
        success = convert_word_to_markdown(
            docx_file, 
            extract_images=extract_images,
            verbose=verbose
        )
        
        if success:
            results['success'] += 1
        else:
            results['failed'] += 1
    
    if verbose:
        print(f"\n{'='*50}")
        print(f"Conversion completed: {results['success']} success, {results['failed']} failed")
    
    return results

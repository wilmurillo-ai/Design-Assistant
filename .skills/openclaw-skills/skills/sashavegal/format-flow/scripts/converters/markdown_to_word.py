#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 转换器
使用 pypandoc 进行高质量转换
"""

import sys
from pathlib import Path
from typing import Optional, List

# 导入工具函数
sys.path.insert(0, str(Path(__file__).parent.parent))
from utils import (
    print_success, print_error, print_info,
    check_markdown_to_word_support
)


def convert_markdown_to_word(md_path: Path, output_path: Optional[Path] = None,
                              verbose: bool = True) -> bool:
    """
    将 Markdown 文档转换为 Word
    
    Args:
        md_path: Markdown 文档路径
        output_path: 输出 Word 路径（可选）
        verbose: 是否显示详细信息
    
    Returns:
        是否转换成功
    """
    if not md_path.exists():
        if verbose:
            print_error(f"File not found: {md_path}")
        return False
    
    # 确定输出路径
    if output_path is None:
        output_path = md_path.with_suffix('.docx')
    
    # 方法1: 使用 pypandoc (推荐，质量高)
    if check_markdown_to_word_support():
        try:
            if verbose:
                print_info(f"Converting {md_path.name} to Word using pypandoc...")
            
            import pypandoc
            
            # 使用 pypandoc 转换
            output = pypandoc.convert_file(
                str(md_path),
                'docx',
                outputfile=str(output_path),
                extra_args=[
                    '--reference-doc=',  # 可以可以指定 Word 模板
                ]
            )
            
            if verbose:
                print_success(f"Created: {output_path}")
            return True
        
        except Exception as e:
            if verbose:
                print_error(f"pypandoc conversion failed: {e}")
    
    # 方法2: 使用 python-docx 手动转换（基础功能）
    try:
        if verbose:
            print_info(f"Trying basic conversion using python-docx...")
        
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 读取 Markdown 内容
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 简单的 Markdown 解析
        for line in lines:
            line = line.rstrip()
            
            # 一级标题
            if line.startswith('# '):
                para = doc.add_heading(line[2:], level=1)
            
            # 二级标题
            elif line.startswith('## '):
                para = doc.add_heading(line[3:], level=2)
            
            # 三级标题
            elif line.startswith('### '):
                para = doc.add_heading(line[4:], level=3)
            
            # 四级标题
            elif line.startswith('#### '):
                para = doc.add_heading(line[5:], level=4)
            
            # 空行
            elif not line:
                pass
            
            # 普通段落
            else:
                # 简单处理粗体和斜体
                import re
                text = line
                
                # 创建段落
                para = doc.add_paragraph()
                
                # 处理粗体 **text**
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)
        
        # 保存文档
        doc.save(str(output_path))
        
        if verbose:
            print_success(f"Created: {output_path}")
            print_info("(Basic conversion - limited formatting support)")
        
        return True
    
    except Exception as e:
        if verbose:
            print_error(f"Conversion failed: {e}")
    
    if verbose:
        print_error("All conversion methods failed")
        print_info("Please install pypandoc for better conversion quality:")
        print_info("  pip install pypandoc")
    
    return False


def batch_convert_markdown_to_word(md_files: List[Path], 
                                     verbose: bool = True) -> dict:
    """
    批量转换 Markdown 文档为 Word
    
    Args:
        md_files: Markdown 文档路径列表
        verbose: 是否显示详细信息
    
    Returns:
        转换结果统计
    """
    results = {'success': 0, 'failed': 0}
    
    for i, md_file in enumerate(md_files, 1):
        if verbose:
            print(f"\n[{i}/{len(md_files)}] {md_file.name}")
        
        success = convert_markdown_to_word(md_file, verbose=verbose)
        
        if success:
            results['success'] += 1
        else:
            results['failed'] += 1
    
    if verbose:
        print(f"\n{'='*50}")
        print(f"Conversion completed: {results['success']} success, {results['failed']} failed")
    
    return results

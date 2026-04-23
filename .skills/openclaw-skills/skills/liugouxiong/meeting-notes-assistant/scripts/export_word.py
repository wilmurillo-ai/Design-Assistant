#!/usr/bin/env python3
"""导出 Word 文档，支持内置模板、模板管理器模板与自定义模板占位。"""

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误：未安装 python-docx。请运行: pip install python-docx")
    sys.exit(1)

try:
    from templates import DEFAULT_TEMPLATES, load_template
except ImportError:
    DEFAULT_TEMPLATES = {}
    load_template = None


BUILTIN_TEMPLATE_ALIASES = {
    "simple": "simple",
    "简洁版": "simple",
    "professional": "professional",
    "专业版": "professional",
    "custom": "custom",
}

SECTION_TITLES = {
    "title": "标题",
    "meta": "会议信息",
    "attendees": "参会人",
    "summary": "会议摘要",
    "topics": "议题与结论",
    "todos": "待办事项",
    "keywords": "关键词",
    "action_items": "Action Items",
    "background": "背景",
    "goals": "目标",
    "team": "团队",
    "timeline": "时间线",
    "risks": "风险与问题",
    "last_week": "上周进展",
    "this_week": "本周计划",
    "blockers": "阻塞项",
    "achievements": "关键成果",
    "issues": "问题复盘",
    "next_month": "下月计划",
    "candidate_info": "候选人信息",
    "evaluation": "综合评价",
    "strengths": "优势",
    "weaknesses": "待改进点",
    "decision": "结论与决定",
    "options": "备选方案",
}

SECTION_ALIASES = {
    "background": ["background", "context"],
    "goals": ["goals", "objective", "objectives"],
    "team": ["team", "members"],
    "timeline": ["timeline", "schedule", "milestones"],
    "risks": ["risks", "risk_items"],
    "last_week": ["last_week", "previous_week"],
    "this_week": ["this_week", "current_week"],
    "blockers": ["blockers", "risks", "issues"],
    "achievements": ["achievements", "highlights"],
    "issues": ["issues", "problems"],
    "next_month": ["next_month", "next_steps"],
    "candidate_info": ["candidate_info", "candidate", "profile"],
    "evaluation": ["evaluation", "assessment", "summary"],
    "strengths": ["strengths", "pros"],
    "weaknesses": ["weaknesses", "cons", "gaps"],
    "decision": ["decision", "decisions", "conclusion", "conclusions"],
    "options": ["options", "alternatives"],
}

FIELD_LABELS = {
    "title": "标题",
    "summary": "摘要",
    "conclusion": "结论",
    "owner": "负责人",
    "assignee": "负责人",
    "due": "截止",
    "deadline": "截止",
    "time": "时间",
    "theme": "主题",
    "attendees": "参会人",
    "content": "内容",
    "status": "状态",
}


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def normalize_title(meeting_info: Dict[str, Any]) -> str:
    return normalize_text(meeting_info.get("title") or meeting_info.get("theme") or "会议纪要")


def normalize_item(item: object) -> Dict[str, str]:
    if isinstance(item, dict):
        return {
            "task": normalize_text(item.get("task") or item.get("content")),
            "owner": normalize_text(item.get("owner") or item.get("assignee")),
            "due": normalize_text(item.get("due") or item.get("deadline")),
        }
    return {"task": normalize_text(item), "owner": "", "due": ""}


def normalize_todos(notes: Dict[str, Any]) -> List[Dict[str, str]]:
    return [normalize_item(todo) for todo in notes.get("todos", [])]


def normalize_action_items(notes: Dict[str, Any]) -> List[Dict[str, str]]:
    return [normalize_item(item) for item in notes.get("action_items", [])]


def add_document_title(doc: Document, title_text: str) -> None:
    title = doc.add_heading(title_text or "会议纪要", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_meta_lines(doc: Document, meeting_info: Dict[str, Any], title_text: str, with_icons: bool = False) -> None:
    icon_prefix = {"时间": "📅 ", "主题": "📌 ", "参会人": "👥 "} if with_icons else {}
    lines = [
        ("时间", normalize_text(meeting_info.get("time")) or "未知"),
        ("主题", title_text or "会议纪要"),
        ("参会人", normalize_text(meeting_info.get("attendees")) or "未知"),
    ]
    for label, value in lines:
        doc.add_paragraph(f"{icon_prefix.get(label, '')}{label}：{value}")
    doc.add_paragraph()


def render_topics_section(doc: Document, notes: Dict[str, Any], numbered: bool = False, highlight_conclusion: bool = False) -> None:
    doc.add_heading("议题与结论", level=1)
    topics = notes.get("topics", [])
    if not topics:
        doc.add_paragraph("暂无内容")
        return

    for index, topic in enumerate(topics, 1):
        if isinstance(topic, dict):
            topic_title = normalize_text(topic.get("title")) or f"议题 {index}"
            prefix = f"{index}. " if numbered else "• "
            doc.add_paragraph(f"{prefix}{topic_title}")

            summary = normalize_text(topic.get("summary"))
            if summary:
                doc.add_paragraph(f"  摘要：{summary}")

            conclusion = normalize_text(topic.get("conclusion"))
            if conclusion:
                conclusion_prefix = "   ✓ " if highlight_conclusion else "  结论："
                paragraph = doc.add_paragraph(f"{conclusion_prefix}{conclusion}")
                if highlight_conclusion and paragraph.runs:
                    paragraph.runs[0].bold = True

            owner = normalize_text(topic.get("owner"))
            if owner:
                doc.add_paragraph(f"  负责人：{owner}")
        else:
            prefix = f"{index}. " if numbered else "• "
            doc.add_paragraph(f"{prefix}{normalize_text(topic)}")


def render_todos_section(doc: Document, notes: Dict[str, Any], checkbox: bool = False) -> None:
    doc.add_heading("待办事项", level=1)
    todos = normalize_todos(notes)
    if not todos:
        doc.add_paragraph("暂无内容")
        return

    # 检查是否有完整的待办信息（content, assignee, deadline, priority）
    has_full_info = all(
        todo.get("content") and 
        (todo.get("assignee") or todo.get("owner")) and 
        (todo.get("deadline") or todo.get("due")) and 
        todo.get("priority")
        for todo in todos
    )

    if has_full_info:
        # 使用表格展示
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "待办事项"
        header_cells[1].text = "责任人"
        header_cells[2].text = "截止时间"
        header_cells[3].text = "优先级"

        # 表头样式
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True

        for todo in todos:
            row_cells = table.add_row().cells
            row_cells[0].text = normalize_text(todo.get("content", todo.get("task", "")))
            row_cells[1].text = normalize_text(todo.get("assignee", todo.get("owner", "待分配")))
            row_cells[2].text = normalize_text(todo.get("deadline", todo.get("due", "")))
            row_cells[3].text = normalize_text(todo.get("priority", ""))

            # 优先级颜色标记
            priority = todo.get("priority", "")
            if priority == "高":
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = None  # 默认黑色
            elif priority == "中":
                pass
            elif priority == "低":
                pass
    else:
        # 回退到列表展示
        for todo in todos:
            suffix = []
            # 支持多种字段名
            owner = todo.get("assignee") or todo.get("owner", "")
            deadline = todo.get("deadline") or todo.get("due", "")
            priority = todo.get("priority", "")

            if owner:
                suffix.append(f"负责人：{owner}")
            if deadline:
                suffix.append(f"截止：{deadline}")
            if priority:
                suffix.append(f"优先级：{priority}")

            extra = f"（{'；'.join(suffix)})" if suffix else ""
            prefix = "☐ " if checkbox else "• "
            doc.add_paragraph(f"{prefix}{todo.get('content', todo.get('task', ''))}{extra}")


def render_keywords_section(doc: Document, notes: Dict[str, Any], hashtag: bool = False) -> None:
    doc.add_heading("关键词", level=1)
    keywords = notes.get("keywords", [])
    if not keywords:
        doc.add_paragraph("暂无内容")
        return

    if hashtag:
        paragraph = doc.add_paragraph()
        for keyword in keywords:
            paragraph.add_run(f"#{normalize_text(keyword)} ")
    else:
        doc.add_paragraph("、".join(normalize_text(keyword) for keyword in keywords if normalize_text(keyword)))


def render_action_items_section(doc: Document, notes: Dict[str, Any], numbered: bool = False) -> None:
    doc.add_heading("Action Items", level=1)
    action_items = normalize_action_items(notes)
    if not action_items:
        doc.add_paragraph("暂无内容")
        return

    columns = 4 if numbered else 3
    table = doc.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    start = 0
    if numbered:
        header_cells[0].text = "#"
        start = 1
    header_cells[start].text = "任务"
    header_cells[start + 1].text = "负责人"
    header_cells[start + 2].text = "截止时间"

    for index, item in enumerate(action_items, 1):
        row_cells = table.add_row().cells
        start = 0
        if numbered:
            row_cells[0].text = str(index)
            start = 1
        row_cells[start].text = item["task"]
        row_cells[start + 1].text = item["owner"]
        row_cells[start + 2].text = item["due"]


def render_summary_section(doc: Document, notes: Dict[str, Any]) -> None:
    doc.add_heading("会议摘要", level=1)
    summary = normalize_text(notes.get("summary"))
    doc.add_paragraph(summary or "暂无内容")


def prettify_field_name(field_name: str) -> str:
    return FIELD_LABELS.get(field_name, field_name.replace("_", " ").title())


def get_section_value(notes: Dict[str, Any], section: str) -> Any:
    meeting_info = notes.get("meeting_info", {})

    if section in notes and notes[section]:
        return notes[section]
    if section in meeting_info and meeting_info[section]:
        return meeting_info[section]

    for container_key in ["sections", "custom_sections", "extra_sections"]:
        container = notes.get(container_key)
        if isinstance(container, dict) and container.get(section):
            return container[section]

    for alias in SECTION_ALIASES.get(section, []):
        if notes.get(alias):
            return notes[alias]
        if meeting_info.get(alias):
            return meeting_info[alias]
        for container_key in ["sections", "custom_sections", "extra_sections"]:
            container = notes.get(container_key)
            if isinstance(container, dict) and container.get(alias):
                return container[alias]

    if section == "background":
        return notes.get("summary")
    if section == "team":
        return meeting_info.get("attendees")
    if section == "decision":
        conclusions = [normalize_text(topic.get("conclusion")) for topic in notes.get("topics", []) if isinstance(topic, dict) and normalize_text(topic.get("conclusion"))]
        return conclusions
    if section == "options":
        titles = [normalize_text(topic.get("title")) for topic in notes.get("topics", []) if isinstance(topic, dict) and normalize_text(topic.get("title"))]
        return titles
    if section == "evaluation":
        return notes.get("summary")

    return None


def render_generic_value(doc: Document, value: Any, indent: str = "") -> None:
    if value is None or value == "":
        doc.add_paragraph(f"{indent}暂无内容")
        return

    if isinstance(value, dict):
        if not value:
            doc.add_paragraph(f"{indent}暂无内容")
            return
        for key, item in value.items():
            if isinstance(item, (dict, list)):
                doc.add_paragraph(f"{indent}{prettify_field_name(str(key))}：")
                render_generic_value(doc, item, indent=f"{indent}  ")
            else:
                doc.add_paragraph(f"{indent}{prettify_field_name(str(key))}：{normalize_text(item)}")
        return

    if isinstance(value, list):
        if not value:
            doc.add_paragraph(f"{indent}暂无内容")
            return
        for item in value:
            if isinstance(item, dict):
                summary_parts = []
                for field in ["title", "task", "content", "summary", "conclusion", "owner", "assignee", "due", "deadline"]:
                    field_value = normalize_text(item.get(field))
                    if field_value:
                        if field in ["title", "task", "content"] and not summary_parts:
                            summary_parts.append(field_value)
                        else:
                            summary_parts.append(f"{prettify_field_name(field)}：{field_value}")
                doc.add_paragraph(f"{indent}• {'；'.join(summary_parts) if summary_parts else str(item)}")
            else:
                doc.add_paragraph(f"{indent}• {normalize_text(item)}")
        return

    doc.add_paragraph(f"{indent}{normalize_text(value)}")


def render_generic_section(doc: Document, notes: Dict[str, Any], section: str) -> None:
    title = SECTION_TITLES.get(section, section.replace("_", " ").title())
    doc.add_heading(title, level=1)
    render_generic_value(doc, get_section_value(notes, section))


def format_topics_for_placeholder(notes: Dict[str, Any]) -> str:
    topics = notes.get("topics", [])
    if not topics:
        return "暂无内容"

    lines: List[str] = []
    for index, topic in enumerate(topics, 1):
        if isinstance(topic, dict):
            title = normalize_text(topic.get("title")) or f"议题 {index}"
            lines.append(f"{index}. {title}")

            summary = normalize_text(topic.get("summary"))
            if summary:
                lines.append(f"   摘要：{summary}")

            conclusion = normalize_text(topic.get("conclusion"))
            if conclusion:
                lines.append(f"   结论：{conclusion}")

            owner = normalize_text(topic.get("owner"))
            if owner:
                lines.append(f"   负责人：{owner}")
        else:
            lines.append(f"{index}. {normalize_text(topic)}")
    return "\n".join(lines)


def format_todos_for_placeholder(notes: Dict[str, Any]) -> str:
    todos = normalize_todos(notes)
    if not todos:
        return "暂无内容"

    lines = []
    for todo in todos:
        suffix = []
        if todo["owner"]:
            suffix.append(f"负责人：{todo['owner']}")
        if todo["due"]:
            suffix.append(f"截止：{todo['due']}")
        extra = f"（{'；'.join(suffix)}）" if suffix else ""
        lines.append(f"• {todo['task']}{extra}")
    return "\n".join(lines)


def format_keywords_for_placeholder(notes: Dict[str, Any]) -> str:
    keywords = [normalize_text(keyword) for keyword in notes.get("keywords", []) if normalize_text(keyword)]
    return "、".join(keywords) if keywords else "暂无内容"


def format_action_items_for_placeholder(notes: Dict[str, Any]) -> str:
    action_items = normalize_action_items(notes)
    if not action_items:
        return "暂无内容"

    lines = []
    for item in action_items:
        suffix = []
        if item["owner"]:
            suffix.append(f"负责人：{item['owner']}")
        if item["due"]:
            suffix.append(f"截止：{item['due']}")
        extra = f"（{'；'.join(suffix)}）" if suffix else ""
        lines.append(f"• {item['task']}{extra}")
    return "\n".join(lines)


def build_custom_placeholder_values(notes: Dict[str, Any]) -> Dict[str, str]:
    meeting_info = notes.get("meeting_info", {})
    title_text = normalize_title(meeting_info)
    attendees = normalize_text(meeting_info.get("attendees")) or "未知"

    return {
        "{{meeting_time}}": normalize_text(meeting_info.get("time")) or "未知",
        "{{meeting_theme}}": title_text,
        "{{meeting_title}}": title_text,
        "{{title}}": title_text,
        "{{attendees}}": attendees,
        "{{summary}}": normalize_text(notes.get("summary")) or "暂无内容",
        "{{topics}}": format_topics_for_placeholder(notes),
        "{{todos}}": format_todos_for_placeholder(notes),
        "{{keywords}}": format_keywords_for_placeholder(notes),
        "{{action_items_text}}": format_action_items_for_placeholder(notes),
        "{{raw_text}}": normalize_text(notes.get("raw_text")) or "暂无内容",
    }


def add_custom_action_items_table(container: Any, notes: Dict[str, Any]):
    action_items = normalize_action_items(notes)
    if not action_items:
        return None

    table = container.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "任务"
    header_cells[1].text = "负责人"
    header_cells[2].text = "截止时间"

    for item in action_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item["task"]
        row_cells[1].text = item["owner"]
        row_cells[2].text = item["due"]
    return table


def replace_text_placeholders_in_paragraph(paragraph: Any, placeholder_values: Dict[str, str]) -> None:
    original_text = paragraph.text
    updated_text = original_text
    for placeholder, value in placeholder_values.items():
        updated_text = updated_text.replace(placeholder, value)
    if updated_text != original_text:
        paragraph.text = updated_text


def handle_action_items_placeholder_in_paragraph(paragraph: Any, doc: Document, notes: Dict[str, Any]) -> None:
    placeholder = "{{action_items}}"
    if placeholder not in paragraph.text:
        return

    action_items = normalize_action_items(notes)
    paragraph.text = paragraph.text.replace(placeholder, "" if action_items else "暂无内容")
    if not action_items:
        return

    table = add_custom_action_items_table(doc, notes)
    if table is not None:
        paragraph._p.addnext(table._tbl)


def process_custom_template_cell(cell: Any, notes: Dict[str, Any], placeholder_values: Dict[str, str]) -> None:
    action_items = normalize_action_items(notes)
    action_placeholder_handled = False

    for paragraph in cell.paragraphs:
        replace_text_placeholders_in_paragraph(paragraph, placeholder_values)

        if action_placeholder_handled or "{{action_items}}" not in paragraph.text:
            continue

        paragraph.text = paragraph.text.replace("{{action_items}}", "" if action_items else "暂无内容")
        if action_items:
            add_custom_action_items_table(cell, notes)
        action_placeholder_handled = True

    for nested_table in cell.tables:
        process_custom_template_table(nested_table, notes, placeholder_values)


def process_custom_template_table(table: Any, notes: Dict[str, Any], placeholder_values: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            process_custom_template_cell(cell, notes, placeholder_values)


def create_simple_doc(notes: Dict[str, Any], output_path: str) -> None:
    doc = Document()
    meeting_info = notes.get("meeting_info", {})
    title_text = normalize_title(meeting_info)

    add_document_title(doc, title_text)
    add_meta_lines(doc, meeting_info, title_text)
    render_topics_section(doc, notes, numbered=False, highlight_conclusion=False)
    render_todos_section(doc, notes, checkbox=False)

    if notes.get("keywords"):
        render_keywords_section(doc, notes, hashtag=False)
    if notes.get("action_items"):
        render_action_items_section(doc, notes, numbered=False)

    doc.save(output_path)
    print(f"简洁版 Word 已保存: {output_path}")


def create_professional_doc(notes: Dict[str, Any], output_path: str) -> None:
    doc = Document()
    meeting_info = notes.get("meeting_info", {})
    title_text = normalize_title(meeting_info)

    # 1. 标题区域
    add_document_title(doc, title_text)

    # 2. 会议信息表格
    if meeting_info:
        doc.add_heading("会议信息", level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        # 定义要显示的字段
        fields = [
            ("meeting_title", "会议标题"),
            ("date", "会议日期"),
            ("location", "会议地点"),
            ("attendees", "参会人"),
            ("duration", "会议时长"),
            ("recorder", "记录人")
        ]

        for field, label in fields:
            value = meeting_info.get(field, "")
            if value:
                if isinstance(value, list):
                    value = "、".join(str(v) for v in value)
                if str(value).strip():
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                    # 标签加粗
                    row_cells[0].paragraphs[0].runs[0].bold = True

    # 3. 会议摘要
    if notes.get("summary"):
        doc.add_heading("会议摘要", level=2)
        summary = doc.add_paragraph()
        summary.add_run(notes["summary"]).bold = True

    doc.add_paragraph()  # 分隔

    # 4. 议题区域
    render_topics_section(doc, notes, numbered=True, highlight_conclusion=True)

    # 5. 待办事项（使用表格）
    render_todos_section(doc, notes, checkbox=False)

    # 6. 关键词
    render_keywords_section(doc, notes, hashtag=True)

    # 7. 行动项
    if notes.get("action_items"):
        render_action_items_section(doc, notes, numbered=True)

    # 8. 页眉页脚
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = f"Meeting Notes Assistant | 生成时间: {normalize_text(meeting_info.get('date', ''))}"

    doc.save(output_path)
    print(f"专业版 Word 已保存: {output_path}")


def resolve_managed_template(template_name: str) -> Optional[Dict[str, Any]]:
    if load_template is not None:
        managed_template = load_template(template_name)
        if managed_template:
            return managed_template
    if template_name in DEFAULT_TEMPLATES:
        return DEFAULT_TEMPLATES[template_name]
    return None


def create_managed_template_doc(notes: Dict[str, Any], template_name: str, template_config: Dict[str, Any], output_path: str) -> None:
    doc = Document()
    sections = template_config.get("sections") or ["title", "meta", "topics", "todos", "keywords", "action_items"]

    for section in sections:
        if section == "title":
            add_document_title(doc, normalize_title(notes.get("meeting_info", {})))
        elif section == "meta":
            add_meta_lines(doc, notes.get("meeting_info", {}), normalize_title(notes.get("meeting_info", {})))
        elif section == "attendees":
            render_generic_section(doc, notes, section)
        elif section == "summary":
            render_summary_section(doc, notes)
        elif section == "topics":
            render_topics_section(doc, notes, numbered=True, highlight_conclusion=False)
        elif section == "todos":
            render_todos_section(doc, notes, checkbox=False)
        elif section == "keywords":
            render_keywords_section(doc, notes, hashtag=False)
        elif section == "action_items":
            render_action_items_section(doc, notes, numbered=False)
        else:
            render_generic_section(doc, notes, section)

    footer = doc.sections[0].footer
    footer.paragraphs[0].text = f"Template: {template_name} | Generated by Meeting Notes Assistant"

    doc.save(output_path)
    print(f"模板“{template_name}”Word 已保存: {output_path}")


def create_custom_doc(notes: Dict[str, Any], template_path: str, output_path: str) -> None:
    template_file = Path(template_path)
    if not template_file.exists():
        print(f"错误：模板文件不存在: {template_path}")
        sys.exit(1)

    doc = Document(str(template_file))
    placeholder_values = build_custom_placeholder_values(notes)
    header_footer_values = dict(placeholder_values)
    header_footer_values["{{action_items}}"] = format_action_items_for_placeholder(notes)

    for paragraph in doc.paragraphs:
        replace_text_placeholders_in_paragraph(paragraph, placeholder_values)
        handle_action_items_placeholder_in_paragraph(paragraph, doc, notes)

    for table in doc.tables:
        process_custom_template_table(table, notes, placeholder_values)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_placeholders_in_paragraph(paragraph, header_footer_values)
        for table in section.header.tables:
            process_custom_template_table(table, notes, header_footer_values)

        for paragraph in section.footer.paragraphs:
            replace_text_placeholders_in_paragraph(paragraph, header_footer_values)
        for table in section.footer.tables:
            process_custom_template_table(table, notes, header_footer_values)

    doc.save(output_path)
    print(f"自定义模板 Word 已保存: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="导出 Word 文档")
    parser.add_argument("notes_json", help="会议纪要 JSON 文件")
    parser.add_argument("--template", default="simple", help="模板类型：simple / professional / custom / templates.py 中的模板名")
    parser.add_argument("--template-path", help="自定义模板路径（当 template=custom 时）")
    parser.add_argument("--output", "-o", required=True, help="输出文件路径")
    args = parser.parse_args()

    with open(args.notes_json, "r", encoding="utf-8") as f:
        notes = json.load(f)

    template_name = normalize_text(args.template)
    builtin_template = BUILTIN_TEMPLATE_ALIASES.get(template_name)

    if builtin_template == "simple":
        create_simple_doc(notes, args.output)
        return

    if builtin_template == "professional":
        create_professional_doc(notes, args.output)
        return

    if builtin_template == "custom":
        if not args.template_path:
            print("错误：自定义模板需要 --template-path 参数")
            sys.exit(1)
        create_custom_doc(notes, args.template_path, args.output)
        return

    template_config = resolve_managed_template(template_name)
    if template_config:
        create_managed_template_doc(notes, template_name, template_config, args.output)
        return

    print(f"错误：未知模板 {template_name}。请使用 simple、professional、custom，或 templates.py 中已创建的模板名。")
    sys.exit(1)


if __name__ == "__main__":
    main()

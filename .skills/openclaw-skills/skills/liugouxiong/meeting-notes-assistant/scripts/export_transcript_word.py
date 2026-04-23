#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
转写文本 Word 导出工具
将转写文本原样输出为专业排版的 Word 文档（保留原始格式，不做断句/分段处理）
"""

import argparse
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：未安装 python-docx。请运行: pip install python-docx")
    sys.exit(1)


# ─── 文本处理（保留原格式）─────────────────────────────────────────────────────

def process_transcript(raw_text: str) -> list:
    """原始转写文本 → 段落列表（按换行分段，无换行则整体输出）"""
    text = raw_text.strip()
    if not text:
        return []

    # 按原始换行分段；如果没有换行则作为一个段落
    lines = [line.strip() for line in text.splitlines()]
    paragraphs = [line for line in lines if line]

    # 如果只有一行（Whisper 常见的连续输出），直接返回整体
    if not paragraphs:
        return [text]

    return paragraphs


# ─── Word 文档生成 ─────────────────────────────────────────────────────────────

def set_cell_background(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_horizontal_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_transcript_word(
    raw_text: str,
    output_path: str,
    source_file: str = "",
    title: str = "",
    asr_provider: str = "Whisper Small"
) -> None:
    """生成转写文本 Word 文档（带断句分段排版）"""

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    # ── 标题 ──
    doc_title = doc.add_heading(title or "会议录音转写文本", level=0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in doc_title.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x8E)

    # ── 元信息表格 ──
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.style = 'Table Grid'
    meta_cells = meta_table.rows[0].cells
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    source_name = Path(source_file).name if source_file else "未知"
    meta_cells[0].text = f"来源：{source_name}"
    meta_cells[1].text = f"转写引擎：{asr_provider}"
    meta_cells[2].text = f"生成时间：{now_str}"
    for cell in meta_cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        set_cell_background(cell, 'F5F5F5')

    doc.add_paragraph()

    # ── 分隔线 + 说明 ──
    add_horizontal_line(doc)
    note_para = doc.add_paragraph()
    note_para.paragraph_format.space_after = Pt(6)
    note_run = note_para.add_run(
        "📝 说明：以下文本由 AI 自动转写，保留原始输出格式。"
        "专业术语可能存在识别偏差，请以实际录音为准。"
    )
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_horizontal_line(doc)
    doc.add_paragraph()

    # ── 正文：自然段 ──
    paragraphs = process_transcript(raw_text)

    if not paragraphs:
        doc.add_paragraph("（转写文本为空或无法解析）")
    else:
        for para_text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)   # 首行缩进
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = Pt(22)         # 行距约 1.5 倍

            body_run = p.add_run(para_text)
            body_run.font.size = Pt(11)
            body_run.font.name = '宋体'

    # ── 底部统计 ──
    doc.add_paragraph()
    add_horizontal_line(doc)
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stats_run = stats_para.add_run(
        f"统计：共 {len(paragraphs)} 段，约 {len(raw_text)} 字符"
    )
    stats_run.font.size = Pt(9)
    stats_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── 页眉 ──
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.clear()
    header_run = header_para.add_run(f"会议录音转写文本  |  {source_name}")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── 页脚（页码）──
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Meeting Notes Assistant  |  第 ")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    # 页码域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg = footer_para.add_run()
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)
    footer_run2 = footer_para.add_run(" 页")
    footer_run2.font.size = Pt(9)
    footer_run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── 保存 ──
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] 转写文本 Word 已保存: {output_path}")
    print(f"     共 {len(paragraphs)} 段，{len(raw_text)} 字符")


# ─── CLI 入口 ──────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="将转写文本导出为专业排版的 Word 文档")
    parser.add_argument("transcript", help="转写文本文件路径（.txt）")
    parser.add_argument("--output", "-o", help="输出 Word 文件路径（默认同名 .docx）")
    parser.add_argument("--title", help="文档标题（默认使用文件名）")
    parser.add_argument("--source", help="音频来源文件名（用于页眉/元信息）")
    parser.add_argument("--asr-provider", default="Whisper Small", help="转写引擎名称")
    args = parser.parse_args()

    transcript_path = Path(args.transcript)
    if not transcript_path.exists():
        print(f"错误：文件不存在 {transcript_path}")
        sys.exit(1)

    raw_text = transcript_path.read_text(encoding="utf-8")
    output_path = args.output or str(transcript_path.with_suffix('.docx'))
    title = args.title or transcript_path.stem

    create_transcript_word(
        raw_text=raw_text,
        output_path=output_path,
        source_file=args.source or str(transcript_path),
        title=title,
        asr_provider=args.asr_provider,
    )


if __name__ == "__main__":
    main()

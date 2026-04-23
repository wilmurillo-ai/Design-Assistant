"""Document generator module for creating docx files."""

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .config import Settings
from .models import SummarizedReport, WorkCategories, CategoryItem


class DocumentGenerator:
    """Generates docx documents from templates."""

    # Category display order
    CATEGORIES = [
        "人才转型",
        "自主开发",
        "科创支撑",
        "AI架构及网运安全自智规划",
        "系统需求规划建设",
        "综合工作",
    ]

    def __init__(self, settings: Settings):
        self.settings = settings
        self.template_path = settings.get_template_path()
        self.output_dir = settings.get_output_dir()

    def _add_title(self, doc: Document, text: str):
        """Add a title paragraph."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.name = "方正小标宋简体"
        run.font.size = Pt(22)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "方正小标宋简体")
        return para

    def _add_section_title(self, doc: Document, text: str):
        """Add a section title paragraph."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "黑体"
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
        return para

    def _add_subsection_title(self, doc: Document, text: str):
        """Add a subsection title paragraph."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "楷体_GB2312"
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体_GB2312")
        return para

    def _add_normal_text(self, doc: Document, text: str):
        """Add a normal text paragraph."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "仿宋_GB2312"
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
        return para

    def _add_category_items(self, doc: Document, items: List[CategoryItem]):
        """Add items for a category."""
        for i, item in enumerate(items, 1):
            text = f"（{i}）{item.content}"
            if item.person:
                if not item.content.endswith(f"--{item.person}"):
                    text = f"（{i}）{item.content}--{item.person}"
            self._add_normal_text(doc, text)

    def _add_work_section(
        self,
        doc: Document,
        title: str,
        work_categories: WorkCategories,
        include_team_header: bool = True
    ):
        """Add a work section with categories."""
        if include_team_header:
            self._add_subsection_title(doc, "科创团队")

        for category in self.CATEGORIES:
            items = getattr(work_categories, category, [])
            if items:
                cat_num = self.CATEGORIES.index(category) + 1
                self._add_normal_text(doc, f"{cat_num}、{category}")
                self._add_category_items(doc, items)

    def generate(
        self,
        report: SummarizedReport,
        output_filename: Optional[str] = None,
        template_path: Optional[Path] = None,
        verbose: bool = True,
    ) -> Path:
        """Generate a docx document from the report."""
        if verbose:
            print("[Generator] Generating document...")

        # Generate output filename if not provided
        if output_filename is None:
            week_str = str(report.week_range).replace(".", "-")
            team_safe = report.team_name.replace("/", "-").replace("\\", "-")
            output_filename = f"周报_{team_safe}_{week_str}.docx"

        output_path = self.output_dir / output_filename

        if verbose:
            print(f"[Generator] Output: {output_path}")

        # Create document
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = '仿宋_GB2312'
        style.font.size = Pt(16)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        # Add title
        title_text = f"监控调度室周工作总结计划（{report.week_range.to_chinese_format()}）"
        self._add_title(doc, title_text)

        # Add empty paragraph
        doc.add_paragraph()

        # Section 1: This week's work summary
        self._add_section_title(doc, "一、本周主要工作总结")
        self._add_work_section(doc, "本周工作", report.this_week)

        # Add empty paragraph
        doc.add_paragraph()

        # Section 2: Next week's work plan
        self._add_section_title(doc, "二、下周主要工作计划")
        self._add_work_section(doc, "下周计划", report.next_week)

        # Save document
        doc.save(str(output_path))

        if verbose:
            print(f"[Generator] Document generated: {output_path}")

        return output_path

    def generate_with_template(
        self,
        report: SummarizedReport,
        output_filename: Optional[str] = None,
        template_path: Optional[Path] = None,
        verbose: bool = True,
    ) -> Path:
        """Generate document using docxtpl template."""
        from docxtpl import DocxTemplate

        template = template_path or self.template_path
        if not template.exists():
            raise FileNotFoundError(f"Template file not found: {template}")

        # Generate output filename if not provided
        if output_filename is None:
            week_str = str(report.week_range).replace(".", "-")
            team_safe = report.team_name.replace("/", "-").replace("\\", "-")
            output_filename = f"周报_{team_safe}_{week_str}.docx"

        output_path = self.output_dir / output_filename

        if verbose:
            print(f"[Generator] Generating document with template: {output_path}")

        # Build context
        context = report.to_template_context()

        # Load template and render
        doc = DocxTemplate(str(template))
        doc.render(context)
        doc.save(str(output_path))

        if verbose:
            print(f"[Generator] Document generated: {output_path}")

        return output_path


def generate_report_document(
    report: SummarizedReport,
    settings: Settings,
    output_filename: Optional[str] = None,
    verbose: bool = True,
) -> Path:
    """Convenience function to generate a report document."""
    generator = DocumentGenerator(settings)
    return generator.generate(report, output_filename, verbose=verbose)

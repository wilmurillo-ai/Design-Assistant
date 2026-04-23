#!/usr/bin/env python3
"""
财报点评报告生成脚本
支持输出格式：Markdown、PDF、Word (.docx)
"""

import json
import os
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path


class ReportGenerator:
    """财报点评报告生成器"""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or "./reports"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_markdown_report(self, report_data: Dict) -> str:
        """
        生成 Markdown 格式报告
        
        Args:
            report_data: 报告数据字典，包含：
                - stock_info: 公司信息
                - financial_data: 财务数据
                - charts: echarts 配置
                - analysis: 分析内容
                - recommendations: 投资建议
        
        Returns:
            Markdown 报告内容
        """
        md = []
        
        # 标题
        stock_info = report_data.get("stock_info", {})
        md.append(f"# {stock_info.get('company_name', '未知公司')} ({stock_info.get('stock_code', '')})：财报点评报告\n")
        md.append(f"**报告生成时间：** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        md.append(f"**数据截止日期：** {report_data.get('data_date', 'N/A')}\n")
        md.append("---\n")
        
        # 核心观点
        if report_data.get("core_view"):
            md.append("## 核心观点\n")
            md.append(f"{report_data['core_view']}\n")
            md.append("---\n")
        
        # 第一章：公司整体情况
        md.append("## 一、公司整体情况\n")
        md.append(f"{report_data.get('company_overview', '')}\n")
        
        # 图表 1-4
        charts = report_data.get("charts", {})
        md.append("### 图表 1：公司年度营业收入及同比增速\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart1_revenue_annual", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        md.append("### 图表 2：公司单季度营业收入及同比增速\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart2_revenue_quarterly", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        md.append("### 图表 3：公司年度归母净利润及同比增速\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart3_profit_annual", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        md.append("### 图表 4：公司单季度归母净利润及同比增速\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart4_profit_quarterly", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        # 第二章：分业务分析
        md.append("## 二、分业务分析\n")
        md.append(f"{report_data.get('business_analysis', '')}\n")
        
        # 第三章：盈利能力与费用分析
        md.append("## 三、盈利能力与费用分析\n")
        md.append(f"{report_data.get('profitability_analysis', '')}\n")
        
        # 图表 5-8
        md.append("### 图表 5：公司年度毛利率、净利率\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart5_margin_annual", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        md.append("### 图表 6：公司单季度毛利率和净利率\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart6_margin_quarterly", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        md.append("### 图表 7：公司年度期间费用率\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart7_expense_annual", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        md.append("### 图表 8：公司单季度期间费用率\n")
        md.append("```echarts\n")
        md.append(json.dumps(charts.get("chart8_expense_quarterly", {}), ensure_ascii=False, indent=2))
        md.append("\n```\n")
        
        # 第四章：经营展望
        md.append("## 四、经营展望\n")
        md.append(f"{report_data.get('outlook', '')}\n")
        
        # 第五章：风险提示
        md.append("## 五、风险提示\n")
        md.append(f"{report_data.get('risks', '')}\n")
        
        # 第六章：投资建议
        md.append("## 六、投资建议\n")
        md.append(f"{report_data.get('recommendation', '')}\n")
        
        # 数据来源与免责声明
        md.append("---\n")
        md.append("**数据来源：** 聚源数据、公司公告、Wind、分析师研报\n")
        md.append(f"**数据截止日期：** {report_data.get('data_date', 'N/A')}\n")
        md.append(f"**报告生成时间：** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        md.append("**免责声明：** 本报告仅供参考，不构成投资建议\n")
        
        # 公告信息溯源
        if report_data.get("announcements"):
            md.append("\n## 附录：重要公告\n")
            for ann in report_data["announcements"]:
                md.append(f"- {ann.get('title', '')} ({ann.get('date', '')})\n")
        
        content = "\n".join(md)
        
        # 保存文件
        stock_code = stock_info.get("stock_code", "unknown")
        report_date = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{stock_code}_earnings_review_{report_date}.md"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        
        print(f"Markdown 报告已保存：{filepath}")
        return filepath
    
    def generate_word_report(self, markdown_path: str) -> str:
        """
        将 Markdown 报告转换为 Word (.docx) 格式
        
        Args:
            markdown_path: Markdown 文件路径
        
        Returns:
            Word 文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            print("错误：需要安装 python-docx 库 (pip install python-docx)")
            print("降级：仅生成 Markdown 格式")
            return None
        
        # 读取 Markdown 文件
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_content = f.read()
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 解析 Markdown（简化版）
        lines = md_content.split("\n")
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 标题处理
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("**") and line.endswith("**"):
                # 粗体文本
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            elif line.startswith("---"):
                doc.add_paragraph()  # 分隔线
            elif line.startswith("```"):
                # 代码块（echarts 配置）
                current_section = "code" if "echarts" in line else None
                if current_section:
                    p = doc.add_paragraph()
                    p.add_run("[echarts 图表配置 - 请在支持 echarts 的环境中查看]").italic = True
            elif current_section == "code":
                if line.startswith("```"):
                    current_section = None
                # 跳过代码块内容
            else:
                # 普通段落
                doc.add_paragraph(line)
        
        # 保存 Word 文件
        base_name = os.path.splitext(os.path.basename(markdown_path))[0]
        word_filename = f"{base_name}.docx"
        word_path = os.path.join(self.output_dir, word_filename)
        
        doc.save(word_path)
        print(f"Word 报告已保存：{word_path}")
        return word_path
    
    def generate_pdf_report(self, markdown_path: str) -> str:
        """
        将 Markdown 报告转换为 PDF 格式
        
        Args:
            markdown_path: Markdown 文件路径
        
        Returns:
            PDF 文件路径
        """
        try:
            # 尝试使用 markdown2pdf
            import markdown
            from weasyprint import HTML, CSS
        except ImportError:
            print("提示：安装 weasyprint 可生成 PDF (pip install weasyprint)")
            print("降级：仅生成 Markdown 格式")
            return None
        
        # 读取 Markdown 文件
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_content = f.read()
        
        # 转换为 HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'toc']
        )
        
        # 添加 CSS 样式
        full_html = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2.5cm;
                }}
                body {{
                    font-family: "SimSun", "宋体", serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    font-size: 18pt;
                    text-align: center;
                    border-bottom: 2px solid #333;
                    padding-bottom: 10px;
                }}
                h2 {{
                    font-size: 14pt;
                    border-bottom: 1px solid #ccc;
                    padding-bottom: 5px;
                    margin-top: 20px;
                }}
                h3 {{
                    font-size: 12pt;
                    font-weight: bold;
                    margin-top: 15px;
                }}
                p {{
                    text-align: justify;
                    margin: 8px 0;
                }}
                .echarts-placeholder {{
                    background: #f5f5f5;
                    border: 1px dashed #ccc;
                    padding: 20px;
                    text-align: center;
                    margin: 15px 0;
                    font-style: italic;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 15px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: center;
                }}
                th {{
                    background-color: #f0f0f0;
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # 生成 PDF
        html_doc = HTML(string=full_html)
        css = CSS(string='@page { size: A4; margin: 2.5cm; }')
        
        base_name = os.path.splitext(os.path.basename(markdown_path))[0]
        pdf_filename = f"{base_name}.pdf"
        pdf_path = os.path.join(self.output_dir, pdf_filename)
        
        html_doc.write_pdf(pdf_path, stylesheets=[css])
        print(f"PDF 报告已保存：{pdf_path}")
        return pdf_path
    
    def generate_all_formats(self, report_data: Dict) -> Dict[str, str]:
        """
        生成所有格式的报告
        
        Args:
            report_data: 报告数据字典
        
        Returns:
            包含各格式文件路径的字典
        """
        output_files = {}
        
        # 1. 生成 Markdown
        md_path = self.generate_markdown_report(report_data)
        output_files["markdown"] = md_path
        
        # 2. 生成 Word
        try:
            word_path = self.generate_word_report(md_path)
            if word_path:
                output_files["word"] = word_path
        except Exception as e:
            print(f"Word 生成失败：{e}")
        
        # 3. 生成 PDF
        try:
            pdf_path = self.generate_pdf_report(md_path)
            if pdf_path:
                output_files["pdf"] = pdf_path
        except Exception as e:
            print(f"PDF 生成失败：{e}")
        
        return output_files


def create_sample_report_data() -> Dict:
    """创建示例报告数据（用于测试）"""
    return {
        "stock_info": {
            "company_name": "宁德时代新能源科技股份有限公司",
            "stock_code": "300750",
            "industry": "电力设备 - 电池",
            "market": "创业板"
        },
        "data_date": "2025-12-31",
        "core_view": """公司 2025 年实现营收 XXXX 亿元，同比 +XX%；归母净利润 XXX 亿元，同比 +XX%。
动力电池全球市占率持续提升，储能业务快速增长。
我们预计公司 2026-2028 年保持稳健增长，维持"买入"评级。""",
        "company_overview": """2025 年公司实现营业收入 XXXX 亿元，同比增长 XX%；
归母净利润 XXX 亿元，同比增长 XX%，扣非净利润 XXX 亿元，同比增长 XX%。
分季度看，Q4 单季度营收 XXX 亿元，同比 +XX%，环比 +XX%。
""",
        "business_analysis": """分业务看：
- 动力电池系统：营收 XXX 亿元，同比 +XX%，占比 XX%
- 储能系统：营收 XXX 亿元，同比 +XX%，占比 XX%
- 电池材料及回收：营收 XXX 亿元，同比 +XX%
""",
        "profitability_analysis": """盈利能力方面：
- 2025 年毛利率 XX.X%，同比 +X.Xpct
- 净利率 XX.X%，同比 +X.Xpct
- 期间费用率 XX.X%，同比 -X.Xpct
""",
        "outlook": """展望未来：
- 产能规划：预计 2026 年产能达到 XXXGWh
- 技术进展：神行电池、麒麟电池量产推进
- 海外市场：欧洲、北美基地建设中
""",
        "risks": """- 原材料价格波动风险
- 行业竞争加剧风险
- 下游需求不及预期风险
- 海外政策风险
""",
        "recommendation": """我们预计公司 2026-2028 年归母净利润分别为 XXX/XXX/XXX 亿元，
对应 EPS 为 X.XX/X.XX/X.XX 元。
参考同行业可比公司，给予 2026 年 XX 倍 PE，目标价 XXX 元，
维持"买入"评级。
""",
        "charts": {
            "chart1_revenue_annual": {"title": {"text": "图 1：年度营业收入及同比增速"}},
            "chart2_revenue_quarterly": {"title": {"text": "图 2：单季度营业收入及同比增速"}},
            "chart3_profit_annual": {"title": {"text": "图 3：年度归母净利润及同比增速"}},
            "chart4_profit_quarterly": {"title": {"text": "图 4：单季度归母净利润及同比增速"}},
            "chart5_margin_annual": {"title": {"text": "图 5：年度毛利率、净利率"}},
            "chart6_margin_quarterly": {"title": {"text": "图 6：单季度毛利率和净利率"}},
            "chart7_expense_annual": {"title": {"text": "图 7：年度期间费用率"}},
            "chart8_expense_quarterly": {"title": {"text": "图 8：单季度期间费用率"}}
        },
        "announcements": [
            {"title": "2025 年年度报告", "date": "2026-03-15"},
            {"title": "2025 年度利润分配预案", "date": "2026-03-15"}
        ]
    }


if __name__ == "__main__":
    import sys
    
    generator = ReportGenerator(output_dir="./reports")
    
    # 测试模式：生成示例报告
    if len(sys.argv) < 2:
        print("用法：python generate_report.py <股票代码> [报告数据 JSON 文件]")
        print("示例：python generate_report.py 300750")
        print("      python generate_report.py 300750 data.json")
        print("\n测试模式：生成示例报告...")
        
        report_data = create_sample_report_data()
        output_files = generator.generate_all_formats(report_data)
        
        print("\n=== 报告生成完成 ===")
        for fmt, path in output_files.items():
            print(f"{fmt}: {path}")
        sys.exit(0)
    
    stock_code = sys.argv[1]
    
    # 从 JSON 文件加载数据或生成空报告
    if len(sys.argv) > 2:
        json_file = sys.argv[2]
        with open(json_file, "r", encoding="utf-8") as f:
            report_data = json.load(f)
    else:
        print(f"警告：未提供报告数据文件，生成空模板报告")
        report_data = {
            "stock_info": {
                "company_name": f"公司（{stock_code}）",
                "stock_code": stock_code
            },
            "data_date": datetime.now().strftime("%Y-%m-%d"),
            "charts": {},
            "announcements": []
        }
    
    output_files = generator.generate_all_formats(report_data)
    
    print("\n=== 报告生成完成 ===")
    for fmt, path in output_files.items():
        print(f"{fmt}: {path}")
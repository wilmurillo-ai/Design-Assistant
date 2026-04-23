#!/usr/bin/env python3
"""
export_reply.py — Convert conversation content to multiple file formats.

Usage:
    python3 export_reply.py --content "..." --format pdf --output ~/Desktop/reply.pdf
    python3 export_reply.py --file input.md --format docx --output ~/Desktop/reply.docx
    python3 export_reply.py --file input.md --format all --output ~/Desktop/

Supported formats: md, txt, html, pdf, docx
"""

import argparse
import os
import sys
from pathlib import Path
from datetime import datetime


def sanitize_filename(title: str) -> str:
    """Convert title to safe filename."""
    import re
    safe = re.sub(r'[^\w\s\-\u4e00-\u9fff]', '', title)
    safe = re.sub(r'\s+', '_', safe.strip())
    return safe[:60] or "export"


def write_markdown(content: str, output_path: Path) -> Path:
    output_path.write_text(content, encoding="utf-8")
    return output_path


def write_txt(content: str, output_path: Path) -> Path:
    # Strip markdown syntax for plain text
    import re
    text = re.sub(r'#{1,6}\s+', '', content)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`{1,3}[^`]*`{1,3}', lambda m: m.group().strip('`'), text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    output_path.write_text(text, encoding="utf-8")
    return output_path


def write_html(content: str, output_path: Path, title: str = "Export") -> Path:
    try:
        import markdown as md_lib
        body_html = md_lib.markdown(content, extensions=['fenced_code', 'tables', 'toc'])
    except ImportError:
        # Fallback: wrap in <pre> if markdown lib unavailable
        import html as html_lib
        body_html = f"<pre>{html_lib.escape(content)}</pre>"

    import html as _html_mod
    title_esc = _html_mod.escape(title)
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title_esc}</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
           max-width: 860px; margin: 48px auto; padding: 0 24px;
           color: #1a1a1a; line-height: 1.7; }}
    h1, h2, h3 {{ border-bottom: 1px solid #e5e7eb; padding-bottom: 6px; }}
    code {{ background: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-size: 0.9em; }}
    pre  {{ background: #f3f4f6; padding: 16px; border-radius: 8px; overflow-x: auto; }}
    pre code {{ background: none; padding: 0; }}
    blockquote {{ border-left: 4px solid #d1d5db; margin: 0; padding-left: 16px; color: #6b7280; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #e5e7eb; padding: 8px 12px; text-align: left; }}
    th {{ background: #f9fafb; }}
    .meta {{ color: #9ca3af; font-size: 0.85em; margin-bottom: 32px; }}
  </style>
</head>
<body>
  <h1>{title}</h1>
  <p class="meta">导出时间 / Exported at：{datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
  {body_html}
</body>
</html>"""
    output_path.write_text(html, encoding="utf-8")
    return output_path


def _find_chrome():  # -> Optional[str]
    """Locate a Chrome/Chromium executable for headless PDF rendering."""
    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/usr/bin/google-chrome",
        "/usr/bin/google-chrome-stable",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser",
    ]
    for c in candidates:
        if os.path.exists(c):
            return c
    # Try PATH
    import shutil
    for name in ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser"):
        found = shutil.which(name)
        if found:
            return found
    return None


def write_pdf(content: str, output_path: Path, title: str = "Export") -> Path:
    import tempfile, subprocess

    # Strategy 1: Chrome / Chromium headless (best quality, CJK native)
    chrome = _find_chrome()
    if chrome:
        html_tmp = Path(tempfile.mktemp(suffix=".html"))
        try:
            write_html(content, html_tmp, title)
            result = subprocess.run(
                [
                    chrome,
                    "--headless", "--disable-gpu", "--no-sandbox",
                    "--run-all-compositor-stages-before-draw",
                    f"--print-to-pdf={output_path}",
                    "--print-to-pdf-no-header",
                    str(html_tmp),
                ],
                capture_output=True, timeout=30,
            )
            if output_path.exists() and output_path.stat().st_size > 1000:
                return output_path
        except Exception:
            pass
        finally:
            html_tmp.unlink(missing_ok=True)

    # Strategy 2: weasyprint
    try:
        from weasyprint import HTML as WH
        html_tmp = Path(tempfile.mktemp(suffix=".html"))
        write_html(content, html_tmp, title)
        WH(filename=str(html_tmp)).write_pdf(str(output_path))
        html_tmp.unlink(missing_ok=True)
        return output_path
    except ImportError:
        pass

    # Strategy 3: pdfkit (wkhtmltopdf wrapper)
    try:
        import pdfkit
        html_tmp = Path(tempfile.mktemp(suffix=".html"))
        write_html(content, html_tmp, title)
        pdfkit.from_file(str(html_tmp), str(output_path))
        html_tmp.unlink(missing_ok=True)
        return output_path
    except (ImportError, Exception):
        pass

    # Strategy 4: fpdf2 — write each character individually to bypass line-break bug
    try:
        from fpdf import FPDF
        import re as _re, unicodedata as _ud

        def _plain(t):
            # Strip Markdown
            t = _re.sub(r'^#{1,6}\s+', '', t, flags=_re.MULTILINE)
            t = _re.sub(r'\*{1,2}(.+?)\*{1,2}', r'\1', t)
            t = _re.sub(r'`+([^`]*)`+', r'\1', t)
            t = _re.sub(r'\[(.+?)\]\(.+?\)', r'\1', t)
            t = _re.sub(r'^>+\s?', '', t, flags=_re.MULTILINE)  # blockquotes
            # Fullwidth / special → ASCII (prevents fpdf2 metric bug)
            for src, dst in [
                ('\uff1a',':'), ('\uff1b',';'), ('\uff0c',','), ('\u3002','.'),
                ('\u300c','['), ('\u300d',']'), ('\u300a','<'), ('\u300b','>'),
                ('\u00b7','-'), ('\u2014','--'), ('\u2013','-'),
                ('\u2019',"'"), ('\u2018',"'"), ('\u201c','"'), ('\u201d','"'),
                ('\u2022','-'), ('\u2026','...'), ('\u00d7','x'), ('\u00b6',''),
                ('\uff08','('), ('\uff09',')'), ('\uff3b','['), ('\uff3d',']'),
            ]:
                t = t.replace(src, dst)
            # Remove control chars
            t = ''.join(c for c in t if _ud.category(c)[0] != 'C' or c in ('\n', '\t'))
            return t

        plain = _plain(content)
        pdf = FPDF()
        pdf.set_margins(20, 20, 20)
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)

        cjk_font_paths = [
            "/Library/Fonts/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        ]
        cjk_font = next((p for p in cjk_font_paths if os.path.exists(p)), None)
        use_cjk = False
        if cjk_font:
            try:
                pdf.add_font("CJK", fname=cjk_font)
                pdf.set_font("CJK", size=11)
                use_cjk = True
            except Exception:
                pass
        if not use_cjk:
            pdf.set_font("Helvetica", size=11)
            # Replace CJK with ASCII placeholder when no CJK font
            plain = _re.sub(r'[\u4e00-\u9fff\u3400-\u4dbf]+', '[CJK]', plain)

        pdf.set_title(title[:60])
        cell_w = pdf.epw
        for line in plain.split("\n"):
            ln = line.strip() or " "
            # Split very long lines to avoid width overflow
            while len(ln) > 120:
                chunk, ln = ln[:120], ln[120:]
                try:
                    pdf.multi_cell(cell_w, 7, chunk)
                except Exception:
                    pass
            if ln:
                try:
                    pdf.multi_cell(cell_w, 7, ln)
                except Exception:
                    try:
                        pdf.multi_cell(cell_w, 7, ln.encode('ascii', 'replace').decode())
                    except Exception:
                        pass
        pdf.output(str(output_path))
        # Verify output has text streams
        raw = output_path.read_bytes()
        if b'BT' not in raw:
            raise RuntimeError("fpdf2 produced empty PDF (no text streams)")
        return output_path
    except ImportError:
        pass

    raise RuntimeError(
        "PDF 导出失败。请安装以下任意一个库：\n"
        "  pip3 install fpdf2          # 纯 Python，推荐\n"
        "  pip3 install weasyprint     # 最佳质量\n"
        "或安装 Google Chrome (用于 headless 渲染)"
    )


def write_docx(content: str, output_path: Path, title: str = "Export") -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        # Title
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        meta = doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        meta.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
        meta.runs[0].font.size = Pt(9)

        for line in content.split("\n"):
            h_match = re.match(r'^(#{1,6})\s+(.*)', line)
            if h_match:
                level = len(h_match.group(1))
                doc.add_heading(h_match.group(2), level=min(level, 9))
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s+', line.strip()):
                doc.add_paragraph(re.sub(r'^\d+\.\s+', '', line.strip()), style='List Number')
            else:
                doc.add_paragraph(line)

        doc.save(str(output_path))
        return output_path
    except ImportError:
        raise RuntimeError(
            "DOCX 导出需要 python-docx\n安装命令：pip3 install python-docx"
        )


FORMAT_MAP = {
    "md":   (".md",   write_markdown),
    "txt":  (".txt",  write_txt),
    "html": (".html", write_html),
    "pdf":  (".pdf",  write_pdf),
    "docx": (".docx", write_docx),
}


def export(content: str, fmt: str, output: str, title: str = "对话导出") -> list[str]:
    """Export content to one or more formats. Returns list of created file paths."""
    output_path = Path(output).expanduser()
    created = []

    formats = list(FORMAT_MAP.keys()) if fmt == "all" else [fmt.lower()]

    for f in formats:
        if f not in FORMAT_MAP:
            print(f"[WARN] 不支持的格式：{f}，跳过", file=sys.stderr)
            continue

        ext, writer = FORMAT_MAP[f]

        if output_path.is_dir() or fmt == "all":
            output_path.mkdir(parents=True, exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M")
            fname = f"{sanitize_filename(title)}_{ts}{ext}"
            file_path = output_path / fname
        else:
            file_path = output_path
            file_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            result = writer(content, file_path, title) if f in ("html", "pdf", "docx") else writer(content, file_path)
            print(f"✅ {f.upper()} 已保存：{result}")
            created.append(str(result))
        except Exception as e:
            print(f"❌ {f.upper()} 导出失败：{e}", file=sys.stderr)

    return created


def main():
    parser = argparse.ArgumentParser(description="Export conversation to file formats")
    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument("--content", help="Direct text content to export")
    src.add_argument("--file", help="Path to input markdown/text file")
    parser.add_argument("--format", required=True,
                        choices=list(FORMAT_MAP.keys()) + ["all"],
                        help="Output format (md/txt/html/pdf/docx/all)")
    parser.add_argument("--output", required=True, help="Output file or directory path")
    parser.add_argument("--title", default="对话导出", help="Document title")
    args = parser.parse_args()

    if args.file:
        content = Path(args.file).expanduser().read_text(encoding="utf-8")
    else:
        content = args.content

    created = export(content, args.format, args.output, args.title)
    if not created:
        sys.exit(1)


if __name__ == "__main__":
    main()

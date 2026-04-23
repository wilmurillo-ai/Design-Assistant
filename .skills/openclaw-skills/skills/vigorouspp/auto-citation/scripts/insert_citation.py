#!/usr/bin/env python3
"""
引用插入工具：将格式化后的引用插入原文档
支持格式：Markdown / LaTeX / Word (.docx)
"""

import re
import sys
import os
import argparse
import shutil
from pathlib import Path
from typing import List, Dict, Optional


class CitationInserter:
    """引用插入基类"""
    
    def __init__(self, document_path: str, citations_path: str, output_path: Optional[str] = None):
        self.document_path = Path(document_path)
        self.citations_path = Path(citations_path)
        self.output_path = Path(output_path) if output_path else None
        self.citations = []
        self.style = 'gb7714'  # 默认格式
    
    def load_citations(self):
        """加载引用内容"""
        with open(self.citations_path, 'r', encoding='utf-8') as f:
            self.citations_text = f.read()
        
        # 自动检测格式
        if self.citations_text.startswith('@'):
            self.style = 'bibtex'
        elif any(re.search(r'\[\d+\]', line) for line in self.citations_text.split('\n')[:10]):
            self.style = 'gb7714'
        elif any('(' in line and ')' in line and any(year in line for year in ['2020', '2021', '2022', '2023', '2024']) for line in self.citations_text.split('\n')[:10]):
            self.style = 'apa'
    
    def insert(self):
        """插入引用（子类实现）"""
        raise NotImplementedError
    
    def backup_original(self):
        """备份原始文件"""
        backup_path = self.document_path.with_suffix(self.document_path.suffix + '.backup')
        shutil.copy2(self.document_path, backup_path)
        return backup_path


class MarkdownInserter(CitationInserter):
    """Markdown 文档引用插入器"""
    
    def __init__(self, document_path: str, citations_path: str, output_path: Optional[str] = None):
        super().__init__(document_path, citations_path, output_path)
    
    def insert(self):
        """在 Markdown 文档末尾插入参考文献"""
        # 备份
        self.backup_original()
        
        # 读取原文档
        with open(self.document_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 检查是否已有参考文献部分
        has_ref_section = re.search(r'^##?\s*(?:参考文献|References|Bibliography)', 
                                    content, re.MULTILINE | re.IGNORECASE)
        
        if has_ref_section:
            # 替换现有参考文献部分
            content = re.sub(
                r'^##?\s*(?:参考文献|References|Bibliography).*?(?=\n##|\Z)',
                f'## 参考文献\n\n{self.citations_text}',
                content,
                flags=re.DOTALL | re.MULTILINE | re.IGNORECASE
            )
        else:
            # 在文档末尾添加
            content = content.rstrip() + '\n\n## 参考文献\n\n' + self.citations_text
        
        # 写入文件
        output_path = self.output_path or self.document_path
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return output_path
    
    def insert_inline_markers(self, marker_positions: List[int], marker_labels: List[str]):
        """在 Markdown 中插入文中引用标记 [1], [2] 等"""
        with open(self.document_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 简单的行级插入（更精确的插入需要段落定位）
        modified_lines = []
        citation_idx = 0
        
        for i, line in enumerate(lines):
            modified_lines.append(line)
            
            # 在段落末尾添加引用（简单策略）
            if i in marker_positions and citation_idx < len(marker_labels):
                # 如果行以标点结尾，在标点之前插入
                stripped = line.rstrip()
                if stripped and not stripped.endswith('```') and not stripped.startswith('#'):
                    label = marker_labels[citation_idx]
                    # 在行末尾添加引用标记
                    if stripped[-1] in '.!?;':
                        modified_lines[-1] = stripped[:-1] + f'[{label}]' + stripped[-1] + '\n'
                    else:
                        modified_lines[-1] = stripped + f'[{label}]\n'
                    citation_idx += 1
        
        output_path = self.output_path or self.document_path
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(modified_lines)
        
        return output_path


class LaTeXInserter(CitationInserter):
    """LaTeX 文档引用插入器"""
    
    def insert(self):
        """在 LaTeX 文档中插入引用"""
        self.backup_original()
        
        with open(self.document_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if self.style == 'bibtex':
            return self._insert_bibtex(content)
        else:
            return self._insert_thebibliography(content)
    
    def _insert_bibtex(self, content: str):
        """使用 BibTeX 方式"""
        # 提取 .bib 文件路径
        bib_match = re.search(r'\\bibliography\{([^}]+)\}', content)
        
        if bib_match:
            # 更新现有 .bib 文件
            bib_file = bib_match.group(1)
            if not bib_file.endswith('.bib'):
                bib_file += '.bib'
            
            # 构建完整路径
            bib_path = self.document_path.parent / bib_file
            
            # 追加或创建 .bib 文件
            with open(bib_path, 'a', encoding='utf-8') as f:
                f.write('\n\n% Added by auto-citation on ' + str(Path(__file__).stat().st_mtime) + '\n')
                f.write(self.citations_text)
            
            output_path = self.output_path or self.document_path
            shutil.copy2(self.document_path, output_path)
            
        else:
            # 没有 \bibliography 命令，需要添加
            # 在 \end{document} 之前添加
            citations_file = self.document_path.stem + '_references.bib'
            bib_path = self.document_path.parent / citations_file
            
            # 写入 .bib 文件
            with open(bib_path, 'w', encoding='utf-8') as f:
                f.write(self.citations_text)
            
            # 修改 .tex 文件添加引用命令
            # 在 \end{document} 之前添加 \bibliography 和 \bibliographystyle
            new_content = re.sub(
                r'(\\end\{document\})',
                f'\\bibliographystyle{{plain}}\n\\bibliography{{{self.document_path.stem}_references}}\n\n\\1',
                content
            )
            
            output_path = self.output_path or self.document_path
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(new_content)
        
        return output_path
    
    def _insert_thebibliography(self, content: str):
        """使用 thebibliography 环境（GB7714/APA 等非 BibTeX 格式）"""
        # 检查是否已有 thebibliography
        has_biblio = re.search(r'\\begin\{thebibliography\}', content)
        
        if has_biblio:
            # 替换现有环境
            content = re.sub(
                r'\\begin\{thebibliography\}.*?\\end\{thebibliography\}',
                f'\\begin{{thebibliography}}{{99}}\n\\small\n\n{self.citations_text}\n\\end{{thebibliography}}',
                content,
                flags=re.DOTALL
            )
        else:
            # 在 \end{document} 之前添加
            new_content = re.sub(
                r'(\\end\{document\})',
                f'\\begin{{thebibliography}}{{99}}\n\\small\n\n{self.citations_text}\n\\end{{thebibliography}}\n\n\\1',
                content
            )
            content = new_content
        
        output_path = self.output_path or self.document_path
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return output_path
    
    def insert_cite_commands(self, cite_keys: List[str], positions: List[int]):
        """在 LaTeX 中插入 \cite{} 命令"""
        with open(self.document_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 简单的行级插入
        modified_lines = []
        cite_idx = 0
        
        for i, line in enumerate(lines):
            modified_lines.append(line)
            
            if i in positions and cite_idx < len(cite_keys):
                # 在段落末尾添加 \cite{}
                stripped = line.rstrip()
                if stripped and not stripped.startswith('%') and not stripped.startswith('\\'):
                    key = cite_keys[cite_idx]
                    if stripped[-1] in '.!?':
                        modified_lines[-1] = stripped[:-1] + f'~\\\\cite{{{key}}}' + stripped[-1] + '\n'
                    else:
                        modified_lines[-1] = stripped + f'~\\\\cite{{{key}}}\n'
                    cite_idx += 1
        
        output_path = self.output_path or self.document_path
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(modified_lines)
        
        return output_path


class WordInserter(CitationInserter):
    """Word 文档引用插入器"""
    
    def insert(self):
        """在 Word 文档末尾插入参考文献"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
            sys.exit(1)
        
        # 备份
        self.backup_original()
        
        # 打开文档
        doc = Document(self.document_path)
        
        # 添加参考文献标题
        doc.add_heading('参考文献', level=1)
        
        # 添加引用内容
        for line in self.citations_text.strip().split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.15
                # 首行缩进
                p.paragraph_format.first_line_indent = Pt(0)
        
        # 保存
        output_path = self.output_path or self.document_path
        doc.save(output_path)
        
        return output_path


def detect_file_type(file_path: str) -> str:
    """检测文件类型"""
    ext = Path(file_path).suffix.lower()
    if ext == '.md':
        return 'markdown'
    elif ext in ['.tex', '.latex']:
        return 'latex'
    elif ext == '.docx':
        return 'word'
    return 'unknown'


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='Insert citations into documents')
    parser.add_argument('--document', type=str, required=True,
                       help='Path to the document file')
    parser.add_argument('--citations', type=str, required=True,
                       help='Path to the formatted citations file')
    parser.add_argument('--output', type=str,
                       help='Output file path (default: overwrite original)')
    parser.add_argument('--style', choices=['bibtex', 'gb7714', 'apa'],
                       help='Citation style (auto-detected if not specified)')
    
    args = parser.parse_args()
    
    # 检查文件存在
    if not os.path.exists(args.document):
        print(f"Error: Document file not found: {args.document}", file=sys.stderr)
        sys.exit(1)
    
    if not os.path.exists(args.citations):
        print(f"Error: Citations file not found: {args.citations}", file=sys.stderr)
        sys.exit(1)
    
    # 检测文件类型
    file_type = detect_file_type(args.document)
    
    # 选择插入器
    inserter_map = {
        'markdown': MarkdownInserter,
        'latex': LaTeXInserter,
        'word': WordInserter,
    }
    
    inserter_class = inserter_map.get(file_type)
    if not inserter_class:
        print(f"Error: Unsupported file type: {file_type}", file=sys.stderr)
        sys.exit(1)
    
    # 创建插入器并执行
    inserter = inserter_class(args.document, args.citations, args.output)
    inserter.load_citations()
    
    # 如果指定了 style，覆盖自动检测的
    if args.style:
        inserter.style = args.style
    
    output_path = inserter.insert()
    
    print(f"Citations inserted successfully.")
    print(f"Output file: {output_path}")
    print(f"Citation style: {inserter.style}")
    
    # 显示摘要
    print(f"\nInserted {len(inserter.citations_text.strip().split(chr(10)))} lines of citations.")


if __name__ == '__main__':
    main()

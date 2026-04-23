#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文字游戏剧本导出 Word 工具
用法: python export_word.py <输出文件名> <剧本标题> <剧本内容>
"""

import sys
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def create_word_document(output_path, title, content):
    """创建Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加生成时间
    time_paragraph = doc.add_paragraph()
    time_run = time_paragraph.add_run(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
    time_run.font.size = Pt(10)
    time_run.font.italic = True
    time_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 按段落分割内容
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # 检查是否是标题（# 开头）
            if para.startswith('# '):
                doc.add_heading(para[2:], level=2)
            elif para.startswith('## '):
                doc.add_heading(para[3:], level=3)
            elif para.startswith('```'):
                # 代码块，用引用格式
                code_content = para.strip('`').strip()
                quote = doc.add_paragraph(code_content)
                quote.style = doc.styles['Quote']
            else:
                # 普通段落，保留换行
                p = doc.add_paragraph()
                lines = para.split('\n')
                for i, line in enumerate(lines):
                    if line.startswith('- '):
                        p.add_run(line + '\n')
                    elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                        p.add_run(line + '\n')
                    else:
                        p.add_run(line)
                        if i < len(lines) - 1:
                            p.add_run('\n')
    
    # 保存文档
    doc.save(output_path)
    return True

def main():
    if len(sys.argv) < 4:
        print("用法: python export_word.py <输出路径> <标题> <内容文件>")
        print("或者: python export_word.py <输出路径> <标题> - <<EOF\n内容...\nEOF")
        sys.exit(1)
    
    output_path = sys.argv[1]
    title = sys.argv[2]
    content_file = sys.argv[3]
    
    if content_file == '-':
        # 从标准输入读取内容
        content = sys.stdin.read()
    else:
        # 从文件读取
        with open(content_file, 'r', encoding='utf-8') as f:
            content = f.read()
    
    success = create_word_document(output_path, title, content)
    if success:
        print(f"✓ Word文档已生成: {output_path}")
        print(f"✓ 文件大小: {os.path.getsize(output_path)} bytes")

if __name__ == '__main__':
    main()

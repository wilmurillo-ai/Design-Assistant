#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word to PDF Converter
Converts Word documents (.docx) to PDF using reportlab library.
Supports Chinese characters, emojis, and proper formatting.
"""

import sys
import os
import argparse
from docx import Document
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# 设置标准输出编码
if sys.platform == 'win32':
    try:
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    except:
        pass


def find_chinese_font():
    """
    Find and register a Chinese font on the system.
    Returns the font name or None if not found.
    """
    font_names = []

    # Windows fonts
    windows_fonts = [
        (r"C:\Windows\Fonts\msyh.ttc", "MicrosoftYaHei"),  # 微软雅黑
        (r"C:\Windows\Fonts\simhei.ttf", "SimHei"),  # 黑体
        (r"C:\Windows\Fonts\simkai.ttf", "KaiTi"),  # 楷体
        (r"C:\Windows\Fonts\simfang.ttf", "FangSong"),  # 仿宋
    ]

    # Try each Windows font
    for font_path, font_name in windows_fonts:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                print(f"[OK] Successfully registered font: {font_name}")
                return font_name
            except Exception as e:
                print(f"[ERROR] Cannot register {font_name}: {e}")

    # macOS fonts
    mac_fonts = [
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti"),
    ]

    for font_path, font_name in mac_fonts:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                print(f"[OK] Successfully registered font: {font_name}")
                return font_name
            except Exception as e:
                print(f"[ERROR] Cannot register {font_name}: {e}")

    # Linux fonts (common locations)
    linux_fonts = [
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WenQuanYiMicroHei"),
        ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", "DroidSansFallback"),
    ]

    for font_path, font_name in linux_fonts:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                print(f"[OK] Successfully registered font: {font_name}")
                return font_name
            except Exception as e:
                print(f"[ERROR] Cannot register {font_name}: {e}")

    print("[WARNING] No Chinese font found. Will use default font. Chinese may appear garbled.")
    return None


def create_heading_style(parent, level, font_name):
    """Create a heading style with the specified font."""
    if level == 1:
        return ParagraphStyle(
            f'Heading{level}',
            parent=parent['Heading1'],
            fontSize=24,
            spaceAfter=12,
            fontName=font_name,
            textColor=(0, 0, 0),
            alignment=TA_CENTER
        )
    elif level == 2:
        return ParagraphStyle(
            f'Heading{level}',
            parent=parent['Heading2'],
            fontSize=18,
            spaceAfter=10,
            fontName=font_name,
            textColor=(0, 0, 0)
        )
    else:
        return parent[f'Heading{level}']


def convert_docx_to_pdf(input_file, output_file=None, font_name=None):
    """
    Convert a Word document to PDF.

    Args:
        input_file: Path to input .docx file
        output_file: Path to output PDF file (optional)
        font_name: Font name to use for Chinese characters (optional)

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Check if input file exists
        if not os.path.exists(input_file):
            print(f"[ERROR] File not found: {input_file}")
            return False

        # Load Word document
        print(f"Loading document: {input_file}")
        doc = Document(input_file)

        # Set output file
        if output_file is None:
            base, ext = os.path.splitext(input_file)
            output_file = base + ".pdf"

        print(f"Generating PDF: {output_file}")

        # Register Chinese font
        if font_name is None:
            font_name = find_chinese_font()

        # Create PDF document
        pdf_file = SimpleDocTemplate(
            output_file,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        story = []
        styles = getSampleStyleSheet()

        # Apply font to all styles
        for name in ['Heading1', 'Heading2', 'Normal']:
            if name in styles.byName:
                style = styles.byName[name]
                if style.fontName == 'Helvetica':
                    style.fontName = font_name or 'Helvetica'

        # Create custom styles
        title_style = create_heading_style(styles, 1, font_name)
        heading1_style = create_heading_style(styles, 1, font_name)
        heading2_style = create_heading_style(styles, 2, font_name)
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leading=16,
            fontName=font_name or 'Helvetica',
            textColor=(0, 0, 0)
        )

        # Add title
        if doc.paragraphs:
            title_text = doc.paragraphs[0].text
            title = Paragraph(title_text, title_style)
            story.append(title)
            story.append(Spacer(1, 0.3 * inch))

        # Process all paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Determine paragraph style
            if para.style.name.startswith('Heading'):
                # Heading
                heading_level = int(para.style.name.replace('Heading', ''))
                if heading_level == 1:
                    p = Paragraph(text, heading1_style)
                elif heading_level == 2:
                    p = Paragraph(text, heading2_style)
                else:
                    p = Paragraph(text, styles.byName[f'Heading{heading_level}'])
            elif para.style.name == 'List Bullet':
                # Bullet list
                p = Paragraph(f"• {text}", body_style)
            elif para.style.name == 'List Number':
                # Numbered list
                p = Paragraph(f"1. {text}", body_style)
            else:
                # Regular paragraph
                p = Paragraph(text, body_style)

            story.append(p)
            story.append(Spacer(1, 0.1 * inch))

        # Generate PDF
        pdf_file.build(story)
        print(f"[OK] PDF generated successfully: {output_file}")
        print(f"      File size: {os.path.getsize(output_file) / 1024:.2f} KB")

        return True

    except Exception as e:
        print(f"[ERROR] PDF generation failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main function for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Convert Word documents to PDF format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s input.docx output.pdf
  %(prog)s document.docx
  %(prog)s input.docx
        """
    )

    parser.add_argument(
        'input_file',
        help='Input Word document path (.docx)'
    )

    parser.add_argument(
        'output_file',
        nargs='?',
        help='Output PDF file path (optional)'
    )

    args = parser.parse_args()

    success = convert_docx_to_pdf(args.input_file, args.output_file)

    if success:
        print("\n[SUCCESS] Conversion completed!")
        sys.exit(0)
    else:
        print("\n[FAILED] Conversion failed!")
        sys.exit(1)


if __name__ == "__main__":
    main()

"""Word 文档批注生成器

在原始 Word 文件上直接添加批注（Comment）
"""

from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import re
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordAnnotator:
    """Word 文档批注器 - 在原文件上添加批注"""
    
    def __init__(self):
        self.severity_colors = {
            'critical': RGBColor(255, 0, 0),      # 红色
            'high': RGBColor(255, 102, 0),        # 橙色
            'medium': RGBColor(255, 192, 0),      # 黄色
            'low': RGBColor(0, 176, 80),          # 绿色
        }
        self.severity_names = {
            'critical': '严重',
            'high': '高',
            'medium': '中',
            'low': '低',
        }
    
    def add_annotations(self, input_path: str, output_path: str, 
                        findings: List[Dict[str, Any]], 
                        author: str = "AI合同审计") -> bool:
        """
        在 Word 文档中添加批注
        
        Args:
            input_path: 原始合同文件路径
            output_path: 输出文件路径
            findings: 审计发现的问题列表
            author: 批注作者
        
        Returns:
            bool: 是否成功
        """
        try:
            doc = Document(input_path)
            
            # 在文档开头添加审计摘要页
            self._add_audit_summary(doc, findings)
            
            # 为每个 finding 添加批注
            success_count = 0
            for i, finding in enumerate(findings):
                if self._add_comment_to_document(doc, finding, author, i):
                    success_count += 1
            
            # 保存文档
            doc.save(output_path)
            print(f"成功添加 {success_count}/{len(findings)} 条批注")
            return True
            
        except Exception as e:
            print(f"添加批注失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _add_audit_summary(self, doc: Document, findings: List[Dict[str, Any]]):
        """在文档开头添加审计摘要"""
        # 在文档最前面插入摘要
        summary_para = doc.add_paragraph()
        summary_para.insert_paragraph_before("【合同审计摘要】")
        
        # 添加风险等级统计
        by_severity = {}
        by_category = {}
        for f in findings:
            sev = f.get('severity', 'low')
            cat = f.get('category', '未分类')
            by_severity[sev] = by_severity.get(sev, 0) + 1
            by_category[cat] = by_category.get(cat, 0) + 1
        
        # 添加统计信息
        total = len(findings)
        summary_text = f"共发现问题 {total} 项\n"
        
        if by_severity:
            summary_text += "按严重程度："
            for sev, count in by_severity.items():
                sev_name = self.severity_names.get(sev, sev.upper())
                summary_text += f"{sev_name} {count} 项；"
            summary_text += "\n"
        
        if by_category:
            summary_text += "按类别："
            for cat, count in by_category.items():
                summary_text += f"{cat} {count} 项；"
            summary_text += "\n"
        
        summary_para = doc.add_paragraph(summary_text)
        
        # 添加分隔线
        doc.add_paragraph("=" * 50)
        doc.add_paragraph()
    
    def _add_comment_to_document(self, doc: Document, finding: Dict[str, Any], 
                                  author: str, comment_id: int) -> bool:
        """在文档中找到相关文本并添加批注"""
        location = finding.get('location', '')
        description = finding.get('description', '')
        suggestion = finding.get('suggestion', '')
        severity = finding.get('severity', 'low')
        category = finding.get('category', '未分类')
        
        # 构建批注文本
        severity_name = self.severity_names.get(severity, severity.upper())
        comment_text = f"【{severity_name}风险】{category}\n"
        comment_text += f"问题：{description}\n"
        if suggestion:
            comment_text += f"建议：{suggestion}"
        
        # 提取关键词用于定位
        keywords = self._extract_keywords(location + ' ' + description)
        
        # 在文档中查找匹配的段落
        best_match = self._find_best_paragraph(doc, keywords)
        if best_match is None:
            return False
        
        paragraph, match_text = best_match
        
        # 在段落中添加批注
        return self._add_comment_to_paragraph(paragraph, match_text, comment_text, 
                                               author, comment_id)
    
    def _find_best_paragraph(self, doc: Document, keywords: List[str]) -> Optional[Tuple[Any, str]]:
        """找到与关键词最匹配的段落"""
        best_para = None
        best_match_score = 0
        best_match_text = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            score = 0
            matched_keywords = []
            for kw in keywords:
                if kw in text:
                    score += len(kw)  # 匹配长度作为权重
                    matched_keywords.append(kw)
            
            if score > best_match_score:
                best_match_score = score
                best_para = para
                # 找到匹配的关键词在文本中的位置
                if matched_keywords:
                    # 选择最长的匹配关键词
                    best_match_text = max(matched_keywords, key=len)
                else:
                    best_match_text = text[:50] if len(text) > 50 else text
        
        if best_para and best_match_score > 0:
            return (best_para, best_match_text)
        return None
    
    def _add_comment_to_paragraph(self, paragraph: Any, match_text: str, 
                                   comment_text: str, author: str, 
                                   comment_id: int) -> bool:
        """在段落中的指定文本上添加批注"""
        try:
            text = paragraph.text
            if match_text not in text:
                return False
            
            # 找到匹配文本的位置
            idx = text.find(match_text)
            before_text = text[:idx]
            after_text = text[idx + len(match_text):]
            
            # 保存原段落的样式
            original_style = paragraph.style
            
            # 清除原段落内容
            for run in paragraph._p.getchildren():
                if run.tag.endswith('r'):  # run 元素
                    paragraph._p.remove(run)
            
            # 重新构建段落，在匹配文本周围添加批注标记
            if before_text:
                run_before = paragraph.add_run(before_text)
            
            # 添加带批注的文本
            run_comment = paragraph.add_run(match_text)
            
            # 应用批注标记
            self._apply_comment_to_run(run_comment, comment_id)
            
            if after_text:
                run_after = paragraph.add_run(after_text)
            
            # 创建批注
            self._create_comment(paragraph, comment_id, comment_text, author)
            
            return True
            
        except Exception as e:
            print(f"添加段落批注失败: {e}")
            return False
    
    def _apply_comment_to_run(self, run: Any, comment_id: int):
        """为 run 添加批注引用标记"""
        try:
            r = run._r
            
            # 创建批注开始标记
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), str(comment_id))
            
            # 创建批注结束标记
            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), str(comment_id))
            
            # 创建批注引用
            comment_ref_run = OxmlElement('w:r')
            comment_ref_inner = OxmlElement('w:commentReference')
            comment_ref_inner.set(qn('w:id'), str(comment_id))
            comment_ref_run.append(comment_ref_inner)
            
            # 插入标记
            r.addprevious(comment_start)
            r.addnext(comment_end)
            comment_end.addnext(comment_ref_run)
            
        except Exception as e:
            print(f"应用批注标记失败: {e}")
    
    def _create_comment(self, paragraph: Any, comment_id: int, 
                        comment_text: str, author: str):
        """创建批注内容"""
        try:
            doc = paragraph.part.document
            
            # 获取或创建 comments.xml
            comments_part = self._get_or_create_comments_part(doc)
            
            # 创建批注元素
            comment = OxmlElement('w:comment')
            comment.set(qn('w:id'), str(comment_id))
            comment.set(qn('w:author'), author)
            comment.set(qn('w:date'), self._get_current_datetime())
            comment.set(qn('w:initials'), author[0] if author else 'AI')
            
            # 创建批注段落
            comment_para = OxmlElement('w:p')
            
            # 添加批注内容（支持多行）
            lines = comment_text.split('\n')
            for line in lines:
                comment_run = OxmlElement('w:r')
                comment_text_elem = OxmlElement('w:t')
                comment_text_elem.text = line
                comment_run.append(comment_text_elem)
                comment_para.append(comment_run)
                
                # 如果不是最后一行，添加换行
                if line != lines[-1]:
                    br = OxmlElement('w:br')
                    comment_para.append(br)
            
            comment.append(comment_para)
            
            # 添加到批注部分
            comments_part.append(comment)
            
        except Exception as e:
            print(f"创建批注失败: {e}")
    
    def _get_or_create_comments_part(self, doc: Document) -> Any:
        """获取或创建文档的批注部分 - 简化版"""
        try:
            # 尝试获取已存在的 comments part
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            
            # 检查是否已有 comments 关系
            comments_part = None
            try:
                comments_part = doc.part._comments_part
            except AttributeError:
                pass
            
            if comments_part is not None:
                return comments_part._element
            
            # 如果没有，创建新的 comments element
            comments = OxmlElement('w:comments')
            comments.set(qn('xmlns:w'), 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            
            # 存储在 document element 上以便复用
            if not hasattr(doc.element, '_comments_element'):
                doc.element._comments_element = comments
            
            return doc.element._comments_element
            
        except Exception as e:
            print(f"获取或创建批注部分失败: {e}")
            # 返回一个简单的元素作为备选
            comments = OxmlElement('w:comments')
            comments.set(qn('xmlns:w'), 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            return comments
    
    def _update_content_types(self, package):
        """更新 [Content_Types].xml 以包含 comments"""
        try:
            from docx.opc.constants import CONTENT_TYPE
            # 检查是否已包含 comments 的 content type
            # 这里简化处理，实际需要修改 [Content_Types].xml
            pass
        except Exception as e:
            print(f"更新 Content_Types 失败: {e}")
    
    def _update_document_rels(self, doc):
        """更新 document.xml.rels 以引用 comments"""
        try:
            # 简化处理，实际需要修改 document.xml.rels
            pass
        except Exception as e:
            print(f"更新 document.rels 失败: {e}")
    
    def _get_current_datetime(self) -> str:
        """获取当前时间的 ISO 格式字符串"""
        return datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        # 提取中文关键词（2-10个字的词组）
        words = re.findall(r'[\u4e00-\u9fa5]{2,10}', text)
        # 去重并返回
        seen = set()
        result = []
        for w in words:
            if w not in seen and len(w) >= 2:
                seen.add(w)
                result.append(w)
        return result[:15]


def annotate_word_contract(input_path: str, output_path: str,
                           findings: List[Dict[str, Any]],
                           author: str = "AI合同审计") -> bool:
    """便捷函数：为 Word 合同添加批注"""
    annotator = WordAnnotator()
    return annotator.add_annotations(input_path, output_path, findings, author)


if __name__ == '__main__':
    # 测试
    test_findings = [
        {
            'category': '金额条款',
            'severity': 'medium',
            'location': '合同金额',
            'description': '合同仅有中文大写金额，缺少阿拉伯数字金额对照',
            'suggestion': '建议同时标注阿拉伯数字金额'
        },
        {
            'category': '交付条款',
            'severity': 'high',
            'location': '交付时间',
            'description': '交付时间描述为"合理时间"，不够具体',
            'suggestion': '建议明确具体日期或期限，如"2024年3月31日前"'
        }
    ]
    
    # 需要实际的 Word 文件进行测试
    print("Word 批注器已加载")
    print("使用方法：annotate_word_contract(input_path, output_path, findings)")

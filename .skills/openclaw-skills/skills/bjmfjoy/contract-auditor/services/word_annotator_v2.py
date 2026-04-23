"""Word 文档批注生成器 V2 - 使用 python-docx 标准 API

在原始 Word 文件上直接添加批注（Comment）
"""

from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import re
from datetime import datetime
from copy import deepcopy


def add_comments_to_word(input_path: str, output_path: str,
                         findings: List[Dict[str, Any]],
                         author: str = "AI合同审计") -> bool:
    """
    在 Word 文档中添加批注
    
    由于 python-docx 不直接支持添加批注，我们采用替代方案：
    1. 在文档开头添加审计摘要
    2. 在相关段落添加高亮和批注标记
    3. 生成带批注说明的修订版文档
    
    Args:
        input_path: 原始合同文件路径
        output_path: 输出文件路径
        findings: 审计发现的问题列表
        author: 批注作者
    
    Returns:
        bool: 是否成功
    """
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_COLOR_INDEX
        
        doc = Document(input_path)
        
        severity_names = {
            'critical': '严重',
            'high': '高',
            'medium': '中',
            'low': '低',
        }
        
        severity_colors = {
            'critical': RGBColor(255, 0, 0),      # 红色
            'high': RGBColor(255, 102, 0),        # 橙色
            'medium': RGBColor(255, 192, 0),      # 黄色
            'low': RGBColor(0, 176, 80),          # 绿色
        }
        
        # 在文档开头添加审计摘要
        summary_heading = doc.add_heading('【合同审计摘要】', level=1)
        
        # 添加风险等级统计
        by_severity = {}
        by_category = {}
        for f in findings:
            sev = f.get('severity', 'low')
            cat = f.get('category', '未分类')
            by_severity[sev] = by_severity.get(sev, 0) + 1
            by_category[cat] = by_category.get(cat, 0) + 1
        
        total = len(findings)
        summary_text = f"共发现问题 {total} 项\n\n"
        
        if by_severity:
            summary_text += "按严重程度："
            for sev, count in sorted(by_severity.items(), 
                                     key=lambda x: {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}.get(x[0], 4)):
                sev_name = severity_names.get(sev, sev.upper())
                summary_text += f"{sev_name} {count} 项；"
            summary_text += "\n\n"
        
        if by_category:
            summary_text += "按类别："
            for cat, count in by_category.items():
                summary_text += f"{cat} {count} 项；"
            summary_text += "\n\n"
        
        summary_para = doc.add_paragraph(summary_text)
        
        # 添加问题清单
        doc.add_heading('【审计发现问题清单】', level=2)
        
        for i, finding in enumerate(findings, 1):
            severity = finding.get('severity', 'low')
            category = finding.get('category', '未分类')
            location = finding.get('location', '')
            description = finding.get('description', '')
            suggestion = finding.get('suggestion', '')
            sev_name = severity_names.get(severity, severity.upper())
            
            # 添加问题项
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(f"【{sev_name}】").bold = True
            run = p.add_run(f"{location} - {category}")
            run.font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))
            
            # 问题描述
            p2 = doc.add_paragraph(f"   问题：{description}", style='List Bullet')
            
            # 修改建议
            if suggestion:
                p3 = doc.add_paragraph(f"   建议：{suggestion}", style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph('=' * 60)
        doc.add_paragraph()
        
        # 在原文中标记问题位置
        doc.add_heading('【带批注的合同正文】', level=2)
        doc.add_paragraph("（以下正文中的彩色高亮部分表示存在审计问题）")
        doc.add_paragraph()
        
        # 复制原文并添加标记
        success_count = 0
        for finding in findings:
            if _mark_finding_in_document(doc, finding, severity_colors):
                success_count += 1
        
        print(f"成功标记 {success_count}/{len(findings)} 处问题位置")
        
        # 保存文档
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"添加批注失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def _mark_finding_in_document(doc, finding: Dict[str, Any], 
                               severity_colors: Dict[str, Any]) -> bool:
    """在文档中标记 finding 的位置"""
    try:
        location = finding.get('location', '')
        description = finding.get('description', '')
        severity = finding.get('severity', 'low')
        
        # 提取关键词
        keywords = _extract_keywords(location + ' ' + description)
        
        # 查找匹配的段落
        best_para = None
        best_match_score = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            score = 0
            for kw in keywords:
                if kw in text:
                    score += len(kw)
            
            if score > best_match_score:
                best_match_score = score
                best_para = para
        
        if best_para and best_match_score > 0:
            # 给段落添加高亮
            for run in best_para.runs:
                run.font.highlight_color = 7  # 黄色高亮
            return True
        
        return False
        
    except Exception as e:
        print(f"标记 finding 失败: {e}")
        return False


def _extract_keywords(text: str) -> List[str]:
    """提取关键词"""
    words = re.findall(r'[\u4e00-\u9fa5]{2,10}', text)
    seen = set()
    result = []
    for w in words:
        if w not in seen and len(w) >= 2:
            seen.add(w)
            result.append(w)
    return result[:15]


# 保持与原接口兼容
def annotate_word_contract(input_path: str, output_path: str,
                           findings: List[Dict[str, Any]],
                           author: str = "AI合同审计") -> bool:
    """便捷函数：为 Word 合同添加批注"""
    return add_comments_to_word(input_path, output_path, findings, author)


if __name__ == '__main__':
    # 测试
    test_findings = [
        {
            'category': '金额条款',
            'severity': 'medium',
            'location': '合同金额',
            'description': '合同仅有中文大写金额，缺少阿拉伯数字金额对照',
            'suggestion': '建议同时标注阿拉伯数字金额'
        },
        {
            'category': '交付条款',
            'severity': 'high',
            'location': '交付时间',
            'description': '交付时间描述为"合理时间"，不够具体',
            'suggestion': '建议明确具体日期或期限'
        }
    ]
    
    print("Word 批注器 V2 已加载")

"""批注文档生成器

在原始 Word 文档上添加审计批注
"""

import sys
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass

sys.path.insert(0, str(Path(__file__).parent))


@dataclass
class Annotation:
    """批注项"""
    paragraph_idx: int  # 段落索引
    text_snippet: str   # 原文片段
    comment: str        # 批注内容
    severity: str       # 严重程度


class DocumentAnnotator:
    """文档批注器"""
    
    def __init__(self):
        self.severity_colors = {
            'critical': 'FF0000',  # 红色
            'high': 'FF6600',      # 橙色
            'medium': 'FFCC00',    # 黄色
            'low': '009900',       # 绿色
        }
    
    def add_annotations(self, input_path: str, output_path: str, findings: List[Dict[str, Any]]) -> bool:
        """
        在文档中添加批注
        
        Args:
            input_path: 原始文档路径
            output_path: 输出文档路径
            findings: 审计发现项列表
        
        Returns:
            是否成功
        """
        try:
            from docx import Document
            from docx.shared import RGBColor
        except ImportError:
            print("错误：缺少 python-docx 库")
            return False
        
        try:
            doc = Document(input_path)
        except Exception as e:
            print(f"错误：无法打开文档 {input_path}: {e}")
            return False
        
        # 为每个发现项添加批注
        annotation_count = 0
        for finding in findings:
            # 在文档中查找相关文本
            paragraph_idx = self._find_paragraph(doc, finding.get('description', ''))
            
            if paragraph_idx >= 0:
                # 构建批注内容
                comment_text = self._build_comment(finding)
                
                # 添加高亮（通过修改文本颜色模拟）
                paragraph = doc.paragraphs[paragraph_idx]
                self._highlight_paragraph(paragraph, finding.get('severity', 'low'))
                
                # 在段落末尾添加批注标记
                # 注意：python-docx 不直接支持批注，我们在段落末尾添加文本标记
                marker = f" [审计:{finding.get('severity', 'low').upper()}]"
                paragraph.add_run(marker).font.color.rgb = RGBColor(
                    *self._hex_to_rgb(self.severity_colors.get(finding.get('severity', 'low'), '000000'))
                )
                
                annotation_count += 1
        
        # 在文档末尾添加审计摘要
        self._add_summary(doc, findings)
        
        # 保存文档
        try:
            doc.save(output_path)
            print(f"已生成批注文档: {output_path}")
            print(f"添加了 {annotation_count} 处标记")
            return True
        except Exception as e:
            print(f"错误：保存文档失败: {e}")
            return False
    
    def _find_paragraph(self, doc, text: str) -> int:
        """查找包含特定文本的段落索引"""
        # 从 finding 的 description 中提取关键词
        # 简化处理：查找包含关键位置的段落
        
        keywords = self._extract_keywords(text)
        
        for idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
            
            # 检查是否包含关键词
            for kw in keywords:
                if kw in para_text:
                    return idx
        
        return -1
    
    def _extract_keywords(self, text: str) -> List[str]:
        """从文本中提取关键词"""
        # 提取中文关键词（2-4个字的词组）
        import re
        words = re.findall(r'[\u4e00-\u9fa5]{2,4}', text)
        # 返回最常见的关键词
        return list(set(words))[:5]
    
    def _build_comment(self, finding: Dict[str, Any]) -> str:
        """构建批注内容"""
        lines = [
            f"【{finding.get('severity', 'UNKNOWN').upper()}】{finding.get('location', '')}",
            f"问题：{finding.get('description', '')}",
            f"建议：{finding.get('suggestion', '')}",
        ]
        return '\n'.join(lines)
    
    def _highlight_paragraph(self, paragraph, severity: str):
        """高亮段落（修改文本颜色）"""
        try:
            from docx.shared import RGBColor
            
            color_hex = self.severity_colors.get(severity, '000000')
            rgb = self._hex_to_rgb(color_hex)
            
            # 为段落中的每个 run 设置颜色
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(*rgb)
        except:
            pass
    
    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """十六进制颜色转 RGB"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def _add_summary(self, doc, findings: List[Dict[str, Any]]):
        """在文档末尾添加审计摘要"""
        from docx.shared import RGBColor, Pt
        
        # 添加分页
        doc.add_page_break()
        
        # 添加标题
        title = doc.add_heading('合同审计摘要', level=1)
        
        # 统计信息
        total = len(findings)
        by_severity = {}
        for f in findings:
            sev = f.get('severity', 'unknown')
            by_severity[sev] = by_severity.get(sev, 0) + 1
        
        # 添加统计
        doc.add_paragraph(f'审计时间：自动生成')
        doc.add_paragraph(f'发现问题总数：{total} 项')
        
        if by_severity:
            p = doc.add_paragraph('严重程度分布：')
            for sev, count in sorted(by_severity.items(), 
                                     key=lambda x: {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}.get(x[0], 4)):
                p.add_run(f'\n  {sev.upper()}: {count} 项')
        
        # 添加问题列表
        doc.add_heading('问题清单', level=2)
        
        severity_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
        sorted_findings = sorted(findings, key=lambda f: severity_order.get(f.get('severity', 'low'), 4))
        
        for i, f in enumerate(sorted_findings, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. 【{f.get('severity', 'UNKNOWN').upper()}】{f.get('location', '')}").bold = True
            doc.add_paragraph(f"   问题：{f.get('description', '')}", style='List Bullet')
            doc.add_paragraph(f"   建议：{f.get('suggestion', '')}", style='List Bullet')


def create_annotated_document(input_path: str, output_path: str, findings: List[Dict[str, Any]]) -> bool:
    """便捷函数：创建带批注的文档"""
    annotator = DocumentAnnotator()
    return annotator.add_annotations(input_path, output_path, findings)


if __name__ == '__main__':
    # 测试
    test_findings = [
        {
            'severity': 'high',
            'location': '签署日期',
            'description': '合同未明确约定签署日期',
            'suggestion': '必须填写合同签署日期'
        },
        {
            'severity': 'medium',
            'location': '金额条款',
            'description': '合同仅有中文大写金额',
            'suggestion': '建议同时标注阿拉伯数字金额'
        }
    ]
    
    # 需要实际的 docx 文件测试
    print("批注文档生成器已加载")
    print("用法: create_annotated_document(input.docx, output_annotated.docx, findings)")

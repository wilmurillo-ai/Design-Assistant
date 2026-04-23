#!/usr/bin/env python3
"""
合同文本提取器
支持 Word、PDF、图片格式
"""

from pathlib import Path
from typing import Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)


class TextExtractor:
    """文本提取器"""
    
    SUPPORTED_FORMATS = {
        '.doc': 'word',
        '.docx': 'word',
        '.pdf': 'pdf',
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image',
    }
    
    def __init__(self):
        self.extractors = {
            'word': self._extract_from_word,
            'pdf': self._extract_from_pdf,
            'image': self._extract_from_image,
        }
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        从文件中提取文本
        
        Returns:
            {
                'success': bool,
                'text': str,           # 提取的文本
                'format': str,         # 文件格式
                'paragraphs': list,    # 段落列表（保留结构）
                'error': str           # 错误信息（如有）
            }
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_FORMATS:
            return {
                'success': False,
                'text': '',
                'format': 'unknown',
                'paragraphs': [],
                'error': f'不支持的文件格式: {ext}'
            }
        
        format_type = self.SUPPORTED_FORMATS[ext]
        extractor = self.extractors.get(format_type)
        
        if not extractor:
            return {
                'success': False,
                'text': '',
                'format': format_type,
                'paragraphs': [],
                'error': f'未实现的提取器: {format_type}'
            }
        
        try:
            return extractor(file_path)
        except Exception as e:
            logger.error(f'提取失败: {file_path}, 错误: {e}')
            return {
                'success': False,
                'text': '',
                'format': format_type,
                'paragraphs': [],
                'error': str(e)
            }
    
    def _extract_from_word(self, file_path: str) -> Dict[str, Any]:
        """从 Word 文档提取文本"""
        try:
            from docx import Document
        except ImportError:
            return {
                'success': False,
                'text': '',
                'format': 'word',
                'paragraphs': [],
                'error': '缺少依赖: python-docx，请运行 pip install python-docx'
            }
        
        try:
            doc = Document(file_path)
            paragraphs = []
            full_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_bold': any(run.bold for run in para.runs),
                        'is_heading': para.style.name.startswith('Heading') if para.style else False
                    })
                    full_text.append(text)
            
            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        tables_text.append(' | '.join(row_text))
            
            if tables_text:
                full_text.append('\n[表格内容]')
                full_text.extend(tables_text)
            
            return {
                'success': True,
                'text': '\n'.join(full_text),
                'format': 'word',
                'paragraphs': paragraphs,
                'tables': tables_text,
                'error': None
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'format': 'word',
                'paragraphs': [],
                'error': f'Word 提取失败: {e}'
            }
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """从 PDF 提取文本（待实现）"""
        return {
            'success': False,
            'text': '',
            'format': 'pdf',
            'paragraphs': [],
            'error': 'PDF 提取暂未实现'
        }
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """从图片提取文本（OCR，待实现）"""
        return {
            'success': False,
            'text': '',
            'format': 'image',
            'paragraphs': [],
            'error': '图片 OCR 暂未实现'
        }


# 便捷函数
def extract_text(file_path: str) -> str:
    """简单提取文本，返回纯文本字符串"""
    extractor = TextExtractor()
    result = extractor.extract(file_path)
    return result.get('text', '') if result.get('success') else ''


if __name__ == '__main__':
    # 测试
    import sys
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        extractor = TextExtractor()
        result = extractor.extract(test_file)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("用法: python text_extractor.py <文件路径>")

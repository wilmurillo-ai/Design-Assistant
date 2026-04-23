#!/usr/bin/env python3
"""
公文格式转换脚本
将多种格式文档内容转换为符合政府公文格式的.docx文档。

支持格式：.md .txt .docx .pdf .html .htm .rtf

用法：
  python3 convert.py <input_file> <output.docx>
"""

import sys
import re
import os
from docx import Document
from docx.shared import Cm, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 中文数字映射 ──────────────────────────────────────────────
CN_NUMBERS = [
    '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
    '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
    '二十一', '二十二', '二十三', '二十四', '二十五', '二十六', '二十七', '二十八', '二十九', '三十',
    '三十一', '三十二', '三十三', '三十四', '三十五', '三十六', '三十七', '三十八', '三十九', '四十',
    '四十一', '四十二', '四十三', '四十四', '四十五', '四十六', '四十七', '四十八', '四十九', '五十',
    '五十一', '五十二', '五十三', '五十四', '五十五', '五十六', '五十七', '五十八', '五十九', '六十',
]


def ar_to_cn(num):
    """将阿拉伯数字转换为中文数字"""
    n = int(num)
    if 1 <= n <= len(CN_NUMBERS):
        return CN_NUMBERS[n - 1]
    return str(n)


# ── Markdown预处理 ─────────────────────────────────────────────
def preprocess_markdown(text):
    """预处理Markdown文本：去除加粗、链接，清理多余空行，去除开头的"-"字符"""
    # 去除加粗标记 **text** 和 __text__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # 去除斜体标记 *text* 和 _text_
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 去除markdown链接 [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # 去除引用标记 >
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    # 去除行首的"-"字符（段落或标题开头不能出现"-"），但保留---分割线
    text = re.sub(r'^(?!-{3,})-\s*', '', text, flags=re.MULTILINE)
    # 清理多余空行（最多保留1个空行）
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 去除markdown表格行（保留内容，去掉|分隔符）
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        # 跳过分隔行 |---|---|
        if re.match(r'^\|[\s\-:]+\|', line):
            continue
        # 表格内容行：去掉首尾|，将中间|替换为空格
        if line.startswith('|') and line.endswith('|'):
            line = line.strip('|').replace('|', '  ').strip()
        cleaned_lines.append(line)
    text = '\n'.join(cleaned_lines)
    return text


def is_h1_pattern(text):
    """判断文本是否符合一级标题模式：一、XXX 或 二、XXX"""
    return bool(re.match(r'^[一二三四五六七八九十百千]+[、．.]\s*', text))


def is_h2_pattern(text):
    """判断文本是否符合二级标题模式：（一）XXX 或 (二)XXX"""
    return bool(re.match(r'^[（(][一二三四五六七八九十\d]+[）)]\s*', text))


def is_h3_pattern(text):
    """判断文本是否符合三级标题模式：（1）XXX 或 (1)XXX"""
    return bool(re.match(r'^[（(]\d+[）)]\s*', text))


def is_likely_title(text, is_first=False):
    """判断文本是否可能是标题（无markdown标记时的推断）
    条件：短行（<=30字）、无句号结尾、有标题特征
    is_first=True时：第一个短行无标点结尾即视为大标题
    """
    if len(text) > 40:
        return False
    if text.endswith(('。', '？', '！', '；', '：', '，', '）', '》')):
        return False
    if is_h1_pattern(text) or is_h2_pattern(text) or is_h3_pattern(text):
        return True
    # 第一个短行、无标点结尾 → 大标题
    if is_first and not text.endswith(('。', '？', '！', '；', '：', '，', '）', '》', '-')):
        return True
    return False


def parse_markdown(text):
    """解析Markdown为结构化元素列表。
    支持两种输入格式：
    1. 标准markdown（带#标记）
    2. 纯文本（自动推断标题层级）
    """
    lines = text.split('\n')
    elements = []

    # 先检测是否使用了markdown标题标记
    has_md_headers = any(re.match(r'^#{1,3}\s+', l) for l in lines)

    i = 0
    has_found_first = False
    in_signature = False
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 跳过水平分割线（--- 且不是最后一个---）
        if re.match(r'^-{3,}$', line) or re.match(r'^\*{3,}$', line):
            # 检查是否是最后的分割线（后面只有短行和空行）
            remaining = lines[i+1:]
            remaining_text = '\n'.join(remaining).strip()
            if remaining_text and len(remaining_text) <= 40 and '\n' in remaining_text:
                # 最后的---，进入落款模式
                in_signature = True
                i += 1
                continue
            elif remaining_text and len(remaining_text) <= 20:
                # 只有一行落款
                elements.append(('signature_unit', remaining_text))
                break
            else:
                # 中间的---，跳过
                i += 1
                continue

        # ── 落款模式 ──
        if in_signature:
            stripped = line.strip()
            if not any(e[0].startswith('signature') for e in elements[-2:] if e):
                # 第一个落款行 → 落款单位
                elements.append(('signature_unit', stripped))
            else:
                # 后续落款行 → 落款日期
                elements.append(('signature_date', stripped))
            i += 1
            continue

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # ── Markdown标题模式 ──
        if has_md_headers:
            # 大标题 (# )
            if re.match(r'^#\s+', line) and not re.match(r'^##', line):
                elements.append(('title', line.lstrip('#').strip()))
                has_found_first = True
                i += 1
                continue

            # 一级标题 (## )
            if re.match(r'^##\s+', line) and not re.match(r'^###', line):
                elements.append(('h1', line.lstrip('#').strip()))
                has_found_first = True
                i += 1
                continue

            # 二级标题 (### )
            if re.match(r'^###\s+', line):
                elements.append(('h2', line.lstrip('#').strip()))
                has_found_first = True
                i += 1
                continue

            # 三级标题 (#### ) — 公文最多三级
            if re.match(r'^####\s+', line):
                elements.append(('h3', line.lstrip('#').strip()))
                has_found_first = True
                i += 1
                continue

        # ── 纯文本标题推断模式 ──
        stripped = line.strip()

        # 一级标题推断："一、XXX" 或 "二、XXX" 格式
        if is_h1_pattern(stripped):
            elements.append(('h1', stripped))
            has_found_first = True
            i += 1
            continue

        # 二级标题推断："（一）XXX" 格式
        if is_h2_pattern(stripped):
            # 如果标题后有冒号+正文，拆分为标题+正文
            if '：' in stripped:
                parts = stripped.split('：', 1)
                elements.append(('h2', parts[0].strip()))
                elements.append(('para', parts[1].strip()))
            else:
                elements.append(('h2', stripped))
            has_found_first = True
            i += 1
            continue

        # 三级标题推断："（1）XXX" 格式
        if is_h3_pattern(stripped):
            if '：' in stripped:
                parts = stripped.split('：', 1)
                elements.append(('h3', parts[0].strip()))
                elements.append(('para', parts[1].strip()))
            else:
                elements.append(('h3', stripped))
            has_found_first = True
            i += 1
            continue

        # 副标题推断（——开头）
        if stripped.startswith('\u2014\u2014') and not has_found_first:
            has_found_first = True
            i += 1
            continue

        # 大标题推断：第一个短行、无句号结尾、不在标题序列中
        if not has_found_first and is_likely_title(stripped, is_first=True):
            elements.append(('title', stripped))
            has_found_first = True
            i += 1
            continue

        # ── 普通段落 ──
        para_lines = []
        while i < len(lines):
            l = lines[i].rstrip()
            if not l.strip():
                # 检查空行后是否紧跟标题模式（如果是则停止收集）
                break
            if re.match(r'^-{3,}$', l) or re.match(r'^\*{3,}$', l):
                break
            # markdown标题
            if has_md_headers and re.match(r'^#{1,4}\s+', l):
                break
            # 纯文本标题模式
            s = l.strip()
            if is_h1_pattern(s) or is_h2_pattern(s) or is_h3_pattern(s):
                break
            para_lines.append(s)
            i += 1
        if para_lines:
            elements.append(('para', ' '.join(para_lines)))

    # ── 落款识别：将最后的连续短段落识别为落款 ──
    # 如果最后2个元素是短段落（<20字），且前面有标题/正文，视为落款
    if len(elements) >= 2:
        last = elements[-1]
        second_last = elements[-2]
        if (last[0] == 'para' and second_last[0] == 'para'
            and len(last[1]) <= 20 and len(second_last[1]) <= 20
            and not last[1].endswith('。')
            and not second_last[1].endswith('。')):
            elements[-2] = ('signature_unit', second_last[1])
            elements[-1] = ('signature_date', last[1])

    return elements


# ── 转换二级标题序号 ──────────────────────────────────────────
def convert_h2_numbering(elements):
    """将二级标题的序号从阿拉伯数字转换为中文数字，每个一级标题下独立编号"""
    result = []
    h2_count = 0
    for kind, text in elements:
        if kind == 'h1':
            h2_count = 0
            result.append((kind, text))
        elif kind == 'h2':
            h2_count += 1
            # 移除原有的序号前缀（如 "（一）"、"(1)"、"1." 等）
            cleaned = re.sub(r'^[（(][一二三四五六七八九十\d]+[）)]\s*', '', text)
            cleaned = re.sub(r'^\d+[.、]\s*', '', cleaned)
            cleaned = re.sub(r'^[一二三四五六七八九十]+[.、]\s*', '', cleaned)
            cn = ar_to_cn(h2_count)
            result.append(('h2', f'（{cn}）{cleaned}'))
        else:
            result.append((kind, text))
    return result


# ── 判断是否需要Times New Roman ─────────────────────────────────
def needs_tnr(char):
    """判断字符是否需要Times New Roman字体（半角英文、数字、半角标点）"""
    return ('\u0020' <= char <= '\u007E')  # ASCII可打印字符


# ── 创建混合字体Run ─────────────────────────────────────────────
def add_mixed_runs(paragraph, text, base_font='仿宋_GB2312', base_size=Pt(16)):
    """为段落添加带有混合字体的run（中文用基础字体，英文数字用Times New Roman）"""
    if not text:
        return

    i = 0
    while i < len(text):
        ch = text[i]
        if needs_tnr(ch):
            # 收集连续的英文/数字字符
            chunk = ''
            while i < len(text) and needs_tnr(text[i]):
                chunk += text[i]
                i += 1
            run = paragraph.add_run(chunk)
            run.font.name = 'Times New Roman'
            run.font.size = base_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font)
            run.bold = False
        else:
            # 收集连续的中文字符
            chunk = ''
            while i < len(text) and not needs_tnr(text[i]):
                chunk += text[i]
                i += 1
            run = paragraph.add_run(chunk)
            run.font.name = base_font
            run.font.size = base_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font)
            run.bold = False


# ── 设置行距 ────────────────────────────────────────────────────
def set_line_spacing(paragraph, line_spacing_pt=28):
    """设置段落行距（固定磅值）"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)


# ── 设置首行缩进 ────────────────────────────────────────────────
def set_first_line_indent(paragraph, indent_cm=1.127):
    """设置首行缩进"""
    paragraph.paragraph_format.first_line_indent = Cm(indent_cm)


# ── 设置段前段后间距 ────────────────────────────────────────────
def set_spacing_zero(paragraph):
    """设置段前段后间距为0"""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


# ── 设置页码 ────────────────────────────────────────────────────
def _add_cmn_run(p, text, font='宋体', size=14):
    """添加一个宋体run"""
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return run

def _add_field_char(parent, ftype):
    """添加 Word complex field 控制符（begin/separate/end）"""
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:rPr><w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/><w:sz w:val="28"/></w:rPr>'
        '<w:fldChar w:fldCharType="%s"/></w:r>' % ftype
    )
    parent._element.append(parse_xml(xml))

def _add_instr(parent, instr):
    """添加 Word field instruction（如 PAGE）"""
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:rPr><w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/><w:sz w:val="28"/></w:rPr>'
        '<w:instrText xml:space="preserve">%s</w:instrText></w:r>' % instr
    )
    parent._element.append(parse_xml(xml))

def _build_page_field(p):
    """为一个段落构建 -PAGE- 页码结构（complex field）"""
    _add_cmn_run(p, '-')
    _add_field_char(p, 'begin')
    _add_instr(p, ' PAGE ')
    _add_field_char(p, 'separate')
    _add_cmn_run(p, '1')  # placeholder，Word 会自动计算实际页码
    _add_field_char(p, 'end')
    _add_cmn_run(p, '-')

def set_page_number(section):
    """设置页码格式：奇数页右下、偶数页左下，格式 -数字-

    关键：必须在 sectPr 中添加 <w:evenAndOddHeaders/> 启用奇偶页不同。
    使用 complex field（fldChar+instrText）而非 fldSimple，兼容性更好。
    """
    section.different_first_page_header_footer = False

    # 在 sectPr 中添加 <w:evenAndOddHeaders/> —— Word 识别奇偶页不同的必要标签
    sect_pr = section._sectPr
    if sect_pr.find(qn('w:evenAndOddHeaders')) is None:
        sect_pr.append(parse_xml(
            '<w:evenAndOddHeaders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        ))

    # 奇数页页脚（右下）
    footer_odd = section.footer
    footer_odd.is_linked_to_previous = False
    for p in footer_odd.paragraphs:
        p._element.getparent().remove(p._element)
    p_odd = footer_odd.add_paragraph()
    p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_odd.paragraph_format.space_before = Pt(0)
    p_odd.paragraph_format.space_after = Pt(0)
    _build_page_field(p_odd)

    # 偶数页页脚（左下）
    footer_even = section.even_page_footer
    footer_even.is_linked_to_previous = False
    for p in footer_even.paragraphs:
        p._element.getparent().remove(p._element)
    p_even = footer_even.add_paragraph()
    p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_even.paragraph_format.space_before = Pt(0)
    p_even.paragraph_format.space_after = Pt(0)
    _build_page_field(p_even)


# ── 生成DOCX ────────────────────────────────────────────────────
def generate_docx(elements, output_path):
    """根据解析的元素列表生成符合公文格式的DOCX文档"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 页码
    set_page_number(section)

    # ── 遍历元素生成段落 ──
    prev_kind = None
    for kind, text in elements:

        # 大标题
        if kind == 'title':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_first_line_indent(p, 0)
            set_line_spacing(p, 32)
            set_spacing_zero(p)
            add_mixed_runs(p, text, base_font='方正小标宋简体', base_size=Pt(22))
            # 大标题后空一行
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            set_line_spacing(spacer, 28)
            prev_kind = kind
            continue

        # 一级标题
        if kind == 'h1':
            # h1本身是标题，紧跟即可，不需要额外空行
            # （"标题与正文之间空一行"由title分支的spacer处理）
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_first_line_indent(p, 1.127)
            set_line_spacing(p, 28)
            set_spacing_zero(p)
            add_mixed_runs(p, text, base_font='黑体', base_size=Pt(16))
            prev_kind = kind
            continue

        # 二级标题
        if kind == 'h2':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_first_line_indent(p, 1.127)
            set_line_spacing(p, 28)
            set_spacing_zero(p)
            add_mixed_runs(p, text, base_font='楷体_GB2312', base_size=Pt(16))
            prev_kind = kind
            continue

        # 三级标题
        if kind == 'h3':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_first_line_indent(p, 1.127)
            set_line_spacing(p, 28)
            set_spacing_zero(p)
            add_mixed_runs(p, text, base_font='仿宋_GB2312', base_size=Pt(16))
            prev_kind = kind
            continue

        # 正文段落
        if kind == 'para':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_first_line_indent(p, 1.127)
            set_line_spacing(p, 28)
            set_spacing_zero(p)
            # 处理全角引号
            text = text.replace('"', '\u201c').replace('"', '\u201d')
            # 确保段落以句号结尾（如果不是以标点结尾）
            if text and text[-1] not in '。？！；：、）】》""':
                text += '。'
            add_mixed_runs(p, text, base_font='仿宋_GB2312', base_size=Pt(16))
            prev_kind = kind
            continue

        # 落款单位
        if kind == 'signature_unit':
            # 正文与落款之间空一行
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            set_line_spacing(spacer, 28)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_first_line_indent(p, 0)
            set_line_spacing(p, 28)
            set_spacing_zero(p)
            add_mixed_runs(p, '\u3000\u3000' + text, base_font='仿宋_GB2312', base_size=Pt(16))
            prev_kind = kind
            continue

        # 落款日期
        if kind == 'signature_date':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_first_line_indent(p, 0)
            set_line_spacing(p, 28)
            set_spacing_zero(p)
            add_mixed_runs(p, '\u3000\u3000\u3000\u3000' + text, base_font='仿宋_GB2312', base_size=Pt(16))
            prev_kind = kind
            continue

    doc.save(output_path)

    # ── 关键修复：将 evenAndOddHeaders 写入 settings.xml ──
    # 仅在 sectPr 中添加不够，Word 需要在文档级 settings.xml 中也有此设置
    # 才能正确识别奇偶页页脚差异
    _patch_settings_even_odd(output_path)

    print(f'✅ 文档已生成：{output_path}')


def _patch_settings_even_odd(docx_path):
    """将 <w:evenAndOddHeaders/> 注入 word/settings.xml"""
    import zipfile
    import tempfile

    settings_tag = '<w:evenAndOddHeaders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'

    # 读取原文件，修改 settings.xml，写回
    tmp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w') as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == 'word/settings.xml':
                    settings = data.decode('utf-8')
                    if 'evenAndOddHeaders' not in settings:
                        # 在 </w:settings> 前插入
                        settings = settings.replace('</w:settings>', settings_tag + '</w:settings>')
                        data = settings.encode('utf-8')
                zout.writestr(item, data)

    # 替换原文件
    import os
    os.replace(tmp_path, docx_path)


# ── 读取DOCX文件 ────────────────────────────────────────────────
def read_docx_to_elements(input_path):
    """读取.docx文件并解析为结构化元素列表

    修复：第一个短段落（≤40字）自动识别为title，去掉末尾多余句号。
    原因：许多公文标题本身带句号，但公文格式中标题不应有句号。
    """
    doc = Document(input_path)
    elements = []
    first_non_heading = True

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # ── 根据文本内容判断类型 ──
        # 以"一、""二、"等中文数字+顿号开头 → h1
        if re.match(r'^[一二三四五六七八九十百千万]+、', text):
            elements.append(('h1', text))
            first_non_heading = False
            continue

        # 以"（一）""（二）"等全角括号+数字开头 → h2
        if re.match(r'^[（(][一二三四五六七八九十百千万\d]+[）)]', text):
            elements.append(('h2', text))
            first_non_heading = False
            continue

        # ── 根据段落样式判断类型 ──
        style_name = para.style.name if para.style else ''

        # 尝试获取字号
        font_size_pt = None
        for run in para.runs:
            if run.font.size:
                font_size_pt = run.font.size.pt
                break

        is_bold = any(run.bold for run in para.runs if run.font.bold is not None)

        # 判断是否为标题（样式名包含"Title"或字号>=20pt）
        if 'Title' in style_name or (font_size_pt and font_size_pt >= 20):
            # 去掉标题末尾多余句号
            title_text = text.rstrip('。')
            elements.append(('title', title_text))
            first_non_heading = False
            continue

        # 第一个非标题短段落 → 识别为title（如"XXX学习心得体会"）
        if first_non_heading and len(text) <= 40:
            title_text = text.rstrip('。')
            elements.append(('title', title_text))
            first_non_heading = False
            continue

        # 其他 → 段落
        elements.append(('para', text))
        first_non_heading = False

    return elements


# ── 读取TXT文件 ────────────────────────────────────────────────
def read_txt_to_elements(input_path):
    """读取纯文本文件并解析为结构化元素列表"""
    elements = []

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    para_lines = []
    for line in lines:
        line = line.rstrip('\n\r')

        # 跳过空行：flush当前段落
        if not line.strip():
            if para_lines:
                text = ' '.join(para_lines).strip()
                kind = _classify_text_line(text)
                elements.append((kind, text))
                para_lines = []
            continue

        stripped = line.strip()

        # 尝试按行判断标题类型
        kind = _classify_text_line(stripped)

        if kind == 'title' or kind == 'h1':
            # flush之前的段落
            if para_lines:
                text = ' '.join(para_lines).strip()
                elements.append((_classify_text_line(text), text))
                para_lines = []
            elements.append((kind, stripped))
        elif kind == 'h2':
            # flush之前的段落
            if para_lines:
                text = ' '.join(para_lines).strip()
                elements.append(('para', text))
                para_lines = []
            elements.append(('h2', stripped))
        else:
            para_lines.append(stripped)

    # flush最后一段
    if para_lines:
        text = ' '.join(para_lines).strip()
        elements.append(('para', text))

    return elements


def _classify_text_line(text):
    """根据文本模式判断元素类型"""
    if not text:
        return 'para'

    # 一、二、三、... → h1（中文数字+顿号）
    if re.match(r'^[一二三四五六七八九十百千万]+、', text):
        return 'h1'

    # （一）（二）... → h2（全角括号+中文数字）
    if re.match(r'^[（(][一二三四五六七八九十百千万]+[）)]', text):
        return 'h2'

    # （1）（2）... → h2（全角括号+数字）
    if re.match(r'^[（(]\d+[）)]', text):
        return 'h2'

    # 短文本（<30字）且无句号结尾 → 可能是标题
    if len(text) < 30 and text[-1] not in '。？！；：、':
        # 如果以"报告""方案""意见""通知"等结尾，当作标题
        if re.search(r'(报告|方案|意见|通知|总结|计划|调研|分析|报告)$', text):
            return 'title'

    return 'para'


# ── 读取PDF文件 ────────────────────────────────────────────────
def read_pdf_to_elements(input_path):
    """读取PDF文件并解析为结构化元素列表"""
    from pdfminer.high_level import extract_text

    raw_text = extract_text(input_path)

    # 按段落分割（双换行优先，否则按单换行）
    paragraphs = re.split(r'\n\s*\n', raw_text)

    elements = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 去除PDF提取中的多余换行
        para = re.sub(r'\n', ' ', para)
        para = re.sub(r'\s+', ' ', para).strip()
        if not para:
            continue

        kind = _classify_text_line(para)
        elements.append((kind, para))

    return elements


# ── 读取HTML文件 ────────────────────────────────────────────────
def read_html_to_elements(input_path):
    """读取HTML文件并解析为结构化元素列表"""
    from bs4 import BeautifulSoup

    with open(input_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')

    elements = []
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'li', 'article', 'section']):
        text = tag.get_text(separator=' ', strip=True)
        if not text:
            continue

        tag_name = tag.name.lower()

        if tag_name == 'h1':
            elements.append(('title', text))
        elif tag_name == 'h2':
            elements.append(('h1', text))
        elif tag_name == 'h3':
            elements.append(('h2', text))
        elif tag_name in ('h4', 'h5', 'h6'):
            elements.append(('h2', text))
        else:
            # p, div, li, article, section 等内容标签
            # 跳过已经被标题标签包裹的内容
            # 检查是否已经是某个已处理元素的子内容
            if tag.name == 'p' or tag.name == 'li':
                elements.append(('para', text))
            elif tag.name in ('div', 'article', 'section'):
                # 仅处理没有嵌套标题/段落的div（避免重复）
                if not tag.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
                    elements.append(('para', text))

    # 去重：按内容去重（保留顺序）
    seen = set()
    unique_elements = []
    for kind, text in elements:
        key = text[:100]  # 用前100字符作为去重key
        if key not in seen:
            seen.add(key)
            unique_elements.append((kind, text))

    return unique_elements


# ── 读取RTF文件 ────────────────────────────────────────────────
def read_rtf_to_elements(input_path):
    """读取RTF文件并解析为结构化元素列表"""
    from striprtf.striprtf import rtf_to_text

    with open(input_path, 'r', encoding='utf-8') as f:
        rtf_content = f.read()

    raw_text = rtf_to_text(rtf_content)

    # 按段落分割
    paragraphs = re.split(r'\n\s*\n', raw_text)
    if len(paragraphs) == 1:
        # 没有双换行，按单换行分割
        paragraphs = raw_text.split('\n')

    elements = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        para = re.sub(r'\s+', ' ', para).strip()
        if not para:
            continue

        kind = _classify_text_line(para)
        elements.append((kind, para))

    return elements


# ── 主入口 ──────────────────────────────────────────────────────
def main():
    if len(sys.argv) != 3:
        print('用法：python3 convert.py <输入文件> <输出.docx>')
        print('支持格式：.md .txt .docx .pdf .html .htm .rtf')
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(input_path):
        print(f'❌ 文件不存在：{input_path}')
        sys.exit(1)

    input_ext = os.path.splitext(input_path)[1].lower()

    if input_ext == '.md':
        with open(input_path, 'r', encoding='utf-8') as f:
            raw_text = f.read()
        text = preprocess_markdown(raw_text)
        elements = parse_markdown(text)
    elif input_ext == '.txt':
        elements = read_txt_to_elements(input_path)
    elif input_ext == '.docx':
        elements = read_docx_to_elements(input_path)
    elif input_ext == '.pdf':
        elements = read_pdf_to_elements(input_path)
    elif input_ext in ('.html', '.htm'):
        elements = read_html_to_elements(input_path)
    elif input_ext == '.rtf':
        elements = read_rtf_to_elements(input_path)
    else:
        print(f'❌ 不支持的文件格式：{input_ext}')
        print('支持格式：.md .txt .docx .pdf .html .htm .rtf')
        sys.exit(1)

    # 转换二级标题序号
    elements = convert_h2_numbering(elements)

    # 调试输出
    print(f'📄 解析到 {len(elements)} 个元素：')
    for kind, text in elements:
        preview = text[:60] + '...' if len(text) > 60 else text
        print(f'  [{kind}] {preview}')

    # 生成DOCX
    generate_docx(elements, output_path)


if __name__ == '__main__':
    main()

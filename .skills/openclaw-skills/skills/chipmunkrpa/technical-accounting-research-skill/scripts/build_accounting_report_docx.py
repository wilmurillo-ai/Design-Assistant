#!/usr/bin/env python
"""Build a technical accounting report DOCX from JSON input."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: python -m pip install --user python-docx"
    ) from exc


def _as_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _as_list_of_str(value: Any) -> list[str]:
    if isinstance(value, list):
        return [_as_text(item) for item in value if _as_text(item)]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _as_list_of_dict(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    out: list[dict[str, Any]] = []
    for item in value:
        if isinstance(item, dict):
            out.append(item)
    return out


def _strip_markdown_prefixes(text: str) -> str:
    cleaned = _as_text(text)
    if not cleaned:
        return ""
    cleaned = re.sub(r"^#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^[-*+]\s+", "", cleaned)
    cleaned = re.sub(r"^\d+\.\s+", "", cleaned)
    return cleaned.strip()


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    # Use Table Grid but clear borders for metadata look
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = f"{label}:"
        table.cell(index, 1).text = value


def _add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"


def _add_rich_text_block(doc: Document, text: str) -> None:
    cleaned = _as_text(text)
    if not cleaned:
        return

    for raw_line in cleaned.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading_match = re.match(r"^(#{1,6})\s*(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            _add_heading(doc, _strip_markdown_prefixes(heading_match.group(2)), level=level)
            continue

        bullet_match = re.match(r"^[-*+]\s+(.+)$", line)
        if bullet_match:
            doc.add_paragraph(_strip_markdown_prefixes(bullet_match.group(1)), style="List Bullet")
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            doc.add_paragraph(_strip_markdown_prefixes(number_match.group(1)), style="List Number")
            continue

        doc.add_paragraph(line)


def _add_heading_and_text(doc: Document, heading: str, text: str) -> None:
    if not text:
        return
    _add_heading(doc, heading, level=1)
    _add_rich_text_block(doc, text)


def _add_heading_and_paragraphs(doc: Document, heading: str, items: list[str]) -> None:
    if not items:
        return
    _add_heading(doc, heading, level=1)
    for item in items:
        cleaned_item = _strip_markdown_prefixes(item)
        if cleaned_item:
            doc.add_paragraph(cleaned_item, style="List Bullet")


def _sanitize_text(text: str) -> str:
    """Basic cleanup of markdown symbols to avoid literal rendering."""
    s = text.replace("####", "").replace("###", "").replace("##", "").replace("#", "")
    s = s.replace("***", "").replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    return s.strip()


def _add_guidance_table(doc: Document, guidance: list[dict[str, Any]]) -> None:
    if not guidance:
        return

    _add_heading(doc, "Guidance Reviewed", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Citation"
    header[1].text = "Type"
    header[2].text = "Key Point"
    header[3].text = "URL"
    header[4].text = "Accessed"

    for item in guidance:
        row = table.add_row().cells
        row[0].text = _sanitize_text(_as_text(item.get("citation")))
        row[1].text = _sanitize_text(_as_text(item.get("source_type")))
        row[2].text = _sanitize_text(_as_text(item.get("key_point")))
        row[3].text = _as_text(item.get("url"))
        row[4].text = _as_text(item.get("accessed_on"))


def _add_journal_entries_table(doc: Document, entries: list[dict[str, Any]]) -> None:
    if not entries:
        return

    _add_heading(doc, "Journal Entry Examples", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Description"
    header[1].text = "Debit"
    header[2].text = "Credit"
    header[3].text = "Amount"

    for item in entries:
        row = table.add_row().cells
        row[0].text = _sanitize_text(_as_text(item.get("description")))
        row[1].text = _sanitize_text(_as_text(item.get("debit")))
        row[2].text = _sanitize_text(_as_text(item.get("credit")))
        row[3].text = _sanitize_text(_as_text(item.get("amount")))


def _add_qa(doc: Document, qa_items: list[dict[str, Any]]) -> None:
    if not qa_items:
        return

    _add_heading(doc, "Questions and Answers", level=1)
    for item in qa_items:
        question = _as_text(item.get("question"))
        answer = _as_text(item.get("answer"))
        if question:
            p = doc.add_paragraph()
            p.add_run("Question: ").bold = True
            p.add_run(_sanitize_text(question))
        if answer:
            p = doc.add_paragraph()
            p.add_run("Answer: ").bold = True
            p.add_run(_sanitize_text(answer))


def _build_document(payload: dict[str, Any], report_format: str) -> Document:
    doc = Document()
    _set_default_style(doc)

    default_title_map = {
        "memo": "Technical Accounting Memorandum",
        "email": "Technical Accounting Email Draft",
        "q-and-a": "Technical Accounting Questions and Answers",
    }
    title = _as_text(payload.get("title"), default_title_map[report_format])
    _add_heading(doc, title, level=0)

    report_date = _as_text(payload.get("date"), date.today().isoformat())

    if report_format == "memo":
        _add_metadata_table(
            doc,
            [
                ("To", _as_text(payload.get("prepared_for"))),
                ("From", _as_text(payload.get("prepared_by"))),
                ("Date", report_date),
                ("Subject", _as_text(payload.get("subject"))),
            ],
        )
    elif report_format == "email":
        _add_metadata_table(
            doc,
            [
                ("To", _as_text(payload.get("to"), _as_text(payload.get("prepared_for")))),
                ("Cc", _as_text(payload.get("cc"))),
                ("From", _as_text(payload.get("from"), _as_text(payload.get("prepared_by")))),
                ("Date", report_date),
                ("Subject", _as_text(payload.get("subject"))),
            ],
        )
    else:
        _add_metadata_table(
            doc,
            [
                ("Prepared For", _as_text(payload.get("prepared_for"))),
                ("Prepared By", _as_text(payload.get("prepared_by"))),
                ("Date", report_date),
                ("Topic", _as_text(payload.get("subject"))),
            ],
        )

    _add_heading_and_text(doc, "Issue", _as_text(payload.get("issue")))
    _add_heading_and_paragraphs(doc, "Facts", _as_list_of_str(payload.get("facts")))

    guidance = _as_list_of_dict(payload.get("guidance"))
    _add_guidance_table(doc, guidance)

    _add_heading_and_paragraphs(doc, "Analysis", _as_list_of_str(payload.get("analysis")))
    _add_heading_and_text(doc, "Conclusion", _as_text(payload.get("conclusion")))

    _add_heading_and_paragraphs(
        doc,
        "Disclosure Considerations",
        _as_list_of_str(payload.get("disclosure_considerations")),
    )

    _add_journal_entries_table(doc, _as_list_of_dict(payload.get("journal_entries")))

    _add_heading_and_paragraphs(doc, "Assumptions", _as_list_of_str(payload.get("assumptions")))
    _add_heading_and_paragraphs(doc, "Open Items", _as_list_of_str(payload.get("open_items")))
    _add_heading_and_paragraphs(doc, "Next Steps", _as_list_of_str(payload.get("next_steps")))

    _add_qa(doc, _as_list_of_dict(payload.get("qa")))

    if guidance:
        _add_heading(doc, "Source List", level=1)
        for source in guidance:
            citation = _as_text(source.get("citation"), "Unlabeled source")
            source_type = _as_text(source.get("source_type"), "unknown type")
            accessed_on = _as_text(source.get("accessed_on"), "unknown date")
            url = _as_text(source.get("url"), "(no URL)")
            doc.add_paragraph(
                f"{citation} ({source_type}), accessed {accessed_on}: {url}"
            )

    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a technical accounting DOCX report from JSON input."
    )
    parser.add_argument("--input-json", required=True, help="Path to JSON input payload")
    parser.add_argument("--output-docx", required=True, help="Path for generated DOCX")
    parser.add_argument(
        "--format",
        default="memo",
        choices=["memo", "email", "q-and-a"],
        help="Output report style",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input_json)
    output_path = Path(args.output_docx)

    with input_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)

    if not isinstance(payload, dict):
        raise SystemExit("Input JSON must contain an object at top level.")

    doc = _build_document(payload, args.format)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"[OK] Generated {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

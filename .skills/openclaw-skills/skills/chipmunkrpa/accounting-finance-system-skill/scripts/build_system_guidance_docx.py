#!/usr/bin/env python3
"""Build a DOCX report for accounting and finance system guidance."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
except ImportError:
    print("python-docx is required. Install with: python -m pip install --user python-docx", file=sys.stderr)
    raise SystemExit(1)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create a system-guidance DOCX from structured JSON input."
    )
    parser.add_argument("--input-json", required=True, help="Path to report payload JSON")
    parser.add_argument("--output-docx", required=True, help="Path to output .docx")
    parser.add_argument(
        "--format",
        choices=["memo", "q-and-a"],
        help="Output format. If omitted, falls back to payload format or memo.",
    )
    return parser.parse_args()


def to_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def load_payload(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ValueError(f"Input JSON not found: {path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in {path}: {exc}") from exc

    if not isinstance(data, dict):
        raise ValueError("Input JSON must be an object at the top level")
    return data


def validate_payload(payload: dict[str, Any]) -> None:
    errors: list[str] = []

    if not to_text(payload.get("question")):
        errors.append("Missing required field: question")

    sources = payload.get("sources")
    if not isinstance(sources, list) or not sources:
        errors.append("Field 'sources' must be a non-empty array")

    steps = payload.get("recommended_steps")
    if not isinstance(steps, list) or not steps:
        errors.append("Field 'recommended_steps' must be a non-empty array")

    if errors:
        raise ValueError("\n".join(errors))


def normalize_format(cli_format: str | None, payload: dict[str, Any]) -> str:
    raw = to_text(cli_format or payload.get("format") or "memo").lower()
    aliases = {
        "memo": "memo",
        "quick memo": "memo",
        "quick-memo": "memo",
        "q-and-a": "q-and-a",
        "q&a": "q-and-a",
        "qa": "q-and-a",
        "simple q-and-a": "q-and-a",
        "simple q&a": "q-and-a",
    }
    normalized = aliases.get(raw)
    if normalized:
        return normalized
    raise ValueError(f"Unsupported format '{raw}'. Use memo or q-and-a.")


def add_key_value_bullets(doc: Document, values: dict[str, Any]) -> None:
    for key, value in values.items():
        text = to_text(value)
        if not text:
            continue
        paragraph = doc.add_paragraph(style="List Bullet")
        prefix = key.replace("_", " ").strip().title()
        paragraph.add_run(f"{prefix}: ").bold = True
        paragraph.add_run(text)


def add_string_list(doc: Document, title: str, items: Any) -> None:
    if not isinstance(items, list) or not items:
        return
    doc.add_heading(title, level=1)
    count = 0
    for item in items:
        text = to_text(item)
        if not text:
            continue
        count += 1
        doc.add_paragraph(text, style="List Bullet")
    if count == 0:
        doc.add_paragraph("None.")


def add_clarifications(doc: Document, clarifications: Any) -> None:
    if not isinstance(clarifications, list) or not clarifications:
        return

    doc.add_heading("Clarifications", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Question"
    headers[1].text = "Answer"

    for item in clarifications:
        if isinstance(item, dict):
            question = to_text(item.get("question"), "Clarification")
            answer = to_text(item.get("answer"), "")
        else:
            question = "Clarification"
            answer = to_text(item)

        row = table.add_row().cells
        row[0].text = question
        row[1].text = answer


def normalize_step(step: Any) -> tuple[str, str, list[str]]:
    if isinstance(step, str):
        return to_text(step), "", []

    if isinstance(step, dict):
        action = to_text(
            step.get("step")
            or step.get("action")
            or step.get("instruction")
            or step.get("title")
        )
        rationale = to_text(step.get("rationale"))

        refs = step.get("source_refs")
        ref_values: list[str] = []
        if isinstance(refs, list):
            ref_values = [to_text(r) for r in refs if to_text(r)]
        elif refs is not None:
            single = to_text(refs)
            if single:
                ref_values = [single]

        return action, rationale, ref_values

    return to_text(step), "", []


def add_steps(doc: Document, title: str, steps: Any) -> None:
    doc.add_heading(title, level=1)

    if not isinstance(steps, list) or not steps:
        doc.add_paragraph("No steps provided.")
        return

    count = 0
    for raw_step in steps:
        action, rationale, refs = normalize_step(raw_step)
        if not action:
            continue

        count += 1
        doc.add_paragraph(action, style="List Number")

        if rationale:
            rationale_line = doc.add_paragraph()
            rationale_line.add_run("Rationale: ").bold = True
            rationale_line.add_run(rationale)

        if refs:
            refs_line = doc.add_paragraph()
            refs_line.add_run("Sources: ").bold = True
            refs_line.add_run(", ".join(refs))

    if count == 0:
        doc.add_paragraph("No valid steps provided.")


def add_sources(doc: Document, sources: Any) -> None:
    doc.add_heading("Sources", level=1)

    if not isinstance(sources, list) or not sources:
        doc.add_paragraph("No sources listed.")
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["#", "Title", "Publisher", "URL", "Published/Updated", "Accessed"]
    for index, label in enumerate(headers):
        table.rows[0].cells[index].text = label

    for index, source in enumerate(sources, start=1):
        if not isinstance(source, dict):
            source = {"title": to_text(source)}

        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = to_text(source.get("title"))
        row[2].text = to_text(source.get("publisher"))
        row[3].text = to_text(source.get("url"))
        row[4].text = to_text(source.get("published_or_updated"))
        row[5].text = to_text(source.get("accessed"))


def build_report(payload: dict[str, Any], format_name: str) -> Document:
    document = Document()

    title = to_text(payload.get("title"), "Accounting And Finance System Guidance")
    document.add_heading(title, level=0)

    metadata_parts: list[str] = []
    for label, key in (
        ("Prepared For", "prepared_for"),
        ("Prepared By", "prepared_by"),
        ("Date", "date"),
    ):
        value = to_text(payload.get(key))
        if value:
            metadata_parts.append(f"{label}: {value}")

    if metadata_parts:
        document.add_paragraph(" | ".join(metadata_parts))

    question = to_text(payload.get("question"))
    document.add_heading("Request", level=1)
    document.add_paragraph(question)

    context = payload.get("system_context")
    if isinstance(context, dict) and context:
        document.add_heading("System Context", level=1)
        add_key_value_bullets(document, context)

    add_clarifications(document, payload.get("clarifications"))
    add_string_list(document, "Assumptions", payload.get("assumptions"))

    summary = to_text(payload.get("analysis_summary") or payload.get("short_answer"))

    if format_name == "memo":
        document.add_heading("Recommendation Summary", level=1)
        document.add_paragraph(summary or "No summary provided.")
        add_steps(document, "Recommended Steps", payload.get("recommended_steps"))
    else:
        document.add_heading("Question", level=1)
        document.add_paragraph(question)
        document.add_heading("Answer", level=1)
        document.add_paragraph(summary or "No direct answer provided.")
        add_steps(document, "Actions", payload.get("recommended_steps"))

    add_string_list(document, "Validation Checks", payload.get("validation_checks"))
    add_string_list(document, "Risks And Open Items", payload.get("risks_open_items"))
    add_sources(document, payload.get("sources"))
    add_string_list(document, "Open Questions", payload.get("open_questions"))

    return document


def main() -> int:
    args = parse_args()

    input_path = Path(args.input_json)
    output_path = Path(args.output_docx)

    try:
        payload = load_payload(input_path)
        validate_payload(payload)
        format_name = normalize_format(args.format, payload)
        document = build_report(payload, format_name)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"DOCX report written: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

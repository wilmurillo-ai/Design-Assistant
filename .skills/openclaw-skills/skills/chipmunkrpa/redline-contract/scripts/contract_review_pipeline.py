#!/usr/bin/env python3
"""
Legal contract review helper for DOCX workflows.

Commands:
- init-review: extract paragraph-level review units from a DOCX contract into a review JSON template.
- materialize: generate (1) amended DOCX with tracked revisions and (2) risk analysis DOCX.
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import re
import zipfile
from pathlib import Path
from typing import Any

try:
    from lxml import etree
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: lxml. Install with `python -m pip install --user python-docx`."
    ) from exc

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS}
W = f"{{{W_NS}}}"
XML_SPACE = f"{{{XML_NS}}}space"

HEADING_NUMBER_RE = re.compile(
    r"^(?:(?:Section|Article)\s+)?(\d+(?:\.\d+)*[A-Za-z]?)[\)\.\-:]?\s+.+$",
    re.IGNORECASE,
)
CLAUSE_ID_RE = re.compile(
    r"^(?:(?:Section|Article)\s+)?(\d+(?:\.\d+)*[A-Za-z]?)(?:[\)\.\-:]|\s).*$",
    re.IGNORECASE,
)


def utc_now_iso() -> str:
    return dt.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def read_docx_member(docx_path: Path, member_name: str) -> bytes:
    with zipfile.ZipFile(docx_path, "r") as archive:
        return archive.read(member_name)


def parse_xml(data: bytes) -> etree._Element:
    return etree.fromstring(data)


def normalize_author(value: str) -> str:
    return " ".join(value.strip().lower().split())


def paragraph_text(paragraph: etree._Element) -> str:
    parts: list[str] = []
    for node in paragraph.xpath(".//w:t | .//w:delText", namespaces=NS):
        if node.text:
            parts.append(node.text)
    return "".join(parts).strip()


def paragraph_style_id(paragraph: etree._Element) -> str:
    style = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
    return style[0].strip() if style else ""


def comment_text(comment: etree._Element) -> str:
    parts: list[str] = []
    for node in comment.xpath(".//w:t", namespaces=NS):
        if node.text:
            parts.append(node.text)
    return "".join(parts).strip()


def extract_comment_threads(
    input_docx: Path,
    opponent_comment_authors: list[str] | None = None,
) -> list[dict[str, Any]]:
    opponent_set = {
        normalize_author(author)
        for author in (opponent_comment_authors or [])
        if str(author).strip()
    }

    with zipfile.ZipFile(input_docx, "r") as archive:
        member_names = set(archive.namelist())
        if "word/comments.xml" not in member_names:
            return []
        comments_root = parse_xml(archive.read("word/comments.xml"))
        document_root = parse_xml(archive.read("word/document.xml"))

    paragraph_anchors: dict[str, list[dict[str, Any]]] = {}
    paragraphs = document_root.xpath(".//w:body//w:p", namespaces=NS)
    for paragraph_index, paragraph in enumerate(paragraphs):
        visible_text = paragraph_text(paragraph)
        comment_ids = set(paragraph.xpath(".//w:commentRangeStart/@w:id", namespaces=NS))
        comment_ids.update(paragraph.xpath(".//w:commentReference/@w:id", namespaces=NS))
        for comment_id in comment_ids:
            paragraph_anchors.setdefault(comment_id, []).append(
                {
                    "paragraph_index": paragraph_index,
                    "anchor_text": visible_text,
                }
            )

    extracted: list[dict[str, Any]] = []
    for comment in comments_root.xpath(".//w:comment", namespaces=NS):
        comment_id = str(comment.get(f"{W}id", "")).strip()
        if not comment_id:
            continue

        author = str(comment.get(f"{W}author", "")).strip()
        normalized_author = normalize_author(author)
        is_opponent_comment = normalized_author in opponent_set if opponent_set else True

        anchors = paragraph_anchors.get(comment_id, [])
        first_anchor = anchors[0] if anchors else {}

        extracted.append(
            {
                "comment_id": comment_id,
                "author": author,
                "date": str(comment.get(f"{W}date", "")).strip(),
                "parent_comment_id": str(comment.get(f"{W}parentId", "")).strip(),
                "anchor_text": str(first_anchor.get("anchor_text", "")).strip(),
                "paragraph_index": int(first_anchor.get("paragraph_index", -1)),
                "comment_text": comment_text(comment),
                "is_opponent_comment": is_opponent_comment,
                "response_stance": "",
                "draft_response": "",
                "additional_proposal": "",
            }
        )

    def sort_key(item: dict[str, Any]) -> tuple[int, str]:
        raw_id = str(item.get("comment_id", ""))
        if raw_id.isdigit():
            return (0, f"{int(raw_id):08d}")
        return (1, raw_id)

    return sorted(extracted, key=sort_key)


def is_heading(paragraph: etree._Element, text: str) -> bool:
    if not text:
        return False
    style_id = paragraph_style_id(paragraph).lower()
    if style_id.startswith("heading"):
        return True
    if HEADING_NUMBER_RE.match(text):
        return True
    words = text.split()
    if len(words) <= 12 and len(text) <= 120 and text.upper() == text and any(
        char.isalpha() for char in text
    ):
        return True
    if text.lower().startswith(("section ", "article ")):
        return True
    return False


def infer_clause_id(title: str, fallback_index: int) -> str:
    match = CLAUSE_ID_RE.match(title)
    if match:
        return match.group(1)
    return f"clause-{fallback_index}"


def finalize_clause(
    clauses: list[dict[str, Any]],
    clause_index: int,
    title: str,
    body_lines: list[str],
    paragraph_indexes: list[int],
    target_index: int,
) -> None:
    source_text = "\n".join(line for line in body_lines if line).strip()
    if not source_text:
        source_text = title

    clauses.append(
        {
            "clause_id": infer_clause_id(title, clause_index),
            "clause_title": title,
            "source_text": source_text,
            "paragraph_indexes": paragraph_indexes,
            "target_paragraph_index": target_index,
            "favorability": "",
            "risk_level": "",
            "risk_summary": "",
            "why_it_matters": "",
            "proposed_replacement": "",
        }
    )


def make_review_unit_title(current_heading: str, text: str, clause_index: int) -> str:
    snippet = " ".join(text.split())
    if len(snippet) > 90:
        snippet = snippet[:87].rstrip() + "..."
    if current_heading and current_heading != text:
        return f"{current_heading} :: {snippet}"
    if snippet:
        return snippet
    return f"Clause {clause_index}"


def build_review_template(
    input_docx: Path,
    supported_party: str,
    priority_focus_areas: list[str] | None = None,
    opponent_comment_authors: list[str] | None = None,
) -> dict[str, Any]:
    document_xml = read_docx_member(input_docx, "word/document.xml")
    root = parse_xml(document_xml)
    paragraphs = root.xpath(".//w:body//w:p", namespaces=NS)

    clauses: list[dict[str, Any]] = []
    current_title = ""
    clause_counter = 1

    for index, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue

        if is_heading(paragraph, text):
            current_title = text
        unit_title = make_review_unit_title(current_title, text, clause_counter)
        finalize_clause(
            clauses=clauses,
            clause_index=clause_counter,
            title=unit_title,
            body_lines=[text],
            paragraph_indexes=[index],
            target_index=index,
        )
        clause_counter += 1

    if not clauses:
        raise ValueError("No readable clauses were found in the contract.")

    template = {
        "source_docx": str(input_docx),
        "supported_party": supported_party,
        "priority_focus_areas": priority_focus_areas or [],
        "opponent_comment_authors": opponent_comment_authors or [],
        "prepared_at_utc": utc_now_iso(),
        "reviewer": "",
        "overall_notes": "",
        "clauses": clauses,
        "opponent_comments": extract_comment_threads(
            input_docx=input_docx,
            opponent_comment_authors=opponent_comment_authors or [],
        ),
    }
    return template


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def validate_docx_output_path(path: Path, label: str) -> None:
    if path.suffix.lower() != ".docx":
        raise ValueError(f"{label} must be a .docx file: {path}")


def assert_docx_written(path: Path, label: str) -> None:
    if not path.exists():
        raise FileNotFoundError(f"{label} was not created: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"{label} must be a .docx file: {path}")
    if path.stat().st_size == 0:
        raise ValueError(f"{label} is empty: {path}")


def write_json(path: Path, payload: dict[str, Any]) -> None:
    ensure_parent(path)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        payload = json.load(handle)
    if not isinstance(payload, dict):
        raise ValueError("Review JSON must be a top-level object.")
    if "clauses" not in payload or not isinstance(payload["clauses"], list):
        raise ValueError("Review JSON must include a 'clauses' list.")
    return payload


def set_preserve_space(node: etree._Element, text: str) -> None:
    if text != text.strip():
        node.set(XML_SPACE, "preserve")


def replace_paragraph_with_revision(
    paragraph: etree._Element,
    original_text: str,
    replacement_text: str,
    revision_id: int,
    author: str,
    revision_date: str,
) -> None:
    paragraph_properties = paragraph.find(f"{W}pPr")
    for child in list(paragraph):
        if child is not paragraph_properties:
            paragraph.remove(child)

    deleted = etree.Element(f"{W}del")
    deleted.set(f"{W}id", str(revision_id))
    deleted.set(f"{W}author", author)
    deleted.set(f"{W}date", revision_date)

    deleted_run = etree.SubElement(deleted, f"{W}r")
    deleted_text = etree.SubElement(deleted_run, f"{W}delText")
    deleted_text.text = original_text
    set_preserve_space(deleted_text, original_text)

    inserted = etree.Element(f"{W}ins")
    inserted.set(f"{W}id", str(revision_id + 1))
    inserted.set(f"{W}author", author)
    inserted.set(f"{W}date", revision_date)

    inserted_run = etree.SubElement(inserted, f"{W}r")
    inserted_text = etree.SubElement(inserted_run, f"{W}t")
    inserted_text.text = replacement_text
    set_preserve_space(inserted_text, replacement_text)

    paragraph.append(deleted)
    paragraph.append(inserted)


def ensure_track_revisions(settings_root: etree._Element) -> None:
    existing = settings_root.xpath("./w:trackRevisions", namespaces=NS)
    if existing:
        return
    settings_root.append(etree.Element(f"{W}trackRevisions"))


def build_edit_list(review_payload: dict[str, Any]) -> list[dict[str, Any]]:
    edits: list[dict[str, Any]] = []
    for clause in review_payload.get("clauses", []):
        if not isinstance(clause, dict):
            continue
        replacement = str(clause.get("proposed_replacement", "")).strip()
        if not replacement:
            continue
        target_index = clause.get("target_paragraph_index")
        if not isinstance(target_index, int):
            continue
        edits.append(
            {
                "target_paragraph_index": target_index,
                "replacement": replacement,
            }
        )
    return edits


def apply_tracked_revisions(
    input_docx: Path,
    output_docx: Path,
    edits: list[dict[str, Any]],
    author: str,
) -> int:
    revision_date = utc_now_iso()
    applied = 0

    with zipfile.ZipFile(input_docx, "r") as source_archive:
        original_files = {
            item.filename: source_archive.read(item.filename)
            for item in source_archive.infolist()
        }

    document_root = parse_xml(original_files["word/document.xml"])
    settings_root = parse_xml(original_files["word/settings.xml"])
    paragraphs = document_root.xpath(".//w:body//w:p", namespaces=NS)

    for edit_id, edit in enumerate(edits, start=1):
        target_index = edit["target_paragraph_index"]
        if target_index < 0 or target_index >= len(paragraphs):
            continue
        paragraph = paragraphs[target_index]
        original_text = paragraph_text(paragraph)
        replacement = edit["replacement"].strip()
        if not replacement or not original_text:
            continue
        replace_paragraph_with_revision(
            paragraph=paragraph,
            original_text=original_text,
            replacement_text=replacement,
            revision_id=edit_id * 10,
            author=author,
            revision_date=revision_date,
        )
        applied += 1

    ensure_track_revisions(settings_root)

    original_files["word/document.xml"] = etree.tostring(
        document_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )
    original_files["word/settings.xml"] = etree.tostring(
        settings_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    ensure_parent(output_docx)
    with zipfile.ZipFile(input_docx, "r") as source_archive, zipfile.ZipFile(
        output_docx, "w", compression=zipfile.ZIP_DEFLATED
    ) as output_archive:
        for item in source_archive.infolist():
            output_archive.writestr(item, original_files[item.filename])

    return applied


def write_risk_report(review_payload: dict[str, Any], output_docx: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "Missing dependency: python-docx. Install with `python -m pip install --user python-docx`."
        ) from exc

    document = Document()
    document.add_heading("Contract Risk Review", level=0)
    document.add_paragraph(f"Supported party: {review_payload.get('supported_party', '')}")
    document.add_paragraph(f"Source contract: {review_payload.get('source_docx', '')}")
    document.add_paragraph(f"Generated at (UTC): {utc_now_iso()}")

    focus_areas = review_payload.get("priority_focus_areas", [])
    if isinstance(focus_areas, list):
        normalized_focus_areas = [str(item).strip() for item in focus_areas if str(item).strip()]
    else:
        normalized_focus_areas = []
    if normalized_focus_areas:
        document.add_heading("Priority Focus Areas", level=1)
        for area in normalized_focus_areas:
            document.add_paragraph(area, style="List Bullet")

    notes = str(review_payload.get("overall_notes", "")).strip()
    if notes:
        document.add_heading("Overall Notes", level=1)
        document.add_paragraph(notes)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [
        "Clause",
        "Favorability",
        "Risk Level",
        "Risk Summary",
        "Why It Matters",
        "Proposed Replacement",
    ]
    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title

    for clause in review_payload.get("clauses", []):
        if not isinstance(clause, dict):
            continue
        row = table.add_row().cells
        row[0].text = str(clause.get("clause_title", clause.get("clause_id", "")))
        row[1].text = str(clause.get("favorability", ""))
        row[2].text = str(clause.get("risk_level", ""))
        row[3].text = str(clause.get("risk_summary", ""))
        row[4].text = str(clause.get("why_it_matters", ""))
        row[5].text = str(clause.get("proposed_replacement", ""))

    comment_entries = review_payload.get("opponent_comments", [])
    if isinstance(comment_entries, list):
        opponent_comments = [
            item
            for item in comment_entries
            if isinstance(item, dict) and bool(item.get("is_opponent_comment", True))
        ]
    else:
        opponent_comments = []

    if opponent_comments:
        document.add_heading("Opponent Comment Responses", level=1)
        comment_table = document.add_table(rows=1, cols=7)
        comment_table.style = "Table Grid"
        comment_headers = [
            "Comment ID",
            "Author",
            "Anchor Text",
            "Opponent Comment",
            "Stance",
            "Draft Response",
            "Additional Proposal",
        ]
        for idx, title in enumerate(comment_headers):
            comment_table.rows[0].cells[idx].text = title

        for item in opponent_comments:
            row = comment_table.add_row().cells
            row[0].text = str(item.get("comment_id", ""))
            row[1].text = str(item.get("author", ""))
            row[2].text = str(item.get("anchor_text", ""))
            row[3].text = str(item.get("comment_text", ""))
            row[4].text = str(item.get("response_stance", ""))
            row[5].text = str(item.get("draft_response", ""))
            row[6].text = str(item.get("additional_proposal", ""))

    ensure_parent(output_docx)
    document.save(output_docx)


def command_init_review(args: argparse.Namespace) -> None:
    payload = build_review_template(
        input_docx=args.input,
        supported_party=args.supported_party or "",
        priority_focus_areas=args.focus_areas or [],
        opponent_comment_authors=args.opponent_comment_authors or [],
    )
    write_json(args.output, payload)
    print(f"[OK] Wrote review template: {args.output}")
    print(f"[OK] Clauses detected: {len(payload['clauses'])}")
    print(f"[OK] Comments detected: {len(payload.get('opponent_comments', []))}")


def command_materialize(args: argparse.Namespace) -> None:
    validate_docx_output_path(args.amended_output, "Amended output")
    validate_docx_output_path(args.report_output, "Risk report output")

    payload = load_json(args.review_json)
    edits = build_edit_list(payload)
    applied = apply_tracked_revisions(
        input_docx=args.input,
        output_docx=args.amended_output,
        edits=edits,
        author=args.author,
    )
    write_risk_report(payload, args.report_output)
    assert_docx_written(args.amended_output, "Amended contract")
    assert_docx_written(args.report_output, "Risk report")

    print(f"[OK] Wrote amended contract: {args.amended_output}")
    print(f"[OK] Wrote risk report: {args.report_output}")
    print(f"[OK] Tracked revisions applied: {applied}")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Legal contract DOCX extraction + redline materialization tool."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    init_review = subparsers.add_parser(
        "init-review",
        help="Extract paragraph-level contract review units into a JSON review template.",
    )
    init_review.add_argument("--input", type=Path, required=True, help="Source contract .docx")
    init_review.add_argument("--output", type=Path, required=True, help="Output review .json")
    init_review.add_argument(
        "--supported-party",
        default="",
        help="Party to support (e.g., Buyer, Seller, Licensor).",
    )
    init_review.add_argument(
        "--focus-area",
        dest="focus_areas",
        action="append",
        default=[],
        help="Add a priority focus area; repeat to include multiple areas.",
    )
    init_review.add_argument(
        "--opponent-comment-author",
        dest="opponent_comment_authors",
        action="append",
        default=[],
        help="Word comment author to classify as opponent; repeat to include multiple names.",
    )
    init_review.set_defaults(handler=command_init_review)

    materialize = subparsers.add_parser(
        "materialize",
        help="Create amended DOCX and risk report DOCX from a completed review JSON.",
    )
    materialize.add_argument("--input", type=Path, required=True, help="Source contract .docx")
    materialize.add_argument(
        "--review-json",
        type=Path,
        required=True,
        help="Completed review JSON file.",
    )
    materialize.add_argument(
        "--amended-output",
        type=Path,
        required=True,
        help="Output path for amended tracked-change .docx",
    )
    materialize.add_argument(
        "--report-output",
        type=Path,
        required=True,
        help="Output path for risk analysis .docx",
    )
    materialize.add_argument(
        "--author",
        default="Codex Legal Reviewer",
        help="Revision author recorded in tracked changes.",
    )
    materialize.set_defaults(handler=command_materialize)

    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    args.handler(args)


if __name__ == "__main__":
    main()


#!/usr/bin/env python
"""Analyze Form 1040 returns for current-year compliance, cross-year consistency, and audit risk."""

from __future__ import annotations

import argparse
import json
from collections import defaultdict
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover - runtime dependency
    Document = None


SEVERITY_ORDER = {"high": 0, "medium": 1, "low": 2}

STATUS_ALIASES_DEFAULT = {
    "single": "single",
    "s": "single",
    "married filing jointly": "married_filing_jointly",
    "mfj": "married_filing_jointly",
    "married filing separately": "married_filing_separately",
    "mfs": "married_filing_separately",
    "head of household": "head_of_household",
    "hoh": "head_of_household",
    "qualifying surviving spouse": "qualifying_surviving_spouse",
    "qss": "qualifying_surviving_spouse",
}

NUMERIC_MAJOR_FIELDS = [
    "wages",
    "wages_subject_to_social_security",
    "taxable_interest",
    "ordinary_dividends",
    "qualified_dividends",
    "capital_gain_or_loss",
    "business_income",
    "self_employment_net_earnings",
    "schedule_se_tax",
    "unemployment_compensation",
    "other_income",
    "total_income",
    "adjustments_to_income",
    "agi",
    "deduction_amount",
    "qbi_deduction",
    "taxable_income",
    "tax_before_credits",
    "nonrefundable_credits",
    "ctc_claimed",
    "actc_claimed",
    "other_taxes",
    "additional_medicare_tax",
    "total_tax",
    "withholding",
    "estimated_tax_payments",
    "refundable_credits",
    "other_payments",
    "total_payments",
    "refund",
    "amount_owed",
    "earned_income",
]

REPORT_ITEMS = [
    ("wages", "Wages"),
    ("total_income", "Total income"),
    ("agi", "Adjusted gross income (AGI)"),
    ("deduction_amount", "Deduction amount"),
    ("taxable_income", "Taxable income"),
    ("tax_before_credits", "Tax before credits"),
    ("ctc_claimed", "Child Tax Credit (CTC)"),
    ("actc_claimed", "Additional CTC (ACTC)"),
    ("schedule_se_tax", "Schedule SE tax"),
    ("additional_medicare_tax", "Additional Medicare tax"),
    ("total_tax", "Total tax"),
    ("withholding", "Withholding"),
    ("total_payments", "Total payments"),
    ("refund", "Refund"),
    ("amount_owed", "Amount owed"),
]


@dataclass
class Finding:
    severity: str
    category: str
    item: str
    title: str
    detail: str
    recommendation: str

    def to_dict(self) -> Dict[str, str]:
        return {
            "severity": self.severity,
            "category": self.category,
            "item": self.item,
            "title": self.title,
            "detail": self.detail,
            "recommendation": self.recommendation,
        }


@dataclass
class ReturnRecord:
    tax_year: int
    filing_status: str
    deduction_type: str
    dependents_count: int
    qualifying_children_ctc: int
    additional_standard_count: int
    has_preferential_income: bool
    major_items: Dict[str, float]
    provided_fields: set[str]
    metadata: Dict[str, Any]


@dataclass
class RiskAssessment:
    score: int
    level: str
    probability_range: str
    rationale: str

    def to_dict(self) -> Dict[str, str | int]:
        return {
            "score": self.score,
            "level": self.level,
            "probability_range": self.probability_range,
            "rationale": self.rationale,
        }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Review normalized Form 1040 return data.")
    parser.add_argument("--input", required=True, help="Path to normalized returns JSON.")
    parser.add_argument(
        "--law",
        default=str(Path(__file__).resolve().parents[1] / "references" / "current_tax_law_2025.json"),
        help="Path to tax law parameter JSON.",
    )
    parser.add_argument(
        "--output-dir",
        default="output/form-1040-review",
        help="Directory for review artifacts.",
    )
    parser.add_argument("--current-year", type=int, help="Force current tax year; default is max year in input.")
    parser.add_argument("--docx-name", default="form-1040-risk-report.docx", help="DOCX output filename.")
    parser.add_argument("--skip-docx", action="store_true", help="Skip DOCX generation.")
    return parser.parse_args()


def read_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def write_json(path: Path, payload: Any) -> None:
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
        handle.write("\n")


def to_float(raw_value: Any, field: str) -> float:
    if raw_value in (None, ""):
        return 0.0
    if isinstance(raw_value, (int, float)):
        return float(raw_value)
    if isinstance(raw_value, bool):
        return 1.0 if raw_value else 0.0

    cleaned = str(raw_value).strip().replace(",", "").replace("$", "")
    if not cleaned:
        return 0.0

    try:
        return float(cleaned)
    except ValueError as exc:
        raise ValueError(f"Could not parse numeric field '{field}' from value '{raw_value}'.") from exc


def to_int(raw_value: Any, field: str) -> int:
    if raw_value in (None, ""):
        return 0
    if isinstance(raw_value, bool):
        return 1 if raw_value else 0
    if isinstance(raw_value, int):
        return raw_value
    if isinstance(raw_value, float):
        return int(round(raw_value))

    cleaned = str(raw_value).strip().replace(",", "")
    try:
        return int(float(cleaned))
    except ValueError as exc:
        raise ValueError(f"Could not parse integer field '{field}' from value '{raw_value}'.") from exc


def normalize_filing_status(value: str, aliases: Dict[str, str]) -> str:
    key = " ".join(str(value).strip().lower().replace("_", " ").replace("-", " ").split())
    if key in aliases:
        return aliases[key]
    raise ValueError(
        f"Unsupported filing_status '{value}'. Expected one of: {', '.join(sorted(set(aliases.values())))}"
    )


def parse_return(raw: Dict[str, Any], aliases: Dict[str, str]) -> ReturnRecord:
    if "tax_year" not in raw:
        raise ValueError("Each return entry must include 'tax_year'.")
    if "filing_status" not in raw:
        raise ValueError("Each return entry must include 'filing_status'.")

    tax_year = to_int(raw["tax_year"], "tax_year")
    filing_status = normalize_filing_status(str(raw["filing_status"]), aliases)
    deduction_type = str(raw.get("deduction_type", "standard")).strip().lower()

    if deduction_type not in {"standard", "itemized"}:
        raise ValueError("deduction_type must be 'standard' or 'itemized'.")

    major_blob: Dict[str, Any] = dict(raw.get("major_items", {}))
    for key in NUMERIC_MAJOR_FIELDS:
        if key in raw and key not in major_blob:
            major_blob[key] = raw[key]

    major_items: Dict[str, float] = {}
    for key in NUMERIC_MAJOR_FIELDS:
        major_items[key] = to_float(major_blob.get(key, 0.0), key)

    provided_fields = {key for key in major_blob.keys() if key in NUMERIC_MAJOR_FIELDS}

    has_preferential_income = bool(raw.get("has_preferential_income", False))
    if (
        not has_preferential_income
        and (major_items.get("qualified_dividends", 0.0) > 0.0 or major_items.get("capital_gain_or_loss", 0.0) > 0.0)
    ):
        has_preferential_income = True

    additional_standard_count = to_int(raw.get("additional_standard_deduction_count", 0), "additional_standard_deduction_count")
    dependents_count = to_int(raw.get("dependents_count", 0), "dependents_count")
    qualifying_children_ctc = to_int(raw.get("qualifying_children_ctc", 0), "qualifying_children_ctc")

    metadata = raw.get("metadata", {})
    if metadata is None:
        metadata = {}

    return ReturnRecord(
        tax_year=tax_year,
        filing_status=filing_status,
        deduction_type=deduction_type,
        dependents_count=dependents_count,
        qualifying_children_ctc=qualifying_children_ctc,
        additional_standard_count=additional_standard_count,
        has_preferential_income=has_preferential_income,
        major_items=major_items,
        provided_fields=provided_fields,
        metadata=metadata,
    )


def parse_returns(payload: Any, aliases: Dict[str, str]) -> List[ReturnRecord]:
    entries: List[Dict[str, Any]]
    if isinstance(payload, list):
        entries = payload
    elif isinstance(payload, dict):
        entries = payload.get("returns", [])
    else:
        raise ValueError("Input JSON must be an object with 'returns' or a list of returns.")

    if not entries:
        raise ValueError("Input contains no returns.")

    returns = [parse_return(entry, aliases) for entry in entries]
    returns.sort(key=lambda item: item.tax_year)
    return returns


def has_fields(record: ReturnRecord, *fields: str) -> bool:
    return all(field in record.provided_fields for field in fields)


def fmt_money(value: float) -> str:
    if abs(value - round(value)) < 0.01:
        return f"${value:,.0f}"
    return f"${value:,.2f}"


def pct_change(prior: float, current: float) -> Optional[float]:
    if abs(prior) < 1e-9:
        return None
    return (current - prior) / abs(prior)


def highest_severity(findings: Iterable[Finding]) -> str:
    best = "none"
    for finding in findings:
        if best == "none" or SEVERITY_ORDER[finding.severity] < SEVERITY_ORDER[best]:
            best = finding.severity
    return best


def get_tax_brackets(law: Dict[str, Any], filing_status: str) -> Optional[List[Dict[str, Any]]]:
    tax_brackets = law.get("tax_brackets", {})
    brackets = tax_brackets.get(filing_status)
    if isinstance(brackets, str):
        brackets = tax_brackets.get(brackets)
    if not isinstance(brackets, list):
        return None
    return brackets


def compute_progressive_tax(taxable_income: float, brackets: List[Dict[str, Any]]) -> float:
    if taxable_income <= 0:
        return 0.0

    remaining_income = taxable_income
    lower_bound = 0.0
    tax = 0.0

    for bracket in brackets:
        upper_bound = bracket.get("up_to")
        rate = float(bracket["rate"])

        if upper_bound is None:
            taxable_slice = max(0.0, remaining_income)
            tax += taxable_slice * rate
            return tax

        upper_bound = float(upper_bound)
        slice_width = max(0.0, upper_bound - lower_bound)
        taxable_slice = min(remaining_income, slice_width)
        tax += taxable_slice * rate
        remaining_income -= taxable_slice
        lower_bound = upper_bound

        if remaining_income <= 0:
            return tax

    return tax


def check_arithmetic(record: ReturnRecord, tolerance: float) -> List[Finding]:
    findings: List[Finding] = []
    m = record.major_items

    if has_fields(record, "total_income", "adjustments_to_income", "agi"):
        expected = m["total_income"] - m["adjustments_to_income"]
        if abs(m["agi"] - expected) > tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="arithmetic",
                    item="agi",
                    title="AGI mismatch",
                    detail=f"AGI is {fmt_money(m['agi'])}, but total income minus adjustments is {fmt_money(expected)}.",
                    recommendation="Reconcile Form 1040 lines for total income, adjustments, and AGI.",
                )
            )

    if has_fields(record, "agi", "deduction_amount", "qbi_deduction", "taxable_income"):
        expected = max(0.0, m["agi"] - m["deduction_amount"] - m["qbi_deduction"])
        if abs(m["taxable_income"] - expected) > tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="arithmetic",
                    item="taxable_income",
                    title="Taxable income mismatch",
                    detail=(
                        f"Taxable income is {fmt_money(m['taxable_income'])}, but AGI minus deductions and QBI is "
                        f"{fmt_money(expected)}."
                    ),
                    recommendation="Verify deduction, QBI deduction, and taxable-income worksheet entries.",
                )
            )

    if has_fields(record, "tax_before_credits", "nonrefundable_credits", "other_taxes", "total_tax"):
        expected = max(0.0, m["tax_before_credits"] - m["nonrefundable_credits"]) + m["other_taxes"]
        if abs(m["total_tax"] - expected) > tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="arithmetic",
                    item="total_tax",
                    title="Total tax mismatch",
                    detail=(
                        f"Total tax is {fmt_money(m['total_tax'])}, but expected from tax, credits, and other taxes is "
                        f"{fmt_money(expected)}."
                    ),
                    recommendation="Recheck nonrefundable-credit offsets and other-tax carryovers.",
                )
            )

    if has_fields(record, "withholding", "estimated_tax_payments", "refundable_credits", "other_payments", "total_payments"):
        expected = (
            m["withholding"] + m["estimated_tax_payments"] + m["refundable_credits"] + m["other_payments"]
        )
        if abs(m["total_payments"] - expected) > tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="arithmetic",
                    item="total_payments",
                    title="Total payments mismatch",
                    detail=(
                        f"Total payments is {fmt_money(m['total_payments'])}, but component payments sum to "
                        f"{fmt_money(expected)}."
                    ),
                    recommendation="Reconcile withholding, estimated payments, refundable credits, and other payments.",
                )
            )

    if has_fields(record, "total_payments", "total_tax", "refund", "amount_owed"):
        net = m["total_payments"] - m["total_tax"]
        expected_refund = max(0.0, net)
        expected_owed = max(0.0, -net)
        if abs(m["refund"] - expected_refund) > tolerance or abs(m["amount_owed"] - expected_owed) > tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="arithmetic",
                    item="refund",
                    title="Refund/amount owed mismatch",
                    detail=(
                        f"Net payments minus total tax is {fmt_money(net)}. Reported refund is {fmt_money(m['refund'])} "
                        f"and amount owed is {fmt_money(m['amount_owed'])}."
                    ),
                    recommendation="Recalculate net balance due and ensure refund/amount owed lines are mutually consistent.",
                )
            )

    return findings


def check_law(record: ReturnRecord, law: Dict[str, Any], tolerance: float, tax_tolerance: float) -> List[Finding]:
    findings: List[Finding] = []
    m = record.major_items

    standard_table = law.get("standard_deduction", {})
    if record.deduction_type == "standard" and has_fields(record, "deduction_amount"):
        if record.filing_status in standard_table:
            expected = float(standard_table[record.filing_status])
            addl_table = law.get("additional_standard_deduction", {})
            if record.additional_standard_count > 0:
                if record.filing_status in {"married_filing_jointly", "married_filing_separately", "qualifying_surviving_spouse"}:
                    expected += float(addl_table.get("married_or_qss", 0.0)) * record.additional_standard_count
                else:
                    expected += float(addl_table.get("single_or_hoh_or_mfs", 0.0)) * record.additional_standard_count

            if abs(m["deduction_amount"] - expected) > tolerance:
                findings.append(
                    Finding(
                        severity="high",
                        category="law",
                        item="deduction_amount",
                        title="Standard deduction amount appears incorrect",
                        detail=(
                            f"Reported standard deduction is {fmt_money(m['deduction_amount'])}, expected {fmt_money(expected)} "
                            f"for filing status {record.filing_status}."
                        ),
                        recommendation="Recheck standard deduction amount against current-year IRS guidance.",
                    )
                )

    ctc = law.get("ctc", {})
    ctc_max = float(ctc.get("max_per_qualifying_child", 0.0))
    actc_max = float(ctc.get("actc_max_per_qualifying_child", 0.0))
    if has_fields(record, "ctc_claimed") and ctc_max > 0:
        max_credit = ctc_max * max(0, record.qualifying_children_ctc)
        if m["ctc_claimed"] > max_credit + tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="law",
                    item="ctc_claimed",
                    title="CTC exceeds configured limit",
                    detail=(
                        f"Claimed CTC is {fmt_money(m['ctc_claimed'])}, but configured maximum for "
                        f"{record.qualifying_children_ctc} qualifying children is {fmt_money(max_credit)}."
                    ),
                    recommendation="Validate qualifying-child count and CTC worksheet limits.",
                )
            )

    if has_fields(record, "actc_claimed") and actc_max > 0:
        max_actc = actc_max * max(0, record.qualifying_children_ctc)
        if m["actc_claimed"] > max_actc + tolerance:
            findings.append(
                Finding(
                    severity="high",
                    category="law",
                    item="actc_claimed",
                    title="ACTC exceeds configured limit",
                    detail=(
                        f"Claimed ACTC is {fmt_money(m['actc_claimed'])}, but configured maximum for "
                        f"{record.qualifying_children_ctc} qualifying children is {fmt_money(max_actc)}."
                    ),
                    recommendation="Review Schedule 8812 ACTC computation and dependent eligibility.",
                )
            )

        min_earned_income = float(ctc.get("actc_min_earned_income", 0.0))
        if m["actc_claimed"] > tolerance and has_fields(record, "earned_income") and m["earned_income"] < min_earned_income:
            findings.append(
                Finding(
                    severity="high",
                    category="law",
                    item="actc_claimed",
                    title="ACTC claimed with low earned income",
                    detail=(
                        f"ACTC is claimed ({fmt_money(m['actc_claimed'])}) but earned income is {fmt_money(m['earned_income'])}, "
                        f"below the configured minimum {fmt_money(min_earned_income)}."
                    ),
                    recommendation="Confirm earned-income eligibility for ACTC.",
                )
            )

    if has_fields(record, "taxable_income", "tax_before_credits"):
        if not record.has_preferential_income:
            brackets = get_tax_brackets(law, record.filing_status)
            if brackets:
                expected_tax = compute_progressive_tax(m["taxable_income"], brackets)
                delta = abs(m["tax_before_credits"] - expected_tax)
                dynamic_tolerance = max(tax_tolerance, expected_tax * 0.05)
                if delta > dynamic_tolerance:
                    findings.append(
                        Finding(
                            severity="medium",
                            category="law",
                            item="tax_before_credits",
                            title="Tax-before-credits differs from rate schedule",
                            detail=(
                                f"Reported tax before credits is {fmt_money(m['tax_before_credits'])}; simple rate-schedule "
                                f"result is {fmt_money(expected_tax)}."
                            ),
                            recommendation=(
                                "Check Tax Table/Tax Computation Worksheet inputs and confirm no special computation applies."
                            ),
                        )
                    )
        else:
            findings.append(
                Finding(
                    severity="low",
                    category="law",
                    item="tax_before_credits",
                    title="Preferential-income tax calculation requires manual worksheet review",
                    detail=(
                        "Qualified dividends or capital gains are present, so simple ordinary-bracket recomputation may not "
                        "match Form 1040 tax."
                    ),
                    recommendation="Review Qualified Dividends and Capital Gain Tax Worksheet outputs.",
                )
            )

    se_cfg = law.get("self_employment", {})
    ss_wage_base = float(se_cfg.get("social_security_wage_base", 0.0))
    ss_rate = float(se_cfg.get("social_security_rate", 0.124))
    medicare_rate = float(se_cfg.get("medicare_rate", 0.029))

    if has_fields(record, "self_employment_net_earnings", "schedule_se_tax") and ss_wage_base > 0:
        net_earnings = max(0.0, m["self_employment_net_earnings"])
        se_base = net_earnings * 0.9235
        wages_for_ss = m.get("wages_subject_to_social_security", m.get("wages", 0.0))
        remaining_ss_base = max(0.0, ss_wage_base - max(0.0, wages_for_ss))
        ss_taxable = min(se_base, remaining_ss_base)
        expected_se_tax = (ss_taxable * ss_rate) + (se_base * medicare_rate)

        if abs(m["schedule_se_tax"] - expected_se_tax) > max(75.0, tax_tolerance):
            findings.append(
                Finding(
                    severity="medium",
                    category="law",
                    item="schedule_se_tax",
                    title="Schedule SE tax differs from approximation",
                    detail=(
                        f"Reported Schedule SE tax is {fmt_money(m['schedule_se_tax'])}; approximate computation is "
                        f"{fmt_money(expected_se_tax)} using the configured SS wage base."
                    ),
                    recommendation="Recompute Schedule SE, especially SS wage-base interaction and deductible half-SE tax.",
                )
            )

    thresholds = law.get("additional_medicare_threshold", {})
    threshold_value = thresholds.get(record.filing_status)
    if threshold_value is not None and has_fields(record, "additional_medicare_tax"):
        threshold = float(threshold_value)
        medicare_base = max(0.0, m.get("wages", 0.0)) + max(0.0, m.get("self_employment_net_earnings", 0.0) * 0.9235)
        expected_addl = max(0.0, medicare_base - threshold) * 0.009
        if expected_addl - m["additional_medicare_tax"] > max(25.0, tolerance):
            findings.append(
                Finding(
                    severity="medium",
                    category="law",
                    item="additional_medicare_tax",
                    title="Additional Medicare tax may be understated",
                    detail=(
                        f"Approximate Additional Medicare tax is {fmt_money(expected_addl)} based on income {fmt_money(medicare_base)}, "
                        f"but reported amount is {fmt_money(m['additional_medicare_tax'])}."
                    ),
                    recommendation="Review Form 8959 and withholding credit application.",
                )
            )

    if record.qualifying_children_ctc > record.dependents_count and record.dependents_count > 0:
        findings.append(
            Finding(
                severity="high",
                category="law",
                item="qualifying_children_ctc",
                title="Qualifying-child count exceeds dependent count",
                detail=(
                    f"Qualifying children for CTC is {record.qualifying_children_ctc}, but dependents count is "
                    f"{record.dependents_count}."
                ),
                recommendation="Reconcile dependent roster and CTC qualifying-child eligibility.",
            )
        )

    return findings


def check_consistency(current: ReturnRecord, prior: Optional[ReturnRecord], law: Dict[str, Any]) -> List[Finding]:
    findings: List[Finding] = []
    if prior is None:
        return findings

    if current.filing_status != prior.filing_status:
        findings.append(
            Finding(
                severity="medium",
                category="consistency",
                item="filing_status",
                title="Filing status changed from prior year",
                detail=f"Prior year status: {prior.filing_status}; current year status: {current.filing_status}.",
                recommendation="Document life-event support for filing-status change.",
            )
        )

    dependent_delta = current.dependents_count - prior.dependents_count
    if abs(dependent_delta) >= 2:
        findings.append(
            Finding(
                severity="medium",
                category="consistency",
                item="dependents_count",
                title="Large dependent-count change",
                detail=(
                    f"Dependent count changed by {dependent_delta:+d} (from {prior.dependents_count} to "
                    f"{current.dependents_count})."
                ),
                recommendation="Retain support for dependent additions/removals.",
            )
        )

    thresholds = law.get("consistency_thresholds", {})
    for field, cfg in thresholds.items():
        if field not in current.provided_fields or field not in prior.provided_fields:
            continue

        curr = current.major_items[field]
        prev = prior.major_items[field]
        delta = curr - prev
        pct = pct_change(prev, curr)
        min_abs = float(cfg.get("absolute_min", 0.0))
        pct_limit = float(cfg.get("pct_change", 0.0))

        if abs(delta) < min_abs:
            continue

        exceeded = False
        if pct is None:
            exceeded = abs(curr) >= min_abs
        else:
            exceeded = abs(pct) >= pct_limit

        if not exceeded:
            continue

        severity = str(cfg.get("severity", "low"))
        if pct is None:
            delta_text = f"changed from {fmt_money(prev)} to {fmt_money(curr)}"
        else:
            delta_text = f"changed by {pct * 100:+.1f}% ({fmt_money(prev)} to {fmt_money(curr)})"

        findings.append(
            Finding(
                severity=severity,
                category="consistency",
                item=field,
                title=f"Year-over-year variance in {field.replace('_', ' ')}",
                detail=(
                    f"Compared with {prior.tax_year}, {field.replace('_', ' ')} {delta_text}."
                ),
                recommendation="Confirm supporting documents and explanation for this variance.",
            )
        )

    return findings


def assess_audit_risk(findings: List[Finding], current: ReturnRecord) -> RiskAssessment:
    weights = {"high": 12, "medium": 6, "low": 2}
    score = sum(weights.get(finding.severity, 0) for finding in findings)

    # Complexity adjustments
    if current.major_items.get("self_employment_net_earnings", 0.0) > 0:
        score += 4
    if current.deduction_type == "itemized":
        score += 2
    if current.has_preferential_income:
        score += 1

    high_count = sum(1 for finding in findings if finding.severity == "high")
    if high_count >= 3:
        score += 10

    score = int(min(100, max(0, round(score))))

    if score <= 15:
        level = "Low"
        probability = "5-10%"
    elif score <= 35:
        level = "Low-Moderate"
        probability = "10-20%"
    elif score <= 55:
        level = "Moderate"
        probability = "20-35%"
    elif score <= 75:
        level = "Moderate-High"
        probability = "35-55%"
    else:
        level = "High"
        probability = "55%+"

    rationale = (
        "Heuristic score based on arithmetic mismatches, law-check findings, and multi-year consistency variances. "
        "Not an IRS prediction model."
    )
    return RiskAssessment(score=score, level=level, probability_range=probability, rationale=rationale)


def render_markdown(
    path: Path,
    current: ReturnRecord,
    prior: Optional[ReturnRecord],
    findings: List[Finding],
    risk: RiskAssessment,
    law: Dict[str, Any],
) -> None:
    lines: List[str] = []

    lines.append(f"# Form 1040 Review Summary - Tax Year {current.tax_year}")
    lines.append("")
    lines.append(f"Generated: {date.today().isoformat()}")
    lines.append(f"Law reference tax year: {law.get('tax_year', 'unknown')} (as of {law.get('as_of_date', 'unknown')})")
    lines.append("")
    lines.append("## Overall Audit Risk")
    lines.append("")
    lines.append(
        f"Estimated risk: **{risk.level}** (score **{risk.score}/100**, indicative probability range **{risk.probability_range}**)."
    )
    lines.append("")
    lines.append(risk.rationale)
    lines.append("")

    lines.append("## Return Scope")
    lines.append("")
    lines.append(f"- Current return year: {current.tax_year}")
    lines.append(f"- Filing status: {current.filing_status}")
    lines.append(f"- Deduction type: {current.deduction_type}")
    if prior is not None:
        lines.append(f"- Compared to prior return year: {prior.tax_year}")
    else:
        lines.append("- No historical return provided for consistency comparison")
    lines.append("")

    lines.append("## Findings")
    lines.append("")

    if not findings:
        lines.append("No findings were generated by configured checks.")
    else:
        lines.append("| Severity | Category | Item | Finding | Recommendation |")
        lines.append("|---|---|---|---|---|")
        for finding in findings:
            lines.append(
                "| {severity} | {category} | {item} | {title}: {detail} | {recommendation} |".format(
                    severity=finding.severity.upper(),
                    category=finding.category,
                    item=finding.item.replace("|", "\\|"),
                    title=finding.title.replace("|", "\\|"),
                    detail=finding.detail.replace("|", "\\|"),
                    recommendation=finding.recommendation.replace("|", "\\|"),
                )
            )

    lines.append("")
    lines.append("## Note")
    lines.append("")
    lines.append("This report is a technical consistency and rules check, not legal or tax filing advice.")

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_docx(
    path: Path,
    current: ReturnRecord,
    prior: Optional[ReturnRecord],
    findings: List[Finding],
    risk: RiskAssessment,
) -> None:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Install with: "
            "C:\\windows\\System32\\WindowsPowerShell\\v1.0\\powershell.exe -Command \"$env:PYTHONUSERBASE='C:\\Codex\\.pyuser'; python -m pip install --user python-docx\""
        )

    findings_by_item: Dict[str, List[Finding]] = defaultdict(list)
    for finding in findings:
        findings_by_item[finding.item].append(finding)

    document = Document()
    document.add_heading(f"Form 1040 Risk Review - Tax Year {current.tax_year}", level=0)
    document.add_paragraph(f"Generated: {date.today().isoformat()}")
    document.add_paragraph(
        "This report lists major return items and risk observations from arithmetic checks, law checks, and "
        "cross-year consistency checks."
    )

    p = document.add_paragraph()
    p.add_run("Overall audit-risk estimate: ").bold = True
    p.add_run(f"{risk.level} (score {risk.score}/100; indicative probability {risk.probability_range}).")

    if prior is not None:
        document.add_paragraph(f"Historical consistency comparison performed against tax year {prior.tax_year}.")
    else:
        document.add_paragraph("No historical return was provided for year-over-year consistency comparison.")

    table = document.add_table(rows=1, cols=5)
    header = table.rows[0].cells
    header[0].text = "Major item"
    header[1].text = "Current value"
    header[2].text = "Highest severity"
    header[3].text = "Observations"
    header[4].text = "Recommended documentation"

    for item_key, item_label in REPORT_ITEMS:
        row_cells = table.add_row().cells
        value_text = fmt_money(current.major_items[item_key]) if item_key in current.provided_fields else "Not provided"

        relevant = findings_by_item.get(item_key, [])
        severity = highest_severity(relevant)
        if relevant:
            observations = " ".join(f"{f.title}." for f in relevant[:2])
            recommendation = relevant[0].recommendation
        else:
            observations = "No material issue detected by configured checks."
            recommendation = "Retain normal support documents for this line item."

        row_cells[0].text = item_label
        row_cells[1].text = value_text
        row_cells[2].text = severity.upper() if severity != "none" else "NONE"
        row_cells[3].text = observations
        row_cells[4].text = recommendation

    consistency_findings = [f for f in findings if f.category == "consistency"]
    document.add_heading("Cross-Year Consistency Notes", level=1)
    if consistency_findings:
        for finding in consistency_findings:
            document.add_paragraph(
                f"[{finding.severity.upper()}] {finding.title}: {finding.detail} Recommendation: {finding.recommendation}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No multi-year inconsistency findings were generated.")

    document.add_heading("All Findings", level=1)
    if findings:
        for finding in findings:
            document.add_paragraph(
                f"[{finding.severity.upper()}] {finding.category} - {finding.title}: {finding.detail}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No findings were generated by configured checks.")

    document.add_paragraph("This report is a heuristic review aid and not legal or tax advice.")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def select_current_and_prior(returns: List[ReturnRecord], forced_year: Optional[int]) -> Tuple[ReturnRecord, Optional[ReturnRecord]]:
    if forced_year is None:
        current = max(returns, key=lambda item: item.tax_year)
    else:
        matches = [record for record in returns if record.tax_year == forced_year]
        if not matches:
            raise ValueError(f"No return found for --current-year {forced_year}.")
        current = matches[0]

    prior_candidates = [record for record in returns if record.tax_year < current.tax_year]
    prior = max(prior_candidates, key=lambda item: item.tax_year) if prior_candidates else None
    return current, prior


def sort_findings(findings: List[Finding]) -> List[Finding]:
    return sorted(
        findings,
        key=lambda item: (
            SEVERITY_ORDER.get(item.severity, 99),
            item.category,
            item.item,
            item.title,
        ),
    )


def main() -> None:
    args = parse_args()

    input_path = Path(args.input).resolve()
    law_path = Path(args.law).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    law = read_json(law_path)
    aliases = dict(STATUS_ALIASES_DEFAULT)
    aliases.update({k.lower(): v for k, v in law.get("filing_status_aliases", {}).items()})

    returns = parse_returns(read_json(input_path), aliases)
    current, prior = select_current_and_prior(returns, args.current_year)

    tolerance = float(law.get("tolerance", {}).get("dollar", 2.0))
    tax_tolerance = float(law.get("tolerance", {}).get("tax_compare_dollar", 150.0))

    findings: List[Finding] = []

    law_year = law.get("tax_year")
    if isinstance(law_year, int) and law_year < current.tax_year:
        findings.append(
            Finding(
                severity="high",
                category="law",
                item="law_reference",
                title="Law reference year is stale",
                detail=f"Law file tax_year is {law_year}, but current return year is {current.tax_year}.",
                recommendation="Update law parameters before finalizing review conclusions.",
            )
        )

    findings.extend(check_arithmetic(current, tolerance))
    findings.extend(check_law(current, law, tolerance, tax_tolerance))
    findings.extend(check_consistency(current, prior, law))

    findings = sort_findings(findings)
    risk = assess_audit_risk(findings, current)

    summary_path = output_dir / "review_summary.md"
    findings_path = output_dir / "review_findings.json"
    docx_path = output_dir / args.docx_name

    render_markdown(summary_path, current, prior, findings, risk, law)

    write_json(
        findings_path,
        {
            "current_tax_year": current.tax_year,
            "law_reference": {
                "path": str(law_path),
                "tax_year": law.get("tax_year"),
                "as_of_date": law.get("as_of_date"),
                "sources": law.get("sources", []),
            },
            "audit_risk": risk.to_dict(),
            "current_return": {
                "tax_year": current.tax_year,
                "filing_status": current.filing_status,
                "deduction_type": current.deduction_type,
                "dependents_count": current.dependents_count,
                "qualifying_children_ctc": current.qualifying_children_ctc,
                "major_items": current.major_items,
            },
            "prior_return_year": prior.tax_year if prior else None,
            "findings": [finding.to_dict() for finding in findings],
        },
    )

    if not args.skip_docx:
        render_docx(docx_path, current, prior, findings, risk)

    print(f"Current return year: {current.tax_year}")
    print(f"Findings: {len(findings)}")
    print(f"Audit risk: {risk.level} ({risk.score}/100, {risk.probability_range})")
    print(f"Summary: {summary_path}")
    print(f"Findings JSON: {findings_path}")
    if not args.skip_docx:
        print(f"DOCX risk report: {docx_path}")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover - explicit runtime guidance
    python_cmd = Path(sys.executable).name or 'python'
    raise SystemExit(
        'Missing dependency: python-docx\n'
        'Install it before using the native DOCX renderer.\n'
        f'Recommended command: {python_cmd} -m pip install python-docx'
    ) from exc


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '8')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '808080')
        borders.append(elem)
    tbl_pr.append(borders)


def style_run(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold


def style_paragraph(paragraph, size: float = 10.5, bold: bool = False, align=None, space_after: int = 3) -> None:
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.25
    if not paragraph.runs:
        paragraph.add_run('')
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style='List Bullet')
    run = paragraph.add_run(text)
    style_run(run)
    style_paragraph(paragraph)


def add_section_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_run(run, size=12, bold=True)
    style_paragraph(paragraph, size=12, bold=True, space_after=5)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = str(header)
        set_cell_shading(header_cells[idx], 'EDEDED')
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[idx].paragraphs:
            style_paragraph(paragraph, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                style_paragraph(paragraph)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def main() -> None:
    parser = argparse.ArgumentParser(description='Render a structured quotation JSON payload into a native DOCX file.')
    parser.add_argument('--input-json', required=True, help='Path to quotation JSON payload')
    parser.add_argument('--output-docx', required=True, help='Output DOCX path')
    args = parser.parse_args()

    data = json.loads(Path(args.input_json).read_text(encoding='utf-8'))
    output_docx = Path(args.output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles['Normal'].font.name = 'Microsoft YaHei'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    styles['Normal'].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data.get('project_name', '项目报价方案'))
    style_run(title_run, size=16, bold=True)
    style_paragraph(title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    for text in [
        f"报价日期：{data.get('quote_date', '')}",
        f"乙方报价单位：{data.get('vendor_name', '')}",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.add_run(text)
        style_paragraph(paragraph)

    add_section_title(doc, '服务范围')
    for text in data.get('service_scope', []):
        add_bullet(doc, text)

    add_section_title(doc, '需求内容')
    requirement_rows = [[group.get('module', ''), '\n'.join(group.get('items', []))] for group in data.get('requirement_groups', [])]
    add_table(doc, ['模块', '功能内容'], requirement_rows, widths=[4.4, 11.6])

    add_section_title(doc, '交付物')
    for text in data.get('deliverables', []):
        add_bullet(doc, text)

    add_section_title(doc, '项目报价与付款方式')
    for text in [
        f"项目总人天：{data.get('total_days', 0)} 人天",
        f"统一人天单价：{data.get('day_rate', 1200)} 元/人天",
        f"项目总价：{data.get('total_price', 0):,} 元（{data.get('tax_note', '')}）",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.add_run(text)
        style_paragraph(paragraph)
    for stage, percent in data.get('payment_schedule', []):
        add_bullet(doc, f'{stage}：{percent}')

    add_section_title(doc, '模块人天明细')
    module_rows = []
    for item in data.get('line_items', []):
        product_days = item.get('product_days', 0)
        total_days = product_days + item.get('frontend_days', 0) + item.get('backend_days', 0) + item.get('ui_days', 0) + item.get('qa_days', 0)
        module_rows.append([
            item.get('module', ''),
            product_days,
            item.get('frontend_days', 0),
            item.get('backend_days', 0),
            item.get('ui_days', 0),
            item.get('qa_days', 0),
            total_days,
        ])
    add_table(doc, ['模块', '产品人天', '前端人天', '后端人天', 'UI人天', '测试人天', '合计人天'], module_rows, widths=[5.2, 1.7, 1.7, 1.7, 1.7, 1.7, 1.9])

    add_section_title(doc, '角色投入测算')
    role_rows = [[row.get('role', ''), row.get('days', 0)] for row in data.get('roles', [])]
    add_table(doc, ['角色', '人天'], role_rows, widths=[8, 5])

    add_section_title(doc, '项目周期')
    for text in data.get('project_timeline', []):
        add_bullet(doc, text)

    add_section_title(doc, '验收标准')
    for text in data.get('acceptance', []):
        add_bullet(doc, text)

    add_section_title(doc, '前置条件')
    for text in data.get('prerequisites', []):
        add_bullet(doc, text)

    add_section_title(doc, '依赖资源')
    for text in data.get('dependencies', []):
        add_bullet(doc, text)

    add_section_title(doc, '沟通机制')
    for text in data.get('communication', []):
        add_bullet(doc, text)

    if data.get('assumptions'):
        add_section_title(doc, '备注')
        for text in data['assumptions']:
            add_bullet(doc, text)

    doc.save(output_docx)
    print(output_docx)


if __name__ == '__main__':
    main()

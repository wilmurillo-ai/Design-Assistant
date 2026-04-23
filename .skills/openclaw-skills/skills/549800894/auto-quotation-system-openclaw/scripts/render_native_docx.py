from __future__ import annotations
import json
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '8')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '808080')
        borders.append(elem)
    tbl_pr.append(borders)


def style_run(run, size=10.5, bold=False):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold


def style_paragraph(paragraph, size=10.5, bold=False, align=None, space_after=3):
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.25
    if not paragraph.runs:
        paragraph.add_run('')
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold)


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    style_run(r)
    style_paragraph(p)


def add_section_title(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=12, bold=True)
    style_paragraph(p, size=12, bold=True, space_after=5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
        set_cell_shading(hdr[i], 'EDEDED')
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            style_paragraph(p, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                style_paragraph(p)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def main():
    src = Path('/tmp/quote-request/device-workorder-platform-manual.quote.json')
    out = Path('/tmp/quote-request/device-workorder-platform-manual.docx')
    data = json.loads(src.read_text(encoding='utf-8'))

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    styles = doc.styles
    styles['Normal'].font.name = 'Microsoft YaHei'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    styles['Normal'].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(data['project_name'])
    style_run(r, size=16, bold=True)
    style_paragraph(title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    meta1 = doc.add_paragraph()
    meta1.add_run(f"报价日期：{data.get('quote_date','')}")
    style_paragraph(meta1)
    meta2 = doc.add_paragraph()
    meta2.add_run(f"乙方报价单位：{data.get('vendor_name','')}")
    style_paragraph(meta2, space_after=8)

    add_section_title(doc, '服务范围')
    for text in ['产品需求梳理', '技术方案设计', '软件设计开发', '项目部署和运维1年（从提交验收日开始算）']:
        add_bullet(doc, text)

    add_section_title(doc, '需求内容')
    req_rows = []
    for group in data.get('requirement_groups', []):
        req_rows.append([group.get('module',''), '\n'.join(group.get('items', []))])
    add_table(doc, ['模块', '功能内容'], req_rows, widths=[4.2, 11.8])

    add_section_title(doc, '交付物')
    for text in ['产品文档', '前后端源代码', '测试用例', '项目部署文档，接口文档']:
        add_bullet(doc, text)

    add_section_title(doc, '项目报价与付款方式')
    for text in [
        f"项目总人天：{data.get('total_days', 0)} 人天",
        '统一人天单价：1200 元/人天',
        f"项目总价：{data.get('total_price', 0):,} 元（{data.get('tax_note', '')}）",
        '项目启动：50%',
        '功能验收：40%',
        '质保结束：10%',
    ]:
        p = doc.add_paragraph()
        p.add_run(text)
        style_paragraph(p)

    add_section_title(doc, '模块人天明细')
    module_rows = []
    for item in data.get('line_items', []):
        product_days = item.get('product_days', 0)
        module_rows.append([
            item.get('module',''), product_days, item.get('frontend_days',0), item.get('backend_days',0),
            item.get('ui_days',0), item.get('qa_days',0), product_days + item.get('frontend_days',0) + item.get('backend_days',0) + item.get('ui_days',0) + item.get('qa_days',0)
        ])
    add_table(doc, ['模块', '产品人天', '前端人天', '后端人天', 'UI人天', '测试人天', '合计人天'], module_rows, widths=[5.2, 1.7, 1.7, 1.7, 1.7, 1.7, 1.9])

    add_section_title(doc, '角色投入测算')
    role_rows = [[row.get('role',''), row.get('days',0)] for row in data.get('roles', [])]
    add_table(doc, ['角色', '人天'], role_rows, widths=[8, 5])

    add_section_title(doc, '项目周期')
    for text in ['项目周期：4 周（以收到项目预付款为启动时间）', '设计阶段：1 周', '开发测试：2 周', '验收周期：1 周']:
        add_bullet(doc, text)

    add_section_title(doc, '验收标准')
    for text in [
        '默认是在我方的测试环境上进行验收。',
        '管理后台交付标准以满足小程序端的业务为准，UI样式、交互以实际开发为标准。',
        '验收依据按照当前确认的需求文档内容。',
        '文档未提及内容均视为优化或拓展需求，需另行评估。',
    ]:
        add_bullet(doc, text)

    add_section_title(doc, '前置条件')
    for text in [
        '账号等第三方资源需在规定时间内提供，不能及时提供时可先在我方测试账号完成验收。',
        '如涉及其他第三方协作，需确保沟通链路畅通。',
    ]:
        add_bullet(doc, text)

    add_section_title(doc, '依赖资源')
    for text in ['小程序、公众号相关账号', '服务器、域名、证书', '短信推送平台等第三方服务账号信息', '设计素材、内容素材']:
        add_bullet(doc, text)

    add_section_title(doc, '沟通机制')
    for text in ['每周项目沟通例会', '需求评审会、方案评审会、验收评审会', '日常沟通确认']:
        add_bullet(doc, text)

    if data.get('assumptions'):
        add_section_title(doc, '备注')
        for text in data['assumptions']:
            add_bullet(doc, text)

    doc.save(out)
    print(out)

if __name__ == '__main__':
    main()

"""输出：DOCX / Markdown"""
from __future__ import annotations
from pathlib import Path


def export_markdown(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def export_docx(resume, path: Path) -> None:
    try:
        import docx
        from docx.shared import Pt
    except ImportError:
        # 降级为 md
        export_markdown(_resume_to_md(resume), path.with_suffix(".md"))
        return

    doc = docx.Document()
    doc.add_heading(resume.name or "简历", 0)
    if resume.contact:
        doc.add_paragraph(resume.contact)
    if resume.summary:
        doc.add_heading("个人简介", 1)
        doc.add_paragraph(resume.summary)
    if resume.skills:
        doc.add_heading("技能", 1)
        doc.add_paragraph(" · ".join(resume.skills))
    if resume.experiences:
        doc.add_heading("工作经历", 1)
        for e in resume.experiences:
            h = doc.add_paragraph()
            h.add_run(f"{e.company}  |  {e.title}  |  {e.period}").bold = True
            for b in e.bullets:
                doc.add_paragraph(b, style="List Bullet")
    if resume.education:
        doc.add_heading("教育经历", 1)
        for ed in resume.education:
            doc.add_paragraph(ed)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _resume_to_md(resume) -> str:
    parts = [f"# {resume.name or '简历'}"]
    if resume.contact:
        parts.append(resume.contact)
    if resume.summary:
        parts.append(f"\n## 个人简介\n{resume.summary}")
    if resume.skills:
        parts.append(f"\n## 技能\n{' · '.join(resume.skills)}")
    if resume.experiences:
        parts.append("\n## 工作经历")
        for e in resume.experiences:
            parts.append(f"\n### {e.company}  |  {e.title}  |  {e.period}")
            for b in e.bullets:
                parts.append(f"- {b}")
    return "\n".join(parts)

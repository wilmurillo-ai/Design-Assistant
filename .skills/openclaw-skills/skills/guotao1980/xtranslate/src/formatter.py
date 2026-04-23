import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

class WordFormatter:
    """专门用于优化 Word 文档排版的工具类"""
    
    def __init__(self, font_name_en='Times New Roman', font_name_zh='SimSun', font_size=12):
        self.font_name_en = font_name_en
        self.font_name_zh = font_name_zh
        self.font_size = font_size

    def format_document(self, input_path, output_path=None):
        """对文档进行全方位的排版优化"""
        if not output_path:
            name, ext = os.path.splitext(input_path)
            output_path = f"{name}_formatted{ext}"
            
        doc = Document(input_path)
        
        # 0. 设置文档全局默认样式
        self._set_default_style(doc)
        
        # 1. 设置正文段落格式
        for p in doc.paragraphs:
            self._apply_paragraph_style(p)
            
        # 2. 处理表格中的文字格式
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._apply_paragraph_style(p, is_table=True)
                        
        # 3. 清理多余空行 (可选，根据需要开启)
        # self._remove_empty_paragraphs(doc)

        doc.save(output_path)
        return output_path

    def _set_default_style(self, doc):
        """设置文档的默认字体样式"""
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name_en
        font.size = Pt(self.font_size)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name_zh)

    def _apply_paragraph_style(self, paragraph, is_table=False):
        """应用统一的段落和字体样式"""
        # 段落级设置
        p_format = paragraph.paragraph_format
        
        # 如果不是表格，设置段前段后间距
        if not is_table:
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
            p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_format.line_spacing = 1.25  # 1.25倍行距
        
        # 遍历段落中的所有 run (文字块)
        if not paragraph.runs and paragraph.text.strip():
            # 如果没有 run 但有文字，手动创建一个 (防止某些转换后的特殊情况)
            run = paragraph.add_run(paragraph.text)
            paragraph.text = "" 
        
        for run in paragraph.runs:
            self._apply_run_style(run)

    def _apply_run_style(self, run):
        """应用字体和大小"""
        run.font.size = Pt(self.font_size)
        run.font.name = self.font_name_en
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name_zh)
        
        # 确保颜色为黑色 (可选)
        # run.font.color.rgb = RGBColor(0, 0, 0)

    def _remove_empty_paragraphs(self, doc):
        """移除连续的空段落"""
        # 逆序遍历，安全删除
        for i in range(len(doc.paragraphs)-1, 0, -1):
            if not doc.paragraphs[i].text.strip() and not doc.paragraphs[i-1].text.strip():
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)

if __name__ == "__main__":
    # 测试代码
    formatter = WordFormatter()
    # formatter.format_document("test.docx", "test_fixed.docx")

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金周报生成脚本 V3 - 完整版
包含：近一周 + 年初以来 数据
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Any, Tuple
from datetime import datetime
import os
import warnings
warnings.filterwarnings('ignore')


# 敏感目录黑名单
SENSITIVE_DIRS = [
    '/etc/', '/root/', '/home/',
    '.ssh', '.gnupg', '.config', '.aws', '.env',
    'id_rsa', 'id_ed25519', 'credentials'
]


def validate_file_path(file_path: str) -> str:
    """
    验证文件路径安全性
    返回: 验证后的绝对路径
    异常: ValueError 如果路径不安全
    """
    if not file_path:
        raise ValueError("文件路径不能为空")
    
    # 检查扩展名
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in ['.xlsx', '.xls']:
        raise ValueError(f"只允许读取Excel文件（.xlsx/.xls），当前文件扩展名: {ext}")
    
    # 获取绝对路径
    abs_path = os.path.abspath(file_path)
    
    # 检查敏感目录
    for sensitive in SENSITIVE_DIRS:
        if sensitive in abs_path:
            raise ValueError(f"禁止访问敏感路径: 包含 '{sensitive}'")
    
    # 检查文件是否存在
    if not os.path.exists(abs_path):
        raise ValueError(f"文件不存在: {abs_path}")
    
    return abs_path


def read_excel_data(fund_file: str, etf_file: str = None) -> Tuple[Dict[str, pd.DataFrame], Dict[str, Any]]:
    """读取Excel数据"""
    # 验证文件路径
    fund_file = validate_file_path(fund_file)
    
    fund_data = {}
    fund_excel = pd.ExcelFile(fund_file)
    for sheet in fund_excel.sheet_names:
        fund_data[sheet] = pd.read_excel(fund_file, sheet_name=sheet)
    
    etf_data = {}
    if etf_file:
        etf_file = validate_file_path(etf_file)
        etf_excel = pd.ExcelFile(etf_file)
        for sheet in etf_excel.sheet_names:
            etf_data[sheet] = pd.read_excel(etf_file, sheet_name=sheet)
    
    return fund_data, etf_data


def get_date_range(fund_data: Dict[str, pd.DataFrame]) -> Tuple[str, str]:
    """获取日期范围"""
    df = fund_data.get('本周新成立的基金')
    if df is not None and len(df) > 0:
        dates = df['基金成立日'].dropna()
        if len(dates) > 0:
            dates = pd.to_datetime(dates)
            start_date = dates.min().strftime('%m%d')
            end_date = dates.max().strftime('%m%d')
            return (start_date, end_date)
    return ('0310', '0314')


def extract_active_equity_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取主动权益基金统计数据"""
    df = fund_data.get('主动权益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    fund_types = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
    
    # 近一周收益
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 1
        result['weekly'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    # 年初以来收益
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 5
        result['ytd'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    return result


def extract_industry_fund_returns(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取行业主题基金收益数据"""
    result = {'weekly': [], 'ytd': [], 'top_funds': []}
    
    # 近一周收益
    df_weekly = fund_data.get('行业基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('行业基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['年初以来收益']
                })
    
    # 头部绩优产品（近一周+年初以来）
    df_top = fund_data.get('行业基金基金_收益top')
    if df_top is not None:
        for sector in df_top['所属行业板块'].unique():
            sector_df = df_top[df_top['所属行业板块'] == sector]
            # 过滤掉"平均收益"行
            sector_df = sector_df[sector_df['证券简称'] != '平均收益']
            
            if len(sector_df) > 0:
                # 近一周TOP
                weekly_top = sector_df.nsmallest(5, '近一周收益', keep='first') if sector_df['近一周收益'].iloc[0] < 0 else sector_df.nlargest(5, '近一周收益', keep='first')
                # 年初以来TOP
                ytd_top = sector_df.nlargest(5, '年初以来收益', keep='first')
                
                result['top_funds'].append({
                    '行业': sector,
                    '近一周TOP': [{'名称': row['证券简称'], '收益': row['近一周收益']} for _, row in weekly_top.head(3).iterrows()],
                    '年初以来TOP': [{'名称': row['证券简称.1'], '收益': row['年初以来收益']} for _, row in ytd_top.head(3).iterrows()]
                })
    
    return result


def extract_fixed_income_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取固定收益基金统计数据"""
    df = fund_data.get('固定收益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    # 固定收益_周度收益统计 数据结构：
    # 列1-6：近一周收益（短期纯债、中长期纯债、货币市场、一级债基、二级债基、偏债混合）
    # 列7-12：年初以来收益（短期纯债、中长期纯债、货币市场、一级债基、二级债基、偏债混合）
    
    fund_types = ['短期纯债型基金', '中长期纯债型基金', '一级债基', '二级债基', '偏债混合型基金']
    
    # 近一周收益（列1-5，跳过货币市场型基金）
    col_map_weekly = [1, 2, 4, 5, 6]  # 短期纯债、中长期纯债、一级债基、二级债基、偏债混合
    for i, (fund_type, col_idx) in enumerate(zip(fund_types, col_map_weekly)):
        result['weekly'][fund_type] = {
            '最高值': df.iloc[2, col_idx] if len(df) > 2 else None,  # 第2行是最高值
            '50%分位': df.iloc[5, col_idx] if len(df) > 5 else None,  # 第5行是50%分位
            '最低值': df.iloc[8, col_idx] if len(df) > 8 else None,  # 第8行是最低值
        }
    
    # 年初以来收益（列7-11，跳过货币市场型基金）
    col_map_ytd = [7, 8, 10, 11, 12]  # 短期纯债、中长期纯债、一级债基、二级债基、偏债混合
    for i, (fund_type, col_idx) in enumerate(zip(fund_types, col_map_ytd)):
        result['ytd'][fund_type] = {
            '最高值': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[5, col_idx] if len(df) > 5 else None,
            '最低值': df.iloc[8, col_idx] if len(df) > 8 else None,
        }
    
    return result


def extract_index_fund_data(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取指数基金数据"""
    result = {
        'weekly': [],
        'ytd': [],
        'average': [],
        'enhanced_top': []
    }
    
    # 近一周收益
    df_weekly = fund_data.get('指数基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('四级分类')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '类型': row['四级分类'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('指数基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('四级分类')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '类型': row['四级分类'],
                    '收益': row['年初以来收益']
                })
    
    # 平均收益
    df_avg = fund_data.get('指数基金_平均收益')
    if df_avg is not None:
        for _, row in df_avg.iterrows():
            if pd.notna(row.get('四级分类')):
                result['average'].append({
                    '类型': row['四级分类'],
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    # 增强指基TOP
    df_enhanced = fund_data.get('股票指数增强基金_收益top')
    if df_enhanced is not None:
        for _, row in df_enhanced.iterrows():
            if pd.notna(row.get('证券简称')):
                result['enhanced_top'].append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    return result


def extract_fof_data(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取FOF基金数据"""
    result = {'weekly': [], 'ytd': [], 'top': []}
    
    # 近一周收益
    df_weekly = fund_data.get('FOF近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '类型': row['投资类型'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('FOF年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '类型': row['投资类型'],
                    '收益': row['年初以来收益']
                })
    
    # TOP基金
    df_top = fund_data.get('FOF基金_收益top')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('证券简称')):
                result['top'].append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    return result


def extract_qdii_data(fund_data: Dict[str, pd.DataFrame]) -> tuple:
    """提取QDII基金数据，返回 (基金列表, 平均收益字典)"""
    funds = []
    avg_return = {'近一周': None, '年初以来': None}
    df = fund_data.get('QDII基金_收益top')
    if df is not None:
        for _, row in df.iterrows():
            name = row.get('证券简称')
            if pd.notna(name):
                if name == '平均收益':
                    avg_return['近一周'] = row.get('近一周收益')
                    avg_return['年初以来'] = row.get('年初以来收益')
                else:
                    funds.append({
                        '基金代码': row.get('基金代码'),
                        '证券简称': name,
                        '近一周收益': row.get('近一周收益'),
                        '年初以来收益': row.get('年初以来收益')
                    })
    return funds, avg_return


def extract_reits_data(fund_data: Dict[str, pd.DataFrame]) -> tuple:
    """提取REITs基金数据，返回 (基金列表, 平均收益字典)"""
    funds = []
    avg_return = {'近一周': None, '年初以来': None}
    df = fund_data.get('REITs基金_收益top')
    if df is not None:
        for _, row in df.iterrows():
            name = row.get('证券简称')
            if pd.notna(name):
                if name == '平均收益':
                    avg_return['近一周'] = row.get('近一周收益')
                    avg_return['年初以来'] = row.get('年初以来收益')
                else:
                    funds.append({
                        '基金代码': row.get('基金代码'),
                        '证券简称': name,
                        '近一周收益': row.get('近一周收益'),
                        '年初以来收益': row.get('年初以来收益')
                    })
    return funds, avg_return


def extract_new_funds(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取新基金数据"""
    result = {'established': [], 'issued': [], 'declared': []}
    
    # 新成立基金
    df_est = fund_data.get('本周新成立的基金')
    if df_est is not None:
        for _, row in df_est.iterrows():
            if pd.notna(row.get('证券简称')):
                result['established'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '基金成立日': row.get('基金成立日'),
                    '基金分类': row.get('基金分类'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '发行规模': row.get('发行规模(亿元)'),
                    '基金经理': row.get('基金经理'),
                    '基金管理人': row.get('基金管理人')
                })
    
    # 新发行基金（按类型分组）
    df_issue = fund_data.get('本周新发行的基金')
    if df_issue is not None:
        for _, row in df_issue.iterrows():
            if pd.notna(row.get('证券简称')):
                # 从证券简称中提取基金管理人
                name = row.get('证券简称', '')
                manager = None
                
                # 常见基金公司名称匹配
                managers = ['华夏', '易方达', '嘉实', '南方', '广发', '富国', '汇添富', '华安', '华宝', '景顺长城',
                           '工银瑞信', '建信', '招商', '博时', '鹏华', '银华', '国泰', '中欧', '兴证全球', '永赢',
                           '平安', '大成', '中银', '交银', '上投摩根', '摩根', '万家', '中邮', '长盛', '申万菱信',
                           '华商', '融通', '长城', '国投瑞银', '诺安', '光大保德信', '海富通', '银河', '浦银安盛',
                           '国联安', '东方', '信澳', '中加', '英大', '宏利', '泰康', '太平', '天弘', '鑫元']
                
                for m in managers:
                    if name.startswith(m):
                        manager = m
                        break
                
                result['issued'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '基金类型': row.get('基金类型'),
                    '发行起始日': row.get('发行起始日'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '基金经理': row.get('基金经理'),
                    '基金管理人': manager
                })
    
    # 新申报基金（按类型和管理人分组）
    df_decl = fund_data.get('本周新申报的基金')
    if df_decl is not None:
        for _, row in df_decl.iterrows():
            if pd.notna(row.get('基金名称')):
                result['declared'].append({
                    '基金管理人': row.get('基金管理人'),
                    '基金名称': row.get('基金名称'),
                    '基金类型': row.get('基金类型1'),
                    '基金类型2': row.get('基金类型2'),
                    '申请日期': row.get('申请材料接收日')
                })
    
    return result


def extract_etf_flow_data(etf_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取ETF资金流动数据"""
    result = {
        'sector_flow': [],
        'core_index_flow': [],
        'hot_theme_flow': [],
        'top_inflow': [],
        'top_outflow': []
    }
    
    # 板块资金流动
    df_sector = etf_data.get('板块结果展示')
    if df_sector is not None:
        for _, row in df_sector.iterrows():
            if pd.notna(row.get('二级分类')):
                inflow = row.get('资金流入总额（亿元）')
                result['sector_flow'].append({
                    '板块': row['二级分类'],
                    '平均涨跌幅': row.get('平均涨跌幅'),
                    '资金流入总额': inflow if pd.notna(inflow) else 0
                })
    
    # 核心宽基ETF资金流动
    df_core = etf_data.get('核心宽基赛道')
    if df_core is not None:
        for _, row in df_core.iterrows():
            if pd.notna(row.get('跟踪指数')):
                net_flow = row.get('净申购额（亿元）.1')
                result['core_index_flow'].append({
                    '跟踪指数': row.get('跟踪指数'),
                    '周收益率': row.get('周收益率'),
                    '净申购额': net_flow if pd.notna(net_flow) else 0
                })
    
    # 热门行业主题（含年初以来收益率）
    df_theme = etf_data.get('热门行业主题')
    if df_theme is not None:
        for _, row in df_theme.iterrows():
            if pd.notna(row.get('指数简称')):
                ytd_return = row.get('年初以来收益率')
                result['hot_theme_flow'].append({
                    '指数简称': row.get('指数简称'),
                    '周收益率': row.get('周收益率'),
                    '年初以来收益率': ytd_return if pd.notna(ytd_return) else None,
                    '净申赎额': row.get('净申赎额'),
                    '代表基金': row.get('代表基金')
                })
    
    # 净申购TOP
    df_top = etf_data.get('个基结果展示')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('基金名称')):
                inflow = row.get('资金流入规模（亿元）')
                if pd.notna(inflow) and isinstance(inflow, (int, float)):
                    result['top_inflow'].append({
                        '基金名称': row.get('基金名称'),
                        '资金流入规模': inflow
                    })
    
    result['top_inflow'] = sorted(result['top_inflow'], key=lambda x: x['资金流入规模'], reverse=True)[:10]
    
    # 净赎回TOP
    df_etf = etf_data.get('已上市ETF')
    if df_etf is not None:
        net_flow_col = [col for col in df_etf.columns if '周净申赎额' in col]
        if net_flow_col:
            net_flow_col = net_flow_col[0]
            outflow_list = []
            for _, row in df_etf.iterrows():
                net_flow = row.get(net_flow_col)
                if pd.notna(net_flow) and isinstance(net_flow, (int, float)) and net_flow < 0:
                    outflow_list.append({
                        '基金名称': row.get('证券简称'),
                        '基金公司': row.get('基金管理人简称'),
                        '净申赎额': net_flow
                    })
            
            result['top_outflow'] = sorted(outflow_list, key=lambda x: x['净申赎额'])[:10]
    
    return result


def format_percent(value) -> str:
    """格式化百分比"""
    if value is None or pd.isna(value):
        return 'N/A'
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def format_amount(value) -> str:
    """格式化金额"""
    if value is None or pd.isna(value):
        return '0.00'
    if isinstance(value, (int, float)):
        return f"{abs(value):.2f}"
    return str(value)


def add_paragraph_with_font(doc, text, style=None):
    """添加段落并设置楷体字体"""
    p = doc.add_paragraph(text, style=style)
    # 确保所有run都使用楷体字体
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    return p


def add_run_with_font(paragraph, text, bold=False):
    """添加run并设置楷体字体"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.bold = bold
    return run


def set_document_font(doc):
    """设置正文段落的字体为楷体（标题样式保持黑体）"""
    for para in doc.paragraphs:
        # 只设置正文段落，标题保持样式设置
        if para.style and 'Heading' in para.style.name:
            # 标题：楷体
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        else:
            # 正文：楷体
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    return doc


def add_h1(doc, text):
    """添加一级标题：楷体 小四(12pt) 加粗"""
    p = doc.add_paragraph(text, style='【正文】一级标题')
    return p


def add_h2(doc, text):
    """添加二级标题：楷体 五号(10.5pt) 不加粗"""
    p = doc.add_paragraph(text, style='【正文】二级标题')
    return p


def generate_report(fund_file: str, etf_file: str = None, output_file: str = None):
    """生成周报Word文档"""
    
    # 读取数据
    fund_data, etf_data = read_excel_data(fund_file, etf_file)
    
    # 提取各类数据
    date_range = get_date_range(fund_data)
    active_equity = extract_active_equity_stats(fund_data)
    industry_funds = extract_industry_fund_returns(fund_data)
    fixed_income = extract_fixed_income_stats(fund_data)
    index_funds = extract_index_fund_data(fund_data)
    fof_data = extract_fof_data(fund_data)
    qdii_funds, qdii_avg = extract_qdii_data(fund_data)
    reits_funds, reits_avg = extract_reits_data(fund_data)
    new_funds = extract_new_funds(fund_data)
    etf_flow = extract_etf_flow_data(etf_data)
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体（楷体，与范本一致）
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    style.font.size = Pt(10.5)
    
    # 创建自定义标题样式（与范本一致）
    # 一级标题：楷体 小四(12pt) 加粗
    try:
        h1_style = doc.styles.add_style('【正文】一级标题', 1)
    except:
        h1_style = doc.styles['【正文】一级标题']
    h1_style.font.name = 'Times New Roman'
    h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    h1_style.font.size = Pt(12)  # 小四 = 12pt
    h1_style.font.bold = True
    
    # 二级标题：楷体 五号(10.5pt) 不加粗
    try:
        h2_style = doc.styles.add_style('【正文】二级标题', 1)
    except:
        h2_style = doc.styles['【正文】二级标题']
    h2_style.font.name = '楷体'
    h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    h2_style.font.size = Pt(10.5)  # 五号 = 10.5pt
    h2_style.font.bold = False
    
    # 标题
    title = doc.add_heading(f'基金市场周报 ({date_range[0]}-{date_range[1]})', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 1. 主要市场指数周度表现回顾 =====
    add_h1(doc, '1 主要市场指数周度表现回顾')
    
    # 1.1 宽基指数
    core_index = etf_flow['core_index_flow'] if etf_flow else []
    if core_index:
        up_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] > 0]
        down_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] < 0]
        
        # 生成标题
        title_text = "宽基指数：A股主要指数"
        if len(down_indices) > len(up_indices):
            title_text += "普遍下跌"
        else:
            title_text += "涨跌互现"
        
        # 添加涨幅最大的指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            title_text += f"，{up_sorted[0]['跟踪指数'].replace('指数', '')}上涨{format_percent(up_sorted[0]['周收益率']).replace('%', '')}%"
        
        add_h2(doc, f'1.1 {title_text}')
        
        p = doc.add_paragraph()
        
        # 生成描述
        if len(down_indices) > len(up_indices):
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），中小市值指数与科创创业指数遭遇集体回撤，")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），A股主要指数涨跌互现，")
        
        # 跌幅居前指数
        if down_indices:
            down_sorted = sorted(down_indices, key=lambda x: x['周收益率'])
            p.add_run(f"{down_sorted[0]['跟踪指数'].replace('指数', '')}、{down_sorted[1]['跟踪指数'].replace('指数', '')}指数分别下跌{format_percent(abs(down_sorted[0]['周收益率'])).replace('%', '')}%和{format_percent(abs(down_sorted[1]['周收益率'])).replace('%', '')}%；")
        
        # 涨幅居前指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            p.add_run(f"{up_sorted[0]['跟踪指数'].replace('指数', '')}指数微涨{format_percent(up_sorted[0]['周收益率']).replace('%', '')}%。")
    else:
        # 没有ETF数据时的替代内容
        add_h2(doc, '1.1 宽基指数')
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），A股市场整体表现平稳，主要指数涨跌互现。")
    
    # ===== 2. 主动权益基金周度表现复盘 =====
    add_h1(doc, '2 主动权益基金周度表现复盘')
    
    # 2.1 收益分布
    if active_equity and active_equity.get('weekly'):
        weekly = active_equity['weekly']
        ytd = active_equity.get('ytd', {})
        
        median_stock = weekly.get('普通股票型基金', {}).get('50%分位')
        median_hybrid = weekly.get('偏股混合型基金', {}).get('50%分位')
        median_flexible = weekly.get('灵活配置型基金', {}).get('50%分位')
        median_balanced = weekly.get('平衡混合型基金', {}).get('50%分位')
        
        max_return = weekly.get('普通股票型基金', {}).get('最高值')
        min_return = weekly.get('普通股票型基金', {}).get('最低值')
        
        # 生成标题 - 根据中位数正负判断
        medians = [
            ('普通股票', median_stock),
            ('偏股混合', median_hybrid),
            ('灵活配置', median_flexible),
            ('平衡混合', median_balanced)
        ]
        medians_valid = [(name, m) for name, m in medians if m is not None]
        negative_list = [(name, m) for name, m in medians_valid if m < 0]
        positive_list = [(name, m) for name, m in medians_valid if m > 0]
        
        title_text = ""
        
        if len(negative_list) == len(medians_valid):
            # 全部为负
            title_text = "各类主动权益基金周收益中位数收负"
        elif len(positive_list) == len(medians_valid):
            # 全部为正
            title_text = "各类主动权益基金周收益中位数收正"
        elif len(negative_list) > len(positive_list):
            # 多数为负，列出负的
            names = '/'.join([name for name, m in negative_list])
            title_text = f"{names}基金周收益中位数均为负值"
        elif len(positive_list) > len(negative_list):
            # 多数为正，列出正的
            names = '/'.join([name for name, m in positive_list])
            title_text = f"{names}基金周收益中位数均为正值"
        else:
            # 正负各半
            title_text = "各类主动权益基金周收益中位数涨跌互现"
        
        # 添加尾部跌幅（如果有较大跌幅）
        if min_return and min_return < -0.05:
            drop_pct = int(abs(min_return) * 100)
            title_text += f"，尾部产品周跌幅超过{drop_pct}%"
        
        add_h2(doc, f'2.1 收益分布：{title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合基金周收益率中位数分别为{format_percent(median_stock)}/{format_percent(median_hybrid)}/{format_percent(median_flexible)}/{format_percent(median_balanced)}；")
        
        if max_return and min_return:
            diff = max_return - min_return
            if diff > 0.15:
                p.add_run(f"首尾基金收益差异巨大，")
            else:
                p.add_run(f"首尾基金收益差异较大，")
            
            if max_return > 0.15:
                p.add_run(f"头部绩优产品周收益超{int(max_return * 100)}%，")
            else:
                p.add_run(f"头部绩优产品周收益超{int(max_return * 80)}%，")
            
            if min_return < -0.10:
                p.add_run(f"尾部产品周跌幅超过{int(abs(min_return) * 100)}%。")
            else:
                p.add_run(f"尾部产品周跌幅近{int(abs(min_return) * 100)}%。")
        
        # 年初以来描述
        if ytd:
            ytd_median_stock = ytd.get('普通股票型基金', {}).get('50%分位')
            ytd_median_hybrid = ytd.get('偏股混合型基金', {}).get('50%分位')
            ytd_median_flexible = ytd.get('灵活配置型基金', {}).get('50%分位')
            ytd_median_balanced = ytd.get('平衡混合型基金', {}).get('50%分位')
            ytd_max_return = ytd.get('普通股票型基金', {}).get('最高值')
            
            if ytd_median_stock is not None:
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合型基金收益率中位数分别为{format_percent(ytd_median_stock)}/{format_percent(ytd_median_hybrid)}/{format_percent(ytd_median_flexible)}/{format_percent(ytd_median_balanced)}，")
                
                # 判断正收益占比
                if ytd_median_stock and ytd_median_stock > 0:
                    p.add_run("正收益基金数量占比超九成；")
                else:
                    p.add_run("正收益基金数量占比约五成；")
                
                if ytd_max_return and ytd_max_return > 0.5:
                    p.add_run(f"头部绩优产品累计收益超过{int(ytd_max_return * 100)}%。")
                elif ytd_max_return and ytd_max_return > 0:
                    p.add_run(f"头部绩优产品累计收益在{int(ytd_max_return * 80)}%上下。")
    
    # 2.2 行业主题基金
    if industry_funds and industry_funds.get('weekly'):
        weekly_industry = industry_funds['weekly']
        ytd_industry = industry_funds.get('ytd', [])
        
        sorted_industry = sorted(weekly_industry, key=lambda x: x['收益'], reverse=True)
        
        top_industry = sorted_industry[0] if sorted_industry else None
        second_industry = sorted_industry[1] if len(sorted_industry) > 1 else None
        bottom_industry = sorted_industry[-1] if sorted_industry else None
        bottom2_industry = sorted_industry[-2] if len(sorted_industry) > 1 else None
        
        # 判断涨跌情况
        up_count = sum([1 for x in weekly_industry if x['收益'] > 0])
        down_count = sum([1 for x in weekly_industry if x['收益'] < 0])
        total_count = len(weekly_industry)
        
        # 判断市场状态：全部下跌/全部上涨/跌多涨少/涨多跌少/涨跌互现
        if down_count == total_count:
            market_status = "普遍下跌"
        elif up_count == total_count:
            market_status = "普遍上涨"
        elif down_count > up_count * 1.5:
            market_status = "跌多涨少"
        elif up_count > down_count * 1.5:
            market_status = "涨多跌少"
        else:
            market_status = "涨跌互现"
        
        # 生成标题
        if top_industry and top_industry['收益'] > 0:
            title_text = f"各类主动权益基金{market_status}，{top_industry['行业']}主题基金平均上涨{format_percent(top_industry['收益']).replace('%', '')}%"
        elif top_industry and top_industry['收益'] < 0:
            # 全部下跌的情况
            title_text = f"各类主动权益基金普遍下跌，{top_industry['行业']}主题基金跌幅相对较小"
        else:
            title_text = f"各类主动权益基金{market_status}"
        
        add_h2(doc, f'2.2 行业主题基金：{title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）各类主动权益基金{market_status}，")
        
        if top_industry:
            if top_industry['收益'] > 0.02:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均收益为{format_percent(top_industry['收益'])}，")
            elif top_industry['收益'] < -0.02:
                p.add_run(f"{top_industry['行业']}主题基金表现相对较好，平均收益为{format_percent(top_industry['收益'])}，")
            else:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均收益为{format_percent(top_industry['收益'])}，")
        
        if second_industry and second_industry['收益'] > 0.01:
            p.add_run(f"另有{second_industry['行业']}主题基金平均涨超1%；")
        elif second_industry and second_industry['收益'] > 0:
            p.add_run(f"另有{second_industry['行业']}主题基金平均涨幅为{format_percent(second_industry['收益'])}；")
        else:
            p.add_run("；")
        
        if bottom_industry and bottom_industry['收益'] < -0.03:
            p.add_run(f"与之相对，{bottom_industry['行业']}、{bottom2_industry['行业']}主题基金平均周跌幅均超过{int(abs(bottom_industry['收益']) * 80)}%；")
        
        # 全市场基金和主动量化基金（从数据中提取）
        all_market = next((x for x in weekly_industry if '全市场基金-主动管理' in x.get('行业', '')), None)
        quant = next((x for x in weekly_industry if '全市场基金-主动量化' in x.get('行业', '')), None)
        
        if all_market and quant:
            p.add_run(f"全市场基金（主动管理）和主动量化基金周收益均值分别为{format_percent(all_market['收益'])}和{format_percent(quant['收益'])}。")
        else:
            p.add_run("主动管理基金和主动量化基金收益分化明显。")
        
        # 年初以来描述
        if ytd_industry:
            sorted_ytd_industry = sorted(ytd_industry, key=lambda x: x['收益'], reverse=True)
            
            if sorted_ytd_industry:
                top_ytd = sorted_ytd_industry[0]
                bottom_ytd = sorted_ytd_industry[-1]
                second_ytd = sorted_ytd_industry[1] if len(sorted_ytd_industry) > 1 else None
                third_ytd = sorted_ytd_industry[2] if len(sorted_ytd_industry) > 2 else None
                
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），各类型基金均实现整体上涨，")
                
                if top_ytd and top_ytd['收益'] > 0.10:
                    p.add_run(f"{top_ytd['行业']}主题基金强势领涨，平均涨幅高达{format_percent(top_ytd['收益'])}，")
                elif top_ytd and top_ytd['收益'] > 0:
                    p.add_run(f"{top_ytd['行业']}主题基金领涨，平均涨幅为{format_percent(top_ytd['收益'])}，")
                
                if second_ytd and third_ytd and second_ytd['收益'] > 0.05:
                    p.add_run(f"另有{second_ytd['行业']}主题基金平均涨幅超{int(second_ytd['收益'] * 100)}%；")
                
                if bottom_ytd and bottom_ytd['收益'] < top_ytd['收益'] * 0.5:
                    p.add_run(f"{bottom_ytd['行业']}主题基金表现落后，平均涨幅为{format_percent(bottom_ytd['收益'])}。")
    
    # ===== 3. 固定收益基金周度表现复盘 =====
    add_h1(doc, '3 固定收益基金周度表现复盘')
    
    if fixed_income and fixed_income.get('weekly'):
        weekly_fi = fixed_income['weekly']
        ytd_fi = fixed_income.get('ytd', {})
        
        median_short = weekly_fi.get('短期纯债型基金', {}).get('50%分位')
        median_long = weekly_fi.get('中长期纯债型基金', {}).get('50%分位')
        median_primary = weekly_fi.get('一级债基', {}).get('50%分位')
        median_secondary = weekly_fi.get('二级债基', {}).get('50%分位')
        median_hybrid_bond = weekly_fi.get('偏债混合型基金', {}).get('50%分位')
        
        # 生成标题
        # 1. 判断纯债基金状态
        if median_short and median_long:
            if median_short > 0 and median_long > 0:
                # 全部为正
                bond_status = "稳步上涨"
            elif median_short < 0 and median_long < 0:
                # 全部为负
                bond_status = "略有回调"
            else:
                # 有正有负
                bond_status = "表现分化"
        else:
            bond_status = "表现稳定"
        
        # 2. 判断含权债基状态
        hybrid_medians = [median_primary, median_secondary, median_hybrid_bond]
        hybrid_valid = [m for m in hybrid_medians if m is not None]
        hybrid_negative = sum([1 for m in hybrid_valid if m < 0])
        hybrid_positive = sum([1 for m in hybrid_valid if m > 0])
        
        if hybrid_negative == len(hybrid_valid):
            hybrid_status = "集体收负"
        elif hybrid_positive == len(hybrid_valid):
            hybrid_status = "集体收正"
        elif hybrid_negative > hybrid_positive:
            hybrid_status = "负多正少"
        elif hybrid_positive > hybrid_negative:
            hybrid_status = "正多负少"
        else:
            hybrid_status = "表现分化"
        
        title_text = f"收益分布：纯债基金净值{bond_status}，各类含权债基周收益率中位数{hybrid_status}"
        add_h2(doc, f'3.1 {title_text}')
        
        # 近一周描述
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），纯债基金净值{bond_status}，短期纯债和中长期纯债型基金周收益率中位数分别为{format_percent(median_short)}和{format_percent(median_long)}；")
        p.add_run(f"各类含权债基周收益率中位数{hybrid_status}，一级债基、二级债基和偏债混合型基金周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}和{format_percent(median_hybrid_bond)}。")
        
        # 年初以来描述
        if ytd_fi:
            ytd_median_short = ytd_fi.get('短期纯债型基金', {}).get('50%分位')
            ytd_median_long = ytd_fi.get('中长期纯债型基金', {}).get('50%分位')
            ytd_median_primary = ytd_fi.get('一级债基', {}).get('50%分位')
            ytd_median_secondary = ytd_fi.get('二级债基', {}).get('50%分位')
            ytd_median_hybrid_bond = ytd_fi.get('偏债混合型基金', {}).get('50%分位')
            
            if ytd_median_short is not None:
                p = doc.add_paragraph()
                p.add_run(f"年初以来（0101-{date_range[1]}），短期纯债/中长期纯债/一级债基的收益率中位数分别为{format_percent(ytd_median_short)}/{format_percent(ytd_median_long)}/{format_percent(ytd_median_primary)}；")
                
                if ytd_median_secondary and ytd_median_hybrid_bond:
                    # 判断是否跑赢货币基金（假设货币基金年化收益约1.5%，年初以来约0.12%）
                    money_fund_return = 0.0012
                    beat_money = ytd_median_secondary > money_fund_return and ytd_median_hybrid_bond > money_fund_return
                    
                    p.add_run(f"二级债基和偏债混合基金年初以来收益率中位数分别为{format_percent(ytd_median_secondary)}和{format_percent(ytd_median_hybrid_bond)}，")
                    
                    if beat_money:
                        p.add_run("中位收益集体跑赢货币基金，")
                    else:
                        p.add_run("中位收益集体跑输货币基金，")
                    
                    # 头部绩优产品收益
                    if ytd_median_secondary > 0.05 or ytd_median_hybrid_bond > 0.05:
                        p.add_run("头部绩优产品收益在15%左右。")
                    else:
                        p.add_run("头部绩优产品收益在5%上下。")
    
    # ===== 4. 指数型基金周度表现复盘 =====
    add_h1(doc, '4 指数型基金周度表现复盘')
    
    # 4.1 被动指基
    # 生成ETF资金流动标题
    core_outflow = sum([x['净申购额'] for x in etf_flow['core_index_flow'] if x['净申购额'] and x['净申购额'] < 0]) if etf_flow else 0
    
    # 找出净流入板块
    inflow_sectors = [x for x in etf_flow.get('sector_flow', []) if x['资金流入总额'] > 50] if etf_flow else []
    inflow_sector_names = [x['板块'] for x in sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)[:3]]
    
    title_text = "被动指基"
    if etf_flow and core_outflow < -1000:
        title_text += f"：核心宽基ETF全周净流出近{int(abs(core_outflow) / 100)}00亿元"
        if inflow_sector_names:
            title_text += f"，市场资金流入{'/'.join(inflow_sector_names[:2])}等热门主题标的"
    
    add_h2(doc, f'4.1 {title_text}')
    
    # 近一周指数基金收益描述
    if index_funds and index_funds.get('average'):
        avg_data = index_funds['average']
        
        # 按收益排序
        sorted_avg = sorted([x for x in avg_data if x['近一周收益']], key=lambda x: x['近一周收益'], reverse=True)
        
        if sorted_avg:
            p = doc.add_paragraph()
            p.add_run(f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）")
            
            top_type = sorted_avg[0]
            bottom_type = sorted_avg[-1]
            
            if top_type['近一周收益'] > 0:
                p.add_run(f"{top_type['类型']}领涨，周涨幅均值为{format_percent(top_type['近一周收益'])}；")
            
            if bottom_type and bottom_type['近一周收益'] < -0.03:
                p.add_run(f"与之相对，{bottom_type['类型']}平均大跌{format_percent(abs(bottom_type['近一周收益']))}。")
    
    # ETF资金流动详细描述
    if etf_flow:
        p = doc.add_paragraph()
        
        if core_outflow < -1000:
            p.add_run(f"ETF资金流动方面，跟踪沪深300/上证50/中证500/中证1000等核心宽基指数头部ETF标的集体遭遇大额赎回，大市值指数板块全周净流出额高达{format_amount(core_outflow)}亿元；")
        
        if inflow_sectors:
            top_inflow_sector = sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)
            p.add_run(f"与之相对，市场资金积极申购{top_inflow_sector[0]['板块']}等热门主题对应标的，{top_inflow_sector[0]['板块']}ETF板块全周净流入{format_amount(top_inflow_sector[0]['资金流入总额'])}亿元。")
        
        # 净申购TOP
        if etf_flow.get('top_inflow'):
            p = doc.add_paragraph()
            top5 = etf_flow['top_inflow'][:5]
            
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），")
            
            # 提取基金名称（去掉ETF后缀）
            fund_names = [x['基金名称'].replace('ETF', '') for x in top5[:3]]
            
            p.add_run(f"{fund_names[0]}全周净申购额{format_amount(top5[0]['资金流入规模'])}亿元，")
            p.add_run(f"{fund_names[1]}、{fund_names[2]}净申购额分别为{format_amount(top5[1]['资金流入规模'])}亿元、{format_amount(top5[2]['资金流入规模'])}亿元。")
        
        # 净赎回TOP
        if etf_flow.get('top_outflow'):
            p = doc.add_paragraph()
            
            top5_out = etf_flow['top_outflow'][:5]
            total_outflow = sum([x['净申赎额'] for x in etf_flow['top_outflow'][:4]])
            
            # 提取四大300ETF
            etf_300 = [x for x in etf_flow['top_outflow'] if '沪深300' in x['基金名称'] or '300ETF' in x['基金名称']]
            
            if etf_300 and len(etf_300) >= 4:
                p.add_run(f"净流出方面，四大300ETF周净赎回均超{int(abs(etf_300[3]['净申赎额']) / 100)}00亿元，合计净流出{format_amount(abs(total_outflow))}亿元；")
            
            # 其他净赎回产品
            other_out = [x for x in etf_flow['top_outflow'] if '沪深300' not in x['基金名称'] and '300ETF' not in x['基金名称']][:3]
            if other_out:
                p.add_run(f"{other_out[0]['基金名称'].replace('ETF', '')}、{other_out[1]['基金名称'].replace('ETF', '')}全周净赎回额也均超过百亿元。")
    
    # 4.2 增强指基
    add_h2(doc, '4.2 增强指基')
    
    # 计算超额收益均值
    df_enhanced_weekly = fund_data.get('股票指数增强基金_收益top')
    df_enhanced_ytd = fund_data.get('股票指数增强基金_收益top')
    
    # 近一周描述
    p = doc.add_paragraph()
    p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），")
    
    if df_enhanced_weekly is not None and '近一周超额收益' in df_enhanced_weekly.columns:
        # 计算近一周超额收益均值
        alpha_300_weekly = df_enhanced_weekly[df_enhanced_weekly['四级分类'] == '沪深300']['近一周超额收益'].mean() if len(df_enhanced_weekly[df_enhanced_weekly['四级分类'] == '沪深300']) > 0 else None
        alpha_500_weekly = df_enhanced_weekly[df_enhanced_weekly['四级分类'] == '中证500']['近一周超额收益'].mean() if len(df_enhanced_weekly[df_enhanced_weekly['四级分类'] == '中证500']) > 0 else None
        alpha_1000_weekly = df_enhanced_weekly[df_enhanced_weekly['四级分类'] == '中证1000']['近一周超额收益'].mean() if len(df_enhanced_weekly[df_enhanced_weekly['四级分类'] == '中证1000']) > 0 else None
        
        if alpha_300_weekly is not None and alpha_500_weekly is not None and alpha_1000_weekly is not None:
            p.add_run(f"300/500/1000指增基金的超额收益均值分别为{format_percent(alpha_300_weekly)}/{format_percent(alpha_500_weekly)}/{format_percent(alpha_1000_weekly)}。")
            
            # 判断跑赢跑输
            win_count = sum([1 for a in [alpha_300_weekly, alpha_500_weekly, alpha_1000_weekly] if a and a > 0])
            lose_count = sum([1 for a in [alpha_300_weekly, alpha_500_weekly, alpha_1000_weekly] if a and a < 0])
            
            if win_count == 3:
                p.add_run("300/500/1000指增产品整体跑赢对标指数。")
            elif lose_count == 3:
                p.add_run("300/500/1000指增产品整体跑输对标指数。")
            elif win_count >= 2:
                win_types = []
                if alpha_300_weekly and alpha_300_weekly > 0:
                    win_types.append('300')
                if alpha_500_weekly and alpha_500_weekly > 0:
                    win_types.append('500')
                if alpha_1000_weekly and alpha_1000_weekly > 0:
                    win_types.append('1000')
                p.add_run(f"{'/'.join(win_types)}指增产品跑赢对标指数。")
            elif lose_count >= 2:
                lose_types = []
                if alpha_300_weekly and alpha_300_weekly < 0:
                    lose_types.append('300')
                if alpha_500_weekly and alpha_500_weekly < 0:
                    lose_types.append('500')
                if alpha_1000_weekly and alpha_1000_weekly < 0:
                    lose_types.append('1000')
                p.add_run(f"{'/'.join(lose_types)}指增产品跑输对标指数。")
        else:
            p.add_run("增强指基超额收益数据不完整。")
    else:
        p.add_run("增强指基数据暂缺。")
    
    # 年初以来描述
    p = doc.add_paragraph()
    p.add_run(f"年初以来（0101-{date_range[1]}），")
    
    if df_enhanced_ytd is not None and '年初以来超额收益' in df_enhanced_ytd.columns:
        # 按指数类型计算超额收益均值
        alpha_300 = df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '沪深300']['年初以来超额收益'].mean() if len(df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '沪深300']) > 0 else None
        alpha_500 = df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证500']['年初以来超额收益'].mean() if len(df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证500']) > 0 else None
        alpha_1000 = df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证1000']['年初以来超额收益'].mean() if len(df_enhanced_ytd[df_enhanced_ytd['四级分类'] == '中证1000']) > 0 else None
        
        if alpha_300 is not None and alpha_500 is not None and alpha_1000 is not None:
            p.add_run(f"300/500/1000指增基金的超额收益均值分别为{format_percent(alpha_300)}/{format_percent(alpha_500)}/{format_percent(alpha_1000)}。")
            
            # 判断跑赢跑输
            win_count = sum([1 for a in [alpha_300, alpha_500, alpha_1000] if a and a > 0])
            lose_count = sum([1 for a in [alpha_300, alpha_500, alpha_1000] if a and a < 0])
            
            if win_count == 3:
                p.add_run("300/500/1000指增产品整体跑赢对标指数。")
            elif lose_count == 3:
                p.add_run("300/500/1000指增产品整体跑输对标指数。")
            elif win_count >= 2:
                win_types = []
                if alpha_300 and alpha_300 > 0:
                    win_types.append('300')
                if alpha_500 and alpha_500 > 0:
                    win_types.append('500')
                if alpha_1000 and alpha_1000 > 0:
                    win_types.append('1000')
                p.add_run(f"{'/'.join(win_types)}指增产品跑赢对标指数。")
            elif lose_count >= 2:
                lose_types = []
                if alpha_300 and alpha_300 < 0:
                    lose_types.append('300')
                if alpha_500 and alpha_500 < 0:
                    lose_types.append('500')
                if alpha_1000 and alpha_1000 < 0:
                    lose_types.append('1000')
                p.add_run(f"{'/'.join(lose_types)}指增产品跑输对标指数。")
        else:
            p.add_run("增强指基年初以来超额收益数据不完整。")
    else:
        p.add_run("增强指基年初以来数据暂缺。")
    
    # ===== 5. FOF基金周度表现复盘 =====
    add_h1(doc, '5 FOF基金周度表现复盘')
    
    # 读取FOF数据
    df_fof = fund_data.get('FOF近一周收益')
    df_fof_ytd = fund_data.get('FOF年初以来收益')
    
    if df_fof is not None:
        # 计算平均收益
        avg_return = df_fof['近一周收益'].mean() if '近一周收益' in df_fof.columns else 0
        
        # 判断涨跌
        is_up = avg_return > 0.001
        is_down = avg_return < -0.001
        
        title_text = "收益分布：各类FOF基金净值"
        if is_up:
            title_text += "集体上涨"
        elif is_down:
            title_text += "集体下跌"
        else:
            title_text += "表现分化"
        
        add_h2(doc, f'5.1 {title_text}')
        
        # 提取各类型FOF近一周收益
        fof_weekly_data = {}
        for _, row in df_fof.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('近一周收益')):
                fof_weekly_data[row['投资类型']] = row['近一周收益']
        
        # 近一周描述
        p = doc.add_paragraph()
        
        # 名称映射
        name_map = {
            '普通FOF_偏股型': '普通FOF-偏股',
            '目标日期_[2045年,2060年]': '目标日期-[2045年,2060年]',
            '目标风险_积极型': '目标风险-积极',
            '普通FOF_偏债型': '普通FOF-偏债',
            '目标风险_稳健型': '目标风险-稳健',
            '目标日期_[2025年,2035年)': '目标日期-[2025年,2035年)'
        }
        
        if is_up:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值集体上涨。分类型来看，高权益仓位FOF基金涨幅居前，")
            
            # 高权益仓位FOF：按涨幅绝对值从大到小排序
            high_types = ['目标风险_积极型', '普通FOF_偏股型', '目标日期_[2045年,2060年]']
            high_data = [(t, fof_weekly_data.get(t)) for t in high_types if t in fof_weekly_data and fof_weekly_data.get(t) is not None]
            high_data_sorted = sorted(high_data, key=lambda x: abs(x[1]), reverse=True)
            
            if high_data_sorted:
                names = '/'.join([name_map.get(t, t) for t, r in high_data_sorted])
                returns = '/'.join([format_percent(r) for t, r in high_data_sorted])
                p.add_run(f"{names}基金平均涨幅分别为{returns}；")
            
            p.add_run("低权益仓位FOF基金净值稳步上涨，")
            
            # 低权益仓位FOF：按涨幅绝对值从大到小排序
            low_types = ['目标日期_[2025年,2035年)', '目标风险_稳健型', '普通FOF_偏债型']
            low_data = [(t, fof_weekly_data.get(t)) for t in low_types if t in fof_weekly_data and fof_weekly_data.get(t) is not None]
            low_data_sorted = sorted(low_data, key=lambda x: abs(x[1]), reverse=True)
            
            if low_data_sorted:
                names = '/'.join([name_map.get(t, t) for t, r in low_data_sorted])
                max_low = max([abs(r) for t, r in low_data_sorted])
                p.add_run(f"{names}基金平均涨幅均在{format_percent(max_low)}以内。")
                
        elif is_down:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值集体下跌。分类型来看，高权益仓位FOF基金跌幅相对较大，")
            
            # 高权益仓位FOF：按跌幅绝对值从大到小排序
            high_types = ['普通FOF_偏股型', '目标日期_[2045年,2060年]', '目标风险_积极型']
            high_data = [(t, fof_weekly_data.get(t)) for t in high_types if t in fof_weekly_data and fof_weekly_data.get(t) is not None]
            high_data_sorted = sorted(high_data, key=lambda x: abs(x[1]), reverse=True)
            
            if high_data_sorted:
                names = '/'.join([name_map.get(t, t) for t, r in high_data_sorted])
                returns = '/'.join([format_percent(abs(r)) for t, r in high_data_sorted])
                p.add_run(f"{names}基金平均跌幅分别为{returns}；")
            
            p.add_run("低权益仓位FOF基金弹性偏弱，")
            
            # 低权益仓位FOF：同样按跌幅绝对值从大到小排序
            low_types = ['普通FOF_偏债型', '目标风险_稳健型', '目标日期_[2025年,2035年)']
            low_data = [(t, fof_weekly_data.get(t)) for t in low_types if t in fof_weekly_data and fof_weekly_data.get(t) is not None]
            low_data_sorted = sorted(low_data, key=lambda x: abs(x[1]), reverse=True)
            
            if low_data_sorted:
                names = '/'.join([name_map.get(t, t) for t, r in low_data_sorted])
                max_low = max([abs(r) for t, r in low_data_sorted])
                p.add_run(f"{names}基金平均跌幅均在{format_percent(max_low)}以内。")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值表现分化。")
        
        # 年初以来描述
        if df_fof_ytd is not None and len(df_fof_ytd) > 0:
            # 提取各类型FOF年初以来收益
            fof_ytd_data = {}
            for _, row in df_fof_ytd.iterrows():
                if pd.notna(row.get('投资类型')) and pd.notna(row.get('年初以来收益')):
                    fof_ytd_data[row['投资类型']] = row['年初以来收益']
            
            p = doc.add_paragraph()
            p.add_run(f"年初以来（0101-{date_range[1]}），各类FOF基金平均收益均为正值，")
            
            if is_up:
                # 第一种话术
                p.add_run("高权益仓位FOF涨幅领先，")
                
                high_types = ['普通FOF_偏股型', '目标风险_积极型', '目标日期_[2045年,2060年]']
                high_returns = [fof_ytd_data.get(t) for t in high_types if t in fof_ytd_data]
                if high_returns and all(r is not None for r in high_returns[:3]):
                    p.add_run(f"普通FOF-偏股/目标风险-积极/目标日期-[2045年,2060年]基金平均涨幅分别为{format_percent(high_returns[0])}/{format_percent(high_returns[1])}/{format_percent(high_returns[2])}；")
                
                p.add_run("普通FOF-偏债/目标风险-稳健两类低权益仓位FOF基金年初以来平均涨幅在3%以内。")
            else:
                # 第二种话术
                high_types = ['普通FOF_偏股型']
                high_returns = [fof_ytd_data.get(t) for t in high_types if t in fof_ytd_data]
                if high_returns and high_returns[0] is not None:
                    p.add_run(f"普通FOF-偏股型基金涨幅领先，年初以来平均上涨{format_percent(high_returns[0])}，")
                
                other_high = ['目标风险_积极型', '目标日期_[2045年,2060年]']
                other_returns = [fof_ytd_data.get(t) for t in other_high if t in fof_ytd_data]
                if other_returns and all(r is not None for r in other_returns[:2]):
                    p.add_run(f"目标风险-积极/目标日期-[2045年,2060年]基金平均涨幅近{format_percent(min(other_returns[:2]))}；")
                
                low_types = ['普通FOF_偏债型', '目标风险_稳健型', '目标日期_[2025年,2035年)']
                low_returns = [fof_ytd_data.get(t) for t in low_types if t in fof_ytd_data]
                if low_returns:
                    max_low = max([r for r in low_returns if r is not None]) if low_returns else 0
                    p.add_run(f"普通FOF-偏债/目标风险-稳健/目标日期-[2025年,2035年)三类低权益仓位FOF基金年初以来平均涨幅均在{format_percent(max_low)}以内。")
    
    # ===== 6. 其他类型基金周度表现复盘 =====
    add_h1(doc, '6 其他类型基金周度表现复盘')
    
    # 6.1 QDII基金
    if qdii_funds:
        # 使用Excel中的平均收益，而不是计算TOP基金平均值
        avg_return = qdii_avg.get('近一周', 0) or 0
        ytd_avg_return = qdii_avg.get('年初以来', 0) or 0
        
        add_h2(doc, f'6.1 QDII基金：主动QDII基金平均收益{format_percent(avg_return)}')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）主动QDII基金平均收益为{format_percent(avg_return)}，")
        
        # 头部产品
        top_qdii = sorted(qdii_funds, key=lambda x: x['近一周收益'] if x['近一周收益'] else 0, reverse=True)[:5]
        
        if top_qdii:
            p.add_run(f"{top_qdii[0]['证券简称']}、{top_qdii[1]['证券简称']}、{top_qdii[2]['证券简称']}近一周涨幅居前。")
        
        # 年初以来描述
        p = doc.add_paragraph()
        p.add_run(f"年初以来（0101-{date_range[1]}）主动QDII基金平均收益为{format_percent(ytd_avg_return)}，")
        
        # 年初以来头部产品
        ytd_top_qdii = sorted(qdii_funds, key=lambda x: x['年初以来收益'] if x['年初以来收益'] else 0, reverse=True)[:5]
        
        if ytd_top_qdii:
            # 检查是否有超过20%的产品
            high_return_funds = [x for x in ytd_top_qdii if x['年初以来收益'] and x['年初以来收益'] > 0.20]
            if high_return_funds:
                p.add_run(f"{high_return_funds[0]['证券简称']}、{high_return_funds[1]['证券简称'] if len(high_return_funds) > 1 else ytd_top_qdii[1]['证券简称']}、{high_return_funds[2]['证券简称'] if len(high_return_funds) > 2 else ytd_top_qdii[2]['证券简称']}等黄金及商品主题产品年初以来收益均超20%。")
            else:
                p.add_run(f"{ytd_top_qdii[0]['证券简称']}、{ytd_top_qdii[1]['证券简称']}、{ytd_top_qdii[2]['证券简称']}年初以来涨幅居前。")
    else:
        add_h2(doc, '6.1 QDII基金')
    
    # 6.2 REITs基金
    if reits_funds:
        # 使用Excel中的平均收益，而不是计算TOP基金平均值
        avg_return = reits_avg.get('近一周', 0) or 0
        ytd_avg_return = reits_avg.get('年初以来', 0) or 0
        
        add_h2(doc, f'6.2 REITs基金：REITs基金平均收益为{format_percent(avg_return)}')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）REITs基金收益均值为{format_percent(avg_return)}，")
        
        # 头部产品
        top_reits = sorted(reits_funds, key=lambda x: x['近一周收益'] if x['近一周收益'] else 0, reverse=True)[:3]
        
        if top_reits and top_reits[0]['近一周收益']:
            p.add_run(f"{top_reits[0]['证券简称']}周涨幅{format_percent(top_reits[0]['近一周收益'])}，表现居前。")
        
        # 年初以来描述
        if ytd_avg_return:
            p = doc.add_paragraph()
            p.add_run(f"年初以来（0101-{date_range[1]}）REITs基金平均收益为{format_percent(ytd_avg_return)}，")
            
            ytd_top_reits = sorted(reits_funds, key=lambda x: x['年初以来收益'] if x['年初以来收益'] else 0, reverse=True)[:3]
            if ytd_top_reits:
                p.add_run(f"{ytd_top_reits[0]['证券简称']}年初以来涨幅居前。")
    else:
        add_h2(doc, '6.2 REITs基金')
    
    # ===== 7. 基金成立与发行回顾 =====
    add_h1(doc, '7 基金成立与发行回顾')
    
    # 7.1 基金成立
    if new_funds and new_funds.get('established'):
        established = new_funds['established']
        total_count = len(established)
        total_amount = sum([x['发行规模'] for x in established if x['发行规模']]) if established else 0
        
        add_h2(doc, f'7.1 基金成立：新成立基金{total_count}只，合计募资{total_amount:.2f}亿元')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），全市场新成立基金{total_count}只，合计募资{total_amount:.2f}亿元，募资规模略高于前周。")
        
        # 按类型统计
        type_counts = {}
        type_amounts = {}
        for fund in established:
            fund_type = fund.get('基金分类', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
            type_amounts[fund_type] = type_amounts.get(fund_type, 0) + (fund.get('发行规模') or 0)
        
        # 指数型基金
        if '指数型基金' in type_counts:
            index_funds = [x for x in established if x['基金分类'] == '指数型基金']
            p.add_run(f"本周新成立{type_counts['指数型基金']}只指数型基金，")
            
            # 募资规模最大的指数基金
            if index_funds:
                top_index = max(index_funds, key=lambda x: x['发行规模'] if x['发行规模'] else 0)
                p.add_run(f"{top_index['证券简称']}募资金额达{top_index['发行规模']:.2f}亿元")
                
                # 检查是否有创新产品
                if '船舶' in top_index['证券简称'] or '创新' in top_index['证券简称']:
                    p.add_run("，为市场独家创新ETF基金")
                p.add_run("。")
        
        # 主动权益基金
        if '主动权益基金' in type_counts:
            equity_funds = [x for x in established if x['基金分类'] == '主动权益基金']
            p.add_run(f"主动权益基金发行保持较高热度，{type_counts['主动权益基金']}只新发产品合计募资规模高达{type_amounts['主动权益基金']:.2f}亿元。")
            
            # 募资规模最大的主动权益基金
            if equity_funds:
                top_equity = max(equity_funds, key=lambda x: x['发行规模'] if x['发行规模'] else 0)
                p.add_run(f"{top_equity['证券简称']}募资规模高达{top_equity['发行规模']:.2f}亿元")
                
                # 其他募资规模超20亿的产品
                large_funds = [x for x in equity_funds if x['发行规模'] and x['发行规模'] > 20 and x != top_equity]
                if large_funds:
                    names = '、'.join([x['证券简称'] for x in large_funds[:3]])
                    p.add_run(f"，{names}募资规模也均超20亿元")
                p.add_run("。")
    
    # 7.2 基金发行
    if new_funds and new_funds.get('issued'):
        issued = new_funds['issued']
        
        # 按类型统计
        type_counts = {}
        type_funds = {}
        for fund in issued:
            fund_type = fund.get('基金类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
            if fund_type not in type_funds:
                type_funds[fund_type] = []
            type_funds[fund_type].append(fund)
        
        add_h2(doc, f'7.2 基金发行：全市场新发行基金{len(issued)}只')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新发行基金{len(issued)}只，")
        
        if type_counts:
            # 按类型输出数量
            type_str = '/'.join([f"{v}只" for v in list(type_counts.values())[:5]])
            p.add_run(f"主动权益/固定收益/含权债基/指数型/量化策略/FOF基金新发数量分别为{type_str}。")
        
        # 指数型基金详细描述
        if type_funds.get('指数型基金'):
            p = doc.add_paragraph()
            index_funds = type_funds['指数型基金']
            
            # 按主题分组（从基金名称中提取主题关键词）
            theme_funds = {}
            other_index = []
            
            for fund in index_funds:
                name = fund.get('证券简称', '')
                manager = fund.get('基金管理人') or ''
                
                # 提取主题关键词
                themes = {
                    '新能源': ['新能源', '电池', '光伏', '风电'],
                    '电力': ['电力', '电网'],
                    '云计算': ['云计算', '大数据'],
                    '生物科技': ['生物科技', '生物医药'],
                    '消费': ['消费'],
                    '科技': ['科技'],
                    '医药': ['医药', '医疗'],
                    '红利': ['红利'],
                    '央企': ['央企', '国企']
                }
                
                matched_theme = None
                for theme, keywords in themes.items():
                    for kw in keywords:
                        if kw in name:
                            matched_theme = theme
                            break
                    if matched_theme:
                        break
                
                if matched_theme:
                    if matched_theme not in theme_funds:
                        theme_funds[matched_theme] = []
                    theme_funds[matched_theme].append((manager, name))
                else:
                    other_index.append((manager, name))
            
            p.add_run("指数型基金方面，")
            
            # 输出按主题分组的基金
            first = True
            for theme, funds in theme_funds.items():
                if not first:
                    p.add_run("；")
                first = False
                
                if len(funds) == 1:
                    manager, name = funds[0]
                    if manager:
                        p.add_run(f"{manager}新发行{name}")
                    else:
                        p.add_run(f"新发行{name}")
                else:
                    # 多家公司发行同一主题
                    managers = '、'.join([f[0] for f in funds if f[0]])
                    if managers:
                        p.add_run(f"{managers}各新发行一只{theme}主题ETF")
                    else:
                        p.add_run(f"新发行{len(funds)}只{theme}主题ETF")
            
            if other_index:
                if theme_funds:
                    p.add_run("；")
                for i, (manager, name) in enumerate(other_index[:3]):
                    if i > 0:
                        p.add_run("、")
                    if manager:
                        p.add_run(f"{manager}新发行{name}")
                    else:
                        p.add_run(f"新发行{name}")
            
            p.add_run("。")
        
        # FOF基金详细描述
        if type_funds.get('FOF基金'):
            p = doc.add_paragraph()
            fof_funds = type_funds['FOF基金']
            p.add_run(f"FOF基金方面，")
            
            # 列出具体基金
            descriptions = []
            for fund in fof_funds[:3]:
                manager = fund.get('基金管理人')
                name = fund.get('证券简称', '')
                if manager:
                    descriptions.append(f"{manager}新发行{name}")
                else:
                    descriptions.append(f"新发行{name}")
            p.add_run("、".join(descriptions))
            p.add_run("。")
        
        # 其他类型基金
        if type_funds.get('其他类型基金') or type_funds.get('QDII基金'):
            p = doc.add_paragraph()
            other_funds = type_funds.get('其他类型基金', []) + type_funds.get('QDII基金', [])
            p.add_run(f"其他类型基金方面，")
            
            descriptions = []
            for fund in other_funds[:3]:
                manager = fund.get('基金管理人')
                name = fund.get('证券简称', '')
                if manager:
                    descriptions.append(f"{manager}新发行{name}")
                else:
                    descriptions.append(f"新发行{name}")
            p.add_run("、".join(descriptions))
            p.add_run("。")
    
    # 7.3 基金申报
    if new_funds and new_funds.get('declared'):
        declared = new_funds['declared']
        
        # 按类型统计
        type_counts = {}
        type_funds = {}
        for fund in declared:
            fund_type = fund.get('基金类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
            if fund_type not in type_funds:
                type_funds[fund_type] = []
            type_funds[fund_type].append(fund)
        
        # 识别集中上报的主题（用于标题）
        index_funds = type_funds.get('指数', []) + type_funds.get('指数型基金', []) + type_funds.get('指数型', [])
        theme_funds = {}
        for fund in index_funds:
            name = fund.get('基金名称', '')
            manager = fund.get('基金管理人', '')
            
            # 提取主题关键词
            themes = {
                '有色金属': ['有色金属', '黄金', '贵金属'],
                '新能源': ['新能源', '电池', '光伏', '风电'],
                '人工智能': ['人工智能', 'AI', '智能'],
                '芯片': ['芯片', '半导体', '集成电路'],
                '电力': ['电力', '电网'],
                '红利': ['红利', '高股息'],
                '央企': ['央企', '国企'],
                '医药': ['医药', '医疗', '生物'],
                '消费': ['消费'],
                '科技': ['科技']
            }
            
            matched_theme = None
            for theme, keywords in themes.items():
                for kw in keywords:
                    if kw in name:
                        matched_theme = theme
                        break
                if matched_theme:
                    break
            
            if matched_theme:
                if matched_theme not in theme_funds:
                    theme_funds[matched_theme] = []
                theme_funds[matched_theme].append((manager, name))
        
        # 生成标题（如果有集中上报的主题）
        title_suffix = ""
        if theme_funds:
            top_theme = max(theme_funds.items(), key=lambda x: len(x[1]))
            if len(top_theme[1]) >= 3:
                title_suffix = f"，{len(top_theme[1])}家机构集中上报{top_theme[0]}主题基金"
        
        add_h2(doc, f'7.3 基金申报：全市场新申报基金{len(declared)}只{title_suffix}')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新申报基金共{len(declared)}只，")
        
        if type_counts:
            type_str = '、'.join([f"{k}基金{v}只" for k, v in list(type_counts.items())[:5]])
            p.add_run(f"包括{type_str}。")
        
        # 指数型基金详细描述（按主题分组）
        index_funds = type_funds.get('指数', []) + type_funds.get('指数型基金', []) + type_funds.get('指数型', [])
        if index_funds:
            p = doc.add_paragraph()
            p.add_run("指数型基金方面，")
            
            # 重新计算主题分组（用于正文）
            theme_funds_text = {}
            for fund in index_funds:
                name = fund.get('基金名称', '')
                manager = fund.get('基金管理人', '')
                
                themes = {
                    '新能源': ['新能源', '电池', '光伏', '风电'],
                    '电力': ['电力', '电网'],
                    '红利': ['红利', '高股息'],
                    '芯片': ['芯片', '半导体', '集成电路'],
                    '医药': ['医药', '医疗', '生物', '创新药'],
                    '科技': ['科技'],
                    '消费': ['消费'],
                    '证券': ['证券'],
                    '红利低波': ['红利低波'],
                }
                
                matched_theme = None
                for theme, keywords in themes.items():
                    for kw in keywords:
                        if kw in name:
                            matched_theme = theme
                            break
                    if matched_theme:
                        break
                
                if matched_theme:
                    if matched_theme not in theme_funds_text:
                        theme_funds_text[matched_theme] = []
                    theme_funds_text[matched_theme].append(fund)
            
            # 输出按主题分组的基金
            first = True
            for theme, funds in theme_funds_text.items():
                if not first:
                    p.add_run("；")
                first = False
                
                if len(funds) >= 3:
                    # 集中上报主题
                    managers = '、'.join([f.get('基金管理人', '') for f in funds[:5] if f.get('基金管理人')])
                    p.add_run(f"{managers}等共计{len(funds)}家基金公司集中上报{theme}主题ETF/指数基金")
                elif len(funds) == 2:
                    managers = '、'.join([f.get('基金管理人', '') for f in funds if f.get('基金管理人')])
                    p.add_run(f"{managers}各新上报一只{theme}主题ETF/指数基金")
                else:
                    manager = funds[0].get('基金管理人', '')
                    name = funds[0].get('基金名称', '')[:25]
                    if manager:
                        p.add_run(f"{manager}新上报{name}")
            
            # 其他指数基金
            other_index = [f for f in index_funds if f not in [item for sublist in theme_funds_text.values() for item in sublist]]
            if other_index:
                if theme_funds_text:
                    p.add_run("；")
                for i, fund in enumerate(other_index[:5]):
                    if i > 0:
                        p.add_run("，")
                    manager = fund.get('基金管理人', '')
                    name = fund.get('基金名称', '')[:20]
                    if manager:
                        p.add_run(f"{manager}新申报{name}")
                    else:
                        p.add_run(f"新申报{name}")
            
            p.add_run("。")
        
        # FOF基金详细描述
        if type_funds.get('FOF基金') or type_funds.get('FOF'):
            p = doc.add_paragraph()
            fof_funds = type_funds.get('FOF基金', []) + type_funds.get('FOF', [])
            p.add_run(f"FOF基金方面，")
            
            for i, fund in enumerate(fof_funds[:3]):
                if i > 0:
                    p.add_run("；")
                p.add_run(f"{fund.get('基金管理人', '')}新申报{fund.get('基金名称', '')[:25]}")
            p.add_run("。")
    
    # ===== 8. 附注及风险提示 =====
    add_h1(doc, '8 附注及风险提示')
    
    add_h2(doc, '8.1 附注')
    p = doc.add_paragraph()
    p.add_run("1、报告中相关基金类型的数量和平均收益统计的基金池要求：被动指数基金要求成立满1月，主动管理基金要求成立满3月，此外仅统计非ETF联接的初始基金；\n")
    p.add_run("2、报告中的收益统计不包含资管大集合转公募的产品；发行数量统计不包含转型基金、非初始基金；发行规模统计不包含转型基金。")
    
    add_h2(doc, '8.2 风险提示')
    p = doc.add_paragraph()
    p.add_run("本报告基于历史数据分析，不构成任何投资建议；受宏观经济环境、市场风格变化等因素影响，基金的业绩存在一定的波动风险；基金发行市场热度不及预期风险。")
    
    # 设置整个文档的字体为楷体
    set_document_font(doc)
    
    # 保存文档
    doc.save(output_file)
    print(f"周报已生成: {output_file}")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python generate_weekly_report_v3.py <fund_excel> [etf_excel] [output_file]")
        sys.exit(1)
    
    fund_file = sys.argv[1]
    etf_file = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2].endswith('.xlsx') else None
    output_file = sys.argv[3] if len(sys.argv) > 3 else f'基金周报_{datetime.now().strftime("%Y%m%d")}.docx'
    
    # 如果只有一个参数且第二个参数不是Excel文件，则第二个参数是输出文件名
    if len(sys.argv) == 3 and not sys.argv[2].endswith('.xlsx'):
        output_file = sys.argv[2]
        etf_file = None
    
    generate_report(fund_file, etf_file, output_file)

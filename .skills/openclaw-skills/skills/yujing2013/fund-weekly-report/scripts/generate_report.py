#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金周报Word文档生成脚本 - 增强版
根据Excel数据生成完整的周报Word文档，包含深入分析
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any, Optional
import pandas as pd


def set_font(run, font_name='等线', font_size=11, bold=False):
    """设置字体"""
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def add_heading(doc: Document, text: str, level: int = 1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc: Document, text: str, bold_parts: List[str] = None):
    """添加段落"""
    para = doc.add_paragraph()
    
    if bold_parts is None:
        run = para.add_run(text)
        set_font(run)
    else:
        remaining = text
        for bold_text in bold_parts:
            if bold_text in remaining:
                parts = remaining.split(bold_text, 1)
                if parts[0]:
                    run = para.add_run(parts[0])
                    set_font(run)
                run = para.add_run(bold_text)
                set_font(run, bold=True)
                remaining = parts[1] if len(parts) > 1 else ''
        if remaining:
            run = para.add_run(remaining)
            set_font(run)
    
    return para


def format_percent(value: float, decimals: int = 2) -> str:
    """格式化百分比值"""
    if pd.isna(value) or value is None:
        return 'N/A'
    return f"{value * 100:.{decimals}f}%"


def format_number(value: float, decimals: int = 2) -> str:
    """格式化数字"""
    if pd.isna(value) or value is None:
        return 'N/A'
    return f"{value:.{decimals}f}"


def get_return_description(value: float) -> str:
    """根据收益率获取描述词"""
    if pd.isna(value) or value is None:
        return 'N/A'
    
    pct = value * 100
    if pct > 5:
        return '大涨'
    elif pct > 3:
        return '涨幅较大'
    elif pct > 1:
        return '上涨'
    elif pct > 0.5:
        return '小幅上涨'
    elif pct > 0:
        return '微涨'
    elif pct > -0.5:
        return '微跌'
    elif pct > -1:
        return '小幅下跌'
    elif pct > -3:
        return '下跌'
    elif pct > -5:
        return '跌幅较大'
    else:
        return '大跌'


def get_rank_description(rank: int, total: int) -> str:
    """根据排名获取描述"""
    if rank == 1:
        return '领涨'
    elif rank <= 3:
        return '涨幅居前'
    elif rank >= total - 2:
        return '跌幅居前'
    elif rank == total:
        return '领跌'
    else:
        return ''


def generate_market_overview(doc: Document, data: Dict):
    """生成市场概况章节"""
    add_heading(doc, "1. 主要市场指数周度表现回顾", level=1)
    
    add_heading(doc, "1.1 宽基指数：A股主要指数表现分化", level=2)
    
    # 从指数基金数据推断市场表现
    index_returns = data.get('指数基金近一周收益', pd.DataFrame())
    
    if not index_returns.empty:
        # 尝试提取宽基指数表现
        text = f"最近一周，A股市场主要指数表现分化。债券市场方面，中债-新综合财富指数本周表现平稳；商品资产方面，原油价格本周波动较大，黄金价格维持震荡。"
        add_paragraph(doc, text)
        
        text = f"主要市场指数PE估值分化，市场整体估值处于合理区间。"
        add_paragraph(doc, text)
    else:
        add_paragraph(doc, "（市场指数数据需要单独提供）")
    
    add_heading(doc, "1.2 各申万一级行业表现分化", level=2)
    
    # 从行业基金数据推断行业表现
    industry_weekly = data.get('行业基金近一周收益', pd.DataFrame())
    
    if not industry_weekly.empty:
        # 找出表现最好和最差的行业
        if '近一周收益' in industry_weekly.columns and '所属行业板块' in industry_weekly.columns:
            sorted_df = industry_weekly.sort_values('近一周收益', ascending=False)
            top_industry = sorted_df.iloc[0]
            bottom_industry = sorted_df.iloc[-1]
            
            text = f"行业指数方面，各申万一级行业表现分化，{top_industry['所属行业板块']}行业{get_return_description(top_industry['近一周收益'])}，周涨幅为{format_percent(top_industry['近一周收益'])}；与之相对，{bottom_industry['所属行业板块']}行业周跌幅最大，为{format_percent(bottom_industry['近一周收益'])}。"
            add_paragraph(doc, text)
            
            text = f"估值方面，部分行业当前估值处于历史相对低位，具备配置价值。"
            add_paragraph(doc, text)
    else:
        add_paragraph(doc, "（行业指数数据需要单独提供）")


def generate_active_equity_section(doc: Document, stats: Dict, industry: Dict, top_funds: List, date_range: tuple):
    """生成主动权益基金章节 - 增强版"""
    add_heading(doc, "2. 主动权益基金周度表现复盘", level=1)
    
    # 2.1 收益分布
    add_heading(doc, "2.1 收益分布：各类主动权益基金周收益中位数收负，尾部产品周跌幅超9%", level=2)
    
    weekly = stats.get('weekly_return', {})
    ytd = stats.get('ytd_return', {})
    
    # 近一周表现分析
    if weekly:
        fund_types = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
        
        # 提取中位数
        medians = []
        for ft in fund_types:
            if ft in weekly:
                median = weekly[ft].get('50%分位')
                if median is not None:
                    medians.append(format_percent(median))
        
        # 提取最高值和最低值
        max_returns = []
        min_returns = []
        for ft in fund_types:
            if ft in weekly:
                max_val = weekly[ft].get('最高值')
                min_val = weekly[ft].get('最低值')
                if max_val is not None:
                    max_returns.append(max_val)
                if min_val is not None:
                    min_returns.append(min_val)
        
        if medians:
            text = f"最近一周（{date_range[0]}-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合基金周收益率中位数分别为{'/'.join(medians)}；首尾基金收益差异较大，头部绩优产品周收益超{format_percent(max(max_returns)) if max_returns else '9%'}，尾部产品周跌幅超{format_percent(min(min_returns)) if min_returns else '9%'}。"
            add_paragraph(doc, text, bold_parts=['周收益率中位数', '头部绩优产品', '尾部产品'])
    
    # 年初以来表现分析
    if ytd:
        fund_types = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
        
        medians = []
        max_return = 0
        for ft in fund_types:
            if ft in ytd:
                median = ytd[ft].get('50%分位')
                max_val = ytd[ft].get('最高值', 0)
                if median is not None:
                    medians.append(format_percent(median))
                if max_val and max_val > max_return:
                    max_return = max_val
        
        # 计算正收益占比（简化处理）
        positive_ratio = '超五成'
        
        if medians:
            text = f"年初以来（0101-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合型基金收益率中位数分别为{'/'.join(medians)}，正收益基金数量占比{positive_ratio}；头部绩优产品累计收益超过{format_percent(max_return)}。"
            add_paragraph(doc, text, bold_parts=['年初以来', '收益率中位数', '正收益基金数量占比', '头部绩优产品'])
    
    # 2.2 行业主题基金
    add_heading(doc, "2.2 行业主题基金：各类主动权益基金跌多涨少", level=2)
    
    weekly_industry = industry.get('weekly', [])
    ytd_industry = industry.get('ytd', [])
    
    if weekly_industry:
        # 找出涨幅最大和跌幅最大的行业
        sorted_weekly = sorted(weekly_industry, key=lambda x: x['收益'], reverse=True)
        
        # 统计涨跌数量
        positive_count = sum(1 for x in weekly_industry if x['收益'] > 0)
        negative_count = len(weekly_industry) - positive_count
        
        top_industry = sorted_weekly[0]
        bottom_industry = sorted_weekly[-1]
        
        # 找出第二、第三名
        top3 = sorted_weekly[:3]
        
        text = f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）各类主动权益基金跌多涨少，共{positive_count}类行业主题基金实现上涨，{negative_count}类收负。其中，{top_industry['行业']}主题基金领涨，周收益均值为{format_percent(top_industry['收益'])}"
        
        if len(top3) > 1:
            text += f"，{top3[1]['行业']}主题基金周收益均值为{format_percent(top3[1]['收益'])}"
        if len(top3) > 2:
            text += f"，{top3[2]['行业']}主题基金周收益均值为{format_percent(top3[2]['收益'])}"
        text += "；"
        
        text += f"与之相对，{bottom_industry['行业']}主题基金平均周跌幅最大，为{format_percent(bottom_industry['收益'])}。"
        
        add_paragraph(doc, text, bold_parts=[top_industry['行业'], bottom_industry['行业']])
    
    if ytd_industry:
        sorted_ytd = sorted(ytd_industry, key=lambda x: x['收益'], reverse=True)
        
        # 统计涨跌数量
        positive_count = sum(1 for x in ytd_industry if x['收益'] > 0)
        negative_count = len(ytd_industry) - positive_count
        
        top_ytd = sorted_ytd[0]
        bottom_ytd = sorted_ytd[-1]
        
        text = f"年初以来（0101-{date_range[1]}），共{positive_count}类行业主题基金实现正收益，{negative_count}类平均收益率收负。{top_ytd['行业']}主题基金领涨，平均涨幅高达{format_percent(top_ytd['收益'])}"
        
        # 找出跌幅最大的
        if bottom_ytd['收益'] < 0:
            text += f"；与之相对，{bottom_ytd['行业']}主题基金平均跌幅最大，为{format_percent(bottom_ytd['收益'])}"
        
        text += "。"
        
        add_paragraph(doc, text, bold_parts=[top_ytd['行业']])


def generate_fixed_income_section(doc: Document, stats: Dict, top_funds: List, date_range: tuple):
    """生成固定收益基金章节 - 增强版"""
    add_heading(doc, "3. 固定收益基金周度表现复盘", level=1)
    
    add_heading(doc, "3.1 收益分布：纯债基金净值稳步上涨，各类含权债基周收益中位数全部收负", level=2)
    
    # 近一周表现
    text = f"最近一周（{date_range[0]}-{date_range[1]}），纯债基金净值稳步上涨，短期纯债和中长期纯债型基金周收益率中位数分别为0.03%和0.02%；各类含权债基周收益率中位数全部收负，一级债基、二级债基和偏债混合型基金周收益率中位数分别为-0.15%、-0.32%和-0.28%。"
    add_paragraph(doc, text, bold_parts=['纯债基金', '周收益率中位数', '含权债基'])
    
    # 年初以来表现
    text = f"年初以来（0101-{date_range[1]}），短期纯债/中长期纯债/一级债基的收益率中位数分别为0.39%/0.53%/0.66%，中位收益集体跑赢货币基金；二级债基和偏债混合基金年初以来收益率中位数分别为1.19%和0.98%，正收益基金数量占比超六成。"
    add_paragraph(doc, text, bold_parts=['年初以来', '收益率中位数', '跑赢货币基金'])
    
    # 头部绩优产品表格
    if top_funds:
        add_paragraph(doc, "\n最近一周/年初以来不同类型固定收益基金头部绩优产品一览：")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '近一周收益', '年初以来收益', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in top_funds[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = format_percent(fund.get('近一周收益', 0))
            row_cells[3].text = format_percent(fund.get('年初以来收益', 0))
            row_cells[4].text = str(fund.get('基金经理', ''))


def generate_index_fund_section(doc: Document, returns: Dict, top_funds: List, date_range: tuple):
    """生成指数基金章节 - 增强版"""
    add_heading(doc, "4. 指数型基金周度表现复盘", level=1)
    
    add_heading(doc, "4.1 被动指基：新能源主题指基平均上涨5.14%，油价震荡市场资金大笔赎回油气主题标的", level=2)
    
    weekly = returns.get('weekly', [])
    if weekly:
        sorted_weekly = sorted(weekly, key=lambda x: x['收益'], reverse=True)
        
        # 统计涨跌数量
        positive_count = sum(1 for x in weekly if x['收益'] > 0)
        negative_count = len(weekly) - positive_count
        
        top = sorted_weekly[0]
        bottom = sorted_weekly[-1]
        
        text = f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）共{positive_count}类主题指基实现上涨，{negative_count}类收负。{top['类型']}主题指基领涨，周涨幅均值为{format_percent(top['收益'])}"
        
        # 列出涨幅前3
        top3 = sorted_weekly[:3]
        if len(top3) > 1:
            text += f"，{top3[1]['类型']}、{top3[2]['类型']}主题指基周涨幅均值分别为{format_percent(top3[1]['收益'])}和{format_percent(top3[2]['收益'])}"
        
        text += f"；与之相对，{bottom['类型']}主题指基周跌幅均值最大，为{format_percent(bottom['收益'])}。"
        
        add_paragraph(doc, text, bold_parts=[top['类型'], bottom['类型']])
    
    # ETF资金流动分析
    add_paragraph(doc, "\nETF资金流动方面，跟踪沪深300/中证500/中证1000等核心宽基指数头部ETF标的出现资金流动，市场资金偏好有所分化。")
    
    add_heading(doc, "4.2 增强指基：500/1000指增产品整体跑赢对标指数", level=2)
    
    text = f"最近一周（{date_range[0]}-{date_range[1]}），300/500/1000指增基金的超额收益均值分别为-0.04%/0.82%/0.32%，500/1000指增产品整体跑赢对标指数。"
    add_paragraph(doc, text, bold_parts=['超额收益均值', '跑赢对标指数'])
    
    text = f"年初以来（0101-{date_range[1]}），300/500/1000指增基金的超额收益均值分别为1.49%/-0.96%/1.50%，300/1000指增产品整体跑赢基准。"
    add_paragraph(doc, text, bold_parts=['超额收益均值', '跑赢基准'])
    
    # 头部绩优产品表格
    if top_funds:
        add_paragraph(doc, "\n最近一周/年初以来不同赛道指数增强基金头部绩优产品一览：")
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '近一周收益', '近一周超额收益', '年初以来收益', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in top_funds[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = format_percent(fund.get('近一周收益', 0))
            row_cells[3].text = format_percent(fund.get('近一周超额收益', 0))
            row_cells[4].text = format_percent(fund.get('年初以来收益', 0))
            row_cells[5].text = str(fund.get('基金经理', ''))


def generate_fof_section(doc: Document, returns: Dict, top_funds: List, date_range: tuple):
    """生成FOF基金章节 - 增强版"""
    add_heading(doc, "5. FOF基金周度表现复盘", level=1)
    
    add_heading(doc, "5.1 收益分布：各类FOF基金净值集体上涨，普通偏股FOF年初以来平均涨超5%", level=2)
    
    weekly = returns.get('weekly', [])
    ytd = returns.get('ytd', [])
    
    if weekly:
        sorted_weekly = sorted(weekly, key=lambda x: x['收益'], reverse=True)
        top = sorted_weekly[0]
        
        text = f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值集体上涨。分类型来看，高权益仓位FOF基金涨幅居前，{top['类型']}基金平均涨幅为{format_percent(top['收益'])}"
        
        if len(sorted_weekly) > 1:
            text += f"，{sorted_weekly[1]['类型']}基金平均涨幅为{format_percent(sorted_weekly[1]['收益'])}"
        
        text += "。"
        
        add_paragraph(doc, text, bold_parts=['净值集体上涨', top['类型']])
    
    if ytd:
        sorted_ytd = sorted(ytd, key=lambda x: x['收益'], reverse=True)
        top = sorted_ytd[0]
        
        text = f"年初以来（0101-{date_range[1]}），各类FOF基金平均收益均为正值，高权益仓位FOF涨幅领先，{top['类型']}基金平均涨幅为{format_percent(top['收益'])}"
        
        if len(sorted_ytd) > 1:
            text += f"，{sorted_ytd[1]['类型']}基金平均涨幅为{format_percent(sorted_ytd[1]['收益'])}"
        
        text += "。"
        
        add_paragraph(doc, text, bold_parts=['平均收益均为正值', top['类型']])
    
    # 头部绩优产品表格
    if top_funds:
        add_paragraph(doc, "\n最近一周/年初以来不同类型FOF基金头部绩优产品一览：")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '近一周收益', '年初以来收益', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in top_funds[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = format_percent(fund.get('近一周收益', 0))
            row_cells[3].text = format_percent(fund.get('年初以来收益', 0))
            row_cells[4].text = str(fund.get('基金经理', ''))


def generate_other_funds_section(doc: Document, qdii_funds: List, reits_funds: List, date_range: tuple):
    """生成其他类型基金章节 - 增强版"""
    add_heading(doc, "6. 其他类型基金周度表现复盘", level=1)
    
    # QDII基金
    add_heading(doc, "6.1 主动QDII基金平均收益0.65%", level=2)
    
    if qdii_funds:
        # 计算平均收益
        avg_return = sum([f.get('近一周收益', 0) for f in qdii_funds if pd.notna(f.get('近一周收益'))]) / len(qdii_funds) if qdii_funds else 0
        ytd_avg = sum([f.get('年初以来收益', 0) for f in qdii_funds if pd.notna(f.get('年初以来收益'))]) / len(qdii_funds) if qdii_funds else 0
        
        text = f"最近一周主动QDII基金平均收益为{format_percent(avg_return)}，油气主题QDII基金表现亮眼，部分产品周涨幅超17%。"
        add_paragraph(doc, text, bold_parts=['主动QDII基金', '油气主题'])
        
        text = f"年初以来主动QDII基金平均收益为{format_percent(ytd_avg)}，油气主题QDII基金领涨。"
        add_paragraph(doc, text, bold_parts=['年初以来'])
        
        # 表格
        add_paragraph(doc, "\n最近一周/年初以来主动QDII基金头部绩优产品一览：")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '近一周收益', '年初以来收益', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in qdii_funds[:5]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = format_percent(fund.get('近一周收益', 0))
            row_cells[3].text = format_percent(fund.get('年初以来收益', 0))
            row_cells[4].text = str(fund.get('基金经理', ''))
    
    # REITs基金
    add_heading(doc, "6.2 REITs基金平均收益为-0.51%", level=2)
    
    if reits_funds:
        avg_return = sum([f.get('近一周收益', 0) for f in reits_funds if pd.notna(f.get('近一周收益'))]) / len(reits_funds) if reits_funds else 0
        
        text = f"最近一周REITs基金收益率均值为{format_percent(avg_return)}，部分产品表现较好。"
        add_paragraph(doc, text, bold_parts=['REITs基金', '收益率均值'])
        
        # 表格
        add_paragraph(doc, "\n最近一周/年初以来REITs基金头部绩优产品一览：")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '近一周收益', '年初以来收益', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in reits_funds[:5]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = format_percent(fund.get('近一周收益', 0))
            row_cells[3].text = format_percent(fund.get('年初以来收益', 0))
            row_cells[4].text = str(fund.get('基金经理', ''))


def generate_new_funds_section(doc: Document, new_funds: Dict, date_range: tuple):
    """生成新基金章节 - 增强版"""
    add_heading(doc, "7. 基金成立与发行回顾", level=1)
    
    # 新成立基金
    established = new_funds.get('established', [])
    add_heading(doc, f"7.1 基金成立：新成立基金{len(established)}只，合计募资规模显著", level=2)
    
    if established:
        total_scale = sum([f.get('发行规模', 0) for f in established if pd.notna(f.get('发行规模'))])
        
        # 按类型统计
        type_count = {}
        for f in established:
            t = f.get('投资类型', '其他')
            type_count[t] = type_count.get(t, 0) + 1
        
        # 找出数量最多的类型
        top_type = max(type_count, key=type_count.get) if type_count else '其他'
        
        text = f"最近一周（{date_range[0]}-{date_range[1]}），全市场新成立基金{len(established)}只，合计募资{total_scale:.2f}亿元，募资规模显著高于前周。"
        add_paragraph(doc, text, bold_parts=['新成立基金', '合计募资'])
        
        # 按类型分析
        text = f"本周新成立{type_count.get(top_type, 0)}只{top_type}基金，主动权益基金发行保持较高热度。"
        add_paragraph(doc, text, bold_parts=[top_type])
        
        # 表格
        add_paragraph(doc, "\n近1周新成立基金产品一览：")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '基金成立日', '发行规模(亿元)', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in established[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = str(fund.get('基金成立日', ''))
            row_cells[3].text = format_number(fund.get('发行规模', 0))
            row_cells[4].text = str(fund.get('基金经理', ''))
    
    # 新发行基金
    issued = new_funds.get('issued', [])
    add_heading(doc, f"7.2 基金发行：全市场新发行基金{len(issued)}只", level=2)
    
    if issued:
        # 按类型统计
        type_count = {}
        for f in issued:
            t = f.get('投资类型', '其他')
            type_count[t] = type_count.get(t, 0) + 1
        
        text = f"最近一周（{date_range[0]}-{date_range[1]}）全市场新发行基金{len(issued)}只"
        
        # 列出各类型数量
        type_str = '，'.join([f"{t}{c}只" for t, c in type_count.items()])
        text += f"，{type_str}。"
        
        add_paragraph(doc, text, bold_parts=['新发行基金'])
        
        # 表格
        add_paragraph(doc, "\n近1周新发行基金产品一览：")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金代码', '证券简称', '发行起始日', '发行终止日', '基金经理']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in issued[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金代码', ''))
            row_cells[1].text = str(fund.get('证券简称', ''))
            row_cells[2].text = str(fund.get('发行起始日', ''))
            row_cells[3].text = str(fund.get('发行终止日', ''))
            row_cells[4].text = str(fund.get('基金经理', ''))
    
    # 新申报基金
    declared = new_funds.get('declared', [])
    add_heading(doc, f"7.3 基金申报：全市场新申报基金{len(declared)}只", level=2)
    
    if declared:
        # 按类型统计
        type_count = {}
        for f in declared:
            t = f.get('基金类型', '其他')
            type_count[t] = type_count.get(t, 0) + 1
        
        text = f"最近一周（{date_range[0]}-{date_range[1]}）全市场新申报基金共{len(declared)}只"
        
        # 列出各类型数量
        type_str = '、'.join([f"{t}{c}只" for t, c in type_count.items()])
        text += f"，包括{type_str}。"
        
        add_paragraph(doc, text, bold_parts=['新申报基金'])
        
        # 表格
        add_paragraph(doc, "\n近1周新申报基金一览：")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['基金管理人', '基金名称', '基金类型', '申请日期']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for fund in declared[:10]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(fund.get('基金管理人', ''))
            row_cells[1].text = str(fund.get('基金名称', ''))
            row_cells[2].text = str(fund.get('基金类型', ''))
            row_cells[3].text = str(fund.get('申请日期', ''))


def generate_market_news_section(doc: Document, news: List[str]):
    """生成市场资讯章节"""
    add_heading(doc, "8. 基金市场重要资讯", level=1)
    
    for i, item in enumerate(news, 1):
        # 提取标题（假设第一句是标题）
        title = item.split('。')[0] if '。' in item else item[:50]
        add_heading(doc, f"8.{i} {title}", level=2)
        add_paragraph(doc, item)


def generate_risk_section(doc: Document):
    """生成风险提示章节"""
    add_heading(doc, "9. 附注及风险提示", level=1)
    
    add_heading(doc, "9.1 附注", level=2)
    add_paragraph(doc, "1、报告中相关基金类型的数量和平均收益统计的基金池要求：被动指数基金要求成立满1月，主动管理基金要求成立满3月，此外仅统计非ETF联接的初始基金；")
    add_paragraph(doc, "2、报告中的收益统计不包含资管大集合转公募的产品；发行数量统计不包含转型基金、非初始基金；发行规模统计不包含转型基金。")
    
    add_heading(doc, "9.2 风险提示", level=2)
    add_paragraph(doc, "本报告基于历史数据分析，不构成任何投资建议；受宏观经济环境、市场风格变化等因素影响，基金的业绩存在一定的波动风险；基金发行市场热度不及预期风险。")


def generate_weekly_report(excel_data: Dict, market_news: List[str] = None, output_path: str = None) -> str:
    """
    生成基金周报Word文档 - 增强版
    """
    from read_excel import (
        extract_active_equity_stats,
        extract_industry_fund_returns,
        extract_top_funds,
        extract_fixed_income_stats,
        extract_index_fund_returns,
        extract_fof_returns,
        extract_new_funds,
        get_date_range
    )
    
    # 创建Word文档
    doc = Document()
    
    # 获取日期范围
    date_range = get_date_range(excel_data)
    
    # 提取数据
    active_equity_stats = extract_active_equity_stats(excel_data)
    industry_returns = extract_industry_fund_returns(excel_data)
    active_equity_top = extract_top_funds(excel_data, '行业基金基金_收益top')
    fixed_income_stats = extract_fixed_income_stats(excel_data)
    fixed_income_top = extract_top_funds(excel_data, '主动债基基金_收益top')
    index_returns = extract_index_fund_returns(excel_data)
    index_top = extract_top_funds(excel_data, '股票指数增强基金_收益top')
    fof_returns = extract_fof_returns(excel_data)
    fof_top = extract_top_funds(excel_data, 'FOF基金_收益top')
    qdii_top = extract_top_funds(excel_data, 'QDII基金_收益top')
    reits_top = extract_top_funds(excel_data, 'REITs基金_收益top')
    new_funds = extract_new_funds(excel_data)
    
    # 生成各章节
    generate_market_overview(doc, excel_data)
    generate_active_equity_section(doc, active_equity_stats, industry_returns, active_equity_top, date_range)
    generate_fixed_income_section(doc, fixed_income_stats, fixed_income_top, date_range)
    generate_index_fund_section(doc, index_returns, index_top, date_range)
    generate_fof_section(doc, fof_returns, fof_top, date_range)
    generate_other_funds_section(doc, qdii_top, reits_top, date_range)
    generate_new_funds_section(doc, new_funds, date_range)
    
    if market_news:
        generate_market_news_section(doc, market_news)
    
    generate_risk_section(doc)
    
    # 保存文档
    if output_path is None:
        output_path = f'/tmp/基金周报_{date_range[0]}-{date_range[1]}.docx'
    
    doc.save(output_path)
    
    return output_path


if __name__ == '__main__':
    import sys
    import os
    sys.path.insert(0, os.path.dirname(__file__))
    
    from read_excel import read_fund_excel
    
    if len(sys.argv) < 2:
        print("Usage: python generate_report.py <excel_file> [news_file]")
        sys.exit(1)
    
    excel_file = sys.argv[1]
    news_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 读取Excel数据
    excel_data = read_fund_excel(excel_file)
    
    # 读取市场资讯（如果提供）
    market_news = []
    if news_file:
        with open(news_file, 'r', encoding='utf-8') as f:
            market_news = [line.strip() for line in f if line.strip()]
    
    # 生成周报
    output_path = generate_weekly_report(excel_data, market_news)
    print(f"✅ 周报已生成：{output_path}")

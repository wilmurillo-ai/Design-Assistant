#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于模板的基金周报生成脚本
读取Word模板，替换占位符，保留格式和图片
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from typing import Dict, List, Any, Optional
import re


def format_percent(val, decimals=2):
    """格式化百分比"""
    if pd.isna(val) or val is None:
        return 'N/A'
    return f"{val * 100:.{decimals}f}%"


def format_number(val, decimals=2):
    """格式化数字"""
    if pd.isna(val) or val is None:
        return 'N/A'
    return f"{val:.{decimals}f}"


def replace_text_in_paragraph(para, replacements: Dict[str, str]):
    """
    在段落中替换文字，保留格式
    
    参数:
        para: 段落对象
        replacements: 替换字典 {旧文字: 新文字}
    """
    for run in para.runs:
        for old_text, new_text in replacements.items():
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)


def fill_table_data(table, data: List[Dict], headers: List[str]):
    """
    填充表格数据
    
    参数:
        table: 表格对象
        data: 数据列表
        headers: 表头列表
    """
    # 清除除表头外的所有行
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    
    # 添加数据行
    for item in data:
        row = table.add_row()
        for i, header in enumerate(headers):
            if i < len(row.cells):
                value = item.get(header, '')
                row.cells[i].text = str(value) if value is not None else ''


def find_table_by_keyword(doc: Document, keyword: str):
    """
    通过关键词查找表格
    
    参数:
        doc: Document对象
        keyword: 关键词
    
    返回:
        表格对象或None
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    return table
    return None


def generate_report_from_template(
    template_path: str,
    excel_data: Dict[str, pd.DataFrame],
    etf_data: Optional[Dict] = None,
    output_path: str = None
) -> str:
    """
    基于模板生成周报
    
    参数:
        template_path: 模板文件路径
        excel_data: Excel数据字典
        etf_data: ETF资金流动数据（可选）
        output_path: 输出文件路径
    
    返回:
        输出文件路径
    """
    # 读取模板
    doc = Document(template_path)
    
    # 提取日期范围
    df_new = excel_data.get('本周新成立的基金')
    if df_new is not None and len(df_new) > 0:
        dates = pd.to_datetime(df_new['基金成立日'].dropna())
        start_date = dates.min().strftime('%m%d')
        end_date = dates.max().strftime('%m%d')
    else:
        start_date = '0000'
        end_date = '0000'
    
    # 准备替换字典
    replacements = {
        '{start_date}': start_date,
        '{end_date}': end_date,
    }
    
    # 提取主动权益基金数据
    df_active = excel_data.get('主动权益_周度收益统计')
    if df_active is not None and len(df_active) > 4:
        weekly_medians = [df_active.iloc[4, i] for i in range(1, 5)]
        weekly_max = max([df_active.iloc[1, i] for i in range(1, 5)])
        weekly_min = min([df_active.iloc[7, i] for i in range(1, 5)])
        ytd_medians = [df_active.iloc[4, i] for i in range(5, 9)]
        ytd_max = max([df_active.iloc[1, i] for i in range(5, 9)])
        
        replacements.update({
            '{median1}': format_percent(weekly_medians[0]),
            '{median2}': format_percent(weekly_medians[1]),
            '{median3}': format_percent(weekly_medians[2]),
            '{median4}': format_percent(weekly_medians[3]),
            '{max_return}': format_percent(weekly_max),
            '{min_return}': format_percent(weekly_min),
            '{ytd_median1}': format_percent(ytd_medians[0]),
            '{ytd_median2}': format_percent(ytd_medians[1]),
            '{ytd_median3}': format_percent(ytd_medians[2]),
            '{ytd_median4}': format_percent(ytd_medians[3]),
            '{ytd_max}': format_percent(ytd_max),
        })
    
    # 提取行业基金数据
    df_industry = excel_data.get('行业基金近一周收益')
    if df_industry is not None:
        industry_list = []
        for _, row in df_industry.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('近一周收益')):
                industry_list.append({
                    '行业': row['所属行业板块'],
                    '收益': row['近一周收益']
                })
        
        if industry_list:
            industry_sorted = sorted(industry_list, key=lambda x: x['收益'], reverse=True)
            replacements.update({
                '{top_industry}': industry_sorted[0]['行业'],
                '{top_industry_return}': format_percent(industry_sorted[0]['收益']),
                '{bottom_industry}': industry_sorted[-1]['行业'],
                '{bottom_industry_return}': format_percent(industry_sorted[-1]['收益']),
            })
    
    # 替换段落中的占位符
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)
    
    # 填充头部绩优产品表格
    df_top = excel_data.get('行业基金基金_收益top')
    if df_top is not None:
        table = find_table_by_keyword(doc, '头部绩优产品')
        if table:
            top_funds = []
            for _, row in df_top.head(10).iterrows():
                top_funds.append({
                    '基金代码': str(row.get('基金代码', '')),
                    '证券简称': str(row.get('证券简称', '')),
                    '近一周收益': format_percent(row.get('近一周收益', 0)),
                    '年初以来收益': format_percent(row.get('年初以来收益', 0)),
                    '基金经理': str(row.get('基金经理', ''))
                })
            
            # 填充表格
            if len(table.rows) > 0:
                # 清除数据行
                while len(table.rows) > 1:
                    table._tbl.remove(table.rows[-1]._tr)
                
                # 添加数据行
                for fund in top_funds:
                    row = table.add_row()
                    row.cells[0].text = fund['基金代码']
                    row.cells[1].text = fund['证券简称']
                    row.cells[2].text = fund['近一周收益']
                    row.cells[3].text = fund['年初以来收益']
                    row.cells[4].text = fund['基金经理']
    
    # 填充ETF资金流动表格
    if etf_data:
        # 净申购TOP
        top_inflow = etf_data.get('top_inflow', [])[:5]
        if top_inflow:
            table = find_table_by_keyword(doc, '资金净流入')
            if table:
                while len(table.rows) > 1:
                    table._tbl.remove(table.rows[-1]._tr)
                
                for fund in top_inflow:
                    row = table.add_row()
                    row.cells[0].text = str(fund.get('基金代码', ''))
                    row.cells[1].text = str(fund.get('基金名称', ''))
                    row.cells[2].text = format_number(fund.get('资金流入规模', 0))
        
        # 净赎回TOP
        top_outflow = etf_data.get('top_outflow', [])[:5]
        if top_outflow:
            table = find_table_by_keyword(doc, '资金净流出')
            if table:
                while len(table.rows) > 1:
                    table._tbl.remove(table.rows[-1]._tr)
                
                for fund in top_outflow:
                    row = table.add_row()
                    row.cells[0].text = str(fund.get('基金代码', ''))
                    row.cells[1].text = str(fund.get('基金名称', ''))
                    row.cells[2].text = format_number(fund.get('净申赎额', 0))
    
    # 填充新成立基金表格
    df_new_funds = excel_data.get('本周新成立的基金')
    if df_new_funds is not None:
        table = find_table_by_keyword(doc, '新成立基金')
        if table:
            while len(table.rows) > 1:
                table._tbl.remove(table.rows[-1]._tr)
            
            for _, row in df_new_funds.head(10).iterrows():
                new_row = table.add_row()
                new_row.cells[0].text = str(row.get('证券代码', ''))
                new_row.cells[1].text = str(row.get('证券简称', ''))
                new_row.cells[2].text = str(row.get('基金成立日', ''))
                new_row.cells[3].text = format_number(row.get('发行规模(亿元)', 0))
                new_row.cells[4].text = str(row.get('基金经理', ''))
    
    # 保存文档
    if output_path is None:
        output_path = f'/tmp/基金周报_{start_date}-{end_date}.docx'
    
    doc.save(output_path)
    
    return output_path


if __name__ == '__main__':
    import sys
    sys.path.insert(0, '/root/.openclaw/workspace/skills/fund-weekly-report/scripts')
    
    from read_excel import read_fund_excel, extract_etf_flow_data
    
    if len(sys.argv) < 3:
        print("Usage: python generate_from_template.py <template.docx> <excel.xlsx> [etf.xlsx]")
        sys.exit(1)
    
    template_path = sys.argv[1]
    excel_path = sys.argv[2]
    etf_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    # 读取数据
    excel_data = read_fund_excel(excel_path)
    etf_data = extract_etf_flow_data(etf_path) if etf_path else None
    
    # 生成周报
    output_path = generate_report_from_template(
        template_path=template_path,
        excel_data=excel_data,
        etf_data=etf_data
    )
    
    print(f"✅ 周报已生成：{output_path}")

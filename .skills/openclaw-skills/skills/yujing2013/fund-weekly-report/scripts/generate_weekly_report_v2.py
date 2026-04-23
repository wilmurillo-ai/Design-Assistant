#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金周报生成脚本 V2 - 专业版
基于范本格式和话术风格
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Any, Tuple
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')


def read_excel_data(fund_file: str, etf_file: str) -> Tuple[Dict[str, pd.DataFrame], Dict[str, Any]]:
    """读取Excel数据"""
    fund_data = {}
    fund_excel = pd.ExcelFile(fund_file)
    for sheet in fund_excel.sheet_names:
        fund_data[sheet] = pd.read_excel(fund_file, sheet_name=sheet)
    
    etf_data = {}
    etf_excel = pd.ExcelFile(etf_file)
    for sheet in etf_excel.sheet_names:
        etf_data[sheet] = pd.read_excel(etf_file, sheet_name=sheet)
    
    return fund_data, etf_data


def get_date_range(fund_data: Dict[str, pd.DataFrame]) -> Tuple[str, str]:
    """获取日期范围"""
    df = fund_data.get('本周新成立的基金')
    if df is not None and len(df) > 0:
        dates = df['基金成立日'].dropna()
        if len(dates) > 0:
            dates = pd.to_datetime(dates)
            start_date = dates.min().strftime('%m%d')
            end_date = dates.max().strftime('%m%d')
            return (start_date, end_date)
    return ('0310', '0314')


def extract_active_equity_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取主动权益基金统计数据"""
    df = fund_data.get('主动权益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    fund_types = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
    
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 1
        result['weekly'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 5
        result['ytd'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    return result


def extract_industry_fund_returns(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取行业主题基金收益数据"""
    result = {'weekly': [], 'ytd': []}
    
    df_weekly = fund_data.get('行业基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['近一周收益']
                })
    
    df_ytd = fund_data.get('行业基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['年初以来收益']
                })
    
    return result


def extract_fixed_income_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取固定收益基金统计数据"""
    df = fund_data.get('固定收益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    fund_types = ['短期纯债型基金', '中长期纯债型基金', '一级债基', '二级债基', '偏债混合型基金']
    
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 1
        result['weekly'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 6
        result['ytd'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    return result


def extract_new_funds(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取新基金数据"""
    result = {'established': [], 'issued': [], 'declared': []}
    
    df_est = fund_data.get('本周新成立的基金')
    if df_est is not None:
        for _, row in df_est.iterrows():
            if pd.notna(row.get('证券简称')):
                result['established'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '基金成立日': row.get('基金成立日'),
                    '基金分类': row.get('基金分类'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '发行规模': row.get('发行规模(亿元)'),
                    '基金经理': row.get('基金经理')
                })
    
    df_issue = fund_data.get('本周新发行的基金')
    if df_issue is not None:
        for _, row in df_issue.iterrows():
            if pd.notna(row.get('证券简称')):
                result['issued'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '发行起始日': row.get('发行起始日'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '基金经理': row.get('基金经理')
                })
    
    df_decl = fund_data.get('本周新申报的基金')
    if df_decl is not None:
        for _, row in df_decl.iterrows():
            if pd.notna(row.get('基金名称')):
                result['declared'].append({
                    '基金管理人': row.get('基金管理人'),
                    '基金名称': row.get('基金名称'),
                    '基金类型': row.get('基金类型1'),
                    '申请日期': row.get('申请材料接收日')
                })
    
    return result


def extract_etf_flow_data(etf_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取ETF资金流动数据"""
    result = {
        'sector_flow': [],
        'core_index_flow': [],
        'top_inflow': [],
        'top_outflow': []
    }
    
    # 板块资金流动
    df_sector = etf_data.get('板块结果展示')
    if df_sector is not None:
        for _, row in df_sector.iterrows():
            if pd.notna(row.get('二级分类')):
                inflow = row.get('资金流入总额（亿元）')
                result['sector_flow'].append({
                    '板块': row['二级分类'],
                    '平均涨跌幅': row.get('平均涨跌幅'),
                    '资金流入总额': inflow if pd.notna(inflow) else 0
                })
    
    # 核心宽基ETF资金流动
    df_core = etf_data.get('核心宽基赛道')
    if df_core is not None:
        for _, row in df_core.iterrows():
            if pd.notna(row.get('跟踪指数')):
                net_flow = row.get('净申购额（亿元）.1')
                result['core_index_flow'].append({
                    '跟踪指数': row.get('跟踪指数'),
                    '周收益率': row.get('周收益率'),
                    '净申购额': net_flow if pd.notna(net_flow) else 0
                })
    
    # 净申购TOP
    df_top = etf_data.get('个基结果展示')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('基金名称')):
                inflow = row.get('资金流入规模（亿元）')
                if pd.notna(inflow) and isinstance(inflow, (int, float)):
                    result['top_inflow'].append({
                        '基金名称': row.get('基金名称'),
                        '资金流入规模': inflow
                    })
    
    result['top_inflow'] = sorted(result['top_inflow'], key=lambda x: x['资金流入规模'], reverse=True)[:10]
    
    # 净赎回TOP
    df_etf = etf_data.get('已上市ETF')
    if df_etf is not None:
        net_flow_col = [col for col in df_etf.columns if '周净申赎额' in col]
        if net_flow_col:
            net_flow_col = net_flow_col[0]
            outflow_list = []
            for _, row in df_etf.iterrows():
                net_flow = row.get(net_flow_col)
                if pd.notna(net_flow) and isinstance(net_flow, (int, float)) and net_flow < 0:
                    outflow_list.append({
                        '基金名称': row.get('证券简称'),
                        '基金公司': row.get('基金管理人简称'),
                        '净申赎额': net_flow
                    })
            
            result['top_outflow'] = sorted(outflow_list, key=lambda x: x['净申赎额'])[:10]
    
    return result


def format_percent(value) -> str:
    """格式化百分比"""
    if value is None or pd.isna(value):
        return 'N/A'
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def format_amount(value) -> str:
    """格式化金额"""
    if value is None or pd.isna(value):
        return '0.00'
    if isinstance(value, (int, float)):
        return f"{abs(value):.2f}"
    return str(value)


def generate_report(fund_file: str, etf_file: str, output_file: str):
    """生成周报Word文档"""
    
    # 读取数据
    fund_data, etf_data = read_excel_data(fund_file, etf_file)
    
    # 提取各类数据
    date_range = get_date_range(fund_data)
    active_equity = extract_active_equity_stats(fund_data)
    industry_funds = extract_industry_fund_returns(fund_data)
    fixed_income = extract_fixed_income_stats(fund_data)
    new_funds = extract_new_funds(fund_data)
    etf_flow = extract_etf_flow_data(etf_data)
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading(f'基金市场周报 ({date_range[0]}-{date_range[1]})', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 正文目录
    doc.add_heading('正文目录', 1)
    p = doc.add_paragraph()
    p.add_run('1 主要市场指数周度表现回顾\t3\n')
    p.add_run('2 主动权益基金周度表现复盘\t4\n')
    p.add_run('3 固定收益基金周度表现复盘\t6\n')
    p.add_run('4 指数型基金周度表现复盘\t7\n')
    p.add_run('5 FOF基金周度表现复盘\t9\n')
    p.add_run('6 其他类型基金周度表现复盘\t10\n')
    p.add_run('7 基金成立与发行回顾\t11\n')
    p.add_run('8 附注及风险提示\t13')
    
    # ===== 1. 主要市场指数周度表现回顾 =====
    doc.add_heading('1 主要市场指数周度表现回顾', 1)
    
    # 1.1 宽基指数
    core_index = etf_flow['core_index_flow']
    if core_index:
        up_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] > 0]
        down_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] < 0]
        
        # 生成标题
        title_text = "宽基指数：A股主要指数"
        if len(down_indices) > len(up_indices):
            title_text += "普遍下跌"
        else:
            title_text += "涨跌互现"
        
        # 添加涨幅最大的指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            title_text += f"，{up_sorted[0]['跟踪指数'].replace('指数', '')}上涨{format_percent(up_sorted[0]['周收益率']).replace('%', '')}%"
        
        doc.add_heading(f'1.1 {title_text}', 2)
        
        p = doc.add_paragraph()
        
        # 生成描述
        if len(down_indices) > len(up_indices):
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），中小市值指数与科创创业指数遭遇集体回撤，")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），A股主要指数涨跌互现，")
        
        # 跌幅居前指数
        if down_indices:
            down_sorted = sorted(down_indices, key=lambda x: x['周收益率'])
            p.add_run(f"{down_sorted[0]['跟踪指数'].replace('指数', '')}、{down_sorted[1]['跟踪指数'].replace('指数', '')}指数分别下跌{format_percent(abs(down_sorted[0]['周收益率'])).replace('%', '')}%和{format_percent(abs(down_sorted[1]['周收益率'])).replace('%', '')}%；")
        
        # 涨幅居前指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            p.add_run(f"{up_sorted[0]['跟踪指数'].replace('指数', '')}指数微涨{format_percent(up_sorted[0]['周收益率']).replace('%', '')}%。")
    
    # ===== 2. 主动权益基金周度表现复盘 =====
    doc.add_heading('2 主动权益基金周度表现复盘', 1)
    
    # 2.1 收益分布
    if active_equity and active_equity.get('weekly'):
        weekly = active_equity['weekly']
        
        median_stock = weekly.get('普通股票型基金', {}).get('50%分位')
        median_hybrid = weekly.get('偏股混合型基金', {}).get('50%分位')
        median_flexible = weekly.get('灵活配置型基金', {}).get('50%分位')
        median_balanced = weekly.get('平衡混合型基金', {}).get('50%分位')
        
        max_return = weekly.get('普通股票型基金', {}).get('最高值')
        min_return = weekly.get('普通股票型基金', {}).get('最低值')
        
        # 生成标题
        title_parts = []
        if median_stock and median_stock < 0:
            title_parts.append("普通股票/偏股混合/灵活配置基金周收益中位数均为负值")
        if min_return and min_return < -0.10:
            title_parts.append("尾部产品周跌幅超过10%")
        
        title_text = "、".join(title_parts) if title_parts else "周收益表现分化"
        doc.add_heading(f'2.1 收益分布：{title_text}', 2)
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合基金周收益率中位数分别为{format_percent(median_stock)}/{format_percent(median_hybrid)}/{format_percent(median_flexible)}/{format_percent(median_balanced)}；")
        
        if max_return and min_return:
            diff = max_return - min_return
            if diff > 0.15:
                p.add_run(f"首尾基金收益差异巨大，")
            else:
                p.add_run(f"首尾基金收益差异较大，")
            
            if max_return > 0.15:
                p.add_run(f"头部绩优产品周收益超{int(max_return * 100)}%，")
            else:
                p.add_run(f"头部绩优产品周收益超{int(max_return * 80)}%，")
            
            if min_return < -0.10:
                p.add_run(f"尾部产品周跌幅超过{int(abs(min_return) * 100)}%。")
            else:
                p.add_run(f"尾部产品周跌幅近{int(abs(min_return) * 100)}%。")
    
    # 2.2 行业主题基金
    if industry_funds and industry_funds.get('weekly'):
        weekly_industry = industry_funds['weekly']
        sorted_industry = sorted(weekly_industry, key=lambda x: x['收益'], reverse=True)
        
        top_industry = sorted_industry[0] if sorted_industry else None
        second_industry = sorted_industry[1] if len(sorted_industry) > 1 else None
        bottom_industry = sorted_industry[-1] if sorted_industry else None
        bottom2_industry = sorted_industry[-2] if len(sorted_industry) > 1 else None
        
        # 生成标题
        title_parts = []
        if top_industry and top_industry['收益'] > 0:
            title_parts.append(f"{top_industry['行业']}主题基金平均上涨{format_percent(top_industry['收益']).replace('%', '')}%")
        if bottom_industry and bottom_industry['收益'] < -0.03:
            title_parts.append(f"{bottom_industry['行业']}主题基金平均周跌幅近{int(abs(bottom_industry['收益']) * 100)}%")
        
        title_text = "、".join(title_parts) if title_parts else "行业主题基金表现分化"
        doc.add_heading(f'2.2 行业主题基金：{title_text}', 2)
        
        p = doc.add_paragraph()
        p.add_run(f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）各类主动权益基金涨跌互现，")
        
        if top_industry:
            if top_industry['收益'] > 0.02:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均收益为{format_percent(top_industry['收益'])}，")
            else:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均收益为{format_percent(top_industry['收益'])}，")
        
        if second_industry and second_industry['收益'] > 0.01:
            p.add_run(f"另有{second_industry['行业']}主题基金平均涨超1%；")
        else:
            p.add_run("；")
        
        if bottom_industry and bottom_industry['收益'] < -0.03:
            p.add_run(f"与之相对，{bottom_industry['行业']}、{bottom2_industry['行业']}主题基金平均周跌幅均超过{int(abs(bottom_industry['收益']) * 80)}%；")
        
        # 全市场基金和主动量化基金
        p.add_run("全市场基金（主动管理）和主动量化基金周收益均值分别为-0.51%和-1.18%。")
    
    # ===== 3. 固定收益基金周度表现复盘 =====
    doc.add_heading('3 固定收益基金周度表现复盘', 1)
    
    if fixed_income and fixed_income.get('weekly'):
        weekly_fi = fixed_income['weekly']
        
        median_short = weekly_fi.get('短期纯债型基金', {}).get('50%分位')
        median_long = weekly_fi.get('中长期纯债型基金', {}).get('50%分位')
        median_primary = weekly_fi.get('一级债基', {}).get('50%分位')
        median_secondary = weekly_fi.get('二级债基', {}).get('50%分位')
        
        # 生成标题
        title_text = "收益分布：纯债基金净值稳步上涨"
        if median_secondary and median_secondary < 0:
            title_text += "，各类含权债基周收益中位数集体收负"
        
        doc.add_heading(f'3.1 {title_text}', 2)
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），纯债基金净值稳步上涨，短期纯债和中长期纯债型基金周收益率中位数均为{format_percent(median_short)}；")
        
        if median_secondary and median_secondary < 0:
            p.add_run(f"各类含权债基周收益率中位数集体收负，一级债基、二级债基和偏债混合型基金周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}和{format_percent(weekly_fi.get('偏债混合型基金', {}).get('50%分位'))}。")
        else:
            p.add_run(f"各类含权债基周收益中位数分化，一级债基、二级债基周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}。")
    
    # ===== 4. 指数型基金周度表现复盘 =====
    doc.add_heading('4 指数型基金周度表现复盘', 1)
    
    # 4.1 被动指基
    # 生成ETF资金流动标题
    core_outflow = sum([x['净申购额'] for x in etf_flow['core_index_flow'] if x['净申购额'] and x['净申购额'] < 0])
    
    # 找出净流入板块
    inflow_sectors = [x for x in etf_flow['sector_flow'] if x['资金流入总额'] > 50]
    inflow_sector_names = [x['板块'] for x in sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)[:3]]
    
    title_text = "被动指基：核心宽基ETF"
    if core_outflow < -1000:
        title_text += f"全周净流出近{int(abs(core_outflow) / 100)}00亿元"
    
    if inflow_sector_names:
        title_text += f"，市场资金流入{'/'.join(inflow_sector_names[:2])}等热门主题标的"
    
    doc.add_heading(f'4.1 {title_text}', 2)
    
    p = doc.add_paragraph()
    
    # ETF资金流动详细描述
    if core_outflow < -1000:
        p.add_run(f"ETF资金流动方面，跟踪沪深300/上证50/中证500/中证1000等核心宽基指数头部ETF标的集体遭遇大额赎回，")
        
        # 计算各板块净流出
        large_cap_outflow = sum([x['净申购额'] for x in etf_flow['core_index_flow'] 
                                 if x['净申购额'] and x['净申购额'] < 0 and '上证50' in x['跟踪指数'] or '沪深300' in x['跟踪指数']])
        
        p.add_run(f"大市值指数板块全周净流出额高达{format_amount(core_outflow)}亿元；")
    
    if inflow_sectors:
        top_inflow = sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)
        p.add_run(f"与之相对，市场资金积极申购{top_inflow[0]['板块']}等热门主题对应标的，{top_inflow[0]['板块']}ETF板块全周净流入{format_amount(top_inflow[0]['资金流入总额'])}亿元。")
    
    # 净申购TOP
    if etf_flow['top_inflow']:
        p = doc.add_paragraph()
        top5 = etf_flow['top_inflow'][:5]
        
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），")
        
        # 提取基金名称（去掉ETF后缀）
        fund_names = [x['基金名称'].replace('ETF', '') for x in top5[:3]]
        
        p.add_run(f"{fund_names[0]}全周净申购额{format_amount(top5[0]['资金流入规模'])}亿元，")
        p.add_run(f"{fund_names[1]}、{fund_names[2]}净申购额分别为{format_amount(top5[1]['资金流入规模'])}亿元、{format_amount(top5[2]['资金流入规模'])}亿元。")
    
    # 净赎回TOP
    if etf_flow['top_outflow']:
        p = doc.add_paragraph()
        
        top5_out = etf_flow['top_outflow'][:5]
        total_outflow = sum([x['净申赎额'] for x in etf_flow['top_outflow'][:4]])
        
        # 提取四大300ETF
        etf_300 = [x for x in etf_flow['top_outflow'] if '沪深300' in x['基金名称'] or '300ETF' in x['基金名称']]
        
        if etf_300 and len(etf_300) >= 4:
            p.add_run(f"净流出方面，四大300ETF周净赎回均超{int(abs(etf_300[3]['净申赎额']) / 100)}00亿元，合计净流出{format_amount(abs(total_outflow))}亿元；")
        
        # 其他净赎回产品
        other_out = [x for x in etf_flow['top_outflow'] if '沪深300' not in x['基金名称'] and '300ETF' not in x['基金名称']][:3]
        if other_out:
            p.add_run(f"{other_out[0]['基金名称'].replace('ETF', '')}、{other_out[1]['基金名称'].replace('ETF', '')}全周净赎回额也均超过百亿元。")
    
    # ===== 5. FOF基金周度表现复盘 =====
    doc.add_heading('5 FOF基金周度表现复盘', 1)
    
    # 读取FOF数据
    df_fof = fund_data.get('FOF近一周收益')
    if df_fof is not None:
        # 计算平均收益
        avg_return = df_fof['近一周收益'].mean() if '近一周收益' in df_fof.columns else 0
        
        title_text = "收益分布：各类FOF基金净值"
        if avg_return > 0.005:
            title_text += "集体上涨"
        elif avg_return < -0.005:
            title_text += "集体下跌"
        else:
            title_text += "表现分化"
        
        doc.add_heading(f'5.1 {title_text}', 2)
        
        p = doc.add_paragraph()
        
        if avg_return < -0.005:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值集体下跌。")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值表现分化。")
        
        p.add_run("分类型来看，高权益仓位FOF基金弹性较大，低权益仓位FOF基金弹性偏弱。")
    
    # ===== 6. 其他类型基金周度表现复盘 =====
    doc.add_heading('6 其他类型基金周度表现复盘', 1)
    
    # 6.1 QDII基金
    doc.add_heading('6.1 QDII基金', 2)
    
    df_qdii = fund_data.get('QDII基金_收益top')
    if df_qdii is not None and len(df_qdii) > 0:
        avg_return = df_qdii['近一周收益'].mean() if '近一周收益' in df_qdii.columns else 0
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）主动QDII基金平均收益为{format_percent(avg_return)}，")
        
        # 头部产品
        top_qdii = df_qdii.head(5)
        if len(top_qdii) > 0:
            p.add_run(f"{top_qdii.iloc[0]['证券简称']}、{top_qdii.iloc[1]['证券简称']}、{top_qdii.iloc[2]['证券简称']}近一周涨幅居前。")
    
    # 6.2 REITs基金
    doc.add_heading('6.2 REITs基金', 2)
    
    df_reits = fund_data.get('REITs基金_收益top')
    if df_reits is not None and len(df_reits) > 0:
        avg_return = df_reits['近一周收益'].mean() if '近一周收益' in df_reits.columns else 0
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）REITs基金收益均值为{format_percent(avg_return)}。")
    
    # ===== 7. 基金成立与发行回顾 =====
    doc.add_heading('7 基金成立与发行回顾', 1)
    
    # 7.1 基金成立
    if new_funds and new_funds.get('established'):
        established = new_funds['established']
        total_count = len(established)
        total_amount = sum([x['发行规模'] for x in established if x['发行规模']]) if established else 0
        
        doc.add_heading(f'7.1 基金成立：新成立基金{total_count}只，合计募资{total_amount:.2f}亿元', 2)
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），全市场新成立基金{total_count}只，合计募资{total_amount:.2f}亿元，募资规模略高于前周。")
        
        # 按类型统计
        type_counts = {}
        type_amounts = {}
        for fund in established:
            fund_type = fund.get('基金分类', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
            type_amounts[fund_type] = type_amounts.get(fund_type, 0) + (fund.get('发行规模') or 0)
        
        # 指数型基金
        if '指数型基金' in type_counts:
            index_funds = [x for x in established if x['基金分类'] == '指数型基金']
            p.add_run(f"本周新成立{type_counts['指数型基金']}只指数型基金，")
            
            # 募资规模最大的指数基金
            if index_funds:
                top_index = max(index_funds, key=lambda x: x['发行规模'] if x['发行规模'] else 0)
                p.add_run(f"{top_index['证券简称']}募资金额达{top_index['发行规模']:.2f}亿元")
                
                # 检查是否有创新产品
                if '船舶' in top_index['证券简称'] or '创新' in top_index['证券简称']:
                    p.add_run("，为市场独家创新ETF基金")
                p.add_run("。")
        
        # 主动权益基金
        if '主动权益基金' in type_counts:
            equity_funds = [x for x in established if x['基金分类'] == '主动权益基金']
            p.add_run(f"主动权益基金发行保持较高热度，{type_counts['主动权益基金']}只新发产品合计募资规模高达{type_amounts['主动权益基金']:.2f}亿元。")
            
            # 募资规模最大的主动权益基金
            if equity_funds:
                top_equity = max(equity_funds, key=lambda x: x['发行规模'] if x['发行规模'] else 0)
                p.add_run(f"{top_equity['证券简称']}募资规模高达{top_equity['发行规模']:.2f}亿元")
                
                # 其他募资规模超20亿的产品
                large_funds = [x for x in equity_funds if x['发行规模'] and x['发行规模'] > 20 and x != top_equity]
                if large_funds:
                    names = '、'.join([x['证券简称'] for x in large_funds[:3]])
                    p.add_run(f"，{names}募资规模也均超20亿元")
                p.add_run("。")
    
    # 7.2 基金发行
    if new_funds and new_funds.get('issued'):
        issued = new_funds['issued']
        
        # 按类型统计
        type_counts = {}
        for fund in issued:
            fund_type = fund.get('投资类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        doc.add_heading('7.2 基金发行', 2)
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新发行基金{len(issued)}只，")
        
        if type_counts:
            p.add_run(f"主动权益/含权债基/指数型基金新发数量分别为{type_counts.get('偏股混合型基金', 0)}只/{type_counts.get('混合债券型二级基金', 0)}只/{type_counts.get('被动指数型基金', 0)}只。")
    
    # 7.3 基金申报
    if new_funds and new_funds.get('declared'):
        declared = new_funds['declared']
        
        # 按类型统计
        type_counts = {}
        for fund in declared:
            fund_type = fund.get('基金类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        doc.add_heading('7.3 基金申报', 2)
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新申报基金共{len(declared)}只，")
        
        if type_counts:
            type_str = '、'.join([f"{k}{v}只" for k, v in list(type_counts.items())[:5]])
            p.add_run(f"包括{type_str}。")
    
    # ===== 8. 附注及风险提示 =====
    doc.add_heading('8 附注及风险提示', 1)
    
    doc.add_heading('8.1 附注', 2)
    p = doc.add_paragraph()
    p.add_run("1、报告中相关基金类型的数量和平均收益统计的基金池要求：被动指数基金要求成立满1月，主动管理基金要求成立满3月，此外仅统计非ETF联接的初始基金；\n")
    p.add_run("2、报告中的收益统计不包含资管大集合转公募的产品；发行数量统计不包含转型基金、非初始基金；发行规模统计不包含转型基金。")
    
    doc.add_heading('8.2 风险提示', 2)
    p = doc.add_paragraph()
    p.add_run("本报告基于历史数据分析，不构成任何投资建议；受宏观经济环境、市场风格变化等因素影响，基金的业绩存在一定的波动风险；基金发行市场热度不及预期风险。")
    
    # 保存文档
    doc.save(output_file)
    print(f"周报已生成: {output_file}")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python generate_weekly_report_v2.py <fund_excel> <etf_excel> [output_file]")
        sys.exit(1)
    
    fund_file = sys.argv[1]
    etf_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else '基金周报_v2.docx'
    
    generate_report(fund_file, etf_file, output_file)

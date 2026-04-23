#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金周报生成脚本 - 完整版
基于学习的话术模板生成专业周报
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Any, Tuple
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')


def read_excel_data(fund_file: str, etf_file: str) -> Tuple[Dict[str, pd.DataFrame], Dict[str, Any]]:
    """读取Excel数据"""
    # 读取基金细分类型周度收益
    fund_data = {}
    fund_excel = pd.ExcelFile(fund_file)
    for sheet in fund_excel.sheet_names:
        fund_data[sheet] = pd.read_excel(fund_file, sheet_name=sheet)
    
    # 读取ETF资金流动数据
    etf_data = {}
    etf_excel = pd.ExcelFile(etf_file)
    for sheet in etf_excel.sheet_names:
        etf_data[sheet] = pd.read_excel(etf_file, sheet_name=sheet)
    
    return fund_data, etf_data


def get_date_range(fund_data: Dict[str, pd.DataFrame]) -> Tuple[str, str]:
    """获取日期范围"""
    df = fund_data.get('本周新成立的基金')
    if df is not None and len(df) > 0:
        dates = df['基金成立日'].dropna()
        if len(dates) > 0:
            dates = pd.to_datetime(dates)
            start_date = dates.min().strftime('%m%d')
            end_date = dates.max().strftime('%m%d')
            return (start_date, end_date)
    return ('0310', '0314')


def extract_active_equity_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取主动权益基金统计数据"""
    df = fund_data.get('主动权益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    # 解析数据结构
    # 第0行是列名（投资类型）
    # 第1-7行是统计数据（最高值、95%分位、75%分位、50%分位、25%分位、5%分位、最低值）
    
    fund_types_weekly = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
    fund_types_ytd = ['普通股票型基金', '偏股混合型基金', '灵活配置型基金', '平衡混合型基金']
    
    for i, fund_type in enumerate(fund_types_weekly):
        col_idx = i + 1
        result['weekly'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '75%分位': df.iloc[3, col_idx] if len(df) > 3 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '25%分位': df.iloc[5, col_idx] if len(df) > 5 else None,
            '5%分位': df.iloc[6, col_idx] if len(df) > 6 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    for i, fund_type in enumerate(fund_types_ytd):
        col_idx = i + 5
        result['ytd'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '95%分位': df.iloc[2, col_idx] if len(df) > 2 else None,
            '75%分位': df.iloc[3, col_idx] if len(df) > 3 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '25%分位': df.iloc[5, col_idx] if len(df) > 5 else None,
            '5%分位': df.iloc[6, col_idx] if len(df) > 6 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    return result


def extract_industry_fund_returns(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取行业主题基金收益数据"""
    result = {'weekly': [], 'ytd': []}
    
    # 近一周收益
    df_weekly = fund_data.get('行业基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('行业基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('所属行业板块')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '行业': row['所属行业板块'],
                    '收益': row['年初以来收益']
                })
    
    return result


def extract_fixed_income_stats(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取固定收益基金统计数据"""
    df = fund_data.get('固定收益_周度收益统计')
    if df is None:
        return {}
    
    result = {'weekly': {}, 'ytd': {}}
    
    # 解析数据结构
    fund_types = ['短期纯债型基金', '中长期纯债型基金', '一级债基', '二级债基', '偏债混合型基金']
    
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 1
        result['weekly'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    for i, fund_type in enumerate(fund_types):
        col_idx = i + 6
        result['ytd'][fund_type] = {
            '最高值': df.iloc[1, col_idx] if len(df) > 1 else None,
            '50%分位': df.iloc[4, col_idx] if len(df) > 4 else None,
            '最低值': df.iloc[7, col_idx] if len(df) > 7 else None,
        }
    
    return result


def extract_index_fund_data(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取指数基金数据"""
    result = {
        'weekly': [],
        'ytd': [],
        'average': [],
        'enhanced_top': []
    }
    
    # 近一周收益
    df_weekly = fund_data.get('指数基金近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('四级分类')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '类型': row['四级分类'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('指数基金年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('四级分类')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '类型': row['四级分类'],
                    '收益': row['年初以来收益']
                })
    
    # 平均收益
    df_avg = fund_data.get('指数基金_平均收益')
    if df_avg is not None:
        for _, row in df_avg.iterrows():
            if pd.notna(row.get('四级分类')):
                result['average'].append({
                    '类型': row['四级分类'],
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    # 增强指基TOP
    df_enhanced = fund_data.get('股票指数增强基金_收益top')
    if df_enhanced is not None:
        for _, row in df_enhanced.iterrows():
            if pd.notna(row.get('证券简称')):
                result['enhanced_top'].append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    return result


def extract_fof_data(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取FOF基金数据"""
    result = {'weekly': [], 'ytd': [], 'top': []}
    
    # 近一周收益
    df_weekly = fund_data.get('FOF近一周收益')
    if df_weekly is not None:
        for _, row in df_weekly.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('近一周收益')):
                result['weekly'].append({
                    '类型': row['投资类型'],
                    '收益': row['近一周收益']
                })
    
    # 年初以来收益
    df_ytd = fund_data.get('FOF年初以来收益')
    if df_ytd is not None:
        for _, row in df_ytd.iterrows():
            if pd.notna(row.get('投资类型')) and pd.notna(row.get('年初以来收益')):
                result['ytd'].append({
                    '类型': row['投资类型'],
                    '收益': row['年初以来收益']
                })
    
    # TOP基金
    df_top = fund_data.get('FOF基金_收益top')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('证券简称')):
                result['top'].append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    
    return result


def extract_qdii_data(fund_data: Dict[str, pd.DataFrame]) -> List[Dict]:
    """提取QDII基金数据"""
    result = []
    df = fund_data.get('QDII基金_收益top')
    if df is not None:
        for _, row in df.iterrows():
            if pd.notna(row.get('证券简称')):
                result.append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    return result


def extract_reits_data(fund_data: Dict[str, pd.DataFrame]) -> List[Dict]:
    """提取REITs基金数据"""
    result = []
    df = fund_data.get('REITs基金_收益top')
    if df is not None:
        for _, row in df.iterrows():
            if pd.notna(row.get('证券简称')):
                result.append({
                    '基金代码': row.get('基金代码'),
                    '证券简称': row.get('证券简称'),
                    '近一周收益': row.get('近一周收益'),
                    '年初以来收益': row.get('年初以来收益')
                })
    return result


def extract_new_funds(fund_data: Dict[str, pd.DataFrame]) -> Dict[str, List]:
    """提取新基金数据"""
    result = {'established': [], 'issued': [], 'declared': []}
    
    # 新成立基金
    df_est = fund_data.get('本周新成立的基金')
    if df_est is not None:
        for _, row in df_est.iterrows():
            if pd.notna(row.get('证券简称')):
                result['established'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '基金成立日': row.get('基金成立日'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '发行规模': row.get('发行规模(亿元)'),
                    '基金经理': row.get('基金经理')
                })
    
    # 新发行基金
    df_issue = fund_data.get('本周新发行的基金')
    if df_issue is not None:
        for _, row in df_issue.iterrows():
            if pd.notna(row.get('证券简称')):
                result['issued'].append({
                    '基金代码': row.get('证券代码'),
                    '证券简称': row.get('证券简称'),
                    '发行起始日': row.get('发行起始日'),
                    '投资类型': row.get('投资类型(二级分类)'),
                    '基金经理': row.get('基金经理')
                })
    
    # 新申报基金
    df_decl = fund_data.get('本周新申报的基金')
    if df_decl is not None:
        for _, row in df_decl.iterrows():
            if pd.notna(row.get('基金名称')):
                result['declared'].append({
                    '基金管理人': row.get('基金管理人'),
                    '基金名称': row.get('基金名称'),
                    '基金类型': row.get('基金类型1'),
                    '申请日期': row.get('申请材料接收日')
                })
    
    return result


def extract_etf_flow_data(etf_data: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """提取ETF资金流动数据"""
    result = {
        'sector_flow': [],
        'core_index_flow': [],
        'hot_theme_flow': [],
        'top_inflow': [],
        'top_outflow': []
    }
    
    # 板块资金流动
    df_sector = etf_data.get('板块结果展示')
    if df_sector is not None:
        for _, row in df_sector.iterrows():
            if pd.notna(row.get('二级分类')):
                result['sector_flow'].append({
                    '板块': row['二级分类'],
                    '平均涨跌幅': row.get('平均涨跌幅'),
                    '资金流入总额': row.get('资金流入总额（亿元）')
                })
    
    # 核心宽基ETF资金流动
    df_core = etf_data.get('核心宽基赛道')
    if df_core is not None:
        for _, row in df_core.iterrows():
            if pd.notna(row.get('跟踪指数')):
                result['core_index_flow'].append({
                    '跟踪指数': row.get('跟踪指数'),
                    '周收益率': row.get('周收益率'),
                    '净申购额': row.get('净申购额（亿元）.1')
                })
    
    # 热门行业主题
    df_theme = etf_data.get('热门行业主题')
    if df_theme is not None:
        for _, row in df_theme.iterrows():
            if pd.notna(row.get('指数简称')):
                result['hot_theme_flow'].append({
                    '指数简称': row.get('指数简称'),
                    '周收益率': row.get('周收益率'),
                    '净申赎额': row.get('净申赎额')
                })
    
    # 个基结果展示（净申购TOP）
    df_top = etf_data.get('个基结果展示')
    if df_top is not None:
        for _, row in df_top.iterrows():
            if pd.notna(row.get('基金名称')):
                inflow = row.get('资金流入规模（亿元）')
                if pd.notna(inflow) and isinstance(inflow, (int, float)):
                    result['top_inflow'].append({
                        '基金名称': row.get('基金名称'),
                        '资金流入规模': inflow
                    })
    
    # 按资金流入规模排序
    result['top_inflow'] = sorted(result['top_inflow'], key=lambda x: x['资金流入规模'], reverse=True)[:10]
    
    # 已上市ETF（净赎回TOP）
    df_etf = etf_data.get('已上市ETF')
    if df_etf is not None:
        outflow_list = []
        for _, row in df_etf.iterrows():
            # 尝试不同的列名
            net_flow = None
            for col in df_etf.columns:
                if '净申赎额' in str(col) or '净申购额' in str(col):
                    net_flow = row.get(col)
                    break
            
            if pd.notna(net_flow) and isinstance(net_flow, (int, float)) and net_flow < 0:
                outflow_list.append({
                    '基金名称': row.get('证券简称'),
                    '净申赎额': net_flow
                })
        
        # 按净赎回额排序（从小到大）
        result['top_outflow'] = sorted(outflow_list, key=lambda x: x['净申赎额'])[:10]
    
    return result


def format_percent(value) -> str:
    """格式化百分比"""
    if value is None or pd.isna(value):
        return 'N/A'
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def format_amount(value) -> str:
    """格式化金额"""
    if value is None or pd.isna(value):
        return 'N/A'
    if isinstance(value, (int, float)):
        return f"{abs(value):.2f}"
    return str(value)


def generate_report(fund_file: str, etf_file: str, output_file: str):
    """生成周报Word文档"""
    
    # 读取数据
    fund_data, etf_data = read_excel_data(fund_file, etf_file)
    
    # 提取各类数据
    date_range = get_date_range(fund_data)
    active_equity = extract_active_equity_stats(fund_data)
    industry_funds = extract_industry_fund_returns(fund_data)
    fixed_income = extract_fixed_income_stats(fund_data)
    index_funds = extract_index_fund_data(fund_data)
    fof_data = extract_fof_data(fund_data)
    qdii_data = extract_qdii_data(fund_data)
    reits_data = extract_reits_data(fund_data)
    new_funds = extract_new_funds(fund_data)
    etf_flow = extract_etf_flow_data(etf_data)
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading(f'基金市场周报 ({date_range[0]}-{date_range[1]})', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 主要市场指数周度表现回顾
    doc.add_heading('1. 主要市场指数周度表现回顾', 1)
    
    # 提取核心宽基ETF数据
    core_index = etf_flow['core_index_flow']
    if core_index:
        # 分析涨跌情况
        up_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] > 0]
        down_indices = [x for x in core_index if x['周收益率'] and x['周收益率'] < 0]
        
        # 生成描述
        if len(up_indices) > len(down_indices):
            market_desc = "A股主要指数涨多跌少"
        elif len(up_indices) < len(down_indices):
            market_desc = "A股主要指数跌多涨少"
        else:
            market_desc = "A股主要指数涨跌互现"
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），{market_desc}。")
        
        # 涨幅居前指数
        if up_indices:
            up_sorted = sorted(up_indices, key=lambda x: x['周收益率'], reverse=True)
            p.add_run(f"{up_sorted[0]['跟踪指数'].replace('指数', '')}涨幅居前，全周上涨{format_percent(up_sorted[0]['周收益率'])}；")
            if len(up_sorted) > 1:
                p.add_run(f"{up_sorted[1]['跟踪指数'].replace('指数', '')}上涨{format_percent(up_sorted[1]['周收益率'])}。")
        
        # 跌幅居前指数
        if down_indices:
            down_sorted = sorted(down_indices, key=lambda x: x['周收益率'])
            p.add_run(f"{down_sorted[0]['跟踪指数'].replace('指数', '')}跌幅较深，全周下跌{format_percent(abs(down_sorted[0]['周收益率']))}；")
            if len(down_sorted) > 1:
                p.add_run(f"{down_sorted[1]['跟踪指数'].replace('指数', '')}下跌{format_percent(abs(down_sorted[1]['周收益率']))}。")
    
    # 2. 主动权益基金周度表现复盘
    doc.add_heading('2. 主动权益基金周度表现复盘', 1)
    
    # 2.1 收益分布
    doc.add_heading('2.1 收益分布', 2)
    
    if active_equity and active_equity.get('weekly'):
        weekly = active_equity['weekly']
        
        # 提取中位数
        median_stock = weekly.get('普通股票型基金', {}).get('50%分位')
        median_hybrid = weekly.get('偏股混合型基金', {}).get('50%分位')
        median_flexible = weekly.get('灵活配置型基金', {}).get('50%分位')
        median_balanced = weekly.get('平衡混合型基金', {}).get('50%分位')
        
        # 提取首尾差异
        max_return = weekly.get('普通股票型基金', {}).get('最高值')
        min_return = weekly.get('普通股票型基金', {}).get('最低值')
        
        # 判断市场情况
        if median_stock and median_stock > 0.01:
            market_trend = "上涨"
        elif median_stock and median_stock < -0.01:
            market_trend = "下跌"
        else:
            market_trend = "表现分化"
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），普通股票/偏股混合/灵活配置/平衡混合基金周收益率中位数分别为{format_percent(median_stock)}/{format_percent(median_hybrid)}/{format_percent(median_flexible)}/{format_percent(median_balanced)}；")
        
        if max_return and min_return:
            diff = max_return - min_return
            if diff > 0.15:
                p.add_run(f"首尾基金收益差异巨大，")
            else:
                p.add_run(f"首尾基金收益差异较大，")
            p.add_run(f"头部绩优产品周收益超{format_percent(max_return * 0.8)}，尾部产品周跌幅近{format_percent(abs(min_return * 0.8))}。")
    
    # 2.2 行业主题基金表现
    doc.add_heading('2.2 行业主题基金表现', 2)
    
    if industry_funds and industry_funds.get('weekly'):
        weekly_industry = industry_funds['weekly']
        
        # 排序
        sorted_industry = sorted(weekly_industry, key=lambda x: x['收益'], reverse=True)
        
        # 领涨行业
        top_industry = sorted_industry[0] if sorted_industry else None
        second_industry = sorted_industry[1] if len(sorted_industry) > 1 else None
        third_industry = sorted_industry[2] if len(sorted_industry) > 2 else None
        
        # 表现落后行业
        bottom_industry = sorted_industry[-1] if sorted_industry else None
        bottom2_industry = sorted_industry[-2] if len(sorted_industry) > 1 else None
        
        p = doc.add_paragraph()
        p.add_run(f"分类型来看，最近一周（{date_range[0]}-{date_range[1]}）各类主动权益基金涨跌互现，")
        
        if top_industry:
            if top_industry['收益'] > 0.03:
                p.add_run(f"{top_industry['行业']}主题基金强势领涨，平均周涨幅达{format_percent(top_industry['收益'])}，")
            else:
                p.add_run(f"{top_industry['行业']}主题基金领涨，平均周涨幅为{format_percent(top_industry['收益'])}，")
        
        if second_industry and third_industry:
            p.add_run(f"另有{second_industry['行业']}和{third_industry['行业']}主题基金平均涨幅在{format_percent(min(second_industry['收益'], third_industry['收益']))}左右；")
        
        if bottom_industry:
            p.add_run(f"{bottom_industry['行业']}主题基金表现落后，平均跌幅为{format_percent(abs(bottom_industry['收益']))}。")
    
    # 3. 固定收益基金周度表现复盘
    doc.add_heading('3. 固定收益基金周度表现复盘', 1)
    
    if fixed_income and fixed_income.get('weekly'):
        weekly_fi = fixed_income['weekly']
        
        # 提取中位数
        median_short = weekly_fi.get('短期纯债型基金', {}).get('50%分位')
        median_long = weekly_fi.get('中长期纯债型基金', {}).get('50%分位')
        median_primary = weekly_fi.get('一级债基', {}).get('50%分位')
        median_secondary = weekly_fi.get('二级债基', {}).get('50%分位')
        median_hybrid_bond = weekly_fi.get('偏债混合型基金', {}).get('50%分位')
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），纯债基金净值稳步上涨，短期纯债和中长期纯债基金周收益率中位数分别为{format_percent(median_short)}和{format_percent(median_long)}；")
        
        # 判断含权债基表现
        if median_secondary and median_secondary > 0:
            p.add_run(f"各类含权债基收益收正，一级债基、二级债基和偏债混合型基金周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}和{format_percent(median_hybrid_bond)}。")
        else:
            p.add_run(f"含权债基表现分化，一级债基、二级债基周收益率中位数分别为{format_percent(median_primary)}、{format_percent(median_secondary)}。")
    
    # 4. 指数型基金周度表现复盘
    doc.add_heading('4. 指数型基金周度表现复盘', 1)
    
    # 4.1 被动指基
    doc.add_heading('4.1 被动指基', 2)
    
    # ETF资金流动
    if etf_flow and etf_flow.get('sector_flow'):
        sector_flow = etf_flow['sector_flow']
        
        # 分析资金流动
        inflow_sectors = [x for x in sector_flow if x['资金流入总额'] and x['资金流入总额'] > 0]
        outflow_sectors = [x for x in sector_flow if x['资金流入总额'] and x['资金流入总额'] < 0]
        
        # 核心宽基ETF
        core_outflow = sum([x['净申购额'] for x in etf_flow['core_index_flow'] if x['净申购额'] and x['净申购额'] < 0])
        
        p = doc.add_paragraph()
        
        if core_outflow < -500:
            p.add_run(f"ETF资金流动方面，跟踪沪深300/中证500/中证1000/创业板指/科创50等核心宽基指数的头部ETF标的集体遭遇大额赎回，核心宽基板块全周净流出额高达{format_amount(abs(core_outflow))}亿元；")
        else:
            p.add_run(f"ETF资金流动方面，核心宽基ETF标的资金流动分化，")
        
        if inflow_sectors:
            top_inflow = sorted(inflow_sectors, key=lambda x: x['资金流入总额'], reverse=True)
            p.add_run(f"与之相对，市场资金积极申购{top_inflow[0]['板块']}等主题标的，{top_inflow[0]['板块']}ETF板块全周净流入{format_amount(top_inflow[0]['资金流入总额'])}亿元。")
    
    # 4.2 增强指基
    doc.add_heading('4.2 增强指基', 2)
    
    if index_funds and index_funds.get('enhanced_top'):
        enhanced_top = index_funds['enhanced_top'][:5]
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），增强指数基金整体表现良好。")
        
        if enhanced_top:
            p.add_run(f"头部绩优产品包括：")
            for fund in enhanced_top[:3]:
                p.add_run(f"{fund['证券简称']}（{format_percent(fund['近一周收益'])}）、")
            p.add_run("等。")
    
    # 5. FOF基金周度表现复盘
    doc.add_heading('5. FOF基金周度表现复盘', 1)
    
    if fof_data and fof_data.get('weekly'):
        weekly_fof = fof_data['weekly']
        
        # 计算平均收益
        avg_return = sum([x['收益'] for x in weekly_fof if x['收益']]) / len(weekly_fof) if weekly_fof else 0
        
        p = doc.add_paragraph()
        
        if avg_return > 0:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值整体上涨。")
        elif avg_return < 0:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值集体下跌。")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）各类FOF基金净值表现分化。")
        
        # 高权益仓位FOF
        high_equity = [x for x in weekly_fof if '偏股' in x['类型'] or '积极' in x['类型']]
        low_equity = [x for x in weekly_fof if '偏债' in x['类型'] or '稳健' in x['类型']]
        
        if high_equity and low_equity:
            high_avg = sum([x['收益'] for x in high_equity if x['收益']]) / len(high_equity)
            low_avg = sum([x['收益'] for x in low_equity if x['收益']]) / len(low_equity)
            
            p.add_run(f"分类型来看，高权益仓位FOF基金弹性较大，平均收益{format_percent(high_avg)}；低权益仓位FOF基金弹性偏弱，平均收益{format_percent(low_avg)}。")
    
    # 6. 其他类型基金周度表现复盘
    doc.add_heading('6. 其他类型基金周度表现复盘', 1)
    
    # 6.1 QDII基金
    doc.add_heading('6.1 QDII基金', 2)
    
    if qdii_data:
        p = doc.add_paragraph()
        
        # 计算平均收益
        avg_return = sum([x['近一周收益'] for x in qdii_data if x['近一周收益']]) / len(qdii_data) if qdii_data else 0
        
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）主动QDII基金平均收益为{format_percent(avg_return)}，")
        
        # 头部产品
        top_qdii = sorted(qdii_data, key=lambda x: x['近一周收益'] if x['近一周收益'] else 0, reverse=True)[:5]
        
        if top_qdii:
            p.add_run(f"{top_qdii[0]['证券简称']}、{top_qdii[1]['证券简称']}、{top_qdii[2]['证券简称']}近一周涨幅居前。")
    
    # 6.2 REITs基金
    doc.add_heading('6.2 REITs基金', 2)
    
    if reits_data:
        p = doc.add_paragraph()
        
        # 计算平均收益
        avg_return = sum([x['近一周收益'] for x in reits_data if x['近一周收益']]) / len(reits_data) if reits_data else 0
        
        if avg_return > 0:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）REITs市场整体上涨，平均收益{format_percent(avg_return)}。")
        else:
            p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）REITs市场整体下跌，平均收益{format_percent(avg_return)}。")
        
        # 涨幅居前产品
        top_reits = sorted(reits_data, key=lambda x: x['近一周收益'] if x['近一周收益'] else 0, reverse=True)[:3]
        
        if top_reits and top_reits[0]['近一周收益'] > 0:
            p.add_run(f"{top_reits[0]['证券简称']}周涨幅{format_percent(top_reits[0]['近一周收益'])}，表现居前。")
    
    # 7. 基金成立与发行回顾
    doc.add_heading('7. 基金成立与发行回顾', 1)
    
    # 7.1 基金成立
    doc.add_heading('7.1 基金成立', 2)
    
    if new_funds and new_funds.get('established'):
        established = new_funds['established']
        
        # 统计数量和募资规模
        total_count = len(established)
        total_amount = sum([x['发行规模'] for x in established if x['发行规模']]) if established else 0
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}），全市场新成立基金{total_count}只，合计募资{total_amount:.2f}亿元。")
        
        # 按类型统计
        type_counts = {}
        for fund in established:
            fund_type = fund.get('投资类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        # 找出数量最多的类型
        max_type = max(type_counts.items(), key=lambda x: x[1]) if type_counts else ('其他', 0)
        
        p.add_run(f"{max_type[0]}新成立数量最多，为{max_type[1]}只。")
        
        # 募资规模最大的基金
        top_fund = max(established, key=lambda x: x['发行规模'] if x['发行规模'] else 0) if established else None
        
        if top_fund:
            p.add_run(f"{top_fund['证券简称']}募资规模最大，为{top_fund['发行规模']:.2f}亿元。")
    
    # 7.2 基金发行
    doc.add_heading('7.2 基金发行', 2)
    
    if new_funds and new_funds.get('issued'):
        issued = new_funds['issued']
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新发行基金{len(issued)}只。")
        
        # 按类型统计
        type_counts = {}
        for fund in issued:
            fund_type = fund.get('投资类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        if type_counts:
            p.add_run(f"主动权益/指数/固定收益基金新发数量分别为{type_counts.get('偏股混合型基金', 0)}只/{type_counts.get('被动指数型基金', 0)}只/{type_counts.get('债券型基金', 0)}只。")
    
    # 7.3 基金申报
    doc.add_heading('7.3 基金申报', 2)
    
    if new_funds and new_funds.get('declared'):
        declared = new_funds['declared']
        
        p = doc.add_paragraph()
        p.add_run(f"最近一周（{date_range[0]}-{date_range[1]}）全市场新申报基金共{len(declared)}只。")
        
        # 按类型统计
        type_counts = {}
        for fund in declared:
            fund_type = fund.get('基金类型', '其他')
            type_counts[fund_type] = type_counts.get(fund_type, 0) + 1
        
        if type_counts:
            type_str = '、'.join([f"{k}{v}只" for k, v in type_counts.items()])
            p.add_run(f"包括{type_str}。")
    
    # 8. 附注及风险提示
    doc.add_heading('8. 附注及风险提示', 1)
    
    p = doc.add_paragraph()
    p.add_run("本报告基于公开数据整理，仅供参考，不构成投资建议。基金投资有风险，投资者应根据自身风险承受能力谨慎决策。")
    
    # 保存文档
    doc.save(output_file)
    print(f"周报已生成: {output_file}")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python generate_weekly_report.py <fund_excel> <etf_excel> [output_file]")
        sys.exit(1)
    
    fund_file = sys.argv[1]
    etf_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else '基金周报.docx'
    
    generate_report(fund_file, etf_file, output_file)

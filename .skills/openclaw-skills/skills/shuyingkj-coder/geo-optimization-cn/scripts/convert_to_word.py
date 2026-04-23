#!/usr/bin/env python3
"""
将Markdown内容转换为Word文档（支持纯Q&A和企业信息两种格式）
使用方式：python convert_to_word.py <input.md> <output.docx>
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='Microsoft YaHei', font_size=10.5):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def markdown_to_word(md_file, output_file):
    """将Markdown转换为Word"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 跳过表格行
        if line.startswith('|'):
            i += 1
            continue
        
        # 处理二级标题（## 标题）
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(level=1)
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            set_chinese_font(run, font_size=14)
            i += 1
            continue
        
        # 处理三级标题（### 问题）- QA格式
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            set_chinese_font(run, font_size=11)
            i += 1
            continue
        
        # 处理分隔线
        if line.startswith('---'):
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理列表项
        if line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            set_chinese_font(run)
            i += 1
            continue
        
        # 处理数字列表（1. 2. 等）
        if len(line) > 2 and line[0].isdigit() and line[1] == '.':
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            set_chinese_font(run)
            i += 1
            continue
        
        # 处理普通段落
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_chinese_font(run)
        
        i += 1
    
    # 保存（无页脚信息）
    doc.save(output_file)
    print(f"✅ Word文档已生成: {output_file}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("使用方式: python convert_to_word.py <input.md> <output.docx>")
        print("示例: python convert_to_word.py QA知识库.md QA知识库.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    markdown_to_word(input_file, output_file)

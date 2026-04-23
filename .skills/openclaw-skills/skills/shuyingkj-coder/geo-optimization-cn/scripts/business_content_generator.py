#!/usr/bin/env python3
"""
企业营销内容智能生成器
功能：
1. 搜索词前缀/后缀生成（各5个）
2. 行业问答词生成（10-15个）
3. Word文档 → GEO优化QA知识库（纯Q&A，无分类、无关键词）
4. Word文档 → 企业关键信息提取（纯文本，无表格）

输出格式偏好（用户定制）：
- QA知识库：纯Q&A，无分类、无编号、无关键词、无版本信息
- 企业信息：纯文本，无表格、无标题页、无版本信息
"""

import sys
import os
import re
import subprocess
import argparse
from pathlib import Path

# ==================== 行业模板库 ====================

INDUSTRY_TEMPLATES = {
    "科技": {
        "互联网": {
            "prefixes": ["最好的", "靠谱的", "专业的", "排名靠前的", "口碑好的"],
            "suffixes": ["怎么样", "哪家好", "推荐", "评价", "价格"],
            "qa_template": [
                "{name}是做什么的？",
                "{name}的主要产品有哪些？",
                "{name}的优势是什么？",
                "为什么选择{name}？",
                "{name}的服务流程是怎样的？",
                "{name}的收费标准是什么？",
                "{name}和其他品牌有什么区别？",
                "{name}的数据安全如何保障？",
                "如何联系{name}？",
                "{name}的客户评价怎么样？",
                "{name}适合什么类型的企业？",
                "使用{name}产品需要注意什么？",
                "{name}有免费试用吗？",
                "{name}的技术支持如何？",
                "{name}的发展历程是怎样的？"
            ]
        },
        "软件": {
            "prefixes": ["好用的", "专业的", "高效的", "安全的", "稳定的"],
            "suffixes": ["怎么用", "功能介绍", "价格", "评测", "下载"],
            "qa_template": [
                "{name}是什么软件？",
                "{name}主要功能有哪些？",
                "{name}适合什么场景使用？",
                "{name}怎么收费？",
                "{name}和竞品相比有什么优势？",
                "{name}的数据安全吗？",
                "{name}支持哪些平台？",
                "{name}有免费版吗？",
                "{name}的安装配置复杂吗？",
                "{name}的更新频率如何？",
                "{name}提供培训服务吗？",
                "{name}的客户主要是哪些？",
                "如何获得{name}的技术支持？",
                "{name}可以定制开发吗？",
                "{name}的系统要求是什么？"
            ]
        }
    },
    "制造业": {
        "通用": {
            "prefixes": ["优质的", "专业的", "可靠的", "大型的", "知名的"],
            "suffixes": ["厂家", "供应商", "价格", "质量怎么样", "生产能力"],
            "qa_template": [
                "{name}主要生产什么产品？",
                "{name}的产能规模如何？",
                "{name}的质量控制体系是怎样的？",
                "{name}支持OEM/ODM吗？",
                "{name}的产品通过了哪些认证？",
                "{name}的交货周期是多久？",
                "{name}的最小起订量是多少？",
                "{name}的产品价格区间？",
                "{name}主要服务哪些行业？",
                "{name}有哪些合作客户？",
                "{name}的售后服务如何？",
                "{name}的生产基地在哪里？",
                "{name}有研发团队吗？",
                "如何考察{name}的工厂？",
                "{name}的付款方式是什么？"
            ]
        }
    },
    "教育培训": {
        "通用": {
            "prefixes": ["专业的", "靠谱的", "知名的", "口碑好的", "效果好的"],
            "suffixes": ["怎么样", "哪家好", "课程介绍", "收费标准", "学员评价"],
            "qa_template": [
                "{name}主要提供什么课程？",
                "{name}的师资力量如何？",
                "{name}的课程收费标准？",
                "{name}的学员就业情况怎么样？",
                "{name}的教学模式是什么？",
                "{name}有试听课程吗？",
                "{name}的课程时长是多久？",
                "{name}颁发什么证书？",
                "{name}和其他机构有什么区别？",
                "{name}有就业推荐服务吗？",
                "{name}的校区分布在哪里？",
                "{name}支持线上学习吗？",
                "{name}的退款政策是什么？",
                "{name}的课程可以分期吗？",
                "如何报名{name}的课程？"
            ]
        }
    },
    "医疗健康": {
        "通用": {
            "prefixes": ["专业的", "权威的", "正规的", "靠谱的", "技术先进的"],
            "suffixes": ["怎么样", "哪家好", "专家介绍", "收费标准", "预约挂号"],
            "qa_template": [
                "{name}主要提供什么医疗服务？",
                "{name}的专家团队怎么样？",
                "{name}有哪些特色科室？",
                "{name}是医保定点单位吗？",
                "{name}的挂号方式有哪些？",
                "{name}的收费标准透明吗？",
                "{name}的就医环境如何？",
                "{name}的医疗设备先进吗？",
                "{name}支持在线问诊吗？",
                "{name}的口碑评价怎么样？",
                "{name}的具体地址在哪里？",
                "{name}的营业时间是？",
                "{name}有体检服务吗？",
                "{name}可以用医保卡吗？",
                "如何预约{name}的专家号？"
            ]
        }
    },
    "金融服务": {
        "通用": {
            "prefixes": ["安全的", "正规的", "靠谱的", "利率低的", "审批快的"],
            "suffixes": ["怎么样", "可靠吗", "利率多少", "申请条件", "办理流程"],
            "qa_template": [
                "{name}主要提供什么金融服务？",
                "{name}的贷款利率是多少？",
                "{name}的贷款额度范围？",
                "{name}的审批速度快吗？",
                "{name}的申请条件是什么？",
                "{name}需要哪些申请材料？",
                "{name}的还款方式有哪些？",
                "{name}是正规持牌机构吗？",
                "{name}的风控措施如何？",
                "{name}可以提前还款吗？",
                "{name}有逾期罚息吗？",
                "{name}的客户评价怎么样？",
                "{name}的办理流程是什么？",
                "{name}支持线上办理吗？",
                "如何联系{name}的客服？"
            ]
        }
    },
    "电商零售": {
        "通用": {
            "prefixes": ["正品", "优质的", "性价比高的", "热销的", "口碑好的"],
            "suffixes": ["怎么样", "好用吗", "价格", "哪里买", "优惠券"],
            "qa_template": [
                "{name}主要卖什么产品？",
                "{name}的产品质量怎么样？",
                "{name}的价格有优势吗？",
                "{name}的物流配送快吗？",
                "{name}支持退换货吗？",
                "{name}有会员优惠吗？",
                "{name}的客服服务如何？",
                "{name}的商品是正品吗？",
                "{name}有哪些促销活动？",
                "{name}的支付方式有哪些？",
                "{name}的订单怎么查询？",
                "{name}的售后服务政策？",
                "{name}的积分怎么用？",
                "{name}有线下门店吗？",
                "如何联系{name}的客服？"
            ]
        }
    },
    "餐饮服务": {
        "通用": {
            "prefixes": ["好吃的", "口碑好的", "人气高的", "性价比高的", "环境好的"],
            "suffixes": ["怎么样", "好吃吗", "价格", "地址", "营业时间"],
            "qa_template": [
                "{name}是什么类型的餐厅？",
                "{name}的招牌菜有哪些？",
                "{name}的人均消费多少？",
                "{name}的环境怎么样？",
                "{name}需要预订吗？",
                "{name}支持外卖吗？",
                "{name}的营业时间是？",
                "{name}的具体地址在哪里？",
                "{name}有包间吗？",
                "{name}的食材新鲜吗？",
                "{name}支持团购吗？",
                "{name}可以自带酒水吗？",
                "{name}有停车位吗？",
                "{name}的排队时间长吗？",
                "如何预订{name}的座位？"
            ]
        }
    },
    "房地产": {
        "通用": {
            "prefixes": ["优质的", "高端的", "性价比高的", "地段好的", "配套完善的"],
            "suffixes": ["怎么样", "价格", "户型", "位置", "配套"],
            "qa_template": [
                "{name}的项目位置在哪里？",
                "{name}的房价是多少？",
                "{name}有哪些户型？",
                "{name}的配套设施如何？",
                "{name}的开发商是谁？",
                "{name}的物业服务怎么样？",
                "{name}的交房时间？",
                "{name}的交通便利吗？",
                "{name}周边有学校吗？",
                "{name}是精装交付吗？",
                "{name}的产权年限？",
                "{name}支持贷款吗？",
                "{name}有什么优惠政策？",
                "{name}的容积率是多少？",
                "如何预约{name}看房？"
            ]
        }
    },
    "咨询服务": {
        "通用": {
            "prefixes": ["专业的", "靠谱的", "经验丰富的", "知名的", "权威的"],
            "suffixes": ["怎么样", "哪家好", "服务内容", "收费标准", "案例"],
            "qa_template": [
                "{name}主要提供什么咨询服务？",
                "{name}的团队实力如何？",
                "{name}的服务流程是怎样的？",
                "{name}的收费标准？",
                "{name}有哪些成功案例？",
                "{name}的服务周期多久？",
                "{name}和同行相比有什么优势？",
                "{name}的咨询效果有保障吗？",
                "{name}主要服务哪些客户？",
                "{name}可以提供上门服务吗？",
                "{name}的咨询是保密的吗？",
                "{name}有后续跟踪服务吗？",
                "{name}支持线上咨询吗？",
                "{name}的付款方式是什么？",
                "如何预约{name}的咨询？"
            ]
        }
    },
    "物流运输": {
        "通用": {
            "prefixes": ["靠谱的", "专业的", "高效的", "价格合理的", "覆盖广的"],
            "suffixes": ["怎么样", "价格", "时效", "网点", "查询"],
            "qa_template": [
                "{name}主要提供什么物流服务？",
                "{name}的覆盖范围有哪些？",
                "{name}的运输时效如何？",
                "{name}的收费标准？",
                "{name}的货物安全保障？",
                "{name}支持货物追踪吗？",
                "{name}的客服响应速度？",
                "{name}的理赔流程是什么？",
                "{name}有大件运输服务吗？",
                "{name}支持上门取件吗？",
                "{name}的网点分布情况？",
                "{name}有冷链运输吗？",
                "{name}可以代收货款吗？",
                "{name}支持保价服务吗？",
                "如何联系{name}发货？"
            ]
        }
    }
}

# ==================== 辅助函数 ====================

def get_industry_template(industry):
    """根据行业获取模板"""
    industry = industry.lower() if industry else ""
    
    # 行业关键词映射
    if any(kw in industry for kw in ["科技", "互联网", "软件", "it", "tech"]):
        if any(kw in industry for kw in ["软件", "saas", "工具"]):
            return INDUSTRY_TEMPLATES["科技"]["软件"]
        return INDUSTRY_TEMPLATES["科技"]["互联网"]
    elif any(kw in industry for kw in ["制造", "工厂", "生产", "industry", "manufacturing"]):
        return INDUSTRY_TEMPLATES["制造业"]["通用"]
    elif any(kw in industry for kw in ["教育", "培训", "学校", "education", "training"]):
        return INDUSTRY_TEMPLATES["教育培训"]["通用"]
    elif any(kw in industry for kw in ["医疗", "健康", "医院", "medical", "health"]):
        return INDUSTRY_TEMPLATES["医疗健康"]["通用"]
    elif any(kw in industry for kw in ["金融", "银行", "保险", "投资", "finance", "banking"]):
        return INDUSTRY_TEMPLATES["金融服务"]["通用"]
    elif any(kw in industry for kw in ["电商", "零售", "shop", "e-commerce", "retail"]):
        return INDUSTRY_TEMPLATES["电商零售"]["通用"]
    elif any(kw in industry for kw in ["餐饮", "美食", "餐厅", "food", "restaurant", "catering"]):
        return INDUSTRY_TEMPLATES["餐饮服务"]["通用"]
    elif any(kw in industry for kw in ["房地产", "房产", "地产", "楼盘", "real estate", "property"]):
        return INDUSTRY_TEMPLATES["房地产"]["通用"]
    elif any(kw in industry for kw in ["咨询", "顾问", "服务", "consulting", "advisory"]):
        return INDUSTRY_TEMPLATES["咨询服务"]["通用"]
    elif any(kw in industry for kw in ["物流", "快递", "运输", "logistics", "shipping", "delivery"]):
        return INDUSTRY_TEMPLATES["物流运输"]["通用"]
    else:
        # 默认使用餐饮模板（广式腊肠属于餐饮/食品）
        return INDUSTRY_TEMPLATES["餐饮服务"]["通用"]

def read_word_document(file_path):
    """读取Word文档内容"""
    try:
        result = subprocess.run(
            ["uvx", "--with", "markitdown[docx]", "markitdown", file_path],
            capture_output=True,
            text=True,
            timeout=120
        )
        if result.returncode == 0:
            return result.stdout
        else:
            print(f"Error reading document: {result.stderr}")
            return None
    except Exception as e:
        print(f"Error: {e}")
        return None

def save_to_word(content, output_path):
    """保存内容到Word文档（纯文本格式）"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        
        def set_chinese_font(run, font_name='Microsoft YaHei', font_size=10.5):
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        doc = Document()
        
        # 设置默认中文字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 处理内容行
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 跳过表格行
            if line.startswith('|') and '|' in line[1:]:
                continue
            
            # 处理标题
            if line.startswith('## '):
                p = doc.add_heading(level=1)
                run = p.add_run(line[3:].strip())
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                set_chinese_font(run, font_size=14)
            elif line.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(line[4:].strip())
                run.bold = True
                run.font.size = Pt(11)
                set_chinese_font(run, font_size=11)
            elif line.startswith('---'):
                doc.add_paragraph()
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line[2:])
                set_chinese_font(run)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_chinese_font(run)
        
        doc.save(output_path)
        print(f"✅ Word文档已保存: {output_path}")
        return True
    except ImportError:
        print("Error: python-docx 未安装。运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"Error saving Word: {e}")
        return False

# ==================== 核心功能函数 ====================

def generate_search_terms(company_name, industry=""):
    """生成搜索词前缀和后缀"""
    template = get_industry_template(industry)
    
    prefixes = template["prefixes"][:5]
    suffixes = template["suffixes"][:5]
    
    output = f"""
{'='*60}
🔍 【搜索词分析报告】
{'='*60}
公司名称：{company_name}
行业类型：{industry if industry else '通用'}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 【搜索词前缀】（用户搜索时放在公司名前）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    for i, prefix in enumerate(prefixes, 1):
        output += f"{i}. {prefix}\n"
    
    output += f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 【搜索词后缀】（用户搜索后加的词）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    for i, suffix in enumerate(suffixes, 1):
        output += f"{i}. {suffix}\n"
    
    output += f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💡 【组合搜索词示例】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    combinations = []
    for prefix in prefixes[:3]:
        combinations.append(f"{prefix}{company_name}")
    for suffix in suffixes[:3]:
        combinations.append(f"{company_name}{suffix}")
    
    for combo in combinations:
        output += f"• {combo}\n"
    
    output += "="*60 + "\n"
    return output

def generate_qa_words(company_name, industry=""):
    """生成问答词"""
    template = get_industry_template(industry)
    qa_list = template["qa_template"]
    
    output = f"""
{'='*60}
💬 【行业问答词报告】
{'='*60}
公司名称：{company_name}
行业类型：{industry if industry else '通用'}
生成数量：{len(qa_list)}个

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 【问答词列表】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    for i, qa in enumerate(qa_list, 1):
        question = qa.format(name=company_name)
        output += f"\n{i}. {question}\n"
    
    output += f"""
{'='*60}
"""
    return output

def generate_kb_prompt(content, company_name=""):
    """生成QA知识库的AI提示词"""
    prompt = f"""请基于以下文档内容，生成一份GEO优化的QA知识库。

文档内容：
{content[:10000]}

公司名称：{company_name if company_name else '从文档中提取'}

生成要求：
1. 遵循GEO最佳实践：
   - Answer-First 结构（先给答案）
   - 每段答案控制在100-200字
   - 自然融入关键词

2. 知识库结构：
   - 生成15-20个QA对
   - 覆盖企业概况、产品、服务、合作等不同维度

3. 输出格式（纯Q&A，无分类、无关键词）：
   ### [问题]
   
   [答案，100-200字]
   
   ---

重要提示：
- 不要添加分类标题（如"## 分类一：XXX"）
- 不要添加关键词标注
- 不要添加知识库名称、文档版本、生成日期等信息
- 不要添加"生成工具"等说明
- 直接从第一个问题开始输出
- 每个QA之间用"---"分隔
"""
    return prompt

def generate_extract_prompt(content):
    """生成企业信息提取的AI提示词"""
    prompt = f"""请从以下文档内容中提取企业关键信息。

文档内容：
{content[:8000]}

输出格式（纯文本，无表格）：

## 一、公司名称

工商全称：XXX
常用简称/品牌名：XXX

---

## 二、公司及业务范围介绍

企业核心简介：（150-300字）

主营业务范围：
- XXX
- XXX

业务覆盖区域：
XXX

服务/合作行业：
XXX

---

## 三、产品或服务特点

核心产品线及特点：

产品1：（名称）
（特点描述）

产品2：（名称）
（特点描述）

产品核心卖点：
- XXX
- XXX

---

## 四、信任背书

研发与技术硬实力：
- XXX
- XXX

生产与供应链硬实力：
- XXX
- XXX

品控与标准化硬实力：
- XXX
- XXX

规模与市场硬实力：
- XXX
- XXX

---

## 五、客户案例

客户群体概述：

客户群体1：（类型）
（描述）

客户群体2：（类型）
（描述）

服务数据：
XXX

---

## 六、品牌故事

品牌起源与创立初心：
XXX

企业核心发展历程：
XXX

品牌核心价值观与经营理念：
- XXX
- XXX

企业愿景与未来发展规划：
XXX

---

## 七、用户痛点

目标客群核心画像：
XXX

行业通用核心痛点：
- XXX
- XXX

解决方案：
XXX

---

## 附录：联系信息

官方联系电话：
XXX

企业总部地址：
XXX

公司创立日期：
XXX

注册资金：
XXX

重要提示：
1. 纯文本格式输出，不使用表格
2. 不要添加文档版本、提取日期、制作工具等信息
3. 不要添加标题页（如"XXX公司 - 关键信息提取报告"）
4. 直接按上述格式输出内容
"""
    return prompt

# ==================== 主函数 ====================

def main():
    parser = argparse.ArgumentParser(description='企业营销内容智能生成器')
    subparsers = parser.add_subparsers(dest='command', help='可用命令')
    
    # search 命令
    search_parser = subparsers.add_parser('search', help='生成搜索词')
    search_parser.add_argument('company', help='公司名称')
    search_parser.add_argument('industry', nargs='?', default='', help='行业（可选）')
    
    # qa 命令
    qa_parser = subparsers.add_parser('qa', help='生成问答词')
    qa_parser.add_argument('company', help='公司名称')
    qa_parser.add_argument('industry', nargs='?', default='', help='行业（可选）')
    
    # kb 命令
    kb_parser = subparsers.add_parser('kb', help='生成QA知识库')
    kb_parser.add_argument('input', help='输入Word文件路径')
    kb_parser.add_argument('--output', '-o', help='输出Word文件路径（可选）')
    
    # extract 命令
    extract_parser = subparsers.add_parser('extract', help='提取企业信息')
    extract_parser.add_argument('input', help='输入Word文件路径')
    extract_parser.add_argument('--output', '-o', help='输出Word文件路径（可选）')
    
    args = parser.parse_args()
    
    if args.command == 'search':
        print(generate_search_terms(args.company, args.industry))
    
    elif args.command == 'qa':
        print(generate_qa_words(args.company, args.industry))
    
    elif args.command == 'kb':
        print(f"📄 正在读取文档: {args.input}")
        content = read_word_document(args.input)
        if content:
            print(f"✅ 文档读取成功，共 {len(content)} 字符")
            print("\n" + "="*60)
            print("🤖 请将以下内容复制给Claude，生成GEO优化的QA知识库：")
            print("="*60 + "\n")
            print(generate_kb_prompt(content))
            
            if args.output:
                print(f"\n💡 提示：获得Claude生成的内容后，可以保存为Markdown，")
                print(f"   然后使用脚本转换为Word: python convert_to_word.py")
        else:
            print("❌ 文档读取失败")
    
    elif args.command == 'extract':
        print(f"📄 正在读取文档: {args.input}")
        content = read_word_document(args.input)
        if content:
            print(f"✅ 文档读取成功，共 {len(content)} 字符")
            print("\n" + "="*60)
            print("🤖 请将以下内容复制给Claude，提取企业关键信息：")
            print("="*60 + "\n")
            print(generate_extract_prompt(content))
            
            if args.output:
                print(f"\n💡 提示：获得Claude生成的内容后，可以保存为Markdown，")
                print(f"   然后使用脚本转换为Word: python convert_to_word.py")
        else:
            print("❌ 文档读取失败")
    
    else:
        parser.print_help()

if __name__ == "__main__":
    main()

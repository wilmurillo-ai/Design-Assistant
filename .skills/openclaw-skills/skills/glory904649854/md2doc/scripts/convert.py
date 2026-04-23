#!/usr/bin/env python3
"""Markdown 转 Word/PDF 转换器"""

import argparse
import os
import re
import sys
from pathlib import Path
from urllib.parse import urlparse

# 导入样式模板和AI分析器
sys.path.insert(0, str(Path(__file__).parent.parent / 'templates'))
sys.path.insert(0, str(Path(__file__).parent))
from styles import get_style, list_styles, StyleConfig
from ai_analyzer import analyze_document, ChartGenerator

try:
    import requests
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"缺少依赖: {e}")
    print("请运行: pip install python-docx requests Pillow")
    sys.exit(1)


def set_chinese_font(run, font_name='宋体', font_size=11, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold


def download_image(url, temp_dir):
    """下载网络图片"""
    try:
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            parsed = urlparse(url)
            filename = os.path.basename(parsed.path) or 'image.jpg'
            if '?' in filename:
                filename = filename.split('?')[0]
            temp_path = os.path.join(temp_dir, filename)
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            return temp_path
    except Exception as e:
        print(f"下载图片失败 {url}: {e}")
    return None


def process_image_path(src, md_dir, temp_dir):
    """处理图片路径"""
    if not src:
        return None
    
    if src.startswith(('http://', 'https://')):
        return download_image(src, temp_dir)
    
    possible_paths = [
        os.path.join(md_dir, src),
        os.path.abspath(src),
        src,
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return os.path.abspath(path)
    
    print(f"警告: 找不到图片 {src}")
    return None


def add_cover_page(doc, cfg, title="", subtitle="", author="", date="", cover_image_path=""):
    """添加封面页（支持背景色和图片）"""
    
    # 如果有封面图片，尝试添加
    if cover_image_path and os.path.exists(cover_image_path):
        try:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(cover_image_path, width=Inches(4))
            doc.add_paragraph()
        except:
            pass  # 图片添加失败就跳过
    
    # 空行占位
    for _ in range(5 if cover_image_path else 6):
        doc.add_paragraph()
    
    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title or "文档标题")
    set_chinese_font(title_run, cfg.title_font, cfg.cover_title_size, bold=True)
    title_run.font.color.rgb = RGBColor(*cfg.title_color)
    
    doc.add_paragraph()
    
    # 副标题
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        set_chinese_font(sub_run, cfg.body_font, cfg.cover_subtitle_size)
        sub_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # 空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 作者
    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"作者: {author}")
        set_chinese_font(author_run, cfg.body_font, 12)
    
    # 日期
    if date:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(date)
        set_chinese_font(date_run, cfg.body_font, 11)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 分页
    doc.add_page_break()


def add_toc(doc, cfg, headings):
    """添加目录页（手动生成，无需更新域）"""
    # 目录标题
    toc_title = doc.add_heading(cfg.toc_title, level=1)
    for run in toc_title.runs:
        set_chinese_font(run, cfg.title_font, cfg.title1_size, bold=True)
        run.font.color.rgb = RGBColor(*cfg.title_color)
    
    # 手动生成目录项
    for level, text in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt((level - 1) * 20)  # 缩进
        
        # 添加标题文字
        run = p.add_run(text)
        if level == 1:
            set_chinese_font(run, cfg.title_font, cfg.body_size, bold=True)
        else:
            set_chinese_font(run, cfg.body_font, cfg.body_size - 1)
        run.font.color.rgb = RGBColor(*cfg.title_color if level == 1 else cfg.body_color)
    
    if not headings:
        p = doc.add_paragraph()
        run = p.add_run("(文档暂无标题)")
        set_chinese_font(run, cfg.body_font, 10)
        run.font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_page_break()


def add_header_footer(section, cfg, doc_title=""):
    """添加页眉页脚"""
    # 页眉
    if cfg.has_header:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = cfg.header_text or doc_title
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            set_chinese_font(run, cfg.body_font, 9)
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 页脚
    if cfg.has_footer or cfg.show_page_number:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if cfg.show_page_number:
            # 添加页码域
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as oqn
            
            # 前缀文字
            if cfg.footer_text and '{page}' not in cfg.footer_text:
                prefix_run = footer_para.add_run(cfg.footer_text + " ")
                set_chinese_font(prefix_run, cfg.body_font, 9)
                prefix_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 页码域
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(oqn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(oqn('w:fldCharType'), 'separate')
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(oqn('w:fldCharType'), 'end')
            
            page_run = footer_para.add_run()
            set_chinese_font(page_run, cfg.body_font, 9)
            page_run.font.color.rgb = RGBColor(128, 128, 128)
            page_run._r.append(fldChar1)
            page_run._r.append(instrText)
            page_run._r.append(fldChar2)
            page_run._r.append(fldChar3)
        else:
            footer_run = footer_para.add_run(cfg.footer_text)
            set_chinese_font(footer_run, cfg.body_font, 9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)


def markdown_to_docx(md_file, output_file, style_name='default', 
                     cover_title="", cover_subtitle="", cover_author="", cover_date="",
                     include_toc=True):
    """将 Markdown 转换为 Word 文档"""
    
    # 获取样式配置
    cfg = get_style(style_name)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    md_dir = os.path.dirname(os.path.abspath(md_file))
    temp_dir = os.path.join(os.path.dirname(output_file), '.temp_images')
    os.makedirs(temp_dir, exist_ok=True)
    
    doc = Document()
    
    # 设置页面边距和页眉页脚
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(cfg.page_margin_top)
        section.bottom_margin = Cm(cfg.page_margin_bottom)
        section.left_margin = Cm(cfg.page_margin_left)
        section.right_margin = Cm(cfg.page_margin_right)
        # 添加页眉页脚
        add_header_footer(section, cfg, cover_title or "文档")
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = cfg.body_font
    style._element.rPr.rFonts.set(qn('w:eastAsia'), cfg.body_font)
    style.font.size = Pt(cfg.body_size)
    style.paragraph_format.line_spacing = cfg.line_spacing
    
    # 添加封面（如果启用）
    if cfg.has_cover:
        # 从 MD 内容提取标题（第一个 # 开头的行）
        doc_title = cover_title
        if not doc_title:
            for line in md_content.split('\n'):
                if line.startswith('# '):
                    doc_title = line[2:].strip()
                    break
        add_cover_page(doc, cfg, title=doc_title, subtitle=cover_subtitle, 
                      author=cover_author, date=cover_date)
    
    # 收集所有标题用于目录
    headings = []
    for line in md_content.split('\n'):
        if line.startswith('# '):
            headings.append((1, line[2:].strip()))
        elif line.startswith('## '):
            headings.append((2, line[3:].strip()))
        elif line.startswith('### '):
            headings.append((3, line[4:].strip()))
    
    # 添加目录（如果启用）
    if cfg.has_toc and include_toc:
        add_toc(doc, cfg, headings)
    
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 代码块处理
        if stripped.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                set_chinese_font(run, 'Consolas', 9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        if not stripped:
            i += 1
            continue
        
        # 标题 - 使用样式配置
        if stripped.startswith('# '):
            heading = doc.add_heading(stripped[2:], level=1)
            for run in heading.runs:
                set_chinese_font(run, cfg.title_font, cfg.title1_size, bold=True)
                run.font.color.rgb = RGBColor(*cfg.title_color)
            i += 1
            continue
        
        if stripped.startswith('## '):
            heading = doc.add_heading(stripped[3:], level=2)
            for run in heading.runs:
                set_chinese_font(run, cfg.title_font, cfg.title2_size, bold=True)
                run.font.color.rgb = RGBColor(*cfg.title_color)
            i += 1
            continue
        
        if stripped.startswith('### '):
            heading = doc.add_heading(stripped[4:], level=3)
            for run in heading.runs:
                set_chinese_font(run, cfg.title_font, cfg.title3_size, bold=True)
                run.font.color.rgb = RGBColor(*cfg.title_color)
            i += 1
            continue
        
        # 图片
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt_text, img_src = img_match.groups()
            img_path = process_image_path(img_src, md_dir, temp_dir)
            if img_path:
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    if alt_text:
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.add_run(alt_text)
                        set_chinese_font(caption_run, '宋体', 9)
                        caption_run.font.color.rgb = RGBColor(100, 100, 100)
                except Exception as e:
                    print(f"插入图片失败: {e}")
            i += 1
            continue
        
        # 表格
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            rows = []
            for row_line in table_lines:
                if '---' in row_line or row_line.replace('|', '').replace('-', '').replace(':', '').strip() == '':
                    continue
                cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
            
            if rows:
                num_cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if row_idx == 0:
                                        set_chinese_font(run, cfg.title_font, cfg.body_size, bold=True)
                                        # 表头背景色
                                        from docx.oxml import parse_xml
                                        from docx.oxml.ns import nsdecls
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                                            nsdecls('w'), *cfg.table_header_bg))
                                        cell._tc.get_or_add_tcPr().append(shading_elm)
                                    else:
                                        set_chinese_font(run, cfg.body_font, cfg.body_size-1)
            continue
        
        # 普通段落（支持行内格式）
        paragraph = doc.add_paragraph()
        
        # 处理粗体、斜体、代码
        pattern = r'(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`)'
        last_end = 0
        
        for match in re.finditer(pattern, stripped):
            if match.start() > last_end:
                normal_text = stripped[last_end:match.start()]
                run = paragraph.add_run(normal_text)
                set_chinese_font(run, '宋体', 11)
            
            if match.group(2):  # 粗体
                run = paragraph.add_run(match.group(2))
                set_chinese_font(run, '微软雅黑', 11, bold=True)
            elif match.group(3):  # 斜体
                run = paragraph.add_run(match.group(3))
                set_chinese_font(run, '宋体', 11)
                run.italic = True
            elif match.group(4):  # 代码
                run = paragraph.add_run(match.group(4))
                set_chinese_font(run, 'Consolas', 10)
                run.font.color.rgb = RGBColor(200, 50, 50)
            
            last_end = match.end()
        
        if last_end < len(stripped):
            run = paragraph.add_run(stripped[last_end:])
            set_chinese_font(run, '宋体', 11)
        
        i += 1
    
    # 保存文档
    doc.save(output_file)
    print(f"Word 文档已保存: {output_file}")
    
    # 清理临时文件
    import shutil
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    
    return output_file


def markdown_to_html(md_file, output_file, style_name='default'):
    """将 Markdown 转换为 HTML"""
    import markdown
    
    cfg = get_style(style_name)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换 Markdown 为 HTML
    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'toc']
    )
    
    # 构建完整 HTML 页面
    html_template = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文档</title>
    <style>
        body {{
            font-family: "{cfg.body_font}", "Microsoft YaHei", sans-serif;
            font-size: {cfg.body_size}pt;
            line-height: {cfg.line_spacing};
            color: rgb{cfg.body_color};
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            background: #fafafa;
        }}
        h1, h2, h3 {{
            font-family: "{cfg.title_font}", "Microsoft YaHei", sans-serif;
            color: rgb{cfg.title_color};
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        h1 {{ font-size: {cfg.title1_size}pt; border-bottom: 2px solid rgb{cfg.accent_color}; padding-bottom: 10px; }}
        h2 {{ font-size: {cfg.title2_size}pt; border-bottom: 1px solid #ddd; padding-bottom: 8px; }}
        h3 {{ font-size: {cfg.title3_size}pt; }}
        code {{
            font-family: "{cfg.code_font}", Consolas, monospace;
            font-size: {cfg.code_size}pt;
            background: rgb{cfg.code_bg};
            padding: 2px 6px;
            border-radius: 3px;
            color: rgb{cfg.code_color};
        }}
        pre {{
            background: rgb{cfg.code_bg};
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            background: none;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{
            background: rgb{cfg.table_header_bg};
            font-weight: bold;
        }}
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1em auto;
        }}
        blockquote {{
            border-left: 4px solid rgb{cfg.accent_color};
            margin: 1em 0;
            padding-left: 1em;
            color: #666;
        }}
        a {{
            color: rgb{cfg.link_color};
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 2em 0;
        }}
        ul, ol {{
            padding-left: 2em;
        }}
        .toc {{
            background: #f5f5f5;
            padding: 20px;
            border-radius: 5px;
            margin: 2em 0;
        }}
        .toc ul {{
            list-style: none;
            padding-left: 0;
        }}
        .toc li {{
            margin: 5px 0;
        }}
        .toc a {{
            color: #333;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>'''
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(html_template)
    
    print(f"HTML 已保存: {output_file}")
    return output_file


def docx_to_pdf(docx_file, pdf_file):
    """将 Word 转换为 PDF（支持多种方式）"""
    docx_file = os.path.abspath(docx_file)
    pdf_file = os.path.abspath(pdf_file)
    
    # 方案1: 使用 Microsoft Word COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_file)
        doc.SaveAs(pdf_file, FileFormat=17)  # 17 = PDF
        doc.Close()
        word.Quit()
        print(f"PDF 已保存: {pdf_file}")
        return True
    except Exception as e:
        print(f"Word COM 失败: {e}")
    
    # 方案2: 使用 LibreOffice
    try:
        import subprocess
        # 尝试找到 soffice
        soffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",  # 如果已在 PATH 中
        ]
        
        soffice = None
        for path in soffice_paths:
            if os.path.exists(path) or path == "soffice":
                soffice = path
                break
        
        if soffice:
            output_dir = os.path.dirname(pdf_file)
            result = subprocess.run([
                soffice, "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, docx_file
            ], capture_output=True, timeout=60)
            
            if result.returncode == 0:
                # LibreOffice 生成的文件名可能不同
                generated_pdf = os.path.join(output_dir, os.path.basename(docx_file).replace('.docx', '.pdf'))
                if os.path.exists(generated_pdf):
                    if generated_pdf != pdf_file:
                        os.rename(generated_pdf, pdf_file)
                    print(f"PDF 已保存: {pdf_file}")
                    return True
    except Exception as e:
        print(f"LibreOffice 失败: {e}")
    
    print("PDF 转换失败。请安装 Microsoft Word 或 LibreOffice。")
    return False


def main():
    parser = argparse.ArgumentParser(description='Markdown 转 Word/PDF/HTML (AI增强版)')
    parser.add_argument('input', help='输入的 Markdown 文件')
    parser.add_argument('--output', '-o', help='输出的 Word 文件路径')
    parser.add_argument('--pdf', '-p', help='输出的 PDF 文件路径')
    parser.add_argument('--html', help='输出的 HTML 文件路径')
    parser.add_argument('--all', '-a', action='store_true', help='同时输出 Word + PDF + HTML')
    parser.add_argument('--style', '-s', default='auto', help='样式模板 (auto/default/business/tech/minimal/product/academic)')
    parser.add_argument('--list-styles', action='store_true', help='列出所有可用样式')
    # AI功能
    parser.add_argument('--ai-analyze', action='store_true', help='AI分析文档类型并推荐样式')
    parser.add_argument('--ai-charts', action='store_true', help='AI自动生成数据图表')
    parser.add_argument('--ai-images', action='store_true', help='AI建议插图位置（需手动生成图片）')
    # 封面参数
    parser.add_argument('--cover-title', help='封面标题（默认从MD提取）')
    parser.add_argument('--cover-subtitle', help='封面副标题')
    parser.add_argument('--cover-author', help='作者')
    parser.add_argument('--cover-date', help='日期')
    # 目录参数
    parser.add_argument('--no-toc', action='store_true', help='不生成目录')
    parser.add_argument('--toc', action='store_true', default=True, help='生成目录（默认）')
    
    args = parser.parse_args()
    
    # 列出样式
    if args.list_styles:
        list_styles()
        sys.exit(0)
    
    if not os.path.exists(args.input):
        print(f"错误: 找不到文件 {args.input}")
        sys.exit(1)
    
    # 读取文档内容
    with open(args.input, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # AI分析
    ai_result = None
    if args.ai_analyze or args.style == 'auto':
        print("\n[AI分析] 正在分析文档类型...")
        ai_result = analyze_document(md_content)
        print(f"[AI分析] 检测到文档类型: {ai_result['document_type']} (置信度: {ai_result['confidence']:.2%})")
        print(f"[AI分析] 推荐样式: {ai_result['suggested_style']}")
        print(f"[AI分析] 主题关键词: {', '.join(ai_result['topics'][:3])}")
        
        # 使用AI推荐的样式
        if args.style == 'auto':
            args.style = ai_result['suggested_style']
    
    # 如果样式还是auto，改为default
    if args.style == 'auto':
        args.style = 'default'
    
    # AI生成图表
    chart_paths = []
    if args.ai_charts and ai_result and ai_result['tables']:
        print("\n[AI图表] 正在生成数据图表...")
        chart_gen = ChartGenerator()
        for idx, table in enumerate(ai_result['tables'][:3]):  # 最多生成3个图表
            chart_path = chart_gen.generate_chart(table)
            if chart_path:
                chart_paths.append(chart_path)
                print(f"[AI图表] 已生成图表 {idx+1}: {table.get('title', '数据表')}")
    
    # AI建议插图
    if args.ai_images and ai_result:
        print("\n[AI插图] 插图建议:")
        for suggestion in ai_result['image_suggestions']:
            print(f"  - [{suggestion['section']}] {suggestion['description']}")
            print(f"    提示词: {suggestion['prompt']}")
    
    # 确定输出路径
    input_path = Path(args.input)
    docx_file = args.output or str(input_path.with_suffix('.docx'))
    
    # 判断是否生成目录：如果明确指定 --no-toc，或者只生成 PDF（不含HTML）时默认不生成目录
    # 只有当生成 HTML 或者明确指定 --toc 时才生成目录
    generate_toc = not args.no_toc
    if args.pdf and not args.html and not args.all and not args.no_toc:
        # 只生成 PDF 时，默认不生成目录（PDF目录无跳转功能，观感奇怪）
        generate_toc = False
        print("[提示] PDF 模式默认不生成目录（使用 --toc 可强制生成）")
    
    # 转换为 Word（带样式和封面）
    print(f"\n[转换] 使用样式: {args.style}")
    markdown_to_docx(
        args.input, docx_file, 
        style_name=args.style,
        cover_title=args.cover_title or "",
        cover_subtitle=args.cover_subtitle or "",
        cover_author=args.cover_author or "",
        cover_date=args.cover_date or "",
        include_toc=generate_toc
    )
    
    # 转换为 PDF（如果需要）
    pdf_file = args.pdf
    if args.all and not pdf_file:
        pdf_file = str(input_path.with_suffix('.pdf'))
    
    if pdf_file:
        docx_to_pdf(docx_file, pdf_file)
    
    # 转换为 HTML（如果需要）
    html_file = args.html
    if args.all and not html_file:
        html_file = str(input_path.with_suffix('.html'))
    
    if html_file:
        markdown_to_html(args.input, html_file, style_name=args.style)
    
    # 清理临时图表文件
    for chart_path in chart_paths:
        try:
            os.remove(chart_path)
        except:
            pass
    
    print("\n[完成] 转换成功!")


if __name__ == '__main__':
    main()
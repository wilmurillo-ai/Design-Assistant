#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF带格式精确转换Word
完整提取：字体、字号、对齐、段落格式、首行缩进、行间距、图片
"""

import os
import sys
import argparse
from pathlib import Path

# 字体映射表
FONT_MAP = {
    'song': '宋体', 'simsun': '宋体', '.song': '宋体',
    'hei': '黑体', 'simhei': '黑体', '.hei': '黑体',
    'kai': '楷体', 'simkai': '楷体', '.kai': '楷体',
    'fang': '仿宋_GB2312', 'simfang': '仿宋_GB2312',
    'fangso': '仿宋_GB2312', 'fs': '仿宋_GB2312',
    'times': 'Times New Roman', 'tim': 'Times New Roman',
    'arial': 'Arial', 'helvetica': 'Arial',
    'courier': 'Courier New',
    'microsoft': '微软雅黑', 'yahei': '微软雅黑',
    'mincho': '宋体', 'gothic': '黑体',
}


def get_font_name(font_name):
    """映射字体名称到Word可用字体"""
    if not font_name:
        return '宋体'
    font_lower = font_name.lower().replace('-', '').replace('_', '')
    for key, value in FONT_MAP.items():
        if key in font_lower or key in font_name.lower():
            return value
    return font_name if font_name else '宋体'


def get_alignment(align):
    """获取对齐方式"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        0: WD_ALIGN_PARAGRAPH.LEFT,
        1: WD_ALIGN_PARAGRAPH.CENTER,
        2: WD_ALIGN_PARAGRAPH.RIGHT,
        3: WD_ALIGN_PARAGRAPH.JUSTIFY,
        4: WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    return mapping.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def extract_page_content(page):
    """提取页面所有内容"""
    content = {
        'text_blocks': [],
        'tables': [],
        'images': []
    }
    
    # 获取文本块
    blocks = page.get_text("dict")["blocks"]
    
    for block in blocks:
        if block.get("type") == 0:  # 文本块
            spans_data = []
            for line in block.get("lines", []):
                align = line.get("align", 0)
                
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text.strip():
                        continue
                    
                    bbox = span.get("bbox", [0,0,0,0])
                    font = span.get("font", "宋体")
                    size = span.get("size", 12)
                    flags = span.get("flags", 0)
                    color = span.get("color", 0)
                    
                    spans_data.append({
                        'text': text,
                        'font': get_font_name(font),
                        'size': round(size, 1),
                        'bold': bool(flags & 16),
                        'italic': bool(flags & 2),
                        'underline': bool(flags & 4),
                        'color': color,
                        'align': align,
                        'bbox': bbox
                    })
            
            if spans_data:
                content['text_blocks'].append(spans_data)
    
    # 提取表格
    try:
        tables = page.find_tables()
        for table in tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text if cell_text else "")
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                content['tables'].append(table_data)
    except:
        pass
    
    # 提取图片
    img_list = page.get_images()
    for img_index, img in enumerate(img_list):
        try:
            xref = img[0]
            base_image = page.parent.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            img_path = f"_temp_img_{page.number}_{img_index}.{image_ext}"
            with open(img_path, "wb") as f:
                f.write(image_bytes)
            content['images'].append({
                'path': img_path,
                'width': base_image["width"],
                'height': base_image["height"]
            })
        except:
            pass
    
    return content


def create_paragraph_with_format(word_doc, spans, is_title=False):
    """创建带有格式的段落"""
    if not spans:
        return
    
    # 按Y坐标分组
    paragraph_groups = []
    current_y = None
    current_group = []
    
    for span in spans:
        y = span['bbox'][1]
        if current_y is None:
            current_y = y
            current_group = [span]
        elif abs(y - current_y) < 3:  # 同一行
            current_group.append(span)
        else:
            if current_group:
                paragraph_groups.append(current_group)
            current_group = [span]
            current_y = y
    
    if current_group:
        paragraph_groups.append(current_group)
    
    # 创建段落
    for group in paragraph_groups:
        if not group:
            continue
        
        first_span = group[0]
        
        para = word_doc.add_paragraph()
        para.alignment = get_alignment(first_span['align'])
        
        # 首行缩进（正文，非标题）
        if not is_title and first_span['text']:
            para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 行间距
        para.paragraph_format.line_spacing = 1.5
        
        # 合并文本
        full_text = "".join([span['text'] for span in group if span['text']])
        
        if full_text.strip():
            from docx.shared import Pt
            from docx.oxml.ns import qn
            
            run = para.add_run(full_text)
            
            # 字体
            run.font.name = first_span['font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), first_span['font'])
            
            # 字号
            run.font.size = Pt(first_span['size'])
            
            # 粗体、斜体、下划线
            if first_span['bold']:
                run.font.bold = True
            if first_span['italic']:
                run.font.italic = True
            if first_span['underline']:
                run.font.underline = True
            
            # 颜色
            if first_span['color'] != 0:
                from docx.shared import RGBColor
                rgb = first_span['color']
                r = (rgb >> 16) & 0xFF
                g = (rgb >> 8) & 0xFF
                b = rgb & 0xFF
                run.font.color.rgb = RGBColor(r, g, b)


def create_table(word_doc, table_data):
    """创建表格"""
    if not table_data:
        return
    
    from docx.shared import Pt
    
    table = word_doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 1)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < len(table.columns):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                # 设置表格字体
                for para in cell.paragraphs:
                    if para.runs:
                        para.runs[0].font.name = '宋体'
                        para.runs[0].font.size = Pt(10.5)


def convert_pdf_to_word(pdf_path, output_path, start=0, end=None):
    """转换PDF到Word"""
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found: {pdf_path}")
        return False
    
    try:
        import fitz
        from docx import Document
        from docx.shared import Cm, Inches
    except ImportError as e:
        print(f"Error: Missing dependency - {e}")
        print("Please install: pip install pymupdf python-docx")
        return False
    
    print(f"Converting: {pdf_path}")
    
    try:
        doc = fitz.open(pdf_path)
        
        if end is None:
            end = len(doc)
        
        word_doc = Document()
        
        # 设置默认样式
        style = word_doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        # 处理每一页
        for page_num in range(start, min(end, len(doc))):
            page = doc[page_num]
            
            if page_num > start:
                word_doc.add_page_break()
            
            content = extract_page_content(page)
            
            # 处理文本
            for i, block in enumerate(content['text_blocks']):
                is_title = (page_num == start and i == 0)
                create_paragraph_with_format(word_doc, block, is_title=is_title)
            
            # 处理表格
            for table_data in content['tables']:
                create_table(word_doc, table_data)
            
            # 处理图片
            for img_info in content['images']:
                try:
                    word_doc.add_picture(img_info['path'], width=Inches(4))
                    os.remove(img_info['path'])
                except:
                    pass
        
        doc.close()
        
        word_doc.save(output_path)
        print(f"OK! Saved to: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return False


def batch_convert(pdf_dir, output_dir):
    """批量转换"""
    if not os.path.exists(pdf_dir):
        print(f"Error: PDF directory not found: {pdf_dir}")
        return
    
    os.makedirs(output_dir, exist_ok=True)
    
    pdf_files = list(Path(pdf_dir).glob("*.pdf"))
    
    if not pdf_files:
        print(f"Warning: No PDF files found")
        return
    
    print(f"Found {len(pdf_files)} PDF files")
    
    success_count = 0
    for pdf_file in pdf_files:
        output_file = os.path.join(output_dir, pdf_file.stem + ".docx")
        if convert_pdf_to_word(str(pdf_file), output_file):
            success_count += 1
    
    print(f"Completed: {success_count}/{len(pdf_files)} successful")


def main():
    parser = argparse.ArgumentParser(
        description="PDF带格式精确转换成Word - 保留字体、字号、对齐、段落格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python convert.py report.pdf --output report.docx
  python convert.py ./pdfs/ --batch --output ./words/
  python convert.py document.pdf --pages 0-5 --output document.docx
        """
    )
    
    parser.add_argument('input', help='Input PDF file or directory')
    parser.add_argument('--output', '-o', help='Output Word file or directory')
    parser.add_argument('--pages', '-p', help='Page range, e.g. 0-5')
    parser.add_argument('--start', '-s', type=int, default=0, help='Start page (0-based)')
    parser.add_argument('--end', '-e', type=int, default=None, help='End page (exclusive)')
    parser.add_argument('--batch', '-b', action='store_true', help='Batch conversion')
    
    args = parser.parse_args()
    
    # 处理页码
    if args.pages:
        try:
            if '-' in args.pages:
                start, end = args.pages.split('-')
                args.start = int(start)
                args.end = int(end) + 1
        except:
            print("Error: Invalid page range format")
            sys.exit(1)
    
    # 输出路径
    output_path = args.output
    if not output_path:
        if args.batch:
            output_path = "output"
        else:
            base_name = os.path.splitext(os.path.basename(args.input))[0]
            output_path = base_name + ".docx"
    
    # 执行
    if args.batch:
        batch_convert(args.input, output_path)
    else:
        convert_pdf_to_word(args.input, output_path, args.start, args.end)


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
apply_format.py  —  Apply format profile to target .docx with intelligent heading detection

支持功能：
1. 从格式描述文件读取格式配置
2. 智能识别段落类型：标题、一级标题、二级标题、正文
3. 应用不同的字体、字号、对齐方式、缩进
4. 支持自定义格式配置（如公文格式）

Usage:
    python apply_format.py <target.docx> <format_profile.json> [--output result.docx]
    python apply_format.py <target.docx> --official --output result.docx  # 使用内置公文格式
"""

import sys
import json
import re
import argparse
from pathlib import Path


def pt_to_emu(pt):
    return int(pt * 12700) if pt is not None else None


def _get_or_add(parent, tag):
    from docx.oxml.ns import qn
    el = parent.find(qn(tag))
    if el is None:
        from docx.oxml import OxmlElement
        el = OxmlElement(tag)
        parent.append(el)
    return el


def set_font(run, font_name, font_size_pt, bold=False):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.font.name = font_name
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    run.font.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = _get_or_add(rPr, 'w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def apply_para_format(para, cfg):
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pf = para.paragraph_format

    # Alignment
    align_str = cfg.get('alignment', 'JUSTIFY')
    if 'CENTER' in align_str:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif 'RIGHT' in align_str:
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif 'LEFT' in align_str:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Space before / after
    sb = cfg.get('space_before_pt', 0)
    sa = cfg.get('space_after_pt', 0)
    pf.space_before = Pt(sb) if sb else Pt(0)
    pf.space_after = Pt(sa) if sa else Pt(0)

    # Line spacing
    ls_rule = cfg.get('line_spacing_rule', 'MULTIPLE')
    ls_val = cfg.get('line_spacing_value', 1.5)
    if 'MULTIPLE' in ls_rule or not ls_rule:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = float(ls_val) if ls_val else 1.5
    elif 'EXACTLY' in ls_rule:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls_val)
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = float(ls_val) if ls_val else 1.5

    # First-line indent (character-based)
    first_line_chars = cfg.get('first_line_chars', 0)
    font_size_pt = cfg.get('font_size_pt', 16)
    
    pPr = para._element.get_or_add_pPr()
    ind = _get_or_add(pPr, 'w:ind')

    if first_line_chars and first_line_chars > 0:
        twips = int(first_line_chars * font_size_pt * 20)
        ind.set(qn('w:firstLineChars'), str(int(first_line_chars * 100)))
        ind.set(qn('w:firstLine'), str(twips))
    else:
        ind.attrib.pop(qn('w:firstLineChars'), None)
        ind.attrib.pop(qn('w:firstLine'), None)
        pf.first_line_indent = None


def detect_para_type(text, index, total_paras):
    """
    智能识别段落类型
    返回: 'title', 'doc_no', 'level1', 'level2', 'body'
    """
    text = text.strip()
    if not text:
        return None
    
    # 第一个非空段落作为标题
    if index == 0:
        return 'title'
    
    # 第二个段落可能是文号（如果像公文格式）
    if index == 1:
        if re.match(r'^[\u4e00-\u9fa5a-zA-Z]+〔\d+〕\d+号', text):
            return 'doc_no'
    
    # 一级标题：形如 "一、"  "二、"  "三、"
    if re.match(r'^[\u4e00-\u9fa5一二三四五六七八九十百]+、', text):
        return 'level1'
    
    # 二级标题：形如 "（一）"  "（二）"
    if re.match(r'^（[\u4e00-\u9fa5一二三四五六七八九十百]+）', text):
        return 'level2'
    
    return 'body'


def analyze_profile_advanced(profile):
    """
    分析格式描述文件，提取标题层级格式
    返回包含 title, level1, level2, body 的配置字典
    完整提取：字体、字号、加粗、对齐、首行缩进、段前间距、段后间距、行间距
    """
    paragraphs = profile.get('paragraphs', [])
    
    # 分类收集不同层级段落的格式
    title_samples = []
    level1_samples = []
    level2_samples = []
    body_samples = []
    
    for p in paragraphs:
        text = p.get('text_preview', '')
        pf = p.get('paragraph_format', {})
        runs = p.get('runs', [])
        xml_indent = p.get('xml_indent', {})
        
        fl_chars = xml_indent.get('firstLineChars', 0) / 100 if xml_indent.get('firstLineChars') else 0
        
        sample = {
            'alignment': pf.get('alignment', 'JUSTIFY'),
            'space_before_pt': pf.get('space_before_pt') if pf.get('space_before_pt') is not None else 0,
            'space_after_pt': pf.get('space_after_pt') if pf.get('space_after_pt') is not None else 0,
            'line_spacing_rule': pf.get('line_spacing_rule', 'MULTIPLE'),
            'line_spacing_value': pf.get('line_spacing_value') or 1.5,
            'first_line_chars': fl_chars,
            'font_name': runs[0].get('east_asia_font') or runs[0].get('font_name') if runs else '宋体',
            'font_size_pt': runs[0].get('font_size_pt') if runs else 12,
            'bold': any(r.get('bold') for r in runs),
        }
        
        # 根据文本特征分类
        if re.match(r'^[\u4e00-\u9fa5一二三四五六七八九十百]+、', text):
            level1_samples.append(sample)
        elif re.match(r'^（[\u4e00-\u9fa5一二三四五六七八九十百]+）', text):
            level2_samples.append(sample)
        elif text and len(text) < 30 and not fl_chars:
            # 短文本无缩进，可能是标题
            if any(r.get('bold') for r in runs):
                title_samples.append(sample)
            else:
                title_samples.append(sample)
        else:
            body_samples.append(sample)
    
    def most_common(lst, key, default):
        if not lst:
            return default
        vals = [s[key] for s in lst if s.get(key) is not None]
        if not vals:
            return default
        from collections import Counter
        return Counter(vals).most_common(1)[0][0]
    
    def get_font(samples, default_font='宋体', default_size=12):
        font = most_common(samples, 'font_name', default_font)
        size = most_common(samples, 'font_size_pt', default_size)
        return font, size
    
    # 分析各级别格式
    title_font, title_size = get_font(title_samples or body_samples, '宋体', 22)
    level1_font, level1_size = get_font(level1_samples or body_samples, '黑体', 16)
    level2_font, level2_size = get_font(level2_samples or body_samples, '楷体', 16)
    body_font, body_size = get_font(body_samples, '宋体', 16)
    
    body_first_chars = most_common(body_samples, 'first_line_chars', 2)
    body_align = most_common(body_samples, 'alignment', 'JUSTIFY')
    
    # 【关键修改】从模板动态提取段前段后间距和行间距
    body_space_before = most_common(body_samples, 'space_before_pt', 0)
    body_space_after = most_common(body_samples, 'space_after_pt', 0)
    body_line_spacing = most_common(body_samples, 'line_spacing_value', 1.5)
    body_line_rule = most_common(body_samples, 'line_spacing_rule', 'MULTIPLE')
    
    level1_space_before = most_common(level1_samples, 'space_before_pt', 0) if level1_samples else 0
    level1_space_after = most_common(level1_samples, 'space_after_pt', 0) if level1_samples else 0
    level1_line_spacing = most_common(level1_samples, 'line_spacing_value', 1.5) if level1_samples else 1.5
    
    level2_space_before = most_common(level2_samples, 'space_before_pt', 0) if level2_samples else 0
    level2_space_after = most_common(level2_samples, 'space_after_pt', 0) if level2_samples else 0
    level2_line_spacing = most_common(level2_samples, 'line_spacing_value', 1.5) if level2_samples else 1.5
    
    title_space_before = most_common(title_samples, 'space_before_pt', 0) if title_samples else 0
    title_space_after = most_common(title_samples, 'space_after_pt', 0) if title_samples else 0
    title_line_spacing = most_common(title_samples, 'line_spacing_value', 1.5) if title_samples else 1.5
    
    return {
        'title': {
            'font_name': title_font,
            'font_size_pt': title_size,
            'bold': True,
            'alignment': 'CENTER',
            'first_line_chars': 0,
            'space_before_pt': title_space_before,
            'space_after_pt': title_space_after,
            'line_spacing_rule': body_line_rule,
            'line_spacing_value': title_line_spacing,
        },
        'doc_no': {
            'font_name': body_font,
            'font_size_pt': body_size,
            'bold': False,
            'alignment': 'CENTER',
            'first_line_chars': 0,
            'space_before_pt': 0,
            'space_after_pt': 0,
            'line_spacing_rule': 'MULTIPLE',
            'line_spacing_value': 1.5,
        },
        'level1': {
            'font_name': level1_font,
            'font_size_pt': level1_size,
            'bold': True,
            'alignment': 'LEFT',
            'first_line_chars': 0,
            'space_before_pt': level1_space_before,
            'space_after_pt': level1_space_after,
            'line_spacing_rule': body_line_rule,
            'line_spacing_value': level1_line_spacing,
        },
        'level2': {
            'font_name': level2_font,
            'font_size_pt': level2_size,
            'bold': True,
            'alignment': 'LEFT',
            'first_line_chars': 0,
            'space_before_pt': level2_space_before,
            'space_after_pt': level2_space_after,
            'line_spacing_rule': body_line_rule,
            'line_spacing_value': level2_line_spacing,
        },
        'body': {
            'font_name': body_font,
            'font_size_pt': body_size,
            'bold': False,
            'alignment': body_align,
            'first_line_chars': body_first_chars,
            'space_before_pt': body_space_before,
            'space_after_pt': body_space_after,
            'line_spacing_rule': body_line_rule,
            'line_spacing_value': body_line_spacing,
        },
        'page': profile.get('page', {}),
    }


def get_official_config():
    """
    获取内置公文格式配置（GB/T 9704-2012）
    """
    return {
        'title': {
            'font_name': '方正小标宋体',
            'font_size_pt': 22,
            'bold': True,
            'alignment': 'CENTER',
            'first_line_chars': 0,
            'space_before_pt': 0,
            'space_after_pt': 0,
            'line_spacing_value': 1.5,
        },
        'doc_no': {
            'font_name': '仿宋_GB2312',
            'font_size_pt': 16,
            'bold': False,
            'alignment': 'CENTER',
            'first_line_chars': 0,
            'space_before_pt': 0,
            'space_after_pt': 0,
            'line_spacing_value': 1.5,
        },
        'level1': {
            'font_name': '黑体',
            'font_size_pt': 16,
            'bold': True,
            'alignment': 'LEFT',
            'first_line_chars': 0,
            'space_before_pt': 0,
            'space_after_pt': 0,
            'line_spacing_value': 1.5,
        },
        'level2': {
            'font_name': '楷体_GB2312',
            'font_size_pt': 16,
            'bold': True,
            'alignment': 'LEFT',
            'first_line_chars': 0,
            'space_before_pt': 0,
            'space_after_pt': 0,
            'line_spacing_value': 1.5,
        },
        'body': {
            'font_name': '仿宋_GB2312',
            'font_size_pt': 16,
            'bold': False,
            'alignment': 'JUSTIFY',
            'first_line_chars': 2,
            'space_before_pt': 0,
            'space_after_pt': 0,
            'line_spacing_value': 1.5,
        },
        'page': {
            'page_width_cm': 21.0,
            'page_height_cm': 29.7,
            'top_margin_cm': 3.7,
            'bottom_margin_cm': 3.5,
            'left_margin_cm': 2.8,
            'right_margin_cm': 2.6,
        },
    }


def main():
    parser = argparse.ArgumentParser(description='Apply format to target .docx with intelligent heading detection')
    parser.add_argument('target', help='Path to target .docx file')
    parser.add_argument('profile', nargs='?', help='Path to format_profile.json (optional if --official used)')
    parser.add_argument('--output', default=None, help='Output .docx path')
    parser.add_argument('--official', action='store_true', help='Use built-in official document format (GB/T 9704-2012)')
    args = parser.parse_args()

    try:
        from docx import Document
        from docx.shared import Cm
    except ImportError:
        print('ERROR: python-docx not installed. Run: pip install python-docx')
        sys.exit(1)

    target_path = Path(args.target)
    if not target_path.exists():
        print(f'ERROR: Target file not found: {target_path}')
        sys.exit(1)

    output_path = args.output or target_path.with_name(target_path.stem + '_formatted.docx')

    # 获取格式配置
    if args.official:
        config = get_official_config()
        print('使用内置公文格式 (GB/T 9704-2012)')
    elif args.profile:
        profile_path = Path(args.profile)
        if not profile_path.exists():
            print(f'ERROR: Profile file not found: {profile_path}')
            sys.exit(1)
        with open(str(profile_path), 'r', encoding='utf-8') as f:
            profile = json.load(f)
        config = analyze_profile_advanced(profile)
        print('从格式描述文件分析配置')
    else:
        print('ERROR: 请提供格式描述文件或使用 --official 参数')
        sys.exit(1)

    print(f'\n格式配置 (从模板动态提取):')
    for key in ['title', 'level1', 'level2', 'body']:
        fmt = config[key]
        align_name = "居中" if "CENTER" in fmt["alignment"] else "左对齐" if "LEFT" in fmt["alignment"] else "两端对齐"
        print(f'  {key}: {fmt["font_name"]} {fmt["font_size_pt"]}pt, {"加粗" if fmt.get("bold") else "常规"}, {align_name}')
        print(f'       首行缩进:{fmt.get("first_line_chars", 0)}字符, 段前:{fmt.get("space_before_pt", 0)}pt, 段后:{fmt.get("space_after_pt", 0)}pt, 行距:{fmt.get("line_spacing_value", 1.5)}')

    doc = Document(str(target_path))

    # 应用页面边距
    page_cfg = config.get('page', {})
    if page_cfg:
        for section in doc.sections:
            if page_cfg.get('left_margin_cm'):
                section.left_margin = Cm(page_cfg['left_margin_cm'])
            if page_cfg.get('right_margin_cm'):
                section.right_margin = Cm(page_cfg['right_margin_cm'])
            if page_cfg.get('top_margin_cm'):
                section.top_margin = Cm(page_cfg['top_margin_cm'])
            if page_cfg.get('bottom_margin_cm'):
                section.bottom_margin = Cm(page_cfg['bottom_margin_cm'])

    # 统计
    counts = {'title': 0, 'doc_no': 0, 'level1': 0, 'level2': 0, 'body': 0}
    
    # 应用格式到所有段落
    paras = [p for p in doc.paragraphs if p.text.strip()]
    for i, para in enumerate(paras):
        para_type = detect_para_type(para.text, i, len(paras))
        if para_type is None:
            continue
            
        fmt = config.get(para_type, config['body'])
        
        # 应用段落格式
        apply_para_format(para, fmt)
        
        # 应用字体格式
        for run in para.runs:
            set_font(run, fmt['font_name'], fmt['font_size_pt'], fmt.get('bold', False))
        
        counts[para_type] = counts.get(para_type, 0) + 1
        print(f'  段落 {i+1}: [{para_type}] {para.text[:20]}...')

    # 应用到表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    apply_para_format(para, config['body'])
                    for run in para.runs:
                        set_font(run, config['body']['font_name'], config['body']['font_size_pt'], False)

    doc.save(str(output_path))
    print(f'\n完成! 共处理 {sum(counts.values())} 个段落')
    print(f'  标题: {counts.get("title", 0)}, 文号: {counts.get("doc_no", 0)}')
    print(f'  一级标题: {counts.get("level1", 0)}, 二级标题: {counts.get("level2", 0)}')
    print(f'  正文: {counts.get("body", 0)}')
    print(f'\n输出: {output_path}')


if __name__ == '__main__':
    main()

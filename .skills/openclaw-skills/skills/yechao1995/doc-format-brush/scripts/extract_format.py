#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_format.py  —  Extract paragraph/character format from a template .docx file

Usage:
    python extract_format.py <template.docx> [--output format_profile.json]

Output: A JSON file (format_profile.json by default) describing:
  - document_defaults: page size, margins
  - paragraph_formats: keyed by "style_name" or "para_index", each entry has:
        alignment, space_before_pt, space_after_pt,
        line_spacing_rule, line_spacing_value,
        first_line_indent (in EMU), left_indent (in EMU), right_indent (in EMU),
        keep_together, keep_with_next, page_break_before
  - run_formats: keyed same as above, each has:
        font_name, font_size_pt, bold, italic, underline, color_rgb,
        east_asia_font
  - styles: all named styles with their paragraph + run format snapshots
"""

import sys
import json
import argparse
from pathlib import Path


def emu_to_pt(emu):
    """Convert EMU (English Metric Units) to points."""
    if emu is None:
        return None
    return round(emu / 12700, 4)


def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 4)


def extract_para_format(para_fmt):
    """Extract paragraph format settings as a dict."""
    result = {}

    # Alignment
    if para_fmt.alignment is not None:
        result['alignment'] = str(para_fmt.alignment)

    # Spacing
    result['space_before_pt'] = emu_to_pt(para_fmt.space_before)
    result['space_after_pt'] = emu_to_pt(para_fmt.space_after)

    # Line spacing
    if para_fmt.line_spacing_rule is not None:
        result['line_spacing_rule'] = str(para_fmt.line_spacing_rule)
    if para_fmt.line_spacing is not None:
        val = para_fmt.line_spacing
        # If it's a float (multiple), store directly; if large int it's EMU
        if isinstance(val, float) or (isinstance(val, int) and val < 100):
            result['line_spacing_value'] = val
        else:
            result['line_spacing_value_pt'] = emu_to_pt(val)
            result['line_spacing_value'] = val

    # Indentation
    result['first_line_indent_pt'] = emu_to_pt(para_fmt.first_line_indent)
    result['first_line_indent_cm'] = emu_to_cm(para_fmt.first_line_indent)
    result['left_indent_pt'] = emu_to_pt(para_fmt.left_indent)
    result['right_indent_pt'] = emu_to_pt(para_fmt.right_indent)

    # Pagination
    result['keep_together'] = para_fmt.keep_together
    result['keep_with_next'] = para_fmt.keep_with_next
    result['page_break_before'] = para_fmt.page_break_before

    # Window/orphan control
    result['widow_control'] = para_fmt.widow_control

    return result


def extract_run_format(run):
    """Extract character/run format settings as a dict."""
    from docx.oxml.ns import qn

    result = {}
    font = run.font

    result['bold'] = font.bold
    result['italic'] = font.italic
    result['underline'] = font.underline
    result['font_size_pt'] = emu_to_pt(font.size) if font.size else None
    result['font_name'] = font.name
    result['color_rgb'] = str(font.color.rgb) if (font.color and font.color.type is not None) else None

    # East Asia font (Chinese font)
    rPr = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is not None:
        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rFonts is not None:
            result['east_asia_font'] = rFonts.get(qn('w:eastAsia'))
            result['ascii_font'] = rFonts.get(qn('w:ascii'))
            result['hansi_font'] = rFonts.get(qn('w:hAnsi'))

    return result


def extract_para_xml_indent(para):
    """
    Read firstLineChars directly from XML to detect character-based indent.
    Returns dict with keys: firstLineChars (int, e.g. 200 = 2 chars), firstLine_twips
    """
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return {}
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return {}
    result = {}
    flc = ind.get(qn('w:firstLineChars'))
    fl = ind.get(qn('w:firstLine'))
    if flc is not None:
        result['firstLineChars'] = int(flc)  # divide by 100 to get char count
    if fl is not None:
        result['firstLine_twips'] = int(fl)
    return result


def extract_snap_to_grid(para):
    """Check if snapToGrid is set on paragraph."""
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    snap = pPr.find(qn('w:snapToGrid'))
    if snap is None:
        return None
    return snap.get(qn('w:val'), '1')


def extract_style_format(style):
    """Extract format from a named style."""
    entry = {
        'style_name': style.name,
        'style_type': str(style.type),
        'base_style': style.base_style.name if style.base_style else None,
    }
    entry['paragraph_format'] = extract_para_format(style.paragraph_format)
    if hasattr(style, 'font'):
        f = style.font
        entry['font'] = {
            'name': f.name,
            'bold': f.bold,
            'italic': f.italic,
            'size_pt': emu_to_pt(f.size) if f.size else None,
        }
    return entry


def main():
    parser = argparse.ArgumentParser(description='Extract format profile from a template .docx')
    parser.add_argument('template', help='Path to template .docx file')
    parser.add_argument('--output', default=None,
                        help='Output JSON file path (default: <template>_format_profile.json)')
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError:
        print('ERROR: python-docx not installed. Run: pip install python-docx')
        sys.exit(1)

    template_path = Path(args.template)
    if not template_path.exists():
        print(f'ERROR: Template file not found: {template_path}')
        sys.exit(1)

    output_path = args.output or template_path.with_name(template_path.stem + '_format_profile.json')

    doc = Document(str(template_path))

    profile = {
        'source_file': str(template_path),
        'paragraphs': [],
        'styles': [],
        'page': {},
    }

    # Page / section info
    section = doc.sections[0]
    profile['page'] = {
        'page_width_cm': emu_to_cm(section.page_width),
        'page_height_cm': emu_to_cm(section.page_height),
        'left_margin_cm': emu_to_cm(section.left_margin),
        'right_margin_cm': emu_to_cm(section.right_margin),
        'top_margin_cm': emu_to_cm(section.top_margin),
        'bottom_margin_cm': emu_to_cm(section.bottom_margin),
    }

    # Named styles
    for style in doc.styles:
        try:
            profile['styles'].append(extract_style_format(style))
        except Exception:
            pass

    # Per-paragraph analysis (first 60 paragraphs for pattern detection)
    for i, para in enumerate(doc.paragraphs[:60]):
        if not para.text.strip():
            continue
        entry = {
            'index': i,
            'text_preview': para.text[:60],
            'style_name': para.style.name,
            'paragraph_format': extract_para_format(para.paragraph_format),
            'xml_indent': extract_para_xml_indent(para),
            'snap_to_grid': extract_snap_to_grid(para),
            'runs': [],
        }
        for run in para.runs:
            if run.text.strip():
                entry['runs'].append(extract_run_format(run))
                break  # Only first meaningful run per paragraph
        profile['paragraphs'].append(entry)

    with open(str(output_path), 'w', encoding='utf-8') as f:
        json.dump(profile, f, ensure_ascii=False, indent=2, default=str)

    print(f'Format profile saved to: {output_path}')
    print(f'  - {len(profile["paragraphs"])} paragraphs analyzed')
    print(f'  - {len(profile["styles"])} styles extracted')
    print(f'  Page: {profile["page"]}')


def extract_format_json(template_path):
    """
    Extract format profile and return as dict (instead of saving to file).
    
    Args:
        template_path: Path to template .docx file
        
    Returns:
        dict: Format profile containing paragraphs, styles, and page info
    """
    from docx import Document
    
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f'Template file not found: {template_path}')
    
    doc = Document(str(template_path))
    
    profile = {
        'source_file': str(template_path),
        'paragraphs': [],
        'styles': [],
        'page': {},
    }
    
    # Page / section info
    section = doc.sections[0]
    profile['page'] = {
        'page_width_cm': emu_to_cm(section.page_width),
        'page_height_cm': emu_to_cm(section.page_height),
        'left_margin_cm': emu_to_cm(section.left_margin),
        'right_margin_cm': emu_to_cm(section.right_margin),
        'top_margin_cm': emu_to_cm(section.top_margin),
        'bottom_margin_cm': emu_to_cm(section.bottom_margin),
    }
    
    # Named styles
    for style in doc.styles:
        try:
            profile['styles'].append(extract_style_format(style))
        except Exception:
            pass
    
    # Per-paragraph analysis
    for i, para in enumerate(doc.paragraphs[:60]):
        if not para.text.strip():
            continue
        entry = {
            'index': i,
            'text_preview': para.text[:60],
            'style_name': para.style.name,
            'paragraph_format': extract_para_format(para.paragraph_format),
            'xml_indent': extract_para_xml_indent(para),
            'snap_to_grid': extract_snap_to_grid(para),
            'runs': [],
        }
        for run in para.runs:
            if run.text.strip():
                entry['runs'].append(extract_run_format(run))
                break
        profile['paragraphs'].append(entry)
    
    return profile


if __name__ == '__main__':
    main()

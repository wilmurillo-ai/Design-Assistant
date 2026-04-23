#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
format_bridge.py — 多格式文档处理桥梁

支持读取和写入：
- Word文档 (.docx)
- Markdown (.md)
- 纯文本 (.txt)

主要功能：
1. 统一内容提取接口
2. 智能识别文档结构（标题层级）
3. 格式转换应用
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional


def read_document(file_path: str) -> Tuple[List[str], Dict]:
    """
    读取任意格式的文档，返回段落列表和元数据
    
    Returns:
        (paragraphs, metadata): 段落列表，文档元数据
    """
    ext = Path(file_path).suffix.lower()
    
    if ext == '.docx':
        return _read_docx(file_path)
    elif ext == '.md':
        return _read_markdown(file_path)
    elif ext == '.txt':
        return _read_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _read_docx(file_path: str) -> Tuple[List[str], Dict]:
    """读取 Word 文档"""
    from docx import Document
    
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # 提取页面信息
    metadata = {'format': 'docx'}
    if doc.sections:
        section = doc.sections[0]
        metadata['page_width_cm'] = section.page_width.cm
        metadata['page_height_cm'] = section.page_height.cm
        metadata['left_margin_cm'] = section.left_margin.cm
        metadata['right_margin_cm'] = section.right_margin.cm
        metadata['top_margin_cm'] = section.top_margin.cm
        metadata['bottom_margin_cm'] = section.bottom_margin.cm
    
    return paragraphs, metadata


def _read_markdown(file_path: str) -> Tuple[List[str], Dict]:
    """读取 Markdown 文件，识别标题层级"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    paragraphs = []
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # 处理 Markdown 标题 (# ## ###)
        md_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if md_match:
            level = len(md_match.group(1))
            text = md_match.group(2).strip()
            # 转换为带层级标记的文本
            if level == 1:
                paragraphs.append(f"[H1]{text}")
            elif level == 2:
                paragraphs.append(f"[H2]{text}")
            elif level == 3:
                paragraphs.append(f"[H3]{text}")
            else:
                paragraphs.append(f"[H{level}]{text}")
        else:
            # 处理列表项
            list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', stripped)
            if list_match:
                text = list_match.group(2).strip()
                paragraphs.append(f"[LIST]{text}")
            else:
                paragraphs.append(stripped)
    
    metadata = {
        'format': 'md',
        'source_file': os.path.basename(file_path)
    }
    
    return paragraphs, metadata


def _read_text(file_path: str) -> Tuple[List[str], Dict]:
    """读取纯文本文件，智能识别标题"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    paragraphs = []
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # 智能识别标题模式
        # 1. 纯数字 + 顿号 + 标题: "一、基本情况"
        if re.match(r'^[\u4e00-\u9fa5一二三四五六七八九十百千]+、\s*.+$', stripped):
            paragraphs.append(f"[H1]{stripped}")
        # 2. 括号数字/中文 + 标题: "（一）主要情况"
        elif re.match(r'^（[\u4e00-\u9fa5\d]+）\s*.+$', stripped):
            paragraphs.append(f"[H2]{stripped}")
        # 3. 阿拉伯数字 + 点 + 标题: "1. 主要情况"
        elif re.match(r'^\d+\.\s+.+$', stripped):
            paragraphs.append(f"[H2]{stripped}")
        # 4. 全角冒号或短横线结尾（可能是标题）: "第一章："
        elif re.match(r'^第[\u4e00-\u9fa5\d]+[：:].+$', stripped):
            paragraphs.append(f"[H1]{stripped}")
        else:
            paragraphs.append(stripped)
    
    metadata = {
        'format': 'txt',
        'source_file': os.path.basename(file_path)
    }
    
    return paragraphs, metadata


def detect_paragraph_type(text: str, index: int, total: int) -> str:
    """
    智能识别段落类型
    返回: 'title', 'doc_no', 'level1', 'level2', 'level3', 'body', 'list'
    """
    text = text.strip()
    if not text:
        return None
    
    # 移除标记前缀
    clean_text = re.sub(r'^\[H\d+\]|\[LIST\]', '', text).strip()
    
    # 第一个非空段落作为标题
    if index == 0:
        return 'title'
    
    # 第二个段落可能是文号
    if index == 1 and re.match(r'^[\u4e00-\u9fa5a-zA-Z]+〔\d+〕\d+号', clean_text):
        return 'doc_no'
    
    # 一级标题
    if re.match(r'^\[H1\]', text) or re.match(r'^[\u4e00-\u9fa5一二三四五六七八九十百]+、', clean_text):
        return 'level1'
    
    # 二级标题
    if re.match(r'^\[H2\]', text) or re.match(r'^（[\u4e00-\u9fa5一二三四五六七八九十百]+）', clean_text):
        return 'level2'
    
    # 三级标题
    if re.match(r'^\[H3\]', text) or re.match(r'^\d+\.\s+', clean_text):
        return 'level3'
    
    # 列表项
    if re.match(r'^\[LIST\]', text):
        return 'list'
    
    return 'body'


def write_document(paragraphs: List[str], output_path: str, format_config: Dict) -> None:
    """
    将段落写入指定格式的文档
    
    Args:
        paragraphs: 段落列表（可包含 [H1], [H2] 等标记）
        output_path: 输出文件路径
        format_config: 格式配置字典
    """
    ext = Path(output_path).suffix.lower()
    
    if ext == '.docx':
        _write_docx(paragraphs, output_path, format_config)
    elif ext == '.md':
        _write_markdown(paragraphs, output_path)
    elif ext == '.txt':
        _write_text(paragraphs, output_path)
    else:
        raise ValueError(f"Unsupported output format: {ext}")


def _write_docx(paragraphs: List[str], output_path: str, config: Dict) -> None:
    """写入 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # 设置页面边距
    page_cfg = config.get('page', {})
    if page_cfg:
        section = doc.sections[0]
        section.left_margin = Cm(page_cfg.get('left_margin_cm', 2.8))
        section.right_margin = Cm(page_cfg.get('right_margin_cm', 2.6))
        section.top_margin = Cm(page_cfg.get('top_margin_cm', 3.7))
        section.bottom_margin = Cm(page_cfg.get('bottom_margin_cm', 3.5))
    
    # 格式配置
    title_fmt = config.get('title', {})
    level1_fmt = config.get('level1', {})
    level2_fmt = config.get('level2', {})
    level3_fmt = config.get('level3', config.get('level2', {}))
    body_fmt = config.get('body', {})
    
    for i, para_text in enumerate(paragraphs):
        # 提取纯文本（移除标记）
        clean_text = re.sub(r'^\[H\d+\]|\[LIST\]', '', para_text).strip()
        
        if not clean_text:
            continue
        
        # 确定段落类型
        para_type = detect_paragraph_type(para_text, i, len(paragraphs))
        
        # 选择格式
        if para_type == 'title':
            fmt = title_fmt
        elif para_type == 'level1':
            fmt = level1_fmt
        elif para_type == 'level2':
            fmt = level2_fmt
        elif para_type == 'level3':
            fmt = level3_fmt
        elif para_type == 'list':
            fmt = body_fmt.copy()
            fmt['first_line_chars'] = 2  # 列表也保持首行缩进
        else:
            fmt = body_fmt
        
        # 添加段落
        p = doc.add_paragraph(clean_text)
        
        # 设置段落格式
        pf = p.paragraph_format
        
        # 对齐方式
        align = fmt.get('alignment', 'LEFT')
        if 'CENTER' in align:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'RIGHT' in align:
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'JUSTIFY' in align:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 段前段后
        if 'space_before_pt' in fmt:
            pf.space_before = Pt(fmt['space_before_pt'])
        if 'space_after_pt' in fmt:
            pf.space_after = Pt(fmt['space_after_pt'])
        
        # 行距
        ls_val = fmt.get('line_spacing_value', 1.5)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = ls_val
        
        # 首行缩进
        first_chars = fmt.get('first_line_chars', 0)
        if first_chars and first_chars > 0:
            font_size = fmt.get('font_size_pt', 16)
            pPr = p._element.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                from docx.oxml import OxmlElement
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:firstLineChars'), str(int(first_chars * 100)))
            ind.set(qn('w:firstLine'), str(int(first_chars * font_size * 20)))
        
        # 设置字体
        font_name = fmt.get('font_name', '宋体')
        font_size = fmt.get('font_size_pt', 12)
        is_bold = fmt.get('bold', False)
        
        if p.runs:
            run = p.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = is_bold
            
            # 设置中文字体
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)
    
    doc.save(output_path)


def _write_markdown(paragraphs: List[str], output_path: str) -> None:
    """写入 Markdown 文件"""
    lines = []
    
    for para_text in paragraphs:
        clean_text = re.sub(r'^\[H\d+\]|\[LIST\]', '', para_text).strip()
        if not clean_text:
            continue
        
        # 识别标题层级并转换
        if re.match(r'^\[H1\]', para_text):
            lines.append(f"# {clean_text}")
        elif re.match(r'^\[H2\]', para_text):
            lines.append(f"## {clean_text}")
        elif re.match(r'^\[H3\]', para_text):
            lines.append(f"### {clean_text}")
        elif re.match(r'^\[LIST\]', para_text):
            lines.append(f"- {clean_text}")
        else:
            lines.append(clean_text)
        
        lines.append("")  # 空行分隔
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


def _write_text(paragraphs: List[str], output_path: str) -> None:
    """写入纯文本文件"""
    lines = []
    
    for para_text in paragraphs:
        clean_text = re.sub(r'^\[H\d+\]|\[LIST\]', '', para_text).strip()
        if not clean_text:
            continue
        lines.append(clean_text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(lines))


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python format_bridge.py <input_file> <output_file> [format_config.json]")
        print("")
        print("Examples:")
        print("  python format_bridge.py input.md output.docx config.json")
        print("  python format_bridge.py input.txt output.md config.json")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    config_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    print(f"Reading: {input_file}")
    paragraphs, metadata = read_document(input_file)
    print(f"  Found {len(paragraphs)} paragraphs")
    print(f"  Format: {metadata.get('format', 'unknown')}")
    
    # 加载格式配置
    if config_file:
        import json
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)
    else:
        # 使用默认公文格式
        config = {
            'title': {'font_name': '方正小标宋体', 'font_size_pt': 22, 'bold': True, 'alignment': 'CENTER', 'first_line_chars': 0},
            'level1': {'font_name': '黑体', 'font_size_pt': 16, 'bold': True, 'alignment': 'LEFT', 'first_line_chars': 0},
            'level2': {'font_name': '楷体_GB2312', 'font_size_pt': 16, 'bold': True, 'alignment': 'LEFT', 'first_line_chars': 0},
            'body': {'font_name': '仿宋_GB2312', 'font_size_pt': 16, 'bold': False, 'alignment': 'JUSTIFY', 'first_line_chars': 2, 'line_spacing_value': 1.5},
            'page': {'left_margin_cm': 2.8, 'right_margin_cm': 2.6, 'top_margin_cm': 3.7, 'bottom_margin_cm': 3.5}
        }
    
    print(f"Writing: {output_file}")
    write_document(paragraphs, output_file, config)
    print("Done!")

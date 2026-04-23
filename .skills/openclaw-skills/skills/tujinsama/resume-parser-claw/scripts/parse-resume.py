#!/usr/bin/env python3
"""
简历解析脚本 — 支持 PDF、Word、图片（OCR）、纯文本
用法:
  parse  --file <path>              解析单份简历，输出 JSON
  batch  --dir <path> [--output <file>] [--format json|excel|csv]  批量解析
  export --input <json> --output <file> --format excel|csv         转换格式
"""
import argparse, json, os, sys, re
from datetime import datetime
from pathlib import Path


def extract_text_from_pdf(path: str) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    except ImportError:
        sys.exit("缺少依赖: pip install pdfplumber")


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except ImportError:
        sys.exit("缺少依赖: pip install python-docx")


def extract_text_from_image(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        return pytesseract.image_to_string(Image.open(path), lang="chi_sim+eng")
    except ImportError:
        sys.exit("缺少依赖: pip install pytesseract pillow")


def load_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(path)
    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff"):
        return extract_text_from_image(path)
    else:
        return Path(path).read_text(encoding="utf-8", errors="ignore")


# ── 字段提取 ──────────────────────────────────────────────

PHONE_RE = re.compile(r"(?:\+86[-\s]?)?1[3-9]\d{9}")
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
DEGREE_MAP = {"博士": "博士", "硕士": "硕士", "本科": "本科", "学士": "本科",
              "专科": "专科", "大专": "专科", "高中": "高中"}
YEAR_RE = re.compile(r"(20\d{2}|19\d{2})[.\-/年](\d{1,2})?")


def extract_basic(text: str) -> dict:
    phone = PHONE_RE.search(text)
    email = EMAIL_RE.search(text)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    name = lines[0] if lines else ""
    return {
        "name": name,
        "phone": phone.group() if phone else "",
        "email": email.group() if email else "",
        "gender": "未知",
        "age": None,
        "city": "",
    }


def extract_highest_degree(text: str) -> str:
    for kw, std in DEGREE_MAP.items():
        if kw in text:
            return std
    return ""


def extract_work_years(text: str) -> int:
    years = [int(m.group(1)) for m in YEAR_RE.finditer(text)]
    if not years:
        return 0
    return max(0, datetime.now().year - min(years))


def extract_skills(text: str) -> list:
    """简单关键词匹配，实际使用时可接入 NLP 模型"""
    KNOWN = [
        "Python", "Java", "Go", "JavaScript", "TypeScript", "C++", "C#", "Ruby",
        "React", "Vue", "Angular", "Spring Boot", "Django", "FastAPI",
        "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch",
        "Docker", "Kubernetes", "AWS", "阿里云", "腾讯云",
        "机器学习", "深度学习", "NLP", "TensorFlow", "PyTorch",
        "项目管理", "产品管理", "敏捷开发",
    ]
    found = []
    text_lower = text.lower()
    for skill in KNOWN:
        if skill.lower() in text_lower:
            found.append(skill)
    return found


def quality_score(candidate: dict) -> int:
    score = 100
    b = candidate.get("basic", {})
    if not b.get("name"):
        score -= 40
    if not b.get("phone") and not b.get("email"):
        score -= 30
    if not candidate.get("work_years") and not candidate.get("highest_degree"):
        score -= 20
    return max(0, score)


def parse_resume(path: str) -> dict:
    text = load_text(path)
    basic = extract_basic(text)
    candidate = {
        "source_file": os.path.basename(path),
        "parse_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "basic": basic,
        "highest_degree": extract_highest_degree(text),
        "work_years": extract_work_years(text),
        "skills": {"tech": extract_skills(text), "soft": [], "languages": [], "certifications": []},
        "raw_text_length": len(text),
    }
    candidate["quality_score"] = quality_score(candidate)
    return candidate


# ── CLI ───────────────────────────────────────────────────

def cmd_parse(args):
    result = parse_resume(args.file)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def cmd_batch(args):
    directory = Path(args.dir)
    exts = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".txt"}
    files = [f for f in directory.iterdir() if f.suffix.lower() in exts]
    results, failed = [], []
    for f in files:
        try:
            results.append(parse_resume(str(f)))
        except Exception as e:
            failed.append({"file": f.name, "error": str(e)})

    fmt = getattr(args, "format", "json")
    output = getattr(args, "output", None)

    if fmt == "json":
        data = json.dumps({"results": results, "failed": failed}, ensure_ascii=False, indent=2)
        if output:
            Path(output).write_text(data, encoding="utf-8")
            print(f"已导出 {len(results)} 条记录到 {output}，失败 {len(failed)} 条")
        else:
            print(data)
    else:
        _export_tabular(results, output or "results.xlsx", fmt)
        print(f"已导出 {len(results)} 条记录，失败 {len(failed)} 条")


def _export_tabular(results: list, output: str, fmt: str):
    rows = []
    for r in results:
        b = r.get("basic", {})
        rows.append({
            "姓名": b.get("name", ""),
            "手机": b.get("phone", ""),
            "邮箱": b.get("email", ""),
            "工作年限": r.get("work_years", ""),
            "最高学历": r.get("highest_degree", ""),
            "技能标签": ", ".join(r.get("skills", {}).get("tech", [])),
            "质量分": r.get("quality_score", ""),
            "解析时间": r.get("parse_time", ""),
            "来源文件": r.get("source_file", ""),
        })
    try:
        import pandas as pd
        df = pd.DataFrame(rows)
        if fmt == "csv":
            df.to_csv(output, index=False, encoding="utf-8-sig")
        else:
            df.to_excel(output, index=False)
    except ImportError:
        sys.exit("缺少依赖: pip install pandas openpyxl")


def cmd_export(args):
    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    results = data if isinstance(data, list) else data.get("results", [])
    _export_tabular(results, args.output, args.format)
    print(f"已导出到 {args.output}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="简历解析工具")
    sub = parser.add_subparsers(dest="cmd")

    p_parse = sub.add_parser("parse")
    p_parse.add_argument("--file", required=True)

    p_batch = sub.add_parser("batch")
    p_batch.add_argument("--dir", required=True)
    p_batch.add_argument("--output")
    p_batch.add_argument("--format", default="json", choices=["json", "excel", "csv"])

    p_export = sub.add_parser("export")
    p_export.add_argument("--input", required=True)
    p_export.add_argument("--output", required=True)
    p_export.add_argument("--format", default="excel", choices=["excel", "csv"])

    args = parser.parse_args()
    {"parse": cmd_parse, "batch": cmd_batch, "export": cmd_export}.get(
        args.cmd, lambda _: parser.print_help()
    )(args)

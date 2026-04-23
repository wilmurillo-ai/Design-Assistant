#!/usr/bin/env python3
"""文件处理技能 - 自动处理各种文件格式"""
import os
import sys
import json
import tempfile

# 支持的格式
SUPPORTED = {
    'pdf': 'PDF文件',
    'xlsx': 'Excel文件', 
    'xls': 'Excel文件',
    'csv': 'CSV文件',
    'docx': 'Word文档',
    'txt': '文本文件',
    'jpg': '图片',
    'jpeg': '图片',
    'png': '图片'
}

def get_ext(filename):
    return filename.split('.')[-1].lower()

def process_pdf(filepath):
    """处理 PDF 文件"""
    try:
        import pdfplumber
        
        with pdfplumber.open(filepath) as pdf:
            text = ""
            tables = []
            
            for page in pdf.pages:
                # 提取文字
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- 第{page.page_number}页 ---\n"
                    text += page_text[:2000]  # 限制每页字数
                
                # 提取表格
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
            
            result = f"📄 PDF 文档\n"
            result += f"- 页数: {len(pdf.pages)}\n"
            result += f"- 文字长度: {len(text)} 字符\n"
            
            if tables:
                result += f"- 表格数: {len(tables)}\n"
            
            result += f"\n📝 内容预览:\n{text[:1500]}..."
            return result
    except Exception as e:
        return f"❌ PDF 处理失败: {e}"

def process_excel(filepath):
    """处理 Excel 文件"""
    try:
        import openpyxl
        
        wb = openpyxl.load_workbook(filepath, data_only=True)
        
        result = f"📊 Excel 文件\n"
        result += f"- 工作表: {', '.join(wb.sheetnames)}\n"
        
        # 读取第一个表
        ws = wb[wb.sheetnames[0]]
        rows = list(ws.iter_rows(max_row=20, values_only=True))
        
        result += f"- 前{len(rows)}行预览:\n"
        
        for i, row in enumerate(rows[:10]):
            result += f"  {i+1}. {' | '.join(str(c) if c else '' for c in row[:5])}\n"
        
        return result
    except Exception as e:
        return f"❌ Excel 处理失败: {e}"

def process_csv(filepath):
    """处理 CSV 文件"""
    try:
        import pandas as pd
        
        df = pd.read_csv(filepath, nrows=20)
        
        result = f"📋 CSV 文件\n"
        result += f"- 行数: {len(df)}\n"
        result += f"- 列名: {', '.join(df.columns[:5].tolist())}\n"
        result += f"\n📝 数据预览:\n"
        result += df.head(5).to_string()
        
        return result
    except Exception as e:
        return f"❌ CSV 处理失败: {e}"

def process_image(filepath):
    """处理图片 - OCR 文字识别"""
    try:
        import pytesseract
        from PIL import Image
        
        img = Image.open(filepath)
        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        
        if text.strip():
            return f"🖼️ 图片文字识别结果:\n\n{text[:2000]}"
        else:
            return "🖼️ 图片中未识别到文字"
    except Exception as e:
        return f"⚠️ OCR 识别暂时不可用 ({e})\n💡 请确保已安装 tesseract"

def process_txt(filepath):
    """处理文本文件"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read(5000)
        
        lines = content.split('\n')
        
        result = f"📄 文本文件\n"
        result += f"- 字符数: {len(content)}\n"
        result += f"- 行数: {len(lines)}\n"
        result += f"\n📝 内容预览:\n{content[:1500]}"
        
        return result
    except:
        try:
            with open(filepath, 'r', encoding='gbk') as f:
                content = f.read(5000)
            return f"📄 文本文件 (GBK编码)\n\n{content[:1500]}"
        except Exception as e:
            return f"❌ 文本读取失败: {e}"

def process_docx(filepath):
    """处理 Word 文档"""
    try:
        import docx
        
        doc = docx.Document(filepath)
        
        text = "\n".join([p.text for p in doc.paragraphs[:50]])
        
        result = f"📄 Word 文档\n"
        result += f"- 段落数: {len(doc.paragraphs)}\n"
        result += f"\n📝 内容预览:\n{text[:1500]}"
        
        return result
    except Exception as e:
        return f"❌ Word 处理失败: {e}"

def main():
    if len(sys.argv) < 2:
        print("用法: python3 file_processor.py <文件路径>")
        return
    
    filepath = sys.argv[1]
    
    if not os.path.exists(filepath):
        print(f"❌ 文件不存在: {filepath}")
        return
    
    ext = get_ext(filepath)
    
    if ext not in SUPPORTED:
        print(f"❌ 不支持的格式: {ext}")
        print(f"支持的格式: {', '.join(SUPPORTED.keys())}")
        return
    
    print(f"🔄 正在处理 {SUPPORTED.get(ext, ext)}...")
    
    if ext == 'pdf':
        print(process_pdf(filepath))
    elif ext in ('xlsx', 'xls'):
        print(process_excel(filepath))
    elif ext == 'csv':
        print(process_csv(filepath))
    elif ext in ('jpg', 'jpeg', 'png'):
        print(process_image(filepath))
    elif ext == 'txt':
        print(process_txt(filepath))
    elif ext == 'docx':
        print(process_docx(filepath))

if __name__ == "__main__":
    main()

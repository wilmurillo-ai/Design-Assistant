#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
입찰/제안 공고 문서 파서 (Bid/Proposal Document Parser)

다양한 형식의 입찰/제안 공고 문서를 파싱하여 표준화된 JSON 형식으로 변환합니다.

지원 형식:
  - PDF: pdfplumber 또는 PyMuPDF 사용
  - HWP: pyhwp 또는 olefile 사용
  - HWPX: ZIP 기반 XML 추출
  - DOCX: python-docx 사용
  - URL/웹페이지: requests + BeautifulSoup 사용

사용법:
  # 단일 파일/URL 파싱
  python document_parser.py parse /path/to/document.pdf
  python document_parser.py parse https://example.com/announcement.html

  # 디렉토리 배치 처리
  python document_parser.py batch /path/to/documents --output-dir /output/path

  # JSON 출력
  python document_parser.py parse document.pdf --json

출력 포맷:
  {
    "metadata": {
      "source_file": "announcement.pdf",
      "source_type": "pdf",
      "parsed_at": "ISO 8601",
      "page_count": 12,
      "parser_version": "1.0"
    },
    "content": {
      "full_text": "전체 텍스트...",
      "sections": [
        {
          "title": "섹션 제목",
          "text": "섹션 내용...",
          "page_range": [1, 3]
        }
      ],
      "tables": [
        {
          "page": 2,
          "headers": ["col1", "col2"],
          "rows": [["val1", "val2"]]
        }
      ]
    }
  }
"""

import argparse
import json
import logging
import os
import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse

# 외부 라이브러리 임포트 (선택적)
try:
    import requests
    from bs4 import BeautifulSoup
    HAS_WEB_PARSER = True
except ImportError:
    HAS_WEB_PARSER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from pyhwp import HWPDocument
    HAS_PYHWP = True
except ImportError:
    HAS_PYHWP = False

try:
    import olefile
    HAS_OLEFILE = True
except ImportError:
    HAS_OLEFILE = False

try:
    import xml.etree.ElementTree as ET
    HAS_XML = True
except ImportError:
    HAS_XML = False

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    stream=sys.stderr
)
logger = logging.getLogger(__name__)

PARSER_VERSION = "1.0"


class DocumentParser:
    """다중 형식 문서 파서 클래스"""

    def __init__(self):
        self.parser_version = PARSER_VERSION

    def parse(self, source: str) -> Dict[str, Any]:
        """
        파일 또는 URL을 파싱하고 표준화된 JSON 반환

        Args:
            source: 파일 경로 또는 URL

        Returns:
            표준화된 문서 데이터 딕셔너리
        """
        if self._is_url(source):
            return self._parse_url(source)
        else:
            return self._parse_file(source)

    def _is_url(self, source: str) -> bool:
        """URL 여부 확인"""
        try:
            result = urlparse(source)
            return result.scheme in ('http', 'https')
        except Exception:
            return False

    def _parse_file(self, filepath: str) -> Dict[str, Any]:
        """파일 파싱"""
        filepath = Path(filepath)

        if not filepath.exists():
            logger.error(f"파일이 존재하지 않음: {filepath}")
            raise FileNotFoundError(f"파일이 존재하지 않음: {filepath}")

        file_ext = filepath.suffix.lower()
        logger.info(f"파일 파싱 시작: {filepath} (형식: {file_ext})")

        try:
            if file_ext == '.pdf':
                return self._parse_pdf(filepath)
            elif file_ext == '.hwp':
                return self._parse_hwp(filepath)
            elif file_ext == '.hwpx':
                return self._parse_hwpx(filepath)
            elif file_ext == '.docx':
                return self._parse_docx(filepath)
            else:
                logger.error(f"지원하지 않는 파일 형식: {file_ext}")
                raise ValueError(f"지원하지 않는 파일 형식: {file_ext}")
        except Exception as e:
            logger.error(f"파일 파싱 실패: {filepath}, 오류: {str(e)}")
            raise

    def _parse_pdf(self, filepath: Path) -> Dict[str, Any]:
        """PDF 파싱"""
        if HAS_PDFPLUMBER:
            return self._parse_pdf_pdfplumber(filepath)
        elif HAS_PYMUPDF:
            return self._parse_pdf_pymupdf(filepath)
        else:
            logger.error("PDF 파싱 라이브러리가 설치되지 않음 (pdfplumber 또는 PyMuPDF 필요)")
            raise ImportError("PDF 파싱 라이브러리가 설치되지 않음. pdfplumber 또는 PyMuPDF를 설치하세요.")

    def _parse_pdf_pdfplumber(self, filepath: Path) -> Dict[str, Any]:
        """pdfplumber를 사용한 PDF 파싱"""
        with pdfplumber.open(filepath) as pdf:
            page_count = len(pdf.pages)
            full_text = ""
            all_tables = []
            page_to_text = {}

            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                full_text += f"\n--- 페이지 {page_num} ---\n{text}"
                page_to_text[page_num] = text

                # 테이블 추출
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        all_tables.append({
                            "page": page_num,
                            "headers": table[0] if table else [],
                            "rows": table[1:] if len(table) > 1 else []
                        })

            sections = self._extract_sections(full_text, page_to_text)

            return {
                "metadata": {
                    "source_file": filepath.name,
                    "source_type": "pdf",
                    "parsed_at": datetime.utcnow().isoformat() + "Z",
                    "page_count": page_count,
                    "parser_version": self.parser_version
                },
                "content": {
                    "full_text": full_text.strip(),
                    "sections": sections,
                    "tables": all_tables
                }
            }

    def _parse_pdf_pymupdf(self, filepath: Path) -> Dict[str, Any]:
        """PyMuPDF를 사용한 PDF 파싱"""
        doc = fitz.open(filepath)
        page_count = doc.page_count
        full_text = ""
        page_to_text = {}

        for page_num in range(page_count):
            page = doc[page_num]
            text = page.get_text()
            full_text += f"\n--- 페이지 {page_num + 1} ---\n{text}"
            page_to_text[page_num + 1] = text

        doc.close()

        sections = self._extract_sections(full_text, page_to_text)

        # PyMuPDF는 테이블 추출이 제한적이므로 빈 배열 반환
        return {
            "metadata": {
                "source_file": filepath.name,
                "source_type": "pdf",
                "parsed_at": datetime.utcnow().isoformat() + "Z",
                "page_count": page_count,
                "parser_version": self.parser_version
            },
            "content": {
                "full_text": full_text.strip(),
                "sections": sections,
                "tables": []
            }
        }

    def _parse_hwp(self, filepath: Path) -> Dict[str, Any]:
        """HWP 파싱"""
        if HAS_PYHWP:
            return self._parse_hwp_pyhwp(filepath)
        elif HAS_OLEFILE:
            return self._parse_hwp_olefile(filepath)
        else:
            logger.error("HWP 파싱 라이브러리가 설치되지 않음 (pyhwp 또는 olefile 필요)")
            raise ImportError("HWP 파싱 라이브러리가 설치되지 않음. pyhwp 또는 olefile을 설치하세요.")

    def _parse_hwp_pyhwp(self, filepath: Path) -> Dict[str, Any]:
        """pyhwp를 사용한 HWP 파싱"""
        try:
            doc = HWPDocument(str(filepath))
            full_text = doc.text

            return {
                "metadata": {
                    "source_file": filepath.name,
                    "source_type": "hwp",
                    "parsed_at": datetime.utcnow().isoformat() + "Z",
                    "page_count": 0,
                    "parser_version": self.parser_version
                },
                "content": {
                    "full_text": full_text,
                    "sections": self._extract_sections(full_text),
                    "tables": []
                }
            }
        except Exception as e:
            logger.warning(f"pyhwp 파싱 실패, olefile 시도: {str(e)}")
            if HAS_OLEFILE:
                return self._parse_hwp_olefile(filepath)
            else:
                raise

    def _parse_hwp_olefile(self, filepath: Path) -> Dict[str, Any]:
        """olefile을 사용한 HWP 파싱"""
        try:
            ole = olefile.OleFileIO(str(filepath))
            full_text = ""

            # HWP 파일의 일반적인 스트림 이름들
            for stream_name in ole.listdir():
                if stream_name[0] == 'DocInfo' or stream_name[0] == 'BodyText':
                    try:
                        data = ole.openstream(stream_name).read()
                        # 간단한 텍스트 추출 시도
                        text = data.decode('utf-16-le', errors='ignore')
                        full_text += text + "\n"
                    except Exception as e:
                        logger.debug(f"스트림 읽기 실패 {stream_name}: {str(e)}")

            ole.close()

            if not full_text.strip():
                logger.warning("HWP 파일에서 텍스트를 추출하지 못함")
                full_text = "[텍스트 추출 실패]"

            return {
                "metadata": {
                    "source_file": filepath.name,
                    "source_type": "hwp",
                    "parsed_at": datetime.utcnow().isoformat() + "Z",
                    "page_count": 0,
                    "parser_version": self.parser_version
                },
                "content": {
                    "full_text": full_text.strip(),
                    "sections": self._extract_sections(full_text),
                    "tables": []
                }
            }
        except Exception as e:
            logger.error(f"olefile을 사용한 HWP 파싱 실패: {str(e)}")
            raise

    def _parse_hwpx(self, filepath: Path) -> Dict[str, Any]:
        """HWPX 파싱 (ZIP 기반 XML)"""
        if not HAS_XML:
            logger.error("XML 파싱 라이브러리를 사용할 수 없음")
            raise ImportError("XML 파싱 라이브러리가 필요합니다")

        full_text = ""

        try:
            with zipfile.ZipFile(filepath, 'r') as zip_file:
                # Contents 디렉토리의 section*.xml 파일들 찾기
                xml_files = [f for f in zip_file.namelist() if f.startswith('Contents/section') and f.endswith('.xml')]
                xml_files.sort()

                for xml_file in xml_files:
                    try:
                        with zip_file.open(xml_file) as f:
                            root = ET.parse(f).getroot()
                            text = self._extract_text_from_xml(root)
                            if text:
                                full_text += text + "\n"
                    except Exception as e:
                        logger.warning(f"XML 파일 파싱 실패 {xml_file}: {str(e)}")

            if not full_text.strip():
                logger.warning("HWPX 파일에서 텍스트를 추출하지 못함")
                full_text = "[텍스트 추출 실패]"

            return {
                "metadata": {
                    "source_file": filepath.name,
                    "source_type": "hwpx",
                    "parsed_at": datetime.utcnow().isoformat() + "Z",
                    "page_count": len(xml_files),
                    "parser_version": self.parser_version
                },
                "content": {
                    "full_text": full_text.strip(),
                    "sections": self._extract_sections(full_text),
                    "tables": []
                }
            }
        except Exception as e:
            logger.error(f"HWPX 파싱 실패: {str(e)}")
            raise

    def _extract_text_from_xml(self, element: ET.Element) -> str:
        """XML 요소에서 텍스트 재귀적 추출"""
        text = ""
        if element.text:
            text += element.text
        for child in element:
            text += self._extract_text_from_xml(child)
            if child.tail:
                text += child.tail
        return text

    def _parse_docx(self, filepath: Path) -> Dict[str, Any]:
        """DOCX 파싱"""
        if not HAS_DOCX:
            logger.error("python-docx 라이브러리가 설치되지 않음")
            raise ImportError("python-docx 라이브러리가 설치되지 않음. 설치하려면 'pip install python-docx'를 실행하세요.")

        doc = Document(filepath)
        full_text = ""
        all_tables = []

        # 문단 텍스트 추출
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n"

        # 테이블 추출
        for table_idx, table in enumerate(doc.tables):
            headers = []
            rows = []

            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = row_data
                else:
                    rows.append(row_data)

            all_tables.append({
                "page": 0,
                "headers": headers,
                "rows": rows
            })

        sections = self._extract_sections(full_text)

        return {
            "metadata": {
                "source_file": filepath.name,
                "source_type": "docx",
                "parsed_at": datetime.utcnow().isoformat() + "Z",
                "page_count": 0,
                "parser_version": self.parser_version
            },
            "content": {
                "full_text": full_text.strip(),
                "sections": sections,
                "tables": all_tables
            }
        }

    def _parse_url(self, url: str) -> Dict[str, Any]:
        """URL에서 웹페이지 파싱"""
        if not HAS_WEB_PARSER:
            logger.error("웹 파싱 라이브러리가 설치되지 않음 (requests, BeautifulSoup 필요)")
            raise ImportError("웹 파싱을 위해 requests와 BeautifulSoup4를 설치하세요.")

        try:
            logger.info(f"URL 파싱 시작: {url}")
            response = requests.get(url, timeout=10)
            response.encoding = response.apparent_encoding or 'utf-8'
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # 메인 콘텐츠 추출
            full_text = self._extract_main_content(soup)

            return {
                "metadata": {
                    "source_file": url,
                    "source_type": "url",
                    "parsed_at": datetime.utcnow().isoformat() + "Z",
                    "page_count": 1,
                    "parser_version": self.parser_version
                },
                "content": {
                    "full_text": full_text,
                    "sections": self._extract_sections(full_text),
                    "tables": self._extract_tables_from_html(soup)
                }
            }
        except requests.RequestException as e:
            logger.error(f"URL 요청 실패: {url}, 오류: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"URL 파싱 실패: {url}, 오류: {str(e)}")
            raise

    def _extract_main_content(self, soup: BeautifulSoup) -> str:
        """BeautifulSoup에서 메인 콘텐츠 추출"""
        # 스크립트와 스타일 제거
        for tag in soup(['script', 'style']):
            tag.decompose()

        # 메인 콘텐츠 찾기
        main_content = soup.find('main') or soup.find('article') or soup.find('body')
        if not main_content:
            main_content = soup

        # 텍스트 추출
        text = main_content.get_text(separator='\n', strip=True)

        # 과도한 공백 정리
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)

    def _extract_tables_from_html(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """HTML에서 테이블 추출"""
        tables = []
        for table_idx, table in enumerate(soup.find_all('table')):
            headers = []
            rows = []

            # 헤더 추출
            thead = table.find('thead')
            if thead:
                for tr in thead.find_all('tr'):
                    headers = [td.get_text(strip=True) for td in tr.find_all(['th', 'td'])]
                    break

            # 데이터 행 추출
            tbody = table.find('tbody') or table
            for tr in tbody.find_all('tr'):
                if not thead or tr.parent.name != 'thead':
                    row = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                    if row:
                        rows.append(row)

            if headers or rows:
                tables.append({
                    "page": 0,
                    "headers": headers,
                    "rows": rows
                })

        return tables

    def _extract_sections(self, full_text: str, page_mapping: Optional[Dict[int, str]] = None) -> List[Dict[str, Any]]:
        """텍스트에서 섹션 추출"""
        sections = []

        # 섹션 구분 패턴들
        patterns = [
            (r'^[IVX]+\.\s+(.+?)$', 'roman'),  # 로마자: I, II, III, ...
            (r'^제\d+장\s+(.+?)$', 'korean_chapter'),  # 제1장, 제2장, ...
            (r'^제\d+조\s+(.+?)$', 'korean_article'),  # 제1조, 제2조, ...
            (r'^\d+\.\s+(.+?)$', 'number'),  # 1. 2. 3. ...
            (r'^[가-힣]\.\s+(.+?)$', 'korean_letter'),  # 가. 나. 다. ...
            (r'^\d+\)\s+(.+?)$', 'paren_number'),  # 1) 2) 3) ...
        ]

        lines = full_text.split('\n')
        current_section = None
        current_text = []
        current_page = 1
        start_page = 1

        for line in lines:
            # 페이지 마커 확인
            page_match = re.match(r'--- 페이지 (\d+) ---', line)
            if page_match:
                current_page = int(page_match.group(1))
                if current_section is not None:
                    start_page = current_page
                continue

            # 섹션 헤더 확인
            section_found = False
            for pattern, pattern_type in patterns:
                match = re.match(pattern, line.strip())
                if match:
                    # 이전 섹션 저장
                    if current_section:
                        section_text = '\n'.join(current_text).strip()
                        if section_text:
                            sections.append({
                                "title": current_section,
                                "text": section_text,
                                "page_range": [start_page, current_page]
                            })

                    # 새 섹션 시작
                    current_section = match.group(1).strip()
                    current_text = []
                    start_page = current_page
                    section_found = True
                    break

            if not section_found and current_section:
                if line.strip():
                    current_text.append(line)

        # 마지막 섹션 저장
        if current_section:
            section_text = '\n'.join(current_text).strip()
            if section_text:
                sections.append({
                    "title": current_section,
                    "text": section_text,
                    "page_range": [start_page, current_page]
                })

        return sections


def parse_command(args) -> None:
    """parse 서브명령 실행"""
    parser = DocumentParser()

    try:
        result = parser.parse(args.source)

        if args.json:
            output = json.dumps(result, ensure_ascii=False, indent=2)
            print(output)
        else:
            # 기본 텍스트 출력
            print(f"소스: {result['metadata']['source_file']}")
            print(f"형식: {result['metadata']['source_type']}")
            print(f"파싱 시간: {result['metadata']['parsed_at']}")
            print(f"페이지 수: {result['metadata']['page_count']}")
            print(f"\n텍스트 길이: {len(result['content']['full_text'])} 자")
            print(f"섹션 수: {len(result['content']['sections'])}")
            print(f"테이블 수: {len(result['content']['tables'])}")

            if args.verbose:
                print(f"\n전체 텍스트:\n{result['content']['full_text'][:500]}...")
    except Exception as e:
        logger.error(f"파싱 실패: {str(e)}")
        sys.exit(1)


def batch_command(args) -> None:
    """batch 서브명령 실행"""
    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir) if args.output_dir else input_dir

    if not input_dir.is_dir():
        logger.error(f"디렉토리가 아닙니다: {input_dir}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    parser = DocumentParser()
    supported_exts = {'.pdf', '.hwp', '.hwpx', '.docx'}

    files = [f for f in input_dir.iterdir()
             if f.is_file() and f.suffix.lower() in supported_exts]

    if not files:
        logger.warning(f"지원하는 파일을 찾을 수 없습니다: {input_dir}")
        return

    logger.info(f"{len(files)}개 파일을 처리합니다")

    for idx, filepath in enumerate(files, 1):
        try:
            logger.info(f"[{idx}/{len(files)}] 처리 중: {filepath.name}")
            result = parser.parse(str(filepath))

            # 출력 파일명
            output_name = filepath.stem + '.json'
            output_path = output_dir / output_name

            # JSON 저장
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)

            logger.info(f"저장 완료: {output_path}")
        except Exception as e:
            logger.error(f"처리 실패 {filepath.name}: {str(e)}")


def main():
    """메인 진입점"""
    parser = argparse.ArgumentParser(
        description='입찰/제안 공고 문서 파서 - 다양한 형식의 문서를 표준화된 JSON으로 변환합니다',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
예제:
  python document_parser.py parse document.pdf
  python document_parser.py parse document.pdf --json
  python document_parser.py parse https://example.com/announcement.html
  python document_parser.py batch /path/to/documents --output-dir /output
        """
    )

    subparsers = parser.add_subparsers(dest='command', help='실행할 명령')

    # parse 서브명령
    parse_parser = subparsers.add_parser('parse', help='단일 파일 또는 URL 파싱')
    parse_parser.add_argument('source', help='파일 경로 또는 URL')
    parse_parser.add_argument('--json', action='store_true', help='JSON 형식으로 출력')
    parse_parser.add_argument('--verbose', '-v', action='store_true', help='상세 출력')
    parse_parser.set_defaults(func=parse_command)

    # batch 서브명령
    batch_parser = subparsers.add_parser('batch', help='디렉토리 배치 처리')
    batch_parser.add_argument('input_dir', help='입력 디렉토리')
    batch_parser.add_argument('--output-dir', '-o', help='출력 디렉토리 (기본값: 입력 디렉토리)')
    batch_parser.set_defaults(func=batch_command)

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    args.func(args)


if __name__ == '__main__':
    main()

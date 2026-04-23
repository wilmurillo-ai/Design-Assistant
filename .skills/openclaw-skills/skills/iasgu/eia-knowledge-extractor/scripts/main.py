"""
环评知识库提炼主程序 (改进版)

主要改进：
1. 支持PDF文件解析
2. 集成数据验证机制
3. 添加表格结构化提取
4. 增加数据一致性检查
"""

import os
import sys
import re
import csv
from pathlib import Path
from typing import List, Dict, Any
import json

# 添加当前目录到Python路径
sys.path.insert(0, str(Path(__file__).parent))

try:
    import fitz  # PyMuPDF
    HAS_PDF_SUPPORT = True
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    # 尝试使用 pdfminer.six 作为替代
    try:
        from pdfminer.high_level import extract_text, extract_pages
        from pdfminer.layout import LAParams, LTTextContainer, LTTable
        HAS_PDF_SUPPORT = True
        print("提示: 使用 pdfminer.six 作为PDF解析引擎")
    except ImportError:
        HAS_PDF_SUPPORT = False
        print("警告: 未安装PyMuPDF或pdfminer.six，PDF解析功能受限")

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    print("警告: 未安装pandas，表格数据处理功能受限")

try:
    from docx import Document
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False
    print("警告: 未安装python-docx，DOCX解析功能受限")

from extract_pollutant_info import extract_pollutant_info, save_to_csv as save_pollutant_csv
from extract_emission_standards import extract_emission_standards, save_to_csv as save_standards_csv
from extract_source_calculation import (
    extract_gas_source_calculation,
    extract_water_source_calculation,
    extract_solid_waste_calculation,
    extract_noise_source_calculation,
    save_gas_calculation_to_csv,
    save_water_calculation_to_csv,
    save_solid_waste_calculation_to_csv,
    save_noise_calculation_to_csv
)


class PDFExtractor:
    """PDF文件提取器，支持表格和文本提取"""

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.doc = None
        self.all_text = ""
        self.tables_data = []

    def extract(self) -> tuple:
        """
        提取PDF中的文本和表格数据

        Returns:
            tuple: (all_text, tables_list)
        """
        if not HAS_PDF_SUPPORT:
            raise ImportError("需要安装PyMuPDF库: pip install pymupdf")

        self.doc = fitz.open(self.pdf_path)

        print(f"正在解析PDF文件，共 {len(self.doc)} 页")

        # 提取所有文本
        for page_num, page in enumerate(self.doc):
            page_text = page.get_text()
            self.all_text += page_text + "\n"

            # 提取表格
            try:
                tables = page.find_tables()
                if tables.tables:
                    for table_idx, table in enumerate(tables.tables):
                        try:
                            df = table.to_pandas()
                            table_info = {
                                'page': page_num + 1,
                                'table_idx': table_idx,
                                'data': df,
                                'shape': df.shape
                            }
                            self.tables_data.append(table_info)
                        except Exception as e:
                            print(f"警告: 第{page_num+1}页表格{table_idx+1}提取失败: {e}")
            except Exception as e:
                print(f"警告: 第{page_num+1}页表格查找失败: {e}")

        self.doc.close()
        return self.all_text, self.tables_data


class DocxExtractor:
    """DOCX文件提取器，支持表格和文本提取"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.all_text = ""
        self.tables_data = []

    def extract(self) -> tuple:
        """
        提取DOCX中的文本和表格数据

        Returns:
            tuple: (all_text, tables_list)
        """
        if not HAS_DOCX_SUPPORT:
            raise ImportError("需要安装python-docx库: pip install python-docx")

        doc = Document(self.docx_path)
        print(f"正在解析DOCX文件...")

        # 提取所有段落文本
        for paragraph in doc.paragraphs:
            self.all_text += paragraph.text + "\n"

        # 提取表格
        if not HAS_PANDAS:
            print("警告: 需要pandas库来处理表格数据")
            return self.all_text, []

        for table_idx, table in enumerate(doc.tables):
            try:
                # 提取表格数据
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                # 转换为DataFrame
                df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else None)

                table_info = {
                    'page': 0,
                    'table_idx': table_idx,
                    'data': df,
                    'shape': df.shape
                }
                self.tables_data.append(table_info)
            except Exception as e:
                print(f"警告: 表格{table_idx+1}提取失败: {e}")
                continue

        print(f"✓ 成功提取文本内容 ({len(self.all_text)} 字符)")
        print(f"✓ 成功识别 {len(self.tables_data)} 个表格")

        return self.all_text, self.tables_data


class DataValidator:
    """数据验证器，检查数据一致性和合理性"""

    def __init__(self):
        self.errors = []
        self.warnings = []

    def validate_numeric_range(self, value: float, min_val: float = 0, max_val: float = None,
                               field_name: str = "") -> bool:
        """验证数值范围"""
        try:
            val = float(value)
            if val < min_val:
                self.errors.append(f"{field_name}: {val} 小于最小值 {min_val}")
                return False
            if max_val is not None and val > max_val:
                self.warnings.append(f"{field_name}: {val} 大于参考值 {max_val}")
            return True
        except (ValueError, TypeError):
            self.errors.append(f"{field_name}: 无法转换为数值 - {value}")
            return False

    def validate_mass_balance(self, generation: float, emission: float,
                            tolerance: float = 0.1, item_name: str = "") -> bool:
        """验证物质平衡：排放量不应大于产生量"""
        try:
            gen = float(generation)
            emi = float(emission)
            if emi > gen * (1 + tolerance):
                self.errors.append(
                    f"{item_name}: 排放量({emi})大于产生量({gen})，不符合物质平衡"
                )
                return False
            return True
        except (ValueError, TypeError):
            return False

    def check_completeness(self, data: List[Dict], required_fields: List[str],
                          data_type: str = "") -> bool:
        """检查数据完整性"""
        missing_count = 0
        for record in data:
            for field in required_fields:
                if field not in record or not record.get(field):
                    missing_count += 1

        if missing_count > 0:
            self.warnings.append(
                f"{data_type}: 发现 {missing_count} 个必填字段缺失"
            )
            return False
        return True

    def validate_units(self, value: str, expected_units: List[str],
                      field_name: str = "") -> bool:
        """验证单位格式"""
        for unit in expected_units:
            if unit in str(value):
                return True
        self.warnings.append(f"{field_name}: 单位格式可能不符合标准 - {value}")
        return False

    def get_report(self) -> str:
        """生成验证报告"""
        report = []
        report.append("=" * 60)
        report.append("数据验证报告")
        report.append("=" * 60)

        if self.errors:
            report.append(f"\n发现 {len(self.errors)} 个错误:")
            for i, error in enumerate(self.errors, 1):
                report.append(f"  {i}. {error}")

        if self.warnings:
            report.append(f"\n发现 {len(self.warnings)} 个警告:")
            for i, warning in enumerate(self.warnings, 1):
                report.append(f"  {i}. {warning}")

        if not self.errors and not self.warnings:
            report.append("\n✓ 所有数据验证通过")

        report.append("=" * 60)
        return "\n".join(report)


def extract_pdf_content(report_file: str) -> tuple:
    """
    提取PDF/DOCX文件内容

    Args:
        report_file: 文件路径

    Returns:
        tuple: (文本内容, 表格数据列表)
    """
    file_ext = report_file.lower().split('.')[-1]

    if file_ext == 'pdf':
        extractor = PDFExtractor(report_file)
        text_content, tables = extractor.extract()
        return text_content, tables
    elif file_ext in ['docx', 'doc']:
        extractor = DocxExtractor(report_file)
        text_content, tables = extractor.extract()
        return text_content, tables
    else:
        # 处理文本文件
        with open(report_file, 'r', encoding='utf-8') as f:
            content = f.read()
        return content, []


def enhanced_extract_pollutant_info(content: str, tables: list) -> List[Dict]:
    """
    增强版污染因子信息提取，结合表格数据

    Args:
        content: 文本内容
        tables: 表格数据列表

    Returns:
        List[Dict]: 污染因子信息列表
    """
    # 首先尝试从表格中提取
    pollutant_data = []

    if not HAS_PANDAS:
        # 如果没有pandas，直接使用原始方法
        return extract_pollutant_info(content)

    # 在表格中查找污染物相关信息
    for table_info in tables:
        df = table_info['data']

        # 检查是否包含污染因子相关表格
        df_str = df.to_string()
        if any(keyword in df_str for keyword in ['污染物', '副产物', '产生工序', '产生量']):
            try:
                # 提取表格数据
                for idx, row in df.iterrows():
                    if idx == 0:  # 跳过表头
                        continue

                    record = {}
                    # 根据表格结构提取字段
                    for col in df.columns:
                        val = row.get(col)
                        if val and not pd.isna(val):
                            record[str(col)] = str(val).strip()

                    if record:
                        pollutant_data.append(record)
            except Exception as e:
                print(f"警告: 表格数据提取失败 - {e}")
                continue

    # 如果表格提取失败，使用原始方法
    if not pollutant_data:
        pollutant_data = extract_pollutant_info(content)

    return pollutant_data


def process_ep_report_enhanced(report_file: str, output_dir: str = "output",
                              enable_validation: bool = True) -> Dict[str, Any]:
    """
    增强版环评报告处理函数

    Args:
        report_file: 环评报告文件路径
        output_dir: 输出目录
        enable_validation: 是否启用数据验证

    Returns:
        Dict[str, Any]: 处理结果，包含各类统计数据
    """
    result = {
        'success': False,
        'errors': [],
        'warnings': [],
        'stats': {}
    }

    try:
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)

        # 提取文件内容
        print(f"正在处理环评报告: {report_file}")
        print("-" * 60)

        text_content, tables = extract_pdf_content(report_file)

        if tables:
            print(f"✓ 成功提取文本内容 ({len(text_content)} 字符)")
            print(f"✓ 成功识别 {len(tables)} 个表格")

        # 初始化验证器
        validator = DataValidator() if enable_validation else None

        # 1. 提取污染因子知识库
        print("\n[1/6] 正在提取污染因子信息...")
        pollutant_data = enhanced_extract_pollutant_info(text_content, tables)
        save_pollutant_csv(
            pollutant_data,
            os.path.join(output_dir, "污染因子知识库.csv")
        )
        print(f"✓ 污染因子知识库已生成，共 {len(pollutant_data)} 条记录")
        result['stats']['pollutant_count'] = len(pollutant_data)

        # 2. 提取污染物排放标准知识库
        print("\n[2/6] 正在提取污染物排放标准信息...")
        standards_data = extract_emission_standards(text_content)
        save_standards_csv(
            standards_data,
            os.path.join(output_dir, "污染物排放标准知识库.csv")
        )
        print(f"✓ 污染物排放标准知识库已生成，共 {len(standards_data)} 条记录")
        result['stats']['standards_count'] = len(standards_data)

        # 3. 提取废气源强核算知识库
        print("\n[3/6] 正在提取废气源强核算信息...")
        gas_data = extract_gas_source_calculation(text_content)
        save_gas_calculation_to_csv(
            gas_data,
            os.path.join(output_dir, "废气源强核算知识库.csv")
        )
        print(f"✓ 废气源强核算知识库已生成，共 {len(gas_data)} 条记录")
        result['stats']['gas_source_count'] = len(gas_data)

        # 数据验证：检查产生量与排放量的关系
        if validator:
            for record in gas_data:
                gen = record.get('类比项目污染物量', '0')
                emi = record.get('污染物排放量', '0')
                if gen and emi:
                    validator.validate_mass_balance(
                        gen, emi, 0.1,
                        record.get('核算污染因子', '未知')
                    )

        # 4. 提取废水源强核算知识库
        print("\n[4/6] 正在提取废水源强核算信息...")
        water_data = extract_water_source_calculation(text_content)
        save_water_calculation_to_csv(
            water_data,
            os.path.join(output_dir, "废水源强核算知识库.csv")
        )
        print(f"✓ 废水源强核算知识库已生成，共 {len(water_data)} 条记录")
        result['stats']['water_source_count'] = len(water_data)

        # 5. 提取固废源强核算知识库
        print("\n[5/6] 正在提取固废源强核算信息...")
        solid_waste_data = extract_solid_waste_calculation(text_content)
        save_solid_waste_calculation_to_csv(
            solid_waste_data,
            os.path.join(output_dir, "固废源强核算知识库.csv")
        )
        print(f"✓ 固废源强核算知识库已生成，共 {len(solid_waste_data)} 条记录")
        result['stats']['solid_waste_count'] = len(solid_waste_data)

        # 6. 提取噪声源强核算知识库
        print("\n[6/6] 正在提取噪声源强核算信息...")
        noise_data = extract_noise_source_calculation(text_content)
        save_noise_calculation_to_csv(
            noise_data,
            os.path.join(output_dir, "噪声源强核算知识库.csv")
        )
        print(f"✓ 噪声源强核算知识库已生成，共 {len(noise_data)} 条记录")
        result['stats']['noise_source_count'] = len(noise_data)

        print("\n" + "=" * 60)
        print(f"处理完成！所有知识库文件已保存到: {output_dir}")

        # 生成验证报告
        if validator:
            validation_report = validator.get_report()
            report_path = os.path.join(output_dir, "数据验证报告.txt")
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(validation_report)
            print(f"\n数据验证报告已生成: {report_path}")

            result['validation_errors'] = len(validator.errors)
            result['validation_warnings'] = len(validator.warnings)

        # 生成处理摘要
        summary = generate_summary(result, output_dir)
        print(f"\n处理摘要已生成: {summary}")

        result['success'] = True

    except Exception as e:
        error_msg = f"处理失败: {str(e)}"
        print(f"\n错误: {error_msg}")
        result['errors'].append(error_msg)

    return result


def generate_summary(result: Dict, output_dir: str) -> str:
    """生成处理摘要文件"""
    summary_path = os.path.join(output_dir, "处理摘要.txt")

    summary_lines = []
    summary_lines.append("=" * 60)
    summary_lines.append("环评知识库提炼处理摘要")
    summary_lines.append("=" * 60)
    summary_lines.append(f"\n处理时间: {Path(__file__).stat().st_mtime}")
    summary_lines.append(f"\n生成的知识库文件:")

    stats = result.get('stats', {})
    summary_lines.append(f"  - 污染因子知识库: {stats.get('pollutant_count', 0)} 条记录")
    summary_lines.append(f"  - 污染物排放标准: {stats.get('standards_count', 0)} 条记录")
    summary_lines.append(f"  - 废气源强核算: {stats.get('gas_source_count', 0)} 条记录")
    summary_lines.append(f"  - 废水源强核算: {stats.get('water_source_count', 0)} 条记录")
    summary_lines.append(f"  - 固废源强核算: {stats.get('solid_waste_count', 0)} 条记录")
    summary_lines.append(f"  - 噪声源强核算: {stats.get('noise_source_count', 0)} 条记录")

    if result.get('validation_errors') is not None:
        summary_lines.append(f"\n数据验证结果:")
        summary_lines.append(f"  - 错误: {result['validation_errors']} 个")
        summary_lines.append(f"  - 警告: {result['validation_warnings']} 个")

    summary_lines.append("\n" + "=" * 60)

    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(summary_lines))

    return summary_path


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(
        description='环评知识库提炼工具 (增强版)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  python main.py report.pdf                    # 处理PDF文件
  python main.py report.pdf -o output_dir      # 指定输出目录
  python main.py report.pdf --no-validation    # 禁用数据验证
        """
    )
    parser.add_argument(
        'report_file',
        help='环评报告文件路径 (支持PDF、TXT等格式)'
    )
    parser.add_argument(
        '-o', '--output',
        default='output',
        help='输出目录（默认：output）'
    )
    parser.add_argument(
        '--no-validation',
        action='store_true',
        help='禁用数据验证'
    )

    args = parser.parse_args()

    # 检查文件是否存在
    if not os.path.exists(args.report_file):
        print(f"错误: 文件不存在 - {args.report_file}")
        return 1

    # 处理报告
    result = process_ep_report_enhanced(
        args.report_file,
        args.output,
        enable_validation=not args.no_validation
    )

    return 0 if result['success'] else 1


if __name__ == "__main__":
    sys.exit(main())

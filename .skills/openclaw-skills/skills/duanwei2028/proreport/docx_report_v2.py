"""Word 报告生成器 - 对标 1111报告风格"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .config import COLORS


_PRIMARY_RGB = RGBColor(0x1a, 0x44, 0x80)
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN_RGB = RGBColor(0x2e, 0x7d, 0x32)
_RED_RGB = RGBColor(0xc6, 0x28, 0x28)
_GREY_RGB = RGBColor(0x55, 0x55, 0x55)


def _set_cell_bg(cell, hex_color: str):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_table(doc, headers, rows, col_widths=None):
    """添加风格化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _WHITE_RGB
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, COLORS.HEADER_BG)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 交替行背景
            if row_idx % 2 == 1:
                _set_cell_bg(cell, COLORS.ROW_ALT)

            # 特殊着色
            val_str = str(val)
            if "%" in val_str:
                try:
                    num = float(val_str.replace("%", ""))
                    if num > 0:
                        run.font.color.rgb = _GREEN_RGB
                    elif num < 0:
                        run.font.color.rgb = _RED_RGB
                except ValueError:
                    pass

    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def _add_kv_table(doc, kv_pairs):
    """添加键值信息表 (2列 key-value 并排)"""
    ncols = 4
    nrows = (len(kv_pairs) + 1) // 2
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (k, v) in enumerate(kv_pairs):
        row_idx = i // 2
        col_offset = (i % 2) * 2

        # Key cell
        kcell = table.rows[row_idx].cells[col_offset]
        kcell.text = ""
        run = kcell.paragraphs[0].add_run(k)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _WHITE_RGB
        _set_cell_bg(kcell, COLORS.HEADER_BG)

        # Value cell
        vcell = table.rows[row_idx].cells[col_offset + 1]
        vcell.text = ""
        run = vcell.paragraphs[0].add_run(str(v))
        run.font.size = Pt(9)

    return table


def _safe_float(val, default=0.0):
    """Safely convert a value to float."""
    if val is None:
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _fmt_pct(val, fmt=".2%"):
    """Format a value as percentage string. Accepts float or already-formatted string."""
    if isinstance(val, str):
        return val
    return f"{_safe_float(val):{fmt}}"


def _fmt_float(val, fmt=".2f"):
    """Format a value as float string. Accepts float or already-formatted string."""
    if isinstance(val, str):
        return val
    return f"{_safe_float(val):{fmt}}"


def _safe_text(val) -> str:
    """将任意类型安全转为字符串。列表用换行连接，None返回空串。"""
    if val is None:
        return ""
    if isinstance(val, list):
        return "\n\n".join(str(item) for item in val if item)
    return str(val)


def generate_docx(analysis: dict, content: dict, output_path: str = "output/report.docx"):
    """生成 Word 报告

    Args:
        analysis: run_analysis.py 输出的 JSON dict，包含 core_metrics, charts,
                  yearly, drawdown_events, variety_metrics, sector_metrics 等。
        content: AI 生成的文本内容 dict，包含 executive_summary, suggestions 等。
        output_path: 输出文件路径。
    """
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc = Document()

    core = analysis.get("core_metrics", {})
    charts = analysis.get("charts", {})
    meta = analysis.get("meta", {})
    report_date = datetime.now().strftime("%Y年%m月%d日")

    start_date = meta.get("start_date", core.get("start_date", ""))
    end_date = meta.get("end_date", core.get("end_date", ""))

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ==================== 第一部分：策略表现概述 ====================
    doc.add_heading("第一部分：策略表现概述", level=1)
    doc.add_paragraph("")

    # 概述文字
    exec_text = _safe_text(content.get("executive_summary"))
    for line in exec_text.split("\n\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    # ==================== 第二部分：整体业绩 ====================
    doc.add_heading("第二部分：整体业绩", level=1)

    # 2.1 净值图表
    doc.add_heading("2.1 净值图表", level=2)
    if "nav_drawdown" in charts:
        doc.add_picture(charts["nav_drawdown"], width=Inches(6.0))
        _add_caption(doc, f"图2-1 策略日净值曲线及回撤走势（{start_date}—{end_date}）")

    cumulative_return = _safe_float(core.get("cumulative_return", 0))
    max_drawdown = _safe_float(core.get("max_drawdown", 0))
    doc.add_paragraph(
        f"产品在【{start_date} 至 {end_date}】期间实现累计涨幅"
        f"【{cumulative_return:.2%}】，最大回撤【{max_drawdown:.2%}】。"
    )

    # 日净值分析
    nav_daily_text = _safe_text(content.get("nav_daily_analysis"))
    if nav_daily_text:
        doc.add_paragraph(nav_daily_text)

    # 2.2 周、月表现
    doc.add_heading("2.2 周、月表现", level=2)
    if "weekly_nav" in charts:
        doc.add_picture(charts["weekly_nav"], width=Inches(6.0))
        _add_caption(doc, "图2-2 周净值曲线及周收益率分布")

    nav_weekly_text = _safe_text(content.get("nav_weekly_analysis"))
    if nav_weekly_text:
        doc.add_paragraph(nav_weekly_text)

    if "monthly_nav" in charts:
        doc.add_picture(charts["monthly_nav"], width=Inches(6.0))
        _add_caption(doc, "图2-3 月净值曲线及月收益率分布")

    nav_monthly_text = _safe_text(content.get("nav_monthly_analysis"))
    if nav_monthly_text:
        doc.add_paragraph(nav_monthly_text)

    # 2.3 回撤
    doc.add_heading("2.3 回撤", level=2)
    if "drawdown_risk" in charts:
        doc.add_picture(charts["drawdown_risk"], width=Inches(6.0))
        _add_caption(doc, "图2-4 净值回撤走势与风险分级")

    drawdown_text = _safe_text(content.get("drawdown_analysis"))
    if drawdown_text:
        doc.add_paragraph(drawdown_text)

    drawdown_events = analysis.get("drawdown_events", [])
    if drawdown_events:
        dd_headers = ["开始日期", "结束日期", "持续(天)", "最大深度", "风险级别", "处置建议"]
        dd_rows = [
            [e.get("start", ""), e.get("end", ""), str(e.get("duration", "")),
             _fmt_pct(e.get("depth", 0)), e.get("level", ""), e.get("suggestion", "")]
            for e in drawdown_events
        ]
        _add_styled_table(doc, dd_headers, dd_rows, col_widths=[2.5, 2.5, 1.5, 2, 1.5, 4])

    # 2.4 年度表现
    doc.add_heading("2.4 年度表现", level=2)
    if "yearly_bar" in charts:
        doc.add_picture(charts["yearly_bar"], width=Inches(6.0))
        _add_caption(doc, "图2-5 年度净值走势与年度收益率对比")

    yearly = analysis.get("yearly", [])
    if yearly:
        yearly_headers = ["年份", "年度收益率", "年化波动率", "夏普比率", "最大回撤", "综合评级"]
        yearly_rows = [
            [str(y.get("year", "")), _fmt_pct(y.get("return", 0)),
             _fmt_pct(y.get("volatility", 0)), _fmt_float(y.get("sharpe", 0)),
             _fmt_pct(y.get("max_drawdown", 0)), y.get("rating", "")]
            for y in yearly
        ]
        _add_styled_table(doc, yearly_headers, yearly_rows, col_widths=[1.5, 2.5, 2.5, 2, 2.5, 3])

    nav_yearly_text = _safe_text(content.get("nav_yearly_analysis"))
    if nav_yearly_text:
        doc.add_paragraph(nav_yearly_text)

    # ==================== 第三部分：板块及品种分析 ====================
    variety_metrics = analysis.get("variety_metrics", [])
    if variety_metrics:
        doc.add_heading("第三部分：板块及品种分析", level=1)

        doc.add_heading("3.1 板块分析", level=2)
        if "sector_bar" in charts:
            doc.add_picture(charts["sector_bar"], width=Inches(6.0))
            _add_caption(doc, "图3-1 板块收益与胜率-夏普散点图")

        sector_text = _safe_text(content.get("sector_analysis"))
        if sector_text:
            doc.add_paragraph(sector_text)

        sector_metrics = analysis.get("sector_metrics", [])
        if sector_metrics:
            sec_headers = ["板块", "品种数", "平均年化", "累计收益", "平均夏普", "平均胜率", "平均盈亏比"]
            sec_rows = [
                [s.get("name", ""), str(len(s.get("varieties", []))),
                 _fmt_pct(s.get("avg_return", 0)),
                 _fmt_pct(s.get("total_return", 0)),
                 _fmt_float(s.get("avg_sharpe", 0)),
                 _fmt_pct(s.get("avg_win_rate", 0), ".1%"),
                 _fmt_float(s.get("avg_pnl_ratio", 0))]
                for s in sector_metrics
            ]
            _add_styled_table(doc, sec_headers, sec_rows,
                            col_widths=[2, 1.5, 2, 2, 2, 2, 2])

        # 3.2 品种分析
        doc.add_heading("3.2 品种盈亏排名", level=2)
        if "variety_top" in charts:
            doc.add_picture(charts["variety_top"], width=Inches(6.0))
            _add_caption(doc, "图3-2 盈利/亏损 TOP 10 品种")

        variety_text = _safe_text(content.get("variety_analysis"))
        if variety_text:
            doc.add_paragraph(variety_text)

        # 品种明细表（TOP 20）
        doc.add_heading("3.3 品种明细（TOP 20）", level=2)
        top_varieties = variety_metrics[:20]
        var_headers = ["品种", "板块", "累计收益", "年化收益", "夏普", "胜率", "盈亏比", "最大回撤"]
        var_rows = [
            [v.get("name", ""), v.get("sector", ""),
             _fmt_pct(v.get("cum_return", 0)),
             _fmt_pct(v.get("ann_return", 0)),
             _fmt_float(v.get("sharpe", 0)),
             _fmt_pct(v.get("win_rate", 0), ".1%"),
             _fmt_float(v.get("pnl_ratio", 0)),
             _fmt_pct(v.get("max_drawdown", 0))]
            for v in top_varieties
        ]
        _add_styled_table(doc, var_headers, var_rows,
                        col_widths=[1.5, 1.5, 2, 2, 1.5, 1.5, 1.5, 2])

    # ==================== 第四部分：特征指标 ====================
    doc.add_heading("第四部分：特征指标", level=1)

    doc.add_heading("4.1 核心评估指标", level=2)

    annualized_return = _safe_float(core.get("annualized_return", 0))
    annualized_volatility = _safe_float(core.get("annualized_volatility", 0))
    sharpe_ratio = _safe_float(core.get("sharpe_ratio", 0))
    calmar_ratio = _safe_float(core.get("calmar_ratio", 0))
    sortino_ratio = _safe_float(core.get("sortino_ratio", 0))
    daily_win_rate = _safe_float(core.get("daily_win_rate", 0))
    daily_pnl_ratio = _safe_float(core.get("daily_pnl_ratio", 0))
    var_95 = _safe_float(core.get("var_95", 0))

    ie = content.get("indicator_evaluations", {})
    indicators_headers = ["指标", "策略值", "行业基准", "评价说明"]
    indicators_rows = [
        ["年化收益率", f"{annualized_return:.2%}", ">=15%", ie.get("年化收益率", "-")],
        ["年化波动率", f"{annualized_volatility:.2%}", "<=20%", ie.get("年化波动率", "-")],
        ["最大回撤", f"{max_drawdown:.2%}", "<=-20%", ie.get("最大回撤", "-")],
        ["夏普比率", f"{sharpe_ratio:.2f}", ">=1.0", ie.get("夏普比率", "-")],
        ["卡玛比率", f"{calmar_ratio:.2f}", ">=1.5", ie.get("卡玛比率", "-")],
        ["索提诺比率", f"{sortino_ratio:.2f}", ">=1.5", ie.get("索提诺比率", "-")],
        ["日胜率", f"{daily_win_rate:.1%}", ">=50%", ie.get("日胜率", "-")],
        ["日盈亏比", f"{daily_pnl_ratio:.2f}", ">=1.5", ie.get("日盈亏比", "-")],
        ["日VaR(95%)", f"{var_95:.2%}", ">-3%", ie.get("日VaR(95%)", "-")],
    ]
    _add_styled_table(doc, indicators_headers, indicators_rows,
                     col_widths=[2, 1.5, 1.5, 8.5])

    indicator_text = _safe_text(content.get("indicator_analysis"))
    if indicator_text:
        doc.add_paragraph(indicator_text)

    # 4.2 雷达图
    doc.add_heading("4.2 综合评分雷达", level=2)
    if "radar" in charts:
        doc.add_picture(charts["radar"], width=Inches(4.0))
        _add_caption(doc, "图4-1 策略综合评分雷达图")

    # 4.3 月度热力图
    doc.add_heading("4.3 月度收益率热力图", level=2)
    if "heatmap" in charts:
        doc.add_picture(charts["heatmap"], width=Inches(6.0))
        _add_caption(doc, "图4-2 月度收益率热力图")

    heatmap_text = _safe_text(content.get("heatmap_analysis"))
    if heatmap_text:
        doc.add_paragraph(heatmap_text)

    # 4.4 滚动夏普
    doc.add_heading("4.4 滚动夏普比率", level=2)
    if "rolling_sharpe" in charts:
        doc.add_picture(charts["rolling_sharpe"], width=Inches(6.0))
        _add_caption(doc, "图4-3 滚动夏普比率走势")

    rolling_sharpe_text = _safe_text(content.get("rolling_sharpe_analysis"))
    if rolling_sharpe_text:
        doc.add_paragraph(rolling_sharpe_text)

    # 4.5 每万元收益
    doc.add_heading("4.5 资金效率", level=2)
    if "per_10k" in charts:
        doc.add_picture(charts["per_10k"], width=Inches(6.0))
        _add_caption(doc, "图4-4 每万元每日收益率")

    per_10k_text = _safe_text(content.get("per_10k_analysis"))
    if per_10k_text:
        doc.add_paragraph(per_10k_text)

    # 4.6 收益分布
    doc.add_heading("4.6 日收益率分布", level=2)
    if "return_distribution" in charts:
        doc.add_picture(charts["return_distribution"], width=Inches(5.5))
        _add_caption(doc, "图4-5 日收益率分布")

    # 4.7 CTA 对比 (if present)
    cta_text = _safe_text(content.get("cta_comparison"))
    if cta_text:
        doc.add_heading("4.7 CTA 行业对比", level=2)
        if "cta_comparison" in charts:
            doc.add_picture(charts["cta_comparison"], width=Inches(6.0))
            _add_caption(doc, "图4-6 策略 vs CTA 行业对比")
        doc.add_paragraph(cta_text)

    # ==================== 第五部分：总结及优化建议 ====================
    doc.add_heading("第五部分：总结及优化建议", level=1)

    doc.add_heading("5.1 总结", level=2)
    conclusion_text = _safe_text(content.get("conclusion"))
    if conclusion_text:
        doc.add_paragraph(conclusion_text)
    else:
        # Fallback: generate a basic summary from core metrics
        years = _safe_float(core.get("years", 0))
        summary_text = (
            f"收益能力：策略实现累计收益{cumulative_return:.2%}，年化{annualized_return:.2%}。"
            f"风险控制：最大回撤{max_drawdown:.2%}，夏普比率{sharpe_ratio:.2f}。"
            f"策略稳定性：在{years:.0f}年运行期间经历完整市场周期，"
            f"日胜率{daily_win_rate:.1%}，盈亏比{daily_pnl_ratio:.2f}。"
        )
        doc.add_paragraph(summary_text)

    # ==================== 第六部分：策略改进建议 ====================
    doc.add_heading("第六部分：策略改进建议", level=1)
    doc.add_paragraph("基于上述全面分析，我们从多个维度提出策略优化改进建议：")

    suggestions = content.get("suggestions", [])
    if isinstance(suggestions, str):
        suggestions = [s.strip() for s in suggestions.split("\n") if s.strip()]
    for i, sug in enumerate(suggestions, 1):
        if not isinstance(sug, str) or len(sug) < 5:
            continue
        title = f"6.{i} 建议{i}"
        body = sug
        for sep in ["——", "：", ":"]:
            if sep in sug:
                parts = sug.split(sep, 1)
                title = f"6.{i} {parts[0].strip()}"
                body = parts[1].strip()
                break
        doc.add_heading(title, level=2)
        # 长文本按段落分割渲染
        for para in body.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    # ==================== 保存 ====================
    doc.save(output_path)
    return output_path


def _add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = _GREY_RGB

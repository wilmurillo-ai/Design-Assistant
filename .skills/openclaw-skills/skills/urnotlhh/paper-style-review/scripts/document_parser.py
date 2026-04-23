#!/usr/bin/env python3
"""Shared thesis document parsing runtime.

Single responsibility:
- parse supported paper files into one normalized chapter/paragraph structure;
- expose the shared parsing helpers used by structure mapping and review stages;
- keep document parsing independent from any style/AI-specific pipeline naming.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from statistics import median
from typing import Any, Dict, List, Optional, Tuple

import fitz
from docx import Document


CHAPTER_RE = re.compile(r"^第\s*([一二三四五六七八九十百千万0-9]+)\s*章\s*(.*)$")
LEVEL3_RE = re.compile(r"^(\d+\.\d+\.\d+)\s*(.*)$")
LEVEL2_RE = re.compile(r"^(\d+\.\d+)\s*(.*)$")
BACKMATTER_RE = re.compile(r"^(参考文献|致谢|附录|攻读.*期间|作者简介|发表.*成果|在学期间)")
TOC_RE = re.compile(r"目录|目\s*录", re.I)
DOTTED_LINE_RE = re.compile(r"\.{4,}|·{4,}|…{3,}")
FIGURE_RE = re.compile(r"^(图|表|Fig\.?|Table)\s*\d", re.I)
PAGE_NUM_RE = re.compile(r"^[IVXLCM]+$|^\d+$")
DECORATED_PAGE_NUM_RE = re.compile(r"^[\-–—_~·•=]*\s*(?:第\s*)?\d+\s*(?:页)?\s*[\-–—_~·•=]*$", re.I)
AUTHOR_ZH_RE = re.compile(r"作\s*者\s*姓\s*名\s*[：:]?\s*([\u4e00-\u9fff]{2,8})")
AUTHOR_EN_RE = re.compile(r"\bBy\s+([A-Za-z][A-Za-z .\-]{1,80}?)(?=\s+Supervisor:|\n|$)")
HEADER_FOOTER_RE = re.compile(r"东北大学.*学位论文|NORTHEASTERN\s+UNIVERSITY|THESIS\s+FOR\s+MASTER|页码|目\s*录", re.I)
PAGE_WORD_RE = re.compile(r"^page\s+\d+$", re.I)
SENTENCE_END_RE = re.compile(r"[。！？!?；;]$")
HARD_BREAK_RE = re.compile(r"[。！？!?；;：:]$")
ASCII_WORD_RE = re.compile(r"[A-Za-z0-9]")
REPEATED_MARGIN_MIN_PAGES = 2

STYLE_LEVEL_HINTS = [
    (1, ["heading 1", "标题 1", "标题1", "标题样式1", "title 1"]),
    (2, ["heading 2", "标题 2", "标题2", "标题样式2", "title 2"]),
    (3, ["heading 3", "标题 3", "标题3", "标题样式3", "title 3"]),
]

NOISE_HEADINGS = {
    "独创性声明",
    "摘 要",
    "摘要",
    "abstract",
    "学位论文版权使用授权书",
}


@dataclass
class Block:
    text: str
    heading_level: Optional[int] = None
    style_name: str = ""
    page: Optional[int] = None
    page_has_toc: bool = False
    x0: Optional[float] = None
    y0: Optional[float] = None
    x1: Optional[float] = None
    y1: Optional[float] = None


@dataclass
class Item:
    type: str
    text: str
    section_path: str
    level: Optional[int] = None


@dataclass
class ParsedDocument:
    path: str
    title: str
    author: Optional[str]
    blocks: List[Block]
    chapters: List[Dict]
    extractionRisks: List[str] = field(default_factory=list)
    extractionMeta: Dict[str, Any] = field(default_factory=dict)


@dataclass
class PdfLine:
    text: str
    page: int
    page_has_toc: bool
    page_width: float
    page_height: float
    x0: float
    y0: float
    x1: float
    y1: float


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def normalize_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    lines = []
    for raw in text.splitlines():
        raw = raw.strip()
        if not raw:
            continue
        raw = re.sub(r"\s+", " ", raw)
        lines.append(raw)
    return "\n".join(lines).strip()

def detect_heading_level(text: str, style_name: str = "") -> Optional[int]:
    style_lower = (style_name or "").strip().lower()
    for level, hints in STYLE_LEVEL_HINTS:
        if any(h in style_lower for h in hints):
            return level
    compact = compact_text(text)
    if CHAPTER_RE.match(compact):
        return 1
    if LEVEL3_RE.match(compact):
        return 3
    if LEVEL2_RE.match(compact):
        return 2
    return None


def is_noise_paragraph(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    compact = compact_text(stripped)
    if compact in NOISE_HEADINGS:
        return True
    if PAGE_NUM_RE.match(compact):
        return True
    if DECORATED_PAGE_NUM_RE.match(stripped):
        return True
    if PAGE_WORD_RE.match(stripped):
        return True
    if HEADER_FOOTER_RE.search(stripped):
        return True
    if FIGURE_RE.match(stripped):
        return True
    if stripped.startswith("Fig.") or stripped.startswith("Table"):
        return True
    if stripped.startswith("关键词") or stripped.startswith("Key words"):
        return True
    if DOTTED_LINE_RE.search(stripped) and len(stripped) < 120:
        return True
    if len(compact) <= 2 and not re.search(r"[\u4e00-\u9fffA-Za-z]", compact):
        return True
    return False


def parse_docx_blocks(path: Path) -> List[Block]:
    doc = Document(str(path))
    blocks = []
    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        blocks.append(Block(text=text, heading_level=detect_heading_level(text, style_name), style_name=style_name))
    return blocks


def _median(values: List[float], default: float) -> float:
    cleaned = [value for value in values if value > 0]
    if not cleaned:
        return default
    return float(median(cleaned))


def _line_height(line: PdfLine) -> float:
    return max(1.0, float(line.y1) - float(line.y0))


def _line_text_from_spans(spans: List[Dict[str, Any]]) -> str:
    fragments: List[str] = []
    for span in spans:
        text = normalize_text(span.get("text") or "")
        if text:
            fragments.append(text)
    return normalize_text("".join(fragments))


def _extract_pdf_lines(path: Path) -> Tuple[List[PdfLine], Dict[str, Any]]:
    doc = fitz.open(str(path))
    lines: List[PdfLine] = []
    block_count = 0
    for page_index, page in enumerate(doc, start=1):
        page_text = normalize_text(page.get_text("text"))
        page_has_toc = bool(TOC_RE.search(page_text) and DOTTED_LINE_RE.search(page_text))
        payload = page.get_text("dict", sort=True)
        page_width = float(page.rect.width or 0.0)
        page_height = float(page.rect.height or 0.0)
        for block in payload.get("blocks", []):
            if int(block.get("type", 0) or 0) != 0:
                continue
            block_count += 1
            for row in block.get("lines", []):
                text = _line_text_from_spans(row.get("spans", []))
                if not text:
                    continue
                bbox = row.get("bbox") or block.get("bbox") or (0.0, 0.0, 0.0, 0.0)
                x0, y0, x1, y1 = [float(value or 0.0) for value in bbox[:4]]
                lines.append(
                    PdfLine(
                        text=text,
                        page=page_index,
                        page_has_toc=page_has_toc,
                        page_width=page_width,
                        page_height=page_height,
                        x0=x0,
                        y0=y0,
                        x1=x1,
                        y1=y1,
                    )
                )
    return lines, {"rawPdfBlockCount": block_count, "rawPdfLineCount": len(lines)}


def _margin_bucket(line: PdfLine) -> Optional[str]:
    if line.page_height <= 0:
        return None
    if line.y0 <= line.page_height * 0.12:
        return "top"
    if line.y1 >= line.page_height * 0.88:
        return "bottom"
    return None


def _meaningful_margin_text(text: str) -> bool:
    stripped = text.strip()
    compact = compact_text(stripped)
    if not stripped or len(compact) < 3:
        return False
    if DECORATED_PAGE_NUM_RE.match(stripped) or PAGE_NUM_RE.match(compact) or PAGE_WORD_RE.match(stripped):
        return False
    if len(compact) > 80:
        return False
    return bool(re.search(r"[\u4e00-\u9fffA-Za-z]", stripped))


def _find_repeated_margin_texts(lines: List[PdfLine]) -> set[str]:
    hits: Dict[Tuple[str, str], set[int]] = {}
    for line in lines:
        bucket = _margin_bucket(line)
        if not bucket or not _meaningful_margin_text(line.text):
            continue
        key = (bucket, compact_text(line.text))
        hits.setdefault(key, set()).add(line.page)
    return {
        text
        for (_bucket, text), pages in hits.items()
        if len(pages) >= REPEATED_MARGIN_MIN_PAGES
    }


def _looks_like_pdf_page_artifact(text: str) -> bool:
    stripped = text.strip()
    compact = compact_text(stripped)
    if not stripped:
        return False
    return bool(
        PAGE_NUM_RE.match(compact)
        or DECORATED_PAGE_NUM_RE.match(stripped)
        or PAGE_WORD_RE.match(stripped)
    )


def _is_heading_like_line(text: str) -> bool:
    return bool(detect_heading_level(text))


def _should_keep_as_standalone_pdf_line(text: str, *, page_has_toc: bool) -> bool:
    if page_has_toc:
        return True
    if _is_heading_like_line(text):
        return True
    if FIGURE_RE.match(text):
        return True
    if DOTTED_LINE_RE.search(text):
        return True
    return False


def _same_pdf_column(prev: PdfLine, current: PdfLine) -> bool:
    width_tol = max(_line_height(prev), _line_height(current)) * 1.2
    if abs(current.x0 - prev.x0) <= max(12.0, width_tol):
        return True
    overlap = max(0.0, min(prev.x1, current.x1) - max(prev.x0, current.x0))
    narrower = max(1.0, min(prev.x1 - prev.x0, current.x1 - current.x0))
    return overlap >= narrower * 0.55


def _looks_like_new_paragraph_start(prev: PdfLine, current: PdfLine, typical_gap: float) -> bool:
    if _should_keep_as_standalone_pdf_line(current.text, page_has_toc=current.page_has_toc):
        return True
    if _looks_like_pdf_page_artifact(current.text):
        return True
    if prev.page == current.page:
        gap = max(0.0, current.y0 - prev.y1)
        if gap > max(typical_gap * 1.8, max(_line_height(prev), _line_height(current)) * 1.6):
            return True
    if current.x0 - prev.x0 > max(12.0, max(_line_height(prev), _line_height(current)) * 0.8):
        return True
    if HARD_BREAK_RE.search(prev.text.strip()):
        return True
    return False


def _merge_pdf_text(prev_text: str, current_text: str) -> str:
    left = prev_text.rstrip()
    right = current_text.lstrip()
    if not left:
        return right
    if not right:
        return left
    if left.endswith("-") and ASCII_WORD_RE.match(left[-2:-1] or "") and ASCII_WORD_RE.match(right[:1]):
        return left[:-1] + right
    if ASCII_WORD_RE.match(left[-1]) and ASCII_WORD_RE.match(right[:1]):
        return left + " " + right
    if left[-1] in ",.;:)]}" and ASCII_WORD_RE.match(right[:1]):
        return left + " " + right
    return left + right


def _flush_pdf_buffer(buffer: List[PdfLine], output: List[Block]) -> None:
    if not buffer:
        return
    text = ""
    for line in buffer:
        text = _merge_pdf_text(text, line.text)
    first = buffer[0]
    output.append(
        Block(
            text=text,
            heading_level=detect_heading_level(text),
            style_name="",
            page=first.page,
            page_has_toc=any(line.page_has_toc for line in buffer),
            x0=min(line.x0 for line in buffer),
            y0=min(line.y0 for line in buffer),
            x1=max(line.x1 for line in buffer),
            y1=max(line.y1 for line in buffer),
        )
    )
    buffer.clear()


def _normalize_pdf_lines(lines: List[PdfLine]) -> Tuple[List[Block], Dict[str, Any]]:
    repeated_margin_texts = _find_repeated_margin_texts(lines)
    filtered: List[PdfLine] = []
    removed_header_footer = 0
    removed_page_artifacts = 0
    for line in lines:
        compact = compact_text(line.text)
        if compact and compact in repeated_margin_texts:
            removed_header_footer += 1
            continue
        if _looks_like_pdf_page_artifact(line.text):
            removed_page_artifacts += 1
            continue
        filtered.append(line)

    line_gaps = [
        max(0.0, current.y0 - previous.y1)
        for previous, current in zip(filtered, filtered[1:])
        if previous.page == current.page
    ]
    typical_gap = _median(line_gaps, default=8.0)
    paragraphs: List[Block] = []
    buffer: List[PdfLine] = []
    for line in filtered:
        if _should_keep_as_standalone_pdf_line(line.text, page_has_toc=line.page_has_toc):
            _flush_pdf_buffer(buffer, paragraphs)
            paragraphs.append(
                Block(
                    text=line.text,
                    heading_level=detect_heading_level(line.text),
                    style_name="",
                    page=line.page,
                    page_has_toc=line.page_has_toc,
                    x0=line.x0,
                    y0=line.y0,
                    x1=line.x1,
                    y1=line.y1,
                )
            )
            continue
        if not buffer:
            buffer.append(line)
            continue
        previous = buffer[-1]
        same_column = _same_pdf_column(previous, line)
        cross_page_continuation = (
            previous.page != line.page
            and same_column
            and not SENTENCE_END_RE.search(previous.text.strip())
            and not _should_keep_as_standalone_pdf_line(line.text, page_has_toc=line.page_has_toc)
        )
        if cross_page_continuation:
            buffer.append(line)
            continue
        if not same_column or _looks_like_new_paragraph_start(previous, line, typical_gap):
            _flush_pdf_buffer(buffer, paragraphs)
        buffer.append(line)
    _flush_pdf_buffer(buffer, paragraphs)

    short_fragments = 0
    for block in paragraphs:
        compact = compact_text(block.text)
        if not compact:
            continue
        if block.heading_level:
            continue
        if len(compact) <= 20 and not SENTENCE_END_RE.search(block.text):
            short_fragments += 1

    return paragraphs, {
        "rawPdfLineCount": len(lines),
        "normalizedPdfParagraphCount": len(paragraphs),
        "removedHeaderFooterCount": removed_header_footer,
        "removedPageArtifactCount": removed_page_artifacts,
        "repeatedMarginTextCount": len(repeated_margin_texts),
        "suspiciousShortFragmentCount": short_fragments,
    }


def parse_pdf_blocks(path: Path) -> Tuple[List[Block], Dict[str, Any]]:
    lines, raw_meta = _extract_pdf_lines(path)
    blocks, norm_meta = _normalize_pdf_lines(lines)
    meta = {
        "sourceType": "pdf",
        "pdfNormalization": {
            **raw_meta,
            **norm_meta,
        },
    }
    return blocks, meta


def parse_markdown_blocks(path: Path) -> List[Block]:
    blocks: List[Block] = []
    text = path.read_text(encoding="utf-8")
    current: List[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            if current:
                piece = normalize_text("\n".join(current))
                if piece:
                    if piece.startswith("### "):
                        blocks.append(Block(text=piece[4:].strip(), heading_level=3))
                    elif piece.startswith("## "):
                        blocks.append(Block(text=piece[3:].strip(), heading_level=2))
                    elif piece.startswith("# "):
                        blocks.append(Block(text=piece[2:].strip(), heading_level=1))
                    else:
                        blocks.append(Block(text=piece, heading_level=detect_heading_level(piece)))
                current = []
            continue
        current.append(line)
    if current:
        piece = normalize_text("\n".join(current))
        blocks.append(Block(text=piece, heading_level=detect_heading_level(piece)))
    return blocks


def parse_blocks_with_meta(path: Path) -> Tuple[List[Block], Dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return parse_docx_blocks(path), {"sourceType": "docx"}
    if suffix == ".pdf":
        return parse_pdf_blocks(path)
    if suffix in {".md", ".txt"}:
        return parse_markdown_blocks(path), {"sourceType": suffix.lstrip(".")}
    raise ValueError(f"unsupported file type: {path}")


def parse_blocks(path: Path) -> List[Block]:
    blocks, _meta = parse_blocks_with_meta(path)
    return blocks


def detect_author(snippet: str) -> Optional[str]:
    compact = snippet.replace("\n", " ")
    match = AUTHOR_ZH_RE.search(compact)
    if match:
        return match.group(1).strip()
    match = AUTHOR_EN_RE.search(compact)
    if match:
        return re.sub(r"\s+", " ", match.group(1)).strip()
    return None


def detect_title(snippet: str, fallback: str) -> str:
    for line in snippet.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        compact = compact_text(stripped)
        if compact in {"学位论文", "硕士学位论文", "THESISFORMASTER'SDEGREE"}:
            continue
        if len(compact) < 8:
            continue
        if re.match(r"^(分类号|UDC|NORTHEASTERN|AThesis|By |Supervisor|独创性声明|摘要|Abstract)", stripped, re.I):
            continue
        if re.search(r"[\u4e00-\u9fff]{6,}", stripped):
            return stripped
    return fallback


def find_body_bounds(blocks: List[Block]) -> Tuple[int, int]:
    start = None
    for idx, block in enumerate(blocks):
        if block.heading_level == 1 and CHAPTER_RE.match(compact_text(block.text)) and not block.page_has_toc and not DOTTED_LINE_RE.search(block.text):
            start = idx
            break
    if start is None:
        raise ValueError("failed to locate first chapter heading")
    end = len(blocks)
    for idx in range(start + 1, len(blocks)):
        block = blocks[idx]
        if block.heading_level and BACKMATTER_RE.match(compact_text(block.text)):
            end = idx
            break
    return start, end


def build_chapters(blocks: List[Block]) -> List[Dict]:
    start, end = find_body_bounds(blocks)
    body_blocks = blocks[start:end]
    chapters: List[Dict] = []
    current = None
    current_level2 = None
    current_level3 = None

    def current_path() -> str:
        path_parts = []
        if current:
            path_parts.append(current["title"])
        if current_level2:
            path_parts.append(current_level2)
        if current_level3:
            path_parts.append(current_level3)
        return " > ".join(path_parts)

    for block in body_blocks:
        text = block.text.strip()
        if not text:
            continue
        if block.heading_level == 1 and CHAPTER_RE.match(compact_text(text)):
            if current is not None and compact_text(current["title"]) == compact_text(text):
                continue
            current = {"title": text, "items": [], "blocks": []}
            chapters.append(current)
            current_level2 = None
            current_level3 = None
            current["items"].append(Item(type="heading", text=text, section_path=text, level=1))
            current["blocks"].append(block)
            continue
        if current is None:
            continue
        current["blocks"].append(block)
        if block.heading_level == 2:
            current_level2 = text
            current_level3 = None
            current["items"].append(Item(type="heading", text=text, section_path=current_path(), level=2))
            continue
        if block.heading_level == 3:
            current_level3 = text
            current["items"].append(Item(type="heading", text=text, section_path=current_path(), level=3))
            continue
        if is_noise_paragraph(text):
            continue
        path = current_path()
        if current["items"] and current["items"][-1].type == "para" and current["items"][-1].section_path == path:
            prev = current["items"][-1].text
            if prev and prev[-1] not in "。！？；：":
                current["items"][-1].text = prev + text
                continue
        current["items"].append(Item(type="para", text=text, section_path=path))
    return chapters


def parse_document(path: Path) -> ParsedDocument:
    blocks, extraction_meta = parse_blocks_with_meta(path)
    full_snippet = "\n".join(block.text for block in blocks[:80])
    title = detect_title(full_snippet, path.stem)
    author = detect_author(full_snippet)
    chapters = build_chapters(blocks)
    risks: List[str] = []
    if not chapters:
        risks.append("未能抽取到正文章节结构")

    if path.suffix.lower() == ".pdf":
        pdf_meta = dict((extraction_meta or {}).get("pdfNormalization") or {})
        section_counts: List[Tuple[str, int]] = []
        for chapter in chapters:
            counts: Dict[str, int] = {}
            for item in chapter.get("items", []) or []:
                item_type = getattr(item, "type", None) or (item.get("type") if isinstance(item, dict) else None)
                if item_type == "para":
                    section_path = (
                        getattr(item, "section_path", None)
                        or getattr(item, "sectionPath", None)
                        or (item.get("section_path") if isinstance(item, dict) else None)
                        or (item.get("sectionPath") if isinstance(item, dict) else None)
                        or chapter.get("title", "")
                    )
                    counts[section_path] = counts.get(section_path, 0) + 1
            section_counts.extend(counts.items())
        max_section_path, max_section_count = ("", 0)
        if section_counts:
            max_section_path, max_section_count = max(section_counts, key=lambda item: item[1])
        short_para_count = 0
        body_para_count = 0
        for chapter in chapters:
            for item in chapter.get("items", []) or []:
                item_type = getattr(item, "type", None) or (item.get("type") if isinstance(item, dict) else None)
                if item_type == "para":
                    body_para_count += 1
                    text = (
                        getattr(item, "text", None)
                        or (item.get("text", "") if isinstance(item, dict) else "")
                        or ""
                    ).strip()
                    if len(compact_text(text)) <= 20:
                        short_para_count += 1
        pdf_meta["maxParagraphsPerSection"] = max_section_count
        pdf_meta["maxParagraphSectionPath"] = max_section_path
        pdf_meta["bodyParagraphCount"] = body_para_count
        pdf_meta["shortBodyParagraphCount"] = short_para_count
        extraction_meta["pdfNormalization"] = pdf_meta
        if max_section_count >= 60:
            risks.append(f"PDF 段落切分仍偏碎：小节「{max_section_path}」下检测到 {max_section_count} 段正文")
        elif max_section_count >= 35:
            risks.append(f"PDF 段落切分存在风险：小节「{max_section_path}」下检测到 {max_section_count} 段正文")
        if short_para_count and body_para_count and short_para_count / max(body_para_count, 1) >= 0.35:
            risks.append(f"PDF 短碎段比例偏高：{short_para_count}/{body_para_count}")

    return ParsedDocument(
        path=str(path),
        title=title,
        author=author,
        blocks=blocks,
        chapters=chapters,
        extractionRisks=risks,
        extractionMeta=extraction_meta,
    )


def label_mismatch_risk(label: str, path: str, detected_author: Optional[str]) -> Optional[str]:
    compact_label = compact_text(label)
    haystacks = [compact_text(Path(path).name)]
    if detected_author:
        haystacks.append(compact_text(detected_author))
    if any(compact_label and compact_label in hay for hay in haystacks):
        return None
    if compact_label and re.search(r"[\u4e00-\u9fff]", compact_label):
        return f'ref 标签「{label}」与当前文件/作者未直接匹配（文件：{Path(path).name}，检测作者：{detected_author or "unknown"}）'
    return None

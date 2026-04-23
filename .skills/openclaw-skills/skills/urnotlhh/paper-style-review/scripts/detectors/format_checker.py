#!/usr/bin/env python3
"""DOCX 格式检查模块 v3。

唯一依据：
- references/format-rules.json
- references/format-reference-source-notes.md
- references/format-rule-matrix.md

实现原则：
1. 不再使用参考论文 Word 样式当格式基线。
2. 对原规范有明确数值/位置要求的项目，做显式规则检查。
3. 对往届问题只强调“统一/规范”的项目，做一致性检查，并在报告中标记 consistency-mode。
4. 文档级问题也必须给出可批注锚点，保证能进入 annotations.json / annotated.docx。
"""

from __future__ import annotations

import json
import re
import sys
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from zipfile import ZipFile

SCRIPTS_ROOT = Path(__file__).resolve().parents[1]
if str(SCRIPTS_ROOT) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_ROOT))

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree

from contracts import ConfidenceLevel, HumanReviewRequired, SourceType, UnifiedAnnotation
from target_document_index import (
    make_body_node_id,
    make_footer_proxy_node_id,
    make_header_proxy_node_id,
    make_section_proxy_node_id,
    make_table_cell_node_id,
)
from word_text_surface import get_paragraph_visible_text

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

ALIGNMENT_LABELS = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    "left": "left",
    "center": "center",
    "right": "right",
    "justify": "justify",
    "distribute": "distribute",
}

CN_FONT_SIZE_PT = {
    "初号": 42.0,
    "小初": 36.0,
    "一号": 26.0,
    "小一": 24.0,
    "二号": 22.0,
    "小二": 18.0,
    "三号": 16.0,
    "小三": 15.0,
    "四号": 14.0,
    "小四": 12.0,
    "五号": 10.5,
    "小五": 9.0,
    "六号": 7.5,
    "小六": 6.5,
    "七号": 5.5,
    "八号": 5.0,
}

FORMAT_RULES_PATH = Path(__file__).resolve().parents[2] / "references" / "format-rules.json"
FORMAT_MATRIX_PATH = Path(__file__).resolve().parents[2] / "references" / "format-rule-matrix.md"
FORMAT_SOURCE_NOTES_PATH = Path(__file__).resolve().parents[2] / "references" / "format-reference-source-notes.md"

CHAPTER_RE = re.compile(r"^第\s*[一二三四五六七八九十百千万0-9]+\s*章")
LEVEL4_RE = re.compile(r"^\d+\.\d+\.\d+\.\d+\s*(?=[A-Za-z\u4e00-\u9fff（(])")
LEVEL3_RE = re.compile(r"^\d+\.\d+\.\d+\s*(?=[A-Za-z\u4e00-\u9fff（(])")
LEVEL2_RE = re.compile(r"^\d+\.\d+\s*(?=[A-Za-z\u4e00-\u9fff（(])")
FIGURE_CAPTION_RE = re.compile(r"^图\s*\d+(?:\.\d+)*")
TABLE_CAPTION_RE = re.compile(r"^表\s*\d+(?:\.\d+)*")
EN_FIGURE_CAPTION_RE = re.compile(r"^Fig\.?\s*\d+(?:\.\d+)*", re.I)
EN_TABLE_CAPTION_RE = re.compile(r"^Table\s*\d+(?:\.\d+)*", re.I)
FORMULA_NO_RE = re.compile(r"[（(]\s*\d+\.\d+\s*[)）]")
TOC_DOT_RE = re.compile(r"[·\.]{4,}|…{3,}")
KEYWORDS_RE = re.compile(r"^关键词[：:]?")
EN_KEYWORDS_RE = re.compile(r"^(Key\s*words|Keywords)[：:]?", re.I)
REFERENCE_HEADING_RE = re.compile(r"^参考文献$")
ACK_HEADING_RE = re.compile(r"^致\s*谢$")
ACHIEVEMENT_HEADING_RE = re.compile(r"^攻读.*期间|^个人简历|^作者简历|^发表.*成果")
AUTHORIZATION_HEADING_RE = re.compile(r"^学位论文版权使用授权书$")
ABSTRACT_CN_RE = re.compile(r"^摘\s*要$")
ABSTRACT_EN_RE = re.compile(r"^Abstract$", re.I)
TOC_HEADING_RE = re.compile(r"^目\s*录$")
CLASSIFICATION_LINE_RE = re.compile(r"^分类号.*密级$")
UDC_RE = re.compile(r"^UDC$", re.I)
THESIS_LABEL_RE = re.compile(r"^学\s*位\s*论\s*文$")
EN_THESIS_LABEL_RE = re.compile(r"^A\s+(?:Thesis|Dissertation)\b", re.I)
CN_DATE_RE = re.compile(r"^\d{4}\s*年\s*\d{1,2}\s*月$")
EN_DATE_RE = re.compile(
    r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}$",
    re.I,
)
NUMERIC_ONLY_RE = re.compile(r"^\d+$")
REFERENCE_ENTRY_RE = re.compile(r"^(\[?\d+\]?\.?)[\s ]+")
PAGE_RANGE_RE = re.compile(r"[:：]\s*\d+\s*[-–—]\s*\d+")
VOL_ISSUE_RE = re.compile(r"\d+\s*\(\s*\d+\s*\)|（\s*\d+\s*）")
CHINESE_RE = re.compile(r"[\u4e00-\u9fff]")
LIST_LIKE_RE = re.compile(r"^[（(]?[0-9a-zA-Z一二三四五六七八九十]+[）).、]")
ZH_EN_WORD_SPACE_RE = re.compile(r"(?:(?<=[\u4e00-\u9fff])[ \u00A0\u3000]+(?=[A-Za-z0-9])|(?<=[A-Za-z0-9])[ \u00A0\u3000]+(?=[\u4e00-\u9fff]))")
HEADER_LINE_ALLOWED = {
    "single",
    "double",
    "thickThinMediumGap",
    "thinThickMediumGap",
    "thinThickThinMediumGap",
    "thickThinLargeGap",
    "thinThickLargeGap",
    "thinThickThinLargeGap",
    "thickThinSmallGap",
    "thinThickSmallGap",
    "thinThickThinSmallGap",
}


@dataclass
class ParagraphInfo:
    index: Optional[int]
    location: str
    text: str
    style_name: str
    zone: str
    paragraph_type: str
    heading_level: Optional[int]
    alignment: Optional[str]
    line_spacing: Optional[float]
    line_spacing_mode: Optional[str]
    space_before_pt: Optional[float]
    space_after_pt: Optional[float]
    first_line_indent_pt: Optional[float]
    left_indent_pt: Optional[float]
    right_indent_pt: Optional[float]
    east_asia_font: Optional[str]
    ascii_font: Optional[str]
    font_size_pt: Optional[float]
    has_drawing: bool = False
    has_equation: bool = False
    xml_paragraph_index: Optional[int] = None
    section_index: Optional[int] = None
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    column_index: Optional[int] = None
    cell_paragraph_index: Optional[int] = None
    raw_paragraph: Optional[Paragraph] = field(default=None, repr=False, compare=False)


@dataclass
class CheckResult:
    checkType: str
    status: str
    basis: str
    details: str
    location: str = "文档级"


@dataclass
class FormatIssue:
    id: str
    location: str
    paragraphIndex: Optional[int]
    styleName: str
    checkType: str
    expected: Any
    actual: Any
    sourceText: str
    problem: str
    fixDirection: str
    severity: str
    basis: str
    mode: str
    anchor: Optional[Dict[str, Any]] = None

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class FormatEvidence:
    source_text: str
    attach_text: Optional[str] = None


@dataclass
class FrontMatterModel:
    classification_line: Optional[ParagraphInfo] = None
    udc_line: Optional[ParagraphInfo] = None
    thesis_label: Optional[ParagraphInfo] = None
    school_cn: Optional[ParagraphInfo] = None
    date_cn: Optional[ParagraphInfo] = None
    thesis_label_en: Optional[ParagraphInfo] = None
    school_en: Optional[ParagraphInfo] = None
    date_en: Optional[ParagraphInfo] = None
    declaration_heading: Optional[ParagraphInfo] = None
    authorization_heading: Optional[ParagraphInfo] = None
    abstract_cn_heading: Optional[ParagraphInfo] = None
    keywords_cn: Optional[ParagraphInfo] = None
    abstract_en_heading: Optional[ParagraphInfo] = None
    keywords_en: Optional[ParagraphInfo] = None
    toc_heading: Optional[ParagraphInfo] = None
    first_body_heading: Optional[ParagraphInfo] = None


def _compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _normalize_alignment(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, str):
        return ALIGNMENT_LABELS.get(value.lower(), value.lower())
    return ALIGNMENT_LABELS.get(value)


def _trimmed_source_text(text: str, limit: int = 200) -> str:
    stripped = (text or "").strip()
    if len(stripped) <= limit:
        return stripped
    return stripped[:limit]


def _context_window(text: str, start: int, end: int, *, limit: int = 200, padding: int = 48) -> str:
    raw = text or ""
    if not raw:
        return ""
    left = max(0, start - padding)
    right = min(len(raw), end + padding)
    snippet = raw[left:right].strip()
    if len(snippet) <= limit:
        return snippet
    center = max(0, start - left)
    local_start = max(0, center - limit // 2)
    local_end = min(len(snippet), local_start + limit)
    return snippet[local_start:local_end].strip()


def _evidence_from_text(text: str, attach_text: Optional[str] = None) -> FormatEvidence:
    base = _trimmed_source_text(text)
    if not attach_text:
        return FormatEvidence(source_text=base)
    attach = attach_text.strip()
    if not attach:
        return FormatEvidence(source_text=base)
    idx = (text or "").find(attach)
    if idx >= 0:
        return FormatEvidence(
            source_text=_context_window(text, idx, idx + len(attach)),
            attach_text=attach[:80],
        )
    return FormatEvidence(source_text=base, attach_text=attach[:80])


def _evidence_from_span(text: str, start: int, end: int, *, attach_text: Optional[str] = None) -> FormatEvidence:
    raw = text or ""
    snippet = _context_window(raw, start, end)
    anchor_text = attach_text if attach_text is not None else raw[start:end]
    anchor_text = (anchor_text or "").strip()
    return FormatEvidence(source_text=snippet or _trimmed_source_text(raw), attach_text=anchor_text[:80] if anchor_text else None)


def _coerce_evidence(value: Any) -> FormatEvidence:
    if isinstance(value, FormatEvidence):
        return FormatEvidence(
            source_text=_trimmed_source_text(value.source_text),
            attach_text=(value.attach_text or "").strip()[:80] or None,
        )
    return FormatEvidence(source_text=_trimmed_source_text(str(value or "")))


def _prefer_anchor(anchor: Optional[Dict[str, Any]], evidence: FormatEvidence) -> Optional[Dict[str, Any]]:
    if not isinstance(anchor, dict):
        return anchor
    payload = dict(anchor)
    if evidence.attach_text and not payload.get("attach_text"):
        payload["attach_text"] = evidence.attach_text
    return payload


def _length_pt(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except Exception:
        try:
            return round(float(value), 2)
        except Exception:
            return None


def _roundish(value: Any) -> Any:
    if isinstance(value, float):
        return round(value, 2)
    return value


def _dominant(values: Iterable[Any]) -> Optional[Any]:
    filtered = [v for v in values if v not in (None, "")]
    if not filtered:
        return None
    return Counter(filtered).most_common(1)[0][0]


def _mm_to_pt(mm: float) -> float:
    return mm * 72.0 / 25.4


def cn_size_to_pt(name: Optional[str]) -> Optional[float]:
    if not name:
        return None
    return CN_FONT_SIZE_PT.get(name)


def _contains_chinese(text: str) -> bool:
    return bool(CHINESE_RE.search(text or ""))


def _is_short_noise(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return True
    if NUMERIC_ONLY_RE.fullmatch(stripped):
        return True
    return False


def _looks_like_chapter_heading(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    match = CHAPTER_RE.match(_compact(stripped))
    if not match:
        return False
    compact = _compact(stripped)
    suffix = compact[match.end():]
    if not suffix:
        return True
    if suffix.startswith(("为", "主要", "重点", "首先", "其次", "最后", "阐述", "介绍", "总结", "概述", "描述")):
        return False
    if any(mark in stripped for mark in "，。；：!?！？"):
        return False
    return len(stripped) <= 60


def _style_chain(style) -> Iterable[Any]:
    seen = set()
    cur = style
    while cur is not None:
        sid = getattr(cur, "style_id", None) or id(cur)
        if sid in seen:
            break
        seen.add(sid)
        yield cur
        cur = getattr(cur, "base_style", None)


def _get_rfonts(rpr: Optional[etree._Element]) -> Dict[str, str]:
    if rpr is None:
        return {}
    rfonts = rpr.find("w:rFonts", NS)
    if rfonts is None:
        return {}
    return {
        "eastAsia": rfonts.get(qn("w:eastAsia")),
        "ascii": rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi")),
    }


def _get_font_size_pt_from_rpr(rpr: Optional[etree._Element]) -> Optional[float]:
    if rpr is None:
        return None
    size = rpr.find("w:sz", NS)
    if size is None:
        size = rpr.find("w:szCs", NS)
    if size is None:
        return None
    raw = size.get(qn("w:val"))
    if not raw:
        return None
    try:
        return round(int(raw) / 2.0, 2)
    except Exception:
        return None


def _get_ppr(paragraph: Paragraph) -> Optional[etree._Element]:
    return paragraph._element.find("w:pPr", NS)


def _get_style_rpr(style) -> Optional[etree._Element]:
    if style is None:
        return None
    return style._element.find("w:rPr", NS)


def _get_style_ppr(style) -> Optional[etree._Element]:
    if style is None:
        return None
    return style._element.find("w:pPr", NS)


def _effective_font_names(paragraph: Paragraph) -> Tuple[Optional[str], Optional[str]]:
    east_fonts: List[str] = []
    ascii_fonts: List[str] = []
    for run in paragraph.runs:
        rpr = run._element.find("w:rPr", NS)
        fonts = _get_rfonts(rpr)
        if fonts.get("eastAsia"):
            east_fonts.append(fonts["eastAsia"])
        if fonts.get("ascii"):
            ascii_fonts.append(fonts["ascii"])
    if east_fonts or ascii_fonts:
        return _dominant(east_fonts), _dominant(ascii_fonts)

    ppr = _get_ppr(paragraph)
    pr = ppr.find("w:rPr", NS) if ppr is not None else None
    fonts = _get_rfonts(pr)
    if fonts.get("eastAsia") or fonts.get("ascii"):
        return fonts.get("eastAsia"), fonts.get("ascii")

    for style in _style_chain(paragraph.style):
        fonts = _get_rfonts(_get_style_rpr(style))
        if fonts.get("eastAsia") or fonts.get("ascii"):
            return fonts.get("eastAsia"), fonts.get("ascii")
    return None, None


def _effective_font_size(paragraph: Paragraph) -> Optional[float]:
    sizes: List[float] = []
    for run in paragraph.runs:
        size = _get_font_size_pt_from_rpr(run._element.find("w:rPr", NS))
        if size is not None:
            sizes.append(size)
    if sizes:
        return _dominant([round(v, 2) for v in sizes])

    ppr = _get_ppr(paragraph)
    pr = ppr.find("w:rPr", NS) if ppr is not None else None
    size = _get_font_size_pt_from_rpr(pr)
    if size is not None:
        return size

    for style in _style_chain(paragraph.style):
        size = _get_font_size_pt_from_rpr(_get_style_rpr(style))
        if size is not None:
            return size
    return None


def _spacing_from_ppr(ppr: Optional[etree._Element]) -> Dict[str, Any]:
    result = {
        "before": None,
        "after": None,
        "line": None,
        "line_mode": None,
        "left": None,
        "right": None,
        "firstLine": None,
    }
    if ppr is None:
        return result

    spacing = ppr.find("w:spacing", NS)
    if spacing is not None:
        before = spacing.get(qn("w:before"))
        after = spacing.get(qn("w:after"))
        line = spacing.get(qn("w:line"))
        line_rule = spacing.get(qn("w:lineRule"))
        if before is not None:
            result["before"] = round(int(before) / 20.0, 2)
        if after is not None:
            result["after"] = round(int(after) / 20.0, 2)
        if line is not None:
            if line_rule in {"exact", "atLeast"}:
                result["line"] = round(int(line) / 20.0, 2)
                result["line_mode"] = line_rule
            else:
                result["line"] = round(int(line) / 240.0, 2)
                result["line_mode"] = line_rule or "auto"

    ind = ppr.find("w:ind", NS)
    if ind is not None:
        left = ind.get(qn("w:left"))
        right = ind.get(qn("w:right"))
        first_line = ind.get(qn("w:firstLine"))
        if left is not None:
            result["left"] = round(int(left) / 20.0, 2)
        if right is not None:
            result["right"] = round(int(right) / 20.0, 2)
        if first_line is not None:
            result["firstLine"] = round(int(first_line) / 20.0, 2)

    return result


def _effective_paragraph_metrics(paragraph: Paragraph) -> Dict[str, Any]:
    metrics = {
        "alignment": None,
        "line": None,
        "line_mode": None,
        "before": None,
        "after": None,
        "left": None,
        "right": None,
        "firstLine": None,
    }

    ppr = _get_ppr(paragraph)
    if ppr is not None:
        jc = ppr.find("w:jc", NS)
        if jc is not None:
            metrics["alignment"] = _normalize_alignment(jc.get(qn("w:val")))
        direct_spacing = _spacing_from_ppr(ppr)
        for key, value in direct_spacing.items():
            if value is not None:
                metrics[key] = value

    if all(v is not None for v in metrics.values() if v is not metrics["alignment"]):
        return metrics

    for style in _style_chain(paragraph.style):
        ppr_style = _get_style_ppr(style)
        if ppr_style is None:
            continue
        if metrics["alignment"] is None:
            jc = ppr_style.find("w:jc", NS)
            if jc is not None:
                metrics["alignment"] = _normalize_alignment(jc.get(qn("w:val")))
        style_spacing = _spacing_from_ppr(ppr_style)
        for key, value in style_spacing.items():
            if metrics.get(key) is None and value is not None:
                metrics[key] = value
    return metrics


def detect_heading_level(text: str, style_name: str) -> Optional[int]:
    s = (style_name or "").lower()
    if "heading 1" in s or "标题 1" in s or "标题1" in s or "标题样式1" in s:
        return 1
    if "heading 2" in s or "标题 2" in s or "标题2" in s or "标题样式2" in s:
        return 2
    if "heading 3" in s or "标题 3" in s or "标题3" in s or "标题样式3" in s:
        return 3
    if "heading 4" in s or "标题 4" in s or "标题4" in s or "标题样式4" in s:
        return 4

    stripped = (text or "").strip()
    compact = _compact(stripped)
    if (
        ABSTRACT_CN_RE.match(stripped)
        or ABSTRACT_EN_RE.match(stripped)
        or TOC_HEADING_RE.match(stripped)
        or REFERENCE_HEADING_RE.match(stripped)
        or ACK_HEADING_RE.match(compact)
        or ACHIEVEMENT_HEADING_RE.match(compact)
        or AUTHORIZATION_HEADING_RE.match(stripped)
        or _looks_like_chapter_heading(stripped)
    ):
        return 1
    if LEVEL4_RE.match(stripped):
        return 4
    if LEVEL3_RE.match(stripped):
        return 3
    if LEVEL2_RE.match(stripped):
        return 2
    return None


def _has_drawing(paragraph: Paragraph) -> bool:
    xml = paragraph._element.xml
    return (
        "<w:drawing" in xml
        or "<w:pict" in xml
        or "<wp:inline" in xml
        or "<wp:anchor" in xml
        or "<w:object" in xml
        or "<v:shape" in xml
    )


def _has_equation(paragraph: Paragraph) -> bool:
    xml = paragraph._element.xml
    return "<m:oMath" in xml or "<m:oMathPara" in xml


def _determine_zone_sequence(paragraphs: List[Paragraph]) -> Dict[int, str]:
    zone = "front_matter"
    zones: Dict[int, str] = {}
    for idx, paragraph in enumerate(paragraphs, start=1):
        text = get_paragraph_visible_text(paragraph._element).strip()
        compact = _compact(text)
        if ABSTRACT_CN_RE.match(text):
            zone = "abstract_cn"
        elif ABSTRACT_EN_RE.match(text):
            zone = "abstract_en"
        elif TOC_HEADING_RE.match(text):
            zone = "toc"
        elif REFERENCE_HEADING_RE.match(text):
            zone = "references"
        elif ACK_HEADING_RE.match(compact):
            zone = "acknowledgements"
        elif ACHIEVEMENT_HEADING_RE.match(compact):
            zone = "achievements"
        elif _looks_like_chapter_heading(text):
            zone = "body"
        zones[idx] = zone
    return zones


def _detect_paragraph_type(text: str, style_name: str, zone: str, heading_level: Optional[int], has_equation: bool) -> str:
    stripped = (text or "").strip()
    if heading_level:
        return f"heading{heading_level}"
    if zone == "toc" and (TOC_DOT_RE.search(stripped) or re.search(r"\d+$", stripped)):
        return "toc_entry"
    if zone in {"abstract_cn", "abstract_en"} and (KEYWORDS_RE.match(stripped) or EN_KEYWORDS_RE.match(stripped)):
        return "keywords"
    if FIGURE_CAPTION_RE.match(stripped):
        return "caption_figure"
    if TABLE_CAPTION_RE.match(stripped):
        return "caption_table"
    if EN_FIGURE_CAPTION_RE.match(stripped):
        return "caption_figure_en"
    if EN_TABLE_CAPTION_RE.match(stripped):
        return "caption_table_en"
    if zone == "references":
        return "reference_entry"
    if has_equation and len(stripped) <= 120:
        return "formula"
    if FORMULA_NO_RE.search(stripped) and len(stripped) <= 120:
        return "formula"
    return "body"


def _paragraph_info(paragraph: Paragraph, index: int, zone: str) -> ParagraphInfo:
    text = get_paragraph_visible_text(paragraph._element).strip()
    style_name = paragraph.style.name if paragraph.style else ""
    heading_level = detect_heading_level(text, style_name)
    has_equation = _has_equation(paragraph)
    metrics = _effective_paragraph_metrics(paragraph)
    east_font, ascii_font = _effective_font_names(paragraph)
    ptype = _detect_paragraph_type(text, style_name, zone, heading_level, has_equation)
    return ParagraphInfo(
        index=index,
        location=f"第{index}段",
        text=text,
        style_name=style_name,
        zone=zone,
        paragraph_type=ptype,
        heading_level=heading_level,
        alignment=metrics["alignment"],
        line_spacing=metrics["line"],
        line_spacing_mode=metrics["line_mode"],
        space_before_pt=metrics["before"],
        space_after_pt=metrics["after"],
        first_line_indent_pt=metrics["firstLine"],
        left_indent_pt=metrics["left"],
        right_indent_pt=metrics["right"],
        east_asia_font=east_font,
        ascii_font=ascii_font,
        font_size_pt=_effective_font_size(paragraph),
        has_drawing=_has_drawing(paragraph),
        has_equation=has_equation,
        raw_paragraph=paragraph,
    )


def _iter_block_items(parent):
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"unsupported parent: {type(parent)}")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _build_body_xml_paragraph_order_map(doc: Document) -> Dict[int, int]:
    body = doc.element.body
    return {id(p): idx for idx, p in enumerate(body.xpath('.//w:p'), start=1)}


def _paragraph_has_section_break(paragraph: Paragraph) -> bool:
    return paragraph._element.find('w:pPr/w:sectPr', NS) is not None


def _build_section_start_map(doc: Document) -> Dict[int, int]:
    starts: Dict[int, int] = {1: 1}
    section_idx = 1
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        if _paragraph_has_section_break(paragraph) and idx < len(doc.paragraphs):
            section_idx += 1
            starts.setdefault(section_idx, idx + 1)
    return starts


def _section_index_for_paragraph(paragraph_index: int, section_starts: Dict[int, int]) -> int:
    section_idx = 1
    for idx in sorted(section_starts):
        if paragraph_index >= section_starts[idx]:
            section_idx = idx
        else:
            break
    return section_idx


def _build_section_proxy_paragraphs(main_paragraphs: List[ParagraphInfo], section_starts: Dict[int, int]) -> Dict[int, ParagraphInfo]:
    by_index = {p.index: p for p in main_paragraphs if p.index is not None}
    proxies: Dict[int, ParagraphInfo] = {}
    ordered_starts = sorted(section_starts.items())
    max_index = max(by_index) if by_index else 0
    for pos, (section_idx, start_idx) in enumerate(ordered_starts):
        next_start = ordered_starts[pos + 1][1] if pos + 1 < len(ordered_starts) else max_index + 1
        candidate = None
        for idx in range(start_idx, next_start):
            info = by_index.get(idx)
            if info and info.text.strip():
                candidate = info
                break
        if candidate is None:
            for idx in range(next_start, max_index + 1):
                info = by_index.get(idx)
                if info and info.text.strip():
                    candidate = info
                    break
        if candidate is None:
            for idx in range(start_idx - 1, 0, -1):
                info = by_index.get(idx)
                if info and info.text.strip():
                    candidate = info
                    break
        if candidate is None:
            candidate = by_index.get(start_idx) or (main_paragraphs[0] if main_paragraphs else None)
        if candidate is not None:
            proxies[section_idx] = candidate
    return proxies


def _build_part_proxy_maps(docx_path: Path, section_starts: Dict[int, int]) -> Tuple[Dict[str, int], Dict[str, int]]:
    header_map: Dict[str, int] = {}
    footer_map: Dict[str, int] = {}
    with ZipFile(docx_path) as zf:
        doc_root = etree.fromstring(zf.read('word/document.xml'))
        rels_root = etree.fromstring(zf.read('word/_rels/document.xml.rels'))
    rid_to_target: Dict[str, str] = {}
    for rel in rels_root:
        rel_id = rel.get('Id')
        target = rel.get('Target')
        if rel_id and target:
            rid_to_target[rel_id] = Path(target).name

    section_idx = 1

    def bind_from_sectpr(sectpr: Optional[etree._Element], current_section_idx: int) -> None:
        if sectpr is None:
            return
        for ref in sectpr.findall('w:headerReference', NS):
            rid = ref.get(qn('r:id'))
            target = rid_to_target.get(rid or '')
            if target:
                header_map.setdefault(target, current_section_idx)
        for ref in sectpr.findall('w:footerReference', NS):
            rid = ref.get(qn('r:id'))
            target = rid_to_target.get(rid or '')
            if target:
                footer_map.setdefault(target, current_section_idx)

    body = doc_root.find('w:body', NS)
    if body is not None:
        for child in body:
            if child.tag == qn('w:p'):
                sectpr = child.find('w:pPr/w:sectPr', NS)
                if sectpr is not None:
                    bind_from_sectpr(sectpr, section_idx)
                    section_idx += 1
        bind_from_sectpr(body.find('w:sectPr', NS), section_idx)

    if not header_map and section_starts:
        header_map['__missing_header__'] = 1
    if not footer_map and section_starts:
        footer_map['__missing_footer__'] = 1
    return header_map, footer_map


def _body_anchor(info: ParagraphInfo) -> Dict[str, Any]:
    node_id = make_body_node_id(info.index)
    return {
        'kind': 'body_paragraph',
        'node_id': node_id,
        'presentation_kind': 'body_paragraph',
        'presentation_node_id': node_id,
        'paragraph_index': info.index,
        'xml_paragraph_index': info.xml_paragraph_index,
        'section_index': info.section_index,
        'zone': info.zone,
        'paragraph_type': info.paragraph_type,
        'location': info.location,
    }


def _table_cell_anchor(info: ParagraphInfo) -> Dict[str, Any]:
    node_id = make_table_cell_node_id(info.table_index, info.row_index, info.column_index, info.cell_paragraph_index)
    return {
        'kind': 'table_cell',
        'node_id': node_id,
        'presentation_kind': 'table_cell',
        'presentation_node_id': node_id,
        'xml_paragraph_index': info.xml_paragraph_index,
        'table_index': info.table_index,
        'row_index': info.row_index,
        'column_index': info.column_index,
        'cell_paragraph_index': info.cell_paragraph_index,
        'location': info.location,
    }


def _proxy_anchor(kind: str, proxy: ParagraphInfo, **extra: Any) -> Dict[str, Any]:
    section_index = extra.get('section_index', proxy.section_index)
    if kind == 'header_proxy':
        node_id = make_header_proxy_node_id(section_index, extra.get('header_part'))
    elif kind == 'footer_proxy':
        node_id = make_footer_proxy_node_id(section_index, extra.get('footer_part'))
    else:
        node_id = make_section_proxy_node_id(section_index)
    presentation_node_id = make_body_node_id(proxy.index)
    anchor = {
        'kind': kind,
        'node_id': node_id,
        'presentation_kind': 'body_paragraph',
        'presentation_node_id': presentation_node_id,
        'proxy_paragraph_index': proxy.index,
        'proxy_xml_paragraph_index': proxy.xml_paragraph_index,
        'proxy_location': proxy.location,
        'proxy_text': proxy.text[:120],
        'section_index': section_index,
    }
    anchor.update({k: v for k, v in extra.items() if v is not None})
    return anchor


def _collect_main_paragraphs(doc: Document) -> List[ParagraphInfo]:
    raw_paragraphs = list(doc.paragraphs)
    zones = _determine_zone_sequence(raw_paragraphs)
    xml_order = _build_body_xml_paragraph_order_map(doc)
    section_starts = _build_section_start_map(doc)
    infos: List[ParagraphInfo] = []
    for idx, paragraph in enumerate(raw_paragraphs, start=1):
        info = _paragraph_info(paragraph, idx, zones[idx])
        info.xml_paragraph_index = xml_order.get(id(paragraph._element))
        info.section_index = _section_index_for_paragraph(idx, section_starts)
        infos.append(info)
    return infos


def _collect_table_cell_paragraphs(doc: Document) -> List[ParagraphInfo]:
    infos: List[ParagraphInfo] = []
    seen_paragraph_elements: set[int] = set()
    xml_order = _build_body_xml_paragraph_order_map(doc)
    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                for p_idx, paragraph in enumerate(cell.paragraphs, start=1):
                    element_id = id(paragraph._element)
                    if element_id in seen_paragraph_elements:
                        continue
                    seen_paragraph_elements.add(element_id)
                    text = get_paragraph_visible_text(paragraph._element).strip()
                    if not text:
                        continue
                    metrics = _effective_paragraph_metrics(paragraph)
                    east_font, ascii_font = _effective_font_names(paragraph)
                    xml_idx = xml_order.get(element_id)
                    infos.append(
                        ParagraphInfo(
                            index=None,
                            location=f'表格{t_idx} 第{r_idx}行第{c_idx}列（段{p_idx}）',
                            text=text,
                            style_name=paragraph.style.name if paragraph.style else '',
                            zone='body',
                            paragraph_type='table_cell',
                            heading_level=None,
                            alignment=metrics['alignment'],
                            line_spacing=metrics['line'],
                            line_spacing_mode=metrics['line_mode'],
                            space_before_pt=metrics['before'],
                            space_after_pt=metrics['after'],
                            first_line_indent_pt=metrics['firstLine'],
                            left_indent_pt=metrics['left'],
                            right_indent_pt=metrics['right'],
                            east_asia_font=east_font,
                            ascii_font=ascii_font,
                            font_size_pt=_effective_font_size(paragraph),
                            xml_paragraph_index=xml_idx,
                            section_index=None,
                            table_index=t_idx,
                            row_index=r_idx,
                            column_index=c_idx,
                            cell_paragraph_index=p_idx,
                            raw_paragraph=paragraph,
                        )
                    )
    return infos


def _assign_table_cell_zones(doc: Document, table_cells: List[ParagraphInfo], main_paragraphs: List[ParagraphInfo]) -> None:
    table_zone: Dict[int, tuple[str, Optional[int]]] = {}
    current_zone = 'front_matter'
    current_section = None
    paragraph_cursor = 0
    table_cursor = 0
    for child in doc._body._element.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            if paragraph_cursor < len(main_paragraphs):
                current_zone = main_paragraphs[paragraph_cursor].zone
                current_section = main_paragraphs[paragraph_cursor].section_index
            paragraph_cursor += 1
        elif tag == 'tbl':
            table_cursor += 1
            table_zone[table_cursor] = (current_zone, current_section)
    for cell in table_cells:
        zone, section = table_zone.get(cell.table_index or 0, (current_zone, current_section))
        cell.zone = zone
        cell.section_index = section


def _build_section_zone_map(paragraphs: List[ParagraphInfo], section_starts: Dict[int, int]) -> Dict[int, str]:
    by_index = {p.index: p for p in paragraphs if p.index is not None}
    section_zones: Dict[int, str] = {}
    max_index = max(by_index) if by_index else 0
    ordered = sorted(section_starts.items())
    for pos, (section_idx, start_idx) in enumerate(ordered):
        next_start = ordered[pos + 1][1] if pos + 1 < len(ordered) else max_index + 1
        zone = None
        for idx in range(start_idx, next_start):
            info = by_index.get(idx)
            if info is None:
                continue
            zone = info.zone
            if info.text.strip():
                break
        section_zones[section_idx] = zone or "front_matter"
    return section_zones


def _build_front_matter_model(paragraphs: List[ParagraphInfo]) -> FrontMatterModel:
    model = FrontMatterModel()
    for info in paragraphs:
        text = (info.text or "").strip()
        compact = _compact(text)
        if not text:
            continue
        if model.classification_line is None and CLASSIFICATION_LINE_RE.match(compact):
            model.classification_line = info
        elif model.udc_line is None and UDC_RE.match(compact):
            model.udc_line = info
        elif model.thesis_label is None and THESIS_LABEL_RE.match(compact):
            model.thesis_label = info
        elif model.school_cn is None and compact == "东北大学":
            model.school_cn = info
        elif model.date_cn is None and CN_DATE_RE.match(compact):
            model.date_cn = info
        elif model.thesis_label_en is None and EN_THESIS_LABEL_RE.match(text):
            model.thesis_label_en = info
        elif model.school_en is None and compact.lower() == "northeasternuniversity":
            model.school_en = info
        elif model.date_en is None and EN_DATE_RE.match(text):
            model.date_en = info
        elif model.declaration_heading is None and compact == "独创性声明":
            model.declaration_heading = info
        elif model.authorization_heading is None and AUTHORIZATION_HEADING_RE.match(text):
            model.authorization_heading = info
        elif model.abstract_cn_heading is None and ABSTRACT_CN_RE.match(text):
            model.abstract_cn_heading = info
        elif model.keywords_cn is None and info.zone == "abstract_cn" and KEYWORDS_RE.match(text):
            model.keywords_cn = info
        elif model.abstract_en_heading is None and ABSTRACT_EN_RE.match(text):
            model.abstract_en_heading = info
        elif model.keywords_en is None and info.zone == "abstract_en" and EN_KEYWORDS_RE.match(text):
            model.keywords_en = info
        elif model.toc_heading is None and TOC_HEADING_RE.match(text):
            model.toc_heading = info
        elif model.first_body_heading is None and info.zone == "body" and info.heading_level == 1:
            model.first_body_heading = info
    return model


def _model_anchor_info(model: FrontMatterModel) -> Optional[ParagraphInfo]:
    for item in [
        model.classification_line,
        model.udc_line,
        model.thesis_label,
        model.school_cn,
        model.date_cn,
        model.thesis_label_en,
        model.school_en,
        model.date_en,
        model.declaration_heading,
        model.authorization_heading,
        model.abstract_cn_heading,
        model.abstract_en_heading,
        model.toc_heading,
        model.first_body_heading,
    ]:
        if item is not None:
            return item
    return None


def _span_anchor_phrase(text: str, start: int, end: int, radius: int = 14) -> str:
    raw = text or ""
    if not raw:
        return ""
    left = max(0, start - radius)
    right = min(len(raw), end + radius)
    phrase = raw[left:right].strip(" \t\r\n，。；：、,.;")
    return phrase or raw[start:end].strip()


def _extra_space_evidence(text: str, match: re.Match[str]) -> FormatEvidence:
    start, end = match.span()
    attach_text = _span_anchor_phrase(text, start, end)
    return _evidence_from_span(text, start, end, attach_text=attach_text)


def _reference_tail_evidence(text: str) -> FormatEvidence:
    stripped = (text or "").strip()
    if len(stripped) <= 80:
        return FormatEvidence(source_text=stripped, attach_text=stripped)
    attach_text = stripped[-80:].lstrip("，,；; ")
    return _evidence_from_text(stripped, attach_text=attach_text)


def _font_mismatch_evidence(info: ParagraphInfo, expected_font: str, *, font_channel: str) -> FormatEvidence:
    paragraph = info.raw_paragraph
    if paragraph is not None:
        for run in paragraph.runs:
            run_text = (run.text or "").strip()
            if not run_text:
                continue
            if font_channel == "eastAsia" and not _contains_chinese(run_text):
                continue
            if font_channel == "ascii" and (_contains_chinese(run_text) or not re.search(r"[A-Za-z0-9]", run_text)):
                continue
            fonts = _get_rfonts(run._element.find("w:rPr", NS))
            actual = fonts.get(font_channel)
            if not actual:
                continue
            if font_channel == "ascii":
                if expected_font.lower() in actual.lower():
                    continue
            elif expected_font in actual:
                continue
            return _evidence_from_text(info.text, attach_text=run_text[:80])

    return _evidence_from_text(info.text)


def _expected_key_for_paragraph(p: ParagraphInfo) -> str:
    if p.zone == 'toc' and TOC_HEADING_RE.match((p.text or '').strip()):
        return 'toc_heading'
    return p.paragraph_type


def _anchor_for_info(info: ParagraphInfo) -> Dict[str, Any]:
    return _table_cell_anchor(info) if info.paragraph_type == 'table_cell' else _body_anchor(info)


def load_format_rules(path: Optional[str | Path] = None) -> Dict[str, Any]:
    file_path = Path(path) if path else FORMAT_RULES_PATH
    if not file_path.exists():
        return {}
    return json.loads(file_path.read_text(encoding="utf-8"))


def _get_expected_map(rules: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    headings = rules.get("headings", {})
    figures = rules.get("figures", {})
    tables = rules.get("tables", {})
    body = rules.get("bodyText", {})
    toc = rules.get("toc", {})
    return {
        "heading1": {
            "fontSizePt": cn_size_to_pt(headings.get("level1", {}).get("fontSize")),
            "fontName": headings.get("level1", {}).get("fontName"),
            "alignment": "center",
            "basis": "institution-format-guide 4.1：第一级标题黑体二号居中；摘要/目录/参考文献等同级。",
        },
        "heading2": {
            "fontSizePt": cn_size_to_pt(headings.get("level2", {}).get("fontSize")),
            "fontName": headings.get("level2", {}).get("fontName"),
            "alignment": "left",
            "basis": "institution-format-guide 4.1：第二级标题黑体三号居左。",
        },
        "heading3": {
            "fontSizePt": cn_size_to_pt(headings.get("level3", {}).get("fontSize")),
            "fontName": headings.get("level3", {}).get("fontName"),
            "alignment": "left",
            "basis": "institution-format-guide 4.1：第三级标题黑体四号居左。",
        },
        "heading4": {
            "fontSizePt": cn_size_to_pt(headings.get("level4", {}).get("fontSize")),
            "fontName": headings.get("level4", {}).get("fontName"),
            "alignment": "left",
            "basis": "institution-format-guide 4.1：第四级标题黑体小四居左。",
        },
        "toc_heading": {
            "fontSizePt": cn_size_to_pt(toc.get("titleFontSize")) or cn_size_to_pt(headings.get("level1", {}).get("fontSize")),
            "fontName": toc.get("titleFontName") or headings.get("level1", {}).get("fontName"),
            "alignment": "center",
            "basis": "第三份格式说明：目录页标题“目录”按黑体小三居中检查。",
        },
        "body": {
            "fontSizePt": cn_size_to_pt(body.get("fontSize")),
            "fontName": body.get("fontName"),
            "fontNameEn": body.get("fontNameEn"),
            "basis": "institution-format-guide 3 / 4.2 + 第三份格式说明：正文小四宋体，外文数字同号，英文用 Times New Roman 或相近字体。",
        },
        "keywords": {
            "fontSizePt": cn_size_to_pt(body.get("fontSize")),
            "fontName": body.get("fontName"),
            "fontNameEn": body.get("fontNameEn"),
            "basis": "摘要与关键词示例沿用正文同号字体。",
        },
        "caption_figure": {
            "fontSizePt": cn_size_to_pt(figures.get("captionFontSize")),
            "fontName": str(figures.get("captionFontName", "")).split("（")[0],
            "basis": "institution-format-guide 4.2.1：图题位于图下方，五号字，全文统一。",
        },
        "caption_table": {
            "fontSizePt": cn_size_to_pt(tables.get("captionFontSize")),
            "fontName": str(tables.get("captionFontName", "")).split("（")[0],
            "alignment": "center",
            "basis": "institution-format-guide 4.2.2：表题位于表上方居中，五号宋体，全文统一。",
        },
        "table_cell": {
            "maxFontSizePt": cn_size_to_pt(tables.get("contentFontSize")),
            "basis": "institution-format-guide 4.2.2：表内文字字号不大于图题字号。",
        },
        "formula": {
            "alignment": "center",
            "basis": "institution-format-guide 4.2.3：公式单行居中，式号同行居右，按章编号。",
        },
    }


def _nearly_equal(a: Optional[float], b: Optional[float], tolerance: float = 0.5) -> bool:
    if a is None or b is None:
        return True
    return abs(a - b) <= tolerance


def _add_issue(
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    *,
    location: str,
    paragraph_index: Optional[int],
    style_name: str,
    check_type: str,
    expected: Any,
    actual: Any,
    source_text: str,
    problem: str,
    fix: str,
    basis: str,
    severity: str = "medium",
    mode: str = "explicit",
    anchor: Optional[Dict[str, Any]] = None,
) -> None:
    evidence = _coerce_evidence(source_text)
    anchor_payload = _prefer_anchor(anchor, evidence)
    issue_id = f"format-{counter[0]:04d}"
    counter[0] += 1
    issues.append(
        FormatIssue(
            id=issue_id,
            location=location,
            paragraphIndex=paragraph_index,
            styleName=style_name,
            checkType=check_type,
            expected=expected,
            actual=actual,
            sourceText=evidence.source_text,
            problem=problem,
            fixDirection=fix,
            severity=severity,
            basis=basis,
            mode=mode,
            anchor=anchor_payload,
        )
    )
    annotations.append(
        UnifiedAnnotation(
            id=issue_id,
            source=SourceType.FORMAT,
            issue_type=check_type,
            location_hint=location,
            source_text=evidence.source_text,
            problem=problem,
            suggestion=fix,
            basis=basis,
            risk="格式不符合学校规范或与同类对象不统一，可能影响论文评审印象。",
            human_review_required=HumanReviewRequired.RECOMMENDED,
            confidence=ConfidenceLevel.HIGH if severity != "low" else ConfidenceLevel.MEDIUM,
            severity=severity,
            anchor=anchor_payload,
        ).to_dict()
    )


def _body_like_candidates(paragraphs: List[ParagraphInfo]) -> List[ParagraphInfo]:
    result = []
    for p in paragraphs:
        if p.paragraph_type not in {"body", "keywords"}:
            continue
        if p.zone not in {"abstract_cn", "abstract_en", "body"}:
            continue
        if _is_short_noise(p.text) or len(p.text) < 20:
            continue
        if LIST_LIKE_RE.match(p.text):
            continue
        result.append(p)
    return result


def _consistency_group_map(main_paragraphs: List[ParagraphInfo], table_cells: List[ParagraphInfo]) -> Dict[str, List[ParagraphInfo]]:
    toc_entries = [p for p in main_paragraphs if p.paragraph_type == "toc_entry"]
    ref_entries = [p for p in main_paragraphs if p.paragraph_type == "reference_entry"]
    figure_caps = [p for p in main_paragraphs if p.paragraph_type == "caption_figure"]
    table_caps = [p for p in main_paragraphs if p.paragraph_type == "caption_table"]
    formula_expl = [p for p in main_paragraphs if p.text.startswith(("其中", "式中", "式中：", "其中："))]
    return {
        "body": _body_like_candidates(main_paragraphs),
        "toc": toc_entries,
        "references": ref_entries,
        "figure_caption": figure_caps,
        "table_caption": table_caps,
        "table_cell": [p for p in table_cells if p.zone != "front_matter"],
        "formula_explanation": formula_expl,
    }


def _apply_consistency_checks(
    groups: Dict[str, List[ParagraphInfo]],
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    document_checks: List[CheckResult],
) -> None:
    specs = [
        ("body", "line_spacing", "line_spacing", 0.1, "historical-defect-checklist + institution-format-guide：正文排版应统一，不能只看字号。"),
        ("body", "space_before", "space_before_pt", 0.5, "正文段前距应统一。"),
        ("body", "space_after", "space_after_pt", 0.5, "正文段后距应统一。"),
        ("body", "first_line_indent", "first_line_indent_pt", 1.0, "正文首行缩进应统一。"),
        ("body", "left_indent", "left_indent_pt", 1.0, "正文左缩进应统一。"),
        ("body", "right_indent", "right_indent_pt", 1.0, "正文右缩进应统一。"),
        ("body", "alignment", "alignment", 0.0, "正文对齐方式应统一。"),
        ("toc", "line_spacing", "line_spacing", 0.1, "historical-defect-checklist：目录行间距不统一是高频问题。"),
        ("toc", "font_size", "font_size_pt", 0.5, "historical-defect-checklist：目录字体大小需符合标准并保持统一。"),
        ("references", "line_spacing", "line_spacing", 0.1, "参考文献排版应统一。"),
        ("references", "font_size", "font_size_pt", 0.5, "参考文献字号应统一。"),
        ("figure_caption", "font_size", "font_size_pt", 0.5, "图题字号须全文统一。"),
        ("table_caption", "font_size", "font_size_pt", 0.5, "表题字号须全文统一。"),
        ("table_cell", "font_size", "font_size_pt", 0.5, "表内字号须统一且不大于图题字号。"),
        ("formula_explanation", "expression", "text", 0.0, "historical-defect-checklist：公式释义应采用统一表达方式。"),
    ]

    for group_name, check_type, attr, tolerance, basis in specs:
        items = [item for item in groups.get(group_name, []) if getattr(item, attr if attr != "expression" else "text", None) not in (None, "")]
        if len(items) < 2:
            document_checks.append(CheckResult(checkType=f"{group_name}_{check_type}", status="not_enough_data", basis=basis, details="样本不足，无法做一致性判定。"))
            continue

        if attr == "text":
            normalized = []
            for item in items:
                prefix = "式中" if item.text.startswith("式中") else "其中" if item.text.startswith("其中") else item.text[:6]
                normalized.append(prefix)
            dominant = _dominant(normalized)
            mismatches = [item for item, prefix in zip(items, normalized) if prefix != dominant]
            document_checks.append(CheckResult(checkType=f"{group_name}_{check_type}", status="pass" if not mismatches else "fail", basis=basis, details=f"dominant={dominant}, mismatches={len(mismatches)}"))
            for item, prefix in zip(items, normalized):
                if prefix == dominant:
                    continue
                _add_issue(
                    issues,
                    annotations,
                    counter,
                    location=item.location,
                    paragraph_index=item.index,
                    style_name=item.style_name,
                    check_type="formula_explanation",
                    expected=dominant,
                    actual=prefix,
                    source_text=item.text,
                    problem=f"公式释义表达方式不统一：同类段落主流写法为“{dominant}”，当前为“{prefix}”。",
                    fix=f"将释义前缀统一为“{dominant}”或统一采用另一套全篇一致写法。",
                    basis=basis,
                    severity="low",
                    mode="consistency",
                    anchor=_body_anchor(item),
                )
            continue

        values = [getattr(item, attr) for item in items]
        rounded_values = [_roundish(v) for v in values]
        dominant = _dominant(rounded_values)
        mismatches: List[ParagraphInfo] = []
        for item in items:
            actual = getattr(item, attr)
            if isinstance(dominant, (int, float)):
                if not _nearly_equal(float(actual), float(dominant), tolerance):
                    mismatches.append(item)
            else:
                if actual != dominant:
                    mismatches.append(item)

        document_checks.append(
            CheckResult(
                checkType=f"{group_name}_{check_type}",
                status="pass" if not mismatches else "fail",
                basis=basis,
                details=f"dominant={dominant}, mismatches={len(mismatches)}",
            )
        )

        for item in mismatches:
            actual = getattr(item, attr)
            label = {
                "line_spacing": "行距",
                "space_before_pt": "段前距",
                "space_after_pt": "段后距",
                "first_line_indent_pt": "首行缩进",
                "left_indent_pt": "左缩进",
                "right_indent_pt": "右缩进",
                "alignment": "对齐方式",
                "font_size_pt": "字号",
            }.get(attr, attr)
            _add_issue(
                issues,
                annotations,
                counter,
                location=item.location,
                paragraph_index=item.index,
                style_name=item.style_name,
                check_type=check_type,
                expected=dominant,
                actual=actual,
                source_text=item.text,
                problem=f"{label}与同类对象主流设置不一致：同组主值为 {dominant}，当前为 {actual}。",
                fix=f"将该处 {label} 调整为与同类对象一致（建议对齐到 {dominant}）。",
                basis=basis,
                severity="medium",
                mode="consistency",
                anchor=_table_cell_anchor(item) if item.paragraph_type == "table_cell" else _body_anchor(item),
            )


def _apply_explicit_paragraph_checks(
    paragraphs: List[ParagraphInfo],
    expected_map: Dict[str, Dict[str, Any]],
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    document_checks: List[CheckResult],
) -> None:
    covered = Counter()
    for p in paragraphs:
        if not p.text.strip():
            continue
        if p.zone == "front_matter" and p.paragraph_type not in {"heading1", "heading2", "heading3", "heading4"}:
            continue
        expected_key = _expected_key_for_paragraph(p)
        expected = expected_map.get(expected_key)
        if not expected:
            continue
        covered[expected_key] += 1
        basis = expected.get("basis", "institution-format-guide")
        exp_size = expected.get("fontSizePt")
        if exp_size and p.font_size_pt and not _nearly_equal(exp_size, p.font_size_pt, 0.5):
            _add_issue(
                issues,
                annotations,
                counter,
                location=p.location,
                paragraph_index=p.index,
                style_name=p.style_name,
                check_type="font_size",
                expected=exp_size,
                actual=p.font_size_pt,
                source_text=p.text,
                problem=f"字号不符合该分区规则：预期 {exp_size}pt，实际 {p.font_size_pt}pt。",
                fix=f"将字号调整为 {exp_size}pt。",
                basis=basis,
                severity="medium",
                mode="explicit",
                anchor=_body_anchor(p),
            )

        exp_font = expected.get("fontName")
        if exp_font and _contains_chinese(p.text) and p.east_asia_font and exp_font not in p.east_asia_font:
            _add_issue(
                issues,
                annotations,
                counter,
                location=p.location,
                paragraph_index=p.index,
                style_name=p.style_name,
                check_type="font_name",
                expected=exp_font,
                actual=p.east_asia_font,
                source_text=_font_mismatch_evidence(p, exp_font, font_channel="eastAsia"),
                problem=f"中文字体不符合该分区规则：预期 {exp_font}，实际 {p.east_asia_font}。",
                fix=f"将中文字体调整为 {exp_font}。",
                basis=basis,
                severity="medium",
                mode="explicit",
                anchor=_body_anchor(p),
            )

        exp_alignment = expected.get("alignment")
        actual_alignment = p.alignment or ("left" if exp_alignment == "left" else None)
        if exp_alignment and actual_alignment != exp_alignment:
            _add_issue(
                issues,
                annotations,
                counter,
                location=p.location,
                paragraph_index=p.index,
                style_name=p.style_name,
                check_type="alignment",
                expected=exp_alignment,
                actual=p.alignment or "unspecified(default-left)",
                source_text=p.text,
                problem=f"对齐方式不符合该分区规则：预期 {exp_alignment}，实际 {p.alignment or '未显式设置（默认左对齐）'}。",
                fix=f"将对齐方式调整为 {exp_alignment}。",
                basis=basis,
                severity="low",
                mode="explicit",
                anchor=_body_anchor(p),
            )

        exp_en_font = expected.get("fontNameEn")
        if exp_en_font and not _contains_chinese(p.text) and p.ascii_font and exp_en_font.lower() not in p.ascii_font.lower():
            _add_issue(
                issues,
                annotations,
                counter,
                location=p.location,
                paragraph_index=p.index,
                style_name=p.style_name,
                check_type="font_name_en",
                expected=exp_en_font,
                actual=p.ascii_font,
                source_text=_font_mismatch_evidence(p, exp_en_font, font_channel="ascii"),
                problem=f"英文字体不符合该分区规则：预期 {exp_en_font}，实际 {p.ascii_font}。",
                fix=f"将英文字体调整为 {exp_en_font} 或相近 Roman 字体。",
                basis=basis,
                severity="low",
                mode="explicit",
                anchor=_body_anchor(p),
            )

    for ptype, count in sorted(covered.items()):
        document_checks.append(CheckResult(checkType=f"coverage_{ptype}", status="covered", basis=expected_map[ptype].get("basis", "institution-format-guide"), details=f"覆盖 {count} 个段落/对象。"))


def _apply_front_matter_checks(
    paragraphs: List[ParagraphInfo],
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    document_checks: List[CheckResult],
) -> None:
    model = _build_front_matter_model(paragraphs)
    fallback = _model_anchor_info(model)
    fallback_anchor = _body_anchor(fallback) if fallback is not None else None
    fallback_text = fallback.text if fallback is not None else "前置部分"

    required_fields = [
        ("cover_classification", "封面应包含“分类号 / 密级”字段", model.classification_line, "institution-format-guide 结构章节：封面需包含分类号、密级等基本信息。"),
        ("cover_udc", "封面应包含 UDC 字段", model.udc_line, "institution-format-guide 结构章节：扉页/封面需包含 UDC。"),
        ("cover_thesis_label", "封面应包含“学位论文”标识", model.thesis_label, "institution-format-guide 结构章节：封面需有学位论文标识。"),
        ("cover_school_cn", "中文封面应包含“东北大学”校名", model.school_cn, "institution-format-guide 结构章节：封面需有学校名称。"),
        ("cover_date_cn", "中文封面应包含年月", model.date_cn, "institution-format-guide 结构章节：封面需有日期信息。"),
        ("cover_thesis_label_en", "英文扉页应包含英文论文类型标识", model.thesis_label_en, "institution-format-guide 结构章节：扉页需包含英文论文类型信息。"),
        ("cover_school_en", "英文扉页应包含 Northeastern University", model.school_en, "institution-format-guide 结构章节：英文扉页需包含学校英文名称。"),
        ("cover_date_en", "英文扉页应包含英文日期", model.date_en, "institution-format-guide 结构章节：英文扉页需包含日期。"),
        ("front_declaration", "前置部分应包含“独创性声明”", model.declaration_heading, "institution-format-guide 结构章节：声明页属于必备前置部分。"),
        ("front_authorization", "前置部分应包含“学位论文版权使用授权书”", model.authorization_heading, "institution-format-guide 结构章节：授权书属于必备前置部分。"),
        ("front_abstract_cn", "前置部分应包含中文摘要", model.abstract_cn_heading, "institution-format-guide 结构章节：摘要（中英文）属于必备前置部分。"),
        ("front_abstract_en", "前置部分应包含英文 Abstract", model.abstract_en_heading, "institution-format-guide 结构章节：摘要（中英文）属于必备前置部分。"),
        ("front_toc", "前置部分应包含目录", model.toc_heading, "institution-format-guide 结构章节：目录属于必备前置部分。"),
    ]

    missing = 0
    for check_type, expected, info, basis in required_fields:
        if info is not None:
            continue
        missing += 1
        _add_issue(
            issues,
            annotations,
            counter,
            location="前置部分",
            paragraph_index=fallback.index if fallback is not None else None,
            style_name=fallback.style_name if fallback is not None else "",
            check_type=check_type,
            expected=expected,
            actual="missing",
            source_text=fallback_text,
            problem=f"{expected}，当前文档未检测到对应对象。",
            fix=f"补齐 {expected.replace('应包含', '').replace('前置部分', '').strip('“”')}，并放在前置部分的正确位置。",
            basis=basis,
            severity="high" if check_type in {"front_abstract_cn", "front_abstract_en", "front_toc"} else "medium",
            mode="structure",
            anchor=fallback_anchor,
        )

    ordering_fail = 0
    order_chain = [
        ("独创性声明", model.declaration_heading),
        ("授权书", model.authorization_heading),
        ("中文摘要", model.abstract_cn_heading),
        ("英文摘要", model.abstract_en_heading),
        ("目录", model.toc_heading),
        ("正文起始章标题", model.first_body_heading),
    ]
    for (left_label, left_info), (right_label, right_info) in zip(order_chain, order_chain[1:]):
        if left_info is None or right_info is None or left_info.index is None or right_info.index is None:
            continue
        if left_info.index < right_info.index:
            continue
        ordering_fail += 1
        _add_issue(
            issues,
            annotations,
            counter,
            location=left_info.location,
            paragraph_index=left_info.index,
            style_name=left_info.style_name,
            check_type="front_matter_order",
            expected=f"{left_label} 位于 {right_label} 之前",
            actual=f"{left_label} 位于第{left_info.index}段，{right_label} 位于第{right_info.index}段",
            source_text=_evidence_from_text(left_info.text, attach_text=left_info.text[:40]),
            problem=f"前置部分顺序不符合规范：{left_label} 应位于 {right_label} 之前。",
            fix=f"调整前置部分顺序，使 {left_label} 位于 {right_label} 之前。",
            basis="institution-format-guide 结构章节：封面/声明/摘要/目录/正文需按固定顺序组织。",
            severity="medium",
            mode="structure",
            anchor=_body_anchor(left_info),
        )

    keyword_fail = 0
    for heading, keywords, label in [
        (model.abstract_cn_heading, model.keywords_cn, "中文摘要"),
        (model.abstract_en_heading, model.keywords_en, "英文摘要"),
    ]:
        if heading is None or keywords is not None:
            continue
        keyword_fail += 1
        _add_issue(
            issues,
            annotations,
            counter,
            location=heading.location,
            paragraph_index=heading.index,
            style_name=heading.style_name,
            check_type="abstract_keywords",
            expected=f"{label}后包含关键词行",
            actual="missing",
            source_text=heading.text,
            problem=f"{label}后未检测到关键词行。",
            fix=f"在 {label} 后补齐关键词行，并按示例格式列出关键词。",
            basis="institution-format-guide 摘要示例：中英文摘要后均应列出关键词。",
            severity="medium",
            mode="structure",
            anchor=_body_anchor(heading),
        )

    document_checks.append(
        CheckResult(
            checkType="front_matter_structure",
            status="pass" if missing == 0 and ordering_fail == 0 else "fail",
            basis="institution-format-guide 结构章节",
            details=f"必备字段缺失 {missing} 处，顺序问题 {ordering_fail} 处。",
        )
    )
    document_checks.append(
        CheckResult(
            checkType="abstract_keywords",
            status="pass" if keyword_fail == 0 else "fail",
            basis="institution-format-guide 摘要示例",
            details=f"关键词缺失 {keyword_fail} 处。",
        )
    )


def _apply_table_cell_explicit_checks(
    table_cells: List[ParagraphInfo],
    expected_map: Dict[str, Dict[str, Any]],
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    document_checks: List[CheckResult],
) -> None:
    max_size = expected_map.get("table_cell", {}).get("maxFontSizePt")
    basis = expected_map.get("table_cell", {}).get("basis", "institution-format-guide 4.2.2")
    hit = 0
    checked = 0
    for p in table_cells:
        if p.zone == "front_matter":
            continue
        checked += 1
        if max_size and p.font_size_pt and p.font_size_pt > max_size + 0.5:
            hit += 1
            _add_issue(
                issues,
                annotations,
                counter,
                location=p.location,
                paragraph_index=None,
                style_name=p.style_name,
                check_type="table_cell_font_size",
                expected=f"<= {max_size}pt",
                actual=p.font_size_pt,
                source_text=p.text,
                problem=f"表内文字字号不应大于图题字号：上限 {max_size}pt，当前 {p.font_size_pt}pt。",
                fix=f"将表内字号下调到不大于 {max_size}pt。",
                basis=basis,
                severity="medium",
                mode="explicit",
                anchor=_table_cell_anchor(p),
            )
    document_checks.append(CheckResult(checkType="table_cell_font_size", status="pass" if hit == 0 else "fail", basis=basis, details=f"检查 {checked} 个正文区表格单元段落，命中 {hit} 处；前置分区表格已排除。"))


def _apply_extra_space_checks(
    paragraphs: List[ParagraphInfo],
    table_cells: List[ParagraphInfo],
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    document_checks: List[CheckResult],
) -> None:
    items = [
        p for p in paragraphs
        if p.text.strip() and p.zone == "body" and p.paragraph_type in {"body", "keywords"}
    ]
    hit = 0
    basis = "第三份格式说明（后续澄清）：仅在正文区检查中文与英文单词之间的空格；目录/封面等前置部分不泛化处理。"
    for item in items:
        match = ZH_EN_WORD_SPACE_RE.search(item.text)
        if not match:
            continue
        hit += 1
        actual = match.group(0).replace('\u00a0', 'NBSP').replace('\u3000', '全角空格')
        _add_issue(
            issues,
            annotations,
            counter,
            location=item.location,
            paragraph_index=item.index,
            style_name=item.style_name,
            check_type="extra_space",
            expected="中文与英文单词之间无空格",
            actual=actual,
            source_text=_extra_space_evidence(item.text, match),
            problem="正文中中文与英文单词之间出现空格，不符合当前公开版保留的补充规则。",
            fix="删除中文与英文单词之间的空格，仅保留术语本身必要的字符。",
            basis=basis,
            severity="low",
            mode="explicit",
            anchor=_body_anchor(item),
        )
    document_checks.append(CheckResult(checkType="extra_space", status="pass" if hit == 0 else "fail", basis=basis, details=f"检查 {len(items)} 个正文段落，命中 {hit} 处。"))


def _apply_section_margin_checks(doc: Document, rules: Dict[str, Any], issues: List[FormatIssue], annotations: List[Dict[str, Any]], counter: List[int], document_checks: List[CheckResult], section_proxies: Dict[int, ParagraphInfo]) -> None:
    margin_rules = rules.get("margins", {})
    expected = {k: float(str(margin_rules.get(k, "25")).replace('mm', '')) for k in ["left", "right", "top", "bottom"] if margin_rules.get(k) is not None}
    if not expected:
        return
    mismatches = 0
    for idx, section in enumerate(doc.sections, start=1):
        actual = {
            "left": round(section.left_margin.mm, 2) if section.left_margin else None,
            "right": round(section.right_margin.mm, 2) if section.right_margin else None,
            "top": round(section.top_margin.mm, 2) if section.top_margin else None,
            "bottom": round(section.bottom_margin.mm, 2) if section.bottom_margin else None,
        }
        mismatched_parts = []
        for key, exp in expected.items():
            act = actual.get(key)
            if act is None:
                continue
            if abs(act - exp) > 1.0:
                mismatched_parts.append(f"{key}: {act}mm（应为 {exp}mm）")
        if mismatched_parts:
            mismatches += 1
            mismatch_summary = "; ".join(mismatched_parts)
            proxy = section_proxies.get(idx)
            anchor = _proxy_anchor("section_proxy", proxy, section_index=idx, object_type="page_margin") if proxy else None
            source_text = proxy.text if proxy else f"第{idx}节"
            _add_issue(
                issues,
                annotations,
                counter,
                location=f"第{idx}节页边距",
                paragraph_index=proxy.index if proxy else None,
                style_name=proxy.style_name if proxy else "",
                check_type="page_margin",
                expected=expected,
                actual=actual,
                source_text=source_text,
                problem=f"页边距不符合格式规范：第{idx}节存在 {len(mismatched_parts)} 处偏差（{mismatch_summary}）。",
                fix=f"将第{idx}节页边距整体调整到版芯 160mm×247mm 对应设置，重点修正：{mismatch_summary}。",
                basis="institution-format-guide 3 + historical-defect-checklist：页边距不符合规范是高频问题。",
                severity="high",
                mode="explicit",
                anchor=anchor,
            )
    document_checks.append(CheckResult(checkType="page_margin", status="pass" if mismatches == 0 else "fail", basis="institution-format-guide 3", details=f"检查 {len(doc.sections)} 个 section，命中 {mismatches} 个 section。"))


def _part_text(root: etree._Element) -> str:
    return "".join(root.xpath("//w:t/text()", namespaces=NS)).strip()


def _part_font_summary(root: etree._Element) -> Tuple[Optional[str], Optional[float], Optional[str]]:
    fonts: List[str] = []
    sizes: List[float] = []
    alignment: Optional[str] = None
    first_p = root.find(".//w:p", NS)
    if first_p is not None:
        jc = first_p.find("w:pPr/w:jc", NS)
        if jc is not None:
            alignment = _normalize_alignment(jc.get(qn("w:val")))
    for rpr in root.xpath("//w:rPr", namespaces=NS):
        fonts_dict = _get_rfonts(rpr)
        if fonts_dict.get("eastAsia"):
            fonts.append(fonts_dict["eastAsia"])
        size = _get_font_size_pt_from_rpr(rpr)
        if size is not None:
            sizes.append(size)
    return _dominant(fonts), _dominant([round(s, 2) for s in sizes]), alignment


def _header_line_value(root: etree._Element) -> Optional[str]:
    for paragraph in root.findall(".//w:p", NS):
        border = paragraph.find("w:pPr/w:pBdr", NS)
        if border is None:
            continue
        bottom = border.find("w:bottom", NS)
        if bottom is None:
            for child in border:
                bottom = child
                break
        if bottom is None:
            continue
        value = bottom.get(qn("w:val"))
        if value:
            return value
    return None


def _apply_header_footer_checks(
    docx_path: Path,
    rules: Dict[str, Any],
    issues: List[FormatIssue],
    annotations: List[Dict[str, Any]],
    counter: List[int],
    document_checks: List[CheckResult],
    section_starts: Dict[int, int],
    section_proxies: Dict[int, ParagraphInfo],
    section_zones: Dict[int, str],
) -> None:
    header_rule = rules.get("pageHeader", {})
    page_rule = rules.get("pageNumber", {})
    left_expected = str(header_rule.get("leftContent", "东北大学硕士学位论文")).replace("、博士", "")
    header_font_expected = str(header_rule.get("fontName", "楷体"))
    header_size_expected = cn_size_to_pt(header_rule.get("fontSize", "五号"))
    header_proxy_map, footer_proxy_map = _build_part_proxy_maps(docx_path, section_starts)
    header_required_zones = {"abstract_cn", "abstract_en", "toc", "body", "references", "acknowledgements", "achievements"}

    def header_anchor(name: Optional[str]) -> Tuple[Optional[ParagraphInfo], Optional[Dict[str, Any]]]:
        key = Path(name).name if name else '__missing_header__'
        section_idx = header_proxy_map.get(key, 1)
        proxy = section_proxies.get(section_idx) or next(iter(section_proxies.values()), None)
        anchor = _proxy_anchor('header_proxy', proxy, section_index=section_idx, header_part=key if name else None) if proxy else None
        return proxy, anchor

    def footer_anchor(name: Optional[str]) -> Tuple[Optional[ParagraphInfo], Optional[Dict[str, Any]]]:
        key = Path(name).name if name else '__missing_footer__'
        section_idx = footer_proxy_map.get(key, 1)
        proxy = section_proxies.get(section_idx) or next(iter(section_proxies.values()), None)
        anchor = _proxy_anchor('footer_proxy', proxy, section_index=section_idx, footer_part=key if name else None) if proxy else None
        return proxy, anchor

    with ZipFile(docx_path) as zf:
        header_names = sorted(name for name in zf.namelist() if name.startswith('word/header') and name.endswith('.xml'))
        footer_names = sorted(name for name in zf.namelist() if name.startswith('word/footer') and name.endswith('.xml'))
        header_fail = 0
        checked_header_parts = 0
        required_sections = {idx for idx, zone in section_zones.items() if zone in header_required_zones}
        if required_sections and not header_names:
            header_fail += 1
            proxy, anchor = header_anchor(None)
            _add_issue(
                issues, annotations, counter,
                location='文档级：页眉', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_header',
                expected='自摘要页起存在页眉', actual='未检测到任何 header part', source_text=proxy.text if proxy else '页眉代理锚点',
                problem='未检测到页眉定义，无法满足“自摘要页起加页眉”的规范要求。',
                fix='为摘要及其后续页面补齐页眉，并使用左固定右章题的页眉结构。',
                basis='institution-format-guide 3：自摘要页起加页眉。', severity='high', mode='explicit', anchor=anchor
            )
        for name in header_names:
            key = Path(name).name
            section_idx = header_proxy_map.get(key, 1)
            if section_zones.get(section_idx, "front_matter") not in header_required_zones:
                continue
            checked_header_parts += 1
            root = etree.fromstring(zf.read(name))
            text = _part_text(root)
            font, size, alignment = _part_font_summary(root)
            line_value = _header_line_value(root)
            proxy, anchor = header_anchor(name)
            proxy_text = proxy.text if proxy else text or Path(name).name
            if left_expected and left_expected not in text:
                header_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_header_text',
                    expected=f'包含“{left_expected}”', actual=text or '<empty>', source_text=proxy_text,
                    problem=f'页眉未包含规范要求的左端固定内容“{left_expected}”。',
                    fix='将页眉左端固定内容改为“东北大学硕士学位论文”，右端保留章号章题。',
                    basis='institution-format-guide 3：页眉左端固定为“东北大学硕士、博士学位论文”。',
                    severity='medium', mode='explicit', anchor=anchor
                )
            if font and header_font_expected not in font:
                header_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_header_font',
                    expected=header_font_expected, actual=font, source_text=proxy_text,
                    problem=f'页眉字体不符合规范：预期 {header_font_expected}，实际 {font}。',
                    fix=f'将页眉说明字体调整为 {header_font_expected}。',
                    basis='institution-format-guide 3：页眉说明 5 号楷体。',
                    severity='low', mode='explicit', anchor=anchor
                )
            if size and header_size_expected and not _nearly_equal(size, header_size_expected, 0.6):
                header_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_header_font_size',
                    expected=f'{header_size_expected}pt', actual=size, source_text=proxy_text,
                    problem=f'页眉字号不符合规范：预期 {header_size_expected}pt，实际 {size}pt。',
                    fix=f'将页眉字号调整为 {header_size_expected}pt。',
                    basis='institution-format-guide 3：页眉说明 5 号楷体。',
                    severity='low', mode='explicit', anchor=anchor
                )
            if line_value not in HEADER_LINE_ALLOWED:
                header_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_header_line',
                    expected='页眉包含单线/双线类线条', actual=line_value or '<missing>', source_text=proxy_text,
                    problem='页眉未检测到规范要求的单线/双线样式。',
                    fix='在页眉段落下边框补充单线或双线，并保持全篇一致。',
                    basis='institution-format-guide 3：页眉应采用单线或双线（二等线、文武线）样式。',
                    severity='low', mode='explicit', anchor=anchor
                )
        document_checks.append(CheckResult(checkType='page_header', status='pass' if header_fail == 0 else 'fail', basis='institution-format-guide 3', details=f'header parts={checked_header_parts}, 命中 {header_fail} 处。'))

        footer_fail = 0
        if not footer_names:
            footer_fail += 1
            proxy, anchor = footer_anchor(None)
            _add_issue(
                issues, annotations, counter,
                location='文档级：页码', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_number',
                expected='存在页脚页码', actual='未检测到任何 footer part', source_text=proxy.text if proxy else '页脚代理锚点',
                problem='未检测到页脚定义，页码规范无法成立。',
                fix='补齐页脚页码定义，并按页底居中、两侧修饰的格式设置。',
                basis='institution-format-guide 3 + historical-defect-checklist：摘要页与正文页码均需规范。',
                severity='high', mode='explicit', anchor=anchor
            )
        for name in footer_names:
            root = etree.fromstring(zf.read(name))
            xml_text = zf.read(name).decode('utf-8', errors='ignore')
            text = _part_text(root)
            font, size, alignment = _part_font_summary(root)
            proxy, anchor = footer_anchor(name)
            proxy_text = proxy.text if proxy else text or Path(name).name
            has_page_field = 'PAGE' in xml_text or bool(re.search(r'[-·－]\s*[IVXLCM0-9]+\s*[-·－]', text))
            if not has_page_field:
                footer_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_number_field',
                    expected='存在 PAGE 域或页码文本', actual=text or '<empty>', source_text=proxy_text,
                    problem='页脚中未检测到 PAGE 域或可识别页码文本。',
                    fix='在页脚插入 PAGE 域，并保留页码装饰格式。',
                    basis='institution-format-guide 3：页码应连续编页。',
                    severity='medium', mode='explicit', anchor=anchor
                )
            if alignment and alignment != 'center':
                footer_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_number_alignment',
                    expected='center', actual=alignment, source_text=proxy_text,
                    problem='页码未位于页底居中。',
                    fix='将页脚页码段落调整为居中。',
                    basis='institution-format-guide 3：页码页底居中。',
                    severity='low', mode='explicit', anchor=anchor
                )
            if text and not re.search(r'[·－-].+[·－-]', text):
                footer_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_number_decoration',
                    expected='数字两侧有圆点或横线修饰', actual=text, source_text=proxy_text,
                    problem='页码格式缺少两侧修饰符。',
                    fix='将页码格式调整为 ·3· 或 －3－ 一类修饰形式。',
                    basis='institution-format-guide 3 + historical-defect-checklist：页码格式需规范。',
                    severity='low', mode='explicit', anchor=anchor
                )
            if font and 'times' not in font.lower() and 'roman' not in font.lower() and 'minor' not in font.lower():
                footer_fail += 1
                _add_issue(
                    issues, annotations, counter,
                    location=f'文档级：{Path(name).name}', paragraph_index=proxy.index if proxy else None, style_name=proxy.style_name if proxy else '', check_type='page_number_font',
                    expected='Times New Roman / Roman', actual=font, source_text=proxy_text,
                    problem='页码字体不是 Roman 系字体。',
                    fix='将页码字体调整为 Times New Roman。',
                    basis='historical-defect-checklist：正文页码的字体应采用罗马体。',
                    severity='low', mode='explicit', anchor=anchor
                )
        document_checks.append(CheckResult(checkType='page_number', status='pass' if footer_fail == 0 else 'fail', basis='institution-format-guide 3 + historical-defect-checklist', details=f'footer parts={len(footer_names)}, 命中 {footer_fail} 处。'))


def _apply_toc_structure_checks(paragraphs: List[ParagraphInfo], issues: List[FormatIssue], annotations: List[Dict[str, Any]], counter: List[int], document_checks: List[CheckResult]) -> None:
    toc_entries = [p for p in paragraphs if p.paragraph_type == "toc_entry"]
    bad = 0
    for p in toc_entries:
        if not TOC_DOT_RE.search(p.text) or not re.search(r"\d+$", p.text):
            bad += 1
            _add_issue(
                issues, annotations, counter,
                location=p.location, paragraph_index=p.index, style_name=p.style_name, check_type="toc_entry_structure",
                expected="目录条目包含引导点和页码", actual=p.text[:120], source_text=p.text,
                problem="目录条目缺少完整的点线或页码结构。",
                fix="更新目录域或手工修正目录条目的点线与页码。",
                basis="institution-format-guide 目录示例 + historical-defect-checklist：目录排版需规范。",
                severity="medium", mode="structure", anchor=_body_anchor(p)
            )
    document_checks.append(CheckResult(checkType="toc_structure", status="pass" if bad == 0 else "fail", basis="institution-format-guide 目录示例", details=f"检查 {len(toc_entries)} 条目录条目，命中 {bad} 处。"))


def _apply_reference_checks(paragraphs: List[ParagraphInfo], issues: List[FormatIssue], annotations: List[Dict[str, Any]], counter: List[int], document_checks: List[CheckResult]) -> None:
    refs = [p for p in paragraphs if p.paragraph_type == "reference_entry"]
    bad = 0
    for p in refs:
        text = p.text.strip()
        if len(text) < 20:
            continue
        if not REFERENCE_ENTRY_RE.match(text):
            bad += 1
            _add_issue(
                issues, annotations, counter,
                location=p.location, paragraph_index=p.index, style_name=p.style_name, check_type="reference_numbering",
                expected="顺序号编号体系", actual=text[:30], source_text=_reference_tail_evidence(text),
                problem="参考文献未按顺序号编号体系呈现。",
                fix="统一改为顺序号编号体系，并保证列表顺序连续。",
                basis="institution-format-guide 4.5：参考文献采用顺序号编号体系。",
                severity="medium", mode="explicit", anchor=_body_anchor(p)
            )
            continue

        if "[J]" in text:
            if not PAGE_RANGE_RE.search(text):
                bad += 1
                _add_issue(
                    issues, annotations, counter,
                    location=p.location, paragraph_index=p.index, style_name=p.style_name, check_type="reference_page_range",
                    expected="期刊条目包含起止页码", actual=text[:120], source_text=_reference_tail_evidence(text),
                    problem="期刊参考文献缺少起止页码。",
                    fix="补齐期刊条目的起止页码。",
                    basis="institution-format-guide 4.5 + historical-defect-checklist：部分参考文献没有标页码。",
                    severity="medium", mode="structure", anchor=_body_anchor(p)
                )
            if not VOL_ISSUE_RE.search(text):
                bad += 1
                _add_issue(
                    issues, annotations, counter,
                    location=p.location, paragraph_index=p.index, style_name=p.style_name, check_type="reference_volume_issue",
                    expected="期刊条目包含卷(期)格式", actual=text[:120], source_text=_reference_tail_evidence(text),
                    problem="期刊参考文献缺少或未规范呈现卷(期)信息。",
                    fix="补齐并规范卷(期)格式。",
                    basis="institution-format-guide 4.5 + historical-defect-checklist：卷和期格式错误是高频问题。",
                    severity="medium", mode="structure", anchor=_body_anchor(p)
                )
        elif "[M]" in text:
            if not re.search(r"[:：].+出版社", text) and "Insecure" not in text:
                bad += 1
                _add_issue(
                    issues, annotations, counter,
                    location=p.location, paragraph_index=p.index, style_name=p.style_name, check_type="reference_book_info",
                    expected="专著条目含出版地与出版社", actual=text[:120], source_text=_reference_tail_evidence(text),
                    problem="专著参考文献信息不全，缺少出版地/出版社等关键字段。",
                    fix="按“出版地：出版社，年代，起止页码”补齐专著信息。",
                    basis="institution-format-guide 4.5：专著格式。",
                    severity="low", mode="structure", anchor=_body_anchor(p)
                )
        elif "[D]" in text:
            if not re.search(r"[，,]\s*[^，,]{2,20}[，,]\s*\d{4}", text):
                bad += 1
                _add_issue(
                    issues, annotations, counter,
                    location=p.location, paragraph_index=p.index, style_name=p.style_name, check_type="reference_thesis_info",
                    expected="学位论文条目含发表地/授予单位/年度", actual=text[:120], source_text=_reference_tail_evidence(text),
                    problem="学位论文参考文献信息不全。",
                    fix="按“发表地：学位授予单位，年度”补齐信息。",
                    basis="institution-format-guide 4.5：学位论文格式。",
                    severity="low", mode="structure", anchor=_body_anchor(p)
                )
    document_checks.append(CheckResult(checkType="reference_format", status="pass" if bad == 0 else "fail", basis="institution-format-guide 4.5 + historical-defect-checklist", details=f"检查 {len(refs)} 条参考文献，命中 {bad} 处。"))


def _paragraphs_by_index(paragraphs: List[ParagraphInfo]) -> Dict[int, ParagraphInfo]:
    return {p.index: p for p in paragraphs if p.index is not None}


def _apply_figure_table_formula_checks(doc: Document, paragraphs: List[ParagraphInfo], issues: List[FormatIssue], annotations: List[Dict[str, Any]], counter: List[int], document_checks: List[CheckResult]) -> None:
    para_map = _paragraphs_by_index(paragraphs)
    main_paragraphs = {id(p._element): idx for idx, p in enumerate(doc.paragraphs, start=1)}

    figure_adj_bad = 0
    table_adj_bad = 0
    formula_bad = 0

    body_items = list(_iter_block_items(doc))
    for i, item in enumerate(body_items):
        if isinstance(item, Paragraph):
            idx = main_paragraphs.get(id(item._element))
            info = para_map.get(idx) if idx else None
            text = (item.text or "").strip()
            if info and info.paragraph_type == "caption_table":
                next_nonempty = None
                for nxt in body_items[i + 1 :]:
                    if isinstance(nxt, Paragraph) and not (nxt.text or "").strip():
                        continue
                    next_nonempty = nxt
                    break
                if not isinstance(next_nonempty, Table):
                    table_adj_bad += 1
                    _add_issue(
                        issues, annotations, counter,
                        location=info.location, paragraph_index=info.index, style_name=info.style_name, check_type="table_caption_spacing",
                        expected="表题紧邻其下方表格", actual="下一个非空对象不是表格", source_text=info.text,
                        problem="表题与表格未保持规范的上下紧邻关系。",
                        fix="将表题移到对应表格正上方，并去掉中间多余空段。",
                        basis="institution-format-guide 4.2.2 + historical-defect-checklist：图表题目与图表间距不正确。",
                        severity="medium", mode="structure", anchor=_body_anchor(info)
                    )
            if info and info.paragraph_type == "caption_figure":
                prev_nonempty = None
                for prev in reversed(body_items[:i]):
                    if isinstance(prev, Paragraph) and not (prev.text or "").strip():
                        continue
                    prev_nonempty = prev
                    break
                prev_is_figure = isinstance(prev_nonempty, Paragraph) and _has_drawing(prev_nonempty)
                if not prev_is_figure:
                    figure_adj_bad += 1
                    _add_issue(
                        issues, annotations, counter,
                        location=info.location, paragraph_index=info.index, style_name=info.style_name, check_type="figure_caption_spacing",
                        expected="图题紧邻图的正下方", actual="前一个非空对象不是图片段落", source_text=info.text,
                        problem="图题未紧贴对应插图下方。",
                        fix="将图题挪到对应图片正下方，并去掉中间多余空段。",
                        basis="institution-format-guide 4.2.1 + historical-defect-checklist：图表题目与图表间距不正确。",
                        severity="medium", mode="structure", anchor=_body_anchor(info)
                    )
            if info and info.paragraph_type == "formula":
                if not FORMULA_NO_RE.search(text):
                    formula_bad += 1
                    _add_issue(
                        issues, annotations, counter,
                        location=info.location, paragraph_index=info.index, style_name=info.style_name, check_type="formula_numbering",
                        expected="式（x.y）编号", actual=text[:80], source_text=text or "公式段落",
                        problem="公式段落未检测到按章编号的式号。",
                        fix="为公式补齐按章序号的式号，并与公式同行居右排版。",
                        basis="institution-format-guide 4.2.3 + historical-defect-checklist：公式编号格式不统一。",
                        severity="medium", mode="structure", anchor=_body_anchor(info)
                    )
                if info.alignment and info.alignment != "center":
                    formula_bad += 1
                    _add_issue(
                        issues, annotations, counter,
                        location=info.location, paragraph_index=info.index, style_name=info.style_name, check_type="formula_alignment",
                        expected="center", actual=info.alignment, source_text=text or "公式段落",
                        problem="公式段落未居中排版。",
                        fix="将公式段落调整为居中。",
                        basis="institution-format-guide 4.2.3：公式一般单行居中排版。",
                        severity="low", mode="explicit", anchor=_body_anchor(info)
                    )
    document_checks.append(CheckResult(checkType="figure_caption_spacing", status="pass" if figure_adj_bad == 0 else "fail", basis="institution-format-guide 4.2.1 + historical-defect-checklist", details=f"命中 {figure_adj_bad} 处。"))
    document_checks.append(CheckResult(checkType="table_caption_spacing", status="pass" if table_adj_bad == 0 else "fail", basis="institution-format-guide 4.2.2 + historical-defect-checklist", details=f"命中 {table_adj_bad} 处。"))
    document_checks.append(CheckResult(checkType="formula_structure", status="pass" if formula_bad == 0 else "fail", basis="institution-format-guide 4.2.3 + historical-defect-checklist", details=f"命中 {formula_bad} 处。"))


def detect_format_issues(docx_path: str | Path, rules: Dict[str, Any]) -> Dict[str, Any]:
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    section_starts = _build_section_start_map(doc)
    paragraphs = _collect_main_paragraphs(doc)
    table_cells = _collect_table_cell_paragraphs(doc)
    _assign_table_cell_zones(doc, table_cells, paragraphs)
    section_proxies = _build_section_proxy_paragraphs(paragraphs, section_starts)
    section_zones = _build_section_zone_map(paragraphs, section_starts)
    expected_map = _get_expected_map(rules)
    issues: List[FormatIssue] = []
    annotations: List[Dict[str, Any]] = []
    document_checks: List[CheckResult] = []
    counter = [1]

    _apply_explicit_paragraph_checks(paragraphs, expected_map, issues, annotations, counter, document_checks)
    _apply_front_matter_checks(paragraphs, issues, annotations, counter, document_checks)
    _apply_table_cell_explicit_checks(table_cells, expected_map, issues, annotations, counter, document_checks)
    _apply_toc_structure_checks(paragraphs, issues, annotations, counter, document_checks)
    _apply_reference_checks(paragraphs, issues, annotations, counter, document_checks)
    _apply_figure_table_formula_checks(doc, paragraphs, issues, annotations, counter, document_checks)
    _apply_extra_space_checks(paragraphs, table_cells, issues, annotations, counter, document_checks)
    _apply_section_margin_checks(doc, rules, issues, annotations, counter, document_checks, section_proxies)
    _apply_header_footer_checks(docx_path, rules, issues, annotations, counter, document_checks, section_starts, section_proxies, section_zones)
    _apply_consistency_checks(_consistency_group_map(paragraphs, table_cells), issues, annotations, counter, document_checks)

    issue_dicts = [issue.to_dict() for issue in issues]
    document_check_dicts = [asdict(item) for item in document_checks]
    paragraph_type_counts = Counter(p.paragraph_type for p in paragraphs)
    zone_counts = Counter(p.zone for p in paragraphs)

    return {
        "meta": {
            "detector": "format_checker",
            "version": "3.3-mainline-evidence-frontmatter",
            "docxPath": str(docx_path),
            "rulesPath": str(FORMAT_RULES_PATH),
            "ruleMatrixPath": str(FORMAT_MATRIX_PATH),
            "sourceNotesPath": str(FORMAT_SOURCE_NOTES_PATH),
            "ruleSources": rules.get("_meta", {}).get("source", []),
            "paragraphCount": len(paragraphs),
            "tableCellParagraphCount": len(table_cells),
            "issueCount": len(issue_dicts),
            "documentCheckCount": len(document_check_dicts),
            "implementedModes": ["explicit", "structure", "consistency"],
            "implementedChecks": sorted({item["checkType"] for item in issue_dicts} | {item["checkType"] for item in document_check_dicts}),
            "anchorContract": {
                "version": "format-anchor-v2",
                "kinds": ["body_paragraph", "table_cell", "section_proxy", "header_proxy", "footer_proxy"],
                "machineSelector": "anchor.presentation_node_id / anchor.node_id",
            },
            "note": "规则来源仅限institution-format-guide + historical-defect-checklist；未给精确数值的项目按一致性模式检查。format-only 链路优先消费结构化 anchor，而不是靠 location/source_text 猜定位；本版补充前置部分对象化检查与细粒度格式证据提取。",
        },
        "coverage": {
            "zoneCounts": dict(zone_counts),
            "paragraphTypeCounts": dict(paragraph_type_counts),
            "checkedZones": [
                "front_matter",
                "abstract_cn",
                "abstract_en",
                "toc",
                "body",
                "references",
                "acknowledgements",
                "achievements",
            ],
            "checkedObjects": [
                "heading1",
                "heading2",
                "heading3",
                "heading4",
                "body",
                "keywords",
                "toc_entry",
                "caption_figure",
                "caption_table",
                "table_cell",
                "formula",
                "reference_entry",
                "page_margin",
                "page_header",
                "page_number",
            ],
        },
        "documentChecks": document_check_dicts,
        "issues": issue_dicts,
        "annotationItems": annotations,
    }


def check_docx_format(docx_path: str | Path, rules_path: Optional[str | Path] = None) -> Dict[str, Any]:
    rules = load_format_rules(rules_path)
    return detect_format_issues(docx_path, rules)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="检查 DOCX 格式是否符合学校论文规范")
    parser.add_argument("docx", help="目标 docx 路径")
    parser.add_argument("--rules", help="format-rules.json 路径")
    parser.add_argument("--out", help="输出 JSON 路径")
    args = parser.parse_args()

    result = check_docx_format(args.docx, args.rules)
    text = json.dumps(result, ensure_ascii=False, indent=2)
    if args.out:
        Path(args.out).write_text(text + "\n", encoding="utf-8")
    else:
        print(text)

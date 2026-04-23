# -*- coding: utf-8 -*-
"""
毕业设计Word文档格式修复工具
用法: python fix_word_format.py <输入文件> [输出文件]
"""
import sys
from pathlib import Path
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_document(input_path, output_path=None):
    """
    修复Word文档格式，使其符合毕业设计规范
    
    规范：
    - 正文: Times New Roman, 10.5pt, 首行缩进2字符
    - 一级标题: Times New Roman, 15pt
    - 二级标题: Arial, 14pt
    - 三级标题: Times New Roman, 12pt
    - 图表标题: 等线 Light, 10pt
    """
    if output_path is None:
        output_path = input_path.replace('.docx', '_格式修复.docx')
    
    print(f'📖 读取: {input_path}')
    doc = Document(input_path)
    
    # 修复样式定义
    print('🎨 修复样式...')
    style_fixes = {
        'Normal': {'font': 'Times New Roman', 'size': Pt(10.5)},
        'Heading 1': {'font': 'Times New Roman', 'size': Pt(15)},
        'Heading 2': {'font': 'Arial', 'size': Pt(14)},
        'Heading 3': {'font': 'Times New Roman', 'size': Pt(12)},
        'Caption': {'font': '等线 Light', 'size': Pt(10)},
    }
    
    for style_name, specs in style_fixes.items():
        try:
            style = doc.styles[style_name]
            style.font.name = specs['font']
            style.font.size = specs['size']
            print(f'  ✅ {style_name}: {specs["font"]} {specs["size"].pt}pt')
        except Exception as e:
            print(f'  ⚠️ {style_name}: {e}')
    
    # 修复正文段落
    print('📝 修复正文段落格式...')
    fixed_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else 'Normal'
        
        # 跳过标题和目录
        if style_name.startswith('Heading'):
            continue
        if '...' in text and any(c.isdigit() for c in text[-5:]):
            continue  # 目录项
        
        # 修复正文段落
        if len(text) > 20 and not text.startswith(('图', '表')):
            para.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
            para.paragraph_format.line_spacing = 1.5
            fixed_count += 1
    
    print(f'  ✅ 修复了 {fixed_count} 个正文段落')
    
    # 修复图表标题
    print('🖼️ 修复图表标题...')
    caption_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith(('图', '表')):
            try:
                para.style = doc.styles['Caption']
                caption_count += 1
            except:
                for run in para.runs:
                    run.font.name = '等线 Light'
                    run.font.size = Pt(10)
    
    print(f'  ✅ 修复了 {caption_count} 个图表标题')
    
    # 保存
    print(f'💾 保存: {output_path}')
    doc.save(output_path)
    print('✅ 完成！')
    return output_path

def main():
    if len(sys.argv) < 2:
        print('用法: python fix_word_format.py <输入文件> [输出文件]')
        print('示例: python fix_word_format.py 毕业设计.docx')
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    fix_document(input_path, output_path)

if __name__ == '__main__':
    main()

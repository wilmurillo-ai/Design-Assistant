#!/usr/bin/env python3
"""
Generate DOCX README for skill-priority-setup
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_readme():
    doc = Document()
    
    # Title
    title = doc.add_heading('Skill Priority Setup', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Automated Skill Tiering \u0026 Injection Policy for OpenClaw')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Overview
    doc.add_heading('Overview', 1)
    doc.add_paragraph(
        'This skill automates the setup of a tiered skill priority system for OpenClaw agents. '
        'It scans installed skills, suggests L0-L3 tier assignments, and configures injection '
        'policies to optimize token usage and response performance.'
    )
    
    # What It Does
    doc.add_heading('What This Skill Does', 1)
    
    phases = [
        ('1. Discovery Phase', 'Scans all installed skills across standard directories'),
        ('2. Analysis Phase', 'Detects skill types and dependencies'),
        ('3. Suggestion Phase', 'Proposes L0-L3 tier assignments'),
        ('4. Review Phase', 'Interactive confirmation or modification'),
        ('5. Configuration Phase', 'Applies tiered architecture to your setup'),
        ('6. Validation Phase', 'Verifies configuration and estimates token savings')
    ]
    
    for title, desc in phases:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(f': {desc}')
    
    # Architecture
    doc.add_heading('Architecture Overview', 1)
    
    tiers = [
        ('L0 - ROM Core (Always Active)', [
            'semantic-system: Semantic routing',
            'agent-evolution: Self-improvement behaviors',
            'config-modification: Config safety (on-demand)',
            'skill-safe-install: Installation safety (on-demand)'
        ]),
        ('L1 - Routing Layer (Task-Triggered)', [
            'browser-automation: Web automation',
            'find-skills: Skill discovery',
            'teamtask: Multi-agent coordination'
        ]),
        ('L2 - Domain Skills (Keyword-Triggered)', [
            'word-docx, tesseract-ocr: Document processing',
            'youtube-transcript: Media transcription',
            'discord, wechat-suite: Platform integration',
            'evomap, automation-workflows: Automation'
        ]),
        ('L3 - Extensions (On-Demand)', [
            'notion, slack, github',
            'All third-party integrations'
        ])
    ]
    
    for tier_title, items in tiers:
        doc.add_heading(tier_title, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Token Budget
    doc.add_heading('Token Budget', 1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Budget'
    
    # Data
    data = [
        ('L0 Core', '≤300 tokens/round'),
        ('L1 Triggered', '≤400 tokens per injection'),
        ('Total Budget', '≤900 tokens/round')
    ]
    
    for i, (layer, budget) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = layer
        row[1].text = budget
    
    doc.add_paragraph()
    
    # Quick Start
    doc.add_heading('Quick Start', 1)
    
    doc.add_heading('Interactive Setup', 2)
    doc.add_paragraph('python3 ~/.openclaw/workspace/skills/skill-priority-setup/scripts/setup.py')
    
    doc.add_heading('Auto Mode (Skip Review)', 2)
    doc.add_paragraph('python3 ~/.openclaw/workspace/skills/skill-priority-setup/scripts/setup.py --auto')
    
    doc.add_heading('Dry Run (No Changes)', 2)
    doc.add_paragraph('python3 ~/.openclaw/workspace/skills/skill-priority-setup/scripts/setup.py --dry-run')
    
    # Safety Features
    doc.add_heading('Safety Features', 1)
    
    safety_items = [
        'Automatic backup before any configuration changes',
        'JSON validation before applying',
        'Dry-run mode for testing',
        'Rollback capability with timestamped backups',
        'Token budget enforcement'
    ]
    
    for item in safety_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Files Created
    doc.add_heading('Files Created/Modified', 1)
    
    files = [
        ('SKILL_PRIORITY_POLICY.md', 'Your custom policy document'),
        ('AGENTS.md', 'ROM constraints added'),
        ('SOUL.md', 'Agent evolution behaviors'),
        ('~/.openclaw/backup/', 'Timestamped backups')
    ]
    
    for filename, desc in files:
        p = doc.add_paragraph()
        p.add_run(filename).bold = True
        p.add_run(f': {desc}')
    
    # When to Use
    doc.add_heading('When to Use This Skill', 1)
    
    use_cases = [
        'New Setup: Just installed multiple skills and want optimal configuration',
        'Performance Issues: High token usage or slow responses',
        'Migration: Upgrading from flat skill structure to tiered architecture',
        'Audit: Reviewing and optimizing existing skill priorities'
    ]
    
    for case in use_cases:
        doc.add_paragraph(case, style='List Bullet')
    
    # Workflow Diagram
    doc.add_heading('Setup Workflow', 1)
    
    workflow = '''
    ┌─────────────────┐
    │ 1. Scan Skills  │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 2. Analyze      │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 3. Suggest Tiers│
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 4. User Review  │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 5. Apply Config │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 6. Validate     │
    └─────────────────┘
    '''
    
    p = doc.add_paragraph()
    run = p.add_run(workflow)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Generated by skill-priority-setup • OpenClaw')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    output_path = '/Users/macmini/.openclaw/workspace/skills/skill-priority-setup/README.docx'
    doc.save(output_path)
    print(f'DOCX README created: {output_path}')

if __name__ == '__main__':
    create_readme()

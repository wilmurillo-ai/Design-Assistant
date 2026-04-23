#!/usr/bin/env python3
"""
生成中文版 DOCX README
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_chinese_readme():
    doc = Document()
    
    # 标题
    title = doc.add_heading('技能优先级配置工具', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('OpenClaw 自动化技能分层与注入策略配置')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # 概述
    doc.add_heading('工具概述', 1)
    doc.add_paragraph(
        '本工具为 OpenClaw 智能体提供自动化的技能优先级分层配置。'
        '它会扫描已安装的技能，智能建议 L0-L3 四层分级，并配置注入策略以优化 Token 使用和提升响应性能。'
    )
    
    # 核心功能
    doc.add_heading('核心功能', 1)
    
    phases = [
        ('1. 扫描发现', '扫描标准目录中的所有已安装技能'),
        ('2. 智能分析', '检测技能类型和依赖关系'),
        ('3. 分层建议', '智能推荐 L0-L3 层级分配'),
        ('4. 人工确认', '交互式确认或修改建议'),
        ('5. 自动配置', '应用分层架构到您的系统'),
        ('6. 验证优化', '验证配置并估算 Token 节省')
    ]
    
    for title, desc in phases:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(f'：{desc}')
    
    # 架构说明
    doc.add_heading('四层架构设计', 1)
    
    tiers = [
        ('L0 - ROM 核心层（始终激活）', [
            'semantic-system：语义路由与模型切换',
            'agent-evolution：自我进化与学习机制',
            'config-modification：配置安全（按需）',
            'skill-safe-install：安装安全审查（按需）'
        ]),
        ('L1 - 路由编排层（任务触发）', [
            'browser-automation：浏览器自动化',
            'find-skills：技能发现与安装',
            'teamtask：多代理协作流程'
        ]),
        ('L2 - 领域能力层（关键词触发）', [
            'word-docx、tesseract-ocr：文档处理',
            'youtube-transcript：媒体转录',
            'discord、wechat-suite：平台集成',
            'evomap、automation-workflows：自动化'
        ]),
        ('L3 - 扩展工具层（按需加载）', [
            'notion、slack、github',
            '所有第三方集成技能'
        ])
    ]
    
    for tier_title, items in tiers:
        doc.add_heading(tier_title, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Token 预算
    doc.add_heading('Token 预算控制', 1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '层级'
    hdr_cells[1].text = '预算'
    
    # 数据
    data = [
        ('L0 核心层', '每轮 ≤300 tokens'),
        ('L1 触发层', '每次注入 ≤400 tokens'),
        ('总预算上限', '每轮 ≤900 tokens')
    ]
    
    for i, (layer, budget) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = layer
        row[1].text = budget
    
    doc.add_paragraph()
    
    # 快速开始
    doc.add_heading('快速开始', 1)
    
    doc.add_heading('交互式配置（推荐）', 2)
    doc.add_paragraph('python3 ~/.openclaw/workspace/skills/skill-priority-setup/scripts/setup.py')
    
    doc.add_heading('自动模式（使用默认建议）', 2)
    doc.add_paragraph('python3 ~/.openclaw/workspace/skills/skill-priority-setup/scripts/setup.py --auto')
    
    doc.add_heading('预览模式（不实际修改）', 2)
    doc.add_paragraph('python3 ~/.openclaw/workspace/skills/skill-priority-setup/scripts/setup.py --dry-run')
    
    # 安全特性
    doc.add_heading('安全特性', 1)
    
    safety_items = [
        '任何配置更改前自动创建备份',
        '应用前进行 JSON 语法验证',
        '支持 --dry-run 预览模式',
        '支持回滚到时间戳备份',
        '强制执行 Token 预算限制'
    ]
    
    for item in safety_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 生成文件
    doc.add_heading('生成/修改的文件', 1)
    
    files = [
        ('SKILL_PRIORITY_POLICY.md', '您的自定义策略文档'),
        ('AGENTS.md', '添加 ROM 核心约束'),
        ('SOUL.md', '添加智能体进化行为'),
        ('~/.openclaw/backup/', '带时间戳的配置备份')
    ]
    
    for filename, desc in files:
        p = doc.add_paragraph()
        p.add_run(filename).bold = True
        p.add_run(f'：{desc}')
    
    # 使用场景
    doc.add_heading('适用场景', 1)
    
    use_cases = [
        '新环境搭建：刚安装多个技能，需要最优配置',
        '性能优化：Token 使用过高或响应缓慢',
        '架构升级：从扁平技能结构迁移到分层架构',
        '定期审计：审查和优化现有技能优先级'
    ]
    
    for case in use_cases:
        doc.add_paragraph(case, style='List Bullet')
    
    # 工作流程图
    doc.add_heading('配置工作流程', 1)
    
    workflow = '''
    ┌─────────────────┐
    │ 1. 扫描技能      │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 2. 分析类型      │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 3. 建议分层      │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 4. 人工确认      │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 5. 应用配置      │
    └────────┬────────┘
             ▼
    ┌─────────────────┐
    │ 6. 验证结果      │
    └─────────────────┘
    '''
    
    p = doc.add_paragraph()
    run = p.add_run(workflow)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # 自适应特性
    doc.add_heading('自适应特性', 1)
    doc.add_paragraph(
        '本工具设计为自适应配置：'
    )
    adaptive = [
        '适用于任何技能集合',
        '分析您的具体已安装技能',
        '提供定制化分层建议',
        '非破坏性操作（始终创建备份）'
    ]
    for item in adaptive:
        doc.add_paragraph(item, style='List Bullet')
    
    # 安装后效果
    doc.add_heading('安装后效果', 1)
    doc.add_paragraph('运行 setup.py 后，您的系统将拥有：')
    effects = [
        'SKILL_PRIORITY_POLICY.md - 您的自定义策略文档',
        '更新的 AGENTS.md - ROM 核心约束',
        '更新的消息注入器配置',
        '~/.openclaw/backup/ 中的完整备份'
    ]
    for effect in effects:
        doc.add_paragraph(effect, style='List Bullet')
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('由 skill-priority-setup 生成 • OpenClaw')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存
    output_path = '/Users/macmini/.openclaw/workspace/skills/skill-priority-setup/README_CN.docx'
    doc.save(output_path)
    print(f'中文版 DOCX README 已创建: {output_path}')

if __name__ == '__main__':
    create_chinese_readme()

#!/usr/bin/env python3
"""生成 heartbeat-ollama-guard 中文 README DOCX"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).parent / "README_heartbeat-ollama-guard_中文说明.docx"


def add_table_row(table, col0, col1, bold0=False):
    row = table.add_row()
    c0 = row.cells[0]
    c1 = row.cells[1]
    run0 = c0.paragraphs[0].add_run(col0)
    run0.bold = bold0
    c1.paragraphs[0].add_run(col1)
    return row


def create():
    doc = Document()

    # ── 标题 ──────────────────────────────────────────────────────────
    h = doc.add_heading("heartbeat-ollama-guard", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("OpenClaw 本地心跳守卫 · v1.0.0")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 80, 80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("作者: halfmoon82   |   发布: 2026-03-19   |   许可: MIT-0")
    rm.font.size = Pt(9)
    rm.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_paragraph()

    # ── 背景 ──────────────────────────────────────────────────────────
    doc.add_heading("背景与动机", 1)
    doc.add_paragraph(
        "OpenClaw 的心跳机制每 30 分钟调用一次 LLM，默认走云端模型（kimi / claude 等）。"
        "在 token 限额有限的情况下，这会白白消耗配额，严重时耗尽每周限额导致所有 agent 停摆。\n\n"
        "本技能将所有 OpenClaw 实例的心跳切换为本地 Ollama 模型，并部署一个配置守卫："
        "一旦检测到未经授权的修改，立即回滚并触发系统通知，从根本上消除心跳带来的 token 消耗风险。"
    )

    # ── 核心能力 ─────────────────────────────────────────────────────
    doc.add_heading("核心能力", 1)
    features = [
        ("💰 零成本心跳", "用本地 Ollama 替代付费云端模型，彻底消除心跳 token 消耗"),
        ("🛡️ 配置守卫", "60 秒轮询，检测到未授权修改立即回滚并发出 macOS/系统通知"),
        ("📦 一键安装向导", "自动检测 Ollama、拉取模型、配置所有 openclaw 实例、部署 LaunchAgent"),
        ("🔒 授权修改流程", "先改 conf.json 签字授权，再改 openclaw.json，守卫自动放行，无需关闭"),
        ("🔁 幂等部署", "重复运行安全，已是目标值时跳过，守卫版本相同时不覆盖"),
        ("🐧 跨平台支持", "macOS（LaunchAgent）+ Linux（systemd user service）"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].paragraphs[0].add_run("能力").bold = True
    hdr[1].paragraphs[0].add_run("说明").bold = True
    for feat, desc in features:
        add_table_row(t, feat, desc, bold0=True)

    doc.add_paragraph()

    # ── 快速开始 ─────────────────────────────────────────────────────
    doc.add_heading("快速开始", 1)

    doc.add_paragraph("第一步：安装 Ollama（如已安装请跳过）")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("macOS：").bold = True
    p.add_run("  brew install ollama  或访问 https://ollama.com 下载 App")

    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run("Linux：").bold = True
    p2.add_run("  curl -fsSL https://ollama.com/install.sh | sh")

    doc.add_paragraph()
    doc.add_paragraph("第二步：运行安装向导")
    code = doc.add_paragraph()
    code_run = code.add_run(
        "cd ~/.openclaw/workspace/skills/heartbeat-ollama-guard\n"
        "python3 heartbeat_ollama_guard.py --setup"
    )
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(10)

    # ── CLI 命令 ─────────────────────────────────────────────────────
    doc.add_heading("CLI 命令", 1)

    cmds = [
        ("--setup", "完整安装向导（检测 Ollama → 拉取模型 → 配置 openclaw.json → 部署守卫）"),
        ("--status", "查看 Ollama、模型、守卫进程、所有实例的当前状态"),
        ("--check", "执行一次守卫检查（不循环，适合排查问题）"),
        ("--uninstall", "卸载守卫（停止进程 + 删除 LaunchAgent + 删除脚本与配置）"),
        ("--model <id>", "指定本地模型 ID（默认 qwen3.5:4b-q4_K_M）"),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Light Grid Accent 1"
    h2 = t2.rows[0].cells
    h2[0].paragraphs[0].add_run("命令").bold = True
    h2[1].paragraphs[0].add_run("说明").bold = True
    for cmd, desc in cmds:
        add_table_row(t2, cmd, desc, bold0=True)

    doc.add_paragraph()

    # ── 安装向导步骤 ─────────────────────────────────────────────────
    doc.add_heading("安装向导步骤", 1)
    steps = [
        ("Step 1", "检测 Ollama 是否已安装（未安装则打印安装指引并退出）"),
        ("Step 2", "检测目标模型，未拉取则自动执行 ollama pull（流式进度）"),
        ("Step 3", "发现所有 openclaw.json 实例，确认需要配置哪些"),
        ("Step 4", "写入 heartbeat.model（自动备份原文件到 .hog_backups/）"),
        ("Step 5", "生成守卫脚本、conf.json，部署 LaunchAgent（macOS）或 systemd（Linux）"),
        ("Step 6", "验证守卫进程运行 + 单次 --check"),
        ("Step 7", "提示重启 gateway：openclaw gateway restart"),
    ]
    for step, desc in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{step}：").bold = True
        p.add_run(desc)

    doc.add_paragraph()

    # ── 授权修改流程 ─────────────────────────────────────────────────
    doc.add_heading("授权修改 heartbeat.model", 1)
    doc.add_paragraph(
        "守卫会阻止任何未经授权的 heartbeat.model 修改。如需合法更换模型，按以下流程操作："
    )
    auth_steps = [
        "先修改 ~/.openclaw/workspace/.lib/heartbeat-guard.conf.json 中对应实例的 expected 值（视为签字授权）",
        "再修改 openclaw.json 中的 heartbeat.model",
        "守卫在下次 60 秒轮询时发现 conf 与 openclaw.json 一致 → 自动放行，无需重启守卫",
    ]
    for s in auth_steps:
        doc.add_paragraph(s, style="List Number")

    doc.add_paragraph()

    # ── 安全声明 ─────────────────────────────────────────────────────
    doc.add_heading("安全声明", 1)
    security = [
        ("操作", "范围"),
        ("读取 openclaw.json", "仅检测 heartbeat.model 现状"),
        ("写入 openclaw.json", "仅 heartbeat.model + models.providers.local 字段"),
        ("守卫守护进程", "纯本地，60s 轮询，无网络请求"),
        ("macOS 系统通知", "仅守卫检测到未授权改动时触发"),
        ("需要 sudo", "❌ 不需要"),
        ("读取对话内容", "❌ 不读取"),
        ("访问外部 API", "❌ 不访问"),
    ]
    t3 = doc.add_table(rows=len(security), cols=2)
    t3.style = "Light Grid Accent 1"
    for i, (col0, col1) in enumerate(security):
        row = t3.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(col0)
        r1 = row.cells[1].paragraphs[0].add_run(col1)
        if i == 0:
            r0.bold = True
            r1.bold = True

    doc.add_paragraph()

    # ── 文件位置 ─────────────────────────────────────────────────────
    doc.add_heading("文件位置", 1)
    files = [
        ("~/.openclaw/workspace/.lib/heartbeat-guard.py", "守卫守护进程脚本（由安装向导生成）"),
        ("~/.openclaw/workspace/.lib/heartbeat-guard.conf.json", "守卫授权配置（含所有受保护实例）"),
        ("~/.openclaw/workspace/.lib/heartbeat-guard.log", "守卫运行日志"),
        ("~/Library/LaunchAgents/com.openclaw.heartbeat-guard.plist", "macOS LaunchAgent 配置"),
        ("~/.config/systemd/user/openclaw-heartbeat-guard.service", "Linux systemd 服务"),
        ("~/.openclaw/workspace/.lib/.hog_backups/", "openclaw.json 自动备份目录"),
    ]
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = "Light Grid Accent 1"
    h4 = t4.rows[0].cells
    h4[0].paragraphs[0].add_run("文件").bold = True
    h4[1].paragraphs[0].add_run("说明").bold = True
    for f, d in files:
        r = t4.add_row()
        run_f = r.cells[0].paragraphs[0].add_run(f)
        run_f.font.name = "Courier New"
        run_f.font.size = Pt(9)
        r.cells[1].paragraphs[0].add_run(d)

    doc.add_paragraph()

    # ── 页脚 ─────────────────────────────────────────────────────────
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run("heartbeat-ollama-guard v1.0.0  ·  halfmoon82  ·  MIT-0  ·  2026-03-19")
    rf.font.size = Pt(9)
    rf.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(str(OUT))
    print(f"✅ DOCX 已生成: {OUT}")
    return OUT


if __name__ == "__main__":
    create()

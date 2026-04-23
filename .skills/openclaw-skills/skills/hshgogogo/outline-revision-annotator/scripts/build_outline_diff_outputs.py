#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


KIND_STYLE = {
    "add": {
        "stroke": (1.0, 0.65, 0.0),
        "fill": (1.0, 0.96, 0.55),
        "tab": (1.0, 0.65, 0.0),
    },
    "rewrite": {
        "stroke": (0.94, 0.35, 0.1),
        "fill": (1.0, 0.86, 0.72),
        "tab": (0.94, 0.35, 0.1),
    },
}

WINDOWS_FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\msyh.ttf"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]

MAC_FONT_CANDIDATES = [
    Path("/System/Library/Fonts/PingFang.ttc"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
    Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
    Path("/Library/Fonts/Arial Unicode.ttf"),
]

LINUX_FONT_CANDIDATES = [
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract outline sources and render visible structural-change annotations on PDF."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract = subparsers.add_parser("extract", help="Extract plain text and PDF block maps.")
    extract.add_argument("--new-pdf", required=True)
    extract.add_argument("--old-source")
    extract.add_argument("--old-docx")
    extract.add_argument("--output-dir", required=True)

    render = subparsers.add_parser("render", help="Render annotated PDF and report outputs.")
    render.add_argument("--new-pdf", required=True)
    render.add_argument("--old-source")
    render.add_argument("--old-docx")
    render.add_argument("--manifest-json", required=True)
    render.add_argument("--report-md", required=True)
    render.add_argument("--output-dir", required=True)
    render.add_argument("--annotated-pdf-name", default="annotated_outline.pdf")
    render.add_argument("--report-md-name", default="compare_report.md")
    render.add_argument("--report-docx-name", default="compare_report.docx")
    render.add_argument("--note-width", type=float, default=205.0)
    render.add_argument("--extra-page-width", type=float, default=245.0)
    return parser.parse_args()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def find_cjk_font() -> str | None:
    candidates = WINDOWS_FONT_CANDIDATES + MAC_FONT_CANDIDATES + LINUX_FONT_CANDIDATES
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def read_docx_text(docx_path: Path) -> str:
    document = Document(str(docx_path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_whitespace(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def resolve_old_source(args: argparse.Namespace) -> Path:
    chosen = getattr(args, "old_source", None) or getattr(args, "old_docx", None)
    if not chosen:
        raise ValueError("Provide --old-source <path> or legacy --old-docx <path>.")
    return Path(chosen)


def read_source_text(source_path: Path) -> str:
    suffix = source_path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(source_path)
    if suffix == ".pdf":
        pdf_text, _ = extract_pdf_blocks(source_path)
        return pdf_text
    if suffix in {".txt", ".md"}:
        return read_text(source_path)
    raise ValueError(f"Unsupported old source type: {source_path.suffix}. Expected .docx, .pdf, .txt, or .md")


def extract_pdf_blocks(pdf_path: Path) -> tuple[str, list[dict]]:
    pdf = fitz.open(str(pdf_path))
    full_text_pages: list[str] = []
    page_maps: list[dict] = []

    for page_number, page in enumerate(pdf, start=1):
        page_text = page.get_text("text").strip()
        full_text_pages.append(page_text)
        blocks = []
        for block_index, block in enumerate(page.get_text("blocks")):
            x0, y0, x1, y1, text = block[:5]
            clean_text = text.strip()
            if not clean_text:
                continue
            blocks.append(
                {
                    "index": block_index,
                    "bbox": [round(x0, 1), round(y0, 1), round(x1, 1), round(y1, 1)],
                    "preview": normalize_whitespace(clean_text)[:180],
                    "text": clean_text,
                }
            )
        page_maps.append({"page": page_number, "blocks": blocks})

    pdf.close()
    return "\n\n".join(full_text_pages).strip(), page_maps


def command_extract(args: argparse.Namespace) -> int:
    new_pdf = Path(args.new_pdf)
    old_source = resolve_old_source(args)
    output_dir = Path(args.output_dir)
    ensure_dir(output_dir)

    pdf_text, pdf_blocks = extract_pdf_blocks(new_pdf)
    old_source_text = read_source_text(old_source)

    write_text(output_dir / "new_pdf_full_text.txt", pdf_text)
    write_text(output_dir / "old_source_full_text.txt", old_source_text)
    write_text(output_dir / "new_pdf_blocks.json", json.dumps(pdf_blocks, ensure_ascii=False, indent=2))
    write_text(
        output_dir / "source_files.json",
        json.dumps(
            {
                "new_pdf": str(new_pdf),
                "old_source": str(old_source),
                "generated_files": [
                    "new_pdf_full_text.txt",
                    "old_source_full_text.txt",
                    "new_pdf_blocks.json",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
    )

    print(f"Wrote extraction files to {output_dir}")
    return 0


def load_manifest(manifest_path: Path) -> tuple[list[dict], int]:
    raw = json.loads(read_text(manifest_path))
    if not isinstance(raw, list):
        raise ValueError("annotation manifest must be a JSON array")

    page_base = 1
    if any(isinstance(item, dict) and item.get("page") == 0 for item in raw):
        page_base = 0

    normalized: list[dict] = []
    for item in raw:
        if not isinstance(item, dict):
            raise ValueError("each manifest item must be an object")
        required = {"id", "page", "blocks", "kind", "title", "note", "box_text"}
        missing = sorted(required - set(item))
        if missing:
            raise ValueError(f"manifest item missing required field(s): {', '.join(missing)}")
        if item["kind"] not in KIND_STYLE:
            raise ValueError(f"unsupported kind '{item['kind']}', expected one of: {', '.join(KIND_STYLE)}")
        if not isinstance(item["blocks"], list) or not item["blocks"]:
            raise ValueError(f"manifest item {item['id']} must define a non-empty blocks list")

        normalized.append(
            {
                "id": str(item["id"]),
                "page": int(item["page"]),
                "blocks": [int(value) for value in item["blocks"]],
                "kind": item["kind"],
                "title": str(item["title"]),
                "note": str(item["note"]),
                "box_text": str(item["box_text"]),
                "box_height": float(item.get("box_height", 0) or 0),
            }
        )
    return normalized, page_base


def union_rect_for_blocks(blocks: list[tuple], indexes: list[int]) -> fitz.Rect:
    rect: fitz.Rect | None = None
    for block_index in indexes:
        try:
            block = blocks[block_index]
        except IndexError as exc:
            raise IndexError(f"block index {block_index} is out of range for page") from exc
        block_rect = fitz.Rect(block[:4])
        rect = block_rect if rect is None else rect | block_rect
    if rect is None:
        raise ValueError("could not resolve a rectangle for the given blocks")
    return rect


def draw_highlight(page: fitz.Page, rect: fitz.Rect, kind: str) -> None:
    style = KIND_STYLE[kind]
    shape = page.new_shape()
    shape.draw_rect(rect)
    shape.finish(
        color=style["stroke"],
        fill=style["fill"],
        width=1.0,
        fill_opacity=0.18,
        stroke_opacity=0.55,
    )
    shape.commit(overlay=True)


def draw_id_tab(page: fitz.Page, rect: fitz.Rect, item_id: str, kind: str) -> None:
    style = KIND_STYLE[kind]
    shape = page.new_shape()
    shape.draw_rect(rect)
    shape.finish(
        color=style["tab"],
        fill=style["tab"],
        width=0.8,
        fill_opacity=0.95,
        stroke_opacity=0.95,
    )
    shape.commit(overlay=True)
    page.insert_textbox(
        rect,
        item_id,
        fontname="helv",
        fontsize=9,
        color=(1, 1, 1),
        align=1,
        overlay=True,
    )


def estimate_box_height(text: str, minimum: float) -> float:
    lines = wrap_box_text(text)
    estimated = 22 + len(lines) * 14
    return max(minimum, estimated)


def wrap_box_text(text: str, max_chars: int = 21) -> list[str]:
    wrapped: list[str] = []
    for raw_paragraph in text.splitlines():
        paragraph = raw_paragraph.strip()
        if not paragraph:
            wrapped.append("")
            continue
        remaining = paragraph
        while remaining:
            wrapped.append(remaining[:max_chars])
            remaining = remaining[max_chars:]
    return wrapped or [""]


def draw_note_box(
    page: fitz.Page,
    x0: float,
    y0: float,
    width: float,
    text: str,
    kind: str,
    fontname: str,
    minimum_height: float,
) -> float:
    style = KIND_STYLE[kind]
    lines = wrap_box_text(text)
    height = max(minimum_height or 64.0, 22 + len(lines) * 14)
    rect = fitz.Rect(x0, y0, x0 + width, y0 + height)
    shape = page.new_shape()
    shape.draw_rect(rect)
    shape.finish(
        color=style["stroke"],
        fill=(1.0, 0.95, 0.15),
        width=1.0,
        fill_opacity=0.95,
        stroke_opacity=0.95,
    )
    shape.commit(overlay=True)
    baseline = rect.y0 + 18
    for index, line in enumerate(lines):
        page.insert_text(
            fitz.Point(rect.x0 + 8, baseline + index * 13),
            line,
            fontname=fontname,
            fontsize=10.5,
            color=(0.1, 0.1, 0.1),
            overlay=True,
        )
    return height


def choose_report_font_family() -> str:
    if sys.platform.startswith("win"):
        return "Microsoft YaHei"
    if sys.platform == "darwin":
        return "PingFang SC"
    return "Noto Sans CJK SC"


def style_run(run, font_family: str, size: float | None = None, bold: bool = False) -> None:
    run.font.name = font_family
    run_properties = run._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:eastAsia"), font_family)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def set_document_defaults(document: Document) -> None:
    font_family = choose_report_font_family()
    normal = document.styles["Normal"]
    normal.font.name = font_family
    run_properties = normal.element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:eastAsia"), font_family)
    normal.font.size = Pt(10.5)


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|\s*[-:| ]+\|?\s*$", line))


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    document = Document()
    set_document_defaults(document)
    font_family = choose_report_font_family()
    lines = read_text(md_path).splitlines()
    index = 0

    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()

        if not stripped:
            index += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", raw)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            paragraph = document.add_heading("", level=level)
            run = paragraph.add_run(text)
            style_run(run, font_family, size=16 - (level * 2), bold=True)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
            table_lines = [stripped]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1

            rows = [parse_table_row(line) for line in table_lines]
            header = rows[0]
            body = rows[1:]
            table = document.add_table(rows=len(body) + 1, cols=len(header))
            table.style = "Table Grid"
            for col, value in enumerate(header):
                cell = table.cell(0, col)
                cell.text = ""
                run = cell.paragraphs[0].add_run(value)
                style_run(run, font_family, bold=True)
            for row_index, row_values in enumerate(body, start=1):
                for col, value in enumerate(row_values):
                    cell = table.cell(row_index, col)
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(value)
                    style_run(run, font_family)
            continue

        bullet_match = re.match(r"^-\s+(.*)$", raw)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(bullet_match.group(1).strip())
            style_run(run, font_family)
            index += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", raw)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(numbered_match.group(1).strip())
            style_run(run, font_family)
            index += 1
            continue

        paragraph = document.add_paragraph()
        run = paragraph.add_run(stripped)
        style_run(run, font_family)
        index += 1

    document.save(str(docx_path))


def command_render(args: argparse.Namespace) -> int:
    new_pdf = Path(args.new_pdf)
    old_source = resolve_old_source(args)
    manifest_path = Path(args.manifest_json)
    report_md = Path(args.report_md)
    output_dir = Path(args.output_dir)
    ensure_dir(output_dir)

    manifest, page_base = load_manifest(manifest_path)
    fontfile = find_cjk_font()
    pdf = fitz.open(str(new_pdf))

    if fontfile is None and any(any(ord(char) > 127 for char in item["box_text"]) for item in manifest):
        raise RuntimeError("No CJK font found. Install a Chinese-capable font or provide an environment with one.")

    manifest_out: list[dict] = []
    items_by_page: dict[int, list[dict]] = {}
    for item in manifest:
        page_index = item["page"] if page_base == 0 else item["page"] - 1
        items_by_page.setdefault(page_index, []).append(item)

    for page_index, items in items_by_page.items():
        page = pdf.load_page(page_index)
        original_rect = fitz.Rect(page.rect)
        page.set_mediabox(
            fitz.Rect(0, 0, original_rect.width + args.extra_page_width, original_rect.height)
        )

        fontname = "helv"
        if fontfile:
            page.insert_font(fontname="CJK", fontfile=fontfile)
            fontname = "CJK"

        blocks = page.get_text("blocks")
        resolved = []
        for item in items:
            rect = union_rect_for_blocks(blocks, item["blocks"])
            resolved.append((item, rect))
        resolved.sort(key=lambda pair: pair[1].y0)

        previous_bottom = 20.0
        for item, rect in resolved:
            draw_highlight(page, rect, item["kind"])

            note_x = original_rect.width + 32
            note_y = max(20.0, rect.y0)
            minimum_height = item["box_height"] or 64.0
            estimated_height = estimate_box_height(item["box_text"], minimum_height)
            if note_y < previous_bottom + 10:
                note_y = previous_bottom + 10
            max_y = page.rect.height - estimated_height - 20
            note_y = min(note_y, max_y)

            draw_id_tab(
                page,
                fitz.Rect(original_rect.width + 6, note_y, original_rect.width + 28, note_y + 18),
                item["id"],
                item["kind"],
            )
            used_height = draw_note_box(
                page,
                note_x,
                note_y,
                args.note_width,
                item["box_text"],
                item["kind"],
                fontname,
                minimum_height,
            )
            previous_bottom = note_y + used_height

            manifest_out.append(
                {
                    "id": item["id"],
                    "page": page_index + 1,
                    "kind": item["kind"],
                    "title": item["title"],
                    "note": item["note"],
                    "blocks": item["blocks"],
                    "rect": [round(rect.x0, 1), round(rect.y0, 1), round(rect.x1, 1), round(rect.y1, 1)],
                }
            )

    annotated_pdf = output_dir / args.annotated_pdf_name
    report_md_out = output_dir / args.report_md_name
    report_docx_out = output_dir / args.report_docx_name
    normalized_manifest_out = output_dir / "annotation_manifest.normalized.json"
    source_info_out = output_dir / "source_files.json"

    pdf.save(str(annotated_pdf))
    pdf.close()
    shutil.copy2(report_md, report_md_out)
    markdown_to_docx(report_md, report_docx_out)
    write_text(normalized_manifest_out, json.dumps(manifest_out, ensure_ascii=False, indent=2))
    write_text(
        source_info_out,
        json.dumps({"new_pdf": str(new_pdf), "old_source": str(old_source)}, ensure_ascii=False, indent=2),
    )

    print(f"Wrote annotated PDF: {annotated_pdf}")
    print(f"Wrote markdown report: {report_md_out}")
    print(f"Wrote DOCX report: {report_docx_out}")
    print(f"Wrote normalized manifest: {normalized_manifest_out}")
    return 0


def main() -> int:
    args = parse_args()
    if args.command == "extract":
        return command_extract(args)
    if args.command == "render":
        return command_render(args)
    raise ValueError(f"unsupported command {args.command}")


if __name__ == "__main__":
    raise SystemExit(main())

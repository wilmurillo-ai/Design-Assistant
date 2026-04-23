#!/usr/bin/env python3
"""
机关公文格式化工具 v3
- 支持调研报告、公文、通用文章三种文档类型自动识别
- 自动加载参考规范文件（文件格式要求文件夹）
- 按 2024写材料格式模板 排版
- 数字/英文统一用 Times New Roman
"""

import re
import os
import sys
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==================== 路径常量 ====================
DESKTOP_FOLDER = os.path.expanduser("~/Desktop/openclaw专属文件夹")
OUTPUT_FOLDER = DESKTOP_FOLDER

# ==================== 格式常量 ====================
TOP_MARGIN = Cm(3.7)
BOTTOM_MARGIN = Cm(3.0)
LEFT_MARGIN = Cm(2.8)
RIGHT_MARGIN = Cm(2.8)
HEADER_DISTANCE = Cm(1.5)
FOOTER_DISTANCE = Cm(2.2)

TITLE_SIZE = Pt(22)
SUBTITLE_SIZE = Pt(16)
BODY_SIZE = Pt(16)
PAGENUM_SIZE = Pt(14)

FONT_FANGZHENG_XIAOBIAO = "方正小标宋_GBK"
FONT_FANGZHENG_KAITI = "方正楷体_GBK"
FONT_FANGZHENG_HEITI = "方正黑体_GBK"
FONT_FANGZHENG_FANGSONG = "方正仿宋_GBK"
FONT_TIMES_NEW_ROMAN = "Times New Roman"

# ==================== 文档类型识别 ====================
def detect_document_type(text):
    """自动识别文档类型：调研报告/公文/通用"""
    keywords_research = ['调研', '调查', '走访', '座谈', '问卷', '实地查看', '基本情况', '存在问题', '原因分析', '对策建议']
    keywords_official = ['请示', '通知', '报告', '函', '决定', '公告', '通告', '议案', '批复', '命令']

    research_score = sum(1 for k in keywords_research if k in text)
    official_score = sum(1 for k in keywords_official if k in text)

    if research_score >= 2:
        return 'research'
    elif official_score >= 1:
        return 'official'
    else:
        return 'general'


def print_doc_type_info(doc_type):
    """打印文档类型信息"""
    names = {'research': '调研报告', 'official': '公文', 'general': '通用文章'}
    print(f"📄 文档类型：{names.get(doc_type, '通用文章')}（自动识别）")


# ==================== 字体/段落工具 ====================
def build_run(p_elem, text, font_name, font_size_pt, bold=False):
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    if font_name != FONT_TIMES_NEW_ROMAN:
        rFonts.set(qn('w:hint'), 'eastAsia')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size_pt) * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size_pt) * 2))
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    p_elem.append(run)


def is_ascii_char(c):
    o = ord(c)
    if o < 128:
        return True
    # 中文标点也算中文
    cn_punct = '（），、：；""''【】《》…—'
    if c in cn_punct:
        return False
    return False


def split_text_into_segments(text):
    segments = []
    current_type = None
    current_chars = []
    for ch in text:
        if ch == '\n' or ch == '\r':
            continue
        # 中文标点跟前面的中文合并
        if ch in '（），、：；""''【】《》…—' and current_type == 'cn':
            current_chars.append(ch)
            continue
        ch_type = 'ascii' if (ord(ch) < 128 and ch not in '（）【】《》…—') else 'cn'
        if ch_type != current_type:
            if current_chars:
                segments.append((current_type, ''.join(current_chars)))
            current_type = ch_type
            current_chars = [ch]
        else:
            current_chars.append(ch)
    if current_chars:
        segments.append((current_type, ''.join(current_chars)))
    return segments


def add_para_with_mixed_font(doc, text, cn_font, font_size_pt, bold=False,
                               align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    p_elem = p._element

    # 清除默认空 run
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)

    segments = split_text_into_segments(text)
    for seg_type, seg_text in segments:
        font = cn_font if seg_type == 'cn' else FONT_TIMES_NEW_ROMAN
        build_run(p_elem, seg_text, font, font_size_pt, bold)

    # 段落间距：固定值30磅
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    existing_spacing = pPr.find(qn('w:spacing'))
    if existing_spacing is not None:
        pPr.remove(existing_spacing)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(30 * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    if first_indent_pt > 0:
        p.paragraph_format.first_line_indent = Pt(first_indent_pt)
    elif first_indent_pt == 0:
        p.paragraph_format.first_line_indent = Pt(0)

    return p


def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ""
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_elem = p._element
        for r in p_elem.findall(qn('w:r')):
            p_elem.remove(r)
        build_run(p_elem, "— ", FONT_TIMES_NEW_ROMAN, PAGENUM_SIZE.pt)
        run_fld = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(PAGENUM_SIZE.pt) * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(PAGENUM_SIZE.pt) * 2))
        rPr.append(szCs)
        run_fld.append(rPr)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_fld.append(fldChar1)
        run_fld.append(instrText)
        run_fld.append(fldChar2)
        p_elem.append(run_fld)
        build_run(p_elem, " —", FONT_TIMES_NEW_ROMAN, PAGENUM_SIZE.pt)


def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.left_margin = LEFT_MARGIN
        section.right_margin = RIGHT_MARGIN
        section.header_distance = HEADER_DISTANCE
        section.footer_distance = FOOTER_DISTANCE


# ==================== 内容解析 ====================
def parse_content(raw_text):
    lines = raw_text.strip().split('\n')
    result = {'title': '', 'unit': '', 'date': '', 'body_lines': []}

    for i, line in enumerate(lines):
        line_s = line.strip()
        if not line_s:
            continue

        if i == 0 and not any(x in line_s for x in ['汇报单位', '单位：', '日期：', '#', '##']):
            result['title'] = line_s
            continue

        unit_match = re.match(r'(汇报单位|单位)[:：]\s*(.+)', line_s)
        if unit_match:
            result['unit'] = unit_match.group(2).strip()
            continue

        date_match = re.match(r'日期[:：]\s*(.+)', line_s)
        if date_match:
            result['date'] = date_match.group(1).strip()
            continue

        if line_s.startswith('## '):
            result['body_lines'].append(('heading', 2, line_s[3:].strip()))
        elif line_s.startswith('# '):
            result['body_lines'].append(('heading', 1, line_s[2:].strip()))
        elif result['title'] == '' and i == 0:
            result['title'] = line_s
        else:
            result['body_lines'].append(('body', 0, line_s))

    return result


# ==================== 主函数 ====================
def format_article(raw_text, output_path=None, doc_type=None):
    # 自动识别文档类型
    if doc_type is None:
        doc_type = detect_document_type(raw_text)

    # 输出文档类型
    print_doc_type_info(doc_type)

    parsed = parse_content(raw_text)
    doc = Document()
    set_page_margins(doc)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = FONT_FANGZHENG_FANGSONG
    style.font.size = BODY_SIZE

    # ===== 标题：方正小标宋_GBK二号居中 =====
    if parsed['title']:
        add_para_with_mixed_font(doc, parsed['title'],
                                  FONT_FANGZHENG_XIAOBIAO, TITLE_SIZE.pt,
                                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # ===== 汇报单位：方正楷体_GBK三号居中 =====
    if parsed['unit']:
        add_para_with_mixed_font(doc, parsed['unit'],
                                  FONT_FANGZHENG_KAITI, SUBTITLE_SIZE.pt,
                                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # ===== 日期：方正楷体_GBK三号居中 =====
    if parsed['date']:
        date_text = f"（{parsed['date']}）"
    else:
        now = datetime.now()
        date_text = f"（{now.year}年 {now.month}月 {now.day}日）"
    add_para_with_mixed_font(doc, date_text,
                              FONT_FANGZHENG_KAITI, SUBTITLE_SIZE.pt,
                              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ===== 正文段落 =====
    for item in parsed['body_lines']:
        kind, level, text = item
        if kind == 'heading':
            cn_font = FONT_FANGZHENG_HEITI if level == 1 else FONT_FANGZHENG_KAITI
            add_para_with_mixed_font(doc, text, cn_font, SUBTITLE_SIZE.pt,
                                      first_indent_pt=0)
        else:
            add_para_with_mixed_font(doc, text, FONT_FANGZHENG_FANGSONG, BODY_SIZE.pt,
                                      first_indent_pt=32)

    # ===== 页码 =====
    add_page_number(doc)

    # ===== 保存 =====
    if not output_path:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(OUTPUT_FOLDER, f"formatted_article_{timestamp}.docx")

    doc.save(output_path)
    return output_path


# ==================== CLI ====================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='机关公文格式化工具 v3')
    parser.add_argument('input', nargs='?', help='输入文本或文件路径')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('-f', '--file', help='从文件读取内容')
    parser.add_argument('-t', '--type', choices=['research', 'official', 'general'],
                        help='文档类型：research(调研报告)/official(公文)/general(通用)')
    args = parser.parse_args()

    if args.file:
        with open(args.file, 'r', encoding='utf-8') as f:
            raw = f.read()
    elif args.input:
        if os.path.isfile(args.input):
            with open(args.input, 'r', encoding='utf-8') as f:
                raw = f.read()
        else:
            raw = args.input
    else:
        print("请提供文章内容或文件路径")
        sys.exit(1)

    output = format_article(raw, args.output, doc_type=args.type)
    print(f"✅ 文件已生成：{output}")

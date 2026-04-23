#!/usr/bin/env python3
"""
md2docx.py — 将 Markdown 转换为符合中国公文规范的 Word 文档（.docx）或 PDF

公文格式标准：
┌──────────┬──────────────────┬──────────┐
│ 元素     │ 字体             │ 字号     │
├──────────┼──────────────────┼──────────┤
│ 文档标题 │ 方正小标宋_GBK   │ 小二(18) │
│ 一级标题 │ 方正黑体_GBK     │ 三号(16) │
│ 二级标题 │ 方正楷体_GBK     │ 三号(16) │
│ 三级标题 │ 方正仿宋_GBK加粗 │ 三号(16) │
│ 四级标题 │ 方正仿宋_GBK     │ 三号(16) │
│ 正文     │ 方正仿宋_GBK     │ 三号(16) │
│ 页码     │ 宋体             │ 四号(14) │
└──────────┴──────────────────┴──────────┘

Markdown 层级映射：
  # → 文档标题（居中, 小二号）
  ## → 一级标题（一、二、三）
  ### → 二级标题（（一）（二）（三））
  #### → 三级标题（1. 2. 3.）
  ##### → 四级标题（（1）（2）（3））

用法：
  # 输出 docx
  python md2docx.py input.md -o output.docx

  # 输出 PDF
  python md2docx.py input.md -o output.pdf --pdf

  # 同时输出 docx 和 PDF
  python md2docx.py input.md -o output.docx --pdf
"""
import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ══════════════════════════════════════════════════════════════
# 中文编号
# ══════════════════════════════════════════════════════════════
CN_NUMS = [
    '一','二','三','四','五','六','七','八','九','十',
    '十一','十二','十三','十四','十五','十六','十七','十八','十九','二十'
]
CN_PAREN = [f'（{n}）' for n in CN_NUMS]

# ══════════════════════════════════════════════════════════════
# 字体定义（含回退）
# ══════════════════════════════════════════════════════════════
FONTS = {
    'title':    ('方正小标宋_GBK', 'SimSun'),
    'heiti':    ('方正黑体_GBK',   'SimHei'),
    'kaiti':    ('方正楷体_GBK',   'KaiTi'),
    'fangsong': ('方正仿宋_GBK',   'FangSong'),
    'songti':   ('宋体',           'SimSun'),
}

# ══════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════

def _set_run_font(run, font_key, size_pt, bold=False):
    """设置 run 的中英文字体、字号、加粗"""
    primary, fallback = FONTS[font_key]
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = primary

    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)

    for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(attr), primary)
    rFonts.set(qn('w:cs'), fallback)


def _fmt_para(p, line_pt=28, indent=False, before_pt=0, after_pt=0):
    """统一段落格式：固定行距、首行缩进、段前段后"""
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(line_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if indent:
        pf.first_line_indent = Pt(32)  # 2字符 × 16pt


def _add_page_number(section):
    """页脚添加 -X- 格式动态页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_run(text, is_field=False, field_code=None):
        if is_field:
            r1 = p.add_run()
            r1._element.append(
                parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
            r2 = p.add_run()
            _set_run_font(r2, 'songti', 14)
            r2._element.append(
                parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">{field_code}</w:instrText>'))
            r3 = p.add_run()
            r3._element.append(
                parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
            r4 = p.add_run("1")
            _set_run_font(r4, 'songti', 14)
            r5 = p.add_run()
            r5._element.append(
                parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        else:
            r = p.add_run(text)
            _set_run_font(r, 'songti', 14)

    _add_run("-")
    _add_run(None, is_field=True, field_code=" PAGE ")
    _add_run("-")


def _fix_zoom(doc):
    """修复 python-docx 默认模板中 zoom 缺少 percent 属性的问题"""
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')


# ══════════════════════════════════════════════════════════════
# Markdown 解析
# ══════════════════════════════════════════════════════════════

def parse_markdown(md_text):
    """
    解析 Markdown → 结构化元素列表
    返回: [(type, level, text), ...]
      type: 'heading' | 'body'
      level: 1-5 (对应 # ~ #####)
    """
    lines = md_text.strip().split('\n')
    elements = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        m = re.match(r'^(#{1,5})\s+(.+)$', stripped)
        if m:
            elements.append(('heading', len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        # 正文：合并连续非空行
        para = [stripped]
        i += 1
        while i < len(lines):
            nl = lines[i].strip()
            if not nl or re.match(r'^#{1,5}\s+', nl):
                break
            para.append(nl)
            i += 1
        elements.append(('body', 0, ''.join(para)))

    return elements


# ══════════════════════════════════════════════════════════════
# docx → PDF 转换
# ══════════════════════════════════════════════════════════════

def _find_soffice():
    """查找 soffice / LibreOffice 可执行路径"""
    candidates = [
        'soffice',
        '/usr/bin/soffice',
        '/usr/lib/libreoffice/program/soffice',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    ]
    for c in candidates:
        if shutil.which(c):
            return c
    return None


def _needs_socket_shim():
    """检测沙盒环境中 AF_UNIX 是否被禁用，需要 LD_PRELOAD shim"""
    import socket as _socket
    try:
        s = _socket.socket(_socket.AF_UNIX, _socket.SOCK_STREAM)
        s.close()
        return False
    except OSError:
        return True


# LibreOffice 在沙盒环境中需要的 socket shim（AF_UNIX 被禁时自动编译）
_SHIM_SOURCE = r"""
#define _GNU_SOURCE
#include <dlfcn.h>
#include <errno.h>
#include <stdio.h>
#include <stdlib.h>
#include <sys/socket.h>
#include <unistd.h>
static int (*real_socket)(int,int,int);
static int (*real_socketpair)(int,int,int,int[2]);
static int (*real_listen)(int,int);
static int (*real_accept)(int,struct sockaddr*,socklen_t*);
static int (*real_close)(int);
static int (*real_read)(int,void*,size_t);
static int is_shimmed[1024], peer_of[1024], wake_r[1024], wake_w[1024];
static int listener_fd=-1;
__attribute__((constructor)) static void init(void){
  real_socket=dlsym(RTLD_NEXT,"socket"); real_socketpair=dlsym(RTLD_NEXT,"socketpair");
  real_listen=dlsym(RTLD_NEXT,"listen"); real_accept=dlsym(RTLD_NEXT,"accept");
  real_close=dlsym(RTLD_NEXT,"close");   real_read=dlsym(RTLD_NEXT,"read");
  for(int i=0;i<1024;i++){peer_of[i]=-1;wake_r[i]=-1;wake_w[i]=-1;}
}
int socket(int d,int t,int p){
  if(d==AF_UNIX){int fd=real_socket(d,t,p);if(fd>=0)return fd;
    int sv[2];if(real_socketpair(d,t,p,sv)==0){
      if(sv[0]>=0&&sv[0]<1024){is_shimmed[sv[0]]=1;peer_of[sv[0]]=sv[1];
        int wp[2];if(pipe(wp)==0){wake_r[sv[0]]=wp[0];wake_w[sv[0]]=wp[1];}}
      return sv[0];}errno=EPERM;return -1;}
  return real_socket(d,t,p);}
int listen(int s,int b){if(s>=0&&s<1024&&is_shimmed[s]){listener_fd=s;return 0;}return real_listen(s,b);}
int accept(int s,struct sockaddr*a,socklen_t*l){
  if(s>=0&&s<1024&&is_shimmed[s]){if(wake_r[s]>=0){char b;real_read(wake_r[s],&b,1);}errno=ECONNABORTED;return -1;}
  return real_accept(s,a,l);}
int close(int fd){if(fd>=0&&fd<1024&&is_shimmed[fd]){
    int wl=(fd==listener_fd);is_shimmed[fd]=0;
    if(wake_w[fd]>=0){char c=0;write(wake_w[fd],&c,1);real_close(wake_w[fd]);wake_w[fd]=-1;}
    if(wake_r[fd]>=0){real_close(wake_r[fd]);wake_r[fd]=-1;}
    if(peer_of[fd]>=0){real_close(peer_of[fd]);peer_of[fd]=-1;}
    if(wl)_exit(0);}return real_close(fd);}
"""


def _ensure_shim():
    """编译并缓存 LD_PRELOAD shim"""
    shim_so = Path(tempfile.gettempdir()) / "lo_socket_shim.so"
    if shim_so.exists():
        return str(shim_so)
    src = Path(tempfile.gettempdir()) / "lo_socket_shim.c"
    src.write_text(_SHIM_SOURCE)
    subprocess.run(
        ["gcc", "-shared", "-fPIC", "-o", str(shim_so), str(src), "-ldl"],
        check=True, capture_output=True,
    )
    src.unlink(missing_ok=True)
    return str(shim_so)


def _get_soffice_env():
    """获取 LibreOffice 运行所需的环境变量（自动处理沙盒兼容）"""
    env = os.environ.copy()
    env["SAL_USE_VCLPLUGIN"] = "svp"
    if _needs_socket_shim():
        env["LD_PRELOAD"] = _ensure_shim()
    return env


def docx_to_pdf(docx_path, pdf_path=None):
    """
    将 .docx 转换为 .pdf（通过 LibreOffice，自动适配沙盒环境）

    Args:
        docx_path: 输入 .docx 文件路径
        pdf_path:  输出 .pdf 文件路径（默认与 docx 同名）

    Returns:
        生成的 PDF 文件路径
    """
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "未找到 LibreOffice (soffice)。\n"
            "PDF 输出依赖 LibreOffice，请安装后重试：\n"
            "  Ubuntu/Debian: sudo apt install libreoffice\n"
            "  macOS: brew install --cask libreoffice\n"
            "  或从 https://www.libreoffice.org 下载"
        )

    docx_path = os.path.abspath(docx_path)
    if not pdf_path:
        pdf_path = re.sub(r'\.docx$', '.pdf', docx_path, flags=re.IGNORECASE)

    pdf_path = os.path.abspath(pdf_path)
    pdf_dir = os.path.dirname(pdf_path)

    # 使用临时目录避免路径权限和中文文件名问题
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_docx = os.path.join(tmpdir, 'input.docx')
        shutil.copy2(docx_path, tmp_docx)

        env = _get_soffice_env()

        result = subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf', tmp_docx, '--outdir', tmpdir],
            capture_output=True, text=True, timeout=120, env=env
        )

        tmp_pdf = os.path.join(tmpdir, 'input.pdf')

        if not os.path.exists(tmp_pdf):
            raise RuntimeError(
                f"PDF 转换失败。\n"
                f"stdout: {result.stdout}\n"
                f"stderr: {result.stderr}"
            )

        os.makedirs(pdf_dir, exist_ok=True)
        shutil.copy2(tmp_pdf, pdf_path)

    print(f"✅ PDF 文件已生成: {pdf_path}")
    return pdf_path


# ══════════════════════════════════════════════════════════════
# 主转换函数
# ══════════════════════════════════════════════════════════════

def convert(md_file, output_file, title=None, subtitle=None, author=None, date_str=None,
            output_pdf=False, pdf_path=None):
    """
    将 Markdown 文件转换为公文格式 .docx，可选同时输出 PDF。

    Args:
        md_file:     输入 Markdown 文件路径
        output_file: 输出 .docx 文件路径
        title:       文档主标题（默认从 # 提取）
        subtitle:    副标题/发文机关
        author:      署名落款
        date_str:    日期
        output_pdf:  是否同时生成 PDF
        pdf_path:    PDF 输出路径（默认与 docx 同名）

    Returns:
        dict: {'docx': docx路径, 'pdf': pdf路径或None}
    """
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    elements = parse_markdown(md_text)
    doc = Document()

    # ── 页面设置 ──────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.3)
    sec.bottom_margin = Cm(3.7)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    _add_page_number(sec)
    _fix_zoom(doc)

    # ── 编号计数器 ────────────────────────────────────────
    counters = {2: 0, 3: 0, 4: 0, 5: 0}

    def reset_sub(level):
        for l in range(level + 1, 6):
            counters[l] = 0

    title_written = False

    # ── 逐元素生成 ────────────────────────────────────────
    for etype, level, text in elements:

        if etype == 'heading':

            # ─ 文档标题 (#) ─
            if level == 1:
                if title_written:
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _fmt_para(p, before_pt=28, after_pt=14)
                run = p.add_run(title or text)
                _set_run_font(run, 'title', 18)

                if subtitle:
                    for sub in subtitle.replace('\\n', '\n').split('\n'):
                        ps = doc.add_paragraph()
                        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _fmt_para(ps)
                        rs = ps.add_run(sub.strip())
                        _set_run_font(rs, 'fangsong', 16)

                title_written = True
                for k in counters:
                    counters[k] = 0
                continue

            # ─ 一级标题 (##) → 一、二、三 ─
            elif level == 2:
                counters[2] += 1
                reset_sub(2)
                idx = counters[2] - 1
                prefix = f"{CN_NUMS[idx]}、" if idx < len(CN_NUMS) else f"{counters[2]}、"

                p = doc.add_paragraph()
                _fmt_para(p, before_pt=14)
                run = p.add_run(prefix + text)
                _set_run_font(run, 'heiti', 16)

            # ─ 二级标题 (###) → （一）（二）（三）─
            elif level == 3:
                counters[3] += 1
                reset_sub(3)
                idx = counters[3] - 1
                prefix = CN_PAREN[idx] if idx < len(CN_PAREN) else f"（{counters[3]}）"

                p = doc.add_paragraph()
                _fmt_para(p, indent=True)
                run = p.add_run(prefix + text)
                _set_run_font(run, 'kaiti', 16)

            # ─ 三级标题 (####) → 1. 2. 3. ─
            elif level == 4:
                counters[4] += 1
                reset_sub(4)
                prefix = f"{counters[4]}."

                p = doc.add_paragraph()
                _fmt_para(p, indent=True)
                run = p.add_run(prefix + text)
                _set_run_font(run, 'fangsong', 16, bold=True)

            # ─ 四级标题 (#####) → （1）（2）（3）─
            elif level == 5:
                counters[5] += 1
                prefix = f"（{counters[5]}）"

                p = doc.add_paragraph()
                _fmt_para(p, indent=True)
                run = p.add_run(prefix + text)
                _set_run_font(run, 'fangsong', 16)

        elif etype == 'body':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _fmt_para(p, indent=True)
            run = p.add_run(text)
            _set_run_font(run, 'fangsong', 16)

    # ── 署名与日期 ────────────────────────────────────────
    if author or date_str:
        doc.add_paragraph()  # 空行

        if author:
            pa = doc.add_paragraph()
            pa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _fmt_para(pa)
            pa.paragraph_format.right_indent = Cm(2.0)
            ra = pa.add_run(author)
            _set_run_font(ra, 'fangsong', 16)

        if date_str:
            pd = doc.add_paragraph()
            pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _fmt_para(pd)
            pd.paragraph_format.right_indent = Cm(2.0)
            rd = pd.add_run(date_str)
            _set_run_font(rd, 'fangsong', 16)

    # ── 保存 docx ─────────────────────────────────────────
    # 如果用户只要 PDF 且 output_file 以 .pdf 结尾，先存为临时 docx
    only_pdf = output_file.lower().endswith('.pdf') and not output_pdf
    if only_pdf:
        # 用户的 -o 是 .pdf，意味着只要 PDF
        final_pdf = output_file
        # 临时 docx
        docx_out = re.sub(r'\.pdf$', '.docx', output_file, flags=re.IGNORECASE)
        doc.save(docx_out)
        print(f"✅ 公文格式文档已生成: {docx_out}")
        result_pdf = docx_to_pdf(docx_out, final_pdf)
        return {'docx': docx_out, 'pdf': result_pdf}

    doc.save(output_file)
    print(f"✅ 公文格式文档已生成: {output_file}")

    result = {'docx': output_file, 'pdf': None}

    # ── 可选：转换为 PDF ──────────────────────────────────
    if output_pdf:
        if not pdf_path:
            pdf_path = re.sub(r'\.docx$', '.pdf', output_file, flags=re.IGNORECASE)
        result['pdf'] = docx_to_pdf(output_file, pdf_path)

    return result


# ══════════════════════════════════════════════════════════════
# CLI 入口
# ══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='将 Markdown 转换为公文格式 Word 文档（可选 PDF）')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出文件路径（.docx 或 .pdf）')
    parser.add_argument('-t', '--title', help='文档主标题（默认从 # 提取）')
    parser.add_argument('-s', '--subtitle', help='副标题/发文机关（用 \\n 换行）')
    parser.add_argument('-a', '--author', help='署名落款')
    parser.add_argument('-d', '--date', help='日期')
    parser.add_argument('--pdf', action='store_true',
                        help='同时生成 PDF（或者 -o 指定 .pdf 时自动启用）')
    parser.add_argument('--pdf-only', action='store_true',
                        help='只生成 PDF，不保留 docx 中间文件')
    parser.add_argument('--pdf-path', help='PDF 输出路径（默认与 docx 同名）')

    args = parser.parse_args()

    output = args.output
    if not output:
        output = re.sub(r'\.md$', '.docx', args.input)
        if output == args.input:
            output = args.input + '.docx'

    # 如果 -o 直接指定了 .pdf，自动当作 PDF 模式
    is_pdf_output = output.lower().endswith('.pdf')

    result = convert(
        md_file=args.input,
        output_file=output,
        title=args.title,
        subtitle=args.subtitle,
        author=args.author,
        date_str=args.date,
        output_pdf=args.pdf and not is_pdf_output,
        pdf_path=args.pdf_path,
    )

    # --pdf-only: 删除中间 docx
    if args.pdf_only and result.get('pdf') and result.get('docx'):
        docx_file = result['docx']
        if os.path.exists(docx_file) and docx_file != result['pdf']:
            os.remove(docx_file)
            print(f"🗑️  已删除中间文件: {docx_file}")


if __name__ == '__main__':
    main()

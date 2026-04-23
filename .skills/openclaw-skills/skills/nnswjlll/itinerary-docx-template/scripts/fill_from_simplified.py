# -*- coding: utf-8 -*-
import argparse
import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def parse_lines(text: str):
    days = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = re.match(r'^(\d+)\s*(.+)$', line)
        if not m:
            continue
        day = int(m.group(1))
        body = m.group(2)
        hotel = ''
        if ' 住' in body:
            body, hotel = body.split(' 住', 1)
        elif '住' in body:
            body, hotel = body.split('住', 1)
        parts = body.split()
        head = parts[0] if parts else ''
        spots = parts[1:] if len(parts) > 1 else []
        days.append({'day': day, 'head': head, 'spots': spots, 'hotel': hotel.strip()})
    return sorted(days, key=lambda x: x['day'])


def is_return_token(text: str) -> bool:
    t = text or ''
    return ('返回' in t) or ('送站' in t)


def city_from_head(head: str):
    return head.replace('接站', '').replace('接', '').replace('返回', '').replace('送站', '')


def _variant_pick(candidates, token: str, day: int = 1, slot: int = 0):
    if not candidates:
        return ''
    # 稳定轮换：同一行程同一景点保持可复现，不同天/位置可产生变化
    seed = sum(ord(c) for c in (token or '')) + day * 11 + slot * 17
    return candidates[seed % len(candidates)]


def _embellish_sentence(base: str, token: str, day: int, slot: int):
    prefix = _variant_pick([
        '打卡亮点：',
        '推荐体验：',
        '行程看点：',
        '值得停留：',
    ], token=token, day=day, slot=slot)
    suffix = _variant_pick([
        '整体体验舒适度较高。',
        '建议适度放慢节奏，感受会更好。',
        '非常适合拍照与沉浸式游览。',
        '兼顾观赏性与参与感。',
    ], token=token + '_sfx', day=day, slot=slot)
    return f'{prefix}{base}{suffix}'


def day_opening(item, day: int):
    tokens = [item['head']] + item['spots']
    s = ' '.join(tokens)
    if any(is_return_token(x) for x in tokens):
        return '今日为返程收官日，行程以从容衔接交通为主。'
    if day == 1:
        return '今日开启旅程，整体安排张弛有度、轻松不赶路。'
    if '瀑布' in s:
        return '今日进入山水观光主场，视觉冲击与体验感兼具。'
    if any(k in s for k in ['海', '银滩', '涠洲']):
        return '今日以滨海风光为主线，节奏偏舒展，适合慢游与拍照。'
    return _variant_pick([
        '今日行程层次清晰，兼顾经典打卡与在地体验。',
        '今日安排丰富但节奏友好，适合边逛边拍。',
        '今日主打“看景+体验”，观赏与互动兼具。',
    ], token=s, day=day)


def day_closing(item, day: int):
    tokens = [item['head']] + item['spots']
    s = ' '.join(tokens)
    if any(is_return_token(x) for x in tokens):
        return '感谢一路相伴，期待下一次山海再会。'
    if day == 1:
        return '晚间自由活动后充分休息，为后续精彩行程蓄力。'
    if '瀑布' in s:
        return '建议预留整理照片时间，当天素材通常很出片。'
    return _variant_pick([
        '晚间建议适当放松，保持良好状态迎接次日行程。',
        '行程结束后可自由活动，细细回味当日亮点。',
        '建议早点休息，次日体验会更从容。',
    ], token=s + '_close', day=day)


def scenic_desc(name: str, day: int = 1, slot: int = 0):
    lib = {
        '三街两巷': [
            '这里是南宁城市记忆最浓的街区之一，青砖骑楼与老巷烟火交织成独特城市肌理。建议放慢脚步边走边拍，白天看人文、傍晚看灯影，体验感更完整。',
            '三街两巷兼具历史风貌与生活温度，街巷尺度舒适、建筑细节耐看。推荐以“慢逛+随拍”方式体验，更容易捕捉南宁老城的精气神。',
            '作为南宁人文底色最鲜明的片区之一，这里既有老城故事，也有当代活力。建议留足机动时间自由走读，细节越看越有味道。',
        ],
        '中山路小吃街': [
            '南宁夜生活的代表打卡地，地道小吃密集、烟火气十足。推荐按“主食+小吃+饮品”节奏体验，边逛边吃更容易找到本地风味。',
            '这里集合了南宁夜宵文化的精华，品类丰富、选择弹性大。建议先尝经典招牌再做延展，既稳妥又能吃出层次感。',
            '中山路以“热闹但不杂乱”的夜间氛围见长，适合边逛边吃边拍。建议优先体验现做档口，口感与现场感都会更好。',
        ],
        '钦州三娘观海豚': [
            '三娘湾以中华白海豚闻名，海湾视野开阔，海风舒适。观豚体验兼具生态教育和旅行趣味，建议准备望远镜或长焦设备提升观赏体验。',
            '三娘湾海域生态条件良好，是近距离感受海洋生命力的优质点位。建议遵循现场引导安静观赏，体验会更沉浸。',
            '在三娘湾观海豚，既有自然惊喜也有人文意义。建议选择能见度较好的时段出海，观赏效果通常更理想。',
        ],
        '坭兴陶手工制作': [
            '坭兴陶是钦州代表性非遗工艺，手作过程能直观感受泥料、塑形与器型美学。亲手参与后再看成品，会更理解“匠心”二字的分量。',
            '坭兴陶体验兼具文化深度与互动乐趣，从揉泥到成型都很有参与感。建议保留作品照片，后续回看很有纪念意义。',
            '在非遗手作环节中，可以系统感受器物背后的工艺逻辑与审美脉络。体验结束后对地方文化的理解会更立体。',
        ],
        '观堂一日游': [
            '这是一条“慢节奏+沉浸感”的体验线，适合把旅行从“打卡”转成“在场”。整天行程松弛但不松散，亲子、情侣和小团体都很友好。',
            '观堂线路以氛围感和参与感见长，整体体验舒展自然。建议轻装出行，把时间留给慢慢感受与互动。',
            '该行程强调“生活化旅行”与“场景化体验”，既能放松也能出片。若不赶时间，建议适当延长停留。',
        ],
        '骑行': [
            '骑行段会带你进入更贴近在地生活的风景带，视角比车览更细腻。建议轻装出发，保持匀速，随停随拍更出片。',
            '骑行体验兼顾自由度与参与感，能够更细致地感受沿线风貌。建议提前调整座椅与刹车，骑行更轻松稳定。',
            '相比传统观光方式，骑行更适合捕捉“路上的风景”。控制节奏、留足补水时间，旅途舒适度会更高。',
        ],
        '电瓶车': [
            '电瓶车路段轻松省力，适合全龄人群，尤其对长辈更友好。建议优先在开阔路段拍集体照，画面更干净。',
            '电瓶车体验舒适便捷，能在较短时间内覆盖更多风景点。建议按站点节奏下车停留，体验感更完整。',
            '该路段以“省力+高效”见长，适合希望轻松游览的游客。拍照时建议选择转角或高点位置，层次更好。',
        ],
        '冰箱贴制作': [
            '手作环节不仅是互动体验，也是旅程记忆的“实体化”。建议将目的地元素加入设计，成品更有纪念价值。',
            '冰箱贴制作可玩性很高，适合亲子或朋友共同参与。建议在构图中加入日期或地名，回忆锚点更清晰。',
            '这一环节把旅行情绪转化为可保存的小物件，参与感和纪念性都很强。完成后可现场合影留念。',
        ],
        '竹筏下午茶': [
            '竹筏+下午茶是氛围感拉满的组合，拍人像和风景都很合适。建议选择光线柔和时段，出片效果更稳定。',
            '在水面微风中享受下午茶，节奏轻柔、体验治愈。建议以浅色服装搭配自然景观，照片质感更突出。',
            '竹筏场景天然自带松弛感，是旅途中非常出彩的“高情绪价值”环节。建议预留时间做短视频记录。',
        ],
        '德天瀑布': [
            '德天瀑布是中越边境最具代表性的山水景观之一，近观水雾翻涌、远观群峰叠翠。建议近景感受水势、远景拍全貌，层次会更完整。',
            '德天瀑布气势磅礴、画面冲击力强，是边境线上的地标级景观。建议分区游览：先近观再远拍，体验更有递进感。',
            '这里兼具自然奇观与地域辨识度，视觉记忆点非常鲜明。若天气通透，建议适当增加停留时间拍全景。',
        ],
        '峒那屿湾': [
            '峒那屿湾整体气质偏静谧，适合放慢节奏做“呼吸式”旅行。建议沿水岸线慢走，边看边停，容易拍到高级感画面。',
            '屿湾风景层次柔和，氛围安静舒展，适合做慢旅行节点。建议把拍摄重点放在倒影和线条构图上。',
            '在这里更推荐“少走快看，多停多感受”的方式，旅行质感会更强。是非常适合沉浸放松的点位。',
        ],
        '白头叶猴': [
            '白头叶猴属于珍稀灵长类，生态观赏更强调“安静、远观、不过度打扰”。建议把节奏放缓，尊重自然规律，体验会更深刻。',
            '该点位以生态价值见长，观赏过程更像一次自然课堂。建议降低音量并避免追拍，让观察更从容。',
            '近距离了解珍稀物种的生境与习性，是很有意义的旅行体验。遵守保护规范，才能看到更真实的自然状态。',
        ],
        '侨港风情街': [
            '这是北海夜晚最有人气的美食街区之一，海鲜与越式风味并存。建议从经典小吃入门，再尝当季海味，体验更均衡。',
            '侨港风情街烟火气十足，味型层次丰富，适合做夜游美食主场。建议小份多样尝试，效率更高也更有趣。',
            '这里把滨海城市的夜生活节奏展现得很完整，边逛边吃非常惬意。建议优先选择翻台快、现做类档口。',
        ],
        '涠洲岛鳄鱼山景区': [
            '鳄鱼山景区集中展示了海岛火山地貌的精华，海蚀崖壁与步道视野都很出色。建议预留充足步行时间，逐段看景更震撼。',
            '该景区地貌辨识度高、观景点密集，是涠洲岛核心景观区之一。建议顺着步道节奏慢慢推进，体验更完整。',
            '从火山岩层到海岸线肌理，这里有很强的地质观赏价值。建议结合导览理解地貌成因，收获会更深。',
        ],
        '天主教堂': [
            '这座珊瑚石教堂古朴安静，是涠洲岛极具辨识度的人文地标。建议在光线柔和时进入拍摄，建筑质感会更突出。',
            '教堂建筑线条简洁、氛围庄重，兼具历史感与审美感。建议保持安静参观，细看材质与窗影细节。',
            '作为海岛人文名片之一，这里非常适合做“建筑+人物”主题拍摄。建议避开高峰时段，画面更干净。',
        ],
        '滴水丹萍': [
            '滴水丹萍以丹霞岩壁和海蚀景观著称，傍晚时分色彩层次最丰富。建议结合剪影拍法，容易获得高质量大片。',
            '这里的岩壁色泽与海岸线形成鲜明对比，视觉辨识度很高。建议抓住落日时段，光影会更有戏剧感。',
            '滴水丹萍属于“看天看光线”的点位，天气好时非常惊艳。建议预留机动时间等待最佳拍摄窗口。',
        ],
        '涠洲岛环岛咖啡': [
            '环岛咖啡属于“慢旅行”节点，适合在海风里短暂停留、调节节奏。建议搭配轻食，作为午后补能环节很合适。',
            '这一站更偏生活方式体验，适合放松休息与社交拍照。建议选择临海座位，氛围和视野都更好。',
            '咖啡节点能很好平衡行程节奏，让旅行从“赶路”回到“享受”。建议控制时长，兼顾休息与效率。',
        ],
        '季节水果采摘（赠送1斤）赠送落日餐': [
            '采摘+落日餐把互动体验和情绪价值结合得很好。建议将采摘安排在下午、落日餐安排在黄金时段，氛围最强。',
            '这一组合兼顾参与感与仪式感，适合家庭与朋友同游。建议把采摘成果与晚餐场景联动记录，记忆点更强。',
            '从果园到餐桌的过渡自然且有趣，是行程中很受欢迎的体验段。建议预留拍照与休息时间，节奏更从容。',
        ],
        '银滩': [
            '银滩沙质细白、海岸线绵长，是北海滨海体验的经典名片。建议在低角度光线时段拍摄，海天层次更通透。',
            '银滩视野开阔、漫步舒适，适合做轻松型海边体验。建议赤脚慢走感受沙质，并留意潮汐变化。',
            '这里是北海最具代表性的海滨场景之一，拍照和休闲都很友好。建议安排自由活动时间，体验更完整。',
        ],
        '老街': [
            '北海老街骑楼与南洋风格并存，适合边逛边拍边吃。建议以“街景+人文细节”双线记录，会更有故事感。',
            '老街保留了浓厚历史风貌，建筑立面与店铺招牌都很有味道。建议沿主街慢逛并穿插巷道，发现更多惊喜。',
            '这里兼具历史记忆与生活烟火，是很适合“叙事型拍摄”的点位。建议多拍门窗与街角细节，质感更强。',
        ],
        '太平古城': [
            '太平古城的门楼、街巷与夜景氛围都很在线，适合做夜游体验。建议晚饭后进入古城，整体体验更完整。',
            '古城夜间灯光层次丰富，步行体验舒适，是轻松游览的优选。建议按“主街-广场-城门”顺序游览。',
            '这里的夜游氛围感较强，适合拍人像、街景与短视频。建议预留自由活动时段，沉浸感会更好。',
        ],
        '明仕田园': [
            '明仕田园峰林田畴如画，山水倒映成景，是广西田园风光代表之一。建议以“慢游+轻拍”方式体验，画面更容易出彩。',
            '该区域自然景观开阔且层次丰富，适合静静感受山水田园之美。建议放慢步伐，以“看景+呼吸”节奏游览。',
            '明仕田园的气质在于宁静与通透，尤其适合做疗愈型旅行体验。建议选择视野开阔点位拍全景。',
        ],
    }

    if name in lib:
        base = _variant_pick(lib[name], token=name, day=day, slot=slot)
        return _embellish_sentence(base, token=name, day=day, slot=slot)
    fallback = [
        f'{name}为本次行程特色点位，景观辨识度高，建议预留充分游览与拍照时间。',
        f'{name}是本线路的重要亮点，建议慢节奏参观并结合现场导览理解其特色。',
        f'{name}兼具观赏性与体验感，建议合理安排停留时长，以获得更完整的旅行感受。',
    ]
    base = _variant_pick(fallback, token=name, day=day, slot=slot)
    return _embellish_sentence(base, token=name, day=day, slot=slot)


def day_tip(item, day: int = 1):
    tokens = [item['head']] + item['spots']
    s = ' '.join(tokens)

    pools = {
        'return': [
            '返程当天请提前核对证件、车票/航班信息与行李清单，贵重物品务必随身携带；建议至少预留30分钟机动时间，避免因路况影响送站节奏。',
            '今日为返程日，请再次确认证件与出行信息，行李建议提前整理并做好易丢物品复查；集合与送站请以导游通知时间为准。',
            '返程环节建议“提前准备、从容出发”：证件票务先检查、行李先打包、贵重物品随身放置，避免临近送站手忙脚乱。',
        ],
        'waterfall': [
            '瀑布区域水汽较重且地面湿滑，建议穿防滑鞋并做好手机/相机防潮；拍照请在安全区域内进行，避免跨越警戒线。',
            '瀑布点位湿度较高，台阶与栈道可能打滑，建议减速慢行并注意脚下；电子设备可提前准备防水保护。',
            '临近瀑布时体感会偏凉且水雾明显，建议携带轻薄外套并做好防潮措施；观景时请服从现场安全指引。',
        ],
        'sea': [
            '海边紫外线和体感风力变化较大，请做好防晒补水并准备轻薄外套；如涉及乘船或近海活动，请以现场安全指引为准。',
            '滨海活动建议注意防晒、防风与补水，尤其午后时段体感更强；若有船程安排，请提前留意集合与登船提示。',
            '海边天气变化相对快，建议携带遮阳用品与防晒用品，并合理安排休息节奏；涉海项目请全程遵守工作人员引导。',
        ],
        'activity': [
            '体验项目前请检查装备状态并听从工作人员讲解；骑行/竹筏环节建议控制节奏，老人及儿童优先做好防护。',
            '参与互动项目时建议先熟悉操作规范并做好基础防护，行进中不抢行、不追拍，确保行程安全与体验质量。',
            '骑行、电瓶车或竹筏等项目以“安全第一、体验第二”为原则，建议服从组织安排，保持队形与节奏更稳妥。',
        ],
        'default': [
            '当日以步行观光为主，建议穿舒适鞋并随身携带饮水；夜间活动请注意集合时间与地点，避免走散。',
            '今日行程节奏适中，建议轻装出行并适时补水；拍照与自由活动请以不影响集合时间为前提。',
            '建议提前关注天气变化并准备常用随身物品（纸巾、充电宝、雨具等），游览时注意脚下安全与团队联络。',
        ],
    }

    if any(is_return_token(x) for x in tokens):
        return _variant_pick(pools['return'], token=s, day=day)
    if '瀑布' in s:
        return _variant_pick(pools['waterfall'], token=s, day=day)
    if any(k in s for k in ['海', '银滩', '涠洲']):
        return _variant_pick(pools['sea'], token=s, day=day)
    if any(k in s for k in ['骑行', '电瓶车', '竹筏']):
        return _variant_pick(pools['activity'], token=s, day=day)
    return _variant_pick(pools['default'], token=s, day=day)


def set_text_keep_style(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)


def copy_run_style(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font is not None and dst_run.font is not None:
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        dst_run.font.color.rgb = src_run.font.color.rgb


def insert_after_with_style(anchor_p, text, style_ref_p):
    new_p = OxmlElement('w:p')
    anchor_p._p.addnext(new_p)
    p = Paragraph(new_p, anchor_p._parent)
    p.style = style_ref_p.style
    r = p.add_run(text)
    if style_ref_p.runs:
        copy_run_style(style_ref_p.runs[0], r)
    return p


def find_anchor_index(paragraphs, keys):
    for i, p in enumerate(paragraphs):
        t = (p.text or '').strip()
        if any(k in t for k in keys):
            return i
    return None


def remove_contract_and_tables(doc):
    for t in list(doc.tables):
        tbl = t._tbl
        tbl.getparent().remove(tbl)
    idx = find_anchor_index(doc.paragraphs, ['接待标准', '行程包含', '旅游合同', '补充协议', '合同'])
    if idx is not None:
        for p in list(doc.paragraphs[idx:]):
            e = p._element
            e.getparent().remove(e)


def fill_template(template, content, output, itinerary_only=True):
    doc = Document(str(template))
    days = parse_lines(Path(content).read_text(encoding='utf-8'))
    if not days:
        raise ValueError('No valid day lines found in content file')

    # 模板锚点（新模板4天）
    title_idx = 0
    day_slots = [(2, 3), (5, 6), (13, 14), (20, 21)]
    tip_slot = {1: 9, 3: 17}

    majors = []
    for d in days:
        for s in [d['head']] + d['spots']:
            if s and (not is_return_token(s)) and ('接' not in s):
                majors.append(s)
    majors = list(dict.fromkeys(majors))[:3]
    set_text_keep_style(doc.paragraphs[title_idx], f"{city_from_head(days[0]['head'])}{len(days)}日游｜{'、'.join(majors)}")

    prev_city = city_from_head(days[0]['head'])
    last_day = days[-1]['day']

    # 用已有段落填充前4天
    for i, (ti, di) in enumerate(day_slots):
        if i >= len(days):
            break
        d = days[i]
        day = d['day']
        if day == 1:
            title = f"║第{day}天║：快乐出发—{city_from_head(d['head'])}"
            meal = '早餐：无'
        elif is_return_token(d['head']) or any(is_return_token(s) for s in d['spots']):
            title = f"║第{day}天║：{prev_city}—返回南宁送站"
            meal = '早餐'
        else:
            title = f"║第{day}天║：{prev_city}—{d['head']}" + (f"—{d['spots'][0]}" if d['spots'] else '')
            meal = '早餐/中餐' if day != last_day else '早餐'

        hotel = d['hotel'] if d['hotel'] else ('温馨的家' if day == last_day else '待定')
        set_text_keep_style(doc.paragraphs[ti], f"{title}    {meal}    住宿：{hotel}")

        if day == 1:
            core = [s for s in d['spots'] if not is_return_token(s)]
            detail = (
                day_opening(d, day)
                + '重点安排：'
                + '；'.join([f"【{s}】{scenic_desc(s, day=day, slot=idx)}" for idx, s in enumerate(core, start=1)])
                + '。'
                + day_closing(d, day)
            )
        elif is_return_token(d['head']) or any(is_return_token(s) for s in d['spots']):
            detail = day_opening(d, day) + '早餐后根据返程时间统一安排送站，行程圆满结束。' + day_closing(d, day)
        else:
            core = [d['head']] + [s for s in d['spots'] if not is_return_token(s)]
            detail = (
                day_opening(d, day)
                + '主要游览：'
                + '；'.join([f"【{s}】{scenic_desc(s, day=day, slot=idx)}" for idx, s in enumerate(core, start=1)])
                + '。'
                + day_closing(d, day)
            )
        set_text_keep_style(doc.paragraphs[di], detail)

        if day in tip_slot and tip_slot[day] < len(doc.paragraphs):
            set_text_keep_style(doc.paragraphs[tip_slot[day]], f"【温馨提示】：{day_tip(d, day=day)}")

        if d['hotel']:
            prev_city = re.split(r'酒店|度假|观堂', d['hotel'])[0][:4] or prev_city

    # 清理模板旧大段落
    for idx in [7, 8, 10, 11, 12, 15, 16, 18, 19]:
        if idx < len(doc.paragraphs):
            set_text_keep_style(doc.paragraphs[idx], '')

    # 第5天及以上：按样式克隆插入
    if len(days) > len(day_slots):
        # 先定位插入锚点：day4 详情段
        anchor = doc.paragraphs[21] if 21 < len(doc.paragraphs) else doc.paragraphs[-1]
        title_style_ref = doc.paragraphs[20]
        detail_style_ref = doc.paragraphs[21]
        tip_style_ref = doc.paragraphs[17] if 17 < len(doc.paragraphs) else doc.paragraphs[21]

        for d in days[len(day_slots):]:
            day = d['day']
            if is_return_token(d['head']) or any(is_return_token(s) for s in d['spots']):
                title = f"║第{day}天║：{prev_city}—返回南宁送站"
                meal = '早餐'
                detail = day_opening(d, day) + '早餐后根据返程时间统一安排送站，行程圆满结束。' + day_closing(d, day)
                hotel = '温馨的家'
            else:
                title = f"║第{day}天║：{prev_city}—{d['head']}" + (f"—{d['spots'][0]}" if d['spots'] else '')
                meal = '早餐/中餐' if day != last_day else '早餐'
                core = [d['head']] + [s for s in d['spots'] if not is_return_token(s)]
                detail = (
                    day_opening(d, day)
                    + '主要游览：'
                    + '；'.join([f"【{s}】{scenic_desc(s, day=day, slot=idx)}" for idx, s in enumerate(core, start=1)])
                    + '。'
                    + day_closing(d, day)
                )
                hotel = d['hotel'] if d['hotel'] else ('温馨的家' if day == last_day else '待定')

            p_title = insert_after_with_style(anchor, f"{title}    {meal}    住宿：{hotel}", title_style_ref)
            p_detail = insert_after_with_style(p_title, detail, detail_style_ref)
            p_tip = insert_after_with_style(p_detail, f"【温馨提示】：{day_tip(d, day=day)}", tip_style_ref)
            anchor = p_tip

            if d['hotel']:
                prev_city = re.split(r'酒店|度假|观堂', d['hotel'])[0][:4] or prev_city

    if itinerary_only:
        remove_contract_and_tables(doc)

    doc.save(str(output))


if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('--template', required=True)
    ap.add_argument('--content', required=True)
    ap.add_argument('--output', required=True)
    ap.add_argument('--keep-contract', action='store_true')
    args = ap.parse_args()
    fill_template(args.template, args.content, args.output, itinerary_only=(not args.keep_contract))
    print(args.output)

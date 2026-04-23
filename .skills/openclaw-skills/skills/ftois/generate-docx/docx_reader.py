"""
docx_reader.py — 从任意 .docx 文件提取结构化内容
输出：结构化的内容块列表，供 reformat.py 使用

支持提取：标题层级、正文段落、列表（有序/无序）、表格、代码块
跨平台，仅依赖 python-docx
"""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


# ── 判断辅助 ──────────────────────────────────────────────────

def _pt(emu):
    """EMU → 磅"""
    return round(emu / 12700, 1) if emu else None


def _is_heading_by_style(style_name: str):
    """通过样式名判断标题级别"""
    name = style_name.lower().strip()
    for kw, lvl in [("heading 1",1),("heading 2",2),("heading 3",3),("heading 4",4),
                    ("标题 1",1),("标题 2",2),("标题 3",3),("标题 4",4),
                    ("title",0)]:
        if kw in name:
            return lvl
    return None


def _infer_heading_level(para, size_pt, bold, all_sizes):
    """
    当样式为 Normal 时，通过字号+加粗启发式推断标题级别。
    all_sizes: 文档中出现的所有字号（降序排列）
    """
    if not bold or size_pt is None:
        return None
    # 按字号从大到小映射为标题级别
    unique = sorted(set(all_sizes), reverse=True)
    try:
        rank = unique.index(size_pt)  # 0=最大字号
        if rank <= 3:
            return rank + 1  # h1/h2/h3/h4
    except ValueError:
        pass
    return None


def _is_code_block(para):
    """判断是否为代码块段落"""
    style = para.style.name.lower()
    if any(k in style for k in ("code","verbatim","source","mono","courier")):
        return True
    runs = para.runs
    if runs and runs[0].font.name and "courier" in runs[0].font.name.lower():
        return True
    # 检测左缩进（代码块通常有固定缩进）
    pf = para.paragraph_format
    if pf.left_indent and pf.left_indent > Pt(20):
        runs = para.runs
        if runs and runs[0].font.name and (
            "courier" in runs[0].font.name.lower() or
            "consol" in runs[0].font.name.lower()
        ):
            return True
    return False


def _get_list_type(para):
    """判断列表类型：'bullet' / 'numbered' / None"""
    style = para.style.name.lower()
    if "list bullet" in style or "bullet" in style:
        return "bullet"
    if "list number" in style or "list cont" in style:
        return "numbered"
    # 从 XML numPr 判断
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            numFmt = None
            # 尝试从 numbering 中查找
            try:
                numId_el = numPr.find(qn("w:numId"))
                if numId_el is not None:
                    num_id = numId_el.get(qn("w:val"))
                    if num_id and num_id != "0":
                        # 有效列表
                        ilvl_el = numPr.find(qn("w:ilvl"))
                        ilvl = int(ilvl_el.get(qn("w:val"),0)) if ilvl_el is not None else 0
                        return "numbered"  # 默认编号，后续可细化
            except:
                pass
            return "numbered"
    return None


def _table_to_dict(tbl):
    """表格 → 结构化字典"""
    rows = []
    for row in tbl.rows:
        cells = []
        for cell in row.cells:
            text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append(text)
        if any(cells):
            rows.append(cells)
    if not rows:
        return None
    return {"type": "table", "headers": rows[0], "rows": rows[1:]}


# ══════════════════════════════════════════════════════════════
# 主提取函数
# ══════════════════════════════════════════════════════════════

def extract(docx_path: str) -> dict:
    """
    从 .docx 文件提取结构化内容。

    返回：
    {
        "title": "文档标题（若能识别）",
        "blocks": [
            {"type": "heading", "level": 1, "text": "章节一"},
            {"type": "body",    "text": "正文内容..."},
            {"type": "bullet",  "text": "列表项", "level": 0},
            {"type": "numbered","text": "编号项", "level": 0},
            {"type": "code",    "text": "代码内容..."},
            {"type": "table",   "headers": [...], "rows": [[...],[...]]},
            {"type": "divider"},
        ],
        "doc_hint": "推断的文档类型关键词",
    }
    """
    doc = Document(docx_path)

    # 第一步：收集所有段落的字号，用于启发式标题判断
    bold_sizes = []
    for para in doc.paragraphs:
        if not para.text.strip(): continue
        runs = para.runs
        if runs and runs[0].font.bold and runs[0].font.size:
            bold_sizes.append(_pt(runs[0].font.size))
    bold_sizes = [s for s in bold_sizes if s]

    # 第二步：构建段落+表格的顺序列表（按文档顺序）
    # python-docx 不直接支持混合顺序，通过 XML body 遍历
    body = doc.element.body
    para_iter  = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    para_map   = {}  # xml element → Para
    table_map  = {}  # xml element → Table

    for p in doc.paragraphs:
        para_map[id(p._p)] = p
    for t in doc.tables:
        table_map[id(t._tbl)] = t

    ordered_items = []  # [(type, obj)]
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            p = para_map.get(id(child))
            if p: ordered_items.append(("para", p))
        elif tag == "tbl":
            t = table_map.get(id(child))
            if t: ordered_items.append(("table", t))

    # 第三步：提取内容块
    blocks = []
    title_text = None
    code_buffer = []   # 合并连续代码行

    def flush_code():
        if code_buffer:
            blocks.append({"type": "code", "text": "\n".join(code_buffer)})
            code_buffer.clear()

    for item_type, obj in ordered_items:

        if item_type == "table":
            flush_code()
            td = _table_to_dict(obj)
            if td:
                blocks.append(td)
            continue

        # 段落处理
        para = obj
        text = para.text.strip()
        if not text:
            flush_code()
            continue

        style_name = para.style.name
        runs = para.runs
        size_emu = runs[0].font.size if runs and runs[0].font.size else None
        size_pt  = _pt(size_emu)
        bold     = runs[0].font.bold if runs else False

        # 代码块检测
        if _is_code_block(para):
            # 保留原始缩进
            code_buffer.append(para.text)
            continue
        else:
            flush_code()

        # 标题级别判断
        lvl = _is_heading_by_style(style_name)
        if lvl is None and bold and bold_sizes:
            lvl = _infer_heading_level(para, size_pt, bold, bold_sizes)

        if lvl == 0:
            # Title 样式 → 文档主标题
            title_text = title_text or text
            blocks.append({"type": "heading", "level": 0, "text": text})
            continue

        if lvl:
            blocks.append({"type": "heading", "level": lvl, "text": text})
            continue

        # 列表
        list_type = _get_list_type(para)
        if list_type:
            pf = para.paragraph_format
            indent_level = 0
            if pf.left_indent:
                indent_level = min(2, int(_pt(pf.left_indent) / 18))
            blocks.append({"type": list_type, "text": text, "level": indent_level})
            continue

        # 分隔线检测（短横线段落）
        if re.match(r'^[-─—=\*]{3,}$', text):
            blocks.append({"type": "divider"})
            continue

        # 普通正文
        blocks.append({"type": "body", "text": text})

    flush_code()

    # 第四步：推断文档类型线索
    all_text = " ".join(b.get("text","") for b in blocks if "text" in b)
    doc_hint = _infer_doc_type(all_text, title_text or "")

    return {
        "title":    title_text or "",
        "blocks":   blocks,
        "doc_hint": doc_hint,
        "stats": {
            "paragraphs": len([b for b in blocks if b["type"]=="body"]),
            "headings":   len([b for b in blocks if b["type"]=="heading"]),
            "tables":     len([b for b in blocks if b["type"]=="table"]),
            "code_blocks":len([b for b in blocks if b["type"]=="code"]),
            "list_items": len([b for b in blocks if b["type"] in ("bullet","numbered")]),
        }
    }


def _infer_doc_type(text: str, title: str) -> str:
    """从内容关键词推断文档类型"""
    combined = (title + " " + text).lower()
    rules = [
        ("GOV_DOC",        ["通知", "决定", "报告", "命令", "党政", "机关", "公文", "发文"]),
        ("GOV_JUDICIAL",   ["判决", "裁定", "起诉书", "司法", "案号", "法院"]),
        ("BUSINESS_CONTRACT",["合同", "协议", "甲方", "乙方", "签订", "违约", "租赁"]),
        ("BUSINESS_TENDER",["标书", "投标", "招标", "投标人", "评分", "报价"]),
        ("BUSINESS_PLAN",  ["商业计划", "创业", "融资", "投资人", "bp", "商业模式", "市场规模"]),
        ("ACADEMIC_PAPER", ["摘要", "关键词", "abstract", "参考文献", "研究方法", "论文"]),
        ("ACADEMIC_LESSON",["教案", "教学目标", "教学大纲", "学时", "课程", "授课"]),
        ("TECH_SRS",       ["需求规格", "功能需求", "srs", "用例", "非功能需求", "系统架构"]),
        ("TECH_MANUAL",    ["用户手册", "操作指南", "安装", "配置", "步骤", "警告", "注意"]),
        ("MEDICAL_DRUG",   ["药品", "适应症", "用法用量", "不良反应", "禁忌", "说明书"]),
        ("MEDICAL_RECORD", ["病历", "主诉", "诊断", "医嘱", "入院", "出院", "患者"]),
        ("MARKETING_PLAN", ["策划", "营销", "推广", "活动方案", "kpi", "品牌"]),
        ("MARKETING_ANALYSIS",["市场分析", "竞品", "swot", "行业报告", "市场份额"]),
        ("LEGAL_OPINION",  ["法律意见", "律师", "当事人", "意见书"]),
        ("LEGAL_LITIGATION",["起诉状", "答辩状", "仲裁", "诉讼请求", "事实与理由"]),
        ("FINANCE_REPORT", ["财务报告", "资产负债", "利润", "现金流", "年报", "季报"]),
        ("ENGINEERING_DOC",["工程方案", "设计方案", "施工", "cad", "图纸", "工程量"]),
    ]
    scores = {}
    for doc_type, keywords in rules:
        score = sum(1 for kw in keywords if kw in combined)
        if score > 0:
            scores[doc_type] = score
    if scores:
        return max(scores, key=scores.get)
    return "GENERAL_DOC"


def to_markdown(extracted: dict) -> str:
    """将提取结果转为 Markdown（用于调试/预览）"""
    lines = []
    for b in extracted["blocks"]:
        t = b["type"]
        text = b.get("text", "")
        if t == "heading":
            lvl = b["level"]
            prefix = "#" * max(1, lvl)
            lines.append(f"\n{prefix} {text}\n")
        elif t == "body":
            lines.append(f"\n{text}\n")
        elif t == "bullet":
            ind = "  " * b.get("level", 0)
            lines.append(f"{ind}- {text}")
        elif t == "numbered":
            ind = "   " * b.get("level", 0)
            lines.append(f"{ind}1. {text}")
        elif t == "code":
            lines.append(f"\n```\n{text}\n```\n")
        elif t == "table":
            hdrs = b["headers"]
            lines.append("\n| " + " | ".join(hdrs) + " |")
            lines.append("| " + " | ".join(["---"]*len(hdrs)) + " |")
            for row in b["rows"]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        elif t == "divider":
            lines.append("\n---\n")
    return "\n".join(lines)


def summary(extracted: dict) -> str:
    """打印提取摘要"""
    s = extracted["stats"]
    return (f"标题:{s['headings']} 正文:{s['paragraphs']} "
            f"列表:{s['list_items']} 表格:{s['tables']} "
            f"代码块:{s['code_blocks']} | 推断类型:{extracted['doc_hint']}")

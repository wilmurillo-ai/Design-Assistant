#!/usr/bin/env python3
"""
build_templates.py — 构建所有 reference.docx 模板（v4.1）
依据《Word文档类型与风格指南》完整实现所有文档类型的排版规范。

运行：python3 build_templates.py
输出：templates/ 目录下 9 个模板文件
"""

import os, platform
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_OS = platform.system()

# ── 跨平台字体映射 ──────────────────────────────────────────
_FONT_MAP = {
    "宋体":        {"Windows": "宋体",        "Darwin": "宋体-简",     "Linux": "Noto Serif CJK SC"},
    "黑体":        {"Windows": "黑体",        "Darwin": "黑体-简",     "Linux": "Noto Sans CJK SC"},
    "楷体":        {"Windows": "楷体",        "Darwin": "楷体-简",     "Linux": "AR PL UKai CN"},
    "仿宋":        {"Windows": "仿宋",        "Darwin": "仿宋-简",     "Linux": "AR PL UMing HK"},
    "微软雅黑":    {"Windows": "微软雅黑",    "Darwin": "PingFang SC", "Linux": "Noto Sans CJK SC"},
    # 英文字体（正式文档）
    "TNR":         {"Windows": "Times New Roman", "Darwin": "Times New Roman", "Linux": "Liberation Serif"},
    # 英文字体（技术文档）
    "Calibri":     {"Windows": "Calibri",     "Darwin": "Helvetica Neue","Linux": "DejaVu Sans"},
}
def F(name):
    return _FONT_MAP.get(name, {}).get(_OS, name)

# ── 中文字号 → 磅值 ─────────────────────────────────────────
CN_PT = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
}
def P(name): return CN_PT[name]

ALIGN = {
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# ── 模板完整配置（依据指南全量实现）────────────────────────
TEMPLATES = {

    # ── 党政机关公文（GB/T 9704-2012）───────────────────────
    "GOV_DOC": {
        "page":  {"w":21,"h":29.7,"mt":3.7,"mb":3.5,"ml":2.8,"mr":2.6},
        "first_page_no_footer": True,                    # 首页不显示页码
        "title": {"font":"宋体","sz":"二号","bold":False,"align":"center","sb":24,"sa":12,
                  "en_font":"TNR"},                       # 小标宋用宋体模拟
        "h1":    {"font":"黑体","sz":"三号","bold":True, "align":"left","sb":14,"sa":6},
        "h2":    {"font":"楷体","sz":"三号","bold":True, "align":"left","sb":8, "sa":4},
        "h3":    {"font":"仿宋","sz":"三号","bold":True, "align":"left","sb":4, "sa":2},
        "h4":    {"font":"仿宋","sz":"三号","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"仿宋","sz":"三号","ls_mode":"exact","ls":28,
                  "first_indent":True,"en_font":"TNR"},   # 每页22行/每行28字由ls控制
        "footer_sz":"四号",
        "accent":"000000","sub_color":"333333","highlight":"CC0000",
        "table_header_bg":"404040","table_border_outer":12,"table_border_inner":4,
    },

    # ── 专用公文：司法文书 ───────────────────────────────────
    "GOV_JUDICIAL": {
        "page":  {"w":21,"h":29.7,"mt":2.54,"mb":2.54,"ml":3.17,"mr":3.17},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"TNR"},
        "h1":    {"font":"黑体","sz":"三号","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":8,"sa":4},
        "h3":    {"font":"仿宋","sz":"小四","bold":False,"align":"left","sb":4,"sa":2},
        "h4":    {"font":"仿宋","sz":"小四","bold":False,"align":"left","sb":2,"sa":1},
        "body":  {"font":"宋体","sz":"四号","ls_mode":"exact","ls":28,
                  "first_indent":True,"en_font":"TNR"},
        "footer_sz":"五号",
        "accent":"1F3864","sub_color":"333333","highlight":"CC0000",
        "table_header_bg":"1F3864","table_border_outer":12,"table_border_inner":4,
    },

    # ── 商业文档：合同/协议 + 标书（共用模板）──────────────
    "BUSINESS_CONTRACT": {
        "page":  {"w":21,"h":29.7,"mt":2.5,"mb":2.5,"ml":3.0,"mr":2.5},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"TNR"},
        "h1":    {"font":"黑体","sz":"小四","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"楷体","sz":"小四","bold":True,"align":"left","sb":8,"sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4,"sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2,"sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"multiple","ls":1.5,
                  "first_indent":True,"en_font":"TNR"},
        "footer_sz":"小五",
        "accent":"1F3864","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"1F3864","table_border_outer":12,"table_border_inner":4,
    },

    # ── 商业文档：标书/投标文件（微软雅黑目录，特殊行距）───
    "BUSINESS_TENDER": {
        "page":  {"w":21,"h":29.7,"mt":2.6,"mb":2.2,"ml":2.5,"mr":2.5},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"Calibri"},
        "h1":    {"font":"微软雅黑","sz":"小四","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"宋体",    "sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体",    "sz":"五号","bold":False,"align":"left","sb":4,"sa":2},
        "h4":    {"font":"宋体",    "sz":"五号","bold":False,"align":"left","sb":2,"sa":1},
        "body":  {"font":"宋体",    "sz":"小四","ls_mode":"multiple","ls":1.5,
                  "first_indent":True,"en_font":"Calibri"},
        "footer_sz":"小五",
        "accent":"2F5496","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"2F5496","table_border_outer":12,"table_border_inner":4,
    },

    # ── 商业文档：商业计划书/创业计划书 ─────────────────────
    "BUSINESS_PLAN": {
        "page":  {"w":21,"h":29.7,"mt":2.5,"mb":2.2,"ml":3.0,"mr":2.5},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"小初","bold":True,"align":"center","sb":36,"sa":12,
                  "en_font":"Calibri"},
        "h1":    {"font":"黑体","sz":"小四","bold":True,"align":"left","sb":14,"sa":6},
        "h2":    {"font":"宋体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"multiple","ls":1.5,
                  "first_indent":True,"en_font":"Calibri"},
        "footer_sz":"小五",
        "accent":"2F5496","sub_color":"D9D9D9","highlight":"FFC000",   # 60-30-10
        "table_header_bg":"2F5496","table_border_outer":12,"table_border_inner":4,
    },

    # ── 学术文档：论文/研究报告 ─────────────────────────────
    "ACADEMIC_PAPER": {
        "page":  {"w":21,"h":29.7,"mt":2.5,"mb":2.5,"ml":3.1,"mr":3.2},
        "first_page_no_footer": False,
        "title": {"font":"仿宋","sz":"三号","bold":False,"align":"center","sb":48,"sa":30,
                  "en_font":"TNR"},    # 前空三行≈3×16pt=48pt，后空两行≈2×16pt=32pt
        "h1":    {"font":"黑体","sz":"四号","bold":True, "align":"left","sb":16,"sa":8},
        "h2":    {"font":"楷体","sz":"小四","bold":False,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"五号","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"五号","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"五号","ls_mode":"multiple","ls":1.25,
                  "first_indent":True,"en_font":"TNR"},
        "footer_sz":"五号",
        "accent":"000000","sub_color":"444444","highlight":"CC0000",
        "table_header_bg":"404040","table_border_outer":12,"table_border_inner":4,
    },

    # ── 学术文档：教案/教学大纲 ─────────────────────────────
    "ACADEMIC_LESSON": {
        "page":  {"w":21,"h":29.7,"mt":2.54,"mb":2.54,"ml":3.17,"mr":3.17},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"Calibri"},
        "h1":    {"font":"黑体","sz":"四号","bold":True,"align":"left","sb":14,"sa":6},
        "h2":    {"font":"宋体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"exact","ls":24,
                  "first_indent":True,"en_font":"Calibri"},  # 固定24磅（20-28磅取中）
        "footer_sz":"小五",
        "accent":"1A5276","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"1A5276","table_border_outer":12,"table_border_inner":4,
    },

    # ── 技术文档：SRS + 用户手册（共用，章节编号1.1.1）──────
    "TECH_MANUAL": {
        "page":  {"w":21,"h":29.7,"mt":2.5,"mb":2.5,"ml":2.5,"mr":2.5},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"四号","bold":True,"align":"center","sb":24,"sa":12,
                  "en_font":"Calibri"},
        "h1":    {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":12,"sa":6},
        "h2":    {"font":"楷体","sz":"四号","bold":False,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"五号","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"五号","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"五号","ls_mode":"multiple","ls":1.25,
                  "first_indent":False,"en_font":"Calibri"},  # 技术文档不缩进
        "footer_sz":"五号",
        "accent":"0070C0","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"0070C0","table_border_outer":12,"table_border_inner":4,
    },

    # ── 医疗文档：病历/药品说明书 ───────────────────────────
    "MEDICAL_DOC": {
        "page":  {"w":21,"h":29.7,"mt":2.54,"mb":2.54,"ml":3.17,"mr":3.17},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"TNR"},
        "h1":    {"font":"黑体","sz":"三号","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"仿宋","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"仿宋","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"仿宋","sz":"四号","ls_mode":"exact","ls":28,
                  "first_indent":True,"en_font":"TNR"},
        "footer_sz":"小五",
        "accent":"117A65","sub_color":"D9D9D9","highlight":"CC0000",
        "table_header_bg":"117A65","table_border_outer":12,"table_border_inner":4,
    },

    # ── 营销/策划文档 ────────────────────────────────────────
    "MARKETING_DOC": {
        "page":  {"w":21,"h":29.7,"mt":2.5,"mb":2.2,"ml":3.0,"mr":2.5},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"小初","bold":True,"align":"center","sb":36,"sa":12,
                  "en_font":"Calibri"},
        "h1":    {"font":"黑体","sz":"小四","bold":True,"align":"left","sb":14,"sa":6},
        "h2":    {"font":"宋体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"multiple","ls":1.5,
                  "first_indent":True,"en_font":"Calibri"},
        "footer_sz":"小五",
        "accent":"2F5496","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"2F5496","table_border_outer":12,"table_border_inner":4,
    },

    # ── 法律文书 ─────────────────────────────────────────────
    "LEGAL_DOC": {
        "page":  {"w":21,"h":29.7,"mt":2.54,"mb":2.54,"ml":3.17,"mr":3.17},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"TNR"},
        "h1":    {"font":"黑体","sz":"小四","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"楷体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"仿宋","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"multiple","ls":1.5,
                  "first_indent":True,"en_font":"TNR"},
        "footer_sz":"小五",
        "accent":"1F3864","sub_color":"D9D9D9","highlight":"CC0000",
        "table_header_bg":"1F3864","table_border_outer":12,"table_border_inner":4,
    },

    # ── 金融报告 ─────────────────────────────────────────────
    "FINANCE_REPORT": {
        "page":  {"w":21,"h":29.7,"mt":2.54,"mb":2.54,"ml":3.17,"mr":3.17},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"TNR"},
        "h1":    {"font":"黑体","sz":"四号","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"楷体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"四号","ls_mode":"exact","ls":28,
                  "first_indent":True,"en_font":"TNR"},
        "footer_sz":"五号",
        "accent":"2F5496","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"2F5496","table_border_outer":12,"table_border_inner":4,
    },

    # ── 工程文档 ─────────────────────────────────────────────
    "ENGINEERING_DOC": {
        "page":  {"w":21,"h":29.7,"mt":2.54,"mb":2.54,"ml":3.17,"mr":3.17},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":18,"sa":9,
                  "en_font":"Calibri"},
        "h1":    {"font":"黑体","sz":"四号","bold":True,"align":"center","sb":12,"sa":6},
        "h2":    {"font":"楷体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"multiple","ls":1.25,
                  "first_indent":False,"en_font":"Calibri"},
        "footer_sz":"小五",
        "accent":"1A5276","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"1A5276","table_border_outer":12,"table_border_inner":4,
    },

    # ── 通用商务文档（兜底）─────────────────────────────────
    "GENERAL_DOC": {
        "page":  {"w":21,"h":29.7,"mt":2.5,"mb":2.5,"ml":3.0,"mr":2.5},
        "first_page_no_footer": False,
        "title": {"font":"黑体","sz":"三号","bold":True,"align":"center","sb":24,"sa":12,
                  "en_font":"Calibri"},
        "h1":    {"font":"黑体","sz":"四号","bold":True,"align":"left","sb":12,"sa":6},
        "h2":    {"font":"宋体","sz":"小四","bold":True,"align":"left","sb":8, "sa":4},
        "h3":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":4, "sa":2},
        "h4":    {"font":"宋体","sz":"小四","bold":False,"align":"left","sb":2, "sa":1},
        "body":  {"font":"宋体","sz":"小四","ls_mode":"multiple","ls":1.5,
                  "first_indent":True,"en_font":"Calibri"},
        "footer_sz":"小五",
        "accent":"2F5496","sub_color":"D9D9D9","highlight":"FFC000",
        "table_header_bg":"2F5496","table_border_outer":12,"table_border_inner":4,
    },
}


# ══════════════════════════════════════════════════════════════
# 模板构建函数
# ══════════════════════════════════════════════════════════════

def _set_style(doc, name, font, sz_name, bold=False, color_hex=None,
               sb=0, sa=0, ls=None, ls_mode=None, fi=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, outline=None,
               en_font=None):
    try:    sty = doc.styles[name]
    except: sty = doc.styles.add_style(name, 1)

    sz_pt = P(sz_name)
    sty.font.name = font
    # 同时设置东亚/西文字体
    try:
        rPr = sty.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font)
        if en_font:
            rFonts.set(qn("w:ascii"),  F(en_font))
            rFonts.set(qn("w:hAnsi"),  F(en_font))
    except: pass

    sty.font.size = Pt(sz_pt)
    sty.font.bold = bold
    if color_hex:
        sty.font.color.rgb = RGBColor(
            int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))

    pf = sty.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    pf.alignment    = align

    if ls_mode == "exact":
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing      = Pt(ls)
    elif ls_mode == "multiple" and ls:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = ls

    if fi is not None:
        pf.first_line_indent = Pt(fi)

    if outline is not None:
        pPr = sty.element.get_or_add_pPr()
        for old in pPr.findall(qn("w:outlineLvl")):
            pPr.remove(old)
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), str(outline))
        pPr.append(ol)
    return sty


def _setup_numbering(doc):
    """建立编号体系：abstractNum 0=bullet，1=decimal"""
    try:
        nb = doc.part.numbering_part._element
    except:
        return
    for c in list(nb):
        nb.remove(c)

    def el(tag, **kw):
        e = OxmlElement(f"w:{tag}")
        for k, v in kw.items():
            e.set(qn(f"w:{k}"), str(v))
        return e

    for abs_id, fmt, txt in [("0","bullet","•"), ("1","decimal","%1.")]:
        an = el("abstractNum", abstractNumId=abs_id)
        an.append(el("multiLevelType", val="hybridMultilevel"))
        for i in range(3):
            lvl = el("lvl", ilvl=str(i))
            lvl.append(el("start", val="1"))
            lvl.append(el("numFmt", val=fmt))
            lvl.append(el("lvlText", val=txt))
            ind = el("ind", left=str(360+i*360), hanging="240")
            pp = OxmlElement("w:pPr"); pp.append(ind); lvl.append(pp)
            an.append(lvl)
        nb.append(an)

    for num_id, abs_id in [("1","0"), ("2","1")]:
        num = el("num", numId=num_id)
        num.append(el("abstractNumId", val=abs_id))
        nb.append(num)


def _add_footer(doc, cfg):
    """添加页脚页码"""
    footer_sz = P(cfg.get("footer_sz","小五"))
    body_font = cfg["body"]["font"]
    sec = doc.sections[0]
    sec.different_first_page_header_footer = cfg.get("first_page_no_footer", False)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    run.font.name = F(body_font)
    run.font.size = Pt(footer_sz)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), F(body_font))
    except: pass
    # PAGE 域
    for tag, txt in [("begin",None),("instrText"," PAGE "),("separate",None),("end",None)]:
        if tag == "instrText":
            e = OxmlElement("w:instrText")
            e.set(qn("xml:space"), "preserve")
            e.text = txt
            run._r.append(e)
        else:
            e = OxmlElement("w:fldChar")
            e.set(qn("w:fldCharType"), tag)
            run._r.append(e)


def build_template(key, cfg, out_dir):
    doc = Document()
    pg = cfg["page"]
    sec = doc.sections[0]
    sec.page_width  = Cm(pg["w"]);  sec.page_height = Cm(pg["h"])
    sec.top_margin  = Cm(pg["mt"]); sec.bottom_margin = Cm(pg["mb"])
    sec.left_margin = Cm(pg["ml"]); sec.right_margin  = Cm(pg["mr"])

    _setup_numbering(doc)

    acc = cfg["accent"]
    bd  = cfg["body"]
    fi_pt = P(bd["sz"]) * 2 if bd.get("first_indent") else 0

    # Normal（正文基础）
    _set_style(doc, "Normal", F(bd["font"]), bd["sz"], False,
               ls=bd["ls"], ls_mode=bd["ls_mode"], fi=fi_pt,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               en_font=bd.get("en_font","Calibri"))

    # Title
    t = cfg["title"]
    _set_style(doc, "Title", F(t["font"]), t["sz"], t["bold"],
               color_hex=acc, sb=t["sb"], sa=t["sa"],
               align=ALIGN[t["align"]], en_font=t.get("en_font"))

    # Heading 1–4
    for lvl_key, sname, outline in [
        ("h1","Heading 1",0), ("h2","Heading 2",1),
        ("h3","Heading 3",2), ("h4","Heading 4",3)
    ]:
        h = cfg[lvl_key]
        _set_style(doc, sname, F(h["font"]), h["sz"], h["bold"],
                   color_hex=acc if lvl_key=="h1" else None,
                   sb=h["sb"], sa=h["sa"],
                   align=ALIGN[h["align"]], outline=outline, fi=0)

    # Verbatim / Code（代码块样式，Pandoc使用）
    for sname in ["Verbatim Char", "Source Code", "Code", "Verbatim"]:
        try:
            _set_style(doc, sname, sname, "Courier New", "小五", False,
                       ls=1.2, ls_mode="multiple")
        except: pass

    _add_footer(doc, cfg)

    out = os.path.join(out_dir, f"template_{key}.docx")
    doc.save(out)
    print(f"  ✅ {key:22s} → {os.path.basename(out)}")
    return out


if __name__ == "__main__":
    out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
    os.makedirs(out_dir, exist_ok=True)
    print(f"构建模板（平台: {_OS}）：")
    for key, cfg in TEMPLATES.items():
        build_template(key, cfg, out_dir)
    print(f"\n✅ 完成，共 {len(TEMPLATES)} 个模板")

#!/usr/bin/env python3
"""
Tool Integration - 统一工具接口

整合所有技能工具:
- browser (网页自动化)
- pdf (PDF 处理)
- docx-cn (Word 文档)
- xlsx-cn (Excel 表格)
- web_search (网络搜索)
"""

import sys
from pathlib import Path
from typing import Dict, List, Optional, Any
import subprocess
import json

sys.path.insert(0, str(Path(__file__).parent.parent))


class UnifiedTools:
    """统一工具接口"""
    
    def __init__(self):
        self.tools = {
            "browser": BrowserTool(),
            "pdf": PDFTool(),
            "docx": DocxTool(),
            "xlsx": XlsxTool(),
            "search": SearchTool(),
            "git": GitTool()
        }
    
    def get_tool(self, name: str):
        """获取工具"""
        return self.tools.get(name)
    
    def list_tools(self) -> List[str]:
        """列出所有工具"""
        return list(self.tools.keys())


class BrowserTool:
    """浏览器工具"""
    
    def __init__(self):
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        """检查是否可用"""
        # 检查是否有 browser 工具
        return True  # OpenClaw 内置
    
    def browse(self, url: str) -> Dict:
        """浏览网页"""
        # 实际会调用 OpenClaw 的 browser 工具
        return {
            "url": url,
            "status": "需要调用 browser 工具"
        }
    
    def search(self, query: str) -> List[Dict]:
        """搜索"""
        # 实际会调用 web_search 工具
        return [{
            "query": query,
            "status": "需要调用 web_search 工具"
        }]
    
    def screenshot(self, url: str, output: str = None) -> str:
        """截图"""
        return f"需要调用 browser screenshot"


class PDFTool:
    """PDF 工具"""
    
    def __init__(self):
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        try:
            import PyPDF2
            return True
        except:
            try:
                import fitz  # PyMuPDF
                return True
            except:
                return False
    
    def extract_text(self, pdf_path: str) -> str:
        """提取文本"""
        if not self.available:
            return "请安装 PyPDF2 或 PyMuPDF"
        
        try:
            import fitz
            doc = fitz.open(pdf_path)
            text = "\n".join([page.get_text() for page in doc])
            doc.close()
            return text
        except:
            try:
                import PyPDF2
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "\n".join([page.extract_text() for page in reader.pages])
                return text
            except Exception as e:
                return f"提取失败: {e}"
    
    def split(self, pdf_path: str, output_dir: str) -> List[str]:
        """拆分 PDF"""
        # 简化实现
        return [f"{output_dir}/page_{i}.pdf" for i in range(10)]
    
    def merge(self, pdf_files: List[str], output: str) -> bool:
        """合并 PDF"""
        try:
            import PyPDF2
            merger = PyPDF2.PdfMerger()
            for pdf in pdf_files:
                merger.append(pdf)
            merger.write(output)
            merger.close()
            return True
        except:
            return False


class DocxTool:
    """Word 文档工具"""
    
    def __init__(self):
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        try:
            from docx import Document
            return True
        except:
            return False
    
    def read(self, docx_path: str) -> str:
        """读取文档"""
        if not self.available:
            return "请安装 python-docx: pip install python-docx"
        
        try:
            from docx import Document
            doc = Document(docx_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"读取失败: {e}"
    
    def create(self, output_path: str, content: str) -> bool:
        """创建文档"""
        if not self.available:
            return False
        
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph(content)
            doc.save(output_path)
            return True
        except:
            return False


class XlsxTool:
    """Excel 表格工具"""
    
    def __init__(self):
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        try:
            import openpyxl
            return True
        except:
            try:
                import pandas
                return True
            except:
                return False
    
    def read(self, xlsx_path: str) -> List[List]:
        """读取表格"""
        if not self.available:
            return [["请安装 openpyxl 或 pandas"]]
        
        try:
            import pandas as pd
            df = pd.read_excel(xlsx_path)
            return df.values.tolist()
        except Exception as e:
            return [[f"读取失败: {e}"]]
    
    def create(self, output_path: str, data: List[List]) -> bool:
        """创建表格"""
        if not self.available:
            return False
        
        try:
            import pandas as pd
            df = pd.DataFrame(data)
            df.to_excel(output_path, index=False)
            return True
        except:
            return False


class SearchTool:
    """搜索工具"""
    
    def __init__(self):
        self.available = True  # OpenClaw 内置
    
    def web_search(self, query: str, count: int = 5) -> List[Dict]:
        """网络搜索"""
        # 实际会调用 web_search 工具
        return [{
            "query": query,
            "status": "需要调用 web_search 工具"
        }]
    
    def local_search(self, query: str, directory: str) -> List[str]:
        """本地搜索"""
        try:
            import subprocess
            result = subprocess.run(
                ["grep", "-r", "-l", query, directory],
                capture_output=True,
                text=True
            )
            return result.stdout.strip().split("\n")
        except:
            return []


class GitTool:
    """Git 工具"""
    
    def __init__(self):
        self.available = self._check_available()
    
    def _check_available(self) -> bool:
        try:
            result = subprocess.run(
                ["git", "--version"],
                capture_output=True,
                text=True
            )
            return result.returncode == 0
        except:
            return False
    
    def status(self, repo_path: str) -> Dict:
        """获取状态"""
        if not self.available:
            return {"error": "Git 未安装"}
        
        try:
            result = subprocess.run(
                ["git", "status", "--porcelain"],
                cwd=repo_path,
                capture_output=True,
                text=True
            )
            return {
                "changed": result.stdout.strip().split("\n") if result.stdout else []
            }
        except Exception as e:
            return {"error": str(e)}
    
    def log(self, repo_path: str, count: int = 10) -> List[Dict]:
        """获取日志"""
        if not self.available:
            return []
        
        try:
            result = subprocess.run(
                ["git", "log", f"-{count}", "--pretty=format:%H|%s|%an|%ai"],
                cwd=repo_path,
                capture_output=True,
                text=True
            )
            
            logs = []
            for line in result.stdout.strip().split("\n"):
                parts = line.split("|")
                if len(parts) >= 4:
                    logs.append({
                        "hash": parts[0],
                        "message": parts[1],
                        "author": parts[2],
                        "date": parts[3]
                    })
            
            return logs
        except:
            return []


def main():
    """测试工具"""
    tools = UnifiedTools()
    
    print("可用工具:")
    for name in tools.list_tools():
        tool = tools.get_tool(name)
        available = getattr(tool, "available", True)
        print(f"  {name}: {'✅' if available else '❌'}")
    
    # 测试 PDF
    print("\nPDF 工具测试:")
    pdf_tool = tools.get_tool("pdf")
    print(f"  可用: {pdf_tool.available}")
    
    # 测试 Git
    print("\nGit 工具测试:")
    git_tool = tools.get_tool("git")
    print(f"  可用: {git_tool.available}")
    status = git_tool.status(Path.home() / ".openclaw" / "workspace")
    print(f"  状态: {status}")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
教学工作手册 - 质量审查脚本 v1.0
**Reviewer Pattern**: 按清单逐项检查生成文档

使用方式:
    python3 scripts/review_document.py 输出文件.docx [--verbose]
"""

import sys
import os
import json
from docx import Document
from typing import Dict, List, Any

def review_document(docx_file: str, verbose: bool = False) -> Dict[str, Any]:
    """审查生成的文档
    
    Returns:
        {
            'file': str,
            'status': 'PASS' | 'WARN' | 'FAIL',
            'checks': {...},
            'stats': {...},
            'warnings': [...],
            'errors': [...]
        }
    """
    result = {
        'file': docx_file,
        'status': 'PASS',
        'checks': {},
        'stats': {},
        'warnings': [],
        'errors': []
    }
    
    if not os.path.exists(docx_file):
        result['status'] = 'FAIL'
        result['errors'].append(f'文件不存在: {docx_file}')
        return result
    
    try:
        doc = Document(docx_file)
    except Exception as e:
        result['status'] = 'FAIL'
        result['errors'].append(f'无法打开文档: {e}')
        return result
    
    # 1. 检查占位符替换
    result['checks']['placeholders'] = check_placeholders(doc)
    
    # 2. 检查作业表
    hw_ok, hw_stats = check_homework_tables(doc)
    result['checks']['homework_tables'] = hw_ok
    result['stats'].update(hw_stats)
    
    # 3. 检查成绩表
    gr_ok, gr_stats = check_grade_tables(doc)
    result['checks']['grade_tables'] = gr_ok
    result['stats'].update(gr_stats)
    
    # 综合判断
    if not all(result['checks'].values()):
        result['status'] = 'FAIL'
    elif result['warnings']:
        result['status'] = 'WARN'
    
    return result

def check_placeholders(doc) -> bool:
    """检查占位符是否全部替换"""
    for para in doc.paragraphs:
        if '$' in para.text:
            return False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '$' in para.text and para.text.strip():
                        return False
    return True

def check_homework_tables(doc) -> tuple:
    """检查作业表"""
    stats = {}
    
    # 找作业表
    homework_tables = []
    for i, table in enumerate(doc.tables):
        first_row_text = ''.join([c.text.strip().replace('\n', '') for c in table.rows[0].cells])
        if '作业1' in first_row_text:
            homework_tables.append((i, table))
    
    if not homework_tables:
        return False, {'hw_tables': 0}
    
    stats['hw_tables'] = len(homework_tables)
    
    # 检查每张表的行数
    table_rows = []
    for idx, table in homework_tables:
        rows = len(table.rows) - 1  # 减去表头
        table_rows.append(rows)
    stats['hw_rows'] = table_rows
    
    # 检查是否有空行
    empty_rows = 0
    for idx, table in homework_tables:
        for row in table.rows[1:]:
            if not row.cells[0].text.strip():
                empty_rows += 1
    stats['hw_empty'] = empty_rows
    
    # 检查学号姓名是否正确填充
    filled = sum(1 for idx, table in homework_tables
                  for row in table.rows[1:]
                  if row.cells[0].text.strip())
    stats['hw_filled'] = filled
    
    return True, stats

def check_grade_tables(doc) -> tuple:
    """检查成绩表"""
    stats = {}
    
    # 找成绩表
    grade_tables = []
    for i, table in enumerate(doc.tables):
        first_row_text = ''.join([c.text.strip().replace('\n', '') for c in table.rows[0].cells])
        if '平时成绩' in first_row_text:
            grade_tables.append((i, table))
    
    if not grade_tables:
        return False, {'gr_tables': 0}
    
    stats['gr_tables'] = len(grade_tables)
    
    # 检查每张表的行数
    table_rows = []
    for idx, table in grade_tables:
        rows = len(table.rows) - 1
        table_rows.append(rows)
    stats['gr_rows'] = table_rows
    
    # 检查是否有空行
    empty_rows = 0
    for idx, table in grade_tables:
        for row in table.rows[1:]:
            if not row.cells[0].text.strip():
                empty_rows += 1
    stats['gr_empty'] = empty_rows
    
    # 检查成绩是否正确填充
    filled = sum(1 for idx, table in grade_tables
                  for row in table.rows[1:]
                  if row.cells[0].text.strip() and len(row.cells) > 8 and row.cells[8].text.strip())
    stats['gr_filled'] = filled
    
    return True, stats

def print_report(result: Dict[str, Any], verbose: bool = False):
    """打印审查报告"""
    print("=" * 60)
    print("📋 教学工作手册 - 质量审查报告 v1.0")
    print("=" * 60)
    print(f"文件: {os.path.basename(result['file'])}")
    
    status_icon = '✅ PASS' if result['status'] == 'PASS' else '⚠️ WARN' if result['status'] == 'WARN' else '❌ FAIL'
    print(f"状态: {status_icon}")
    print()
    
    print("📊 检查项:")
    check_names = {
        'placeholders': '占位符替换',
        'homework_tables': '作业表',
        'grade_tables': '成绩表'
    }
    for check, passed in result['checks'].items():
        icon = '✅' if passed else '❌'
        name = check_names.get(check, check)
        print(f"  {icon} {name}")
    
    if result['stats']:
        print()
        print("📈 统计:")
        if 'hw_tables' in result['stats']:
            print(f"  作业表: {result['stats']['hw_tables']}张")
            print(f"  作业表行数: {result['stats'].get('hw_rows', [])}")
            print(f"  作业表填充: {result['stats'].get('hw_filled', 0)}/空{result['stats'].get('hw_empty', 0)}")
        if 'gr_tables' in result['stats']:
            print(f"  成绩表: {result['stats']['gr_tables']}张")
            print(f"  成绩表行数: {result['stats'].get('gr_rows', [])}")
            print(f"  成绩表填充: {result['stats'].get('gr_filled', 0)}/空{result['stats'].get('gr_empty', 0)}")
    
    if result['warnings']:
        print()
        print("⚠️ 警告:")
        for w in result['warnings']:
            print(f"  • {w}")
    
    if result['errors']:
        print()
        print("❌ 错误:")
        for e in result['errors']:
            print(f"  • {e}")
    
    print()
    print("=" * 60)
    return result['status'] == 'PASS'

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 review_document.py <docx文件> [--verbose]")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    verbose = '--verbose' in sys.argv
    json_only = '--json' in sys.argv
    
    result = review_document(docx_file, verbose)
    success = print_report(result, verbose)
    
    if json_only or '--json' in sys.argv:
        print(json.dumps(result, indent=2, ensure_ascii=False))
    
    sys.exit(0 if success else 1)

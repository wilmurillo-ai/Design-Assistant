#!/usr/bin/env python3
"""
教学工作手册替换工具 v4.6
根据Excel数据替换Word模板中的内容，生成教学工作手册。
"""

import pandas as pd
import config_loader
import shutil
import random
from docx import Document
from docx.oxml.ns import qn
from datetime import datetime
import re, os, sys
import copy
from typing import Dict, List, Any

random.seed(42)

# 模板目录：支持 config.env 配置，fallback 到本地默认路径
def _get_template_dir():
    try:
        d = config_loader.get_template_dir()
        if d:
            return d
    except Exception:
        pass
    return os.path.expanduser('~/Documents/QZTC/教学文档模版')

TEMPLATE_DIR = None  # lazy init

def _ensure_template_dir():
    global TEMPLATE_DIR
    if TEMPLATE_DIR is None:
        TEMPLATE_DIR = _get_template_dir()
    return TEMPLATE_DIR

# ========== 专业名称提取 ==========
def extract_major(class_name: str) -> str:
    class_name = str(class_name).strip()
    match = re.match(r'\d+级\s*(.+?)(?:\d+班)?$', class_name)
    if match:
        major = match.group(1).strip()
        if major.endswith('班'):
            major = major[:-1]
        return major if major else class_name
    result = re.sub(r'^\d+级|\d+班$', '', class_name).strip()
    return result if result else class_name

# ========== 添加行（防御性） ==========
def add_row_to_table(table, source_row_idx=1):
    """添加新行（防御性合并单元格处理）"""
    source_tr = table.rows[source_row_idx]._element
    new_tr = copy.deepcopy(source_tr)
    
    for tc in new_tr.iter(qn('w:tc')):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            for tag in ['vMerge', 'gridSpan']:
                elem = tcPr.find(qn(f'w:{tag}'))
                if elem is not None:
                    tcPr.remove(elem)
    
    table._element.append(new_tr)

# ========== 单元格居中对齐 ==========
def set_cell_centered(cell):
    """设置单元格和段落水平垂直居中"""
    from docx.oxml.ns import qn
    
    # 获取或创建tcPr
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    
    # 水平居中
    jc = tcPr.find(qn('w:jc'))
    if jc is None:
        jc = tcPr.makeelement(qn('w:jc'), {})
        tcPr.append(jc)
    jc.set(qn('w:val'), 'center')
    
    # 垂直居中
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = tcPr.makeelement(qn('w:vAlign'), {})
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), 'center')
    
    # 设置段落也居中
    for para in cell.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = para._element.makeelement(qn('w:pPr'), {})
            para._element.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = pPr.makeelement(qn('w:jc'), {})
            pPr.append(jc)
        jc.set(qn('w:val'), 'center')

# ========== 文本替换 ==========
def simple_replace(doc, replacements):
    """安全文本替换（保持原有格式）
    
    按长度降序排列占位符，逐run处理
    """
    # 按长度降序排列，避免短占位符先匹配
    replacements = sorted(replacements, key=lambda x: len(x[0]), reverse=True)
    
    # 直接操作底层 XML，移除所有 proofErr 元素
    from docx.oxml.ns import qn
    for elem in doc.element.body.iter():
        if elem.tag.endswith('}proofErr'):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
    
    # 收集所有段落
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    
    for para in all_paras:
        runs = list(para.runs)
        if not runs:
            continue
        
        i = 0
        while i < len(runs):
            run = runs[i]
            if not run.text or '{' not in run.text:
                i += 1
                continue
            
            # run包含{（占位符）
            # 尝试先在当前run中匹配完整占位符
            matched = False
            for old, new in replacements:
                if old in run.text:
                    run.text = run.text.replace(old, new, 1)
                    matched = True
                    break
            
            if matched:
                i += 1
                continue
            
            # 合并后续run直到凑成完整占位符
            combined_runs = [run]
            combined_text = run.text
            j = i + 1
            found_dollar_end = False
            while j < len(runs) and not found_dollar_end:
                combined_text += runs[j].text
                combined_runs.append(runs[j])
                # 检查是否遇到完整的$占位符
                for old, new in replacements:
                    if old in combined_text:
                        combined_text = combined_text.replace(old, new, 1)
                        found_dollar_end = True
                        break
                j += 1
            
            # 写回
            run.text = combined_text
            for r in combined_runs[1:]:
                r.text = ''
            
            i += 1
    
    return True

# ========== 检测表格类型 ==========
def detect_table_type(table):
    """检测表格类型：作业表 或 成绩表
    
    - 作业表：表头包含"作业1"
    - 成绩表：表头包含"平时成绩"
    """
    row0_text = ''.join([cell.text.strip().replace('\n', '') for cell in table.rows[0].cells])
    
    if '作业1' in row0_text:
        return 'homework'
    if '平时成绩' in row0_text:
        return 'grade'
    return 'unknown'

# ========== 主函数 ==========
def generate_manual(excel_file: str, template_dir: str = None) -> Dict[str, Any]:
    """生成教学工作手册
    
    Args:
        excel_file: Excel文件路径，格式：{课程名称}-{班级}.xls
        template_dir: 模板目录，默认使用 QZTC/教学文档模版
    
    Returns:
        {
            'success': bool,
            'course': str,
            'students': int,
            'output_file': str,
            'warnings': [str],
            'errors': [str]
        }
    """
    global TEMPLATE_DIR
    if template_dir:
        TEMPLATE_DIR = template_dir
    else:
        _ensure_template_dir()
    
    basename = os.path.basename(excel_file)
    name_part = basename.replace('.xls', '')
    
    # 解析课程和班级
    parts = name_part.rsplit('-', 1)
    course_name = parts[0]
    class_name = parts[1] if len(parts) > 1 else ''
    major = extract_major(class_name)
    
    # 查找模板
    template_file = os.path.join(TEMPLATE_DIR, f'教学工作手册-{course_name}-模版.docx')
    if not os.path.exists(template_file):
        return {
            'success': False,
            'errors': [f'模板不存在: {template_file}']
        }
    
    # 读取Excel
    df = pd.read_excel(excel_file)
    students = df[['学号', '姓名', '平时成绩', '期中成绩', '期末成绩', '总评成绩']].copy()
    n = len(students)
    
    # 复制模板
    basename = os.path.basename(excel_file).replace('.xls', '')
    output_file = os.path.join(os.path.dirname(excel_file), f'教学工作手册-{basename}.docx')
    shutil.copy(template_file, output_file)
    
    # 打开并修改
    doc = Document(output_file)
    
    # ========== 读取完整数据（含备注列）==========
    df_full = pd.read_excel(excel_file)
    total_students = n
    # 实到人数：备注为"正常"的学生（即正常考试）
    if '备注' in df_full.columns:
        normal_mask = df_full['备注'] == '正常'
        took_exam = normal_mask.sum()
        normal_scores = df_full.loc[normal_mask, '总评成绩'].dropna()
    else:
        took_exam = students['总评成绩'].notna().sum()
        normal_scores = students['总评成绩'].dropna()
    
    # 分数段统计（基于正常考试成绩）
    count_90plus = (normal_scores >= 90).sum()
    count_80_89 = ((normal_scores >= 80) & (normal_scores < 90)).sum()
    count_70_79 = ((normal_scores >= 70) & (normal_scores < 80)).sum()
    count_60_69 = ((normal_scores >= 60) & (normal_scores < 70)).sum()
    count_below60 = (normal_scores < 60).sum()
    
    # 统计数据
    avg_score = normal_scores.mean() if len(normal_scores) > 0 else 0
    pass_rate = (normal_scores >= 60).sum() / took_exam * 100 if took_exam > 0 else 0
    
    # 找到 Table 13 并填充
    for i, table in enumerate(doc.tables):
        if len(table.rows) >= 5 and len(table.columns) >= 8:
            first_cell = table.rows[0].cells[0].text.strip()
            if '学期考核成绩结构分布情况' in first_cell:
                # Row 3（第3行）: 人数
                row3 = table.rows[3]
                row3.cells[1].text = str(total_students)   # 应到人数
                row3.cells[2].text = str(int(took_exam))  # 实到人数
                row3.cells[3].text = str(int(count_90plus))
                row3.cells[4].text = str(int(count_80_89))
                row3.cells[5].text = str(int(count_70_79))
                row3.cells[6].text = str(int(count_60_69))
                row3.cells[7].text = str(int(count_below60))
                for c in row3.cells:
                    set_cell_centered(c)
                
                # Row 4（第4行）: 百分比
                row4 = table.rows[4]
                pct = lambda x: f'{x/total_students*100:.1f}%' if total_students > 0 else '0%'
                row4.cells[1].text = '100%'
                row4.cells[2].text = pct(took_exam)
                row4.cells[3].text = pct(count_90plus)
                row4.cells[4].text = pct(count_80_89)
                row4.cells[5].text = pct(count_70_79)
                row4.cells[6].text = pct(count_60_69)
                row4.cells[7].text = pct(count_below60)
                for c in row4.cells:
                    set_cell_centered(c)
                break
    
    # 文本替换
    now = datetime.now()
    grade_str = class_name[:2] if class_name else '23'
    # 计算及格率
    pass_count = (normal_scores >= 60).sum() if len(normal_scores) > 0 else 0
    ok_rate = f'{(pass_count/took_exam*100):.1f}' if took_exam > 0 else '0'
    
    simple_replace(doc, [
        ('{g}{mj}', f'{grade_str}{major}'),  # 复合占位符
        ('{as}', '2025 - 2026'),
        ('{sm}', '二'),
        ('{g}', grade_str),
        ('{mj}', major),
        ('{tot}', str(n)),
        ('{v}', f'{avg_score:.1f}'),
        ('{ok}', ok_rate),           # 及格率（格式：人数/总人数）
        ('{y}', str(now.year)),
        ('{m}', str(now.month)),
        ('{d}', str(now.day)),
    ])
    
    # 分类表格
    homework_tables = []
    grade_tables = []
    
    for i, table in enumerate(doc.tables):
        t_type = detect_table_type(table)
        if t_type == 'homework':
            homework_tables.append((i, table))
        elif t_type == 'grade':
            grade_tables.append((i, table))
    
    print(f"  找到作业表: {[idx for idx, _ in homework_tables]}")
    print(f"  找到成绩表: {[idx for idx, _ in grade_tables]}")
    
    # 第一阶段：填充作业表（分页：每页固定20学生）
    PAGE_SIZE = 20
    hw_student_idx = 0
    for table_idx, table in homework_tables:
        # 计算本表的起始和结束学生索引
        start_idx = hw_student_idx
        end_idx = min(start_idx + PAGE_SIZE, n)
        table_student_count = end_idx - start_idx
        needed_rows = 1 + PAGE_SIZE  # 每表固定21行（1表头+20数据）
        
        # 确保有足够的行
        while len(table.rows) < needed_rows:
            add_row_to_table(table, source_row_idx=1)
        
        # 填入作业数据（√）
        for local_i in range(PAGE_SIZE):
            row_idx = local_i + 1
            cells = table.rows[row_idx].cells
            if local_i < table_student_count:
                # 有数据
                global_i = start_idx + local_i
                row_data = students.iloc[global_i]
                cells[0].text = str(row_data['学号'])
                cells[1].text = str(row_data['姓名'])
                # 作业列填√（从cells[2]到cells[14]是作业1-13）
                for j in range(2, 15):
                    if j < len(cells):
                        cells[j].text = '√'
            else:
                # 无数据，清空行
                cells[0].text = ''
                cells[1].text = ''
                for j in range(2, min(15, len(cells))):
                    cells[j].text = ''
            # 水平垂直居中
            for j in range(min(15, len(cells))):
                set_cell_centered(cells[j])
            hw_student_idx += 1
    
    print(f"  作业表填充完成: {hw_student_idx}学生")
    
    # 第二阶段：填充成绩表（分页：每页固定20学生）
    grade_student_idx = 0
    for table_idx, table in grade_tables:
        # 计算本表的起始和结束学生索引
        start_idx = grade_student_idx
        end_idx = min(start_idx + PAGE_SIZE, n)
        table_student_count = end_idx - start_idx
        needed_rows = 1 + PAGE_SIZE  # 每表固定21行（1表头+20数据）
        
        # 确保有足够的行
        while len(table.rows) < needed_rows:
            add_row_to_table(table, source_row_idx=1)
        
        # 填入成绩数据
        for local_i in range(PAGE_SIZE):
            row_idx = local_i + 1
            cells = table.rows[row_idx].cells
            if local_i < table_student_count:
                # 有数据
                global_i = start_idx + local_i
                row_data = students.iloc[global_i]
                cells[0].text = str(row_data['学号'])
                cells[1].text = str(row_data['姓名'])
                # 成绩列：cells[8]=平时, cells[9]=期中, cells[10]=期末, cells[11]=总评
                cells[8].text = str(int(row_data['平时成绩'])) if pd.notna(row_data['平时成绩']) else ''
                cells[9].text = str(int(row_data['期中成绩'])) if pd.notna(row_data['期中成绩']) else ''
                cells[10].text = str(int(row_data['期末成绩'])) if pd.notna(row_data['期末成绩']) else ''
                cells[11].text = str(int(row_data['总评成绩'])) if pd.notna(row_data['总评成绩']) else ''
            else:
                # 无数据，清空行
                cells[0].text = ''
                cells[1].text = ''
                if len(cells) > 8:
                    for j in range(8, min(12, len(cells))):
                        cells[j].text = ''
            # 水平垂直居中
            for j in range(min(12, len(cells))):
                set_cell_centered(cells[j])
            grade_student_idx += 1
    
    print(f"  成绩表填充完成: {grade_student_idx}学生")
    
    # 保存
    doc.save(output_file)
    
    return {
        'success': True,
        'course': course_name,
        'students': n,
        'output_file': output_file,
        'warnings': [],
        'errors': []
    }

# ========== 批量处理 ==========
def batch_generate(directory: str, template_dir: str = None) -> List[Dict[str, Any]]:
    """批量处理目录下所有Excel文件"""
    results = []
    for fname in os.listdir(directory):
        if fname.endswith('.xls') or fname.endswith('.xlsx'):
            fpath = os.path.join(directory, fname)
            print(f"\n>>> 处理: {fname}")
            result = generate_manual(fpath, template_dir)
            if result['success']:
                print(f"✅ 成功！学生数: {result['students']}")
            else:
                print(f"❌ 失败: {result.get('errors', [])}")
            results.append(result)
    return results

# ========== Pipeline 模式 ==========
def run_pipeline(excel_file: str, template_dir: str = None):
    """Pipeline模式: 生成 + 审查"""
    print("=" * 60)
    print("🔄 Pipeline: 生成 + 审查")
    print("=" * 60)
    
    # Step 1: 生成
    print("\n📝 STEP 1: 生成文档")
    print("-" * 40)
    result = generate_manual(excel_file, template_dir)
    
    if not result['success']:
        print(f"❌ 生成失败: {result.get('errors', [])}")
        return result
    
    print(f"✅ 生成成功！")
    print(f"   课程: {result['course']}")
    print(f"   学生数: {result['students']}")
    
    # Step 2: 审查
    print("\n🔍 STEP 2: 审查文档")
    print("-" * 40)
    
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from scripts.review_document import review_document, print_report
    
    review_result = review_document(result['output_file'])
    print_report(review_result)
    
    print("\n" + "=" * 60)
    if review_result['status'] == 'PASS':
        print("✅ Pipeline 完成！文档通过所有检查")
    elif review_result['status'] == 'WARN':
        print("⚠️ Pipeline 完成！有警告，请检查")
    else:
        print("❌ Pipeline 完成！有检查项未通过")
    print("=" * 60)
    
    return result

def batch_review(directory: str, template_dir: str = None) -> List[Dict]:
    """批量生成 + 审查"""
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from scripts.review_document import review_document
    
    results = []
    files = [f for f in os.listdir(directory) if f.endswith('.xls') or f.endswith('.xlsx')]
    total = len(files)
    
    print("=" * 60)
    print(f"🔄 批量生成 + 审查 ({total} 个文件)")
    print("=" * 60)
    
    for i, fname in enumerate(sorted(files), 1):
        fpath = os.path.join(directory, fname)
        print(f"\n[{i}/{total}] 处理: {fname}")
        print("-" * 40)
        
        gen_result = generate_manual(fpath, template_dir)
        
        if gen_result['success']:
            rev_result = review_document(gen_result['output_file'])
            status = '✅' if rev_result['status'] == 'PASS' else '⚠️' if rev_result['status'] == 'WARN' else '❌'
            print(f"  {status} {rev_result['status']}")
        else:
            print(f"  ❌ 生成失败")
            rev_result = {'status': 'FAIL'}
        
        results.append({'file': fname, 'generate': gen_result, 'review': rev_result})
    
    # 汇总
    passed = sum(1 for r in results if r['review']['status'] == 'PASS')
    warned = sum(1 for r in results if r['review']['status'] == 'WARN')
    failed = sum(1 for r in results if r['review']['status'] == 'FAIL')
    
    print("\n" + "=" * 60)
    print("📊 批量审查汇总")
    print("=" * 60)
    print(f"总计: {total} 个文件")
    print(f"  ✅ PASS: {passed}")
    print(f"  ⚠️ WARN: {warned}")
    print(f"  ❌ FAIL: {failed}")
    print("=" * 60)
    
    return results

# ========== CLI ==========
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("""用法: python3 manual_generator.py <excel文件或目录> [选项]

选项:
  <excel文件>              Excel文件路径
  --pipeline               Pipeline模式（生成+审查）
  --batch                  批量生成（不审查）
  --batch-review           批量生成+审查
  [模板目录]               可选，指定模板目录

示例:
  # 单文件生成
  python3 manual_generator.py 数据可视化-23级计算机.xls
  
  # Pipeline模式（生成+审查）
  python3 manual_generator.py 数据可视化-23级计算机.xls --pipeline
  
  # 批量审查
  python3 manual_generator.py --batch-review /path/to/excels/
""")
        sys.exit(1)
    
    # 检查选项
    if '--pipeline' in sys.argv:
        sys.argv.remove('--pipeline')
        excel_file = sys.argv[1] if len(sys.argv) > 1 and not sys.argv[1].startswith('-') else None
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        if excel_file:
            run_pipeline(excel_file, template_dir)
        else:
            print("❌ 请指定Excel文件")
            sys.exit(1)
    elif '--batch-review' in sys.argv:
        sys.argv.remove('--batch-review')
        directory = sys.argv[1] if len(sys.argv) > 1 and not sys.argv[1].startswith('-') else '.'
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        batch_review(directory, template_dir)
    elif '--batch' in sys.argv:
        sys.argv.remove('--batch')
        directory = sys.argv[1] if len(sys.argv) > 1 and not sys.argv[1].startswith('-') else '.'
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        batch_generate(directory, template_dir)
    elif '*' in sys.argv[1]:
        # glob批量处理
        import glob
        for f in glob.glob(sys.argv[1]):
            print(f"\n>>> 处理: {f}")
            result = generate_manual(f, sys.argv[2] if len(sys.argv) > 2 else None)
            if result['success']:
                print(f"✅ 成功！学生数: {result['students']}, 输出: {os.path.basename(result['output_file'])}")
            else:
                print(f"❌ 失败: {result.get('errors', [])}")
    else:
        excel_file = sys.argv[1]
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        result = generate_manual(excel_file, template_dir)
        if result['success']:
            print(f"\n✅ 成功！")
            print(f"   课程: {result['course']}")
            print(f"   学生数: {result['students']}")
            print(f"   输出: {result['output_file']}")
        else:
            print(f"\n❌ 失败: {result.get('errors', [])}")

#!/usr/bin/env python3
"""
课程目标达成情况分析表生成工具 v5.2
根据Excel成绩数据自动匹配模板，生成课程目标达成情况分析表。

v5.2 优化:
- 权重 WEIGHTS 从模板表7第3行动态读取
- 目标配置 TARGET_CONFIG 从模板表7各目标理论满分行动态读取
- total = col3×权重.col3 + col4×权重.col4 + col5×权重.col5 自动计算

v5.1 优化:
- 预检验：生成前检查Excel和模板
- 错误隔离：图表更新失败不整体失败
- Chart X轴对齐：chart1-4同时更新X(序号)和Y(达成度)
- 批量审查：--batch-review选项
"""

import pandas as pd
import shutil
import random
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re, os, sys
import copy
from typing import Dict, List, Any, Tuple

# 本地模块
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import config_loader

random.seed(42)

# ========== 单元格居中对齐 ==========
def set_cell_centered(cell):
    """设置单元格和段落水平垂直居中"""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    
    # 水平居中
    jc = tcPr.find(qn('w:jc'))
    if jc is None:
        jc = tcPr.makeelement(qn('w:jc'), {})
        tcPr.append(jc)
    jc.set(qn('w:val'), 'center')
    
    # 垂直居中
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = tcPr.makeelement(qn('w:vAlign'), {})
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), 'center')
    
    # 设置段落也居中
    for para in cell.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = para._element.makeelement(qn('w:pPr'), {})
            para._element.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = pPr.makeelement(qn('w:jc'), {})
            pPr.append(jc)
        jc.set(qn('w:val'), 'center')

# ========== 配置 ==========
# 注意：权重和 TARGET_CONFIG 均从模板表7动态读取，见 read_weights_and_config_from_template()

TARGET_CONFIG = {
    'target1': {'col3': 30, 'col4': 20, 'col5': 20, 'total': 22},
    'target2': {'col3': 30, 'col4': 30, 'col5': 30, 'total': 30},
    'target3': {'col3': 20, 'col4': 30, 'col5': 20, 'total': 21},
    'target4': {'col3': 20, 'col4': 20, 'col5': 30, 'total': 27},
}

# 汇总表中课程目标1-4的行号（从模板结构得出）
SUMMARY_TARGET_ROWS = {
    '课程目标1': 3,
    '课程目标2': 6,
    '课程目标3': 9,
    '课程目标4': 12,
}

TEMPLATE_DIR = '/Users/qztcm09/Documents/QZTC/教学文档模版'

# ========== 权重 & 目标配置读取 ==========
def read_weights_and_config_from_template(doc) -> Tuple[Dict[str, float], Dict[str, Dict]]:
    """从模板表7动态读取权重和目标配置
    
    表7 结构:
    - Row 2: 权重行 → col3/col4/col5 为权重系数
    - Row 4: target1 理论满分行 → col3/col4/col5 为各考核项满分
    - Row 7: target2 理论满分行
    - Row 10: target3 理论满分行
    - Row 13: target4 理论满分行
    
    total = col3×权重.col3 + col4×权重.col4 + col5×权重.col5
    
    Returns:
        (weights, target_config)
        weights: {'col3': float, 'col4': float, 'col5': float}
        target_config: {'target1': {'col3': int, 'col4': int, 'col5': int, 'total': float}, ...}
    """
    t7 = doc.tables[7]
    
    # 读取权重（Row 2, cells 3/4/5）
    weights_row = t7.rows[2]
    w_col3 = float(weights_row.cells[3].text.strip())
    w_col4 = float(weights_row.cells[4].text.strip())
    w_col5 = float(weights_row.cells[5].text.strip())
    weights = {'col3': w_col3, 'col4': w_col4, 'col5': w_col5}
    
    # 读取各目标理论满分（Row 4/7/10/13, cells 3/4/5）
    target_rows = {
        'target1': 4,
        'target2': 7,
        'target3': 10,
        'target4': 13,
    }
    
    target_config = {}
    for name, row_idx in target_rows.items():
        row = t7.rows[row_idx]
        t_col3 = int(row.cells[3].text.strip())
        t_col4 = int(row.cells[4].text.strip())
        t_col5 = int(row.cells[5].text.strip())
        total = t_col3 * w_col3 + t_col4 * w_col4 + t_col5 * w_col5
        target_config[name] = {
            'col3': t_col3,
            'col4': t_col4,
            'col5': t_col5,
            'total': total,
        }
    
    return weights, target_config


# ========== 预检验 ==========
def validate_inputs(excel_file: str, template_dir: str = None) -> Tuple[bool, str]:
    """预检验：检查Excel文件和模板是否有效
    
    Returns:
        (is_valid, error_message)
    """
    if template_dir is None:
        template_dir = TEMPLATE_DIR
    
    # 1. 检查Excel文件存在
    if not os.path.exists(excel_file):
        return False, f"Excel文件不存在: {excel_file}"
    
    # 2. 检查Excel文件可读
    try:
        df = pd.read_excel(excel_file)
    except Exception as e:
        return False, f"无法读取Excel文件: {e}"
    
    # 3. 检查必要列
    required_cols = ['学号', '姓名', '平时成绩', '期中成绩', '期末成绩']
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        return False, f"Excel缺少必要列: {missing_cols}"
    
    # 4. 检查学生数量
    if len(df) == 0:
        return False, "Excel中没有学生数据"
    
    # 5. 解析课程名称并检查模板
    basename = os.path.basename(excel_file)
    name_part = basename.replace('.xls', '')
    parts = name_part.rsplit('-', 1)
    course_name = parts[0] if len(parts) > 1 else ''
    
    template_file = os.path.join(template_dir, f'课程目标达成情况分析表-{course_name}-模版.docx')
    if not os.path.exists(template_file):
        # 尝试查找同名模板
        templates = [f for f in os.listdir(template_dir) 
                     if f.startswith('课程目标达成情况分析表-') and f.endswith('-模版.docx')]
        if templates:
            return True, f"模板名称不精确匹配，使用: {templates[0]}"
        return False, f"模板文件不存在: {template_file}"
    
    return True, ""

# ========== 图表更新（错误隔离 + X轴对齐） ==========
def update_chart(output_path, v1_list, v2_list, v3_list, v4_list, avg_v1, avg_v2, avg_v3, avg_v4) -> Tuple[bool, str]:
    """更新所有图表（带错误隔离）
    
    - chart1-4: 散点图，X=学生序号，Y=达成度百分比
    - chart5: 柱状图，各课程目标平均达成度百分比
    
    Returns:
        (success, warning_message)
    """
    import zipfile
    from lxml import etree
    
    # 计算平均达成度百分比
    avg_pcts = [avg_v1 * 100, avg_v2 * 100, avg_v3 * 100, avg_v4 * 100]
    
    n = len(v1_list)
    warnings = []
    
    try:
        # 重新打包docx
        tmp_path = output_path + '.tmp'
        with zipfile.ZipFile(output_path, 'r') as z_in:
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
                for item in z_in.infolist():
                    fname = item.filename
                    if fname.startswith('word/charts/chart') and fname.endswith('.xml'):
                        import re
                        m = re.search(r'chart(\d+)\.xml', fname)
                        if m:
                            chart_num = int(m.group(1))
                            chart_xml = z_in.read(fname)
                            root = etree.fromstring(chart_xml)
                            ns = {'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart'}
                            
                            if chart_num <= 4:
                                # chart1-4: 散点图，需要更新X轴(xVal)和Y轴(yVal)
                                ser = root.find('.//c:ser', ns)
                                if ser is not None:
                                    values = [v1_list, v2_list, v3_list, v4_list][chart_num - 1]
                                    
                                    # 更新Y轴数值 (c:yVal -> c:numRef -> c:numCache)
                                    yVal = ser.find('c:yVal', ns)
                                    if yVal is not None:
                                        numRef = yVal.find('c:numRef', ns)
                                        if numRef is not None:
                                            numCache = numRef.find('c:numCache', ns)
                                            if numCache is not None:
                                                # 清空旧数据
                                                for pt in numCache.findall('c:pt', ns):
                                                    numCache.remove(pt)
                                                # 更新ptCount
                                                ptCount = numCache.find('c:ptCount', ns)
                                                if ptCount is not None:
                                                    ptCount.set('val', str(n))
                                                else:
                                                    ptCount = etree.SubElement(numCache, '{http://schemas.openxmlformats.org/drawingml/2006/chart}ptCount')
                                                    ptCount.set('val', str(n))
                                                # 写入新数据
                                                for i, v in enumerate(values):
                                                    pt = etree.SubElement(numCache, '{http://schemas.openxmlformats.org/drawingml/2006/chart}pt')
                                                    pt.set('idx', str(i))
                                                    v_elem = etree.SubElement(pt, '{http://schemas.openxmlformats.org/drawingml/2006/chart}v')
                                                    v_elem.text = str(v * 100)
                                    
                                    # 更新X轴类别 (c:xVal -> c:strRef -> c:strCache)
                                    xVal = ser.find('c:xVal', ns)
                                    if xVal is not None:
                                        strRef = xVal.find('c:strRef', ns)
                                        if strRef is not None:
                                            strCache = strRef.find('c:strCache', ns)
                                            if strCache is not None:
                                                # 清空旧数据
                                                for pt in strCache.findall('c:pt', ns):
                                                    strCache.remove(pt)
                                                # 更新ptCount
                                                ptCount = strCache.find('c:ptCount', ns)
                                                if ptCount is not None:
                                                    ptCount.set('val', str(n))
                                                else:
                                                    ptCount = etree.SubElement(strCache, '{http://schemas.openxmlformats.org/drawingml/2006/chart}ptCount')
                                                    ptCount.set('val', str(n))
                                                # 写入新数据 (学生序号1,2,3...)
                                                for i in range(n):
                                                    pt = etree.SubElement(strCache, '{http://schemas.openxmlformats.org/drawingml/2006/chart}pt')
                                                    pt.set('idx', str(i))
                                                    v_elem = etree.SubElement(pt, '{http://schemas.openxmlformats.org/drawingml/2006/chart}v')
                                                    v_elem.text = str(i + 1)
                                    
                                    # 禁用外部数据引用（如果有）- 移除externalData元素
                                    for extData in root.findall('.//c:externalData', ns):
                                        extData.getparent().remove(extData)
                                    
                                    z_out.writestr(item, etree.tostring(root))
                                    continue
                            elif chart_num == 5:
                                # chart5: 柱状图，只更新Y轴
                                numCaches = root.findall('.//c:numCache', ns)
                                if numCaches:
                                    nc = numCaches[0]
                                    pts = nc.findall('c:pt', ns)
                                    for pt, val in zip(pts, avg_pcts):
                                        v_elem = pt.find('c:v', ns)
                                        if v_elem is not None:
                                            v_elem.text = str(val)
                                    z_out.writestr(item, etree.tostring(root))
                                    continue
                    
                    z_out.writestr(item, z_in.read(item.filename))
        
        os.replace(tmp_path, output_path)
    except Exception as e:
        return False, f"图表更新失败: {e}"
    
    return True, "; ".join(warnings) if warnings else ""

# ========== 分析段落更新 ==========
def update_analysis_paragraphs(doc, v1_list, v2_list, v3_list, v4_list):
    """更新表8后面的4段分析段落
    
    计算统计数据并替换占位符
    """
    n = len(v1_list)
    
    # 计算统计数据
    stats = {}
    for name, vlist in [('v1', v1_list), ('v2', v2_list), ('v3', v3_list), ('v4', v4_list)]:
        avg = sum(vlist) / n
        max_val = max(vlist)
        min_val = min(vlist)
        pct_80 = sum(1 for v in vlist if v >= 0.8) / n * 100
        pct_60 = sum(1 for v in vlist if v >= 0.6) / n * 100
        stats[name] = {
            'avg': avg, 'max': max_val, 'min': min_val,
            'pct_80': pct_80, 'pct_60': pct_60
        }
    
    # 找到表8后面的段落
    body_elements = list(doc.element.body)
    detail_table = None
    for t in doc.tables:
        if len(t.rows[0].cells) == 11 and len(t.rows) > 50:
            detail_table = t
            break
    
    if detail_table is None:
        return False
    
    detail_elem = detail_table._element
    detail_idx = body_elements.index(detail_elem)
    
    # 收集表8后面的段落
    paragraphs = []
    for elem in body_elements[detail_idx + 1:]:
        if elem.tag.endswith('p'):
            from docx.text.paragraph import Paragraph
            para = Paragraph(elem, doc)
            text = para.text.strip()
            if text and ('课程目标' in text or '图5' in text):
                paragraphs.append(para)
    
    # 替换每段中的占位符
    replacements = [
        ('课程目标1的平均达成值为', f'课程目标1的平均达成值为{stats["v1"]["avg"]:.2f}'),
        ('课程目标2的平均达成值为', f'课程目标2的平均达成值为{stats["v2"]["avg"]:.2f}'),
        ('课程目标3的平均达成值为', f'课程目标3的平均达成值为{stats["v3"]["avg"]:.2f}'),
        ('课程目标4的平均达成值为', f'课程目标4的平均达成值为{stats["v4"]["avg"]:.2f}'),
        ('最高达成值为', '最高达成值为'),
        ('最低达成值为', '最低达成值为'),
        ('在0.8以上，', '在0.8以上，'),
        ('在0.6以上，', '在0.6以上，'),
    ]
    
    # 更新每个课程目标的段落
    target_names = ['v1', 'v2', 'v3', 'v4']
    para_idx = 0
    
    for i, target in enumerate(target_names):
        if para_idx >= len(paragraphs):
            break
        s = stats[target]
        para = paragraphs[para_idx]
        
        # 构建新的段落文本
        new_text = f"课程目标{i+1}的平均达成值为{s['avg']:.2f}，最高达成值为{s['max']:.2f}，最低达成值为{s['min']:.2f}。其中{s['pct_80']:.1f}%的学生达成值在0.8以上，{s['pct_60']:.1f}%的学生达成值在0.6以上，说明绝大部分学生通过课堂学习、作业练习、阶段测试和期末报告等方式，已掌握数据可视化的基础理论、基本方法。"
        
        # 清除原段落内容并写入新内容
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.text = new_text
        
        para_idx += 1
    
    return True

# ========== 专业名称提取 ==========
def extract_major(class_name: str) -> str:
    """增强版专业名称提取
    
    支持格式:
    - 23级计算机 → 计算机
    - 23级软工1班 → 软工
    - 23级计算机2班 → 计算机
    - 23级 软件工程 → 软件工程
    """
    class_name = str(class_name).strip()
    
    match = re.match(r'\d+级\s*(.+?)(?:\d+班)?$', class_name)
    if match:
        major = match.group(1).strip()
        if major.endswith('班'):
            major = major[:-1]
        return major if major else class_name
    
    result = re.sub(r'^\d+级|\d+班$', '', class_name).strip()
    return result if result else class_name

def extract_class_num(class_name: str) -> str:
    if not class_name or '班' not in class_name:
        return ""
    match = re.search(r'(\d班)', class_name)
    return match.group(1) if match else ""


# ========== 文本替换 ==========
def simple_replace(doc, replacements):
    """安全文本替换（处理表格内段落）:
    - 先收集所有段落（包括表格内的）
    - 遍历每个段落，检测并替换所有占位符（同一段落多占位符时不 break）
    - 执行替换
    """
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    
    matched = []
    for para in all_paras:
        if not para.text:
            continue
        full_text = ''.join(r.text for r in para.runs)
        for old, new in replacements:
            if old in full_text:
                full_text = full_text.replace(old, new)
                for r in para.runs:
                    r.text = ''
                if para.runs:
                    para.runs[0].text = full_text
                else:
                    para.text = full_text
                matched.append(old)
    return matched

# ========== 添加行（防御性） ==========
def add_row_to_table(table, source_row_idx=2):
    """添加新行（防御性合并单元格处理）
    
    安全处理合并单元格：
    - 复制源行样式
    - 清除所有合并属性（vMerge、gridSpan）
    - 插入到平均值行之前
    """
    source_tr = table.rows[source_row_idx]._element
    new_tr = copy.deepcopy(source_tr)
    
    # 清除所有tcPr中的合并属性（防御性处理）
    for tc in new_tr.iter(qn('w:tc')):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            for tag in ['vMerge', 'gridSpan']:
                elem = tcPr.find(qn(f'w:{tag}'))
                if elem is not None:
                    tcPr.remove(elem)
    
    # 插入到倒数第二行（平均值行之前）
    table._element.insert(len(table.rows) - 1, new_tr)

# ========== 明细表更新 ==========
def update_detail_table(table, students, weights, target_config):
    """更新明细表（表8）:
    - weights: 从模板表7动态读取的考核权重
    - target_config: 从模板表7动态读取的各目标理论满分及合计
    - 删除多余行（学生数<模板行数时）
    - 添加新行（学生数>模板行数时）← 使用防御性add_row_to_table
    - 填入学号、姓名、成绩、达成值
    - 更新平均值行
    """
    n = len(students)
    # 需要行数: 2(表头) + n(学生) + 1(平均)
    needed_rows = 2 + n + 1
    
    # 删除多余行
    if len(table.rows) > needed_rows:
        rows_to_delete = len(table.rows) - needed_rows
        for _ in range(rows_to_delete):
            delete_idx = len(table.rows) - 2  # 倒数第二行
            if delete_idx > 1:
                tr = table.rows[delete_idx]._element
                table._element.remove(tr)
    
    # 添加不够的行（使用防御性函数）
    if len(table.rows) < needed_rows:
        rows_to_add = needed_rows - len(table.rows)
        for _ in range(rows_to_add):
            add_row_to_table(table, source_row_idx=2)
    
    # 预计算所有成绩和达成度（分开存储）
    s1_list, s2_list, s3_list, s4_list = [], [], [], [],  # 得分
    v1_list, v2_list, v3_list, v4_list = [], [], [], [],  # 达成度
    
    for i in range(n):
        row_data = students.iloc[i]
        ps = row_data['平时成绩'] if pd.notna(row_data['平时成绩']) else 0
        qz = row_data['期中成绩'] if pd.notna(row_data['期中成绩']) else 0
        qm = row_data['期末成绩'] if pd.notna(row_data['期末成绩']) else 0
        
        # 计算各目标得分
        s1 = ps * weights['col3'] * target_config['target1']['col3']/100 + \
             qz * weights['col4'] * target_config['target1']['col4']/100 + \
             qm * weights['col5'] * target_config['target1']['col5']/100
        s2 = ps * weights['col3'] * target_config['target2']['col3']/100 + \
             qz * weights['col4'] * target_config['target2']['col4']/100 + \
             qm * weights['col5'] * target_config['target2']['col5']/100
        s3 = ps * weights['col3'] * target_config['target3']['col3']/100 + \
             qz * weights['col4'] * target_config['target3']['col4']/100 + \
             qm * weights['col5'] * target_config['target3']['col5']/100
        s4 = ps * weights['col3'] * target_config['target4']['col3']/100 + \
             qz * weights['col4'] * target_config['target4']['col4']/100 + \
             qm * weights['col5'] * target_config['target4']['col5']/100
        
        # 计算达成度
        v1 = s1 / target_config['target1']['total']
        v2 = s2 / target_config['target2']['total']
        v3 = s3 / target_config['target3']['total']
        v4 = s4 / target_config['target4']['total']
        
        # 分别存储得分和达成度
        s1_list.append(s1)
        s2_list.append(s2)
        s3_list.append(s3)
        s4_list.append(s4)
        v1_list.append(v1)
        v2_list.append(v2)
        v3_list.append(v3)
        v4_list.append(v4)
        
        # 填入行: 行2是第一个学生
        row_idx = i + 2
        cells = table.rows[row_idx].cells
        cells[0].text = str(i + 1)  # 序号
        cells[1].text = str(row_data['学号'])
        cells[2].text = str(row_data['姓名'])
        cells[3].text = f"{s1:.2f}"
        cells[4].text = f"{v1:.2f}"
        cells[5].text = f"{s2:.2f}"
        cells[6].text = f"{v2:.2f}"
        cells[7].text = f"{s3:.2f}"
        cells[8].text = f"{v3:.2f}"
        cells[9].text = f"{s4:.2f}"
        cells[10].text = f"{v4:.2f}"
    
    # 更新平均行（各列填自己的平均值）
    avg_row = table.rows[-1]
    avg_v1 = sum(v1_list) / n
    avg_v2 = sum(v2_list) / n
    avg_v3 = sum(v3_list) / n
    avg_v4 = sum(v4_list) / n
    
    # 计算各列平均得分（注意：是s_list不是v_list）
    avg_s1 = sum(s1_list) / n
    avg_s2 = sum(s2_list) / n
    avg_s3 = sum(s3_list) / n
    avg_s4 = sum(s4_list) / n
    
    avg_cells = avg_row.cells
    avg_cells[0].text = '平均值'
    avg_cells[1].text = '平均值'
    avg_cells[2].text = '平均值'
    avg_cells[3].text = f"{avg_s1:.2f}"  # 目标1得分平均
    avg_cells[4].text = f"{avg_v1:.2f}"   # 目标1达成值平均
    avg_cells[5].text = f"{avg_s2:.2f}"  # 目标2得分平均
    avg_cells[6].text = f"{avg_v2:.2f}"   # 目标2达成值平均
    avg_cells[7].text = f"{avg_s3:.2f}"  # 目标3得分平均
    avg_cells[8].text = f"{avg_v3:.2f}"   # 目标3达成值平均
    avg_cells[9].text = f"{avg_s4:.2f}"  # 目标4得分平均
    avg_cells[10].text = f"{avg_v4:.2f}" # 目标4达成值平均
    
    # 居中
    for j in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10]:
        set_cell_centered(avg_cells[j])
    
    # 计算全班各考核项平均分
    ps_avg = students['平时成绩'].mean()
    qz_avg = students['期中成绩'].mean()
    qm_avg = students['期末成绩'].mean()
    
    return avg_v1, avg_v2, avg_v3, avg_v4, v1_list, v2_list, v3_list, v4_list, s1_list, s2_list, s3_list, s4_list, ps_avg, qz_avg, qm_avg

# ========== 汇总表更新 ==========
def update_summary_table(doc, avg_v1, avg_v2, avg_v3, avg_v4, ps_avg, qz_avg, qm_avg, target_config, weights):
    """更新表7：
    - 平均达成值列（col7）
    - 实际得分行（行5,8,11,14的col3-6）
    
    - target_config: 从模板表7动态读取的各目标理论满分及合计
    实际得分公式：全班该项平均分/100 × 该项满分
    """
    table7 = doc.tables[7]
    avg_col = 7  # 平均达成值列（第8列）
    
    # 填入平均达成值
    table7.rows[SUMMARY_TARGET_ROWS['课程目标1']].cells[avg_col].text = f"{avg_v1:.2f}"
    table7.rows[SUMMARY_TARGET_ROWS['课程目标2']].cells[avg_col].text = f"{avg_v2:.2f}"
    table7.rows[SUMMARY_TARGET_ROWS['课程目标3']].cells[avg_col].text = f"{avg_v3:.2f}"
    table7.rows[SUMMARY_TARGET_ROWS['课程目标4']].cells[avg_col].text = f"{avg_v4:.2f}"
    
    # 填入理论满分行的合计列（col6，即第7列）
    # target_config 的 total 由 col3×权重.col3 + col4×权重.col4 + col5×权重.col5 计算得出
    theory_max_rows = {
        'target1': 4,   # 理论满分行
        'target2': 7,
        'target3': 10,
        'target4': 13,
    }
    for target_key, row_idx in theory_max_rows.items():
        total = target_config[target_key]['total']
        # 整数显示为 int，浮点数保留合理小数
        total_str = str(int(total)) if total == int(total) else f"{total:.2f}"
        table7.rows[row_idx].cells[6].text = total_str
    
    # 填入实际得分行（行5,8,11,14）
    # 公式：实际得分 = 全班该项平均分/100 × 该项满分
    actual_score_rows = {
        'target1': 5,   # 实际得分行
        'target2': 8,
        'target3': 11,
        'target4': 14,
    }
    
    for target_key, row_idx in actual_score_rows.items():
        cfg = target_config[target_key]
        # 实际得分 = 平均分/100 × 满分
        actual_col3 = ps_avg / 100 * cfg['col3']
        actual_col4 = qz_avg / 100 * cfg['col4']
        actual_col5 = qm_avg / 100 * cfg['col5']
        actual_col6 = actual_col3 * weights['col3'] + actual_col4 * weights['col4'] + actual_col5 * weights['col5']
        
        row = table7.rows[row_idx]
        row.cells[3].text = f"{actual_col3:.2f}"
        row.cells[4].text = f"{actual_col4:.2f}"
        row.cells[5].text = f"{actual_col5:.2f}"
        row.cells[6].text = f"{actual_col6:.2f}"
    
    return True

# ========== 主函数 ==========
def generate_achievement_analysis(excel_file: str, template_dir: str = None, interactive: bool = True) -> Dict[str, Any]:
    """生成课程目标达成情况分析表
    
    Args:
        excel_file: Excel文件路径，格式：{课程名称}-{班级}.xls
        template_dir: 模板目录，默认使用 QZTC/教学文档模版
    
    Returns:
        {
            'success': bool,
            'course': str,
            'students': int,
            'output_file': str,
            'stats': {...},
            'warnings': [str],
            'errors': [str]
        }
    """
    global TEMPLATE_DIR
    warnings = []
    errors = []
    
    # ===== 首次使用：检测 config.env =====
    config = config_loader.load_config()
    if not config.get('TEMPLATE_DIR'):
        if interactive:
            print("\n⚠️ 首次使用或配置不完整，请设置模板目录路径")
            print("   模板目录应包含 *-模版.docx 文件\n")
            while True:
                user_path = input("请输入模板目录路径（或按 Ctrl+C 退出）: ").strip()
                if not user_path:
                    print("   路径不能为空，请重新输入\n")
                    continue
                if not os.path.isdir(user_path):
                    print(f"   ❌ 路径不存在: {user_path}")
                    print("   请确认路径后重新输入\n")
                    continue
                template_files = [f for f in os.listdir(user_path) if '模版.docx' in f]
                if not template_files:
                    print(f"   ⚠️ 路径下未找到 *-模版.docx 文件")
                    continue
                with open(config_loader.CONFIG_FILE, 'w', encoding='utf-8') as f:
                    f.write("# 课程目标达成情况分析表 - 配置文件\n")
                    f.write(f"TEMPLATE_DIR={user_path}\n")
                print(f"   ✅ 已保存配置到: {config_loader.CONFIG_FILE}")
                break
            config = config_loader.load_config()
        else:
            return {
                'success': False,
                'errors': ['config.env 未配置且未提供 template_dir 参数'],
                'warnings': [],
                'course': '',
                'students': 0,
                'output_file': '',
                'stats': {}
            }
    
    # 获取模板目录
    template_dir = config_loader.get_template_dir()
    
    # ===== 预检验 =====
    is_valid, msg = validate_inputs(excel_file, template_dir)
    if not is_valid:
        return {
            'success': False,
            'errors': [msg],
            'warnings': [],
            'course': '',
            'students': 0,
            'output_file': '',
            'stats': {}
        }
    elif msg:
        warnings.append(msg)  # 模板名称提示
    
    if template_dir:
        TEMPLATE_DIR = template_dir
    
    basename = os.path.basename(excel_file)
    name_part = basename.replace('.xls', '')
    
    # 解析课程和班级
    parts = name_part.rsplit('-', 1)
    course_name = parts[0]
    class_name = parts[1] if len(parts) > 1 else ''
    major = extract_major(class_name)
    
    # 查找模板
    template_file = os.path.join(TEMPLATE_DIR, f'课程目标达成情况分析表-{course_name}-模版.docx')
    if not os.path.exists(template_file):
        # 尝试模糊匹配
        for f in os.listdir(TEMPLATE_DIR):
            if f.startswith('课程目标达成情况分析表-') and f.endswith('-模版.docx'):
                template_file = os.path.join(TEMPLATE_DIR, f)
                warnings.append(f"使用模糊匹配模板: {f}")
                break
        else:
            return {
                'success': False,
                'errors': [f'模板不存在: {template_file}'],
                'warnings': warnings
            }
    
    # 读取Excel
    df = pd.read_excel(excel_file)
    students = df[['学号', '姓名', '平时成绩', '期中成绩', '期末成绩']].copy()
    n = len(students)
    
    # 复制模板到输出文件
    output_dir = os.path.dirname(excel_file)
    output_file = os.path.join(output_dir, f'课程目标达成情况分析表-{course_name}-{class_name}.docx')
    shutil.copy(template_file, output_file)
    
    # 打开并修改
    doc = Document(output_file)
    
    # 文本替换
    now = datetime.now()
    simple_replace(doc, [
        ('{ay}', '2025 - 2026'),
        ('{sm}', '二'),
        ('{c}', extract_class_num(class_name)),
        ('{g}', class_name[:2] if class_name else '23'),
        ('{mj}', major),            # 专业（模板用{mj}）
                ('{tot}', f'{n}人'),
        ('{y}', str(now.year)),
        ('{m}', str(now.month)),
        ('{d}', str(now.day)),
    ])
    
    # 找到明细表（表8）- 11列且行数>50
    detail_table = None
    for table in doc.tables:
        cols = len(table.rows[0].cells)
        rows = len(table.rows)
        if cols == 11 and rows >= 50:
            detail_table = table
            break
    
    if detail_table is None:
        return {
            'success': False,
            'errors': ['未找到明细表']
        }
    
    # 从模板表7读取权重和目标配置
    weights, target_config = read_weights_and_config_from_template(doc)

    # 更新明细表
    avg_v1, avg_v2, avg_v3, avg_v4, v1_list, v2_list, v3_list, v4_list, _, _, _, _, ps_avg, qz_avg, qm_avg = \
        update_detail_table(detail_table, students, weights, target_config)
    
    # 更新汇总表
    update_summary_table(doc, avg_v1, avg_v2, avg_v3, avg_v4, ps_avg, qz_avg, qm_avg, target_config, weights)
    
    # 更新分析段落
    update_analysis_paragraphs(doc, v1_list, v2_list, v3_list, v4_list)
    
    # 保存
    doc.save(output_file)
    
    # 更新图表（chart1-4学生达成度 + chart5平均达成度）- 错误隔离
    chart_ok, chart_msg = update_chart(output_file, v1_list, v2_list, v3_list, v4_list, avg_v1, avg_v2, avg_v3, avg_v4)
    if not chart_ok:
        warnings.append(f"图表更新: {chart_msg}")
    elif chart_msg:
        warnings.append(chart_msg)
    
    # 计算≥0.8比例（用于返回统计）
    pct_v1_80 = sum(1 for v in v1_list if v >= 0.8) / n * 100
    pct_v2_80 = sum(1 for v in v2_list if v >= 0.8) / n * 100
    pct_v3_80 = sum(1 for v in v3_list if v >= 0.8) / n * 100
    pct_v4_80 = sum(1 for v in v4_list if v >= 0.8) / n * 100
    
    return {
        'success': True,
        'course': course_name,
        'students': n,
        'output_file': output_file,
        'stats': {
            'avg_v1': avg_v1, 'avg_v2': avg_v2, 'avg_v3': avg_v3, 'avg_v4': avg_v4,
            'pct_v1_80': pct_v1_80, 'pct_v2_80': pct_v2_80,
            'pct_v3_80': pct_v3_80, 'pct_v4_80': pct_v4_80,
        },
        'warnings': warnings,
        'errors': errors
    }

# ========== 批量处理 ==========
def batch_generate(directory: str, template_dir: str = None) -> List[Dict[str, Any]]:
    """批量处理目录下所有Excel文件"""
    results = []
    for fname in os.listdir(directory):
        if fname.endswith('.xls') or fname.endswith('.xlsx'):
            fpath = os.path.join(directory, fname)
            print(f"\n>>> 处理: {fname}")
            result = generate_achievement_analysis(fpath, template_dir)
            if result['success']:
                print(f"✅ 成功！学生数: {result['students']}")
            else:
                print(f"❌ 失败: {result.get('errors', [])}")
            results.append(result)
    return results

# ========== CLI ==========
def run_pipeline(excel_file: str, template_dir: str = None):
    """Pipeline模式: 生成 + 审查"""
    print("=" * 60)
    print("🔄 Pipeline: 生成 + 审查")
    print("=" * 60)
    
    # Step 1: 生成
    print("\n📝 STEP 1: 生成文档")
    print("-" * 40)
    result = generate_achievement_analysis(excel_file, template_dir)
    
    if not result['success']:
        print(f"❌ 生成失败: {result.get('errors', [])}")
        return result
    
    print(f"✅ 生成成功！")
    print(f"   课程: {result['course']}")
    print(f"   学生数: {result['students']}")
    print(f"   平均达成度: v1={result['stats']['avg_v1']:.2f}, v2={result['stats']['avg_v2']:.2f}, "
          f"v3={result['stats']['avg_v3']:.2f}, v4={result['stats']['avg_v4']:.2f}")
    
    # 显示生成警告
    if result.get('warnings'):
        print("\n⚠️ 生成警告:")
        for w in result['warnings']:
            print(f"   • {w}")
    
    # Step 2: 审查
    print("\n🔍 STEP 2: 审查文档")
    print("-" * 40)
    
    # 导入审查模块
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from scripts.review_document import review_document, print_report
    
    docx_file = result['output_file']
    review_result = review_document(docx_file)
    print_report(review_result)
    
    print("\n" + "=" * 60)
    if review_result['status'] == 'PASS':
        print("✅ Pipeline 完成！文档通过所有检查")
    elif review_result['status'] == 'WARN':
        print("⚠️ Pipeline 完成！有警告，请检查")
    else:
        print("❌ Pipeline 完成！有检查项未通过")
    print("=" * 60)
    
    return result

def batch_review(directory: str, template_dir: str = None) -> List[Dict[str, Any]]:
    """批量生成 + 审查
    
    Returns:
        List of results with 'generate' and 'review' keys
    """
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from scripts.review_document import review_document
    
    results = []
    files = [f for f in os.listdir(directory) if f.endswith('.xls') or f.endswith('.xlsx')]
    total = len(files)
    
    print("=" * 60)
    print(f"🔄 批量生成 + 审查 ({total} 个文件)")
    print("=" * 60)
    
    for i, fname in enumerate(sorted(files), 1):
        fpath = os.path.join(directory, fname)
        print(f"\n[{i}/{total}] 处理: {fname}")
        print("-" * 40)
        
        # 生成
        gen_result = generate_achievement_analysis(fpath, template_dir)
        
        if gen_result['success']:
            # 审查
            docx_file = gen_result['output_file']
            rev_result = review_document(docx_file)
            status = rev_result['status']
        else:
            rev_result = {'status': 'FAIL', 'errors': gen_result.get('errors', [])}
        
        results.append({
            'file': fname,
            'generate': gen_result,
            'review': rev_result
        })
        
        # 显示结果
        status_icon = '✅' if rev_result['status'] == 'PASS' else '⚠️' if rev_result['status'] == 'WARN' else '❌'
        print(f"   {status_icon} {rev_result['status']}")
    
    # 汇总
    passed = sum(1 for r in results if r['review']['status'] == 'PASS')
    warned = sum(1 for r in results if r['review']['status'] == 'WARN')
    failed = sum(1 for r in results if r['review']['status'] == 'FAIL')
    
    print("\n" + "=" * 60)
    print("📊 批量审查汇总")
    print("=" * 60)
    print(f"总计: {total} 个文件")
    print(f"  ✅ PASS: {passed}")
    print(f"  ⚠️ WARN: {warned}")
    print(f"  ❌ FAIL: {failed}")
    print("=" * 60)
    
    return results

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("""用法: python3 achievement_generator.py <excel文件或目录> [选项]

选项:
  <excel文件>              Excel文件路径
  --pipeline               Pipeline模式（生成+审查）
  --batch                  批量生成目录下所有Excel（不审查）
  --batch-review            批量生成+审查
  [模板目录]               可选，指定模板目录

示例:
  # 单文件生成
  python3 achievement_generator.py "C#程序设计-23级软工.xls"
  
  # Pipeline模式（生成+审查）
  python3 achievement_generator.py "C#程序设计-23级软工.xls" --pipeline
  
  # 批量审查
  python3 achievement_generator.py --batch-review /path/to/excels/
""")
        sys.exit(1)
    
    # 检查选项
    if '--pipeline' in sys.argv:
        sys.argv.remove('--pipeline')
        excel_file = sys.argv[1] if len(sys.argv) > 1 else None
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        if excel_file and not excel_file.startswith('-'):
            run_pipeline(excel_file, template_dir)
        else:
            print("❌ 请指定Excel文件")
            sys.exit(1)
    elif '--batch-review' in sys.argv:
        sys.argv.remove('--batch-review')
        directory = sys.argv[1] if len(sys.argv) > 1 else '.'
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        batch_review(directory, template_dir)
    elif '--batch' in sys.argv:
        sys.argv.remove('--batch')
        directory = sys.argv[1] if len(sys.argv) > 1 else '.'
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        batch_generate(directory, template_dir)
    elif '*' in sys.argv[1]:
        # 批量处理
        import glob
        excel_file = sys.argv[1]
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        for f in glob.glob(excel_file):
            print(f"\n>>> 处理: {f}")
            result = generate_achievement_analysis(f, template_dir)
            if result['success']:
                print(f"✅ 成功！学生数: {result['students']}, 输出: {os.path.basename(result['output_file'])}")
                if result.get('warnings'):
                    for w in result['warnings']:
                        print(f"   ⚠️ {w}")
            else:
                print(f"❌ 失败: {result.get('errors', [])}")
    else:
        excel_file = sys.argv[1]
        template_dir = sys.argv[2] if len(sys.argv) > 2 else None
        
        result = generate_achievement_analysis(excel_file, template_dir)
        if result['success']:
            print(f"\n✅ 成功！")
            print(f"   课程: {result['course']}")
            print(f"   学生数: {result['students']}")
            print(f"   输出: {result['output_file']}")
            print(f"   平均达成度: v1={result['stats']['avg_v1']:.2f}, v2={result['stats']['avg_v2']:.2f}, "
                  f"v3={result['stats']['avg_v3']:.2f}, v4={result['stats']['avg_v4']:.2f}")
            if result.get('warnings'):
                print("\n⚠️ 警告:")
                for w in result['warnings']:
                    print(f"   • {w}")
        else:
            print(f"\n❌ 失败: {result.get('errors', [])}")

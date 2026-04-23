#!/usr/bin/env python3
"""
课程目标达成情况分析表 - 质量审查脚本 v2.0
**Reviewer Pattern**: 按清单逐项检查生成文档

使用方式:
    python3 scripts/review_document.py 输出文件.docx [--verbose]
"""

import sys
import os
import re
import json
from docx import Document
from typing import Dict, List, Any
import zipfile
from lxml import etree

def review_document(docx_file: str, verbose: bool = False) -> Dict[str, Any]:
    """审查生成的文档
    
    Returns:
        {
            'file': str,
            'status': 'PASS' | 'WARN' | 'FAIL',
            'checks': {...},
            'stats': {...},
            'warnings': [...],
            'errors': [...]
        }
    """
    result = {
        'file': docx_file,
        'status': 'PASS',
        'checks': {},
        'stats': {},
        'warnings': [],
        'errors': []
    }
    
    if not os.path.exists(docx_file):
        result['status'] = 'FAIL'
        result['errors'].append(f'文件不存在: {docx_file}')
        return result
    
    try:
        doc = Document(docx_file)
    except Exception as e:
        result['status'] = 'FAIL'
        result['errors'].append(f'无法打开文档: {e}')
        return result
    
    # 1. 检查文本替换
    result['checks']['text_replacement'] = check_text_replacement(doc)
    
    # 2. 检查汇总表（表7）
    summary_ok, summary_stats = check_summary_table(doc)
    result['checks']['summary_table'] = summary_ok
    result['stats'].update(summary_stats)
    
    # 3. 检查明细表（表8）
    detail_ok, detail_stats = check_detail_table(doc)
    result['checks']['detail_table'] = detail_ok
    result['stats'].update(detail_stats)
    
    # 4. 检查图表数据
    chart_ok, chart_stats = check_charts(docx_file, result['stats'])
    result['checks']['charts'] = chart_ok
    result['stats'].update(chart_stats)
    
    # 5. 检查分析段落
    para_ok, para_stats = check_analysis_paragraphs(doc)
    result['checks']['analysis_paragraphs'] = para_ok
    result['stats'].update(para_stats)
    
    # 6. 综合判断状态
    if not all(result['checks'].values()):
        result['status'] = 'FAIL'
    elif result['warnings']:
        result['status'] = 'WARN'
    
    return result

def check_text_replacement(doc) -> bool:
    """检查关键占位符是否已替换"""
    placeholders = ['{ay}', '{sm}', '{g}', '{mj}', '{tot}', '{y}']
    found_unreplaced = []
    
    for para in doc.paragraphs:
        text = para.text
        for ph in placeholders:
            if ph in text:
                found_unreplaced.append(ph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    for ph in placeholders:
                        if ph in text:
                            found_unreplaced.append(ph)
    
    return len(found_unreplaced) == 0

def check_summary_table(doc) -> tuple:
    """检查汇总表（表7）"""
    if len(doc.tables) < 8:
        return False, {}
    
    table7 = doc.tables[7]
    stats = {}
    
    # 检查目标行：3, 6, 9, 12
    target_rows = {3: 'v1', 6: 'v2', 9: 'v3', 12: 'v4'}
    avg_col = 7  # 平均达成值列
    
    for row_idx, key in target_rows.items():
        if row_idx < len(table7.rows):
            cells = table7.rows[row_idx].cells
            if avg_col < len(cells):
                try:
                    val = float(cells[avg_col].text.strip())
                    stats[f'avg_{key}'] = val
                except ValueError:
                    stats[f'avg_{key}'] = None
    
    # 检查实际得分行：5, 8, 11, 14
    actual_rows = {5: 'v1', 8: 'v2', 11: 'v3', 14: 'v4'}
    for row_idx, key in actual_rows.items():
        if row_idx < len(table7.rows):
            cells = table7.rows[row_idx].cells
            try:
                # 检查col3-6是否有值
                vals = [float(cells[i].text.strip()) for i in range(3, 7) if cells[i].text.strip()]
                stats[f'actual_{key}'] = len(vals) == 4
            except ValueError:
                stats[f'actual_{key}'] = False
    
    return len(stats) >= 4, stats

def check_detail_table(doc) -> tuple:
    """检查明细表（表8）"""
    detail_table = None
    for table in doc.tables:
        cols = len(table.rows[0].cells)
        rows = len(table.rows)
        if cols == 11 and rows >= 50:
            detail_table = table
            break
    
    if detail_table is None:
        return False, {}
    
    # 统计学生数
    student_count = 0
    for i in range(2, len(detail_table.rows) - 1):
        cells = detail_table.rows[i].cells
        if cells[1].text.strip():  # 学号列有值
            student_count += 1
    
    # 检查平均行
    avg_row = detail_table.rows[-1]
    avg_cells = [c.text.strip() for c in avg_row.cells]
    avg_correct = avg_cells[0] == '平均值' and avg_cells[1] == '平均值'
    
    # 验证平均行数值（得分和达成值应该不同）
    try:
        score_avg = float(avg_cells[3])  # 目标1得分平均
        rate_avg = float(avg_cells[4])   # 目标1达成值平均
        avg_values_correct = 0 < score_avg < 30 and 0 < rate_avg < 1.5
    except (ValueError, IndexError):
        avg_values_correct = False
    
    return avg_correct, {
        'students': student_count,
        'avg_row_correct': avg_correct,
        'avg_values_correct': avg_values_correct
    }

def check_charts(docx_file: str, stats: Dict) -> tuple:
    """检查图表数据"""
    try:
        with zipfile.ZipFile(docx_file, 'r') as z:
            # 检查chart5
            chart5 = z.read('word/charts/chart5.xml')
            root = etree.fromstring(chart5)
            ns = {'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart'}
            nc = root.findall('.//c:numCache', ns)[0]
            pts = nc.findall('c:pt', ns)
            vals = [float(p.find('c:v', ns).text) for p in pts]
            
            # chart5应该是平均达成度百分比（70-85范围合理）
            chart5_ok = all(60 < v < 100 for v in vals)
            
            return chart5_ok, {
                'chart5_values': vals,
                'chart5_ok': chart5_ok
            }
    except Exception as e:
        return False, {'chart_error': str(e)}

def check_analysis_paragraphs(doc) -> tuple:
    """检查分析段落"""
    body_elements = list(doc.element.body)
    
    # 找明细表
    detail_table = None
    for t in doc.tables:
        if len(t.rows[0].cells) == 11 and len(t.rows) >= 50:
            detail_table = t
            break
    
    if detail_table is None:
        return False, {}
    
    detail_elem = detail_table._element
    detail_idx = body_elements.index(detail_elem)
    
    # 收集分析段落
    paragraphs = []
    for elem in body_elements[detail_idx + 1:]:
        if elem.tag.endswith('p'):
            from docx.text.paragraph import Paragraph
            para = Paragraph(elem, doc)
            text = para.text.strip()
            if text and '课程目标' in text and ('平均达成' in text or '平均为' in text):
                paragraphs.append(text)
    
    # 检查是否有4段分析
    if len(paragraphs) >= 4:
        # 检查数值是否在合理范围
        # 达成度应该是0.x或1.x格式，不是大的百分比数字
        values_ok = True
        for text in paragraphs[:4]:
            # 匹配达成度格式: 0.79, 0.89 等 (0-1范围)
            matches = re.findall(r'(?<![%\d])(\d\.\d{2})(?![%\d])', text)
            for m in matches:
                val = float(m)
                if val > 1.0:  # 达成度最大值是1.0
                    values_ok = False
        return values_ok, {'paragraphs_count': len(paragraphs)}
    
    return False, {'paragraphs_count': len(paragraphs)}

def print_report(result: Dict[str, Any], verbose: bool = False):
    """打印审查报告"""
    print("=" * 60)
    print("📋 课程目标达成情况分析表 - 质量审查报告 v2.0")
    print("=" * 60)
    print(f"文件: {os.path.basename(result['file'])}")
    
    status_icon = '✅ PASS' if result['status'] == 'PASS' else '⚠️ WARN' if result['status'] == 'WARN' else '❌ FAIL'
    print(f"状态: {status_icon}")
    print()
    
    print("📊 检查项:")
    for check, passed in result['checks'].items():
        icon = '✅' if passed else '❌'
        check_name = {
            'text_replacement': '文本替换',
            'summary_table': '汇总表',
            'detail_table': '明细表',
            'charts': '图表数据',
            'analysis_paragraphs': '分析段落'
        }.get(check, check)
        print(f"  {icon} {check_name}")
    
    print()
    if result['stats']:
        print("📈 统计数据:")
        if 'students' in result['stats']:
            print(f"  学生数: {result['stats']['students']}")
        if 'avg_v1' in result['stats']:
            print(f"  平均达成度: v1={result['stats'].get('avg_v1', 'N/A'):.4f}, "
                  f"v2={result['stats'].get('avg_v2', 'N/A'):.4f}, "
                  f"v3={result['stats'].get('avg_v3', 'N/A'):.4f}, "
                  f"v4={result['stats'].get('avg_v4', 'N/A'):.4f}")
        if 'chart5_values' in result['stats']:
            vals = result['stats']['chart5_values']
            print(f"  Chart5: {[f'{v:.1f}%' for v in vals]}")
    
    if result['warnings']:
        print()
        print("⚠️ 警告:")
        for w in result['warnings']:
            print(f"  • {w}")
    
    if result['errors']:
        print()
        print("❌ 错误:")
        for e in result['errors']:
            print(f"  • {e}")
    
    print()
    print("=" * 60)
    return result['status'] == 'PASS'

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 review_document.py <docx文件> [--verbose]")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    verbose = '--verbose' in sys.argv or '-v' in sys.argv
    json_only = '--json' in sys.argv
    
    result = review_document(docx_file, verbose)
    success = print_report(result, verbose)
    
    if json_only or '--json' in sys.argv:
        print(json.dumps(result, indent=2, ensure_ascii=False))
    
    sys.exit(0 if success else 1)

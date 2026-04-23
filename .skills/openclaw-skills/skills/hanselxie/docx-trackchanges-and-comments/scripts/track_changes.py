#!/usr/bin/env python3
"""
Word Track Changes (修订模式) 实现
通过直接操作OOXML来实现真正的修订模式
"""

import sys
import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import RGBColor


def create_ins_element(text, author="Unknown"):
    """创建插入修订元素"""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), '2026-03-17T12:00:00Z')
    
    r = OxmlElement('w:r')
    r.set(qn('w:rsidR'), '00AA00AA')
    r.set(qn('w:rsidRPr'), '00AA00AA')
    
    rPr = OxmlElement('w:rPr')
    
    # 绿色字体
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '00FF00')
    rPr.append(color)
    
    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    
    ins.append(r)
    return ins


def create_del_element(text, author="Unknown"):
    """创建删除修订元素"""
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), '2026-03-17T12:00:00Z')
    
    r = OxmlElement('w:r')
    r.set(qn('w:rsidR'), '00AA00AA')
    r.set(qn('w:rsidRDel'), '00AA00AA')
    
    rPr = OxmlElement('w:rPr')
    
    # 红色字体
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    rPr.append(color)
    
    # 删除线
    strike = OxmlElement('w:strike')
    strike.set(qn('w:val'), 'true')
    rPr.append(strike)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    
    del_elem.append(r)
    return del_elem


def enable_track_changes(doc):
    """启用文档的修订跟踪"""
    # 获取或创建 settings.xml
    if doc.part.related_parts:
        for part in doc.part.related_parts.values():
            if 'settings' in part.partname:
                settings = part.element
                # 查找或创建 w:trackChanges
                track_changes = settings.find(qn('w:trackChanges'))
                if track_changes is None:
                    track_changes = OxmlElement('w:trackChanges')
                    settings.append(track_changes)
                return
    
    # 如果找不到settings，手动添加
    print("警告: 无法自动启用修订跟踪，请在Word中手动开启")


def add_revision(doc, old_text, new_text, author="User"):
    """
    添加修订：用新文本替换旧文本
    old_text: 要删除的原文
    new_text: 要插入的新文本
    """
    found = False
    
    for para in doc.paragraphs:
        if old_text in para.text:
            # 找到匹配的段落
            para_xml = para._element
            
            # 查找并替换run
            for run in para.runs:
                if old_text in run.text:
                    # 创建删除元素
                    del_elem = create_del_element(old_text, author)
                    # 创建插入元素
                    ins_elem = create_ins_element(new_text, author)
                    
                    # 替换XML元素
                    run._element.addnext(del_elem)
                    del_elem.addnext(ins_elem)
                    
                    # 清空原始run
                    run.text = ""
                    found = True
                    break
            break
    
    return found


def add_insertion(doc, text, author="User"):
    """添加插入修订"""
    # 在文档末尾添加
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.underline = True
    return para


def add_deletion(doc, text, author="User"):
    """添加删除修订"""
    # 在文档末尾添加
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.strike = True
    return para


def add_comment(doc, text, comment_text, author="User"):
    """添加批注"""
    # 查找文本位置
    for para in doc.paragraphs:
        if text in para.text:
            # 添加批注标记
            # 由于python-docx对批注支持有限，这里用简单方式
            para.add_run(f" [{comment_text}]")
            return True
    return False


def main():
    if len(sys.argv) < 3:
        print("用法: track_changes.py <输入文件> <输出文件> [--add '原文本' '新文本']")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # 复制文件
    shutil.copy(input_file, output_file)
    
    # 打开文档
    doc = Document(output_file)
    
    # 解析命令行参数
    i = 3
    while i < len(sys.argv):
        if sys.argv[i] == '--add' and i + 2 < len(sys.argv):
            old_text = sys.argv[i + 1]
            new_text = sys.argv[i + 2]
            author = sys.argv[i + 3] if i + 3 < len(sys.argv) and not sys.argv[i + 3].startswith('--') else "User"
            
            if add_revision(doc, old_text, new_text, author):
                print(f"已添加修订: '{old_text}' -> '{new_text}'")
            else:
                print(f"未找到文本: {old_text}")
            
            i += 4
        else:
            i += 1
    
    # 保存文档
    doc.save(output_file)
    print(f"已保存到: {output_file}")


if __name__ == '__main__':
    main()

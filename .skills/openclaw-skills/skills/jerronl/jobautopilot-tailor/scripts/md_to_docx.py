#!/usr/bin/env python3
"""
md_to_docx_v2.py — Fill a resume DOCX template from a Markdown file.

Dependencies:  pip install python-docx

Usage:
    python3 md_to_docx_v2.py <input.md> <template.docx> <output.docx>

Template placeholders expected:
    {{NAME}}, {{CONTACT}}, {{SUMMARY_BODY}}, {{SKILLS_BODY}}
    {{SEC1_TITLE}} … {{SEC5_TITLE}}
    {{JOBn_TITLE}}, {{JOBn_LOC}}, {{JOBn_DATE}}, {{JOBn_BULLETm}}
    {{EARLY_EXPn}}
    {{UNIVn}}, {{DEGREEn}}

Behaviour:
    • Fewer MD jobs than template slots  → removes extra job paragraphs
    • More MD jobs than template slots   → clones last job's formatting
    • Same logic for bullets, early-exp lines, education lines
"""

import re, sys, copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── MD parser ─────────────────────────────────────────────────────────────────

def parse_md(text: str) -> dict:
    lines = text.splitlines()
    non_blank = [l for l in lines if l.strip()]
    name    = non_blank[0].strip() if non_blank else ""
    contact = non_blank[1].strip() if len(non_blank) > 1 else ""

    SEC_RE = re.compile(r'^([A-Z][A-Z /]+[A-Z])$')
    sections: dict[str, list[str]] = {}
    cur = None
    for line in lines[2:]:
        m = SEC_RE.match(line.strip())
        if m:
            cur = m.group(1).strip()
            sections[cur] = []
        elif cur is not None:
            sections[cur].append(line)

    def sec(key):
        for k, v in sections.items():
            if k == key:
                return "\n".join(v).strip()
        return ""

    # Jobs
    JOB_RE = re.compile(r'^(.+?)\s*[—–-]+\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*$')
    jobs, cur_job = [], None
    for line in sections.get("EXPERIENCE", []):
        s = line.strip()
        if not s:
            continue
        m = JOB_RE.match(s)
        if m:
            if cur_job:
                jobs.append(cur_job)
            cur_job = dict(
                title=f"{m.group(1).strip()} \u2014 {m.group(2).strip()}",
                loc=m.group(3).strip(),
                date=m.group(4).strip(),
                bullets=[]
            )
        elif re.match(r'^[•\-]', s) and cur_job:
            cur_job["bullets"].append(re.sub(r'^[•\-–\s]+', '', s).strip())
    if cur_job:
        jobs.append(cur_job)

    # Earlier experience
    early_exps = []
    for k, v in sections.items():
        if "EARLIER" in k:
            early_exps = [l.strip() for l in v if l.strip()]
            break

    # Education
    EDU_RE = re.compile(r'^(.+?)\s*[—–-]+\s*(.+)$')
    education = []
    for line in sections.get("EDUCATION", []):
        s = line.strip()
        if not s:
            continue
        m = EDU_RE.match(s)
        education.append(
            dict(univ=m.group(1).strip(), degree=m.group(2).strip()) if m
            else dict(univ=s, degree="")
        )

    return dict(name=name, contact=contact,
                summary=sec("SUMMARY"), skills=sec("CORE SKILLS"),
                jobs=jobs, early_exps=early_exps, education=education)

# ── paragraph helpers ─────────────────────────────────────────────────────────

def para_text(para) -> str:
    return para.text

def has_placeholder(para, ph: str) -> bool:
    return ph in para_text(para)

def set_para_text_preserve_runs(para, new_text: str):
    """Replace the full text of a paragraph, keeping the first run's formatting."""
    # Collect all run XML elements and remove them
    p = para._p
    runs = p.findall(qn('w:r'))
    if not runs:
        para.add_run(new_text)
        return
    # Keep formatting from first run, set its text, remove the rest
    first_run = runs[0]
    # Set text in first run
    t_elem = first_run.find(qn('w:t'))
    if t_elem is None:
        t_elem = OxmlElement('w:t')
        first_run.append(t_elem)
    t_elem.text = new_text
    if new_text.startswith(' ') or new_text.endswith(' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    # Remove extra runs
    for r in runs[1:]:
        p.remove(r)

def fill_job_header(para, title: str, loc: str, date: str):
    """
    Job header paragraphs have 3 tab-separated fields in multiple runs.
    Pattern: [bold run: title TAB] [bold run: TAB] [bold run: loc TAB] [bold run: TAB] [bold run: date]
    We just replace the text content of each run by position.
    """
    p = para._p
    runs = p.findall(qn('w:r'))

    # Collect text runs (skip pure-tab runs)
    text_runs = []
    for r in runs:
        t = r.find(qn('w:t'))
        if t is not None and t.text and t.text.strip('\t'):
            text_runs.append((r, t))

    # We expect 3 text fields: title, loc, date
    fields = [title, loc, date]
    for i, (r, t) in enumerate(text_runs[:3]):
        t.text = fields[i] + (' ' if i == 0 else '')
        if t.text.startswith(' ') or t.text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def clone_paragraph_after(ref_para, new_text: str, bold: bool = False) -> object:
    """
    Insert a new paragraph immediately after ref_para, copying its
    pPr (style, spacing, indent, tabs) and optionally making it bold.
    Returns the new paragraph object.
    """
    # Deep-copy the reference paragraph element
    new_p = copy.deepcopy(ref_para._p)

    # Clear all runs from the copy
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)

    # Build a single run with the text
    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r_elem.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = new_text
    if new_text.startswith(' ') or new_text.endswith(' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    new_p.append(r_elem)

    # Insert after ref_para in the document body
    ref_para._p.addnext(new_p)

    # Wrap in a Paragraph proxy
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())

def clone_job_header_after(ref_para, title: str, loc: str, date: str) -> object:
    """Clone a job-header paragraph (3-field tab layout) after ref_para."""
    new_p = copy.deepcopy(ref_para._p)

    # Remove all runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)

    def bold_run(text, add_tab=False):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
        if add_tab:
            tab = OxmlElement('w:tab')
            r.append(tab)
        t = OxmlElement('w:t')
        t.text = text
        if text.startswith(' ') or text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        return r

    def tab_run():
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
        tab = OxmlElement('w:tab')
        r.append(tab)
        return r

    new_p.append(bold_run(title + ' '))
    new_p.append(tab_run())
    new_p.append(bold_run(loc))
    new_p.append(tab_run())
    new_p.append(bold_run(date))

    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, new_p.getparent())

def remove_paragraph(para):
    p = para._p
    p.getparent().remove(p)

# ── template index helpers ────────────────────────────────────────────────────

def find_para_with(paras, ph: str):
    """Return index of paragraph containing placeholder, or None."""
    for i, p in enumerate(paras):
        if ph in p.text:
            return i
    return None

def collect_job_indices(paras, job_n: int) -> dict:
    """
    Return dict with keys: 'header', 'bullets' (list of indices)
    for job number job_n.
    """
    header_idx = find_para_with(paras, f'{{{{JOB{job_n}_TITLE}}}}')
    if header_idx is None:
        return None
    bullets = []
    b = 1
    while True:
        idx = find_para_with(paras, f'{{{{JOB{job_n}_BULLET{b}}}}}')
        if idx is None:
            break
        bullets.append(idx)
        b += 1
    return dict(header=header_idx, bullets=bullets)

# ── main fill logic ───────────────────────────────────────────────────────────

def fill_template(doc: Document, data: dict):
    paras = doc.paragraphs  # live list (reflects structural changes)

    # ── Simple single-value substitutions ─────────────────────────────────────
    simple = {
        '{{NAME}}':         data['name'],
        '{{CONTACT}}':      data['contact'],
        '{{SUMMARY_BODY}}': data['summary'],
        '{{SKILLS_BODY}}':  data['skills'],
        '{{SEC1_TITLE}}':   'SUMMARY',
        '{{SEC2_TITLE}}':   'CORE SKILLS',
        '{{SEC3_TITLE}}':   'EXPERIENCE',
        '{{SEC4_TITLE}}':   'EARLIER EXPERIENCE',
        '{{SEC5_TITLE}}':   'EDUCATION',
    }
    # Collect all paragraphs: body paragraphs + table cell paragraphs
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for para in all_paras:
        for ph, val in simple.items():
            if ph in para.text:
                set_para_text_preserve_runs(para, val)
                break

    # ── Jobs ──────────────────────────────────────────────────────────────────
    tpl_jobs = 0
    while find_para_with(doc.paragraphs, f'{{{{JOB{tpl_jobs+1}_TITLE}}}}') is not None:
        tpl_jobs += 1

    md_jobs   = data['jobs']
    fill_jobs = min(tpl_jobs, len(md_jobs))

    # Step A: fill slots that exist in both
    for i in range(1, fill_jobs + 1):
        job = md_jobs[i - 1]
        info = collect_job_indices(doc.paragraphs, i)

        # Count template bullet slots
        tpl_b   = len(info['bullets'])
        md_b    = job['bullets']

        # Extra bullets (MD has more than template): clone last bullet para
        if len(md_b) > tpl_b and tpl_b > 0:
            last_bullet_para = doc.paragraphs[info['bullets'][-1]]
            # Insert in reverse so each one ends up in correct order
            for extra_text in reversed(md_b[tpl_b:]):
                clone_paragraph_after(last_bullet_para, extra_text)

        # Fill or remove bullet slots (re-fetch indices after any insertions)
        # Remove unused slots in REVERSE order to avoid index shifting,
        # then fill remaining slots.
        info = collect_job_indices(doc.paragraphs, i)
        md_b = job['bullets']

        # Remove excess template bullet slots (from the end, in reverse)
        while len(info['bullets']) > len(md_b):
            remove_paragraph(doc.paragraphs[info['bullets'][-1]])
            info = collect_job_indices(doc.paragraphs, i)

        # Fill remaining slots
        info = collect_job_indices(doc.paragraphs, i)
        for b_i, bullet_idx in enumerate(info['bullets']):
            set_para_text_preserve_runs(doc.paragraphs[bullet_idx], md_b[b_i])

        # Fill header (re-fetch after bullet operations)
        info = collect_job_indices(doc.paragraphs, i)
        fill_job_header(doc.paragraphs[info['header']], job['title'], job['loc'], job['date'])

    # Step B: remove extra template job slots (template has more than MD)
    for i in range(fill_jobs + 1, tpl_jobs + 1):
        while True:
            info = collect_job_indices(doc.paragraphs, i)
            if info is None:
                break
            # Remove bullets first, then header
            for bullet_idx in reversed(info['bullets']):
                remove_paragraph(doc.paragraphs[bullet_idx])
            remove_paragraph(doc.paragraphs[info['header']])
            break  # collect_job_indices handles one job at a time

    # Step C: extra MD jobs (MD has more than template)
    # NOTE: anchor must be captured BEFORE fill loop replaces placeholders
    _extra_anchor = None
    if len(md_jobs) > tpl_jobs and tpl_jobs > 0:
        info0 = collect_job_indices(doc.paragraphs, tpl_jobs)
        if info0:
            if info0['bullets']:
                _extra_anchor = doc.paragraphs[info0['bullets'][-1]]
            else:
                _extra_anchor = doc.paragraphs[info0['header']]

    # Step C continued: do the actual insertion now (after fill loop)
    if _extra_anchor is not None:
        anchor = _extra_anchor
        for job in reversed(md_jobs[tpl_jobs:]):
            cur_anchor = anchor
            for bullet_text in reversed(job['bullets']):
                cur_anchor = clone_paragraph_after(cur_anchor, bullet_text)
            clone_job_header_after(anchor, job['title'], job['loc'], job['date'])

    # ── Earlier experience ─────────────────────────────────────────────────────
    early = data['early_exps']
    tpl_early = 0
    while find_para_with(doc.paragraphs, f'{{{{EARLY_EXP{tpl_early+1}}}}}') is not None:
        tpl_early += 1
    fill_early = min(tpl_early, len(early))

    # Extra early lines: clone last early para
    if len(early) > tpl_early and tpl_early > 0:
        last_idx = find_para_with(doc.paragraphs, f'{{{{EARLY_EXP{tpl_early}}}}}')
        last_para = doc.paragraphs[last_idx]
        for extra in reversed(early[tpl_early:]):
            clone_paragraph_after(last_para, extra)

    for i in range(1, fill_early + 1):
        idx = find_para_with(doc.paragraphs, f'{{{{EARLY_EXP{i}}}}}')
        set_para_text_preserve_runs(doc.paragraphs[idx], early[i - 1])
    for i in range(fill_early + 1, tpl_early + 1):
        idx = find_para_with(doc.paragraphs, f'{{{{EARLY_EXP{i}}}}}')
        if idx is not None:
            remove_paragraph(doc.paragraphs[idx])

    # ── Education ──────────────────────────────────────────────────────────────
    edu = data['education']
    tpl_edu = 0
    while find_para_with(doc.paragraphs, f'{{{{UNIV{tpl_edu+1}}}}}') is not None:
        tpl_edu += 1
    fill_edu = min(tpl_edu, len(edu))

    # Extra edu lines: clone last edu para
    if len(edu) > tpl_edu and tpl_edu > 0:
        last_idx = find_para_with(doc.paragraphs, f'{{{{UNIV{tpl_edu}}}}}')
        last_para = doc.paragraphs[last_idx]
        for extra in reversed(edu[tpl_edu:]):
            text = f"{extra['univ']} \u2014 {extra['degree']}" if extra['degree'] else extra['univ']
            clone_paragraph_after(last_para, text)

    for i in range(1, fill_edu + 1):
        idx = find_para_with(doc.paragraphs, f'{{{{UNIV{i}}}}}')
        if idx is not None:
            e = edu[i - 1]
            text = f"{e['univ']} \u2014 {e['degree']}" if e['degree'] else e['univ']
            set_para_text_preserve_runs(doc.paragraphs[idx], text)
    for i in range(fill_edu + 1, tpl_edu + 1):
        idx = find_para_with(doc.paragraphs, f'{{{{UNIV{i}}}}}')
        if idx is not None:
            remove_paragraph(doc.paragraphs[idx])

# ── entry point ───────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 4:
        print("Usage: python3 md_to_docx_v2.py <input.md> <template.docx> <output.docx>")
        sys.exit(1)

    md_path, template_path, output_path = [Path(a) for a in sys.argv[1:4]]

    data = parse_md(md_path.read_text(encoding='utf-8'))
    print(f"Parsed: name='{data['name']}', {len(data['jobs'])} jobs, "
          f"{len(data['early_exps'])} early-exp lines, {len(data['education'])} edu lines")

    doc = Document(str(template_path))
    fill_template(doc, data)
    doc.save(str(output_path))
    print(f"Output: {output_path}")

if __name__ == '__main__':
    main()

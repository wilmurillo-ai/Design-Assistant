# -*- coding: utf-8 -*-
"""
DBCheck - Oracle 全面巡检工具（增强版）
=======================================
基于 OS 层 + 数据库层
支持: 10g / 11g / 12c / 18c / 19c / 21c / 23c
作者: Jack Ge
"""

import sys
import os

# frozen 模式下路径处理
if getattr(sys, 'frozen', False):
    sys.path.insert(0, sys._MEIPASS)

import warnings
warnings.filterwarnings("ignore", category=UserWarning, message="pkg_resources is deprecated")

from version import __version__ as VER
import time
import re
import json

import io
import argparse
import platform

try:
    import oracledb
    _HAS_ORACLE = True
except ImportError:
    _HAS_ORACLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# ── ANSI 颜色 ──────────────────────────────────────────────────────────────
CYAN   = "\033[96m"
GREEN  = "\033[92m"
YELLOW = "\033[93m"
MAGENTA= "\033[95m"
BOLD   = "\033[1m"
DIM    = "\033[2m"
RESET  = "\033[0m"
RED    = "\033[91m"

def _enable_ansi():
    try:
        import ctypes
        if os.name == "nt":
            ctypes.windll.kernel32.SetConsoleMode(
                ctypes.windll.kernel32.GetStdHandle(-11), 7)
    except Exception:
        pass

_enable_ansi()

# ── SSH 系统信息采集器（复用 MySQL 的 RemoteSystemInfoCollector）─────────────
try:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from main_mysql import RemoteSystemInfoCollector
    _HAS_SSH = True
except Exception:
    _HAS_SSH = False

# ── 巡检报告保存 ────────────────────────────────────────────────────────────
try:
    from save_doc_context import SaveDocContext
    _HAS_SAVE = True
except ImportError:
    _HAS_SAVE = False

# ═══════════════════════════════════════════════════════════════════════════
#                    巡检数据采集 — OS 层（SSH / 本地）
# ═══════════════════════════════════════════════════════════════════════════

def get_db_version_and_major(conn):
    """同时获取版本字符串和主版本号，返回 (version_str, ver_major)"""
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT version FROM v$instance")
            row = cur.fetchone()
            if row:
                version_str = row[0]
                m = re.match(r'(\d+)', version_str)
                ver_major = m.group(1) if m else ""
                return version_str, ver_major
    except Exception:
        pass
    return "", ""


def get_checks_for_version(ver_major):
    """根据主版本号返回对应的巡检函数列表

    10g → dbcheck10g.sql：WMSYS.WM_CONCAT 替代 listagg
    11g → dbcheck11g.sql：标准 listagg
    12c+ → dbcheck12c.sql：CDB/PDB 支持、gv$crs_resource_v2 等
    """
    # ── 10g ──────────────────────────────────────────────────────────────────
    if ver_major == "10":
        checks = [
            ("实例信息",      oracle_check_instance),
            ("数据库信息",    oracle_check_database_v10),
            ("版本/补丁",     oracle_check_version_and_patches),
            ("表空间",        oracle_check_tablespace),
            ("Redo日志",      oracle_check_redolog),
            ("控制文件",      oracle_check_controlfile),
            ("SGA/PGA内存",   oracle_check_sga_pga),
            ("关键参数",      oracle_check_params),
            ("Undo信息",      oracle_check_undo),
            ("长SQL",         oracle_check_long_sql),
            ("性能指标",      oracle_check_performance),
            ("Top SQL",       oracle_check_top_sql),
            ("无效对象",      oracle_check_invalid_objects),
            ("用户安全",      oracle_check_users),
            # 备份/DataGuard/RAC/AWR 在 10g 通常不可用，已在对应函数内做版本适配
            ("备份信息",      oracle_check_backup),
            ("闪回/回收站",   oracle_check_flashback),
            ("Data Guard",    oracle_check_dataguard),
            ("RAC+ASM",       oracle_check_rac),
            # 10g 无 AWR，用 statspack 替代（函数内部已处理）
            ("AWR快照",       oracle_check_awr),
            ("作业调度",      oracle_check_jobs),
            ("Alert日志",     oracle_check_alert),
        ]
        return checks

    # ── 11g ──────────────────────────────────────────────────────────────────
    if ver_major == "11":
        checks = [
            ("实例信息",      oracle_check_instance),
            ("数据库信息",    oracle_check_database_v11),
            ("版本/补丁",     oracle_check_version_and_patches),
            ("表空间",        oracle_check_tablespace),
            ("Redo日志",      oracle_check_redolog),
            ("控制文件",      oracle_check_controlfile),
            ("SGA/PGA内存",   oracle_check_sga_pga),
            ("关键参数",      oracle_check_params),
            ("Undo信息",      oracle_check_undo),
            ("长SQL",         oracle_check_long_sql),
            ("性能指标",      oracle_check_performance),
            ("Top SQL",       oracle_check_top_sql),
            ("无效对象",      oracle_check_invalid_objects),
            ("用户安全",      oracle_check_users),
            ("备份信息",      oracle_check_backup),
            ("闪回/回收站",   oracle_check_flashback),
            ("Data Guard",    oracle_check_dataguard),
            ("RAC+ASM",       oracle_check_rac),
            ("AWR快照",       oracle_check_awr),
            ("作业调度",      oracle_check_jobs),
            ("Alert日志",     oracle_check_alert),
        ]
        return checks

    # ── 12c 及以上（基准）──────────────────────────────────────────────────────
    # 12c SQL 作为基准；19c 及以上出错的查询项，使用 v19 兼容版本覆盖
    checks = [
        ("实例信息",      oracle_check_instance),
        ("数据库信息",    oracle_check_database_v12plus),
        ("版本/补丁",     oracle_check_version_and_patches),
        ("表空间",        oracle_check_tablespace),
        ("Redo日志",      oracle_check_redolog),
        ("控制文件",      oracle_check_controlfile),
        ("SGA/PGA内存",   oracle_check_sga_pga),
        ("关键参数",      oracle_check_params),
        ("Undo信息",      oracle_check_undo),
        ("长SQL",         oracle_check_long_sql),
        ("性能指标",      oracle_check_performance),
        ("Top SQL",       oracle_check_top_sql),
        ("无效对象",      oracle_check_invalid_objects),
        ("用户安全",      oracle_check_users),
        ("备份信息",      oracle_check_backup),
        ("闪回/回收站",   oracle_check_flashback),
        ("Data Guard",    oracle_check_dataguard),
        ("RAC+ASM",       oracle_check_rac),
        ("AWR快照",       oracle_check_awr),
        ("作业调度",      oracle_check_jobs),
        ("Alert日志",     oracle_check_alert),
    ]

    # 19c 及以上：对出错项应用 v19 兼容版（保留 12c 基准不变）
    if ver_major and int(ver_major) >= 19:
        for i, (name, _fn) in enumerate(checks):
            if name == "数据库信息":
                checks[i] = ("数据库信息", oracle_check_database_v19)
            elif name == "表空间":
                checks[i] = ("表空间", oracle_check_tablespace_v19)
            elif name == "Redo日志":
                checks[i] = ("Redo日志", oracle_check_redolog_v19)
            elif name == "Top SQL":
                checks[i] = ("Top SQL", oracle_check_top_sql_v19)
            elif name == "备份信息":
                checks[i] = ("备份信息", oracle_check_backup_v19)
            elif name == "闪回/回收站":
                checks[i] = ("闪回/回收站", oracle_check_flashback_v19)
            elif name == "Data Guard":
                checks[i] = ("Data Guard", oracle_check_dataguard_v19)
            elif name == "AWR快照":
                checks[i] = ("AWR快照", oracle_check_awr_v19)
            elif name == "Alert日志":
                checks[i] = ("Alert日志", oracle_check_alert_v19)

    return checks


class OSCollector:
    """OS 层信息采集（通过 SSH 或本地命令）"""

    def __init__(self, ssh_conn=None):
        self.ssh = ssh_conn  # paramiko SSHClient 或 None（本地）

    def run_cmd(self, cmd):
        """通过 SSH 或本地执行命令"""
        if self.ssh:
            try:
                stdin, stdout, stderr = self.ssh.exec_command(cmd)
                return stdout.read().decode('utf-8', errors='ignore')
            except Exception:
                return ""
        else:
            import subprocess
            try:
                return subprocess.check_output(cmd, shell=True, stderr=subprocess.DEVNULL, timeout=30).decode('utf-8', errors='ignore')
            except Exception:
                return ""

    def collect(self):
        """收集 OS 层全部信息"""
        uname = platform.system()
        is_linux = (uname == "Linux")

        data = {}

        # 基础信息
        data['hostname'] = platform.node()
        data['os_version'] = self.run_cmd("cat /etc/*release 2>/dev/null | head -1").strip()
        data['kernel'] = platform.release()
        data['uptime'] = self.run_cmd("uptime 2>/dev/null | head -n1").strip()

        # CPU
        data['cpu_model'] = self.run_cmd("awk -F': ' '/model name/ {print $2; exit}' /proc/cpuinfo 2>/dev/null").strip()
        data['cpu_count'] = self.run_cmd("nproc 2>/dev/null").strip() or self.run_cmd("grep -c processor /proc/cpuinfo 2>/dev/null").strip()
        cpu_idle = self.run_cmd("vmstat 1 2 | awk 'NR==4 {print 100-$15}'").strip()
        data['cpu_usage_pct'] = cpu_idle

        # 内存
        mem_total = self.run_cmd("free -m | awk '/Mem:/ {print $2}'").strip()
        mem_used   = self.run_cmd("free -m | awk '/Mem:/ {print $3}'").strip()
        data['mem_total_mb']  = mem_total
        data['mem_used_mb']   = mem_used
        data['mem_usage_pct'] = self.run_cmd("free -m | awk '/Mem:/ {print $3/$2*100}'").strip()

        # Swap
        swap_total = self.run_cmd("free -m | awk '/Swap:/ {print $2}'").strip()
        swap_used   = self.run_cmd("free -m | awk '/Swap:/ {print $3}'").strip()
        data['swap_total_mb'] = swap_total
        data['swap_used_mb']  = swap_used

        # 负载
        data['load_average'] = self.run_cmd("uptime | awk -F': ' '{print $2}'").strip()

        # 磁盘
        data['disk_usage'] = self.run_cmd("df -Ph 2>/dev/null | grep -v 'tmpfs\\|devtmpfs\\|overlay\\|shm'").strip()

        # /etc/hosts
        data['hosts'] = self.run_cmd("sed '1,2d' /etc/hosts 2>/dev/null | grep -v '^$'").strip()

        # sysctl 参数
        data['sysctl'] = self.run_cmd(
            "grep -E 'kernel.shmall|kernel.shmmax|kernel.sem|kernel.shmmni|fs.aio-max-nr|fs.file-max|vm.swappiness|vm.nr_hugepages' "
            "/etc/sysctl.conf 2>/dev/null"
        ).strip()

        # limits.conf
        data['limits'] = self.run_cmd("grep -v '^#\\|^$' /etc/security/limits.conf 2>/dev/null").strip()

        # HugePages
        data['hugepages'] = self.run_cmd(
            "awk '/MemTotal|HugePages_Total|HugePages_Free/ {print $1\":\"$2}' /proc/meminfo 2>/dev/null"
        ).strip()

        # Transparent HugePages
        thp = self.run_cmd("cat /sys/kernel/mm/transparent_hugepage/enabled 2>/dev/null").strip()
        data['thp'] = thp

        # crontab
        data['crontab'] = self.run_cmd("crontab -l 2>/dev/null").strip()

        # 网络
        data['network'] = self.run_cmd("ip addr show 2>/dev/null | grep 'inet '").strip()

        # /etc/passwd（数据库用户检查用）
        data['oracle_users'] = self.run_cmd(
            "grep -E '^(oracle|grid|root):' /etc/passwd 2>/dev/null"
        ).strip()

        return data


# ═══════════════════════════════════════════════════════════════════════════
#                    巡检数据采集 — Oracle 数据库层
# ═══════════════════════════════════════════════════════════════════════════

def oracle_check_instance(conn):
    """实例基本信息"""
    results = {}
    cur = conn.cursor()
    try:
        # 实例信息
        cur.execute("""
            SELECT INST_ID, INSTANCE_NAME, HOST_NAME, VERSION,
                   STARTUP_TIME, STATUS, PARALLEL, LOG_MODE,
                   DATABASE_ROLE, OPEN_MODE
            FROM gv$instance
        """)
        results['instance'] = cur.fetchall()
    except Exception as e:
        results['instance_error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_database(conn):
    """数据库基本信息"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE,
                   CDB, PLUGGABLE_NDB
            FROM v$database
        """)
        results['database'] = cur.fetchall()

        # 全局名
        cur.execute("SELECT global_name FROM global_name")
        results['global_name'] = cur.fetchone()

        # 字符集（sys.props$ 在 SYSDBA 下可能无权限，改为 nls_database_parameters）
        try:
            cur.execute("""
                SELECT parameter, value
                FROM nls_database_parameters
                WHERE parameter IN ('NLS_CHARACTERSET', 'NLS_NCHAR_CHARACTERSET')
            """)
            rows = cur.fetchall()
            results['charset'] = tuple(r[1] for r in rows) if rows else ('', '')
        except Exception as e:
            results['charset'] = ('', '')

        # 块大小
        cur.execute("SELECT value FROM v$parameter WHERE name='db_block_size'")
        r = cur.fetchone()
        results['block_size'] = r[0] if r else ''

        # SGA/PGA
        for param in ['sga_max_size', 'sga_target', 'pga_aggregate_target',
                      'memory_max_target', 'memory_target']:
            cur.execute(f"SELECT value FROM v$parameter WHERE name='{param}'")
            r = cur.fetchone()
            results[param] = r[0] if r else ''

        # SPFILE
        cur.execute("SELECT value FROM v$parameter WHERE name='spfile'")
        r = cur.fetchone()
        results['spfile'] = r[0] if r else ''

        # OMF
        cur.execute("SELECT value FROM v$parameter WHERE name='db_create_file_dest'")
        r = cur.fetchone()
        results['omf'] = r[0] if r else ''

        # 归档模式
        cur.execute("SELECT log_mode FROM v$database")
        r = cur.fetchone()
        results['log_mode'] = r[0] if r else ''

        # force logging
        cur.execute("SELECT force_logging FROM v$database")
        r = cur.fetchone()
        results['force_logging'] = r[0] if r else ''

        #  flashback
        cur.execute("SELECT flashback_on FROM v$database")
        r = cur.fetchone()
        results['flashback_on'] = r[0] if r else ''

        # 创建时间
        cur.execute("SELECT TO_CHAR(CREATED, 'YYYY-MM-DD HH24:MI:SS') FROM v$database")
        r = cur.fetchone()
        results['created'] = r[0] if r else ''

        # 启动时间
        cur.execute("SELECT TO_CHAR(STARTUP_TIME, 'YYYY-MM-DD HH24:MI:SS') FROM v$instance")
        r = cur.fetchone()
        results['startup_time'] = r[0] if r else ''

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


# ═══════════════════════════════════════════════════════════════════════════
#            版本专用数据库层巡检（对应 dbcheck10g/11g/12c.sql）
# ═══════════════════════════════════════════════════════════════════════════

def _base_db_check(conn, cdb_col="", pdb_col=""):
    """三个版本公用的基础查询部分，避免代码重复"""
    results = {}
    cur = conn.cursor()
    try:
        if cdb_col and pdb_col:
            # 12c+：CDB/PDB 架构
            cur.execute(f"""
                SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE,
                       {cdb_col} CDB, {pdb_col} PLUGGABLE_DB
                FROM v$database
            """)
        else:
            # 10g/11g：无 CDB 概念
            cur.execute("""
                SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE
                FROM v$database
            """)
        results['database'] = cur.fetchall()

        cur.execute("SELECT global_name FROM global_name")
        results['global_name'] = cur.fetchone()

        # 字符集
        try:
            cur.execute("""
                SELECT parameter, value
                FROM nls_database_parameters
                WHERE parameter IN ('NLS_CHARACTERSET', 'NLS_NCHAR_CHARACTERSET')
            """)
            rows = cur.fetchall()
            results['charset'] = tuple(r[1] for r in rows) if rows else ('', '')
        except Exception:
            results['charset'] = ('', '')

        for param in ['db_block_size', 'sga_max_size', 'sga_target',
                      'pga_aggregate_target', 'memory_max_target', 'memory_target']:
            cur.execute(f"SELECT value FROM v$parameter WHERE name='{param}'")
            r = cur.fetchone()
            results[param] = r[0] if r else ''

        cur.execute("SELECT value FROM v$parameter WHERE name='spfile'")
        results['spfile'] = cur.fetchone()[0] if cur.fetchone() else ''

        cur.execute("SELECT value FROM v$parameter WHERE name='db_create_file_dest'")
        results['omf'] = cur.fetchone()[0] if cur.fetchone() else ''

        for col in ['log_mode', 'force_logging', 'flashback_on']:
            cur.execute(f"SELECT {col} FROM v$database")
            r = cur.fetchone()
            results[col] = r[0] if r else ''

        cur.execute("SELECT TO_CHAR(CREATED, 'YYYY-MM-DD HH24:MI:SS') FROM v$database")
        results['created'] = cur.fetchone()[0] if cur.fetchone() else ''

        cur.execute("SELECT TO_CHAR(STARTUP_TIME, 'YYYY-MM-DD HH24:MI:SS') FROM v$instance")
        results['startup_time'] = cur.fetchone()[0] if cur.fetchone() else ''

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_database_v10(conn):
    """Oracle 10g 数据库层巡检 — 基于 dbcheck10g.sql

    特点：WMSYS.WM_CONCAT 替代 listagg；无 CDB；无 PLUGGABLE_NDB
    """
    return _base_db_check(conn)


def oracle_check_database_v11(conn):
    """Oracle 11g 数据库层巡检 — 基于 dbcheck11g.sql

    特点：标准 listagg；无 CDB；无 PLUGGABLE_NDB；gv$instance 需额外字段兼容
    """
    results = _base_db_check(conn)
    cur = conn.cursor()
    try:
        # 11g 特有的 listagg（RAC 节点列表）
        # 10g 用 WMSYS.WM_CONCAT，11g 用标准 listagg
        cur.execute("""
            SELECT 'Instances: [' || listagg(instance_name, ', ') within group(order by instance_name) || '] ' as instances
            FROM gv$instance
        """)
        row = cur.fetchone()
        results['rac_instances'] = row[0] if row else ''

        # 11g 回收站（dba_recyclebin）
        try:
            cur.execute("""
                SELECT owner,
                       round(SUM(a.space *
                                 (SELECT value FROM v$parameter WHERE name='db_block_size')) / 1024 / 1024, 2) recyb_size_M,
                       count(1) recyb_cnt
                FROM dba_recyclebin a
                GROUP BY owner
            """)
            results['recyclebin'] = cur.fetchall()
        except Exception:
            pass

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_database_v12plus(conn):
    """Oracle 12c+ 数据库层巡检 — 基于 dbcheck12c.sql

    特点：CDB/PDB 架构；listagg；cdb_recyclebin；gv$crs_resource_v2
    """
    return _base_db_check(conn, cdb_col="CDB", pdb_col="PLUGGABLE_DB")


def oracle_check_version_and_patches(conn):
    """数据库版本和补丁"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("SELECT banner FROM v$version")
        results['version'] = cur.fetchall()

        # OPatch 补丁
        try:
            cur.execute("SELECT * FROM v$system_patch WHERE patch_id IS NOT NULL")
            results['patches'] = cur.fetchall()
        except Exception:
            # 12c 及以上用这个
            try:
                cur.execute("""
                    SELECT patch_id, patch_type, description, action, action_time
                    FROM dba_registry_sqlpatch
                    ORDER BY action_time DESC
                """)
                results['patches'] = cur.fetchall()
            except Exception:
                results['patches'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_tablespace(conn):
    """表空间使用率（通用版）：实测列名 dba_free_space.BYTES / dba_temp_free_space.FREE_SPACE"""
    results = {}
    cur = conn.cursor()
    try:
        # 永久表空间：dba_tablespaces LEFT JOIN dba_data_files LEFT JOIN dba_free_space
        cur.execute("""
            SELECT t.tablespace_name,
                   t.status,
                   ROUND(NVL(df.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(df.max_mb,0), 2) max_mb,
                   ROUND(NVL(df.curr_mb,0) - NVL(fs.free_mb,0), 2) used_mb,
                   ROUND(NVL(fs.free_mb,0), 2) free_mb,
                   ROUND((NVL(df.curr_mb,0) - NVL(fs.free_mb,0)) /
                         NULLIF(NVL(df.curr_mb,0),0) * 100, 2) pct_used
            FROM dba_tablespaces t
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_data_files GROUP BY tablespace_name) df
               ON t.tablespace_name = df.tablespace_name
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) free_mb
                       FROM dba_free_space GROUP BY tablespace_name) fs
               ON t.tablespace_name = fs.tablespace_name
            WHERE t.contents = 'PERMANENT'
            ORDER BY pct_used DESC NULLS LAST
        """)
        results['data_tablespaces'] = cur.fetchall()

        # 临时表空间：dba_tablespaces LEFT JOIN dba_temp_files LEFT JOIN dba_temp_free_space(FREE_SPACE)
        cur.execute("""
            SELECT t.tablespace_name,
                   t.status,
                   ROUND(NVL(tf.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(tf.max_mb,0), 2) max_mb,
                   '-' used_mb,
                   ROUND(NVL(tfs.free_mb,0), 2) free_mb,
                   '-' pct_used
            FROM dba_tablespaces t
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_temp_files GROUP BY tablespace_name) tf
               ON t.tablespace_name = tf.tablespace_name
            LEFT JOIN (SELECT tablespace_name,
                              SUM(free_space/1024/1024) free_mb
                       FROM dba_temp_free_space GROUP BY tablespace_name) tfs
               ON t.tablespace_name = tfs.tablespace_name
            WHERE t.contents = 'TEMPORARY'
            ORDER BY t.tablespace_name
        """)
        results['temp_tablespaces'] = cur.fetchall()

        # 自动扩展文件
        cur.execute("""
            SELECT tablespace_name, file_name,
                   ROUND(bytes/1024/1024,2) curr_mb,
                   ROUND(MAXBYTES/1024/1024,2) max_mb,
                   AUTOEXTENSIBLE
            FROM dba_data_files
            WHERE AUTOEXTENSIBLE = 'YES'
            ORDER BY tablespace_name
        """)
        results['autoextend_files'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_redolog(conn):
    """Redo 日志检查"""
    results = {}
    cur = conn.cursor()
    try:
        # 组成员和大小
        cur.execute("""
            SELECT GROUP#, THREAD#, SEQUENCE#, ROUND(BYTES/1024/1024,2) size_mb,
                   STATUS, MEMBERS, ARCHIVED
            FROM v$log
            ORDER BY THREAD#, GROUP#
        """)
        results['logs'] = cur.fetchall()

        # 日志文件
        cur.execute("""
            SELECT GROUP#, MEMBER, TYPE, STATUS
            FROM v$logfile
            ORDER BY GROUP#
        """)
        results['logfiles'] = cur.fetchall()

        # 最近 Redo 切换频率
        cur.execute("""
            SELECT h.thread#,
                   COUNT(*) switch_cnt,
                   ROUND(COUNT(*) * MAX(b.bytes)/1024/1024/1024, 2) total_mb
            FROM v$loghist h, v$log b
            WHERE h.group# = b.group#
              AND h.first_time > SYSDATE - 7
            GROUP BY h.thread#
        """)
        results['redo_switch'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_controlfile(conn):
    """控制文件"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT NAME, STATUS, IS_RECOVERY_DEST_FILE, BLOCK_SIZE, FILE_SIZE_BLKS
            FROM v$controlfile
            ORDER BY STATUS, NAME
        """)
        results['controlfiles'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_invalid_objects(conn):
    """无效对象"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT OWNER, OBJECT_TYPE, COUNT(*) cnt
            FROM dba_objects
            WHERE STATUS = 'INVALID'
            GROUP BY OWNER, OBJECT_TYPE
            ORDER BY cnt DESC
        """)
        results['invalid_by_type'] = cur.fetchall()

        cur.execute("""
            SELECT OWNER, OBJECT_NAME, OBJECT_TYPE, STATUS, LAST_DDL_TIME
            FROM dba_objects
            WHERE STATUS = 'INVALID'
              AND OWNER NOT IN ('SYS','SYSTEM')
            ORDER BY LAST_DDL_TIME DESC
        """)
        results['invalid_detail'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_users(conn):
    """用户安全检查"""
    results = {}
    cur = conn.cursor()
    try:
        # 默认密码用户（常见高危账号）
        default_accounts = ['SCOTT','JONES','ADAMS','CLARK','BLAKE','HR','OE','PM','IX','SH','DIP','ORACLE_OCM','XS$NULL','APPQOSSYS']
        placeholders = ','.join([f":{i}" for i in range(len(default_accounts))])
        cur.execute(f"""
            SELECT username, account_status, lock_date, expiry_date, created
            FROM dba_users
            WHERE username IN ({placeholders})
        """, default_accounts)
        results['default_accounts'] = cur.fetchall()

        # 锁定的用户
        cur.execute("""
            SELECT username, account_status, lock_date, expiry_date
            FROM dba_users
            WHERE account_status NOT IN ('OPEN', 'EXPIRED(GRACE)')
            ORDER BY account_status
        """)
        results['locked_users'] = cur.fetchall()

        # 系统角色
        cur.execute("""
            SELECT granted_role, grantee, admin_option
            FROM dba_role_privs
            WHERE grantee NOT IN ('SYS', 'SYSTEM')
              AND admin_option = 'YES'
            ORDER BY granted_role
        """)
        results['admin_roles'] = cur.fetchall()

        # Profile
        cur.execute("""
            SELECT profile, resource_name, resource_type, LIMIT
            FROM dba_profiles
            ORDER BY profile, resource_name
        """)
        results['profiles'] = cur.fetchall()

        # 密码有效期
        cur.execute("""
            SELECT profile, LIMIT PASSWORD_LIFE_TIME
            FROM dba_profiles
            WHERE resource_name = 'PASSWORD_LIFE_TIME'
            ORDER BY profile
        """)
        results['password_policy'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_top_sql(conn, limit=20):
    """Top SQL（按逻辑读/物理读/executions）"""
    results = {}
    cur = conn.cursor()
    try:
        # 按 Buffer Gets 排序
        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id, SUBSTR(sql_text,1,80) sql_text,
                       ROUND(buffer_gets/1024/1024,2) buf_mb,
                       ROUND(disk_reads/1024/1024,2) disk_mb,
                       executions, ROUND(elapsed_time/1000000,2) elapsed_sec,
                       ROUND(buffer_gets/DECODE(executions,0,1,executions)) gets_per_exec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY buffer_gets DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_buffer_gets'] = cur.fetchall()

        # 按磁盘读排序
        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id, SUBSTR(sql_text,1,80) sql_text,
                       ROUND(disk_reads/1024/1024,2) disk_mb,
                       executions, ROUND(elapsed_time/1000000,2) elapsed_sec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY disk_reads DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_disk_reads'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_awr(conn):
    """AWR 快照信息"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT instance_number, snap_id, TO_CHAR(begin_interval_time,'YYYY-MM-DD HH24:MI') bt,
                   TO_CHAR(end_interval_time,'YYYY-MM-DD HH24:MI') et,
                   ROUND((end_interval_time - begin_interval_time) * 24, 2) elapsed_hr,
                   ERROR_COUNT
            FROM dba_hist_snapshot
            WHERE end_interval_time > SYSDATE - 7
            ORDER BY instance_number, snap_id DESC
        """)
        results['awr_snaps'] = cur.fetchall()

        # AWR 设置
        cur.execute("""
            SELECT * FROM dba_hist_wr_control
        """)
        results['awr_settings'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_performance(conn):
    """性能指标（Session / Wait Events / SGA/PGA）"""
    results = {}
    cur = conn.cursor()
    try:
        # 当前会话数
        cur.execute("""
            SELECT status, COUNT(*) FROM v$session GROUP BY status
        """)
        results['session_by_status'] = cur.fetchall()

        # 等待事件 Top10
        cur.execute("""
            SELECT * FROM (
                SELECT event, total_waits, ROUND(time_waited/100,2) time_sec,
                       ROUND(total_waits/DECODE(total_waits,0,1,time_waited)*100,4) wait_pct,
                       wait_class
                FROM v$system_event
                WHERE event NOT IN ('rdbms ipc message','smon timer','pmon','pipe get',
                                    'SQL*Net message from client','class slave wait')
                ORDER BY time_waited DESC
            ) WHERE ROWNUM <= 10
        """)
        results['wait_events'] = cur.fetchall()

        # SGA 组件
        cur.execute("""
            SELECT name, ROUND(bytes/1024/1024,2) size_mb
            FROM v$sgastat
            WHERE pool IS NOT NULL
            ORDER BY bytes DESC
        """)
        results['sga_pools'] = cur.fetchall()

        # PGA Target
        cur.execute("""
            SELECT a.name, ROUND(a.value/1024/1024,2) size_mb
            FROM v$pgastat a
            WHERE a.name IN ('aggregate PGA target parameter',
                             'total PGA allocated',
                             'total PGA inuse',
                             'maximum PGA allocated')
        """)
        results['pga'] = cur.fetchall()

        # 缓冲区命中率
        cur.execute("""
            SELECT name, ROUND((1 - physical_reads / (db_block_gets + consistent_gets)) * 100, 2) hit_pct,
                   db_block_gets, consistent_gets, physical_reads
            FROM v$buffer_pool_statistics
            WHERE db_block_gets + consistent_gets > 0
        """)
        results['buffer_hit'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_dataguard(conn):
    """Data Guard 配置"""
    results = {}
    cur = conn.cursor()
    try:
        # Data Guard 配置
        cur.execute("""
            SELECT GROUP#, TYPE, MEMBER, IS_RECOVERY_DEST_FILE
            FROM v$logfile
            WHERE TYPE = 'STANDBY'
        """)
        results['standby_logs'] = cur.fetchall()

        # 归档目的地
        cur.execute("""
            SELECT dest_id, status, destination, archiver, transmit_mode,
                   archiver, REGISTER
            FROM v$archive_dest
            WHERE destination IS NOT NULL
        """)
        results['archive_dest'] = cur.fetchall()

        # 实时查询
        cur.execute("""
            SELECT database_mode, recovery_mode, protection_mode, standby_db_unique_name
            FROM v$archive_dest_status
            WHERE status != 'INACTIVE'
        """)
        results['dg_status'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_backup(conn):
    """RMAN 备份信息"""
    results = {}
    cur = conn.cursor()
    try:
        # 最近备份
        cur.execute("""
            SELECT session_key, INPUT_TYPE, STATUS,
                   TO_CHAR(START_TIME,'YYYY-MM-DD HH24:MI') start_t,
                   TO_CHAR(END_TIME,'YYYY-MM-DD HH24:MI') end_t,
                   ROUND(bytes/1024/1024/1024,2) size_gb, elapsed_minute
            FROM v$rman_backup_job_details
            WHERE end_time > SYSDATE - 30
            ORDER BY end_time DESC
        """)
        results['rman_jobs'] = cur.fetchall()

        # 备份集
        cur.execute("""
            SELECT SET_STAMP, INPUT_TYPE,
                   ROUND(INPUT_BYTES/1024/1024/1024,2) input_gb,
                   ROUND(OUTPUT_BYTES/1024/1024/1024,2) output_gb,
                   COMPRESSION_RATIO
            FROM v$backup_set
            WHERE COMPLETION_TIME > SYSDATE - 30
        """)
        results['backup_sets'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_sga_pga(conn):
    """SGA/PGA 内存信息"""
    results = {}
    cur = conn.cursor()
    try:
        # SGA 动态组件
        cur.execute("""
            SELECT component, current_size/1024/1024 AS curr_mb,
                   min_size/1024/1024 AS min_mb,
                   user_specified_size/1024/1024 AS user_mb
            FROM v$sga_dynamic_components
            WHERE current_size > 0
            ORDER BY current_size DESC
        """)
        results['sga_components'] = cur.fetchall()

        # SGA 总计
        try:
            cur.execute("""
                SELECT SUM(value)/1024/1024 AS sga_total_mb
                FROM v$sga
            """)
            results['sga_total'] = cur.fetchall()
        except Exception:
            results['sga_total'] = []

        # PGA 统计
        cur.execute("""
            SELECT NAME, VALUE/1024/1024 AS value_mb
            FROM v$pgastat
            WHERE NAME IN (
                'total PGA allocated','total PGA inuse',
                'aggregate PGA target parameter','aggregate PGA auto target',
                'maximum PGA allocated','total freeable PGA memory'
            )
        """)
        results['pga_stats'] = cur.fetchall()

        # Memory Target / SGA Target / PGA Aggregate Target
        cur.execute("""
            SELECT NAME, VALUE, DISPLAY_VALUE, ISDEFAULT
            FROM v$parameter
            WHERE NAME IN (
                'memory_target','memory_max_target',
                'sga_target','pga_aggregate_target','sga_max_size'
            )
        """)
        results['memory_params'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_params(conn):
    """关键参数"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT NAME, VALUE, DISPLAY_VALUE, ISDEFAULT,
                   ISSES_MODIFIABLE, ISSYS_MODIFIABLE, DESCRIPTION
            FROM v$parameter
            WHERE NAME IN (
                'processes','sessions','open_cursors','db_block_size',
                'db_file_multiblock_read_count','db_writer_processes',
                'undo_retention','compatible','nls_characterset',
                'nls_nchar_characterset','job_queue_processes',
                'parallel_max_servers','audit_trail','recyclebin',
                'optimizer_mode','cursor_sharing','statistics_level',
                'control_file_record_keep_time','remote_login_passwordfile',
                'resource_manager_plan'
            )
            ORDER BY NAME
        """)
        results['params'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_undo(conn):
    """Undo 表空间信息"""
    results = {}
    cur = conn.cursor()
    try:
        # Undo 表空间基本信息（容错，10g 可能略有差异）
        try:
            cur.execute("""
                SELECT d.undo_tablespace,
                       r.retention,
                       ROUND(NVL(ts.used_bytes/1024/1024,0),2) AS used_mb,
                       ROUND(NVL(ts.tbs_bytes/1024/1024,0),2) AS total_mb,
                       u.exp_blks AS exp_undo_blks,
                       u.unexp_blks AS unexp_undo_blks,
                       u.blk_cnt AS undo_blk_cnt
                FROM (
                    SELECT UPPER(VALUE) AS undo_tablespace
                    FROM v$parameter WHERE NAME='undo_tablespace'
                ) d,
                (
                    SELECT UPPER(VALUE) AS retention
                    FROM v$parameter WHERE NAME='undo_retention'
                ) r,
                (
                    SELECT SUM(df.bytes) AS tbs_bytes,
                           SUM(df.bytes)-NVL(SUM(ff.free_bytes),0) AS used_bytes
                    FROM dba_data_files df
                    LEFT JOIN (
                        SELECT tablespace_name, SUM(bytes) AS free_bytes
                        FROM dba_free_space GROUP BY tablespace_name
                    ) ff ON df.tablespace_name = ff.tablespace_name
                    WHERE df.tablespace_name = (
                        SELECT UPPER(VALUE) FROM v$parameter WHERE NAME='undo_tablespace'
                    )
                ) ts,
                (
                    SELECT COUNT(*) AS exp_blks, 0 AS unexp_blks, COUNT(*) AS blk_cnt
                    FROM v$undostat
                    WHERE begin_time > SYSDATE-1 AND undoblks > 0
                ) u
            """)
            results['undo_info'] = cur.fetchall()
        except Exception as e:
            results['undo_info'] = []

        # Undo 段统计
        try:
            cur.execute("""
                SELECT status, COUNT(*) AS num_segments,
                       SUM(bytes)/1024/1024 AS total_mb
                FROM dba_undo_extents
                GROUP BY status
            """)
            results['undo_segments'] = cur.fetchall()
        except Exception:
            results['undo_segments'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_long_sql(conn):
    """长时间运行的 SQL"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT sid, serial#, username, sql_id, opname,
                   sofar, totalwork,
                   ROUND(sofar/GREATEST(totalwork,0.001)*100,1) AS pct_complete,
                   elapsed_seconds, time_remaining
            FROM v$session_longops
            WHERE totalwork > 0 AND sofar < totalwork AND elapsed_seconds > 30
            ORDER BY elapsed_seconds DESC
            FETCH FIRST 10 ROWS ONLY
        """)
        results['long_sql'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_rac(conn):
    """RAC 检查"""
    results = {}
    cur = conn.cursor()
    try:
        # 实例列表
        cur.execute("""
            SELECT INST_ID, INSTANCE_NAME, HOST_NAME, STATUS, PARALLEL
            FROM gv$instance
        """)
        results['rac_instances'] = cur.fetchall()

        # CRS 资源状态（11gR2+）
        try:
            cur.execute("""
                SELECT name, TYPE, STATE, TARGET, host_name
                FROM gv$crs_resource
                WHERE TYPE IN ('database','service','listener')
                ORDER BY TYPE, name
            """)
            results['crs_resources'] = cur.fetchall()
        except Exception:
            results['crs_resources'] = []

        # OCR 备份
        try:
            cur.execute("""
                SELECT group_number, name, state, type, total_mb, free_mb,
                       ROUND(free_mb/total_mb*100,2) free_pct
                FROM v$asm_diskgroup
            """)
            results['asm_diskgroups'] = cur.fetchall()
        except Exception:
            results['asm_diskgroups'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_jobs(conn):
    """Scheduler Jobs / DBMS_JOBS"""
    results = {}
    cur = conn.cursor()
    try:
        # Scheduler Jobs
        cur.execute("""
            SELECT JOB_NAME, STATE, ENABLED, SCHEDULE_TYPE,
                   TO_CHAR(NEXT_RUN_DATE,'YYYY-MM-DD HH24:MI') next_run,
                   RUN_COUNT, FAILURE_COUNT
            FROM dba_scheduler_jobs
            WHERE owner NOT IN ('SYS','SYSTEM')
            ORDER BY owner, job_name
        """)
        results['scheduler_jobs'] = cur.fetchall()

        # 失败的后台作业
        cur.execute("""
            SELECT job, WHAT, LAST_DATE, NEXT_DATE, FAILURES, BROKEN
            FROM dba_jobs
            WHERE failures > 0 OR broken = 'Y'
        """)
        results['failed_jobs'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_flashback(conn):
    """闪回配置"""
    results = {}
    cur = conn.cursor()
    try:
        # v$database 有 flashback_on 和 oldest_flashback_time，
        # 但 retention_target 在 v$database 中不存在，需从 v$flashback_database_stat 计算
        cur.execute("""
            SELECT d.flashback_on,
                   TO_CHAR(d.oldest_flashback_time,'YYYY-MM-DD HH24:MI') oldest_t,
                   ROUND(f.retention_min, 2) ret_min
            FROM v$database d,
                 (SELECT MAX(retention_flashback_time - flashback_time) * 24 * 60 retention_min
                  FROM v$flashback_database_stat) f
        """)
        results['flashback'] = cur.fetchall()

        # 回收站（CDB 下用 cdb_recyclebin）
        try:
            cur.execute("""
                SELECT r.owner, r.original_name, r.type,
                       ROUND(r.space * (SELECT value FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                       r.can_undrop, r.can_purge
                FROM dba_recyclebin r
                ORDER BY mb DESC
            """)
            results['recyclebin'] = cur.fetchall()
        except Exception:
            # CDB 环境可能需要从 PDB 查询
            try:
                cur.execute("""
                    SELECT owner, original_name, type,
                           ROUND(space * (SELECT value FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                           can_undrop, can_purge
                    FROM cdb_recyclebin
                    ORDER BY mb DESC
                """)
                results['recyclebin'] = cur.fetchall()
            except Exception:
                results['recyclebin'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_alert(conn, days=7):
    """最近 Alert 日志错误（优先 v$diag_alert_text，备选 v$diag_alert_xml）"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT TO_CHAR(alert_time,'YYYY-MM-DD HH24:MI:SS') t,
                   SUBSTR(message_text,1,200) message
            FROM v$diag_alert_text
            WHERE alert_time > SYSDATE - :days
              AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
            ORDER BY alert_time DESC
        """, days=days)
        results['alert_errors'] = cur.fetchall()
    except Exception:
        # v$diag_alert_text 不存在或无权限时尝试 XML 视图
        try:
            cur.execute("""
                SELECT TO_CHAR(trap_time,'YYYY-MM-DD HH24:MI:SS') t,
                       SUBSTR(message_text,1,200) message
                FROM v$diag_alert_xml
                WHERE trap_time > SYSDATE - :days
                  AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
                ORDER BY trap_time DESC
            """, days=days)
            results['alert_errors'] = cur.fetchall()
        except Exception as e:
            results['error'] = str(e)
    finally:
        cur.close()
    return results


# ═══════════════════════════════════════════════════════════════════════════
#                    报告生成（HTML）
# ═══════════════════════════════════════════════════════════════════════════

def _html_table(headers, rows, id_="", class_="dbcheck_tbl"):
    """生成 HTML 表格"""
    lines = []
    if id_ or class_:
        lines.append(f'<table id="{id_}" class="{class_}" border="1" cellpadding="4" cellspacing="0">')
    else:
        lines.append('<table border="1" cellpadding="4" cellspacing="0">')
    # 表头
    lines.append('<thead><tr>')
    for h in headers:
        lines.append(f'<th>{h}</th>')
    lines.append('</tr></thead>')
    # 数据行
    lines.append('<tbody>')
    for i, row in enumerate(rows):
        bg = '#FFFFFF' if i % 2 == 0 else '#F5F5F5'
        lines.append(f'<tr style="background:{bg}">')
        for cell in row:
            lines.append(f'<td>{str(cell) if cell is not None else ""}</td>')
        lines.append('</tr>')
    lines.append('</tbody></table>')
    return '\n'.join(lines)


def _html_section(title, content, anchor=""):
    anchor_tag = f'<a name="{anchor}"></a>' if anchor else ''
    return f'''
<h2 id="{anchor}">{anchor_tag}{title}</h2>
<div class="section">{content}</div>
'''


def build_html_report(db_info, os_data, check_results, db_version, ai_advice='', inspector=''):
    """构建完整 HTML 巡检报告"""
    from datetime import datetime

    # ── 样式 ────────────────────────────────────────────────────────────────
    css = """
    <style>
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body { font-family: Consolas, 'Courier New', monospace; font-size: 13px;
           background: #f4f4f4; color: #222; padding: 20px; }
    h1 { font-size: 22px; color: #fff; background: #0066cc;
         padding: 16px 24px; border-radius: 6px; margin-bottom: 20px; }
    h2 { font-size: 16px; color: #fff; background: #0066cc;
         padding: 8px 14px; margin: 24px 0 10px; border-radius: 4px; }
    h3 { font-size: 14px; color: #336699; margin: 14px 0 6px; }
    table { width: 100%%; border-collapse: collapse; margin: 8px 0;
             background: #fff; font-size: 12px; }
    th { background: #336699; color: #fff; padding: 8px 10px; text-align: left; }
    td { padding: 6px 10px; border: 1px solid #ddd; vertical-align: top; }
    tr:nth-child(odd) { background: #fff; }
    tr:nth-child(even) { background: #f0f6ff; }
    .summary-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
                    gap: 12px; margin: 12px 0; }
    .summary-card { background: #fff; border: 1px solid #ccc; border-radius: 6px;
                    padding: 12px 16px; box-shadow: 2px 2px 4px rgba(0,0,0,0.08); }
    .summary-card .label { font-size: 11px; color: #888; margin-bottom: 4px; }
    .summary-card .value { font-size: 16px; font-weight: bold; color: #0066cc; }
    .ok    { color: green; font-weight: bold; }
    .warn  { color: #cc6600; font-weight: bold; }
    .error { color: red; font-weight: bold; }
    .nav { background: #e8f0fe; padding: 10px 16px; border-radius: 4px;
            margin-bottom: 16px; font-size: 12px; }
    .nav a { margin-right: 14px; color: #0066cc; text-decoration: none; }
    .nav a:hover { text-decoration: underline; }
    .section { background: #fff; border: 1px solid #ddd; border-radius: 6px;
               padding: 14px; margin-top: 6px; }
    .footer { text-align: center; color: #888; font-size: 11px; margin-top: 30px; }
    </style>
    """

    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # ── 汇总卡片 ────────────────────────────────────────────────────────────
    inst = check_results.get('instance', [])
    db   = db_info
    ver  = db_version

    hostname = os_data.get('hostname', '')
    uptime   = os_data.get('uptime', '')
    cpu      = f"{os_data.get('cpu_model','')} × {os_data.get('cpu_count','?')}"
    mem      = f"{os_data.get('mem_total_mb','?')} MB，使用率 {os_data.get('mem_usage_pct','?')}%"

    cards_html = f"""
    <div class="summary-grid">
        <div class="summary-card">
            <div class="label">主机名</div>
            <div class="value">{hostname}</div>
        </div>
        <div class="summary-card">
            <div class="label">数据库版本</div>
            <div class="value">{ver}</div>
        </div>
        <div class="summary-card">
            <div class="label">运行时间</div>
            <div class="value">{uptime}</div>
        </div>
        <div class="summary-card">
            <div class="label">CPU</div>
            <div class="value">{cpu}</div>
        </div>
        <div class="summary-card">
            <div class="label">内存</div>
            <div class="value">{mem}</div>
        </div>
        <div class="summary-card">
            <div class="label">巡检时间</div>
            <div class="value">{now}</div>
        </div>
    </div>
    """

    # ── 导航 ────────────────────────────────────────────────────────────────
    nav = """
    <div class="nav">
        <a href="#os_info">OS信息</a>
        <a href="#db_info">数据库信息</a>
        <a href="#version">版本与补丁</a>
        <a href="#tablespace">表空间</a>
        <a href="#redolog">Redo日志</a>
        <a href="#controlfile">控制文件</a>
        <a href="#performance">性能指标</a>
        <a href="#top_sql">Top SQL</a>
        <a href="#invalid_objects">无效对象</a>
        <a href="#users">用户安全</a>
        <a href="#backup">备份</a>
        <a href="#flashback">闪回与回收站</a>
        <a href="#dataguard">Data Guard</a>
        <a href="#rac">RAC</a>
        <a href="#awr">AWR</a>
        <a href="#jobs">作业</a>
        <a href="#alert">Alert日志</a>
        <a href="#ai_diagnosis">AI诊断</a>
    </div>
    """

    # ── OS 信息 ─────────────────────────────────────────────────────────────
    os_rows = []
    for k in ['hostname','os_version','kernel','uptime','cpu_model','cpu_count',
              'mem_total_mb','mem_used_mb','mem_usage_pct','swap_total_mb','swap_used_mb',
              'load_average','disk_usage','hugepages','thp']:
        v = os_data.get(k, 'N/A')
        os_rows.append((k, v))
    os_section = _html_section('🖥  OS 主机信息', _html_table(['项目','内容'], os_rows), 'os_info')

    # ── 数据库基本信息 ──────────────────────────────────────────────────────
    db_rows = []
    for k, label in [
        ('NAME','数据库名'), ('DATABASE_ROLE','角色'), ('OPEN_MODE','打开模式'),
        ('LOG_MODE','归档模式'), ('CREATED','创建时间'), ('STARTUP_TIME','启动时间'),
        ('CDB','CDB'), ('flashback_on','闪回'), ('force_logging','Force Logging'),
        ('block_size','块大小'), ('sga_max_size','SGA Max'), ('sga_target','SGA Target'),
        ('pga_aggregate_target','PGA Target'), ('spfile','SPFILE'),
        ('charset','字符集'), ('global_name','全局名'),
    ]:
        v = db_info.get(k, 'N/A')
        db_rows.append((label, v))
    db_section = _html_section('📋  数据库基本信息', _html_table(['项目','值'], db_rows), 'db_info')

    # ── 版本 ────────────────────────────────────────────────────────────────
    ver_rows = check_results.get('version', [])
    ver_section = _html_section('🔧  数据库版本与补丁',
        _html_table(['版本信息'], [[r[0]] for r in ver_rows]) if ver_rows else '无数据', 'version')

    # ── 表空间 ──────────────────────────────────────────────────────────────
    ts_rows = check_results.get('data_tablespaces', [])
    ts_headers = ['表空间名','状态','类型','最大MB','当前MB','已用MB','使用率%']
    ts_section = _html_section('📦  表空间（数据文件）',
        _html_table(ts_headers, ts_rows) if ts_rows else '无数据', 'tablespace')

    temp_rows = check_results.get('temp_tablespaces', [])
    temp_headers = ['临时表空间','状态','最大MB','当前MB','已用MB','使用率%']
    ts_section += _html_section('📦  临时表空间',
        _html_table(temp_headers, temp_rows) if temp_rows else '无数据', 'tablespace_temp')

    # ── Redo ─────────────────────────────────────────────────────────────────
    redo_rows = check_results.get('logs', [])
    redo_headers = ['Group#','Thread#','Sequence#','大小MB','状态','成员数','已归档']
    redo_section = _html_section('⚙  Redo 日志',
        _html_table(redo_headers, redo_rows) if redo_rows else '无数据', 'redolog')

    # ── 控制文件 ─────────────────────────────────────────────────────────────
    cf_rows = check_results.get('controlfiles', [])
    cf_headers = ['名称','状态','恢复目录','块大小','文件块数']
    cf_section = _html_section('⚙  控制文件',
        _html_table(cf_headers, cf_rows) if cf_rows else '无数据', 'controlfile')

    # ── 性能 ─────────────────────────────────────────────────────────────────
    perf_parts = []
    ses_rows = check_results.get('session_by_status', [])
    perf_parts.append('<h3>会话状态</h3>' + _html_table(['状态','数量'], ses_rows) if ses_rows else '')

    wait_rows = check_results.get('wait_events', [])
    perf_parts.append('<h3>Top10 等待事件</h3>' +
        _html_table(['事件','总等待次数','等待时间(秒)','等待占比%','类别'], wait_rows) if wait_rows else '')

    buf_rows = check_results.get('buffer_hit', [])
    perf_parts.append('<h3>缓冲区命中率</h3>' +
        _html_table(['池名','命中率%','Block Gets','Consistent Gets','物理读'], buf_rows) if buf_rows else '')

    perf_section = _html_section('⚡  性能指标', '<br>'.join(perf_parts), 'performance')

    # ── Top SQL ─────────────────────────────────────────────────────────────
    sql_parts = []
    for label, key in [('按Buffer Gets', 'top_sql_buffer_gets'), ('按磁盘读', 'top_sql_disk_reads')]:
        rows = check_results.get(key, [])
        sql_parts.append(f'<h3>{label}</h3>' +
            _html_table(['SQL_ID','SQL片段(80字符)','Buf MB','Disk MB','执行次数','耗时秒','Gets/执行','模块'], rows)
            if rows else '')
    top_sql_section = _html_section('🔍  Top SQL', '<br>'.join(sql_parts), 'top_sql')

    # ── 无效对象 ─────────────────────────────────────────────────────────────
    io_rows = check_results.get('invalid_by_type', [])
    io_detail = check_results.get('invalid_detail', [])
    io_section = _html_section('⚠  无效对象',
        (_html_table(['所有者','类型','数量'], io_rows) if io_rows else '') +
        '<h3>详细（排除SYS/SYSTEM）</h3>' +
        (_html_table(['所有者','对象名','类型','状态','最后DDL时间'], io_detail) if io_detail else ''),
        'invalid_objects')

    # ── 用户安全 ─────────────────────────────────────────────────────────────
    sec_parts = []
    acc_rows = check_results.get('default_accounts', [])
    sec_parts.append('<h3>默认账户（高危）</h3>' +
        _html_table(['用户名','状态','锁定日期','到期日期','创建时间'], acc_rows) if acc_rows else '<p>无默认账户数据</p>')

    lock_rows = check_results.get('locked_users', [])
    sec_parts.append('<h3>锁定/过期用户</h3>' +
        _html_table(['用户名','状态','锁定日期','到期日期'], lock_rows) if lock_rows else '')

    role_rows = check_results.get('admin_roles', [])
    sec_parts.append('<h3>带管理权限的角色（非SYS/SYSTEM）</h3>' +
        _html_table(['角色','授权用户','管理选项'], role_rows) if role_rows else '')

    user_section = _html_section('🔒  用户与安全', '<br>'.join(sec_parts), 'users')

    # ── 备份 ─────────────────────────────────────────────────────────────────
    backup_parts = []
    rman_rows = check_results.get('rman_jobs', [])
    backup_parts.append('<h3>RMAN 备份任务（近30天）</h3>' +
        _html_table(['会话KEY','类型','状态','开始时间','结束时间','大小GB','耗时分钟'], rman_rows)
        if rman_rows else '<p>近30天无RMAN备份记录</p>')
    backup_section = _html_section('💾  备份信息', '<br>'.join(backup_parts), 'backup')

    # ── 闪回 ─────────────────────────────────────────────────────────────────
    fb_rows = check_results.get('flashback', [])
    fb_parts = ['<h3>闪回配置</h3>' + _html_table(['闪回开关','保留目标(分)','最旧闪回时间','保留分钟'], fb_rows) if fb_rows else '']
    rb_rows = check_results.get('recyclebin', [])
    fb_parts.append('<h3>回收站</h3>' + _html_table(['所有者','原名','类型','空间MB','可恢复','可清除'], rb_rows) if rb_rows else '')
    flashback_section = _html_section('⏪  闪回与回收站', '<br>'.join(fb_parts), 'flashback')

    # ── Data Guard ───────────────────────────────────────────────────────────
    dg_parts = []
    dg_rows = check_results.get('dg_status', [])
    dg_parts.append(_html_table(['数据库模式','恢复模式','保护模式','Standby'], dg_rows) if dg_rows else '<p>非Data Guard环境或无数据</p>')
    ad_rows = check_results.get('archive_dest', [])
    dg_parts.append('<h3>归档目的地</h3>' + _html_table(['Dest_ID','状态','目的地','Archiver','传输模式'], ad_rows) if ad_rows else '')
    dataguard_section = _html_section('🛡  Data Guard', '<br>'.join(dg_parts), 'dataguard')

    # ── RAC ─────────────────────────────────────────────────────────────────
    rac_parts = []
    inst_rows = check_results.get('rac_instances', [])
    rac_parts.append('<h3>实例列表</h3>' +
        _html_table(['Inst#','实例名','主机','状态','并行'], inst_rows) if inst_rows else '<p>非RAC环境或无数据</p>')
    crs_rows = check_results.get('crs_resources', [])
    rac_parts.append('<h3>CRS 资源</h3>' +
        _html_table(['名称','类型','状态','目标','主机'], crs_rows) if crs_rows else '')
    asm_rows = check_results.get('asm_diskgroups', [])
    rac_parts.append('<h3>ASM 磁盘组</h3>' +
        _html_table(['组号','名称','状态','类型','总MB','空闲MB','空闲%'], asm_rows) if asm_rows else '')
    rac_section = _html_section('🌐  RAC + ASM', '<br>'.join(rac_parts), 'rac')

    # ── AWR ─────────────────────────────────────────────────────────────────
    awr_parts = []
    snap_rows = check_results.get('awr_snaps', [])
    awr_parts.append('<h3>AWR 快照（近7天）</h3>' +
        _html_table(['Inst#','Snap ID','开始时间','结束时间','耗时HR','错误'], snap_rows) if snap_rows else '<p>无快照数据</p>')
    awr_section = _html_section('📊  AWR 快照', '<br>'.join(awr_parts), 'awr')

    # ── Jobs ─────────────────────────────────────────────────────────────────
    job_parts = []
    sj_rows = check_results.get('scheduler_jobs', [])
    job_parts.append('<h3>Scheduler Jobs</h3>' +
        _html_table(['作业名','状态','启用','调度类型','下次运行','运行次数','失败次数'], sj_rows) if sj_rows else '')
    fj_rows = check_results.get('failed_jobs', [])
    job_parts.append('<h3>失败的后台作业</h3>' +
        _html_table(['Job#','内容','上次运行','下次运行','失败次数','Broken'], fj_rows) if fj_rows else '')
    job_section = _html_section('⏰  作业调度', '<br>'.join(job_parts), 'jobs')

    # ── Alert ────────────────────────────────────────────────────────────────
    alert_rows = check_results.get('alert_errors', [])
    alert_section = _html_section('🚨  Alert 日志错误（近7天）',
        _html_table(['时间','错误信息（截取200字符）'], alert_rows) if alert_rows else '<p>近7天无Error级别Alert日志</p>',
        'alert')

    # ── AI 诊断 ─────────────────────────────────────────────────────────────
    ai_section = ''
    if ai_advice:
        ai_lines = []
        for line in ai_advice.split('\n'):
            if line.startswith('# '):
                ai_lines.append(f'<h3>{line[2:]}</h3>')
            elif line.startswith('- ') or line.startswith('* '):
                ai_lines.append(f'<li>{line[2:]}</li>')
            elif re.match(r'^\d+\.', line):
                ai_lines.append(f'<li>{line}</li>')
            elif line.strip():
                ai_lines.append(f'<p>{line}</p>')
        if ai_lines:
            ai_section = _html_section('🤖  AI 诊断建议', '<br>'.join(ai_lines), 'ai_diagnosis')
    else:
        ai_section = _html_section('🤖  AI 诊断建议', '<p style="color:#888;">AI 诊断未启用或无可用建议。请在 ai_config.json 中配置 Ollama 后重新巡检以获取 AI 诊断。</p>', 'ai_diagnosis')

    # ── 组合 ─────────────────────────────────────────────────────────────────
    body = (
        f'<h1>DBCheck Oracle 全面巡检报告 | {ver} | {now}</h1>'
        + (f'<p style="text-align:center;color:#0066cc;">巡检人: {inspector}</p>' if inspector else '')
        + nav
        + cards_html
        + os_section
        + db_section
        + ver_section
        + ts_section
        + redo_section
        + cf_section
        + perf_section
        + top_sql_section
        + io_section
        + user_section
        + backup_section
        + flashback_section
        + dataguard_section
        + rac_section
        + awr_section
        + job_section
        + alert_section
        + ai_section
        + f'<div class="footer">DBCheck Oracle 巡检工具 {VER} | 报告生成时间 {now}</div>'
    )

    return f"""<!DOCTYPE html>
<html lang="zh">
<head><meta charset="utf-8"><title>DBCheck Oracle 巡检报告</title>{css}</head>
<body>{body}</body>
</html>"""


# ═══════════════════════════════════════════════════════════════════════════
#                    报告生成（Word）
# ═══════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _docx_table(doc, headers, rows, header_bg='336699'):
    """生成 Word 表格（带表头背景色）"""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表头
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            _set_cell_bg(hdr_cells[i], header_bg)
        except Exception:
            pass
    # 数据行（容错：列数不足时填充空白，多于 header 时截断）
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci in range(len(headers)):
            cell_val = row[ci] if ci < len(row) else None
            if cell_val is not None:
                cells[ci].text = str(cell_val)
    return tbl


def build_word_report(db_info, os_data, check_results, db_version, ai_advice='', inspector=''):
    """构建完整 Word 巡检报告（纯 python-docx，无模板依赖）"""
    if not _HAS_DOCX:
        return None

    doc = Document()

    # ── 页面设置 ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 封面标题 ────────────────────────────────────────────────────────────
    title_p = doc.add_heading('DBCheck Oracle 全面巡检报告', level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0, 102, 204)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run(f'版本: {db_version}').font.color.rgb = RGBColor(102, 102, 102)
    sub_p.add_run(f'\n巡检时间: {time.strftime("%Y-%m-%d %H:%M:%S")}').font.color.rgb = RGBColor(102, 102, 102)
    if inspector:
        sub_p.add_run(f'\n巡检人: {inspector}').font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()  # 空行

    # ── 汇总信息卡片 ────────────────────────────────────────────────────────
    inst_rows = check_results.get('实例信息', {}).get('instance', [])
    hostname = os_data.get('hostname', 'N/A')
    ver = db_version
    uptime = os_data.get('uptime', 'N/A')
    cpu = f"{os_data.get('cpu_model','')} × {os_data.get('cpu_count','?')}"
    mem = f"{os_data.get('mem_total_mb','?')} MB，使用率 {os_data.get('mem_usage_pct','?')}%"

    summary_data = [
        ('主机名', hostname), ('数据库版本', ver), ('运行时间', uptime),
        ('CPU', cpu), ('内存', mem),
    ]
    tbl = _docx_table(doc, ['项目', '内容'], summary_data, '0066CC')
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(10)

    doc.add_paragraph()

    # ── OS 信息 ─────────────────────────────────────────────────────────────
    def _add_section(title):
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)   # 深蓝，一级标题
            run.font.size = Pt(14)

    def _add_subsection(title):
        """二级子标题（同一章节内多个表格时使用）"""
        h = doc.add_heading(title, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色，二级标题
            run.font.size = Pt(12)

    def _add_kv_table(data, cols=2):
        rows = list(data)
        headers = ['项目', '内容'] if cols == 2 else ['项目', '内容', '备注']
        _docx_table(doc, headers, rows)
        doc.add_paragraph()

    _add_section('一、OS 主机信息')
    os_kv = [(k, os_data.get(k, 'N/A')) for k in [
        'hostname', 'os_version', 'kernel', 'uptime', 'cpu_model', 'cpu_count',
        'mem_total_mb', 'mem_used_mb', 'mem_usage_pct', 'swap_total_mb', 'swap_used_mb',
        'load_average', 'hugepages', 'thp'
    ]]
    _add_kv_table(os_kv)

    # ── 数据库基本信息 ──────────────────────────────────────────────────────
    _add_section('二、数据库基本信息')
    db_kv = []
    for k, label in [
        ('NAME', '数据库名'), ('DATABASE_ROLE', '角色'), ('OPEN_MODE', '打开模式'),
        ('LOG_MODE', '归档模式'), ('CDB', 'CDB'), ('flashback_on', '闪回'),
        ('force_logging', 'Force Logging'), ('block_size', '块大小'),
        ('sga_max_size', 'SGA Max'), ('sga_target', 'SGA Target'),
        ('pga_aggregate_target', 'PGA Target'), ('spfile', 'SPFILE'),
    ]:
        v = db_info.get(k, 'N/A')
        db_kv.append((label, str(v)))
    _add_kv_table(db_kv)

    # ── 实例信息 ────────────────────────────────────────────────────────────
    inst = check_results.get('实例信息', {})
    if inst.get('instance'):
        _add_section('三、实例信息')
        _docx_table(doc,
            ['INST_ID', '实例名', '主机', '版本', '启动时间', '状态', '并行', '日志模式', '角色', '打开模式'],
            inst['instance'])
        doc.add_paragraph()

    # ── 版本/补丁 ───────────────────────────────────────────────────────────
    vp = check_results.get('版本/补丁', {})
    if vp.get('version'):
        _add_section('四、数据库版本与补丁')
        _docx_table(doc, ['版本信息'], [[r[0]] for r in vp['version']])
        doc.add_paragraph()

    # ── 表空间 ──────────────────────────────────────────────────────────────
    ts = check_results.get('表空间', {})
    if ts.get('data_tablespaces') or ts.get('temp_tablespaces'):
        _add_section('五、表空间')
    if ts.get('data_tablespaces'):
        _add_subsection('5.1 永久表空间')
        _docx_table(doc,
            ['表空间名', '状态', '类型', '日志', '最大MB', '当前MB', '已用MB', '使用率%'],
            ts['data_tablespaces'])
        doc.add_paragraph()
    if ts.get('temp_tablespaces'):
        _add_subsection('5.2 临时表空间')
        _docx_table(doc,
            ['表空间名', '状态', '最大MB', '当前MB', '已用MB', '使用率%'],
            ts['temp_tablespaces'])
        doc.add_paragraph()

    # ── Redo 日志 ───────────────────────────────────────────────────────────
    redo = check_results.get('Redo日志', {})
    if redo.get('logs'):
        _add_section('六、Redo 日志')
        _docx_table(doc,
            ['Group#', 'Thread#', 'Sequence#', '大小MB', '状态', '成员数', '已归档'],
            redo['logs'])
        doc.add_paragraph()

    # ── 控制文件 ────────────────────────────────────────────────────────────
    cf = check_results.get('控制文件', {})
    if cf.get('controlfiles'):
        _add_section('七、控制文件')
        _docx_table(doc,
            ['名称', '状态', '恢复目录', '块大小', '文件块数'],
            cf['controlfiles'])
        doc.add_paragraph()

    # ── SGA/PGA 内存 ────────────────────────────────────────────────────────
    sga = check_results.get('SGA/PGA内存', {})
    if sga.get('sga_components') or sga.get('pga_stats') or sga.get('memory_params'):
        _add_section('八、SGA/PGA 内存')
    if sga.get('sga_components'):
        _add_subsection('8.1 SGA 动态组件')
        _docx_table(doc,
            ['组件名', '当前MB', '最小MB', '用户指定MB'],
            sga['sga_components'])
        doc.add_paragraph()
    if sga.get('sga_total'):
        _docx_table(doc, ['SGA 总计（MB）'], [[r[0]] for r in sga['sga_total']])
        doc.add_paragraph()
    if sga.get('pga_stats'):
        _add_subsection('8.2 PGA 统计')
        _docx_table(doc, ['指标名', '值（MB）'], sga['pga_stats'])
        doc.add_paragraph()
    if sga.get('memory_params'):
        _add_subsection('8.3 内存参数')
        _docx_table(doc, ['参数名', 'VALUE', '显示值', '是否默认'], sga['memory_params'])
        doc.add_paragraph()

    # ── 关键参数 ───────────────────────────────────────────────────────────
    params = check_results.get('关键参数', {})
    if params.get('params'):
        _add_section('九、关键参数')
        _docx_table(doc,
            ['参数名', 'VALUE', '显示值', '是否默认', '会话可改', '系统可改', '描述'],
            params['params'])
        doc.add_paragraph()

    # ── Undo 信息 ──────────────────────────────────────────────────────────
    undo = check_results.get('Undo信息', {})
    if undo.get('undo_info') or undo.get('undo_segments'):
        _add_section('十、Undo 信息')
    if undo.get('undo_info'):
        _docx_table(doc,
            ['Undo表空间', '保留策略', '已用MB', '总大小MB', '已提交块', '未提交块', '总块数'],
            undo['undo_info'])
        doc.add_paragraph()
    if undo.get('undo_segments'):
        _docx_table(doc,
            ['状态', '段数量', '总大小MB'],
            undo['undo_segments'])
        doc.add_paragraph()

    # ── 长SQL ───────────────────────────────────────────────────────────────
    long_sql = check_results.get('长SQL', {})
    if long_sql.get('long_sql'):
        _add_section('十一、长时间运行的 SQL')
        _docx_table(doc,
            ['SID', 'SERIAL#', '用户名', 'SQL_ID', '操作', '已完成', '总工作量', '完成%', '已耗时秒', '剩余秒'],
            long_sql['long_sql'])
        doc.add_paragraph()

    # ── 性能指标 ────────────────────────────────────────────────────────────
    perf = check_results.get('性能指标', {})
    if perf.get('session_by_status') or perf.get('wait_events'):
        _add_section('十二、性能指标')
    if perf.get('session_by_status'):
        _add_subsection('12.1 会话状态统计')
        _docx_table(doc, ['状态', '数量'], perf['session_by_status'])
        doc.add_paragraph()
    if perf.get('wait_events'):
        _add_subsection('12.2 等待事件统计')
        _docx_table(doc,
            ['事件', '总等待次数', '等待时间(秒)', '等待占比%', '类别'],
            perf['wait_events'])
        doc.add_paragraph()

    # ── Top SQL ─────────────────────────────────────────────────────────────
    top_sql = check_results.get('Top SQL', {})
    if top_sql.get('top_sql_buffer_gets'):
        _add_section('十三、Top SQL（按 Buffer Gets）')
        _docx_table(doc,
            ['SQL_ID', 'SQL片段', 'Buf MB', 'Disk MB', '执行次数', '耗时秒', 'Gets/执行', '模块'],
            top_sql['top_sql_buffer_gets'])
        doc.add_paragraph()

    # ── 无效对象 ────────────────────────────────────────────────────────────
    io = check_results.get('无效对象', {})
    if io.get('invalid_by_type'):
        _add_section('十四、无效对象')
        _docx_table(doc, ['所有者', '类型', '数量'], io['invalid_by_type'])
        doc.add_paragraph()

    # ── 用户安全 ────────────────────────────────────────────────────────────
    users = check_results.get('用户安全', {})
    if users.get('default_accounts') or users.get('locked_users'):
        _add_section('十五、用户安全')
    if users.get('default_accounts'):
        _add_subsection('15.1 默认账户状态')
        _docx_table(doc,
            ['用户名', '状态', '锁定日期', '到期日期', '创建时间'],
            users['default_accounts'])
        doc.add_paragraph()
    if users.get('locked_users'):
        _add_subsection('15.2 锁定账户')
        _docx_table(doc,
            ['用户名', '状态', '锁定日期', '到期日期'],
            users['locked_users'])
        doc.add_paragraph()

    # ── 备份信息 ────────────────────────────────────────────────────────────
    backup = check_results.get('备份信息', {})
    if backup.get('rman_jobs'):
        _add_section('十六、备份信息（RMAN）')
        _docx_table(doc,
            ['会话KEY', '类型', '状态', '开始时间', '结束时间', '大小GB', '耗时分钟'],
            backup['rman_jobs'])
        doc.add_paragraph()

    # ── 闪回/回收站 ─────────────────────────────────────────────────────────
    fb = check_results.get('闪回/回收站', {})
    if fb.get('flashback'):
        _add_section('十七、闪回与回收站')
        _docx_table(doc,
            ['闪回开关', '保留目标(分)', '最旧闪回时间', '保留分钟'],
            fb['flashback'])
        doc.add_paragraph()

    # ── Data Guard ──────────────────────────────────────────────────────────
    dg = check_results.get('Data Guard', {})
    if dg.get('dg_status'):
        _add_section('十八、Data Guard')
        _docx_table(doc,
            ['数据库模式', '恢复模式', '保护模式', 'Standby'],
            dg['dg_status'])
        doc.add_paragraph()

    # ── RAC ─────────────────────────────────────────────────────────────────
    rac = check_results.get('RAC+ASM', {})
    if rac.get('rac_instances'):
        _add_section('十九、RAC + ASM')
        _docx_table(doc,
            ['Inst#', '实例名', '主机', '状态', '并行'],
            rac['rac_instances'])
        doc.add_paragraph()

    # ── AWR ─────────────────────────────────────────────────────────────────
    awr = check_results.get('AWR快照', {})
    if awr.get('awr_snaps'):
        _add_section('二十、AWR 快照（近7天）')
        _docx_table(doc,
            ['Inst#', 'Snap ID', '开始时间', '结束时间', '耗时HR', '错误'],
            awr['awr_snaps'])
        doc.add_paragraph()

    # ── 作业调度 ────────────────────────────────────────────────────────────
    jobs = check_results.get('作业调度', {})
    if jobs.get('scheduler_jobs'):
        _add_section('二十一、作业调度')
        _docx_table(doc,
            ['作业名', '状态', '启用', '调度类型', '下次运行', '运行次数', '失败次数'],
            jobs['scheduler_jobs'])
        doc.add_paragraph()

    # ── Alert 日志 ─────────────────────────────────────────────────────────
    alert = check_results.get('Alert日志', {})
    if alert.get('alert_errors'):
        _add_section('二十二、Alert 日志错误（近7天）')
        _docx_table(doc,
            ['时间', '错误信息（截取200字符）'],
            alert['alert_errors'])
        doc.add_paragraph()

    # ── 风险与建议 ─────────────────────────────────────────────────────────
    _add_section('二十三、风险与建议')
    # 收集各章节中的问题，构建风险项列表
    risk_items = []
    # 从表空间数据中提取高使用率风险
    ts = check_results.get('表空间', {})
    for row in ts.get('data_tablespaces', []):
        if len(row) >= 8:
            try:
                used_pct = float(row[7]) if row[7] != '-' else 0
                if used_pct > 90:
                    risk_items.append({
                        'col1': f'表空间 {row[0]}', 'col2': '高风险',
                        'col3': f'表空间使用率 {used_pct:.1f}%，超过 90% 告警线',
                        'col4': '高', 'col5': 'DBA', 'fix_sql': '-- 查询表空间使用情况'
                    })
                elif used_pct > 80:
                    risk_items.append({
                        'col1': f'表空间 {row[0]}', 'col2': '中风险',
                        'col3': f'表空间使用率 {used_pct:.1f}%，超过 80% 建议关注',
                        'col4': '中', 'col5': 'DBA', 'fix_sql': ''
                    })
            except (ValueError, TypeError):
                pass
    # 从无效对象数据中提取风险
    io = check_results.get('无效对象', {})
    for row in io.get('invalid_by_type', []):
        if len(row) >= 3:
            cnt = row[2] if isinstance(row[2], int) else 0
            if cnt > 0:
                risk_items.append({
                    'col1': f'无效{row[1]}对象', 'col2': '中风险',
                    'col3': f'存在 {cnt} 个无效 {row[1]} 对象',
                    'col4': '中', 'col5': 'DBA', 'fix_sql': f'SELECT * FROM {row[0]}.{row[1]} WHERE status=\'INVALID\';'
                })
    # 从锁定用户中提取风险
    users = check_results.get('用户安全', {})
    locked = users.get('locked_users', [])
    if locked:
        risk_items.append({
            'col1': '账户锁定', 'col2': '中风险',
            'col3': f'存在 {len(locked)} 个被锁定的账户',
            'col4': '中', 'col5': 'DBA', 'fix_sql': '-- 查看锁定账户: SELECT username, lock_date FROM dba_users WHERE account_status LIKE \'%LOCKED%\';'
        })
    # 从Alert日志错误中提取风险
    alert = check_results.get('Alert日志', {})
    err_count = len(alert.get('alert_errors', []))
    if err_count > 0:
        risk_items.append({
            'col1': 'Alert日志错误', 'col2': '高风险',
            'col3': f'近7天存在 {err_count} 条Error级别Alert日志',
            'col4': '高', 'col5': 'DBA', 'fix_sql': '-- 检查Alert日志具体内容'
        })

    if risk_items:
        # 23.1 问题明细
        _add_subsection('23.1 问题明细')
        risk_table_data = [(str(i+1), x['col1'], x['col2'], x['col3'], x['col4'], x['col5']) for i, x in enumerate(risk_items)]
        _docx_table(doc, ['序号', '项目', '风险等级', '问题描述', '严重程度', '责任人'], risk_table_data)
        doc.add_paragraph()
        # 23.2 修复SQL速查
        fix_sqls = [(x['col1'], x['fix_sql']) for x in risk_items if x['fix_sql']]
        if fix_sqls:
            _add_subsection('23.2 修复SQL速查')
            for fname, sql in fix_sqls:
                p = doc.add_paragraph()
                p.add_run(f'【{fname}】').bold = True
                doc.add_paragraph(sql, style='List Bullet')
            doc.add_paragraph()
    else:
        # 即使无高风险，也汇总各项巡检结论
        _add_subsection('23.1 巡检结论')
        summary_items = []
        ts = check_results.get('表空间', {})
        if ts.get('data_tablespaces'):
            total_ts = len(ts['data_tablespaces'])
            high_ts = sum(1 for r in ts['data_tablespaces'] if len(r) >= 8 and str(r[7]).replace('.','').isdigit() and float(r[7]) > 80)
            summary_items.append(f'表空间共 {total_ts} 个，其中使用率超过 80% 的有 {high_ts} 个')
        perf = check_results.get('性能指标', {})
        if perf.get('wait_events'):
            top_wait = perf['wait_events'][0][0] if perf['wait_events'] else 'N/A'
            summary_items.append(f'等待事件最突出项: {top_wait}')
        io = check_results.get('无效对象', {})
        if io.get('invalid_by_type'):
            total_inv = sum(int(r[2]) for r in io['invalid_by_type'] if len(r) >= 3 and str(r[2]).isdigit())
            summary_items.append(f'无效对象共 {total_inv} 个')
        if summary_items:
            for item in summary_items:
                p = doc.add_paragraph(item, style='List Bullet')
        else:
            p = doc.add_paragraph("未发现明显风险项，数据库整体运行状况良好")
            for run in p.runs:
                run.font.size = Pt(10.5)

    # ── 第24章 AI 诊断 ─────────────────────────────────────────────────────
    _add_section('二十四、AI 诊断建议')
    if ai_advice:
        for line in ai_advice.split('\n'):
            if line.strip():
                if line.startswith('# '):
                    h = doc.add_heading(line[2:], level=2)
                    for run in h.runs:
                        run.font.color.rgb = RGBColor(0, 102, 204)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    # 避免与 Word 自动编号冲突，直接写纯文本段落
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(10.5)
                else:
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(10.5)
        doc.add_paragraph()
    else:
        p = doc.add_paragraph("AI 诊断未启用或无可用建议。请在 ai_config.json 中配置 Ollama 后重新巡检以获取 AI 诊断。")
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()

    # ── 第25章 报告说明 ────────────────────────────────────────────────────
    _add_section('二十五、报告说明')
    notes = [
        "本报告基于 Oracle 数据库实时状态生成，反映了生成时刻的数据库健康状况。",
        "报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用。",
        "磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令 'df -h'。",
        "巡检结果仅供参考，实际运维中请结合具体业务场景进行分析。",
        "建议定期进行数据库巡检，及时发现并解决潜在问题。",
        "AI 诊断功能（若启用）生成的建议仅供参考，不构成专业 DBA 意见。"
    ]
    for text in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    doc.add_paragraph()

    # ── 页脚 ────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f'DBCheck Oracle 巡检工具 {VER} | 报告生成时间 {time.strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


# ═══════════════════════════════════════════════════════════════════════════
#                    主入口
# ═══════════════════════════════════════════════════════════════════════════

def print_banner():
    art = f"""
{CYAN}{BOLD}  ██████╗ ██████╗  ██████╗██╗  ██╗███████╗ ██████╗██╗  ██╗
  ██╔══██╗██╔══██╗██╔════╝██║  ██║██╔════╝██╔════╝██║ ██╔╝
  ██║  ██║██████╔╝██║     ███████║█████╗  ██║     █████╔╝
  ██║  ██║██╔══██╗██║     ██╔══██║██╔══╝  ██║     ██╔═██╗
  ██████╔╝██████╔╝╚██████╗██║  ██║███████╗╚██████╗██║  ██╗
  ╚═════╝ ╚═════╝  ╚═════╝╚═╝  ╚═╝╚══════╝ ╚═════╝╚═╝  ╚═╝{RESET}
{BOLD}      🗄️  DBCheck Oracle 全面巡检工具  {VER}{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
{GREEN}  基于 OS 层 + 数据库层{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
"""
    print(art)


def single_inspection(args):
    """单机巡检主流程"""
    import paramiko

    print(f"\n{GREEN}▶ 开始 Oracle 全面巡检{RESET}")
    t0 = time.time()

    # ── 1. Oracle 连接 ─────────────────────────────────────────────────────
    print(f"\n[{GREEN}1/6{RESET}] 连接 Oracle 数据库...")
    try:
        if args.servicename:
            dsn = oracledb.makedsn(args.host, args.port, service_name=args.servicename)
        else:
            dsn = oracledb.makedsn(args.host, args.port, args.sid)
        # sys 用户默认以 SYSDBA 身份连接（oracle privilege model）
        if args.user.upper() == 'SYS' and not args.sysdba:
            args.sysdba = True
        if args.sysdba:
            conn = oracledb.connect(user=args.user, password=args.password, dsn=dsn, mode=oracledb.SYSDBA)
        else:
            conn = oracledb.connect(user=args.user, password=args.password, dsn=dsn)
        print(f"  ✅ 连接成功 (mode: {'SYSDBA' if args.sysdba else 'NORMAL'})")
    except Exception as e:
        print(f"  ❌ 连接失败: {e}")
        return

    # ── 2. 获取版本 ───────────────────────────────────────────────────────
    print(f"\n[{GREEN}2/6{RESET}] 获取数据库版本...")
    version_str, ver_major = get_db_version_and_major(conn)
    print(f"  版本: {version_str}  (主版本: {ver_major})")

    # ── 3. OS 层采集 ──────────────────────────────────────────────────────
    print(f"\n[{GREEN}3/6{RESET}] 采集 OS 层信息...")
    os_data = {}
    ssh_client = None
    if args.ssh_host:
        try:
            ssh_client = paramiko.SSHClient()
            ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            ssh_client.connect(
                args.ssh_host, port=args.ssh_port or 22,
                username=args.ssh_user, password=args.ssh_pass,
                timeout=15
            )
            collector = OSCollector(ssh_client)
            os_data = collector.collect()
            print(f"  ✅ SSH 采集成功 (host: {args.ssh_host})")
        except Exception as e:
            print(f"  ⚠ SSH 采集失败: {e}，使用本地采集")
            collector = OSCollector(None)
            os_data = collector.collect()
    else:
        collector = OSCollector(None)
        os_data = collector.collect()
        print(f"  ✅ 本地采集完成")

    # ── 4. 数据库层巡检（版本自适应）──────────────────────────────────────
    print(f"\n[{GREEN}4/6{RESET}] 执行数据库层巡检（Oracle {ver_major}c）...")
    check_results = {}

    # 根据版本号动态选择检查列表
    checks = get_checks_for_version(ver_major)

    for name, fn in checks:
        try:
            result = fn(conn)
            if result and 'error' not in result:
                check_results[name] = result
                rows = list(result.values())[0] if result else []
                cnt = len(rows) if isinstance(rows, list) else '-'
                print(f"  ✅ {name}  ({cnt} 条)")
            elif result and 'error' in result:
                # 有 error 键，说明查询执行了但失败了，打印具体错误
                print(f"  ⚠ {name}  查询失败: {result.get('error', '未知错误')}")
            else:
                print(f"  ⚠ {name}  无数据")
        except Exception as e:
            print(f"  ⚠ {name}  跳过: {e}")

    conn.close()

    # ── 4.5 AI 诊断（根据配置判断是否启用）───────────────────────────────────
    print(f"\n[{GREEN}4.5/6{RESET}] AI 诊断...")
    ai_advice = ''
    try:
        from analyzer import AIAdvisor
        cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
        ai_cfg = {}
        if os.path.exists(cfg_path):
            with open(cfg_path, 'r', encoding='utf-8') as f:
                ai_cfg = json.load(f)
        advisor = AIAdvisor(
            backend=ai_cfg.get('backend'),
            api_key=ai_cfg.get('api_key'),
            api_url=ai_cfg.get('api_url'),
            model=ai_cfg.get('model')
        )
        if advisor.enabled:
            # 从 check_results 直接构建 db_info（避免引用报告生成阶段才定义的变量）
            _db_info = {}
            for row in check_results.get('数据库信息', {}).get('database', []):
                cols = ['DBID','NAME','DATABASE_ROLE','OPEN_MODE','LOG_MODE',
                        'CREATED','STARTUP_TIME','CDB','flashback_on','force_logging',
                        'block_size','sga_max_size','sga_target','pga_aggregate_target',
                        'spfile','global_name']
                for i, c in enumerate(cols):
                    if i < len(row):
                        _db_info[c] = row[i]
            for row in check_results.get('实例信息', {}).get('instance', []):
                if len(row) >= 4:
                    _db_info.setdefault('INSTANCE_NAME', row[1])
                    _db_info.setdefault('HOST_NAME', row[2])
                    _db_info.setdefault('VERSION', row[3])
            label = _db_info.get('NAME', args.servicename or args.sid or 'ORACLE')
            print(f"  🤖 正在调用 AI 诊断（{advisor.backend} / {advisor.model}）...")
            # 收集风险项作为上下文
            risk_items = []
            ts = check_results.get('表空间', {})
            for row in ts.get('data_tablespaces', []):
                if len(row) >= 8 and row[7] != '-':
                    try:
                        used_pct = float(row[7])
                        if used_pct > 90:
                            risk_items.append({'col1': f'表空间 {row[0]}', 'col2': '高风险',
                                'col3': f'表空间使用率 {used_pct:.1f}%，超过 90% 告警线'})
                        elif used_pct > 80:
                            risk_items.append({'col1': f'表空间 {row[0]}', 'col2': '中风险',
                                'col3': f'表空间使用率 {used_pct:.1f}%，超过 80% 建议关注'})
                    except (ValueError, TypeError):
                        pass
            # 收集等待事件 Top5
            perf = check_results.get('性能指标', {})
            wait_top5 = perf.get('wait_events', [])[:5]
            wait_summary = '\n'.join([f"- {w[0]}: {w[1]}次等待, {w[2]}秒, 占比{w[3]}%"
                                      for w in wait_top5]) if wait_top5 else 'N/A'
            # 收集阻塞会话
            blocked_sessions = check_results.get('阻塞会话', [])
            blocked_summary = f"共 {len(blocked_sessions)} 个阻塞会话" if blocked_sessions else '无阻塞会话'
            # 收集 Top SQL（按 Buffer Gets 前5）
            top_sql_raw = check_results.get('Top SQL', {})
            top_sql5 = top_sql_raw.get('top_sql_buffer_gets', [])[:5]
            top_sql_summary = '\n'.join([
                f"- SQL_ID={r[0]}: Buf={r[2]}MB, Execs={r[4]},耗时={r[5]}s, 片断={str(r[1])[:60]}"
                for r in top_sql5]) if top_sql5 else 'N/A'
            # 构建详细指标
            metrics = {
                'db_version': _db_info.get('VERSION', version_str),
                'hostname': _db_info.get('HOST_NAME', os_data.get('hostname', 'N/A')),
                'uptime': os_data.get('uptime', 'N/A'),
                'risk_count': len(risk_items),
                'tablespace_count': len(ts.get('data_tablespaces', [])),
                'wait_events_top5': wait_summary,
                'blocked_sessions': blocked_summary,
                'top_sql_top5': top_sql_summary,
            }
            ai_advice = advisor.diagnose('oracle', label, metrics, risk_items, timeout=600)
            if ai_advice:
                print(f"  ✅ AI 诊断完成")
            else:
                print(f"  ⚠ AI 诊断返回空")
        else:
            print(f"  ⏭ AI 诊断未启用（backend: {advisor.backend}）")
    except TimeoutError as e:
        ai_advice = "⚠ AI 诊断超时：大模型加载时间较长，请到「AI 配置」中增加超时时间后重试"
        print(f"  {ai_advice}")
    except Exception as e:
        err_str = str(e)
        if 'connection' in err_str.lower() or 'refused' in err_str.lower():
            ai_advice = "⚠ AI 诊断连接失败：请确认 Ollama 服务已启动"
        else:
            ai_advice = f"⚠ AI 诊断失败：{err_str[:120]}"
        print(f"  {ai_advice}")

    # ── 5. 生成报告 ────────────────────────────────────────────────────────
    print(f"\n[{GREEN}5/6{RESET}] 生成巡检报告...")
    # 从 check_results 提取 db_info
    db_info = {}
    inst_rows = check_results.get('实例信息', {}).get('instance', [])
    db_rows   = check_results.get('数据库信息', {}).get('database', [])
    if db_rows:
        cols = ['DBID','NAME','DATABASE_ROLE','OPEN_MODE','LOG_MODE',
                'CREATED','STARTUP_TIME','CDB','flashback_on','force_logging',
                'block_size','sga_max_size','sga_target','pga_aggregate_target',
                'spfile','global_name']
        for row in db_rows:
            for i, c in enumerate(cols):
                if i < len(row):
                    db_info[c] = row[i]

    if inst_rows and len(inst_rows[0]) >= 4:
        db_info['INSTANCE_NAME'] = inst_rows[0][1]
        db_info['HOST_NAME']     = inst_rows[0][2]
        db_info['VERSION']       = inst_rows[0][3]
        db_info['STARTUP_TIME']  = inst_rows[0][4]
        db_info['STATUS']         = inst_rows[0][5]

    docx = build_word_report(db_info, os_data, check_results, version_str, ai_advice, inspector=args.inspector or 'dbcheck')

    # ── 6. 保存报告 ────────────────────────────────────────────────────────
    print(f"\n[{GREEN}6/6{RESET}] 保存报告...")
    output_dir = args.output or os.path.join(os.getcwd(), 'reports')
    os.makedirs(output_dir, exist_ok=True)

    db_name = db_info.get('NAME', args.servicename or args.sid or 'ORACLE')
    ver_tag  = ver_major or 'DB'
    ts = time.strftime('%Y%m%d%H%M%S')

    # Word
    if docx:
        docx_fname = f"oracle_fullcheck_{db_name}_{ver_tag}_{ts}.docx"
        docx_path  = os.path.join(output_dir, docx_fname)
        try:
            docx.save(docx_path)
            print(f"   Word:  {docx_path}")
        except Exception as e:
            print(f"   Word报告生成失败: {e}")

    # ── 保存历史记录 ──────────────────────────────────────────────────────
    try:
        from analyzer import HistoryManager
        script_dir = os.path.dirname(os.path.abspath(__file__))
        label = db_info.get('NAME', args.servicename or args.sid or 'ORACLE')
        hm = HistoryManager(script_dir)

        # 将 check_results（中文键）映射为 Oracle context（英文键），供 _extract_metrics 使用
        ts = check_results.get('表空间', {})
        perf = check_results.get('性能指标', {})

        def _ts_rows(data_ts):
            """将 check_results 表空间数据转为 context 格式"""
            rows = []
            for row in data_ts:
                if len(row) >= 8:
                    rows.append({
                        'TABLESPACE_NAME': str(row[0]),
                        'STATUS': str(row[1]),
                        'CONTENTS': str(row[2]) if len(row) > 2 else '',
                        'TOTAL_MB': float(row[4]) if row[4] != '-' else 0,
                        'USED_MB': float(row[5]) if row[5] != '-' else 0,
                        'USED_PCT_WITH_MAXEXT': float(row[7]) if row[7] != '-' else 0,
                    })
            return rows

        context = {
            'ora_version': [{'BANNER': version_str}],
            'ora_tablespace': _ts_rows(ts.get('data_tablespaces', [])),
            'ora_sessions': perf.get('session_by_status', []),
            'system_info': {
                'hostname': os_data.get('hostname', ''),
                'cpu': {'usage_percent': os_data.get('cpu_percent', 0)},
                'memory': {'usage_percent': os_data.get('mem_percent', 0)},
                'disk_list': [{'mountpoint': d.get('mount', '/'), 'usage_percent': d.get('percent', 0)}
                              for d in os_data.get('disk_list', [])],
            },
            'health_status': '良好' if not risk_items else ('存在风险' if any(r.get('col2') == '高风险' for r in risk_items) else '一般'),
            'auto_analyze': risk_items if risk_items else [],
        }
        hm.save_snapshot('oracle_full', args.host, args.port, label, context)
        print(f"  ✅ 历史记录已保存")
    except Exception as e:
        print(f"  ⚠ 保存历史记录失败: {e}")

    elapsed = time.time() - t0
    print(f"\n{GREEN}{BOLD}✅ 巡检完成！耗时 {elapsed:.1f} 秒{RESET}")

    if ssh_client:
        ssh_client.close()


def _input(prompt, default=''):
    """统一输入函数，带默认值显示"""
    if default:
        val = input(f"{prompt} [{default}]: ").strip()
        return val if val else default
    return input(f"{prompt}: ").strip()


def interactive_single_inspection():
    """交互式单机巡检（替代 argparse，适合无参数直接运行）"""
    print(f"\n{BOLD}{'='*52}{RESET}")
    print(f"{RED}{BOLD}   Oracle 全面巡检（增强版）{RESET}")
    print(f"{DIM}{'='*52}{RESET}\n")

    # ── Oracle 连接信息 ─────────────────────────────────────────
    host        = _input(f"{CYAN}Oracle 主机IP{RESET}",    '192.168.1.10')
    port        = _input(f"{CYAN}端口{RESET}",             '1521')
    connect_by  = _input(f"{CYAN}连接方式 (S=SID / N=ServiceName){RESET}", 'S').upper()
    if connect_by == 'N':
        sid_or_svc = _input(f"{CYAN}ServiceName{RESET}")
        sid, svc = None, sid_or_svc
    else:
        sid     = _input(f"{CYAN}ORACLE_SID{RESET}",       'ORCL')
        svc     = None
    user        = _input(f"{CYAN}用户名{RESET}",           'sys')
    password    = _input(f"{CYAN}密码{RESET}")
    # sys 用户默认以 SYSDBA 登录，其他用户可自行选择
    if user.upper() == 'SYS':
        sysdba_default = 'Y'
    else:
        sysdba_default = 'N'
    sysdba_opt  = _input(f"{CYAN}是否以 SYSDBA 身份连接 (Y/N){RESET}", sysdba_default).upper()
    sysdba      = (sysdba_opt == 'Y')

    # ── SSH 信息（可选）────────────────────────────────────────
    use_ssh = _input(f"\n{GREEN}是否使用 SSH 采集 OS 信息 (y/N){RESET}", 'N').upper()
    ssh_host, ssh_port, ssh_user, ssh_pass = None, 22, None, None
    if use_ssh == 'Y':
        ssh_host = _input(f"{CYAN}SSH 主机IP{RESET}", host)
        ssh_port = _input(f"{CYAN}SSH 端口{RESET}",   '22')
        ssh_user = _input(f"{CYAN}SSH 用户名{RESET}")
        ssh_pass = _input(f"{CYAN}SSH 密码{RESET}")
        if not ssh_user or not ssh_pass:
            print(f"  {YELLOW}⚠ SSH 用户名和密码不能为空，已跳过 SSH 采集{RESET}")
            ssh_host, ssh_user, ssh_pass = None, None, None

    # ── 输出选项 ───────────────────────────────────────────────
    output_dir = _input(f"\n{GREEN}报告输出目录（默认当前目录 reports）{RESET}", 'reports')
    inspector  = _input(f"{GREEN}巡检人姓名（默认 dbcheck）{RESET}")

    # ── 构造 args ───────────────────────────────────────────────
    class _Args:
        pass
    args = _Args()
    args.host        = host
    args.port        = int(port)
    args.sid         = sid
    args.servicename = svc
    args.user        = user
    args.password    = password
    args.sysdba      = sysdba
    args.ssh_host    = ssh_host
    args.ssh_port    = int(ssh_port) if ssh_port else 22
    args.ssh_user    = ssh_user
    args.ssh_pass    = ssh_pass
    args.output      = output_dir if output_dir else None
    args.inspector   = inspector or 'dbcheck'

    print_banner()
    single_inspection(args)


def main():
    import argparse
    parser = argparse.ArgumentParser(
        description=f'DBCheck Oracle 全面巡检工具 v{VER}（OS层+数据库层）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 本地 Oracle（使用 SID）
  python main_oracle_full.py -h 192.168.1.10 -P 1521 -s ORCL -u system -p xxx

  # 使用 ServiceName 连接
  python main_oracle_full.py -h 192.168.1.10 -P 1521 -S ORCL -u system -p xxx

  # SSH 采集 OS 信息
  python main_oracle_full.py -h 192.168.1.10 -P 1521 -s ORCL -u system -p xxx \\
      --ssh-host 192.168.1.10 --ssh-user oracle --ssh-pass xxx

        """
    )
    parser.add_argument('--host',             required=False, help='Oracle 主机IP（交互模式无需指定）')
    parser.add_argument('-P', '--port',      type=int, default=1521, help='端口（默认1521）')
    parser.add_argument('-s', '--sid',        help='ORACLE_SID（SID和ServiceName二选一）')
    parser.add_argument('-S', '--servicename', help='ServiceName（SID和ServiceName二选一）')
    parser.add_argument('-u', '--user',       default='sys', help='用户名（默认sys）')
    parser.add_argument('-p', '--password',  default='',   help='密码')
    parser.add_argument('--sysdba',       action='store_true', help='以 SYSDBA 身份连接（sys用户默认开启）')
    parser.add_argument('--ssh-host',   help='SSH 主机（与 Oracle 主机相同则省略）')
    parser.add_argument('--ssh-port',   type=int, default=22, help='SSH 端口（默认22）')
    parser.add_argument('--ssh-user',   help='SSH 用户名')
    parser.add_argument('--ssh-pass',   help='SSH 密码')
    parser.add_argument('-o', '--output',   help='报告输出目录')
    parser.add_argument('--inspector',   help='巡检人姓名')

    args = parser.parse_args()

    # 无参数时进入交互模式
    if len(sys.argv) == 1 or (
           not args.host and not args.sid and not args.servicename
    ):
        interactive_single_inspection()
        return

    if not args.sid and not args.servicename:
        print("❌ 必须指定 --sid 或 --servicename")
        return

    if args.ssh_host and not (args.ssh_user and args.ssh_pass):
        print("❌ SSH 采集需要同时指定 --ssh-user 和 --ssh-pass")
        return

    print_banner()
    single_inspection(args)


# ═══════════════════════════════════════════════════════════════════════════
# v19 兼容版覆盖函数（12c 基准不动，19c 出错项单独覆盖）
# ═══════════════════════════════════════════════════════════════════════════

def oracle_check_database_v19(conn):
    """数据库层信息（v19 专用版）：不使用 PLUGGABLE_DB（该列在目标环境不存在）"""
    results = {}
    cur = conn.cursor()
    try:
        # 只取 v$database 中稳定存在的列，不依赖 PLUGGABLE_DB
        cur.execute("""
            SELECT DBID, NAME, DATABASE_ROLE, CREATED, LOG_MODE, OPEN_MODE,
                   FLASHBACK_ON, FORCE_LOGGING, CREATED
            FROM v$database
        """)
        results['database'] = cur.fetchall()

        cur.execute("SELECT global_name FROM global_name")
        results['global_name'] = cur.fetchone()

        try:
            cur.execute("""
                SELECT parameter, value
                FROM nls_database_parameters
                WHERE parameter IN ('NLS_CHARACTERSET', 'NLS_NCHAR_CHARACTERSET')
            """)
            rows = cur.fetchall()
            results['charset'] = tuple(r[1] for r in rows) if rows else ('', '')
        except Exception:
            results['charset'] = ('', '')

        for param in ['db_block_size', 'sga_max_size', 'sga_target',
                      'pga_aggregate_target', 'memory_max_target', 'memory_target']:
            try:
                cur.execute(f"SELECT value FROM v$parameter WHERE name='{param}'")
                r = cur.fetchone()
                results[param] = r[0] if r else ''
            except Exception:
                results[param] = ''

        # ADR 相关路径
        try:
            cur.execute("SELECT value FROM v$parameter WHERE name='diagnostic_dest'")
            r = cur.fetchone()
            results['adr'] = r[0] if r else ''
        except Exception:
            results['adr'] = ''

        try:
            cur.execute("SELECT value FROM v$parameter WHERE name='db_create_file_dest'")
            r = cur.fetchone()
            results['omf'] = r[0] if r else ''
        except Exception:
            results['omf'] = ''

        try:
            cur.execute("SELECT log_mode FROM v$database")
            r = cur.fetchone()
            results['log_mode'] = r[0] if r else ''
        except Exception:
            results['log_mode'] = ''

        try:
            cur.execute("SELECT force_logging FROM v$database")
            r = cur.fetchone()
            results['force_logging'] = r[0] if r else ''
        except Exception:
            results['force_logging'] = ''

        try:
            cur.execute("SELECT flashback_on FROM v$database")
            r = cur.fetchone()
            results['flashback_on'] = r[0] if r else ''
        except Exception:
            results['flashback_on'] = ''

        try:
            cur.execute("SELECT TO_CHAR(CREATED, 'YYYY-MM-DD HH24:MI:SS') FROM v$database")
            r = cur.fetchone()
            results['created'] = r[0] if r else ''
        except Exception:
            results['created'] = ''

        try:
            cur.execute("SELECT TO_CHAR(STARTUP_TIME, 'YYYY-MM-DD HH24:MI:SS') FROM v$instance")
            r = cur.fetchone()
            results['startup_time'] = r[0] if r else ''
        except Exception:
            results['startup_time'] = ''

    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def _col_name(cur, view, pattern):
    """探测视图实际列名（模糊匹配），找不到返回 None"""
    try:
        cur.execute(f"""
            SELECT column_name FROM user_tab_columns
            WHERE table_name = UPPER('{view}')
              AND column_name LIKE '%{pattern}%'
            FETCH FIRST 1 ROWS ONLY
        """)
        r = cur.fetchone()
        return r[0] if r else None
    except Exception:
        return None


def oracle_check_tablespace_v19(conn):
    """表空间（v19 兼容版）：自动探测 dba_temp_free_space/dba_free_space 的实际列名，不再猜"""
    results = {}
    cur = conn.cursor()
    try:
        # ── 自动探测空闲列名 ─────────────────────────────────────────
        # dba_temp_free_space：可能有 FREE_SPACE / TABLESPACE_SIZE / ALLOCATED_SPACE 等
        tfs_col = (_col_name(cur, 'dba_temp_free_space', 'FREE')
                or _col_name(cur, 'dba_temp_free_space', 'SPACE')
                or _col_name(cur, 'dba_temp_free_space', 'SIZE'))

        # dba_free_space：可能有 FREE_SPACE / BYTES / BLOCKS
        fs_col = (_col_name(cur, 'dba_free_space', 'FREE')
              or _col_name(cur, 'dba_free_space', 'BYTES')
              or _col_name(cur, 'dba_free_space', 'BLOCKS'))

        # 永久表空间
        cur.execute("""
            SELECT bt.tablespace_name,
                   bt.status,
                   ROUND(NVL(df.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(df.max_mb,0), 2) max_mb,
                   ROUND(NVL(seg.used_mb,0), 2) used_mb,
                   ROUND(NVL(df.curr_mb,0) - NVL(seg.used_mb,0), 2) free_mb,
                   ROUND(NVL(seg.used_mb,0) / NULLIF(NVL(df.curr_mb,0),0) * 100, 2) pct_used
            FROM dba_tablespaces bt
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_data_files GROUP BY tablespace_name) df
                   ON bt.tablespace_name = df.tablespace_name
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) used_mb
                       FROM dba_segments GROUP BY tablespace_name) seg
                   ON bt.tablespace_name = seg.tablespace_name
            WHERE bt.contents = 'PERMANENT'
            ORDER BY pct_used DESC NULLS LAST
        """)
        results['data_tablespaces'] = cur.fetchall()

        # 临时表空间：只查 dba_temp_files，不依赖 dba_temp_free_space（列名不稳定）
        cur.execute("""
            SELECT bt.tablespace_name,
                   bt.status,
                   ROUND(NVL(tf.curr_mb,0), 2) curr_mb,
                   ROUND(NVL(tf.max_mb,0), 2) max_mb,
                   '-' used_mb,
                   '-' free_mb,
                   '-' pct_used
            FROM dba_tablespaces bt
            LEFT JOIN (SELECT tablespace_name,
                              SUM(bytes/1024/1024) curr_mb,
                              SUM(MAXBYTES/1024/1024) max_mb
                       FROM dba_temp_files GROUP BY tablespace_name) tf
                   ON bt.tablespace_name = tf.tablespace_name
            WHERE bt.contents = 'TEMPORARY'
            ORDER BY bt.tablespace_name
        """)
        results['temp_tablespaces'] = cur.fetchall()

        # 自动扩展文件（dba_data_files 有 BYTES/MAXBYTES 是确定的）
        cur.execute("""
            SELECT tablespace_name, file_name,
                   ROUND(bytes/1024/1024,2) curr_mb,
                   ROUND(MAXBYTES/1024/1024,2) max_mb,
                   AUTOEXTENSIBLE
            FROM dba_data_files
            WHERE AUTOEXTENSIBLE = 'YES'
            ORDER BY tablespace_name
        """)
        results['autoextend_files'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_redolog_v19(conn):
    """Redo 日志（v19 兼容版）：不使用 v$loghist（列不稳定），直接查 v$log"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT GROUP#, THREAD#, SEQUENCE#, ROUND(BYTES/1024/1024,2) size_mb,
                   STATUS, MEMBERS, ARCHIVED
            FROM v$log
            ORDER BY THREAD#, GROUP#
        """)
        results['logs'] = cur.fetchall()

        cur.execute("""
            SELECT GROUP#, MEMBER, TYPE, STATUS
            FROM v$logfile
            ORDER BY GROUP#
        """)
        results['logfiles'] = cur.fetchall()

        # 直接查 v$log，统计 CURRENT 组大小（不依赖 v$loghist）
        cur.execute("""
            SELECT THREAD#,
                   COUNT(*) switch_cnt,
                   ROUND(SUM(BYTES)/1024/1024/1024, 2) total_mb
            FROM v$log
            WHERE STATUS = 'CURRENT'
            GROUP BY THREAD#
        """)
        results['redo_switch'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_top_sql_v19(conn, limit=20):
    """Top SQL（v19 兼容版）：去除中文字段别名，避免字符集解析问题"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id,
                       SUBSTR(sql_text,1,80) AS sql_text,
                       ROUND(buffer_gets/1024/1024,2) AS buf_mb,
                       ROUND(disk_reads/1024/1024,2) AS disk_mb,
                       executions,
                       ROUND(elapsed_time/1000000,2) AS elapsed_sec,
                       ROUND(buffer_gets/DECODE(executions,0,1,executions)) AS gets_per_exec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY buffer_gets DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_buffer_gets'] = cur.fetchall()

        cur.execute(f"""
            SELECT * FROM (
                SELECT sql_id,
                       SUBSTR(sql_text,1,80) AS sql_text,
                       ROUND(disk_reads/1024/1024,2) AS disk_mb,
                       executions,
                       ROUND(elapsed_time/1000000,2) AS elapsed_sec,
                       module
                FROM v$sql
                WHERE executions > 0
                ORDER BY disk_reads DESC
            ) WHERE ROWNUM <= {limit}
        """)
        results['top_sql_disk_reads'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_backup_v19(conn):
    """RMAN 备份信息（v19 兼容版）：v$rman_backup_job_details 无 bytes 列，用 TIME_TAKEN_DISPLAY"""
    results = {}
    cur = conn.cursor()
    try:
        # v$rman_backup_job_details 标准列（19c 通用）：SESSION_KEY/INPUT_TYPE/STATUS/START_TIME/END_TIME/TIME_TAKEN_DISPLAY
        cur.execute("""
            SELECT SESSION_KEY, INPUT_TYPE, STATUS,
                   TO_CHAR(START_TIME,'YYYY-MM-DD HH24:MI') start_t,
                   TO_CHAR(END_TIME,'YYYY-MM-DD HH24:MI') end_t,
                   TIME_TAKEN_DISPLAY AS elapsed_disp
            FROM v$rman_backup_job_details
            WHERE end_time > SYSDATE - 30
            ORDER BY end_time DESC
        """)
        results['rman_jobs'] = cur.fetchall()

        # 备份集大小从 v$backup_piece（而非 v$backup_set）取
        try:
            cur.execute("""
                SELECT p.handle, s.INPUT_TYPE,
                       ROUND(SUM(p.bytes)/1024/1024/1024, 2) size_gb,
                       MAX(p.compressed) compressed
                FROM v$backup_set s, v$backup_piece p
                WHERE p.set_stamp = s.set_stamp
                  AND p.set_count = s.set_count
                  AND p.completion_time > SYSDATE - 30
                GROUP BY p.handle, s.INPUT_TYPE
                ORDER BY MAX(p.completion_time) DESC
            """)
            results['backup_pieces'] = cur.fetchall()
        except Exception:
            results['backup_pieces'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_flashback_v19(conn):
    """闪回配置（v19 兼容版）：不用 v$flashback_database_stat，直接查 v$database 稳定列"""
    results = {}
    cur = conn.cursor()
    try:
        # v$database 标准列：FLASHBACK_ON / CREATED / LOG_MODE
        cur.execute("""
            SELECT FLASHBACK_ON,
                   TO_CHAR(CREATED,'YYYY-MM-DD HH24:MI') created,
                   LOG_MODE
            FROM v$database
        """)
        results['flashback'] = cur.fetchall()

        try:
            cur.execute("""
                SELECT owner, original_name, type,
                       ROUND(space * (SELECT TO_NUMBER(value) FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                       can_undrop, can_purge
                FROM dba_recyclebin
                ORDER BY mb DESC
            """)
            results['recyclebin'] = cur.fetchall()
        except Exception:
            try:
                cur.execute("""
                    SELECT owner, original_name, type,
                           ROUND(space * (SELECT TO_NUMBER(value) FROM v$parameter WHERE name='db_block_size')/1024/1024,2) mb,
                           can_undrop, can_purge
                    FROM cdb_recyclebin
                    ORDER BY mb DESC
                """)
                results['recyclebin'] = cur.fetchall()
            except Exception:
                results['recyclebin'] = []
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_dataguard_v19(conn):
    """Data Guard（v19 兼容版）：不使用 STANDBY_DB_UNIQUE_NAME（可能不存在），改用安全列"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT GROUP#, TYPE, MEMBER, IS_RECOVERY_DEST_FILE
            FROM v$logfile
            WHERE TYPE = 'STANDBY'
        """)
        results['standby_logs'] = cur.fetchall()
    except Exception:
        results['standby_logs'] = []

    try:
        cur.execute("""
            SELECT dest_id, status, destination, archiver, transmit_mode
            FROM v$archive_dest
            WHERE destination IS NOT NULL
        """)
        results['archive_dest'] = cur.fetchall()
    except Exception:
        results['archive_dest'] = []

    try:
        # 不用 STANDBY_DB_UNIQUE_NAME（列名在部分环境不存在），只取存在的列
        cur.execute("""
            SELECT dest_id, database_mode, recovery_mode, protection_mode, status
            FROM v$archive_dest_status
            WHERE status != 'INACTIVE'
        """)
        results['dg_status'] = cur.fetchall()
    except Exception as e:
        results['dg_status'] = []
        results['dg_error'] = str(e)

    return results


def oracle_check_awr_v19(conn):
    """AWR 快照（v19 兼容版）：INTERVAL DAY TO SECOND 不能直接 * 24，用 EXTRACT 转换"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT instance_number, snap_id,
                   TO_CHAR(begin_interval_time,'YYYY-MM-DD HH24:MI') bt,
                   TO_CHAR(end_interval_time,'YYYY-MM-DD HH24:MI') et,
                   ROUND(EXTRACT(DAY FROM (end_interval_time - begin_interval_time)) * 24 +
                         EXTRACT(HOUR FROM (end_interval_time - begin_interval_time)) +
                         EXTRACT(MINUTE FROM (end_interval_time - begin_interval_time)) / 60, 2) elapsed_hr,
                   ERROR_COUNT
            FROM dba_hist_snapshot
            WHERE end_interval_time > SYSDATE - 7
            ORDER BY instance_number, snap_id DESC
        """)
        results['awr_snaps'] = cur.fetchall()

        cur.execute("SELECT * FROM dba_hist_wr_control")
        results['awr_settings'] = cur.fetchall()
    except Exception as e:
        results['error'] = str(e)
    finally:
        cur.close()
    return results


def oracle_check_alert_v19(conn, days=7):
    """Alert 日志（v19 兼容版）：三重容错 v$diag_alert_text → v$diag_alert_xml → 直接读 ADR"""
    results = {}
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT TO_CHAR(alert_time,'YYYY-MM-DD HH24:MI:SS') t,
                   SUBSTR(message_text,1,200) message
            FROM v$diag_alert_text
            WHERE alert_time > SYSDATE - :days
              AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
            ORDER BY alert_time DESC
        """, days=days)
        results['alert_errors'] = cur.fetchall()
    except Exception:
        try:
            cur.execute("""
                SELECT TO_CHAR(trap_time,'YYYY-MM-DD HH24:MI:SS') t,
                       SUBSTR(message_text,1,200) message
                FROM v$diag_alert_xml
                WHERE trap_time > SYSDATE - :days
                  AND (message_text LIKE '%ORA-%' OR message_text LIKE '%ERROR%')
                ORDER BY trap_time DESC
            """, days=days)
            results['alert_errors'] = cur.fetchall()
        except Exception:
            # ADR HOME 路径查不到时直接标记为无权限/不可访问
            results['alert_errors'] = []
            results['alert_hint'] = 'ADR视图不可访问（需 SYSDBA 权限或 CDB 环境）'
    finally:
        cur.close()
    return results


if __name__ == '__main__':
    main()

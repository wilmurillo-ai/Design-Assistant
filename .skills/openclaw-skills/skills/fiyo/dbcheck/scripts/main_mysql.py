#!/usr/bin/env python3
# -*- coding:utf-8 -*-
from version import __version__ as VER

# 磁盘采集时忽略的外接 ISO / Media 挂载点前缀
IGNORE_MOUNTS = {'/mnt/iso', '/media', '/run/media', '/iso', '/cdrom'}

import warnings
warnings.filterwarnings("ignore")
import itertools
import math
import sys
import datetime
import argparse
import subprocess
import logging
import logging.handlers
import socket
import re
import time
from pathlib import Path
import pymysql
import sys, getopt, os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate
import configparser
import importlib
import subprocess
import json
import hashlib
import base64
from datetime import datetime, timedelta
import platform
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import tempfile
import io
import psutil
import shutil
import paramiko

importlib.reload(sys)

# 内置SQL模板配置
SQL_TEMPLATES_CONTENT = """
[report]
name = HeathCheck report
template = ./templates/wordtemplates_v2.0.docx
output = /tmp/MySQLCheckReport.docx

[variables]
datadir = show global variables like 'datadir';
myversion = SELECT VERSION() as version;
threads_connected = SHOW STATUS LIKE 'Threads_connected';
back_log = show variables like 'back_log%';
max_allowed_packet = show variables like 'max_allowed_packet%';
interactive_timeout = show variables like 'interactive_timeout%';
skip_name_resolve = show variables like 'skip_name_resolve%';
max_used_connections = show global status like 'max_used_connections';
max_connections = show global variables like 'max_connections';
current_connections = show global status like 'Threads_connected';
aborted_connections = show global status like 'aborted%';
expire_logs_days = show variables like 'expire_logs_days%';
open_files_limit = show variables like 'open_files_limit%';
thread_cache_size = show variables like 'thread_cache_size%';
sort_buffer_size = show variables like 'sort_buffer_size%';
join_buffer_size = show variables like 'join_buffer_size%';
innodb_buffer_pool_size = show global variables like 'innodb_buffer_pool_size';
innodb_io_capacity = show global variables like 'innodb_io_capacity';
opened_tables = show global status like '%opened_tables%';
table_open_cache = show variables like '%table_open_cache%';
innodb_file_per_table = show variables like 'innodb_file_per_table%';
innodb_open_files = show variables like 'innodb_open_files%';
innodb_thread_concurrency = show variables like 'innodb_thread_concurrency%';
innodb_flush_log_at_trx_commit = show variables like 'innodb_flush_log_at_trx_commit%';
sync_binlog = show variables like 'sync_binlog%';
innodb_log_buffer_size = show variables like 'innodb_log_buffer_size%';
innodb_log_file_size = show variables like 'innodb_log_file_size%';
innodb_log_files_in_group = show variables like 'innodb_log_files_in_group%';
queries = show status like 'queries';
character_set_database = show variables like 'character_set_database';
basedir = show variables like 'basedir';
slow_query_log = show variables like 'slow_query_log';
table_locks_immediate = show status like 'Table_locks_immediate';
table_locks_waited = show status like 'Table_locks_waited';
db_size = SELECT table_schema 'Database_name', sum(table_rows) 'No_of_rows', 
          round(sum(data_length) / 1024 / 1024, 2) 'Size_data_MB', 
          round(sum(index_length)/ 1024 / 1024, 2) 'Size_index_MB' 
          FROM information_schema.TABLES 
          GROUP BY table_schema;
processlist = SHOW FULL PROCESSLIST;
log_bin = SHOW VARIABLES LIKE 'log_bin%';
query_cache = SHOW VARIABLES LIKE '%query_cache%';
slave_status = show slave status;
mysql_users = select user as col1,host as col2,Grant_priv as col3,plugin as col4,account_locked as col5 from mysql.user where user not in ('mysql.infoschema','mysql.session','mysql.sys');
instancetime = select DATE_FORMAT(date_sub(now(), INTERVAL variable_value SECOND),"%Y-%m-%d %H:%i:%s") started_at from performance_schema.global_status where variable_name='Uptime';
platform_info = select variable_name, variable_value from performance_schema.session_variables where variable_name in ('version_compile_os','version_compile_machine');
"""

class RemoteSystemInfoCollector:
    """远程系统信息收集器 - 通过SSH连接获取远程主机信息"""
    
    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        """
        初始化远程系统信息收集器。

        :param host: 远程主机 IP 地址或主机名
        :param port: SSH 端口，默认 22
        :param username: SSH 登录用户名，默认 root
        :param password: SSH 登录密码（与 key_file 二选一）
        :param key_file: SSH 私钥文件路径（与 password 二选一）
        """
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None
    
    def connect(self):
        """
        建立 SSH 连接。

        优先使用私钥文件认证，若无私钥则使用密码认证。
        自动接受远程主机密钥（AutoAddPolicy）。

        :return: 连接成功返回 True，失败返回 False
        """
        try:
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if self.key_file:
                private_key = paramiko.RSAKey.from_private_key_file(self.key_file)
                self.ssh_client.connect(hostname=self.host, port=self.port, username=self.username, pkey=private_key, timeout=10)
            else:
                self.ssh_client.connect(hostname=self.host, port=self.port, username=self.username, password=self.password, timeout=10)
            return True
        except Exception as e:
            print(f"SSH连接失败 {self.host}:{self.port}: {e}")
            return False
    
    def disconnect(self):
        """
        断开 SSH 连接，释放资源。
        """
        if self.ssh_client:
            self.ssh_client.close()
    
    def execute_command(self, command):
        """
        在远程主机上执行 Shell 命令。

        :param command: 要执行的 Shell 命令字符串
        :return: (stdout 输出内容, stderr 错误内容) 的元组；执行失败时返回 ("", 错误信息)
        """
        try:
            stdin, stdout, stderr = self.ssh_client.exec_command(command, timeout=10)
            output = stdout.read().decode('utf-8').strip()
            error = stderr.read().decode('utf-8').strip()
            return output, error
        except Exception as e:
            print(f"执行命令失败: {command}, 错误: {e}")
            return "", str(e)
    
    def get_cpu_info(self):
        """
        通过远程 Shell 命令采集 CPU 信息。

        采集项目包括：CPU 使用率（top）、逻辑核心数（nproc）、
        物理核心数（lscpu）、当前频率及最大频率（lscpu）。

        :return: 包含 CPU 信息的字典，字段：
                 usage_percent（使用率%）、physical_cores（物理核心数）、
                 logical_cores（逻辑核心数）、current_frequency（当前频率 GHz）、
                 max_frequency（最大频率 GHz）；失败时返回空字典
        """
        try:
            cmd = "top -bn1 | grep 'Cpu(s)' | awk '{print $2}' | cut -d'%' -f1"
            output, _ = self.execute_command(cmd)
            cpu_percent = float(output) if output else 0.0
            cmd = "nproc"
            output, _ = self.execute_command(cmd)
            logical_cores = int(output) if output else 0
            cmd = "lscpu | grep 'Core(s) per socket' | awk '{print $4}'"
            output, _ = self.execute_command(cmd)
            cores_per_socket = int(output) if output else 1
            cmd = "lscpu | grep 'Socket(s)' | awk '{print $2}'"
            output, _ = self.execute_command(cmd)
            sockets = int(output) if output else 1
            physical_cores = cores_per_socket * sockets
            cmd = "lscpu | grep 'CPU MHz' | awk '{print $3}'"
            output, _ = self.execute_command(cmd)
            current_frequency = float(output) / 1000 if output else 0.0
            cmd = "lscpu | grep 'CPU max MHz' | awk '{print $4}'"
            output, _ = self.execute_command(cmd)
            max_frequency = float(output) / 1000 if output else current_frequency
            return {
                'usage_percent': cpu_percent,
                'physical_cores': physical_cores,
                'logical_cores': logical_cores,
                'current_frequency': round(current_frequency, 2),
                'max_frequency': round(max_frequency, 2)
            }
        except Exception as e:
            print(f"获取CPU信息失败: {e}")
            return {}
    
    def get_memory_info(self):
        """
        通过远程 Shell 命令采集内存信息。

        使用 `free -b` 命令分别采集物理内存和 Swap 信息，
        所有容量单位均转换为 GB。

        :return: 包含内存信息的字典，字段：
                 total_gb（总量）、available_gb（可用量）、used_gb（已用量）、
                 usage_percent（使用率%）、swap_total_gb、swap_used_gb、
                 swap_usage_percent；失败时返回空字典
        """
        try:
            cmd = "free -b | grep Mem"
            output, _ = self.execute_command(cmd)
            if output:
                parts = output.split()
                total_bytes = int(parts[1])
                used_bytes = int(parts[2])
                free_bytes = int(parts[3])
                available_bytes = int(parts[6]) if len(parts) > 6 else free_bytes
                usage_percent = (used_bytes / total_bytes) * 100 if total_bytes > 0 else 0
                memory_info = {
                    'total_gb': round(total_bytes / (1024**3), 2),
                    'available_gb': round(available_bytes / (1024**3), 2),
                    'used_gb': round(used_bytes / (1024**3), 2),
                    'usage_percent': round(usage_percent, 2)
                }
                cmd = "free -b | grep Swap"
                output, _ = self.execute_command(cmd)
                if output:
                    parts = output.split()
                    swap_total = int(parts[1])
                    swap_used = int(parts[2])
                    swap_usage_percent = (swap_used / swap_total) * 100 if swap_total > 0 else 0
                    memory_info.update({
                        'swap_total_gb': round(swap_total / (1024**3), 2),
                        'swap_used_gb': round(swap_used / (1024**3), 2),
                        'swap_usage_percent': round(swap_usage_percent, 2)
                    })
                return memory_info
            return {}
        except Exception as e:
            print(f"获取内存信息失败: {e}")
            return {}
    
    def get_disk_info(self):
        """
        通过远程 Shell 命令采集磁盘使用信息。

        使用 `df -h` 命令，自动过滤 tmpfs、devtmpfs、overlay 等虚拟文件系统，
        并将 K/M/G/T 单位统一换算为 GB。

        :return: 磁盘分区信息列表，每项为字典，字段：
                 device（设备名）、mountpoint（挂载点）、fstype（文件系统类型）、
                 total_gb、used_gb、free_gb、usage_percent；失败时返回空列表
        """
        try:
            IGNORE_PATTERN = "|".join([
                "/mnt/iso", "/iso", "/media", "/run/media", "/cdrom",
                "/mnt/iso/", "/mnt/media/", "/run/media/", "/iso/", "/cdrom/"
            ])
            ISO_FILTER = "mnt/iso|/iso|/media/|/run/media/|/cdrom"
            cmd = f"df -h | grep -vE 'tmpfs|devtmpfs' | grep -vE '{ISO_FILTER}' | tail -n +2"
            output, _ = self.execute_command(cmd)
            disk_data = []
            if output:
                lines = output.strip().split('\n')
                for line in lines:
                    parts = line.split()
                    if len(parts) >= 6:
                        device = parts[0]
                        mountpoint = parts[5]
                        if any(vfs in device for vfs in ['tmpfs', 'devtmpfs', 'overlay']):
                            continue
                        size_str = parts[1]
                        used_str = parts[2]
                        avail_str = parts[3]
                        usage_percent_str = parts[4].rstrip('%')
                        def to_gb(s):
                            # 将带单位的字符串（K/M/G/T）统一转换为 GB 浮点数
                            if not s: return 0.0
                            s = s.strip().upper()
                            if s.endswith('G'): return round(float(s[:-1]), 2)
                            elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                            elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                            elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                            else:
                                try: return round(float(s), 2)
                                except: return 0.0
                        total_gb = to_gb(size_str)
                        used_gb = to_gb(used_str)
                        free_gb = to_gb(avail_str)
                        try: usage_percent = float(usage_percent_str)
                        except: usage_percent = 0.0
                        disk_data.append({
                            'device': device, 'mountpoint': mountpoint, 'fstype': "ext4",
                            'total_gb': total_gb, 'used_gb': used_gb, 'free_gb': free_gb, 'usage_percent': usage_percent
                        })
            return disk_data
        except Exception as e:
            print(f"获取磁盘信息失败: {e}")
            return []
        except Exception as e:
            print(f"获取磁盘信息失败: {e}")
            return []

    def get_mysql_datadir(self):
        """
        通过远程 Shell 命令采集 MySQL datadir 路径。

        使用 `mysql -e "show global variables like 'datadir'"` 命令查询。

        :return: 包含 datadir 路径的字典，字段：datadir；失败时返回空字典
        """
        try:
            cmd = 'mysql -e "show global variables like \'datadir\';" 2>/dev/null | tail -n 1 | awk \'{print $2}\''
            output, _ = self.execute_command(cmd)
            if output:
                return {'datadir': output.strip()}
            return {}
        except Exception as e:
            print(f"获取MySQL datadir失败: {e}")
            return {}

    def get_system_info(self):
        """
        聚合采集远程主机的全部系统信息。

        依次调用 connect()、get_cpu_info()、get_memory_info()、get_disk_info()，
        并额外采集主机名（hostname）、内核版本（uname -a）、启动时间（who -b）。
        无论成功与否，最终都会调用 disconnect() 断开 SSH 连接。

        :return: 包含系统信息的字典，字段：cpu、memory、disk、hostname、platform、boot_time；
                 SSH 连接失败时返回空字典
        """
        if not self.connect():
            return {}
        try:
            system_info = {
                'cpu': self.get_cpu_info(),
                'memory': self.get_memory_info(),
                'disk': self.get_disk_info(),
                'hostname': "",
                'platform': "",
                'boot_time': "",
                'mysql_datadir': ""
            }
            cmd = "hostname"
            output, _ = self.execute_command(cmd)
            if output: system_info['hostname'] = output.strip()
            cmd = "uname -a"
            output, _ = self.execute_command(cmd)
            if output: system_info['platform'] = output.strip()
            cmd = "who -b | awk '{print $3 \" \" $4}'"
            output, _ = self.execute_command(cmd)
            if output: system_info['boot_time'] = output.strip()
            # 采集 MySQL datadir
            datadir_result = self.get_mysql_datadir()
            if datadir_result:
                system_info['mysql_datadir'] = datadir_result.get('datadir', '')
            return system_info
        finally:
            self.disconnect()

class LocalSystemInfoCollector:
    """本地系统信息收集器 - 使用 psutil 库采集当前主机系统信息，无需 SSH"""

    def __init__(self):
        """初始化本地系统信息收集器（无需任何参数）。"""
        pass

    def get_cpu_info(self):
        """
        采集本机 CPU 信息。

        使用 psutil 获取 CPU 使用率（采样间隔 1 秒）、
        物理/逻辑核心数、当前频率及最大频率。

        :return: 包含 CPU 信息的字典，字段：
                 usage_percent、physical_cores、logical_cores、
                 current_frequency（MHz）、max_frequency（MHz）；
                 失败时返回空字典
        """
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            cpu_count = psutil.cpu_count(logical=False)
            cpu_count_logical = psutil.cpu_count(logical=True)
            cpu_freq = psutil.cpu_freq()
            return {
                'usage_percent': cpu_percent,
                'physical_cores': cpu_count,
                'logical_cores': cpu_count_logical,
                'current_frequency': round(cpu_freq.current, 2) if cpu_freq else 'N/A',
                'max_frequency': round(cpu_freq.max, 2) if cpu_freq else 'N/A'
            }
        except Exception as e:
            print(f"获取CPU信息失败: {e}")
            return {}

    def get_memory_info(self):
        """
        采集本机内存及 Swap 使用情况。

        使用 psutil.virtual_memory() 和 psutil.swap_memory()，
        所有容量单位统一转换为 GB。

        :return: 包含内存信息的字典，字段：
                 total_gb、available_gb、used_gb、usage_percent、
                 swap_total_gb、swap_used_gb、swap_usage_percent；
                 失败时返回空字典
        """
        try:
            memory = psutil.virtual_memory()
            swap = psutil.swap_memory()
            return {
                'total_gb': round(memory.total / (1024**3), 2),
                'available_gb': round(memory.available / (1024**3), 2),
                'used_gb': round(memory.used / (1024**3), 2),
                'usage_percent': memory.percent,
                'swap_total_gb': round(swap.total / (1024**3), 2),
                'swap_used_gb': round(swap.used / (1024**3), 2),
                'swap_usage_percent': swap.percent
            }
        except Exception as e:
            print(f"获取内存信息失败: {e}")
            return {}

    def get_disk_info(self):
        """
        采集本机磁盘分区信息，并额外检查常见 MySQL 数据目录。

        遍历所有已挂载分区（跳过 loop 设备和无文件系统类型的分区），
        同时检测 /var/lib/mysql、/data/mysql、/usr/local/mysql/data
        等常见 MySQL 数据路径，若存在则一并采集。

        :return: 以挂载点为键的磁盘信息字典，每值包含：
                 device、mountpoint、fstype、total_gb、used_gb、
                 free_gb、usage_percent；失败时返回空字典
        """
        try:
            disk_info = {}
            partitions = psutil.disk_partitions()
            IGNORE_PREFIXES = ('/mnt/', '/media', '/run/media', '/snap', '/iso', '/cdrom')
            for partition in partitions:
                mp = partition.mountpoint
                if partition.fstype and 'loop' not in partition.device and not mp.startswith(IGNORE_PREFIXES):
                    try:
                        usage = psutil.disk_usage(partition.mountpoint)
                        disk_info[partition.mountpoint] = {
                            'device': partition.device, 'mountpoint': partition.mountpoint, 'fstype': partition.fstype,
                            'total_gb': round(usage.total / (1024**3), 2), 'used_gb': round(usage.used / (1024**3), 2),
                            'free_gb': round(usage.free / (1024**3), 2), 'usage_percent': usage.percent
                        }
                    except PermissionError:
                        continue
            # 额外检查常见 MySQL 数据目录
            mysql_paths = ['/var/lib/mysql', '/data/mysql', '/usr/local/mysql/data']
            for path in mysql_paths:
                if os.path.exists(path):
                    try:
                        usage = psutil.disk_usage(path)
                        disk_info[f'mysql_data_{path}'] = {
                            'device': 'MySQL Data', 'mountpoint': path, 'fstype': 'N/A',
                            'total_gb': round(usage.total / (1024**3), 2), 'used_gb': round(usage.used / (1024**3), 2),
                            'free_gb': round(usage.free / (1024**3), 2), 'usage_percent': usage.percent
                        }
                    except Exception: pass
            return disk_info
        except Exception as e:
            print(f"获取磁盘信息失败: {e}")
            return {}

    def get_system_info(self):
        """
        聚合采集本机全部系统信息。

        整合 CPU、内存、磁盘信息，以及主机名、操作系统平台描述和系统启动时间。

        :return: 包含系统信息的字典，字段：
                 cpu、memory、disk、hostname、platform、boot_time
        """
        return {
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info(),
            'hostname': socket.gethostname(),
            'platform': platform.platform(),
            'boot_time': datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')
        }

class SystemInfoCollector:
    """系统信息收集器工厂类 - 根据主机类型创建对应的采集器实例"""

    @staticmethod
    def create_collector(host_type='local', **kwargs):
        """
        工厂方法：根据 host_type 创建并返回合适的系统信息采集器。

        :param host_type: 采集器类型，'local' 使用本地 psutil 采集，
                          'remote' 使用 SSH 远程采集，默认为 'local'
        :param kwargs: 当 host_type='remote' 时，透传给 RemoteSystemInfoCollector
                       的连接参数（host、port、username、password、key_file）
        :return: LocalSystemInfoCollector 或 RemoteSystemInfoCollector 实例
        """
        if host_type == 'remote':
            return RemoteSystemInfoCollector(**kwargs)
        else:
            return LocalSystemInfoCollector()

def get_host_disk_usage():
    """
    跨平台获取主机磁盘使用情况的备用函数。

    - Windows 系统：调用 wmic logicaldisk 命令，解析 CSV 输出获取各逻辑盘信息。
    - Linux/macOS 系统：调用 df -h 命令，优先采集 /、/boot、/home、/var、
      /usr、/opt、/tmp 等重要挂载点；若无匹配则回退至采集全部非虚拟文件系统分区。
    所有容量单位统一转换为 GB（浮点数）。

    :return: 磁盘分区信息列表，每项为字典，字段：
             device（设备名/盘符）、mountpoint（挂载点）、fstype（文件系统类型）、
             total_gb、used_gb、free_gb、usage_percent；
             失败时返回空列表
    """
    try:
        disk_data = []
        if platform.system() == "Windows":
            # Windows：使用 wmic 获取逻辑盘信息（CSV 格式）
            result = subprocess.Popen(["wmic", "logicaldisk", "get", "deviceid,size,freespace", "/format:csv"], 
                                    stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = result.communicate(timeout=10)
            output = stdout.decode('utf-8', errors='ignore')
            lines = output.strip().split('\n')[1:]
            for line in lines:
                if line.strip():
                    parts = line.split(',')
                    if len(parts) >= 4:
                        device = parts[1]
                        total_bytes = int(parts[2]) if parts[2] else 0
                        free_bytes = int(parts[3]) if parts[3] else 0
                        used_bytes = total_bytes - free_bytes
                        total_gb = round(total_bytes / (1024**3), 2)
                        used_gb = round(used_bytes / (1024**3), 2)
                        free_gb = round(free_bytes / (1024**3), 2)
                        usage_percent = round((used_bytes / total_bytes) * 100, 2) if total_bytes > 0 else 0
                        disk_data.append({
                            'device': device, 'mountpoint': device + "\\", 'fstype': "NTFS",
                            'total_gb': total_gb, 'used_gb': used_gb, 'free_gb': free_gb, 'usage_percent': usage_percent
                        })
        else:
            # Linux/macOS：使用 df -h 获取磁盘信息
            result = subprocess.Popen(["df", "-h"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = result.communicate(timeout=10)
            output = stdout.decode('utf-8', errors='ignore')
            lines = output.strip().split('\n')[1:]
            # 优先采集重要挂载点
            important_mounts = ['/', '/boot', '/home', '/var', '/usr', '/opt', '/tmp']
            for line in lines:
                parts = line.split()
                if len(parts) >= 6:
                    device = parts[0]
                    mountpoint = parts[5]
                    if mountpoint in important_mounts:
                        size_str = parts[1]
                        used_str = parts[2]
                        avail_str = parts[3]
                        usage_percent_str = parts[4].rstrip('%')
                        def to_gb(s):
                            # 将带单位字符串（K/M/G/T）转换为 GB 浮点数
                            if not s: return 0.0
                            s = s.strip().upper()
                            if s.endswith('G'): return round(float(s[:-1]), 2)
                            elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                            elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                            elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                            else:
                                try: return round(float(s), 2)
                                except: return 0.0
                        total_gb = to_gb(size_str)
                        used_gb = to_gb(used_str)
                        free_gb = to_gb(avail_str)
                        try: usage_percent = float(usage_percent_str)
                        except: usage_percent = 0.0
                        disk_data.append({
                            'device': device, 'mountpoint': mountpoint, 'fstype': "ext4",
                            'total_gb': total_gb, 'used_gb': used_gb, 'free_gb': free_gb, 'usage_percent': usage_percent
                        })
            # 若没有匹配到重要挂载点，则回退采集全量（跳过虚拟文件系统和 ISO/Media 分区）
            if not disk_data:
                for line in lines:
                    parts = line.split()
                    if len(parts) >= 6:
                        device = parts[0]
                        mountpoint = parts[5]
                        if any(vfs in device for vfs in ['tmpfs', 'devtmpfs', 'overlay']): continue
                        # 跳过外接 ISO / Media 挂载点
                        if mountpoint in IGNORE_MOUNTS or any(mountpoint.startswith(p) for p in IGNORE_MOUNTS): continue
                        size_str = parts[1]
                        used_str = parts[2]
                        avail_str = parts[3]
                        usage_percent_str = parts[4].rstrip('%')
                        def to_gb(s):
                            if not s: return 0.0
                            s = s.strip().upper()
                            if s.endswith('G'): return round(float(s[:-1]), 2)
                            elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                            elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                            elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                            else:
                                try: return round(float(s), 2)
                                except: return 0.0
                        total_gb = to_gb(size_str)
                        used_gb = to_gb(used_str)
                        free_gb = to_gb(avail_str)
                        try: usage_percent = float(usage_percent_str)
                        except: usage_percent = 0.0
                        disk_data.append({
                            'device': device, 'mountpoint': mountpoint, 'fstype': "ext4",
                            'total_gb': total_gb, 'used_gb': used_gb, 'free_gb': free_gb, 'usage_percent': usage_percent
                        })
        return disk_data
    except Exception as e:
        print(f"获取磁盘使用率失败: {str(e)}")
        return []

class WordTemplateGenerator:
    """Word 模板生成器 - 动态生成包含 Jinja2 模板语法的巡检报告 Word 模板文件"""

    def __init__(self, inspector_name="Jack"):
        """
        初始化 Word 模板生成器。

        创建新的 Document 对象并调用 _setup_document() 设置页面边距。
        """
        self.doc = Document()
        self.inspector_name = inspector_name
        self._setup_document()

    def _setup_document(self):
        """
        设置 Word 文档的页面边距。

        将文档所有节（Section）的上、下、左、右边距统一设置为 2.54 cm（标准 A4 边距）。
        """
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def create_template(self):
        """
        按顺序组装完整的巡检报告模板，包含封面及 7 个章节。

        依次调用以下方法：
          封面 → 健康状态概览 → 系统资源检查 → MySQL 配置检查
          → 性能分析 → 数据库信息 → 安全信息 → 报告说明

        :return: 组装完成的 Document 对象
        """
        self._add_title_page()
        self._add_summary_section()
        self._add_system_info_section()
        self._add_mysql_config_section()
        self._add_performance_section()
        self._add_database_info_section()
        self._add_security_section()
        self._add_notes_section()
        return self.doc

    def _add_title_page(self):
        """
        生成报告封面页。

        包含标题（MySQL数据库健康巡检报告）和一个 8 行 2 列的信息表，
        表格中使用 Jinja2 模板变量填充：数据库名称、服务器地址、MySQL版本、
        服务器主机名、实例启动时间、巡检人员、服务器平台、报告生成时间。
        封面末尾插入分页符。
        """
        title = self.doc.add_heading('MySQL数据库健康巡检报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        self.doc.add_paragraph()
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        cells = table.rows[0].cells
        cells[0].text = "数据库名称"
        cells[1].text = "{{ co_name[0]['CO_NAME'] }}"
        cells = table.rows[1].cells
        cells[0].text = "服务器地址"
        cells[1].text = "{{ ip[0]['IP'] }}:{{ port[0]['PORT'] }}"
        cells = table.rows[2].cells
        cells[0].text = "MySQL版本"
        cells[1].text = "{{ myversion[0]['version'] }}"
        cells = table.rows[3].cells
        cells[0].text = "服务器主机名"
        cells[1].text = "{{ system_info.hostname }}"
        cells = table.rows[4].cells
        cells[0].text = "实例启动时间"
        cells[1].text = "{% if instancetime %}{{ instancetime[0]['started_at'] }}{% else %}N/A{% endif %}"
        cells = table.rows[5].cells
        cells[0].text = "巡检人员"
        cells[1].text = "{{ inspector_name }}"
        cells = table.rows[6].cells
        cells[0].text = "服务器平台"
        cells[1].text = "{% if platform_info and platform_info|length > 0 %}{% for item in platform_info %}{% if item.variable_name == 'version_compile_os' %}{{ item.variable_value }}{% endif %}{% endfor %}{% else %}N/A{% endif %}"
        cells = table.rows[7].cells
        cells[0].text = "报告生成时间"
        cells[1].text = "{{ report_time }}"
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_page_break()

    def _add_summary_section(self):
        """
        生成第 1 章「健康状态概览」。

        包含一个 2 行 2 列的状态表（总体健康状态、发现问题数量）
        以及健康总结段落，所有值均使用 Jinja2 模板变量占位。
        """
        heading = self.doc.add_heading('1. 健康状态概览', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        cells = table.rows[0].cells
        cells[0].text = "总体健康状态"
        cells[1].text = "{{ health_status }}"
        cells = table.rows[1].cells
        cells[0].text = "发现问题数量"
        cells[1].text = "{{ problem_count }} 个"
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        p = self.doc.add_paragraph("健康总结: ")
        p.add_run("{{ health_summary[0]['health_summary'] }}").bold = True
        p.runs[0].font.size = Pt(11)
        p.runs[1].font.size = Pt(11)
    def _add_system_info_section(self):
        """
        生成第 2 章「系统资源检查」，包含三个小节：

        - 2.1 CPU 信息：4 列表格（使用率、物理/逻辑核心数、当前频率）
        - 2.2 内存信息：4 列表格（总量、已用、可用、使用率）
        - 2.3 磁盘信息：2 列表格（挂载点、使用率），使用 Jinja2 循环语法动态生成行
        """
        heading = self.doc.add_heading('2. 系统资源检查', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading('2.1 CPU信息', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        for i in range(4):
            table.columns[i].width = Cm(3.5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'CPU使用率'
        hdr_cells[1].text = '物理核心数'
        hdr_cells[2].text = '逻辑核心数'
        hdr_cells[3].text = '当前频率(GHz)'
        data_cells = table.rows[1].cells
        data_cells[0].text = "{% if system_info.cpu and system_info.cpu.usage_percent is defined %}{{ '%.2f'|format(system_info.cpu.usage_percent) }}%{% else %}未获取{% endif %}"
        data_cells[1].text = "{% if system_info.cpu and system_info.cpu.physical_cores is defined %}{{ system_info.cpu.physical_cores }}{% else %}未获取{% endif %}"
        data_cells[2].text = "{% if system_info.cpu and system_info.cpu.logical_cores is defined %}{{ system_info.cpu.logical_cores }}{% else %}未获取{% endif %}"
        data_cells[3].text = "{% if system_info.cpu and system_info.cpu.current_frequency != 'N/A' %}{{ '%.2f'|format(system_info.cpu.current_frequency/1000) }}{% else %}未获取{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('2.2 内存信息', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        for i in range(4):
            table.columns[i].width = Cm(3.5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '总内存(GB)'
        hdr_cells[1].text = '已使用(GB)'
        hdr_cells[2].text = '可用内存(GB)'
        hdr_cells[3].text = '使用率'
        data_cells = table.rows[1].cells
        data_cells[0].text = "{% if system_info.memory and system_info.memory.total_gb is defined %}{{ '%.2f'|format(system_info.memory.total_gb) }}{% else %}未获取{% endif %}"
        data_cells[1].text = "{% if system_info.memory and system_info.memory.used_gb is defined %}{{ '%.2f'|format(system_info.memory.used_gb) }}{% else %}未获取{% endif %}"
        data_cells[2].text = "{% if system_info.memory and system_info.memory.available_gb is defined %}{{ '%.2f'|format(system_info.memory.available_gb) }}{% else %}未获取{% endif %}"
        data_cells[3].text = "{% if system_info.memory and system_info.memory.usage_percent is defined %}{{ '%.2f'|format(system_info.memory.usage_percent) }}%{% else %}未获取{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('2.3 磁盘信息', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Cm(8)
        table.columns[1].width = Cm(4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '挂载点'
        hdr_cells[1].text = '使用率'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells = table.add_row().cells
        row_cells[0].text = "{% for disk in system_info.disk_list %}{{ disk.mountpoint }}{% if not loop.last %}\n{% endif %}{% endfor %}"
        row_cells[1].text = "{% for disk in system_info.disk_list %}{% if disk.usage_percent is defined %}{{ '%.2f'|format(disk.usage_percent) }}%{% else %}未获取{% endif %}{% if not loop.last %}\n{% endif %}{% endfor %}"
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
    def _add_mysql_config_section(self):
        """
        生成第 3 章「MySQL 配置检查」，包含三个小节：

        - 3.1 连接配置：最大连接数、当前连接数、交互超时、文件打开限制
        - 3.2 内存配置：InnoDB 缓冲池、排序缓冲区、连接缓冲区、线程缓存
        - 3.3 日志配置：慢查询日志、Binlog 保留天数、InnoDB 日志文件大小、日志刷新设置
        每小节均使用「配置项 / 当前值」二列键值表，值通过 Jinja2 模板变量填充。
        """
        heading = self.doc.add_heading('3. MySQL配置检查', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading('3.1 连接配置', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '配置项'
        hdr_cells[1].text = '当前值'
        data_cells = table.rows[1].cells
        data_cells[0].text = '最大连接数'
        data_cells[1].text = "{% if max_connections and max_connections[0]['Value'] %}{{ max_connections[0]['Value'] }}{% endif %}"
        data_cells = table.rows[2].cells
        data_cells[0].text = '当前连接数'
        data_cells[1].text = "{% if current_connections and current_connections[0]['Value'] %}{{ current_connections[0]['Value'] }}{% endif %}"
        data_cells = table.rows[3].cells
        data_cells[0].text = '交互超时'
        data_cells[1].text = "{% if interactive_timeout and interactive_timeout[0]['Value'] %}{{ interactive_timeout[0]['Value'] }}{% endif %}"
        data_cells = table.rows[4].cells
        data_cells[0].text = '文件打开限制'
        data_cells[1].text = "{% if open_files_limit and open_files_limit[0]['Value'] %}{{ open_files_limit[0]['Value'] }}{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('3.2 内存配置', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '配置项'
        hdr_cells[1].text = '当前值'
        data_cells = table.rows[1].cells
        data_cells[0].text = 'InnoDB缓冲池'
        data_cells[1].text = "{% if innodb_buffer_pool_size and innodb_buffer_pool_size[0]['Value'] %}{{ innodb_buffer_pool_size[0]['Value'] }}{% endif %}"
        data_cells = table.rows[2].cells
        data_cells[0].text = '排序缓冲区'
        data_cells[1].text = "{% if sort_buffer_size and sort_buffer_size[0]['Value'] %}{{ sort_buffer_size[0]['Value'] }}{% endif %}"
        data_cells = table.rows[3].cells
        data_cells[0].text = '连接缓冲区'
        data_cells[1].text = "{% if join_buffer_size and join_buffer_size[0]['Value'] %}{{ join_buffer_size[0]['Value'] }}{% endif %}"
        data_cells = table.rows[4].cells
        data_cells[0].text = '线程缓存'
        data_cells[1].text = "{% if thread_cache_size and thread_cache_size[0]['Value'] %}{{ thread_cache_size[0]['Value'] }}{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('3.3 日志配置', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '配置项'
        hdr_cells[1].text = '当前值'
        data_cells = table.rows[1].cells
        data_cells[0].text = '慢查询日志'
        data_cells[1].text = "{% if slow_query_log and slow_query_log[0]['Value'] %}{{ slow_query_log[0]['Value'] }}{% endif %}"
        data_cells = table.rows[2].cells
        data_cells[0].text = 'Binlog保留天数'
        data_cells[1].text = "{% if expire_logs_days and expire_logs_days[0]['Value'] %}{{ expire_logs_days[0]['Value'] }}{% endif %}"
        data_cells = table.rows[3].cells
        data_cells[0].text = 'InnoDB日志文件大小'
        data_cells[1].text = "{% if innodb_log_file_size and innodb_log_file_size[0]['Value'] %}{{ innodb_log_file_size[0]['Value'] }}{% endif %}"
        data_cells = table.rows[4].cells
        data_cells[0].text = '日志刷新设置'
        data_cells[1].text = "{% if innodb_flush_log_at_trx_commit and innodb_flush_log_at_trx_commit[0]['Value'] %}{{ innodb_flush_log_at_trx_commit[0]['Value'] }}{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def _add_performance_section(self):
        """
        生成第 4 章「性能分析」，包含三个小节：

        - 4.1 QPS 检查：总查询数表格
        - 4.2 锁信息：立即锁表 / 等待锁表两行数据
        - 4.3 异常连接：异常客户端连接 / 异常连接尝试两行数据
        所有值均通过 Jinja2 模板变量填充。
        """
        heading = self.doc.add_heading('4. 性能分析', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading('4.1 QPS检查', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '查询总数'
        hdr_cells[1].text = '值'
        data_cells = table.rows[1].cells
        data_cells[0].text = '总查询数'
        data_cells[1].text = "{% if queries and queries[0]['Value'] %}{{ queries[0]['Value'] }}{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('4.2 锁信息', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '锁类型'
        hdr_cells[1].text = '值'
        data_cells1 = table.rows[1].cells
        data_cells1[0].text = '立即锁表'
        data_cells1[1].text = "{% if table_locks_immediate and table_locks_immediate[0]['Value'] %}{{ table_locks_immediate[0]['Value'] }}{% endif %}"
        data_cells2 = table.rows[2].cells
        data_cells2[0].text = '等待锁表'
        data_cells2[1].text = "{% if table_locks_waited and table_locks_waited[0]['Value'] %}{{ table_locks_waited[0]['Value'] }}{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('4.3 异常连接', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '异常类型'
        hdr_cells[1].text = '值'
        data_cells1 = table.rows[1].cells
        data_cells1[0].text = '异常客户端连接'
        data_cells1[1].text = "{% if aborted_connections and aborted_connections|length > 0 and aborted_connections[0]['Value'] %}{{ aborted_connections[0]['Value'] }}{% endif %}"
        data_cells2 = table.rows[2].cells
        data_cells2[0].text = '异常连接尝试'
        data_cells2[1].text = "{% if aborted_connections and aborted_connections|length > 1 and aborted_connections[1]['Value'] %}{{ aborted_connections[1]['Value'] }}{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def _add_database_info_section(self):
        """
        生成第 5 章「数据库信息」，包含两个小节：

        - 5.1 数据库大小：4 列表格（数据库名、表行数、数据大小 MB、索引大小 MB），
          通过 Jinja2 按索引访问 db_size 列表，最多展示 10 行
        - 5.2 当前进程列表：6 列表格（ID、用户、数据库、状态、命令、时间），
          通过 Jinja2 按索引访问 processlist 列表，最多展示 10 行
        """
        heading = self.doc.add_heading('5. 数据库信息', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading('5.1 数据库大小', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(3)
        table.columns[3].width = Cm(3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '数据库名'
        hdr_cells[1].text = '表行数'
        hdr_cells[2].text = '数据大小(MB)'
        hdr_cells[3].text = '索引大小(MB)'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(10):
            row_cells = table.add_row().cells
            db_name_template = "{{{{ db_size[{}].Database_name if db_size and db_size[{}] else '' }}}}".format(i, i)
            rows_template = "{{{{ db_size[{}].No_of_rows if db_size and db_size[{}] else '' }}}}".format(i, i)
            data_size_template = "{{{{ db_size[{}].Size_data_MB if db_size and db_size[{}] else '' }}}}".format(i, i)
            index_size_template = "{{{{ db_size[{}].Size_index_MB if db_size and db_size[{}] else '' }}}}".format(i, i)
            row_cells[0].text = db_name_template
            row_cells[1].text = rows_template
            row_cells[2].text = data_size_template
            row_cells[3].text = index_size_template
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading('5.2 当前进程列表', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(2)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(2)
        table.columns[5].width = Cm(3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = '用户'
        hdr_cells[2].text = '数据库'
        hdr_cells[3].text = '状态'
        hdr_cells[4].text = '命令'
        hdr_cells[5].text = '时间'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(10):
            row_cells = table.add_row().cells
            id_template = "{{{{ processlist[{}].Id if processlist and processlist[{}] else '' }}}}".format(i, i)
            user_template = "{{{{ processlist[{}].User if processlist and processlist[{}] else '' }}}}".format(i, i)
            db_template = "{{{{ processlist[{}].db if processlist and processlist[{}] else '' }}}}".format(i, i)
            state_template = "{{{{ processlist[{}].State if processlist and processlist[{}] else '' }}}}".format(i, i)
            command_template = "{{{{ processlist[{}].Command if processlist and processlist[{}] else '' }}}}".format(i, i)
            time_template = "{{{{ processlist[{}].Time if processlist and processlist[{}] else '' }}}}".format(i, i)
            row_cells[0].text = id_template
            row_cells[1].text = user_template
            row_cells[2].text = db_template
            row_cells[3].text = state_template
            row_cells[4].text = command_template
            row_cells[5].text = time_template
    def _add_security_section(self):
        """
        生成第 6 章「安全信息」。

        包含 6.1「数据库用户信息」小节：5 列表格
        （用户名、主机、授权权限、认证插件、账户锁定），
        通过 Jinja2 按索引访问 mysql_users 列表，最多展示 15 行。
        """
        heading = self.doc.add_heading('6. 安全信息', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading('6.1 数据库用户信息', level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(3)
        table.columns[4].width = Cm(3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '用户名'
        hdr_cells[1].text = '主机'
        hdr_cells[2].text = '授权权限'
        hdr_cells[3].text = '认证插件'
        hdr_cells[4].text = '账户锁定'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(15):
            row_cells = table.add_row().cells
            user_template = "{{{{ mysql_users[{}].col1 if mysql_users and mysql_users[{}] else '' }}}}".format(i, i)
            host_template = "{{{{ mysql_users[{}].col2 if mysql_users and mysql_users[{}] else '' }}}}".format(i, i)
            grant_template = "{{{{ mysql_users[{}].col3 if mysql_users and mysql_users[{}] else '' }}}}".format(i, i)
            plugin_template = "{{{{ mysql_users[{}].col4 if mysql_users and mysql_users[{}] else '' }}}}".format(i, i)
            locked_template = "{{{{ mysql_users[{}].col5 if mysql_users and mysql_users[{}] else '' }}}}".format(i, i)
            row_cells[0].text = user_template
            row_cells[1].text = host_template
            row_cells[2].text = grant_template
            row_cells[3].text = plugin_template
            row_cells[4].text = locked_template
    def _add_notes_section(self):
        """
        生成第 7 章「报告说明」。

        以段落形式输出 5 条固定说明文字，包含：
        报告生成说明、空白项说明、磁盘信息范围说明、巡检结果免责说明、定期巡检建议。
        """
        heading = self.doc.add_heading('7. 报告说明', level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        notes = [
            "1. 本报告基于MySQL数据库实时状态生成，反映了生成时刻的数据库健康状况",
            "2. 报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用",
            "3. 磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令'df -h'",
            "4. 巡检结果仅供参考，实际运维中请结合具体业务场景进行分析",
            "5. 建议定期进行数据库巡检，及时发现并解决潜在问题"
        ]
        for note in notes:
            p = self.doc.add_paragraph()
            p.add_run(note)
            p.runs[0].font.size = Pt(10)

class SimpleCrypto:
    """简单加密解密工具类 - 基于 XOR + SHA256 密钥 + Base64 编码实现对称加密"""

    def __init__(self, secret="ODB_SECRET_2024"):
        """
        初始化加密工具。

        :param secret: 加密用的原始密钥字符串，默认为 "ODB_SECRET_2024"，
                       会被编码为 bytes 并通过 SHA256 派生实际加密密钥
        """
        self.secret = secret.encode('utf-8')

    def _xor_encrypt_decrypt(self, data, key):
        """
        使用 XOR 对数据进行加密或解密（XOR 可逆，加解密使用同一方法）。

        :param data: 需要处理的字节数据（bytes）
        :param key: 加密/解密使用的密钥字节（bytes），循环重复使用
        :return: 经过 XOR 处理后的字节数据（bytes）
        """
        key_len = len(key)
        result = bytearray()
        for i, byte in enumerate(data):
            result.append(byte ^ key[i % key_len])
        return bytes(result)

    def encrypt(self, text):
        """
        加密文本字符串。

        流程：UTF-8 编码 → SHA256 派生 32 字节密钥 → XOR 加密 → Base64 编码。

        :param text: 需要加密的明文字符串
        :return: Base64 编码后的密文字符串
        """
        data = text.encode('utf-8')
        key = hashlib.sha256(self.secret).digest()[:32]
        encrypted = self._xor_encrypt_decrypt(data, key)
        return base64.b64encode(encrypted).decode('utf-8')

    def decrypt(self, token):
        """
        解密 Base64 编码的密文字符串。

        流程：Base64 解码 → SHA256 派生 32 字节密钥 → XOR 解密 → UTF-8 解码。

        :param token: Base64 编码的密文字符串
        :return: 解密后的明文字符串
        :raises ValueError: 解密过程中发生异常时抛出，包含具体错误信息
        """
        try:
            encrypted = base64.b64decode(token.encode('utf-8'))
            key = hashlib.sha256(self.secret).digest()[:32]
            decrypted = self._xor_encrypt_decrypt(encrypted, key)
            return decrypted.decode('utf-8')
        except Exception as e:
            raise ValueError(f"解密失败: {e}")

class LicenseValidator:
    """许可证验证类 - 负责许可证文件的创建、读取、解密和有效性验证"""

    def __init__(self):
        """
        初始化许可证验证器。

        设置许可证文件路径（mysql_inspector.lic）、试用期天数（36500 天，约 100 年），
        创建 SimpleCrypto 加密实例，并调用 _init_license_system() 确保许可证文件存在。
        """
        self.license_file = "mysql_inspector.lic"
        self.trial_days = 36500
        self.crypto = SimpleCrypto()
        self._init_license_system()

    def _parse_datetime(self, date_string):
        """
        解析日期时间字符串，兼容 Python 3.6 及以上版本。

        优先使用 Python 3.7+ 的 datetime.fromisoformat()，
        若不支持则手动解析 ISO 格式（含 'T' 分隔符）或空格分隔格式。

        :param date_string: ISO 格式的日期时间字符串，如 "2026-01-01T00:00:00"
        :return: datetime 对象；解析失败时返回当前时间
        """
        try:
            return datetime.fromisoformat(date_string)
        except AttributeError:
            try:
                if 'T' in date_string:
                    date_part, time_part = date_string.split('T')
                    year, month, day = map(int, date_part.split('-'))
                    time_parts = time_part.split(':')
                    hour, minute = int(time_parts[0]), int(time_parts[1])
                    second = int(time_parts[2].split('.')[0]) if len(time_parts) > 2 else 0
                    return datetime(year, month, day, hour, minute, second)
                else:
                    date_part, time_part = date_string.split(' ')
                    year, month, day = map(int, date_part.split('-'))
                    hour, minute, second = map(int, time_part.split(':'))
                    return datetime(year, month, day, hour, minute, second)
            except Exception as e:
                print(f"日期解析错误: {e}")
                return datetime.now()

    def _format_datetime(self, dt):
        """
        将 datetime 对象格式化为 ISO 格式字符串（兼容 Python 3.6）。

        :param dt: datetime 对象
        :return: 格式为 "YYYY-MM-DDTHH:MM:SS" 的字符串
        """
        return dt.strftime('%Y-%m-%dT%H:%M:%S')

    def _init_license_system(self):
        """
        初始化许可证系统。

        检查许可证文件是否存在，若不存在则自动调用 _create_trial_license() 创建。
        """
        if not os.path.exists(self.license_file):
            self._create_trial_license()

    def _create_trial_license(self):
        """
        创建永久许可证文件（试用期为 100 年）。

        许可证数据包含：类型（PERMANENT）、创建时间、过期时间、机器 ID、数字签名。
        数据序列化为 JSON 后通过 SimpleCrypto 加密，写入 license_file。
        若过期时间年份超过 9999 则自动修正为 9999-12-31 或 2099-12-31。
        """
        create_time = datetime.now()
        try:
            expire_time = create_time + timedelta(days=self.trial_days)
            if expire_time.year > 9999:
                expire_time = datetime(9999, 12, 31)
                print("⚠️  许可证日期超出范围，已调整为9999-12-31")
        except OverflowError:
            expire_time = datetime(2099, 12, 31)
            print("⚠️  许可证日期溢出，已调整为2099-12-31")
        license_data = {
            "type": "PERMANENT",
            "create_time": self._format_datetime(create_time),
            "expire_time": self._format_datetime(expire_time),
            "machine_id": self._get_machine_id(),
            "signature": self._generate_signature("PERMANENT")
        }
        encrypted_data = self.crypto.encrypt(json.dumps(license_data))
        with open(self.license_file, 'w') as f:
            f.write(encrypted_data)

    def _get_machine_id(self):
        """
        获取当前主机的唯一标识符。

        通过组合主机名、操作系统名称和版本号生成字符串，
        再取 MD5 哈希的前 16 位作为机器 ID。

        :return: 16 位十六进制机器 ID 字符串；获取失败时返回 "unknown_machine"
        """
        try:
            machine_info = f"{platform.node()}-{platform.system()}-{platform.release()}"
            return hashlib.md5(machine_info.encode()).hexdigest()[:16]
        except:
            return "unknown_machine"

    def _generate_signature(self, license_type):
        """
        根据许可证类型生成数字签名。

        签名数据由许可证类型、当前日期和固定密钥拼接而成，
        使用 SHA256 哈希生成最终签名。

        :param license_type: 许可证类型字符串，"PERMANENT" 或其他值
        :return: SHA256 哈希的十六进制字符串
        """
        if license_type == "PERMANENT":
            key = "ODB2024PERM"
        else:
            key = "ODB2024TRL"
        signature_data = f"{license_type}-{datetime.now().strftime('%Y%m%d')}-{key}"
        return hashlib.sha256(signature_data.encode()).hexdigest()

    def _verify_signature(self, license_data):
        """
        验证许可证数据中的数字签名是否有效。

        重新生成预期签名并与许可证中的签名对比。

        :param license_data: 解密后的许可证数据字典，需包含 "type" 和 "signature" 字段
        :return: 签名匹配返回 True，不匹配或发生异常返回 False
        """
        try:
            expected_signature = self._generate_signature(license_data["type"])
            return license_data["signature"] == expected_signature
        except:
            return False

    def validate_license(self):
        """
        验证许可证文件的完整性和有效性。

        依次执行：文件存在性检查 → 读取并解密 → 签名验证 → 过期时间检查。

        :return: 三元组 (is_valid, message, remaining_days)
                 - is_valid (bool): 许可证是否有效
                 - message (str): 验证结果描述信息
                 - remaining_days (int): 剩余有效天数（永久版固定返回 99999）
        """
        try:
            if not os.path.exists(self.license_file):
                return False, "许可证文件不存在", 0
            with open(self.license_file, 'r') as f:
                encrypted_data = f.read().strip()
            decrypted_data = self.crypto.decrypt(encrypted_data)
            license_data = json.loads(decrypted_data)
            if not self._verify_signature(license_data):
                return False, "许可证签名无效", 0
            expire_time = self._parse_datetime(license_data["expire_time"])
            remaining_days = (expire_time - datetime.now()).days
            if remaining_days < 0:
                return False, "许可证已过期", 0
            license_type = license_data.get("type", "PERMANENT")
            if license_type == "PERMANENT":
                return True, "永久版许可证有效", 99999
            else:
                return True, f"{license_type}版许可证有效，剩余 {remaining_days} 天", remaining_days
        except Exception as e:
            return False, f"许可证验证失败: {str(e)}", 0

def getlogger():
    """
    获取全局日志记录器。

    创建名为 "mysql_check" 的 Logger，日志级别为 INFO，
    若尚未添加处理器则附加一个输出到控制台（StreamHandler）的格式化处理器。
    该函数保证不重复添加处理器（幂等调用安全）。

    :return: 配置完成的 logging.Logger 实例
    """
    logger = logging.getLogger('mysql_check')
    logger.setLevel(logging.INFO)
    if not logger.handlers:
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
    return logger

logger = getlogger()

class passArgu(object):
    """命令行参数解析类 - 负责解析程序启动时传入的命令行参数"""

    def get_argus(self):
        """
        解析命令行参数并返回解析结果对象。

        支持以下参数：
        - -C / --sqltemplates: SQL 模板文件路径，默认 templates/sqltemplates.ini
        - -L / --label: 单机巡检时的数据库标签名称
        - -B / --batch: 批量模式开关（store_true）

        :return: argparse.Namespace 对象，包含各参数的值
        """
        all_info = argparse.ArgumentParser(
            description="--example: python3 mysql_autoDOC.py -C templates/sqltemplates.ini -L '标签名称'")
        all_info.add_argument('-C', '--sqltemplates', required=False, default='templates/sqltemplates.ini',
                              help='SQL sqltemplates.')
        all_info.add_argument('-L', '--label', required=False, help='Label used when health check single database.')
        all_info.add_argument('-B', '--batch', action='store_true', help='Batch mode (use interactive input for multiple DBs)')
        all_para = all_info.parse_args()
        return all_para

class ExcelTemplateManager:
    """Excel 模板管理器 - 负责批量巡检配置 Excel 模板的创建和读取"""

    def __init__(self):
        """
        初始化 Excel 模板管理器。

        设置默认模板文件名为 "mysql_batch_template.xlsx"。
        """
        self.template_file = "mysql_batch_template.xlsx"

    def create_template(self):
        """
        创建批量巡检 Excel 配置模板文件。

        模板包含两个工作表：
        - 「MySQL数据库配置」：13 列表头，包含 MySQL 连接信息和 SSH 连接信息，
          并预填 2 行示例数据；密码列用红色字体提醒
        - 「使用说明」：字段说明及注意事项

        :return: 成功返回 True，失败返回 False
        """
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "MySQL数据库配置"
            headers = [
                "序号", "数据库标签", "主机地址", "端口", "用户名", 
                "密码", "数据库名称", 
                "SSH主机", "SSH端口", "SSH用户名", "SSH密码", "SSH密钥文件路径",
                "备注"
            ]
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            column_widths = [8, 20, 15, 8, 15, 20, 15, 15, 8, 15, 20, 25, 20]
            for col, width in enumerate(column_widths, 1):
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width
            example_data = [
                [1, "生产数据库", "192.168.1.100", 3306, "root", "password", "mysql",
                 "192.168.1.100", 22, "root", "ssh_password", "/path/to/private_key", "主数据库"],
                [2, "测试数据库", "localhost", 3306, "test_user", "test123", "test_db",
                 "", 22, "", "", "", "测试环境"],
            ]
            for row, data in enumerate(example_data, 2):
                for col, value in enumerate(data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    if col == 6 or col == 11:
                        cell.font = Font(color="FF0000")
            ws2 = wb.create_sheet("使用说明")
            instructions = [
                ["MySQL批量巡检配置模板使用说明"],
                [""],
                ["字段说明:"],
                ["序号", "自动生成的序号，用于标识"],
                ["数据库标签", "用于报告标识的名称，如'生产数据库'"],
                ["主机地址", "MySQL服务器IP地址或主机名"],
                ["端口", "MySQL服务端口，默认3306"],
                ["用户名", "连接数据库的用户名"],
                ["密码", "连接数据库的密码"],
                ["数据库名称", "要连接的数据库名称(可选)"],
                ["SSH主机", "远程主机的IP地址或主机名（用于获取系统信息），如果为空则使用MySQL主机地址"],
                ["SSH端口", "SSH服务端口，默认22"],
                ["SSH用户名", "SSH连接用户名"],
                ["SSH密码", "SSH连接密码（与密钥文件二选一）"],
                ["SSH密钥文件路径", "SSH私钥文件路径（与密码二选一）"],
                ["备注", "额外的说明信息"],
                [""],
                ["注意事项:"],
                ["1. 请确保MySQL服务器允许远程连接"],
                ["2. 建议使用只读权限的账户进行巡检"],
                ["3. 密码会以明文保存，请妥善保管Excel文件"],
                ["4. 支持同时巡检多个不同的MySQL实例"],
                ["5. 如需获取系统信息（CPU、内存、磁盘等），请填写SSH连接信息"],
                ["6. SSH连接支持密码和密钥两种认证方式，优先使用密钥"],
            ]
            for row, instruction in enumerate(instructions, 1):
                if len(instruction) > 1:
                    ws2.cell(row=row, column=1, value=instruction[0])
                    ws2.cell(row=row, column=2, value=instruction[1])
                else:
                    ws2.cell(row=row, column=1, value=instruction[0])
                if row == 1:
                    ws2.cell(row=row, column=1).font = Font(bold=True, size=14)
            wb.save(self.template_file)
            print(f"✅ Excel模板已创建: {self.template_file}")
            print("📝 请填写模板中的数据库连接信息后使用批量巡检功能")
            print("🔐 如需获取系统信息，请填写SSH连接信息")
            return True
        except Exception as e:
            print(f"❌ 创建Excel模板失败: {e}")
            return False
    def read_template(self, file_path=None):
        """
        读取 Excel 配置模板，解析并返回数据库连接信息列表。

        从「MySQL数据库配置」工作表的第 2 行开始逐行读取，跳过标签列（B列）为空的行。
        自动处理端口、密码、SSH 信息的默认值及类型转换。
        若 SSH 主机列为空，则默认使用 MySQL 主机地址作为 SSH 主机。

        :param file_path: Excel 文件路径，默认使用 self.template_file
        :return: 数据库配置字典列表，每项包含：
                 name、ip、port、user、password、database、
                 ssh_host、ssh_port、ssh_user、ssh_password、ssh_key_file、remark；
                 文件不存在或读取失败时返回 None
        """
        if file_path is None:
            file_path = self.template_file
        if not os.path.exists(file_path):
            print(f"❌ Excel模板文件不存在: {file_path}")
            return None
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb["MySQL数据库配置"]
            db_list = []
            for row in range(2, ws.max_row + 1):
                if not ws.cell(row=row, column=2).value:
                    continue
                port_value = ws.cell(row=row, column=4).value
                if port_value is None:
                    port_value = 3306
                else:
                    try:
                        port_value = int(port_value)
                    except (ValueError, TypeError):
                        port_value = 3306
                password_value = ws.cell(row=row, column=6).value
                if password_value is None:
                    password_value = ""
                else:
                    password_value = str(password_value)
                ssh_host = ws.cell(row=row, column=8).value
                if not ssh_host:
                    ssh_host = str(ws.cell(row=row, column=3).value) if ws.cell(row=row, column=3).value else "localhost"
                else:
                    ssh_host = str(ssh_host)
                ssh_port = ws.cell(row=row, column=9).value
                if ssh_port is None:
                    ssh_port = 22
                else:
                    try:
                        ssh_port = int(ssh_port)
                    except (ValueError, TypeError):
                        ssh_port = 22
                ssh_user = ws.cell(row=row, column=10).value
                if ssh_user is None:
                    ssh_user = "root"
                else:
                    ssh_user = str(ssh_user)
                ssh_password = ws.cell(row=row, column=11).value
                if ssh_password is None:
                    ssh_password = ""
                else:
                    ssh_password = str(ssh_password)
                ssh_key_file = ws.cell(row=row, column=12).value
                if ssh_key_file is None:
                    ssh_key_file = ""
                else:
                    ssh_key_file = str(ssh_key_file)
                db_info = {
                    'name': str(ws.cell(row=row, column=2).value) if ws.cell(row=row, column=2).value else f"DB_{row-1}",
                    'ip': str(ws.cell(row=row, column=3).value) if ws.cell(row=row, column=3).value else "localhost",
                    'port': port_value,
                    'user': str(ws.cell(row=row, column=5).value) if ws.cell(row=row, column=5).value else "root",
                    'password': password_value,
                    'database': str(ws.cell(row=row, column=7).value) if ws.cell(row=row, column=7).value else "",
                    'ssh_host': ssh_host,
                    'ssh_port': ssh_port,
                    'ssh_user': ssh_user,
                    'ssh_password': ssh_password,
                    'ssh_key_file': ssh_key_file,
                    'remark': str(ws.cell(row=row, column=13).value) if ws.cell(row=row, column=13).value else ""
                }
                db_list.append(db_info)
            if not db_list:
                print("❌ Excel模板中没有有效的数据库配置")
                return None
            print(f"✅ 从Excel模板读取到 {len(db_list)} 个数据库配置")
            ssh_count = sum(1 for db in db_list if db['ssh_host'] and (db['ssh_password'] or db['ssh_key_file']))
            print(f"🔐 其中 {ssh_count} 个配置包含SSH连接信息，可以获取系统信息")
            return db_list
        except Exception as e:
            print(f"❌ 读取Excel模板失败: {e}")
            return None

def input_db_info():
    """
    通过命令行交互式输入数据库连接信息。

    依次提示用户输入：主机地址、端口、用户名、密码、数据库标签名称，
    以及可选的 SSH 连接信息（主机、端口、用户名、认证方式：密码或私钥文件）。
    输入完成后自动验证 MySQL 连接；若配置了 SSH 信息，同时验证 SSH 连接。
    连接验证失败时可选择重新输入或退出。

    :return: 包含数据库连接信息的字典（含 SSH 信息字段）；
             用户放弃输入或连接验证失败且不重试时返回 None
    """
    print("\n请输入数据库连接信息:")
    host = input("主机地址 [localhost]: ").strip() or "localhost"
    port_input = input("端口 [3306]: ").strip()
    if not port_input:
        port = 3306
    else:
        try:
            port = int(port_input)
        except ValueError:
            print("⚠️  端口输入无效，使用默认值3306")
            port = 3306
    user = input("用户名 [root]: ").strip() or "root"
    import getpass
    password = getpass.getpass("密码: ").strip()
    db_name = input("数据库名称(用于报告标识) [MySQL_Server]: ").strip() or "MySQL_Server"
    print("\n🔐 系统信息收集配置:")
    print("如需获取系统信息（CPU、内存、磁盘等），请配置SSH连接")
    enable_ssh = input("是否配置SSH连接? (y/n) [n]: ").strip().lower()
    ssh_info = {}
    if enable_ssh in ['y', 'yes']:
        ssh_host = input(f"SSH主机地址 [{host}]: ").strip() or host
        ssh_port_input = input("SSH端口 [22]: ").strip()
        if not ssh_port_input:
            ssh_port = 22
        else:
            try:
                ssh_port = int(ssh_port_input)
            except ValueError:
                print("⚠️  SSH端口输入无效，使用默认值22")
                ssh_port = 22
        ssh_user = input("SSH用户名 [root]: ").strip() or "root"
        auth_choice = input("SSH认证方式: 1.密码 2.密钥文件 [1]: ").strip()
        if auth_choice == '2':
            ssh_key_file = input("SSH私钥文件路径: ").strip()
            ssh_password = ""
            if not os.path.exists(ssh_key_file):
                print(f"❌ 密钥文件不存在: {ssh_key_file}")
                retry = input("是否重新输入? (y/n) [y]: ").strip().lower()
                if retry in ['', 'y', 'yes']:
                    return input_db_info()
                else:
                    ssh_key_file = ""
        else:
            ssh_password = getpass.getpass("SSH密码: ").strip()
            ssh_key_file = ""
        ssh_info = {
            'ssh_host': ssh_host,
            'ssh_port': ssh_port,
            'ssh_user': ssh_user,
            'ssh_password': ssh_password,
            'ssh_key_file': ssh_key_file
        }
    print(f"\n🔍 正在验证MySQL连接 {host}:{port}...")
    try:
        conn = pymysql.connect(host=host, port=port, user=user, password=password, charset='utf8mb4', connect_timeout=10)
        conn.close()
        print(f"✅ 成功连接到MySQL {host}:{port}")
    except Exception as e:
        print(f"❌ MySQL连接失败: {e}")
        retry = input("是否重新输入? (y/n) [n]: ").strip().lower()
        if retry == 'y':
            return input_db_info()
        else:
            return None
    if ssh_info:
        print(f"🔍 正在验证SSH连接 {ssh_info['ssh_host']}:{ssh_info['ssh_port']}...")
        try:
            collector = RemoteSystemInfoCollector(
                host=ssh_info['ssh_host'], port=ssh_info['ssh_port'], username=ssh_info['ssh_user'],
                password=ssh_info['ssh_password'] if ssh_info['ssh_password'] else None,
                key_file=ssh_info['ssh_key_file'] if ssh_info['ssh_key_file'] else None
            )
            if collector.connect():
                print(f"✅ 成功连接到SSH {ssh_info['ssh_host']}:{ssh_info['ssh_port']}")
                collector.disconnect()
            else:
                print(f"❌ SSH连接失败")
        except Exception as e:
            print(f"❌ SSH连接失败: {e}")
    db_info = {'name': db_name, 'ip': host, 'port': port, 'user': user, 'password': password}
    db_info.update(ssh_info)
    return db_info

def show_main_menu():
    """
    显示程序主菜单并等待用户选择。

    打印 MySQL 数据库巡检工具 的主菜单，
    菜单选项：1 单机巡检、2 批量巡检、3 创建 Excel 模板、4 退出。
    循环接受输入，直到用户输入有效选项（1-4）为止。

    :return: 用户选择的菜单项字符串（"1"/"2"/"3"/"4"）
    """
    print("\n" + "=" * 60)
    print("            DBCheck - MySQL 巡检工具 " + VER)
    print("=" * 60)
    print("1. 单机巡检")
    print("2. 批量巡检(从Excel导入)")
    print("3. 创建Excel配置模板")
    print("4. 退出")
    print("=" * 60)
    while True:
        choice = input("请选择巡检模式 (1-4): ").strip()
        if choice in ['1', '2', '3', '4']:
            return choice
        else:
            print("❌ 无效选择，请输入1-4之间的数字")

class getData(object):
    """数据采集类 - 负责连接 MySQL 数据库并执行全量巡检 SQL，同步采集系统信息和风险分析"""

    def __init__(self, ip, port, user, password, ssh_info=None):
        """
        初始化数据采集实例并建立 MySQL 连接。

        通过 passArgu 解析命令行参数获取标签名，
        使用 pymysql 建立数据库连接；连接失败时 conn_db2 置为 None。
        初始化空的 context 字典，用于存储所有巡检结果。

        :param ip: MySQL 服务器 IP 地址或主机名
        :param port: MySQL 服务端口
        :param user: MySQL 登录用户名
        :param password: MySQL 登录密码
        :param ssh_info: SSH 连接信息字典（可选），含 ssh_host、ssh_port、ssh_user、
                         ssh_password、ssh_key_file 字段；为空则使用本地采集模式
        """
        self.label = str(infos.label)
        self.H = ip
        self.P = int(port)
        self.user = user
        self.password = password
        self.ssh_info = ssh_info or {}
        try:
            self.conn_db2 = pymysql.connect(host=self.H, port=self.P, user=self.user, password=self.password, charset='utf8mb4', connect_timeout=10)
        except Exception as e:
            print(f"❌ 数据库连接失败: {e}")
            self.conn_db2 = None
        self.context = {}
    def print_progress_bar(self, iteration, total, prefix='', suffix='', decimals=1, length=50, fill='█'):
        """
        在终端打印文字版进度条（覆盖当前行）。

        :param iteration: 当前步骤数（从 0 开始）
        :param total: 总步骤数
        :param prefix: 进度条前缀文字
        :param suffix: 进度条后缀文字
        :param decimals: 百分比小数位数，默认 1
        :param length: 进度条字符长度，默认 50
        :param fill: 已完成部分的填充字符，默认 '█'
        """
        percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
        filled_length = int(length * iteration // total)
        bar = fill * filled_length + '-' * (length - filled_length)
        print(f'\r{prefix} |{bar}| {percent}% {suffix}', end='\r')
        if iteration == total:
            print()
    def checkdb(self, sqlfile=''):
        """
        执行 MySQL 数据库健康巡检，采集系统信息并进行风险分析。

        主要流程：
        1. 读取 SQL 模板（内置 builtin 模式 或 外部 .ini 文件）
        2. 逐条执行模板中的 SQL 语句，结果以字典列表形式存入 context
        3. 采集系统信息（SSH 远程或本地 psutil）并规范化磁盘信息格式
        4. 自动分析风险项（连接数使用率 > 80%、内存使用率 > 90%、磁盘使用率 > 90%）
        5. 全程打印进度条

        :param sqlfile: SQL 模板文件路径，传入 'builtin' 时使用内置模板（SQL_TEMPLATES_CONTENT），
                        传入空字符串或文件路径时从文件加载
        :return: 包含所有巡检结果的 context 字典；连接异常或读取模板失败时返回当前已有内容
        """
        print("\n开始巡检...")
        total_steps = 15
        current_step = 0
        cfg = configparser.RawConfigParser()
        try:
            if sqlfile == 'builtin':
                cfg.read_string(SQL_TEMPLATES_CONTENT)
            else:
                cfg.read(sqlfile, encoding='utf-8')
        except Exception as e:
            print(f"❌ 读取SQL模板失败: {e}")
            return self.context
        init_keys = ["myversion", "threads_connected", "back_log", "max_allowed_packet", 
                    "interactive_timeout", "skip_name_resolve", "max_used_connections", 
                    "max_connections", "current_connections", "aborted_connections", 
                    "expire_logs_days", "open_files_limit", "thread_cache_size", 
                    "sort_buffer_size", "join_buffer_size", "innodb_buffer_pool_size", 
                    "innodb_io_capacity", "opened_tables", "table_open_cache", 
                    "innodb_file_per_table", "innodb_open_files", "innodb_thread_concurrency", 
                    "innodb_flush_log_at_trx_commit", "sync_binlog", "innodb_log_buffer_size", 
                    "innodb_log_file_size", "innodb_log_files_in_group", "queries", 
                    "character_set_database", "basedir", "slow_query_log", 
                    "table_locks_immediate", "table_locks_waited", "db_size", 
                    "processlist", "log_bin", "query_cache", "slave_status",
                    "mysql_users", "instancetime", "platform_info"]
        for key in init_keys:
            self.context.update({key: []})
        try:
            cursor_ver = self.conn_db2.cursor()
            cursor_ver.execute("SELECT VERSION()")
            version_result = cursor_ver.fetchone()
            mysql_version = version_result[0] if version_result else "Unknown"
            cursor_ver.close()
            self.context.update({"myversion": [{'version': mysql_version}]})
            self.context.update({"health_summary": [{'health_summary': '运行良好'}]})
        except Exception as e:
            print(f"❌ 获取版本信息失败: {e}")
            self.context.update({"myversion": [{'version': 'Unknown'}]})
            self.context.update({"health_summary": [{'health_summary': '运行良好'}]})
        try:
            cursor2 = self.conn_db2.cursor()
            variables_items = list(cfg.items("variables"))
            for i, (name, stmt) in enumerate(variables_items):
                try:
                    current_step = int((i / len(variables_items)) * total_steps)
                    self.print_progress_bar(current_step, total_steps, prefix='巡检进度:', suffix=f'步骤 {i+1}/{len(variables_items)}')
                    cursor2.execute(stmt.replace('\n', ' ').replace('\r', ' '))
                    result = [dict((cursor2.description[i][0], value) for i, value in enumerate(row)) for row in cursor2.fetchall()]
                    self.context[name] = result
                    time.sleep(0.05)
                except Exception as e:
                    print(f"\n⚠️  步骤 {name} 执行失败: {e}")
                    self.context[name] = []
        except Exception as e:
            print(f'\n❌ 数据库查询失败: {e}')
        finally:
            if 'cursor2' in locals():
                cursor2.close()
        current_step = total_steps - 2
        self.print_progress_bar(current_step, total_steps, prefix='巡检进度:', suffix='收集系统信息')
        try:
            if self.ssh_info and self.ssh_info.get('ssh_host'):
                print(f"\n🔍 通过SSH收集系统信息: {self.ssh_info['ssh_host']}")
                collector = RemoteSystemInfoCollector(
                    host=self.ssh_info['ssh_host'], port=self.ssh_info.get('ssh_port', 22),
                    username=self.ssh_info.get('ssh_user', 'root'),
                    password=self.ssh_info.get('ssh_password'), key_file=self.ssh_info.get('ssh_key_file')
                )
            else:
                print(f"\n🔍 收集本地系统信息")
                collector = LocalSystemInfoCollector()
            system_info = collector.get_system_info()
            if isinstance(system_info.get('disk'), dict):
                disk_list = list(system_info['disk'].values())
                system_info['disk_list'] = disk_list
            elif isinstance(system_info.get('disk'), list):
                system_info['disk_list'] = system_info['disk']
            else:
                disk_info = get_host_disk_usage()
                system_info['disk_list'] = disk_info
            self.context.update({"system_info": system_info})
            # 如果通过SSH获取到MySQL datadir，覆盖SQL查询结果（SSH更精准）
            if self.ssh_info and self.ssh_info.get('ssh_host') and system_info.get('mysql_datadir'):
                ssh_datadir = system_info.get('mysql_datadir', '')
                if ssh_datadir:
                    self.context['datadir'] = [{'Value': ssh_datadir}]
                    print(f"\n✅ 通过SSH获取MySQL datadir: {ssh_datadir}")
        except Exception as e:
            print(f"\n❌ 收集系统信息失败: {e}")
            self.context.update({"system_info": {
                'hostname': '未知', 'platform': '未知', 'boot_time': '未知',
                'cpu': {}, 'memory': {},
                'disk_list': [{'device': '/dev/sda1', 'mountpoint': '/', 'fstype': 'ext4', 'total_gb': 0, 'used_gb': 0, 'free_gb': 0, 'usage_percent': 0}]
            }})
        current_step = total_steps - 1
        self.print_progress_bar(current_step, total_steps, prefix='巡检进度:', suffix='分析风险和建议')
        self.context.update({"auto_analyze": []})
        try:
            # 使用增强智能分析模块（15+ 条规则）
            try:
                from analyzer import smart_analyze_mysql
                issues = smart_analyze_mysql(self.context)
                self.context['auto_analyze'] = issues
            except ImportError:
                # 降级：使用内置基础规则
                if self.context.get('max_used_connections') and self.context.get('max_connections'):
                    max_used = int(self.context['max_used_connections'][0]['Value'])
                    max_conn = int(self.context['max_connections'][0]['Value'])
                    conn_usage = (max_used / max_conn) * 100 if max_conn > 0 else 0
                    if conn_usage > 80:
                        self.context['auto_analyze'].append({
                            'col1': "连接数使用率", "col2": "高风险",
                            "col3": f"连接数使用率高达 {conn_usage:.1f}%，接近最大连接数限制",
                            "col4": "高", "col5": "DBA", "fix_sql": ""
                        })
                if self.context.get('system_info', {}).get('memory', {}).get('usage_percent', 0) > 90:
                    self.context['auto_analyze'].append({
                        'col1': "系统内存使用率", "col2": "高风险",
                        "col3": f"系统内存使用率超过90%",
                        "col4": "高", "col5": "系统管理员", "fix_sql": ""
                    })
        except Exception as e:
            print(f"\n❌ 风险分析失败: {e}")

        # AI 智能诊断（从 ai_config.json 读取配置，传递给 analyzer.AIAdvisor）
        self.context['ai_advice'] = ''
        try:
            from analyzer import AIAdvisor
            import json as _json
            cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
            ai_cfg = {}
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    ai_cfg = _json.load(f)
            advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
            if advisor.enabled:
                label = self.context.get('co_name', [{}])[0].get('CO_NAME', 'MySQL')
                print(f"\n🤖 正在调用 AI 诊断（{advisor.backend} / {advisor.model}）...")
                ai_advice = advisor.diagnose('mysql', label, self.context, issues)
                self.context['ai_advice'] = ai_advice
        except Exception as e:
            self.context['ai_advice'] = ''

        self.print_progress_bar(total_steps, total_steps, prefix='巡检进度:', suffix='完成')
        return self.context

class saveDoc(object):
    """报告保存类 - 将巡检数据渲染到 Word 模板并输出最终报告文件"""

    def __init__(self, context, ofile, ifile, inspector_name="Jack"):
        """
        初始化报告保存实例。

        :param context: 包含所有巡检数据的字典（由 getData.checkdb() 返回）
        :param ofile: 输出的 Word 报告文件路径（目标路径）
        :param ifile: Word 模板文件路径（包含 Jinja2 占位符的 .docx 文件）
        :param inspector_name: 巡检人员姓名，默认 "Jack"
        """
        self.context = context
        self.ofile = ofile
        self.ifile = ifile
        self.inspector_name = inspector_name

    def contextsave(self):
        """
        将巡检数据渲染并保存为 Word 报告文件。

        主要流程：
        1. 补全 context 中缺失的必要字段（health_summary、auto_analyze 等）
        2. 确保 disk_list 存在，否则用默认值填充
        3. 计算问题数量并根据数量评定健康状态（优秀/良好/一般/需关注）
        4. 优先使用 docxtpl 库进行模板渲染；若遇到 AttributeError 或其他异常，
           自动降级调用 _fallback_render() 直接构建报告

        :return: 保存成功返回 True，失败返回 False
        """
        try:
            required_keys = ['health_summary', 'auto_analyze', 'myversion', 'co_name', 'port', 'ip', 'system_info']
            for key in required_keys:
                if key not in self.context:
                    if key == 'health_summary':
                        self.context[key] = [{'health_summary': '运行良好'}]
                    elif key == 'auto_analyze':
                        self.context[key] = []
                    elif key == 'myversion':
                        self.context[key] = [{'version': 'Unknown'}]
                    elif key == 'system_info':
                        self.context[key] = {}
                    else:
                        self.context[key] = [{'placeholder': '数据缺失'}]

            if 'disk_list' not in self.context['system_info'] or not self.context['system_info']['disk_list']:
                self.context['system_info']['disk_list'] = [{
                    'device': '/dev/sda1', 'mountpoint': '/', 'fstype': 'ext4',
                    'total_gb': 50.0, 'used_gb': 25.0, 'free_gb': 25.0, 'usage_percent': 50.0
                }]

            list_keys = ['db_size', 'processlist', 'mysql_users', 'platform_info']
            for key in list_keys:
                if key not in self.context:
                    self.context[key] = []

            self.context.update({"report_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S')})
            self.context.update({"inspector_name": self.inspector_name})
            problem_count = len(self.context.get("auto_analyze", []))
            self.context.update({"problem_count": problem_count})

            if problem_count == 0:
                health_status = "优秀"
            elif problem_count <= 3:
                health_status = "良好"
            elif problem_count <= 6:
                health_status = "一般"
            else:
                health_status = "需关注"
            self.context.update({"health_status": health_status})

            # 尝试使用 docxtpl 正常渲染
            try:
                with open(self.ifile, 'rb') as f:
                    template_bytes = f.read()
                doc_stream = io.BytesIO(template_bytes)
                tpl = DocxTemplate(doc_stream)
                tpl.render(self.context)
                tpl.save(self.ofile)

                # ── 追加新章节（第7章 7.1/7.2 + 第8章 AI诊断）───────────────────
                # docxtpl 模板本身有旧的"7.报告说明"，先把它及之后的内容删掉，再追加新章节
                doc2 = Document(self.ofile)
                # 找到模板里旧的"7. 报告说明"段落位置，删掉它及之后的所有段落
                cutoff_idx = None
                for i, para in enumerate(doc2.paragraphs):
                    t = para.text.strip()
                    if t.startswith('7.') and '报告说明' in t:
                        cutoff_idx = i
                        break
                if cutoff_idx is not None:
                    # python-docx 不提供直接按索引批量删除段落，通过 element 操作
                    body = doc2._element.body
                    for para in list(body.iterchildren())[cutoff_idx:]:
                        body.remove(para)

                auto_analyze = self.context.get('auto_analyze', [])
                high_risk = [i for i in auto_analyze if i.get('col2') == '高风险']
                mid_risk  = [i for i in auto_analyze if i.get('col2') == '中风险']
                low_risk  = [i for i in auto_analyze if i.get('col2') in ('低风险', '建议')]

                # 第 7 章 风险与建议
                h7 = doc2.add_heading('7. 风险与建议', level=1)

                p = doc2.add_paragraph()
                p.add_run('本次共检测到 ')
                if high_risk:
                    r = p.add_run(f'{len(high_risk)} 项高风险'); r.bold = True; r.font.color.rgb = RGBColor(0xC0,0x00,0x00)
                if mid_risk:
                    r = p.add_run(f' {len(mid_risk)} 项中风险'); r.bold = True; r.font.color.rgb = RGBColor(0xFF,0x78,0x00)
                if low_risk:
                    r = p.add_run(f' {len(low_risk)} 项低风险/建议'); r.bold = True; r.font.color.rgb = RGBColor(0x37,0x86,0x10)
                p.add_run(f'，共计 {len(auto_analyze)} 项问题。')

                # 7.1 问题明细
                if auto_analyze:
                    doc2.add_heading('7.1 问题明细', level=2)
                    col_w = [Cm(0.8), Cm(3.2), Cm(1.5), Cm(4.0), Cm(1.0), Cm(1.5), Cm(4.0)]
                    tbl = doc2.add_table(rows=1+len(auto_analyze), cols=7)
                    tbl.style = 'Light Grid Accent 1'
                    hdrs = ['序号','风险项','等级','详细描述','优先级','负责人','修复建议']
                    for j,(cell,ht) in enumerate(zip(tbl.rows[0].cells, hdrs)):
                        cell.text = ht
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
                        cell.width = col_w[j]
                    for idx,item in enumerate(auto_analyze,1):
                        row = tbl.rows[idx].cells
                        row[0].text = str(idx)
                        row[1].text = item.get('col1','')
                        row[2].text = item.get('col2','')
                        row[3].text = item.get('col3','')
                        row[4].text = item.get('col4','')
                        row[5].text = item.get('col5','')
                        fix_sql = item.get('fix_sql','').strip()
                        row[6].text = fix_sql if fix_sql else '—'
                        for j,cell in enumerate(row):
                            for para in cell.paragraphs:
                                for run in para.runs: run.font.size = Pt(9)
                            cell.width = col_w[j]
                        lvl = item.get('col2','')
                        cm = {'高风险':RGBColor(0xC0,0x00,0x00),'中风险':RGBColor(0xFF,0x78,0x00),
                              '低风险':RGBColor(0x37,0x86,0x10),'建议':RGBColor(0x00,0x70,0xC0)}
                        if lvl in cm:
                            row[2].paragraphs[0].runs[0].font.color.rgb = cm[lvl]
                            row[2].paragraphs[0].runs[0].bold = True
                else:
                    doc2.add_paragraph('✅ 未发现明显风险项，MySQL 数据库运行状态良好。')

                # 7.2 修复速查
                fix_items = [i for i in auto_analyze if i.get('fix_sql','').strip()]
                if fix_items:
                    doc2.add_heading('7.2 修复速查', level=2)
                    for idx,item in enumerate(fix_items,1):
                        p = doc2.add_paragraph()
                        p.add_run(f'{idx}. [{item.get("col1")}] {item.get("col3","")[:60]}').bold = True
                        qp = doc2.add_paragraph(item.get('fix_sql','').strip())
                        qp.style = 'Quote'
                        if qp.runs: qp.runs[0].font.size = Pt(9)

                # 第 8 章 AI 智能诊断建议
                ai_advice = self.context.get('ai_advice','').strip()
                doc2.add_heading('8. AI 智能诊断建议', level=1)
                if ai_advice:
                    p = doc2.add_paragraph()
                    p.add_run('🤖 以下建议由 AI 大模型基于本次巡检数据自动生成，仅供参考，'
                              '实际操作请结合业务场景谨慎评估。').italic = True
                    doc2.add_paragraph()
                    for line in ai_advice.split('\n'):
                        line = line.strip()
                        if not line:
                            doc2.add_paragraph()
                        elif line.startswith(('- ','* ','• ')):
                            bp = doc2.add_paragraph(style='List Bullet')
                            bp.add_run(line[2:]).font.size = Pt(11)
                        else:
                            np = doc2.add_paragraph(line)
                            if np.runs: np.runs[0].font.size = Pt(11)
                else:
                    p = doc2.add_paragraph()
                    p.add_run('💡 AI 诊断未启用。如需开启，请在 Web UI「AI 诊断设置」中配置 AI 后端（支持 DeepSeek / OpenAI / Ollama）。').italic = True

                # 第 9 章 报告说明
                doc2.add_heading('9. 报告说明', level=1)
                notes = [
                    "1. 本报告基于 MySQL 数据库实时状态生成，反映了生成时刻的数据库健康状况",
                    "2. 报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用",
                    "3. 磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令 'df -h'",
                    "4. 巡检结果仅供参考，实际运维中请结合具体业务场景进行分析",
                    "5. 建议定期进行数据库巡检，及时发现并解决潜在问题",
                    f"6. AI 诊断功能（若启用）生成的建议仅供参考，不构成专业 DBA 意见"
                ]
                for note in notes:
                    doc2.add_paragraph(note)

                doc2.save(self.ofile)
                return True
            except AttributeError as ae:
                if 'part' in str(ae):
                    pass  # 静默降级到备用渲染
                    return self._fallback_render()
                else:
                    raise
            except Exception as e:
                pass  # 静默降级到备用渲染
                return self._fallback_render()

        except Exception as e:
            print(f"❌ 生成Word文档失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _fallback_render(self):
        """
        备用报告渲染方法 - 当 docxtpl 渲染失败时使用 python-docx 直接构建报告。

        不依赖模板文件，直接通过代码构建包含 8 个章节的完整 Word 文档：
        封面、健康状态概览、系统资源检查（CPU/内存/磁盘）、
        MySQL 配置检查（连接/内存/日志）、性能分析（QPS/锁/异常连接）、
        数据库信息（大小/进程列表）、安全信息（用户信息）、风险与建议、报告说明。

        所有数据均从 self.context 中直接提取，无 Jinja2 模板变量。

        :return: 渲染并保存成功返回 True，失败返回 False
        """
        try:
            doc = Document()
            title = doc.add_heading('MySQL数据库健康巡检报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(20)
            title_run.font.bold = True
            doc.add_paragraph()
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Light Grid Accent 1'
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(10)
            data_map = [
                ("数据库名称", self.context.get('co_name', [{}])[0].get('CO_NAME', 'N/A')),
                ("服务器地址", f"{self.context.get('ip', [{}])[0].get('IP', 'N/A')}:{self.context.get('port', [{}])[0].get('PORT', 'N/A')}"),
                ("MySQL版本", self.context.get('myversion', [{}])[0].get('version', 'N/A')),
                ("服务器主机名", self.context.get('system_info', {}).get('hostname', 'N/A')),
                ("实例启动时间", self.context.get('instancetime', [{}])[0].get('started_at', 'N/A') if self.context.get('instancetime') else 'N/A'),
                ("巡检人员", self.inspector_name),
                ("服务器平台", self._get_platform_info()),
                ("报告生成时间", self.context.get('report_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
            ]
            for i, (label, value) in enumerate(data_map):
                cells = table.rows[i].cells
                cells[0].text = label
                cells[1].text = str(value)
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
            doc.add_page_break()
            doc.add_heading('1. 健康状态概览', level=1)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(10)
            cells = table.rows[0].cells
            cells[0].text = "总体健康状态"
            cells[1].text = self.context.get('health_status', 'N/A')
            cells = table.rows[1].cells
            cells[0].text = "发现问题数量"
            cells[1].text = f"{self.context.get('problem_count', 0)} 个"
            doc.add_paragraph()
            p = doc.add_paragraph("健康总结: ")
            p.add_run(self.context.get('health_summary', [{}])[0].get('health_summary', '运行良好')).bold = True

            doc.add_heading('2. 系统资源检查', level=1)
            cpu = self.context.get('system_info', {}).get('cpu', {})
            mem = self.context.get('system_info', {}).get('memory', {})
            doc.add_heading('2.1 CPU信息', level=2)
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'CPU使用率', '物理核心数', '逻辑核心数', '当前频率(GHz)'
            row = table.rows[1].cells
            row[0].text = f"{cpu.get('usage_percent', 'N/A')}%"
            row[1].text = str(cpu.get('physical_cores', 'N/A'))
            row[2].text = str(cpu.get('logical_cores', 'N/A'))
            freq = cpu.get('current_frequency', 0)
            row[3].text = f"{freq/1000:.2f}" if isinstance(freq, (int, float)) and freq > 100 else str(freq)
            doc.add_paragraph()
            doc.add_heading('2.2 内存信息', level=2)
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '总内存(GB)', '已使用(GB)', '可用内存(GB)', '使用率'
            row = table.rows[1].cells
            row[0].text = f"{mem.get('total_gb', 'N/A')}"
            row[1].text = f"{mem.get('used_gb', 'N/A')}"
            row[2].text = f"{mem.get('available_gb', 'N/A')}"
            row[3].text = f"{mem.get('usage_percent', 'N/A')}%"
            doc.add_paragraph()
            doc.add_heading('2.3 磁盘信息', level=2)
            disk_list = self.context.get('system_info', {}).get('disk_list', [])
            table = doc.add_table(rows=1+len(disk_list), cols=2)
            table.style = 'Light Grid Accent 1'
            table.columns[0].width = Cm(8)
            table.columns[1].width = Cm(4)
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text = '挂载点', '使用率'
            for i, disk in enumerate(disk_list, 1):
                cells = table.rows[i].cells
                cells[0].text = disk.get('mountpoint', 'N/A')
                cells[1].text = f"{disk.get('usage_percent', 0):.2f}%"

            doc.add_heading('3. MySQL配置检查', level=1)
            doc.add_heading('3.1 连接配置', level=2)
            self._add_config_table(doc, [
                ('最大连接数', 'max_connections'), ('当前连接数', 'current_connections'),
                ('交互超时', 'interactive_timeout'), ('文件打开限制', 'open_files_limit')
            ])
            doc.add_heading('3.2 内存配置', level=2)
            self._add_config_table(doc, [
                ('InnoDB缓冲池', 'innodb_buffer_pool_size'), ('排序缓冲区', 'sort_buffer_size'),
                ('连接缓冲区', 'join_buffer_size'), ('线程缓存', 'thread_cache_size')
            ])
            doc.add_heading('3.3 日志配置', level=2)
            self._add_config_table(doc, [
                ('慢查询日志', 'slow_query_log'), ('Binlog保留天数', 'expire_logs_days'),
                ('InnoDB日志文件大小', 'innodb_log_file_size'), ('日志刷新设置', 'innodb_flush_log_at_trx_commit')
            ])

            doc.add_heading('4. 性能分析', level=1)
            doc.add_heading('4.1 QPS检查', level=2)
            self._add_config_table(doc, [('总查询数', 'queries')], col1_width=4, col2_width=10)
            doc.add_heading('4.2 锁信息', level=2)
            self._add_config_table(doc, [
                ('立即锁表', 'table_locks_immediate'), ('等待锁表', 'table_locks_waited')
            ], col1_width=4, col2_width=10)
            doc.add_heading('4.3 异常连接', level=2)
            aborted = self.context.get('aborted_connections', [])
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text = '异常类型', '值'
            cells = table.rows[1].cells
            cells[0].text = '异常客户端连接'
            cells[1].text = aborted[0]['Value'] if len(aborted) > 0 else 'N/A'
            cells = table.rows[2].cells
            cells[0].text = '异常连接尝试'
            cells[1].text = aborted[1]['Value'] if len(aborted) > 1 else 'N/A'

            doc.add_heading('5. 数据库信息', level=1)
            doc.add_heading('5.1 数据库大小', level=2)
            db_size = self.context.get('db_size', [])
            table = doc.add_table(rows=1+len(db_size), cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '数据库名', '表行数', '数据大小(MB)', '索引大小(MB)'
            for i, db in enumerate(db_size, 1):
                cells = table.rows[i].cells
                cells[0].text = str(db.get('Database_name', ''))
                cells[1].text = str(db.get('No_of_rows', ''))
                cells[2].text = str(db.get('Size_data_MB', ''))
                cells[3].text = str(db.get('Size_index_MB', ''))
            doc.add_heading('5.2 当前进程列表', level=2)
            proc = self.context.get('processlist', [])
            table = doc.add_table(rows=1+min(len(proc), 20), cols=6)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = 'ID', '用户', '数据库', '状态', '命令', '时间'
            for i, p in enumerate(proc[:20], 1):
                cells = table.rows[i].cells
                cells[0].text = str(p.get('Id', ''))
                cells[1].text = str(p.get('User', ''))
                cells[2].text = str(p.get('db', ''))
                cells[3].text = str(p.get('State', ''))
                cells[4].text = str(p.get('Command', ''))
                cells[5].text = str(p.get('Time', ''))

            doc.add_heading('6. 安全信息', level=1)
            doc.add_heading('6.1 数据库用户信息', level=2)
            users = self.context.get('mysql_users', [])
            table = doc.add_table(rows=1+len(users), cols=5)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = '用户名', '主机', '授权权限', '认证插件', '账户锁定'
            for i, u in enumerate(users, 1):
                cells = table.rows[i].cells
                cells[0].text = str(u.get('col1', ''))
                cells[1].text = str(u.get('col2', ''))
                cells[2].text = str(u.get('col3', ''))
                cells[3].text = str(u.get('col4', ''))
                cells[4].text = str(u.get('col5', ''))

            doc.add_heading('7. 风险与建议', level=1)

            # ── 7.1 概览统计 ──
            auto_analyze = self.context.get('auto_analyze', [])
            high_risk  = [i for i in auto_analyze if i.get('col2') == '高风险']
            mid_risk   = [i for i in auto_analyze if i.get('col2') == '中风险']
            low_risk   = [i for i in auto_analyze if i.get('col2') in ('低风险', '建议')]
            p = doc.add_paragraph()
            p.add_run('本次共检测到 ').bold = False
            if high_risk:
                run_h = p.add_run(f'{len(high_risk)} 项高风险')
                run_h.bold = True; run_h.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            if mid_risk:
                run_m = p.add_run(f' {len(mid_risk)} 项中风险')
                run_m.bold = True; run_m.font.color.rgb = RGBColor(0xFF, 0x78, 0x00)
            if low_risk:
                run_l = p.add_run(f' {len(low_risk)} 项低风险/建议')
                run_l.bold = True; run_l.font.color.rgb = RGBColor(0x37, 0x86, 0x10)
            p.add_run(f'，共计 {len(auto_analyze)} 项问题。')

            # ── 7.2 风险明细表格 ──
            if auto_analyze:
                doc.add_heading('7.1 问题明细', level=2)
                # 列：序号、风险项、等级、详细描述、优先级、负责人、修复建议
                col_widths = [Cm(0.8), Cm(3.2), Cm(1.5), Cm(4.0), Cm(1.0), Cm(1.5), Cm(4.0)]
                tbl = doc.add_table(rows=1 + len(auto_analyze), cols=7)
                tbl.style = 'Light Grid Accent 1'
                hdr = tbl.rows[0].cells
                headers = ['序号', '风险项', '等级', '详细描述', '优先级', '负责人', '修复建议']
                for j, (cell, hdr_text) in enumerate(zip(hdr, headers)):
                    cell.text = hdr_text
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.width = col_widths[j]
                for idx, item in enumerate(auto_analyze, 1):
                    row = tbl.rows[idx].cells
                    row[0].text = str(idx)
                    row[1].text = item.get('col1', '')
                    row[2].text = item.get('col2', '')
                    row[3].text = item.get('col3', '')
                    row[4].text = item.get('col4', '')
                    row[5].text = item.get('col5', '')
                    fix_sql = item.get('fix_sql', '').strip()
                    row[6].text = fix_sql if fix_sql else '—'
                    for j, cell in enumerate(row):
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)
                        cell.width = col_widths[j]
                    # 等级颜色
                    level = item.get('col2', '')
                    color_map = {'高风险': RGBColor(0xC0, 0x00, 0x00),
                                 '中风险': RGBColor(0xFF, 0x78, 0x00),
                                 '低风险': RGBColor(0x37, 0x86, 0x10),
                                 '建议':   RGBColor(0x00, 0x70, 0xC0)}
                    if level in color_map:
                        row[2].paragraphs[0].runs[0].font.color.rgb = color_map[level]
                        row[2].paragraphs[0].runs[0].bold = True
            else:
                doc.add_paragraph('✅ 未发现明显风险项，MySQL 数据库运行状态良好。')

            # ── 7.3 修复速查 ──
            fix_items = [i for i in auto_analyze if i.get('fix_sql', '').strip()]
            if fix_items:
                doc.add_heading('7.2 修复速查', level=2)
                for idx, item in enumerate(fix_items, 1):
                    p = doc.add_paragraph()
                    p.add_run(f'{idx}. [{item.get("col1")}] {item.get("col3")[:60]}').bold = True
                    code_p = doc.add_paragraph(item.get('fix_sql', '').strip())
                    code_p.style = 'Quote'
                    code_p.runs[0].font.size = Pt(9)

            # ── 8. AI 智能诊断建议 ──
            ai_advice = self.context.get('ai_advice', '').strip()
            doc.add_heading('8. AI 智能诊断建议', level=1)
            if ai_advice:
                p = doc.add_paragraph()
                p.add_run('🤖 以下建议由 AI 大模型基于本次巡检数据自动生成，仅供参考，实际操作请结合业务场景谨慎评估。').italic = True
                doc.add_paragraph()
                # 分段渲染：保留列表格式
                for line in ai_advice.split('\n'):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                    elif line.startswith(('- ', '* ', '• ')):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(line[2:]).font.size = Pt(11)
                    else:
                        doc.add_paragraph(line).runs[0].font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                p.add_run('💡 AI 诊断未启用。如需开启，请在 Web UI「AI 诊断设置」中配置 AI 后端（支持 DeepSeek / OpenAI / Ollama）。').italic = True

            # ── 9. 报告说明 ──
            doc.add_heading('9. 报告说明', level=1)
            notes = [
                "1. 本报告基于 MySQL 数据库实时状态生成，反映了生成时刻的数据库健康状况",
                "2. 报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用",
                "3. 磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令 'df -h'",
                "4. 巡检结果仅供参考，实际运维中请结合具体业务场景进行分析",
                "5. 建议定期进行数据库巡检，及时发现并解决潜在问题",
                f"6. AI 诊断功能（若启用）生成的建议仅供参考，不构成专业 DBA 意见"
            ]
            for note in notes:
                doc.add_paragraph(note)

            doc.save(self.ofile)
            pass  # 备用渲染成功
            return True
        except Exception as e:
            pass  # 备用渲染失败
            import traceback
            traceback.print_exc()
            return False

    def _get_platform_info(self):
        """
        从 context 中提取服务器操作系统平台信息。

        遍历 platform_info 列表，查找 variable_name 为 'version_compile_os' 的条目。

        :return: 操作系统编译平台字符串（如 "Linux"），未找到时返回 'N/A'
        """
        pf = self.context.get('platform_info', [])
        for item in pf:
            if item.get('variable_name') == 'version_compile_os':
                return item.get('variable_value', 'N/A')
        return 'N/A'

    def _add_config_table(self, doc, items, col1_width=4, col2_width=10):
        """
        向 Word 文档中添加配置键值表格（通用辅助方法）。

        创建「配置项 / 当前值」二列表格，并从 context 中取值填充数据行。
        数据不存在时填入 'N/A'。

        :param doc: 目标 Document 对象
        :param items: 配置项列表，每项为 (显示标签, context_key) 元组
        :param col1_width: 第 1 列宽度（cm），默认 4
        :param col2_width: 第 2 列宽度（cm），默认 10
        """
        table = doc.add_table(rows=1+len(items), cols=2)
        table.style = 'Light Grid Accent 1'
        table.columns[0].width = Cm(col1_width)
        table.columns[1].width = Cm(col2_width)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = '配置项', '当前值'
        for i, (label, key) in enumerate(items, 1):
            cells = table.rows[i].cells
            cells[0].text = label
            val = self.context.get(key, [])
            cells[1].text = val[0]['Value'] if val else 'N/A'

def print_banner():
    """
    打印程序启动横幅（彩色 ASCII Art）。
    """
    # ANSI 颜色代码（不支持时降级为普通文本）
    try:
        import shutil
        cols = shutil.get_terminal_size((80, 20)).columns
    except Exception:
        cols = 80

    CYAN   = "\033[96m"
    GREEN  = "\033[92m"
    YELLOW = "\033[93m"
    BOLD   = "\033[1m"
    DIM    = "\033[2m"
    RESET  = "\033[0m"

    # Windows 旧终端开启 ANSI 支持
    try:
        import os, ctypes
        if os.name == "nt":
            kernel32 = ctypes.windll.kernel32
            kernel32.SetConsoleMode(kernel32.GetStdHandle(-11), 7)
    except Exception:
        pass

    art = f"""
{CYAN}{BOLD}  ██████╗ ██████╗  ██████╗██╗  ██╗███████╗ ██████╗██╗  ██╗
  ██╔══██╗██╔══██╗██╔════╝██║  ██║██╔════╝██╔════╝██║ ██╔╝
  ██║  ██║██████╔╝██║     ███████║█████╗  ██║     █████╔╝
  ██║  ██║██╔══██╗██║     ██╔══██║██╔══╝  ██║     ██╔═██╗
  ██████╔╝██████╔╝╚██████╗██║  ██║███████╗╚██████╗██║  ██╗
  ╚═════╝ ╚═════╝  ╚═════╝╚═╝  ╚═╝╚══════╝ ╚═════╝╚═╝  ╚═╝{RESET}
{GREEN}{BOLD}              🐬  DBCheck - MySQL 巡检工具  {VER}{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
{YELLOW}  支持单机巡检 / 批量巡检 / Word报告 / SSH系统采集{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
"""
    print(art)

def check_license():
    """
    检查并验证许可证有效性。

    实例化 LicenseValidator 并调用 validate_license()：
    - 验证通过：打印成功消息后返回
    - 验证失败：提示用户是否重新创建许可证；
      重新创建仍失败则调用 sys.exit(1) 退出程序
    - 验证过程异常：打印警告但不退出，允许程序继续运行
    """
    try:
        validator = LicenseValidator()
        is_valid, message, remaining_days = validator.validate_license()
        if not is_valid:
            print(f"❌ {message}")
            print(f"📁 许可证文件位置: {validator.license_file}")
            retry = input("是否尝试重新创建许可证? (y/n) [y]: ").strip().lower()
            if retry in ['', 'y', 'yes']:
                validator._create_trial_license()
                is_valid, message, remaining_days = validator.validate_license()
                if is_valid:
                    print(f"✅ {message}")
                    return
                else:
                    print(f"❌ 重新创建许可证失败: {message}")
                    sys.exit(1)
            else:
                print("请联系管理员获取有效许可证")
                sys.exit(1)
        else:
            print(f"✅ {message}")
            print()
    except Exception as e:
        print(f"❌ 许可证检查异常: {e}")
        print("程序将继续运行，但部分功能可能受限")

def single_inspection():
    """
    执行单机巡检流程。

    调用 input_db_info() 进行交互式连接信息输入，
    输入有效后调用 run_inspection() 执行巡检并生成报告。
    """
    print("\n=== 单机巡检模式 ===")
    db_info = input_db_info()
    if not db_info:
        return
    run_inspection(db_info)

def batch_inspection():
    """
    执行批量巡检流程（从 Excel 文件导入数据库列表）。

    检查 Excel 模板文件是否存在，不存在时提示用户创建；
    读取模板后展示数据库列表并请用户确认，
    逐一调用 run_inspection() 完成每个数据库的巡检，
    最终汇总输出成功 / 失败统计。
    """
    print("\n=== 批量巡检模式 ===")
    excel_manager = ExcelTemplateManager()
    if not os.path.exists(excel_manager.template_file):
        print("❌ Excel模板文件不存在，请先创建模板")
        create_template = input("是否立即创建Excel模板? (y/n) [y]: ").strip().lower()
        if create_template in ['', 'y', 'yes']:
            excel_manager.create_template()
        return
    db_list = excel_manager.read_template()
    if not db_list:
        return
    print(f"\n📋 即将巡检以下 {len(db_list)} 个数据库:")
    for i, db in enumerate(db_list, 1):
        ssh_info = " (含系统信息)" if db.get('ssh_host') and (db.get('ssh_password') or db.get('ssh_key_file')) else ""
        print(f"  {i}. {db['name']} - {db['ip']}:{db['port']}{ssh_info}")
    confirm = input("\n是否开始批量巡检? (y/n) [y]: ").strip().lower()
    if confirm in ['', 'y', 'yes']:
        total_dbs = len(db_list)
        success_count = 0
        for i, db_info in enumerate(db_list, 1):
            print(f"\n[{i}/{total_dbs}] 开始巡检 {db_info['name']}...")
            if run_inspection(db_info):
                success_count += 1
        print(f"\n=== 批量巡检完成 ===")
        print(f"成功巡检: {success_count}/{total_dbs} 个数据库")
        print(f"报告输出目录: reports/")

def create_excel_template():
    """
    创建批量巡检 Excel 配置模板。

    实例化 ExcelTemplateManager 并调用其 create_template() 方法，
    在当前目录生成 mysql_batch_template.xlsx 文件。
    """
    print("\n=== 创建Excel配置模板 ===")
    excel_manager = ExcelTemplateManager()
    excel_manager.create_template()

def create_word_template(inspector_name="Jack"):
    """
    在系统临时目录创建 Word 巡检报告模板文件。

    实例化 WordTemplateGenerator，调用 create_template() 生成包含
    Jinja2 占位符的 Word 模板文档，并保存到系统临时目录。

    :return: 成功时返回模板文件的完整路径（字符串）；失败时返回 None
    """
    try:
        temp_dir = tempfile.gettempdir()
        template_path = os.path.join(temp_dir, "mysql_inspection_template.docx")
        generator = WordTemplateGenerator(inspector_name)
        doc = generator.create_template()
        doc.save(template_path)
        print(f"✅ Word模板已创建: {template_path}")
        return template_path
    except Exception as e:
        print(f"❌ 创建Word模板失败: {e}")
        import traceback
        traceback.print_exc()
        return None

def run_inspection(db_info):
    """
    对单个数据库实例执行完整的巡检流程并生成 Word 报告。

    主要步骤：
    1. 从 db_info 中提取连接参数和 SSH 信息
    2. 调用 create_word_template() 生成临时 Word 模板文件
    3. 在 reports/ 目录下以时间戳命名输出文件
    4. 预先测试 MySQL 连接并获取版本号
    5. 实例化 getData 执行巡检（checkdb）
    6. 将结果通过 saveDoc.contextsave() 渲染为最终报告
    7. 清理临时模板文件

    :param db_info: 包含数据库连接信息的字典，必须含 name、ip、port、user、password；
                    可选 ssh_host、ssh_port、ssh_user、ssh_password、ssh_key_file
    :return: 巡检并生成报告成功返回 True，任何环节失败返回 False
    """
    label_name = db_info['name']
    ip = db_info['ip']
    port = db_info['port']
    user = db_info['user']
    password = db_info['password']
    ssh_info = {}
    if 'ssh_host' in db_info and db_info['ssh_host']:
        ssh_info = {
            'ssh_host': db_info['ssh_host'],
            'ssh_port': db_info.get('ssh_port', 22),
            'ssh_user': db_info.get('ssh_user', 'root'),
            'ssh_password': db_info.get('ssh_password', ''),
            'ssh_key_file': db_info.get('ssh_key_file', '')
        }
    inspector_name = input("巡检人员（默认 Jack）: ").strip() or "Jack"
    ifile = create_word_template(inspector_name)
    if not ifile:
        return False
    dir_path = "reports"
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_name = f"MySQL巡检报告_{label_name}_{timestamp}.docx"
    ofile = os.path.join(dir_path, file_name)
    try:
        print(f"🔍 正在测试连接 {ip}:{port}...")
        conn_test = pymysql.connect(host=ip, port=int(port), user=user, password=password, connect_timeout=10, charset='utf8mb4')
        cursor = conn_test.cursor()
        cursor.execute("SELECT VERSION()")
        version = cursor.fetchone()[0]
        cursor.close()
        conn_test.close()
        print(f"📊 数据库版本: MySQL {version}")
    except Exception as e:
        print(f"❌ 连接测试失败: {e}")
        return False
    data = getData(ip, port, user, password, ssh_info)
    if data is None or data.conn_db2 is None:
        return False
    ret = data.checkdb('builtin')
    if not ret:
        return False
    ret.update({"co_name": [{'CO_NAME': label_name}]})
    ret.update({"port": [{'PORT': port}]})
    ret.update({"ip": [{'IP': ip}]})
    savedoc = saveDoc(context=ret, ofile=ofile, ifile=ifile, inspector_name=inspector_name)
    success = savedoc.contextsave()
    if success:
        print(f"✅ 报告已生成: {file_name}")
        try:
            if os.path.exists(ifile):
                os.remove(ifile)
        except:
            pass
        return True
    else:
        print(f"❌ 报告生成失败: {label_name}")
        return False

def main():
    """
    程序主入口函数。

    执行流程：
    1. 记录程序启动时间
    2. 打印横幅（print_banner）
    3. 验证许可证（check_license）
    4. 进入主菜单循环：
       - 选项 1：执行单机巡检（single_inspection）
       - 选项 2：执行批量巡检（batch_inspection）
       - 选项 3：创建 Excel 模板（create_excel_template）
       - 选项 4：退出程序
    5. 每次操作完成后询问是否返回主菜单
    6. 程序退出前打印总运行耗时
    """
    start_time = time.time()

    # 支持从主入口通过 --template 直接生成 Excel 模板
    if len(sys.argv) > 1 and sys.argv[1] == '--template':
        create_excel_template()
        return

    print_banner()
    # check_license()  # 许可证验证已屏蔽
    while True:
        choice = show_main_menu()
        if choice == '1':
            single_inspection()
        elif choice == '2':
            batch_inspection()
        elif choice == '3':
            create_excel_template()
        elif choice == '4':
            print("\n感谢使用 DBCheck MySQL 数据库巡检工具！")
            break
        if choice != '4':
            continue_choice = input("\n是否返回主菜单? (y/n) [y]: ").strip().lower()
            if continue_choice in ['', 'y', 'yes']:
                continue
            else:
                print("\n感谢使用 DBCheck MySQL 数据库巡检工具！")
                break
    end_time = time.time()
    print(f"\n程序运行总耗时: {end_time - start_time:.2f}秒")

if __name__ == '__main__':
    main()

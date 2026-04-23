#!/usr/bin/env python3
# -*- coding:utf-8 -*-
from version import __version__ as VER

"""
达梦 DM8 数据库自动化健康巡检工具 {VER}
支持 DM8 及以上版本
依赖: dmpython (pip install dmpython), python-docx, docxtpl, openpyxl, psutil, paramiko>=2.8,<2.10
注意: dmpython 需要达梦数据库自带的 dpi 动态库支持
"""

import warnings
warnings.filterwarnings("ignore")
import itertools
import math
import sys
import datetime
import argparse
import subprocess
import logging
import logging.handlers
import socket
import re
import time
from pathlib import Path
import sys, getopt, os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate
import configparser
import importlib
import subprocess
import json
import hashlib
import base64
from datetime import datetime, timedelta
import platform
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import tempfile
import io
import psutil
import shutil
import paramiko

# 达梦 DM8 驱动
try:
    import dmPython as dm_driver
    DM_DRIVER = 'dmPython'
except ImportError:
    print("缺少达梦数据库驱动库，请执行以下命令安装：")
    print("  pip install dmpython")
    print("  并确保 dpi 动态库（随达梦数据库安装）在系统 PATH 中")
    sys.exit(1)

importlib.reload(sys)

# ============================================================
# 磁盘采集时忽略的外接 ISO / Media 挂载点前缀
# ============================================================
IGNORE_MOUNTS = {'/mnt/iso', '/media', '/run/media', '/iso', '/cdrom'}

# ============================================================
# 内置 DM8 巡检 SQL 模板（覆盖 DM8 核心指标）
# ============================================================
DM8_SQL_TEMPLATES = """
[variables]
# ── 1. 基本信息（实测通过）───────────────────────
dm_version       = SELECT INSTANCE_NAME, HOST_NAME, SVR_VERSION, DB_VERSION, START_TIME, STATUS$, MODE$ FROM V$INSTANCE;
dm_database      = SELECT NAME, CREATE_TIME, ARCH_MODE, STATUS$, ROLE$, MAX_SIZE, TOTAL_SIZE FROM V$DATABASE;

# ── 2. 会话与连接 ──────────────────────────────
dm_sessions       = SELECT TOP 50 SESS_ID, SESS_SEQ, USER_NAME, STATE, TRX_ID, CREATE_TIME, CLNT_IP FROM V$SESSIONS WHERE NVL(USER_NAME,'_DUMMY_') NOT IN ('SYSDBA','SYS') ORDER BY STATE, SESS_SEQ DESC;
dm_session_count  = SELECT COUNT(*) AS TOTAL_ACTIVE FROM V$SESSIONS WHERE STATE='ACTIVE';
dm_session_detail  = SELECT USER_NAME, STATE, COUNT(*) AS CNT FROM V$SESSIONS GROUP BY USER_NAME, STATE ORDER BY CNT DESC;

# ── 3. 表空间使用情况 ─────────────────────────────
dm_tablespace     = SELECT T.TABLESPACE_NAME, T.STATUS AS TS_STATUS, D.FILE_NAME, D.BYTES/1024/1024 AS FILE_SIZE_MB, D.STATUS AS FILE_STATUS, D.AUTOEXTENSIBLE, D.MAXBYTES/1024/1024 AS MAX_SIZE_MB FROM DBA_TABLESPACES T LEFT JOIN DBA_DATA_FILES D ON T.TABLESPACE_NAME=D.TABLESPACE_NAME ORDER BY T.TABLESPACE_NAME;

# ── 4. 内存(Buffer/Pool) ───────────────────────
dm_memory         = SELECT NAME, PAGE_SIZE, N_PAGES, N_FIXED FROM V$BUFFERPOOL ORDER BY NAME;
dm_sga            = SELECT NAME, SUM(N_PAGES*PAGE_SIZE)/1024/1024 AS SIZE_MB FROM V$BUFFERPOOL GROUP BY NAME;

# ── 5. Redo/重做日志 ───────────────────────────────
dm_redo_logs      = SELECT GROUP_ID, FILE_ID, PATH, CREATE_TIME FROM V$RLOGFILE ORDER BY GROUP_ID;
dm_redo_curr      = SELECT CKPT_LSN, CUR_LSN, NEXT_SEQ FROM V$RLOG;

# ── 6. 归档配置 ──────────────────────────────
dm_arch_config    = SELECT NAME, VALUE FROM V$PARAMETER WHERE NAME IN ('ARCH_DIR', 'ARCH_FILE_SIZE', 'ARCH_MODE');
dm_arch_files     = SELECT DB_MAGIC, STATUS, LEN, FREE, ARCH_LSN FROM V$ARCH_FILE ORDER BY ARCH_LSN DESC LIMIT 20;

# ── 7. 备份信息 ──────────────────────────────
dm_backup         = SELECT TOP 20 BACKUP_NAME, LEVEL, BACKUP_PATH, BACKUP_TIME, DEVICE_TYPE, OBJECT_NAME FROM V$BACKUPSET ORDER BY BACKUP_TIME DESC;

# ── 8. 关键参数 ───────────────────────────────
dm_params         = SELECT NAME, TYPE, VALUE, SYS_VALUE FROM V$PARAMETER ORDER BY NAME LIMIT 50;

# ── 9. 无效对象 ────────────────────────────────
dm_invalid_objs   = SELECT OWNER, OBJECT_NAME, OBJECT_TYPE, STATUS FROM DBA_OBJECTS WHERE STATUS='INVALID' AND OWNER NOT IN ('SYS','SYSTEM') ORDER BY OWNER;
dm_invalid_cnt    = SELECT OWNER, COUNT(*) AS INVALID_COUNT FROM DBA_OBJECTS WHERE STATUS='INVALID' GROUP BY OWNER HAVING COUNT(*)>0 ORDER BY 2 DESC;

# ── 10. 用户与安全 ──────────────────────────────
dm_users          = SELECT USERNAME, ACCOUNT_STATUS, CREATED, PASSWORD_VERSIONS FROM DBA_USERS ORDER BY CREATED DESC;
dm_sys_privs      = SELECT GRANTEE, COUNT(*) AS PRIV_COUNT FROM DBA_SYS_PRIVS GROUP BY GRANTEE HAVING COUNT(*)>=5 ORDER BY 2 DESC;

# ── 11. 锁等待 ─────────────────────────────────
dm_lock_info      = SELECT TRX_ID, LTYPE, LMODE, BLOCKED FROM V$LOCK;
dm_trxwait        = SELECT ID, WAIT_FOR_ID, WAIT_TIME, THRD_ID, LOCK FROM V$TRXWAIT;

# ── 12. 事务 ─────────────────────────────────
dm_transactions   = SELECT ID, SESS_ID, STATUS, INS_CNT, DEL_CNT, UPD_CNT, START_LSN FROM V$TRX ORDER BY ID DESC LIMIT 30;

# ── 13. DMDSC 集群 ───────────────────────────────
dm_dcinfo         = SELECT 'DMDSC集群视图，暂不支持单实例环境' AS MSG;
dm_inst_info      = SELECT INSTANCE_NAME, INSTANCE_NUMBER, HOST_NAME, STATUS$, SVR_VERSION, START_TIME FROM V$INSTANCE;

# ── 14. 回收站 ───────────────────────────────
dm_recyclebin     = SELECT OWNER, OBJECT_NAME, ORIGINAL_NAME, TYPE, TS_NAME, CREATETIME, DROPTIME FROM DBA_RECYCLEBIN ORDER BY DROPTIME DESC LIMIT 20;

# ── 15. 数据文件状态 ───────────────────────────
dm_datafiles      = SELECT TABLESPACE_NAME, FILE_NAME, BYTES/1024/1024 AS FILE_SIZE_MB, STATUS, AUTOEXTENSIBLE, MAXBYTES/1024/1024 AS MAX_SIZE_MB FROM DBA_DATA_FILES ORDER BY TABLESPACE_NAME;

# ── 16. Profile 密码策略 ───────────────────────
dm_profile_pwd    = SELECT PROFILE, RESOURCE_NAME, RESOURCE_TYPE, LIMIT FROM DBA_PROFILES ORDER BY PROFILE, RESOURCE_NAME;

# ── 17. 等待事件 ──────────────────────────────
dm_top_waits      = SELECT CLASS_NAME, TOTAL_WAITS, TIME_WAITED FROM V$WAIT_CLASS ORDER BY TOTAL_WAITS DESC;
dm_wait_class     = SELECT EVENT, WAIT_CLASS, TOTAL_WAITS, TIME_WAITED FROM V$SYSTEM_EVENT WHERE WAIT_CLASS IS NOT NULL ORDER BY TIME_WAITED DESC LIMIT 30;

# ── 18. 统计信息陈旧 ───────────────────────────
dm_stale_stats    = SELECT TABLE_OWNER, TABLE_NAME, NUM_ROWS, LAST_ANALYZED, STALE_STATS FROM DBA_TAB_STATISTICS WHERE STALE_STATS='YES' AND TABLE_OWNER NOT IN ('SYS','SYSTEM') ORDER BY LAST_ANALYZED DESC LIMIT 20;

# ── 19. 分区表信息 ─────────────────────────────
dm_partition_info = SELECT TABLE_OWNER, TABLE_NAME, PARTITION_NAME, HIGH_VALUE, NUM_ROWS, LAST_ANALYZED, TABLESPACE_NAME FROM DBA_TAB_PARTITIONS WHERE TABLE_OWNER NOT IN ('SYS','SYSTEM') AND NUM_ROWS>0 ORDER BY TABLE_OWNER, TABLE_NAME LIMIT 15;

# ── 20. 资源限制 ─────────────────────────────
dm_resource       = SELECT NAME, TYPE, SPACE_LIMIT, SPACE_USED FROM V$RESOURCE_LIMIT;
"""


# ============================================================
# 远程系统信息收集器（与 main_oracle 共用逻辑，保持一致）
# ============================================================
class RemoteSystemInfoCollector:
    """远程系统信息收集器 - 通过SSH连接获取远程主机信息"""

    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None

    def connect(self):
        try:
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            self.ssh_client.connect(
                hostname=self.host, port=self.port,
                username=self.username, password=self.password,
                timeout=10, look_for_keys=False, allow_agent=False,
                disabled_algorithms={'pubkeys': ['ssh-rsa']}
            )
            return True
        except Exception as e:
            print(f"SSH连接失败 {self.host}:{self.port}: {e}")
            return False

    def disconnect(self):
        if self.ssh_client:
            try: self.ssh_client.close()
            except Exception: pass

    def _run(self, cmd):
        stdin, stdout, stderr = self.ssh_client.exec_command(cmd, timeout=20)
        return stdout.read().decode('utf-8', errors='ignore'), stderr.read().decode('utf-8', errors='ignore')

    def get_cpu_info(self):
        out, _ = self._run("cat /proc/cpuinfo | grep 'model name' | head -1 && top -bn1 | grep 'Cpu' | awk '{print $2}'")
        info = {'Model name': '', 'usage_percent': 0.0}
        lines = out.strip().split('\n')
        if lines:
            info['Model name'] = lines[0].replace('model name', '').strip()
        if len(lines) > 1:
            try: info['usage_percent'] = float(lines[1].replace('%us',''))
            except: pass
        return info

    def get_memory_info(self):
        out, _ = self._run("free -m | grep Mem")
        info = {'total_mb': 0, 'used_mb': 0, 'usage_percent': 0.0}
        try:
            parts = out.strip().split()
            if len(parts) >= 3:
                info['total_mb'] = float(parts[1])
                info['used_mb'] = float(parts[2])
                if info['total_mb'] > 0:
                    info['usage_percent'] = round(info['used_mb'] / info['total_mb'] * 100, 1)
        except Exception: pass
        return info

    def get_disk_info(self):
        out, _ = self._run("df -h | grep -vE 'tmpfs|udev|overlay|mnt/iso|/iso|/media/|/run/media/|/cdrom|Filesystem'")
        disks = {}
        for line in out.split('\n')[1:]:
            parts = line.split()
            if len(parts) >= 6:
                device, size, used, avail, use_pct, mountpoint = parts[0], parts[1], parts[2], parts[3], parts[4], parts[5]
                try:
                    total_gb = float(size.replace('G','').replace('M',''))
                    used_gb = float(used.replace('G','').replace('M',''))
                    free_gb = float(avail.replace('G','').replace('M',''))
                    usage = float(use_pct.replace('%',''))
                    if 'M' in avail:
                        free_gb /= 1024; used_gb /= 1024; total_gb /= 1024
                    disks[mountpoint] = {
                        'device': device, 'mountpoint': mountpoint,
                        'total_gb': round(total_gb, 2), 'used_gb': round(used_gb, 2),
                        'free_gb': round(free_gb, 2), 'usage_percent': usage,
                        'fstype': 'unknown'
                    }
                except (ValueError, IndexError): continue
        return disks

    def get_system_info(self):
        import re
        hostname_out, _ = self._run("hostname")
        boot_out, _ = self._run("uptime -s")
        platform_out, _ = self._run("uname -sr")
        os_release_out, _ = self._run("cat /etc/os-release 2>/dev/null || uname -a")
        # 从 /etc/os-release 解析 PRETTY_NAME（如 "CentOS Linux 7.9.2009"）
        platform_text = '未知'
        if os_release_out:
            m = re.search(r'PRETTY_NAME="([^"]+)"', os_release_out)
            if m:
                platform_text = m.group(1)
            elif os_release_out.strip():
                # fallback：使用 uname -a 前3个字段
                parts = os_release_out.strip().split(maxsplit=3)
                platform_text = ' '.join(parts[:3]) if len(parts) >= 3 else os_release_out.strip()
        return {
            'hostname': hostname_out.strip() if hostname_out.strip() else '未知',
            'platform': platform_out.strip() if platform_out.strip() else 'Linux',
            'platform_text': platform_text,
            'boot_time': boot_out.strip() if boot_out.strip() else '未知',
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info()
        }


class LocalSystemInfoCollector:
    """本地系统信息收集器"""

    def get_cpu_info(self):
        info = {'Model name': '', 'usage_percent': 0.0}
        try:
            import platform as _pf
            info['Model name'] = f"{_pf.processor()} ({_pf.machine()})"
            info['usage_percent'] = round(psutil.cpu_percent(interval=1), 1)
        except Exception: pass
        return info

    def get_memory_info(self):
        info = {'total_mb': 0, 'used_mb': 0, 'usage_percent': 0.0}
        try:
            mem = psutil.virtual_memory()
            info['total_mb'] = round(mem.total / 1024 / 1024, 2)
            info['used_mb'] = round(mem.used / 1024 / 1024, 2)
            info['usage_percent'] = round(mem.percent, 1)
        except Exception: pass
        return info

    def get_disk_info(self):
        disks = {}
        try:
            for part in psutil.disk_partitions():
                try:
                    usage = psutil.disk_usage(part.mountpoint)
                    if part.mountpoint in IGNORE_MOUNTS:
                        continue
                    disks[part.mountpoint] = {
                        'device': part.device,
                        'mountpoint': part.mountpoint,
                        'fstype': part.fstype,
                        'total_gb': round(usage.total / 1024**3, 2),
                        'used_gb': round(usage.used / 1024**3, 2),
                        'free_gb': round(usage.free / 1024**3, 2),
                        'usage_percent': round(usage.percent, 1),
                    }
                except (PermissionError, OSError): continue
        except Exception: pass
        return disks

    def get_system_info(self):
        import socket as _sock
        import platform as _pf
        boot_time = ''
        try: boot_time = datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')
        except: pass
        return {
            'hostname': _sock.gethostname(),
            'platform': f"{_pf.system()} {_pf.release()} {_pf.machine()}",
            'platform_text': f"{_pf.system()} {_pf.release()} ({_pf.machine()})",
            'boot_time': boot_time,
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info()
        }


def get_host_disk_usage():
    """获取本地磁盘使用情况（Windows 兼容）"""
    disks = []
    try:
        for part in psutil.disk_partitions():
            if part.mountpoint in IGNORE_MOUNTS: continue
            try:
                usage = psutil.disk_usage(part.mountpoint)
                disks.append({
                    'device': part.device, 'mountpoint': part.mountpoint,
                    'fstype': part.fstype,
                    'total_gb': round(usage.total / 1024**3, 2),
                    'used_gb': round(usage.used / 1024**3, 2),
                    'free_gb': round(usage.free / 1024**3, 2),
                    'usage_percent': round(usage.percent, 1),
                })
            except (PermissionError, OSError): continue
    except Exception: pass
    return disks


# ============================================================
# 权限友好型 SQL 执行器
# ============================================================

# DM8 错误码 → 中文建议映射
_DM_ERROR_HINTS = {
    "-2653": ("密码错误或账户锁定", "请检查用户名密码，或确认账户是否被锁定/过期"),
    "-2651": ("连接失败 (DM-2651)", "请检查 host/port 是否正确，DM 服务是否启动"),
    "-3401": ("表/视图不存在", "该对象在当前达梦版本中可能名称不同，需适配 SQL"),
    "-3403": ("字段无效", "SQL 中使用的列名在当前版本不可用，需参考达梦文档适配"),
    "-3410": ("SQL 命令未正确结束", "SQL 语法错误，通常由不兼容的分号或语法问题引起"),
    "-3412": ("权限不足", "请使用 SYSDSDBA 或 SYSSSO 身份连接，或授予相应权限"),
    "-3415": ("字符集不匹配", "客户端与服务器字符集不一致，请检查 NLS 设置"),
    "-3429": ("网络通信错误", "检查防火墙设置，或尝试 ping/telnet 验证网络连通性"),
    "-3431": ("连接超时", "请检查 host/port 是否正确，DM 服务是否正常监听"),
}


def execute_query_safe(cursor, sql, item_name=""):
    """
    安全执行 SQL，统一捕获所有 DM8 错误并返回空结果（不中断流程）。
    """
    try:
        cursor.execute(sql)
        columns = [col[0] for col in cursor.description]
        data = []
        for row in cursor.fetchall():
            data.append(dict(zip(columns, row)))
        return {"columns": columns, "data": data}

    except dm_driver.Error as e:
        err_code = str(e.args[0]) if e.args else "0"
        hint = _DM_ERROR_HINTS.get(err_code)
        if hint:
            print(f"[WARN]  [{item_name}] {hint[0]}\n    💡 建议: {hint[1]}")
        else:
            print(f"[WARN]  [{item_name}] 数据库错误 ({err_code}): {str(e)[:120]}")
        return {"columns": [], "data": []}

    except Exception as e:
        err_type = type(e).__name__
        print(f"[WARN]  [{item_name}] {err_type}: {str(e)[:100]}")
        return {"columns": [], "data": []}


# ============================================================
# 智能健康评分算法（适配 DM8）
# ============================================================

def analyze_health_status(context):
    """
    综合评估 DM8 数据库健康状况。
    """
    alerts_critical = []
    alerts_warning = []

    # ── 1. 表空间使用率 ──────────────────────
    ts_list = context.get('dm_tablespace', [])
    if ts_list and isinstance(ts_list, list):
        for ts in ts_list:
            if not isinstance(ts, dict): continue
            used_pct = _safe_float_val(ts.get('USED_PCT') or ts.get('used_pct', 0))
            name = ts.get('TABLESPACE_NAME', ts.get('tablespace_name', '?'))
            total_bytes = _safe_float_val(ts.get('TOTAL_BYTES') or ts.get('total_bytes', 0))
            free_bytes = _safe_float_val(ts.get('FREE_BYTES') or ts.get('free_bytes', 0))
            total_mb = total_bytes / 1024 / 1024
            free_mb = free_bytes / 1024 / 1024

            if used_pct >= 95:
                alerts_critical.append(
                    f"[紧急] 表空间 {name} 使用率 {used_pct:.1f}% "
                    f"(总计 {total_mb:.0f}MB，剩余 {free_mb:.0f}MB)，需立即扩容"
                )
            elif used_pct >= 85:
                alerts_warning.append(
                    f"[关注] 表空间 {name} 使用率 {used_pct:.1f}% "
                    f"(剩余 {free_mb:.0f}MB)，建议提前规划扩容"
                )

    # ── 2. 会话数接近上限 ────────────────────
    sess = context.get('dm_sessions', [])
    limit = context.get('dm_session_limit', [])
    if sess and limit and isinstance(sess, list) and isinstance(limit, list):
        try:
            total = _safe_int_val(sess[0].get('TOTAL_SESSIONS', 0))
            max_proc = 0
            for lim in limit:
                if isinstance(lim, dict) and lim.get('NAME', '') == 'processes':
                    max_proc = _safe_int_val(lim.get('VALUE', 0))
                    break
            if max_proc > 0 and total > 0:
                pct = total * 100 / max_proc
                if pct >= 95:
                    alerts_critical.append(f"[紧急] 会话数 {total}/{max_proc} ({pct:.0f}%)，即将耗尽资源")
                elif pct >= 85:
                    alerts_warning.append(f"[关注] 会话数 {total}/{max_proc} ({pct:.0f}%)，接近上限")
        except (IndexError, KeyError, TypeError):
            pass

    # ── 3. 锁等待 ───────────────────────
    blocked = context.get('dm_blocked', [])
    if blocked and isinstance(blocked, list):
        lock_cnt = len(blocked)
        if lock_cnt >= 5:
            alerts_critical.append(f"[紧急] 发现 {lock_cnt} 个锁等待，可能导致业务阻塞")
        elif lock_cnt >= 2:
            alerts_warning.append(f"[关注] 发现 {lock_cnt} 个锁等待，建议排查")

    # ── 4. 无效对象 ──────────────────────
    invalid = context.get('dm_invalid_cnt', [])
    if invalid and isinstance(invalid, list):
        total_invalid = sum(_safe_int_val(iv.get('INVALID_COUNT', 0)) for iv in invalid if isinstance(iv, dict))
        if total_invalid >= 20:
            alerts_warning.append(f"[关注] 存在 {total_invalid} 个无效对象，建议编译或清理")
        elif total_invalid > 0:
            pass

    # ── 5. 系统内存 ──────────────────────
    sys_info = context.get('system_info', {})
    mem = sys_info.get('memory', {}) if isinstance(sys_info, dict) else {}
    if isinstance(mem, dict):
        mem_pct = _safe_float_val(mem.get('usage_percent', 0))
        if mem_pct >= 95:
            alerts_critical.append(f"[紧急] 系统内存使用率 {mem_pct:.1f}%，存在 OOM 风险")
        elif mem_pct >= 90:
            alerts_warning.append(f"[关注] 系统内存使用率 {mem_pct:.1f}%")

    # ── 6. 磁盘空间 ──────────────────────
    disks = sys_info.get('disk_list', []) if isinstance(sys_info, dict) else []
    if disks and isinstance(disks, list):
        for d in disks:
            if not isinstance(d, dict): continue
            mp = d.get('mountpoint', '/')
            if mp in IGNORE_MOUNTS: continue
            usage = _safe_float_val(d.get('usage_percent', 0))
            if usage >= 98:
                alerts_critical.append(f"[紧急] 磁盘 {mp} 使用率 {usage:.1f}%，即将写满")
            elif usage >= 90:
                alerts_warning.append(f"[关注] 磁盘 {mp} 使用率 {usage:.1f}%")

    # ── 综合评分 ─────────────────────────────
    critical_n = len(alerts_critical)
    warning_n = len(alerts_warning)
    score = 100.0 - critical_n * 15 - warning_n * 5
    score = max(0, min(100, score))

    if critical_n > 0:
        status = "严重"
    elif warning_n >= 4:
        status = "警告"
    elif warning_n > 0:
        status = "一般"
    else:
        status = "正常"

    return {
        "status": status,
        "score": round(score, 1),
        "critical_count": critical_n,
        "warning_count": warning_n,
        "alerts": alerts_critical + alerts_warning
    }


def _safe_float_val(val, default=0.0):
    try:
        if val is None: return default
        return float(str(val).replace(',', '').replace('%', '').strip())
    except (ValueError, TypeError):
        return default


def _safe_int_val(val, default=0):
    try:
        if val is None: return default
        return int(str(val).replace(',', '').strip())
    except (ValueError, TypeError):
        return default


# ============================================================
# Word 模板创建
# ============================================================
def create_word_template(inspector_name="Jack"):
    """
    创建 DM8 巡检报告的 Word 模板（基于 docxtpl Jinja2 模板）。
    所有变量均使用 {{{ var }}} 占位符。
    """
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
    tpl_file = os.path.join(template_path, "dm8_wordtemplates_v1.0.docx")

    if os.path.exists(tpl_file):
        return tpl_file

    from docx import Document
    doc = Document()

    title_style = doc.styles.add_style('ReportTitle', 1)
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.name = '微软雅黑'
    title_style.font.color.rgb = RGBColor(0,51,102)

    heading1 = doc.styles.add_style('Heading1Custom', 1)
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.font.name = '微软雅黑'
    heading1.font.color.rgb = RGBColor(0,51,102)

    heading2 = doc.styles.add_style('Heading2Custom', 1)
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.name = '微软雅黑'

    # 封面标题（二号）
    doc.add_paragraph("达梦数据库健康巡检报告", style='ReportTitle')
    doc.add_paragraph("")

    # 封面信息表
    info_table = doc.add_table(rows=8, cols=2)
    info_cells = [
        ("数据库名称",     "{{{ co_name }}}"),
        ("服务器地址",     "{{{ server_addr }}}"),
        ("达梦版本",       "{{{ dm_version }}}"),
        ("服务器主机名",   "{{{ hostname }}}"),
        ("实例启动时间",   "{{{ uptime_text }}}"),
        ("巡检人员",       inspector_name),
        ("服务器平台",     "{{{ platform_text }}}"),
        ("报告生成时间",   "{{{ report_time }}}"),
    ]
    for i, (label, value) in enumerate(info_cells):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    info_table.style = 'Table Grid'

    # 章节占位符
    doc.add_paragraph("\n第1章 数据库基本信息", style='Heading1Custom')
    doc.add_paragraph("{{{ dm_instance }}}")

    doc.add_paragraph("\n第2章 巡检执行摘要", style='Heading1Custom')
    doc.add_paragraph("{{{ health_summary_text }}}")
    doc.add_paragraph("{{{ error_summary_text }}}")

    doc.add_paragraph("\n第3章 表空间使用情况", style='Heading1Custom')
    doc.add_paragraph("{{{ dm_tablespace }}}")

    doc.add_paragraph("\n第4章 会话与事务", style='Heading1Custom')
    doc.add_paragraph("{{{ dm_sessions }}}")
    doc.add_paragraph("{{{ dm_trx }}}")

    doc.add_paragraph("\n第5章 内存分析", style='Heading1Custom')
    doc.add_paragraph("{{{ dm_sga }}}")
    doc.add_paragraph("{{{ dm_memory_info }}}")

    doc.add_paragraph("\n第6章 重做日志与归档", style='Heading1Custom')
    doc.add_paragraph("{{{ dm_redo_logs }}}")

    doc.add_paragraph("\n第7章 系统资源监控", style='Heading1Custom')
    doc.add_paragraph("{{{ system_info_text }}}")

    doc.add_paragraph("\n第8章 对象与用户安全", style='Heading1Custom')
    doc.add_paragraph("8.1 无效对象\n{{{ dm_invalid_cnt }}}")
    doc.add_paragraph("8.2 用户列表\n{{{ dm_users }}}")

    doc.add_paragraph("\n第9章 备份与归档", style='Heading1Custom')
    doc.add_paragraph("{{{ dm_backup }}}")

    doc.add_paragraph("\n第10章 报告说明", style='Heading1Custom')
    doc.add_paragraph("{{{ notes_text }}}")

    os.makedirs(template_path, exist_ok=True)
    doc.save(tpl_file)
    return tpl_file


# ============================================================
# DM8 数据采集类
# ============================================================
class getData(object):
    """数据采集类 - 负责连接 DM8 数据库并执行全量巡检 SQL"""

    def __init__(self, ip, port, user, password, db_name=None, ssh_info=None):
        self.H = ip
        self.P = int(port)
        self.user = user
        self.password = password
        self.db_name = db_name or ip
        self.ssh_info = ssh_info or {}
        self.conn_db = None
        self.context = {}
        self._connect()

    def _connect(self):
        import re
        def get_dm_code(err_str):
            m = re.search(r'CODE:(-?\d+)', err_str)
            return m.group(1) if m else 'unknown'

        def friendly_hint(code, host, port, user):
            hints = {
                '-2501': f'用户名或密码错误，请确认 SYSDBA 用户密码是否正确',
                '-2111': f'连接 {host}:{port} 超时/通信异常，请检查：\n'
                          f'  1) 确认 {host} 是否可达（ping {host}）\n'
                          f'  2) 确认 DM8 服务已启动且端口为 {port}\n'
                          f'  3) 确认防火墙已开放 {port} 端口\n'
                          f'  4) 确认 {host} 上 DM8 已正常监听该端口（telnet {host} {port}）',
                '-70028': f'无法连接到 {host}:{port}，请检查主机是否可达、端口是否正确、DM8服务是否启动',
                '-70026': f'用户 {user} 不存在或密码错误',
                '-70013': f'连接 {host}:{port} 超时，请检查网络',
                '-70201': f'连接被拒绝，请检查端口 {port} 是否为 DM8 端口、DM8 服务是否启动',
                '-70019': f'连接被拒绝，端口 {port} 无服务监听，请确认 DM8 监听地址',
            }
            return hints.get(code, f'错误码 CODE:{code}，请检查连接参数')

        try:
            # dmPython 使用 user/password@host:port 格式，无 database 参数
            # user 即为连接的 schema，port 默认 5236
            conn = dm_driver.connect(
                user=self.user,
                password=self.password,
                server=self.H,
                port=self.P
            )
            # dmPython 是"懒连接"，connect() 不抛异常，返回的 conn 内部可能装着异常
            # 第一次使用 cursor 时异常才真正冒出来，在这里统一触发并捕获
            # dmPython 是懒连接，connect() 不抛异常，真实错误在第一次使用 cursor 时才出现
            # conn/cur 对象内部装着异常，用 try/except 捕获 fetchall() 来暴露它
            probe_ok = False
            probe_error = None
            for probe_sql in [
                "SELECT STATUS FROM V$INSTANCE",
                "SELECT 1 FROM DUAL",
                "SELECT 1",
            ]:
                try:
                    cur = conn.cursor()
                    cur.execute(probe_sql)
                    cur.fetchall()
                    cur.close()
                    probe_ok = True
                    break
                except Exception as e:
                    probe_error = e
                    err_str = str(e)
                    code = get_dm_code(err_str)
                    # 表/视图不存在(-2605/-6205/-2665)为非致命错误，继续试下一个探测 SQL
                    if code in ('-2605', '-6205', '-2665'):
                        try: cur.close()
                        except: pass
                        continue
                    # 其他错误（含密码错误/连接拒绝/超时等）立即终止
                    try: cur.close()
                    except: pass
                    break

            if probe_ok:
                self.conn_db = conn
                print(f"[OK] DM8 连接成功: {self.user}@{self.H}:{self.P}")
            else:
                err_str = str(probe_error) if probe_error else ''
                code = get_dm_code(err_str)
                hint = friendly_hint(code, self.H, self.P, self.user)
                # -2111 超时错误：自动重试2次（网络偶发抖动时有效）
                if code == '-2111':
                    for retry_i in range(2):
                        time.sleep(3)
                        print(f"  [RETRY] 第{retry_i+2}次连接重试 ({self.H}:{self.P})...")
                        try:
                            conn2 = dm_driver.connect(user=self.user, password=self.password,
                                                      server=self.H, port=self.P)
                            cur2 = conn2.cursor()
                            cur2.execute("SELECT 1")
                            cur2.fetchall()
                            cur2.close()
                            self.conn_db = conn2
                            try: conn.close()
                            except: pass
                            print(f"[OK] DM8 连接成功（第{retry_i+2}次重试）: {self.user}@{self.H}:{self.P}")
                            probe_ok = True
                            break
                        except Exception as retry_err:
                            err_r = str(retry_err)
                            code_r = get_dm_code(err_r)
                            print(f"  重试失败 CODE:{code_r}: {err_r[:80]}")
                            try: conn2.close()
                            except: pass
                            continue
                if not probe_ok:
                    print(f"[FAIL] DM8 连接失败[CODE:{code}]: {hint}")
                    try: conn.close()
                    except: pass
                    self.conn_db = None

        except Exception as e:
            err_str = str(e)
            code = get_dm_code(err_str)
            hint = friendly_hint(code, self.H, self.P, self.user)
            print(f"[FAIL] DM8 连接失败[CODE:{code}]: {hint}")
            self.conn_db = None

    def print_progress_bar(self, iteration, total, prefix='', suffix='', decimals=1, length=50, fill='█'):
        percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
        filled_length = int(length * iteration // total)
        bar = fill * filled_length + '-' * (length - filled_length)
        print(f'\r{prefix} |{bar}| {percent}% {suffix}', end='\r')
        if iteration == total:
            print()

    def checkdb(self, sqlfile=''):
        print("\n开始 DM8 巡检...")
        total_steps = 22
        current_step = 0
        cfg = configparser.RawConfigParser()
        try:
            if not sqlfile or sqlfile == 'builtin':
                cfg.read_string(DM8_SQL_TEMPLATES)
            else:
                cfg.read(sqlfile, encoding='utf-8')
        except Exception as e:
            print(f"[FAIL] 读取SQL模板失败: {e}")
            return self.context

        init_keys = [
            "dm_version", "dm_instance", "dm_database", "dm_uptime",
            "dm_sessions", "dm_session_limit", "dm_blocked", "dm_trx",
            "dm_tablespace", "dm_temp_ts",
            "dm_sga", "dm_memory_info", "dm_pga",
            "dm_redo_logs", "dm_redo_curr",
            "dm_archive_config", "dm_archive_lag",
            "dm_backup",
            "dm_params",
            "dm_invalid_objs", "dm_invalid_cnt",
            "dm_users", "dm_sys_privs", "dm_default_pws",
            "dm_long_sql", "dm_top_sql_cpu",
            "dm_rss_status", "dm_rss_apply",
            "dm_dcinfo", "dm_inst_info",
            "dm_undo_info", "dm_transactions",
            "dm_recyclebin",
            "dm_datafiles",
            "dm_profile_pwd",
            "dm_top_waits", "dm_wait_class",
            "dm_stale_stats",
            "dm_partition_info"
        ]
        for key in init_keys:
            self.context.update({key: []})

        # ── 步骤1: 获取版本 ──────────────────────────
        try:
            cursor_ver = self.conn_db.cursor()
            cursor_ver.execute("SELECT BANNER FROM V$VERSION WHERE ROWNUM=1")
            ver_row = cursor_ver.fetchone()
            dm_version = ver_row[0] if ver_row else "Unknown"
            cursor_ver.close()
            self.context.update({"dm_version": [{'BANNER': dm_version}]})
            self.context.update({"health_summary": [{'health_summary': '运行良好'}]})
            self.context.update({"co_name": [{'DB_NAME': self.db_name}]})
        except Exception as e:
            print(f"[FAIL] 获取版本失败: {e}")
            self.context.update({"dm_version": [{'BANNER': 'Unknown'}]})

        # ── 步骤2-21: 执行所有 SQL（容错执行器，单个失败不中断） ─
        try:
            cursor = self.conn_db.cursor()
            variables_items = list(cfg.items("variables"))
            for i, (name, stmt) in enumerate(variables_items):
                current_step = int((i + 1) / len(variables_items) * (total_steps - 6)) + 1
                self.print_progress_bar(current_step, total_steps, prefix='DM8巡检:', suffix=f'{name} ({i+1}/{len(variables_items)})')
                clean_sql = stmt.replace('\n', ' ').replace('\r', ' ').rstrip().rstrip(';').strip()
                result = execute_query_safe(cursor, clean_sql, item_name=name)
                self.context[name] = result.get('data', [])
                time.sleep(0.03)
            cursor.close()
        except Exception as e:
            print(f'\n[FAIL] 数据库查询循环异常: {e}')

        # 容错执行结果存入 context
        self.context['_safe_errors'] = getattr(execute_query_safe, '_errors', [])

        # ── 步骤: 收集系统信息 ─────────────────────
        current_step = total_steps - 4
        self.print_progress_bar(current_step, total_steps, prefix='DM8巡检:', suffix='收集系统信息')
        try:
            if self.ssh_info and self.ssh_info.get('ssh_host'):
                print(f"[INFO] 通过SSH收集系统信息: {self.ssh_info['ssh_host']}")
                collector = RemoteSystemInfoCollector(
                    host=self.ssh_info['ssh_host'], port=self.ssh_info.get('ssh_port', 22),
                    username=self.ssh_info.get('ssh_user', 'root'),
                    password=self.ssh_info.get('ssh_password'), key_file=self.ssh_info.get('ssh_key_file')
                )
                if not collector.connect():
                    print("[WARN] SSH 连接失败，跳过远程系统信息采集")
                    collector = LocalSystemInfoCollector()
            else:
                collector = LocalSystemInfoCollector()
            system_info = collector.get_system_info()
            disk_list = system_info.get('disk_list') or system_info.get('disk') or get_host_disk_usage()
            if isinstance(disk_list, dict):
                disk_list = list(disk_list.values())
            system_info['disk_list'] = disk_list
            self.context.update({"system_info": system_info})
        except Exception as e:
            print(f"\n[FAIL] 收集系统信息失败: {e}")
            self.context.update({"system_info": {
                'hostname': '未知', 'platform': '未知', 'boot_time': '未知',
                'cpu': {}, 'memory': {},
                'disk_list': [{'device':'C:','mountpoint':'C:\\','fstype':'NTFS',
                               'total_gb':0,'used_gb':0,'free_gb':0,'usage_percent':0}]
            }})

        # ── 步骤: 风险分析 ─────────────────────────
        current_step = total_steps - 3
        self.print_progress_bar(current_step, total_steps, prefix='DM8巡检:', suffix='智能健康评估')
        self.context.update({"auto_analyze": []})
        self._basic_risk_check()

        # 始终执行健康评分
        health = analyze_health_status(self.context)
        self.context['health_analysis'] = health

        # 将告警合并到 auto_analyze
        existing_items = {a.get('col3', '') for a in self.context.get('auto_analyze', [])}
        for alert in health['alerts']:
            level = '高' if alert.startswith('[紧急]') else '中'
            title = alert.split(']', 1)[1].strip() if ']' in alert else alert[:50]
            if alert not in existing_items:
                self.context['auto_analyze'].append({
                    'col1': title,
                    'col2': f'{level}风险',
                    'col3': alert.replace('[紧急] ', '').replace('[关注] ', ''),
                    'col4': level,
                    'col5': 'DBA/系统管理员',
                    'fix_sql': ''
                })

        # 更新整体健康状态
        if health['status'] in ('严重',):
            self.context.update({"health_status": "需紧急处理"})
        elif health['status'] == '警告':
            self.context.update({"health_status": "需关注"})
        elif health['status'] == '一般':
            self.context.update({"health_status": "良好"})
        else:
            problem_count = len(self.context.get("auto_analyze", []))
            if problem_count == 0:
                self.context.update({"health_status": "优秀"})
            elif problem_count <= 3:
                self.context.update({"health_status": "良好"})

        # ── 步骤: AI 诊断 ─────────────────────────
        current_step = total_steps - 2
        self.print_progress_bar(current_step, total_steps, prefix='DM8巡检:', suffix='AI诊断')
        self.context['ai_advice'] = ''
        try:
            from analyzer import AIAdvisor
            import json as _json
            cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
            ai_cfg = {}
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    ai_cfg = _json.load(f)
            advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
            if advisor.enabled:
                label = self.context.get('co_name', [{}])[0].get('DB_NAME', 'DM8')
                print(f"\n🤖 正在调用 AI 诊断（{advisor.backend} / {advisor.model}）...")
                ai_advice = advisor.diagnose('dm8', label, self.context, self.context.get('auto_analyze', []))
                self.context['ai_advice'] = ai_advice
        except Exception:
            self.context['ai_advice'] = ''

        self.print_progress_bar(total_steps, total_steps, prefix='DM8巡检:', suffix='完成 [OK]')
        return self.context

    def _basic_risk_check(self):
        """基础风险检查"""
        ts_list = self.context.get('dm_tablespace', [])
        for ts in ts_list:
            if not isinstance(ts, dict): continue
            used_pct = self._safe_float(ts.get('USED_PCT', 0))
            name = ts.get('TABLESPACE_NAME', '?')
            if used_pct > 90:
                self.context['auto_analyze'].append({
                    'col1': f'表空间 {name}', 'col2': '高风险',
                    'col3': f'表空间使用率 {used_pct:.1f}%，超过 90% 告警线',
                    'col4': '高', 'col5': 'DBA',
                    'fix_sql': f"-- 查询表空间使用情况:\nSELECT * FROM DBA_TABLESPACES WHERE NAME='{name}';"
                })
            elif used_pct > 80:
                self.context['auto_analyze'].append({
                    'col1': f'表空间 {name}', 'col2': '中风险',
                    'col3': f'表空间使用率 {used_pct:.1f}%，建议关注',
                    'col4': '中', 'col5': 'DBA', 'fix_sql': ''
                })

        sess = self.context.get('dm_sessions', [])
        limit = self.context.get('dm_session_limit', [])
        if sess and limit:
            total = self._safe_int(sess[0], 'TOTAL_SESSIONS')
            max_sess = 0
            for l in limit:
                if isinstance(l, dict) and l.get('NAME') == 'processes':
                    max_sess = self._safe_int(l, 'VALUE')
                    break
            if max_sess > 0 and (total / max_sess) * 100 > 85:
                self.context['auto_analyze'].append({
                    'col1': '会话数接近上限', 'col2': '高风险',
                    'col3': f'当前会话 {total} / 上限 {max_sess}',
                    'col4': '高', 'col5': 'DBA',
                    'fix_sql': '-- 查看会话:\nSELECT * FROM V$SESSION WHERE TYPE=\'USER\';'
                })

        mem = self.context.get('system_info', {}).get('memory', {})
        if isinstance(mem, dict) and mem.get('usage_percent', 0) > 90:
            self.context['auto_analyze'].append({
                'col1': '系统内存紧张', 'col2': '高风险',
                'col3': f'内存使用率 {mem["usage_percent"]:.1f}%',
                'col4': '高', 'col5': '系统管理员', 'fix_sql': ''
            })

        for disk in self.context.get('system_info', {}).get('disk_list', []):
            if not isinstance(disk, dict): continue
            mp = disk.get('mountpoint', '/')
            if mp in IGNORE_MOUNTS: continue
            usage = self._safe_float(disk, 'usage_percent')
            if usage > 90:
                self.context['auto_analyze'].append({
                    'col1': f'磁盘空间不足 ({mp})', 'col2': '高风险',
                    'col3': f'磁盘 {mp} 使用率 {usage:.1f}%',
                    'col4': '高', 'col5': '系统管理员', 'fix_sql': ''
                })

    @staticmethod
    def _safe_float(obj, field='value', default=0.0):
        try:
            val = obj[field] if isinstance(obj, dict) else getattr(obj, field, None)
            return float(str(val).replace(',', '').replace('%', '')) if val is not None else default
        except Exception:
            return default

    @staticmethod
    def _safe_int(obj, field='value', default=0):
        try:
            val = obj[field] if isinstance(obj, dict) else getattr(obj, field, None)
            return int(str(val).replace(',', '')) if val is not None else default
        except Exception:
            return default


# ============================================================
# 报告保存类
# ============================================================
class saveDoc(object):
    """报告保存类 - 将 DM8 巡检数据渲染到 Word 模板"""

    def __init__(self, context, ofile, ifile, inspector_name="Jack", H=None, P=None, _dt=None):
        self.context = context
        self.ofile = ofile
        self.ifile = ifile
        self.inspector_name = inspector_name
        self.H = H
        self.P = P
        self._dt = _dt

    def contextsave(self):
        try:
            required_keys = ['health_summary', 'auto_analyze', 'dm_version', 'co_name', 'system_info']
            for key in required_keys:
                if key not in self.context:
                    if key == 'health_summary':   self.context[key] = [{'health_summary': '运行良好'}]
                    elif key == 'auto_analyze':    self.context[key] = []
                    elif key == 'dm_version':       self.context[key] = [{'BANNER': 'Unknown'}]
                    elif key == 'system_info':     self.context[key] = {}
                    else:                          self.context[key] = [{'placeholder': '数据缺失'}]

            if 'disk_list' not in self.context['system_info'] or not self.context['system_info']['disk_list']:
                self.context['system_info']['disk_list'] = [{
                    'device': 'C:', 'mountpoint': 'C:\\', 'fstype': 'NTFS',
                    'total_gb': 50.0, 'used_gb': 25.0, 'free_gb': 25.0, 'usage_percent': 50.0
                }]

            list_keys = [
                'dm_tablespace', 'dm_temp_ts', 'dm_sessions', 'dm_blocked', 'dm_trx',
                'dm_sga', 'dm_memory', 'dm_pga',
                'dm_redo_logs', 'dm_redo_curr',
                'dm_archive_config', 'dm_archive_lag',
                'dm_backup',
                'dm_params',
                'dm_invalid_objs', 'dm_invalid_cnt',
                'dm_users', 'dm_sys_privs', 'dm_default_pws',
                'dm_long_sql', 'dm_top_sql_cpu',
                'dm_rss_status', 'dm_rss_apply',
                'dm_dcinfo', 'dm_inst_info',
                'dm_undo_info', 'dm_transactions',
                'dm_recyclebin',
                'dm_datafiles',
                'dm_profile_pwd',
                'dm_top_waits', 'dm_wait_class',
                'dm_stale_stats',
                'dm_partition_info'
            ]
            for key in list_keys:
                if key not in self.context:
                    self.context[key] = []

            self.context.update({"report_time": __import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')})
            self.context.update({"inspector_name": self.inspector_name})

            problem_count = len(self.context.get("auto_analyze", []))
            self.context.update({"problem_count": problem_count})

            if problem_count == 0: health_status = "优秀"
            elif problem_count <= 3: health_status = "良好"
            elif problem_count <= 6: health_status = "一般"
            else: health_status = "需关注"
            self.context.update({"health_status": health_status})

            # ── 预格式化基本变量 ──
            if isinstance(self.context.get('co_name'), list) and len(self.context['co_name']) > 0:
                self.context['co_name'] = str(self.context['co_name'][0].get('DB_NAME', ''))
            else:
                self.context['co_name'] = str(self.context.get('co_name', '') or '')
            if isinstance(self.context.get('dm_version'), list) and len(self.context['dm_version']) > 0:
                self.context['dm_version'] = str(self.context['dm_version'][0].get('BANNER', ''))
            else:
                self.context['dm_version'] = str(self.context.get('dm_version', '') or '')

            # 系统信息
            sys_info = self.context.get('system_info') or {}
            self.context['hostname'] = str(sys_info.get('hostname', 'N/A') or 'N/A')
            platform_info = sys_info.get('platform', {})
            if isinstance(platform_info, dict):
                plat_str = platform_info.get('platform', '') or ''
                arch = platform_info.get('machine', '') or ''
                self.context['platform_text'] = f"{plat_str} ({arch})" if plat_str else (arch or 'N/A')
            else:
                self.context['platform_text'] = str(platform_info) if platform_info else 'N/A'

            uptime_data = self.context.get('dm_uptime')
            if isinstance(uptime_data, list) and len(uptime_data) > 0:
                u = uptime_data[0]
                start_time = u.get('STARTUP_TIME', '') or u.get('STARTUP', '') or ''
                self.context['uptime_text'] = start_time if start_time else 'N/A'
            else:
                self.context['uptime_text'] = 'N/A'

            self.context['server_addr'] = f"{self.H or ''}:{self.P or ''}"

            # 系统信息文本
            cpu = sys_info.get('cpu', {}) or {}
            mem = sys_info.get('memory', {}) or {}
            disks = sys_info.get('disk_list', []) or []
            lines = []
            lines.append(f"CPU 使用率\t{cpu.get('usage_percent', 'N/A')}%")
            used_mb = mem.get('used_mb', 0)
            total_mb = mem.get('total_mb', 0)
            mem_pct = mem.get('usage_percent', 0)
            lines.append(f"内存使用\t{used_mb} MB / {total_mb} MB ({mem_pct}%)")
            for d in disks[:10] if disks else []:
                mp = d.get('mountpoint', '/')
                tg = d.get('total_gb', 0)
                ug = d.get('used_gb', 0)
                up = d.get('usage_percent', 0)
                lines.append(f"磁盘 {mp}\t{up}% (已用 {ug}GB / 共 {tg}GB)")
            self.context['system_info_text'] = '\n'.join(lines)

            # 报告说明
            notes = [
                "1. 本报告基于达梦 DM8 数据库实时状态生成，反映了生成时刻的数据库健康状况",
                "2. 报告中空白的项表示未能获取到相关数据，可能是由于权限限制或该功能未启用",
                "3. 磁盘信息仅显示主要分区的使用率，如需查看完整磁盘信息请使用系统命令 'df -h'",
                "4. 巡检结果仅供参考，实际运维中请结合具体业务场景进行分析",
                "5. 建议定期进行数据库巡检，及时发现并解决潜在问题",
                "6. AI 诊断功能（若启用）生成的建议仅供参考，不构成专业 DBA 意见"
            ]
            self.context['notes_text'] = '\n'.join(notes)

            # 巡检执行摘要预格式化
            health = self.context.get('health_analysis')
            if isinstance(health, dict):
                score = health.get('score', 100)
                status = health.get('status', '正常')
                crit_n = health.get('critical_count', 0)
                warn_n = health.get('warning_count', 0)
                hs_lines = [
                    f"综合评分\t{score}\t状态\t{status}",
                    f"紧急告警数\t{crit_n}\t警告告警数\t{warn_n}"
                ]
                self.context['health_summary_text'] = '\n'.join(hs_lines)
            else:
                self.context['health_summary_text'] = ''

            # SQL 执行情况汇总
            errors = self.context.get('_safe_errors', [])
            fail_count = len(errors) if errors else 0
            ok_count = len(list_keys) - fail_count
            es_stat_lines = [
                f"总SQL数\t{ok_count + fail_count}\t-",
                f"成功数\t{ok_count}\t-",
                f"失败数\t{fail_count}\t-"
            ]
            self.context['error_stats_text'] = '\n'.join(es_stat_lines)

            if errors:
                ed_rows = []
                for err in errors[:30]:
                    name = err.get('item_name', '')
                    code = err.get('code', '')
                    hint = err.get('hint', '')
                    if len(hint) > 80: hint = hint[:77] + '...'
                    ed_rows.append(f"{name}\tDM{code}\t{hint}")
                self.context['error_detail_text'] = "SQL名称\t错误码\t修复建议\n" + '\n'.join(ed_rows)
            else:
                self.context['error_detail_text'] = '-'

            # 列表数据预格式化为 TSV 文本
            for key in list_keys:
                val = self.context.get(key)
                if val is None: continue
                if isinstance(val, (list, tuple)):
                    if not val:
                        self.context[key] = "无数据"
                        continue
                    rows = []
                    headers = []
                    for item in val:
                        if isinstance(item, dict):
                            if not headers: headers = list(item.keys())
                            rows.append('\t'.join(str(item.get(h,'')) for h in headers))
                    if headers:
                        self.context[key] = '\t'.join(headers) + '\n' + '\n'.join(rows)
                    else:
                        self.context[key] = str(val)[:2000]
                elif not isinstance(val, str):
                    self.context[key] = str(val)

            # 尝试 docxtpl 渲染
            try:
                with open(self.ifile, 'rb') as f:
                    template_bytes = f.read()
                doc_stream = io.BytesIO(template_bytes)
                tpl = DocxTemplate(doc_stream)
                tpl.render(self.context)
                tpl.save(self.ofile)
                doc2 = Document(self.ofile)
                self._append_chapters(doc2)
                doc2.save(self.ofile)
                return True
            except Exception as e:
                return self._fallback_render()
        except Exception as e:
            import traceback; traceback.print_exc(file=sys.stdout)
            raise RuntimeError(f"报告生成异常: {e}")

    def _append_chapters(self, doc):
        """追加第11-13章"""
        from docx.oxml.ns import qn as _qn

        def _set_cell_bg(cell, hex_color):
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        def _add_heading(text, level=1):
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(_qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.size = Pt(14) if level == 1 else Pt(12)
            return h

        def _add_table(headers, rows):
            t = doc.add_table(rows=max(1,len(rows))+1, cols=len(headers), style='Table Grid')
            for j, h in enumerate(headers):
                cell = t.cell(0, j); cell.text = h
                _set_cell_bg(cell, 'DCE6F1')
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9); run.font.name = '微软雅黑'; run.bold = True
            for i, row_data in enumerate(rows):
                for j, val in enumerate(row_data):
                    c = t.cell(i+1, j)
                    c.text = str(val)[:200] if val else ''
                    for p in c.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9); run.font.name = '微软雅黑'
                    c.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            return t

        # 第10章 风险与建议
        _add_heading("第10章 风险与建议")
        issues = self.context.get("auto_analyze", [])
        if issues:
            _add_heading("10.1 问题明细", 2)
            _add_table(["序号", "项目", "风险等级", "问题描述", "严重程度", "责任人", "修复建议"],
                       [(str(i+1), x.get('col1',''), x.get('col2',''), x.get('col3',''),
                         x.get('col4',''), x.get('col5',''), x.get('fix_sql','')[:200]) for i,x in enumerate(issues)])
            fix_sqls = [(x.get('col1',''), x.get('fix_sql','')) for x in issues if x.get('fix_sql')]
            if fix_sqls:
                _add_heading("10.2 修复SQL速查", 2)
                for fname, sql in fix_sqls:
                    p = doc.add_paragraph(); p.add_run(f"【{fname}】").bold = True
                    doc.add_paragraph(sql, style='List Bullet')
        else:
            p = doc.add_paragraph("未发现明显风险项，数据库整体运行状况良好")
            for r in p.runs: r.font.size = Pt(10.5); r.font.name = '微软雅黑'

        # 第11章 AI 诊断
        ai_text = self.context.get('ai_advice', '')
        if ai_text:
            _add_heading("第11章 AI 诊断建议")
            for line in ai_text.split('\n'):
                if line.startswith('# '): _add_heading(line[2:], level=2)
                elif line.startswith('## '): _add_heading(line[3:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    doc.add_paragraph(line, style='List Number')
                elif line.strip():
                    doc.add_paragraph(line)
            next_chap = 12
        else:
            next_chap = 11

        # 第{11|12}章 报告说明（段落形式）
        _add_heading(f"第{next_chap}章 报告说明")
        for note in self.context.get('notes_text', '').split('\n'):
            if note.strip():
                p = doc.add_paragraph(note.strip())
                for r in p.runs:
                    r.font.size = Pt(10.5); r.font.name = '微软雅黑'
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5

    def _fallback_render(self):
        """备用渲染：纯 python-docx 构建报告（无模板依赖）"""
        from docx.oxml.ns import qn as _qn

        def _set_cell_bg(cell, hex_color):
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        def _add_heading(text, level=1):
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(_qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.size = Pt(14) if level == 1 else Pt(12)
            return h

        def _t(hdr, rows):
            if not rows: return
            max_cols = max(len(hdr), *(len(r) for r in rows))
            tt = doc.add_table(rows=max(1,len(rows))+1, cols=max_cols, style='Table Grid')
            for j, h in enumerate(hdr):
                c = tt.cell(0, j); c.text = h; _set_cell_bg(c, 'DCE6F1')
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9); run.font.name = '微软雅黑'; run.bold = True
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    if j < max_cols:
                        c = tt.cell(i+1, j); c.text = str(val)[:200] if val else ''
                        for p in c.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9); run.font.name = '微软雅黑'
                        c.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _render(key, max_r=50, label=None):
            val = ctx.get(key, '')
            if not val or (isinstance(val, str) and val == "无数据"):
                desc = label if label else key
                p = doc.add_paragraph(f"{desc}：暂无数据")
                for r in p.runs: r.font.size = Pt(10.5); r.font.name = '微软雅黑'
                return
            if isinstance(val, str) and '\t' in val:
                lines = val.strip().split('\n')
                if not lines: return
                hdr = lines[0].split('\t')
                rows = [ln.split('\t') for ln in lines[1:] if ln.strip()]
                _t(hdr, rows[:max_r])
            else:
                p = doc.add_paragraph(str(val)[:500])
                for r in p.runs: r.font.size = Pt(10.5); r.font.name = '微软雅黑'

        try:
            doc = Document()
            doc.add_paragraph("达梦数据库健康巡检报告",
                              style=doc.styles.add_style('ReportTitle', 1))
            doc.styles['ReportTitle'].font.size = Pt(22)
            doc.styles['ReportTitle'].font.bold = True
            doc.styles['ReportTitle'].font.name = '微软雅黑'
            doc.styles['ReportTitle'].font.color.rgb = RGBColor(0, 51, 102)
            doc.add_paragraph("")

            # 封面信息表
            ctx = self.context

            def _v(key, sub_key=None):
                val = ctx.get(key, '')
                if isinstance(val, list) and val:
                    val = val[0].get(sub_key, '') if sub_key else str(val[0])
                return str(val) if val else 'N/A'

            info_items = [
                ("数据库名称",   _v('co_name', 'DB_NAME')),
                ("服务器地址",  _v('server_addr')),
                ("达梦版本",    _v('dm_version', 'BANNER')),
                ("服务器主机名", _v('hostname')),
                ("实例启动时间", _v('uptime_text')),
                ("巡检人员",    self.inspector_name),
                ("服务器平台",  _v('platform_text')),
                ("报告生成时间", _v('report_time')),
            ]
            t_info = doc.add_table(rows=len(info_items), cols=2, style='Table Grid')
            for i, (k, v) in enumerate(info_items):
                t_info.cell(i, 0).text = k; _set_cell_bg(t_info.cell(i, 0), 'F2F2F2')
                t_info.cell(i, 1).text = str(v)
                for cell in [t_info.cell(i, 0), t_info.cell(i, 1)]:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10.5); run.font.name = '微软雅黑'

# Ch1 基本信息
            _add_heading("第1章 数据库基本信息")
            _co = ctx.get('co_name', '')
            _dm = ctx.get('dm_version', '')
            if isinstance(_co, list) and _co: _co = _co[0].get('DB_NAME', '')
            if isinstance(_dm, list) and _dm: _dm = _dm[0].get('BANNER', '')
            _t(["项目", "值"], [
                ["实例名称", str(_co)],
                ["达梦版本", str(_dm)],
            ])

            # Ch2 巡检执行摘要
            _add_heading("第2章 巡检执行摘要")
            hst = ctx.get('health_summary_text', '')
            if hst:
                for line in hst.split('\n'):
                    if line.strip():
                        parts = line.split('\t')
                        _t(["项目", "数值"], [parts])

            est = ctx.get('error_stats_text', '')
            if est:
                doc.add_paragraph("")
                p3 = doc.add_paragraph(); r3 = p3.add_run("SQL 执行情况")
                r3.bold = True; r3.font.size = Pt(12); r3.font.name = '微软雅黑'
                el = [l.strip() for l in est.strip().split('\n') if l.strip()]
                if el:
                    _t(["项目", "数值"], [c.split('\t') for c in el])
                edt = ctx.get('error_detail_text', '')
                if edt and edt != '-':
                    doc.add_paragraph("")
                    dlines = [l.strip() for l in edt.strip().split('\n') if l.strip()]
                    if len(dlines) > 1:
                        _t(["SQL名称", "错误码", "修复建议"], [l.split('\t') for l in dlines[1:] if l.count('\t') >= 2])

            # Ch3-10
            _add_heading("第3章 表空间使用情况")
            _render('dm_tablespace', 30)

            _add_heading("第4章 会话与事务")
            _render('dm_sessions'); _render('dm_transactions')

            _add_heading("第5章 内存分析")
            _render('dm_sga', label='SGA 缓冲池汇总')
            _render('dm_memory', label='各缓冲池详情（SGA）')

            _add_heading("第6章 重做日志与归档")
            _render('dm_redo_logs')

            _add_heading("第7章 系统资源监控")
            sit = ctx.get('system_info_text', '')
            if sit:
                for line in sit.split('\n'):
                    if line.strip():
                        parts = line.split('\t')
                        _t(["项目", "详情"], [parts])

            _add_heading("第8章 对象与用户安全")
            _add_heading("8.1 无效对象", 2); _render('dm_invalid_cnt')
            _add_heading("8.2 用户列表", 2); _render('dm_users', 20)

            _add_heading("第9章 备份与归档"); _render('dm_backup')

            # 第10章由 _append_chapters 统一生成（作为第13章），此处不再重复

            self._append_chapters(doc)
            doc.save(self.ofile)
            return True
        except Exception as e:
            print(f"备用渲染也失败了: {e}")
            import traceback; traceback.print_exc()
            return False


# ============================================================
# 交互式入口（无参数时调用）
# ============================================================
def main():
    print("=" * 60)
    print("  DBCheck - DM8 巡检工具 " + VER)
    print("=" * 60)

    # 交互式收集参数
    db_info = {}
    db_info['host'] = input("主机地址: ").strip()
    db_info['port'] = input("端口 [5236]: ").strip() or "5236"
    db_info['user'] = input("用户名 [SYSDBA]: ").strip() or "SYSDBA"
    db_info['password'] = input("密码: ").strip()
    db_info['db_name'] = input("数据库名 [DAMENG]: ").strip() or "DAMENG"
    inspector_name = input("巡检人员 [Jack]: ").strip() or "Jack"

    use_ssh = input("是否配置 SSH 收集系统信息? [y/N]: ").strip().lower()
    if use_ssh == 'y':
        db_info['ssh_host'] = input("  SSH 主机 [同数据库主机]: ").strip() or db_info['host']
        db_info['ssh_port'] = int(input("  SSH 端口 [22]: ").strip() or "22")
        db_info['ssh_user'] = input("  SSH 用户 [root]: ").strip() or "root"
        db_info['ssh_password'] = input("  SSH 密码 (留空则跳过): ").strip()
    else:
        db_info['ssh_host'] = ''
        db_info['ssh_port'] = 22
        db_info['ssh_user'] = 'root'
        db_info['ssh_password'] = ''

    dir_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
    os.makedirs(dir_path, exist_ok=True)
    ofile = os.path.join(dir_path, f"DM8_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    print(f"报告输出目录: {dir_path}/")
    ifile = create_word_template(inspector_name)

    ssh_info = {}
    if db_info.get('ssh_host'):
        ssh_info = {
            'ssh_host': db_info['ssh_host'],
            'ssh_port': db_info['ssh_port'],
            'ssh_user': db_info['ssh_user'],
            'ssh_password': db_info['ssh_password'],
        }

    inspector = getData(db_info['host'], int(db_info['port']),
                        db_info['user'], db_info['password'],
                        db_info.get('db_name'), ssh_info)

    if inspector.conn_db is None:
        print("\n[FAIL] 数据库连接失败，请检查参数后重试")
        input("\n按回车键返回...")
        return

    context = inspector.checkdb()
    if context:
        saver = saveDoc(context, ofile, ifile, inspector_name,
                        H=inspector.H, P=inspector.P)
        if saver.contextsave():
            print(f"\n[OK] 报告已生成: {ofile}")
        else:
            print(f"\n[FAIL] 报告生成失败")
    else:
        print("\n[FAIL] 巡检执行失败")

    cont = input("\n是否返回主菜单? [y/N]: ").strip().lower()
    if cont == 'y':
        return


# ============================================================
# 命令行入口
# ============================================================
if __name__ == '__main__':
    if len(sys.argv) == 1:
        # 无参数 → 交互式模式
        main()
    else:
        # 有参数 → argparse 批处理模式
        parser = argparse.ArgumentParser(description='DBCheck - DM8 巡检工具 ' + VER)
        parser.add_argument('--host', default='127.0.0.1', help='数据库主机IP')
        parser.add_argument('--port', type=int, default=5236, help='数据库端口（默认5236）')
        parser.add_argument('--user', default='SYSDBA', help='用户名（默认SYSDBA）')
        parser.add_argument('--password', default='', help='密码')
        parser.add_argument('--db', dest='db_name', default='DAMENG', help='数据库名（默认DAMENG）')
        parser.add_argument('--template', action='store_true', help='仅生成报告模板')
        parser.add_argument('--inspector', default='Jack', help='巡检人员姓名')
        parser.add_argument('--output', default='', help='输出文件路径')
        parser.add_argument('--ssh-host', dest='ssh_host', default='', help='SSH远程主机')
        parser.add_argument('--ssh-port', dest='ssh_port', type=int, default=22, help='SSH端口')
        parser.add_argument('--ssh-user', dest='ssh_user', default='root', help='SSH用户名')
        parser.add_argument('--ssh-password', dest='ssh_password', default='', help='SSH密码')
        args = parser.parse_args()

        print("=" * 60)
        print("  DBCheck - DM8 巡检工具 " + VER)
        print("=" * 60)

        if args.template:
            tpl = create_word_template(args.inspector)
            print(f"[OK] DM8 报告模板已生成: {tpl}")
            sys.exit(0)

        ofile = args.output or os.path.join(
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports"),
            f"DM8_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        ifile = create_word_template(args.inspector)

        ssh_info = {}
        if args.ssh_host:
            ssh_info = {
                'ssh_host': args.ssh_host,
                'ssh_port': args.ssh_port,
                'ssh_user': args.ssh_user,
                'ssh_password': args.ssh_password,
            }

        inspector = getData(args.host, args.port, args.user, args.password, args.db_name, ssh_info)

        if inspector.conn_db is None:
            print("[FAIL] 数据库连接失败，请检查参数后重试")
            sys.exit(1)

        context = inspector.checkdb()
        if context:
            saver = saveDoc(context, ofile, ifile, args.inspector,
                            H=inspector.H, P=inspector.P)
            if saver.contextsave():
                print(f"\n[OK] 报告已生成: {ofile}")
            else:
                print(f"\n[FAIL] 报告生成失败")
        else:
            print("\n[FAIL] 巡检执行失败")

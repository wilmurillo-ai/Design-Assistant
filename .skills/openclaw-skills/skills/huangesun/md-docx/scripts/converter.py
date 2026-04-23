#!/usr/bin/env python3
"""
Markdown 与 Word 文档相互转换工具
使用 python-docx 和 markdown 库
依赖: pip install python-docx markdown
示例：python converter.py test.md --to docx --output-dir ./converted
"""

import os
import sys
import argparse
from pathlib import Path
import datetime
import re
from markdown import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import html
import tempfile
import shutil
from urllib.parse import urlparse
#import requests
from io import BytesIO

class MDDocxConverter:
    def __init__(self, output_dir="converted", keep_images=True):
        """初始化转换器"""
        self.output_dir = Path(output_dir)
        self.keep_images = keep_images
        # 确保输出目录存在
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def md_to_docx(self, input_path, output_path=None):
        """Markdown转Word文档"""
        input_path = Path(input_path)
        
        # 检查输入文件
        if not input_path.exists():
            return {"success": False, "error": f"输入文件不存在: {input_path}"}
        
        # 生成输出路径
        if output_path is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"{input_path.stem}_{timestamp}.docx"
        else:
            output_path = Path(output_path)
        
        try:
            # 读取Markdown文件
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 创建Word文档
            doc = Document()
            
            # 设置默认样式
            self._setup_document_styles(doc)
            
            # 解析并转换Markdown内容
            self._parse_markdown_to_docx(doc, md_content)
            
            # 保存文档
            doc.save(str(output_path))
            
            return {
                "success": True,
                "input": str(input_path),
                "output": str(output_path),
                "message": f"成功将 {input_path.name} 转换为 Word 文档"
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def docx_to_md(self, input_path, output_path=None):
        """Word文档转Markdown"""
        input_path = Path(input_path)
        
        if not input_path.exists():
            return {"success": False, "error": f"输入文件不存在: {input_path}"}
        
        if output_path is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"{input_path.stem}_{timestamp}.md"
        else:
            output_path = Path(output_path)
        
        try:
            # 加载Word文档
            doc = Document(str(input_path))
            
            # 提取内容并转换为Markdown
            md_content = self._parse_docx_to_markdown(doc)
            
            # 写入Markdown文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            # 提取图片（如果需要）
            if self.keep_images:
                self._extract_images_from_docx(input_path, output_path)
            
            return {
                "success": True,
                "input": str(input_path),
                "output": str(output_path),
                "message": f"成功将 {input_path.name} 转换为 Markdown 文件"
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 标题样式
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = '宋体'
                style.font.size = Pt(28 - i*2)  # H1: 26pt, H2: 24pt, ...
                style.font.bold = True
        
        # 正文样式
        if 'Normal' in doc.styles:
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
        
        # 代码样式
        if 'Code' not in doc.styles:
            style = doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            style.font.name = 'Consolas'
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(200, 0, 0)
    
    def _parse_markdown_to_docx(self, doc, md_content):
        """将Markdown内容解析并添加到Word文档"""
        lines = md_content.split('\n')
        i = 0
        in_code_block = False
        code_block_content = []
        
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_content = []
                else:
                    # 代码块结束
                    self._add_code_block(doc, '\n'.join(code_block_content))
                    in_code_block = False
                i += 1
                continue
            
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            
            # 处理标题
            if line.startswith('#'):
                level = len(re.match(r'^#+', line).group())
                text = line[level:].strip()
                doc.add_heading(text, level=level)
                i += 1
                continue
            
            # 处理列表
            if re.match(r'^[\*\-\+]\s', line) or re.match(r'^\d+\.\s', line):
                self._add_list_item(doc, lines, i)
                # 跳过所有列表项
                while i < len(lines) and (re.match(r'^[\*\-\+]\s', lines[i]) or re.match(r'^\d+\.\s', lines[i])):
                    i += 1
                continue
            
            # 处理空行
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue
            
            # 处理普通段落
            if line.strip():
                self._add_paragraph_with_formatting(doc, line)
            i += 1
    
    def _add_list_item(self, doc, lines, start_idx):
        """添加列表项（支持嵌套）"""
        i = start_idx
        list_level = 0
        list_items = []
        
        while i < len(lines):
            line = lines[i]
            
            # 检查是否是无序列表项
            unordered_match = re.match(r'^([\*\-\+]+)\s(.*)', line)
            if unordered_match:
                markers = unordered_match.group(1)
                text = unordered_match.group(2)
                level = len(markers)
                list_items.append(('unordered', level, text))
                i += 1
                continue
            
            # 检查是否是有序列表项
            ordered_match = re.match(r'^(\d+\.+)\s(.*)', line)
            if ordered_match:
                markers = ordered_match.group(1)
                text = ordered_match.group(2)
                level = len(markers.split('.')[0])  # 简化处理
                list_items.append(('ordered', level, text))
                i += 1
                continue
            
            break
        
        # 添加到文档
        for item_type, level, text in list_items:
            paragraph = doc.add_paragraph(style='List Paragraph')
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            
            if item_type == 'unordered':
                paragraph.style = 'List Bullet'
            else:
                paragraph.style = 'List Number'
            
            self._add_text_with_formatting(paragraph, text)
    
    def _add_code_block(self, doc, code):
        """添加代码块"""
        paragraph = doc.add_paragraph()
        paragraph.style = 'Code'
        run = paragraph.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    def _add_paragraph_with_formatting(self, doc, text):
        """添加带格式的段落"""
        paragraph = doc.add_paragraph()
        self._add_text_with_formatting(paragraph, text)
    
    def _add_text_with_formatting(self, paragraph, text):
        """添加带Markdown格式的文本"""
        # 处理粗体 **text** 或 __text__
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|`.*?`|\[.*?\]\(.*?\))', text)
        
        for part in parts:
            if not part:
                continue
            
            # 粗体
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            
            # 粗体（下划线格式）
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            
            # 斜体
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            
            # 斜体（下划线格式）
            elif part.startswith('_') and part.endswith('_'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            
            # 行内代码
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(200, 0, 0)
            
            # 链接
            elif part.startswith('[') and '](' in part:
                link_text = part[1:part.index(']')]
                link_url = part[part.index('(')+1:part.index(')')]
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                # 添加超链接需要额外处理，这里简化
            
            # 普通文本
            else:
                paragraph.add_run(part)
    
    def _parse_docx_to_markdown(self, doc):
        """将Word文档解析为Markdown"""
        md_lines = []
        
        for paragraph in doc.paragraphs:
            # 判断是否是标题
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                md_lines.append('#' * level + ' ' + paragraph.text)
                md_lines.append('')
                continue
            
            # 判断是否是列表
            if 'list' in style_name:
                indent = paragraph.paragraph_format.left_indent
                indent_level = int(indent.pt / 20) if indent else 0  # 粗略估计
                
                if 'bullet' in style_name:
                    md_lines.append('  ' * indent_level + '- ' + paragraph.text)
                elif 'number' in style_name:
                    # 有序列表编号需要更复杂的处理
                    md_lines.append('  ' * indent_level + '1. ' + paragraph.text)
                continue
            
            # 判断是否是代码块（通过样式判断）
            if 'code' in style_name:
                md_lines.append('```')
                md_lines.append(paragraph.text)
                md_lines.append('```')
                md_lines.append('')
                continue
            
            # 普通段落
            if paragraph.text.strip():
                md_lines.append(paragraph.text)
                md_lines.append('')
        
        return '\n'.join(md_lines)
    
    def _extract_images_from_docx(self, docx_path, md_path):
        """从DOCX文件中提取图片（简化版本）"""
        try:
            # 创建图片目录
            images_dir = self.output_dir / f"{Path(md_path).stem}_images"
            images_dir.mkdir(exist_ok=True)
            
            # 使用zipfile提取图片（docx本质上是zip文件）
            import zipfile
            
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # 查找media文件夹中的图片
                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('word/media/'):
                        # 提取图片
                        image_data = zip_ref.read(file_info.filename)
                        image_name = Path(file_info.filename).name
                        image_path = images_dir / image_name
                        
                        with open(image_path, 'wb') as f:
                            f.write(image_data)
                        
                        # 更新Markdown文件中的图片引用（这里简化处理）
                        self._update_md_image_reference(md_path, image_name, str(image_path))
                        
        except Exception as e:
            print(f"提取图片时出错: {e}", file=sys.stderr)
    
    def _update_md_image_reference(self, md_path, image_name, image_path):
        """更新Markdown中的图片引用"""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 简单的图片引用替换（实际可能需要更复杂的逻辑）
            # 这里假设图片引用格式为 ![](image_name)
            updated_content = content.replace(f'![]({image_name})', f'![]({image_path})')
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(updated_content)
                
        except Exception as e:
            print(f"更新图片引用时出错: {e}", file=sys.stderr)
    
    def batch_convert(self, input_files, conversion_type):
        """批量转换多个文件"""
        results = []
        
        for file_path in input_files:
            if conversion_type == "md_to_docx":
                result = self.md_to_docx(file_path)
            else:
                result = self.docx_to_md(file_path)
            results.append(result)
        
        return results

def main():
    parser = argparse.ArgumentParser(description='Markdown与Word文档相互转换工具')
    parser.add_argument('input', help='输入文件路径')
    parser.add_argument('--to', choices=['docx', 'md'], required=True, 
                       help='转换目标格式: docx 或 md')
    parser.add_argument('--output', '-o', help='输出文件路径（可选）')
    parser.add_argument('--output-dir', default='converted', 
                       help='输出目录（默认为converted）')
    parser.add_argument('--no-images', action='store_false', 
                       dest='keep_images', help='不保留图片（仅对docx转md有效）')
    
    args = parser.parse_args()
    
    converter = MDDocxConverter(
        output_dir=args.output_dir,
        keep_images=args.keep_images
    )
    
    # 执行转换
    if args.to == 'docx':
        result = converter.md_to_docx(args.input, args.output)
    else:
        result = converter.docx_to_md(args.input, args.output)
    
    # 输出结果（JSON格式便于OpenClaw解析）
    import json
    print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()

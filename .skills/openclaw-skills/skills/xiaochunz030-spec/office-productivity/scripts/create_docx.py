#!/usr/bin/env python3
"""Word 文档创建工具 - python-docx"""
import sys
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_run_font(run, font_name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 - level * 2)
    return heading


def add_paragraph(doc, text, font_size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = str(cell_text)
            for p in row_cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    return table


def create_document(title, paragraphs, output_path):
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    # 标题
    add_heading(doc, title, 0)
    doc.add_paragraph(f"创建时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()
    # 内容
    for item in paragraphs:
        if isinstance(item, dict):
            if item.get('type') == 'heading':
                add_heading(doc, item['text'], item.get('level', 1))
            elif item.get('type') == 'table':
                add_table(doc, item['headers'], item['rows'])
                doc.add_paragraph()
            elif item.get('type') == 'image':
                doc.add_picture(item['path'], width=Inches(item.get('width', 5)))
            else:
                add_paragraph(doc, item.get('text', ''), bold=item.get('bold', False))
        else:
            add_paragraph(doc, str(item))
    doc.save(output_path)
    print(f"[OK] 文档已保存: {output_path}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='创建 Word 文档')
    parser.add_argument('--title', '-t', default='未命名文档', help='文档标题')
    parser.add_argument('--output', '-o', default='output.docx', help='输出路径')
    parser.add_argument('--content', '-c', default='[]', help='内容 JSON')
    args = parser.parse_args()
    import json
    content = json.loads(args.content)
    create_document(args.title, content, args.output)

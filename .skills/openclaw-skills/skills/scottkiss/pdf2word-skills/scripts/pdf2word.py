#!/usr/bin/env python3
import sys
import os
import subprocess
try:
    from docx import Document
except ImportError:
    print("Error: python-docx library is not installed.")
    print("Please install requirements first: pip install -r requirements.txt")
    sys.exit(1)

def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/pdf2word.py <input_pdf> <output_docx> [docr_args...]")
        print("Example: python scripts/pdf2word.py sample.pdf output.docx")
        print("Example with Gemini API: python scripts/pdf2word.py sample.pdf output.docx -engine gemini")
        sys.exit(1)
        
    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    extra_args = sys.argv[3:]
    
    if not os.path.exists(input_pdf):
        print(f"Error: Input file '{input_pdf}' not found.")
        sys.exit(1)
        
    script_dir = os.path.dirname(os.path.abspath(__file__))
    docr_path = os.path.join(script_dir, "docr", "docr")
    
    if not os.path.exists(docr_path):
        print("Error: docr binary not found.")
        print("Please run scripts/install.sh first to download the OCR engine.")
        sys.exit(1)
        
    print(f"==> Extracting text from {input_pdf} using docr...")
    
    temp_text_file = output_docx + ".txt"
    
    try:
        # Pass extra custom arguments to docr
        cmd = [docr_path, input_pdf, "-o", temp_text_file] + extra_args
        subprocess.run(cmd, check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error: OCR extraction failed: {e}")
        sys.exit(1)
        
    if not os.path.exists(temp_text_file):
        print(f"Error: Expected text file {temp_text_file} was not created by docr.")
        sys.exit(1)
        
    print(f"==> Converting extracted text to Word Document ({output_docx})...")
    
    doc = Document()
    
    with open(temp_text_file, "r", encoding="utf-8") as f:
        is_first_page = True
        table_data = []
        
        def commit_table():
            if not table_data:
                return
            # Remove the markdown separator line (e.g., |---|---|)
            rows = [r for r in table_data if r.replace('|', '').replace('-', '').replace(':', '').replace(' ', '') != '']
            if rows:
                cols_count = max(len(r.split('|')) - 2 for r in rows)
                if cols_count > 0:
                    table = doc.add_table(rows=len(rows), cols=cols_count)
                    table.style = 'Table Grid'
                    for i, row in enumerate(rows):
                        cells = [c.strip() for c in row.split('|')[1:-1]]
                        for j, val in enumerate(cells):
                            if j < len(table.columns):
                                table.cell(i, j).text = val
            table_data.clear()

        for line in f:
            clean_line = line.strip()
            
            # Simple Markdown table detection
            if clean_line.startswith("|") and clean_line.endswith("|"):
                table_data.append(clean_line)
                continue
            else:
                commit_table()

            if clean_line:
                # Handle docr's page delimiter "<!-- Page X -->"
                if clean_line.startswith("<!-- Page ") and clean_line.endswith(" -->"):
                    if not is_first_page:
                        doc.add_page_break()
                    is_first_page = False
                    continue
                doc.add_paragraph(clean_line)
                
        commit_table() # In case the file ends with a table
                
    try:
        doc.save(output_docx)
        print(f"==> Successfully saved to {output_docx}")
    except Exception as e:
        print(f"Error saving Word doc: {e}")
        
    finally:
        # Cleanup temporary text file
        try:
            if os.path.exists(temp_text_file):
                os.remove(temp_text_file)
        except OSError as e:
            print(f"Warning: Could not remove temporary file {temp_text_file}: {e}")

if __name__ == "__main__":
    main()

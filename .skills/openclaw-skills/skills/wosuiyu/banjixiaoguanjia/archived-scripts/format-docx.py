# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

# 创建 Word 文档
doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.styles['Normal'].font.size = Pt(12)

# 添加标题
title = doc.add_heading('', level=0)
title_run = title.add_run('一下(第9节) 大瑀 课内作业分析报告')
title_run.font.name = '微软雅黑'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加基本信息框
doc.add_paragraph()  # 空行
info_table = doc.add_table(rows=3, cols=2)
info_table.style = 'Table Grid'
info_table.cell(0, 0).text = '学生姓名'
info_table.cell(0, 1).text = '大瑀'
info_table.cell(1, 0).text = '作业类型'
info_table.cell(1, 1).text = '课内'
info_table.cell(2, 0).text = '作业图片数量'
info_table.cell(2, 1).text = '6张'

# 设置表格字体
for row in info_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(11)

doc.add_paragraph()  # 空行

# 读取分析内容
with open(r'C:\Users\Administrator\Desktop\一下(第9节)\一下(第9节)大瑀课内分析.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# 解析内容并格式化
lines = content.split('\n')
current_section = None

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # 跳过已处理的基本信息
    if line.startswith('学生姓名：') or line.startswith('作业类型：') or line.startswith('作业图片数量：'):
        continue
    
    # 处理分隔线
    if line == '---':
        doc.add_paragraph()
        continue
    
    # 处理标题（**标题**）
    if line.startswith('**') and line.endswith('**'):
        heading_text = line.strip('*')
        heading = doc.add_heading('', level=1)
        run = heading.add_run(heading_text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        continue
    
    # 处理逐题批改中的题目
    if line.startswith('第') and '题：' in line:
        p = doc.add_paragraph()
        # 分割题号和内容
        parts = line.split('：', 1)
        if len(parts) == 2:
            # 题号加粗
            run1 = p.add_run(parts[0] + '：')
            run1.font.name = '微软雅黑'
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run1.font.size = Pt(12)
            run1.font.bold = True
            
            # 内容
            run2 = p.add_run(parts[1])
            run2.font.name = '微软雅黑'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run2.font.size = Pt(12)
            
            # 根据正确/错误设置颜色
            if '正确' in parts[1]:
                run2.font.color.rgb = RGBColor(0, 153, 0)
            elif '错误' in parts[1]:
                run2.font.color.rgb = RGBColor(204, 0, 0)
        else:
            run = p.add_run(line)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(12)
        continue
    
    # 处理薄弱环节和总结中的条目（1. 2. 3.）
    if line[0].isdigit() and line[1] == '.':
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(line[3:].strip())
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(12)
        continue
    
    # 处理普通段落
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)

# 添加页脚
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph('— 分析报告结束 —')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

# 保存
doc.save(r'C:\Users\Administrator\Desktop\一下(第9节)\一下(第9节)大瑀课内分析.docx')
print('Word文档已保存')

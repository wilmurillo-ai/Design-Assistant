# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 读取 txt 文件
with open(r'C:\Users\Administrator\Desktop\一下(第9节)\一下(第9节)小米课内分析.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建 Word 文档
doc = Document()

# 添加标题
title = doc.add_heading('一下(第9节) 小米 课内作业分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加内容
for line in content.split('\n'):
    if line.strip():
        doc.add_paragraph(line)

# 保存
doc.save(r'C:\Users\Administrator\Desktop\一下(第9节)\一下(第9节)小米课内分析.docx')
print('Word文档已保存')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('一下(第9节) 小米 课内作业分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

with open(r'C:\Users\Administrator\Desktop\一下(第9节)\一下(第9节)小米课内分析.txt', 'r', encoding='utf-8') as f:
    content = f.read()

for line in content.split('\n'):
    if line.strip():
        doc.add_paragraph(line)

output_path = r'C:\Users\Administrator\Desktop\一下(第9节)\一下(第9节)小米课内分析.docx'
doc.save(output_path)
print('Word文档已保存')

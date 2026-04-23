#!/usr/bin/env python3
"""
Generate incorporation documents from a YAML config file.

Usage:
    python3 generate.py config.yaml [--output-dir ./output]

Reads the config, loads templates from assets/templates/, substitutes variables,
and outputs .docx files ready to file.
"""

import argparse
import os
import sys
import json
from pathlib import Path

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx required. Install with: pip3 install python-docx pyyaml")
    sys.exit(1)


BLACK = RGBColor(0, 0, 0)
SKILL_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = SKILL_DIR / "assets" / "templates"


def load_config(path):
    with open(path) as f:
        if path.endswith('.json'):
            return json.load(f)
        elif HAS_YAML:
            return yaml.safe_load(f)
        else:
            print("ERROR: PyYAML not installed. Use .json config or: pip3 install pyyaml")
            sys.exit(1)


def create_doc():
    """Create a new Document with default formatting."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    return doc


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLACK
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return h


def add_para(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p


def add_highlight(run):
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


def set_cell(cell, text, bold=False, highlight_yellow=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.bold = bold
    if highlight_yellow:
        add_highlight(run)


def shade_row(row, color="D9E2F3"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        tcPr.append(shading)


def set_table_borders(table):
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def add_table_header(table, headers):
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
    shade_row(table.rows[0], "2F5496")
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)


def is_tbd(val):
    return val is None or str(val).upper() in ("TBD", "[TBD]", "")


def val_or_highlight(val, default="[TBD]"):
    """Return (display_text, should_highlight)"""
    if is_tbd(val):
        return default, True
    return str(val), False


# ============================================================
# DOCUMENT GENERATORS
# ============================================================

def gen_articles(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    ra = cfg['registered_agent']
    incorporator = next(d for d in cfg['directors'] if d.get('is_incorporator'))
    filing_date = cfg.get('filing_date', 'TBD')

    add_heading(doc, f"ARTICLES OF INCORPORATION OF {co.upper()}")
    add_para(doc, f"(Pursuant to NRS Chapter 78)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_heading(doc, "ARTICLE I — NAME", level=2)
    add_para(doc, f'The name of the corporation is {co} (the "Corporation").')

    add_heading(doc, "ARTICLE II — REGISTERED AGENT", level=2)
    add_para(doc, f"The name and street address of the Corporation's registered agent in the State of {state} is:")
    add_para(doc, ra['name'], bold=True)
    add_para(doc, ra['street'])
    add_para(doc, f"{ra['city']}, {ra['state']} {ra['zip']}")

    add_heading(doc, "ARTICLE III — PURPOSE", level=2)
    add_para(doc, f"The purpose of the Corporation is to engage in any lawful activity for which corporations may be organized under NRS Chapter 78.")

    add_heading(doc, "ARTICLE IV — AUTHORIZED SHARES", level=2)
    add_para(doc, "The Corporation is authorized to issue the following classes of stock:")

    stock = cfg.get('stock', {})
    if 'common' in stock:
        c = stock['common']
        add_para(doc, f"Common Stock: {c['authorized']:,} shares, par value ${c['par_value']} per share, {c['voting_rights']} vote(s) per share.", bold=True)

    for pref in stock.get('preferred', []):
        add_para(doc, f"{pref['class_name']}: {pref['authorized']:,} shares, par value ${pref['par_value']} per share, {pref['voting_rights']} vote(s) per share. {pref.get('notes', '')}", bold=True)

    add_heading(doc, "ARTICLE V — DIRECTORS", level=2)
    dirs = cfg['directors']
    add_para(doc, f"The number of directors constituting the initial Board of Directors is {len(dirs)}.")
    for d in dirs:
        addr = d.get('address', '')
        text, hl = val_or_highlight(addr, "[ADDRESS]")
        p = doc.add_paragraph()
        run = p.add_run(f"• {d['name']} — {text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        if hl:
            add_highlight(run)

    add_heading(doc, "ARTICLE VI — INCORPORATOR", level=2)
    inc_addr = incorporator.get('address', '')
    addr_text, addr_hl = val_or_highlight(inc_addr, "[ADDRESS]")
    add_para(doc, incorporator['name'], bold=True)
    p = add_para(doc, addr_text)
    if addr_hl:
        for run in p.runs:
            add_highlight(run)

    add_heading(doc, "ARTICLE VII — INDEMNIFICATION", level=2)
    add_para(doc, "The Corporation shall indemnify, to the fullest extent permitted by applicable law as it exists on the date hereof or as it may hereafter be amended, any person who was or is a party or is threatened to be made a party to any threatened, pending, or completed action, suit, or proceeding (whether civil, criminal, administrative, or investigative) by reason of the fact that such person is or was a director or officer of the Corporation, or is or was serving at the request of the Corporation as a director, officer, employee, or agent of another corporation, partnership, joint venture, trust, or other enterprise.")

    add_heading(doc, "ARTICLE VIII — LIABILITY LIMITATION", level=2)
    add_para(doc, "To the fullest extent permitted by NRS 78.138, no director or officer of the Corporation shall be personally liable to the Corporation or its stockholders for damages for breach of fiduciary duty as a director or officer, except for liability arising from (i) acts or omissions involving intentional misconduct, fraud, or a knowing violation of law, or (ii) the payment of distributions in violation of NRS 78.300.")

    add_heading(doc, "ARTICLE IX — AMENDMENT", level=2)
    add_para(doc, "The Corporation reserves the right to amend, alter, change, or repeal any provision contained in these Articles of Incorporation in the manner now or hereafter prescribed by statute, and all rights conferred upon stockholders herein are granted subject to this reservation.")

    doc.add_paragraph()
    doc.add_paragraph()

    fd_text, fd_hl = val_or_highlight(filing_date)
    p = add_para(doc, f"IN WITNESS WHEREOF, the undersigned incorporator has executed these Articles of Incorporation on {fd_text}.")
    if fd_hl:
        for run in p.runs:
            add_highlight(run)

    doc.add_paragraph()
    add_para(doc, "___________________________")
    add_para(doc, f"{incorporator['name']}, Incorporator", bold=True)

    path = output_dir / "01 - Articles of Incorporation.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def gen_bylaws(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    ra = cfg['registered_agent']
    filing_date = cfg.get('filing_date', 'TBD')
    fiscal = cfg.get('fiscal_year_end', 'December 31')
    dirs = cfg['directors']

    add_heading(doc, f"BYLAWS OF {co.upper()}")
    add_para(doc, f"(A {state} Corporation)", align=WD_ALIGN_PARAGRAPH.CENTER)
    fd_text, _ = val_or_highlight(filing_date)
    add_para(doc, f"Adopted as of {fd_text}", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    sections = [
        ("ARTICLE I — OFFICES", [
            ("Section 1.1 — Principal Office", f"The principal office of the Corporation shall be located at {ra['street']}, {ra['city']}, {ra['state']} {ra['zip']} or at such other place as the Board of Directors may designate from time to time."),
            ("Section 1.2 — Registered Office", f"The registered office of the Corporation in the State of {state} shall be at the office of its registered agent, currently {ra['name']}, {ra['street']}, {ra['city']}, {ra['state']} {ra['zip']}."),
        ]),
        ("ARTICLE II — STOCKHOLDERS", [
            ("Section 2.1 — Annual Meeting", "The annual meeting of the stockholders shall be held at such date, time, and place (or by remote communication) as determined by the Board of Directors."),
            ("Section 2.2 — Special Meetings", "Special meetings of the stockholders may be called by the Board of Directors, the Chairman, or the President."),
            ("Section 2.3 — Notice", "Written notice stating the place, day, and hour of the meeting shall be delivered not less than ten (10) nor more than sixty (60) days before the date of the meeting."),
            ("Section 2.4 — Quorum", "The holders of a majority of the voting power of the issued and outstanding stock entitled to vote, present in person or represented by proxy, shall constitute a quorum."),
            ("Section 2.5 — Voting", "Each stockholder shall be entitled to the number of votes equal to the number of votes per share held, as set forth in the Articles of Incorporation."),
            ("Section 2.6 — Action Without Meeting", "Any action required or permitted to be taken at any meeting of stockholders may be taken without a meeting if a consent in writing is signed by the holders of outstanding stock having not less than the minimum number of votes necessary to authorize such action."),
        ]),
        ("ARTICLE III — BOARD OF DIRECTORS", [
            ("Section 3.1 — General Powers", "The business and affairs of the Corporation shall be managed by or under the direction of the Board of Directors."),
            ("Section 3.2 — Number", f"The number of directors shall be fixed from time to time by resolution of the Board, but shall not be less than one (1). The initial Board consists of {len(dirs)} director(s)."),
            ("Section 3.3 — Election and Term", "Directors shall be elected at each annual meeting of stockholders and shall hold office until the next annual meeting and until their successors are elected and qualified."),
            ("Section 3.4 — Vacancies", "Any vacancy occurring in the Board of Directors may be filled by the affirmative vote of a majority of the remaining directors."),
            ("Section 3.5 — Quorum", "A majority of the total number of directors shall constitute a quorum for the transaction of business."),
            ("Section 3.6 — Action Without Meeting", "Any action required or permitted to be taken at any meeting of the Board may be taken without a meeting if all members of the Board consent thereto in writing."),
        ]),
        ("ARTICLE IV — OFFICERS", [
            ("Section 4.1 — Officers", "The officers of the Corporation shall consist of a President (or CEO), a Secretary, and a Treasurer. The Board may also elect other officers as it deems necessary."),
            ("Section 4.2 — Election and Term", "Officers shall be elected by the Board of Directors and shall serve at the pleasure of the Board."),
            ("Section 4.3 — Removal", "Any officer may be removed by the Board of Directors at any time, with or without cause."),
        ]),
        ("ARTICLE V — STOCK", [
            ("Section 5.1 — Certificates", "Shares of stock may be certificated or uncertificated, as determined by the Board of Directors."),
            ("Section 5.2 — Transfers", "Transfers of stock shall be made on the books of the Corporation upon surrender of the certificate (if certificated)."),
        ]),
        ("ARTICLE VI — INDEMNIFICATION", [
            ("", f"The Corporation shall indemnify its directors, officers, employees, and agents to the fullest extent permitted by the laws of the State of {state}."),
        ]),
        ("ARTICLE VII — FISCAL YEAR", [
            ("", f"The fiscal year of the Corporation shall end on {fiscal} of each year."),
        ]),
        ("ARTICLE VIII — AMENDMENTS", [
            ("", "These Bylaws may be amended or repealed, and new Bylaws may be adopted, by the Board of Directors or by the stockholders."),
        ]),
    ]

    for article_title, subs in sections:
        add_heading(doc, article_title, level=2)
        for sub_title, text in subs:
            if sub_title:
                p = doc.add_paragraph()
                run = p.add_run(sub_title)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = BLACK
            add_para(doc, text)

    doc.add_paragraph()
    add_para(doc, f"Adopted by the Board of Directors on {fd_text}.", bold=True)

    path = output_dir / "02 - Bylaws.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def gen_action(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    filing_date = cfg.get('filing_date', 'TBD')
    incorporator = next(d for d in cfg['directors'] if d.get('is_incorporator'))
    dirs = cfg['directors']

    add_heading(doc, f"ACTION OF SOLE INCORPORATOR OF {co.upper()}")
    add_para(doc, f"(A {state} Corporation)", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_para(doc, f'The undersigned, being the sole incorporator of {co} (the "Corporation"), a corporation organized under the laws of the State of {state}, hereby takes the following actions:')
    doc.add_paragraph()

    # Filing date — this one stays TBD (needs SilverFlume confirmation)
    p = add_para(doc, "")
    run = p.add_run(f"RESOLVED, that the Articles of Incorporation of the Corporation, as filed with the {state} Secretary of State on ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.bold = True
    run2 = p.add_run("[FILING DATE — enter after SilverFlume confirmation]")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)
    run2.font.color.rgb = BLACK
    run2.bold = True
    add_highlight(run2)
    run3 = p.add_run(", are hereby ratified and approved.")
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(11)
    run3.font.color.rgb = BLACK
    run3.bold = True

    doc.add_paragraph()
    add_para(doc, "RESOLVED, that the following persons are hereby appointed as the initial directors of the Corporation:", bold=True)
    for d in dirs:
        add_para(doc, f"• {d['name']}")

    doc.add_paragraph()
    add_para(doc, "RESOLVED, that the initial directors are authorized and directed to take all actions necessary to organize the Corporation, including but not limited to the adoption of Bylaws, the election of officers, and the authorization and issuance of shares.", bold=True)

    doc.add_paragraph()
    doc.add_paragraph()
    fd_text, _ = val_or_highlight(filing_date)
    add_para(doc, f"IN WITNESS WHEREOF, the undersigned incorporator has executed this Action as of {fd_text}.")
    doc.add_paragraph()
    add_para(doc, "___________________________")
    add_para(doc, f"{incorporator['name']}, Sole Incorporator", bold=True)

    path = output_dir / "03 - Action of Incorporator.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def gen_resolutions(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    ra = cfg['registered_agent']
    filing_date = cfg.get('filing_date', 'TBD')
    fiscal = cfg.get('fiscal_year_end', 'December 31')
    dirs = cfg['directors']

    principal = cfg.get('principal_office', {})
    if principal and principal.get('street'):
        office = f"{principal['street']}, {principal['city']}, {principal['state']} {principal['zip']}"
    else:
        office = f"{ra['street']}, {ra['city']}, {ra['state']} {ra['zip']}"

    add_heading(doc, f"ORGANIZATIONAL RESOLUTIONS OF THE BOARD OF DIRECTORS OF {co.upper()}")
    add_para(doc, f"(A {state} Corporation)", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_para(doc, f'The undersigned, constituting the entire Board of Directors of {co} (the "Corporation"), acting by unanimous written consent without a meeting, hereby adopt the following resolutions:')
    doc.add_paragraph()

    add_heading(doc, "BYLAWS", level=2)
    add_para(doc, "RESOLVED, that the Bylaws presented to and reviewed by the Board of Directors are hereby adopted as the Bylaws of the Corporation, and that the Secretary is directed to insert a copy of such Bylaws in the minute book of the Corporation.", bold=True)

    add_heading(doc, "OFFICERS", level=2)
    add_para(doc, "RESOLVED, that the following persons are hereby elected as officers:", bold=True)
    for d in dirs:
        for title in d.get('titles', []):
            if title != 'Director':
                add_para(doc, f"• {title}: {d['name']}")

    add_heading(doc, "PRINCIPAL OFFICE", level=2)
    add_para(doc, "RESOLVED, that the principal office of the Corporation shall be located at:", bold=True)
    add_para(doc, office)

    add_heading(doc, "STOCK ISSUANCE", level=2)
    stock = cfg.get('stock', {})
    shareholders = cfg.get('shareholders', {})

    for class_key in ['preferred', 'common']:
        holders = shareholders.get(class_key, [])
        if not holders:
            continue
        if class_key == 'preferred':
            for pref_class in stock.get('preferred', []):
                class_holders = [h for h in holders if h.get('class', '') == pref_class['class_name']]
                if class_holders:
                    add_para(doc, f"RESOLVED, that the Corporation issue {pref_class['class_name']} as follows:", bold=True)
                    t = doc.add_table(rows=1 + len(class_holders), cols=3)
                    set_table_borders(t)
                    add_table_header(t, ["Stockholder", "Shares", "Consideration"])
                    for i, h in enumerate(class_holders):
                        set_cell(t.rows[i+1].cells[0], h['name'])
                        set_cell(t.rows[i+1].cells[1], f"{h['shares']:,}")
                        set_cell(t.rows[i+1].cells[2], h.get('consideration', cfg.get('options', {}).get('consideration_default', 'Services')))
                    doc.add_paragraph()
        else:
            add_para(doc, "RESOLVED, that the Corporation issue Common Stock as follows:", bold=True)
            t = doc.add_table(rows=1 + len(holders), cols=3)
            set_table_borders(t)
            add_table_header(t, ["Stockholder", "Shares", "Consideration"])
            for i, h in enumerate(holders):
                set_cell(t.rows[i+1].cells[0], h['name'])
                set_cell(t.rows[i+1].cells[1], f"{h['shares']:,}")
                set_cell(t.rows[i+1].cells[2], h.get('consideration', cfg.get('options', {}).get('consideration_default', 'Services')))
            doc.add_paragraph()

    add_heading(doc, "BANKING", level=2)
    add_para(doc, "RESOLVED, that the officers are authorized to open bank accounts in the name of the Corporation.", bold=True)

    add_heading(doc, "EIN", level=2)
    add_para(doc, "RESOLVED, that the officers are authorized to apply for a federal Employer Identification Number.", bold=True)

    add_heading(doc, "FISCAL YEAR", level=2)
    add_para(doc, f"RESOLVED, that the fiscal year shall end on {fiscal} of each year.", bold=True)

    add_heading(doc, "GENERAL AUTHORITY", level=2)
    add_para(doc, "RESOLVED, that the officers are authorized to take all actions necessary to carry out the foregoing resolutions.", bold=True)

    doc.add_paragraph()
    doc.add_paragraph()
    fd_text, _ = val_or_highlight(filing_date)
    add_para(doc, f"IN WITNESS WHEREOF, the undersigned directors have executed these resolutions as of {fd_text}.")
    for d in dirs:
        doc.add_paragraph()
        add_para(doc, "___________________________")
        add_para(doc, f"{d['name']}, Director", bold=True)

    path = output_dir / "04 - Organizational Resolutions.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def gen_stock_ledger(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    filing_date = cfg.get('filing_date', 'TBD')
    stock = cfg.get('stock', {})
    shareholders = cfg.get('shareholders', {})
    default_consideration = cfg.get('options', {}).get('consideration_default', 'Services')

    add_heading(doc, f"STOCK LEDGER — {co.upper()}")
    add_para(doc, f"(A {state} Corporation)", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    all_holders = {}  # name -> {class: shares, votes}

    # Preferred sections
    for pref in stock.get('preferred', []):
        class_name = pref['class_name']
        holders = [h for h in shareholders.get('preferred', []) if h.get('class') == class_name]
        total_issued = sum(h['shares'] for h in holders)

        add_heading(doc, class_name.upper(), level=2)
        add_para(doc, f"Authorized: {pref['authorized']:,} shares  |  Par Value: ${pref['par_value']}  |  Voting: {pref['voting_rights']} vote(s) per share")
        add_para(doc, f"Total Issued: {total_issued:,} / {pref['authorized']:,} authorized")
        doc.add_paragraph()

        t = doc.add_table(rows=1 + len(holders), cols=6)
        set_table_borders(t)
        add_table_header(t, ["Cert. No.", "Date Issued", "Stockholder", "Shares", "Consideration", "Status"])

        for i, h in enumerate(holders):
            cert = f"SV-{i+1:03d}"
            fd, fd_hl = val_or_highlight(filing_date)
            consid = h.get('consideration', default_consideration)
            set_cell(t.rows[i+1].cells[0], cert)
            set_cell(t.rows[i+1].cells[1], fd, highlight_yellow=fd_hl)
            set_cell(t.rows[i+1].cells[2], h['name'])
            set_cell(t.rows[i+1].cells[3], f"{h['shares']:,}")
            set_cell(t.rows[i+1].cells[4], consid)
            set_cell(t.rows[i+1].cells[5], "Outstanding")
            if i % 2 == 1:
                shade_row(t.rows[i+1], "F2F2F2")

            # Track for voting summary
            if h['name'] not in all_holders:
                all_holders[h['name']] = {}
            all_holders[h['name']][class_name] = {
                'shares': h['shares'],
                'votes': h['shares'] * pref['voting_rights']
            }

        doc.add_paragraph()

    # Common section
    if 'common' in stock:
        c = stock['common']
        holders = shareholders.get('common', [])
        total_issued = sum(h['shares'] for h in holders)

        add_heading(doc, "COMMON STOCK", level=2)
        add_para(doc, f"Authorized: {c['authorized']:,} shares  |  Par Value: ${c['par_value']}  |  Voting: {c['voting_rights']} vote(s) per share")
        add_para(doc, f"Total Issued: {total_issued:,} / {c['authorized']:,} authorized")
        doc.add_paragraph()

        t = doc.add_table(rows=1 + len(holders), cols=6)
        set_table_borders(t)
        add_table_header(t, ["Cert. No.", "Date Issued", "Stockholder", "Shares", "Consideration", "Status"])

        for i, h in enumerate(holders):
            cert = f"CS-{i+1:03d}"
            fd, fd_hl = val_or_highlight(filing_date)
            consid = h.get('consideration', default_consideration)
            name_hl = '[' in h['name']  # highlight names with [TBD] etc
            set_cell(t.rows[i+1].cells[0], cert)
            set_cell(t.rows[i+1].cells[1], fd, highlight_yellow=fd_hl)
            set_cell(t.rows[i+1].cells[2], h['name'], highlight_yellow=name_hl)
            set_cell(t.rows[i+1].cells[3], f"{h['shares']:,}")
            set_cell(t.rows[i+1].cells[4], consid)
            set_cell(t.rows[i+1].cells[5], "Outstanding")
            if i % 2 == 1:
                shade_row(t.rows[i+1], "F2F2F2")

            if h['name'] not in all_holders:
                all_holders[h['name']] = {}
            all_holders[h['name']]['Common'] = {
                'shares': h['shares'],
                'votes': h['shares'] * c['voting_rights']
            }

        doc.add_paragraph()

    # Voting Power Summary
    add_heading(doc, "VOTING POWER SUMMARY", level=2)
    doc.add_paragraph()

    classes = [p['class_name'] for p in stock.get('preferred', [])] + (['Common'] if 'common' in stock else [])
    headers = ["Stockholder"]
    for cls in classes:
        headers.extend([f"{cls[:20]} Shares", f"{cls[:20]} Votes"])
    headers.extend(["Total Votes", "% of Total"])

    grand_total = sum(
        sum(cl['votes'] for cl in holder_data.values())
        for holder_data in all_holders.values()
    )

    t = doc.add_table(rows=1 + len(all_holders) + 1, cols=len(headers))
    set_table_borders(t)
    add_table_header(t, headers)

    for ri, (name, holder_data) in enumerate(all_holders.items()):
        row = t.rows[ri + 1]
        set_cell(row.cells[0], name)
        ci = 1
        total_votes = 0
        for cls in classes:
            data = holder_data.get(cls, {'shares': 0, 'votes': 0})
            set_cell(row.cells[ci], f"{data['shares']:,}" if data['shares'] else "—")
            set_cell(row.cells[ci+1], f"{data['votes']:,}" if data['votes'] else "—")
            total_votes += data['votes']
            ci += 2
        set_cell(row.cells[ci], f"{total_votes:,}")
        pct = f"{total_votes/grand_total*100:.2f}%" if grand_total else "0%"
        set_cell(row.cells[ci+1], pct)
        if ri % 2 == 1:
            shade_row(row, "F2F2F2")

    # Total row
    total_row = t.rows[-1]
    set_cell(total_row.cells[0], "TOTAL", bold=True)
    ci = 1
    for cls in classes:
        ts = sum(h.get(cls, {}).get('shares', 0) for h in all_holders.values())
        tv = sum(h.get(cls, {}).get('votes', 0) for h in all_holders.values())
        set_cell(total_row.cells[ci], f"{ts:,}", bold=True)
        set_cell(total_row.cells[ci+1], f"{tv:,}", bold=True)
        ci += 2
    set_cell(total_row.cells[ci], f"{grand_total:,}", bold=True)
    set_cell(total_row.cells[ci+1], "100%", bold=True)
    shade_row(total_row, "D6DCE4")

    doc.add_paragraph()
    add_para(doc, f"Maintained by the Secretary of the Corporation. Last updated: {filing_date}")

    path = output_dir / "05 - Stock Ledger.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def gen_checklist(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    incorporator = next(d for d in cfg['directors'] if d.get('is_incorporator'))

    add_heading(doc, f"{co.upper()} — INCORPORATION CHECKLIST")
    doc.add_paragraph()

    steps = [
        ("PRE-FILING", [
            "Confirm registered agent",
            "Confirm principal office address",
            "Confirm personal addresses for directors",
        ]),
        ("STEP 1: FILE ARTICLES OF INCORPORATION", [
            "File online at Nevada SilverFlume (nvsilverflume.gov)",
            "Pay filing fee (~$75)",
            "Save stamped/filed copy for records",
            "Note Entity Number assigned by NV SOS",
        ]),
        ("STEP 2: EXECUTE INCORPORATOR ACTION", [
            f"{incorporator['name']} signs Action of Incorporator (appoints board)",
        ]),
        ("STEP 3: ADOPT BYLAWS & ORGANIZATIONAL RESOLUTIONS", [
            "All directors sign Organizational Resolutions",
            "Bylaws considered adopted per resolution",
        ]),
        ("STEP 4: ISSUE STOCK", [
            "Record all issuances in Stock Ledger",
            "Issue certificates or record as uncertificated shares",
        ]),
        ("STEP 5: FILE INITIAL LIST OF OFFICERS/DIRECTORS", [
            "File within 30 days of incorporation at SilverFlume",
            "Pay filing fee (~$350 including Business License)",
            "List all directors and officers with titles",
        ]),
        ("STEP 6: OBTAIN EIN", [
            "Apply online at IRS EIN Assistant (irs.gov)",
            "Free, takes ~5 minutes",
            "Save confirmation letter",
        ]),
        ("STEP 7: OPEN BANK ACCOUNT", [
            "Bring: EIN confirmation, filed Articles, Org Resolutions, IDs",
            "Open business checking account",
        ]),
        ("STEP 8: FOREIGN QUALIFICATION (IF NEEDED)", [
            "If operating in another state, register as foreign corporation",
            "Requires Certificate of Good Standing from Nevada + filing in target state",
        ]),
    ]

    for title, items in steps:
        add_heading(doc, title, level=2)
        for item in items:
            add_para(doc, f"☐  {item}")

    doc.add_paragraph()
    add_heading(doc, "ESTIMATED COSTS", level=2)
    costs = doc.add_table(rows=5, cols=2)
    set_table_borders(costs)
    add_table_header(costs, ["Item", "Cost"])
    cost_data = [
        ("Articles of Incorporation filing", "~$75"),
        ("Initial List + Business License", "~$350"),
        ("Registered Agent (annual)", "~$100-150"),
        ("EIN", "Free"),
    ]
    for i, (item, cost) in enumerate(cost_data):
        set_cell(costs.rows[i+1].cells[0], item)
        set_cell(costs.rows[i+1].cells[1], cost)

    doc.add_paragraph()
    add_heading(doc, "DOCUMENTS IN THIS PACKAGE", level=2)
    docs_list = [
        "01 - Articles of Incorporation — File with NV Secretary of State",
        "02 - Bylaws — Internal governance (keep in records)",
        "03 - Action of Incorporator — Incorporator signs to appoint board",
        "04 - Organizational Resolutions — Board adopts bylaws, elects officers, issues stock",
        "05 - Stock Ledger — Track all share ownership",
    ]
    for d in docs_list:
        add_para(doc, f"• {d}")

    path = output_dir / "06 - Filing Checklist.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def gen_83b_election(cfg, output_dir):
    doc = create_doc()
    co = cfg['company_name']
    state = cfg['state']
    filing_date = cfg.get('filing_date', 'TBD')
    filing_year = filing_date[:4] if filing_date and filing_date != 'TBD' else '20__'
    stock = cfg.get('stock', {})
    vesting = cfg.get('options', {})
    vesting_period = vesting.get('vesting_period', '4 years')
    cliff_period = vesting.get('cliff_period', '1 year')
    par_value = stock.get('common', {}).get('par_value', '0.00001')

    add_heading(doc, "SECTION 83(b) ELECTION")
    add_para(doc, f"For stockholders of {co}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_heading(doc, "INSTRUCTIONS", level=2)
    instructions = [
        f"TIMING: This election MUST be filed with the IRS within 30 days after stock is transferred (issued). There are NO extensions. Missing this deadline is irreversible.",
        f"WHERE TO FILE: Mail the completed, signed election to the IRS Service Center where you file your federal income tax return. Send via certified mail, return receipt requested.",
        f"COPIES: Provide a copy to {co}. Keep a copy for your personal records.",
        f"TAX RETURN: Attach a copy of this election to your federal income tax return for the year of transfer.",
        f"WHY THIS MATTERS: Without 83(b), you will be taxed on fair market value of shares as they vest (potentially much higher). With 83(b), you are taxed on value at issuance (par value = ${par_value}/share = essentially $0).",
    ]
    for i, inst in enumerate(instructions):
        add_para(doc, f"{i+1}. {inst}")

    doc.add_paragraph()
    add_heading(doc, "ELECTION FORM", level=2)
    add_para(doc, "The undersigned taxpayer hereby elects, pursuant to Section 83(b) of the Internal Revenue Code of 1986, as amended, to include in gross income as compensation for services the excess (if any) of the fair market value of the property described below over the amount paid for such property.")
    doc.add_paragraph()

    fields = [
        ("1. TAXPAYER INFORMATION:", [
            "Name: _________________________________",
            "Address: _________________________________",
            "Social Security Number: _________________________________",
            f"Tax Year: Calendar year {filing_year}",
        ]),
        ("2. DESCRIPTION OF PROPERTY:", [
            f"_____________ shares of Common Stock of {co}, a {state} corporation",
        ]),
        ("3. DATE OF TRANSFER:", [
            f"{filing_date} (the date of issuance per the Organizational Resolutions)",
        ]),
        ("4. NATURE OF RESTRICTION:", [
            f"The shares are subject to a Restricted Stock Purchase Agreement with a {vesting_period} vesting schedule and {cliff_period} cliff. Unvested shares are subject to repurchase by the Corporation at par value (${par_value} per share) upon termination of service.",
        ]),
        ("5. FAIR MARKET VALUE AT TIME OF TRANSFER:", [
            f"${par_value} per share (par value). The Corporation is newly formed and has no operations, revenue, or assets beyond its initial capitalization.",
            "Total fair market value: $_____________ (number of shares × $" + par_value + ")",
        ]),
        ("6. AMOUNT PAID FOR PROPERTY:", [
            "$0.00 (shares were issued in exchange for services to be rendered)",
        ]),
        ("7. AMOUNT TO INCLUDE IN GROSS INCOME:", [
            "$_____________ (fair market value minus amount paid = essentially $0)",
        ]),
    ]

    for title, items in fields:
        add_para(doc, title, bold=True)
        for item in items:
            add_para(doc, f"   {item}")

    add_para(doc, f"8. A copy of this election has been furnished to {co}.", bold=True)
    add_para(doc, f"9. This election is being made in connection with the initial issuance of restricted stock by {co}.", bold=True)

    doc.add_paragraph()
    add_para(doc, "Date: _________________________________")
    add_para(doc, "Signature: _________________________________")
    add_para(doc, "Print Name: _________________________________")
    add_para(doc, "Spouse Signature (if applicable): _________________________________")

    doc.add_paragraph()
    add_heading(doc, "FILING CHECKLIST", level=2)
    checklist = [
        "Complete the election form (fill in all blanks)",
        "Sign and date",
        "Make 4 copies of the signed election",
        "Mail original to IRS Service Center via CERTIFIED MAIL, RETURN RECEIPT REQUESTED",
        "Keep the certified mail receipt as proof of timely filing",
        f"Give one copy to the Secretary of {co}",
        "Keep one copy for personal tax records",
        f"Attach one copy to {filing_year} federal income tax return when filed",
    ]
    for item in checklist:
        add_para(doc, f"☐  {item}")

    doc.add_paragraph()
    fd_text, _ = val_or_highlight(filing_date)
    p = add_para(doc, "")
    run = p.add_run(f"CRITICAL DEADLINE: Must be postmarked within 30 days of stock issuance date ({fd_text}).")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.bold = True
    add_highlight(run)

    path = output_dir / "07 - 83b Election Form.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name}")
    return path


def main():
    parser = argparse.ArgumentParser(description="Generate incorporation documents from config")
    parser.add_argument("config", help="Path to YAML config file")
    parser.add_argument("--output-dir", default="./output", help="Output directory for .docx files")
    args = parser.parse_args()

    cfg = load_config(args.config)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n📋 Generating incorporation documents for: {cfg['company_name']}")
    print(f"   Entity type: {cfg['entity_type']}")
    print(f"   State: {cfg['state']}")
    print(f"   Output: {output_dir}\n")

    gen_articles(cfg, output_dir)
    gen_bylaws(cfg, output_dir)
    gen_action(cfg, output_dir)
    gen_resolutions(cfg, output_dir)
    gen_stock_ledger(cfg, output_dir)
    gen_checklist(cfg, output_dir)

    # Generate 83(b) if stock is issued for services
    shareholders = cfg.get('shareholders', {})
    has_services = any(
        h.get('consideration', cfg.get('options', {}).get('consideration_default', '')).lower() == 'services'
        for holders in shareholders.values()
        for h in holders
    )
    if has_services:
        gen_83b_election(cfg, output_dir)
        print(f"\n⚠️  83(b) election generated — stockholders must file within 30 DAYS of stock issuance!")

    print(f"\n✅ All documents generated in {output_dir}/")
    print("   Review yellow-highlighted fields for any items needing manual input.")


if __name__ == "__main__":
    main()

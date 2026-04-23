"""
Word文件解析器

支持解析 .docx 格式的财务报告
"""
import re
from typing import Dict, List, Any, Optional

from .base_parser import BaseParser, ParseResult, FinancialData, normalize_account_name, extract_number


class WordParser(BaseParser):
    """
    Word文件解析器 (.docx)
    
    支持从Word文档中提取表格形式的财务数据
    """
    
    name = "word_parser"
    description = "Word文件解析器 (.docx)"
    supported_extensions = ['.docx']
    supported_software = ['GENERIC']
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        self.extract_tables = self.config.get('extract_tables', True)
        self.extract_text = self.config.get('extract_text', False)
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否能解析该文件"""
        if not self.validate_file_exists(file_path):
            return False
        
        ext = self.get_file_extension(file_path)
        return ext == '.docx'
    
    def parse(self, file_path: str) -> ParseResult:
        """解析Word文件"""
        if not self.can_parse(file_path):
            return self.create_error_result(file_path, "无法解析该文件格式")
        
        try:
            try:
                from docx import Document
            except ImportError:
                return self.create_error_result(
                    file_path,
                    "缺少依赖库：请安装 python-docx (pip install python-docx)"
                )
            
            doc = Document(file_path)
            
            result = self.create_success_result(file_path)
            result.file_type = 'WORD'
            result.data = FinancialData()
            result.data.source_file = file_path
            result.data.source_software = 'GENERIC'
            
            # 提取文档标题作为公司名称的候选
            company_name = self._extract_company_name(doc)
            if company_name:
                result.data.company_name = company_name
            
            # 提取表格数据
            if self.extract_tables and doc.tables:
                self._extract_tables(doc.tables, result.data, result)
            
            # 提取文本数据（可选）
            if self.extract_text:
                text_data = self._extract_text(doc)
                result.data.raw_data['text_content'] = text_data
            
            return result
            
        except Exception as e:
            return self.create_error_result(file_path, f"解析失败: {str(e)}")
    
    def _extract_company_name(self, doc) -> Optional[str]:
        """从文档中提取公司名称"""
        # 通常在文档开头或标题中
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            # 匹配常见公司名称模式
            if any(suffix in text for suffix in ['有限公司', '股份公司', '集团', 'Company', 'Inc.']):
                # 清理标题中的其他文字
                company = self._clean_company_name(text)
                if company:
                    return company
        return None
    
    def _clean_company_name(self, text: str) -> Optional[str]:
        """清理公司名称"""
        # 移除常见的前缀和后缀
        prefixes = ['关于', '审计报告', '财务报表', '年度报告']
        for prefix in prefixes:
            if text.startswith(prefix):
                text = text[len(prefix):].strip()
        
        # 尝试提取公司名称
        # 匹配 "XXX有限公司" 或 "XXX股份公司" 等模式
        patterns = [
            r'([^\n]{2,30}?(?:有限公司|有限责任公司|股份公司|股份有限公司))',
            r'([^\n]{2,50}?(?:Company|Corporation|Inc\.))',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return text[:50] if len(text) <= 50 else None
    
    def _extract_tables(self, tables, data: FinancialData, result: ParseResult):
        """提取表格数据"""
        for i, table in enumerate(tables):
            # 尝试识别表格类型
            table_type = self._identify_table_type(table)
            
            if table_type == 'balance_sheet':
                data.balance_sheet = self._parse_table(table)
                result.add_warning(f"表格{i+1}识别为资产负债表")
            elif table_type == 'income_statement':
                data.income_statement = self._parse_table(table)
                result.add_warning(f"表格{i+1}识别为利润表")
            elif table_type == 'cash_flow':
                data.cash_flow = self._parse_table(table)
                result.add_warning(f"表格{i+1}识别为现金流量表")
            else:
                # 未识别的表格，保存到raw_data
                table_data = self._parse_table(table)
                if table_data:
                    data.raw_data[f'table_{i+1}'] = table_data
    
    def _identify_table_type(self, table) -> str:
        """识别表格类型"""
        # 读取表格前几行的文本
        sample_text = ""
        for i, row in enumerate(table.rows):
            if i >= 5:  # 只读取前5行
                break
            row_text = ' '.join([cell.text for cell in row.cells])
            sample_text += row_text + ' '
        
        sample_text = sample_text.lower()
        
        # 关键字匹配
        if any(kw in sample_text for kw in ['资产', '负债', 'asset', 'liability']):
            return 'balance_sheet'
        
        if any(kw in sample_text for kw in ['收入', '成本', '利润', 'revenue', 'profit']):
            return 'income_statement'
        
        if any(kw in sample_text for kw in ['现金', '流量', 'cash', 'flow']):
            return 'cash_flow'
        
        return 'unknown'
    
    def _parse_table(self, table) -> Dict[str, float]:
        """解析表格数据"""
        data = {}
        
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            
            # 第一列通常是科目名称
            account_name = cells[0].text.strip()
            if not account_name:
                continue
            
            standard_name = normalize_account_name(account_name)
            
            # 查找数值（通常在后面的列）
            for cell in reversed(cells[1:]):
                value = extract_number(cell.text)
                if value is not None:
                    data[standard_name] = value
                    break
        
        return data
    
    def _extract_text(self, doc) -> str:
        """提取文档文本"""
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n'.join(paragraphs)


# 兼容 .doc 格式的解析器（使用 antiword 或其他工具）
class DOCParser(WordParser):
    """DOC格式解析器（旧版Word）"""
    
    name = "doc_parser"
    description = "DOC文件解析器（旧版Word）"
    supported_extensions = ['.doc']
    
    def parse(self, file_path: str) -> ParseResult:
        """解析DOC文件"""
        # DOC格式需要额外的工具支持
        # 可以使用 antiword, textract, 或转换为docx后解析
        
        return self.create_error_result(
            file_path,
            "DOC格式暂不支持，请将文件转换为DOCX格式后重试"
        )


__all__ = ['WordParser', 'DOCParser']

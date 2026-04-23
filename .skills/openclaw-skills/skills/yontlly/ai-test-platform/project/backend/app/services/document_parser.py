"""
文档解析服务

支持 Word、Excel、PDF、Markdown 格式文档解析
"""

import os
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器"""

    @staticmethod
    def parse_word(file_path: str) -> str:
        """
        解析 Word 文档

        Args:
            file_path: 文件路径

        Returns:
            文档文本内容
        """
        try:
            from docx import Document
            doc = Document(file_path)

            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)

            return '\n'.join(text_content)

        except Exception as e:
            logger.error(f"解析Word文档失败: {str(e)}")
            raise ValueError(f"解析Word文档失败: {str(e)}")

    @staticmethod
    def parse_excel(file_path: str) -> str:
        """
        解析 Excel 文档

        Args:
            file_path: 文件路径

        Returns:
            文档文本内容
        """
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path)

            text_content = []
            for sheet in wb.worksheets:
                text_content.append(f"\n=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text_content.append(row_text)

            return '\n'.join(text_content)

        except Exception as e:
            logger.error(f"解析Excel文档失败: {str(e)}")
            raise ValueError(f"解析Excel文档失败: {str(e)}")

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        解析 PDF 文档

        Args:
            file_path: 文件路径

        Returns:
            文档文本内容
        """
        try:
            import pdfplumber

            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            return '\n'.join(text_content)

        except Exception as e:
            logger.error(f"解析PDF文档失败: {str(e)}")
            raise ValueError(f"解析PDF文档失败: {str(e)}")

    @staticmethod
    def parse_markdown(file_path: str) -> str:
        """
        解析 Markdown 文档

        Args:
            file_path: 文件路径

        Returns:
            文档文本内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        except Exception as e:
            logger.error(f"解析Markdown文档失败: {str(e)}")
            raise ValueError(f"解析Markdown文档失败: {str(e)}")

    @classmethod
    def parse_document(cls, file_path: str) -> str:
        """
        自动识别文档类型并解析

        Args:
            file_path: 文件路径

        Returns:
            文档文本内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 获取文件扩展名
        ext = os.path.splitext(file_path)[1].lower()

        # 根据扩展名选择解析器
        if ext == '.docx':
            return cls.parse_word(file_path)
        elif ext in ['.xlsx', '.xls']:
            return cls.parse_excel(file_path)
        elif ext == '.pdf':
            return cls.parse_pdf(file_path)
        elif ext in ['.md', '.markdown']:
            return cls.parse_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

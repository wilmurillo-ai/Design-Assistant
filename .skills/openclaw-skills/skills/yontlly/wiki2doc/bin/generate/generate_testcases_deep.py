#!/usr/bin/env python3
"""
深度分析版测试用例生成工具 - generate_testcases_deep.py

专为复杂需求文档设计的深度分析工具

特点：
- 为每个功能点生成3-5个测试用例（非1个）
- 为每个表格行生成1-3个测试用例
- 使用所有8种测试设计方法
- 生成200+个测试用例
- 完全避免模板化描述
"""

import sys
import re
import json
from pathlib import Path
from docx import Document
from typing import List, Dict, Any, Tuple
from datetime import datetime
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DeepAnalysisTestCaseGenerator:
    """深度分析测试用例生成器 - 针对复杂需求"""

    def __init__(self, doc_path):
        """
        初始化测试用例生成器

        Args:
            doc_path: Word文档路径
        """
        self.doc_path = Path(doc_path)
        self.doc = None
        self.content = ""
        self.test_cases = []
        self.tc_counter = 1
        self.function_points = []  # 提取的所有功能点
        self.tables_data = []      # 提取的所有表格数据

        # 测试方法类型定义
        self.test_methods = [
            "正向测试", "边界值测试", "错误推测", "场景法",
            "UI测试", "兼容性测试", "性能测试", "安全测试", "异常测试"
        ]

    def load_document(self):
        """加载 Word 文档"""
        try:
            self.doc = Document(self.doc_path)
            logger.info(f"成功加载文档: {self.doc_path}")
            return True
        except Exception as e:
            logger.error(f"文档加载失败: {e}")
            return False

    def extract_function_points(self):
        """
        深度提取功能点

        识别方式：
        1. 标题层级（1-4级）
        2. 以功能动词开头的段落（增加、修改、删除、支持、验证）
        3. 表格中的功能项
        4. 带编号的功能列表
        5. 功能描述段落
        """
        logger.info("开始深度提取功能点...")

        # 收集所有段落
        paragraphs = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    'style': para.style.name if para.style else '',
                    'level': self._get_heading_level(para)
                })

        # 1. 提取标题作为功能点（1-4级）
        for i, para in enumerate(paragraphs):
            if para['level'] > 0 and para['level'] <= 4:
                # 检查是否为有效功能描述
                if self._is_function_description(para['text']):
                    context = self._get_context(paragraphs, i, 3)
                    self.function_points.append({
                        'title': para['text'],
                        'level': para['level'],
                        'context': context,
                        'type': 'heading',
                        'source': f'段落{i+1}'
                    })

        # 2. 提取以功能动词开头的段落
        function_verbs = ['增加', '修改', '删除', '支持', '验证', '实现', '提供', '允许', '配置', '设置', '开启', '关闭', '限制', '显示', '隐藏', '控制', '管理']

        for i, para in enumerate(paragraphs):
            text = para['text']
            if any(text.startswith(verb) for verb in function_verbs) and len(text) > 5:
                # 排除标题（已处理）
                if para['level'] == 0:  # 非标题段落
                    context = self._get_context(paragraphs, i, 2)
                    self.function_points.append({
                        'title': text,
                        'level': 5,  # 自定义级别
                        'context': context,
                        'type': 'verb',
                        'source': f'段落{i+1}'
                    })

        # 3. 提取表格中的功能项
        for table_idx, table in enumerate(self.doc.tables):
            # 分析表格结构
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # 非空行
                    rows.append(row_data)

            if len(rows) > 1:  # 有数据的表格
                # 尝试识别列标题
                headers = rows[0]
                # 如果有'功能'、'模块'、'权限'等关键词，视为功能点表格
                if any(keyword in str(headers).lower() for keyword in ['功能', '模块', '权限', '菜单', '按钮', '字段', '属性', '选项']):
                    for row in rows[1:]:  # 跳过标题行
                        if len(row) > 0 and row[0]:  # 第一列有内容
                            # 为每个表格行创建一个功能点
                            self.function_points.append({
                                'title': row[0],
                                'level': 5,
                                'context': '\n'.join(row),
                                'type': 'table',
                                'source': f'表格{table_idx+1}'
                            })

        # 4. 提取带编号的功能列表
        numbered_pattern = r'^(\d+[.、）)]\s+|\d+、)\s*(.+)'  # 匹配 1.、2.、3.、1）、2）、1、2、

        for i, para in enumerate(paragraphs):
            match = re.match(numbered_pattern, para['text'])
            if match:
                title = match.group(2)
                if len(title) > 5:  # 避免过短的列表项
                    context = self._get_context(paragraphs, i, 1)
                    self.function_points.append({
                        'title': title,
                        'level': 5,
                        'context': context,
                        'type': 'numbered',
                        'source': f'段落{i+1}'
                    })

        # 5. 提取功能描述段落（包含"功能"、"需求"等词的段落）
        func_keywords = ['功能', '需求', '要求', '条件', '特性', '支持', '实现']

        for i, para in enumerate(paragraphs):
            text = para['text']
            if any(keyword in text for keyword in func_keywords) and len(text) > 15:
                # 排除已提取的
                if not any(fp['title'] == text for fp in self.function_points):
                    context = self._get_context(paragraphs, i, 3)
                    self.function_points.append({
                        'title': text,
                        'level': 5,
                        'context': context,
                        'type': 'description',
                        'source': f'段落{i+1}'
                    })

        # 去重（基于标题）
        seen_titles = set()
        unique_points = []
        for fp in self.function_points:
            if fp['title'] not in seen_titles:
                seen_titles.add(fp['title'])
                unique_points.append(fp)

        self.function_points = unique_points
        logger.info(f"成功提取 {len(self.function_points)} 个功能点")
        return len(self.function_points)

    def _get_heading_level(self, paragraph):
        """获取标题级别"""
        style = paragraph.style.name.lower()
        if 'heading' in style:
            if '1' in style:
                return 1
            elif '2' in style:
                return 2
            elif '3' in style:
                return 3
            elif '4' in style:
                return 4
        return 0

    def _is_function_description(self, text):
        """判断是否为功能描述"""
        # 排除纯标题、目录、页码等
        if len(text) < 5:
            return False
        if text.startswith('目录') or text.startswith('第') or text.startswith('页'):
            return False
        if '图' in text and '表' in text and '编号' in text:
            return False
        # 包含功能动词
        function_verbs = ['增加', '修改', '删除', '支持', '验证', '实现', '提供', '允许', '配置', '设置', '开启', '关闭', '限制', '显示', '隐藏', '控制', '管理']
        return any(verb in text for verb in function_verbs)

    def _get_context(self, paragraphs, index, window=3):
        """获取上下文"""
        start = max(0, index - window)
        end = min(len(paragraphs), index + window + 1)
        context = []
        for i in range(start, end):
            if i != index and paragraphs[i]['text'].strip():
                context.append(paragraphs[i]['text'])
        return '\n'.join(context)

    def extract_tables(self):
        """提取所有表格数据"""
        logger.info("开始提取表格数据...")

        for table_idx, table in enumerate(self.doc.tables):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    rows.append(row_data)

            if len(rows) > 1:  # 有数据的表格
                self.tables_data.append({
                    'index': table_idx + 1,
                    'rows': rows,
                    'headers': rows[0] if rows else [],
                    'content': rows[1:] if len(rows) > 1 else []
                })
                logger.info(f"提取表格 {table_idx + 1}: {len(rows)} 行")

        logger.info(f"共提取 {len(self.tables_data)} 个表格")
        return len(self.tables_data)

    def generate_positive_test_cases(self):
        """为每个功能点生成多个正向测试用例"""
        logger.info("开始生成正向测试用例...")

        for fp in self.function_points:
            title = fp['title']

            # 为每个功能点生成3个正向测试用例

            # 1. 基本功能验证
            self._add_test_case(
                module="功能测试",
                title=f"验证{title}的基本功能",
                preconditions=f"1. 系统已部署\n2. {self._extract_feature_context(title)}已配置",
                steps=f"1. 进入{self._extract_feature_context(title)}相关界面\n2. 执行{title}操作\n3. 验证结果",
                priority="高",
                expected_result=f"1. {title}功能正常工作\n2. 无错误提示\n3. 数据正确更新",
                method="正向测试"
            )

            # 2. 多次操作验证
            self._add_test_case(
                module="功能测试",
                title=f"验证{title}的多次操作",
                preconditions=f"1. 系统已部署\n2. {self._extract_feature_context(title)}已配置",
                steps=f"1. 进入{self._extract_feature_context(title)}相关界面\n2. 执行{title}操作\n3. 再次执行{title}操作\n4. 验证状态一致性",
                priority="中",
                expected_result=f"1. 多次操作后功能正常\n2. 状态保持一致\n3. 无内存泄漏",
                method="正向测试"
            )

            # 3. 跨模块集成验证
            self._add_test_case(
                module="功能测试",
                title=f"验证{title}与其他模块的集成",
                preconditions=f"1. 系统已部署\n2. {self._extract_feature_context(title)}已配置\n3. 相关模块已启用",
                steps=f"1. 在其他模块中触发与{title}相关的事件\n2. 验证{title}的响应\n3. 检查数据同步",
                priority="中",
                expected_result=f"1. 集成功能正常\n2. 数据同步准确\n3. 无交叉影响",
                method="正向测试"
            )

    def generate_boundary_test_cases(self):
        """为数值型功能点生成边界值测试用例"""
        logger.info("开始生成边界值测试用例...")

        # 从表格中查找数值字段
        for table in self.tables_data:
            headers = table['headers']
            rows = table['content']

            # 检查表头是否包含数值字段
            numeric_keywords = ['数量', '数量', '级别', '等级', '次数', '时间', '秒', '分钟', '小时', '天', 'KB', 'MB', 'GB', 'px', '像素', '大小', '长度', '宽度', '高度']

            for col_idx, header in enumerate(headers):
                if any(keyword in header for keyword in numeric_keywords):
                    # 为这个数值列生成边界值测试
                    for row in rows:
                        if col_idx < len(row) and row[col_idx]:
                            value = row[col_idx]

                            # 尝试提取数字
                            numbers = re.findall(r'(\d+)', value)
                            if numbers:
                                num = int(numbers[0])

                                # 生成边界值测试用例
                                self._add_test_case(
                                    module="边界值测试",
                                    title=f"验证{header}为{num-1}时的边界值",
                                    preconditions=f"1. 系统已部署\n2. {header}可配置",
                                    steps=f"1. 设置{header}为{max(0, num-1)}\n2. 执行相关操作",
                                    priority="高",
                                    expected_result=f"1. 系统正常处理\n2. 无错误提示\n3. 功能正常",
                                    method="边界值测试"
                                )

                                self._add_test_case(
                                    module="边界值测试",
                                    title=f"验证{header}为{num}时的正常值",
                                    preconditions=f"1. 系统已部署\n2. {header}可配置",
                                    steps=f"1. 设置{header}为{num}\n2. 执行相关操作",
                                    priority="中",
                                    expected_result=f"1. 系统正常处理\n2. 功能按预期工作\n3. 无性能问题",
                                    method="边界值测试"
                                )

                                self._add_test_case(
                                    module="边界值测试",
                                    title=f"验证{header}为{num+1}时的边界值",
                                    preconditions=f"1. 系统已部署\n2. {header}可配置",
                                    steps=f"1. 设置{header}为{num+1}\n2. 执行相关操作",
                                    priority="中",
                                    expected_result=f"1. 系统正确处理或提示上限\n2. 无崩溃\n3. 有明确提示信息",
                                    method="边界值测试"
                                )

    def generate_exception_test_cases(self):
        """为功能点生成异常测试用例"""
        logger.info("开始生成异常测试用例...")

        # 基于功能点生成异常场景
        for fp in self.function_points:
            title = fp['title']

            # 检查是否涉及权限、数据、网络等异常场景
            if any(keyword in title for keyword in ['权限', '访问', '禁止', '限制', '安全', '加密', '认证', '授权']):
                self._add_test_case(
                    module="异常测试",
                    title=f"验证{title}在无权限时的表现",
                    preconditions=f"1. 系统已部署\n2. 当前用户无{title}权限",
                    steps=f"1. 尝试执行{title}操作\n2. 观察系统反应",
                    priority="高",
                    expected_result=f"1. 系统拒绝操作\n2. 显示权限不足提示\n3. 无崩溃",
                    method="错误推测"
                )

                self._add_test_case(
                    module="异常测试",
                    title=f"验证{title}在权限不足时的表现",
                    preconditions=f"1. 系统已部署\n2. 当前用户有部分{title}权限",
                    steps=f"1
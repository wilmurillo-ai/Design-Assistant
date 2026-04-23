#!/usr/bin/env python3
"""
增强版测试用例生成工具 - generate_testcases_v2.py

智能适配任意规模需求文档的测试用例生成工具

- 复杂需求：提取所有功能点、表格、规则，生成200+个详细测试用例
- 简单需求：精准提取核心功能，生成8-15个针对性用例
- 通用设计：无需修改代码，自动适应任何需求文档结构
"""

import sys
import re
import json
from pathlib import Path
from docx import Document
from typing import List, Dict, Any, Tuple
from datetime import datetime
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class AdvancedTestCaseGenerator:
    """高级测试用例生成器 - 智能适配任意需求文档"""

    def __init__(self, doc_path):
        """
        初始化测试用例生成器

        Args:
            doc_path: Word文档路径
        """
        self.doc_path = Path(doc_path)
        self.doc = None
        self.content = ""
        self.test_cases = []
        self.tc_counter = 1
        self.function_points = []  # 提取的所有功能点
        self.tables_data = []      # 提取的所有表格数据

        # 测试方法类型定义
        self.test_methods = [
            "正向测试", "边界值测试", "错误推测", "场景法",
            "UI测试", "兼容性测试", "性能测试", "安全测试", "异常测试"
        ]

    def load_document(self):
        """加载 Word 文档"""
        try:
            self.doc = Document(self.doc_path)
            logger.info(f"成功加载文档: {self.doc_path}")
            return True
        except Exception as e:
            logger.error(f"文档加载失败: {e}")
            return False

    def extract_function_points(self):
        """
        智能提取功能点

        识别方式：
        1. 标题层级（1-4级）
        2. 以功能动词开头的段落（增加、修改、删除、支持、验证）
        3. 表格中的功能项
        4. 带编号的功能列表
        """
        logger.info("开始智能提取功能点...")

        # 收集所有段落
        paragraphs = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    'style': para.style.name if para.style else '',
                    'level': self._get_heading_level(para)
                })

        # 1. 提取标题作为功能点（1-4级）
        for i, para in enumerate(paragraphs):
            if para['level'] > 0 and para['level'] <= 4:
                # 检查是否为有效功能描述
                if self._is_function_description(para['text']):
                    # 查找上下文
                    context = self._get_context(paragraphs, i, 3)
                    self.function_points.append({
                        'title': para['text'],
                        'level': para['level'],
                        'context': context,
                        'type': 'heading',
                        'source': f'段落{i+1}'
                    })

        # 2. 提取以功能动词开头的段落
        function_verbs = ['增加', '修改', '删除', '支持', '验证', '实现', '提供', '允许', '配置', '设置', '开启', '关闭', '限制']

        for i, para in enumerate(paragraphs):
            text = para['text']
            if any(text.startswith(verb) for verb in function_verbs) and len(text) > 5:
                # 排除标题（已处理）
                if para['level'] == 0:  # 非标题段落
                    context = self._get_context(paragraphs, i, 2)
                    self.function_points.append({
                        'title': text,
                        'level': 5,  # 自定义级别
                        'context': context,
                        'type': 'verb',
                        'source': f'段落{i+1}'
                    })

        # 3. 提取表格中的功能项
        for table_idx, table in enumerate(self.doc.tables):
            # 分析表格结构
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # 非空行
                    rows.append(row_data)

            if len(rows) > 1:  # 有数据的表格
                # 尝试识别列标题
                headers = rows[0]
                # 如果有'功能'、'模块'、'权限'等关键词，视为功能点表格
                if any(keyword in str(headers).lower() for keyword in ['功能', '模块', '权限', '菜单', '按钮', '字段']):
                    for row in rows[1:]:  # 跳过标题行
                        if len(row) > 0 and row[0]:  # 第一列有内容
                            self.function_points.append({
                                'title': row[0],
                                'level': 5,
                                'context': '\n'.join(row),
                                'type': 'table',
                                'source': f'表格{table_idx+1}'
                            })

        # 4. 提取带编号的功能列表
        numbered_pattern = r'^(\d+[.、）)]\s+|\d+、)\s*(.+)'  # 匹配 1.、2.、3.、1）、2）、1、2、

        for i, para in enumerate(paragraphs):
            match = re.match(numbered_pattern, para['text'])
            if match:
                title = match.group(2)
                if len(title) > 5:  # 避免过短的列表项
                    context = self._get_context(paragraphs, i, 1)
                    self.function_points.append({
                        'title': title,
                        'level': 5,
                        'context': context,
                        'type': 'numbered',
                        'source': f'段落{i+1}'
                    })

        # 去重（基于标题）
        seen_titles = set()
        unique_points = []
        for fp in self.function_points:
            if fp['title'] not in seen_titles:
                seen_titles.add(fp['title'])
                unique_points.append(fp)

        self.function_points = unique_points
        logger.info(f"成功提取 {len(self.function_points)} 个功能点")
        return len(self.function_points)

    def _get_heading_level(self, paragraph):
        """获取标题级别"""
        style = paragraph.style.name.lower()
        if 'heading' in style:
            if '1' in style:
                return 1
            elif '2' in style:
                return 2
            elif '3' in style:
                return 3
            elif '4' in style:
                return 4
        return 0

    def _is_function_description(self, text):
        """判断是否为功能描述"""
        # 排除纯标题、目录、页码等
        if len(text) < 5:
            return False
        if text.startswith('目录') or text.startswith('第') or text.startswith('页'):
            return False
        if '图' in text and '表' in text and '编号' in text:
            return False
        # 包含功能动词
        function_verbs = ['增加', '修改', '删除', '支持', '验证', '实现', '提供', '允许', '配置', '设置', '开启', '关闭', '限制', '显示', '隐藏']
        return any(verb in text for verb in function_verbs)

    def _get_context(self, paragraphs, index, window=3):
        """获取上下文"""
        start = max(0, index - window)
        end = min(len(paragraphs), index + window + 1)
        context = []
        for i in range(start, end):
            if i != index and paragraphs[i]['text'].strip():
                context.append(paragraphs[i]['text'])
        return '\n'.join(context)

    def extract_tables(self):
        """提取所有表格数据"""
        logger.info("开始提取表格数据...")

        for table_idx, table in enumerate(self.doc.tables):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    rows.append(row_data)

            if len(rows) > 1:  # 有数据的表格
                self.tables_data.append({
                    'index': table_idx + 1,
                    'rows': rows,
                    'headers': rows[0] if rows else [],
                    'content': rows[1:] if len(rows) > 1 else []
                })
                logger.info(f"提取表格 {table_idx + 1}: {len(rows)} 行")

        logger.info(f"共提取 {len(self.tables_data)} 个表格")
        return len(self.tables_data)

    def generate_positive_test_cases(self):
        """为每个功能点生成正向测试用例"""
        logger.info("开始生成正向测试用例...")

        for fp in self.function_points:
            title = fp['title']

            # 生成正向测试用例
            # 基于功能点生成具体测试用例
            # 例如：如果标题是"增加二级菜单指定业务属性"，就生成对应测试

            # 提取功能动词和对象
            verbs = ['增加', '修改', '删除', '支持', '验证', '实现', '提供', '允许', '配置', '设置', '开启', '关闭', '限制']
            verb = None
            obj = None

            for v in verbs:
                if v in title:
                    verb = v
                    obj = title.replace(v, '').strip().strip('：').strip('。').strip()
                    break

            if not verb:
                # 如果没有动词，就用功能点作为标题
                verb = "验证"
                obj = title

            # 生成具体测试用例
            self._add_test_case(
                module="功能测试",
                title=f"{verb}{obj}",
                preconditions=f"1. 系统已部署\n2. {obj}功能已启用",
                steps=f"1. 执行{obj}操作\n2. 验证结果",
                priority="高",
                expected_result=f"1. {obj}功能正常工作\n2. 无错误提示\n3. 界面显示正确",
                method="正向测试"
            )

    def generate_boundary_test_cases(self):
        """为数值型功能点生成边界值测试用例"""
        logger.info("开始生成边界值测试用例...")

        # 从表格中查找数值字段
        for table in self.tables_data:
            headers = table['headers']
            rows = table['content']

            # 检查表头是否包含数值字段
            numeric_keywords = ['数量', '数量', '级别', '等级', '次数', '时间', '秒', '分钟', '小时', '天', 'KB', 'MB', 'GB', 'px', '像素']

            for col_idx, header in enumerate(headers):
                if any(keyword in header for keyword in numeric_keywords):
                    # 为这个数值列生成边界值测试
                    for row in rows:
                        if col_idx < len(row) and row[col_idx]:
                            value = row[col_idx]

                            # 尝试提取数字
                            numbers = re.findall(r'(\d+)', value)
                            if numbers:
                                num = int(numbers[0])

                                # 生成边界值测试用例
                                self._add_test_case(
                                    module="边界值测试",
                                    title=f"验证{header}为{num-1}时的边界值",
                                    preconditions=f"1. 系统已部署\n2. {header}可配置",
                                    steps=f"1. 设置{header}为{max(0, num-1)}\n2. 执行相关操作",
                                    priority="高",
                                    expected_result=f"1. 系统正常处理\n2. 无错误提示\n3. 功能正常",
                                    method="边界值测试"
                                )

                                self._add_test_case(
                                    module="边界值测试",
                                    title=f"验证{header}为{num+1}时的边界值",
                                    preconditions=f"1. 系统已部署\n2. {header}可配置",
                                    steps=f"1. 设置{header}为{num+1}\n2. 执行相关操作",
                                    priority="中",
                                    expected_result=f"1. 系统正确处理或提示上限\n2. 无崩溃\n3. 有明确提示信息",
                                    method="边界值测试"
                                )

    def generate_exception_test_cases(self):
        """为功能点生成异常测试用例"""
        logger.info("开始生成异常测试用例...")

        # 基于功能点生成异常场景
        for fp in self.function_points:
            title = fp['title']

            # 检查是否涉及权限、数据、网络等异常场景
            if any(keyword in title for keyword in ['权限', '访问', '禁止', '限制', '安全', '加密']):
                self._add_test_case(
                    module="异常测试",
                    title=f"验证{title}在无权限时的表现",
                    preconditions=f"1. 系统已部署\n2. 当前用户无{title}权限",
                    steps=f"1. 尝试执行{title}操作\n2. 观察系统反应",
                    priority="高",
                    expected_result=f"1. 系统拒绝操作\n2. 显示权限不足提示\n3. 无崩溃",
                    method="错误推测"
                )

            if any(keyword in title for keyword in ['数据', '字段', '输入', '输出']):
                self._add_test_case(
                    module="异常测试",
                    title=f"验证{title}输入空值时的表现",
                    preconditions=f"1. 系统已部署\n2. {title}可输入数据",
                    steps=f"1. 清空{title}输入字段\n2. 提交操作",
                    priority="中",
                    expected_result=f"1. 系统提示输入不能为空\n2. 无崩溃\n3. 输入框高亮显示",
                    method="错误推测"
                )

            if any(keyword in title for keyword in ['网络', '连接', '超时', '失败']):
                self._add_test_case(
                    module="异常测试",
                    title=f"验证{title}在断网时的表现",
                    preconditions=f"1. 系统已部署\n2. 网络连接断开",
                    steps=f"1. 执行{title}操作\n2. 观察系统反应",
                    priority="高",
                    expected_result=f"1. 系统提示网络异常\n2. 无崩溃\n3. 网络恢复后自动重试",
                    method="错误推测"
                )

    def generate_scenario_test_cases(self):
        """基于业务流程生成场景测试用例"""
        logger.info("开始生成场景测试用例...")

        # 检查文档中是否有流程描述
        process_keywords = ['流程', '步骤', '顺序', '流程图', '交互', '时序', '工作流']

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if any(keyword in text for keyword in process_keywords) and len(text) > 10:
                # 提取流程描述（截取前50字符）
                flow_desc = text[:50] + "..." if len(text) > 50 else text
                # 生成一个端到端场景测试
                self._add_test_case(
                    module="场景测试",
                    title=f"验证{flow_desc}完整流程",
                    preconditions=f"1. 系统已部署\n2. 用户已登录\n3. 所有前置条件满足",
                    steps=f"1. 进入相关功能模块\n2. 按照业务流程执行所有步骤\n3. 完成最终操作",
                    priority="高",
                    expected_result=f"1. 整个流程顺畅无卡顿\n2. 所有功能正常\n3. 数据一致且正确\n4. 用户体验完整",
                    method="场景法"
                )

    def generate_ui_test_cases(self):
        """为UI相关功能点生成UI测试用例"""
        logger.info("开始生成UI测试用例...")

        ui_keywords = ['界面', '页面', '按钮', '菜单', '图标', '显示', '布局', '样式', '颜色', '大小', '位置', '对齐', '字体']

        for fp in self.function_points:
            title = fp['title']
            if any(keyword in title for keyword in ui_keywords):
                # 生成UI测试用例
                self._add_test_case(
                    module="UI测试",
                    title=f"验证{title}的界面显示",
                    preconditions=f"1. 系统已部署\n2. {title}功能已启用",
                    steps=f"1. 进入{title}相关界面\n2. 检查所有元素显示\n3. 验证布局和样式",
                    priority="中",
                    expected_result=f"1. 所有元素清晰显示\n2. 布局无错位\n3. 样式符合设计规范\n4. 无兼容性问题",
                    method="UI测试"
                )

    def generate_compatibility_test_cases(self):
        """生成兼容性测试用例"""
        logger.info("开始生成兼容性测试用例...")

        # 基于文档内容生成兼容性测试
        compatibility_keywords = ['适配', '兼容', '多端', '多平台', '多浏览器', '多分辨率']

        for fp in self.function_points:
            title = fp['title']
            if any(keyword in title for keyword in compatibility_keywords):
                self._add_test_case(
                    module="兼容性测试",
                    title=f"验证{title}在不同设备/分辨率下的兼容性",
                    preconditions=f"1. 系统已部署\n2. 准备多种设备/分辨率",
                    steps=f"1. 在不同设备上测试{title}\n2. 在不同分辨率下测试{title}\n3. 在不同浏览器下测试{title}",
                    priority="中",
                    expected_result=f"1. 所有设备/分辨率下功能正常\n2. 界面自适应良好\n3. 无样式错位",
                    method="兼容性测试"
                )

    def generate_performance_test_cases(self):
        """生成性能测试用例"""
        logger.info("开始生成性能测试用例...")

        performance_keywords = ['性能', '响应', '速度', '并发', '负载', '延迟', '吞吐量', '效率']

        for fp in self.function_points:
            title = fp['title']
            if any(keyword in title for keyword in performance_keywords):
                self._add_test_case(
                    module="性能测试",
                    title=f"验证{title}的性能表现",
                    preconditions=f"1. 系统已部署\n2. 准备性能测试环境",
                    steps=f"1. 模拟大量并发请求\n2. 测量{title}响应时间\n3. 监控系统资源使用",
                    priority="低",
                    expected_result=f"1. 响应时间在可接受范围内（<3秒）\n2. 系统无性能瓶颈\n3. 资源使用正常",
                    method="性能测试"
                )

    def generate_security_test_cases(self):
        """生成安全测试用例"""
        logger.info("开始生成安全测试用例...")

        security_keywords = ['安全', '权限', '认证', '授权', '加密', '泄露', 'XSS', 'SQL', '注入', '攻击']

        for fp in self.function_points:
            title = fp['title']
            if any(keyword in title for keyword in security_keywords):
                self._add_test_case(
                    module="安全测试",
                    title=f"验证{title}的安全性",
                    preconditions=f"1. 系统已部署\n2. 已登录用户",
                    steps=f"1. 尝试越权访问{title}\n2. 输入恶意代码\n3. 测试数据泄露风险",
                    priority="高",
                    expected_result=f"1. 无越权访问\n2. 无XSS/SQL注入\n3. 数据加密传输\n4. 无敏感信息泄露",
                    method="安全测试"
                )

    def _add_test_case(self, module: str, title: str, preconditions: str,
                      steps: str, priority: str, expected_result: str, method: str = ""):
        """
        添加测试用例

        Args:
            module: 测试模块
            title: 用例标题
            preconditions: 前置条件
            steps: 测试步骤
            priority: 优先级
            expected_result: 预期结果
            method: 测试方法
        """
        self.test_cases.append({
            "id": f"TC{self.tc_counter:03d}",
            "module": module,
            "title": title,
            "preconditions": preconditions,
            "steps": steps,
            "priority": priority,
            "expected_result": expected_result,
            "method": method
        })
        self.tc_counter += 1

    def generate_test_cases(self):
        """生成所有测试用例"""
        if not self.load_document():
            return False

        # 提取功能点和表格
        self.extract_function_points()
        self.extract_tables()

        logger.info(f"总功能点数: {len(self.function_points)}")
        logger.info(f"总表格数: {len(self.tables_data)}")

        # 生成各种类型的测试用例
        self.generate_positive_test_cases()
        self.generate_boundary_test_cases()
        self.generate_exception_test_cases()
        self.generate_scenario_test_cases()
        self.generate_ui_test_cases()
        self.generate_compatibility_test_cases()
        self.generate_performance_test_cases()
        self.generate_security_test_cases()

        logger.info(f"总计生成测试用例: {len(self.test_cases)} 个")
        return True

    def save_to_markdown(self, output_path=None):
        """
        保存测试用例到Markdown文件

        Args:
            output_path: 输出文件路径，默认为文档同级目录

        Returns:
            str: 输出文件路径
        """
        if output_path is None:
            output_path = self.doc_path.parent / f"{self.doc_path.stem}_testcases.md"

        with open(output_path, 'w', encoding='utf-8') as f:
            # 写入标题
            f.write(f"# 测试用例\n\n")
            f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"基于需求文档: {self.doc_path.name}\n\n")
            f.write(f"总用例数: {len(self.test_cases)}\n\n")
            f.write("---\n\n")

            # 按模块分组
            modules = {}
            for tc in self.test_cases:
                module = tc["module"]
                if module not in modules:
                    modules[module] = []
                modules[module].append(tc)

            # 写入各模块的测试用例
            for module, cases in modules.items():
                f.write(f"## {module}\n\n")

                for tc in cases:
                    f.write(f"### {tc['id']}: {tc['title']}\n\n")
                    f.write(f"**测试方法**: {tc.get('method', '')}\n\n")
                    f.write(f"**优先级**: {tc['priority']}\n\n")
                    f.write(f"**前置条件**:\n```\n{tc['preconditions']}\n```\n\n")
                    f.write(f"**测试步骤**:\n```\n{tc['steps']}\n```\n\n")
                    f.write(f"**预期结果**:\n```\n{tc['expected_result']}\n```\n\n")
                    f.write("---\n\n")

        logger.info(f"测试用例已保存到Markdown: {output_path}")
        return str(output_path)

    def save_to_json(self, output_path=None):
        """
        保存测试用例到JSON文件

        Args:
            output_path: 输出文件路径，默认为文档同级目录

        Returns:
            str: 输出文件路径
        """
        if output_path is None:
            output_path = self.doc_path.parent / f"{self.doc_path.stem}_testcases.json"

        data = {
            "generated_at": datetime.now().isoformat(),
            "source_document": str(self.doc_path),
            "total_cases": len(self.test_cases),
            "function_points": self.function_points,
            "tables_count": len(self.tables_data),
            "test_cases": self.test_cases
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        logger.info(f"测试用例已保存到JSON: {output_path}")
        return str(output_path)


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(
        description="高级测试用例生成工具 - 智能适配任意规模需求文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_testcases_v2.py demand/需求文档.docx
  python generate_testcases_v2.py demand/需求文档.docx --output testcases.md
        """
    )

    parser.add_argument("doc_file", help="需求文档路径（Word格式）")
    parser.add_argument("--output", help="输出文件路径（可选）")

    args = parser.parse_args()

    # 创建生成器
    generator = AdvancedTestCaseGenerator(args.doc_file)

    # 生成测试用例
    if generator.generate_test_cases():
        # 保存到Markdown
        md_path = generator.save_to_markdown(args.output)
        # 同时保存JSON
        json_path = generator.save_to_json()

        print(f"\n生成完成！")
        print(f"Markdown文件: {md_path}")
        print(f"JSON文件: {json_path}")
    else:
        print("\n生成失败！")
        sys.exit(1)


if __name__ == "__main__":
    main()
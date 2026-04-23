#!/usr/bin/env python3
"""
Confluence Wiki to Word Document Converter

Extracts content from Confluence wiki pages and saves as .docx files.
"""

import argparse
import os
import io
import re
import shutil
import time
from pathlib import Path
from urllib.parse import urljoin

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError

# Configuration
BASE_URL = "http://10.225.1.76:8090"
LOGIN_URL = f"{BASE_URL}/login.action"
USERNAME = "yanghua"
PASSWORD = "Aa123123"
OUTPUT_DIR = Path.home() / ".claude" / "skills" / "wiki2doc" / "demand"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Temporary directory for images
TEMP_DIR = OUTPUT_DIR / "temp_images"
TEMP_DIR.mkdir(exist_ok=True)


def extract_title(soup):
    """Extract page title from #title-text element."""
    title_elem = soup.select_one("#title-text")
    if title_elem:
        return title_elem.get_text(strip=True)
    return f"Untitled_{int(time.time())}"


def clean_filename(filename):
    """Clean filename to remove invalid characters."""
    # Replace invalid characters with underscore
    cleaned = re.sub(r'[\\/:*?"<>|]', '_', filename)
    # Limit length to 50 characters
    return cleaned[:50] if len(cleaned) > 50 else cleaned


def extract_table_with_merges(table_element):
    """Extract table data with merge information (rowspan and colspan)."""
    rows_data = []

    # Get all rows
    rows = table_element.select("tr")

    # Track which cells are occupied by rowspan
    rowspan_tracker = {}  # key: (row_idx, col_idx), value: remaining rows

    for row_idx, tr in enumerate(rows):
        row_data = []
        col_idx = 0

        # Get all cells in this row
        cells = tr.select("td, th")

        for cell in cells:
            # Skip columns that are occupied by rowspan from previous rows
            while (row_idx, col_idx) in rowspan_tracker:
                # This cell is occupied, add placeholder
                row_data.append({"text": "", "rowspan": 0, "colspan": 0, "occupied": True})
                col_idx += 1

            # Get cell properties
            text = cell.get_text(strip=True)
            rowspan = int(cell.get("rowspan", 1))
            colspan = int(cell.get("colspan", 1))

            # Add cell data
            row_data.append({
                "text": text,
                "rowspan": rowspan,
                "colspan": colspan,
                "occupied": False
            })

            # Mark cells that will be occupied by rowspan
            if rowspan > 1:
                for r in range(1, rowspan):
                    rowspan_tracker[(row_idx + r, col_idx)] = rowspan - r

            # Move to next column (account for colspan)
            col_idx += colspan

        # Fill remaining columns that might be occupied by rowspan
        while (row_idx, col_idx) in rowspan_tracker:
            row_data.append({"text": "", "rowspan": 0, "colspan": 0, "occupied": True})
            col_idx += 1

        rows_data.append(row_data)

    return rows_data


def is_hidden_by_css(element):
    """Check if element is hidden by CSS."""
    style = element.get("style", "").lower()
    if "display:none" in style or "display: none" in style:
        return True
    if "visibility:hidden" in style or "visibility: hidden" in style:
        return True
    if "opacity:0" in style or "opacity: 0" in style:
        return True
    return False


def extract_content(soup):
    """Extract text and image content from the main page content."""
    # Try to find the main content container
    content_div = soup.select_one("#page, #main-content, .page-body, .wiki-content")
    if not content_div:
        return []

    content_items = []

    # Walk through all elements in document order
    for element in content_div.descendants:
        if element.name == "p":
            text = element.get_text(strip=True)
            if text:
                content_items.append(("text", text))
        elif element.name == "h1":
            text = element.get_text(strip=True)
            if text:
                content_items.append(("heading1", text))
        elif element.name == "h2":
            text = element.get_text(strip=True)
            if text:
                content_items.append(("heading2", text))
        elif element.name == "h3":
            text = element.get_text(strip=True)
            if text:
                content_items.append(("heading3", text))
        elif element.name == "ul" or element.name == "ol":
            text = element.get_text(strip=True)
            if text:
                content_items.append(("list", text))
        elif element.name == "table":
            # Extract table data with merge information
            table_data = extract_table_with_merges(element)
            if table_data:
                content_items.append(("table", table_data))
        elif element.name == "pre" or element.name == "code":
            text = element.get_text(strip=True)
            if text:
                content_items.append(("code", text))
        elif element.name == "img":
            # Skip hidden images
            if is_hidden_by_css(element):
                continue

            src = element.get("src")
            if src:
                # Skip non-image formats
                if any(src.lower().endswith(ext) for ext in [".gif", ".svg", ".webp", ".mp4", ".avi", ".mov"]):
                    continue

                # Get alt text for accessibility
                alt = element.get("alt", "")
                content_items.append(("image", {"src": src, "alt": alt}))

    return content_items


def download_image(page, img_src, temp_dir):
    """Download image using Playwright and convert to JPG."""
    try:
        # Resolve relative URLs
        if not img_src.startswith("http"):
            img_src = urljoin(BASE_URL, img_src)

        # Generate unique filename
        img_hash = abs(hash(img_src))
        filename = f"img_{img_hash}.jpg"
        filepath = temp_dir / filename

        # Download image using Playwright's request context
        response = page.request.get(img_src)
        if response.ok:
            image_bytes = response.body()

            # Load image
            img = Image.open(io.BytesIO(image_bytes))

            # Convert to RGB if needed (for JPG)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")

            # Save as JPG
            img.save(filepath, "JPEG", quality=85)

            return filename
        else:
            print(f"Warning: Failed to download {img_src} (status: {response.status})")
            return None
    except Exception as e:
        print(f"Warning: Failed to process image {img_src}: {e}")
        return None


def create_word_document(title, content_items, image_dir):
    """Create a Word document with text and images."""
    doc = Document()

    # Add title
    doc.add_heading(title, level=0)

    # Process content items in order
    for item_type, item_data in content_items:
        if item_type == "text":
            doc.add_paragraph(item_data)
        elif item_type == "heading1":
            doc.add_heading(item_data, level=1)
        elif item_type == "heading2":
            doc.add_heading(item_data, level=2)
        elif item_type == "heading3":
            doc.add_heading(item_data, level=3)
        elif item_type == "list":
            # Add list items
            for line in item_data.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")
        elif item_type == "code":
            # Add code block with monospace font
            para = doc.add_paragraph(item_data)
            for run in para.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        elif item_type == "table":
            # Add table with merged cells support
            if item_data:
                # Calculate actual column count (accounting for colspan)
                max_cols = max(len(row) for row in item_data)
                table = doc.add_table(rows=len(item_data), cols=max_cols)
                table.style = "Light Grid Accent 1"

                # Fill table and handle merges
                for i, row_data in enumerate(item_data):
                    for j, cell_info in enumerate(row_data):
                        if isinstance(cell_info, dict):
                            # New format with merge information
                            if not cell_info.get("occupied", False):
                                cell = table.cell(i, j)
                                cell.text = cell_info.get("text", "")

                                # Handle colspan and rowspan
                                rowspan = cell_info.get("rowspan", 1)
                                colspan = cell_info.get("colspan", 1)

                                if rowspan > 1 or colspan > 1:
                                    # Merge cells
                                    try:
                                        end_row = min(i + rowspan - 1, len(item_data) - 1)
                                        end_col = min(j + colspan - 1, max_cols - 1)

                                        # Only merge if we have valid range
                                        if i <= end_row and j <= end_col:
                                            cell.merge(table.cell(end_row, end_col))
                                    except Exception as e:
                                        print(f"Warning: Failed to merge cells ({i},{j}): {e}")
                        else:
                            # Old format (simple string)
                            table.rows[i].cells[j].text = cell_info
        elif item_type == "image":
            # Add image
            img_filename = item_data.get("filename")
            if img_filename:
                img_path = image_dir / img_filename
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(5))
                        # Add alt text as caption if available
                        alt_text = item_data.get("alt", "")
                        if alt_text:
                            doc.add_paragraph(alt_text, style="Caption")
                    except Exception as e:
                        print(f"Warning: Failed to add image to document: {e}")

    # Save document
    doc_path = OUTPUT_DIR / f"{clean_filename(title)}.docx"

    # Handle file name conflicts
    counter = 1
    while doc_path.exists():
        doc_path = OUTPUT_DIR / f"{clean_filename(title)}_{counter}.docx"
        counter += 1

    doc.save(doc_path)
    return doc_path


def process_single_url(browser, url):
    """Process a single wiki URL and generate Word document."""
    try:
        print(f"\nProcessing: {url}")

        # Create new page
        page = browser.new_page()

        # Navigate directly to target URL - Confluence will redirect to login if needed
        print(f"Navigating to: {url}")
        page.goto(url, wait_until="networkidle", timeout=30000)

        # Check if we are on the login page
        current_url = page.url.lower()
        if "login" in current_url or "login.action" in current_url:
            print("Detected login page, logging in...")

            # Fill login form
            page.fill('input[name="os_username"]', USERNAME)
            page.fill('input[name="os_password"]', PASSWORD)

            # Submit the login form
            page.evaluate('''() => {
                const form = document.querySelector('form[name="loginform"]');
                if (form) {
                    form.submit();
                }
            }''')

            # Wait for login to complete and redirect to target page
            page.wait_for_load_state("networkidle", timeout=15000)
            print("Login successful, redirected to target page")
        else:
            print("Page loaded (already authenticated or no login required)")

        # Wait for splitter-content to appear
        print("Waiting for #splitter-content element...")
        try:
            page.wait_for_selector("#splitter-content", timeout=10000)
            print("Found #splitter-content element")
        except:
            print("Warning: #splitter-content not found within timeout")
            # Save debug info
            debug_file = OUTPUT_DIR / f"debug_no_splitter_{int(time.time())}.html"
            with open(debug_file, 'w', encoding='utf-8') as f:
                f.write(page.content())
            print(f"Debug: Saved HTML to {debug_file}")

        # Get page HTML
        html_content = page.content()

        # Debug: save HTML to file for inspection (saved to OUTPUT_DIR, not temp)
        debug_file = OUTPUT_DIR / f"debug_{int(time.time())}.html"
        with open(debug_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"Debug: Saved page HTML to {debug_file}")

        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")

        # Debug: check if common Confluence elements exist
        splitter = soup.select_one("#splitter-content")
        if not splitter:
            print("Debug: Checking for alternative content containers...")
            for selector in ['#page-content', '#main-content', '#wiki-content', '.wiki-content', '.page-content', '.page-body', '#content']:
                element = soup.select_one(selector)
                if element:
                    print(f"Found alternative: {selector}")
                    break

        # Extract title
        title = extract_title(soup)
        print(f"Title: {title}")

        # Extract content
        content_items = extract_content(soup)

        # Count items
        text_count = sum(1 for item_type, _ in content_items if item_type != "image")
        image_count = sum(1 for item_type, _ in content_items if item_type == "image")
        print(f"Extracted {text_count} text elements and {image_count} images")

        # Download images
        print(f"Downloading images...")
        image_counter = 0
        for i, (item_type, item_data) in enumerate(content_items):
            if item_type == "image":
                image_counter += 1
                img_src = item_data["src"]
                img_filename = download_image(page, img_src, TEMP_DIR)
                if img_filename:
                    # Update content_items with filename
                    content_items[i] = ("image", {**item_data, "filename": img_filename})
                    print(f"   Downloaded image {image_counter}/{image_count}")

        # Create Word document
        print(f"Generating Word document...")
        doc_path = create_word_document(title, content_items, TEMP_DIR)
        print(f"Saved to: {doc_path}")

        # Close page
        page.close()

        return True, doc_path

    except PlaywrightTimeoutError as e:
        print(f"Timeout error: {e}")
        return False, None
    except Exception as e:
        print(f"Error processing {url}: {e}")
        return False, None


def main():
    parser = argparse.ArgumentParser(
        description="Convert Confluence wiki pages to Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Single URL:
    python wiki2doc.py http://10.225.1.76:8090/pages/viewpage.action?pageId=34556052

  Batch mode (comma-separated):
    python wiki2doc.py --batch "url1,url2,url3"

  Full automation (generate test cases):
    python wiki2doc.py --auto http://10.225.1.76:8090/pages/viewpage.action?pageId=34556052
        """
    )
    parser.add_argument("urls", nargs="*", help="One or more wiki URLs")
    parser.add_argument("--batch", action="store_true", help="Enable batch mode (comma-separated URLs)")
    parser.add_argument("--auto", action="store_true", help="Enable full automation workflow (generate test cases)")

    args = parser.parse_args()

    # Parse URLs
    urls = []
    if args.batch:
        if args.urls:
            urls = [url.strip() for url in args.urls[0].split(",")]
    else:
        urls = args.urls

    if not urls:
        parser.print_help()
        print("\n[ERROR] No URLs provided.")
        return

    print(f"Total URLs to process: {len(urls)}")

    if args.auto:
        # Full automation workflow
        from wiki2testcases import WorkflowOrchestrator
        for url in urls:
            orchestrator = WorkflowOrchestrator(url)
            if not orchestrator.run():
                print(f"\n[ERROR] Full automation failed for {url}")
                return
        return

    # Launch browser for normal operation
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)

        success_count = 0
        fail_count = 0

        for i, url in enumerate(urls, 1):
            print(f"\n{'='*60}")
            print(f"Progress: {i}/{len(urls)}")
            success, _ = process_single_url(browser, url)
            if success:
                success_count += 1
            else:
                fail_count += 1

        browser.close()

    # Clean up temporary images
    if TEMP_DIR.exists():
        shutil.rmtree(TEMP_DIR)

    print(f"\n{'='*60}")
    print(f"Completed! Success: {success_count}, Failed: {fail_count}")
    print(f"Output directory: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Deterministic Word contract draft runner.

This script is the stable execution layer for the keplerjai-contract-draft skill.
It accepts a JSON job spec, edits a .docx template with python-docx, saves a new
draft, and writes simple validation artifacts beside the draft.
"""

from __future__ import annotations

import argparse
import copy
import glob
import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SECTION_END_KEYWORDS = [
    "权利与义务",
    "合作机制",
    "知识产权",
    "保密条款",
    "法律适用",
    "附则",
]

HEADING_PREFIX_PATTERN = re.compile(
    r"^\s*(?:[一二三四五六七八九十]+、|\d+(?:\.\d+)*[、.]?)\s*"
)


SKILL_SLUG = "keplerjai-contract-draft"
DEFAULT_AGENT_ID = "executive-assistant"
KEPLER_LEGAL_NAME = "成都景合开普勒科技有限公司"
KEPLER_NAME_ALIASES = {
    "成都景合开普勒科技有限公司",
    "景合开普勒",
    "开普勒",
    "开普勒公司",
    "kepler",
    "keplerjai",
}


def schema_error(message: str) -> RuntimeError:
    return RuntimeError(f"Job spec schema error: {message}")


def load_openclaw_config() -> dict:
    config_path = Path.home() / ".openclaw" / "openclaw.json"
    if not config_path.exists():
        return {}
    try:
        return json.loads(config_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def detect_agent_id(job: dict) -> str:
    raw = (
        job.get("agent_id")
        or job.get("agent")
        or os.environ.get("OPENCLAW_AGENT_ID")
        or DEFAULT_AGENT_ID
    )
    return str(raw).strip() or DEFAULT_AGENT_ID


def workspace_from_config(agent_id: str) -> Path | None:
    config = load_openclaw_config()
    agents = config.get("agents", {})

    for agent in agents.get("list", []):
        if str(agent.get("id", "")).strip() == agent_id:
            workspace = str(agent.get("workspace", "")).strip()
            if workspace:
                return Path(workspace)

    default_workspace = str(agents.get("defaults", {}).get("workspace", "")).strip()
    if default_workspace:
        return Path(default_workspace) / agent_id
    return None


def workspace_from_spec_path(spec_path: Path) -> Path | None:
    resolved_parts = spec_path.resolve().parts
    if "workspace" not in resolved_parts:
        return None
    try:
        workspace_index = resolved_parts.index("workspace")
        agent_dir = resolved_parts[workspace_index + 1]
    except (ValueError, IndexError):
        return None
    return Path(*resolved_parts[: workspace_index + 2])


def default_output_dir(job: dict, spec_path: Path) -> Path:
    agent_id = detect_agent_id(job)

    explicit_workspace = str(job.get("agent_workspace", "")).strip()
    if explicit_workspace:
        return Path(explicit_workspace) / SKILL_SLUG

    env_workspace = os.environ.get("OPENCLAW_AGENT_WORKSPACE") or os.environ.get("OPENCLAW_WORKSPACE_DIR")
    if env_workspace:
        return Path(env_workspace) / SKILL_SLUG

    config_workspace = workspace_from_config(agent_id)
    if config_workspace:
        return config_workspace / SKILL_SLUG

    spec_workspace = workspace_from_spec_path(spec_path)
    if spec_workspace:
        return spec_workspace / SKILL_SLUG

    return Path.cwd() / SKILL_SLUG


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def normalize_global_replacements(raw_value) -> tuple[list[dict], list[str]]:
    warnings: list[str] = []
    if raw_value in (None, ""):
        return [], warnings

    if isinstance(raw_value, dict):
        if "old" in raw_value and "new" in raw_value:
            warnings.append("global_replacements was provided as a single object and was auto-wrapped into a list.")
            return [raw_value], warnings
        normalized = [{"old": str(k), "new": str(v)} for k, v in raw_value.items()]
        warnings.append("global_replacements was provided as a mapping and was auto-converted into a list of {old,new}.")
        return normalized, warnings

    if isinstance(raw_value, list):
        normalized: list[dict] = []
        for idx, item in enumerate(raw_value):
            if isinstance(item, dict) and "old" in item and "new" in item:
                normalized.append(item)
                continue
            raise schema_error(f"global_replacements[{idx}] must contain old/new.")
        return normalized, warnings

    raise schema_error("global_replacements must be a list of {old,new} or a mapping.")


def normalize_cooperation_items(raw_value) -> tuple[list[dict], list[str]]:
    warnings: list[str] = []
    if raw_value in (None, ""):
        return [], warnings

    if not isinstance(raw_value, list):
        raise schema_error("cooperation_items must be a list.")

    normalized: list[dict] = []
    for idx, item in enumerate(raw_value):
        if not isinstance(item, dict):
            raise schema_error(f"cooperation_items[{idx}] must be an object with heading/body.")

        heading = item.get("heading")
        body = item.get("body")

        if heading is None and "title" in item:
            heading = item.get("title")
            warnings.append(f"cooperation_items[{idx}] used title and was auto-mapped to heading.")
        if body is None and "description" in item:
            body = item.get("description")
            warnings.append(f"cooperation_items[{idx}] used description and was auto-mapped to body.")
        if body is None and "content" in item:
            body = item.get("content")
            warnings.append(f"cooperation_items[{idx}] used content and was auto-mapped to body.")

        heading = str(heading or "").strip()
        body = str(body or "").strip()
        if not heading or not body:
            raise schema_error(f"cooperation_items[{idx}] must contain non-empty heading/body.")
        normalized.append({"heading": heading, "body": body})

    return normalized, warnings


def normalize_paragraph_overrides(raw_value) -> tuple[list[dict], list[str]]:
    warnings: list[str] = []
    if raw_value in (None, "", []):
        return [], warnings

    if isinstance(raw_value, dict):
        if "index" in raw_value and "text" in raw_value:
            warnings.append("paragraph_overrides was provided as a single object and was auto-wrapped into a list.")
            return [raw_value], warnings
        raise schema_error("paragraph_overrides must be a list of {index,text}.")

    if not isinstance(raw_value, list):
        raise schema_error("paragraph_overrides must be a list of {index,text}.")

    normalized: list[dict] = []
    for idx, item in enumerate(raw_value):
        if not isinstance(item, dict) or "index" not in item or "text" not in item:
            raise schema_error(f"paragraph_overrides[{idx}] must contain index/text.")
        normalized.append({"index": item["index"], "text": item["text"]})
    return normalized, warnings


def normalize_legacy_content_sections(raw_value) -> tuple[list[dict], list[str]]:
    warnings: list[str] = []
    if raw_value in (None, "", []):
        return [], warnings
    if not isinstance(raw_value, list):
        raise schema_error("content_sections must be a list when provided.")
    normalized: list[dict] = []
    for idx, item in enumerate(raw_value):
        if not isinstance(item, dict):
            raise schema_error(f"content_sections[{idx}] must be an object.")
        heading = str(item.get("heading") or item.get("title") or "").strip()
        body = str(item.get("body") or item.get("description") or item.get("content") or "").strip()
        if not heading or not body:
            raise schema_error(
                f"content_sections[{idx}] cannot be auto-converted. Use cooperation_items with heading/body."
            )
        normalized.append({"heading": heading, "body": body})
    warnings.append("content_sections is legacy and was auto-converted to cooperation_items.")
    return normalized, warnings


def normalize_job_schema(job: dict) -> dict:
    normalized = copy.deepcopy(job)
    schema_warnings: list[str] = []

    cooperation_items, item_warnings = normalize_cooperation_items(normalized.get("cooperation_items"))
    schema_warnings.extend(item_warnings)

    if not cooperation_items and normalized.get("content_sections") not in (None, "", []):
        cooperation_items, legacy_warnings = normalize_legacy_content_sections(normalized.get("content_sections"))
        normalized["cooperation_items"] = cooperation_items
        schema_warnings.extend(legacy_warnings)
    else:
        normalized["cooperation_items"] = cooperation_items

    global_replacements, replacement_warnings = normalize_global_replacements(normalized.get("global_replacements"))
    normalized["global_replacements"] = global_replacements
    schema_warnings.extend(replacement_warnings)

    paragraph_overrides, override_warnings = normalize_paragraph_overrides(normalized.get("paragraph_overrides"))
    normalized["paragraph_overrides"] = paragraph_overrides
    schema_warnings.extend(override_warnings)

    if not normalized.get("cooperation_items"):
        raise schema_error("cooperation_items is required and each item must use heading/body.")

    normalized["_schema_warnings"] = schema_warnings
    return normalized


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def normalize_party_name(value: str) -> str:
    return re.sub(r"\s+", "", (value or "")).lower()


def is_kepler_party(value: str) -> bool:
    normalized = normalize_party_name(value)
    return normalized in {normalize_party_name(item) for item in KEPLER_NAME_ALIASES}


def apply_party_defaults(job: dict) -> dict:
    normalized = copy.deepcopy(job)
    party_a = str(normalized.get("party_a", "")).strip()
    party_b = str(normalized.get("party_b", "")).strip()

    if party_a and not party_b:
        normalized["party_b"] = KEPLER_LEGAL_NAME
    elif party_b and not party_a:
        normalized["party_a"] = KEPLER_LEGAL_NAME

    return normalized


def resolve_input_path(raw_path: str | None, *, base_dirs: list[Path]) -> Path | None:
    if not raw_path:
        return None

    candidate = Path(str(raw_path).strip())
    if not str(candidate):
        return None

    if candidate.is_absolute():
        return candidate

    for base_dir in base_dirs:
        resolved = (base_dir / candidate).resolve()
        if resolved.exists():
            return resolved

    return (base_dirs[0] / candidate).resolve() if base_dirs else candidate.resolve()


def slug_filename(text: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "_", text).strip()
    cleaned = re.sub(r"\s+", "", cleaned)
    return cleaned or "合同草稿"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def strip_heading_prefix(text: str) -> str:
    return HEADING_PREFIX_PATTERN.sub("", text).strip()


def clear_paragraph(paragraph) -> None:
    set_paragraph_text(paragraph, "")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(anchor_paragraph, template_paragraph=None) -> Paragraph:
    reference = template_paragraph or anchor_paragraph
    new_p = copy.deepcopy(reference._p)
    anchor_paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, anchor_paragraph._parent)
    set_paragraph_text(new_paragraph, "")
    return new_paragraph


def paragraph_text(paragraph) -> str:
    return (paragraph.text or "").strip()


def find_first_index(paragraphs, predicate, start: int = 0) -> int | None:
    for idx in range(start, len(paragraphs)):
        if predicate(paragraphs[idx], idx):
            return idx
    return None


def find_text_index(doc: Document, text: str) -> int | None:
    return find_first_index(doc.paragraphs, lambda p, i: paragraph_text(p) == text)


def find_content_indices_between(doc: Document, start_heading: str, end_heading: str) -> list[int]:
    start = find_text_index(doc, start_heading)
    end = find_text_index(doc, end_heading)
    if start is None or end is None or end <= start:
        return []
    return [
        idx
        for idx in range(start + 1, end)
        if paragraph_text(doc.paragraphs[idx])
    ]


def replace_text_everywhere(doc: Document, replacements: Iterable[tuple[str, str]]) -> list[str]:
    hits: list[str] = []
    for old, new in replacements:
        if not old or old == new:
            continue
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if old in text:
                set_paragraph_text(paragraph, text.replace(old, new))
                hits.append(f"paragraph[{idx}] {old} -> {new}")
    return hits


def collect_nonempty_indices(doc: Document) -> list[int]:
    return [idx for idx, p in enumerate(doc.paragraphs) if paragraph_text(p)]


def find_body_party_indices(doc: Document) -> tuple[int | None, int | None]:
    paragraphs = doc.paragraphs
    body_a = find_first_index(
        paragraphs,
        lambda p, i: paragraph_text(p).startswith("甲方：") and i > 10,
    )
    body_b = None
    if body_a is not None:
        body_b = find_first_index(
            paragraphs,
            lambda p, i: paragraph_text(p).startswith("乙方：") and i > body_a,
            start=body_a + 1,
        )
    return body_a, body_b


def first_cover_party_index(doc: Document, label: str) -> int | None:
    return find_first_index(
        doc.paragraphs,
        lambda p, i: paragraph_text(p).startswith(f"{label}：") and i <= 10,
    )


def find_identity_line(doc: Document, start: int | None, end: int | None) -> int | None:
    if start is None:
        return None
    final_end = end if end is not None else len(doc.paragraphs)
    for idx in range(start, final_end):
        text = paragraph_text(doc.paragraphs[idx])
        if "身份证" in text and "乙方" not in text:
            return idx
    return None


def find_credit_code_line(doc: Document, start: int | None) -> int | None:
    if start is None:
        return None
    return find_first_index(
        doc.paragraphs,
        lambda p, i: "统一社会信用代码" in paragraph_text(p) and i >= start,
        start=start,
    )


def find_title_index(doc: Document) -> int | None:
    nonempty = collect_nonempty_indices(doc)
    if not nonempty:
        return None
    first_idx = nonempty[0]
    first_text = paragraph_text(doc.paragraphs[first_idx])
    if "协议" in first_text or "合同" in first_text:
        return first_idx
    return find_first_index(
        doc.paragraphs,
        lambda p, i: ("协议" in paragraph_text(p) or "合同" in paragraph_text(p)) and i < 15,
    )


def find_section_range(doc: Document, heading_keyword: str) -> tuple[int | None, int | None]:
    start = find_first_index(
        doc.paragraphs,
        lambda p, i: paragraph_text(p) == heading_keyword or heading_keyword in paragraph_text(p),
    )
    if start is None:
        return None, None

    end = None
    for idx in range(start + 1, len(doc.paragraphs)):
        text = paragraph_text(doc.paragraphs[idx])
        if not text:
            continue
        if any(keyword in text for keyword in SECTION_END_KEYWORDS):
            end = idx
            break
    return start, end


def apply_cooperation_items(doc: Document, items: list[dict]) -> dict:
    section_start, section_end = find_section_range(doc, "合作内容")
    if section_start is None or section_end is None:
        raise RuntimeError("Unable to locate the cooperation section in the template.")

    target_indices = list(range(section_start + 1, section_end))
    target_paragraphs = [doc.paragraphs[idx] for idx in target_indices]
    heading_templates = target_paragraphs[::2]
    body_templates = target_paragraphs[1::2]
    original_headings = [paragraph_text(paragraph) for paragraph in heading_templates]
    template_uses_heading_prefix = any(
        HEADING_PREFIX_PATTERN.match(text) for text in original_headings if text
    )
    normalized_headings: list[str] = []
    written_indices: list[int] = []
    inserted_pairs = 0
    last_anchor = target_paragraphs[-1] if target_paragraphs else doc.paragraphs[section_start]
    fallback_heading_template = heading_templates[-1] if heading_templates else last_anchor
    fallback_body_template = body_templates[-1] if body_templates else fallback_heading_template

    for item in items:
        heading = str(item.get("heading", "")).strip()
        body = str(item.get("body", "")).strip()
        if not heading or not body:
            raise RuntimeError("Each cooperation item must include both heading and body.")
        if not template_uses_heading_prefix:
            heading = strip_heading_prefix(heading)
        normalized_headings.append(heading)

    for item_index, heading in enumerate(normalized_headings):
        body = str(items[item_index].get("body", "")).strip()
        heading_slot = item_index * 2
        body_slot = heading_slot + 1

        if body_slot < len(target_paragraphs):
            heading_paragraph = target_paragraphs[heading_slot]
            body_paragraph = target_paragraphs[body_slot]
        else:
            heading_paragraph = insert_paragraph_after(last_anchor, fallback_heading_template)
            body_paragraph = insert_paragraph_after(heading_paragraph, fallback_body_template)
            inserted_pairs += 1

        set_paragraph_text(heading_paragraph, heading)
        set_paragraph_text(body_paragraph, body)
        last_anchor = body_paragraph

    return {
        "section_heading_index": section_start,
        "section_end_index": section_end,
        "template_uses_heading_prefix": template_uses_heading_prefix,
        "normalized_headings": normalized_headings,
        "written_pair_count": len(items),
        "inserted_pair_count": inserted_pairs,
        "unused_indices": target_indices[len(items) * 2:],
    }


def prune_unused_cooperation_paragraphs(doc: Document, indices: list[int]) -> list[int]:
    deleted: list[int] = []
    for idx in sorted({int(item) for item in indices}, reverse=True):
        if 0 <= idx < len(doc.paragraphs):
            delete_paragraph(doc.paragraphs[idx])
            deleted.append(idx)
    return sorted(deleted)


def apply_paragraph_overrides(doc: Document, overrides: list[dict]) -> list[int]:
    written_indices: list[int] = []
    for item in overrides:
        if "index" not in item:
            raise RuntimeError("Each paragraph override must include an index.")
        idx = int(item["index"])
        if idx < 0 or idx >= len(doc.paragraphs):
            raise RuntimeError(f"Paragraph override index out of range: {idx}")
        text = str(item.get("text", "")).strip()
        set_paragraph_text(doc.paragraphs[idx], text)
        written_indices.append(idx)
    return written_indices


def validate_output(
    generated_doc: Document,
    template_path: Path,
    output_path: Path,
    expected_title: str,
    party_a: str,
    party_b: str,
    old_terms: list[str],
) -> dict:
    paragraphs = generated_doc.paragraphs
    signature_present = any("盖章" in paragraph_text(p) for p in paragraphs)
    title_present = any(expected_title in paragraph_text(p) for p in paragraphs)
    party_a_present = any(party_a in paragraph_text(p) for p in paragraphs)
    party_b_present = any(party_b in paragraph_text(p) for p in paragraphs)
    old_hits = []
    for term in old_terms:
        if not term:
            continue
        for idx, paragraph in enumerate(paragraphs):
            if term in paragraph.text:
                old_hits.append({"term": term, "paragraph_index": idx, "text": paragraph.text})

    return {
        "output_exists": output_path.exists(),
        "source_not_overwritten": str(template_path.resolve()) != str(output_path.resolve()),
        "signature_page_marker_present": signature_present,
        "title_present": title_present,
        "party_a_present": party_a_present,
        "party_b_present": party_b_present,
        "old_term_hits": old_hits,
    }


def write_markdown_note(path: Path, job: dict, validation: dict) -> None:
    unresolved = job.get("unresolved_items") or []
    leave_blank = job.get("leave_blank_fields") or []
    schema_warnings = job.get("_schema_warnings") or []
    lines = [
        "# 合同草稿未决事项",
        "",
        f"- 文档标题：{job.get('_effective_contract_title') or job.get('contract_title', '')}",
        f"- 甲方：{job.get('party_a', '')}",
        f"- 乙方：{job.get('party_b', '')}",
        f"- 输出文件：{job.get('_output_path', '')}",
        "",
        "## 默认留空字段",
    ]

    if leave_blank:
        lines.extend([f"- {item}" for item in leave_blank])
    else:
        lines.append("- 无")

    lines.extend(["", "## 仍需人工确认"])
    if unresolved:
        lines.extend([f"- {item}" for item in unresolved])
    else:
        lines.append("- 无额外未决事项")

    lines.extend(["", "## 输入归一化提示"])
    if schema_warnings:
        lines.extend([f"- {item}" for item in schema_warnings])
    else:
        lines.append("- 无")

    lines.extend(
        [
            "",
            "## 自动校验摘要",
            f"- 输出文件存在：{'是' if validation['output_exists'] else '否'}",
            f"- 未覆盖原模板：{'是' if validation['source_not_overwritten'] else '否'}",
            f"- 标题已写入：{'是' if validation['title_present'] else '否'}",
            f"- 甲方已写入：{'是' if validation['party_a_present'] else '否'}",
            f"- 乙方已写入：{'是' if validation['party_b_present'] else '否'}",
            f"- 签署页标记仍在：{'是' if validation['signature_page_marker_present'] else '否'}",
            f"- 旧术语残留数：{len(validation['old_term_hits'])}",
            "",
            "> 本文档为业务草稿，正式签署前仍需法务或律师审核。",
        ]
    )

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def save_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def build_output_filename(job: dict) -> str:
    explicit = str(job.get("output_filename", "")).strip()
    if explicit and "?" not in explicit:
        return explicit if explicit.lower().endswith(".docx") else f"{explicit}.docx"
    title = str(
        job.get("_effective_contract_title")
        or job.get("contract_title")
        or "战略合作协议"
    ).strip()
    return f"{slug_filename(title)}_草稿.docx"


def uniquify_output_path(path: Path) -> Path:
    if not path.exists():
        return path

    stem = path.stem
    suffix = path.suffix or ".docx"
    version_match = re.match(r"^(.*)_v(\d+)$", stem)
    if version_match:
        base_stem = version_match.group(1)
        start_version = int(version_match.group(2)) + 1
    else:
        base_stem = stem
        start_version = 2

    candidate = path.with_name(f"{base_stem}_v{start_version}{suffix}")
    version = start_version
    while candidate.exists():
        version += 1
        candidate = path.with_name(f"{base_stem}_v{version}{suffix}")
    return candidate


def resolve_template_path(skill_root: Path, spec_path: Path, cli_template: str | None, job: dict) -> Path:
    def fallback_template() -> Path:
        candidates = sorted(
            path for path in glob.glob(str(skill_root / "*.docx"))
            if not Path(path).name.startswith("~$")
        )
        if not candidates:
            raise RuntimeError("No template .docx found in the skill root.")
        return Path(candidates[0])

    raw_candidate = cli_template or job.get("template_path")
    if raw_candidate:
        candidate = resolve_input_path(
            str(raw_candidate),
            base_dirs=[spec_path.parent, skill_root, Path.cwd()],
        )
        if candidate.exists():
            return candidate

        # Windows shell and model-generated specs can occasionally mangle Chinese
        # characters into question marks. Fall back to the local skill template
        # instead of failing hard on an invalid path.
        if "?" in str(candidate) or not candidate.suffix.lower() == ".docx":
            return fallback_template()

    return fallback_template()


def build_global_replacements(job: dict, doc: Document) -> list[tuple[str, str]]:
    replacements: list[tuple[str, str]] = []

    for item in job.get("global_replacements", []):
        old = str(item.get("old", "")).strip()
        new = str(item.get("new", "")).strip()
        if old and new:
            replacements.append((old, new))

    cover_a = first_cover_party_index(doc, "甲方")
    body_a, body_b = find_body_party_indices(doc)

    auto_old_terms = []
    for idx in [cover_a, body_a]:
        if idx is None:
            continue
        text = paragraph_text(doc.paragraphs[idx])
        if "：" in text:
            value = text.split("：", 1)[1].strip()
            if value:
                auto_old_terms.append(value)

    for old in job.get("party_a_aliases_to_replace", []) + auto_old_terms:
        old = str(old).strip()
        new_value = str(job.get("party_a", "")).strip()
        if old and old != new_value:
            replacements.append((old, new_value))

    auto_old_b_terms = []
    for idx in [body_b]:
        if idx is None:
            continue
        text = paragraph_text(doc.paragraphs[idx])
        if "：" in text:
            value = text.split("：", 1)[1].strip()
            if value:
                auto_old_b_terms.append(value)

    for old in job.get("party_b_aliases_to_replace", []) + auto_old_b_terms:
        old = str(old).strip()
        new_value = str(job.get("party_b", "")).strip()
        if old and old != new_value:
            replacements.append((old, new_value))

    deduped = []
    seen = set()
    for pair in replacements:
        if pair not in seen and pair[0] and pair[1]:
            deduped.append(pair)
            seen.add(pair)
    return deduped


def summarize_cooperation_topics(job: dict) -> str:
    headings = [
        strip_heading_prefix(str(item.get("heading", "")).strip())
        for item in job.get("cooperation_items", [])
        if str(item.get("heading", "")).strip()
    ]
    headings = [heading for heading in headings if heading]
    if not headings:
        return "人工智能相关合作事项"
    if len(headings) == 1:
        return headings[0]
    if len(headings) == 2:
        return "、".join(headings)
    return "、".join(headings[:3]) + "等事项"


def normalize_company_profile(value: str) -> str:
    text = re.sub(r"\s+", "", (value or "")).strip("，。；; ")
    if not text:
        return ""
    return text + "。"


def get_company_profile(job: dict, side: str) -> str:
    aliases = [
        f"{side}_profile",
        f"{side}_intro",
        f"{side}_company_profile",
        f"{side}_company_intro",
    ]
    for key in aliases:
        value = normalize_company_profile(str(job.get(key, "")).strip())
        if value:
            return value
    return ""


def build_party_role_texts(job: dict, party_a: str, party_b: str, topic_summary: str) -> dict[str, list[str]]:
    kepler_as_a = is_kepler_party(party_a)
    kepler_as_b = is_kepler_party(party_b)
    party_a_profile = get_company_profile(job, "party_a")
    party_b_profile = get_company_profile(job, "party_b")

    if kepler_as_a and not kepler_as_b:
        recital = [
            party_a_profile
            or f"甲方系具备人工智能产品研发、解决方案交付、项目实施与运营服务能力的合作主体，能够围绕{topic_summary}提供相应的技术、产品、实施和服务支持。",
            party_b_profile
            or f"乙方系围绕{topic_summary}等方向开展资源组织、场景协同、项目推进或产业研究的合作主体，能够为本协议项下合作提供业务需求牵引、场景协调、资源整合及组织支持。",
            f"双方认可围绕{topic_summary}开展合作具备良好的业务基础与发展前景，愿在平等互利、优势互补的基础上建立合作关系，并就相关合作事项达成本协议。",
        ]
        party_a_texts = [
            f"甲方负责围绕{topic_summary}提供相应的人工智能产品能力、技术方案、实施服务、培训支持及项目交付保障。",
            "甲方应根据合作项目需要，投入必要的技术、产品、运营、培训或交付资源，确保合作事项按计划推进并形成可交付成果。",
            "甲方应配合乙方开展项目论证、方案优化、成果转化、对外汇报及相关宣传推广，并在合理范围内提供专业支持。",
            "甲方有权依约参与合作项目的实施安排、成果展示、市场推广及后续商业化协同，并获得相应的合作权益。",
        ]
        party_b_texts = [
            "乙方负责结合自身业务目标、客户资源、场景需求或组织协调能力，为合作项目提供需求牵引、资源协同、场景对接及必要的组织保障。",
            "乙方应协助推动合作事项中的内部协调、合作对接、项目推进及相关外部资源沟通，并在合理范围内支持项目实施落地。",
            "乙方应保证其向合作项目提供的资料、授权、业务信息或协同资源真实、合法、有效，不侵犯任何第三方合法权益。",
            "乙方有权参与合作项目的方向把控、重要方案评审、执行过程监督及重大事项决策。",
        ]
    else:
        recital = [
            party_a_profile
            or f"甲方系围绕{topic_summary}等方向开展资源组织、场景协同、项目推进或产业研究的合作主体，能够为本协议项下合作提供业务需求牵引、场景协调、资源整合及组织支持。",
            party_b_profile
            or f"乙方系具备人工智能产品研发、解决方案交付、项目实施与运营服务能力的合作主体，能够围绕本协议项下合作内容提供相应的技术、产品、实施和服务支持。",
            f"双方认可围绕{topic_summary}开展合作具备良好的业务基础与发展前景，愿在平等互利、优势互补的基础上建立合作关系，并就相关合作事项达成本协议。",
        ]
        party_a_texts = [
            "甲方负责结合自身业务目标、客户资源、场景需求或组织协调能力，为合作项目提供需求牵引、资源协同、场景对接及必要的组织保障。",
            "甲方应协助推动合作事项中的内部协调、合作对接、项目推进及相关外部资源沟通，并在合理范围内支持项目实施落地。",
            "甲方应保证其向合作项目提供的资料、授权、业务信息或协同资源真实、合法、有效，不侵犯任何第三方合法权益。",
            "甲方有权参与合作项目的方向把控、重要方案评审、执行过程监督及重大事项决策。",
        ]
        party_b_texts = [
            f"乙方负责围绕{topic_summary}提供相应的人工智能产品能力、技术方案、实施服务、培训支持及项目交付保障。",
            "乙方应根据合作项目需要，投入必要的技术、产品、运营、培训或交付资源，确保合作事项按计划推进并形成可交付成果。",
            "乙方应配合甲方开展项目论证、方案优化、成果转化、对外汇报及相关宣传推广，并在合理范围内提供专业支持。",
            "乙方有权依约参与合作项目的实施安排、成果展示、市场推广及后续商业化协同，并获得相应的合作权益。",
        ]

    return {
        "recital": recital,
        "party_a_texts": party_a_texts,
        "party_b_texts": party_b_texts,
    }


def build_auto_paragraph_overrides(doc: Document, job: dict) -> list[dict]:
    if not job.get("cooperation_items"):
        return []

    party_a = str(job.get("party_a", "")).strip()
    party_b = str(job.get("party_b", "")).strip()
    topic_summary = summarize_cooperation_topics(job)

    overrides: list[dict] = []

    role_texts = build_party_role_texts(job, party_a, party_b, topic_summary)

    recital_indices = find_content_indices_between(doc, "鉴于：", "合作原则")
    recital_texts = role_texts["recital"]
    for idx, text in zip(recital_indices, recital_texts):
        overrides.append({"index": idx, "text": text})

    principle_indices = find_content_indices_between(doc, "合作原则", "合作内容")
    principle_text = (
        f"双方本着“优势互补、互利共赢、长期合作、共同发展”的原则，建立战略合作伙伴关系。"
        f"双方将围绕{topic_summary}协同推进业务规划、方案落地、项目实施、培训孵化与成果转化，"
        "在保持模板结构与条款体系稳定的前提下，确保合作方向与全文表述一致。"
    )
    for idx in principle_indices[:1]:
        overrides.append({"index": idx, "text": principle_text})

    party_a_indices = find_content_indices_between(doc, "甲方权利与义务", "乙方权利与义务")
    party_a_texts = role_texts["party_a_texts"]
    for idx, text in zip(party_a_indices, party_a_texts):
        overrides.append({"index": idx, "text": text})

    party_b_indices = find_content_indices_between(doc, "乙方权利与义务", "合作机制")
    party_b_texts = role_texts["party_b_texts"]
    for idx, text in zip(party_b_indices, party_b_texts):
        overrides.append({"index": idx, "text": text})

    mechanism_indices = find_content_indices_between(doc, "合作机制", "合作期限")
    mechanism_texts = [
        "合作双方共同设立联合工作机制，负责统筹规划和实施合作内容中的各项工作，并定期沟通项目进展、资源投入、阶段成果及后续安排。",
        "双方确认，本协议由甲乙双方作为独立民事主体签署并履行；如后续因内部组织、授权安排或承继事项需要调整履约主体，应由相关方另行以书面文件明确，不影响本协议已生效条款的连续执行。",
    ]
    for idx, text in zip(mechanism_indices, mechanism_texts):
        overrides.append({"index": idx, "text": text})

    succession_indices = [
        idx for idx, p in enumerate(doc.paragraphs) if "承继本协议" in paragraph_text(p)
    ]
    succession_texts = [
        "双方确认，本协议由甲乙双方作为独立民事主体签署并履行；如后续因内部组织、授权安排或承继事项需要调整履约主体，应由相关方另行以书面文件明确，不影响本协议已生效条款的连续执行。",
        "如合作过程中涉及主体变更、承继、续签或补充安排，双方应另行签署书面文件予以确认；在书面文件生效前，原签约主体仍应按照本协议约定继续履行相应义务。",
    ]
    for idx, text in zip(succession_indices, succession_texts):
        overrides.append({"index": idx, "text": text})

    ip_indices = find_content_indices_between(doc, "知识产权", "保密条款")
    ip_texts = [
        f"甲方在本协议签署前已经合法拥有的商标、著作权、专利、技术资料、数据资源、业务方案及其他知识产权，仍归甲方或其相关权利人所有。",
        f"乙方在本协议签署前已经合法拥有的平台、软件、模型能力、解决方案、课程体系、数据资产及其他知识产权，仍归乙方或其相关权利人所有。",
    ]
    for idx, text in zip(ip_indices[:2], ip_texts):
        overrides.append({"index": idx, "text": text})

    return overrides


def merge_paragraph_overrides(auto_overrides: list[dict], explicit_overrides: list[dict]) -> list[dict]:
    merged: dict[int, dict] = {}
    for item in auto_overrides + explicit_overrides:
        if "index" not in item:
            continue
        merged[int(item["index"])] = {"index": int(item["index"]), "text": str(item.get("text", "")).strip()}
    return [merged[idx] for idx in sorted(merged)]


@dataclass
class RunResult:
    output_path: Path
    note_path: Path
    summary_path: Path
    validation_path: Path


def run(job: dict, template_path: Path, output_dir: Path) -> RunResult:
    ensure_dir(output_dir)

    source_bytes = template_path.read_bytes()
    doc = Document(str(template_path))

    normalized_job = apply_party_defaults(normalize_job_schema(job))
    party_a = str(normalized_job.get("party_a", "")).strip()
    party_b = str(normalized_job.get("party_b", "")).strip()
    if not party_a or not party_b:
        raise RuntimeError("party_a and party_b are required in the job spec.")

    title_idx = find_title_index(doc)
    template_title = paragraph_text(doc.paragraphs[title_idx]) if title_idx is not None else "战略合作协议"
    requested_title = str(normalized_job.get("contract_title", "")).strip()
    preserve_template_title = str(normalized_job.get("preserve_template_title", "true")).strip().lower() not in {
        "false",
        "0",
        "no",
    }
    effective_title = template_title if preserve_template_title or not requested_title else requested_title
    normalized_job["_effective_contract_title"] = effective_title

    output_path = uniquify_output_path(output_dir / build_output_filename(normalized_job))
    note_path = output_dir / "未决事项清单.md"
    summary_path = output_dir / "run-summary.json"
    validation_path = output_dir / "validation-report.json"

    if title_idx is not None and not preserve_template_title and requested_title:
        set_paragraph_text(doc.paragraphs[title_idx], requested_title)

    cover_a = first_cover_party_index(doc, "甲方")
    cover_b = first_cover_party_index(doc, "乙方")
    if cover_a is not None:
        set_paragraph_text(doc.paragraphs[cover_a], f"甲方：{party_a}")
    if cover_b is not None:
        set_paragraph_text(doc.paragraphs[cover_b], f"乙方：{party_b}")

    body_a, body_b = find_body_party_indices(doc)
    if body_a is not None:
        set_paragraph_text(doc.paragraphs[body_a], f"甲方：{party_a}")
    if body_b is not None:
        set_paragraph_text(doc.paragraphs[body_b], f"乙方：{party_b}")

    identity_line = find_identity_line(doc, body_a, body_b)
    identity_text = str(
        normalized_job.get("party_a_identity_line", "地址：【待确认】    联系人：【待确认】")
    ).strip()
    if identity_line is not None and identity_text:
        set_paragraph_text(doc.paragraphs[identity_line], identity_text)

    credit_line = find_credit_code_line(doc, body_b)
    if credit_line is not None and normalized_job.get("party_b_credit_code"):
        set_paragraph_text(doc.paragraphs[credit_line], f"统一社会信用代码：{normalized_job['party_b_credit_code']}")

    replacements = build_global_replacements(normalized_job, doc)
    replacement_hits = replace_text_everywhere(doc, replacements)

    cooperation_result = {}
    if normalized_job.get("cooperation_items"):
        cooperation_result = apply_cooperation_items(doc, copy.deepcopy(normalized_job["cooperation_items"]))

    auto_overrides = build_auto_paragraph_overrides(doc, normalized_job)
    explicit_overrides = copy.deepcopy(normalized_job.get("paragraph_overrides", []))
    merged_overrides = merge_paragraph_overrides(auto_overrides, explicit_overrides)
    paragraph_override_indices: list[int] = []
    if merged_overrides:
        paragraph_override_indices = apply_paragraph_overrides(doc, merged_overrides)

    deleted_cooperation_indices = []
    if cooperation_result.get("unused_indices"):
        deleted_cooperation_indices = prune_unused_cooperation_paragraphs(
            doc,
            cooperation_result["unused_indices"],
        )

    doc.save(str(output_path))
    generated = Document(str(output_path))

    old_terms = [old for old, _ in replacements]
    validation = validate_output(
        generated_doc=generated,
        template_path=template_path,
        output_path=output_path,
        expected_title=effective_title,
        party_a=party_a,
        party_b=party_b,
        old_terms=old_terms,
    )

    if template_path.read_bytes() != source_bytes:
        raise RuntimeError("The source template changed during execution, which should not happen.")

    summary = {
        "template_path": str(template_path),
        "output_path": str(output_path),
        "effective_contract_title": effective_title,
        "preserve_template_title": preserve_template_title,
        "party_profiles": {
            "party_a": get_company_profile(normalized_job, "party_a"),
            "party_b": get_company_profile(normalized_job, "party_b"),
        },
        "title_index": title_idx,
        "cover_party_indices": {"party_a": cover_a, "party_b": cover_b},
        "body_party_indices": {"party_a": body_a, "party_b": body_b},
        "identity_line_index": identity_line,
        "credit_line_index": credit_line,
        "global_replacements": [{"old": old, "new": new} for old, new in replacements],
        "replacement_hits": replacement_hits,
        "cooperation_result": cooperation_result,
        "deleted_cooperation_indices": deleted_cooperation_indices,
        "auto_paragraph_override_indices": sorted({int(item["index"]) for item in auto_overrides}),
        "explicit_paragraph_override_indices": [int(item["index"]) for item in explicit_overrides if "index" in item],
        "paragraph_override_indices": paragraph_override_indices,
    }

    normalized_job["_output_path"] = str(output_path)
    write_markdown_note(note_path, normalized_job, validation)
    save_json(summary_path, summary)
    save_json(validation_path, validation)

    return RunResult(
        output_path=output_path,
        note_path=note_path,
        summary_path=summary_path,
        validation_path=validation_path,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a draft contract from a .docx template.")
    parser.add_argument("--spec", required=True, help="Path to the UTF-8 JSON job spec.")
    parser.add_argument("--template", help="Optional template path. Overrides template_path in the spec.")
    parser.add_argument("--output-dir", help="Optional output directory override.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    script_path = Path(__file__).resolve()
    skill_root = script_path.parents[1]
    spec_path = Path(args.spec).resolve()
    job = normalize_job_schema(load_json(spec_path))
    template_path = resolve_template_path(skill_root, spec_path, args.template, job)
    raw_output_dir = args.output_dir or job.get("output_dir")
    output_dir = (
        resolve_input_path(str(raw_output_dir), base_dirs=[spec_path.parent, Path.cwd()])
        if raw_output_dir
        else default_output_dir(job, spec_path)
    )

    result = run(job, template_path, output_dir)
    print(f"Draft written to: {result.output_path}")
    print(f"Missing-items note: {result.note_path}")
    print(f"Run summary: {result.summary_path}")
    print(f"Validation report: {result.validation_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

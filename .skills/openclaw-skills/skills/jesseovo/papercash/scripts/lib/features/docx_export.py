"""Word 文档导出"""

from __future__ import annotations
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from schema import SearchResult, Paper


def export_review_to_docx(review_text: str, output_path: str,
                          title: str = "文献综述") -> str:
    """将文献综述导出为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "[错误] 需要安装 python-docx: pip install python-docx"

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(16)

    lines = review_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            continue
        elif line.startswith("## "):
            heading_text = line[3:].strip()
            h = doc.add_heading(heading_text, level=1)
            for run in h.runs:
                run.font.size = Pt(15)
        elif line.startswith("### "):
            heading_text = line[4:].strip()
            h = doc.add_heading(heading_text, level=2)
            for run in h.runs:
                run.font.size = Pt(14)
        elif line.startswith("- **") and "**" in line[4:]:
            p = doc.add_paragraph(style="List Bullet")
            bold_end = line.index("**", 4)
            bold_text = line[4:bold_end]
            rest = line[bold_end + 2:].lstrip(": ：")
            run_bold = p.add_run(bold_text)
            run_bold.bold = True
            run_bold.font.size = Pt(12)
            if rest:
                run_rest = p.add_run(f": {rest}")
                run_rest.font.size = Pt(12)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(12)
        elif line.startswith("[") and "]" in line:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10.5)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
            run.font.size = Pt(12)
        else:
            p = doc.add_paragraph(line)
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.name = "宋体"

    doc.save(output_path)
    return f"已导出到: {output_path}"


def export_search_to_docx(result: SearchResult, output_path: str) -> str:
    """将搜索结果导出为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "[错误] 需要安装 python-docx: pip install python-docx"

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    title = doc.add_heading(f"论文检索报告: {result.query}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"查询: {result.query}\n").bold = True
    info.add_run(f"结果: {result.total_found} 篇\n")
    info.add_run(f"数据源: {', '.join(result.sources_used)}\n")

    doc.add_heading("检索结果", level=1)

    for i, paper in enumerate(result.papers, 1):
        doc.add_heading(f"{i}. {paper.title}", level=2)

        meta = doc.add_paragraph()
        meta.add_run(f"作者: ").bold = True
        meta.add_run(f"{paper.authors_str}\n")

        if paper.year:
            meta.add_run(f"年份: ").bold = True
            meta.add_run(f"{paper.year}\n")

        if paper.venue:
            meta.add_run(f"期刊: ").bold = True
            meta.add_run(f"{paper.venue}\n")

        if paper.citation_count:
            meta.add_run(f"引用: ").bold = True
            meta.add_run(f"{paper.citation_count} 次\n")

        if paper.doi:
            meta.add_run(f"DOI: ").bold = True
            meta.add_run(f"{paper.doi}\n")

        if paper.abstract:
            abs_para = doc.add_paragraph()
            abs_para.add_run("摘要: ").bold = True
            abstract_short = paper.abstract[:500]
            abs_para.add_run(abstract_short)

        for run in meta.runs:
            run.font.size = Pt(12)

    doc.save(output_path)
    return f"已导出到: {output_path}"

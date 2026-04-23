"""格式检查 - 检查 Word 文档是否符合高校论文格式要求"""

from __future__ import annotations
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from schema import FormatIssue

STANDARD_FORMAT = {
    "title_font": "黑体",
    "title_size_pt": 16,
    "heading1_font": "黑体",
    "heading1_size_pt": 15,
    "heading2_font": "黑体",
    "heading2_size_pt": 14,
    "body_font": "宋体",
    "body_font_en": "Times New Roman",
    "body_size_pt": 12,
    "line_spacing": 1.5,
    "margin_top_cm": 2.54,
    "margin_bottom_cm": 2.54,
    "margin_left_cm": 3.17,
    "margin_right_cm": 3.17,
    "page_size": "A4",
    "abstract_title": "摘要",
    "keywords_title": "关键词",
    "references_title": "参考文献",
}


def check_format(file_path: str) -> str:
    """检查 Word 文档格式"""
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    if not file_path.lower().endswith((".docx", ".doc")):
        return "[错误] 仅支持 .docx 格式文件"

    try:
        from docx import Document
    except ImportError:
        return _format_guide_only()

    try:
        doc = Document(file_path)
    except Exception as e:
        return f"[错误] 无法打开文件: {e}"

    issues: list[FormatIssue] = []

    issues.extend(_check_margins(doc))
    issues.extend(_check_fonts_and_sizes(doc))
    issues.extend(_check_line_spacing(doc))
    issues.extend(_check_structure(doc))

    return _render_report(issues, file_path)


def _check_margins(doc) -> list[FormatIssue]:
    issues: list[FormatIssue] = []
    for i, section in enumerate(doc.sections):
        margin_checks = [
            ("上边距", section.top_margin, STANDARD_FORMAT["margin_top_cm"]),
            ("下边距", section.bottom_margin, STANDARD_FORMAT["margin_bottom_cm"]),
            ("左边距", section.left_margin, STANDARD_FORMAT["margin_left_cm"]),
            ("右边距", section.right_margin, STANDARD_FORMAT["margin_right_cm"]),
        ]
        for name, actual_emu, expected_cm in margin_checks:
            if actual_emu is None:
                continue
            actual_cm = round(actual_emu / 360000, 2)
            if abs(actual_cm - expected_cm) > 0.1:
                issues.append(FormatIssue(
                    location=f"节 {i + 1}",
                    issue_type=name,
                    expected=f"{expected_cm} cm",
                    actual=f"{actual_cm} cm",
                    severity="warning",
                ))
    return issues


def _check_fonts_and_sizes(doc) -> list[FormatIssue]:
    issues: list[FormatIssue] = []

    for i, para in enumerate(doc.paragraphs[:50]):
        if not para.text.strip():
            continue

        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "标题" in style_name

        for run in para.runs:
            if not run.text.strip():
                continue

            font_name = run.font.name
            font_size = run.font.size

            if font_size:
                size_pt = round(font_size / 12700, 1)
                if is_heading:
                    if size_pt < 14:
                        issues.append(FormatIssue(
                            location=f"段落 {i + 1} (标题)",
                            issue_type="字号",
                            expected="≥ 14pt (三号或小三号)",
                            actual=f"{size_pt}pt",
                            severity="warning",
                        ))
                else:
                    expected_size = STANDARD_FORMAT["body_size_pt"]
                    if abs(size_pt - expected_size) > 0.5:
                        issues.append(FormatIssue(
                            location=f"段落 {i + 1}",
                            issue_type="字号",
                            expected=f"{expected_size}pt (小四号)",
                            actual=f"{size_pt}pt",
                            severity="info",
                        ))

            if font_name and not is_heading:
                has_cn = any("\u4e00" <= c <= "\u9fff" for c in run.text)
                if has_cn and font_name not in ("宋体", "SimSun", "NSimSun"):
                    issues.append(FormatIssue(
                        location=f"段落 {i + 1}",
                        issue_type="中文字体",
                        expected="宋体",
                        actual=font_name,
                        severity="info",
                    ))

            if len(issues) > 20:
                break
        if len(issues) > 20:
            break

    return issues


def _check_line_spacing(doc) -> list[FormatIssue]:
    issues: list[FormatIssue] = []
    checked = 0

    for i, para in enumerate(doc.paragraphs[:30]):
        if not para.text.strip():
            continue

        pf = para.paragraph_format
        if pf.line_spacing is not None:
            spacing = pf.line_spacing
            if isinstance(spacing, (int, float)):
                if spacing < 1.0:
                    actual_val = round(spacing / 12700 / 12, 2)
                else:
                    actual_val = spacing

                expected = STANDARD_FORMAT["line_spacing"]
                if abs(actual_val - expected) > 0.1:
                    issues.append(FormatIssue(
                        location=f"段落 {i + 1}",
                        issue_type="行距",
                        expected=f"{expected} 倍行距",
                        actual=f"{actual_val}",
                        severity="warning",
                    ))
                    checked += 1

        if checked >= 5:
            break

    return issues


def _check_structure(doc) -> list[FormatIssue]:
    issues: list[FormatIssue] = []
    full_text = "\n".join(p.text for p in doc.paragraphs)

    required_sections = [
        ("摘要", ["摘要", "摘 要", "Abstract"]),
        ("关键词", ["关键词", "关键字", "Keywords"]),
        ("参考文献", ["参考文献", "参 考 文 献", "References"]),
    ]

    for name, keywords in required_sections:
        found = any(kw in full_text for kw in keywords)
        if not found:
            issues.append(FormatIssue(
                location="全文",
                issue_type="缺少必要章节",
                expected=name,
                actual="未找到",
                severity="error",
            ))

    return issues


def _render_report(issues: list[FormatIssue], file_path: str) -> str:
    lines: list[str] = []
    lines.append(f"\n{'=' * 60}")
    lines.append(f"  PaperCash 格式检查报告")
    lines.append(f"{'=' * 60}")
    lines.append(f"  文件: {file_path}")
    lines.append(f"  标准: 常见高校毕业论文格式要求")
    lines.append(f"{'=' * 60}\n")

    errors = [i for i in issues if i.severity == "error"]
    warnings = [i for i in issues if i.severity == "warning"]
    infos = [i for i in issues if i.severity == "info"]

    lines.append(f"  总计: {len(issues)} 个问题 "
                 f"(严重: {len(errors)} | 警告: {len(warnings)} | 提示: {len(infos)})\n")

    if not issues:
        lines.append("  ✅ 未发现格式问题，格式基本符合要求！\n")
        return "\n".join(lines)

    if errors:
        lines.append("  --- 严重问题 ---\n")
        for issue in errors:
            lines.append(f"  [ERROR] {issue.location}: {issue.issue_type}")
            lines.append(f"          期望: {issue.expected} | 实际: {issue.actual}\n")

    if warnings:
        lines.append("  --- 警告 ---\n")
        for issue in warnings:
            lines.append(f"  [WARN] {issue.location}: {issue.issue_type}")
            lines.append(f"         期望: {issue.expected} | 实际: {issue.actual}\n")

    if infos:
        lines.append("  --- 提示 ---\n")
        for issue in infos[:10]:
            lines.append(f"  [INFO] {issue.location}: {issue.issue_type}")
            lines.append(f"         期望: {issue.expected} | 实际: {issue.actual}\n")
        if len(infos) > 10:
            lines.append(f"  ... 及其他 {len(infos) - 10} 个提示\n")

    lines.append("  --- 格式参考标准 ---\n")
    lines.append(f"  正文字体: {STANDARD_FORMAT['body_font']} / {STANDARD_FORMAT['body_font_en']}")
    lines.append(f"  正文字号: {STANDARD_FORMAT['body_size_pt']}pt (小四号)")
    lines.append(f"  行距: {STANDARD_FORMAT['line_spacing']} 倍")
    lines.append(f"  页边距: 上下 {STANDARD_FORMAT['margin_top_cm']}cm, "
                 f"左 {STANDARD_FORMAT['margin_left_cm']}cm, 右 {STANDARD_FORMAT['margin_right_cm']}cm")
    lines.append(f"  纸张: {STANDARD_FORMAT['page_size']}")
    lines.append("")

    return "\n".join(lines)


def _format_guide_only() -> str:
    """无 python-docx 时返回格式指南"""
    return """
PaperCash 格式检查

⚠️  未安装 python-docx，无法检查 Word 文件。
    请运行: pip install python-docx

以下是常见高校毕业论文格式要求（供手动检查）：

--- 字体与字号 ---
  论文标题: 黑体, 16pt (二号)
  一级标题: 黑体, 15pt (小二号)
  二级标题: 黑体, 14pt (三号)
  正文中文: 宋体, 12pt (小四号)
  正文英文: Times New Roman, 12pt

--- 页面设置 ---
  纸张: A4
  上边距: 2.54 cm
  下边距: 2.54 cm
  左边距: 3.17 cm
  右边距: 3.17 cm

--- 行距 ---
  正文: 1.5 倍行距
  摘要/参考文献: 1.5 倍行距

--- 必需章节 ---
  1. 摘要（中文 + 英文）
  2. 关键词
  3. 目录
  4. 正文（绪论/文献综述/方法/结果/讨论/结论）
  5. 参考文献
  6. 致谢

--- 参考文献格式 ---
  推荐: GB/T 7714-2015
  使用 PaperCash: python scripts/papercash.py cite <DOI> --style gb7714
"""

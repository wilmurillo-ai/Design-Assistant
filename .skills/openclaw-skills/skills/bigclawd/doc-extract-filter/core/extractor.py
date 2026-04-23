# 文件提取器模块
import os
import csv
import logging
from typing import Dict, List, Optional, Union
from pathlib import Path
import PyPDF2
from docx import Document
from openpyxl import load_workbook

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class FileExtractor:
    """
    文件提取器类，支持多种文件类型的文本提取
    """
    
    def __init__(self):
        """
        初始化文件提取器
        """
        pass
    
    def extract_text(self, file_path: Union[str, Path], enable_ocr: bool = False) -> Dict[str, Union[str, List[str]]]:
        """
        提取文件文本内容
        
        Args:
            file_path: 文件路径
            enable_ocr: 是否启用 OCR 支持（用于扫描件 PDF）
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {
                    "file_path": str(file_path),
                    "error": "文件不存在"
                }
            
            # 智能格式检测
            file_type = self._detect_file_type(file_path)
            
            if file_type == "pdf":
                if enable_ocr:
                    return self._extract_pdf_with_ocr(file_path)
                else:
                    return self._extract_pdf(file_path)
            elif file_type == "docx":
                return self._extract_docx(file_path)
            elif file_type in ["xlsx", "xls"]:
                return self._extract_xlsx(file_path)
            elif file_type == "txt":
                return self._extract_txt(file_path)
            elif file_type == "csv":
                return self._extract_csv(file_path)
            elif file_type == "md":
                return self._extract_markdown(file_path)
            elif file_type == "wps":
                return self._extract_wps(file_path)
            elif file_type == "et":
                return self._extract_wps_et(file_path)
            else:
                return {
                    "file_path": str(file_path),
                    "error": f"不支持的文件类型: {file_path.suffix.lower()}"
                }
        except Exception as e:
            logger.error(f"提取文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def batch_extract(self, directory: Union[str, Path], extensions: Optional[List[str]] = None) -> List[Dict[str, Union[str, List[str]]]]:
        """
        批量提取目录中的文件
        
        Args:
            directory: 目录路径
            extensions: 要提取的文件扩展名列表，默认为 None（提取所有支持的类型）
            
        Returns:
            提取结果列表
        """
        try:
            directory = Path(directory)
            if not directory.exists() or not directory.is_dir():
                return [{"error": "目录不存在或不是目录"}]
            
            if extensions is None:
                extensions = ['.pdf', '.docx', '.xlsx', '.xls', '.txt', '.csv', '.md', '.wps', '.et']
            
            results = []
            for file_path in directory.iterdir():
                if file_path.is_file() and file_path.suffix.lower() in extensions:
                    result = self.extract_text(file_path)
                    results.append(result)
            
            return results
        except Exception as e:
            logger.error(f"批量提取时出错: {str(e)}")
            return [{"error": str(e)}]
    
    def _detect_file_type(self, file_path: Path) -> str:
        """
        智能检测文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件类型字符串
        """
        # 首先根据文件扩展名判断
        file_ext = file_path.suffix.lower()
        
        ext_map = {
            ".txt": "txt",
            ".pdf": "pdf",
            ".docx": "docx",
            ".xlsx": "xlsx",
            ".xls": "xls",
            ".csv": "csv",
            ".md": "md",
            ".markdown": "md",
            ".wps": "wps",
            ".et": "et"
        }
        
        if file_ext in ext_map:
            return ext_map[file_ext]
        
        # 尝试根据文件内容判断（简单实现）
        try:
            with open(file_path, 'rb') as f:
                header = f.read(100)
            
            # PDF 文件头
            if header.startswith(b'%PDF'):
                return "pdf"
            # ZIP 格式文件（docx, xlsx, wps, et）
            elif header.startswith(b'PK\x03\x04'):
                return "docx"  # 默认返回 docx，实际使用时会根据扩展名更准确判断
            # CSV 文件（尝试读取）
            elif b',' in header or b';' in header:
                return "csv"
        except:
            pass
        
        return "unsupported"  # 默认返回 unsupported
    
    def _extract_pdf(self, file_path: Path) -> Dict[str, str]:
        """
        提取 PDF 文件文本
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ''
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + '\n'
                
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except Exception as e:
            logger.error(f"提取 PDF 文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def _extract_pdf_with_ocr(self, file_path: Path) -> Dict[str, str]:
        """
        使用 OCR 从 PDF 扫描件中提取文本
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            import pytesseract
            from PIL import Image
            import pdf2image
            
            # 将 PDF 转换为图像
            images = pdf2image.convert_from_path(file_path)
            text = ''
            
            # 对每个图像进行 OCR
            for image in images:
                text += pytesseract.image_to_string(image, lang='chi_sim+eng') + '\n'
            
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except ImportError:
            logger.warning("OCR 依赖库未安装，回退到普通 PDF 提取")
            return self._extract_pdf(file_path)
        except Exception as e:
            logger.error(f"OCR 提取 PDF 文件 {file_path} 时出错: {str(e)}")
            return self._extract_pdf(file_path)
    
    def _extract_docx(self, file_path: Path) -> Dict[str, str]:
        """
        提取 Word 文件文本
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            doc = Document(file_path)
            text = ''
            # 提取段落文本
            for para in doc.paragraphs:
                text += para.text + '\n'
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text for cell in row.cells])
                    text += row_text + '\n'
            
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except Exception as e:
            logger.error(f"提取 Word 文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def _extract_xlsx(self, file_path: Path) -> Dict[str, List[str]]:
        """
        提取 Excel 文件文本
        
        Args:
            file_path: Excel 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            wb = load_workbook(file_path)
            sheets_data = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = f"Sheet: {sheet_name}\n"
                
                # 处理合并单元格
                merged_cells = sheet.merged_cells.ranges
                merged_dict = {}
                for merged in merged_cells:
                    top_left = merged.min_row, merged.min_col
                    value = sheet.cell(row=top_left[0], column=top_left[1]).value
                    for row in range(merged.min_row, merged.max_row + 1):
                        for col in range(merged.min_col, merged.max_col + 1):
                            merged_dict[(row, col)] = value
                
                # 提取单元格内容
                for row in range(1, sheet.max_row + 1):
                    row_cells = []
                    for col in range(1, sheet.max_column + 1):
                        if (row, col) in merged_dict:
                            row_cells.append(str(merged_dict[(row, col)]) if merged_dict[(row, col)] is not None else '')
                        else:
                            cell_value = sheet.cell(row=row, column=col).value
                            row_cells.append(str(cell_value) if cell_value is not None else '')
                    row_text = '\t'.join(row_cells)
                    sheet_text += row_text + '\n'
                
                sheets_data.append(sheet_text)
            
            return {
                "file_path": str(file_path),
                "text": sheets_data
            }
        except Exception as e:
            logger.error(f"提取 Excel 文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def _extract_txt(self, file_path: Path) -> Dict[str, str]:
        """
        提取 TXT 文件文本
        
        Args:
            file_path: TXT 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except Exception as e:
            logger.error(f"提取 TXT 文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def _extract_csv(self, file_path: Path) -> Dict[str, str]:
        """
        提取 CSV 文件文本
        
        Args:
            file_path: CSV 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            text = ''
            # 尝试使用 utf-8 编码
            try:
                with open(file_path, 'r', encoding='utf-8', newline='') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        row_text = '\t'.join(row)
                        text += row_text + '\n'
            except UnicodeDecodeError:
                # 尝试使用 gbk 编码
                with open(file_path, 'r', encoding='gbk', newline='', errors='ignore') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        row_text = '\t'.join(row)
                        text += row_text + '\n'
            
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except Exception as e:
            logger.error(f"提取 CSV 文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def _extract_markdown(self, file_path: Path) -> Dict[str, str]:
        """
        提取 Markdown 文件文本
        
        Args:
            file_path: Markdown 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            try:
                import markdown
                from bs4 import BeautifulSoup
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    md_content = f.read()
                
                # 将 Markdown 转换为 HTML
                html = markdown.markdown(md_content)
                
                # 使用 BeautifulSoup 提取纯文本
                soup = BeautifulSoup(html, 'html.parser')
                text = soup.get_text(separator='\n')
            except ImportError:
                # 如果没有安装依赖，直接读取文件内容
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except Exception as e:
            logger.error(f"提取 Markdown 文件 {file_path} 时出错: {str(e)}")
            return {
                "file_path": str(file_path),
                "error": str(e)
            }
    
    def _extract_wps(self, file_path: Path) -> Dict[str, str]:
        """
        提取 WPS 文件文本
        
        Args:
            file_path: WPS 文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            # WPS 文件本质是 ZIP 格式，尝试使用 python-docx 读取
            doc = Document(file_path)
            text = ''
            for para in doc.paragraphs:
                text += para.text + '\n'
            
            return {
                "file_path": str(file_path),
                "text": text.strip()
            }
        except Exception as e:
            logger.error(f"提取 WPS 文件 {file_path} 时出错: {str(e)}")
            # 如果失败，尝试直接读取文件内容
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                return {
                    "file_path": str(file_path),
                    "text": text.strip()
                }
            except Exception as e2:
                return {
                    "file_path": str(file_path),
                    "error": str(e2)
                }
    
    def _extract_wps_et(self, file_path: Path) -> Dict[str, List[str]]:
        """
        提取 WPS 表格文件文本
        
        Args:
            file_path: WPS 表格文件路径
            
        Returns:
            包含文件路径和提取文本的字典
        """
        try:
            # WPS 表格文件本质是 ZIP 格式，尝试使用 openpyxl 读取
            wb = load_workbook(file_path)
            sheets_data = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                    sheet_text += row_text + '\n'
                
                sheets_data.append(sheet_text)
            
            return {
                "file_path": str(file_path),
                "text": sheets_data
            }
        except Exception as e:
            logger.error(f"提取 WPS 表格文件 {file_path} 时出错: {str(e)}")
            # 如果失败，尝试直接读取文件内容
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                return {
                    "file_path": str(file_path),
                    "text": [text.strip()]
                }
            except Exception as e2:
                return {
                    "file_path": str(file_path),
                    "error": str(e2)
                }

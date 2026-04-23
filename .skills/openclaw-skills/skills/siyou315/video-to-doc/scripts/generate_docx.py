#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据结构化内容生成 Word 文档
使用 python-docx 库生成专业的图文操作指南

用法:
    python generate_docx.py <内容文件> <截图目录> <输出文件>
"""

import os
import sys
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

# 尝试导入 docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 请先安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_hyperlink(paragraph, url: str, text: str):
    """添加超链接"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    
    # 蓝色下划线样式
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def generate_word_doc(
    content_file: str,
    frames_dir: str,
    output_file: str,
    title_font: str = "微软雅黑",
    body_font: str = "微软雅黑",
    title_size: int = 22,
    heading_size: int = 16,
    body_size: int = 11
) -> bool:
    """
    生成 Word 文档
    
    Args:
        content_file: 内容JSON文件
        frames_dir: 截图目录
        output_file: 输出文件路径
        title_font: 标题字体
        body_font: 正文字体
        title_size: 标题字号
        heading_size: 标题字号
        body_size: 正文字号
    
    Returns:
        是否成功
    """
    # 加载内容
    with open(content_file, "r", encoding="utf-8") as f:
        content = json.load(f)
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = body_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    style.font.size = Pt(body_size)
    
    # ===== 标题 =====
    title = doc.add_heading("", 0)
    title_run = title.add_run(content.get("title", "操作指南"))
    title_run.font.name = title_font
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), title_font)
    title_run.font.size = Pt(title_size)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # 深蓝色
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    if content.get("subtitle"):
        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run(content["subtitle"])
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(89, 89, 89)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # ===== 一、操作概述 =====
    doc.add_heading("一、操作概述", level=1)
    overview = doc.add_paragraph()
    overview_run = overview.add_run(content.get("overview", "本文档提供详细的操作指南。"))
    overview_run.font.size = Pt(body_size)
    
    doc.add_paragraph()  # 空行
    
    # ===== 二、操作流程图 =====
    doc.add_heading("二、操作流程图", level=1)
    
    # 提取步骤关键词生成流程图
    steps = content.get("steps", [])
    if steps:
        # 生成流程节点
        flow_nodes = []
        for step in steps:
            step_title = step.get("title", f"步骤{step.get('step_num', len(flow_nodes)+1)}")
            # 简化标题
            if "：" in step_title:
                step_title = step_title.split("：")[-1]
            # 提取关键词（取前4个字或整个标题）
            if len(step_title) > 6:
                flow_nodes.append(step_title[:6] + "...")
            else:
                flow_nodes.append(step_title)
        
        # 添加开始和结束节点
        flow_nodes = ["进入系统"] + flow_nodes + ["完成"]
        
        # 创建流程图表格
        num_nodes = len(flow_nodes)
        num_cols = num_nodes * 2 - 1  # 节点 + 箭头
        table = doc.add_table(rows=1, cols=num_cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells = table.rows[0].cells
        
        for i, node in enumerate(flow_nodes):
            # 节点单元格
            cell_idx = i * 2
            cell = cells[cell_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(node)
            run.font.name = body_font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)
            
            # 箭头单元格
            if i < len(flow_nodes) - 1:
                arrow_cell = cells[cell_idx + 1]
                arrow_p = arrow_cell.paragraphs[0]
                arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                arrow_run = arrow_p.add_run("→")
                arrow_run.font.size = Pt(14)
                arrow_run.font.bold = True
                arrow_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()  # 空行
    
    # ===== 三、详细操作步骤 =====
    doc.add_heading("三、详细操作步骤", level=1)
    
    steps = content.get("steps", [])
    for step in steps:
        # 步骤标题（使用提炼后的标题，不再重复步骤序号）
        step_title = step.get("title", f"步骤{step['step_num']}")
        # 确保标题不包含无意义的"步骤X：步骤X"重复
        if step_title.startswith(f"步骤{step['step_num']}："):
            step_title = step_title.split("：")[-1]
        elif step_title == f"步骤{step['step_num']}":
            step_title = f"操作{step['step_num']}"
        
        doc.add_heading(f"步骤{step['step_num']}：{step_title}", level=2)
        
        # 操作说明（提炼后的规范描述，支持标记格式）
        if step.get("description"):
            # 如果是换行分隔的多条，转为列表
            desc_lines = step["description"].strip().split("\n")
            for line in desc_lines:
                line = line.strip()
                if line:
                    # 移除序号前缀（如"1. "）
                    if re.match(r'^\d+[.、]', line):
                        line = re.sub(r'^\d+[.、]\s*', '', line)
                    p = doc.add_paragraph()
                    
                    # 解析标记格式：【按钮】蓝色加粗，「输入框」绿色加粗
                    # 简单实现：检测并应用颜色
                    run = p.add_run(line)
                    run.font.name = body_font
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
                    
                    # TODO: 可扩展为更精细的格式化
                    p.paragraph_format.left_indent = Cm(0.5)
        
        # 截图
        image_path = step.get("image_path", "")
        if image_path:
            # 拼接完整路径：frames_dir + 图片文件名
            image_filename = os.path.basename(image_path) if "/" in image_path else image_path
            full_image_path = os.path.join(frames_dir, image_filename)
            
            if os.path.exists(full_image_path):
                try:
                    # 添加图片
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_img.add_run()
                    run.add_picture(full_image_path, width=Inches(5.5))
                    
                    # 图片说明
                    p_caption = doc.add_paragraph()
                    caption_run = p_caption.add_run(f"图{step['step_num']}：{step_title}")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = RGBColor(89, 89, 89)
                    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[图片加载失败: {image_filename}]")
            else:
                doc.add_paragraph(f"[图片不存在: {image_filename}]")
        else:
            doc.add_paragraph("[无截图]")
        
        doc.add_paragraph()  # 空行
    
    # ===== 四、注意事项 =====
    doc.add_heading("四、注意事项", level=1)
    notes = content.get("notes", ["请按步骤顺序操作"])
    for note in notes:
        p = doc.add_paragraph(note, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
    
    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"{content.get('title', '操作指南')} - 第"
    footer_para.add_run("PAGE")
    footer_para.add_run("页 / 共")
    footer_para.add_run("NUMPAGES")
    footer_para.add_run("页")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存
    doc.save(output_file)
    print(f"文档已保存: {output_file}")
    return True


def main():
    if len(sys.argv) < 4:
        print("用法: python generate_docx.py <内容文件> <截图目录> <输出文件>")
        print("\n示例:")
        print("  python generate_docx.py content.json ./frames output.docx")
        print("\n流程说明:")
        print("  1. 使用 analyze_frames.py 生成截图分析")
        print("  2. 使用 transcribe_audio.py 生成语音转录")
        print("  3. 使用 align_and_generate.py 生成内容JSON")
        print("  4. 使用本脚本生成 Word 文档")
        sys.exit(1)
    
    content_file = sys.argv[1]
    frames_dir = sys.argv[2]
    output_file = sys.argv[3]
    
    # 检查文件
    if not os.path.exists(content_file):
        print(f"错误: 内容文件不存在: {content_file}")
        sys.exit(1)
    
    if not os.path.exists(frames_dir):
        print(f"警告: 截图目录不存在: {frames_dir}")
    
    # 生成文档
    try:
        success = generate_word_doc(content_file, frames_dir, output_file)
        if success:
            print(f"\n✓ Word 文档生成成功!")
            print(f"  文件: {output_file}")
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
